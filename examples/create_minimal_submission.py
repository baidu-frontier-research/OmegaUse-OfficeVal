"""Create a tiny synthetic submission ZIP for smoke testing the CLI."""

from __future__ import annotations

import argparse
import tempfile
import zipfile
from pathlib import Path

from docx import Document


TASK_ID = "002"


def create_submission(output_path: Path) -> Path:
    """Create one synthetic DOCX under officeval_002 and archive it."""
    output_path = output_path.expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="omegause_synthetic_") as temp_dir:
        document_path = Path(temp_dir) / "synthetic.docx"
        document = Document()
        document.add_heading("Synthetic OmegaUse-OfficeVal Smoke Test", level=1)
        document.add_paragraph(
            "This document is generated locally for package and pipeline testing."
        )
        document.save(document_path)

        with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.write(
                document_path,
                f"officeval_{TASK_ID}/{document_path.name}",
            )

    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Create a minimal synthetic OfficeVal submission ZIP."
    )
    parser.add_argument(
        "--output",
        type=Path,
        required=True,
        help="Output ZIP path, for example ./synthetic-submission.zip",
    )
    args = parser.parse_args()
    print(create_submission(args.output))


if __name__ == "__main__":
    main()
