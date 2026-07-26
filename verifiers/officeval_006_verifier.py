# -*- coding: utf-8 -*-
"""
高中数学作业自动评分脚本
对目录内的 Word 文档进行两维度自动评测.

统一接口:
    evaluate(dir_path: str) -> dict
        - 接收脚本所在目录的路径, 脚本自己在该目录里定位并打开被评估的文档
        - 返回结构化字典 (含维度一是否通过、维度二逐项得分、总分)
"""
import os
import re
import sys
import json

from docx import Document

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

SCRIPT_ID = '006'


# ============================================================
# 辅助函数
# ============================================================

def half_pt_to_pt(val):
    return int(val) / 2.0


def get_para_alignment(para):
    ppr = para._element.find('{%s}pPr' % W_NS)
    if ppr is None:
        return None
    jc = ppr.find('{%s}jc' % W_NS)
    if jc is None:
        return None
    return jc.get('{%s}val' % W_NS)


def get_para_line_spacing(para):
    ppr = para._element.find('{%s}pPr' % W_NS)
    if ppr is None:
        return None, None
    sp = ppr.find('{%s}spacing' % W_NS)
    if sp is None:
        return None, None
    return sp.get('{%s}line' % W_NS), sp.get('{%s}lineRule' % W_NS)


def line_spacing_multiplier(line, lineRule):
    if line is None:
        return None
    line = int(line)
    if lineRule is None or lineRule == 'auto':
        return line / 240.0
    return line / 240.0


def get_run_font_info(run):
    rpr = run._element.find('{%s}rPr' % W_NS)
    if rpr is None:
        return {}
    info = {}
    fonts = rpr.find('{%s}rFonts' % W_NS)
    if fonts is not None:
        info['eastAsia'] = fonts.get('{%s}eastAsia' % W_NS)
        info['ascii'] = fonts.get('{%s}ascii' % W_NS)
        info['hAnsi'] = fonts.get('{%s}hAnsi' % W_NS)
    sz = rpr.find('{%s}sz' % W_NS)
    if sz is not None:
        info['size_halfpt'] = int(sz.get('{%s}val' % W_NS))
    b = rpr.find('{%s}b' % W_NS)
    if b is None:
        info['bold'] = False
    else:
        val = b.get('{%s}val' % W_NS)
        # OOXML: <w:b/> 或 val in (true/1/on) 表示加粗;
        #        <w:b w:val="0"/> (或 false/off) 表示显式取消加粗
        info['bold'] = val is None or str(val).lower() in ('true', '1', 'on')
    return info


def get_first_run_font(para):
    for r in para.runs:
        if r.text.strip():
            return get_run_font_info(r)
    return {}


def has_omml_equation(para):
    element = para._element
    if element.findall('.//{%s}oMath' % M_NS) or element.findall('.//{%s}oMathPara' % M_NS):
        return True
    return False


def check_omml_in_doc(doc):
    for p in doc.paragraphs:
        if has_omml_equation(p):
            return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if has_omml_equation(p):
                        return True
    return False


def has_sqrt_text(text):
    return '√' in text or '\u221a' in text


def has_fraction_text(text):
    return bool(re.search(r'\d\s*/\s*\d', text))


def get_footer_info(section):
    footer = section.footer
    if footer is None:
        return None, False, False, ''
    has_page = False
    has_numpages = False
    footer_text = ''
    for p in footer.paragraphs:
        footer_text += p.text
        for fld in p._element.findall('.//{%s}fldSimple' % W_NS):
            instr = fld.get('{%s}instr' % W_NS)
            if instr == 'PAGE':
                has_page = True
            elif instr == 'NUMPAGES':
                has_numpages = True
    return footer_text.strip(), has_page, has_numpages, footer_text


def is_empty_para(para):
    return not para.text.strip()


def count_images(doc):
    return sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)


def count_text_runs(doc):
    return sum(1 for p in doc.paragraphs for r in p.runs if r.text.strip())


# ============================================================
# 评分主逻辑
# ============================================================

# 维度二全部评分项 (标签, 分值). 标签需与后续 hit_items.append(...) 完全一致.
ALL_RULES = [
    ('+1 标题行字体: 黑体小二加粗中文, Times New Roman小二加粗数字', 1),
    ('+1 标题行段落格式: 居中对齐, 1.1倍行距', 1),
    ('+1 第二行字体: 宋体五号加粗', 1),
    ('+1 第二行段落格式: 居中对齐, 1.1倍行距', 1),
    ('+1 章节标题字体: 宋体五号加粗', 1),
    ('+1 章节标题段落格式: 两端对齐, 1.35倍行距', 1),
    ('+1 正文字体: 中文宋体五号, 英文及数字Times New Roman五号', 1),
    ('+1 正文段落格式: 两端对齐, 1.34倍行距', 1),
    ('+3 页码: "第x页 共y页" 页面底部居中', 3),
    ('-5 根号显示不全: 使用 "√x" 文本形式而非 OMML "$\\sqrt{x}$"', -5),
    ('-5 试卷第一页、第二页题目下方出现空白行', -5),
    ('-5 选项不对齐: 上一题选项A与下一题选项A/C位置不一致', -5),
    ('-5 分数显示不对应: 使用 "a/b" 文本形式而非 OMML "$\\dfrac{a}{b}$"', -5),
]


def _rule_name(label):
    """把 '+1 xxx' / '-5 xxx' 前缀去掉, 只保留规则名."""
    m = re.match(r'^[+-]\d+\s*', label)
    return label[m.end():] if m else label


def _build_dim2_items(hit_items):
    """把命中标签列表拼成结构化 dim2_items (含未命中)."""
    hit_set = set(hit_items)
    items = []
    for label, max_delta in ALL_RULES:
        hit = label in hit_set
        items.append({
            'rule': _rule_name(label),
            'max_delta': max_delta,
            'delta': max_delta if hit else 0,
            'hit': hit,
            'detail': '',
        })
    return items


def _max_score():
    # 满分 = 所有加分项 max_delta 之和 (不含扣分项)
    return sum(md for _, md in ALL_RULES if md > 0)


def _open_doc_as_document(file_path: str):
    """打开 .docx 文档 (python-docx 仅支持 OOXML/.docx)."""
    return Document(file_path)


def _find_target_file(dir_path):
    """在给定目录内定位待评估的 Word 文档 (.docx)."""
    if not os.path.isdir(dir_path):
        return None
    for name in sorted(os.listdir(dir_path)):
        if name.startswith('~$'):
            continue
        if name.lower().endswith('.docx'):
            return os.path.join(dir_path, name)
    return None


def _error_result(dir_path, message):
    file_name = ''
    if dir_path:
        candidate = _find_target_file(dir_path)
        if candidate:
            file_name = os.path.basename(candidate)
    return {
        'id': SCRIPT_ID,
        'file_name': file_name,
        'status': 'error',
        'error': message,
        'dim1_pass': False,
        'dim1_reason': message,
        'dim2_items': [],
        'total_score': 0,
        'max_score': _max_score(),
    }


def _gate_fail_result(file_name, reason):
    return {
        'id': SCRIPT_ID,
        'file_name': file_name,
        'status': 'ok',
        'error': None,
        'dim1_pass': False,
        'dim1_reason': reason,
        'dim2_items': _build_dim2_items([]),
        'total_score': 0,
        'max_score': _max_score(),
    }


def evaluate(dir_path: str) -> dict:
    """对指定目录内的 Word 文档进行两维度评测.

    参数:
        dir_path: 脚本所在目录路径, 脚本自动在其中定位 .docx 文件.
    返回:
        结构化字典, 字段说明见 "脚本接口差异与统一建议.md" §2.2.
    """
    try:
        return _evaluate_impl(dir_path)
    except Exception as exc:
        return _error_result(dir_path, f'{type(exc).__name__}: {exc}')


def _evaluate_impl(dir_path: str) -> dict:
    file_path = _find_target_file(dir_path)
    if file_path is None:
        return _error_result(dir_path, '目录下未找到 .docx 文件')
    file_name = os.path.basename(file_path)

    # ---- 维度一: 一票否决项 (Gate) ----
    gate_pass = True
    dim1_reason = ''

    if not file_path.lower().endswith('.docx'):
        gate_pass = False
        dim1_reason = '文件扩展名不是 .docx'

    try:
        doc = _open_doc_as_document(file_path)
    except Exception as exc:
        return _error_result(dir_path, f'无法打开文档: {exc}')

    if gate_pass:
        # 空段落过多 / 乱码
        total_paras = len(doc.paragraphs)
        empty_count = sum(1 for p in doc.paragraphs if is_empty_para(p))
        empty_ratio = empty_count / max(total_paras, 1)

        garbled = False
        # Word 中合法的空白控制字符 (制表 / 换行 / 回车) 不属于乱码
        _allowed_ctrl = {0x09, 0x0A, 0x0D}
        for p in doc.paragraphs:
            for ch in p.text:
                code = ord(ch)
                if code in _allowed_ctrl:
                    continue
                if code < 0x20 or (0x7F <= code <= 0x9F):
                    garbled = True
                    break
            if garbled:
                break

        overlap_issue = False
        for p in doc.paragraphs[:50]:
            line, _ = get_para_line_spacing(p)
            if line and int(line) < 0:
                overlap_issue = True
                break

        if empty_ratio > 0.5 or garbled or overlap_issue:
            gate_pass = False
            if empty_ratio > 0.5:
                dim1_reason = f'空段落比例过高 ({empty_ratio:.2f})'
            elif garbled:
                dim1_reason = '文档含乱码控制字符'
            else:
                dim1_reason = '行距异常(<0)导致重叠'

        # PDF 截图检测
        text_run_count = count_text_runs(doc)
        image_count = count_images(doc)
        if text_run_count == 0 or (image_count > 0 and text_run_count < 10):
            gate_pass = False
            dim1_reason = 'PDF 截图/无文字内容'

        # 答案内容检测
        full_text = ''
        for p in doc.paragraphs:
            full_text += p.text + '\n'
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        full_text += p.text + '\n'

        if any(kw in full_text for kw in ['答案', '解析', '详解', '参考答案']):
            gate_pass = False
            dim1_reason = '文档包含答案/解析内容'
        if '0.83' in full_text or '3.68' in full_text:
            gate_pass = False
            dim1_reason = '文档包含疑似答案数值'

    if not gate_pass:
        return _gate_fail_result(file_name, dim1_reason or '维度一未通过')

    # ---- 维度二: 得分项检查 ----
    hit_items = []

    paragraphs = doc.paragraphs
    title_para = paragraphs[0]
    second_para = paragraphs[1]

    # 识别章节标题
    #   - 兼容两种落盘方式:
    #     (a) 文本直接含 "一、单选题" 之类 -> 正则 ^[一二三四五六七八九十]、
    #     (b) "一、二、三、四" 由 w:pPr/w:numPr 自动编号生成, 文本仅剩 "单选题" 等 -> 走 numPr + 名称白名单
    pattern_chinese_num = re.compile(r'^[一二三四五六七八九十]、')
    _section_names = {'单选题', '多选题', '填空题', '解答题'}

    def _para_has_numpr(para):
        ppr = para._element.find('{%s}pPr' % W_NS)
        return ppr is not None and ppr.find('{%s}numPr' % W_NS) is not None

    section_header_paras = []
    section_header_indices = set()
    for i, p in enumerate(paragraphs):
        stripped = p.text.strip()
        norm = re.sub(r'\s+', '', p.text)
        if pattern_chinese_num.match(stripped) or (norm in _section_names and _para_has_numpr(p)):
            section_header_paras.append(p)
            section_header_indices.add(i)

    # 正文段落 (排除标题、第二行、章节标题)
    body_paras = []
    skip_idx = {0, 1} | section_header_indices
    for i, p in enumerate(paragraphs):
        if i not in skip_idx and p.text.strip():
            body_paras.append(p)

    # --- +1: 标题行字体 (黑体小二加粗中文, Times New Roman小二加粗数字) ---
    # 细则: 文档顶部第一行内容为 "2026 年 3 月 17 日高中数学作业"
    #       中文文本 -> 黑体 / 小二(18pt=36 halfpt) / 加粗
    #       数字文本 -> Times New Roman / 小二 / 加粗
    def _is_cjk(ch):
        return '一' <= ch <= '鿿'

    expected_title_norm = '2026年3月17日高中数学作业'
    title_text_norm = re.sub(r'\s+', '', title_para.text)
    title_content_ok = title_text_norm == expected_title_norm

    chinese_ok = True
    digit_ok = True
    found_cn = False
    found_digit = False
    for run in title_para.runs:
        if not run.text:
            continue
        f = get_run_font_info(run)
        size_ok = f.get('size_halfpt') == 36  # 小二 = 18pt = 36 half-pt
        bold_ok = f.get('bold') is True
        ascii_font = (f.get('ascii') or '').lower()
        hansi_font = (f.get('hAnsi') or '').lower()
        east_font = f.get('eastAsia')
        for ch in run.text:
            if _is_cjk(ch):
                found_cn = True
                if not (east_font == '黑体' and size_ok and bold_ok):
                    chinese_ok = False
            elif ch.isdigit():
                found_digit = True
                if not ((ascii_font == 'times new roman' or hansi_font == 'times new roman')
                        and size_ok and bold_ok):
                    digit_ok = False

    if title_content_ok and found_cn and found_digit and chinese_ok and digit_ok:
        hit_items.append('+1 标题行字体: 黑体小二加粗中文, Times New Roman小二加粗数字')

    # --- +1: 标题行段落格式 (居中对齐, 1.1倍行距) ---
    # 细则: 文档顶部第一行内容为 "2026 年 3 月 17 日高中数学作业"
    #       段落格式 -> 居中对齐 + 1.1 倍行距
    # 基于 Word 段落属性:
    #   w:pPr/w:jc/@w:val = "center"
    #   w:pPr/w:spacing/@w:lineRule = "auto" 且 @w:line = 264 (1.1 * 240)
    title_align = get_para_alignment(title_para)
    title_line, title_rule = get_para_line_spacing(title_para)
    align_ok = title_align == 'center'
    spacing_ok = False
    if title_line is not None and (title_rule is None or title_rule == 'auto'):
        try:
            spacing_ok = int(title_line) == 264
        except (TypeError, ValueError):
            spacing_ok = False
    if title_content_ok and align_ok and spacing_ok:
        hit_items.append('+1 标题行段落格式: 居中对齐, 1.1倍行距')

    # --- +1: 第二行字体 (宋体五号加粗) ---
    # 细则: 文档顶部第二行内容为包含
    #       "学校：____________　姓名：____________　班级：____________　考号：____________"
    #       汉字文本为宋体五号加粗
    # 判定: 先校验第二行内容包含四个字段且各带下划线占位 (前置条件),
    #       再遍历该段落所有含汉字的 run, 确认均为 宋体 / 五号(10.5pt, 21 half-pt 附近) / 加粗;
    #       内容不对或任一汉字 run 字体不符均不得分.
    s2_text = second_para.text
    # 内容包含性判定: 四个字段 (学校/姓名/班级/考号) 且每个字段后都有下划线占位
    required_fields = ['学校', '姓名', '班级', '考号']
    fields_ok = all(fd in s2_text for fd in required_fields)
    underline_ok = s2_text.count('_') >= 4 * 12  # 每个字段 12 个下划线, 共 4 个字段
    s2_content_ok = fields_ok and underline_ok

    s2_font_ok = False
    if s2_content_ok:
        s2_font_ok = True
        found_s2_cn = False
        for run in second_para.runs:
            if not any(_is_cjk(ch) for ch in run.text):
                continue
            found_s2_cn = True
            f = get_run_font_info(run)
            size_ok = (f.get('size_halfpt') is not None
                       and abs(half_pt_to_pt(f['size_halfpt']) - 10.5) <= 1)
            bold_ok = f.get('bold') is True
            if not (f.get('eastAsia') == '宋体' and size_ok and bold_ok):
                s2_font_ok = False
                break
        if not found_s2_cn:
            s2_font_ok = False

    if s2_content_ok and s2_font_ok:
        hit_items.append('+1 第二行字体: 宋体五号加粗')

    # --- +1: 第二行段落格式 (居中对齐, 1.1倍行距) ---
    # 细则: 文档顶部第二行内容包含 "学校：____________　姓名：____________　班级：____________　考号：____________"
    #       段落格式 -> 居中对齐 + 1.1 倍行距
    # 基于 Word 段落属性:
    #   w:pPr/w:jc/@w:val = "center"
    #   w:pPr/w:spacing/@w:lineRule = "auto" 且 @w:line = 264 (1.1 * 240)
    s2_align = get_para_alignment(second_para)
    s2_line, s2_rule = get_para_line_spacing(second_para)
    s2_align_ok = s2_align == 'center'
    s2_spacing_ok = False
    if s2_line is not None and (s2_rule is None or s2_rule == 'auto'):
        try:
            s2_spacing_ok = int(s2_line) == 264
        except (TypeError, ValueError):
            s2_spacing_ok = False
    if s2_content_ok and s2_align_ok and s2_spacing_ok:
        hit_items.append('+1 第二行段落格式: 居中对齐, 1.1倍行距')

    # --- +1: 章节标题字体 (宋体五号加粗) ---
    # 细则: "一、单选题"、"二、多选题"、"三、填空题"、"四、解答题" 格式为 宋体 / 五号 / 加粗
    # 基于 Word 属性 (含继承链, 由高到低):
    #   run rPr  →  段落 pPr/rPr  →  段落 style rPr (递归 basedOn)  →  docDefaults/rPrDefault
    #   - "一、二、三、四" 由 w:pPr/w:numPr 自动编号生成 (numId=1), 不在 p.text 中
    #     -> 章节段落识别: 段落文本 ∈ {单选题, 多选题, 填空题, 解答题} 且带 w:numPr
    #   - 宋体: eastAsia == '宋体'; 五号: sz val == 21; 加粗: w:b (无 val 或 val ∈ true/1/on)

    _styles_el = doc.styles.element
    _style_by_id = {
        s.get('{%s}styleId' % W_NS): s
        for s in _styles_el.findall('{%s}style' % W_NS)
        if s.get('{%s}styleId' % W_NS)
    }

    def _rpr_to_font(rpr):
        info = {'eastAsia': None, 'ascii': None, 'hAnsi': None,
                'size_halfpt': None, 'bold': None}
        if rpr is None:
            return info
        fonts = rpr.find('{%s}rFonts' % W_NS)
        if fonts is not None:
            info['eastAsia'] = fonts.get('{%s}eastAsia' % W_NS)
            info['ascii'] = fonts.get('{%s}ascii' % W_NS)
            info['hAnsi'] = fonts.get('{%s}hAnsi' % W_NS)
        sz = rpr.find('{%s}sz' % W_NS)
        if sz is not None:
            val = sz.get('{%s}val' % W_NS)
            if val:
                info['size_halfpt'] = int(val)
        b = rpr.find('{%s}b' % W_NS)
        if b is not None:
            v = b.get('{%s}val' % W_NS)
            info['bold'] = v is None or str(v).lower() in ('true', '1', 'on')
        return info

    def _style_font_chain(style_id, seen=None):
        """沿 basedOn 链累积 style rPr, 高优先级在前."""
        if seen is None:
            seen = set()
        if not style_id or style_id in seen or style_id not in _style_by_id:
            return {'eastAsia': None, 'ascii': None, 'hAnsi': None,
                    'size_halfpt': None, 'bold': None}
        seen.add(style_id)
        s = _style_by_id[style_id]
        cur = _rpr_to_font(s.find('{%s}rPr' % W_NS))
        based = s.find('{%s}basedOn' % W_NS)
        base_id = based.get('{%s}val' % W_NS) if based is not None else None
        base = _style_font_chain(base_id, seen)
        return {k: cur.get(k) if cur.get(k) is not None else base.get(k)
                for k in cur}

    _doc_default_font = None
    def _doc_defaults_font():
        nonlocal _doc_default_font
        if _doc_default_font is not None:
            return _doc_default_font
        dd = _styles_el.find('{%s}docDefaults' % W_NS)
        if dd is None:
            _doc_default_font = _rpr_to_font(None)
            return _doc_default_font
        rprd = dd.find('{%s}rPrDefault' % W_NS)
        rpr = rprd.find('{%s}rPr' % W_NS) if rprd is not None else None
        _doc_default_font = _rpr_to_font(rpr)
        return _doc_default_font

    def _default_para_style_id():
        for sid, s in _style_by_id.items():
            if s.get('{%s}type' % W_NS) == 'paragraph' and s.get('{%s}default' % W_NS) == '1':
                return sid
        return None

    _default_pstyle_id = _default_para_style_id()

    def _para_style_id(para):
        ppr = para._element.find('{%s}pPr' % W_NS)
        if ppr is None:
            return _default_pstyle_id
        ps = ppr.find('{%s}pStyle' % W_NS)
        if ps is None:
            return _default_pstyle_id
        return ps.get('{%s}val' % W_NS) or _default_pstyle_id

    def _get_para_rpr_font(para):
        ppr = para._element.find('{%s}pPr' % W_NS)
        if ppr is None:
            return _rpr_to_font(None)
        return _rpr_to_font(ppr.find('{%s}rPr' % W_NS))

    def _effective_font(para, run):
        rf = get_run_font_info(run)
        pf = _get_para_rpr_font(para)
        sf = _style_font_chain(_para_style_id(para))
        df = _doc_defaults_font()
        # 继承合并: run > pPr/rPr > style chain > docDefaults
        merged = {}
        for k in ('eastAsia', 'ascii', 'hAnsi', 'size_halfpt', 'bold'):
            for src in (rf, pf, sf, df):
                v = src.get(k) if isinstance(src, dict) else None
                if v is not None:
                    merged[k] = v
                    break
            else:
                merged[k] = None
        return merged

    def _has_numpr(para):
        ppr = para._element.find('{%s}pPr' % W_NS)
        return ppr is not None and ppr.find('{%s}numPr' % W_NS) is not None

    required_sections = ['一、单选题', '二、多选题', '三、填空题', '四、解答题']
    required_names = {'单选题', '多选题', '填空题', '解答题'}
    name_to_full = {'单选题': '一、单选题', '多选题': '二、多选题',
                    '填空题': '三、填空题', '解答题': '四、解答题'}

    section_map = {}
    for p in paragraphs:
        t = re.sub(r'\s+', '', p.text)
        if t in required_sections and t not in section_map:
            section_map[t] = p
            continue
        if t in required_names and _has_numpr(p):
            full = name_to_full[t]
            if full not in section_map:
                section_map[full] = p

    if all(name in section_map for name in required_sections):
        all_ok = True
        for name in required_sections:
            p = section_map[name]
            para_ok = False
            for run in p.runs:
                if not run.text.strip():
                    continue
                f = _effective_font(p, run)
                if not (f.get('eastAsia') == '宋体'
                        and f.get('size_halfpt') == 21
                        and f.get('bold') is True):
                    para_ok = False
                    break
                para_ok = True
            if not para_ok:
                all_ok = False
                break
        if all_ok:
            hit_items.append('+1 章节标题字体: 宋体五号加粗')

    # --- +1: 章节标题段落格式 (两端对齐, 1.35倍行距) ---
    # 细则: "一、单选题"、"二、多选题"、"三、填空题"、"四、解答题" 段落格式为 两端对齐 + 1.35倍行距
    # 基于 Word 段落属性:
    #   w:pPr/w:jc/@w:val = "both" (两端对齐)
    #   w:pPr/w:spacing/@w:lineRule = "auto" 且 @w:line = 324 (1.35 * 240)
    if all(name in section_map for name in required_sections):
        all_ok = True
        for name in required_sections:
            p = section_map[name]
            align = get_para_alignment(p)
            line, rule = get_para_line_spacing(p)
            # 两端对齐: w:jc/@w:val = "both"
            align_ok = align == 'both'
            # 1.35 倍行距: line=324, lineRule=auto
            spacing_ok = False
            if line is not None and (rule is None or rule == 'auto'):
                try:
                    spacing_ok = int(line) == 324
                except (TypeError, ValueError):
                    spacing_ok = False
            if not (align_ok and spacing_ok):
                all_ok = False
                break
        if all_ok:
            hit_items.append('+1 章节标题段落格式: 两端对齐, 1.35倍行距')

    # --- +1: 正文字体 (中文宋体五号, 英文及阿拉伯数字 Times New Roman 五号) ---
    # 细则: 试卷正文
    #   中文        -> 宋体 / 五号 (10.5pt = 21 half-pt)
    #   英文 / 数字 -> Times New Roman / 五号
    # 基于 Word run 属性:
    #   w:rPr/w:rFonts/@w:eastAsia = "宋体"          (中文)
    #   w:rPr/w:rFonts/@w:ascii    = "Times New Roman" (英文/数字)
    #   w:rPr/w:rFonts/@w:hAnsi    = "Times New Roman" (英文/数字)
    #   w:rPr/w:sz/@w:val = 21
    if body_paras:
        all_ok = True
        for p in body_paras:
            for run in p.runs:
                if not run.text:
                    continue
                f = get_run_font_info(run)
                # 五号: 21 half-pt (精确)
                if f.get('size_halfpt') != 21:
                    all_ok = False
                    break
                east_font = f.get('eastAsia')
                ascii_font = (f.get('ascii') or '')
                hansi_font = (f.get('hAnsi') or '')
                for ch in run.text:
                    if _is_cjk(ch):
                        if east_font != '宋体':
                            all_ok = False
                            break
                    elif ('A' <= ch <= 'Z') or ('a' <= ch <= 'z') or ch.isdigit():
                        if not (ascii_font == 'Times New Roman'
                                and hansi_font == 'Times New Roman'):
                            all_ok = False
                            break
                if not all_ok:
                    break
            if not all_ok:
                break
        if all_ok:
            hit_items.append('+1 正文字体: 中文宋体五号, 英文及数字Times New Roman五号')

    # --- +1: 正文段落格式 (两端对齐, 1.34倍行距) ---
    # 细则: 试卷正文段落格式为 两端对齐 + 1.34 倍行距
    # 基于 Word 段落属性:
    #   w:pPr/w:jc/@w:val = "both" (两端对齐)
    #   w:pPr/w:spacing/@w:lineRule = "auto" 且 @w:line = round(1.34 * 240) = 322
    if body_paras:
        all_ok = True
        for p in body_paras:
            align = get_para_alignment(p)
            line, rule = get_para_line_spacing(p)
            align_ok = align == 'both'
            spacing_ok = False
            if line is not None and (rule is None or rule == 'auto'):
                try:
                    # 1.34 * 240 = 321.6, Word 存整数 (321 或 322)
                    spacing_ok = int(line) in (321, 322)
                except (TypeError, ValueError):
                    spacing_ok = False
            if not (align_ok and spacing_ok):
                all_ok = False
                break
        if all_ok:
            hit_items.append('+1 正文段落格式: 两端对齐, 1.34倍行距')

    # --- +3: 页码 (每页都有, 样式 "第x页 共y页", 页面底部居中) ---
    # 细则:
    #   1) 试卷每页都有页码       -> 每个 section 的 footer 均须包含页码域
    #   2) 样式为 "第x页 共y页"    -> footer 文本能匹配 ^第.*页共.*页$
    #   3) 页码放置在页面底部居中  -> 承载页码的 footer 段落水平居中
    # 基于 Word footer 属性:
    #   页码域: w:fldSimple/@w:instr 含 "PAGE" / "NUMPAGES"
    #          或 w:instrText 文本含 "PAGE" / "NUMPAGES" (复杂 field)
    #   居中: 两种合法方式均识别
    #     (A) 承载段落 w:pPr/w:jc/@w:val = "center"
    #     (B) 页码位于 w:drawing 里的锚定文本框, 且水平位置相对页边距居中:
    #         w:drawing//wp:anchor/wp:positionH[@relativeFrom='margin']/wp:align = 'center'
    # 说明: 当页码放在 (B) 中时, 承载段落的 p.text 为空,
    #       需要递归收集 w:t (含 drawing 内部) 作为样式匹配文本

    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

    def _drawing_horiz_center(el):
        anchors = el.findall('.//{%s}anchor' % WP_NS)
        for anchor in anchors:
            ph = anchor.find('{%s}positionH' % WP_NS)
            if ph is None:
                continue
            if ph.get('relativeFrom') != 'margin':
                continue
            align = ph.find('{%s}align' % WP_NS)
            if align is not None and (align.text or '').strip() == 'center':
                return True
        return False

    def _collect_all_text(el):
        parts = []
        for wt in el.findall('.//{%s}t' % W_NS):
            if wt.text:
                parts.append(wt.text)
        return ''.join(parts)

    def _check_footer(section):
        footer = section.footer
        if footer is None:
            return False
        for p in footer.paragraphs:
            has_page = False
            has_numpages = False
            for fld in p._element.findall('.//{%s}fldSimple' % W_NS):
                instr = fld.get('{%s}instr' % W_NS) or ''
                if 'PAGE' in instr and 'NUMPAGES' not in instr:
                    has_page = True
                if 'NUMPAGES' in instr:
                    has_numpages = True
            for it in p._element.findall('.//{%s}instrText' % W_NS):
                txt = (it.text or '').strip()
                if 'PAGE' in txt and 'NUMPAGES' not in txt:
                    has_page = True
                if 'NUMPAGES' in txt:
                    has_numpages = True
            if not (has_page and has_numpages):
                continue
            # 样式匹配文本: 递归收集所有 w:t (覆盖 drawing 内部承载的页码字样)
            all_text = _collect_all_text(p._element)
            cleaned = re.sub(r'\s+', '', all_text)
            style_ok = bool(re.match(r'^第.*页共.*页$', cleaned))
            # 居中: 段落 jc=center 或 drawing 锚定 wp:align=center relativeFrom=margin
            align = get_para_alignment(p)
            align_ok = align == 'center' or _drawing_horiz_center(p._element)
            if style_ok and align_ok:
                return True
        return False

    if doc.sections and all(_check_footer(sec) for sec in doc.sections):
        hit_items.append('+3 页码: "第x页 共y页" 页面底部居中')

    # --- -5: 根号显示不全 (√x 文本形式, 非 OMML 数学公式) ---
    # 细则: 试卷中的根号显示不全, 如根号x应显示为 "$\sqrt{x}$" (OMML 数学公式)
    #       而不是 "√x" (普通文本字符 U+221A)
    # 基于 Word XML:
    #   - OMML 公式渲染的根号在 <m:rad> / <m:radPr> 结构中, 其 <m:t> 里的字符
    #     由 Word 排版为根号符号, 不会以 U+221A 作为普通 run 文本存在
    #   - "√x" 文本形式表现为普通 run 里的 <w:t> 直接包含 U+221A 字符
    # 判定: 只要文档中存在任一 <w:t> 文本里含有 U+221A 且该 <w:t> 不在 <m:oMath> 内,
    #       即触发扣分
    def _has_sqrt_text_outside_omml(container):
        for wt in container._element.findall('.//{%s}t' % W_NS):
            if not wt.text or '√' not in wt.text:
                continue
            # 检查是否位于 oMath / oMathPara 内
            in_omml = False
            anc = wt.getparent()
            while anc is not None:
                tag = anc.tag
                if tag == '{%s}oMath' % M_NS or tag == '{%s}oMathPara' % M_NS:
                    in_omml = True
                    break
                anc = anc.getparent()
            if not in_omml:
                return True
        return False

    sqrt_bad = False
    for p in doc.paragraphs:
        if _has_sqrt_text_outside_omml(p):
            sqrt_bad = True
            break
    if not sqrt_bad:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if _has_sqrt_text_outside_omml(p):
                            sqrt_bad = True
                            break
                    if sqrt_bad:
                        break
                if sqrt_bad:
                    break
            if sqrt_bad:
                break
    if sqrt_bad:
        hit_items.append('-5 根号显示不全: 使用 "√x" 文本形式而非 OMML "$\\sqrt{x}$"')

    # --- -5: 试卷第一页、第二页题目下方出现空白行 ---
    # 细则: 试卷第一页、第二页题目下方出现空白行
    # 基于 Word 属性:
    #   - 页边界: <w:lastRenderedPageBreak/> (Word 保存时记录的最近一次渲染分页点)
    #            或显式 <w:br w:type="page"/> 分页符
    #   - "题目" 段落: 段落文本以 "数字." 或 "数字、" 起始, 如 "1.", "12、"
    #   - "空白行"    : 段落 w:p 内不含任何 "可见内容":
    #                   - p.text strip 后无文字
    #                   - 且不含图形对象 (w:drawing / w:pict / v:shape / v:group)
    #                   - 且不含公式对象 (m:oMath / m:oMathPara)
    #                   仅此三类都缺失时才视为真正的排版空行
    # 判定: 在第 1、2 页范围内, 若存在 "题目段落" 之后的下一个段落是空段落, 则触发扣分

    V_NS = 'urn:schemas-microsoft-com:vml'

    def _para_has_visual_object(para):
        el = para._element
        for tag in ('{%s}drawing' % W_NS,
                    '{%s}pict' % W_NS,
                    '{%s}shape' % V_NS,
                    '{%s}group' % V_NS,
                    '{%s}oMath' % M_NS,
                    '{%s}oMathPara' % M_NS):
            if el.findall('.//' + tag):
                return True
        return False

    def _is_true_empty_para(para):
        return is_empty_para(para) and not _para_has_visual_object(para)

    # 1) 定位第 1、2 页涉及的段落索引范围: 前两个 page break 之前的段落
    page_break_para_idx = []
    for i, p in enumerate(paragraphs):
        el = p._element
        if el.findall('.//{%s}lastRenderedPageBreak' % W_NS):
            page_break_para_idx.append(i)
            continue
        found_page_br = False
        for br in el.findall('.//{%s}br' % W_NS):
            if br.get('{%s}type' % W_NS) == 'page':
                found_page_br = True
                break
        if found_page_br:
            page_break_para_idx.append(i)

    if len(page_break_para_idx) >= 2:
        upper = page_break_para_idx[1]
    elif len(page_break_para_idx) == 1:
        upper = len(paragraphs) - 1
    else:
        upper = len(paragraphs) - 1

    question_pattern = re.compile(r'^\s*\d+\s*[\.\．、]')
    empty_after_question = False
    for i in range(upper):
        p = paragraphs[i]
        if question_pattern.match(p.text or ''):
            if i + 1 <= upper and _is_true_empty_para(paragraphs[i + 1]):
                empty_after_question = True
                break
    if empty_after_question:
        hit_items.append('-5 试卷第一页、第二页题目下方出现空白行')

    # --- -5: 试卷中每道题的选项不对齐 ---
    # 细则: 上一题的选项 A 应与下一题的选项 A 或选项 C 对齐
    # 思路: 对每道题, 提取其选项 A / C 在页面上的实际水平位置 (以 twips 表示),
    #       再比较相邻题目间 A-A 或 A-C 的位置差是否为 0 (或在容差内).
    #       覆盖三种常见排版:
    #         (1) 表格布局 (4 列×1行 或 2 列×2行): 位置 = 表格左边界 + 列偏移
    #         (2) 制表位(tab)分隔: 位置 = 触发该选项文本前的 tab 停靠位置(w:tabs/w:tab/@w:pos),
    #             若没有显式 tabs 定义则用默认制表位 (Word 默认 720 twips 一档) 估算
    #         (3) 普通段落 (纯空格/文本缩进): 位置 = 段落左缩进 + 选项前空白字符数 * 单字符宽度估算
    #       "对齐" 判定: 位置值以整数(twips)比较, 差值在容差 (20 twips, 约 0.03cm) 内视为对齐.
    EMU_PER_TWIP = 635  # 1 twip = 635 EMU, 用于统一表格 (twips) 与 drawing (EMU) 单位
    ALIGN_TOLERANCE_TWIPS = 20

    def _table_col_positions(tbl):
        """返回表格中 A 列 / C 列 (若存在) 的左边界水平位置 (twips)."""
        el = tbl._element
        tblPr = el.find('{%s}tblPr' % W_NS)
        jc_val = None
        ind_w = 0
        if tblPr is not None:
            jc = tblPr.find('{%s}jc' % W_NS)
            if jc is not None:
                jc_val = jc.get('{%s}val' % W_NS)
            tblInd = tblPr.find('{%s}tblInd' % W_NS)
            if tblInd is not None:
                w = tblInd.get('{%s}w' % W_NS)
                ind_w = int(w) if w else 0
        # 居中/居右表格没有固定左边界基准, 仅左对齐(或默认)表格的位置才有跨表可比性
        base = ind_w if jc_val in (None, 'left', 'start') else None

        tblGrid = el.find('{%s}tblGrid' % W_NS)
        cols = []
        if tblGrid is not None:
            for c in tblGrid.findall('{%s}gridCol' % W_NS):
                w = c.get('{%s}w' % W_NS)
                cols.append(int(w) if w else 0)

        row0_cells = tbl.rows[0].cells if tbl.rows else []
        n_cols = len(cols) or len(row0_cells)

        pos_a = base
        pos_c = None
        if base is not None:
            if n_cols >= 4:
                # 4 列布局: A 在列0, C 在列2 -> 偏移 = 列0宽 + 列1宽
                if len(cols) >= 2:
                    pos_c = base + cols[0] + cols[1]
            elif n_cols == 2:
                # 2 列 x 2 行布局: A 在 (行0,列0), C 在 (行1,列0) -> 与 A 同列
                pos_c = base
        return pos_a, pos_c

    def _is_option_table(tbl):
        # 只考察 A/B/C/D 选项表: 首行首格文本以 A．/ A. 起始
        if not tbl.rows:
            return False
        first_text = tbl.rows[0].cells[0].text.strip()
        return bool(re.match(r'^A[．.]', first_text))

    def _para_default_left_indent(para):
        """段落左缩进 (twips), 取 w:ind/@w:left 或 @w:start, 缺省为 0."""
        ppr = para._element.find('{%s}pPr' % W_NS)
        if ppr is None:
            return 0
        ind = ppr.find('{%s}ind' % W_NS)
        if ind is None:
            return 0
        v = ind.get('{%s}left' % W_NS) or ind.get('{%s}start' % W_NS)
        try:
            return int(v) if v else 0
        except (TypeError, ValueError):
            return 0

    def _para_tab_stops(para):
        """段落自定义制表位列表 (twips), 按位置升序."""
        ppr = para._element.find('{%s}pPr' % W_NS)
        if ppr is None:
            return []
        tabs = ppr.find('{%s}tabs' % W_NS)
        if tabs is None:
            return []
        stops = []
        for tab in tabs.findall('{%s}tab' % W_NS):
            pos = tab.get('{%s}pos' % W_NS)
            if pos is not None:
                try:
                    stops.append(int(pos))
                except (TypeError, ValueError):
                    pass
        return sorted(stops)

    DEFAULT_TAB_INTERVAL = 720  # Word 默认制表位间隔 (twips, 0.5 英寸)

    def _next_tab_stop(custom_stops, after):
        for pos in custom_stops:
            if pos > after:
                return pos
        # 无自定义制表位覆盖该区间, 用默认间隔估算下一档
        return ((after // DEFAULT_TAB_INTERVAL) + 1) * DEFAULT_TAB_INTERVAL

    def _option_positions_in_paragraph_text(para):
        """从纯段落文本中定位选项 A / C 的水平位置 (twips), 兼容 tab 分隔与空格缩进.

        - 若字符 A/C 前是若干个 tab (\\t), 用制表位序列逐档推进定位.
        - 否则用 "段落左缩进 + 前导空白字符数 * 估算字宽" 近似定位 (空格布局).
        """
        indent = _para_default_left_indent(para)
        tab_stops = _para_tab_stops(para)
        SPACE_WIDTH_TWIPS = 100  # 五号字下单个空格宽度的粗略估算 (仅用于同文档内相对比较)

        text = para.text
        result = {}
        for label, pattern in (('A', r'A[．.、]'), ('C', r'C[．.、]')):
            m = re.search(pattern, text)
            if not m:
                continue
            prefix = text[:m.start()]
            if '\t' in prefix:
                pos = indent
                for _ in prefix.split('\t')[:-1]:
                    pos = _next_tab_stop(tab_stops, pos)
                result[label] = pos
            else:
                leading_spaces = len(prefix) - len(prefix.lstrip(' 　'))
                result[label] = indent + leading_spaces * SPACE_WIDTH_TWIPS
        return result.get('A'), result.get('C')

    def _extract_question_option_positions():
        """按文档顺序遍历, 为每道选择题提取其选项 A / C 的水平位置.

        返回 [(pos_a, pos_c), ...], 每题一条记录 (缺失的位置为 None).
        表格布局按表格出现顺序处理; 非表格布局按 "选项段落所在的题目段落" 顺序处理,
        两者根据在文档 body 中的出现顺序统一排序.
        """
        body = doc.element.body
        records = []

        # 表格 -> python-docx Table 对象映射 (按 body 出现顺序)
        table_by_element = {t._element: t for t in doc.tables}
        # 段落 -> python-docx Paragraph 对象映射
        para_by_element = {p._element: p for p in doc.paragraphs}

        for child in body.iterchildren():
            tag = child.tag
            if tag == '{%s}tbl' % W_NS and child in table_by_element:
                tbl = table_by_element[child]
                if _is_option_table(tbl):
                    pos_a, pos_c = _table_col_positions(tbl)
                    if pos_a is not None or pos_c is not None:
                        records.append((pos_a, pos_c))
            elif tag == '{%s}p' % W_NS and child in para_by_element:
                p = para_by_element[child]
                if re.search(r'(?:^|\s)A[．.、]', p.text):
                    pos_a, pos_c = _option_positions_in_paragraph_text(p)
                    if pos_a is not None or pos_c is not None:
                        records.append((pos_a, pos_c))
        return records

    def _aligned(prev_rec, cur_rec, tol=ALIGN_TOLERANCE_TWIPS):
        """判定上一题选项 A 与下一题选项 A 或 C 是否对齐 (位置差在容差内)."""
        prev_a, _ = prev_rec
        cur_a, cur_c = cur_rec
        if prev_a is None:
            return True  # 缺失基准位置时不误判为不对齐
        for cand in (cur_a, cur_c):
            if cand is not None and abs(cand - prev_a) <= tol:
                return True
        return cur_a is None and cur_c is None

    option_records = _extract_question_option_positions()
    misaligned = False
    for i in range(1, len(option_records)):
        if not _aligned(option_records[i - 1], option_records[i]):
            misaligned = True
            break
    if misaligned:
        hit_items.append('-5 选项不对齐: 上一题选项A与下一题选项A/C位置不一致')

    # --- -5: 分数显示不对应 (a/b 文本形式, 非 OMML 分数) ---
    # 细则: 试卷中的分数显示不对应, 如三分之一应显示为 "$\dfrac{1}{3}$" (OMML 数学公式)
    #       而不是 "1/3" (普通文本 "数字/数字")
    # 基于 Word XML:
    #   - OMML 分数存放在 <m:f> / <m:frac> 结构中, 分子分母由 <m:num> / <m:den> 包裹,
    #     其字符不会以 "数字/数字" 形式出现在普通 <w:t> 里
    #   - "a/b" 文本形式: 普通 <w:t> 里直接出现 "\d+/\d+" 或跨相邻 <w:t> 拼接后出现该模式
    # 判定: 只要文档中存在任一 <w:t>(或同一 run/相邻 run 拼接) 匹配 "\d+/\d+" 且该 <w:t>
    #       不位于 <m:oMath> / <m:oMathPara> 内, 即触发扣分
    def _has_frac_text_outside_omml(container):
        # 逐段落检查: 收集不在 OMML 内的 w:t 文本, 拼接后正则匹配 "\d+/\d+"
        buf = []
        for wt in container._element.findall('.//{%s}t' % W_NS):
            if not wt.text:
                continue
            in_omml = False
            anc = wt.getparent()
            while anc is not None:
                tag = anc.tag
                if tag == '{%s}oMath' % M_NS or tag == '{%s}oMathPara' % M_NS:
                    in_omml = True
                    break
                anc = anc.getparent()
            if not in_omml:
                buf.append(wt.text)
        text = ''.join(buf)
        return bool(re.search(r'\d+\s*/\s*\d+', text))

    frac_bad = False
    for p in doc.paragraphs:
        if _has_frac_text_outside_omml(p):
            frac_bad = True
            break
    if not frac_bad:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if _has_frac_text_outside_omml(p):
                            frac_bad = True
                            break
                    if frac_bad:
                        break
                if frac_bad:
                    break
            if frac_bad:
                break
    if frac_bad:
        hit_items.append('-5 分数显示不对应: 使用 "a/b" 文本形式而非 OMML "$\\dfrac{a}{b}$"')

    # ---- 汇总输出 (结构化) ----
    total = 0
    for item in hit_items:
        m = re.match(r'^([+-]\d+)', item)
        if m:
            total += int(m.group(1))

    return {
        'id': SCRIPT_ID,
        'file_name': file_name,
        'status': 'ok',
        'error': None,
        'dim1_pass': True,
        'dim1_reason': '',
        'dim2_items': _build_dim2_items(hit_items),
        'total_score': total,
        'max_score': _max_score(),
    }


if __name__ == '__main__':
    # 本地调试: 默认在脚本所在目录里定位文档; 也可用 sys.argv[1] 显式覆盖.
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    _result = evaluate(_dir)
    # 用 ensure_ascii=True 避免在 Windows 默认代码页下的终端编码报错;
    # 也遵守 "不允许修改 sys.stdout" 的统一约定.
    print(json.dumps(_result, ensure_ascii=True, indent=2))
