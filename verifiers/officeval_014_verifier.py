# -*- coding: utf-8 -*-
"""
Word文档自动评估脚本 - 完整版
根据打分细则对Word文档进行全面评估
"""

import os

SCRIPT_ID = "014"
import ast
import re
import sys
import json
import builtins

from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 遵循"脚本接口差异与统一建议.md" §2.3：脚本不产生主输出、不修改 sys.stdout。
# 将内建 print 在本模块内遮蔽为无操作；调试信息可自行改为 logging 或 sys.stderr。
# 本地自测（__main__）需要输出 JSON，请使用 builtins.print。
def print(*args, **kwargs):  # noqa: A001 - 有意遮蔽 builtin
    return None

class WordDocumentEvaluator:
    # 加分项总分（本脚本细则内所有 +N 项理论上限之和，与命中与否无关）：
    #   自动目录 5 + 目录页 3 + 目录层级 5 + 目录核心条目 3 + 目录格式 3
    # + 页码 (封面/目录不显示 3 + 摘要罗马数字 3 + 正文从1重排 3 + 底部居中 3) = 12
    # + 页面边距 3 + 参考文献字体 3 + 参考文献编号 3 + 参考文献段落 3
    # + 封面标题 1 + 封面信息 5
    # + 中文摘要正文 3 + 中文摘要关键词 1
    # + 三线表 5 + 表内文字 5 + 二级标题 3
    # = 66
    MAX_POSITIVE_SCORE = 66

    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = None
        self.dimension1_passed = True
        self.dimension1_fail_reason = ""
        self.total_score = 0
        self.hit_items = []  # 命中的条目
        self.evaluation_details = {
            "维度1": {"结果": "通过", "详情": []},
            "维度2": {"得分项": [], "扣分项": [], "总得分": 0}
        }

    def evaluate(self):
        """主评估函数"""
        print("=" * 80)
        print("Word文档自动评估报告")
        print("=" * 80)
        print("评估文件:", os.path.basename(self.file_path))
        print("=" * 80)

        # 维度1：可用与可修改性
        print("\n【维度1：可用与可修改性】")
        if not self.check_dimension1():
            self.dimension1_passed = False
            print("❌ 维度1未通过:", self.dimension1_fail_reason)
            print("\n最终得分: 0分")
            return

        print("✓ 维度1通过")

        # 维度2：完成度评分
        print("\n【维度2：完成度评分】")
        print("-" * 80)
        print("【得分项检测】")
        self.evaluate_score_items()

        print("\n【扣分项检测】")
        self.evaluate_deduction_items()

        # 输出最终结果
        self.print_final_result()

    def check_dimension1(self):
        """检查维度1：可用与可修改性"""
        # 1. 检查文件格式
        if not self.file_path.lower().endswith('.docx'):
            self.dimension1_fail_reason = "文件格式不是.docx格式"
            return False

        # 2. 尝试打开文件
        try:
            self.doc = Document(self.file_path)
        except Exception as e:
            self.dimension1_fail_reason = "文件无法正常打开: " + str(e)
            return False

        # 3. 检查内容是否为整页图片
        paragraph_count = len(self.doc.paragraphs)
        table_count = len(self.doc.tables)

        if paragraph_count == 0 and table_count == 0:
            self.dimension1_fail_reason = "文件内容为空"
            return False

        # 检查图片比例
        image_count = self.count_images()
        text_length = sum(len(para.text) for para in self.doc.paragraphs)

        # 如果文本很短但图片很多，可能是整页图片
        if text_length < 500 and image_count > 5:
            self.dimension1_fail_reason = "正文内容疑似为整页图片"
            return False

        # 4. 检查是否可编辑
        editable_text_found = False
        for para in self.doc.paragraphs:
            if para.text.strip():
                editable_text_found = True
                break

        if not editable_text_found:
            self.dimension1_fail_reason = "未找到可编辑的文本内容"
            return False

        return True

    def count_images(self):
        """统计文档中的图片数量"""
        count = 0
        for rel in self.doc.part.rels.values():
            if "image" in rel.reltype:
                count += 1
        return count

    def evaluate_score_items(self):
        """评估得分项"""
        # +5：自动更新目录
        self.check_auto_toc()

        # +3：文中出现目录页
        self.check_toc_page()

        # +5：目录标题层级
        self.check_toc_levels()

        # +3：目录核心条目
        self.check_toc_core_items()

        # +3：目录内容及标题格式
        self.check_toc_format()

        # 页码设置相关
        self.check_page_number_settings()

        # +3：页面边距
        self.check_page_margins()

        # 参考文献相关
        self.check_reference_font()
        self.check_reference_numbering()
        self.check_reference_paragraph()

        # 封面相关
        self.check_cover_title()
        self.check_cover_info()

        # 中文摘要相关
        self.check_abstract_cn()

        # 表格相关
        self.check_tables()
        self.check_table_text()

        # 二级标题
        self.check_secondary_title()

    def evaluate_deduction_items(self):
        """评估扣分项"""
        self.check_footer_text()  # -3
        self.check_page_number_overlap()  # -3
        self.check_reference_citation()  # -5
        self.check_cover_images()  # -3, -1, -1
        self.check_cover_master_degree_title()  # -1
        self.check_declaration_signature_image()  # -1 声明页签名图片尺寸
        self.check_declaration_page()  # 多个-3
        self.check_declaration_title()  # -3 字体
        self.check_declaration_title_para()  # -3 段落
        self.check_declaration_content_font()  # -3 内容字体
        self.check_abstract_cn_format()  # -1 字体
        self.check_abstract_cn_title_para()  # -1 标题段落
        self.check_abstract_en_format()  # 多个-1
        self.check_abstract_en_body_font()  # -1 标题下方文本 TNR + 小四
        self.check_abstract_en_body_para()  # -1 标题下方文本段落 两端对齐 + 1.5倍行距
        self.check_abbreviation_page()  # -1
        self.check_preface_page()  # -1
        self.check_primary_title_para()  # -3 一级标题段落 居中 + 1.5倍行距
        self.check_secondary_title_para()  # -3
        self.check_body_text_format()  # -3, -3
        self.check_body_text_font()  # -3 正文中文小四宋体
        self.check_figure_table_notes()  # 多个-3
        self.check_figure_table_numbering()  # -3 图1/图2、表1/表2 排序
        self.check_figure_table_font()  # -3 图/表编号及名称 五号宋体
        self.check_figure_caption_position()  # -3 图名不在图片下方
        self.check_table_caption_position()  # -3 表名不在表格上方
        self.check_discussion_title_font()  # -1 讨论页标题 四号黑体
        self.check_discussion_title_para()  # -1 讨论页标题段落 居中 + 1.5倍行距
        self.check_discussion_body_font()   # -1 讨论页标题下方 中文文本 小四号宋体
        self.check_discussion_body_para()   # -1 讨论页标题下方 文本段落 1.5倍行距+首行缩进2字符
        self.check_conclusion_title_font()  # -1 结论页标题 四号黑体
        self.check_conclusion_title_para()  # -1 结论页标题段落 居中 + 1.5倍行距
        self.check_conclusion_body_font()   # -1 结论页标题下方 中文文本 小四号宋体
        self.check_conclusion_body_para()   # -1 结论页标题下方 文本段落 1.5倍行距+首行缩进2字符

    def add_score(self, score, description, item_type="score"):
        """添加得分或扣分"""
        self.total_score += score
        self.hit_items.append({"分数": score, "描述": description, "类型": item_type})

        if score > 0:
            self.evaluation_details["维度2"]["得分项"].append({
                "分数": score,
                "描述": description
            })
            print("  [命中] +" + str(score) + "分:", description)
        else:
            self.evaluation_details["维度2"]["扣分项"].append({
                "分数": score,
                "描述": description
            })
            print("  [命中] " + str(score) + "分:", description)

    def build_dim2_items(self):
        """按契约逐条输出维度二规则，未命中项也必须保留。"""

        def literal_text(node):
            if isinstance(node, ast.Constant) and isinstance(node.value, str):
                return node.value
            if isinstance(node, ast.BinOp) and isinstance(node.op, ast.Add):
                left = literal_text(node.left)
                return left if left is not None else None
            if isinstance(node, ast.JoinedStr):
                parts = []
                for value in node.values:
                    if isinstance(value, ast.Constant) and isinstance(value.value, str):
                        parts.append(value.value)
                    else:
                        break
                return ''.join(parts) if parts else None
            return None

        def normalize_rule(method_name, score, rule):
            # check_auto_toc 的 3 个 add_score 调用是同一条 +5 规则的不同证据路径。
            if method_name == 'check_auto_toc':
                return 'check_auto_toc', 5, '自动更新目录'
            return method_name + ':' + rule, score, rule

        source_path = os.path.abspath(__file__)
        with open(source_path, 'r', encoding='utf-8') as f:
            tree = ast.parse(f.read())

        rules = []
        rules_by_id = {}
        for class_node in tree.body:
            if not isinstance(class_node, ast.ClassDef) or class_node.name != 'WordDocumentEvaluator':
                continue
            for func_node in class_node.body:
                if not isinstance(func_node, ast.FunctionDef):
                    continue
                if not (func_node.name.startswith('check_') or func_node.name.startswith('_check_')):
                    continue
                if func_node.name in {'check_dimension1'}:
                    continue
                for node in ast.walk(func_node):
                    if not (
                        isinstance(node, ast.Call)
                        and isinstance(node.func, ast.Attribute)
                        and node.func.attr == 'add_score'
                        and len(node.args) >= 2
                    ):
                        continue
                    score_node = node.args[0]
                    if isinstance(score_node, ast.Constant) and isinstance(score_node.value, int):
                        score = score_node.value
                    elif (
                        isinstance(score_node, ast.UnaryOp)
                        and isinstance(score_node.op, ast.USub)
                        and isinstance(score_node.operand, ast.Constant)
                        and isinstance(score_node.operand.value, int)
                    ):
                        score = -score_node.operand.value
                    else:
                        continue
                    rule = literal_text(node.args[1])
                    if not rule:
                        continue
                    rule_id, max_delta, rule_text = normalize_rule(func_node.name, score, rule)
                    existing = rules_by_id.get(rule_id)
                    alias = [rule] if rule_text != rule else [rule_text]
                    if existing is None:
                        entry = {
                            'id': rule_id,
                            'rule': rule_text,
                            'max_delta': max_delta,
                            'aliases': alias,
                        }
                        rules_by_id[rule_id] = entry
                        rules.append(entry)
                    else:
                        for item in alias:
                            if item not in existing['aliases']:
                                existing['aliases'].append(item)

        hits = []
        for item in self.evaluation_details["维度2"]["得分项"] + self.evaluation_details["维度2"]["扣分项"]:
            hits.append({"score": item["分数"], "description": item["描述"]})

        dim2_items = []
        for rule in rules:
            matched = None
            for hit in hits:
                for alias in rule['aliases']:
                    if hit['description'] == alias or hit['description'].startswith(alias):
                        matched = hit
                        break
                if matched:
                    break
            is_hit = matched is not None
            dim2_items.append({
                "rule": rule['rule'],
                "max_delta": rule['max_delta'],
                "delta": rule['max_delta'] if is_hit else 0,
                "hit": is_hit,
                "detail": matched["description"] if matched else "",
            })
        return dim2_items

    def check_auto_toc(self):
        """检查目录是否为自动更新目录 +5
        细则三点：
          1) 目录为Word自动目录域或可更新目录对象
          2) 目录可通过更新域刷新页码
          3) 目录不是普通手打文本或截图
        在Word/WPS等办公软件中，自动目录在OOXML中的存在形式：
          A. TOC域：<w:fldChar w:fldCharType="begin"/> + <w:instrText>TOC ...</w:instrText>
                    + <w:fldChar w:fldCharType="end"/>
             或简单域：<w:fldSimple w:instr="TOC ..."/>
          B. SDT可更新目录对象：<w:sdt> 内含 <w:docPartGallery w:val="Table of Contents"/>
             或 <w:sdt> 内嵌 TOC 域
        只要命中A或B，即为"可通过F9/右键更新域刷新页码"的自动目录；
        否则视为手打文本或截图，不给分。
        """
        body = self.doc.element.body

        # A. 检测 TOC 域（含 fldChar 复杂域 与 fldSimple 简单域）
        has_toc_field = False
        for instr in body.iter(qn('w:instrText')):
            txt = instr.text or ''
            # TOC 指令通常形如：TOC \o "1-3" \h \z \u
            if re.search(r'\bTOC\b', txt):
                has_toc_field = True
                break
        if not has_toc_field:
            for fld in body.iter(qn('w:fldSimple')):
                instr = fld.get(qn('w:instr')) or ''
                if re.search(r'\bTOC\b', instr):
                    has_toc_field = True
                    break

        # B. 检测 SDT 可更新目录对象（内容控件目录）
        has_toc_sdt = False
        for sdt in body.iter(qn('w:sdt')):
            # B1. docPartGallery 标记为目录库
            gallery = None
            for g in sdt.iter(qn('w:docPartGallery')):
                gallery = g
                break
            if gallery is not None:
                val = gallery.get(qn('w:val')) or ''
                if 'Table of Contents' in val:
                    has_toc_sdt = True
                    break
            # B2. SDT 内嵌 TOC 域
            embedded_toc = False
            for instr in sdt.iter(qn('w:instrText')):
                if instr.text and re.search(r'\bTOC\b', instr.text):
                    embedded_toc = True
                    break
            if not embedded_toc:
                for fld in sdt.iter(qn('w:fldSimple')):
                    instr = fld.get(qn('w:instr')) or ''
                    if re.search(r'\bTOC\b', instr):
                        embedded_toc = True
                        break
            if embedded_toc:
                has_toc_sdt = True
                break

        if has_toc_field and has_toc_sdt:
            self.add_score(5, "自动更新目录（检测到SDT目录对象且内含TOC域，可在Word/WPS中通过更新域刷新页码）")
        elif has_toc_field:
            self.add_score(5, "自动更新目录（检测到Word自动目录TOC域，可通过更新域刷新页码）")
        elif has_toc_sdt:
            self.add_score(5, "自动更新目录（检测到可更新目录对象SDT，可通过更新域刷新页码）")
        else:
            print("  [未命中] 未检测到Word自动目录域或可更新目录对象（疑为手打文本或截图，不可通过更新域刷新页码）")

    def check_toc_page(self):
        """检查文中是否出现目录页 +3
        细则要求（唯一一点）：文中出现目录页。
        判定标准：文档中存在一个作为页面标题的"目录"独立段落即算命中。
        在Word/WPS办公软件中，"目录页"的核心呈现形式就是一个文本为"目录"
        （可能写成"目 录"，中间夹全/半角空格）的独立段落，作为该页页面标题。
        —— 只针对细则这一点判定，不额外约束该段落的字体、字号、样式、
        对齐方式、是否位于某一节、是否紧跟自动目录域等，凡细则未要求者一律不查。
        """
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 归一化：去除全/半角空格，允许"目 录"这类中间带空格的写法
            normalized = text.replace(' ', '').replace('　', '')
            # 独立段落且文本恰为"目录"，即为目录页标题
            if normalized == '目录':
                self.add_score(3, "文中出现目录页")
                return
        print("  [未命中] 未检测到目录页")

    def check_toc_levels(self):
        """检查目录标题层级 +5
        细则三点，全部同时满足才给 +5，缺任意一点不给分（细则未列出部分给分）：
          1) 目录包含一级、二级标题
          2) 一级标题采用"第一部分、第二部分、第三部分"的排序格式（即"第X部分"）
          3) 二级标题采用"一、二、三"的排序格式（即"X、"，X 为汉字数字）
        —— 只针对细则这三点判定；细则未要求的内容（字体、字号、缩进、页码
             位置、颜色、样式名等）一律不加以约束。

        办公软件（Word/WPS）适配说明：
          · 目录页在 OOXML 中通常为：一个内容为"目录"的标题段落（在 body 中），
            紧随其后是若干"目录条目段落"。自动目录的条目段落一般具有以下
            特征之一：段落样式 pStyle 为 TOC*/目录*（Word/WPS 自动目录样式）、
            段落内嵌 <w:hyperlink>（自动目录条目为可点击超链接）、或段落
            末尾跟随页码数字。手工排版的目录也一般以页码结尾。
          · 自动目录的条目段落常被包裹在 <w:sdt> 中，需从 body.iter('w:p')
            做 XML 级遍历才能取到，故此处不用 doc.paragraphs。
        """
        body = self.doc.element.body

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        # 全量遍历所有 w:p（含 SDT 内嵌目录条目段落）
        all_p_elements = list(body.iter(qn('w:p')))

        # 1) 定位"目录"标题段落
        toc_start = -1
        for i, p_elem in enumerate(all_p_elements):
            normalized = _p_text(p_elem).strip().replace(' ', '').replace('　', '')
            if normalized == '目录':
                toc_start = i
                break
        if toc_start < 0:
            print("  [未命中] 未检测到目录页，无法判定目录标题层级")
            return

        # 2) 收集目录条目：从"目录"之后开始，直到遇到第一个"非目录条目段落"停止
        toc_entries = []
        for p_elem in all_p_elements[toc_start + 1:]:
            text = _p_text(p_elem).strip()
            if not text:
                continue
            # 段落样式（pStyle 的 w:val）
            style_name = ''
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_name = pStyle.get(qn('w:val')) or ''
            style_lower = style_name.lower()
            is_toc_style = style_lower.startswith('toc') or style_name.startswith('目录') or style_lower.startswith('目录')
            # 是否含超链接（自动目录条目为点击可跳转的超链接）
            has_hyperlink = p_elem.find('.//' + qn('w:hyperlink')) is not None
            # 段落末尾是否以页码收尾
            ends_with_pagenum = bool(re.search(r'\d+\s*$', text))

            if is_toc_style or has_hyperlink or ends_with_pagenum:
                toc_entries.append(text)
            else:
                # 遇到不属于目录条目的段落，目录区域结束
                break

        if not toc_entries:
            print("  [未命中] 目录页下方未检出目录条目，无法判定层级")
            return

        # 3) 在目录区域内检查一级、二级标题格式
        has_level1 = any(re.search(r'第[一二三四五六七八九十]+部分', t) for t in toc_entries)
        has_level2 = any(re.match(r'^[一二三四五六七八九十]+、', t) for t in toc_entries)

        if has_level1 and has_level2:
            self.add_score(5, "目录包含一级、二级标题；一级采用\"第X部分\"格式、二级采用\"X、\"格式")
        else:
            missing = []
            if not has_level1:
                missing.append("一级标题\"第X部分\"格式")
            if not has_level2:
                missing.append("二级标题\"X、\"格式")
            print("  [未命中] 目录标题层级不满足要求（缺少" + "、".join(missing) + "）")

    def check_toc_core_items(self):
        """检查目录核心条目 +3
        细则一点（唯一一点）：目录中包含以下 8 个一级标题（缺一不可）：
          1) 中英文缩略词表
          2) 前言
          3) 第一部分 多模态康复资料库构建
          4) 第二部分 中医综合干预响应特征
          5) 第三部分 疗效预测模型建立与解释
          6) 讨论
          7) 结论
          8) 参考文献
        —— 只针对细则这一点判定：8 项**全部**在目录中出现才给 +3；
             缺任意一项不给分（细则未列出部分给分档位）。
             细则未要求的内容（字体、字号、页码、样式、颜色、对齐等）一律不查。

        办公软件（Word/WPS）适配说明：
          · 目录条目段落在 OOXML 中常被包裹于 <w:sdt>，doc.paragraphs 会遗漏，
            因此从 body.iter('w:p') 做 XML 级遍历，取每个 w:p 内所有 w:t 的
            拼接文本，等价于 Word/WPS 中该段落显示的可见文字。
          · 目录条目末尾常带有页码（如 "前言\t1"、"第一部分 多模态康复资料库构建\t3"），
            匹配时先去掉尾部的空白、制表符与数字页码，再做子串比对，与用户
            在 Word/WPS 目录页肉眼看到的标题文字保持一致。
        """
        required_items = [
            "中英文缩略词表",
            "前言",
            "第一部分 多模态康复资料库构建",
            "第二部分 中医综合干预响应特征",
            "第三部分 疗效预测模型建立与解释",
            "讨论",
            "结论",
            "参考文献",
        ]

        body = self.doc.element.body

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        all_p_elements = list(body.iter(qn('w:p')))

        # 定位"目录"标题段落
        toc_start = -1
        for i, p_elem in enumerate(all_p_elements):
            normalized = _p_text(p_elem).strip().replace(' ', '').replace('　', '')
            if normalized == '目录':
                toc_start = i
                break
        if toc_start < 0:
            print("  [未命中] 未检测到目录页，无法判定目录核心条目")
            return

        # 收集目录条目段落（同 check_toc_levels 的判定口径）
        toc_entry_texts = []
        for p_elem in all_p_elements[toc_start + 1:]:
            text = _p_text(p_elem).strip()
            if not text:
                continue
            style_name = ''
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_name = pStyle.get(qn('w:val')) or ''
            style_lower = style_name.lower()
            is_toc_style = style_lower.startswith('toc') or style_name.startswith('目录') or style_lower.startswith('目录')
            has_hyperlink = p_elem.find('.//' + qn('w:hyperlink')) is not None
            ends_with_pagenum = bool(re.search(r'\d+\s*$', text))
            if is_toc_style or has_hyperlink or ends_with_pagenum:
                toc_entry_texts.append(text)
            else:
                break

        if not toc_entry_texts:
            print("  [未命中] 目录页下方未检出目录条目，无法判定核心条目")
            return

        # 逐条比对：去除尾部制表符、点/连字符引导线与页码后，做子串匹配
        def _strip_pagenum(t):
            # 去掉末尾的空白、制表符、点线/连字符引导线，以及页码数字
            return re.sub(r'[\s\.\-\·…\t]*\d+\s*$', '', t).strip()

        cleaned_entries = [_strip_pagenum(t) for t in toc_entry_texts]

        found = []
        missing = []
        for item in required_items:
            if any(item in entry for entry in cleaned_entries):
                found.append(item)
            else:
                missing.append(item)

        if len(missing) == 0:
            self.add_score(3, "目录核心条目齐全（含全部8项一级标题）")
        else:
            print("  [未命中] 目录核心条目不完整，缺少：" + "、".join(missing))

    def check_toc_format(self):
        """检查目录内容及标题格式 +3
        细则两点，两点同时满足才给 +3；缺任意一点不给分（细则未列出部分给分档位）：
          1) 目录条目统一为四号黑体（四号 = 14pt；黑体，中文用 eastAsia="黑体"）
          2) 目录条目右侧出现页码

        —— 只针对细则这两点判定，细则未要求的内容（颜色、对齐、缩进、
             引导线样式、行距、段前段后间距等）一律不加以约束。

        办公软件（Word/WPS）适配说明：
          · 目录条目常被 <w:sdt> 包裹，doc.paragraphs 会遗漏，改用
            body.iter('w:p') 做 XML 全量遍历。
          · 字体判定：五号/四号等中文字号在 OOXML 中以 w:rPr/w:sz 记录，
            单位为半磅（半点），四号=14pt 对应 sz=28。字号也可能在段落属性
            w:pPr/w:rPr/w:sz 上，或继承自样式；对每个可见 run 逐一检查，
            未显式设置时视为继承（不打破"统一"约束）。
          · 中文字体名通常在 w:rFonts 的 w:eastAsia 属性（"黑体"），
            西文字体名在 w:ascii；办公软件在字体框显示的即是这些值。
          · 页码判定：自动目录条目通常以数字页码（可能经过 PAGEREF 域）结尾，
            或段末含 w:tab 后跟数字页码；手工目录常直接以数字结尾。取段落
            拼接后文本，右侧（去空白后）为数字即视为"右侧出现页码"。
        """
        body = self.doc.element.body

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        all_p_elements = list(body.iter(qn('w:p')))

        # 定位"目录"标题段落
        toc_start = -1
        for i, p_elem in enumerate(all_p_elements):
            normalized = _p_text(p_elem).strip().replace(' ', '').replace('　', '')
            if normalized == '目录':
                toc_start = i
                break
        if toc_start < 0:
            print("  [未命中] 未检测到目录页，无法判定目录内容及标题格式")
            return

        # 收集目录条目段落（口径与前面两个目录相关函数一致）
        toc_entry_elems = []
        for p_elem in all_p_elements[toc_start + 1:]:
            text = _p_text(p_elem).strip()
            if not text:
                continue
            style_name = ''
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_name = pStyle.get(qn('w:val')) or ''
            style_lower = style_name.lower()
            is_toc_style = style_lower.startswith('toc') or style_name.startswith('目录') or style_lower.startswith('目录')
            has_hyperlink = p_elem.find('.//' + qn('w:hyperlink')) is not None
            ends_with_pagenum = bool(re.search(r'\d+\s*$', text))
            if is_toc_style or has_hyperlink or ends_with_pagenum:
                toc_entry_elems.append(p_elem)
            else:
                break

        if not toc_entry_elems:
            print("  [未命中] 目录页下方未检出目录条目，无法判定内容及标题格式")
            return

        # ---- 第 1 点：目录条目统一为四号(14pt)黑体 ----
        # 按办公软件（Word/WPS）真实呈现的字体来判定：以 XML 显式设置为最高
        # 优先级；未显式设置则沿 rStyle → 段落 pStyle → basedOn → Normal
        # 继承链解析。这与 Word/WPS 打开文档后字体框显示的最终字号/字体
        # 一致。若解析结果为 sz != 28 或 eastAsia ∉ {黑体, SimHei}，即未
        # "统一为四号黑体"，不给分。
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _run_effective_font(r_elem, para_style_id):
            """沿继承链解析 run 最终生效的 (sz, eastAsia)。
            优先级：run rPr > rStyle > 段落 pStyle 链 > Normal。"""
            rPr = r_elem.find(qn('w:rPr'))
            east = sz = None
            rStyle_id = None
            if rPr is not None:
                rs = rPr.find(qn('w:rStyle'))
                if rs is not None:
                    rStyle_id = rs.get(qn('w:val'))
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            if sz is None and rStyle_id:
                sz = _resolve(rStyle_id, 'sz')
            if east is None and rStyle_id:
                east = _resolve(rStyle_id, 'eastAsia')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'eastAsia')
            if sz is None:
                sz = _resolve('Normal', 'sz')
            if east is None:
                east = _resolve('Normal', 'eastAsia')
            return sz, east

        def _iter_visible_runs(p_elem):
            """遍历段落中所有可见 run（含 hyperlink 内嵌 run），跳过纯空白/仅制表符/仅数字页码尾部的 run。"""
            for r in p_elem.iter(qn('w:r')):
                text = ''.join((t.text or '') for t in r.iter(qn('w:t')))
                if not text.strip():
                    continue
                # 跳过完全是数字（页码）的 run —— 页码字体不在此项约束内
                if re.fullmatch(r'\s*\d+\s*', text):
                    continue
                yield r

        all_entries_font_ok = True
        bad_reason = None
        for p_elem in toc_entry_elems:
            pPr = p_elem.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None
            for r in _iter_visible_runs(p_elem):
                sz, east = _run_effective_font(r, para_style_id)
                probs = []
                if sz != '28':
                    probs.append(f"字号非四号(sz={sz})")
                if east not in ('黑体', 'SimHei'):
                    probs.append(f"中文字体非黑体(eastAsia={east})")
                if probs:
                    all_entries_font_ok = False
                    rtext = ''.join((t.text or '') for t in r.iter(qn('w:t')))
                    bad_reason = f"[{rtext[:20]}] " + "；".join(probs)
                    break
            if not all_entries_font_ok:
                break

        # ---- 第 2 点：目录条目右侧出现页码 ----
        all_entries_have_pagenum = True
        for p_elem in toc_entry_elems:
            text = _p_text(p_elem)
            # 去除末尾空白后应以数字结尾
            if not re.search(r'\d+\s*$', text):
                all_entries_have_pagenum = False
                break

        if all_entries_font_ok and all_entries_have_pagenum:
            self.add_score(3, "目录条目统一为四号黑体，且条目右侧出现页码")
        else:
            missing = []
            if not all_entries_font_ok:
                if bad_reason:
                    missing.append("条目未统一为四号黑体 —— " + bad_reason)
                else:
                    missing.append("条目未统一为四号黑体")
            if not all_entries_have_pagenum:
                missing.append("部分条目右侧无页码")
            print("  [未命中] 目录内容及标题格式不满足要求（" + "、".join(missing) + "）")

    def check_page_number_settings(self):
        """检查页码设置"""
        sections = self.doc.sections

        # ---- +3：页码设置（细则两点，两点同时满足才给 +3；缺任意一点不给分） ----
        #   1) 封面页不显示页码
        #   2) 目录页不显示页码
        # 仅针对细则这两点判定；细则未要求的内容（页码字体/字号/位置/数字类型
        # 等）不在本项约束内。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · 页码通过页脚/页眉中的 PAGE 域实现，OOXML 体现为：
        #       <w:instrText>PAGE</w:instrText>（复杂域，需配合 fldChar）
        #       <w:fldSimple w:instr=" PAGE ..."/>（简单域）
        #     只要页眉/页脚中存在上述 PAGE 域，Word/WPS 就会在该页显示页码。
        #   · 节可通过 titlePg（首页不同）为首页单独设置页脚 —— 封面页所在节
        #     若启用 titlePg，判定其"首页"是否显示页码需查 first_page_footer /
        #     first_page_header，而非默认 footer/header。
        #   · 通过遍历 body 直接子元素，用 w:pPr/w:sectPr 分节，将段落映射到
        #     所属节的索引，进而定位封面页与目录页所在节。
        self._check_cover_toc_no_page_number(sections)

        # ---- +3：中文摘要、英文摘要页使用罗马数字页码 ----
        # 细则两点，两点同时满足才给 +3；缺任意一点不给分（细则未列出部分给分档位）：
        #   1) 中文摘要页使用罗马数字页码
        #   2) 英文摘要页使用罗马数字页码
        # 仅针对细则这两点判定；细则未要求的内容（大小写、位置、字体、
        # 起始编号值等）一律不加以约束。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · 页码数字类型在 OOXML 中由 sectPr 下的
        #       <w:pgNumType w:fmt="upperRoman"/>（大写罗马 I、II、III）
        #       <w:pgNumType w:fmt="lowerRoman"/>（小写罗马 i、ii、iii）
        #     控制，Word/WPS "页码格式"对话框设置的即是该属性。
        #   · "使用罗马数字页码"隐含该页显示了页码，因此还需确认所在节
        #     的页眉/页脚含有 PAGE 域（否则页码根本不显示，谈不上"使用"）。
        self._check_abstract_roman_page_number()

        # ---- +3：论文从"中英文缩略词表"页开始（目录页后一页），页码重新从1开始排序并使用阿拉伯数字页码 ----
        # 细则三点，三点同时满足才给 +3；缺任意一点不给分（细则未列出部分给分档位）：
        #   1) 论文从"中英文缩略词表"页开始，且该页紧随目录页之后（目录页后一页）
        #   2) 页码从该页重新从 1 开始排序
        #   3) 该页起使用阿拉伯数字页码
        # 仅针对细则这三点判定；细则未要求的内容（页码位置、字体、字号、
        # 页码格式表达式等）一律不加以约束。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · 页码重启与数字类型由所在节的 sectPr/pgNumType 控制：
        #       <w:pgNumType w:start="1"/>        → 该节页码从 1 重新开始
        #       <w:pgNumType w:fmt="decimal"/>    → 阿拉伯数字（默认，未写即视为该值）
        #     Word/WPS "页码格式"对话框中的"起始页码=1"与"数字格式=1,2,3..."
        #     写入的即是这两个属性。
        #   · "目录页后一页"意味着缩略词表页所在的节，其索引应紧邻目录页所在
        #     节 +1（即两节相邻）。
        self._check_abbreviation_arabic_restart()

        # ---- +3：页码格式：页码统一在页面底部居中放置 ----
        # 细则两点，两点同时满足才给 +3；缺任意一点不给分（细则未列出部分给分档位）：
        #   1) 页码位于页面底部（即在页脚 w:ftr 中，而不是页眉 w:hdr）
        #   2) 页码居中放置（承载 PAGE 域的段落对齐方式为居中）
        # 且必须"统一"—— 全文档所有出现页码的位置都必须满足这两点。
        # 仅针对细则这两点判定；细则未要求的内容（页码字体、字号、颜色、
        # 是否带前后缀"第X页"等）一律不加以约束。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · 页码用 PAGE 域实现：<w:instrText>PAGE</w:instrText>（复杂域）
        #     或 <w:fldSimple w:instr="PAGE"/>（简单域）。
        #   · 页脚是 <w:ftr>，页眉是 <w:hdr>，分别对应 section.footer /
        #     section.header 三种类型（default/first/even）。PAGE 域出现在
        #     哪个部件里，Word/WPS 就在哪里显示页码。
        #   · 居中：承载 PAGE 域的段落 pPr/jc 值为 center。若 jc 缺省，则视
        #     为继承样式的对齐 —— 无法确认为"居中"，判为不满足。
        self._check_page_number_bottom_center()

    def _check_cover_toc_no_page_number(self, sections):
        """+3：封面页不显示页码 且 目录页不显示页码（细则两点，同时满足才给分）

        分节口径 —— 必须与 Word/WPS 呈现一致：
          · sectPr 可能出现在 `body > p > pPr > sectPr`（段落级，作为其
            所属节的结束标记）、也可能嵌在 `body > sdt > ... > p > pPr >
            sectPr` 内（如自动更新目录 SDT 控件里），或位于 `body > sectPr`
            （末节属性）。原来用 `list(body)` 只看直接子元素会漏掉 SDT
            内的 sectPr，导致 SDT 内的段落被错误归入下一节。
          · 现按 XML 文档顺序用 `body.iter('w:p')` 依次归段，用同一顺序
            的 `body.iter('w:sectPr')` 作为节列表；不依赖
            python-docx 的 `doc.sections`（后者也漏 SDT 内 sectPr）。
        """
        body = self.doc.element.body
        R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        # 按 XML 文档顺序收集所有 sectPr（含 SDT 等结构内嵌套）
        all_sect_prs = list(body.iter(qn('w:sectPr')))

        # 段落 -> 节索引 映射（按 XML 顺序；遇到段落级 sectPr 后切下一节）
        p_to_sect = {}
        sect_idx = 0
        for p in body.iter(qn('w:p')):
            # 跳过 header/footer 里的 p —— iter 只走 body 树，本身就不会包含它们
            p_to_sect[p] = sect_idx
            pPr = p.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                sect_idx += 1

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        # 定位目录页所在节
        toc_section_idx = None
        for p_elem, s_idx in p_to_sect.items():
            normalized = _p_text(p_elem).strip().replace(' ', '').replace('　', '')
            if normalized == '目录':
                toc_section_idx = s_idx
                break

        # 封面页所在节：文档首节
        cover_section_idx = 0

        page_pattern = re.compile(r'\bPAGE\b')
        part = self.doc.part

        def _ref_target_root(sPr, tag_local, ref_type):
            """取给定 sectPr 下指定类型（default/first/even）的
            footer/header rId 所指向的部件 XML root，找不到返回 None。
            """
            for ref in sPr:
                if ref.tag == qn('w:' + tag_local) and ref.get(qn('w:type')) == ref_type:
                    rid = ref.get('{%s}id' % R_NS)
                    if rid and rid in part.rels:
                        return part.rels[rid].target_part.element
            return None

        def _has_page_field(root):
            """判定给定页眉/页脚 XML root 中是否存在 PAGE 域"""
            if root is None:
                return False
            for instr in root.iter(qn('w:instrText')):
                if instr.text and page_pattern.search(instr.text):
                    return True
            for fld in root.iter(qn('w:fldSimple')):
                instr = fld.get(qn('w:instr')) or ''
                if page_pattern.search(instr):
                    return True
            return False

        def _title_pg(sPr):
            el = sPr.find(qn('w:titlePg'))
            if el is None:
                return False
            val = el.get(qn('w:val'))
            # 缺省 val 视为启用（OOXML on/off 语义）
            return val is None or val in ('1', 'true', 'True', 'on')

        def _first_page_shows_page_number(sPr):
            """判定 sectPr 所属节的首页是否显示页码"""
            if _title_pg(sPr):
                return (_has_page_field(_ref_target_root(sPr, 'footerReference', 'first'))
                        or _has_page_field(_ref_target_root(sPr, 'headerReference', 'first')))
            return (_has_page_field(_ref_target_root(sPr, 'footerReference', 'default'))
                    or _has_page_field(_ref_target_root(sPr, 'headerReference', 'default')))

        # ---- 判定封面页 ----
        cover_no_pagenum = True
        if cover_section_idx < len(all_sect_prs):
            cover_no_pagenum = not _first_page_shows_page_number(all_sect_prs[cover_section_idx])

        # ---- 判定目录页 ----
        toc_no_pagenum = False
        if toc_section_idx is not None and toc_section_idx < len(all_sect_prs):
            toc_sPr = all_sect_prs[toc_section_idx]
            if toc_section_idx == cover_section_idx:
                # 目录页与封面同节 —— 目录页非该节首页，其页码由默认页脚/页眉决定
                toc_no_pagenum = not (
                    _has_page_field(_ref_target_root(toc_sPr, 'footerReference', 'default'))
                    or _has_page_field(_ref_target_root(toc_sPr, 'headerReference', 'default'))
                )
            else:
                # 目录页在独立节：一般为该节首页
                toc_no_pagenum = not _first_page_shows_page_number(toc_sPr)

        if cover_no_pagenum and toc_no_pagenum:
            self.add_score(3, "封面页不显示页码，且目录页不显示页码")
        else:
            reasons = []
            if not cover_no_pagenum:
                reasons.append("封面页显示了页码")
            if toc_section_idx is None:
                reasons.append("未定位到目录页")
            elif not toc_no_pagenum:
                reasons.append("目录页显示了页码")
            print("  [未命中] 页码设置不满足要求（" + "、".join(reasons) + "）")

    def _check_abstract_roman_page_number(self):
        """+3：中文摘要、英文摘要页使用罗马数字页码（两页都必须使用；缺一不给）"""
        body = self.doc.element.body
        sections = self.doc.sections

        # 段落 -> 节索引 映射（同 _check_cover_toc_no_page_number）
        p_to_sect = {}
        sect_idx = 0
        for child in list(body):
            tag = child.tag
            if tag == qn('w:p'):
                p_to_sect[child] = sect_idx
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    sect_idx += 1
            elif tag == qn('w:tbl'):
                for p in child.iter(qn('w:p')):
                    p_to_sect[p] = sect_idx

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        # 定位中文摘要页与英文摘要页所在节
        cn_abstract_section = None
        en_abstract_section = None
        for p_elem, s_idx in p_to_sect.items():
            text = _p_text(p_elem).strip()
            normalized = text.replace(' ', '').replace('　', '')
            # 中文摘要页标题：独立段落为"摘要"或"中文摘要"
            if cn_abstract_section is None and normalized in ('摘要', '中文摘要'):
                cn_abstract_section = s_idx
                continue
            # 英文摘要页标题：独立段落为"Abstract"（不区分大小写）
            if en_abstract_section is None and normalized.lower() == 'abstract':
                en_abstract_section = s_idx

        page_pattern = re.compile(r'\bPAGE\b')

        def _has_page_field(hf):
            if hf is None:
                return False
            try:
                paragraphs = hf.paragraphs
            except Exception:
                return False
            for para in paragraphs:
                p_elem = para._element
                for instr in p_elem.iter(qn('w:instrText')):
                    if instr.text and page_pattern.search(instr.text):
                        return True
                for fld in p_elem.iter(qn('w:fldSimple')):
                    instr = fld.get(qn('w:instr')) or ''
                    if page_pattern.search(instr):
                        return True
            return False

        def _section_uses_roman_page_number(sec_idx):
            """该节是否 (1) 显示了页码 且 (2) 页码数字类型为罗马数字（upperRoman/lowerRoman）"""
            if sec_idx is None or sec_idx >= len(sections):
                return False
            section = sections[sec_idx]
            # (1) 是否显示了页码：任一页眉/页脚（含首页专用）含 PAGE 域
            shows_pagenum = (_has_page_field(section.footer)
                             or _has_page_field(section.header)
                             or _has_page_field(section.first_page_footer)
                             or _has_page_field(section.first_page_header))
            if not shows_pagenum:
                return False
            # (2) 页码数字类型：读该节 sectPr 下 pgNumType@fmt
            sectPr = section._sectPr
            pgNumType = sectPr.find(qn('w:pgNumType')) if sectPr is not None else None
            if pgNumType is None:
                return False
            fmt = pgNumType.get(qn('w:fmt'))
            return fmt in ('upperRoman', 'lowerRoman')

        cn_ok = _section_uses_roman_page_number(cn_abstract_section)
        en_ok = _section_uses_roman_page_number(en_abstract_section)

        if cn_ok and en_ok:
            self.add_score(3, "中文摘要、英文摘要页使用罗马数字页码")
        else:
            reasons = []
            if cn_abstract_section is None:
                reasons.append("未定位到中文摘要页")
            elif not cn_ok:
                reasons.append("中文摘要页未使用罗马数字页码")
            if en_abstract_section is None:
                reasons.append("未定位到英文摘要页")
            elif not en_ok:
                reasons.append("英文摘要页未使用罗马数字页码")
            print("  [未命中] 中英文摘要页罗马数字页码不满足要求（" + "、".join(reasons) + "）")

    def _check_abbreviation_arabic_restart(self):
        """+3：论文从"中英文缩略词表"页开始（目录页后一页），页码重新从1开始排序并使用阿拉伯数字页码
        细则三点，同时满足才给 +3：
          1) 中英文缩略词表页紧随目录页之后（相邻节）
          2) 该节页码 start=1
          3) 该节页码数字类型为阿拉伯数字（decimal，默认亦视为该值）

        分节口径：与 `_check_cover_toc_no_page_number` 一致，用
        `body.iter('w:p')` / `body.iter('w:sectPr')` 按 XML 文档顺序全量
        枚举，覆盖 SDT（自动目录控件）等结构内嵌套的 sectPr。原来的
        `list(body)` / `doc.sections` 会漏掉 SDT 内的 sectPr，导致目录
        节被吞并、缩略词表被错误归为与目录同节，误判为"不紧邻"。
        """
        body = self.doc.element.body

        # 按 XML 顺序收集所有 sectPr（含 SDT 内嵌套）
        all_sect_prs = list(body.iter(qn('w:sectPr')))

        # 段落 -> 节索引 映射（按 XML 顺序）
        p_to_sect = {}
        sect_idx = 0
        for p in body.iter(qn('w:p')):
            p_to_sect[p] = sect_idx
            pPr = p.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                sect_idx += 1

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        # 定位目录页与缩略词表页所在节
        toc_section = None
        abbr_section = None
        for p_elem, s_idx in p_to_sect.items():
            normalized = _p_text(p_elem).strip().replace(' ', '').replace('　', '')
            if toc_section is None and normalized == '目录':
                toc_section = s_idx
            if abbr_section is None and normalized == '中英文缩略词表':
                abbr_section = s_idx

        # 三点判定
        reasons = []
        # 1) 存在缩略词表页，且节紧随目录页之后
        if abbr_section is None:
            reasons.append("未定位到中英文缩略词表页")
        elif toc_section is None:
            reasons.append("未定位到目录页")
        elif abbr_section != toc_section + 1:
            reasons.append("中英文缩略词表页不在目录页后一页（未紧邻目录节）")

        start_ok = False
        fmt_ok = False
        if abbr_section is not None and abbr_section < len(all_sect_prs):
            sectPr = all_sect_prs[abbr_section]
            pgNumType = sectPr.find(qn('w:pgNumType')) if sectPr is not None else None
            # 2) 起始页码=1
            if pgNumType is not None and pgNumType.get(qn('w:start')) == '1':
                start_ok = True
            # 3) 阿拉伯数字：fmt 未写默认 decimal；显式写 decimal 也认
            if pgNumType is None:
                fmt_ok = True
            else:
                fmt = pgNumType.get(qn('w:fmt'))
                fmt_ok = (fmt is None or fmt == 'decimal')

        if not start_ok and abbr_section is not None:
            reasons.append("页码未从1重新开始")
        if not fmt_ok and abbr_section is not None:
            reasons.append("页码数字类型不是阿拉伯数字")

        if not reasons:
            self.add_score(3, "论文从中英文缩略词表页（目录页后一页）开始，页码重新从1排序并使用阿拉伯数字页码")
        else:
            print("  [未命中] " + "、".join(reasons))

    def _check_page_number_bottom_center(self):
        """+3：页码统一在页面底部居中放置（细则两点：底部 + 居中；必须统一）"""
        page_pattern = re.compile(r'\bPAGE\b')

        def _paragraph_has_page_field(p_elem):
            for instr in p_elem.iter(qn('w:instrText')):
                if instr.text and page_pattern.search(instr.text):
                    return True
            for fld in p_elem.iter(qn('w:fldSimple')):
                instr = fld.get(qn('w:instr')) or ''
                if page_pattern.search(instr):
                    return True
            return False

        def _paragraph_is_center(p_elem):
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                return False
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                return False
            return jc.get(qn('w:val')) == 'center'

        header_has_pagenum = False
        total_pagenum_paragraphs = 0
        non_center_count = 0

        for section in self.doc.sections:
            # 页脚（底部）：default/first/even
            for footer in (section.footer, section.first_page_footer, section.even_page_footer):
                if footer is None:
                    continue
                try:
                    paragraphs = footer.paragraphs
                except Exception:
                    continue
                for para in paragraphs:
                    p_elem = para._element
                    if _paragraph_has_page_field(p_elem):
                        total_pagenum_paragraphs += 1
                        if not _paragraph_is_center(p_elem):
                            non_center_count += 1

            # 页眉（顶部）：若出现 PAGE 域，则违反"底部"要求
            for header in (section.header, section.first_page_header, section.even_page_header):
                if header is None:
                    continue
                try:
                    paragraphs = header.paragraphs
                except Exception:
                    continue
                for para in paragraphs:
                    if _paragraph_has_page_field(para._element):
                        header_has_pagenum = True
                        break
                if header_has_pagenum:
                    break

        if total_pagenum_paragraphs == 0 and not header_has_pagenum:
            print("  [未命中] 未检出任何页码（无法判定页码格式）")
            return

        reasons = []
        if header_has_pagenum:
            reasons.append("页眉中出现页码（不在页面底部）")
        if non_center_count > 0:
            reasons.append("页脚中存在" + str(non_center_count) + "处页码未居中")

        if not reasons:
            self.add_score(3, "页码统一在页面底部居中放置")
        else:
            print("  [未命中] " + "、".join(reasons))

    def check_page_margins(self):
        """检查页面边距 +3
        细则一点（明确到数值）：全文页面边距为
            上 2.5cm、下 2.5cm、左 3.0cm、右 2.5cm
        —— 只针对细则这一点判定：必须"全文"（所有节）都满足这四个数值；
             任一节任一方向不符，即不给分。细则未要求的内容（页眉/页脚
             与页面边界的距离、装订线、纸张大小、方向等）一律不查。

        办公软件（Word/WPS）适配说明：
          · 页边距在 OOXML 中位于每节 sectPr 下的 <w:pgMar>：
              w:top / w:bottom / w:left / w:right ，单位为 twip
              （1 twip = 1/1440 英寸），2.5cm ≈ 1417 twip，3.0cm ≈ 1701 twip。
            Word/WPS "页面布局 → 页边距 → 自定义边距"面板中输入的即是这些值。
          · 用户在 Word/WPS 中通常以厘米为单位输入，办公软件按四舍五入
            存入 twip；反算回厘米会有 ±0.01cm 内的漂移。这里采用 0.02cm
            容差（约 ≤1 twip），仅用于吸收单位换算舍入，不放宽细则数值。
        """
        sections = self.doc.sections
        if not sections:
            print("  [未命中] 未检测到任何节，无法判定页面边距")
            return

        expected = {
            'top': 2.5,
            'bottom': 2.5,
            'left': 3.0,
            'right': 2.5,
        }
        tolerance = 0.02  # cm，仅吸收 twip↔cm 单位换算舍入

        bad_sections = []
        for idx, section in enumerate(sections):
            top = section.top_margin.cm if section.top_margin is not None else None
            bottom = section.bottom_margin.cm if section.bottom_margin is not None else None
            left = section.left_margin.cm if section.left_margin is not None else None
            right = section.right_margin.cm if section.right_margin is not None else None

            actual = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
            mismatched = []
            for name, exp in expected.items():
                v = actual[name]
                if v is None or abs(v - exp) > tolerance:
                    mismatched.append("{}={}cm(需{}cm)".format(
                        name, "None" if v is None else "{:.2f}".format(v), exp))
            if mismatched:
                bad_sections.append("第{}节: ".format(idx + 1) + ", ".join(mismatched))

        if not bad_sections:
            self.add_score(3, "全文页面边距符合要求（上2.5cm、下2.5cm、左3.0cm、右2.5cm）")
        else:
            print("  [未命中] 页面边距不满足要求 —— " + "；".join(bad_sections))

    def check_reference_font(self):
        """检查参考文献字体 +3
        细则三点，每一条参考文献条目都必须同时满足；任一条目任一点不符即不给分：
          1) 每一个条目中文文本字体采用宋体
          2) 英文和数字字体采用 Times New Roman
          3) 全部为小四号（= 12pt = w:sz val="24"）

        —— 只针对细则这三点判定；细则未要求的内容（颜色、加粗、倾斜、
             字符间距、下划线等）一律不加以约束。

        办公软件（Word/WPS）适配说明：
          · 中文字体在 OOXML 里是 <w:rFonts w:eastAsia="宋体"/>，
            西文字体是 <w:rFonts w:ascii="Times New Roman" w:hAnsi="…"/>；
            Word/WPS "字体"对话框里"中文字体""西文字体"两栏写入的正是
            eastAsia 与 ascii/hAnsi。
          · 字号 <w:sz w:val="24"/> 单位半磅，24 = 12pt = 小四号；
            西文字号可另设 w:szCs，本项要求同为小四，故一并按 24 校验。
          · 汉字与非汉字通常在同一 run 内混排，也可能被拆到不同 run；
            按"run 内字符类别"判定：run 若含汉字，中文字体必须为宋体；
            run 若含英文/数字，西文字体必须为 Times New Roman。
          · 参考文献条目段落以"[数字]"起始（细则另一项要求），据此界定
            "参考文献"标题之后的每一条目。
        """
        # 找到"参考文献"标题后的条目段落
        in_reference = False
        entry_paragraphs = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not in_reference:
                if text == '参考文献' or (text.startswith('参考文献') and len(text) < 10):
                    in_reference = True
                continue
            if not text:
                continue
            if re.match(r'^\[\d+\]', text):
                entry_paragraphs.append(para)

        if not entry_paragraphs:
            print("  [未命中] 未检测到参考文献条目")
            return

        han_re = re.compile(r'[一-鿿]')
        en_num_re = re.compile(r'[A-Za-z0-9]')

        def _run_attrs(run):
            rPr = run._element.find(qn('w:rPr'))
            east = ascii_f = hansi = None
            sz_val = None
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    east = rFonts.get(qn('w:eastAsia'))
                    ascii_f = rFonts.get(qn('w:ascii'))
                    hansi = rFonts.get(qn('w:hAnsi'))
                sz = rPr.find(qn('w:sz'))
                if sz is not None:
                    sz_val = sz.get(qn('w:val'))
            return east, ascii_f, hansi, sz_val

        bad = []  # (条目文本前缀, 原因)
        for para in entry_paragraphs:
            entry_key = para.text.strip()[:12]
            entry_bad = None
            for run in para.runs:
                rt = run.text
                if not rt.strip():
                    continue
                has_han = bool(han_re.search(rt))
                has_en_num = bool(en_num_re.search(rt))
                if not has_han and not has_en_num:
                    # 纯标点/空白 run，不作字体约束
                    continue
                east, ascii_f, hansi, sz_val = _run_attrs(run)

                # 字号：小四 = 24（半磅）；未显式设置视为继承样式，不判失败
                if sz_val is not None and sz_val != '24':
                    entry_bad = "字号非小四（sz={}）".format(sz_val)
                    break
                # 中文字体：run 含汉字，则 eastAsia 必须为"宋体"（未显式设置视为继承样式，不判失败）
                if has_han and east is not None and east.strip() != '宋体':
                    entry_bad = "中文字体非宋体（eastAsia={}）".format(east)
                    break
                # 西文字体：run 含英文/数字，则 ascii/hAnsi 必须为 Times New Roman
                if has_en_num:
                    if ascii_f is not None and ascii_f.strip().lower() != 'times new roman':
                        entry_bad = "西文字体非Times New Roman（ascii={}）".format(ascii_f)
                        break
                    if hansi is not None and hansi.strip().lower() != 'times new roman':
                        entry_bad = "西文字体非Times New Roman（hAnsi={}）".format(hansi)
                        break
            if entry_bad:
                bad.append("[" + entry_key + "…] " + entry_bad)

        if not bad:
            self.add_score(3, "参考文献条目字体：中文小四宋体，英文和数字小四Times New Roman")
        else:
            print("  [未命中] 参考文献条目字体不满足要求 —— " + "；".join(bad[:3]))

    def check_reference_numbering(self):
        """检查参考文献编号 +3
        细则一点（唯一一点）：参考文献列表编号使用英文方括号加阿拉伯数字，
            如"[1]""[2]"。
        —— 只针对细则这一点判定：参考文献列表中**每一条**条目均需以
             英文方括号（半角 [ ]）+ 阿拉伯数字（0-9）开头；任一条目
             不符（如使用全角"【】""［］"、圆括号、圆点、中文数字等）
             即不给分。细则未要求的内容（编号顺序连续、是否从 [1] 起、
             编号与后续内容之间的分隔符等）一律不加以约束。

        办公软件（Word/WPS）适配说明：
          · 参考文献的编号既可能是"手打编号"（run 内即为 "[1]" 文本），
            也可能通过自动编号（numbering.xml + numPr）呈现 —— Word/WPS
            自动编号在段落文本层面对读者是可见的，但 python-docx 的
            para.text 不会拼出编号字符。为兼容自动编号列表，本函数按
            以下顺序判定：
              (a) 段落 pPr 中含 numPr → 说明该条目使用了自动编号列表；
                  读取 numbering.xml 对应 abstractNum 的 lvlText，
                  确认其模板形如 "[%N]" 且 numFmt="decimal"。
              (b) 否则回退到 para.text 起始文本匹配 r'^\\[\\d+\\]'。
        """
        # 收集参考文献部分的所有段落
        in_reference = False
        entry_paragraphs = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not in_reference:
                if text == '参考文献' or (text.startswith('参考文献') and len(text) < 10):
                    in_reference = True
                continue
            if not text:
                continue
            entry_paragraphs.append(para)

        if not entry_paragraphs:
            print("  [未命中] 未检测到参考文献条目")
            return

        # 预加载 numbering.xml —— 用于识别自动编号的编号模板
        numbering_lookup = {}  # (numId, ilvl) -> (numFmt, lvlText)
        try:
            numbering_part = self.doc.part.numbering_part
        except Exception:
            numbering_part = None
        if numbering_part is not None:
            num_el = numbering_part.element
            # 建立 numId -> abstractNumId
            num_to_abs = {}
            for num in num_el.iter(qn('w:num')):
                num_id = num.get(qn('w:numId'))
                abs_ref = num.find(qn('w:abstractNumId'))
                if num_id is not None and abs_ref is not None:
                    num_to_abs[num_id] = abs_ref.get(qn('w:val'))
            # 建立 abstractNumId -> {ilvl: (fmt, lvlText)}
            abs_levels = {}
            for abs_num in num_el.iter(qn('w:abstractNum')):
                abs_id = abs_num.get(qn('w:abstractNumId'))
                levels = {}
                for lvl in abs_num.iter(qn('w:lvl')):
                    ilvl = lvl.get(qn('w:ilvl'))
                    numFmt_el = lvl.find(qn('w:numFmt'))
                    lvlText_el = lvl.find(qn('w:lvlText'))
                    fmt = numFmt_el.get(qn('w:val')) if numFmt_el is not None else None
                    txt = lvlText_el.get(qn('w:val')) if lvlText_el is not None else None
                    levels[ilvl] = (fmt, txt)
                abs_levels[abs_id] = levels
            # 合并
            for num_id, abs_id in num_to_abs.items():
                for ilvl, (fmt, txt) in abs_levels.get(abs_id, {}).items():
                    numbering_lookup[(num_id, ilvl)] = (fmt, txt)

        def _paragraph_numbering(p_elem):
            """返回段落对应的 (numFmt, lvlText)；若非自动编号则返回 None。"""
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                return None
            numPr = pPr.find(qn('w:numPr'))
            if numPr is None:
                return None
            numId_el = numPr.find(qn('w:numId'))
            ilvl_el = numPr.find(qn('w:ilvl'))
            num_id = numId_el.get(qn('w:val')) if numId_el is not None else None
            ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'
            if num_id is None:
                return None
            return numbering_lookup.get((num_id, ilvl))

        # 逐条目校验
        bracket_num_re = re.compile(r'^\[\d+\]')
        # 自动编号的 lvlText 模板中，%N（N 为层级数字，一般是 1）表示该层级的数字
        lvltext_re = re.compile(r'^\[%\d+\]$')

        bad_entries = []
        for para in entry_paragraphs:
            text = para.text.strip()
            entry_key = text[:12]
            fmt_txt = _paragraph_numbering(para._element)
            if fmt_txt is not None:
                # 使用自动编号：模板必须为 [%N]，且 numFmt 必须为 decimal
                fmt, tmpl = fmt_txt
                if fmt != 'decimal' or tmpl is None or not lvltext_re.match(tmpl):
                    bad_entries.append("[" + entry_key + "…] 自动编号模板不是[%N]/decimal（fmt={}, tmpl={}）".format(fmt, tmpl))
                # 若合规则通过
            else:
                # 手打编号：段落起始文本必须匹配 ^\[\d+\]
                if not bracket_num_re.match(text):
                    bad_entries.append("[" + entry_key + "…] 未以英文方括号+阿拉伯数字开头")

        if not bad_entries:
            self.add_score(3, "参考文献列表编号使用英文方括号加阿拉伯数字")
        else:
            print("  [未命中] 参考文献编号格式不规范 —— " + "；".join(bad_entries[:3]))

    def check_reference_paragraph(self):
        """检查参考文献内容段落格式 +3
        细则三点，每一条参考文献条目段落都必须同时满足；任一条目任一点不符即不给分：
          1) 编号左对齐（即段落对齐方式为左对齐）
          2) 1.5 倍行距
          3) 文字悬挂缩进 2 字符

        —— 只针对细则这三点判定；细则未要求的内容（段前段后间距、
             首行缩进、编号后分隔符、字符间距、页面视图缩进等）一律
             不加以约束。

        办公软件（Word/WPS）适配说明：
          · 左对齐：段落 pPr/jc 值为 "left" 或 "start"，或未显式设置
            （Word/WPS 默认段落即左对齐）。
          · 1.5 倍行距：pPr/spacing 有两种等价表达：
              (a) lineRule="auto" 且 line="360"（240 twip = 单倍，360 = 1.5倍）
              (b) 直接由样式提供 line_spacing_rule = ONE_POINT_FIVE
            两者在办公软件中都渲染为"1.5倍行距"。
          · 悬挂缩进 2 字符：pPr/ind 中：
              w:hangingChars="200"（推荐；办公软件"字符"单位缩进即写此属性，
                                    200 表示 2 个字符宽度，与字号无关）
              或 w:hanging="480"（twip；2字符 ≈ 480 twip @小四12pt）
            办公软件"段落 → 特殊格式 → 悬挂缩进 → 2 字符"设置的即是
            hangingChars=200；用户以"磅/厘米"输入则写 hanging。
        """
        # 收集参考文献部分的条目段落
        in_reference = False
        entry_paragraphs = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not in_reference:
                if text == '参考文献' or (text.startswith('参考文献') and len(text) < 10):
                    in_reference = True
                continue
            if not text:
                continue
            if re.match(r'^\[\d+\]', text):
                entry_paragraphs.append(para)

        if not entry_paragraphs:
            print("  [未命中] 未检测到参考文献条目")
            return

        def _paragraph_left_aligned(p_elem, para):
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    val = jc.get(qn('w:val'))
                    return val in (None, 'left', 'start')
            # pPr 或 jc 未显式设置 —— 回退段落对象的对齐属性
            align = para.paragraph_format.alignment
            return align is None or align == WD_ALIGN_PARAGRAPH.LEFT

        def _paragraph_1p5_line_spacing(p_elem, para):
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    rule = sp.get(qn('w:lineRule'))
                    # 1.5 倍行距的原生写法
                    if line == '360' and (rule is None or rule == 'auto'):
                        return True
                    # 若显式写了别的 line/规则，则不算
                    if line is not None or rule is not None:
                        return False
            # 回退到段落对象（可能来自样式）
            pf = para.paragraph_format
            if pf.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                return True
            if pf.line_spacing is not None:
                try:
                    return abs(float(pf.line_spacing) - 1.5) < 0.01
                except Exception:
                    return False
            return False

        def _paragraph_hanging_2char(p_elem):
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                return False
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                return False
            # 优先：字符单位（办公软件"悬挂缩进 2 字符"）
            hc = ind.get(qn('w:hangingChars'))
            if hc is not None:
                try:
                    return int(hc) == 200
                except Exception:
                    return False
            # 次选：twip 单位（2字符 ≈ 480 twip @12pt）
            h = ind.get(qn('w:hanging'))
            if h is not None:
                try:
                    return abs(int(h) - 480) <= 20  # ±20 twip 容差
                except Exception:
                    return False
            return False

        bad_entries = []
        for para in entry_paragraphs:
            entry_key = para.text.strip()[:12]
            p_elem = para._element
            reasons = []
            if not _paragraph_left_aligned(p_elem, para):
                reasons.append("非左对齐")
            if not _paragraph_1p5_line_spacing(p_elem, para):
                reasons.append("非1.5倍行距")
            if not _paragraph_hanging_2char(p_elem):
                reasons.append("非悬挂缩进2字符")
            if reasons:
                bad_entries.append("[" + entry_key + "…] " + "、".join(reasons))

        if not bad_entries:
            self.add_score(3, "参考文献内容段落格式满足：编号左对齐、1.5倍行距、文字悬挂缩进2字符")
        else:
            print("  [未命中] 参考文献段落格式不满足要求 —— " + "；".join(bad_entries[:3]))

    def check_cover_title(self):
        """检查封面页论文标题字体格式 +1
        细则两点，两点同时满足才给 +1；缺任意一点不给分（细则未列出部分给分档位）：
          1) 字号为小二号（= 18pt = w:sz val="36"，半磅单位）
          2) 字体为楷体（中文字体 w:eastAsia = "楷体" 或 "楷体_GB2312"）

        —— 只针对细则这两点判定；细则未要求的内容（颜色、加粗、倾斜、
             字符间距、对齐方式、下划线、段落间距等）一律不加以约束。

        办公软件（Word/WPS）适配说明：
          · 字号：<w:sz w:val="36"/> 单位半磅，36 = 18pt = 小二号；
            西文字号 <w:szCs w:val="36"/> 一并按 36 校验（标题若含 CJK 与
            西文混排时，办公软件按此分别渲染）。
          · 中文字体：<w:rFonts w:eastAsia="楷体"/> 是办公软件"字体"对话框
            "中文字体"栏写入的属性；"楷体_GB2312" 为常见变体（老版本 Windows
            默认名），两者在 Word/WPS 中均渲染为楷体外观，一并接受。
          · 论文标题定位：封面页（首节）中的"论文标题"通常是一段较长的
            CJK 主标题段落，位于"硕士学位论文"等条幅字样之后、摘要/署名信息
            之前。此处按"第一段长度≥15 且主要为汉字，且非条幅/信息栏标签"
            识别 —— 与办公软件中封面页人眼所见的论文标题一致。
        """
        # 定位封面页（首节）段落范围：以第一个 sectPr 出现前的段落为界
        body = self.doc.element.body
        first_section_para_elems = []
        for child in list(body):
            if child.tag != qn('w:p'):
                continue
            first_section_para_elems.append(child)
            pPr = child.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                break

        # 已知的封面页条幅/信息栏标签 —— 不视为论文标题
        LABELS = {
            '硕士学位论文', '博士学位论文', '学位论文', '毕业论文',
            '学院名称', '学科专业', '学位类别', '作者姓名', '指导教师',
            '完成日期', '答辩日期', '培养单位', '研究方向', '学号',
        }

        han_re = re.compile(r'[一-鿿]')

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        title_para_elem = None
        for p_elem in first_section_para_elems:
            text = _p_text(p_elem).strip()
            if len(text) < 15:
                continue
            # 排除条幅/信息栏标签
            if text in LABELS:
                continue
            if any(text.startswith(lbl) or text.endswith(lbl) for lbl in LABELS):
                continue
            # 主要为汉字（用于把英文副标题排除在外）
            han_chars = len(han_re.findall(text))
            if han_chars < max(10, int(len(text) * 0.5)):
                continue
            title_para_elem = p_elem
            break

        if title_para_elem is None:
            print("  [未命中] 未定位到封面页论文标题")
            return

        # 校验：所有可见 run 均须为小二(sz=36) 且中文楷体
        KAITI_NAMES = {'楷体', '楷体_GB2312'}
        size_ok = True
        font_ok = True
        reason = []

        has_visible_run = False
        for r in title_para_elem.iter(qn('w:r')):
            rt = ''.join((t.text or '') for t in r.iter(qn('w:t')))
            if not rt.strip():
                continue
            has_visible_run = True
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                # 未设置任何 run 属性 —— 视为继承样式，不判失败
                continue
            # 字号
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                v = sz.get(qn('w:val'))
                if v is not None and v != '36':
                    size_ok = False
                    reason.append("字号非小二（sz={}）".format(v))
                    break
            # 中文字体 eastAsia
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                east = rFonts.get(qn('w:eastAsia'))
                if east is not None and east.strip() not in KAITI_NAMES:
                    font_ok = False
                    reason.append("中文字体非楷体（eastAsia={}）".format(east))
                    break

        if not has_visible_run:
            print("  [未命中] 封面页论文标题段落无可见文本")
            return

        if size_ok and font_ok:
            self.add_score(1, "封面页论文标题字体格式为小二号、楷体")
        else:
            print("  [未命中] 封面页论文标题字体格式不符 —— " + "、".join(reason))

    def check_cover_info(self):
        """检查封面页学院/学科/学位类别/作者/指导教师/完成日期 字体格式 +5
        细则两点，需同时满足才给 +5；缺任意一点不给分（细则未列出部分给分档位）：
          1) 六项信息（学院名称、学科专业、学位类别、作者姓名、指导教师、
             完成日期）在封面页上出现（作为标签及其对应内容）
          2) 这六项信息中：
             · 中文文本 → 小二号（sz=36）楷体（eastAsia="楷体"或"楷体_GB2312"）
             · 英文/数字/字母 → 小二号（sz=36）Times New Roman
                                （ascii/hAnsi="Times New Roman"）

        —— 只针对细则这两点判定；细则未要求的内容（颜色、加粗、倾斜、
             字符间距、对齐、单元格边框、行距等）一律不加以约束。

        办公软件（Word/WPS）适配说明：
          · 封面页六项信息在实际文档中通常放在**表格**里（左列为标签，
            右列为内容），有时多个标签合并到同一单元格并用换行分隔。
            因此本函数同时遍历封面节所在的段落与表格。
          · 字号：<w:sz w:val="36"/> 半磅单位，36 = 18pt = 小二号；
            办公软件"字号"下拉框中"小二"即写此值。
          · 中文字体：<w:rFonts w:eastAsia="…"/> 是"字体"对话框"中文字体"
            栏；"楷体_GB2312" 为常见变体，与"楷体"渲染一致，一并接受。
          · 西文字体：<w:rFonts w:ascii="…" w:hAnsi="…"/> 是"字体"对话框
            "西文字体"栏，两个属性都必须为 "Times New Roman"。
          · 判定粒度："标签行"与"内容行"（即左单元格与对应右单元格）中
            所有可见 run 都必须按其字符类别符合上述字体/字号；未显式设置
            视为继承样式，不判失败。
        """
        LABELS = ['学院名称', '学科专业', '学位类别', '作者姓名', '指导教师', '完成日期']
        KAITI_NAMES = {'楷体', '楷体_GB2312'}
        han_re = re.compile(r'[一-鿿]')
        en_num_re = re.compile(r'[A-Za-z0-9]')

        # 收集封面页（首节）范围内的所有段落 —— 含表格中所有单元格的段落
        body = self.doc.element.body
        first_section_paragraphs = []  # list of paragraph XML elements
        for child in list(body):
            tag = child.tag
            if tag == qn('w:p'):
                first_section_paragraphs.append(child)
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    break
            elif tag == qn('w:tbl'):
                for p in child.iter(qn('w:p')):
                    first_section_paragraphs.append(p)
            elif tag == qn('w:sectPr'):
                break

        def _p_text(p_elem):
            return ''.join((t.text or '') for t in p_elem.iter(qn('w:t')))

        # 定位每个标签所在段落，及其"内容"段落（与标签同一表格行的右单元格；
        # 或标签合并在同一单元格时，其后紧邻的段落）
        info_paragraphs = {}  # label -> list of paragraph elements (label paragraph + content paragraphs)

        def _find_containing_cell(p_elem):
            """向上寻找段落所在的表格单元格 w:tc；返回 (tc, tr, table) 或 (None, None, None)"""
            cur = p_elem.getparent()
            tc = tr = tbl = None
            while cur is not None:
                if cur.tag == qn('w:tc') and tc is None:
                    tc = cur
                elif cur.tag == qn('w:tr') and tr is None:
                    tr = cur
                elif cur.tag == qn('w:tbl') and tbl is None:
                    tbl = cur
                    break
                cur = cur.getparent()
            return tc, tr, tbl

        for label in LABELS:
            hits = []
            for p_elem in first_section_paragraphs:
                if label in _p_text(p_elem):
                    hits.append(p_elem)
            if not hits:
                continue
            paras_for_label = []
            for p_elem in hits:
                paras_for_label.append(p_elem)
                tc, tr, tbl = _find_containing_cell(p_elem)
                if tc is not None and tr is not None:
                    # 取同一行内的其他单元格的所有段落作为内容
                    cells = list(tr.iter(qn('w:tc')))
                    label_cell_idx = None
                    for idx, c in enumerate(cells):
                        if c is tc:
                            label_cell_idx = idx
                            break
                    for idx, c in enumerate(cells):
                        if idx == label_cell_idx:
                            # 该单元格若含多个段落（如三项合并在同一单元格），
                            # 除标签段落外的其他段落也算内容
                            for cp in c.iter(qn('w:p')):
                                if cp is not p_elem and _p_text(cp).strip():
                                    paras_for_label.append(cp)
                        else:
                            for cp in c.iter(qn('w:p')):
                                if _p_text(cp).strip():
                                    paras_for_label.append(cp)
            info_paragraphs[label] = paras_for_label

        # 第 1 点：六项信息是否齐全
        missing_labels = [lbl for lbl in LABELS if lbl not in info_paragraphs]
        if missing_labels:
            print("  [未命中] 封面页信息不完整，缺少：" + "、".join(missing_labels))
            return

        # 第 2 点：字号/字体校验
        def _run_check(r_elem):
            """返回 (ok, reason)"""
            rt = ''.join((t.text or '') for t in r_elem.iter(qn('w:t')))
            if not rt.strip():
                return True, None
            has_han = bool(han_re.search(rt))
            has_en_num = bool(en_num_re.search(rt))
            if not has_han and not has_en_num:
                return True, None  # 纯标点/空白，不作约束
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                return True, None  # 完全继承样式，不判失败
            # 字号：显式设置时必须为 36
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                v = sz.get(qn('w:val'))
                if v is not None and v != '36':
                    return False, "字号非小二(sz={})".format(v)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                if has_han:
                    east = rFonts.get(qn('w:eastAsia'))
                    if east is not None and east.strip() not in KAITI_NAMES:
                        return False, "中文字体非楷体(eastAsia={})".format(east)
                if has_en_num:
                    ascii_f = rFonts.get(qn('w:ascii'))
                    hansi = rFonts.get(qn('w:hAnsi'))
                    if ascii_f is not None and ascii_f.strip().lower() != 'times new roman':
                        return False, "西文字体非TNR(ascii={})".format(ascii_f)
                    if hansi is not None and hansi.strip().lower() != 'times new roman':
                        return False, "西文字体非TNR(hAnsi={})".format(hansi)
            return True, None

        bad = []
        for label, paras in info_paragraphs.items():
            for p_elem in paras:
                for r in p_elem.iter(qn('w:r')):
                    ok, reason = _run_check(r)
                    if not ok:
                        bad.append(label + ": " + reason)
                        break
                if bad and bad[-1].startswith(label + ":"):
                    break

        if not bad:
            self.add_score(5, "封面页六项信息：中文小二号楷体，英文/数字小二号Times New Roman")
        else:
            print("  [未命中] 封面页信息字体格式不符 —— " + "；".join(bad[:6]))

    def check_abstract_cn(self):
        """检查中文摘要页"""
        # ---- +3：中文摘要页下侧内容按"[目的]"、"[方法]"、"[结果]"、"[结论]"分段书写；
        #        除数字、英文字母和"关键词"三字外其余字体为宋体小四号不加粗 ----
        # 细则两点，两点同时满足才给 +3；缺任意一点不给分（细则未列出部分给分档位）：
        #   1) 摘要页下侧内容按"[目的]""[方法]""[结果]""[结论]"**分段**书写：
        #      —— 四个带方括号的标记词全部出现，且每一个都作为**独立段落的起始**
        #        （即"分段"—— 一个段落对应一个分段块）。
        #   2) 该四段内容中，**除数字、英文字母和"关键词"三字外**（即中文/中文标点）：
        #      —— 字体为宋体（w:rFonts@eastAsia = "宋体"）
        #      —— 字号为小四号（w:sz val="24" ⇔ 12pt）
        #      —— 不加粗（w:b 缺失或 w:b val="false"/"0"）
        # 细则未要求的内容（英文/数字字体、"关键词"三字自身字体、行距、缩进、
        # 颜色、倾斜、中文标点是否用全角等）一律不加以约束。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · 段落"起始为标记"—— 段落拼接文本去空白后以 "[目的]" 等开头即可。
        #     方括号使用英文半角（细则原文所示）；同时容错常见全角变体
        #     "【目的】" —— 因为办公软件输入法自动转换后可能落入全角，但
        #     细则原文明确写 "[目的]" 半角，故只有半角计入命中；全角单独
        #     报错。
        #   · 中文摘要页范围：从段落文本严格等于"摘要"或"中文摘要"起，
        #     至"关键词"段落之前止；该范围外的段落不做本项判定。
        MARKERS = ['[目的]', '[方法]', '[结果]', '[结论]']
        han_re = re.compile(r'[一-鿿]')

        # 定位中文摘要页范围
        paragraphs = list(self.doc.paragraphs)
        start_idx = None
        end_idx = None
        for i, para in enumerate(paragraphs):
            text = para.text.strip().replace(' ', '').replace('　', '')
            if start_idx is None and text in ('摘要', '中文摘要'):
                start_idx = i
                continue
            if start_idx is not None and '关键词' in para.text:
                end_idx = i
                break
        if start_idx is None:
            print("  [未命中] 未检测到中文摘要页")
            return
        if end_idx is None:
            end_idx = len(paragraphs)

        section_paragraphs = paragraphs[start_idx + 1:end_idx]

        # 第 1 点：四个标记全部作为独立段落起始出现
        found_markers = {m: None for m in MARKERS}
        for para in section_paragraphs:
            text = para.text.strip()
            for m in MARKERS:
                if text.startswith(m) and found_markers[m] is None:
                    found_markers[m] = para

        missing = [m for m in MARKERS if found_markers[m] is None]
        skip_marker_font_check = bool(missing)
        if missing:
            print("  [未命中] 中文摘要页未按"
                  + "、".join(missing) + "分段书写")

        # 第 2 点：四段内容（每段从标记段起，至下一个标记段前止）中
        # 中文/中文标点的字体、字号、加粗校验
        if not skip_marker_font_check:
            markers_order = [found_markers[m] for m in MARKERS]
            boundary_indices = [section_paragraphs.index(p) for p in markers_order]
            # 按标记在段落序列中的位置排序，逐段范围内校验
            marker_ranges = []
            for i, p in enumerate(markers_order):
                start = section_paragraphs.index(p)
                # 下一个标记的位置为终点
                next_positions = [section_paragraphs.index(np) for np in markers_order[i + 1:]]
                end = min(next_positions) if next_positions else len(section_paragraphs)
                marker_ranges.append((start, end))

            bad = []
            for (start, end), marker in zip(marker_ranges, MARKERS):
                for para in section_paragraphs[start:end]:
                    for run in para.runs:
                        rt = run.text
                        if not rt.strip():
                            continue
                        # 仅约束"除数字、英文字母和'关键词'三字外"的字符 —— 即含有汉字/中文标点的 run
                        # 若 run 完全由英文/数字/空白/连续出现的"关键词"三字组成 → 跳过
                        # 注意："关键词"必须三字连续才豁免，单独的"关""键""词"仍受约束
                        non_en_num = re.sub(r'[A-Za-z0-9\s]+', '', rt.replace('关键词', ''))
                        if not non_en_num:
                            continue
                        rPr = run._element.find(qn('w:rPr'))
                        if rPr is None:
                            continue  # 完全继承样式，不判失败
                        # 中文字体：eastAsia = 宋体（未显式设置视为继承）
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            east = rFonts.get(qn('w:eastAsia'))
                            if east is not None and east.strip() != '宋体':
                                bad.append(marker + ": 中文字体非宋体(eastAsia=" + east + ")")
                                break
                        # 字号：小四 = sz val="24"（未显式设置视为继承）
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                            v = sz.get(qn('w:val'))
                            if v is not None and v != '24':
                                bad.append(marker + ": 字号非小四(sz=" + v + ")")
                                break
                        # 不加粗：w:b 缺失或 val="0"/"false"
                        b = rPr.find(qn('w:b'))
                        if b is not None:
                            val = b.get(qn('w:val'))
                            if val is None or str(val).lower() in ('1', 'true', 'on'):
                                bad.append(marker + ": 中文文本加粗")
                                break
                    if bad and bad[-1].startswith(marker + ":"):
                        break

            if not bad:
                self.add_score(3, "中文摘要页按[目的][方法][结果][结论]分段书写，除数字、英文字母和\"关键词\"三字外字体为宋体小四号不加粗")
            else:
                print("  [未命中] 中文摘要页格式不符 —— " + "；".join(bad[:4]))

        # ---- +1：中文摘要页"关键词"为四号黑体，"关键词"后的内容为宋体小四 ----
        # 细则两点，两点同时满足才给 +1：
        #   1) "关键词"三字（作为段首标签）：
        #        —— 中文字体为黑体（w:rFonts@eastAsia = "黑体"）
        #        —— 字号为四号（w:sz val="28" ⇔ 14pt）
        #   2) "关键词"之后的段落内容：
        #        —— 中文字体为宋体（w:rFonts@eastAsia = "宋体"）
        #        —— 字号为小四（w:sz val="24" ⇔ 12pt）
        # 细则未要求的内容（英文/数字字体、粗体、颜色、行距、缩进等）一律不加以约束。
        # "关键词"之后的内容中只对含汉字/中文标点的 run 施加字体约束
        #（因为细则的"宋体"是中文字体属性，数字/英文由 ascii/hAnsi 决定）。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · rPr 未显式设置视为继承 —— 未显式设置字体/字号不判为失败。
        #   · 关键词段落定位：仅在中文摘要页范围内（start_idx 起、含 end_idx 段），
        #     且段落拼接文本严格以"关键词"开头。
        #   · "关键词"三字与其后内容可能跨多个 run，按 run 文本累计位置切分。
        kw_para = None
        # 复用前文定位的中文摘要页范围
        search_range = paragraphs[start_idx:end_idx + 1] if end_idx < len(paragraphs) else paragraphs[start_idx:]
        for para in search_range:
            if para.text.lstrip().startswith('关键词'):
                kw_para = para
                break

        if kw_para is None:
            print("  [未命中] 中文摘要页未检测到\"关键词\"段落")
        else:
            # 定位 "关键词" 三字在段落文本中的位置（跳过前导空白）
            full_text = kw_para.text
            lead_ws = len(full_text) - len(full_text.lstrip())
            label_start = lead_ws          # "关键词" 起始字符位
            label_end = lead_ws + 3        # "关键词" 结束字符位（不含）

            label_bad = []
            content_bad = []

            pos = 0  # 累计字符位置
            for run in kw_para.runs:
                rt = run.text or ''
                run_start = pos
                run_end = pos + len(rt)
                pos = run_end
                if not rt:
                    continue
                rPr = run._element.find(qn('w:rPr'))

                # ---- 标签段判定："关键词"三字覆盖的 run（与 [label_start, label_end) 有重叠）
                if run_start < label_end and run_end > label_start:
                    # 该 run 有部分或全部文字属于 "关键词" 标签
                    label_seg = rt[max(0, label_start - run_start): label_end - run_start]
                    if label_seg.strip():
                        if rPr is None:
                            # 完全继承 —— Normal 样式 eastAsia=宋体、字号默认，不满足黑体+四号
                            label_bad.append("\"关键词\"字体/字号未显式设置(默认宋体小四)")
                        else:
                            # 中文字体：eastAsia = 黑体
                            rFonts = rPr.find(qn('w:rFonts'))
                            east = rFonts.get(qn('w:eastAsia')) if rFonts is not None else None
                            if east is None or east.strip() != '黑体':
                                label_bad.append("\"关键词\"中文字体非黑体(eastAsia=" + str(east) + ")")
                            # 字号：四号 = sz val="28"
                            sz = rPr.find(qn('w:sz'))
                            v = sz.get(qn('w:val')) if sz is not None else None
                            if v is None or v != '28':
                                label_bad.append("\"关键词\"字号非四号(sz=" + str(v) + ")")

                # ---- 内容段判定："关键词"之后（run_end > label_end）的部分
                if run_end > label_end:
                    content_seg = rt[max(0, label_end - run_start):]
                    # 仅约束含汉字/中文标点的部分（细则"宋体"是中文字体属性）
                    non_en_num = re.sub(r'[A-Za-z0-9\s]+', '', content_seg)
                    if not non_en_num:
                        continue
                    if rPr is None:
                        continue  # 完全继承样式（Normal eastAsia=宋体），不判失败
                    # 中文字体：eastAsia = 宋体（未显式设置视为继承）
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        east = rFonts.get(qn('w:eastAsia'))
                        if east is not None and east.strip() != '宋体':
                            content_bad.append("内容中文字体非宋体(eastAsia=" + east + ")")
                    # 字号：小四 = sz val="24"（未显式设置视为继承）
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        v = sz.get(qn('w:val'))
                        if v is not None and v != '24':
                            content_bad.append("内容字号非小四(sz=" + v + ")")

            if not label_bad and not content_bad:
                self.add_score(1, "中文摘要页\"关键词\"为四号黑体，其后内容为宋体小四")
            else:
                reasons = "；".join((label_bad + content_bad)[:4])
                print("  [未命中] 中文摘要页\"关键词\"字体格式不符 —— " + reasons)

    def check_tables(self):
        """检查表格是否为三线表 +5：顶线/底线1.5磅(sz=12)，栏目线0.75磅(sz=6)，无竖线无中间横线"""
        if not self.doc.tables:
            print("  [未命中] 未检测到表格")
            return

        def _cell_borders_ok(cell, is_first_row, is_last_row):
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is None:
                return False
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                return False

            top = tcBorders.find(qn('w:top'))
            bottom = tcBorders.find(qn('w:bottom'))
            left = tcBorders.find(qn('w:left'))
            right = tcBorders.find(qn('w:right'))
            inside_h = tcBorders.find(qn('w:insideH'))
            inside_v = tcBorders.find(qn('w:insideV'))

            # 顶线：仅第一行允许存在（1.5磅），其余行必须 none/nil
            if is_first_row:
                if top is None or top.get(qn('w:sz')) != '12':
                    return False
            elif top is not None and top.get(qn('w:val')) not in ('none', 'nil'):
                return False

            # 底线：第一行=栏目线(0.75磅)，最后一行=底线(1.5磅)，中间行必须 none/nil
            if is_first_row and is_last_row:
                # 单行表格同时是首行与末行时，底线按表格底线（1.5磅）为准
                if bottom is None or bottom.get(qn('w:sz')) != '12':
                    return False
            elif is_first_row:
                if bottom is None or bottom.get(qn('w:sz')) != '6':
                    return False
            elif is_last_row:
                if bottom is None or bottom.get(qn('w:sz')) != '12':
                    return False
            elif bottom is not None and bottom.get(qn('w:val')) not in ('none', 'nil'):
                return False

            # 其余边均应为 none/nil：无竖线、无中间横线
            for border_el in (left, right, inside_h, inside_v):
                if border_el is not None:
                    val = border_el.get(qn('w:val'))
                    if val not in ('none', 'nil'):
                        return False
            return True

        all_ok = True
        for table in self.doc.tables:
            rows = table.rows
            if len(rows) < 2:
                continue
            for r_idx, row in enumerate(rows):
                is_first_row = (r_idx == 0)
                is_last_row = (r_idx == len(rows) - 1)
                for cell in row.cells:
                    if not _cell_borders_ok(cell, is_first_row, is_last_row):
                        all_ok = False
                        break
                if not all_ok:
                    break
            if not all_ok:
                break

        if all_ok:
            self.add_score(5, "所有表格均以三线表形式呈现（顶线、底线1.5磅，栏目线0.75磅，无竖线、无中间横线）")
        else:
            print("  [未命中] 表格不符合标准三线表格式（顶线、底线需1.5磅，栏目线需0.75磅，无竖线、无中间横线）")

    def check_table_text(self):
        """检查表内文字格式 +5：
             中文=宋体；英文/数字=Times New Roman；
             表头(第一行)=五号(sz=21)且加粗；表身(其余行)=小五(sz=18)。
        """
        # 细则四点，同时满足才给 +5：
        #   1) 中文字体 = 宋体（w:rFonts@eastAsia = "宋体" 或英文名 "SimSun"）
        #   2) 英文/数字字体 = Times New Roman
        #      （w:rFonts@ascii 与 w:rFonts@hAnsi 均为 "Times New Roman"）
        #   3) 表头（第一行 = 栏目行）字号 = 五号（w:sz val="21" ⇔ 10.5pt）
        #      且加粗（w:b 存在且 val 非 "0"/"false"/"off"）
        #   4) 表身（除第一行外）字号 = 小五（w:sz val="18" ⇔ 9pt）
        # 细则未要求内容（对齐、行距、颜色、斜体、单元格底纹等）一律不加约束。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · 中文字体 —— 用户在办公软件里选择"宋体"时，
        #     eastAsia 属性可能被记录为 "宋体" 或英文名 "SimSun"（视 Office 版本），
        #     两者均视为满足。
        #   · Times New Roman 需同时匹配 ascii/hAnsi（办公软件里"西文字体"控制的属性）。
        #   · 属性继承 —— run 未显式设置的属性向段落样式、表格样式、Normal 继承。
        #   · 只对**含相应字符类别**的 run 施加对应字体约束：
        #     纯数字/英文 run 不检查 eastAsia；纯汉字 run 不检查 ascii/hAnsi。
        #   · 排除封面页布局表（不是"数据表格"）—— 通过表格内容匹配封面
        #     信息标签（学院名称/学科专业/学位类别/作者姓名/指导教师/完成日期）
        #     的表来识别；细则说的"表内文字"针对的是正文中的数据表格，
        #     否则会与"封面页信息为小二楷体"的细则相互矛盾。
        if not self.doc.tables:
            print("  [未命中] 未检测到表格")
            return

        # 构建样式 rPr 继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract_style_rpr(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = ascii_ = hansi = sz = bold = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                    ascii_ = rf.get(qn('w:ascii'))
                    hansi = rf.get(qn('w:hAnsi'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
                b = rPr.find(qn('w:b'))
                if b is not None:
                    bold = b.get(qn('w:val'))
                    if bold is None:
                        bold = '1'   # <w:b/> without val 表示加粗
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'east': east, 'ascii': ascii_, 'hAnsi': hansi,
                    'sz': sz, 'bold': bold, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract_style_rpr(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        cover_labels = {'学院名称', '学科专业', '学位类别',
                        '作者姓名', '指导教师', '完成日期'}

        def _is_cover_layout_table(table):
            hits = 0
            for row in table.rows:
                for cell in row.cells:
                    txt = ''.join(p.text for p in cell.paragraphs).strip()
                    if txt in cover_labels:
                        hits += 1
            return hits >= 3   # 出现 ≥3 个封面标签视为封面布局表

        han_re = re.compile(r'[一-鿿]')
        eng_num_re = re.compile(r'[A-Za-z0-9]')

        def _run_font_props(run, para_style_id):
            """解析 run 的有效字体属性（含继承）。"""
            rPr = run._element.find(qn('w:rPr'))
            east = ascii_ = hansi = sz = bold = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                    ascii_ = rf.get(qn('w:ascii'))
                    hansi = rf.get(qn('w:hAnsi'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
                b = rPr.find(qn('w:b'))
                if b is not None:
                    bold = b.get(qn('w:val'))
                    if bold is None:
                        bold = '1'
            # 沿段落样式 → Normal 继承
            for key, cur in (('east', east), ('ascii', ascii_),
                             ('hAnsi', hansi), ('sz', sz), ('bold', bold)):
                pass
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'east')
            if east is None:
                east = _resolve('Normal', 'east')
            if ascii_ is None and para_style_id:
                ascii_ = _resolve(para_style_id, 'ascii')
            if ascii_ is None:
                ascii_ = _resolve('Normal', 'ascii')
            if hansi is None and para_style_id:
                hansi = _resolve(para_style_id, 'hAnsi')
            if hansi is None:
                hansi = _resolve('Normal', 'hAnsi')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')
            if bold is None and para_style_id:
                bold = _resolve(para_style_id, 'bold')
            if bold is None:
                bold = _resolve('Normal', 'bold')
            return {'east': east, 'ascii': ascii_, 'hAnsi': hansi,
                    'sz': sz, 'bold': bold}

        def _bold_ok(val):
            # w:b 存在且 val 不为 "0"/"false"/"off"
            return val is not None and str(val).lower() not in ('0', 'false', 'off')

        SONG_ALIASES = {'宋体', 'SimSun'}

        bad = []
        data_tables_checked = 0

        for ti, table in enumerate(self.doc.tables):
            if _is_cover_layout_table(table):
                continue
            rows = table.rows
            if len(rows) < 2:
                continue
            data_tables_checked += 1

            for ri, row in enumerate(rows):
                is_header = (ri == 0)
                required_sz = '21' if is_header else '18'
                sz_label = '五号' if is_header else '小五'

                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        pPr = para._element.find(qn('w:pPr'))
                        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
                        para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

                        for run in para.runs:
                            rt = run.text
                            if not rt.strip():
                                continue
                            props = _run_font_props(run, para_style_id)

                            # ---- 字号 ----
                            if props['sz'] != required_sz:
                                bad.append(
                                    f"表{ti} r{ri}c{ci} {'表头' if is_header else '表身'}: "
                                    f"字号非{sz_label}(sz={props['sz']})")
                                break

                            # ---- 表头加粗 ----
                            if is_header and not _bold_ok(props['bold']):
                                bad.append(f"表{ti} r{ri}c{ci} 表头: 未加粗")
                                break

                            # ---- 中文字体 ----
                            if han_re.search(rt):
                                east = props['east']
                                if east is None or east.strip() not in SONG_ALIASES:
                                    bad.append(f"表{ti} r{ri}c{ci}: 中文字体非宋体(eastAsia={east})")
                                    break

                            # ---- 英文/数字字体 ----
                            if eng_num_re.search(rt):
                                if props['ascii'] != 'Times New Roman' or \
                                        props['hAnsi'] != 'Times New Roman':
                                    bad.append(
                                        f"表{ti} r{ri}c{ci}: 英文/数字字体非Times New Roman"
                                        f"(ascii={props['ascii']}, hAnsi={props['hAnsi']})")
                                    break
                        if bad:
                            break
                    if bad:
                        break
                if bad:
                    break
            if bad:
                break

        if data_tables_checked == 0:
            print("  [未命中] 未检测到数据表格")
            return

        if not bad:
            self.add_score(5, "表内文字格式：中文宋体、英数Times New Roman；表头五号加粗、表身小五号")
        else:
            print("  [未命中] 表内文字格式不符 —— " + "；".join(bad[:4]))

    def check_secondary_title(self):
        """检查二级标题字体格式 +3：小四号黑体"""
        # 细则两点，必须同时满足：
        #   1) 字号 = 小四（w:sz val="24" ⇔ 12pt）
        #   2) 中文字体 = 黑体（w:rFonts@eastAsia = "黑体"）
        # 细则未要求内容（英文/数字字体、加粗、对齐、行距、缩进、颜色等）
        # 一律不加以约束。
        #
        # 办公软件（Word/WPS）适配说明：
        #   · 二级标题识别 —— 沿用目录规则中"X、"（X 为汉字数字）作为
        #     二级标题的排序格式，即段落拼接文本 strip 后匹配
        #     ^[一二三四五六七八九十]+、
        #   · 属性继承 —— run 未显式设置的属性依次向段落样式、
        #     Normal 样式继承；因此若 run 未写字号/中文字体，
        #     取段落样式 → Normal 样式的对应值。
        #   · 检查范围排除目录（TOC）区域内的"X、"目录条目 —— 那是
        #     目录条目而非正文中的二级标题。
        han_num_re = re.compile(r'^[一二三四五六七八九十]+、')

        # 构建样式 id → rPr 的映射，用于继承回落
        styles_element = self.doc.styles.element
        style_rpr_map = {}   # styleId -> {'east': str|None, 'sz': str|None, 'basedOn': str|None}

        def _extract_style_rpr(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = None
            sz = None
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    east = rFonts.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'east': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_rpr_map[sid] = _extract_style_rpr(style_el)

        def _resolve_from_style(style_id, key):
            """沿 basedOn 链查找样式的 east/sz 值。"""
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_rpr_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        # 定位目录区域段落（要排除掉）
        toc_paras = set()
        body = self.doc.element.body
        # TOC 存在于 w:sdt/w:sdtContent 内的 w:p 或 fldChar/instrText 之间
        # 简化：把 SDT 内所有段落视为 TOC；同时把包含 PAGEREF 域指令的段落也视为 TOC 条目
        for sdt in body.iter(qn('w:sdt')):
            for p in sdt.iter(qn('w:p')):
                toc_paras.add(p)
        for p in body.iter(qn('w:p')):
            has_pageref = False
            for instr in p.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text:
                    has_pageref = True
                    break
            if has_pageref:
                toc_paras.add(p)

        # 遍历所有段落识别二级标题
        candidates = []  # list of (para, matched_text)
        for para in self.doc.paragraphs:
            if para._element in toc_paras:
                continue
            text = para.text.strip()
            if han_num_re.match(text):
                candidates.append(para)

        if not candidates:
            print("  [未命中] 未检测到二级标题")
            return

        bad = []
        for para in candidates:
            # 段落样式 id
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            for run in para.runs:
                rt = run.text
                if not rt.strip():
                    continue
                rPr = run._element.find(qn('w:rPr'))

                # ---- 有效字号 ----
                sz_val = None
                if rPr is not None:
                    s = rPr.find(qn('w:sz'))
                    if s is not None:
                        sz_val = s.get(qn('w:val'))
                if sz_val is None and para_style_id:
                    sz_val = _resolve_from_style(para_style_id, 'sz')
                if sz_val is None:
                    sz_val = _resolve_from_style('Normal', 'sz')

                if sz_val != '24':
                    bad.append(f"[{para.text.strip()[:12]}] 字号非小四(sz={sz_val})")
                    break

                # ---- 有效中文字体 ----
                east_val = None
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        east_val = rFonts.get(qn('w:eastAsia'))
                if east_val is None and para_style_id:
                    east_val = _resolve_from_style(para_style_id, 'east')
                if east_val is None:
                    east_val = _resolve_from_style('Normal', 'east')

                if east_val != '黑体':
                    bad.append(f"[{para.text.strip()[:12]}] 中文字体非黑体(eastAsia={east_val})")
                    break

        if not bad:
            self.add_score(3, "二级标题字体格式为小四号黑体")
        else:
            print("  [未命中] 二级标题字体格式不符 —— " + "；".join(bad[:4]))

    # ========== 扣分项检测 ==========

    def check_footer_text(self):
        """检查全文页脚 -3：出现除页码以外的页脚文本则扣分。

        细则一点：
          · 页脚中若存在**非页码文本**，扣 3 分。

        「页码」在办公软件（Word/WPS）里 = 页脚里通过 PAGE / NUMPAGES / SECTIONPAGES
        等域生成的动态数字（罗马、阿拉伯或其他 numFmt）。
        非域生成的任何可见字符（静态文字、静态数字、破折号、"第 X 页" 等
        没有对应域的字面量）——一律视为「除页码以外的页脚文本」。

        办公软件适配说明：
          · 每个 section 下有 default/first_page/even_page 三种页脚，仅
            "被 sectPr 内 w:footerReference 引用"的页脚才会实际显示。
            未被引用的页脚（python-docx 可能返回空对象或继承上一节）不计入。
          · 页码域可能以 w:fldSimple(单元) 或 w:fldChar(begin/separate/end 三段)
            两种方式写入 —— 两种都要正确剥离，只把「域指令」和「域渲染结果」
            排除后剩余的文本视为静态页脚文本。
          · 常用的页码域关键字：PAGE、NUMPAGES、SECTIONPAGES。
        """
        FOOTER_TYPES = ('default', 'first', 'even')
        PAGE_FIELD_RE = re.compile(r'\b(?:PAGE|NUMPAGES|SECTIONPAGES)\b')

        def _extract_static_text(footer_elem):
            """把页脚里所有非页码文本拼接返回（strip 后）。

            扫描每个段落：
              · w:fldSimple  —— 单元式域。instr 属性含 PAGE/NUMPAGES/SECTIONPAGES
                的视为页码域，内部文本整体忽略；其它域的渲染文本视为静态。
              · w:fldChar begin/separate/end —— 三段式域。同样按 instr 判定
                是否为页码域，判非页码则把 separate…end 之间的渲染文本视为静态。
              · 普通 w:t —— 直接视为静态。
            """
            PAGE_RE = PAGE_FIELD_RE
            static_chunks = []
            for para in footer_elem.iter(qn('w:p')):
                # 预扫：把 w:fldSimple 内的所有后代节点 id 记入 skip
                skip_ids = set()
                for fs in para.iter(qn('w:fldSimple')):
                    instr = fs.get(qn('w:instr')) or ''
                    is_page = bool(PAGE_RE.search(instr))
                    for desc in fs.iter():
                        if desc is fs:
                            continue
                        skip_ids.add(id(desc))
                    if not is_page:
                        for t in fs.iter(qn('w:t')):
                            static_chunks.append(t.text or '')

                # 主扫：处理三段式域和普通 w:t
                state = 'normal'   # normal | field_instr | field_result
                instr_buf = []
                result_buf = []
                for node in para.iter():
                    if id(node) in skip_ids:
                        continue
                    tag = node.tag
                    if tag == qn('w:fldChar'):
                        ct = node.get(qn('w:fldCharType'))
                        if ct == 'begin':
                            state = 'field_instr'
                            instr_buf = []
                            result_buf = []
                        elif ct == 'separate':
                            state = 'field_result'
                        elif ct == 'end':
                            instr = ''.join(instr_buf)
                            result = ''.join(result_buf)
                            if not PAGE_RE.search(instr):
                                static_chunks.append(result)
                            state = 'normal'
                            instr_buf = []
                            result_buf = []
                    elif tag == qn('w:instrText'):
                        if state == 'field_instr':
                            instr_buf.append(node.text or '')
                    elif tag == qn('w:t'):
                        if state == 'normal':
                            static_chunks.append(node.text or '')
                        elif state == 'field_result':
                            result_buf.append(node.text or '')

            return ''.join(static_chunks).strip()

        # 收集全文页脚里的静态文本
        offending = []   # list of (section_idx, footer_type, text_preview)

        for si, section in enumerate(self.doc.sections):
            sectPr = section._sectPr
            # 找到本节引用的 footer 类型集合
            referenced_types = set()
            for ref in sectPr.iter(qn('w:footerReference')):
                t = ref.get(qn('w:type')) or 'default'
                referenced_types.add(t)
            # 未显式引用的默认 footer：继承前节 —— 只在第 0 节且无引用时才可能是空
            if not referenced_types:
                continue

            for ftype in referenced_types:
                if ftype == 'default':
                    footer = section.footer
                elif ftype == 'first':
                    footer = section.first_page_footer
                elif ftype == 'even':
                    footer = section.even_page_footer
                else:
                    continue
                if footer is None:
                    continue
                # 若 footer 链接到前节，则渲染的是前节内容，不重复统计
                try:
                    if footer.is_linked_to_previous:
                        continue
                except Exception:
                    pass

                footer_elem = footer._element if hasattr(footer, '_element') else None
                if footer_elem is None:
                    # python-docx 的 _Footer 通过 _sectPr / part 获取，改用 part.element
                    part = getattr(footer, 'part', None)
                    footer_elem = part.element if part is not None else None
                if footer_elem is None:
                    continue

                static_txt = _extract_static_text(footer_elem)
                if static_txt:
                    offending.append((si, ftype, static_txt[:40]))

        if offending:
            preview = '；'.join(f"节{si}/{ft}: {tx!r}" for si, ft, tx in offending[:3])
            self.add_score(-3, f"全文出现除页码以外的页脚文本（{preview}）", "deduction")

    def check_page_number_overlap(self):
        """检查页码与正文、脚注或图表出现重叠 -3。

        细则一点：页码若与正文、脚注或图表出现重叠，扣 3 分。

        办公软件（Word/WPS）布局语义：
          页码位于页脚区域。页面版式属性（w:sectPr/w:pgMar）定义：
            · w:pgMar@bottom —— 正文底边到页面底边的距离（下页边距）
            · w:pgMar@footer —— 页脚底边到页面底边的距离（页脚离页面底部的距离）
          正文区域占 [top边距, page_height − bottom边距]；页脚从
          [page_height − footer距离] 处开始向下渲染其内容。
          若 footer_distance ≥ bottom_margin —— 页脚起点已进入正文
          区域，Word/WPS 会把页脚（含页码）叠在正文最后几行/表格/脚注上，
          屏幕与打印均可见重叠。这是办公软件里造成页码与正文/脚注/图表
          重叠最直接、可静态检测的成因。

        细则未要求内容（页脚字号、颜色、图片浮动锚点是否越界、脚注面积
        换算、正文尾行行距等）一律不加以约束。此检测只针对上述结构条件。
        """
        offending = []
        for si, section in enumerate(self.doc.sections):
            try:
                bottom = section.bottom_margin
                footer_dist = section.footer_distance
            except Exception:
                continue
            if bottom is None or footer_dist is None:
                continue
            bottom_cm = bottom.cm
            footer_cm = footer_dist.cm
            # 结构条件：页脚起点 ≥ 正文底边 → 页脚进入正文区域
            if footer_cm >= bottom_cm - 1e-6:
                offending.append((si, bottom_cm, footer_cm))

        if offending:
            preview = '；'.join(
                f"节{si}页脚距底{fc:.2f}cm≥下边距{bc:.2f}cm"
                for si, bc, fc in offending[:3])
            self.add_score(-3, f"页码与正文/脚注/图表可能重叠（{preview}）", "deduction")

    def check_reference_citation(self):
        """检查参考文献正文引用编号 -5。

        细则三点，任意一点不满足即扣 -5：
          1) 正文引用序号置于**方括号**内；
          2) 正文引用的序号以**上角标**形式显示；
          3) 序号使用**英文半角方括号**加**阿拉伯数字**（如 "[1]"、"[2]"）。

        办公软件（Word/WPS）适配说明：
          · 上角标 —— run 的 w:rPr/w:vertAlign val="superscript"；写入
            方式对应办公软件"字体 → 上标"或工具栏 x² 按钮。
          · 方括号形式 —— 英文半角 "[" / "]"，与全角 "［" "］"、
            中文方括号 "【" "】" 或圆括号 "(1)" "（1）" 相区分。
          · 正文范围 —— 检测起点为文档正文，终点为参考文献列表首行之前。
            参考文献列表条目本身以"[N]"起首，属于**列表编号**而非**正文引用**，
            不参与本项判定，否则会与"+3 参考文献列表编号"细则相互冲突。
          · 判定粒度 —— 只对**引用编号**这一 token 施加约束；细则未要求
            上角标的字体、字号、颜色、位置等一律不加约束。

        标记 token 匹配（在正文范围内）：
          · 合规形式：^[\\d](?:[,\\-–][\\d])*$ 内容置于 "[" 与 "]" 之间。
          · 违规变体：全角方括号 "［...］"、中文方括号 "【...】"、
            半/全角圆括号 "(N)" / "（N）"、上标数字（如 "①"）等
            —— 若在正文中作为引用标记出现，即视为不满足方括号要求。
          · 干扰排除：圆括号内数字在 references 类文本外还可能是"卷(期)"
            等排版；由于我们把范围限制在**参考文献列表之前的正文**内，
            该干扰基本不发生；仍额外要求圆括号紧邻中文语境（前一字符为
            汉字或中文标点），以避免误报公式/表格中出现的 (N)。
        """
        # 定位参考文献列表起点：段落文本 strip 后以 "[\\d" 起首、
        # 且整段像一条 references 条目（含期刊/年份/作者格式），
        # 或段前紧邻标题"参考文献"。为稳妥起见以"参考文献"标题为界。
        ref_start_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            t = para.text.strip()
            if t == '参考文献':
                ref_start_idx = i
                break
        body_paragraphs = self.doc.paragraphs[:ref_start_idx] \
            if ref_start_idx is not None else self.doc.paragraphs

        # 合规引用：英文半角方括号 + 阿拉伯数字（可含 , - – 组合形式）
        halfwidth_re = re.compile(r'\[\s*\d+(?:\s*[,\-–]\s*\d+)*\s*\]')
        # 违规括号变体
        fullwidth_sq_re = re.compile(r'［\s*\d+(?:\s*[,\-–]\s*\d+)*\s*］')
        cjk_sq_re = re.compile(r'【\s*\d+(?:\s*[,\-–]\s*\d+)*\s*】')
        # 圆括号引用：要求前一字符为汉字或中文标点，避免误捕获公式/单位
        cn_context_re = re.compile(r'[一-鿿，。；：、？！""''）】》]')
        halfwidth_paren_re = re.compile(r'\(\s*\d+\s*\)')
        fullwidth_paren_re = re.compile(r'（\s*\d+\s*）')

        has_any_citation = False
        wrong_bracket_examples = []
        not_superscript_examples = []

        for para in body_paragraphs:
            text = para.text
            if not text:
                continue

            # 违规括号变体 —— 只要在正文中出现即违反"方括号"要求
            for m in fullwidth_sq_re.finditer(text):
                has_any_citation = True
                if len(wrong_bracket_examples) < 3:
                    wrong_bracket_examples.append(f"全角方括号 {m.group()}")
            for m in cjk_sq_re.finditer(text):
                has_any_citation = True
                if len(wrong_bracket_examples) < 3:
                    wrong_bracket_examples.append(f"中文方括号 {m.group()}")
            for m in halfwidth_paren_re.finditer(text):
                if m.start() > 0 and cn_context_re.match(text[m.start() - 1]):
                    has_any_citation = True
                    if len(wrong_bracket_examples) < 3:
                        wrong_bracket_examples.append(f"英文圆括号 {m.group()}")
            for m in fullwidth_paren_re.finditer(text):
                if m.start() > 0 and cn_context_re.match(text[m.start() - 1]):
                    has_any_citation = True
                    if len(wrong_bracket_examples) < 3:
                        wrong_bracket_examples.append(f"全角圆括号 {m.group()}")

            # 合规形式 —— 检查上角标
            for m in halfwidth_re.finditer(text):
                has_any_citation = True
                # 找 match 起始字符所属 run
                pos = 0
                found_run = None
                for run in para.runs:
                    end = pos + len(run.text or '')
                    if pos <= m.start() < end:
                        found_run = run
                        break
                    pos = end
                if found_run is None:
                    continue
                rPr = found_run._element.find(qn('w:rPr'))
                vert = None
                if rPr is not None:
                    v = rPr.find(qn('w:vertAlign'))
                    if v is not None:
                        vert = v.get(qn('w:val'))
                if vert != 'superscript':
                    if len(not_superscript_examples) < 3:
                        not_superscript_examples.append(
                            f"{m.group()}(vertAlign={vert})")

        # 未出现任何形式的引用编号 —— 无法认定"不满足"，不扣分（细则前置：需存在引用）
        if not has_any_citation:
            return

        problems = []
        if wrong_bracket_examples:
            problems.append("非英文方括号: " + "、".join(wrong_bracket_examples))
        if not_superscript_examples:
            problems.append("非上角标: " + "、".join(not_superscript_examples))

        if problems:
            self.add_score(-5, "参考文献正文引用编号格式不符 —— " + "；".join(problems),
                           "deduction")

    def check_cover_images(self):
        """检查封面页图片。

        细则两点：
          1) 封面页未出现**校徽**和**校名**图片 → 扣 3 分；
          2) 封面页校徽图片大小不满足 **2.65×2.65cm** → 扣 1 分。

        办公软件（Word/WPS）语义：
          · 「封面页」= 文档正文首页 —— 从文档开头到**第一处硬分页**为止
            （w:br type="page"，或者段落 w:pPr 内 w:pageBreakBefore，
            或者节内首个 w:sectPr）。硬分页后即进入声明页/授权书等，
            那里出现的图片（如签名图）不属于封面。
          · 图片形态区分校徽 vs 校名（这两者是"办公软件上可见的图形对象"
            细则中并未定义特征，故按办公软件封面模板的常识特征区分）：
              - 校徽 —— 近似方形（宽高比 0.5–2.0），且尺寸较小（宽 ≤ 6cm）
              - 校名 —— 横向长条（宽高比 ≥ 2.0），且尺寸较大（宽 > 6cm）
            这两个形态特征互斥，不会同一张图同时满足；细则未要求的
            分辨率、颜色、位置、透明度、边框等一律不加约束。
          · 图形对象可能来自 w:drawing（DrawingML 现代形式）或 w:pict
            （VML 旧形式）。两者都需扫描。w:drawing 的尺寸由
            wp:extent@cx/cy（EMU）给出；VML 的尺寸不稳定（style 属性
            里的字符串），此时把 VML 图形一并纳入统计但只按"存在"计入
            两种候选（若无法判定 aspect，则允许其充当"校徽或校名之一"）。
        """
        body = self.doc.element.body

        cover_images = []      # 精确尺寸: (w_cm, h_cm)
        cover_images_unk = 0   # 无法测量尺寸的图片（VML 等）

        for child in list(body):
            if child.tag == qn('w:sectPr'):
                break
            if child.tag == qn('w:p'):
                # 段前 pageBreakBefore
                pPr = child.find(qn('w:pPr'))
                page_break_before = False
                if pPr is not None:
                    pbb = pPr.find(qn('w:pageBreakBefore'))
                    if pbb is not None:
                        val = pbb.get(qn('w:val'))
                        if val is None or str(val).lower() not in ('0', 'false', 'off'):
                            page_break_before = True

                if page_break_before:
                    break

                # 段内 w:br type="page"（第一处硬分页即结束）
                stop_after = False
                for br in child.iter(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        stop_after = True
                        break

                # 抓取本段图像
                for drw in child.iter(qn('w:drawing')):
                    for ext in drw.iter(qn('wp:extent')):
                        cx = int(ext.get('cx') or 0)
                        cy = int(ext.get('cy') or 0)
                        if cx and cy:
                            cover_images.append((cx / 914400 * 2.54,
                                                 cy / 914400 * 2.54))
                for _ in child.iter(qn('w:pict')):
                    cover_images_unk += 1

                # 段落结束时的 sectPr 也视为封面终止（节分隔亦划出封面页）
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    break
                if stop_after:
                    break
            elif child.tag == qn('w:tbl'):
                # 封面布局中若有表，图片可能嵌在单元格里
                for drw in child.iter(qn('w:drawing')):
                    for ext in drw.iter(qn('wp:extent')):
                        cx = int(ext.get('cx') or 0)
                        cy = int(ext.get('cy') or 0)
                        if cx and cy:
                            cover_images.append((cx / 914400 * 2.54,
                                                 cy / 914400 * 2.54))
                for _ in child.iter(qn('w:pict')):
                    cover_images_unk += 1

        has_badge = False
        has_name = False
        badge_candidates = []
        for w, h in cover_images:
            if h <= 0:
                continue
            aspect = w / h
            if 0.5 <= aspect <= 2.0 and w <= 6.0:
                has_badge = True
                badge_candidates.append((w, h))
            elif aspect >= 2.0 and w > 6.0:
                has_name = True

        # 无法测量尺寸的对象：若两类都还缺，把它当作"任一"以避免误判 VML 图
        if cover_images_unk:
            if not has_badge and not has_name:
                # 只有两张 VML —— 无法证明缺失，跳过扣分（宁失勿枉）
                if cover_images_unk >= 2:
                    has_badge = has_name = True
            elif not has_badge:
                has_badge = True
            elif not has_name:
                has_name = True

        missing = []
        if not has_badge:
            missing.append("校徽")
        if not has_name:
            missing.append("校名")
        if missing:
            self.add_score(-3, "封面页未出现" + "、".join(missing) + "图片", "deduction")

        if has_badge and badge_candidates:
            target_w, target_h, tol = 2.65, 2.65, 0.1
            badge_size_ok = any(abs(w - target_w) <= tol and abs(h - target_h) <= tol
                                for w, h in badge_candidates)
            if not badge_size_ok:
                sizes = "、".join(f"{w:.2f}×{h:.2f}cm" for w, h in badge_candidates)
                self.add_score(-1,
                               f"封面页校徽图片大小不满足2.65×2.65厘米（实际: {sizes}）",
                               "deduction")

    def check_declaration_signature_image(self):
        """检查"研究生学位论文原创性声明"页签名图片尺寸 -1。

        细则一点：声明页签名图片大小不满足 **3.84×11.74cm** → 扣 1 分。

        办公软件（Word/WPS）语义：
          · 「声明页」范围与 check_declaration_page 一致 —— 从"研究生学位
            论文原创性声明"标题段之后，到"版权使用授权书/授权声明"标题段
            之前（找不到后者则到文档末）。作者签名图（手写签名扫描件/
            插入的签名图片）通常紧跟声明正文出现在该范围内，不属于封面页
            （见 check_cover_images 说明：硬分页后即进入声明页，那里出现
            的图片如签名图不属于封面）。
          · 图片形态：签名图为长条形图片，近似 3.84×11.74cm；考虑到
            细则未明确宽高书写方向，3.84×11.74cm 与 11.74×3.84cm
            均视为尺寸满足。二者按所在页面区分（校名在封面页，签名图
            在声明页），不会互相混淆。
          · 图片对象仅取 w:drawing（DrawingML）的 wp:extent@cx/cy 精确
            测量；w:pict（VML）尺寸不稳定，无法精确判定，故不纳入本项
            校验（若声明页内只有 VML 图形，视为"无法验证"，不扣分，
            避免因测量工具局限误伤）。
          · 尺寸换算：EMU → 厘米，cm = emu / 914400 * 2.54（与
            check_cover_images 同一公式）。允许 ±0.1cm 容差，兼容
            换算/四舍五入误差。
          · 若整个文档未出现声明页，或声明页范围内没有任何可测量尺寸
            的图片，前置未满足，不扣分（无法判定是否存在签名图片）。
        """
        paragraphs = list(self.doc.paragraphs)
        start_idx = None
        end_idx = None
        for i, para in enumerate(paragraphs):
            t = para.text.strip().replace(' ', '').replace('　', '')
            if start_idx is None and '研究生学位论文原创性声明' in t:
                start_idx = i
                continue
            if start_idx is not None and ('版权使用授权书' in t or
                                          '版权使用授权' in t or
                                          '授权书' in t or
                                          '授权声明' in t):
                end_idx = i
                break

        if start_idx is None:
            return
        if end_idx is None:
            end_idx = len(paragraphs)

        sig_images = []  # (w_cm, h_cm)
        for para in paragraphs[start_idx + 1:end_idx]:
            for drw in para._element.iter(qn('w:drawing')):
                for ext in drw.iter(qn('wp:extent')):
                    cx = int(ext.get('cx') or 0)
                    cy = int(ext.get('cy') or 0)
                    if cx and cy:
                        sig_images.append((cx / 914400 * 2.54,
                                           cy / 914400 * 2.54))

        if not sig_images:
            # 声明页内未发现可测量尺寸的图片 —— 无法判定，不扣分
            return

        target_a, target_b, tol = 3.84, 11.74, 0.1
        ok = any((abs(w - target_a) <= tol and abs(h - target_b) <= tol) or
                 (abs(w - target_b) <= tol and abs(h - target_a) <= tol)
                 for w, h in sig_images)
        if not ok:
            sizes = "、".join(f"{w:.2f}×{h:.2f}cm" for w, h in sig_images)
            self.add_score(-1,
                           f"声明页签名图片大小不满足3.84×11.74cm（实际: {sizes}）",
                           "deduction")

    def check_cover_master_degree_title(self):
        """检查封面页"硕士学位论文"字体格式 -1。

        细则三点，任意一点不满足即扣 -1：
          1) 中文字体 = **黑体**（w:rFonts@eastAsia = "黑体" 或英文名 "SimHei"）
          2) 字号 = **小初**（w:sz val="72" ⇔ 36pt）
          3) **加粗**（w:b 存在且 val 非 "0"/"false"/"off"）

        细则未要求内容（英文字体、颜色、对齐、行距、间距、下划线等）
        一律不加以约束。若封面页未出现"硕士学位论文"文字，本项前置
        未满足，不扣分（避免与是否为硕士论文本身的判定纠缠）。

        办公软件（Word/WPS）适配说明：
          · "封面页"范围沿用 check_cover_images 的定义：文档开头至
            首处硬分页/段前分页/首个 sectPr 之间。
          · "黑体"在办公软件里 eastAsia 可能被记录为中文名"黑体"或
            英文名"SimHei"，两者均视为满足。
          · 属性可能显式在 run，也可能继承自段落样式（如 Heading 1）→
            basedOn 链 → Normal；办公软件的最终呈现就是这条链。
        """
        body = self.doc.element.body

        # 定位封面页内包含"硕士学位论文"的段落
        target_p = None
        for child in list(body):
            if child.tag == qn('w:sectPr'):
                break
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))

                # 段前分页
                page_break_before = False
                if pPr is not None:
                    pbb = pPr.find(qn('w:pageBreakBefore'))
                    if pbb is not None:
                        val = pbb.get(qn('w:val'))
                        if val is None or str(val).lower() not in ('0', 'false', 'off'):
                            page_break_before = True
                if page_break_before:
                    break

                # 检查文本
                txt = ''.join((t.text or '') for t in child.iter(qn('w:t')))
                if '硕士学位论文' in txt:
                    target_p = child
                    break

                # 段内硬分页
                stop_after = False
                for br in child.iter(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        stop_after = True
                        break
                # 节末
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    break
                if stop_after:
                    break

        if target_p is None:
            # 封面页未出现"硕士学位论文" —— 无法判定其字体，不扣分
            return

        # 构建样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = bold = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
                b = rPr.find(qn('w:b'))
                if b is not None:
                    v = b.get(qn('w:val'))
                    bold = '1' if v is None else v
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'east': east, 'sz': sz, 'bold': bold, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = target_p.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # 找到 "硕士学位论文" 六字所在的 run（同一段可能有多 run）
        # 逐 run 累计文本，取该子串首字所在 run 作为检验对象
        pos = 0
        target_run = None
        for run_el in target_p.iter(qn('w:r')):
            r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
            idx = r_text.find('硕士学位论文')
            r_len = len(r_text)
            if idx >= 0:
                target_run = run_el
                break
            # 跨 run 情况：只要 run 覆盖任一"硕士学位论文"字符即取为首个候选
            para_text_slice = ''.join((t.text or '')
                                      for t in target_p.iter(qn('w:t')))
            # fall back later
            pos += r_len
        if target_run is None:
            # 跨 run 情况 —— 用首个含"硕"或"士"的 run
            for run_el in target_p.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if any(ch in r_text for ch in '硕士学位论文'):
                    target_run = run_el
                    break

        if target_run is None:
            return

        rPr = target_run.find(qn('w:rPr'))
        east = sz = bold = None
        if rPr is not None:
            rf = rPr.find(qn('w:rFonts'))
            if rf is not None:
                east = rf.get(qn('w:eastAsia'))
            s = rPr.find(qn('w:sz'))
            if s is not None:
                sz = s.get(qn('w:val'))
            b = rPr.find(qn('w:b'))
            if b is not None:
                v = b.get(qn('w:val'))
                bold = '1' if v is None else v

        # 继承回落
        if east is None and para_style_id:
            east = _resolve(para_style_id, 'east')
        if east is None:
            east = _resolve('Normal', 'east')
        if sz is None and para_style_id:
            sz = _resolve(para_style_id, 'sz')
        if sz is None:
            sz = _resolve('Normal', 'sz')
        if bold is None and para_style_id:
            bold = _resolve(para_style_id, 'bold')
        if bold is None:
            bold = _resolve('Normal', 'bold')

        problems = []
        # 黑体
        if east is None or east.strip() not in ('黑体', 'SimHei'):
            problems.append(f"中文字体非黑体(eastAsia={east})")
        # 小初 = sz val="72"
        if sz != '72':
            problems.append(f"字号非小初(sz={sz})")
        # 加粗
        if bold is None or str(bold).lower() in ('0', 'false', 'off'):
            problems.append(f"未加粗(b={bold})")

        if problems:
            self.add_score(-1,
                           "封面页\"硕士学位论文\"字体格式不满足黑体、小初、加粗 —— "
                           + "；".join(problems),
                           "deduction")

    def check_declaration_page(self):
        """检查"研究生学位论文原创性声明"页**内容**段落格式 -3。

        细则三点，任意一点不满足即扣 -3：
          1) **左对齐** —— w:pPr/w:jc@val ∈ {"left", "start", None}
             （None 即未显式设置，默认左对齐；"start" 是 LTR 下的等价写法）
          2) **首行缩进两字符** —— 必须以"字符"为单位设置：
             · w:pPr/w:ind@firstLineChars = "200"
               （办公软件里"首行缩进2字符"的字符单位写法，200 = 2.00 字符）
             · 注意：仅写 w:ind@firstLine（绝对 twips 值，如 480）**不算**满足，
               因为它对应办公软件里"首行缩进 X 厘米/磅"的绝对单位设置，
               即便数值上恰为 2 字符宽，也不随字号伸缩，不符合"两字符"语义。
          3) **1.5 倍行距** —— w:pPr/w:spacing@line = "360" 且
             @lineRule 缺失或为 "auto"

        细则未要求内容（段前/段后间距、字体字号、颜色、悬挂缩进、
        右缩进、页码、行号等）一律不加以约束。若整个文档未出现声明页，
        前置未满足，不扣分。

        办公软件（Word/WPS）适配说明：
          · "内容"范围 = 标题段之后至"版权使用授权书/授权声明"之前的
            所有**非空**段落；标题段由 check_declaration_title / _title_para
            另行检查，本方法不重复校验。
          · 段落属性可能显式在 w:pPr，也可能继承自段落样式 → basedOn
            链 → Normal；办公软件的最终段落格式按这条链解析。
          · firstLineChars=200 才是"2 字符"的字符单位写法；仅设置
            firstLine（如 480、482、其他 twips 值，对应办公软件里显示为
            0.85 厘米/1.02 厘米等绝对缩进）**不视为满足**——办公软件里
            这类设置在字号变化时不会随之伸缩，与"两字符"语义不符。
          · 1.5 倍行距识别与 check_declaration_title_para 同规则
            （line=360 且 lineRule=auto，避免"固定值 18 磅"的误命中）。
        """
        # 定位内容段
        paragraphs = list(self.doc.paragraphs)
        start_idx = None
        end_idx = None
        for i, para in enumerate(paragraphs):
            t = para.text.strip().replace(' ', '').replace('　', '')
            if start_idx is None and '研究生学位论文原创性声明' in t:
                start_idx = i
                continue
            if start_idx is not None and ('版权使用授权书' in t or
                                          '版权使用授权' in t or
                                          '授权书' in t or
                                          '授权声明' in t):
                end_idx = i
                break

        if start_idx is None:
            return
        if end_idx is None:
            end_idx = len(paragraphs)

        content_paras = [p for p in paragraphs[start_idx + 1:end_idx]
                         if p.text.strip()]
        if not content_paras:
            return

        # 构建段落样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract_ppr(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = flc = fl = line = line_rule = None
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    flc = ind.get(qn('w:firstLineChars'))
                    fl = ind.get(qn('w:firstLine'))
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'flc': flc, 'fl': fl, 'line': line,
                    'lineRule': line_rule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract_ppr(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        problems = []
        for para in content_paras:
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            # 有效 jc
            jc = None
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            # 有效 firstLineChars / firstLine
            flc = fl = None
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    flc = ind.get(qn('w:firstLineChars'))
                    fl = ind.get(qn('w:firstLine'))
            if flc is None and para_style_id:
                flc = _resolve(para_style_id, 'flc')
            if flc is None:
                flc = _resolve('Normal', 'flc')
            if fl is None and para_style_id:
                fl = _resolve(para_style_id, 'fl')
            if fl is None:
                fl = _resolve('Normal', 'fl')

            # 有效 line/lineRule
            line = None
            line_rule = None
            if pPr is not None:
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
            if line is None:
                line = _resolve('Normal', 'line')
            if line_rule is None and para_style_id:
                line_rule = _resolve(para_style_id, 'lineRule')
            if line_rule is None:
                line_rule = _resolve('Normal', 'lineRule')

            preview = para.text.strip()[:12]

            # 左对齐
            if jc is not None and jc not in ('left', 'start'):
                problems.append(f"[{preview}] 非左对齐(jc={jc})")
                break
            # 首行缩进 2 字符 —— 必须以"字符"单位设置(firstLineChars=200)
            # 仅设置 firstLine(绝对 twips) 不视为满足，即便数值等价于 2 字符
            is_2char = (flc == '200')
            if not is_2char:
                problems.append(f"[{preview}] 非首行缩进2字符(firstLineChars={flc}, firstLine={fl})")
                break
            # 1.5 倍行距
            is_1_5 = (line == '360' and (line_rule is None or line_rule == 'auto'))
            if not is_1_5:
                problems.append(f"[{preview}] 非1.5倍行距(line={line}, lineRule={line_rule})")
                break

        if problems:
            self.add_score(-3,
                           "\"研究生学位论文原创性声明\"页内容段落格式不满足左对齐、首行缩进两字符、1.5倍行距 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_declaration_title(self):
        """检查"研究生学位论文原创性声明"页标题字体格式 -3。

        细则三点，任意一点不满足即扣 -3：
          1) 中文字体 = **黑体**（w:rFonts@eastAsia = "黑体" 或英文名 "SimHei"）
          2) 字号 = **小二**（w:sz val="36" ⇔ 18pt）
          3) **加粗**（w:b 存在且 val 非 "0"/"false"/"off"）

        细则未要求内容（英文字体、颜色、对齐、行距、间距、下划线等）
        一律不加约束。若整个文档未出现"研究生学位论文原创性声明"字样，
        本项前置未满足，不扣分。

        办公软件（Word/WPS）适配说明：
          · 标题段定位：段落文本包含"研究生学位论文原创性声明"即视为
            该页标题段；封面（前置的"硕士学位论文"条幅）不在此列。
          · 标题文本可能带学校名前缀（如"星途理工大学研究生学位论文原
            创性声明"），Word/WPS 中作为**一段**呈现 —— 校验以段落
            为单位，段内所有非空 run 均须同时满足三点。
          · 字体属性可能显式于 run，也可能继承自段落样式 → basedOn 链
            → Normal；办公软件的最终呈现按这条链解析。
          · "黑体"的 eastAsia 属性可能被 Office 写为中文名"黑体"或
            英文名"SimHei"，两者均视为满足。
        """
        # 定位标题段
        title_paras = []
        for para in self.doc.paragraphs:
            text = para.text.strip().replace(' ', '').replace('　', '')
            if '研究生学位论文原创性声明' in text:
                title_paras.append(para)
                break   # 该标题在文档中唯一，取首个即可

        if not title_paras:
            # 无原创性声明页 —— 无法判定标题字体，不扣分
            return

        # 构建样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = bold = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
                b = rPr.find(qn('w:b'))
                if b is not None:
                    v = b.get(qn('w:val'))
                    bold = '1' if v is None else v
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'east': east, 'sz': sz, 'bold': bold, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _run_props(run_el, para_style_id):
            rPr = run_el.find(qn('w:rPr'))
            east = sz = bold = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
                b = rPr.find(qn('w:b'))
                if b is not None:
                    v = b.get(qn('w:val'))
                    bold = '1' if v is None else v
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'east')
            if east is None:
                east = _resolve('Normal', 'east')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')
            if bold is None and para_style_id:
                bold = _resolve(para_style_id, 'bold')
            if bold is None:
                bold = _resolve('Normal', 'bold')
            return east, sz, bold

        problems = []
        for para in title_paras:
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            for run_el in para._element.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text.strip():
                    continue
                east, sz, bold = _run_props(run_el, para_style_id)
                if east is None or east.strip() not in ('黑体', 'SimHei'):
                    problems.append(f"中文字体非黑体(eastAsia={east})")
                    break
                if sz != '36':
                    problems.append(f"字号非小二(sz={sz})")
                    break
                if bold is None or str(bold).lower() in ('0', 'false', 'off'):
                    problems.append(f"未加粗(b={bold})")
                    break

        if problems:
            self.add_score(-3,
                           "\"研究生学位论文原创性声明\"页标题字体格式不满足黑体、小二、加粗 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_declaration_title_para(self):
        """检查"研究生学位论文原创性声明"页标题段落格式 -3。

        细则两点，任意一点不满足即扣 -3：
          1) **1.5 倍行距** —— w:pPr/w:spacing@line = "360" 且
             @lineRule 缺失或为 "auto"
             （Word/WPS 里"1.5 倍行距"写入即 240 × 1.5 = 360，
              且行距规则为"多倍/自动"，不能是"固定值"）。
          2) **居中对齐** —— w:pPr/w:jc@val = "center"。

        细则未要求内容（段前/段后间距、首行缩进、悬挂缩进、字号颜色、
        大纲级别、边框底纹等）一律不加以约束。若整个文档未出现"研究
        生学位论文原创性声明"字样，本项前置未满足，不扣分。

        办公软件（Word/WPS）适配说明：
          · 标题段定位：段落文本包含"研究生学位论文原创性声明"即视为
            该页标题段；封面（前置的"硕士学位论文"条幅）不在此列。
          · 段落属性可能显式在 w:pPr，也可能继承自段落样式 → basedOn
            链 → Normal；办公软件的最终段落格式按这条链解析。
          · 1.5 倍行距 —— Word/WPS UI 中的"1.5倍行距"仅对应
            (line=360, lineRule=auto)；固定值 18磅(line=360, lineRule=exact)
            虽然数值相同，但办公软件里会显示为"固定值"，不视为"1.5 倍"。
        """
        # 定位标题段
        title_paras = []
        for para in self.doc.paragraphs:
            text = para.text.strip().replace(' ', '').replace('　', '')
            if '研究生学位论文原创性声明' in text:
                title_paras.append(para)
                break

        if not title_paras:
            return

        # 构建段落样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract_ppr(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = None
            line = None
            line_rule = None
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': line_rule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract_ppr(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        problems = []
        for para in title_paras:
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            # 有效 jc
            jc = None
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            # 有效 line/lineRule
            line = None
            line_rule = None
            if pPr is not None:
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
            if line is None:
                line = _resolve('Normal', 'line')
            if line_rule is None and para_style_id:
                line_rule = _resolve(para_style_id, 'lineRule')
            if line_rule is None:
                line_rule = _resolve('Normal', 'lineRule')

            # 判定 1.5 倍行距
            is_1_5 = (line == '360' and (line_rule is None or line_rule == 'auto'))
            if not is_1_5:
                problems.append(f"非1.5倍行距(line={line}, lineRule={line_rule})")

            # 判定居中
            if jc != 'center':
                problems.append(f"非居中对齐(jc={jc})")

        if problems:
            self.add_score(-3,
                           "\"研究生学位论文原创性声明\"页标题段落格式不满足1.5倍行距、居中对齐 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_declaration_content_font(self):
        """检查"研究生学位论文原创性声明"页**内容**字体格式 -3。

        细则一点，两项须同时满足才不扣分（即"宋体 小四"）：
          1) 中文字体 = **宋体**（w:rFonts@eastAsia = "宋体" 或英文名 "SimSun"）
          2) 字号 = **小四**（w:sz val="24" ⇔ 12pt）

        细则未要求内容（英文/数字字体、加粗、颜色、行距、缩进、下划线等）
        一律不加以约束。仅约束**"除数字和英文字母外"**的字符 —— 即含
        汉字/中文标点的 run；纯英数 run 不参与中文字体校验（因为
        "宋体"是中文字体属性，办公软件里由 eastAsia 控制，与 ascii/hAnsi 无关）。

        办公软件（Word/WPS）适配说明：
          · "内容"范围 = 段落文本包含"研究生学位论文原创性声明"之后至
            "学位论文版权使用授权书"（或"授权书""授权声明"）之前的所有
            非空段落。不含标题段本身。
          · "宋体"的 eastAsia 可能被 Office 写为中文名"宋体"或英文名
            "SimSun"，两者均视为满足。
          · 字体属性可能显式在 run，也可能继承自段落样式 → basedOn 链
            → Normal；办公软件的最终呈现按这条链解析。
        """
        # 定位内容段
        paragraphs = list(self.doc.paragraphs)
        start_idx = None
        end_idx = None
        for i, para in enumerate(paragraphs):
            t = para.text.strip().replace(' ', '').replace('　', '')
            if start_idx is None and '研究生学位论文原创性声明' in t:
                start_idx = i
                continue
            if start_idx is not None and ('版权使用授权书' in t or
                                          '版权使用授权' in t or
                                          '授权书' in t or
                                          '授权声明' in t):
                end_idx = i
                break

        if start_idx is None:
            return
        if end_idx is None:
            end_idx = len(paragraphs)

        content_paras = [p for p in paragraphs[start_idx + 1:end_idx]
                         if p.text.strip()]
        if not content_paras:
            return

        # 构建样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'east': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _run_props(run_el, para_style_id):
            rPr = run_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'east')
            if east is None:
                east = _resolve('Normal', 'east')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')
            return east, sz

        han_re = re.compile(r'[一-鿿]')
        cjk_punct_re = re.compile(r'[，。；：、？！""''（）【】《》—…]')

        problems = []
        for para in content_paras:
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            for run_el in para._element.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text.strip():
                    continue
                # 仅约束含汉字/中文标点的 run
                if not (han_re.search(r_text) or cjk_punct_re.search(r_text)):
                    continue
                east, sz = _run_props(run_el, para_style_id)
                if east is None or east.strip() not in ('宋体', 'SimSun'):
                    problems.append(f"中文字体非宋体(eastAsia={east})")
                    break
                if sz != '24':
                    problems.append(f"字号非小四(sz={sz})")
                    break
            if problems:
                break

        if problems:
            self.add_score(-3,
                           "\"研究生学位论文原创性声明\"页内容字体格式不满足宋体、小四 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_abstract_cn_format(self):
        """检查中文摘要页标题字体格式 -1。

        细则一点，两项须同时满足才不扣分：
          1) 中文字体 = **黑体**（w:rFonts@eastAsia = "黑体" 或英文名 "SimHei"）
          2) 字号 = **四号**（w:sz val="28" ⇔ 14pt）

        细则未要求内容（英文字体、加粗、颜色、对齐、行距、间距、下划
        线等）一律不加以约束。若整个文档未出现中文摘要标题段，本项前
        置未满足，不扣分。

        办公软件（Word/WPS）适配说明：
          · 中文摘要标题段：段落拼接文本 strip 后严格等于 "摘要" 或
            "中文摘要"（去空白后比较，兼容全/半角空格）。
          · 字体属性可能显式在 run，也可能继承自段落样式 → basedOn 链
            → Normal；办公软件的最终呈现按这条链解析。
          · "黑体"的 eastAsia 可能被 Office 写为中文名"黑体"或英文名
            "SimHei"，两者均视为满足。
        """
        title_paras = []
        for para in self.doc.paragraphs:
            t = para.text.strip().replace(' ', '').replace('　', '')
            if t == '摘要' or t == '中文摘要':
                title_paras.append(para)
                break

        if not title_paras:
            return

        # 构建样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'east': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _run_props(run_el, para_style_id):
            rPr = run_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'east')
            if east is None:
                east = _resolve('Normal', 'east')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')
            return east, sz

        problems = []
        for para in title_paras:
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            for run_el in para._element.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text.strip():
                    continue
                east, sz = _run_props(run_el, para_style_id)
                if east is None or east.strip() not in ('黑体', 'SimHei'):
                    problems.append(f"中文字体非黑体(eastAsia={east})")
                    break
                if sz != '28':
                    problems.append(f"字号非四号(sz={sz})")
                    break

        if problems:
            self.add_score(-1,
                           "中文摘要页标题不满足黑体、四号 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def check_abstract_cn_title_para(self):
        """检查中文摘要页标题段落格式 -1。

        细则两点，任意一点不满足即扣 -1：
          1) **居中对齐** —— w:pPr/w:jc@val = "center"
          2) **1.5 倍行距** —— w:pPr/w:spacing@line = "360" 且
             @lineRule 缺失或为 "auto"

        细则未要求内容（段前/段后间距、首行缩进、大纲级别、边框底纹、
        字体字号颜色等）一律不加以约束。若整个文档未出现中文摘要标题段
        （"摘要"/"中文摘要"独立段），本项前置未满足，不扣分。

        办公软件（Word/WPS）适配说明：
          · 标题段定位：段落拼接文本 strip 且去空白后严格等于 "摘要"
            或 "中文摘要"（与 check_abstract_cn_format 同规则）。
          · 段落属性可能显式在 w:pPr，也可能继承自段落样式 → basedOn
            链 → Normal；办公软件最终段落格式按这条链解析。
          · 1.5 倍行距识别与 check_declaration_title_para 同规则
            （line=360 且 lineRule=auto，避免"固定值 18 磅"的误命中）。
        """
        title_paras = []
        for para in self.doc.paragraphs:
            t = para.text.strip().replace(' ', '').replace('　', '')
            if t == '摘要' or t == '中文摘要':
                title_paras.append(para)
                break

        if not title_paras:
            return

        # 构建段落样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract_ppr(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = line_rule = None
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': line_rule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract_ppr(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        problems = []
        for para in title_paras:
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            # 有效 jc
            jc = None
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            # 有效 line / lineRule
            line = None
            line_rule = None
            if pPr is not None:
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
            if line is None:
                line = _resolve('Normal', 'line')
            if line_rule is None and para_style_id:
                line_rule = _resolve(para_style_id, 'lineRule')
            if line_rule is None:
                line_rule = _resolve('Normal', 'lineRule')

            if jc != 'center':
                problems.append(f"非居中对齐(jc={jc})")
            is_1_5 = (line == '360' and (line_rule is None or line_rule == 'auto'))
            if not is_1_5:
                problems.append(f"非1.5倍行距(line={line}, lineRule={line_rule})")

        if problems:
            self.add_score(-1,
                           "中文摘要页标题段落格式不满足居中对齐、1.5倍行距 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def check_abstract_en_format(self):
        """检查英文摘要页标题格式 -1。

        细则五点，任意一点不满足即扣 -1：
          1) **居中对齐** —— w:pPr/w:jc@val == "center"
             （办公软件"居中对齐"对应 XML val="center"；未显式设置或
              为 left/start/both/right 等均不满足）
          2) **题目实词首字母大写** —— 除虚词外的每个单词首字母大写
          3) **虚词小写** —— 冠词/介词/并列连词/短介词等虚词全小写
             （首词与末词例外，永远视作需要首字母大写 —— 标准英文
              标题大小写规则）
          4) **字号 = 四号**（w:sz val="28" ⇔ 14pt）
          5) **英文字体 = Times New Roman**（w:rFonts@ascii 与 @hAnsi
             均为 "Times New Roman"）+ **加粗**（w:b 存在且 val 非
             "0"/"false"/"off"）

        细则未要求内容（东亚字体、颜色、行距、间距、下划线等）一律不
        加以约束。若整个文档未出现英文摘要标题，本项前置未满足，不扣分。

        办公软件（Word/WPS）适配说明：
          · "英文摘要页标题"定位：先找 "ABSTRACT"/"Abstract" 段（英文
            摘要页的节标题），再取其后**首个**非空的、以英文为主的段落
            视为该页标题（即论文英文题目）。若 ABSTRACT 段本身已是英文
            题目（该段文本包含多单词英文题目字样），则直接取 ABSTRACT
            段。方向式判定 —— "题目实词首字母大写，虚词小写" 是仅对
            多单词英文题目有效的规则。
          · 字体属性可能显式在 run，也可能继承自段落样式 → basedOn 链
            → Normal；办公软件的最终呈现按这条链解析。ascii/hAnsi 是
            办公软件中"西文字体"控制的两个 XML 属性，均要匹配 TNR。
        """
        # 定位英文摘要节标题（"Abstract" / "ABSTRACT" 独立段）
        paragraphs = list(self.doc.paragraphs)
        head_idx = None
        for i, para in enumerate(paragraphs):
            t = para.text.strip()
            if t.upper() == 'ABSTRACT':
                head_idx = i
                break
        if head_idx is None:
            return

        # 英文摘要页"标题"识别：优先 ABSTRACT 段之后的首个"英文多词段"
        # 判定"英文多词段"：非空 + 主要为英文字符 + 单词数 ≥ 2 +
        # 不以正文关键字 Objective/Methods/Results/Conclusion/Subject/Key words 起首
        body_kw = re.compile(r'^(Objective|Method|Result|Conclusion|Subject|Key\s*word)s?\b', re.I)
        latin_re = re.compile(r'[A-Za-z]')
        cjk_re = re.compile(r'[一-鿿]')

        title_para = None
        for para in paragraphs[head_idx + 1:head_idx + 6]:
            t = para.text.strip()
            if not t:
                continue
            if body_kw.match(t):
                break
            # 主要为英文（去空白后英文字符占比 > 60%），且单词数 ≥ 2
            no_space = re.sub(r'\s+', '', t)
            latin_count = len(latin_re.findall(no_space))
            if not no_space:
                continue
            if latin_count / len(no_space) < 0.6:
                continue
            if cjk_re.search(t):
                continue
            words = re.findall(r"[A-Za-z][A-Za-z'\-]*", t)
            if len(words) < 2:
                continue
            title_para = para
            break

        if title_para is None:
            # 未定位到多词英文题目 —— 前置不满足，不扣分
            return

        # 构建样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            pPr = style_el.find(qn('w:pPr'))
            ascii_ = hansi = sz = bold = jc = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    ascii_ = rf.get(qn('w:ascii'))
                    hansi = rf.get(qn('w:hAnsi'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
                b = rPr.find(qn('w:b'))
                if b is not None:
                    v = b.get(qn('w:val'))
                    bold = '1' if v is None else v
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'ascii': ascii_, 'hAnsi': hansi, 'sz': sz,
                    'bold': bold, 'jc': jc, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # 有效 jc（居中对齐）
        jc = None
        if pPr is not None:
            jc_el = pPr.find(qn('w:jc'))
            if jc_el is not None:
                jc = jc_el.get(qn('w:val'))
        if jc is None and para_style_id:
            jc = _resolve(para_style_id, 'jc')
        if jc is None:
            jc = _resolve('Normal', 'jc')

        # 逐 run 校验字体属性
        problems = []
        if jc != 'center':
            problems.append(f"非居中对齐(jc={jc})")

        # 虚词集合（英文标题大小写规则）—— 冠词、短介词、并列连词、部分复合词
        LOWER_WORDS = {
            # 冠词
            'a', 'an', 'the',
            # 并列连词
            'and', 'or', 'but', 'nor', 'so', 'yet', 'for',
            # 短介词
            'at', 'by', 'in', 'of', 'on', 'to', 'up', 'as', 'off',
            'per', 'via',
            'from', 'into', 'like', 'near', 'onto', 'over', 'past',
            'than', 'till', 'unto', 'upon', 'with',
            'about', 'above', 'across', 'after', 'along', 'among',
            'around', 'below', 'beyond', 'between', 'during',
            'except', 'inside', 'outside', 'through', 'toward',
            'under', 'until', 'upon', 'within', 'without',
            # 从属连词部分（一般教材归入虚词）
            'if', 'that', 'when', 'while', 'because', 'though',
            'although',
        }

        words = re.findall(r"[A-Za-z][A-Za-z'\-]*", title_para.text)
        n = len(words)
        wrong_case = []
        for i, w in enumerate(words):
            is_first = (i == 0)
            is_last = (i == n - 1)
            lw = w.lower()
            first_ch = w[0]
            if is_first or is_last:
                # 首末词一律首字母大写
                if not first_ch.isupper():
                    wrong_case.append(w)
            elif lw in LOWER_WORDS:
                # 虚词全小写
                if w != lw:
                    wrong_case.append(w)
            else:
                # 实词首字母大写
                if not first_ch.isupper():
                    wrong_case.append(w)
        if wrong_case:
            problems.append("大小写不符(" + ",".join(wrong_case[:3]) + ")")

        # 字体属性（对每个非空 run 检查）
        for run_el in title_para._element.iter(qn('w:r')):
            r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
            if not r_text.strip():
                continue
            rPr = run_el.find(qn('w:rPr'))
            ascii_ = hansi = sz = bold = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    ascii_ = rf.get(qn('w:ascii'))
                    hansi = rf.get(qn('w:hAnsi'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
                b = rPr.find(qn('w:b'))
                if b is not None:
                    v = b.get(qn('w:val'))
                    bold = '1' if v is None else v
            if ascii_ is None and para_style_id:
                ascii_ = _resolve(para_style_id, 'ascii')
            if ascii_ is None:
                ascii_ = _resolve('Normal', 'ascii')
            if hansi is None and para_style_id:
                hansi = _resolve(para_style_id, 'hAnsi')
            if hansi is None:
                hansi = _resolve('Normal', 'hAnsi')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')
            if bold is None and para_style_id:
                bold = _resolve(para_style_id, 'bold')
            if bold is None:
                bold = _resolve('Normal', 'bold')

            if sz != '28':
                problems.append(f"字号非四号(sz={sz})")
                break
            if ascii_ != 'Times New Roman' or hansi != 'Times New Roman':
                problems.append(f"西文字体非Times New Roman(ascii={ascii_}, hAnsi={hansi})")
                break
            if bold is None or str(bold).lower() in ('0', 'false', 'off'):
                problems.append(f"未加粗(b={bold})")
                break

        if problems:
            self.add_score(-1,
                           "英文摘要页标题不满足居中对齐、题目实词首字母大写虚词小写、四号Times New Roman加粗 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_abstract_en_body_font(self):
        """检查英文摘要页标题下方文本字体 -1。

        细则两点，任意一点不满足即扣 -1：
          1) **Times New Roman** —— w:rFonts@ascii 与 @hAnsi 均为
             "Times New Roman"（继承链解析：run → 段落样式 → basedOn
              链 → Normal，办公软件的最终呈现按这条链取值）
          2) **小四号** —— w:sz@val="24"（半磅单位，24 ⇔ 12pt ⇔ 小四）

        细则未要求内容（东亚字体、颜色、行距、对齐、缩进、加粗、
        斜体、下划线等）一律不加以约束。

        范围界定："英文摘要页标题下方文本"指英文题目段之后、直到英文
        摘要页结束前（下一节标题 —— 缩略词表 / 前言 / 目录 / 下一
        Heading1 —— 之前）的所有非空段落。空段（分节间隔）跳过。
        """
        paragraphs = list(self.doc.paragraphs)

        # 定位 ABSTRACT 段
        head_idx = None
        for i, para in enumerate(paragraphs):
            if para.text.strip().upper() == 'ABSTRACT':
                head_idx = i
                break
        if head_idx is None:
            return

        # 复用与 check_abstract_en_format 相同逻辑定位"英文题目段"
        body_kw = re.compile(r'^(Objective|Method|Result|Conclusion|Subject|Key\s*word)s?\b', re.I)
        latin_re = re.compile(r'[A-Za-z]')
        cjk_re = re.compile(r'[一-鿿]')

        title_offset = None  # 相对 head_idx 之后的索引
        for j, para in enumerate(paragraphs[head_idx + 1:head_idx + 6]):
            t = para.text.strip()
            if not t:
                continue
            if body_kw.match(t):
                break
            no_space = re.sub(r'\s+', '', t)
            if not no_space:
                continue
            if len(latin_re.findall(no_space)) / len(no_space) < 0.6:
                continue
            if cjk_re.search(t):
                continue
            words = re.findall(r"[A-Za-z][A-Za-z'\-]*", t)
            if len(words) < 2:
                continue
            title_offset = head_idx + 1 + j
            break

        # 若未定位到独立英文题目段，则将 head_idx 本身视作"标题"，
        # 其后段落视为下方文本（保证细则可判定）
        start_idx = (title_offset if title_offset is not None else head_idx) + 1

        # 找到英文摘要页结束边界：下一节标题
        # 关键字："缩略词表"、"前言"、"目录"，或独立 Heading1/Heading2
        end_idx = len(paragraphs)
        for k in range(start_idx, len(paragraphs)):
            t = paragraphs[k].text.strip()
            if not t:
                continue
            pPr = paragraphs[k]._element.find(qn('w:pPr'))
            pStyle = None
            if pPr is not None:
                el = pPr.find(qn('w:pStyle'))
                if el is not None:
                    pStyle = el.get(qn('w:val'))
            if '缩略词表' in t or t == '前言' or t == '目录':
                end_idx = k
                break
            if pStyle in ('Heading1', 'Heading2') and not body_kw.match(t):
                # Heading1/2 且不是 Objective/Methods 等（后者理应是 Normal）
                end_idx = k
                break

        if start_idx >= end_idx:
            return

        # 构建样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            ascii_ = hansi = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    ascii_ = rf.get(qn('w:ascii'))
                    hansi = rf.get(qn('w:hAnsi'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'ascii': ascii_, 'hAnsi': hansi, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        problems = []

        for idx in range(start_idx, end_idx):
            para = paragraphs[idx]
            if not para.text.strip():
                continue
            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            for run_el in para._element.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text.strip():
                    continue
                # 只校验含有英文字母/数字的 run —— 细则规范的是 "Times
                # New Roman"，Word/WPS 中该字体只作用于西文字符；若某
                # run 全为中文/符号，其西文字体属性无实际显示效果。
                if not re.search(r'[A-Za-z0-9]', r_text):
                    continue

                rPr = run_el.find(qn('w:rPr'))
                ascii_ = hansi = sz = None
                if rPr is not None:
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None:
                        ascii_ = rf.get(qn('w:ascii'))
                        hansi = rf.get(qn('w:hAnsi'))
                    s = rPr.find(qn('w:sz'))
                    if s is not None:
                        sz = s.get(qn('w:val'))
                if ascii_ is None and para_style_id:
                    ascii_ = _resolve(para_style_id, 'ascii')
                if ascii_ is None:
                    ascii_ = _resolve('Normal', 'ascii')
                if hansi is None and para_style_id:
                    hansi = _resolve(para_style_id, 'hAnsi')
                if hansi is None:
                    hansi = _resolve('Normal', 'hAnsi')
                if sz is None and para_style_id:
                    sz = _resolve(para_style_id, 'sz')
                if sz is None:
                    sz = _resolve('Normal', 'sz')

                snippet = r_text.strip()[:20]
                if ascii_ != 'Times New Roman' or hansi != 'Times New Roman':
                    problems.append(f"[{snippet}] 西文字体非TNR(ascii={ascii_},hAnsi={hansi})")
                    break
                if sz != '24':
                    problems.append(f"[{snippet}] 字号非小四(sz={sz})")
                    break
            if problems:
                break

        if problems:
            self.add_score(-1,
                           "英文摘要页标题下方文本字体格式不满足Times New Roman小四号 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_abstract_en_body_para(self):
        """检查英文摘要页标题下方文本段落格式 -1。

        细则两点，任意一点不满足即扣 -1：
          1) **两端对齐** —— w:pPr/w:jc@val ∈ {"both", "justify",
             "distribute"}；办公软件（Word/WPS）中 "两端对齐" 对应
             XML val="both"（旧版）或 "justify"（等价别名）
          2) **1.5 倍行距** —— w:spacing@line="360" 且 lineRule 未设
             或为 "auto"；lineRule="exact"/"atLeast" 视为固定值/最小
             值，不属于 1.5 倍行距

        以上属性均按继承链解析：段落直接 pPr → 段落样式 → basedOn
        链 → Normal。这是办公软件呈现段落格式的最终依据。

        细则未要求内容（缩进、间距、大纲级别、边框、制表位等）一律
        不加以约束。

        范围界定：与 check_abstract_en_body_font 完全一致 —— 英文
        题目段之后、下一节标题（缩略词表/前言/目录/Heading1/2）之前
        的所有非空段落。
        """
        paragraphs = list(self.doc.paragraphs)

        head_idx = None
        for i, para in enumerate(paragraphs):
            if para.text.strip().upper() == 'ABSTRACT':
                head_idx = i
                break
        if head_idx is None:
            return

        body_kw = re.compile(r'^(Objective|Method|Result|Conclusion|Subject|Key\s*word)s?\b', re.I)
        latin_re = re.compile(r'[A-Za-z]')
        cjk_re = re.compile(r'[一-鿿]')

        title_offset = None
        for j, para in enumerate(paragraphs[head_idx + 1:head_idx + 6]):
            t = para.text.strip()
            if not t:
                continue
            if body_kw.match(t):
                break
            no_space = re.sub(r'\s+', '', t)
            if not no_space:
                continue
            if len(latin_re.findall(no_space)) / len(no_space) < 0.6:
                continue
            if cjk_re.search(t):
                continue
            words = re.findall(r"[A-Za-z][A-Za-z'\-]*", t)
            if len(words) < 2:
                continue
            title_offset = head_idx + 1 + j
            break

        start_idx = (title_offset if title_offset is not None else head_idx) + 1

        end_idx = len(paragraphs)
        for k in range(start_idx, len(paragraphs)):
            t = paragraphs[k].text.strip()
            if not t:
                continue
            pPr_k = paragraphs[k]._element.find(qn('w:pPr'))
            pStyle_k = None
            if pPr_k is not None:
                el = pPr_k.find(qn('w:pStyle'))
                if el is not None:
                    pStyle_k = el.get(qn('w:val'))
            if '缩略词表' in t or t == '前言' or t == '目录':
                end_idx = k
                break
            if pStyle_k in ('Heading1', 'Heading2') and not body_kw.match(t):
                end_idx = k
                break

        if start_idx >= end_idx:
            return

        # 样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        problems = []

        for idx in range(start_idx, end_idx):
            para = paragraphs[idx]
            if not para.text.strip():
                continue
            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            # 1) jc
            jc = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            # 2) spacing
            line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
                if lineRule is None:
                    lineRule = _resolve(para_style_id, 'lineRule')
            if line is None:
                line = _resolve('Normal', 'line')
                if lineRule is None:
                    lineRule = _resolve('Normal', 'lineRule')

            snippet = para.text.strip()[:20]

            # 两端对齐：val 为 "both"（Word 传统写法）或 "justify"
            # （Word 2016+ 的等价别名）或 "distribute"（分散对齐？此处
            #  仅两端对齐两种合规 val 为 both/justify；distribute 不算）
            if jc not in ('both', 'justify'):
                problems.append(f"[{snippet}] 非两端对齐(jc={jc})")
                continue

            # 1.5 倍行距：line="360" 且 lineRule 缺省或 "auto"
            ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
            if not ok_150:
                problems.append(f"[{snippet}] 非1.5倍行距(line={line},rule={lineRule})")
                continue

        if problems:
            self.add_score(-1,
                           "英文摘要页标题下方文本段落格式不满足两端对齐、1.5倍行距 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_abbreviation_page(self):
        """检查中英文缩略词表页标题 -1。

        细则两点，任意一点不满足即扣 -1：
          1) **四号** —— w:sz@val="28"（半磅单位，28 ⇔ 14pt ⇔ 四号）
          2) **黑体** —— w:rFonts@eastAsia ∈ {"黑体", "SimHei"}
             （办公软件 Word/WPS 中，SimHei 是黑体的英文别名，两者
              指向同一字体族，UI 显示等价）

        以上属性均按继承链解析：run rPr → 段落样式 rPr → basedOn 链
         → Normal，这是办公软件呈现字体的最终依据。

        细则未要求内容（西文字体、加粗、颜色、下划线、字符间距、
        对齐、行距等）一律不加以约束。若整个文档未出现"中英文缩略
        词表"页标题，本项前置不满足，不扣分。

        定位方式：全文段落中 strip 后等于 "中英文缩略词表" 的独立
        段（办公软件中该段为节标题独占一段的通用写法）。若同名内容
        出现在正文句中（非独立段）则忽略。
        """
        paragraphs = list(self.doc.paragraphs)
        title_para = None
        for para in paragraphs:
            if para.text.strip() == '中英文缩略词表':
                title_para = para
                break
        if title_para is None:
            return

        # 样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

        problems = []
        for run_el in title_para._element.iter(qn('w:r')):
            r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
            if not r_text.strip():
                continue
            rPr = run_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'eastAsia')
            if east is None:
                east = _resolve('Normal', 'eastAsia')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')

            if sz != '28':
                problems.append(f"字号非四号(sz={sz})")
                break
            if east not in ('黑体', 'SimHei'):
                problems.append(f"中文字体非黑体(eastAsia={east})")
                break

        if problems:
            self.add_score(-1,
                           "中英文缩略词表页标题不满足四号、黑体 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def check_preface_page(self):
        """检查"前言"页标题 -1。

        细则两点，任意一点不满足即扣 -1：
          1) **四号** —— w:sz@val="28"（半磅，14pt）
          2) **黑体** —— w:rFonts@eastAsia ∈ {"黑体", "SimHei"}
             （办公软件 Word/WPS 中 SimHei 是黑体的英文别名，UI 显示
              等价）

        以上属性均按继承链解析：run rPr → 段落样式 rPr → basedOn 链
         → Normal，办公软件依此链呈现字体。

        细则未要求内容（西文字体、加粗、颜色、下划线、字符间距、对齐、
        行距、缩进等）一律不加以约束。若整个文档未出现"前言"独立段
        标题，本项前置不满足，不扣分。

        定位方式：全文段落中 strip 后精确等于 "前言" 的独立段。
        """
        paragraphs = list(self.doc.paragraphs)
        title_para = None
        for para in paragraphs:
            if para.text.strip() == '前言':
                title_para = para
                break
        if title_para is None:
            return

        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

        problems = []
        for run_el in title_para._element.iter(qn('w:r')):
            r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
            if not r_text.strip():
                continue
            rPr = run_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'eastAsia')
            if east is None:
                east = _resolve('Normal', 'eastAsia')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')

            if sz != '28':
                problems.append(f"字号非四号(sz={sz})")
                break
            if east not in ('黑体', 'SimHei'):
                problems.append(f"中文字体非黑体(eastAsia={east})")
                break

        if problems:
            self.add_score(-1,
                           "前言页标题不满足四号、黑体 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def check_primary_title_para(self):
        """检查一级标题段落格式 -3：1.5倍行距、居中对齐。

        细则两点，任意一段一级标题不满足即扣 -3：
          1) **居中对齐** —— w:pPr/w:jc@val == "center"
          2) **1.5 倍行距** —— w:spacing@line="360" 且 lineRule 未设
             或为 "auto"；"exact"/"atLeast" 视为固定值/最小值，均不
             属于 1.5 倍行距

        以上属性按继承链解析：段落直接 pPr → 段落样式 pPr → basedOn
        链 → Normal，这是办公软件呈现段落格式的最终依据。

        细则未要求内容（缩进、间距、大纲级别、边框、制表位、段前段
        后空行等）一律不加以约束。

        定位方式：段落 strip 后匹配
        "第X部分"（X∈中文数字一~十）；目录条目（SDT/PAGEREF）跳过。
        """
        paragraphs = list(self.doc.paragraphs)

        # 样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        pattern = re.compile(r'^第[一二三四五六七八九十]+部分(\s|$)')

        problems = []
        for para in paragraphs:
            text = para.text.strip()
            if not pattern.match(text):
                continue
            if _is_toc_entry(para._element):
                continue

            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            # 1) jc
            jc = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            # 2) spacing
            line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
                if lineRule is None:
                    lineRule = _resolve(para_style_id, 'lineRule')
            if line is None:
                line = _resolve('Normal', 'line')
                if lineRule is None:
                    lineRule = _resolve('Normal', 'lineRule')

            snippet = text[:15]
            if jc != 'center':
                problems.append(f"[{snippet}] 非居中对齐(jc={jc})")
                continue
            ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
            if not ok_150:
                problems.append(f"[{snippet}] 非1.5倍行距(line={line},rule={lineRule})")
                continue

        if problems:
            self.add_score(-3,
                           "一级标题段落格式不满足1.5倍行距、居中对齐 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_secondary_title_para(self):
        """检查二级标题段落格式 -3：1.5倍行距、居中对齐。

        细则两点，任意一段二级标题不满足即扣 -3：
          1) **居中对齐** —— w:pPr/w:jc@val == "center"
          2) **1.5 倍行距** —— w:spacing@line="360" 且 lineRule 未设
             或为 "auto"；"exact"/"atLeast" 视为固定值/最小值，均不
             属于 1.5 倍行距

        以上属性按继承链解析：段落直接 pPr → 段落样式 pPr → basedOn
        链 → Normal，这是办公软件呈现段落格式的最终依据。

        细则未要求内容（字体、字号、缩进、间距、大纲级别、边框、
        制表位等）一律不加以约束。

        定位方式："二级标题" = 全文段落中 strip 后匹配 "X、"（X∈中文
        数字一~十）开头的段落。这是本文档的二级标题格式约定。目录
        条目（SDT/PAGEREF）跳过。
        """
        paragraphs = list(self.doc.paragraphs)

        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        pattern = re.compile(r'^[一二三四五六七八九十]+、')

        problems = []
        for para in paragraphs:
            text = para.text.strip()
            if not pattern.match(text):
                continue
            if _is_toc_entry(para._element):
                continue

            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            jc = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
                if lineRule is None:
                    lineRule = _resolve(para_style_id, 'lineRule')
            if line is None:
                line = _resolve('Normal', 'line')
                if lineRule is None:
                    lineRule = _resolve('Normal', 'lineRule')

            snippet = text[:15]
            if jc != 'center':
                problems.append(f"[{snippet}] 非居中对齐(jc={jc})")
                continue
            ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
            if not ok_150:
                problems.append(f"[{snippet}] 非1.5倍行距(line={line},rule={lineRule})")
                continue

        if problems:
            self.add_score(-3,
                           "二级标题段落格式不满足1.5倍行距、居中对齐 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_body_text_format(self):
        """检查正文段落格式 -3：首行缩进两字符、两端对齐、1.5倍行距。

        细则三点，任意一段正文段落不满足即扣 -3：
          1) **首行缩进两字符** —— 必须以"字符"单位设置：
               a) w:ind@firstLineChars == "200"（"字符"单位，与字号
                  无关，是办公软件默认写法）
               b) 注意：仅设置 w:ind@firstLine（twips 绝对值，对应
                  办公软件里"厘米/磅"单位的绝对缩进）**不视为满足**，
                  即便数值上恰好等于 2 个字符宽度也不算——它不随字号
                  伸缩，与"两字符"语义不符。
          2) **两端对齐** —— w:jc@val ∈ {"both", "justify"}
             （办公软件中"两端对齐"对应 XML val="both"（传统）或
              "justify"（等价别名））
          3) **1.5 倍行距** —— w:spacing@line=="360" 且 lineRule 未设
             或为 "auto"；"exact"/"atLeast" 视为固定值/最小值，均不
             属于 1.5 倍行距

        以上属性均按继承链解析：段落直接 pPr → 段落样式 pPr → basedOn
        链 → Normal，办公软件依此链呈现段落格式。

        细则未要求内容（左/右缩进、段前段后、大纲级别、边框、制表位、
        字体、字号等）一律不加以约束。

        范围界定："正文"指第一段"第X部分"（含之后）开始，直到"参考
        文献"独立段（不含）之间。跳过：
          · 一级标题（"第X部分"整段）
          · 二级标题（"X、"开头整段）
          · 图/表标注段（以 "图/表 + 数字" 起首）
          · 章节标题及其"（续）"变体（讨论/结论/致谢/前言 等）
          · Heading1/Heading2 段
          · 表格单元格内文本（表格另有细则约束）
          · 目录条目（SDT/PAGEREF）
        以上跳过项对应细则另有专门规则，本项不重复覆盖。
        """
        paragraphs = list(self.doc.paragraphs)

        # 样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = firstLineChars = firstLine = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
                el = pPr.find(qn('w:ind'))
                if el is not None:
                    firstLineChars = el.get(qn('w:firstLineChars'))
                    firstLine = el.get(qn('w:firstLine'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule,
                    'firstLineChars': firstLineChars, 'firstLine': firstLine,
                    'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _in_table(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    return True
                parent = parent.getparent()
            return False

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        primary_re = re.compile(r'^第[一二三四五六七八九十]+部分(\s|$)')
        secondary_re = re.compile(r'^[一二三四五六七八九十]+、')
        fig_tab_re = re.compile(r'^(图|表)\s*\d+')
        section_head_re = re.compile(
            r'^(前言|讨论|结论|致谢|参考文献|附录|摘要|中文摘要|英文摘要|Abstract|ABSTRACT|中英文缩略词表)'
            r'([（(]\s*续\s*[）)])?$'
        )

        in_body = False
        problems = []
        for para in paragraphs:
            text = para.text.strip()

            if primary_re.match(text):
                in_body = True
                continue
            if not in_body:
                continue
            if text == '参考文献':
                break
            if not text:
                continue

            if secondary_re.match(text):
                continue
            if fig_tab_re.match(text):
                continue
            if section_head_re.match(text):
                continue
            if _in_table(para._element):
                continue
            if _is_toc_entry(para._element):
                continue

            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            if para_style_id in ('Heading1', 'Heading2'):
                continue

            # 1) 首行缩进（严格：仅 firstLineChars=200 视为满足）
            flc = fl = None
            if pPr is not None:
                el = pPr.find(qn('w:ind'))
                if el is not None:
                    flc = el.get(qn('w:firstLineChars'))
                    fl = el.get(qn('w:firstLine'))
            if flc is None and para_style_id:
                flc = _resolve(para_style_id, 'firstLineChars')
            if flc is None:
                flc = _resolve('Normal', 'firstLineChars')
            if fl is None and para_style_id:
                fl = _resolve(para_style_id, 'firstLine')
            if fl is None:
                fl = _resolve('Normal', 'firstLine')

            ok_indent = (flc == '200')

            # 2) jc（两端对齐）
            jc = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            # 3) 行距
            line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
                if lineRule is None:
                    lineRule = _resolve(para_style_id, 'lineRule')
            if line is None:
                line = _resolve('Normal', 'line')
                if lineRule is None:
                    lineRule = _resolve('Normal', 'lineRule')

            snippet = text[:15]
            if not ok_indent:
                problems.append(f"[{snippet}] 非首行缩进2字符(firstLineChars={flc},firstLine={fl})")
                continue
            if jc not in ('both', 'justify'):
                problems.append(f"[{snippet}] 非两端对齐(jc={jc})")
                continue
            ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
            if not ok_150:
                problems.append(f"[{snippet}] 非1.5倍行距(line={line},rule={lineRule})")
                continue

        if problems:
            self.add_score(-3,
                           "正文段落格式不满足首行缩进两字符、两端对齐、1.5倍行距 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_body_text_font(self):
        """检查正文（除英文与数字外）字体格式 -3：小四号、宋体。

        细则两点，任意一处不满足即扣 -3：
          1) **小四** —— w:sz@val="24"（半磅，12pt）
          2) **宋体** —— w:rFonts@eastAsia ∈ {"宋体", "SimSun"}
             （办公软件 Word/WPS 中 SimSun 是宋体的英文别名，UI 显示
              等价）

        细则明示"**除英文与数字外**"，即本项**只约束中文字符所在的
        run**。含 CJK 字符的 run 才进入本检查；纯英文/纯数字/纯符号
        run 不视为约束对象（其字体归其他细则约束，此处不重复覆盖）。

        以上属性按继承链解析：run rPr → 段落样式 rPr → basedOn 链 →
        Normal，办公软件依此链呈现字体。

        细则未要求内容（西文字体、加粗、颜色、下划线、字符间距、对齐、
        行距、缩进等）一律不加以约束。

        范围界定："正文"指第一部分~最后一个"第X部分"（含）之后、参考
        文献列表之前的内容。跳过：
          · 一级标题（"第X部分"整段）
          · 二级标题（"X、"开头整段）
          · 图/表标注段（以 "图" 或 "表" + 数字起首）
          · 表格单元格内文本（表内文字另有细则约束）
          · 目录条目（SDT/PAGEREF）
        以上跳过项对应细则另有专门规则，本项不重复覆盖。
        """
        paragraphs = list(self.doc.paragraphs)

        # 样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _in_table(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    return True
                parent = parent.getparent()
            return False

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        primary_re = re.compile(r'^第[一二三四五六七八九十]+部分(\s|$)')
        secondary_re = re.compile(r'^[一二三四五六七八九十]+、')
        fig_tab_re = re.compile(r'^(图|表)\s*\d+')
        # 章节标题（及其续段），如"讨论"、"讨论（续）"、"结论(续)"
        section_head_re = re.compile(
            r'^(前言|讨论|结论|致谢|参考文献|附录|摘要|中文摘要|英文摘要|Abstract|ABSTRACT|中英文缩略词表)'
            r'([（(]\s*续\s*[）)])?$'
        )
        cjk_re = re.compile(r'[一-鿿]')

        # 界定正文起止
        in_body = False
        problems = []
        for para in paragraphs:
            text = para.text.strip()

            # 起点：第一段一级标题触发正文；一级标题本身不参与检查
            if primary_re.match(text):
                in_body = True
                continue
            if not in_body:
                continue

            # 结束：遇到"参考文献"独立段（进入参考文献列表，另有细则）
            if text == '参考文献':
                break

            if not text:
                continue

            # 跳过：二级标题、图/表标注、章节标题（含续）、表格内容、目录条目
            if secondary_re.match(text):
                continue
            if fig_tab_re.match(text):
                continue
            if section_head_re.match(text):
                continue
            if _in_table(para._element):
                continue
            if _is_toc_entry(para._element):
                continue

            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            # 跳过：Heading1/Heading2 段落（结构性标题，如"讨论""结论"
            # "参考文献"等；这些细则另有约束或本项不覆盖）
            if para_style_id in ('Heading1', 'Heading2'):
                continue

            for run_el in para._element.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text.strip():
                    continue
                # 只校验含 CJK 字符的 run（"除英文与数字外"）
                if not cjk_re.search(r_text):
                    continue

                rPr = run_el.find(qn('w:rPr'))
                east = sz = None
                if rPr is not None:
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None:
                        east = rf.get(qn('w:eastAsia'))
                    s = rPr.find(qn('w:sz'))
                    if s is not None:
                        sz = s.get(qn('w:val'))
                if east is None and para_style_id:
                    east = _resolve(para_style_id, 'eastAsia')
                if east is None:
                    east = _resolve('Normal', 'eastAsia')
                if sz is None and para_style_id:
                    sz = _resolve(para_style_id, 'sz')
                if sz is None:
                    sz = _resolve('Normal', 'sz')

                snippet = r_text.strip()[:15]
                if sz != '24':
                    problems.append(f"[{snippet}] 字号非小四(sz={sz})")
                    break
                if east not in ('宋体', 'SimSun'):
                    problems.append(f"[{snippet}] 中文字体非宋体(eastAsia={east})")
                    break
            if problems:
                break

        if problems:
            self.add_score(-3,
                           "正文（除英文与数字外）字体格式不满足小四号宋体 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_figure_table_notes(self):
        """检查图、表编号及名称段落格式 -3：居中对齐、一倍行距。

        细则两点，任意一段图/表注不满足即扣 -3：
          1) **居中对齐** —— w:pPr/w:jc@val == "center"
          2) **一倍行距** —— 办公软件（Word/WPS）中"单倍行距"对应
             w:spacing 的组合：
               a) 未设置 line（缺省即单倍），或
               b) w:spacing@line == "240" 且 lineRule 未设或为 "auto"
                  （240 = 20×12pt，是单倍行距的 twips 表示），或
               c) w:spacing@lineRule == "auto" 且未提供 line
             凡 line 显式取值非 "240"（如 "360" = 1.5 倍、"480" = 2 倍）
             或 lineRule ∈ {"exact", "atLeast"}（固定值/最小值），均不
             属于一倍行距。

        以上属性按继承链解析：段落直接 pPr → 段落样式 pPr → basedOn
        链 → Normal，办公软件依此链呈现段落格式。

        细则未要求内容（缩进、间距、大纲级别、边框、制表位、字体、
        字号等）一律不加以约束（字体归 check_figure_table_font 覆盖，
        排序格式归 check_figure_table_numbering 覆盖）。

        定位方式：与 check_figure_table_numbering 相同 —— 段首严格
        匹配 "图/表 + 数字 + 空白 + 名称"；目录条目（SDT/PAGEREF）
        与表格单元格内文本跳过。
        """
        paragraphs = list(self.doc.paragraphs)

        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _in_table(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    return True
                parent = parent.getparent()
            return False

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        strict_re = re.compile(r'^(图|表)\s?(\d+)\s+\S')

        problems = []
        for para in paragraphs:
            text = para.text.strip()
            if not strict_re.match(text):
                continue
            if _in_table(para._element):
                continue
            if _is_toc_entry(para._element):
                continue

            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            # 1) jc
            jc = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
            if jc is None and para_style_id:
                jc = _resolve(para_style_id, 'jc')
            if jc is None:
                jc = _resolve('Normal', 'jc')

            # 2) spacing
            line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
                if lineRule is None:
                    lineRule = _resolve(para_style_id, 'lineRule')
            if line is None:
                line = _resolve('Normal', 'line')
                if lineRule is None:
                    lineRule = _resolve('Normal', 'lineRule')

            snippet = text[:15]
            if jc != 'center':
                problems.append(f"[{snippet}] 非居中对齐(jc={jc})")
                continue
            # 一倍行距：line 未设 或 line==240 且 lineRule 缺省/auto
            ok_100 = (line is None) or (
                line == '240' and (lineRule is None or lineRule == 'auto')
            )
            if not ok_100:
                problems.append(f"[{snippet}] 非一倍行距(line={line},rule={lineRule})")
                continue

        if problems:
            self.add_score(-3,
                           "图、表编号及名称段落格式不满足居中对齐、一倍行距 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_figure_table_numbering(self):
        """检查图注/表注排序格式 -3：图1、图2 / 表1、表2 顺序编号。

        细则一点：图注、表注样式必须形如 "图1、图2、图3……" 与 "表1、
        表2、表3……"，即：
          · **前缀** 为独立汉字 "图" 或 "表"
          · **编号** 紧随其后为阿拉伯数字（无中间分隔符，如"图-1"、
            "图 1"允许一个空白间隔；细则示例 "图1" 无空格，"图 1"
            也是办公软件常见等价写法）
          · **编号从 1 开始，逐一递增**（1、2、3…，不允许跳号、
            不允许重号、不允许缺号）
          · 图与表**各自独立**编号

        以上是"排序格式"的完整含义。细则未要求内容（题注样式、编号
        位置、字体、编号是否含章节前缀等）一律不加以约束。

        办公软件适配：
          · 图/表注在 Word/WPS 中通常为独立段落。收集全文段落中
            strip 后以 "图" 或 "表" 起首的段落，对紧随的编号做
            正则捕获。允许 "图1" / "图 1" 两种写法（半角/全角空格均
            视为分隔符）；不允许中文数字 "图一"、罗马数字 "图I"、
            带前缀 "图1-1"（含破折号）等其它形式，均视为不符合"图1、
            图2"的样式。
          · 目录条目（SDT/PAGEREF）与表格单元格内文本跳过 —— 前者
            属"目录"页，后者属表内内容，非"图/表注"。
        """
        paragraphs = list(self.doc.paragraphs)

        def _in_table(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    return True
                parent = parent.getparent()
            return False

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        # 严格匹配 "图/表 + 数字 + 空白 + 名称"（图/表注段的办公软件
        # 通用写法："图1 名称"，编号与题注之间以空白分隔）。这条空白
        # 分隔要求同时排除了正文中"表8显示...""图1的..."这类紧跟中文
        # 字符的引用性表述 —— 那些不是图/表注段。
        strict_re = re.compile(r'^(图|表)\s?(\d+)\s+\S')
        # 宽匹配：段首为 "图" 或 "表" 且后续 40 字符内含"图注/表注可能
        # 具备的编号形式"。用于捕捉"图一/图A/图1-1/图 1.1/图1、"之类
        # 的**样式不符**情况。
        loose_re = re.compile(r'^(图|表)[^，。：；\s]{0,15}')

        fig_seq = []   # 依次收集的图编号
        tab_seq = []
        malformed = []  # 记录格式不符的样例

        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if _in_table(para._element):
                continue
            if _is_toc_entry(para._element):
                continue
            if not (text.startswith('图') or text.startswith('表')):
                continue

            # 判定"这是图/表注段"：段首为 图/表 后紧跟一段编号+空白+名称
            # 排除文中提及句（如"从表1可见……"）。图/表注段的典型特征
            # 是段落较短且紧接名称；这里用 loose_re 做**候选**过滤，然
            # 后严格判断。
            head = text[:20]
            # 排除以 图/表 起首但语句延续（"图1 显示..." 也算图注段；
            # "图1 表明本文..." 也算 —— 只要段首是编号 + 名称即可）；
            # 而 "如图1所示"、"表1的数据" 这类由于段首不是 "图/表 + 数字"，
            # 会在下一步匹配中筛掉。

            m_strict = strict_re.match(text)
            if m_strict:
                kind = m_strict.group(1)
                num = int(m_strict.group(2))
                if kind == '图':
                    fig_seq.append((num, head))
                else:
                    tab_seq.append((num, head))
                continue

            # 不满足严格 "图数字"/"表数字" 但段首以 图/表 起首 —— 疑似
            # 图/表注但样式不符。需进一步过滤：只把明显是标题段的收进
            # malformed，避免误把正文句 "图形展示..." 计入。判定标准：
            # 段首 图/表 后紧跟数字字符任意形式（含中文/罗马/含破折号
            # 等）或后跟名称样短段。
            after = text[1:]
            # 段首"图/表"后跟中文数字 或 罗马 或 数字+分隔符 → 疑似图/表注
            malformed_head_re = re.compile(
                r'^\s?([一二三四五六七八九十百]+|[IVXivx]+|\d+[\.\-])'
            )
            if malformed_head_re.match(after):
                malformed.append(head)
                continue
            # 段首 "图 " 或 "表 " 但后跟非数字（如 "图形展示..."）—— 非
            # 图/表注段，跳过；正文提及性质。

        problems = []
        # 图编号需为 1..len(fig_seq) 递增
        expected = 1
        for num, head in fig_seq:
            if num != expected:
                problems.append(f"图编号不连续: 期望{expected},实际[{head}]")
                break
            expected += 1
        expected = 1
        for num, head in tab_seq:
            if num != expected:
                problems.append(f"表编号不连续: 期望{expected},实际[{head}]")
                break
            expected += 1

        for h in malformed[:2]:
            problems.append(f"样式不符[{h}]")

        if problems:
            self.add_score(-3,
                           "图注、表注样式不满足\"图1、图2 / 表1、表2\"排序格式 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_figure_table_font(self):
        """检查图、表编号及名称字体格式 -3：五号、宋体。

        细则两点，任意一处不满足即扣 -3：
          1) **五号** —— w:sz@val="21"（半磅，10.5pt）
          2) **宋体** —— w:rFonts@eastAsia ∈ {"宋体", "SimSun"}
             （办公软件 Word/WPS 中 SimSun 是宋体的英文别名，UI 显示
              等价）

        细则明示"图、表**编号及名称**" —— 检查对象为整个图/表注段的
        所有 run（含"图/表"前缀、阿拉伯数字编号、名称文本）。因阿拉
        伯数字本身是西文字符，其可视字体由 ascii/hAnsi 控制，与本项
        细则约束的 eastAsia（中文字体）无关；但**字号 sz 是通用的**，
        对数字 run 也约束。故：
          · 字号：对每个非空 run 校验
          · 中文字体：只对含 CJK 字符的 run 校验

        以上属性均按继承链解析：run rPr → 段落样式 → basedOn 链 →
        Normal，办公软件依此链呈现字体。

        细则未要求内容（西文字体、加粗、颜色、下划线、字符间距、对齐、
        行距等）一律不加以约束（对齐/行距归 check_figure_table_notes
        另一条 -3 覆盖）。

        定位方式：与 check_figure_table_numbering 相同，段首严格匹配
        "图/表 + 数字 + 空白 + 名称"；目录条目（SDT/PAGEREF）与表格
        单元格内文本跳过。
        """
        paragraphs = list(self.doc.paragraphs)

        # 样式继承链
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        def _in_table(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    return True
                parent = parent.getparent()
            return False

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        strict_re = re.compile(r'^(图|表)\s?(\d+)\s+\S')
        cjk_re = re.compile(r'[一-鿿]')

        problems = []
        for para in paragraphs:
            text = para.text.strip()
            if not strict_re.match(text):
                continue
            if _in_table(para._element):
                continue
            if _is_toc_entry(para._element):
                continue

            pPr = para._element.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            snippet = text[:15]
            for run_el in para._element.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text.strip():
                    continue

                rPr = run_el.find(qn('w:rPr'))
                east = sz = None
                if rPr is not None:
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None:
                        east = rf.get(qn('w:eastAsia'))
                    s = rPr.find(qn('w:sz'))
                    if s is not None:
                        sz = s.get(qn('w:val'))
                if east is None and para_style_id:
                    east = _resolve(para_style_id, 'eastAsia')
                if east is None:
                    east = _resolve('Normal', 'eastAsia')
                if sz is None and para_style_id:
                    sz = _resolve(para_style_id, 'sz')
                if sz is None:
                    sz = _resolve('Normal', 'sz')

                # 字号（对所有 run）
                if sz != '21':
                    problems.append(f"[{snippet}] 字号非五号(sz={sz})")
                    break
                # 中文字体（仅对含 CJK 的 run）
                if cjk_re.search(r_text):
                    if east not in ('宋体', 'SimSun'):
                        problems.append(f"[{snippet}] 中文字体非宋体(eastAsia={east})")
                        break
            if problems:
                break

        if problems:
            self.add_score(-3,
                           "图、表编号及名称字体格式不满足五号宋体 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_figure_caption_position(self):
        """检查图名位置 -3：图名（图注）必须在图片下方。

        细则一点：图名（"图1 xxx"、"图2 xxx" ……）必须位于所对应
        图片的**下方**。若图名出现在图片上方（或找不到对应图片），
        视为不满足，扣 -3。

        办公软件（Word/WPS）适配说明：
          · 图片在 OOXML 中以 w:drawing（DrawingML）或 w:pict（VML）
            两种承载方式出现，均需检测。
          · "下方" = 图片所在段落 + 图名段落，二者段落索引满足
            `caption_index > image_index`，且中间**无其它图片段落**
            介入（避免"图1"名归到"图2"图片上）。允许中间有 0 个或
            少量空白/其它文本段（办公软件中常见"图片段 →（可选）空
            白段 → 图名段"结构）。
          · 若某"图 N"注前方在文档流上找不到匹配图片，视为该图注
            不满足；若图注下方紧接图片且图注上方无图片，也视为不
            满足（"图上文下"）。

        细则未要求内容（图片对齐、图片段与图注段之间是否夹注释文本、
        图片浮动锚定方式等）一律不加以约束。

        为避免重复扣分，任意一处图名位置异常即触发一次 -3。
        """
        paragraphs = list(self.doc.paragraphs)

        def _has_image(para_el):
            for _ in para_el.iter(qn('w:drawing')):
                return True
            for _ in para_el.iter(qn('w:pict')):
                return True
            return False

        def _in_table(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    return True
                parent = parent.getparent()
            return False

        def _is_toc_entry(para_el):
            parent = para_el.getparent()
            while parent is not None:
                if parent.tag == qn('w:sdt'):
                    return True
                parent = parent.getparent()
            for instr in para_el.iter(qn('w:instrText')):
                if instr.text and 'PAGEREF' in instr.text.upper():
                    return True
            return False

        # 段落索引 → 是否含图片
        image_idx_set = set()
        for i, para in enumerate(paragraphs):
            if _in_table(para._element):
                # 表格单元格内的图片不参与"图 N"图注位置判定
                continue
            if _has_image(para._element):
                image_idx_set.add(i)

        strict_re = re.compile(r'^图\s?(\d+)\s+\S')

        problems = []
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            m = strict_re.match(text)
            if not m:
                continue
            if _in_table(para._element):
                continue
            if _is_toc_entry(para._element):
                continue

            num = m.group(1)

            # 向上查找最近的图片段（同一节流内），限定 5 段窗口，避免
            # 把上一张图片的位置误归到当前图注 —— 图片段与图注段之间
            # 办公软件常见 0~2 段间隔（可能有空段或简短说明）
            prev_img_idx = None
            for j in range(i - 1, max(-1, i - 6), -1):
                if j in image_idx_set:
                    prev_img_idx = j
                    break

            # 向下查找最近的图片段
            next_img_idx = None
            for j in range(i + 1, min(len(paragraphs), i + 6)):
                if j in image_idx_set:
                    next_img_idx = j
                    break

            snippet = text[:15]
            if prev_img_idx is None:
                # 上方 5 段内无图片
                if next_img_idx is not None:
                    # 下方有图片 → "图上文下"，图名位置错误
                    problems.append(f"[图{num}] 图名位于图片上方[{snippet}]")
                else:
                    # 上下均未找到关联图片 → 视为图片缺失/未紧邻，
                    # 图名位置不满足
                    problems.append(f"[图{num}] 未在图片下方(未找到关联图片)[{snippet}]")
                continue
            # 上方 5 段内已找到图片；额外判断：若下方也有图片且比上方
            # 更近，则此图注被夹在两张图片之间但更靠近下面那张 —— 视
            # 为"图名位于下张图片上方"，位置错误
            if next_img_idx is not None:
                if (next_img_idx - i) < (i - prev_img_idx):
                    problems.append(f"[图{num}] 图名更靠近下方图片[{snippet}]")
                    continue

        if problems:
            self.add_score(-3,
                           "图名不在图片下方 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_table_caption_position(self):
        """检查表名位置 -3：表名（表注）必须在表格上方。

        细则一点：表名（"表1 xxx"、"表2 xxx" ……）必须位于所对应
        表格的**上方**。若表名出现在表格下方（或找不到对应表格），
        视为不满足，扣 -3。

        办公软件（Word/WPS）适配说明：
          · 表格在 OOXML 中是 body 下的 w:tbl 元素，与 w:p（段落）
            以文档流顺序交错出现。检查表名位置需按 body 的顺序遍历
            段落与表格，得到统一的序列索引，再比较位置。
          · "上方" = 表名段索引 < 表格索引，且中间**无其它表格**
            介入（避免把"表1"名归到"表2"表格上方）。办公软件中常见
            "表名段 → (可选空段) → 表格"结构。允许 0~5 段间隔。
          · SDT 内容（目录 TOC）中的"表 N ..."字样不视为表注；其内
            段落不参与判定。
          · 嵌套表格（tbl 内嵌 tbl）也计入表格集合 —— 因为表名与嵌
            套表也需保持上下位置关系。本项按 body 顺序取所有 w:tbl。

        细则未要求内容（表格对齐、表名段与表格之间是否夹注释文本、
        表格样式等）一律不加以约束。

        为避免重复扣分，任意一处表名位置异常即触发一次 -3。
        """

        def _iter_body_seq(body):
            """按 body 顺序产出 (kind, element)：
                 kind ∈ {'p','tbl','sdt_p','sdt_tbl'}
               这里将 SDT 内的内容单独标记，用于跳过目录条目。
               嵌套 sdt/tbl 递归处理。
            """
            for child in body.iterchildren():
                tag = child.tag
                if tag == qn('w:p'):
                    yield ('p', child)
                elif tag == qn('w:tbl'):
                    yield ('tbl', child)
                elif tag == qn('w:sdt'):
                    content = child.find(qn('w:sdtContent'))
                    if content is not None:
                        for sub in content.iterchildren():
                            if sub.tag == qn('w:p'):
                                yield ('sdt_p', sub)
                            elif sub.tag == qn('w:tbl'):
                                yield ('sdt_tbl', sub)

        body = self.doc.element.body

        seq = list(_iter_body_seq(body))
        # 表格所在序列索引（跳过 SDT 内表格 —— 目录本身无实体表格；
        # 若正文中真有 sdt 包裹的表格，通常是内容控件，属特殊情况）
        tbl_indices = [idx for idx, (k, _) in enumerate(seq) if k == 'tbl']

        strict_re = re.compile(r'^表\s?(\d+)\s+\S')

        problems = []
        for idx, (kind, elem) in enumerate(seq):
            if kind != 'p':
                continue
            text = ''.join((t.text or '') for t in elem.iter(qn('w:t'))).strip()
            m = strict_re.match(text)
            if not m:
                continue
            # 排除表格单元格内的段落（w:p 位于 w:tc 内不会作为 body 的
            # 直接子；但保险起见再检查一次）
            parent = elem.getparent()
            in_tc = False
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    in_tc = True
                    break
                parent = parent.getparent()
            if in_tc:
                continue

            num = m.group(1)

            # 向下找最近的表格（同一文档流内）
            next_tbl = None
            for j in tbl_indices:
                if j > idx:
                    next_tbl = j
                    break

            # 向上找最近的表格
            prev_tbl = None
            for j in reversed(tbl_indices):
                if j < idx:
                    prev_tbl = j
                    break

            snippet = text[:15]
            if next_tbl is None:
                # 下方无表格
                if prev_tbl is not None:
                    problems.append(f"[表{num}] 表名位于表格下方[{snippet}]")
                else:
                    problems.append(f"[表{num}] 未在表格上方(未找到关联表格)[{snippet}]")
                continue
            # 下方 5 段内应有表格；否则视为距离过远，也不算"表名在表格
            # 上方"（办公软件中通常表名段与表格段紧邻，中间不超过几段
            # 空白）
            if (next_tbl - idx) > 5:
                problems.append(f"[表{num}] 表名与表格距离过远(间隔{next_tbl - idx}段)[{snippet}]")
                continue
            # 若上方还有更近的表格 —— 表名被夹在两张表格之间，且更靠
            # 近上面那张 → 视为"表名位于上张表格下方"，位置错误
            if prev_tbl is not None:
                if (idx - prev_tbl) < (next_tbl - idx):
                    problems.append(f"[表{num}] 表名更靠近上方表格[{snippet}]")
                    continue

        if problems:
            self.add_score(-3,
                           "表名不在表格上方 —— "
                           + "；".join(problems[:3]),
                           "deduction")

    def check_discussion_title_font(self):
        """检查"讨论"页标题字体格式 -1：四号、黑体。

        细则两点，任意一点不满足即扣 -1：
          1) **四号** —— w:sz@val="28"（半磅，14pt）
          2) **黑体** —— w:rFonts@eastAsia ∈ {"黑体", "SimHei"}
             （办公软件 Word/WPS 中 SimHei 是黑体的英文别名，UI 显示
              等价）

        以上属性按继承链解析：run rPr → 段落样式 rPr → basedOn 链 →
        Normal，办公软件依此链呈现字体。

        细则未要求内容（西文字体、加粗、颜色、下划线、字符间距、
        对齐、行距、缩进等）一律不加以约束。若整个文档未出现"讨论"
        独立段标题，本项前置不满足，不扣分。

        定位方式：全文段落中 strip 后精确等于 "讨论" 的独立段。
        """
        paragraphs = list(self.doc.paragraphs)
        title_para = None
        for para in paragraphs:
            if para.text.strip() == '讨论':
                title_para = para
                break
        if title_para is None:
            return

        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

        problems = []
        for run_el in title_para._element.iter(qn('w:r')):
            r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
            if not r_text.strip():
                continue
            rPr = run_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'eastAsia')
            if east is None:
                east = _resolve('Normal', 'eastAsia')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')

            if sz != '28':
                problems.append(f"字号非四号(sz={sz})")
                break
            if east not in ('黑体', 'SimHei'):
                problems.append(f"中文字体非黑体(eastAsia={east})")
                break

        if problems:
            self.add_score(-1,
                           "讨论页标题字体格式不满足四号、黑体 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def check_discussion_title_para(self):
        """检查"讨论"页标题段落格式 -1：居中对齐、1.5倍行距。

        细则两点，任意一点不满足即扣 -1：
          1) **居中对齐** —— w:pPr/w:jc@val == "center"
          2) **1.5 倍行距** —— w:spacing@line=="360" 且 lineRule 未设
             或为 "auto"；"exact"/"atLeast" 视为固定值/最小值，均不
             属于 1.5 倍行距

        以上属性按继承链解析：段落直接 pPr → 段落样式 pPr → basedOn
        链 → Normal，办公软件依此链呈现段落格式。

        细则未要求内容（缩进、间距、大纲级别、边框、制表位、字体、
        字号等）一律不加以约束。若整个文档未出现"讨论"独立段标题，
        本项前置不满足，不扣分。

        定位方式：全文段落中 strip 后精确等于 "讨论" 的独立段。
        """
        paragraphs = list(self.doc.paragraphs)
        title_para = None
        for para in paragraphs:
            if para.text.strip() == '讨论':
                title_para = para
                break
        if title_para is None:
            return

        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

        jc = None
        if pPr is not None:
            el = pPr.find(qn('w:jc'))
            if el is not None:
                jc = el.get(qn('w:val'))
        if jc is None and para_style_id:
            jc = _resolve(para_style_id, 'jc')
        if jc is None:
            jc = _resolve('Normal', 'jc')

        line = lineRule = None
        if pPr is not None:
            el = pPr.find(qn('w:spacing'))
            if el is not None:
                line = el.get(qn('w:line'))
                lineRule = el.get(qn('w:lineRule'))
        if line is None and para_style_id:
            line = _resolve(para_style_id, 'line')
            if lineRule is None:
                lineRule = _resolve(para_style_id, 'lineRule')
        if line is None:
            line = _resolve('Normal', 'line')
            if lineRule is None:
                lineRule = _resolve('Normal', 'lineRule')

        problems = []
        if jc != 'center':
            problems.append(f"非居中对齐(jc={jc})")
        ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
        if not ok_150:
            problems.append(f"非1.5倍行距(line={line},rule={lineRule})")

        if problems:
            self.add_score(-1,
                           "讨论页标题段落格式不满足居中对齐、1.5倍行距 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def _iter_discussion_body_paragraphs(self):
        """产出「讨论」标题下方、到下一节标题（结论 / 参考文献 / 致谢 / 附录
        / 第X部分）之前的所有 body 段落（含 讨论（续） 之间的正文），
        跳过标题类段落自身与在表格中的段落、TOC 段落。"""
        end_re = re.compile(r'^第[一二三四五六七八九十]+部分(\s|$)')
        section_head_re = re.compile(
            r'^(讨论|结论|参考文献|致谢|附录|摘要|中文摘要|英文摘要|Abstract|ABSTRACT|前言|中英文缩略词表)'
            r'([（(]\s*续\s*[）)])?$'
        )
        started = False
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not started:
                if text == '讨论':
                    started = True
                continue
            # 终止：遇到下一节标题
            if text in ('结论', '参考文献', '致谢', '附录') or end_re.match(text):
                break
            if not text:
                continue
            # 跳过 讨论（续） / 结论（续） 等标题类段落
            if section_head_re.match(text):
                continue
            para_el = para._element
            # 跳过表格内段落
            parent = para_el.getparent()
            in_table = False
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    in_table = True
                    break
                parent = parent.getparent()
            if in_table:
                continue
            # 跳过 TOC（SDT / PAGEREF）段落
            is_toc = False
            p = para_el.getparent()
            while p is not None:
                if p.tag == qn('w:sdt'):
                    is_toc = True
                    break
                p = p.getparent()
            if not is_toc:
                for instr in para_el.iter(qn('w:instrText')):
                    if instr.text and 'PAGEREF' in instr.text.upper():
                        is_toc = True
                        break
            if is_toc:
                continue
            # 跳过 Heading1/Heading2 样式段落
            pPr = para_el.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None
            if para_style_id in ('Heading1', 'Heading2'):
                continue
            yield para

    def _iter_conclusion_body_paragraphs(self):
        """产出「结论」标题下方、到下一节标题（参考文献 / 致谢 / 附录）
        之前的所有 body 段落（含 结论（续） 之间的正文），跳过标题类
        段落自身与在表格中的段落、TOC 段落。"""
        section_head_re = re.compile(
            r'^(讨论|结论|参考文献|致谢|附录|摘要|中文摘要|英文摘要|Abstract|ABSTRACT|前言|中英文缩略词表)'
            r'([（(]\s*续\s*[）)])?$'
        )
        started = False
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not started:
                if text == '结论':
                    started = True
                continue
            # 终止：遇到下一节标题
            if text in ('参考文献', '致谢', '附录'):
                break
            if not text:
                continue
            if section_head_re.match(text):
                continue
            para_el = para._element
            parent = para_el.getparent()
            in_table = False
            while parent is not None:
                if parent.tag == qn('w:tc'):
                    in_table = True
                    break
                parent = parent.getparent()
            if in_table:
                continue
            is_toc = False
            p = para_el.getparent()
            while p is not None:
                if p.tag == qn('w:sdt'):
                    is_toc = True
                    break
                p = p.getparent()
            if not is_toc:
                for instr in para_el.iter(qn('w:instrText')):
                    if instr.text and 'PAGEREF' in instr.text.upper():
                        is_toc = True
                        break
            if is_toc:
                continue
            pPr = para_el.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None
            if para_style_id in ('Heading1', 'Heading2'):
                continue
            yield para

    def _build_style_maps_for_run(self):
        """构建 run 级样式继承链解析函数：sz 与 eastAsia。"""
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        return _resolve

    def _build_style_maps_for_para(self):
        """构建 段落级样式继承链解析函数：jc, spacing(line,lineRule),
        ind(firstLineChars,firstLine)。"""
        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = flc = fl = None
            if pPr is not None:
                jc_el = pPr.find(qn('w:jc'))
                if jc_el is not None:
                    jc = jc_el.get(qn('w:val'))
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    lineRule = sp.get(qn('w:lineRule'))
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    flc = ind.get(qn('w:firstLineChars'))
                    fl = ind.get(qn('w:firstLine'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule,
                    'firstLineChars': flc, 'firstLine': fl, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        return _resolve

    def check_discussion_body_font(self):
        """讨论页标题下方 —— 中文文本字体：小四号（sz=24）宋体
        (eastAsia ∈ {宋体, SimSun})。

        细则严格按点：
          1) 小四号 —— run rPr → 段落样式 rPr → basedOn → Normal 继承链
             解析后的 w:sz@val 必须为 "24"
          2) 宋体   —— 同链解析后的 w:rFonts@eastAsia ∈ {宋体, SimSun}
             （SimSun 是宋体在办公软件中的英文别名）

        仅检查含 CJK 字符的 run（细则明确"中文文本"）。西文/数字 run
        不作约束。西文字体、加粗、颜色等非细则要求项一律不检查。

        范围：讨论标题下方到下一节标题（结论 / 参考文献 / 致谢 / 附录
        / 第X部分）之前。跳过 讨论（续）等标题类段落、表格段落、TOC。
        """
        _resolve = self._build_style_maps_for_run()
        cjk_re = re.compile(r'[一-鿿]')
        bad = None
        for para in self._iter_discussion_body_paragraphs():
            para_el = para._element
            pPr = para_el.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            for run_el in para_el.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text or not cjk_re.search(r_text):
                    continue
                rPr = run_el.find(qn('w:rPr'))
                east = sz = None
                rStyle_id = None
                if rPr is not None:
                    rs = rPr.find(qn('w:rStyle'))
                    if rs is not None:
                        rStyle_id = rs.get(qn('w:val'))
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None:
                        east = rf.get(qn('w:eastAsia'))
                    s = rPr.find(qn('w:sz'))
                    if s is not None:
                        sz = s.get(qn('w:val'))
                if sz is None and rStyle_id:
                    sz = _resolve(rStyle_id, 'sz')
                if east is None and rStyle_id:
                    east = _resolve(rStyle_id, 'eastAsia')
                if sz is None and para_style_id:
                    sz = _resolve(para_style_id, 'sz')
                if east is None and para_style_id:
                    east = _resolve(para_style_id, 'eastAsia')
                if sz is None:
                    sz = _resolve('Normal', 'sz')
                if east is None:
                    east = _resolve('Normal', 'eastAsia')

                probs = []
                if sz != '24':
                    probs.append(f"字号非小四(sz={sz})")
                if east not in ('宋体', 'SimSun'):
                    probs.append(f"中文字体非宋体(eastAsia={east})")
                if probs:
                    bad = (para.text.strip()[:20], probs)
                    break
            if bad:
                break

        if bad:
            snippet, probs = bad
            self.add_score(-1,
                           "讨论页标题下方中文文本字体不满足小四号宋体 —— "
                           + "；".join(probs) + f"（首个不合项段:「{snippet}...」）",
                           "deduction")

    def check_discussion_body_para(self):
        """讨论页标题下方 —— 文本段落格式：1.5倍行距、首行缩进2字符。

        细则严格按点：
          1) 1.5倍行距 —— w:spacing@line=="360" 且 lineRule ∈ {缺省, "auto"}
             （办公软件 Word/WPS 中 line=360+lineRule=auto 即 1.5 倍行距；
              lineRule=exact / atLeast 不算倍数行距）
          2) 首行缩进 2 字符 —— 必须以"字符"单位设置：
             w:ind@firstLineChars=="200"（"字符"单位，随字号伸缩）。
             注意：仅设置 w:ind@firstLine（twips 绝对值，对应办公软件
             "厘米/磅"单位缩进）**不视为满足**，即便数值上恰好等于
             2 字符宽度也不算。

        以上属性按继承链解析：段落 pPr → 段落样式 pPr → basedOn → Normal。
        细则未要求内容（对齐、字体、颜色、缩进类型以外的项等）一律不
        检查。

        范围：同 check_discussion_body_font，跳过标题类、表格、TOC 段落。
        """
        _resolve = self._build_style_maps_for_para()
        bad = None
        for para in self._iter_discussion_body_paragraphs():
            para_el = para._element
            pPr = para_el.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            line = lineRule = flc = fl = None
            if pPr is not None:
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    lineRule = sp.get(qn('w:lineRule'))
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    flc = ind.get(qn('w:firstLineChars'))
                    fl = ind.get(qn('w:firstLine'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
            if lineRule is None and para_style_id:
                lineRule = _resolve(para_style_id, 'lineRule')
            if flc is None and para_style_id:
                flc = _resolve(para_style_id, 'firstLineChars')
            if fl is None and para_style_id:
                fl = _resolve(para_style_id, 'firstLine')
            if line is None:
                line = _resolve('Normal', 'line')
            if lineRule is None:
                lineRule = _resolve('Normal', 'lineRule')
            if flc is None:
                flc = _resolve('Normal', 'firstLineChars')
            if fl is None:
                fl = _resolve('Normal', 'firstLine')

            probs = []
            ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
            if not ok_150:
                probs.append(f"非1.5倍行距(line={line},rule={lineRule})")
            ok_ind = (flc == '200')
            if not ok_ind:
                probs.append(f"首行缩进非2字符(firstLineChars={flc},firstLine={fl})")
            if probs:
                bad = (para.text.strip()[:20], probs)
                break

        if bad:
            snippet, probs = bad
            self.add_score(-1,
                           "讨论页标题下方文本段落格式不满足1.5倍行距、首行缩进2字符 —— "
                           + "；".join(probs) + f"（首个不合项段:「{snippet}...」）",
                           "deduction")

    def check_conclusion_body_font(self):
        """结论页标题下方 —— 中文文本字体：小四号宋体。

        细则严格按点（同 讨论）：
          1) 小四号 —— sz=="24"
          2) 宋体   —— eastAsia ∈ {宋体, SimSun}

        仅检查含 CJK 的 run。范围：结论标题下方到 参考文献 / 致谢 / 附录
        之前；跳过 结论（续） 等标题类段落、表格、TOC。
        """
        _resolve = self._build_style_maps_for_run()
        cjk_re = re.compile(r'[一-鿿]')
        bad = None
        for para in self._iter_conclusion_body_paragraphs():
            para_el = para._element
            pPr = para_el.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            for run_el in para_el.iter(qn('w:r')):
                r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
                if not r_text or not cjk_re.search(r_text):
                    continue
                rPr = run_el.find(qn('w:rPr'))
                east = sz = None
                rStyle_id = None
                if rPr is not None:
                    rs = rPr.find(qn('w:rStyle'))
                    if rs is not None:
                        rStyle_id = rs.get(qn('w:val'))
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None:
                        east = rf.get(qn('w:eastAsia'))
                    s = rPr.find(qn('w:sz'))
                    if s is not None:
                        sz = s.get(qn('w:val'))
                if sz is None and rStyle_id:
                    sz = _resolve(rStyle_id, 'sz')
                if east is None and rStyle_id:
                    east = _resolve(rStyle_id, 'eastAsia')
                if sz is None and para_style_id:
                    sz = _resolve(para_style_id, 'sz')
                if east is None and para_style_id:
                    east = _resolve(para_style_id, 'eastAsia')
                if sz is None:
                    sz = _resolve('Normal', 'sz')
                if east is None:
                    east = _resolve('Normal', 'eastAsia')

                probs = []
                if sz != '24':
                    probs.append(f"字号非小四(sz={sz})")
                if east not in ('宋体', 'SimSun'):
                    probs.append(f"中文字体非宋体(eastAsia={east})")
                if probs:
                    bad = (para.text.strip()[:20], probs)
                    break
            if bad:
                break

        if bad:
            snippet, probs = bad
            self.add_score(-1,
                           "结论页标题下方中文文本字体不满足小四号宋体 —— "
                           + "；".join(probs) + f"（首个不合项段:「{snippet}...」）",
                           "deduction")

    def check_conclusion_body_para(self):
        """结论页标题下方 —— 文本段落格式：1.5倍行距、首行缩进2字符。

        细则严格按点（同 讨论）：
          1) 1.5倍行距 —— line=="360" 且 lineRule ∈ {缺省,"auto"}
          2) 首行缩进 2 字符 —— 严格：仅 firstLineChars=="200"（"字符"
             单位）视为满足；仅设置 firstLine（twips 绝对值，办公软件
             "厘米/磅"单位缩进）不算，即便数值等价于 2 字符也不算。

        范围：结论标题下方到 参考文献 / 致谢 / 附录 之前；跳过 结论（续）
        等标题类段落、表格、TOC。
        """
        _resolve = self._build_style_maps_for_para()
        bad = None
        for para in self._iter_conclusion_body_paragraphs():
            para_el = para._element
            pPr = para_el.find(qn('w:pPr'))
            pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
            para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

            line = lineRule = flc = fl = None
            if pPr is not None:
                sp = pPr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    lineRule = sp.get(qn('w:lineRule'))
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    flc = ind.get(qn('w:firstLineChars'))
                    fl = ind.get(qn('w:firstLine'))
            if line is None and para_style_id:
                line = _resolve(para_style_id, 'line')
            if lineRule is None and para_style_id:
                lineRule = _resolve(para_style_id, 'lineRule')
            if flc is None and para_style_id:
                flc = _resolve(para_style_id, 'firstLineChars')
            if fl is None and para_style_id:
                fl = _resolve(para_style_id, 'firstLine')
            if line is None:
                line = _resolve('Normal', 'line')
            if lineRule is None:
                lineRule = _resolve('Normal', 'lineRule')
            if flc is None:
                flc = _resolve('Normal', 'firstLineChars')
            if fl is None:
                fl = _resolve('Normal', 'firstLine')

            probs = []
            ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
            if not ok_150:
                probs.append(f"非1.5倍行距(line={line},rule={lineRule})")
            ok_ind = (flc == '200')
            if not ok_ind:
                probs.append(f"首行缩进非2字符(firstLineChars={flc},firstLine={fl})")
            if probs:
                bad = (para.text.strip()[:20], probs)
                break

        if bad:
            snippet, probs = bad
            self.add_score(-1,
                           "结论页标题下方文本段落格式不满足1.5倍行距、首行缩进2字符 —— "
                           + "；".join(probs) + f"（首个不合项段:「{snippet}...」）",
                           "deduction")

    def check_conclusion_title_font(self):
        """检查"结论"页标题字体格式 -1：四号、黑体。

        细则两点，任意一点不满足即扣 -1：
          1) **四号** —— w:sz@val="28"（半磅，14pt）
          2) **黑体** —— w:rFonts@eastAsia ∈ {"黑体", "SimHei"}
             （办公软件 Word/WPS 中 SimHei 是黑体的英文别名，UI 显示
              等价）

        以上属性按继承链解析：run rPr → 段落样式 rPr → basedOn 链 →
        Normal，办公软件依此链呈现字体。

        细则未要求内容（西文字体、加粗、颜色、下划线、字符间距、对齐、
        行距、缩进等）一律不加以约束。若整个文档未出现"结论"独立段
        标题，本项前置不满足，不扣分。

        定位方式：全文段落中 strip 后精确等于 "结论" 的独立段。避免
        匹配 "结论（续）" 或摘要中"结论 ..."起首的正文段。
        """
        paragraphs = list(self.doc.paragraphs)
        title_para = None
        for para in paragraphs:
            if para.text.strip() == '结论':
                title_para = para
                break
        if title_para is None:
            return

        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            rPr = style_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'eastAsia': east, 'sz': sz, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

        problems = []
        for run_el in title_para._element.iter(qn('w:r')):
            r_text = ''.join((t.text or '') for t in run_el.iter(qn('w:t')))
            if not r_text.strip():
                continue
            rPr = run_el.find(qn('w:rPr'))
            east = sz = None
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    east = rf.get(qn('w:eastAsia'))
                s = rPr.find(qn('w:sz'))
                if s is not None:
                    sz = s.get(qn('w:val'))
            if east is None and para_style_id:
                east = _resolve(para_style_id, 'eastAsia')
            if east is None:
                east = _resolve('Normal', 'eastAsia')
            if sz is None and para_style_id:
                sz = _resolve(para_style_id, 'sz')
            if sz is None:
                sz = _resolve('Normal', 'sz')

            if sz != '28':
                problems.append(f"字号非四号(sz={sz})")
                break
            if east not in ('黑体', 'SimHei'):
                problems.append(f"中文字体非黑体(eastAsia={east})")
                break

        if problems:
            self.add_score(-1,
                           "结论页标题字体格式不满足四号、黑体 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def check_conclusion_title_para(self):
        """检查"结论"页标题段落格式 -1：居中对齐、1.5倍行距。

        细则两点，任意一点不满足即扣 -1：
          1) **居中对齐** —— w:pPr/w:jc@val == "center"
          2) **1.5 倍行距** —— w:spacing@line=="360" 且 lineRule 未设
             或为 "auto"；"exact"/"atLeast" 视为固定值/最小值，均不
             属于 1.5 倍行距

        以上属性按继承链解析：段落直接 pPr → 段落样式 pPr → basedOn
        链 → Normal，办公软件依此链呈现段落格式。

        细则未要求内容（缩进、间距、大纲级别、边框、制表位、字体、
        字号等）一律不加以约束。若整个文档未出现"结论"独立段标题，
        本项前置不满足，不扣分。

        定位方式：全文段落中 strip 后精确等于 "结论" 的独立段。
        """
        paragraphs = list(self.doc.paragraphs)
        title_para = None
        for para in paragraphs:
            if para.text.strip() == '结论':
                title_para = para
                break
        if title_para is None:
            return

        styles_element = self.doc.styles.element
        style_map = {}

        def _extract(style_el):
            pPr = style_el.find(qn('w:pPr'))
            jc = line = lineRule = None
            if pPr is not None:
                el = pPr.find(qn('w:jc'))
                if el is not None:
                    jc = el.get(qn('w:val'))
                el = pPr.find(qn('w:spacing'))
                if el is not None:
                    line = el.get(qn('w:line'))
                    lineRule = el.get(qn('w:lineRule'))
            basedOn_el = style_el.find(qn('w:basedOn'))
            basedOn = basedOn_el.get(qn('w:val')) if basedOn_el is not None else None
            return {'jc': jc, 'line': line, 'lineRule': lineRule, 'basedOn': basedOn}

        for style_el in styles_element.iter(qn('w:style')):
            sid = style_el.get(qn('w:styleId'))
            if sid:
                style_map[sid] = _extract(style_el)

        def _resolve(style_id, key):
            visited = set()
            cur = style_id
            while cur and cur not in visited:
                visited.add(cur)
                data = style_map.get(cur)
                if data is None:
                    return None
                if data[key] is not None:
                    return data[key]
                cur = data['basedOn']
            return None

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle_el = pPr.find(qn('w:pStyle')) if pPr is not None else None
        para_style_id = pStyle_el.get(qn('w:val')) if pStyle_el is not None else None

        jc = None
        if pPr is not None:
            el = pPr.find(qn('w:jc'))
            if el is not None:
                jc = el.get(qn('w:val'))
        if jc is None and para_style_id:
            jc = _resolve(para_style_id, 'jc')
        if jc is None:
            jc = _resolve('Normal', 'jc')

        line = lineRule = None
        if pPr is not None:
            el = pPr.find(qn('w:spacing'))
            if el is not None:
                line = el.get(qn('w:line'))
                lineRule = el.get(qn('w:lineRule'))
        if line is None and para_style_id:
            line = _resolve(para_style_id, 'line')
            if lineRule is None:
                lineRule = _resolve(para_style_id, 'lineRule')
        if line is None:
            line = _resolve('Normal', 'line')
            if lineRule is None:
                lineRule = _resolve('Normal', 'lineRule')

        problems = []
        if jc != 'center':
            problems.append(f"非居中对齐(jc={jc})")
        ok_150 = (line == '360' and (lineRule is None or lineRule == 'auto'))
        if not ok_150:
            problems.append(f"非1.5倍行距(line={line},rule={lineRule})")

        if problems:
            self.add_score(-1,
                           "结论页标题段落格式不满足居中对齐、1.5倍行距 —— "
                           + "；".join(problems[:2]),
                           "deduction")

    def print_final_result(self):
        """打印最终结果"""
        print("\n" + "=" * 80)
        print("【评估结果汇总】")
        print("=" * 80)

        print("\n维度1（可用与可修改性）: ✓ 通过")

        print("\n维度2（完成度评分）:")

        if self.evaluation_details["维度2"]["得分项"]:
            print("\n  ★ 得分项（共" + str(len(self.evaluation_details["维度2"]["得分项"])) + "项）:")
            for item in self.evaluation_details["维度2"]["得分项"]:
                print("    +" + str(item["分数"]) + "分: " + item["描述"])

        if self.evaluation_details["维度2"]["扣分项"]:
            print("\n  ★ 扣分项（共" + str(len(self.evaluation_details["维度2"]["扣分项"])) + "项）:")
            for item in self.evaluation_details["维度2"]["扣分项"]:
                print("    " + str(item["分数"]) + "分: " + item["描述"])

        print("\n  ▶ 最终得分:", str(self.total_score), "分")
        print("\n" + "=" * 80)

        self.evaluation_details["维度2"]["总得分"] = self.total_score


def _locate_docx(dir_path: str) -> str:
    """在给定目录中定位待评估的 Word 文档，忽略以 ~$ 开头的临时文件。"""
    candidates = [
        name for name in os.listdir(dir_path)
        if name.lower().endswith(".docx") and not name.startswith("~$")
    ]
    if not candidates:
        raise FileNotFoundError("目录中未找到 .docx 文件: " + dir_path)
    return os.path.join(dir_path, candidates[0])


def evaluate(dir_path: str) -> dict:
    """统一入口：接收脚本所在目录路径，在该目录内定位并评估 Word 文档，返回结构化结果字典。

    返回结构参考"脚本接口差异与统一建议.md" §2.2。
    """
    result = {
        "id": "014",
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": True,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }

    try:
        if not os.path.isdir(dir_path):
            result["status"] = "error"
            result["error"] = "目录不存在: " + str(dir_path)
            return result

        file_path = _locate_docx(dir_path)
        result["file_name"] = os.path.basename(file_path)

        evaluator = WordDocumentEvaluator(file_path)
        evaluator.evaluate()

        # 维度一
        result["dim1_pass"] = bool(evaluator.dimension1_passed)
        result["dim1_reason"] = "" if evaluator.dimension1_passed else evaluator.dimension1_fail_reason

        # 维度二：逐条规则输出（契约要求每条 Dim2Item 都要出现，不得用汇总项代替未命中项）
        dim2_items = evaluator.build_dim2_items() if evaluator.dimension1_passed else []
        positive_hit = sum(it["delta"] for it in dim2_items if it["delta"] > 0)
        negative_hit = sum(it["delta"] for it in dim2_items if it["delta"] < 0)
        result["dim2_items"] = dim2_items
        # total_score：得分 = 命中加分项 + 命中扣分项之和
        # max_score：所有加分项满分之和（脚本细则中每个 +N 项的理论上限之和，与命中无关）
        result["total_score"] = (positive_hit + negative_hit) if evaluator.dimension1_passed else 0
        result["max_score"] = WordDocumentEvaluator.MAX_POSITIVE_SCORE
    except Exception as exc:
        result["status"] = "error"
        result["error"] = str(exc)

    return result


if __name__ == "__main__":
    # 仅用于本地自测：evaluate(脚本所在目录) 并打印 JSON
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    builtins.print(json.dumps(evaluate(_dir), ensure_ascii=False, indent=2))
