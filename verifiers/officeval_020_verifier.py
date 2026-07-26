"""
教案格式评估器 - 教案格式调整版.docx (脚本编号 020)
====================================

评分逻辑:
  维度一(可用与可修改性)  3 项 -- 任一不通过 → 总分 = 0, 不再评估维度二
  维度二(完成度)         18 项 -- 命中得分点 +N, 命中扣分点 -N, 累加求和

对外接口: 仅暴露 evaluate(dir_path: str) -> dict
  - dir_path: 脚本所在目录, 由脚本自行在其中定位并打开 .docx
  - 返回结构见文件末尾, 与《脚本接口差异与统一建议》§2.2 对齐
    (dim2_items[*].detail 统一为空字符串)

运行环境: pip install python-docx
"""

from __future__ import annotations

import re

SCRIPT_ID = "020"

import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Set, Tuple

from docx import Document

# =========================================================================
# 常量
# =========================================================================
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

SECTION_HEADERS = [
    "教材分析", "学情分析", "设计理念", "教学目标", "教学重难点",
    "教学准备", "课时安排", "教学思路说明", "教学过程", "板书设计",
    "作业设计", "教学反思",
]
DOC_TITLE = "《浅赏叙事古文，读懂处世道理》教学案例"
HALF_WIDTH_BRACKET = re.compile(r"[\[\]]")
RE_L1 = re.compile(r"^([一二三四五六七八九十]+)、")
RE_L2 = re.compile(r"^[（(]([一二三四五六七八九十]+)[)）]")
RE_L3 = re.compile(r"^(\d+)[．.\.]")
TWO_CHAR_CM_LO = 0.6
TWO_CHAR_CM_HI = 1.0

# =========================================================================
# XML 工具
# =========================================================================
def qn(tag: str) -> str:
    return f"{{{W}}}{tag}"

def emu2cm(emu) -> Optional[float]:
    return None if emu is None else emu / 360000

def halfpt2pt(hp) -> Optional[float]:
    return None if hp is None else hp / 2

def run_eastasia_font(run) -> Optional[str]:
    rPr = run._element.find(qn("rPr"))
    if rPr is None: return None
    rFonts = rPr.find(qn("rFonts"))
    if rFonts is None: return None
    return rFonts.get(qn("eastAsia")) or rFonts.get(qn("ascii"))

def run_size_pt(run) -> Optional[float]:
    rPr = run._element.find(qn("rPr"))
    if rPr is None: return None
    sz = rPr.find(qn("sz"))
    return halfpt2pt(int(sz.get(qn("val")))) if sz is not None else None

def run_is_bold(run) -> Optional[bool]:
    rPr = run._element.find(qn("rPr"))
    if rPr is None: return None
    b = rPr.find(qn("b"))
    if b is None: return None
    val = b.get(qn("val"))
    if val is None: return True
    return val not in ("0", "false")

def para_alignment(p) -> Optional[str]:
    pPr = p._element.find(qn("pPr"))
    if pPr is None: return None
    jc = pPr.find(qn("jc"))
    if jc is None: return None
    val = jc.get(qn("val"))
    return {"left": "LEFT", "right": "RIGHT", "center": "CENTER",
            "both": "JUSTIFY", "distribute": "DISTRIBUTE"}.get(val, val)

def para_line_spacing(p) -> Optional[Tuple[float, str]]:
    pPr = p._element.find(qn("pPr"))
    if pPr is None: return None
    sp = pPr.find(qn("spacing"))
    if sp is None: return None
    line = sp.get(qn("line"))
    rule = sp.get(qn("lineRule")) or "auto"
    if line is None: return None
    try:
        line = int(line)
    except ValueError:
        return None
    if rule == "auto":
        return (line / 240.0, "multiple")
    elif rule in ("exact", "atLeast"):
        return (line / 20.0, "pt")
    return (line, rule)

def para_indents_cm(p) -> Tuple[Optional[float], Optional[float]]:
    pf = p.paragraph_format
    return emu2cm(pf.first_line_indent), emu2cm(pf.left_indent)

def para_first_line_indent_chars(p) -> Optional[float]:
    """读取首行缩进的字符数，仅认 firstLineChars（WPS/Word 显示的字符单位）"""
    pPr = p._element.find(qn("pPr"))
    if pPr is None: return None
    ind = pPr.find(qn("ind"))
    if ind is None: return None
    chars = ind.get(qn("firstLineChars"))
    if chars is not None:
        return int(chars) / 100.0
    return None

def para_first_line_indent_chars_or_twips(p, font_pt: float = 12.0) -> Optional[float]:
    """读取首行缩进字符数，优先 firstLineChars>0，否则用 firstLine twips 换算
    返回 None 表示未显式设置(继承默认)"""
    pPr = p._element.find(qn("pPr"))
    if pPr is None: return None
    ind = pPr.find(qn("ind"))
    if ind is None: return None
    chars = ind.get(qn("firstLineChars"))
    if chars is not None and int(chars) > 0:
        return int(chars) / 100.0
    fl = ind.get(qn("firstLine"))
    if fl is not None and int(fl) > 0:
        return int(fl) / (font_pt * 20)
    # firstLineChars=0 或 firstLine=0: 视为未主动设置缩进值
    return None

def para_has_no_indent(p) -> bool:
    """检查段落是否顶格(无任何缩进), 任何缩进属性非零即判为有缩进"""
    pPr = p._element.find(qn("pPr"))
    if pPr is None: return True
    ind = pPr.find(qn("ind"))
    if ind is None: return True
    for attr in ("firstLine", "left", "firstLineChars", "leftChars"):
        val = ind.get(qn(attr))
        if val is not None and int(val) > 0:
            return False
    return True

# =========================================================================
# 章节切片
# =========================================================================
def collect_blocks(doc: Document) -> dict:
    blocks = {h: [] for h in SECTION_HEADERS}
    current: Optional[str] = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in SECTION_HEADERS:
            current = t
            blocks[current].append(p)
        elif current is not None:
            blocks[current].append(p)
    return blocks

def find_paragraphs_by_text(doc: Document, text: str) -> List:
    return [p for p in doc.paragraphs if p.text.strip() == text]

# =========================================================================
# 评分点
#   score: 正数 = 得分, 负数 = 扣分
#   desc :  命中时打印的简短描述
#   hit  :  命中条件 (True 则加入累加)
# =========================================================================
@dataclass
class Point:
    score: int
    desc: str
    hit: bool

# =========================================================================
# 维度一
# =========================================================================
def check_dim1(path: str) -> Tuple[bool, Optional[Document], str]:
    # 1.1 文件格式 / 可打开
    #   只识别 OOXML 的 .docx; 二进制 .doc 直接判定失败
    ext = Path(path).suffix.lower()
    ext_ok = ext == ".docx"
    try:
        doc = Document(path)
        _ = doc.paragraphs
        can_open = True
    except Exception:
        doc = None
        can_open = False
    if not (ext_ok and can_open):
        return False, None, f"文件 {ext!r} 解析{'失败' if not can_open else '正常但扩展名不符 (仅支持 .docx)'}"

    return True, doc, ""

# =========================================================================
# 维度二 — 每项返回 (Point, 一些副产物供后续项使用)
# =========================================================================
def style_match(runs, font_kw, size_pt, bold):
    if not runs: return False
    for r in runs:
        f = run_eastasia_font(r)
        s = run_size_pt(r)
        b = run_is_bold(r)
        if f is None or font_kw not in f: return False
        if s is None or abs(s - size_pt) > 0.5: return False
        if bold is not None and b != bold: return False
    return True

def check_dim2(doc: Document) -> List[Point]:
    blocks = collect_blocks(doc)
    proc_paras = blocks.get("教学过程", [])
    proc_body = proc_paras[1:] if proc_paras else []

    # ---- 自动编号支持: 若段落走 Word 编号列表, 编号前缀不写入 p.text,
    # 需读取 numbering.xml 里对应级别的 numFmt/lvlText 才能识别 ----
    CHINESE_AUTO_FMTS = {"chineseCounting", "chineseCountingThousand",
                         "ideographTraditional", "ideographZodiac"}

    def _numbering_lvl_element(p):
        """返回段落自动编号所对应的 <w:lvl> 元素; 无自动编号则 None"""
        pPr = p._element.find(qn("pPr"))
        if pPr is None: return None
        numPr = pPr.find(qn("numPr"))
        if numPr is None: return None
        numId_el = numPr.find(qn("numId"))
        ilvl_el = numPr.find(qn("ilvl"))
        if numId_el is None: return None
        try:
            numId = int(numId_el.get(qn("val")))
        except (TypeError, ValueError):
            return None
        try:
            ilvl = int(ilvl_el.get(qn("val"))) if ilvl_el is not None else 0
        except (TypeError, ValueError):
            ilvl = 0
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part is None: return None
            numbering_el = numbering_part.element
        except Exception:
            return None
        # numId -> abstractNumId
        abs_id = None
        for num in numbering_el.findall(qn("num")):
            if num.get(qn("numId")) == str(numId):
                aeid = num.find(qn("abstractNumId"))
                if aeid is not None:
                    abs_id = aeid.get(qn("val"))
                break
        if abs_id is None: return None
        # abstractNum -> lvl[ilvl]
        for absnum in numbering_el.findall(qn("abstractNum")):
            if absnum.get(qn("abstractNumId")) != abs_id: continue
            for lvl in absnum.findall(qn("lvl")):
                if lvl.get(qn("ilvl")) == str(ilvl):
                    return lvl
        return None

    def _numbering_lvl_format(p):
        """若段落使用自动编号, 返回 (numFmt, lvlText); 否则 None"""
        lvl = _numbering_lvl_element(p)
        if lvl is None: return None
        numFmt_el = lvl.find(qn("numFmt"))
        lvlText_el = lvl.find(qn("lvlText"))
        numFmt = numFmt_el.get(qn("val")) if numFmt_el is not None else None
        lvlText = lvlText_el.get(qn("val")) if lvlText_el is not None else None
        return (numFmt, lvlText)

    def _numbering_lvl_ind(p):
        """若段落自动编号级别定义了段落缩进, 返回其 <w:ind> 元素; 否则 None.
        自动编号标题的首行缩进/左缩进常写在 numbering.xml 的 lvl/pPr/ind,
        而非段落自身 pPr, 需一并读取才能正确判定缩进。"""
        lvl = _numbering_lvl_element(p)
        if lvl is None: return None
        lvl_pPr = lvl.find(qn("pPr"))
        if lvl_pPr is None: return None
        return lvl_pPr.find(qn("ind"))

    def _is_auto_l1(p) -> bool:
        """段落是否为 numPr 渲染出的一级标题 (中文数字 + 顿号, 如"一、二、三、")"""
        fmt = _numbering_lvl_format(p)
        if fmt is None:
            return False
        numFmt, lvlText = fmt
        return numFmt in CHINESE_AUTO_FMTS and lvlText == "%1、"

    def _is_auto_l2(p) -> bool:
        """段落是否为 numPr 渲染出的二级标题 (全角括号中文数字, 如"（一）（二）")"""
        fmt = _numbering_lvl_format(p)
        if fmt is None:
            return False
        numFmt, lvlText = fmt
        return numFmt in CHINESE_AUTO_FMTS and lvlText == "（%1）"

    def _is_auto_l3(p) -> bool:
        """段落是否为 numPr 渲染出的三级标题 (阿拉伯数字 + 点号, 如"1." "2.")

        - numFmt="decimal" 对应"1,2,3,..." 自动编号
        - lvlText 兼容半角点 "%1." 与全角点 "%1．" (Word/WPS 保存差异)
        """
        fmt = _numbering_lvl_format(p)
        if fmt is None:
            return False
        numFmt, lvlText = fmt
        return numFmt == "decimal" and lvlText in ("%1.", "%1．")

    # 一级标题 l1_paras: 兼容"文本手工输入"与"Word 自动编号"两种编号落地方式
    #   A) 手工: 段落文本以"中文数字 + 顿号"开头 -> RE_L1 匹配
    #   B) 自动: 段落 numPr 对应的 numFmt=chineseCounting 类, lvlText="%1、"
    l1_paras = [p for p in proc_body
                if RE_L1.match(p.text.strip()) or _is_auto_l1(p)]
    # 二级标题 l2_paras: 同样兼容手工文本与自动编号
    #   A) 手工: 段落文本以全角/半角括号中文数字开头 -> RE_L2 匹配
    #   B) 自动: 段落 numPr 对应的 numFmt=chineseCounting 类, lvlText="（%1）"
    l2_paras = [p for p in proc_body
                if RE_L2.match(p.text.strip()) or _is_auto_l2(p)]
    # 三级标题 l3_paras: 同样兼容手工文本与自动编号
    #   A) 手工: 段落文本以"数字 + 点(半角/全角)"开头 -> RE_L3 匹配
    #   B) 自动: 段落 numPr 对应的 numFmt=decimal, lvlText="%1." 或 "%1．"
    l3_paras = [p for p in proc_body
                if RE_L3.match(p.text.strip()) or _is_auto_l3(p)]

    points: List[Point] = []

    # ---- +3 页面 A4 纵向 + 页边距 上 2.6 / 下 2.6 / 左 2.4 / 右 2.4 cm ----
    # 直接读取 sectPr/pgSz、sectPr/pgMar 的 twips 值 (与 Word/WPS "页面设置"
    # 对话框显示的数值同源), 对所有分节都进行校验
    def _twips2cm(v):
        return int(v) / 1440.0 * 2.54 if v is not None else None

    page_ok = bool(doc.sections)
    for _sec in doc.sections:
        sectPr = _sec._sectPr
        pgSz = sectPr.find(qn("pgSz"))
        pgMar = sectPr.find(qn("pgMar"))
        if pgSz is None or pgMar is None:
            page_ok = False; break
        w_cm = _twips2cm(pgSz.get(qn("w")))
        h_cm = _twips2cm(pgSz.get(qn("h")))
        # Word/WPS: orient 属性缺省即为 portrait (纵向)
        orient = pgSz.get(qn("orient"))
        portrait = orient in (None, "portrait")
        # A4 纵向: 宽 21.0 cm, 高 29.7 cm
        a4_ok = (w_cm is not None and h_cm is not None
                 and abs(w_cm - 21.0) < 0.1 and abs(h_cm - 29.7) < 0.1)
        t_cm = _twips2cm(pgMar.get(qn("top")))
        b_cm = _twips2cm(pgMar.get(qn("bottom")))
        l_cm = _twips2cm(pgMar.get(qn("left")))
        r_cm = _twips2cm(pgMar.get(qn("right")))
        margin_ok = all(v is not None and abs(v - t) < 0.1 for v, t in
                        [(t_cm, 2.6), (b_cm, 2.6), (l_cm, 2.4), (r_cm, 2.4)])
        if not (a4_ok and portrait and margin_ok):
            page_ok = False; break
    points.append(Point(3,
        "页面 A4 纵向, 页边距 上 2.6 / 下 2.6 / 左 2.4 / 右 2.4 cm",
        page_ok))

    # ---- +1 学案标题《浅赏叙事古文，读懂处世道理》教学案例 字体 黑体/小二号/加粗 ----
    # 定位标题段落: 文本精确匹配 DOC_TITLE
    title = next((p for p in doc.paragraphs if p.text.strip() == DOC_TITLE), None)

    def _title_font_ok(p) -> bool:
        if p is None:
            return False
        # 只检查含实际文字的 run, 空白/占位 run 不参与 (办公软件里空 run 不显示格式)
        text_runs = [r for r in p.runs if r.text.strip()]
        if not text_runs:
            return False
        for r in text_runs:
            # 字体: 中文字符走 rFonts@eastAsia; Word/WPS 中"黑体"可存为 "黑体" 或 "SimHei"
            font = run_eastasia_font(r) or ""
            if not ("黑体" in font or font.lower() == "simhei"):
                return False
            # 字号: 小二号 = 18 磅 (与 Word/WPS "字号"下拉里 "小二" 对应)
            size = run_size_pt(r)
            if size is None or abs(size - 18) > 0.5:
                return False
            # 加粗: 存在 <w:b/> 且 val 未被显式关闭
            if run_is_bold(r) is not True:
                return False
        return True

    points.append(Point(1,
        "学案标题《浅赏叙事古文，读懂处世道理》教学案例 字体 黑体/小二号/加粗",
        _title_font_ok(title)))

    # ---- +1 学案标题《浅赏叙事古文，读懂处世道理》教学案例
    #      段落格式为居中对齐、行间距 38 磅 ----
    # 说明:
    #   - "居中对齐": Word/WPS 的 <w:jc w:val="center"/>; 若段落未直接设置,
    #     需回溯到段落所属样式 (pStyle -> style.paragraph_format) 里查, 因为
    #     办公软件显示时按样式继承结果呈现
    #   - "行间距 38 磅": Word/WPS 中"行距"选择"固定值"并输入 38 磅
    #     → <w:spacing w:line="760" w:lineRule="exact"/> (760 = 38*20)
    #     "最小值 38 磅" 也在办公软件里呈现为 ≥38 磅, 但细则明确写"行间距 38 磅"
    #     指精确 38 磅, 故只认 exact
    def _resolve_align(p):
        """返回段落最终生效的对齐 (含样式继承), 未设置则返回 None"""
        a = para_alignment(p)
        if a is not None:
            return a
        # 回溯样式链
        style = p.style
        while style is not None:
            pf = style.paragraph_format
            if pf is not None and pf.alignment is not None:
                # WD_PARAGRAPH_ALIGNMENT: CENTER=1, LEFT=0, RIGHT=2, JUSTIFY=3
                return {0: "LEFT", 1: "CENTER", 2: "RIGHT",
                        3: "JUSTIFY", 4: "DISTRIBUTE"}.get(int(pf.alignment))
            style = getattr(style, "base_style", None)
        return None

    def _resolve_line_spacing(p):
        """返回 (value, unit) 或 None; 含样式继承"""
        ls = para_line_spacing(p)
        if ls is not None:
            return ls
        style = p.style
        while style is not None:
            sp_el = style.element.find(qn("pPr"))
            if sp_el is not None:
                spacing = sp_el.find(qn("spacing"))
                if spacing is not None and spacing.get(qn("line")) is not None:
                    line = int(spacing.get(qn("line")))
                    rule = spacing.get(qn("lineRule")) or "auto"
                    if rule == "auto":
                        return (line / 240.0, "multiple")
                    return (line / 20.0, "pt")
            style = getattr(style, "base_style", None)
        return None

    if title is not None:
        align = _resolve_align(title)
        ls = _resolve_line_spacing(title)
        center_ok = align == "CENTER"
        # 行间距 38 磅: 严格要求 exact 规则下 38 磅 (允许 ±0.5 磅舍入)
        # 需要通过 XML 直接读 lineRule 以区分 exact / atLeast
        line_ok = False
        pPr = title._element.find(qn("pPr"))
        if pPr is not None:
            sp = pPr.find(qn("spacing"))
            if sp is not None:
                line = sp.get(qn("line"))
                rule = sp.get(qn("lineRule"))
                if line is not None and rule == "exact":
                    try:
                        if abs(int(line) / 20.0 - 38.0) <= 0.5:
                            line_ok = True
                    except ValueError:
                        pass
        title_p_ok = center_ok and line_ok
    else:
        title_p_ok = False
    points.append(Point(1,
        "学案标题《浅赏叙事古文，读懂处世道理》教学案例 "
        "段落居中对齐, 行间距 38 磅",
        title_p_ok))

    # ---- +5 12 个大板块标题 (教材分析/学情分析/设计理念/教学目标/教学重难点/
    #      教学准备/课时安排/教学思路说明/教学过程/板书设计/作业设计/教学反思)
    #      黑体/小四号/不加粗 ----
    # 说明:
    #   - 字体: 中文字符走 rFonts@eastAsia; "黑体" 在 Word/WPS 中可存为
    #     中文 "黑体" 或英文 "SimHei", 两者办公软件显示一致
    #   - 字号: 小四号 = 12 磅 (Word/WPS "字号"下拉里的"小四")
    #   - 不加粗: <w:b/> 不存在, 或存在但 val 为 "0"/"false"
    #     若段落级样式默认加粗但 run 未显式关掉, 办公软件仍显示为加粗,
    #     故还需回溯样式链, 确认最终生效为"不加粗"
    def _resolve_run_bold(run, para) -> bool:
        """返回 run 在办公软件里最终呈现的加粗状态"""
        b = run_is_bold(run)
        if b is not None:
            return b
        # run 未设置 -> 回溯段落 rPr (pPr/rPr)
        pPr = para._element.find(qn("pPr"))
        if pPr is not None:
            rPr = pPr.find(qn("rPr"))
            if rPr is not None:
                bel = rPr.find(qn("b"))
                if bel is not None:
                    val = bel.get(qn("val"))
                    return val not in ("0", "false") if val is not None else True
        # 回溯段落样式链
        style = para.style
        while style is not None:
            sfont = style.font
            if sfont is not None and sfont.bold is not None:
                return bool(sfont.bold)
            style = getattr(style, "base_style", None)
        return False  # 均未设置 -> 默认不加粗

    def _section_heading_ok(p) -> bool:
        text_runs = [r for r in p.runs if r.text.strip()]
        if not text_runs:
            return False
        for r in text_runs:
            font = run_eastasia_font(r) or ""
            if not ("黑体" in font or font.lower() == "simhei"):
                return False
            size = run_size_pt(r)
            if size is None or abs(size - 12) > 0.5:  # 小四号 = 12 磅
                return False
            if _resolve_run_bold(r, p):  # 最终生效为加粗 → 不符合"不加粗"
                return False
        return True

    sec_font_ok = True
    for h in SECTION_HEADERS:
        ps = blocks.get(h, [])
        # 必须存在该板块标题段落, 且其字体格式满足要求
        if not ps or not _section_heading_ok(ps[0]):
            sec_font_ok = False
            break
    points.append(Point(5,
        "12 个大板块标题 (教材分析/学情分析/设计理念/教学目标/教学重难点/"
        "教学准备/课时安排/教学思路说明/教学过程/板书设计/作业设计/教学反思) "
        "黑体/小四号/不加粗",
        sec_font_ok))

    # ---- +3 12 个大板块标题 (教材分析/学情分析/设计理念/教学目标/教学重难点/
    #      教学准备/课时安排/教学思路说明/教学过程/板书设计/作业设计/教学反思)
    #      顶格排列, 不设置任何缩进 ----
    # 说明:
    #   Word/WPS "段落"对话框的"缩进"区包含: 左侧、右侧、特殊格式(首行/悬挂)
    #   对应 XML 属性 (twips 或 字符百分比 chars):
    #     left / leftChars       —— 左缩进
    #     right / rightChars     —— 右缩进
    #     firstLine / firstLineChars —— 首行缩进
    #     hanging / hangingChars —— 悬挂缩进
    #   任一被显式设置且不为 0 都会在办公软件里呈现出缩进 → 不满足"顶格"
    #   还需回溯段落样式链, 因为通过样式(如"标题 1")继承的缩进同样会显示出来
    INDENT_ATTRS = ("left", "leftChars", "right", "rightChars",
                    "firstLine", "firstLineChars", "hanging", "hangingChars")

    def _ind_all_zero(ind_el) -> bool:
        """<w:ind> 元素内所有缩进属性都为 0 (或未设置) 才返回 True"""
        if ind_el is None:
            return True
        for a in INDENT_ATTRS:
            v = ind_el.get(qn(a))
            if v is None:
                continue
            try:
                if int(v) != 0:
                    return False
            except ValueError:
                return False
        return True

    def _para_no_indent_final(p) -> bool:
        """段落在办公软件里最终呈现是否无任何缩进 (含样式继承)"""
        # 1) 段落直接设置
        pPr = p._element.find(qn("pPr"))
        if pPr is not None:
            ind = pPr.find(qn("ind"))
            if ind is not None and not _ind_all_zero(ind):
                return False
        # 2) 回溯段落样式链, 任一祖先样式设置了非零缩进则最终仍有缩进
        style = p.style
        while style is not None:
            spPr = style.element.find(qn("pPr"))
            if spPr is not None:
                s_ind = spPr.find(qn("ind"))
                if s_ind is not None and not _ind_all_zero(s_ind):
                    # 只有当当前段落自己"显式清零"才能覆盖样式缩进,
                    # 否则办公软件按样式呈现有缩进
                    self_ind = pPr.find(qn("ind")) if pPr is not None else None
                    if self_ind is None:
                        return False
                    # 段落显式将全部缩进属性覆盖为 0 才算清零
                    for a in INDENT_ATTRS:
                        sv = s_ind.get(qn(a))
                        if sv is None:
                            continue
                        try:
                            if int(sv) == 0:
                                continue
                        except ValueError:
                            return False
                        # 样式该属性非零, 段落必须显式设为 0 才覆盖
                        selfv = self_ind.get(qn(a))
                        if selfv is None or int(selfv) != 0:
                            return False
            style = getattr(style, "base_style", None)
        return True

    sec_indent_ok = True
    for h in SECTION_HEADERS:
        ps = blocks.get(h, [])
        # 板块标题段落不存在即视为整项不通过 (细则要求这 12 个板块标题都顶格)
        if not ps:
            sec_indent_ok = False
            break
        if not _para_no_indent_final(ps[0]):
            sec_indent_ok = False
            break
    points.append(Point(3,
        "12 个大板块标题 (教材分析/学情分析/设计理念/教学目标/教学重难点/"
        "教学准备/课时安排/教学思路说明/教学过程/板书设计/作业设计/教学反思) "
        "顶格排列, 不设置任何缩进",
        sec_indent_ok))

    # ---- +3 教学过程一级标题 字体 ----
    l1_font_ok = bool(l1_paras) and all(
        style_match(p.runs, "黑体", 12, False) for p in l1_paras)
    points.append(Point(3, "教学过程一级标题 黑体/小四号/不加粗", l1_font_ok))

    # ---- +3 教学过程板块一级标题 段落不顶格, 首行缩进两字符 ----
    # 说明:
    #   - "不顶格": 段落在办公软件里最终呈现为有缩进 (与"顶格"相对)
    #     即 <w:ind> 中至少存在一个非零缩进属性 (或通过样式链继承出非零缩进)
    #   - "首行缩进两字符": Word/WPS "段落"对话框 "特殊格式" 选"首行缩进",
    #     值为"2 字符"
    #     XML 层面对应 <w:ind w:firstLineChars="200"/> (100 = 1 字符)
    #     若仅用 firstLine (twips) 表达, 换算需按段落中文字号: 2 字符 = 2*字号*20 twips
    #     Word/WPS 保存"2 字符"缩进时通常会写 firstLineChars="200",
    #     故优先认 firstLineChars=200; 否则用 firstLine + 字号换算兼容
    def _para_first_line_2chars(p) -> bool:
        pPr = p._element.find(qn("pPr"))
        ind = pPr.find(qn("ind")) if pPr is not None else None
        # 自动编号标题的缩进常写在 numbering.xml 的 lvl/pPr/ind, 段落自身无 ind 时回退读取
        if ind is None:
            ind = _numbering_lvl_ind(p)
        if ind is None:
            return False
        # 优先 firstLineChars: 100 = 1 字符, 允许 ±5 覆盖办公软件舍入
        flc = ind.get(qn("firstLineChars"))
        if flc is not None:
            try:
                if abs(int(flc) - 200) <= 5:
                    return True
            except ValueError:
                pass
        # 退化到 firstLine (twips) + run 字号换算
        fl = ind.get(qn("firstLine"))
        if fl is not None:
            # 取段落首个含文本的 run 字号, 无则退到 12 磅 (小四, 教学过程一级标题细则字号)
            size_pt = None
            for r in p.runs:
                if r.text.strip():
                    size_pt = run_size_pt(r)
                    if size_pt is not None:
                        break
            if size_pt is None:
                size_pt = 12.0
            try:
                chars = int(fl) / (size_pt * 20.0)
                if abs(chars - 2.0) <= 0.1:
                    return True
            except ValueError:
                pass
        return False

    def _para_not_top_aligned(p) -> bool:
        """段落在办公软件里呈现为非顶格 (存在任一非零缩进)"""
        # 段落自身/样式链无缩进时, 若自动编号级别定义了非零缩进, 仍视为非顶格
        if not _para_no_indent_final(p):
            return True
        lvl_ind = _numbering_lvl_ind(p)
        if lvl_ind is not None and not _ind_all_zero(lvl_ind):
            return True
        return False

    l1_ind_ok = bool(l1_paras) and all(
        _para_not_top_aligned(p) and _para_first_line_2chars(p)
        for p in l1_paras)
    points.append(Point(3,
        "教学过程板块一级标题段落不顶格, 首行缩进两字符",
        l1_ind_ok))

    # ---- +3 教学过程板块二级标题 (如 (一)(二)(三)……)
    #      字体为 黑体 / 小四号 / 不加粗 ----
    # 说明:
    #   - 二级标题定位: 已有 l2_paras (段落文本以 "(一)"/"（一）" 等中文数字
    #     圆括号编号开头, 括号可全角/半角), 与细则示例 "(一)(二)(三)……" 对应
    #   - 字体: rFonts@eastAsia 命中 "黑体" 或 "SimHei" (Word/WPS 里显示都是黑体)
    #   - 字号: 小四号 = 12 磅 (Word/WPS "字号"下拉里的"小四"), ±0.5 磅覆盖半磅整数化
    #   - 不加粗: 按办公软件的解析顺序回溯 run→段落rPr→样式链,
    #     最终生效为加粗即判为不符合 (与前面 12 板块标题的口径一致, 复用 _resolve_run_bold)
    def _l2_heading_ok(p) -> bool:
        text_runs = [r for r in p.runs if r.text.strip()]
        if not text_runs:
            return False
        for r in text_runs:
            font = run_eastasia_font(r) or ""
            if not ("黑体" in font or font.lower() == "simhei"):
                return False
            size = run_size_pt(r)
            if size is None or abs(size - 12) > 0.5:  # 小四号 = 12 磅
                return False
            if _resolve_run_bold(r, p):  # 最终生效为加粗 → 不符合"不加粗"
                return False
        return True

    l2_font_ok = bool(l2_paras) and all(_l2_heading_ok(p) for p in l2_paras)
    points.append(Point(3,
        "教学过程板块二级标题 (如 (一)(二)(三)……) 字体 黑体/小四号/不加粗",
        l2_font_ok))

    # ---- +3 教学过程板块二级标题 段落不顶格, 首行缩进两字符 ----
    # 说明:
    #   - "不顶格": 段落在办公软件里最终呈现为有缩进 (与"顶格"相对),
    #     即 <w:ind> 中至少存在一个非零缩进属性 (或通过样式链继承出非零缩进)
    #     复用前面已定义的 _para_not_top_aligned (基于 _para_no_indent_final 取反)
    #   - "首行缩进两字符": Word/WPS "段落"对话框 "特殊格式" 选"首行缩进",
    #     值为"2 字符"; 优先认 <w:ind w:firstLineChars="200"/>,
    #     否则用 firstLine twips + 字号换算兼容 (复用 _para_first_line_2chars)
    l2_ind_ok = bool(l2_paras) and all(
        _para_not_top_aligned(p) and _para_first_line_2chars(p)
        for p in l2_paras)
    points.append(Point(3,
        "教学过程板块二级标题段落不顶格, 首行缩进两字符",
        l2_ind_ok))

    # ---- +3 教学过程板块三级标题 (如 1、2、3……)
    #      字体为 宋体 / 小四号 / 加粗 ----
    # 说明:
    #   - 三级标题定位: 已有 l3_paras (段落文本以阿拉伯数字加 "." 或 "．" 开头,
    #     即 RE_L3), 与细则示例 "1、2、3……" 的编号语义对应
    #     (细则用顿号泛指编号形式, 文档实际采用 "1." / "1．", 已在 l3_paras 中匹配)
    #   - 字体: rFonts@eastAsia 命中 "宋体" 或 "SimSun" (Word/WPS 里显示都是宋体)
    #   - 字号: 小四号 = 12 磅 (Word/WPS "字号"下拉里的"小四"), ±0.5 磅覆盖半磅整数化
    #   - 加粗: 按办公软件解析顺序回溯 run→段落rPr→样式链, 最终生效为加粗
    #     (与前面标题项口径一致, 复用 _resolve_run_bold)
    def _l3_heading_ok(p) -> bool:
        text_runs = [r for r in p.runs if r.text.strip()]
        if not text_runs:
            return False
        for r in text_runs:
            font = run_eastasia_font(r) or ""
            if not ("宋体" in font or font.lower() == "simsun"):
                return False
            size = run_size_pt(r)
            if size is None or abs(size - 12) > 0.5:  # 小四号 = 12 磅
                return False
            if not _resolve_run_bold(r, p):  # 最终生效必须为加粗
                return False
        return True

    l3_font_ok = bool(l3_paras) and all(_l3_heading_ok(p) for p in l3_paras)
    points.append(Point(3,
        "教学过程板块三级标题 (如 1、2、3……) 字体 宋体/小四号/加粗",
        l3_font_ok))

    # ---- 准备 body_paras: 除标题之外的所有文本段落 ----
    # 标题范围 (细则前面已单独约束, 此处排除):
    #   1) 学案标题 DOC_TITLE
    #   2) 12 个大板块标题 SECTION_HEADERS
    #   3) 教学过程一级标题 l1_paras (一、二、……)
    #   4) 教学过程二级标题 l2_paras ((一)(二)……)
    #   5) 教学过程三级标题 l3_paras (1. 2. ……)
    # 注: 不再用"全段黑体"启发式判定标题, 该启发式会误伤正文里用黑体强调的
    #     词句, 且细则未把此类段落归为标题
    all_paras = list(doc.paragraphs)
    heading_elements = {pp._element for pp in l1_paras + l2_paras + l3_paras}
    body_paras = []
    for p in all_paras:
        t = p.text.strip()
        if not t:
            continue
        if t == DOC_TITLE or t in SECTION_HEADERS:
            continue
        if p._element in heading_elements:
            continue
        body_paras.append(p)

    # ---- +5 除标题之外的所有文本: 字体为 宋体 / 小四号 ----
    # 说明:
    #   - 字体: rFonts@eastAsia 命中 "宋体" 或 "SimSun" (办公软件里显示都是宋体)
    #     若 run 未设置, 按办公软件解析顺序回溯: 段落 rPr → 段落样式链 → docDefaults
    #   - 字号: 小四号 = 12 磅 (Word/WPS "字号"下拉里的"小四"), ±0.5 磅覆盖半磅整数化
    #     同样支持样式/docDefaults 继承
    #   - 只针对细则两点 (字体、字号), 加粗/斜体/颜色/下划线一律不检查
    def _run_final_east_asia_font(run, para):
        """返回 run 在办公软件里最终生效的中文字体名 (含样式继承 & docDefaults)"""
        f = run_eastasia_font(run)
        if f:
            return f
        # 段落级 rPr
        pPr = para._element.find(qn("pPr"))
        if pPr is not None:
            rPr = pPr.find(qn("rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("rFonts"))
                if rFonts is not None:
                    v = rFonts.get(qn("eastAsia")) or rFonts.get(qn("ascii"))
                    if v: return v
        # 段落样式链 (含字符样式)
        for st in (getattr(run, "style", None), para.style):
            style = st
            while style is not None:
                rPr = style.element.find(qn("rPr"))
                if rPr is not None:
                    rFonts = rPr.find(qn("rFonts"))
                    if rFonts is not None:
                        v = rFonts.get(qn("eastAsia")) or rFonts.get(qn("ascii"))
                        if v: return v
                style = getattr(style, "base_style", None)
        # docDefaults
        styles_el = doc.styles.element
        docDefaults = styles_el.find(qn("docDefaults"))
        if docDefaults is not None:
            rPrDefault = docDefaults.find(qn("rPrDefault"))
            if rPrDefault is not None:
                rPr = rPrDefault.find(qn("rPr"))
                if rPr is not None:
                    rFonts = rPr.find(qn("rFonts"))
                    if rFonts is not None:
                        v = rFonts.get(qn("eastAsia")) or rFonts.get(qn("ascii"))
                        if v: return v
        return None

    def _run_final_size_pt(run, para):
        """返回 run 最终生效的字号 (磅), 含样式继承 & docDefaults"""
        s = run_size_pt(run)
        if s is not None:
            return s
        pPr = para._element.find(qn("pPr"))
        if pPr is not None:
            rPr = pPr.find(qn("rPr"))
            if rPr is not None:
                sz = rPr.find(qn("sz"))
                if sz is not None and sz.get(qn("val")) is not None:
                    return halfpt2pt(int(sz.get(qn("val"))))
        for st in (getattr(run, "style", None), para.style):
            style = st
            while style is not None:
                rPr = style.element.find(qn("rPr"))
                if rPr is not None:
                    sz = rPr.find(qn("sz"))
                    if sz is not None and sz.get(qn("val")) is not None:
                        return halfpt2pt(int(sz.get(qn("val"))))
                style = getattr(style, "base_style", None)
        styles_el = doc.styles.element
        docDefaults = styles_el.find(qn("docDefaults"))
        if docDefaults is not None:
            rPrDefault = docDefaults.find(qn("rPrDefault"))
            if rPrDefault is not None:
                rPr = rPrDefault.find(qn("rPr"))
                if rPr is not None:
                    sz = rPr.find(qn("sz"))
                    if sz is not None and sz.get(qn("val")) is not None:
                        return halfpt2pt(int(sz.get(qn("val"))))
        return None

    body_font_ok = bool(body_paras)
    for p in body_paras:
        if not body_font_ok:
            break
        for r in p.runs:
            if not r.text.strip():
                continue
            f = _run_final_east_asia_font(r, p) or ""
            if not ("宋体" in f or f.lower() == "simsun"):
                body_font_ok = False; break
            s = _run_final_size_pt(r, p)
            if s is None or abs(s - 12) > 0.5:  # 小四号 = 12 磅
                body_font_ok = False; break
    points.append(Point(5,
        "除标题之外的所有文本: 字体为 宋体 / 小四号",
        body_font_ok))

    # ---- +5 除标题之外的所有文本: 段落格式为
    #      1.5 倍行距 / 首行缩进 2 字符 / 对齐方式为两端对齐或左对齐 ----
    # 说明:
    #   body_paras 已在前面按细则严格排除 5 类标题, 此处针对其余段落判定
    #   三点必须同时满足, 缺一不可 (细则用顿号并列):
    #
    #   1) 1.5 倍行距: Word/WPS "段落" → "行距" 选"1.5 倍行距"
    #      XML: <w:spacing w:line="360" w:lineRule="auto"/> (240*1.5=360)
    #      也接受"多倍行距 1.5" 的等价写法; 需按办公软件继承规则回溯样式链
    #      不接受 "固定值/最小值 18 磅" 等其他表达 —— 细则明确写"1.5 倍"
    #
    #   2) 首行缩进 2 字符: Word/WPS "段落" → "特殊格式 首行缩进 2 字符"
    #      XML 优先 <w:ind w:firstLineChars="200"/>; 兼容 firstLine twips + 字号换算
    #      需含样式链继承 —— 段落未设置时按样式呈现出来的缩进也算
    #
    #   3) 对齐方式为两端对齐或左对齐:
    #      Word/WPS "段落" → "常规 对齐方式" 选 "两端对齐" 或 "左对齐"
    #      XML: <w:jc w:val="both"/> (两端对齐) 或 <w:jc w:val="left"/> (左对齐)
    #      未显式设置 <w:jc> 时, Word/WPS 中英文段落默认按"两端对齐"呈现, 视为符合
    def _resolve_line_spacing_final(p):
        """返回 (value, unit) 或 None; 含样式继承"""
        ls = para_line_spacing(p)
        if ls is not None:
            return ls
        style = p.style
        while style is not None:
            spPr = style.element.find(qn("pPr"))
            if spPr is not None:
                sp = spPr.find(qn("spacing"))
                if sp is not None and sp.get(qn("line")) is not None:
                    try:
                        line = int(sp.get(qn("line")))
                    except ValueError:
                        return None
                    rule = sp.get(qn("lineRule")) or "auto"
                    if rule == "auto":
                        return (line / 240.0, "multiple")
                    return (line / 20.0, "pt")
            style = getattr(style, "base_style", None)
        return None

    def _resolve_first_line_chars_final(p):
        """返回段落最终生效的首行缩进字符数, 未设置返回 None; 含样式继承"""
        def _read_from_ind(ind_el, size_pt):
            if ind_el is None:
                return None
            flc = ind_el.get(qn("firstLineChars"))
            if flc is not None:
                try:
                    v = int(flc)
                    if v > 0:
                        return v / 100.0
                except ValueError:
                    pass
            fl = ind_el.get(qn("firstLine"))
            if fl is not None:
                try:
                    v = int(fl)
                    if v > 0:
                        return v / (size_pt * 20.0)
                except ValueError:
                    pass
            return None

        # 取段落首个含文本 run 的最终字号 (用于 firstLine twips 换算), 无则用 12 磅
        size_pt = None
        for r in p.runs:
            if r.text.strip():
                size_pt = _run_final_size_pt(r, p)
                if size_pt is not None:
                    break
        if size_pt is None:
            size_pt = 12.0

        pPr = p._element.find(qn("pPr"))
        if pPr is not None:
            v = _read_from_ind(pPr.find(qn("ind")), size_pt)
            if v is not None:
                return v
        style = p.style
        while style is not None:
            spPr = style.element.find(qn("pPr"))
            if spPr is not None:
                v = _read_from_ind(spPr.find(qn("ind")), size_pt)
                if v is not None:
                    return v
            style = getattr(style, "base_style", None)
        return None

    def _resolve_align_final(p):
        """返回段落最终生效的对齐 (含样式继承); 未设置返回 None"""
        a = para_alignment(p)
        if a is not None:
            return a
        style = p.style
        while style is not None:
            spPr = style.element.find(qn("pPr"))
            if spPr is not None:
                jc = spPr.find(qn("jc"))
                if jc is not None:
                    val = jc.get(qn("val"))
                    return {"left": "LEFT", "right": "RIGHT", "center": "CENTER",
                            "both": "JUSTIFY", "distribute": "DISTRIBUTE"}.get(val, val)
            style = getattr(style, "base_style", None)
        return None

    body_fmt_ok = bool(body_paras)
    for p in body_paras:
        if not body_fmt_ok:
            break
        # 1) 1.5 倍行距: 只认 "multiple 1.5" (±0.05 覆盖舍入)
        ls = _resolve_line_spacing_final(p)
        ls_ok = ls is not None and ls[1] == "multiple" and abs(ls[0] - 1.5) <= 0.05
        # 2) 首行缩进 2 字符 (±0.1 字符覆盖办公软件舍入)
        ind_chars = _resolve_first_line_chars_final(p)
        ind_ok = ind_chars is not None and abs(ind_chars - 2.0) <= 0.1
        # 3) 两端对齐或左对齐 (未设置时办公软件默认按两端对齐呈现, 视为符合)
        align = _resolve_align_final(p)
        al_ok = align in ("JUSTIFY", "LEFT", None)
        if not (ls_ok and ind_ok and al_ok):
            body_fmt_ok = False
    points.append(Point(5,
        "除标题之外的所有文本: 段落 1.5 倍行距 / 首行缩进 2 字符 / 两端对齐或左对齐",
        body_fmt_ok))

    # ---- -3 "教学目标"板块下的目标编号类型非 "1." / "2." / "3." ----
    # 说明:
    #   办公软件 (Word/WPS) 中编号有两种落地方式, 都需覆盖:
    #     A) 自动编号 (段落应用了编号列表): XML 里段落有 <w:numPr>,
    #        文字前缀不写入 p.text, 但办公软件里显示为 "1." "2." "3."
    #        需要读取 numbering.xml 里对应级别的 numFmt 与 lvlText,
    #        只有 numFmt=decimal 且 lvlText="%1." (半角点号)
    #        才对应"1./2./3." 类型
    #     B) 手工输入: 段落文本以数字加分隔符开头, 需匹配 "1." 半角点号
    #        (不接受全角"．"、顿号"、"、右括号")"、下划线等其它分隔)
    #   任意目标项使用了 A) 或 B) 中不符合"1." 的编号形式 -> 触发扣分
    #   (复用 check_dim2 顶部定义的 _numbering_lvl_format)
    target_paras = [p for p in blocks.get("教学目标", [])[1:] if p.text.strip()]
    goal_num_bad = False
    for p in target_paras:
        # A) 自动编号
        fmt = _numbering_lvl_format(p)
        if fmt is not None:
            numFmt, lvlText = fmt
            # 只有 decimal + "%1." 才是"1." 类型 (Word/WPS 编号库里"1. 2. 3." 项)
            if not (numFmt == "decimal" and lvlText == "%1."):
                goal_num_bad = True
                break
            continue
        # B) 手工编号: 段落以数字开头才判定
        t = p.text.strip()
        if re.match(r"^\d", t) and not re.match(r"^\d+\.(?!\d)", t):
            # 以数字开头但不是"数字+半角点"格式 (排除小数如 "3.14")
            goal_num_bad = True
            break
    points.append(Point(-3,
        "“教学目标”板块下的目标编号类型非 “1.”/“2.”/“3.”",
        goal_num_bad))

    # ---- -5 文中出现半角方括号 ----
    # 说明:
    #   "半角方括号" = ASCII 的 "[" (U+005B) 与 "]" (U+005D)
    #   办公软件 (Word/WPS) 里能承载文字的位置远不止 doc.paragraphs:
    #     - 主体段落 (含普通段落、列表段落)
    #     - 表格单元格内段落 (可嵌套表格)
    #     - 文本框 / 形状内文字 (<w:txbxContent>)
    #     - 页眉 / 页脚段落
    #     - 脚注 / 尾注 (footnotes.xml / endnotes.xml)
    #   原代码只查 doc.paragraphs, 会漏掉表格/文本框/页眉页脚里的半角方括号,
    #   与"文中出现半角方括号"的语义 (只要文档里出现即扣分) 不符
    #
    #   直接扫描所有 <w:t> 文本节点; 只认 ASCII "[" 与 "]",
    #   全角 "【】" "［］" 不触发扣分 (细则明确"半角")
    def _iter_all_text_elements():
        # 主文档 + 页眉/页脚/脚注/尾注 相关 part
        parts = [doc.part]
        for sec in doc.sections:
            for hf in (sec.header, sec.first_page_header, sec.even_page_header,
                       sec.footer, sec.first_page_footer, sec.even_page_footer):
                try:
                    if hf is not None and hf.part is not None:
                        parts.append(hf.part)
                except Exception:
                    pass
        # 脚注 / 尾注 part (若存在)
        for rel in list(doc.part.rels.values()):
            try:
                if rel.reltype.endswith("/footnotes") or rel.reltype.endswith("/endnotes"):
                    parts.append(rel.target_part)
            except Exception:
                pass
        seen = set()
        for pt in parts:
            if id(pt) in seen: continue
            seen.add(id(pt))
            root = getattr(pt, "element", None)
            if root is None: continue
            # 所有 <w:t> 文本节点 (含表格、文本框、任意嵌套层级)
            for t in root.iter(qn("t")):
                yield t.text or ""

    bracket_hit = any(("[" in s) or ("]" in s) for s in _iter_all_text_elements())
    points.append(Point(-5, "文中出现半角方括号", bracket_hit))

    # ---- -3 12 个大板块标题 (教材分析/学情分析/设计理念/教学目标/教学重难点/
    #      教学准备/课时安排/教学思路说明/教学过程/板书设计/作业设计/教学反思)
    #      不满足单独成行, 或标题一行出现非标题文本 ----
    # 说明:
    #   办公软件 (Word/WPS) 中 "单独成行" 的落地含义:
    #     段落 (以 ¶ 结束) 是排版意义上的一行 (可视换行不改变段落归属),
    #     故 "标题单独成行" = 存在一个段落, 去空白后文本恰等于板块名, 且
    #     没有任何段落以板块名开头后又跟随其它非空白文字
    #   触发扣分的两种情况 (细则用"或"并列):
    #     A) 某板块标题从未以"整段等于板块名"的形式出现  → 未单独成行
    #     B) 某段落以板块名开头, 后续还有别的字符 (如 "教材分析: xxx",
    #        "教材分析 具体内容")  → 标题一行出现非标题文本
    #   扫描范围含正文段落 + 表格单元格 (含嵌套表格) —— 办公软件里表格
    #   单元格里的段落同样是"行", 需要一并检查
    def _iter_body_paragraphs():
        for p in doc.paragraphs:
            yield p
        def _walk(tables):
            for tbl in tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for pp in cell.paragraphs:
                            yield pp
                        for pp in _walk(cell.tables):
                            yield pp
        for pp in _walk(doc.tables):
            yield pp

    all_body_paras = list(_iter_body_paragraphs())
    sec_alone_bad = False
    for h in SECTION_HEADERS:
        alone_seen = False
        mixed_seen = False
        for p in all_body_paras:
            t = p.text.strip()
            if not t:
                continue
            if t == h:
                alone_seen = True
                continue
            # 以板块名开头, 后跟其他字符 → 标题一行混入非标题文本
            # 用 startswith(h) 后判长度, 避免误伤 "教学目标是..." 这类子串场景
            # 只匹配起始位置 (细则语义: 该行"是"板块标题却混了其他内容)
            if t.startswith(h) and len(t) > len(h):
                # 排除"教学目标"是"教学目标要求"这类同前缀但语义不同的其它板块名
                # (SECTION_HEADERS 内无相互前缀重叠, 这里的判定安全)
                mixed_seen = True
                break
        if mixed_seen or not alone_seen:
            sec_alone_bad = True
            break
    points.append(Point(-3,
        "12 个大板块标题 (教材分析/学情分析/设计理念/教学目标/教学重难点/"
        "教学准备/课时安排/教学思路说明/教学过程/板书设计/作业设计/教学反思) "
        "不满足单独成行, 或标题一行出现非标题文本",
        sec_alone_bad))

    return points

# =========================================================================
# 批量评估入口 — 与《脚本接口差异与统一建议》§2.1/§2.2 对齐
#   - 对外只暴露 evaluate(dir_path: str) -> dict
#   - dir_path: 脚本所在目录, 由脚本自行在其中定位并打开 .docx
#   - 返回结构中 dim2_items[*].detail 统一为空字符串
# =========================================================================
def evaluate(dir_path: str) -> dict:
    """入口: dir_path 为脚本所在目录, 自行在该目录内定位并打开 .docx"""
    import os

    file_name = ""
    # 维度二加分项满分 (与 check_dim2 中 append 顺序对应的正分值枚举):
    #   +3, +1, +1, +5, +3, +3, +3, +3, +3, +3, +5, +5, 其余为扣分项
    POSITIVE_SCORES = [3, 1, 1, 5, 3, 3, 3, 3, 3, 3, 5, 5]
    positive_max = sum(POSITIVE_SCORES)

    try:
        if not os.path.isdir(dir_path):
            raise FileNotFoundError(f"目录不存在: {dir_path}")
        docx_files = [
            f for f in os.listdir(dir_path)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        ]
        if not docx_files:
            raise FileNotFoundError(f"目录内未找到 .docx 文件: {dir_path}")
        docx_files.sort()
        file_name = docx_files[0]
        doc_path = os.path.join(dir_path, file_name)

        # 维度一
        dim1_ok, doc, dim1_msg = check_dim1(doc_path)
        if not dim1_ok or doc is None:
            return {
                "id": SCRIPT_ID,
                "file_name": file_name,
                "status": "ok",
                "error": None,
                "dim1_pass": False,
                "dim1_reason": dim1_msg,
                "dim2_items": [],
                "total_score": 0,
                "max_score": positive_max,
            }

        # 维度二
        points = check_dim2(doc)
        dim2_items = []
        total_score = 0
        for p in points:
            if p.score >= 0:
                # 加分项: 命中即得 score
                delta = p.score if p.hit else 0
                max_delta = p.score
                hit = bool(p.hit)
            else:
                # 扣分项 (score < 0): 命中违规条件才扣分，未命中为 0。
                max_delta = p.score
                delta = max_delta if p.hit else 0
                hit = bool(p.hit)
            dim2_items.append({
                "rule": p.desc,
                "max_delta": max_delta,
                "delta": delta,
                "hit": hit,
                "detail": "",
            })
            total_score += delta

        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": True,
            "dim1_reason": "",
            "dim2_items": dim2_items,
            "total_score": total_score,
            "max_score": positive_max,
        }
    except Exception as exc:
        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "error",
            "error": f"{type(exc).__name__}: {exc}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": positive_max,
        }


if __name__ == "__main__":
    # 本地调试: 默认取脚本所在目录; 也可通过命令行参数传入其它目录
    import json
    import os
    import sys

    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    print(json.dumps(evaluate(_dir), ensure_ascii=False, indent=2))
