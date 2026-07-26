#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# pyright: reportUnknownMemberType=false, reportUnknownVariableType=false, reportUnknownArgumentType=false, reportUnknownParameterType=false, reportMissingParameterType=false, reportPrivateUsage=false, reportUnknownLambdaType=false, reportGeneralTypeIssues=false

import os
import re
import sys
import json
import math
import shutil
import tempfile
import traceback
import zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

try:
    import pdf_backend
except ImportError:
    from verifiers import pdf_backend

# 说明：按接口统一约定，本模块不修改全局 sys.stdout，也不 print 主结果。
# 主结果统一由 evaluate(dir_path) -> dict 返回；入参为"脚本所在目录的路径"，
# 由脚本自己在该目录内定位并打开被评估的 .docx。本文件仅在 __main__ 下作为
# 本地调试入口，才把 dict 序列化为 JSON 打印。

SCRIPT_ID = "008"
# 维度二所有评分项 max_delta 之和；用于维度一未通过 / 脚本异常时兜底
MAX_SCORE_TOTAL = 117

# 维度二所有评分项按"结果追加顺序"的 max_delta 列表；末端把 results 里的
# "+N：..." 字符串反解为结构化 dim2_items 时用。顺序须与 _evaluate_impl 中
# results.append(...) 的顺序严格一致；共 41 项，累加 = MAX_SCORE_TOTAL。
_MAX_DELTAS_ORDER = [
    3,  # 01 首行缩进2字符
    3,  # 02 行距27-30磅
    5,  # 03 数字/字母 Times New Roman
    5,  # 04 页码：底部居中/宋体小五/阿拉伯
    5,  # 05 页码距下边线1.40cm
    3,  # 06 页眉文本内容
    5,  # 07 页眉字体（方正姚体）
    5,  # 08 页眉位置：横线上方居中/距上边线1.5cm
    5,  # 09 页眉双实线/浅灰/2.25磅
    5,  # 10 页面设置 A4
    3,  # 11 第1页标题：黑体小二居中
    3,  # 12 英文著录信息 TNR 小四 居中
    1,  # 13 著录信息与"摘要"之间空一行
    3,  # 14 摘要标题：黑体三号居中钢蓝
    3,  # 15 摘要正文：宋体小四两端对齐
    3,  # 16 关键词：黑体五号 全角分号
    5,  # 17 一级标题
    3,  # 18 标题内序号与文字间距
    3,  # 19 正文宋体小四两端对齐
    3,  # 20 脚注紧贴/上标
    3,  # 21 脚注编号到40 TNR 小四
    3,  # 22 表格编号连续/续表头
    3,  # 23 表格标题
    3,  # 24 表格中文本
    3,  # 25 表格三线格
    3,  # 26 图注宋体五号加粗
    3,  # 27 图片与图注不分页
    3,  # 28 图片嵌入型
    1, 1, 1, 1, 1, 1, 1, 1, 1, 1,  # 29-38 十个章节图片数量
    3,  # 39 参考文献标题
    3,  # 40 参考文献编号格式
    3,  # 41 参考文献条目格式
]
assert sum(_MAX_DELTAS_ORDER) == MAX_SCORE_TOTAL


def _resolve_effective_eastasia(doc, run):
    """按 Word 的字体继承链，解析 run 中文字符实际生效的东亚字体（打开办公软件看到的）：
       run 直接设定(w:eastAsia) → 段落样式(含 basedOn 链) → docDefaults。
       任一层拿到即返回；都没有则返回 None。"""
    # 1) run 级
    rPr = run._element.rPr
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            ea = rFonts.get(qn('w:eastAsia'))
            if ea:
                return ea
    # 2) 段落样式链（w:pStyle → styles.xml，沿 basedOn 上溯）
    styles_el = doc.styles.element
    para = run._parent
    style_id = None
    try:
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_id = pStyle.get(qn('w:val'))
    except Exception:
        style_id = None

    def _style_by_id(sid):
        for st in styles_el.findall(qn('w:style')):
            if st.get(qn('w:styleId')) == sid:
                return st
        return None

    seen = set()
    while style_id and style_id not in seen:
        seen.add(style_id)
        st = _style_by_id(style_id)
        if st is None:
            break
        rpr = st.find(qn('w:rPr'))
        if rpr is not None:
            rf = rpr.find(qn('w:rFonts'))
            if rf is not None:
                ea = rf.get(qn('w:eastAsia'))
                if ea:
                    return ea
        base = st.find(qn('w:basedOn'))
        style_id = base.get(qn('w:val')) if base is not None else None

    # 3) docDefaults
    dd = styles_el.find(qn('w:docDefaults'))
    if dd is not None:
        rpd = dd.find(qn('w:rPrDefault'))
        if rpd is not None:
            rpr = rpd.find(qn('w:rPr'))
            if rpr is not None:
                rf = rpr.find(qn('w:rFonts'))
                if rf is not None:
                    ea = rf.get(qn('w:eastAsia'))
                    if ea:
                        return ea
    return None


def _resolve_effective_ascii(doc, run):
    """按 Word 的字体继承链，解析 run 对西文/数字实际生效的字体（打开办公软件看到的）：
       run 直接设定(w:ascii → 回落 w:hAnsi) → 段落样式(含 basedOn 链) → docDefaults。
       任一层拿到即返回；都没有则返回 None。"""
    def _ascii_of(rFonts):
        if rFonts is None:
            return None
        return rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))

    # 1) run 级
    rPr = run._element.rPr
    if rPr is not None:
        v = _ascii_of(rPr.find(qn('w:rFonts')))
        if v:
            return v
    # 2) 段落样式链
    styles_el = doc.styles.element
    para = run._parent
    style_id = None
    try:
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_id = pStyle.get(qn('w:val'))
    except Exception:
        style_id = None

    def _style_by_id(sid):
        for st in styles_el.findall(qn('w:style')):
            if st.get(qn('w:styleId')) == sid:
                return st
        return None

    seen = set()
    while style_id and style_id not in seen:
        seen.add(style_id)
        st = _style_by_id(style_id)
        if st is None:
            break
        rpr = st.find(qn('w:rPr'))
        if rpr is not None:
            v = _ascii_of(rpr.find(qn('w:rFonts')))
            if v:
                return v
        base = st.find(qn('w:basedOn'))
        style_id = base.get(qn('w:val')) if base is not None else None

    # 3) docDefaults
    dd = styles_el.find(qn('w:docDefaults'))
    if dd is not None:
        rpd = dd.find(qn('w:rPrDefault'))
        if rpd is not None:
            rpr = rpd.find(qn('w:rPr'))
            if rpr is not None:
                v = _ascii_of(rpr.find(qn('w:rFonts')))
                if v:
                    return v
    return None


def _open_word_document(file_path):
    """打开 .docx 文档（python-docx 仅支持 OOXML/.docx）."""
    return Document(file_path)


def _render_to_pdf(file_path, tmp_dir):
    """把 .docx 文档渲染为 PDF，用于获取真实分页结果。

    仅通过本机 Microsoft Word (COM 自动化) 的 ExportAsFixedFormat 渲染；
    未安装 pywin32 或调用失败则抛出异常。
    """
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    out_pdf = os.path.join(tmp_dir, base_name + '.pdf')
    errors = []

    try:
        import win32com.client  # pywin32
    except ImportError as exc:
        errors.append('未安装 pywin32, 无法调用 Word 渲染 PDF: ' + str(exc))
    else:
        word = None
        doc = None
        try:
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=True)
            doc.ExportAsFixedFormat(
                OutputFileName=os.path.abspath(out_pdf),
                ExportFormat=17,  # wdExportFormatPDF
            )
            if os.path.isfile(out_pdf):
                return out_pdf
            errors.append('Word PDF 渲染未生成目标文件')
        except Exception as exc:
            errors.append('Word PDF 渲染失败: ' + str(exc))
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass

    raise RuntimeError('; '.join(errors) or '无法渲染 PDF')


def _repair_fragment_hyperlink_relationships(file_path):
    """在临时 DOCX 中修正被错误声明为包内文件的书签超链接关系。"""
    rel_ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
    ET.register_namespace('', rel_ns)
    fd, temp_path = tempfile.mkstemp(prefix='officeval_008_', suffix='.docx')
    os.close(fd)
    changed = False
    try:
        with zipfile.ZipFile(file_path, 'r') as source, zipfile.ZipFile(
            temp_path, 'w'
        ) as target:
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename.endswith('.rels'):
                    root = ET.fromstring(data)
                    rels_changed = False
                    for rel in root.findall(f'{{{rel_ns}}}Relationship'):
                        if (
                            rel.get('Type', '').endswith('/hyperlink')
                            and rel.get('Target', '').startswith('#')
                            and rel.get('TargetMode') != 'External'
                        ):
                            rel.set('TargetMode', 'External')
                            rels_changed = True
                            changed = True
                    if rels_changed:
                        data = ET.tostring(
                            root, encoding='utf-8', xml_declaration=True
                        )
                target.writestr(info, data)
        if changed:
            return temp_path
        os.remove(temp_path)
        return None
    except Exception:
        try:
            os.remove(temp_path)
        except OSError:
            pass
        raise


def evaluate(dir_path):
    """按 §2.2 约定，返回结构化 dict；不 print、不改 sys.stdout、不 sys.exit。
    入参 dir_path 为"脚本所在目录的路径"，脚本自己在该目录里定位并打开被评估的
    .docx（忽略以 ~$ 开头的 Word 临时锁文件）。脚本自身崩溃（含目录不存在、
    未发现文档等）由本函数捕获并返回 status='error'。"""
    result = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": True,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": MAX_SCORE_TOTAL,
    }
    try:
        # 在 dir_path 目录内定位被评估的 .docx（忽略 ~$ 打头的 Word 临时锁文件）
        if not dir_path or not os.path.isdir(dir_path):
            raise FileNotFoundError(f"目录不存在：{dir_path}")
        cands = sorted(
            f for f in os.listdir(dir_path)
            if f.lower().endswith('.docx') and not f.startswith('~$')
        )
        if not cands:
            raise FileNotFoundError(f"目录中未发现 .docx：{dir_path}")
        file_path = os.path.join(dir_path, cands[0])
        result["file_name"] = os.path.basename(file_path)
        try:
            return _evaluate_impl(file_path, result)
        except KeyError as exc:
            if 'word/#' not in str(exc):
                raise
            repaired_path = _repair_fragment_hyperlink_relationships(file_path)
            if repaired_path is None:
                raise
            try:
                return _evaluate_impl(repaired_path, result)
            finally:
                try:
                    os.remove(repaired_path)
                except OSError:
                    pass
    except Exception as exc:
        result["status"] = "error"
        result["error"] = f"{type(exc).__name__}: {exc}"
        result["dim2_items"] = []
        result["total_score"] = 0
        return result


def _evaluate_impl(file_path, result):
    """维度一 + 维度二的评估主体；对内部构造 results 列表沿用原逻辑，
    末尾按 _MAX_DELTAS_ORDER 反解成结构化 dim2_items 填入 result。"""
    doc = _open_word_document(file_path)
    results = []

    # 全局常量：严格"宋体"判定（仅认这几种东亚字体名）
    STRICT_SONGTI = ('宋体', 'SimSun', 'NSimSun')

    # ========== 维度1：可用与可修改性 ==========
    # 当前脚本不再把连续空白页、乱码/文字重叠面积或全文页数范围作为维度一门禁。

    # ========== 维度2：完成度评分 ==========

    # --- 辅助函数 ---
    def is_title(para):
        text = para.text.strip()
        if re.match(r'^\d+\.?\d*\.?\d*\s', text):
            return True
        if text in ["摘要", "关键词", "参考文献", "致谢"]:
            return True
        return False

    def is_caption(para):
        text = para.text.strip()
        return bool(re.match(r'^(图|表)\s*\d+', text))

    def is_in_table(para):
        parent = para._element.getparent()
        while parent is not None:
            if parent.tag.endswith('tbl'):
                return True
            parent = parent.getparent()
        return False

    # +3：除表格、参考文献外，文章中出现超过2行的文本内容整体缩进均为：首行缩进2字符
    # 严格对齐细则，仅踩细则明确写出的点，不额外约束：
    #   ① 范围排除项：只排除"表格内段落"和"参考文献区"（细则原文仅此两项）
    #   ② "出现超过2行"：按 Word 实际版式估算该段占的行数 > 2 才纳入判定
    #        行数 = ceil( 段落字符视觉宽度 / 每行可容纳字符宽度 )
    #        每行可容纳宽度 = (页面宽 - 左边距 - 右边距) / 单字符宽(≈字号磅值)，
    #        并从第一行扣掉"首行缩进"占位；此换算与 Word 排版口径一致。
    #   ③ 首行缩进2字符：Word 的两种落地方式，命中其一即算合规：
    #        (a) w:ind/@w:firstLineChars == 200（Word"2 字符"的原生存储，与字号无关）
    #        (b) w:ind/@w:firstLine 绝对磅值 ÷ 段落字号 ≈ 2 字符
    #   ④ "整体均为"：所有纳入判定的段落必须全部合规。

    # 定位参考文献区（"参考文献"标题及其之后的所有段落都排除）
    _ref_start_for_indent = len(doc.paragraphs)
    for _i, _p in enumerate(doc.paragraphs):
        if _p.text.strip() == '参考文献':
            _ref_start_for_indent = _i
            break

    def _para_main_size(para):
        """段落主体字号（最长 run 的字号），无则回落 12pt（小四）。1 字符宽≈字号磅值。"""
        main = None
        maxlen = 0
        for r in para.runs:
            if r.text and len(r.text) > maxlen and r.font.size:
                maxlen = len(r.text)
                main = r.font.size.pt
        return main if main else 12.0

    def _visual_width_units(text):
        """段落文本的视觉宽度（以"字符宽"为单位）：中文/全角=1，半角字母数字/空格≈0.5。"""
        w = 0.0
        for ch in text:
            if ch in ('\t',):
                w += 2.0
            elif '一' <= ch <= '鿿' or '＀' <= ch <= '￯' or ch == '　':
                w += 1.0   # 中日韩全角
            else:
                w += 0.5   # 半角
        return w

    def _para_line_count(para):
        """按 Word 版式估算该段占的行数。"""
        text = para.text
        if not text.strip():
            return 0
        sec = doc.sections[0]
        size = _para_main_size(para)                      # pt，一个全角字符宽≈size
        # 页面正文可用宽度（pt）：页宽 - 左右边距（EMU；1pt = 12700 EMU）
        pw = sec.page_width
        lm = sec.left_margin
        rm = sec.right_margin
        if pw is None or lm is None or rm is None or size <= 0:
            return 1
        usable_pt = (int(pw) - int(lm) - int(rm)) / 12700.0
        if usable_pt <= 0:
            return 1
        chars_per_line = usable_pt / size                 # 每行可容纳"字符宽"数
        if chars_per_line <= 0:
            return 1
        # 首行被"首行缩进"占位（换算成字符宽）
        pPr = para._element.find(qn('w:pPr'))
        first_indent_units = 0.0
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                fl = ind.get(qn('w:firstLine'))           # 单位 twips(1/20 pt)
                if flc is not None and flc.isdigit():
                    first_indent_units = int(flc) / 100.0  # 100 = 1 字符
                elif fl is not None and fl.isdigit():
                    first_indent_units = (int(fl) / 20.0) / size
        total_units = _visual_width_units(text)
        # 首行有效容量减去缩进；其余行按满行算
        first_cap = max(chars_per_line - first_indent_units, 1.0)
        if total_units <= first_cap:
            return 1
        remaining = total_units - first_cap
        return 1 + math.ceil(remaining / chars_per_line)

    def _first_line_indent_2chars(para):
        """首行缩进是否为 2 字符（Word 口径）。"""
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                # (a) firstLineChars==200 即 Word 的"2 字符"
                flc = ind.get(qn('w:firstLineChars'))
                if flc is not None and flc.isdigit():
                    return int(flc) == 200
                # (b) 绝对磅值折算：firstLine(twips) / 20 / 字号 ≈ 2 字符
                fl = ind.get(qn('w:firstLine'))
                if fl is not None and fl.isdigit():
                    pt = int(fl) / 20.0
                    chars = pt / _para_main_size(para)
                    return 1.8 <= chars <= 2.2
        return False

    indent_correct = 0
    indent_total = 0
    for i, para in enumerate(doc.paragraphs):
        if i >= _ref_start_for_indent:        # ① 排除参考文献区
            continue
        if is_in_table(para):                 # ① 排除表格
            continue
        if not para.text.strip():             # 空段落无"文本内容"，跳过
            continue
        if _para_line_count(para) <= 2:       # ② 仅"超过2行"的文本
            continue
        indent_total += 1
        if _first_line_indent_2chars(para):   # ③ 首行缩进2字符
            indent_correct += 1
    # ④ 整体均为
    if indent_total > 0 and indent_correct == indent_total:
        results.append(f"+3：除表格、参考文献外，文章中出现超过2行的文本内容整体缩进均为首行缩进2字符（{indent_correct}/{indent_total}）")
    else:
        results.append(f"+0：除表格、参考文献外，文章中出现超过2行的文本内容整体缩进均为首行缩进2字符（{indent_correct}/{indent_total}）")

    # +3：除表格、标题、图注、参考文献，文章所有正文文本行距均为27-30磅
    # 严格对齐细则，仅踩细则明确写出的点，不额外约束：
    #   ① 排除项：表格内段落、标题、图注、参考文献区（细则原文列出的四项）
    #   ② 判定对象：剩余的"正文文本"——所有非空的正文段落
    #   ③ 行距范围：27–30 磅（含端点）
    #
    # 与办公软件对齐的读取方式（Word 行距在 OOXML 里存 w:pPr/w:spacing）：
    #   - w:lineRule="exact" 或 "atLeast"：w:line 单位为 1/20 磅（twips）→ 磅值 = line/20
    #     （对应 Word"固定值/最小值"，直接就是磅数，这正是能落到"27-30磅"的设置）
    #   - w:lineRule="auto"（或缺省）：w:line 单位为 1/240 倍行距（倍数 = line/240），
    #     Word 的"单倍/1.5倍/多倍"行距，非固定磅值 → 不是"27-30磅"的行距，判不合规
    #   - 完全未设置行距：继承默认，非明确的 27–30 磅 → 不合规
    def _line_spacing_pt(para):
        """返回 (磅值 or None, 是否为固定/最小磅值行距)。
        仅 exact/atLeast 才是真正的"磅值行距"；auto 是倍数行距，返回 (None, False)。"""
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            return (None, False)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            return (None, False)
        line = sp.get(qn('w:line'))
        rule = sp.get(qn('w:lineRule'))
        if line is None or not line.lstrip('-').isdigit():
            return (None, False)
        if rule in ('exact', 'atLeast'):
            return (int(line) / 20.0, True)   # twips → pt
        # auto / 缺省 rule ⇒ 倍数行距，不是磅值行距
        return (None, False)

    # 参考文献区起点（"参考文献"标题及其之后的段落都排除）
    _ref_start_for_ls = len(doc.paragraphs)
    for _i, _p in enumerate(doc.paragraphs):
        if _p.text.strip() == '参考文献':
            _ref_start_for_ls = _i
            break

    line_spacing_all_ok = True
    checked_any = False
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():          # 无文本内容，不算"文本"
            continue
        if is_in_table(para):              # ① 排除表格
            continue
        if is_title(para):                 # ① 排除标题
            continue
        if is_caption(para):               # ① 排除图注
            continue
        if i >= _ref_start_for_ls:         # ① 排除参考文献
            continue
        checked_any = True                 # ② 剩余正文文本
        pt, is_point_spacing = _line_spacing_pt(para)
        # ③ 必须是磅值行距且落在 27–30 磅（含端点）
        if not is_point_spacing or pt is None or not (27.0 <= pt <= 30.0):
            line_spacing_all_ok = False
            break
    if checked_any and line_spacing_all_ok:
        results.append("+3：除表格、标题、图注、参考文献，文章所有正文文本行距均为27-30磅")
    else:
        results.append("+0：除表格、标题、图注、参考文献，文章所有正文文本行距均为27-30磅")

    # +5：除页码外，文章内出现的所有数字与字母均为新罗马字体
    # 严格对齐细则，仅踩细则明确写出的点：
    #   ① 判定对象：文章内出现的"数字与字母"（ASCII 的 0-9 / A-Z / a-z）
    #   ② 排除项：只排除"页码"——即页脚中承载 PAGE 域（自动页码）的段落
    #   ③ 字体要求：这些数字/字母字符的西文字体为"新罗马"(Times New Roman)
    #
    # 与办公软件对齐（Word 里西文字体存在 run 的 w:rPr/w:rFonts）：
    #   - ASCII 字符实际生效的字体：优先 w:ascii，回落 w:hAnsi（Word 的"西文字体"设置）
    #   - run 未显式设时，沿"段落样式→docDefaults"继承链解析实际生效字体（打开办公软件看到的）
    ASCII_RE = re.compile(r'[A-Za-z0-9]')

    def _iter_paragraphs_except_pagenumber():
        """遍历全文段落（正文 + 表格 + 页眉 + 页脚），但跳过页脚里承载页码的段落。"""
        for p in doc.paragraphs:
            yield p
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for section in doc.sections:
            for p in section.header.paragraphs:
                yield p
            for p in section.footer.paragraphs:
                # ② 排除页码：页脚中含 PAGE / NUMPAGES 域的段落
                xml = p._element.xml
                if 'PAGE' in xml or 'NUMPAGES' in xml:
                    continue
                yield p

    correct_font = True
    for para in _iter_paragraphs_except_pagenumber():
        for run in para.runs:
            # ① 只看含"数字或字母"的 run
            if not run.text or not ASCII_RE.search(run.text):
                continue
            # ③ 西文字体必须是 Times New Roman（新罗马）——按办公软件继承链解析实际生效字体
            font = _resolve_effective_ascii(doc, run)
            if font != 'Times New Roman':
                correct_font = False
                break
        if not correct_font:
            break
    if correct_font:
        results.append("+5：除页码外，文章内出现的所有数字与字母均为新罗马字体")
    else:
        results.append("+0：除页码外，文章内出现的所有数字与字母均为新罗马字体")

    # +5：文档所有页码：在页面底部居中显示；字体为宋体，小五；格式类型例如：1，2，3…
    # 严格对齐细则，仅踩细则明确写出的点（"距下边线"是另一评分项，这里不判）：
    #   ① 在页面底部：页码为页脚(footer)中的 PAGE 域（Word 的自动页码域）
    #   ② 居中：承载 PAGE 域的段落 w:jc == center
    #   ③ 字体为宋体：承载页码的 run 字体为宋体（数字走西文 ascii/hAnsi，回落 eastAsia）
    #   ④ 小五：字号 9pt（小五号 = 9pt；sz 以半磅存储，18 = 9pt）
    #   ⑤ 格式类型 1,2,3…：阿拉伯数字编号
    #        - 节级：w:sectPr/w:pgNumType/@w:fmt 为空或 "decimal"
    #        - 域级：PAGE 域指令若带 \* 格式开关，必须是 arabic 系（不能是 roman/字母等）
    #   ⑥ "文档所有页码"：所有含页码的页脚段落都必须合规，且至少存在一个页码。
    SONGTI_NAMES = {'宋体', 'SimSun', 'NSimSun'}
    ARABIC_FMTS = {None, 'decimal', 'arabic', 'arabicDash'}

    def _field_instr(para):
        """拼出段落内 PAGE 域的指令文本（用于识别 \\* 格式开关）。"""
        texts = []
        for it in para._element.iter(qn('w:instrText')):
            texts.append(it.text or '')
        return ''.join(texts)

    def _page_format_is_arabic(section, para):
        """判断该页码是否为阿拉伯数字 1,2,3… 格式。"""
        # 域级：PAGE \* roman / \* alphabetic 等非阿拉伯开关
        instr = _field_instr(para)
        m = re.search(r'\\\*\s*([A-Za-z]+)', instr)
        if m:
            sw = m.group(1).lower()
            if sw not in ('arabic', 'arabicdash', 'mergeformat', 'charformat'):
                # roman / alphabetic / ordinal / cardtext 等 ⇒ 非 1,2,3
                if sw in ('roman', 'alphabetic', 'ordinal', 'ordtext',
                          'cardtext', 'dollartext', 'hex'):
                    return False
        # 节级：sectPr/pgNumType/@w:fmt
        sectPr = section._sectPr
        if sectPr is not None:
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                fmt = pgNumType.get(qn('w:fmt'))
                if fmt not in ARABIC_FMTS:
                    return False
        return True

    def _run_font_is_songti(run):
        """页码为数字：按办公软件口径解析实际生效字体（西文优先、回落中文），命中宋体即可。"""
        v = _resolve_effective_ascii(doc, run)
        if v and any(n in v for n in SONGTI_NAMES):
            return True
        ea = _resolve_effective_eastasia(doc, run)
        return bool(ea and any(n in ea for n in SONGTI_NAMES))

    page_number_seen = False
    page_number_all_ok = True
    for section in doc.sections:
        for para in section.footer.paragraphs:
            # ① 底部页脚中的 PAGE 域
            if 'PAGE' not in para._element.xml:
                continue
            page_number_seen = True
            # ② 居中
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                page_number_all_ok = False
                break
            # ③ 宋体 + ④ 小五（承载 PAGE 域的 run 需满足）
            font_size_ok = False
            for run in para.runs:
                if not _run_font_is_songti(run):
                    continue
                if run.font.size and 8.5 <= run.font.size.pt <= 9.5:
                    font_size_ok = True
                    break
            if not font_size_ok:
                page_number_all_ok = False
                break
            # ⑤ 阿拉伯数字格式 1,2,3…
            if not _page_format_is_arabic(section, para):
                page_number_all_ok = False
                break
        if not page_number_all_ok:
            break
    if page_number_seen and page_number_all_ok:
        results.append("+5：文档所有页码：在页面底部居中显示；字体为宋体，小五；格式类型例如：1，2，3…")
    else:
        results.append("+0：文档所有页码：在页面底部居中显示；字体为宋体，小五；格式类型例如：1，2，3…")

    # +5：文档所有页码距离页面下边线1.40cm
    # 严格对齐细则，只踩细则唯一的点——"页码距页面下边线 = 1.40cm"，不额外约束居中/字体等。
    #
    # 与办公软件对齐：Word「页面设置 → 版式 → 距边界·页脚」的值，即页脚(页码)到页面
    # 底边的距离，存于 w:sectPr/w:pgMar/@w:footer（EMU），对应 python-docx 的
    # section.footer_distance。要求"文档所有节"的该值都等于 1.40cm。
    #   - 1.40cm 判定：四舍五入到 0.01cm 后 == 1.40（严格对齐"1.40"，不做宽区间放水）
    page_dist_ok = True
    page_dist_seen = False
    for section in doc.sections:
        fd = section.footer_distance
        if fd is None:
            page_dist_ok = False
            break
        page_dist_seen = True
        if round(fd.cm, 2) != 1.40:
            page_dist_ok = False
            break
    if page_dist_seen and page_dist_ok:
        results.append("+5：文档所有页码距离页面下边线1.40cm")
    else:
        results.append("+0：文档所有页码距离页面下边线1.40cm")

    # +3：整篇文档页眉均为"华语青月大学本科生毕业设计（论文）用纸"
    # 严格对齐细则，只踩细则的点——"整篇文档"的"页眉"内容"均为"该文本：
    #   ① 页眉内容 = 期望文本（strip 后严格相等，允许前后空白，不允许其他多余非空白字符）
    #   ② "整篇文档" = 所有节(section)。且每节若启用了"首页不同/奇偶页不同"，
    #      被启用且未链接到上一节的那几种页眉也必须都等于期望文本；
    #      未启用或链接继承的页眉沿用生效页眉，不重复判。
    #   ③ "均为" = 逐一校验，任一生效页眉不符即判 0，且至少要存在页眉。
    #
    # 与办公软件对齐：Word 的页眉分 default / first_page / even_page 三种，
    #   受"首页不同""奇偶页不同"开关控制；python-docx 对应 section.header /
    #   first_page_header / even_page_header，is_linked_to_previous 表示继承上一节。
    expected_header = "华语青月大学本科生毕业设计（论文）用纸"

    def _header_text_ok(hdr):
        """页眉整体文本 strip 后是否严格等于期望文本。"""
        txt = ''.join(p.text for p in hdr.paragraphs).strip()
        return txt == expected_header

    header_all_ok = True
    header_seen = False
    for section in doc.sections:
        # 需要校验的生效页眉集合
        headers_to_check = [section.header]
        if section.different_first_page_header_footer:
            headers_to_check.append(section.first_page_header)
        # 奇偶页不同由文档级 evenAndOddHeaders 控制；启用时偶数页页眉单独生效
        headers_to_check.append(section.even_page_header)

        for hdr in headers_to_check:
            # 链接到上一节 ⇒ 继承，不单独判（避免把"空的继承页眉"误判为不符）
            if hdr.is_linked_to_previous:
                continue
            header_seen = True
            if not _header_text_ok(hdr):
                header_all_ok = False
                break
        if not header_all_ok:
            break

    if header_seen and header_all_ok:
        results.append('+3：整篇文档页眉均为"华语青月大学本科生毕业设计（论文）用纸"')
    else:
        results.append('+0：整篇文档页眉均为"华语青月大学本科生毕业设计（论文）用纸"')

    # +5：文档中所有页眉：页眉文本为浅灰色，方正姚体，小二号，1倍行距
    # 逻辑（按确认口径）：对文档中"实际出现的页眉文本"判定字体/字号/颜色，对其所在
    # 段落判定行距——每个有文本的页眉 run，需满足：
    #   ① 中文字体(w:eastAsia，按继承链解析)必须为"方正姚体"
    #   ② 字号为小二（18pt = 36 半磅，容差 17–19pt）
    #   ③ 有效颜色为浅灰色（灰阶 R≈G≈B，亮度落在浅灰范围 0x80–0xE0）
    #   段落级：④ 1倍行距（w:pPr/w:spacing 未设置，或 lineRule=auto 且 line=240）
    #   - 前提：页眉文本必须是"出来的"（非空 run）；没有页眉文本的（空/继承）不限定、跳过
    #   - 全文没有任何出来的页眉文本 ⇒ 无判定对象 ⇒ 判 +0
    #
    # 与办公软件对齐：中文字体存 run 的 w:rPr/w:rFonts/@w:eastAsia；字号存 w:rPr/w:sz
    #   （半磅，继承链同字体）；颜色存 w:rPr/w:color/@w:val（固定 RGB）；行距存
    #   w:pPr/w:spacing/@w:line + @w:lineRule（auto 时 line/240 = 倍数，line=240 即 1倍）。
    def _run_font_is_yaoti(run):
        """页眉文本为中文，判中文字体（按办公软件继承链解析实际生效字体）是否为方正姚体。"""
        ea = _resolve_effective_eastasia(doc, run)
        return bool(ea and '方正姚体' in ea)

    def _run_effective_size_halfpt(run):
        """按继承链解析 run 实际生效字号（半磅）：run 级 → 段落样式(basedOn 链) → docDefaults。"""
        rPr = run._element.rPr
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                if val and val.lstrip('-').isdigit():
                    return int(val)
        styles_el = doc.styles.element
        para = run._parent
        style_id = None
        try:
            pPr = para._p.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_id = pStyle.get(qn('w:val'))
        except Exception:
            style_id = None

        def _style_by_id(sid):
            for st in styles_el.findall(qn('w:style')):
                if st.get(qn('w:styleId')) == sid:
                    return st
            return None

        seen = set()
        while style_id and style_id not in seen:
            seen.add(style_id)
            st = _style_by_id(style_id)
            if st is None:
                break
            rpr = st.find(qn('w:rPr'))
            if rpr is not None:
                sz = rpr.find(qn('w:sz'))
                if sz is not None:
                    val = sz.get(qn('w:val'))
                    if val and val.lstrip('-').isdigit():
                        return int(val)
            base = st.find(qn('w:basedOn'))
            style_id = base.get(qn('w:val')) if base is not None else None
        dd = styles_el.find(qn('w:docDefaults'))
        if dd is not None:
            rpd = dd.find(qn('w:rPrDefault'))
            if rpd is not None:
                rpr = rpd.find(qn('w:rPr'))
                if rpr is not None:
                    sz = rpr.find(qn('w:sz'))
                    if sz is not None:
                        val = sz.get(qn('w:val'))
                        if val and val.lstrip('-').isdigit():
                            return int(val)
        return None

    def _run_size_is_xiaoer(run):
        """小二号 = 18pt = 36 半磅，容差 17–19pt。"""
        half_pt = _run_effective_size_halfpt(run)
        if half_pt is None:
            return False
        return 17.0 <= half_pt / 2.0 <= 19.0

    def _run_color_is_lightgray(run):
        """有效颜色为浅灰色：固定 RGB 灰阶（R≈G≈B）且亮度落在 0x80–0xE0。
        未显式设色（继承默认黑色/自动色）视为不满足。"""
        rPr = run._element.rPr
        if rPr is None:
            return False
        color = rPr.find(qn('w:color'))
        if color is None:
            return False
        val = color.get(qn('w:val'))
        if not val or not re.fullmatch(r'[0-9A-Fa-f]{6}', val):
            return False
        r = int(val[0:2], 16); g = int(val[2:4], 16); b = int(val[4:6], 16)
        if not (abs(r - g) <= 16 and abs(g - b) <= 16 and abs(r - b) <= 16):
            return False
        brightness = (r + g + b) / 3.0
        return 0x80 <= brightness <= 0xE0

    def _para_is_single_line_spacing(para):
        """1倍行距：w:pPr/w:spacing 未设置（默认单倍），或 lineRule=auto 且 line=240。"""
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            return True
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            return True
        line = sp.get(qn('w:line'))
        rule = sp.get(qn('w:lineRule'))
        if line is None:
            return True
        if rule not in (None, 'auto'):
            return False
        return line.lstrip('-').isdigit() and int(line) == 240

    header_format_ok = True
    header_fmt_seen = False
    for section in doc.sections:
        headers_to_check = [section.header]
        if section.different_first_page_header_footer:
            headers_to_check.append(section.first_page_header)
        headers_to_check.append(section.even_page_header)

        for hdr in headers_to_check:
            if hdr.is_linked_to_previous:
                continue
            for para in hdr.paragraphs:
                para_has_text = any(r.text.strip() for r in para.runs)
                if para_has_text and not _para_is_single_line_spacing(para):
                    header_format_ok = False
                    break
                for run in para.runs:
                    if not run.text.strip():     # 前提：只判"出来的"页眉文本
                        continue
                    header_fmt_seen = True
                    # 字体（方正姚体）+ 字号（小二）+ 颜色（浅灰）
                    if not (_run_font_is_yaoti(run)
                            and _run_size_is_xiaoer(run)
                            and _run_color_is_lightgray(run)):
                        header_format_ok = False
                        break
                if not header_format_ok:
                    break
            if not header_format_ok:
                break
        if not header_format_ok:
            break

    if header_fmt_seen and header_format_ok:
        results.append("+5：文档中所有页眉：页眉文本为浅灰色，方正姚体，小二号，1倍行距")
    else:
        results.append("+0：文档中所有页眉：页眉文本为浅灰色，方正姚体，小二号，1倍行距")

    # +5：整篇文档页眉文本的位置：位于粗横线上方居中显示；页眉距离页面上边线1.5cm
    # 严格对齐细则，仅踩细则明确写出的三个点，对"整篇文档"所有出来的页眉文本判定：
    #   ① 位于粗横线上方：承载页眉文本的段落 w:pPr/w:pBdr 含 w:bottom（下边框）
    #      —— 边框在段落"下方"，故文字在横线"上方"。（"粗/横线样式"是另一评分项，此处只判"上方"关系）
    #   ② 居中显示：该段落 w:jc == center
    #   ③ 页眉距页面上边线 1.5cm：section.header_distance == 1.5cm（四舍五入到 0.01cm）
    #   前提：只对"出来的页眉文本"（非空 run 所在段落）判；空/继承的页眉不限定、跳过。
    #   "整篇文档" ⇒ 所有节的生效页眉里、每个有文本的段落都必须满足①②，且其所属节 ③ 成立。
    #   全文没有任何出来的页眉文本 ⇒ 无判定对象 ⇒ 判 +0。
    #
    # 与办公软件对齐：页眉距上边线存 w:sectPr/w:pgMar/@w:header（对应 section.header_distance，
    #   Word「页面设置→版式→距边界·页眉」）；居中存段落 w:jc；页眉横线是段落下边框 w:pBdr/w:bottom
    #   （Word 里"页眉"段落加的"下框线"）。均为 Word 可设可读的属性。
    header_pos_ok = True
    header_pos_seen = False
    for section in doc.sections:
        # ③ 页眉距上边线 1.5cm（节级）
        hd = section.header_distance
        dist_ok = (hd is not None and round(hd.cm, 2) == 1.50)

        headers_to_check = [section.header]
        if section.different_first_page_header_footer:
            headers_to_check.append(section.first_page_header)
        headers_to_check.append(section.even_page_header)

        for hdr in headers_to_check:
            if hdr.is_linked_to_previous:
                continue
            for para in hdr.paragraphs:
                # 前提：只判"出来的"页眉文本段落
                if not any(r.text.strip() for r in para.runs):
                    continue
                header_pos_seen = True
                # ③ 该段所属节的页眉距上边线须为 1.5cm
                if not dist_ok:
                    header_pos_ok = False
                    break
                # ② 居中
                if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    header_pos_ok = False
                    break
                # ① 位于横线上方：段落下边框存在
                pPr = para._element.find(qn('w:pPr'))
                pBdr = pPr.find(qn('w:pBdr')) if pPr is not None else None
                if pBdr is None or pBdr.find(qn('w:bottom')) is None:
                    header_pos_ok = False
                    break
            if not header_pos_ok:
                break
        if not header_pos_ok:
            break

    if header_pos_seen and header_pos_ok:
        results.append("+5：整篇文档页眉文本的位置：位于粗横线上方居中显示；页眉距离页面上边线1.5cm")
    else:
        results.append("+0：整篇文档页眉文本的位置：位于粗横线上方居中显示；页眉距离页面上边线1.5cm")

    # +5：文档中所有页眉处出现直线：类型为双实线，颜色为浅灰色，2.25磅
    # 严格对齐细则，仅踩细则明确写出的点，对"文档中所有出来的页眉"判定其直线：
    #   ① 出现直线：页眉段落 w:pPr/w:pBdr 存在边框（w:bottom 等）
    #   ② 类型为双实线：w:val 属于双实线家族——标准双线 "double"，及 Word 的
    #      "细粗/粗细双线" thinThick* / thickThin* 系列
    #   ③ 颜色为浅灰色：w:color 为灰阶（R≈G≈B）且亮度落在浅灰范围 0x80–0xE0
    #   ④ 2.25磅：w:sz 以 1/8 磅为单位，2.25磅 = 18；严格判 == 18
    #   前提：只对"出来的页眉"（含非空文本的页眉）判其直线；空/继承页眉不限定、跳过。
    #   "所有页眉""出现直线" ⇒ 每个有文本的页眉都必须存在一条满足②③④的直线，任一缺失/不符判 0。
    #
    # 与办公软件对齐：页眉横线是段落"下框线"，存 w:pPr/w:pBdr/w:bottom；
    #   线型 @w:val、颜色 @w:color、磅数 @w:sz（1/8 磅）——均为 Word「边框和底纹」里可设可读的属性。
    has_header_line = True
    header_line_seen = False
    DOUBLE_LINE_VALS = {
        'double',
        'thinThickSmallGap', 'thinThickMediumGap', 'thinThickLargeGap',
        'thickThinSmallGap', 'thickThinMediumGap', 'thickThinLargeGap',
    }

    def _edge_is_double_lightgray_225(edge):
        val = edge.get(qn('w:val'))
        color = edge.get(qn('w:color'))
        sz = edge.get(qn('w:sz'))
        # ② 双实线
        if val not in DOUBLE_LINE_VALS:
            return False
        # ③ 浅灰色
        if not (color and re.fullmatch(r'[0-9A-Fa-f]{6}', color)):
            return False
        r = int(color[0:2], 16); g = int(color[2:4], 16); b = int(color[4:6], 16)
        if not (abs(r - g) <= 16 and abs(g - b) <= 16 and abs(r - b) <= 16):
            return False
        if not (0x80 <= r <= 0xE0):
            return False
        # ④ 2.25磅（sz = 18，1/8 磅）
        if not (sz and sz.isdigit() and int(sz) == 18):
            return False
        return True

    for section in doc.sections:
        headers_to_check = [section.header]
        if section.different_first_page_header_footer:
            headers_to_check.append(section.first_page_header)
        headers_to_check.append(section.even_page_header)

        for hdr in headers_to_check:
            if hdr.is_linked_to_previous:
                continue
            # 前提：只判"出来的"页眉（含非空文本）
            if not any(r.text.strip() for p in hdr.paragraphs for r in p.runs):
                continue
            header_line_seen = True
            # 该页眉内需存在一条满足②③④的直线（①：pBdr 存在 + edge 命中）
            line_found = False
            for para in hdr.paragraphs:
                pPr = para._element.find(qn('w:pPr'))
                pBdr = pPr.find(qn('w:pBdr')) if pPr is not None else None
                if pBdr is None:
                    continue
                for edge in pBdr:
                    if _edge_is_double_lightgray_225(edge):
                        line_found = True
                        break
                if line_found:
                    break
            if not line_found:
                has_header_line = False
                break
        if not has_header_line:
            break

    if header_line_seen and has_header_line:
        results.append("+5：文档中所有页眉处出现直线：类型为双实线，颜色为浅灰色，2.25磅")
    else:
        results.append("+0：文档中所有页眉处出现直线：类型为双实线，颜色为浅灰色，2.25磅")

    # +5：页面设置（A4，上3.3cm，下2.6cm，左2.1cm，右2.2cm）
    page_setup_ok = True
    for section in doc.sections:
        page_width = section.page_width.cm if section.page_width else 0
        page_height = section.page_height.cm if section.page_height else 0
        if not (20.8 <= page_width <= 21.2 and 29.5 <= page_height <= 29.9):
            page_setup_ok = False
            break
        top = section.top_margin.cm if section.top_margin else 0
        bottom = section.bottom_margin.cm if section.bottom_margin else 0
        left = section.left_margin.cm if section.left_margin else 0
        right = section.right_margin.cm if section.right_margin else 0
        if not (3.1 <= top <= 3.5 and 2.4 <= bottom <= 2.8 and
                1.9 <= left <= 2.3 and 2.0 <= right <= 2.4):
            page_setup_ok = False
            break
    if page_setup_ok:
        results.append("+5：页面设置（A4，上3.3cm，下2.6cm，左2.1cm，右2.2cm）")
    else:
        results.append("+0：页面设置（A4，上3.3cm，下2.6cm，左2.1cm，右2.2cm）")

    # +3：第1页顶部出现非空标题文本：字体为黑体，小二号，居中显示
    # 只要第1页顶部有一段非空标题文本，就校验其格式——黑体 + 小二号 + 居中；三项都满足才 +3。
    #   ① 位置：第 1 页顶部 = 文档第一个非空段落（存在非空 run）
    #   ② 居中：段落 alignment == CENTER
    #   ③ 黑体：标题段落所有非空 run 的中文字体（按办公软件继承链解析实际生效字体）均须含 黑体/SimHei
    #   ④ 小二号：标题段落所有非空 run 的字号均须 = 小二（18pt，容差 17–19pt）
    #   前提：标题文本必须"非空"（存在非空 run）；③④须对该段所有非空 run 成立，任一 run 不符即判 +0。
    #
    # 与办公软件对齐：中文字体存 run 的 w:rFonts/@w:eastAsia（未设时沿样式/docDefaults 继承）；
    #   字号存 w:sz（半磅）；居中存段落 w:jc。均为 Word 可设可读的属性。
    HEITI_KEYS = ('黑体', 'SimHei')

    title_page_ok = False
    first_title_para = None
    for para in doc.paragraphs:
        if any(r.text.strip() for r in para.runs):
            first_title_para = para
            break

    if first_title_para is not None:
        # ② 居中
        if first_title_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            # ③+④ 所有非空 run 均须同时满足黑体 + 小二
            title_page_ok = True
            for run in first_title_para.runs:
                if not run.text.strip():
                    continue
                ea = _resolve_effective_eastasia(doc, run)
                is_heiti = bool(ea and any(k in ea for k in HEITI_KEYS))
                is_xiaoer = bool(run.font.size and 17 <= run.font.size.pt <= 19)
                if not (is_heiti and is_xiaoer):
                    title_page_ok = False
                    break
    if title_page_ok:
        results.append("+3：第1页顶部出现非空标题文本：字体为黑体，小二号，居中显示")
    else:
        results.append("+0：第1页顶部出现非空标题文本：字体为黑体，小二号，居中显示")

    # +3：英文状态下的作者名字、机构名称、收稿/修订/录用/发表日期等信息：
    #     文本字体为Times New Roman、小四、居中
    # 按确认口径：判定对象 = 标题页区域（大标题之后、"摘要"之前）里所有"著录信息段"，
    #   即作者/机构/收稿日期等英文著录信息。对这些段逐段校验三点，全部合规才 +3：
    #   ① 文本字体为 Times New Roman：西文字体 w:ascii（回落 w:hAnsi）== "Times New Roman"
    #   ② 小四：字号 12pt（小四 = 12pt，容差 11.5–12.5）
    #   ③ 居中：段落 alignment == CENTER
    #   前提：这些著录信息文本必须"出来"（存在非空著录段）；一段都没有则判 +0。
    #
    # 范围界定：从大标题的下一段起，到"摘要"标题之前；排除大标题本身与空段。
    # 与办公软件对齐：西文字体存 run 的 w:rFonts/@w:ascii、@w:hAnsi；字号 w:sz（半磅）；
    #   居中存段落 w:jc。均为 Word 可设可读的属性。

    # 1) 标题页边界："摘要"之前
    title_page_end = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().replace(' ', '').replace('　', '') == '摘要':
            title_page_end = i
            break

    def _check_para_format(para):
        """校验段落：TNR + 小四(11.5–12.5pt) + CENTER。段落内所有非空 run 都必须满足。"""
        # ③ 居中
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            return False
        seen = False
        for run in para.runs:
            if not run.text.strip():
                continue
            seen = True
            # ① 西文字体 Times New Roman（按办公软件继承链解析实际生效字体）
            if _resolve_effective_ascii(doc, run) != 'Times New Roman':
                return False
            # ② 小四
            if not (run.font.size and 11.5 <= run.font.size.pt <= 12.5):
                return False
        return seen

    # 2) 标题页内的"著录信息段"：大标题(第一个非空段)之后、摘要之前的所有非空段
    first_nonempty_idx = None
    for i in range(title_page_end):
        if doc.paragraphs[i].text.strip():
            first_nonempty_idx = i
            break

    info_paras = []
    if first_nonempty_idx is not None:
        for i in range(first_nonempty_idx + 1, title_page_end):
            if doc.paragraphs[i].text.strip():
                info_paras.append(doc.paragraphs[i])

    # 3) 所有著录段逐段校验，全部合规且至少存在一段才给分
    english_info_ok = bool(info_paras) and all(_check_para_format(p) for p in info_paras)
    if english_info_ok:
        results.append("+3：英文状态下的作者名字、机构名称、收稿/修订/录用/发表日期等信息：文本字体为Times New Roman、小四、居中")
    else:
        results.append("+0：英文状态下的作者名字、机构名称、收稿/修订/录用/发表日期等信息：文本字体为Times New Roman、小四、居中")

    # +1：英文状态下的作者名字、机构名称、收稿/修订/录用/发表日期等文献信息与下方"摘要"两字之间空一行
    # 按确认口径："空一行" = "摘要"标题段紧邻的上一段是空段落（Word 里多敲的一个回车）。
    #   ① 定位"摘要"标题段：strip 并去空格后等于"摘要"（容忍"摘  要"这类中间空格写法）
    #   ② 空一行：其紧邻上一段(i-1)为空段落（无任何非空白文本）
    #   前提：必须存在"摘要"标题段且它上方还有段落；否则判 +0。
    #
    # 与办公软件对齐：Word 的"空一行"就是一个空的段落(<w:p> 无文本)；此判定与在 Word 里
    #   于文献信息和"摘要"之间敲一个回车形成的空行完全对应。
    abstract_spacing_ok = False
    for i, para in enumerate(doc.paragraphs):
        norm = para.text.strip().replace(' ', '').replace('　', '')
        if norm == "摘要":                       # ① 摘要标题段
            if i > 0:
                prev = doc.paragraphs[i - 1]
                if not prev.text.strip():        # ② 上一段为空段落
                    abstract_spacing_ok = True
            break
    if abstract_spacing_ok:
        results.append("+1：英文状态下的作者名字、机构名称、收稿/修订/录用/发表日期等文献信息与下方\"摘要\"两字之间空一行")
    else:
        results.append("+0：英文状态下的作者名字、机构名称、收稿/修订/录用/发表日期等文献信息与下方\"摘要\"两字之间空一行")

    # +3：摘要页"摘要"标题文本：内容为中文；字体为黑体、三号，居中显示，文本颜色为钢蓝（着色1、深色25%）
    # 严格对齐细则，五个点都踩到；标题段所有非空 run 均须同时满足字体/字号/颜色，段落满足居中，内容为中文"摘要"：
    #   ① 内容为中文："摘要"（容忍中间空格，如"摘  要"）
    #   ② 字体为黑体：run 中文字体（按办公软件继承链解析实际生效字体）含 黑体/SimHei
    #   ③ 三号：三号 = 16pt（容差 15.5–16.5pt）
    #   ④ 居中显示：段落 alignment == CENTER
    #   ⑤ 文本颜色为钢蓝（着色1、深色25%）：
    #        (a) 主题色：w:color/@w:themeColor = accent1（着色1），且 @w:themeShade≈BF（深色25%）
    #        (b) 或固定 RGB = 着色1深色25% 的钢蓝色值 #2E74B5（含常见等价 1F4E79/2F5496 等）
    #   前提：必须存在中文"摘要"标题段；②③⑤须对该段所有非空 run 成立，任一 run 不符即判 +0。
    #
    # 与办公软件对齐：中文字体 w:rFonts/@w:eastAsia；字号 w:sz（半磅）；居中 w:jc；
    #   颜色 w:color——"着色1 深色25%"在 Word 里存为 themeColor=accent1 + themeShade=BF，
    #   若手选标准色则存固定 RGB。均为 Word「字体颜色」可设可读的属性。
    def _heiti_of(run):
        return _resolve_effective_eastasia(doc, run)

    def _is_steel_blue(run):
        """钢蓝（着色1、深色25%）：主题色 accent1 + themeShade 深色，或对应固定 RGB。"""
        rPr = run._element.rPr
        if rPr is None:
            return False
        color = rPr.find(qn('w:color'))
        if color is None:
            return False
        # (a) 主题色：着色1(accent1) + 深色 25%(themeShade≈BF)
        theme = color.get(qn('w:themeColor'))
        shade = color.get(qn('w:themeShade'))
        if theme == 'accent1' and shade:
            try:
                sv = int(shade, 16)
                # 深色25% ⇒ 保留 75% 亮度 ⇒ shade≈BF(191)，容差 [0xB0,0xCF]
                if 0xB0 <= sv <= 0xCF:
                    return True
            except ValueError:
                pass
        # (b) 固定 RGB：着色1深色25% 的钢蓝 = #2E74B5（及常见等价深钢蓝）
        val = color.get(qn('w:val'))
        if val and re.fullmatch(r'[0-9A-Fa-f]{6}', val):
            STEEL_BLUES = {'2E74B5', '1F4E79', '2F5496', '1F3864', '305496'}
            if val.upper() in STEEL_BLUES:
                return True
        return False

    abstract_title_ok = False
    for para in doc.paragraphs:
        norm = para.text.strip().replace(' ', '').replace('　', '')
        if norm == "摘要":                       # ① 内容为中文"摘要"
            # ④ 居中
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                break
            # ②③⑤ 所有非空 run 均须同时满足黑体 + 三号 + 钢蓝
            abstract_title_ok = True
            for run in para.runs:
                if not run.text.strip():
                    continue
                ea = _heiti_of(run)
                is_heiti = bool(ea and ('黑体' in ea or 'SimHei' in ea))   # ② 黑体
                is_sanhao = bool(run.font.size and 15.5 <= run.font.size.pt <= 16.5)  # ③ 三号
                is_steel = _is_steel_blue(run)                             # ⑤ 钢蓝
                if not (is_heiti and is_sanhao and is_steel):
                    abstract_title_ok = False
                    break
            break
    if abstract_title_ok:
        results.append("+3：摘要页\"摘要\"标题文本：内容为中文；字体为黑体、三号，居中显示，文本颜色为钢蓝（着色1、深色25%）")
    else:
        results.append("+0：摘要页\"摘要\"标题文本：内容为中文；字体为黑体、三号，居中显示，文本颜色为钢蓝（着色1、深色25%）")

    # +3：摘要页标题下方的文本：内容为中文；字体为宋体、小四号、两端对齐
    # 严格对齐细则的三点，对摘要正文所有段、所有非空 run 判定（本条不涉及字体颜色）：
    #   ① 内容为中文（含中文字符）
    #   ② 字体为宋体：含中文的 run，其"实际生效"的东亚字体为 宋体/SimSun/NSimSun
    #      —— 按办公软件口径：run 未显式设时，沿"段落样式→docDefaults"继承链解析出真正生效的字体
    #   ③ 小四：11.5–12.5pt
    #   ④ 两端对齐：alignment == JUSTIFY
    # 范围：摘要标题（容忍 "摘  要"）的下一段，到"关键词"之前
    abstract_content_ok = False
    abs_start = -1
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().replace(' ', '').replace('　', '') == "摘要":
            abs_start = i + 1
            break
    if abs_start > 0:
        # 摘要正文范围：直到遇到"关键词"或下一个标题
        abs_paras = []
        for j in range(abs_start, len(doc.paragraphs)):
            t = doc.paragraphs[j].text.strip()
            if not t:
                continue
            if t.startswith('关键词') or is_title(doc.paragraphs[j]):
                break
            abs_paras.append(doc.paragraphs[j])
        if abs_paras:
            all_ok = True
            for para in abs_paras:
                # ④ 两端对齐
                if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    all_ok = False
                    break
                has_chinese = False
                for run in para.runs:
                    t = run.text or ''
                    if not t.strip():
                        continue
                    # ③ 小四
                    if not (run.font.size and 11.5 <= run.font.size.pt <= 12.5):
                        all_ok = False
                        break
                    # ① + ② 含中文的 run → 实际生效字体为宋体
                    if re.search(r'[一-鿿]', t):
                        has_chinese = True
                        ea = _resolve_effective_eastasia(doc, run)
                        if not (ea and any(k in ea for k in STRICT_SONGTI)):
                            all_ok = False
                            break
                if not all_ok:
                    break
                # ① 段落必须含中文
                if not has_chinese:
                    all_ok = False
                    break
            abstract_content_ok = all_ok
    if abstract_content_ok:
        results.append("+3：摘要页标题下方的文本：内容为中文；字体为宋体、小四号、两端对齐")
    else:
        results.append("+0：摘要页标题下方的文本：内容为中文；字体为宋体、小四号、两端对齐")

    # +3：关键词及其冒号后面的词语：内容均为中文；冒号后面的词语用全角分号分隔，
    #     最后一个词后面没有标点符号；字体为黑体、五号、两端对齐
    # 严格对齐细则，仅踩细则明确写出的点：
    #   ① 内容均为中文：整段（"关键词"标签 + 冒号后所有词语）除全角分隔符外均为中文字符
    #   ② 冒号后面的词语用全角分号"；"分隔
    #   ③ 最后一个词后面没有标点符号（body 结尾不是任何标点）
    #   ④ 字体为黑体：所有非空 run 的中文字体（办公软件实际生效）含 黑体/SimHei
    #   ⑤ 五号：所有非空 run 字号 = 五号（10.5pt，容差 10–11pt）
    #   ⑥ 两端对齐：段落 alignment == JUSTIFY
    #
    # 与办公软件对齐：中文字体 w:rFonts/@w:eastAsia（沿样式/docDefaults 继承）；
    #   字号 w:sz（半磅）；两端对齐 w:jc=both。均为 Word 可设可读的属性。
    PUNCT_SET = set('；;。.，,、!！?？：:～~…—-（）()【】[]《》<>"\'"“”‘’')
    keywords_ok = False
    for para in doc.paragraphs:
        if not para.text.strip().startswith('关键词'):
            continue
        text = para.text.strip()
        ok = True

        # ⑥ 两端对齐
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            ok = False

        # 拆"关键词[：]xxx；yyy；zzz"（细则未规定冒号全/半角，两者都接受）
        m = re.match(r'^关键词[：:](.*)$', text)
        if not m:
            ok = False
        else:
            body = m.group(1)
            # ② 全角分号分隔：不允许出现半角 ; 或逗号 ，,（即不能用别的分隔符）
            if any(c in body for c in (';', ',', '，')):
                ok = False
            # ③ 最后一个词后无标点：body 结尾不能是标点
            if body and body[-1] in PUNCT_SET:
                ok = False
            # ① 内容均为中文：整段去掉"关键词"标签、冒号、全角分号后，必须全是中文字符
            content = text[len('关键词'):]         # 去"关键词"
            content = content.lstrip('：:')          # 去冒号
            content = content.replace('；', '')      # 去全角分号分隔符
            if not content or not all('一' <= c <= '鿿' for c in content):
                ok = False

        # ④⑤ 黑体 + 五号：所有非空 run 都必须满足
        run_seen = False
        for run in para.runs:
            if not run.text.strip():
                continue
            run_seen = True
            ea = _resolve_effective_eastasia(doc, run)
            if not (ea and ('黑体' in ea or 'SimHei' in ea)):
                ok = False
                break
            if not (run.font.size and 10 <= run.font.size.pt <= 11):
                ok = False
                break
        if not run_seen:
            ok = False

        keywords_ok = ok
        break
    if keywords_ok:
        results.append("+3：关键词及其冒号后面的词语：内容均为中文；冒号后面的词语用全角分号分隔，最后一个词后面没有标点符号；字体为黑体、五号、两端对齐")
    else:
        results.append("+0：关键词及其冒号后面的词语：内容均为中文；冒号后面的词语用全角分号分隔，最后一个词后面没有标点符号；字体为黑体、五号、两端对齐")

    # +5：一级标题文本：中文文本采用黑体、三号，数字或序号采用Times New Roman，三号；所有标题均为居中对齐
    # 严格对齐细则，仅踩细则明确写出的点；针对"一级标题"（单数字编号，如"1 引言""2 结果与讨论"）：
    #   ① 中文文本：所在 run 中文字体（办公软件实际生效）含 黑体/SimHei，且三号（16pt，容差15.5–16.5）
    #   ② 数字或序号：所在 run 西文字体（办公软件实际生效）== Times New Roman，且三号
    #   ③ 所有标题均为居中对齐：段落 alignment == CENTER
    #   所有一级标题都必须满足①②③才给分；至少识别到一个一级标题。
    #
    # 一级标题识别：按完整的标题编号规则匹配"单个数字编号"（如"1""2.""3、"等，
    #   后接空白和标题正文），不与二级"X.Y"、三级"X.Y.Z"编号混淆；不设标题文本长度限制
    #   （rubric 未规定长度上限，过滤长度会漏检合法的长标题）。
    #
    # 与办公软件对齐：中文字体 w:rFonts/@w:eastAsia、西文/数字字体 w:rFonts/@w:ascii
    #   （均沿样式/docDefaults 继承）；字号 w:sz（半磅）；居中 w:jc。均为 Word 可设可读的属性。
    L1_TITLE_RE = re.compile(r'^(\d+)[.、]?[ \t　\xa0]+(\S.*)$')
    l1_ok = True
    l1_seen = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        m = L1_TITLE_RE.match(text)
        if not m:
            continue
        l1_seen += 1
        # ③ 居中
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            l1_ok = False
            break
        for run in para.runs:
            t = run.text or ''
            if not t.strip():
                continue
            has_zh = bool(re.search(r'[一-鿿]', t))   # 含中文文本
            has_num = bool(re.search(r'\d', t))       # 含数字/序号
            # 三号（含中文或数字的 run 都要满足）
            if (has_zh or has_num) and not (run.font.size and 15.5 <= run.font.size.pt <= 16.5):
                l1_ok = False
                break
            # ① 中文文本 → 黑体
            if has_zh:
                ea = _resolve_effective_eastasia(doc, run)
                if not (ea and ('黑体' in ea or 'SimHei' in ea)):
                    l1_ok = False
                    break
            # ② 数字或序号 → Times New Roman
            if has_num:
                if _resolve_effective_ascii(doc, run) != 'Times New Roman':
                    l1_ok = False
                    break
        if not l1_ok:
            break
    if l1_seen > 0 and l1_ok:
        results.append("+5：一级标题文本：中文文本采用黑体、三号，数字或序号采用Times New Roman，三号；所有标题均为居中对齐")
    else:
        results.append("+0：一级标题文本：中文文本采用黑体、三号，数字或序号采用Times New Roman，三号；所有标题均为居中对齐")

    # +3：标题内序号或数字与文字之间的间距满足：0.2-0.5字符
    # 字符宽度换算：半角空格≈0.5，全角空格≈1.0，不间断空格( )≈0.5；
    # 制表符(\t)取决于具体制表位，难以静态估算 → 保守按 1.0 处理。
    # 所有编号标题段都必须落在 [0.2, 0.5] 区间内才给分。
    SPACE_WIDTH = {
        ' ': 0.5,  # 半角空格
        ' ': 0.5,  # NBSP
        '　': 1.0,  # 全角空格
        '\t':     1.0,  # Tab，保守估
    }
    TITLE_NUM_RE = re.compile(r'^(\d+(?:\.\d+){0,2})([ \t 　]+)(\S.*)$')
    spacing_ok = True
    checked_any = False
    for para in doc.paragraphs:
        text = para.text.strip()
        m = TITLE_NUM_RE.match(text)
        if not m:
            continue
        sep = m.group(2)
        # 求和；未识别字符按半角空格估
        width = sum(SPACE_WIDTH.get(c, 0.5) for c in sep)
        checked_any = True
        if not (0.2 <= width <= 0.5):
            spacing_ok = False
            break
    if checked_any and spacing_ok:
        results.append("+3：标题内序号或数字与文字之间的间距满足：0.2-0.5字符")
    else:
        results.append("+0：标题内序号或数字与文字之间的间距满足：0.2-0.5字符")

    # +3：除标题、图注、表格中的字体、关键词行、无编号小标题行以及作者行，
    #     其余所有中文文本字体均为宋体、小四，两端对齐
    # 排除细则明确写出的六类：标题、图注、表格、关键词行、无编号小标题行、作者行。
    #   - 标题/图注/表格：is_title / is_caption / is_in_table（既有判定）
    #   - 关键词行：段落文本以"关键词"开头
    #   - 无编号小标题行：单独成段、不含标点的短句，且不属于 is_title 的编号标题
    #     （如“摘要”“致谢”外的无编号小标题；此处按"段落文本不含任何标点符号"识别）
    #   - 作者行：标题页"著录信息段"（大标题之后、"摘要"之前的著录信息，即 info_paras）
    # 对剩余的所有中文文本段落逐一判定，三点全部满足才算合规（任一段违反即判 0）：
    #   ⓐ 两端对齐：段落 alignment == JUSTIFY
    #   ⓑ 字体为宋体：段落内含中文的 run，其中文字体（办公软件实际生效）含 宋体/SimSun/NSimSun
    #   ⓒ 小四：含中文的 run 字号 = 小四（12pt，容差 11.5–12.5）
    #   （"字体为宋体、小四"是对中文文本的排版要求，故按含中文 run 判定；纯英文/数字 run 不约束。）
    #
    # 与办公软件对齐：中文字体 w:rFonts/@w:eastAsia（沿样式/docDefaults 继承）；
    #   字号 w:sz（半磅）；两端对齐 w:jc=both。
    NO_PUNCT_RE = re.compile(
        r'^[^，,。.！!？?；;：:、"\'"“”‘’（）()【】\[\]《》<>～~…—-]+$'
    )

    def _is_unnumbered_subtitle(para):
        """无编号小标题行：单独成段、不含任何标点的短句，且不是已识别的编号标题。"""
        text = para.text.strip()
        if not text or is_title(para):
            return False
        return bool(NO_PUNCT_RE.match(text)) and len(text) <= 15

    info_para_ids = {id(p) for p in info_paras}

    body_font_ok = True
    body_seen = 0
    for para in doc.paragraphs:
        # 排除：标题 / 图注 / 表格
        if is_title(para) or is_caption(para) or is_in_table(para):
            continue
        if not para.text.strip():           # 空段落无文本
            continue
        # 排除：关键词行
        if para.text.strip().startswith('关键词'):
            continue
        # 排除：作者行（标题页著录信息段）
        if id(para) in info_para_ids:
            continue
        # 排除：无编号小标题行
        if _is_unnumbered_subtitle(para):
            continue
        # 只判定含中文的段落（rubric 限定"中文文本"）
        if not re.search(r'[一-鿿]', para.text):
            continue
        body_seen += 1
        # ⓐ 两端对齐
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            body_font_ok = False
            break
        # ⓑⓒ 含中文的 run：宋体 + 小四
        for run in para.runs:
            t = run.text or ''
            if not re.search(r'[一-鿿]', t):
                continue
            ea = _resolve_effective_eastasia(doc, run)
            if not (ea and any(k in ea for k in STRICT_SONGTI)):
                body_font_ok = False
                break
            if not (run.font.size and 11.5 <= run.font.size.pt <= 12.5):
                body_font_ok = False
                break
        if not body_font_ok:
            break
    if body_seen > 0 and body_font_ok:
        results.append("+3：除标题、图注、表格中的字体、关键词行、无编号小标题行以及作者行，其余所有中文文本字体均为宋体、小四，两端对齐")
    else:
        results.append("+0：除标题、图注、表格中的字体、关键词行、无编号小标题行以及作者行，其余所有中文文本字体均为宋体、小四，两端对齐")

    # +3：全文脚注要紧贴在文字的右侧，无额外空格；格式为方括号 + 阿拉伯数字，整体为上标
    # 严格对齐细则，仅踩细则写出的四点（未写出的不额外约束）：
    #   ① 紧贴在文字的右侧：脚注左侧紧挨正文文字，左侧不能是空白字符
    #   ② 无额外空格：脚注右侧不能有空白字符
    #   ③ 格式为方括号 + 阿拉伯数字：形如 [\d+]
    #   ④ 整体为上标：承载该脚注的所有 run 全部为上标（办公软件里显示为上标）
    #
    # "脚注"即引用标记——全文任意位置（含表格内）形如 [N] 的标记；排除段首以 [N]
    #   开头的参考文献条目段（那是文末条目而非正文脚注引用）。
    # 与办公软件对齐：上标读取 run 的 w:vertAlign=superscript 生效值（即 Word/WPS 里
    #   看到的上标效果）；左右紧贴按段落可见文本的相邻字符判定。
    WHITESPACE = set(' 　\t\xa0')
    footnote_total = 0
    footnote_violations = 0

    def _check_footnote_fmt_para(para):
        nonlocal footnote_total, footnote_violations
        text = para.text or ''
        # 跳过参考文献条目：以 [\d+] 开头的段落（编号在段首，不是脚注引用）
        if re.match(r'^\s*\[\d+\]', text):
            return
        # 先把段落里 run 的累积起止位置算出
        run_spans = []  # [(start, end, run)]
        cursor = 0
        for r in para.runs:
            rt = r.text or ''
            run_spans.append((cursor, cursor + len(rt), r))
            cursor += len(rt)

        for m in re.finditer(r'\[\d+\]', text):
            footnote_total += 1
            s, e = m.start(), m.end()
            # ① 紧贴文字右侧：左侧不能是空白字符（段首也算紧贴）
            if s > 0 and text[s - 1] in WHITESPACE:
                footnote_violations += 1
                continue
            # ② 无额外空格：右侧不能有空白字符
            if e < len(text) and text[e] in WHITESPACE:
                footnote_violations += 1
                continue
            # ③ 格式为方括号+阿拉伯数字（由 finditer 的正则 \[\d+\] 已保证）
            # ④ 整体为上标：覆盖该区间的所有 run 全部 superscript
            covering = [r for (rs, re_, r) in run_spans if rs < e and re_ > s]
            if not covering or not all(r.font.superscript for r in covering):
                footnote_violations += 1
                continue

    # 正文段落（表格外全部）
    for para in doc.paragraphs:
        _check_footnote_fmt_para(para)
    # 表格内段落（细则口径：表格里的引用标记也算脚注）
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _check_footnote_fmt_para(para)

    if footnote_total > 0 and footnote_violations == 0:
        results.append("+3：全文脚注紧贴文字右侧无额外空格；方括号+阿拉伯数字，整体上标")
    else:
        results.append("+0：全文脚注紧贴文字右侧无额外空格；方括号+阿拉伯数字，整体上标")

    # +3：全文脚注编号总共到40，字体为Times New Roman、小四
    # 严格对齐细则，仅踩细则写出的三点（未写出的不额外约束）：
    #   ① 编号总共到40：全文脚注出现过的最大编号达到 40（脚注编号总共编到 40）
    #   ② 字体为 Times New Roman：所有脚注 run 的西文字体为 Times New Roman
    #   ③ 小四：所有脚注 run 字号为小四（12pt，容差 11.5–12.5）
    # "脚注"即引用标记——全文任意位置（含表格内）形如 [N] 的标记；排除段首以 [N]
    #   开头的参考文献条目段（那是文末条目而非正文引用）。
    # 与办公软件对齐：其字体/字号读取 run 生效值（w:ascii 沿样式/docDefaults 继承、
    #   w:sz 半磅），即打开 Word/WPS 看到的效果。
    BRACKET_RE = re.compile(r'\[(\d+(?:[,\-–]\d+)*)\]')
    cited_numbers = set()
    fn_font_ok = True
    fn_size_ok = True
    fn_seen = False

    def _scan_footnote_para(para):
        nonlocal fn_seen, fn_font_ok, fn_size_ok
        text = para.text or ''
        # 跳过参考文献条目段（段首 [N]）
        if re.match(r'^\s*\[\d+\]', text):
            return
        # run 起止位置
        cur = 0
        spans = []
        for r in para.runs:
            t = r.text or ''
            spans.append((cur, cur + len(t), r))
            cur += len(t)
        for m in BRACKET_RE.finditer(text):
            fn_seen = True
            inner = m.group(1)
            for piece in re.split(r'[,]', inner):
                if '-' in piece or '–' in piece:
                    a, b = re.split(r'[-–]', piece)
                    cited_numbers.update(range(int(a), int(b) + 1))
                else:
                    cited_numbers.add(int(piece))
            s, e = m.start(), m.end()
            covering = [r for (rs, re_, r) in spans if rs < e and re_ > s]
            for r in covering:
                if _resolve_effective_ascii(doc, r) != 'Times New Roman':
                    fn_font_ok = False
                if not (r.font.size and 11.5 <= r.font.size.pt <= 12.5):
                    fn_size_ok = False

    # 正文段落（表格外全部）
    for para in doc.paragraphs:
        _scan_footnote_para(para)
    # 表格内段落（细则口径：表格里的引用标记也算脚注）
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_footnote_para(para)

    # ① 编号总共到40：全文脚注出现过的最大编号达到 40
    fn_count_ok = fn_seen and (max(cited_numbers) >= 40 if cited_numbers else False)
    if fn_count_ok and fn_font_ok and fn_size_ok:
        results.append("+3：全文脚注编号总共到40，字体为Times New Roman、小四")
    else:
        results.append("+0：全文脚注编号总共到40，字体为Times New Roman、小四")

    # +3：表格格式满足：全文表序编号连续，表格分页时后页的续表有和原表内容格式均一样的表头
    # 严格对齐细则，仅踩细则写出的两点（未写出的不额外约束）：
    #   ① 全文表序编号连续：全文表题（表格上方居中的"表N ..."标题段）的编号
    #        去重排序后必须为连续的 1..N
    #   ② 表格分页时后页的续表有和原表内容格式均一样的表头：识别相邻 <w:tbl>
    #        之间无正文文字的"续表对"（即同一张表被分页），要求续表首行表头单元格
    #        与原表首行表头单元格文本一致；若不存在续表对（无分页续表），此点 N/A 通过。
    #
    # 与办公软件对齐：表题按段落对齐=居中（区分正文里的"表N..."引用句）；表格分页
    #   在 OOXML 中表现为相邻 <w:tbl> 且中间无正文段落文字，与 Word/WPS 里跨页续表一致。
    table_format_ok = True
    # ① 全文表序编号连续（只取真正的表题：居中的"表N ..."段，排除正文引用句）
    table_nums = []
    for para in doc.paragraphs:
        text = para.text.strip()
        m = re.match(r'^表\s*(\d+)', text)
        if m and para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            table_nums.append(int(m.group(1)))
    unique_nums = sorted(set(table_nums))
    if not unique_nums:
        table_format_ok = False
    elif unique_nums != list(range(1, len(unique_nums) + 1)):
        table_format_ok = False

    # ② 续表头：按 body 顺序找相邻 tbl 对，中间段落全为空才视为续表对（分页续表）
    body = doc.element.body
    body_children = list(body.iterchildren())
    tbl_positions = [i for i, c in enumerate(body_children)
                     if c.tag.split('}')[-1] == 'tbl']

    def _row_texts(tbl_elem):
        """取 OOXML <w:tbl> 第一行所有单元格的纯文本列表（strip 后）"""
        first_tr = tbl_elem.find(qn('w:tr'))
        if first_tr is None:
            return None
        cells = []
        for tc in first_tr.findall(qn('w:tc')):
            cells.append(''.join(tc.itertext()).strip())
        return cells

    cont_pairs_total = 0
    cont_pairs_violations = 0
    for k in range(len(tbl_positions) - 1):
        i1, i2 = tbl_positions[k], tbl_positions[k + 1]
        between = body_children[i1 + 1:i2]
        # 中间所有段落必须无任何非空白文字（仅分页符/空段）才算续表对
        is_cont_pair = True
        for elem in between:
            tag = elem.tag.split('}')[-1]
            if tag != 'p':
                # 中间夹了别的元素（如 sectPr）一般也无文字，跳过判断
                continue
            text = ''.join(elem.itertext()).strip()
            if text:
                is_cont_pair = False
                break
        if not is_cont_pair:
            continue
        cont_pairs_total += 1
        first1 = _row_texts(body_children[i1])
        first2 = _row_texts(body_children[i2])
        if first1 is None or first2 is None or first1 != first2:
            cont_pairs_violations += 1

    if cont_pairs_violations > 0:
        table_format_ok = False

    if table_format_ok:
        if cont_pairs_total > 0:
            results.append(
                f"+3：表格格式（表序编号连续，共{len(unique_nums)}个表；"
                f"识别到{cont_pairs_total}对续表，续表头与原表一致）"
            )
        else:
            results.append(
                f"+3：表格格式（表序编号连续，共{len(unique_nums)}个表；无跨页续表）")
    else:
        results.append("+0：表格格式（表序编号连续，续表头与原表一致）")

    # +3：全文所有的表格标题：位于表格上方居中显示，距离表格上边框线0-0.3字符；字体采用宋体五号、加粗
    # 严格对齐细则，仅踩细则写出的点（未写出的不额外约束）：
    #   ① 位于表格上方：取每个 <w:tbl> 紧邻的上一个非空段落，且为 "表N ..." 标题文本
    #   ② 居中显示：段落对齐 == CENTER
    #   ③ 距离表格上边框线 0-0.3 字符：标题段与表格间距（段后距 space_after）在
    #        0 到 0.3 字符之间（0.3 字符按五号 10.5pt 折算 ≈ 3.15pt）
    #   ④ 字体宋体：标题段含中文的 run，其东亚字体（办公软件实际生效）为 宋体/SimSun/NSimSun
    #   ⑤ 五号：标题段非空 run 字号为五号（10.5pt，容差 10–11）
    #   ⑥ 加粗：标题段非空 run 全部加粗（font.bold == True）
    #
    # 与办公软件对齐：居中读 w:jc=center；标题与表格间距读标题段 w:spacing/@w:after（无则为0）；
    #   中文字体沿 w:rFonts/@w:eastAsia 继承链解析、字号 w:sz（半磅）、加粗 w:b，
    #   即打开 Word/WPS 看到的效果。所有表格都必须有合规标题才给分。
    TABLE_TITLE_RE = re.compile(r'^表\s*\d+')
    CHAR_PT = 10.5           # 五号字号磅值（标题字号）
    MAX_AFTER_PT = 0.3 * CHAR_PT   # 0.3 字符 ≈ 3.15pt

    # 按 body 顺序建立"元素序列"，便于定位每个表格的上一段落
    body_children = list(doc.element.body.iterchildren())
    # 段落 element -> Paragraph 映射
    para_by_elem = {p._element: p for p in doc.paragraphs}

    table_title_ok = True
    table_count_for_title = 0
    for idx, child in enumerate(body_children):
        if not child.tag.endswith('}tbl'):
            continue
        table_count_for_title += 1
        # 找紧邻的上一个段落
        prev_para = None
        for k in range(idx - 1, -1, -1):
            prev = body_children[k]
            if prev.tag.endswith('}p'):
                p_obj = para_by_elem.get(prev)
                if p_obj is not None and p_obj.text.strip():
                    prev_para = p_obj
                break
        # ① 上方必须是 "表N ..." 标题
        if prev_para is None or not TABLE_TITLE_RE.match(prev_para.text.strip()):
            table_title_ok = False
            break
        # ② 居中显示
        if prev_para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            table_title_ok = False
            break
        # ③ 距离表格上边框线 0-0.3 字符：标题段与表格间距（段后距）在 [0, 0.3字符]
        #    未设置段后距时按 0 处理（0 在 0-0.3 范围内，合规）
        sa = prev_para.paragraph_format.space_after
        sa_pt = sa.pt if sa is not None else 0.0
        if not (0.0 <= sa_pt <= MAX_AFTER_PT):
            table_title_ok = False
            break
        # ④⑤⑥ 宋体 + 五号 + 加粗
        para_ok = True
        for run in prev_para.runs:
            t = run.text or ''
            if not t.strip():
                continue
            # ⑥ 加粗
            if not run.font.bold:
                para_ok = False
                break
            # ⑤ 五号（10.5pt，容差 10–11）
            if not (run.font.size and 10 <= run.font.size.pt <= 11):
                para_ok = False
                break
            # ④ 中文 run 宋体
            if re.search(r'[一-鿿]', t):
                ea = _resolve_effective_eastasia(doc, run)
                if not (ea and any(k in ea for k in STRICT_SONGTI)):
                    para_ok = False
                    break
        if not para_ok:
            table_title_ok = False
            break
    if table_count_for_title > 0 and table_title_ok:
        results.append("+3：表格标题（表上方居中，距上边框0-0.3字符，宋体五号加粗）")
    else:
        results.append("+0：表格标题（表上方居中，距上边框0-0.3字符，宋体五号加粗）")

    # +3：全文所有表格中的文本：中文字体采用宋体、五号，英文和阿拉伯数字采用Times New Roman、五号
    # 严格对齐细则，仅踩细则写出的点（未写出的不额外约束）。对表格内每个非空 run 判定：
    #   ① 中文字体采用宋体：run 含中文字符时，其东亚字体（办公软件实际生效）为 宋体/SimSun/NSimSun
    #   ② 英文和阿拉伯数字采用 Times New Roman：run 含英文字母或阿拉伯数字时，
    #        其西文字体（办公软件实际生效）为 Times New Roman
    #   ③ 五号：run 字号为五号（10.5pt，容差 10–11）
    # 任一 run 违反即判 0；至少要有一个表格内非空 run 参与判定。
    # 与办公软件对齐：中文字体沿 w:rFonts/@w:eastAsia 继承链解析、西文字体沿
    #   w:ascii→hAnsi 继承链解析、字号 w:sz（半磅），即打开 Word/WPS 看到的效果。
    ZH_RE = re.compile(r'[一-鿿]')
    EN_RE = re.compile(r'[A-Za-z0-9]')
    SONGTI_KEYS = ('宋体', 'SimSun', 'NSimSun')

    table_content_ok = True
    table_checked_any = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        t = run.text or ''
                        if not t.strip():
                            continue
                        table_checked_any = True
                        ea = _resolve_effective_eastasia(doc, run)
                        en_font = _resolve_effective_ascii(doc, run)
                        # ③ 五号
                        if not (run.font.size and 10.0 <= run.font.size.pt <= 11.0):
                            table_content_ok = False
                            break
                        # ① 中文字符 → 宋体
                        if ZH_RE.search(t):
                            if not (ea and any(k in ea for k in SONGTI_KEYS)):
                                table_content_ok = False
                                break
                        # ② 英文/阿拉伯数字 → Times New Roman
                        if EN_RE.search(t):
                            if en_font != 'Times New Roman':
                                table_content_ok = False
                                break
                    if not table_content_ok:
                        break
                if not table_content_ok:
                    break
            if not table_content_ok:
                break
        if not table_content_ok:
            break
    if table_checked_any and table_content_ok:
        results.append("+3：全文所有表格中文本：中文宋体五号，英文/阿拉伯数字Times New Roman五号")
    else:
        results.append("+0：全文所有表格中文本：中文宋体五号，英文/阿拉伯数字Times New Roman五号")

    # +3：全文所有表格：表格中只有三条边框线，上下边框单实线1.5磅，栏目线单实线0.75磅
    # 严格对齐（已按用户修订的）细则，仅踩细则写出的点（未写出的不额外约束）：
    #   ① 表格中只有三条边框线（三线格）：整张表只显示三条横线——顶线、栏目线、底线；
    #        所有竖线（左/右/内竖）无线，除三条横线外的内部横线也无线
    #   ② 上下边框单实线1.5磅：顶线（首行上边框）、底线（末行下边框）为 single，1.5磅
    #        （w:sz=12，容差 11–13）
    #   ③ 栏目线单实线0.75磅：表头下方那条横线（首行下边框）为 single，0.75磅
    #        （w:sz=6，容差 5–7）
    #
    # 与办公软件对齐：本文档三线用单元格级边框(tcBorders)实现，Word/WPS 里看到的即
    #   每个 tc 的 top/bottom/left/right 生效值。故按"办公软件实际渲染的每条格线"判定：
    #     顶线   = 首行各 tc 的 top
    #     栏目线 = 首行各 tc 的 bottom
    #     底线   = 末行各 tc 的 bottom
    #     竖线   = 所有 tc 的 left/right（须无线）
    #     其余内部横线 = 非首行的 top、非末行的 bottom（须无线）
    def _get_border(elem, name):
        """返回 (val, sz_int_or_None)；elem 是 tcBorders/tblBorders，name 是 'top'/'bottom'/..."""
        if elem is None:
            return (None, None)
        sub = elem.find(qn(f'w:{name}'))
        if sub is None:
            return (None, None)
        val = sub.get(qn('w:val'))
        sz_str = sub.get(qn('w:sz'))
        sz = int(sz_str) if sz_str and sz_str.isdigit() else None
        return (val, sz)

    def _tc_borders(cell):
        tcPr = cell._tc.find(qn('w:tcPr'))
        return tcPr.find(qn('w:tcBorders')) if tcPr is not None else None

    NO_LINE = {None, 'nil', 'none'}

    def _is_single(val, sz, lo, hi):
        return val == 'single' and sz is not None and lo <= sz <= hi

    border_ok = True
    if not doc.tables:
        border_ok = False
    for table in doc.tables:
        rows = table.rows
        if not rows:
            border_ok = False
            break
        n = len(rows)

        # ① 竖线全部无线 & 除三条横线外的内部横线全部无线
        no_extra_line = True
        for ri, row in enumerate(rows):
            for cell in row.cells:
                tb = _tc_borders(cell)
                lv, _ = _get_border(tb, 'left')
                rv, _ = _get_border(tb, 'right')
                if lv not in NO_LINE or rv not in NO_LINE:
                    no_extra_line = False
                    break
                # 非首行的上边框应无线（首行上边框=顶线，单独判）
                if ri > 0:
                    tv, _ = _get_border(tb, 'top')
                    if tv not in NO_LINE:
                        no_extra_line = False
                        break
                # 非末行的下边框应无线（首行下边框=栏目线、末行下边框=底线，单独判）
                if ri < n - 1 and ri != 0:
                    bv, _ = _get_border(tb, 'bottom')
                    if bv not in NO_LINE:
                        no_extra_line = False
                        break
            if not no_extra_line:
                break
        if not no_extra_line:
            border_ok = False
            break

        # ② 顶线（首行各 tc 上边框）= single 1.5磅
        top_ok = True
        for cell in rows[0].cells:
            v, s = _get_border(_tc_borders(cell), 'top')
            if not _is_single(v, s, 11, 13):
                top_ok = False
                break
        # ② 底线（末行各 tc 下边框）= single 1.5磅
        bottom_ok = True
        for cell in rows[-1].cells:
            v, s = _get_border(_tc_borders(cell), 'bottom')
            if not _is_single(v, s, 11, 13):
                bottom_ok = False
                break
        if not (top_ok and bottom_ok):
            border_ok = False
            break

        # ③ 栏目线（首行各 tc 下边框）= single 0.75磅
        col_ok = True
        if n < 2:
            # 只有一行则无表头/栏目分隔线
            col_ok = False
        else:
            for cell in rows[0].cells:
                v, s = _get_border(_tc_borders(cell), 'bottom')
                if not _is_single(v, s, 5, 7):
                    col_ok = False
                    break
        if not col_ok:
            border_ok = False
            break

    if border_ok:
        results.append("+3：全文所有表格为三线格（上下1.5磅，栏目线0.75磅，单实线）")
    else:
        results.append("+0：全文所有表格为三线格（上下1.5磅，栏目线0.75磅，单实线）")

    # +3：所有图注文本：位于图片下方居中显示；字体采用宋体五号、加粗
    # 严格对齐细则，仅踩细则写出的点（未写出的不额外约束）：
    #   ① 位于图片下方：图注段（"图N ..."）的上一段落必须含图片（<w:drawing> 或 <w:pict>），
    #        即办公软件里图注紧接在图片下方一行
    #   ② 居中显示：图注段落对齐 == CENTER
    #   ③ 字体采用宋体：图注段含中文的 run，其东亚字体（办公软件实际生效）为 宋体/SimSun/NSimSun
    #   ④ 五号：图注段非空 run 字号为五号（10.5pt，容差 10–11）
    #   ⑤ 加粗：图注段非空 run 全部加粗（font.bold == True）
    #
    # 图注识别：段落形如 "图N ..." 且居中——用居中区分正文里的"图N..."句首引用
    #   （引用句为两端对齐，不居中），与办公软件里真正的图注一致。
    # 至少识别到 1 段图注；任一图注违反即判 0。
    # 与办公软件对齐：居中读 w:jc=center；上方图片读上一段的 <w:drawing>/<w:pict>；
    #   中文字体沿 w:rFonts/@w:eastAsia 继承链解析、字号 w:sz、加粗 w:b，即 Word/WPS 看到的效果。
    CAPTION_RE = re.compile(r'^图\s*\d+')

    def _has_drawing(p_elem):
        return (p_elem.find('.//' + qn('w:drawing')) is not None or
                p_elem.find('.//' + qn('w:pict')) is not None)

    def _eastasia(run):
        return _resolve_effective_eastasia(doc, run)

    caption_ok = True
    caption_seen = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not CAPTION_RE.match(text):
            continue
        # ② 居中显示（同时用于区分"图N..."正文引用 vs 真图注）
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            continue
        caption_seen += 1
        # ① 位于图片下方：上一段落须含图片
        if i == 0 or not _has_drawing(doc.paragraphs[i - 1]._element):
            caption_ok = False
            break
        # ③④⑤ 逐 run 校验字体/字号/加粗
        para_format_ok = True
        for run in para.runs:
            t = run.text or ''
            if not t.strip():
                continue
            # ⑤ 加粗
            if not run.font.bold:
                para_format_ok = False
                break
            # ④ 五号（10.5pt，容差 10–11）
            if not (run.font.size and 10.0 <= run.font.size.pt <= 11.0):
                para_format_ok = False
                break
            # ③ 中文字符 → 宋体（严格只认 宋体/SimSun/NSimSun）
            if re.search(r'[一-鿿]', t):
                ea = _eastasia(run)
                if not (ea and any(k in ea for k in STRICT_SONGTI)):
                    para_format_ok = False
                    break
        if not para_format_ok:
            caption_ok = False
            break
    if caption_seen > 0 and caption_ok:
        results.append("+3：所有图注文本：图片下方居中，宋体五号加粗")
    else:
        results.append("+0：所有图注文本：图片下方居中，宋体五号加粗")

    # +3：图片与图注没有出现分页显示，图注的图序号与标题之间出现0.2-0.5字符的距离
    # 严格对齐细则，仅踩细则写出的两点（未写出的不额外约束）：
    #   ① 图片与图注没有出现分页显示：图片与它的图注必须在同一页——是否分页取决于
    #        实际排版结果（页边距/字号/上下文内容都会影响分页位置），静态读取分页符
    #        （w:pageBreakBefore/手动分页）无法发现"自然分页"把图片和图注挤到两页的
    #        情况，因此改为把文档渲染为 PDF，按渲染后的真实页码判断图片段与图注段
    #        是否落在同一页。
    #   ② 图注的图序号与标题之间出现0.2-0.5字符的距离：图注"图N"与其后标题文字之间的
    #        分隔间距落在 0.2–0.5 字符。
    #
    # 与办公软件对齐：用 LibreOffice/Word 把文档渲染成 PDF（与打印/预览的分页结果一致），
    #   再用 pdfplumber 按段落文本定位其渲染后所在页码；间距按分隔字符的字符宽求和
    #   （半角/NBSP≈0.5，全角≈1.0，Tab≈1.0），与"标题序号与文字间距"口径一致。
    CAPTION_FULL_RE = re.compile(r'^(图\s*\d+)(\s*)(.*)$')

    def _build_page_locator(src_file_path):
        """把文档渲染为 PDF，返回一个函数：给定文本片段，返回其首次出现的页码（0-based）。
        渲染或 pdfplumber 文本解析失败时返回 None（调用方按"无法判定"处理）。"""
        render_tmp_dir = tempfile.mkdtemp(prefix='officeval_008_pdf_')
        try:
            pdf_path = _render_to_pdf(src_file_path, render_tmp_dir)
        except Exception:
            shutil.rmtree(render_tmp_dir, ignore_errors=True)
            return None

        try:
            with pdf_backend.open_pdf(pdf_path) as pdf_doc:
                page_texts = [
                    pdf_doc.page_text(index)
                    for index in range(pdf_doc.page_count)
                ]
        except Exception:
            shutil.rmtree(render_tmp_dir, ignore_errors=True)
            return None
        shutil.rmtree(render_tmp_dir, ignore_errors=True)

        # 逐页拼接查找游标：同一段文本可能在多页重复（如页眉），用"从上次命中页开始往后找"
        # 的方式，尽量按文档顺序把段落文本对应到递增的页码序列上。
        state = {'cursor': 0}

        def _locate(snippet):
            snippet = snippet.strip()
            if not snippet:
                return None
            for p in range(state['cursor'], len(page_texts)):
                if snippet in page_texts[p]:
                    state['cursor'] = p
                    return p
            # 未在游标之后找到，放宽到全文重新找一次（不推进游标，避免顺序错乱）
            for p, txt in enumerate(page_texts):
                if snippet in txt:
                    return p
            return None

        return _locate

    _locate_page = _build_page_locator(file_path)

    fig_layout_ok = True
    fig_layout_seen = 0
    page_locate_failed = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        m = CAPTION_FULL_RE.match(text)
        if not m or para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            continue
        # 上一段必须含图片（与上一条评分项保持口径一致：图片在上方）
        if i == 0 or not _has_drawing(doc.paragraphs[i - 1]._element):
            continue
        fig_layout_seen += 1
        prev_para = doc.paragraphs[i - 1]
        # ① 图片与图注是否同页：按渲染后的真实分页结果判断
        if _locate_page is None:
            page_locate_failed = True
            fig_layout_ok = False
            break
        prev_text = prev_para.text.strip()
        # 图片段本身可能没有文字；用图注前一段的文字锚点，缺失时退化为图注段自身定位
        img_page = _locate_page(prev_text) if prev_text else None
        cap_page = _locate_page(text)
        if img_page is None:
            img_page = cap_page  # 图片段无可锚定文本时，视作与图注同段落区域
        if img_page is None or cap_page is None or img_page != cap_page:
            fig_layout_ok = False
            break
        # ② 图序号与标题间距 ∈ [0.2, 0.5] 字符
        sep = m.group(2)
        body = m.group(3)
        if not body:
            # 只有"图N"没标题，间距无意义；按不合规处理
            fig_layout_ok = False
            break
        width = sum(SPACE_WIDTH.get(c, 0.5) for c in sep)
        if not (0.2 <= width <= 0.5):
            fig_layout_ok = False
            break
    if fig_layout_seen > 0 and fig_layout_ok and not page_locate_failed:
        results.append("+3：图片与图注不分页（同页）；图序号与标题间距0.2-0.5字符")
    else:
        results.append("+0：图片与图注不分页（同页）；图序号与标题间距0.2-0.5字符")

    # +3：文档内插入的图片均采用嵌入型
    # 严格对齐细则，仅踩细则写出的这一点（未写出的不额外约束）：
    #   · 文档内插入的图片均采用嵌入型：文档中每一张插入的图片都是"嵌入型"
    #     （办公软件里"嵌入型/嵌入到文字层"= 随文字排版，非浮动环绕）。
    #     至少要有 1 张图片，且全部为嵌入型才给分。
    #
    # 与办公软件对齐——图片在 OOXML 中的两种承载与"嵌入型"判定：
    #   · <w:drawing>：子节点 <wp:inline> = 嵌入型（Word/WPS 的"嵌入型"）；
    #                 <wp:anchor> = 浮动型（四周环绕/衬于文字等）
    #   · <w:pict>（旧式 VML）：v:shape 的 style 含 position:absolute = 浮动型，否则嵌入型
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    img_total = 0
    img_inline = 0
    for drawing in doc.element.body.iter(qn('w:drawing')):
        img_total += 1
        # 嵌入型：<wp:inline>；浮动型：<wp:anchor>
        if drawing.find(f'{{{WP_NS}}}inline') is not None:
            img_inline += 1
    for pict in doc.element.body.iter(qn('w:pict')):
        img_total += 1
        # 默认按嵌入型处理；若发现 v:shape 浮动样式（position:absolute），则不算嵌入型
        is_inline = True
        for shape in pict.iter():
            if shape.tag.split('}')[-1] == 'shape':
                style = (shape.get('style') or '').replace(' ', '').lower()
                if 'position:absolute' in style:
                    is_inline = False
                break
        if is_inline:
            img_inline += 1
    if img_total > 0 and img_inline == img_total:
        results.append("+3：文档内插入的图片均采用嵌入型")
    else:
        results.append("+0：文档内插入的图片均采用嵌入型")

    # +1 系列：各章节图片数量检查
    # 口径（按细则本意——细则只规定"该编号标题下方出现 N 张图片"）：
    #   1) 标题匹配：只确认段落编号（如 2.2 / 2.1.1）是否等于细则编号，标题文本不检查
    #   2) 章节区间：从匹配到的标题段下一段，到下一个"编号标题行"（形如 X.Y[.Z] 标题、
    #      单独成行）为止
    #   3) 区间内 <w:drawing> 数量恰好等于要求数即通过
    #
    # 细则编号 -> 要求图片数（标题文本仅作展示用途，不参与匹配）
    sections_required = {
        "2.1.1": ("物相特征", 1),
        "2.1.2": ("形貌与元素分布", 1),
        "2.2":   ("磁回收行为", 2),
        "2.3.1": ("接触时间影响", 1),
        "2.3.2": ("初始染料浓度影响", 1),
        "2.3.3": ("温度依赖性", 1),
        "2.3.4": ("溶液 pH 的影响", 1),
        "2.4":   ("吸附等温线分析", 1),
        "2.5":   ("动力学分析", 1),
        "2.6":   ("再生与基质耐受性", 1),
    }

    # 编号 + (编号与标题名之间的空白) + 标题名；空白允许半角/全角/NBSP/Tab，数量任意
    HEADING_RE = re.compile(r'^(\d+(?:\.\d+){1,2})([ \t　\xa0]+)(\S.*)$')

    def _heading_number(text):
        """取出"X.Y[.Z] 标题"的编号部分；不是标题行则返回 None。"""
        m = HEADING_RE.match(text)
        if not m:
            return None
        return m.group(1)

    # 1) 收集所有"编号标题行"（单独成行、长度 ≤ 30），作为章节边界
    section_index = []  # [(para_idx, number)]
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if len(text) > 30:
            continue
        num = _heading_number(text)
        if num is not None:
            section_index.append((i, num))

    # 2) 计算每个目标编号的图片数：本节标题下一段 → 下一个编号标题行
    section_image_count = {}
    for spec_num in sections_required:
        # 找到编号等于该细则编号的边界
        match_k = None
        for k, (idx, num) in enumerate(section_index):
            if num == spec_num:
                match_k = k
                break
        if match_k is None:
            section_image_count[spec_num] = None  # 没找到该编号
            continue
        start = section_index[match_k][0] + 1
        end = (section_index[match_k + 1][0]
               if match_k + 1 < len(section_index) else len(doc.paragraphs))
        cnt = 0
        for j in range(start, end):
            cnt += len(list(doc.paragraphs[j]._element.iter(qn('w:drawing'))))
        section_image_count[spec_num] = cnt

    # 3) 出具评分：区间内图片数恰好等于要求数即通过
    for num, (title_text, required) in sections_required.items():
        actual = section_image_count.get(num)
        label = f"{num} {title_text}"
        if actual == required:
            results.append(f"+1：\"{label}\"标题下方文本内包含{actual}张图片")
        else:
            shown = actual if actual is not None else 0
            results.append(f"+0：\"{label}\"标题下方文本内包含{shown}张图片")

    # +3：参考文献部分的标题文本"参考文献"："文献"两字为黑体、四号；"参考"两字为蓝色，整体居中显示
    # 严格对齐细则，仅踩细则写出的三点（未写出的不额外约束）：
    #   ① "文献"两字为黑体、四号：承载"文献"的 run，东亚字体（办公软件实际生效）为 黑体/SimHei，
    #        且字号为四号（14pt，容差 13.5–14.5）
    #   ② "参考"两字为蓝色：承载"参考"的 run，字体颜色（办公软件实际生效）为蓝色
    #   ③ 整体居中显示：标题段落对齐 == CENTER
    #
    # 与办公软件对齐：字体沿 w:rFonts/@w:eastAsia、字号沿 w:sz、颜色沿 w:color 的
    #   run→段落样式(含 basedOn 链)→docDefaults 继承链解析，即打开 Word/WPS 看到的效果；
    #   居中读 w:jc=center。
    def _resolve_effective_color(doc, run):
        """解析 run 实际生效的字体颜色，返回 (val_or_None, themeColor_or_None)。"""
        def _color_of(rPr):
            if rPr is None:
                return None
            return rPr.find(qn('w:color'))

        # 1) run 级
        c = _color_of(run._element.rPr)
        if c is not None:
            return (c.get(qn('w:val')), c.get(qn('w:themeColor')))
        # 2) 段落样式链
        styles_el = doc.styles.element
        style_id = None
        try:
            pPr = run._parent._p.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_id = pStyle.get(qn('w:val'))
        except Exception:
            style_id = None

        def _style_by_id(sid):
            for st in styles_el.findall(qn('w:style')):
                if st.get(qn('w:styleId')) == sid:
                    return st
            return None

        seen = set()
        while style_id and style_id not in seen:
            seen.add(style_id)
            st = _style_by_id(style_id)
            if st is None:
                break
            c = _color_of(st.find(qn('w:rPr')))
            if c is not None:
                return (c.get(qn('w:val')), c.get(qn('w:themeColor')))
            base = st.find(qn('w:basedOn'))
            style_id = base.get(qn('w:val')) if base is not None else None
        # 3) docDefaults
        dd = styles_el.find(qn('w:docDefaults'))
        if dd is not None:
            rpd = dd.find(qn('w:rPrDefault'))
            if rpd is not None:
                c = _color_of(rpd.find(qn('w:rPr')))
                if c is not None:
                    return (c.get(qn('w:val')), c.get(qn('w:themeColor')))
        return (None, None)

    def _run_is_blue(run):
        val, theme = _resolve_effective_color(doc, run)
        # 主题色蓝色系（Word 默认强调色/超链接蓝）
        if theme and ('accent1' in theme or 'accent5' in theme or theme == 'hyperlink'):
            return True
        # 显式 RGB：蓝分量明显大于红、绿
        if val and re.fullmatch(r'[0-9A-Fa-f]{6}', val):
            r = int(val[0:2], 16)
            g = int(val[2:4], 16)
            b = int(val[4:6], 16)
            if b >= 0x80 and b > r + 0x20 and b > g + 0x10:
                return True
        return False

    def _char_runs(para, ch):
        """返回承载字符串 ch 的 run 列表（按段落文本定位）"""
        text = para.text or ''
        pos = text.find(ch)
        if pos < 0:
            return []
        s, e = pos, pos + len(ch)
        cur = 0
        out = []
        for r in para.runs:
            t = r.text or ''
            rs, re_ = cur, cur + len(t)
            if rs < e and re_ > s:
                out.append(r)
            cur = re_
        return out

    ref_title_ok = False
    for para in doc.paragraphs:
        norm = para.text.strip().replace(' ', '').replace('　', '')
        if norm != "参考文献":
            continue
        ok = True
        # ③ 整体居中显示
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            ok = False
        # ① "文献"两字黑体、四号
        wenxian_runs = _char_runs(para, '文献') or _char_runs(para, '文')
        if not wenxian_runs:
            ok = False
        else:
            for r in wenxian_runs:
                ea = _resolve_effective_eastasia(doc, r)
                if not (ea and ('黑体' in ea or 'SimHei' in ea)):
                    ok = False
                    break
                # 四号 = 14pt（容差 13.5–14.5）
                if not (r.font.size and 13.5 <= r.font.size.pt <= 14.5):
                    ok = False
                    break
        # ② "参考"两字蓝色
        cankao_runs = _char_runs(para, '参考') or _char_runs(para, '参')
        if not cankao_runs or not all(_run_is_blue(r) for r in cankao_runs):
            ok = False
        ref_title_ok = ok
        break
    if ref_title_ok:
        results.append("+3：参考文献标题（\"文献\"黑体四号，\"参考\"蓝色，整体居中）")
    else:
        results.append("+0：参考文献标题（\"文献\"黑体四号，\"参考\"蓝色，整体居中）")

    # +3：参考文献条目编号用方括号进行编号，例如"[1]"，整体采用连续编号，编号与文献内容之间无多余空格
    # 严格对齐细则，仅踩细则写出的三点（未写出的不额外约束）：
    #   ① 条目编号用方括号进行编号（例如"[1]"）：参考文献区内每一条目段都以 [数字] 开头
    #   ② 整体采用连续编号：条目编号序列为连续的 1, 2, …, N
    #   ③ 编号与文献内容之间无多余空格：紧接 ] 之后到正文之间无"多余"空格
    #        （允许 0 或 1 个常规半角空格/NBSP 作正常分隔；≥2 个、或全角空格/Tab 即为多余）
    #
    # 与办公软件对齐：先定位"参考文献"标题段，其后的每个非空段落即办公软件里看到的
    #   一条参考文献；编号 [N] 为条目文本首部（非自动编号域），故按段落可见文本判定。
    REF_HDR_TITLE = '参考文献'
    REF_ENTRY_RE = re.compile(r'^\[(\d+)\](\s*)(.*)$')

    # 1) 定位参考文献区起始段
    ref_start = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == REF_HDR_TITLE:
            ref_start = i
            break

    ref_format_ok = False
    entries = []      # [(num, sep), ...]
    if ref_start >= 0:
        format_violation = False
        for j in range(ref_start + 1, len(doc.paragraphs)):
            text = doc.paragraphs[j].text.strip()
            if not text:
                continue
            m = REF_ENTRY_RE.match(text)
            if not m:
                # ① 参考文献区内的非空段落若不是 [N] 开头，说明有条目未用方括号编号 → 违规
                format_violation = True
                break
            entries.append((int(m.group(1)), m.group(2)))
        # ① 至少有一条条目
        if entries and not format_violation:
            nums = [e[0] for e in entries]
            # ② 连续编号 = [1..N]
            seq_ok = (nums == list(range(1, len(nums) + 1)))
            # ③ 编号与内容间无多余空格：每条 sep 只能是 0 或 1 个半角空格/NBSP
            sep_ok = True
            for _, sep in entries:
                if len(sep) > 1:
                    sep_ok = False
                    break
                if len(sep) == 1 and sep not in (' ', '\xa0'):
                    sep_ok = False
                    break
            if seq_ok and sep_ok:
                ref_format_ok = True
    if ref_format_ok:
        results.append(f"+3：参考文献编号格式（方括号，连续编号，编号与内容间无多余空格；共{len(entries)}条）")
    else:
        results.append("+0：参考文献编号格式（方括号，连续编号，编号与内容间无多余空格）")

    # +3：参考文献条目除编号字体采用Times New Roman、10号之外，中文文本统一采用宋体小四、整体悬挂缩进2字符
    # 严格对齐细则，仅踩细则写出的四点（未写出的不额外约束）。对参考文献区内每一条目：
    #   ① 编号字体采用 Times New Roman：承载编号 [N] 的 run，西文字体（办公软件实际生效）
    #        为 Times New Roman
    #   ② 编号 10号：承载编号 [N] 的 run 字号为10号（10pt，容差 9.5–10.5）
    #   ③ 中文文本统一采用宋体小四：条目中含中文字符的 run，其东亚字体为 宋体/SimSun/NSimSun，
    #        字号为小四（12pt，容差 11.5–12.5）；条目无中文 run 则该项无中文文本可约束
    #   ④ 整体悬挂缩进2字符：段落设悬挂缩进，缩进量为 2 字符
    #        （中文文本基准小四 12pt，2 字符=24pt，容差 20–28pt；或 hangingChars=2 字符）
    # 任一条目违反任一项即判 0；至少要有一条条目。
    #
    # 与办公软件对齐：西文/中文字体沿 w:ascii→hAnsi、w:eastAsia 继承链解析、字号 w:sz；
    #   悬挂缩进读段落 w:ind/@w:hanging（twips，绝对值）或 @w:hangingChars（1/100 字符），
    #   即打开 Word/WPS 看到的效果。
    REF_FORMAT_TITLE = '参考文献'
    REF_NUM_RE = re.compile(r'^\[(\d+)\]')

    # 1) 定位参考文献区
    ref_fmt_start = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == REF_FORMAT_TITLE:
            ref_fmt_start = i
            break

    def _hanging_indent(para):
        """返回 (hanging_pt, hanging_chars)：段落悬挂缩进的绝对磅值与字符数（无则为 None）。"""
        pPr = para._element.find(qn('w:pPr'))
        ind = pPr.find(qn('w:ind')) if pPr is not None else None
        if ind is None:
            return (None, None)
        hp = None
        h = ind.get(qn('w:hanging'))
        if h is not None and h.lstrip('-').isdigit():
            hp = abs(int(h)) / 20.0        # twips → pt
        hc = None
        hcs = ind.get(qn('w:hangingChars'))
        if hcs is not None and hcs.lstrip('-').isdigit():
            hc = abs(int(hcs)) / 100.0      # 1/100 字符 → 字符
        return (hp, hc)

    ref_format2_ok = False
    ref_entry_count = 0
    if ref_fmt_start >= 0:
        all_ok = True
        for j in range(ref_fmt_start + 1, len(doc.paragraphs)):
            para = doc.paragraphs[j]
            text = para.text.strip()
            if not text:
                continue
            bm = REF_NUM_RE.match(text)
            if not bm:
                continue
            ref_entry_count += 1

            # 准备 run 起止位置
            cur = 0
            spans = []
            for r in para.runs:
                t = r.text or ''
                spans.append((cur, cur + len(t), r))
                cur += len(t)

            # ① + ②：覆盖 [N] 区间的第一个 run（编号 run）
            num_end = bm.end()
            num_run = next((r for (rs, re_, r) in spans if rs < num_end and re_ > 0), None)
            if num_run is None:
                all_ok = False
                break
            # ① 编号字体 Times New Roman
            if _resolve_effective_ascii(doc, num_run) != 'Times New Roman':
                all_ok = False
                break
            # ② 编号 10号（10pt，容差 9.5–10.5）
            if not (num_run.font.size and 9.5 <= num_run.font.size.pt <= 10.5):
                all_ok = False
                break

            # ③ 中文文本统一宋体小四：仅约束含中文字符的 run
            for r in para.runs:
                t = r.text or ''
                if not re.search(r'[一-鿿]', t):
                    continue
                ea = _resolve_effective_eastasia(doc, r)
                if not (ea and any(k in ea for k in STRICT_SONGTI)):
                    all_ok = False
                    break
                if not (r.font.size and 11.5 <= r.font.size.pt <= 12.5):
                    all_ok = False
                    break
            if not all_ok:
                break

            # ④ 整体悬挂缩进 2 字符：hangingChars=2 字符，或 hanging 绝对量 ≈ 24pt（小四 2 字符）
            hp, hc = _hanging_indent(para)
            hanging_ok = False
            if hc is not None and 1.8 <= hc <= 2.2:
                hanging_ok = True
            elif hp is not None and 20 <= hp <= 28:
                hanging_ok = True
            if not hanging_ok:
                all_ok = False
                break
        ref_format2_ok = (all_ok and ref_entry_count > 0)
    if ref_format2_ok:
        results.append(f"+3：参考文献格式（编号TNR 10号；中文宋体小四；悬挂缩进2字符；共{ref_entry_count}条）")
    else:
        results.append("+0：参考文献格式（编号TNR 10号；中文宋体小四；悬挂缩进2字符）")

    # ========== 将 results 反解为结构化 dim2_items ==========
    # 逐项按追加顺序对齐 _MAX_DELTAS_ORDER：解析每条 "+N：说明" 的 N 与 rule。
    if len(results) != len(_MAX_DELTAS_ORDER):
        # 顺序表与实际条目不一致：视为脚本内部错误，交给外层兜底
        raise RuntimeError(
            f"dim2 项数不一致：results={len(results)} vs order={len(_MAX_DELTAS_ORDER)}"
        )
    dim2_items = []
    total_score = 0
    for raw, max_delta in zip(results, _MAX_DELTAS_ORDER):
        # 格式约定：'+N：' 前缀 + 规则说明；命中 N=max_delta，未命中 N=0
        head, _, rule = raw.partition("：")
        try:
            delta = int(head.replace("+", "").strip())
        except ValueError:
            delta = 0
        hit = delta > 0
        dim2_items.append({
            "rule": rule,
            "max_delta": max_delta,
            "delta": max_delta if hit else 0,
            "hit": hit,
            "detail": "",
        })
        total_score += (max_delta if hit else 0)

    result["dim2_items"] = dim2_items
    result["total_score"] = total_score
    result["max_score"] = MAX_SCORE_TOTAL
    return result


if __name__ == "__main__":
    # 本地调试入口：evaluate 只接收"脚本所在目录的路径"，由脚本自己在该目录里
    # 定位并打开被评估的 .docx。命令行第 1 个参数可覆盖目录路径；未指定时使用
    # 脚本自身所在目录。
    # 注：evaluate() 的接口约定就是"入参为目录路径"——路径始终由外部显式传入；
    #     这里的默认值仅为本地调试便利，不影响 evaluate 的接口约定。
    if len(sys.argv) >= 2:
        _dir = sys.argv[1]
    else:
        _dir = os.path.dirname(os.path.abspath(__file__))
    _out = evaluate(_dir)
    print(json.dumps(_out, ensure_ascii=False, indent=2))
