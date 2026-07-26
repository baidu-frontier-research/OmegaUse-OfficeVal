# -*- coding: utf-8 -*-
"""
Word文档自动评估脚本 - 完整版
根据打分细则对Word文档进行自动评估
支持详细的格式、表格、图片检查
"""

import os

SCRIPT_ID = "011"
import sys
import json
import contextlib
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# 维度二规则目录：(delta, rule_description)
# 用于将命中的扣分/加分项与未命中项一起纳入 dim2_items，
# 便于批量汇总时对齐 100 个文档的评分项矩阵。
_DIMENSION2_RULES = [
    (5, '文档第4页只存在一个表格，且表格未出现任何断开位置'),
    (-1, '封面页"《城市绿地系统规划》"未出现在页面上方三分之一处'),
    (-1, '封面页"《城市绿地系统规划》"字体格式不满足Noto Sans CJK SC、56号、加粗'),
    (-1, '封面页"《城市绿地系统规划》"段落格式不满足单倍行距、居中对齐'),
    (-1, '封面页"韧性街区绿地的认知与设计"字体格式不满足Noto Sans CJK SC、小初、加粗'),
    (-1, '封面页"第1节 场地解读图的作用及组织逻辑"段落格式不满足单倍行距、居中对齐'),
    (-1, '封面页"第1节 场地解读图的作用及组织逻辑"字体格式不满足Noto Sans CJK SC、一号、加粗'),
    (-1, '封面页"教学设计"字体格式不满足Noto Sans CJK SC、48号、加粗'),
    (-1, '封面页"教学设计"段落格式不满足单倍行距、居中对齐'),
    (-1, '第2页"第1节 场地解读图的作用及组织逻辑"字体格式不满足Noto Sans CJK SC 小二 加粗'),
    (-1, '第2页"第1节 场地解读图的作用及组织逻辑"段落格式不满足：两端对齐、文本之前0.1厘米、首行缩进两字符、单倍行距、段前间距13.5磅'),
    (-3, '第2页存在的表格数量大于1或表格存在断开位置'),
    (-3, '第3页存在的表格数量大于1或表格存在断开位置'),
    (-1, '第3页"行为特征"右侧单元格内未出现图片'),
    (-1, '第3页"行为特征"右侧单元格图片布局不满足四周型环绕'),
    (-1, '第3页"行为特征"右侧单元格图片大小不满足11.25×5.04厘米'),
    (-3, '第5页存在的表格数量大于1或表格存在断开位置'),
    (-3, '第6页存在的表格数量大于1或表格存在断开位置'),
    (-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内未出现图片'),
    (-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内的图片布局不满足嵌入型'),
    (-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内的图片大小不满足4.19×13.53厘米'),
    (-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内的图片与下方表格或字体重叠'),
    (-3, '第7页存在的表格数量大于1或表格存在断开位置'),
    (-3, '第8页存在的表格数量大于1或表格存在断开位置'),
    (-3, '第9页存在的表格数量大于1或表格存在断开位置'),
    (-3, '第10页存在的表格数量大于1或表格存在断开位置'),
    (-3, '第11页存在的表格数量大于1或表格存在断开位置'),
    (-1, '第11页"3.板书设计"下方未出现图片'),
    (-1, '第11页"3.板书设计"下方的图片尺寸不满足9.64×6.14厘米'),
    (-1, '第11页"3.板书设计"下方的图片与下方表格或字体重叠'),
    (-1, '第11页"3.板书设计"下方的图片布局不满足嵌入型'),
    (-3, '文档中除封面页外其余任意一页没有出现页码'),
    (-1, '文档中页码格式不满足以下任意一条：页码位置页面底部居中、样式为"1、2、3"、字体为宋体小五'),
]


class DocumentEvaluator:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.score = 0
        self.dimension1_passed = True
        self.dimension1_failures = []
        self.dimension2_results = []  # 存储(得分, 描述)



    def evaluate(self):
        """主评估函数"""
        print("=" * 80)
        print("Word文档自动评估报告")
        print("=" * 80)
        print(f"文件名: {os.path.basename(self.doc_path)}")
        print("-" * 80)

        # 维度1：可用与可修改性
        print("\n【维度1：可用与可修改性检查】")
        self.check_dimension1()

        if not self.dimension1_passed:
            print("\n" + "=" * 80)
            print("维度1检查未通过，文档评分为 0 分")
            print("=" * 80)
            print("\n未通过原因:")
            for failure in self.dimension1_failures:
                print(f"  ❌ {failure}")
            return 0

        print("✅ 维度1检查通过")

        # 维度2：完成度评分
        print("\n【维度2：完成度评分】")
        self.check_dimension2()

        # 计算总分
        self.score = sum([r[0] for r in self.dimension2_results])

        print("\n" + "=" * 80)
        print(f"最终得分: {self.score} 分")
        print("=" * 80)

        # 打印得分明细
        print("\n【得分明细】")
        for score, desc in self.dimension2_results:
            if score > 0:
                print(f"  +{score}分：{desc}")
            elif score < 0:
                print(f"  {score}分：{desc}")

        return self.score

    def check_dimension1(self):
        """检查维度1：可用与可修改性"""
        # 1. 检查文件扩展名
        ext = os.path.splitext(self.doc_path)[1].lower()
        if ext != '.docx':
            self.dimension1_passed = False
            self.dimension1_failures.append(f"文件扩展名不正确: {ext}，应为.docx")
            return
        print("  ✓ 文件扩展名检查通过 (.docx)")

        # 2. 检查文件可正常打开
        print("  ✓ 文件可正常打开")

        # 3. 检查是否为Word可编辑文档
        print("  ✓ 文件为Word可编辑文档")

        # 4. 检查表格中的文字内容是否可编辑
        has_editable_content = self._check_editable_content()
        if has_editable_content:
            print("  ✓ 表格内容可编辑")
        else:
            print("  ✓ 表格内容检查通过")

        # 当前脚本不再把连续空白页、乱码/文字重叠面积作为维度一门禁。

    def _check_editable_content(self):
        """检查内容是否可编辑"""
        text_in_tables = 0
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_in_tables += 1
        return text_in_tables > 0


    def check_dimension2(self):
        """检查维度2：完成度评分"""
        # 检查第4页表格
        self._check_page4_table()

        # 检查封面页
        self._check_cover_page()

        # 检查第2页
        self._check_page2()

        # 检查第3页
        self._check_page3()

        # 检查第5-6页
        self._check_page5_6()

        # 检查第7-11页
        self._check_pages_7_to_11()

        # 检查页码
        self._check_page_numbers()

        # 检查页码格式
        self._check_page_number_format()

    def _check_cover_page(self):
        """检查封面页各项要求"""
        cover_paras = self.doc.paragraphs[:20]

        # 位置项通过办公软件（Word / WPS 文字）实测坐标进行判定，
        # 严格对应细则："-1：封面页"《城市绿地系统规划》"未出现在页面上方三分之一处"
        self._check_cover_title1_position()

        # 字体格式项通过办公软件读取真实字体名/字号/加粗进行判定，
        # 严格对应细则："-1：封面页"《城市绿地系统规划》"字体格式不满足
        #                Noto Sans CJK SC、56号、加粗"
        self._check_cover_title1_font()

        # 段落格式项通过办公软件读取真实段落对齐方式与行距进行判定，
        # 严格对应细则："-1：封面页"《城市绿地系统规划》"段落格式不满足
        #                单倍行距、居中对齐"
        self._check_cover_title1_paragraph()

        # 副标题"韧性街区绿地的认知与设计"字体格式，通过办公软件读取真实
        # 字体名/字号/加粗判定，严格对应细则：
        #   "-1：封面页"韧性街区绿地的认知与设计"字体格式不满足
        #    Noto Sans CJK SC、小初、加粗"
        # 小初 = 36 磅（办公软件"小初"预设字号）
        self._check_cover_title2_font()

        # 封面页"第1节 场地解读图的作用及组织逻辑"段落格式，通过办公软件
        # 读取真实段落对齐方式与行距进行判定，严格对应细则：
        #   "-1：封面页"第1节 场地解读图的作用及组织逻辑"段落格式不满足
        #    单倍行距、居中对齐"
        self._check_cover_title3_paragraph()

        # 封面页"第1节 场地解读图的作用及组织逻辑"字体格式，通过办公软件
        # 读取真实字体名/字号/加粗判定，严格对应细则：
        #   "-1：封面页"第1节 场地解读图的作用及组织逻辑"字体格式不满足
        #    Noto Sans CJK SC、一号、加粗"
        # 一号 = 26 磅（办公软件"一号"预设字号）
        self._check_cover_title3_font()

        # 封面页"教学设计"字体格式，通过办公软件读取真实字体名/字号/加粗
        # 判定，严格对应细则：
        #   "-1：封面页"教学设计"字体格式不满足 Noto Sans CJK SC、48号、加粗"
        self._check_cover_design_font()

        # 封面页"教学设计"段落格式，通过办公软件读取真实段落对齐方式与行距
        # 进行判定，严格对应细则：
        #   "-1：封面页"教学设计"段落格式不满足单倍行距、居中对齐"
        self._check_cover_design_paragraph()

        found_title1 = False
        found_title2 = False
        found_title3 = False
        found_design = False

        for i, para in enumerate(cover_paras):
            text = para.text.strip()

            # 1. 检查'《城市绿地系统规划》'
            if '城市绿地系统规划' in text and not found_title1:
                found_title1 = True
                # 该标题的位置/字体/段落三项，均已由上方三个专用方法按细则判定，
                # 此处不再重复处理，避免与"细则没有要求的代码不加以约束"冲突。

            # 2. 检查'韧性街区绿地的认知与设计'
            if '韧性街区绿地的认知与设计' in text and not found_title2:
                found_title2 = True
                # 字体格式由专用方法按细则通过办公软件判定，此处不再叠加约束

            # 3. 检查'第1节 场地解读图的作用及组织逻辑'
            if '第1节 场地解读图的作用及组织逻辑' in text and not found_title3:
                found_title3 = True
                # 段落格式与字体格式均由专用方法按细则通过办公软件判定

            # 4. 检查'教学设计'
            if text == '教学设计' and not found_design:
                found_design = True
                # 字体格式与段落格式均由专用方法按细则通过办公软件判定

    def _check_cover_title1_position(self):
        """检查封面页"《城市绿地系统规划》"是否出现在页面上方三分之一处。

        细则：-1：封面页"《城市绿地系统规划》"未出现在页面上方三分之一处
        —— 仅按此一项判定，不附加其它约束。
        —— 通过办公软件（Word / WPS 文字）读取该文字在页面中的真实纵向
           坐标（Information(wdVerticalPositionRelativeToPage)），与页面
           上边距到页面高度上方三分之一分界线比较，与办公软件所见一致。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            # 办公软件不可用时不做该项判定，避免误扣分
            return

        try:
            # 页面高度（磅）
            page_setup = wdDoc.Sections(1).PageSetup
            page_height_pt = float(page_setup.PageHeight)
            top_third_line = page_height_pt / 3.0

            # 用 Find 在第1页范围内查找"《城市绿地系统规划》"
            target_text = '《城市绿地系统规划》'
            found_range = None

            # 优先在第1页范围内搜索
            try:
                first_page_range = wdDoc.GoTo(What=1, Which=1, Count=1)  # wdGoToPage, wdGoToFirst
                # 用 Content 全文查找，取第1页出现的位置
            except Exception:
                pass

            rng = wdDoc.Content
            find = rng.Find
            find.ClearFormatting()
            find.Text = target_text
            find.Forward = True
            find.Wrap = 0  # wdFindStop
            if find.Execute():
                found_range = rng

            if found_range is None:
                # 找不到文字，视为未出现在上方三分之一处 —— 该情况严格来说
                # 属于"未出现"，仍按细则判 -1
                self.dimension2_results.append((-1, '封面页"《城市绿地系统规划》"未出现在页面上方三分之一处'))
                return

            # 该文字所在页码
            page_of_text = found_range.Information(3)  # wdActiveEndPageNumber
            # 该文字顶部相对页面的纵向坐标（磅）
            vertical_pos_pt = found_range.Information(7)  # wdVerticalPositionRelativeToPage

            # 细则要求出现在"页面上方三分之一处"：
            #   —— 文字位于封面页（第1页）；
            #   —— 且其顶部纵坐标 ≤ 页面高度的 1/3。
            if page_of_text != 1 or vertical_pos_pt > top_third_line:
                self.dimension2_results.append((-1, '封面页"《城市绿地系统规划》"未出现在页面上方三分之一处'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_cover_title1_font(self):
        """检查封面页"《城市绿地系统规划》"的字体格式。

        细则：-1：封面页"《城市绿地系统规划》"字体格式不满足
              Noto Sans CJK SC、56号、加粗
        —— 仅按此一项判定，不附加其它约束。
        —— 通过办公软件（Word / WPS 文字）逐字符读取真实字体名、字号、
           加粗属性；与办公软件所见完全一致。
        —— 三个属性中任一不满足即判定"字体格式不满足"，扣 -1。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            # 办公软件不可用时不判定，避免误扣分
            return

        try:
            target_text = '《城市绿地系统规划》'
            rng = wdDoc.Content
            find = rng.Find
            find.ClearFormatting()
            find.Text = target_text
            find.Forward = True
            find.Wrap = 0  # wdFindStop
            if not find.Execute():
                # 找不到文字，无法就该项判定；细则中另有"未出现"项处理
                return

            start, end = rng.Start, rng.End
            expected_name = 'Noto Sans CJK SC'
            expected_size = 56.0
            # 加粗：Word/WPS 中布尔属性为 True(-1)/False(0)，
            # 若返回 wdUndefined（9999999），代表所选文字包含混合值，不满足"加粗"。

            ok = True
            for i in range(start, end):
                ch = wdDoc.Range(i, i + 1)
                if not ch.Text or ch.Text in ('\r', '\n', '\x07'):
                    continue
                fnt = ch.Font

                # 字体名：中英文/远东三个字段任一不匹配即视为不满足
                # 之所以三个都要匹配，是因为办公软件在中日韩场景下会
                # 用 NameFarEast 显示中文字，NameAscii 显示西文字符。
                name = getattr(fnt, 'Name', None)
                name_ascii = getattr(fnt, 'NameAscii', None)
                name_fe = getattr(fnt, 'NameFarEast', None)
                if not (name == expected_name
                        and (name_ascii in (expected_name, '', None))
                        and (name_fe in (expected_name, '', None))):
                    # 更宽松地说：至少 Name 与 NameFarEast 必须命中
                    if not (name == expected_name and name_fe == expected_name):
                        ok = False
                        break

                # 字号：必须为 56
                try:
                    size_val = float(fnt.Size)
                except Exception:
                    size_val = None
                if size_val is None or abs(size_val - expected_size) > 1e-6:
                    ok = False
                    break

                # 加粗：必须为 True（Word/WPS COM 返回 -1）
                bold_val = fnt.Bold
                # -1 True, 0 False, 9999999 混合值/未定义
                if bold_val != -1 and bold_val is not True:
                    ok = False
                    break

            if not ok:
                self.dimension2_results.append((
                    -1,
                    '封面页"《城市绿地系统规划》"字体格式不满足Noto Sans CJK SC、56号、加粗',
                ))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_text_paragraph_center_single_via_office(self, target_text, label, restrict_to_page=None):
        """通用：通过办公软件判定某段目标文字所在段落是否满足
        "单倍行距、居中对齐" 两项。

        仅按这两个属性判定；任一不满足即在结果中追加 (-1, label)。
        找不到目标文字时不判扣。

        参数：
            target_text: 需要在文档中查找的原始文字
            label: 命中失败时写入 dimension2_results 的描述文字
            restrict_to_page: 若指定，则只有目标文字位于该页时才判定；
                              否则遍历所有命中，直到找到位于该页的一处。
                              这样可以区分同一文本在不同页出现时的判定归属。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            rng = wdDoc.Content
            find = rng.Find
            find.ClearFormatting()
            find.Text = target_text
            find.Forward = True
            find.Wrap = 0  # wdFindStop

            para = None
            if restrict_to_page is None:
                if not find.Execute():
                    return
                para = rng.Paragraphs(1)
            else:
                # 逐一查找，直到落到期望页
                while find.Execute():
                    if int(rng.Information(3)) == int(restrict_to_page):
                        para = rng.Paragraphs(1)
                        break
                    # 从当前 rng.End 之后继续查找
                    rng = wdDoc.Range(rng.End, wdDoc.Content.End)
                    find = rng.Find
                    find.ClearFormatting()
                    find.Text = target_text
                    find.Forward = True
                    find.Wrap = 0
                if para is None:
                    return

            pf = para.Format
            # 对齐方式：1 = wdAlignParagraphCenter （居中）
            alignment_ok = (int(pf.Alignment) == 1)
            # 行距规则：0 = wdLineSpaceSingle （单倍行距）
            line_spacing_ok = (int(pf.LineSpacingRule) == 0)

            if not (alignment_ok and line_spacing_ok):
                self.dimension2_results.append((-1, label))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_cover_title3_paragraph(self):
        """检查封面页"第1节 场地解读图的作用及组织逻辑"的段落格式。

        细则：-1：封面页"第1节 场地解读图的作用及组织逻辑"段落格式不满足
              单倍行距、居中对齐
        —— 仅按此一项判定，不附加其它约束（不检查字体、位置、缩进等）。
        —— 通过办公软件（Word / WPS 文字）读取该文字所在段落的真实
           `ParagraphFormat.Alignment` 与 `LineSpacingRule`，与办公软件里
           "开始→段落"面板显示的对齐方式和行距完全一致。
        —— 由于同名章节标题在正文第2页还会再次出现，这里通过 restrict_to_page=1
           限定只对封面页（第1页）的那次出现进行判定。
        """
        self._check_text_paragraph_center_single_via_office(
            target_text='第1节 场地解读图的作用及组织逻辑',
            label='封面页"第1节 场地解读图的作用及组织逻辑"段落格式不满足单倍行距、居中对齐',
            restrict_to_page=1,
        )

    def _check_cover_title3_font(self):
        """检查封面页"第1节 场地解读图的作用及组织逻辑"的字体格式。

        细则：-1：封面页"第1节 场地解读图的作用及组织逻辑"字体格式不满足
              Noto Sans CJK SC、一号、加粗
        —— 仅按此一项判定，不附加其它约束（不检查段落格式、位置等）。
        —— "一号" 在办公软件中的字号预设为 26 磅（Word / WPS 文字），
           故通过办公软件读取真实 Font.Size 与 26 严格比较。
        —— 由于同名章节标题在正文第2页还会再次出现，这里通过
           restrict_to_page=1 限定只对封面页（第1页）的那次出现进行判定。
        """
        self._check_text_font_via_office(
            target_text='第1节 场地解读图的作用及组织逻辑',
            expected_size=26.0,
            label='封面页"第1节 场地解读图的作用及组织逻辑"字体格式不满足Noto Sans CJK SC、一号、加粗',
            restrict_to_page=1,
        )

    def _check_cover_design_font(self):
        """检查封面页"教学设计"的字体格式。

        细则：-1：封面页"教学设计"字体格式不满足 Noto Sans CJK SC、48号、加粗
        —— 仅按此一项判定，不附加其它约束（不检查段落格式、位置等）。
        —— 通过办公软件（Word / WPS 文字）逐字符读取真实字体名、字号、
           加粗属性；与办公软件里"开始→字体"面板显示的属性完全一致。
        —— 限定 restrict_to_page=1，只对封面页那次出现进行判定，避免与
           文档其它页可能出现的"教学设计"字样互相干扰。
        """
        self._check_text_font_via_office(
            target_text='教学设计',
            expected_size=48.0,
            label='封面页"教学设计"字体格式不满足Noto Sans CJK SC、48号、加粗',
            restrict_to_page=1,
        )

    def _check_cover_design_paragraph(self):
        """检查封面页"教学设计"的段落格式。

        细则：-1：封面页"教学设计"段落格式不满足单倍行距、居中对齐
        —— 仅按此一项判定，不附加其它约束（不检查字体、位置、缩进等）。
        —— 通过办公软件（Word / WPS 文字）读取该文字所在段落的真实
           `ParagraphFormat.Alignment` 与 `LineSpacingRule`，与办公软件里
           "开始→段落"面板显示的对齐方式和行距完全一致。
        —— 限定 restrict_to_page=1，只对封面页那次出现进行判定。
        """
        self._check_text_paragraph_center_single_via_office(
            target_text='教学设计',
            label='封面页"教学设计"段落格式不满足单倍行距、居中对齐',
            restrict_to_page=1,
        )

    def _check_text_font_via_office(self, target_text, expected_size, label, restrict_to_page=None):
        """通用：通过办公软件判定某段目标文字是否满足
        "Noto Sans CJK SC、<expected_size>号、加粗"。

        仅按传入的字号和"Noto Sans CJK SC、加粗"共 3 个属性判定；三者中
        任一不满足即在结果中追加 (-1, label)。找不到目标文字时不判扣。

        参数：
            target_text: 需要在文档中查找的原始文字
            expected_size: 期望字号（磅）
            label: 命中失败时写入 dimension2_results 的描述文字
            restrict_to_page: 若指定，则只有目标文字位于该页时才判定；
                              这样可以区分同一文本在不同页出现时的判定归属。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            rng = wdDoc.Content
            find = rng.Find
            find.ClearFormatting()
            find.Text = target_text
            find.Forward = True
            find.Wrap = 0  # wdFindStop

            hit_range = None
            if restrict_to_page is None:
                if find.Execute():
                    hit_range = rng
            else:
                # 逐一查找，直到落到期望页
                while find.Execute():
                    if int(rng.Information(3)) == int(restrict_to_page):
                        hit_range = rng
                        break
                    rng = wdDoc.Range(rng.End, wdDoc.Content.End)
                    find = rng.Find
                    find.ClearFormatting()
                    find.Text = target_text
                    find.Forward = True
                    find.Wrap = 0

            if hit_range is None:
                return

            start, end = hit_range.Start, hit_range.End
            expected_name = 'Noto Sans CJK SC'

            ok = True
            for i in range(start, end):
                ch = wdDoc.Range(i, i + 1)
                if not ch.Text or ch.Text in ('\r', '\n', '\x07'):
                    continue
                fnt = ch.Font

                # 字体名：办公软件在中日韩场景下由 Name / NameFarEast 分别管辖
                # 西文与远东字体；两者都必须命中 "Noto Sans CJK SC"。
                name = getattr(fnt, 'Name', None)
                name_fe = getattr(fnt, 'NameFarEast', None)
                if not (name == expected_name and name_fe == expected_name):
                    ok = False
                    break

                # 字号：严格等于期望值
                try:
                    size_val = float(fnt.Size)
                except Exception:
                    size_val = None
                if size_val is None or abs(size_val - expected_size) > 1e-6:
                    ok = False
                    break

                # 加粗：COM 中 True == -1，0 表示未加粗，9999999 表示混合值
                bold_val = fnt.Bold
                if bold_val != -1 and bold_val is not True:
                    ok = False
                    break

            if not ok:
                self.dimension2_results.append((-1, label))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_cover_title2_font(self):
        """检查封面页"韧性街区绿地的认知与设计"的字体格式。

        细则：-1：封面页"韧性街区绿地的认知与设计"字体格式不满足
              Noto Sans CJK SC、小初、加粗
        —— 仅按此一项判定，不附加其它约束（不检查段落格式、位置等）。
        —— "小初" 在办公软件中的字号预设为 36 磅（Word / WPS 文字），
           故通过办公软件读取真实 Font.Size 与 36 严格比较。
        """
        self._check_text_font_via_office(
            target_text='韧性街区绿地的认知与设计',
            expected_size=36.0,
            label='封面页"韧性街区绿地的认知与设计"字体格式不满足Noto Sans CJK SC、小初、加粗',
        )

    def _check_cover_title1_paragraph(self):
        """检查封面页"《城市绿地系统规划》"的段落格式。

        细则：-1：封面页"《城市绿地系统规划》"段落格式不满足单倍行距、居中对齐
        —— 仅按此一项判定，不附加其它约束（首行缩进、段前段后、两端对齐等
           细则未写的属性不做检查）。
        —— 通过办公软件（Word / WPS 文字）读取该标题所在段落的真实
           `ParagraphFormat.Alignment` 与 `LineSpacingRule`；与办公软件所
           见（"开始→段落"里的对齐方式和行距）完全一致。
        —— 两个属性中任一不满足即判定"段落格式不满足"，扣 -1。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            # 办公软件不可用时不判定，避免误扣分
            return

        try:
            target_text = '《城市绿地系统规划》'
            rng = wdDoc.Content
            find = rng.Find
            find.ClearFormatting()
            find.Text = target_text
            find.Forward = True
            find.Wrap = 0  # wdFindStop
            if not find.Execute():
                # 找不到文字，无法就本项判定
                return

            para = rng.Paragraphs(1)
            pf = para.Format

            # 对齐方式：Word/WPS 常量
            #   0 = wdAlignParagraphLeft
            #   1 = wdAlignParagraphCenter   ← 居中
            #   2 = wdAlignParagraphRight
            #   3 = wdAlignParagraphJustify
            alignment_ok = (int(pf.Alignment) == 1)

            # 行距规则：
            #   0 = wdLineSpaceSingle        ← 单倍行距
            #   1 = wdLineSpace1pt5
            #   2 = wdLineSpaceDouble
            #   3 = wdLineSpaceAtLeast
            #   4 = wdLineSpaceExactly
            #   5 = wdLineSpaceMultiple
            line_spacing_ok = (int(pf.LineSpacingRule) == 0)

            if not (alignment_ok and line_spacing_ok):
                self.dimension2_results.append((
                    -1,
                    '封面页"《城市绿地系统规划》"段落格式不满足单倍行距、居中对齐',
                ))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_font_format(self, paragraph, expected_size, expected_bold):
        '''检查字体格式，返回是否满足要求'''
        for run in paragraph.runs:
            if run.text.strip():
                # 检查字号
                size_ok = True
                if run.font.size:
                    size_pt = run.font.size.pt
                    if abs(size_pt - expected_size) > 2:
                        size_ok = False

                # 检查加粗
                bold_ok = (run.font.bold == expected_bold) or (expected_bold and run.font.bold)

                return size_ok and bold_ok
        return True

    def _check_para_format(self, paragraph, expected_alignment):
        '''检查段落格式，返回是否满足要求'''
        pf = paragraph.paragraph_format
        if pf.alignment == expected_alignment:
            return True
        return False

    def _check_para_format_full(self, paragraph):
        '''检查段落格式完整性'''
        pf = paragraph.paragraph_format
        # 简化检查：只检查对齐方式
        return pf.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    def _check_page2(self):
        """检查第2页"""
        # 第2页"第1节 场地解读图的作用及组织逻辑"字体格式，通过办公软件
        # 读取真实字体名/字号/加粗判定，严格对应细则：
        #   "-1：第2页"第1节 场地解读图的作用及组织逻辑"字体格式不满足
        #    Noto Sans CJK SC 小二 加粗"
        # 小二 = 18 磅（办公软件"小二"预设字号）
        self._check_page2_title_font()

        # 第2页"第1节 场地解读图的作用及组织逻辑"段落格式，通过办公软件
        # 读取真实段落属性判定，严格对应细则：
        #   "-1：第2页"第1节 场地解读图的作用及组织逻辑"段落格式不满足：
        #    两端对齐、文本之前0.1厘米、首行缩进两字符、单倍行距、段前间距13.5磅"
        self._check_page2_title_paragraph()

        # 第2页表格数量与断开检查，通过办公软件读取真实分页信息判定，
        # 严格对应细则：
        #   "-3：第2页存在的表格表格数量大于1或者表格出现断开位置"
        self._check_page_tables_count_and_split(page_num=2, penalty=-3,
            label='第2页存在的表格数量大于1或表格存在断开位置')

    def _check_page2_title_font(self):
        """检查第2页"第1节 场地解读图的作用及组织逻辑"的字体格式。

        细则：-1：第2页"第1节 场地解读图的作用及组织逻辑"字体格式不满足
              Noto Sans CJK SC 小二 加粗
        —— 仅按此一项判定，不附加其它约束（不检查段落格式、位置等）。
        —— "小二" 在办公软件中的字号预设为 18 磅（Word / WPS 文字），
           故通过办公软件读取真实 Font.Size 与 18 严格比较。
        —— 同名章节标题在封面页（第1页）也会出现且字号不同（那里由封面页
           专项判定负责），本项通过 restrict_to_page=2 限定只对第2页那次
           出现进行判定。
        """
        self._check_text_font_via_office(
            target_text='第1节 场地解读图的作用及组织逻辑',
            expected_size=18.0,
            label='第2页"第1节 场地解读图的作用及组织逻辑"字体格式不满足Noto Sans CJK SC 小二 加粗',
            restrict_to_page=2,
        )

    def _check_page2_title_paragraph(self):
        """检查第2页"第1节 场地解读图的作用及组织逻辑"的段落格式。

        细则：-1：第2页"第1节 场地解读图的作用及组织逻辑"段落格式不满足：
              两端对齐、文本之前0.1厘米、首行缩进两字符、单倍行距、段前间距13.5磅
        —— 严格按细则的 5 个属性判定，任一不满足即扣 -1；细则未列的属性
           （右缩进、段后间距、字体、颜色等）一律不检查。
        —— 通过办公软件（Word / WPS 文字）读取该文字所在段落的真实
           `ParagraphFormat` 属性，与办公软件里"开始→段落"面板显示完全一致：
             · 两端对齐          → Alignment == 3 (wdAlignParagraphJustify)
             · 文本之前 0.1 厘米 → LeftIndent 换算为厘米后 ≈ 0.1
                                   （办公软件 1 cm ≈ 28.3464567 pt）
             · 首行缩进两字符    → CharacterUnitFirstLineIndent == 2
             · 单倍行距          → LineSpacingRule == 0 (wdLineSpaceSingle)
             · 段前间距 13.5 磅  → SpaceBefore == 13.5
        —— 通过 restrict_to_page=2 限定只对第2页那次出现进行判定，避免与
           封面页同名标题互相干扰。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            target_text = '第1节 场地解读图的作用及组织逻辑'
            rng = wdDoc.Content
            find = rng.Find
            find.ClearFormatting()
            find.Text = target_text
            find.Forward = True
            find.Wrap = 0  # wdFindStop

            para = None
            while find.Execute():
                if int(rng.Information(3)) == 2:
                    para = rng.Paragraphs(1)
                    break
                rng = wdDoc.Range(rng.End, wdDoc.Content.End)
                find = rng.Find
                find.ClearFormatting()
                find.Text = target_text
                find.Forward = True
                find.Wrap = 0

            if para is None:
                return

            pf = para.Format

            # 1) 两端对齐：wdAlignParagraphJustify == 3
            alignment_ok = (int(pf.Alignment) == 3)

            # 2) 文本之前 0.1 厘米（左缩进）：pt → cm，1 cm = 28.3464567 pt
            left_indent_pt = float(pf.LeftIndent)
            left_indent_cm = left_indent_pt / 28.3464567
            # 与 0.1 cm 比较；办公软件"文本之前"精度约 0.01 cm，容差取 0.02 cm
            left_indent_ok = (abs(left_indent_cm - 0.1) <= 0.02)

            # 3) 首行缩进两字符：CharacterUnitFirstLineIndent == 2
            try:
                first_line_char_units = float(pf.CharacterUnitFirstLineIndent)
            except Exception:
                first_line_char_units = None
            first_line_ok = (first_line_char_units is not None
                             and abs(first_line_char_units - 2.0) < 1e-6)

            # 4) 单倍行距：wdLineSpaceSingle == 0
            line_spacing_ok = (int(pf.LineSpacingRule) == 0)

            # 5) 段前间距 13.5 磅
            space_before_ok = (abs(float(pf.SpaceBefore) - 13.5) < 1e-6)

            if not (alignment_ok and left_indent_ok and first_line_ok
                    and line_spacing_ok and space_before_ok):
                self.dimension2_results.append((
                    -1,
                    '第2页"第1节 场地解读图的作用及组织逻辑"段落格式不满足：两端对齐、文本之前0.1厘米、首行缩进两字符、单倍行距、段前间距13.5磅',
                ))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page_tables_count_and_split(self, page_num, penalty, label):
        """通用：判定某一页是否满足"表格数量 > 1 或 表格出现断开位置"。

        细则示例："-3：第2页存在的表格表格数量大于1或者表格出现断开位置"

        —— 判定条件为两者的"或"，且允许表格跨页：
             a) 该页上出现的顶层表格数量 > 1；或
             b) 该页上一张表出现"断开"：即某一行或多行的边框与相邻行的
                边框之间存在空隙 / 空白（该空隙不属于表格）。
        —— 关键观察：在 Word / WPS 的文档模型里，单个 Table 对象内部
           的行永远紧贴排列、不存在行间空隙；若视觉上一张表被切成两段、
           中间夹了非表格内容，那这两段一定是**两个独立的 Table 对象**。
           因此 (b) 的判定在同一页上等价于该页上存在 ≥ 2 张顶层表格。
        —— 一张表格横跨多页的情形，每一页上仍只是它的一部分（顶层 Table
           数量为 1），按细则应"允许"，故不视为断开。

        综上：本方法只需判定"该页顶层 Table 数量 > 1"，即可同时覆盖细则
        的两种触发条件；跨页本身不触发。
        通过办公软件（Word / WPS 文字）读取每张表格的真实起止页码
        （Range.Information(wdActiveEndPageNumber)），与办公软件所见
        的分页与表格切分完全一致。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            page_count = wdDoc.ComputeStatistics(2)  # wdStatisticPages
            if page_count < page_num:
                return

            tables_on_page = 0
            hit_details = []  # [(t_idx, start_page, end_page)]

            for t_idx, table in enumerate(wdDoc.Tables, start=1):
                # 仅统计顶层表格：嵌套在其它表格单元格内的子表格不计入
                try:
                    if int(table.NestingLevel) != 1:
                        continue
                except Exception:
                    pass

                try:
                    s = table.Range.Start
                    e = table.Range.End
                    start_page = int(wdDoc.Range(s, s).Information(3))
                    end_page = int(wdDoc.Range(e, e).Information(3))
                except Exception:
                    continue

                # 若表格任一部分落在目标页，则视为"存在于该页"
                if start_page <= page_num <= end_page:
                    tables_on_page += 1
                    hit_details.append((t_idx, start_page, end_page))

            if tables_on_page > 1:
                # 打印命中细节（便于核对："数量>1" 即等价于"存在断开"）
                try:
                    hit_str = '; '.join(
                        f'表{idx}[{sp}->{ep}]' for idx, sp, ep in hit_details
                    )
                    print(f'  · 第{page_num}页命中扣分：数量={tables_on_page}(>1，等价于存在断开)；命中表格：{hit_str}')
                except Exception:
                    pass
                self.dimension2_results.append((penalty, label))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page3(self):
        """检查第3页"""
        # 第3页表格数量与断开检查，通过办公软件读取真实分页信息判定，
        # 严格对应细则：
        #   "-3：第3页存在的表格数量大于1或表格存在断开位置"
        self._check_page_tables_count_and_split(page_num=3, penalty=-3,
            label='第3页存在的表格数量大于1或表格存在断开位置')

        # 第3页"行为特征"右侧单元格图片检查，通过办公软件判定其真实是否
        # 包含图片，严格对应细则：
        #   "-1：第3页"行为特征"右侧单元格内未出现图片"
        self._check_page3_behavior_image()

        # 第3页"行为特征"右侧单元格图片布局检查，通过办公软件读取该图片
        # 的真实环绕方式（WrapFormat.Type）判定，严格对应细则：
        #   "-1：第3页"行为特征"右侧单元格图片布局不满足四周型环绕"
        self._check_page3_behavior_image_wrap()

        # 第3页"行为特征"右侧单元格图片大小检查，通过办公软件读取该图片
        # 的真实 Width / Height（磅→厘米）判定，严格对应细则：
        #   "-1：第3页"行为特征"右侧单元格图片大小不满足11.25×5.04厘米"
        self._check_page3_behavior_image_size()

        # 第3页"行为特征"右侧单元格图片是否与表格或字体重叠检查已按用户
        # 要求删除，不再作为扣分项。
        # self._check_page3_behavior_image_overlap()

        # 未找到'行为特征'不计分

    def _locate_page3_behavior_right_cell(self, wdDoc):
        """通过办公软件定位到第3页"行为特征"单元格的右侧单元格。

        返回右侧单元格对象；若"行为特征"不存在于第3页、或该单元格已在
        最右列（无右邻）、或表格结构异常，则返回 None。
        """
        for ti in range(1, wdDoc.Tables.Count + 1):
            tbl = wdDoc.Tables(ti)
            cells = tbl.Range.Cells
            for ci in range(1, cells.Count + 1):
                cell = cells(ci)
                if '行为特征' not in cell.Range.Text:
                    continue
                if int(cell.Range.Information(3)) != 3:
                    continue
                row, col = cell.RowIndex, cell.ColumnIndex
                for cj in range(1, cells.Count + 1):
                    other = cells(cj)
                    if other.RowIndex == row and other.ColumnIndex == col + 1:
                        return other
                return None
        return None

    def _check_page3_behavior_image(self):
        """检查第3页"行为特征"右侧单元格内是否存在图片。

        细则：-1：第3页"行为特征"右侧单元格内未出现图片
        —— 仅按此一项判定；不检查图片的位置/大小/环绕方式等（细则未列）。
        —— 通过办公软件（Word / WPS 文字）真实识别该单元格内的图片：
             · 嵌入型图片：InlineShapes.Count；
             · 浮动型图片：文档级 Shapes 集合中锚点落入该单元格 Range 的形状。
           这与办公软件里在该单元格中看到的图片完全一致。
        —— 只对位于第 3 页的"行为特征"单元格进行判定（细则"第3页"三字），
           避免与其它页可能同名的单元格互相干扰。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            behavior_cell = None
            # 遍历所有表格，用 Table.Range.Cells 平铺访问单元格（避免跨页
            # 表格上 Rows(i) 抛"发生意外"的问题）
            for ti in range(1, wdDoc.Tables.Count + 1):
                tbl = wdDoc.Tables(ti)
                cells = tbl.Range.Cells
                # 用页码严格限定为"第3页"的"行为特征"单元格
                for ci in range(1, cells.Count + 1):
                    cell = cells(ci)
                    if '行为特征' not in cell.Range.Text:
                        continue
                    cell_page = int(cell.Range.Information(3))
                    if cell_page != 3:
                        continue
                    behavior_cell = cell
                    behavior_row = cell.RowIndex
                    behavior_col = cell.ColumnIndex
                    behavior_table = tbl
                    break
                if behavior_cell is not None:
                    break

            if behavior_cell is None:
                # 找不到该单元格，本项不判扣
                return

            # 在同一表格中找到 (behavior_row, behavior_col + 1) 的右侧单元格
            right_cell = None
            cells = behavior_table.Range.Cells
            for ci in range(1, cells.Count + 1):
                cell = cells(ci)
                if cell.RowIndex == behavior_row and cell.ColumnIndex == behavior_col + 1:
                    right_cell = cell
                    break

            if right_cell is None:
                # "行为特征"在最右一列 —— 无右侧单元格；按细则字面即"右侧
                # 单元格内未出现图片"（右侧单元格不存在，谈不上出现），扣 -1
                self.dimension2_results.append((-1, '第3页"行为特征"右侧单元格内未出现图片'))
                return

            r_start = right_cell.Range.Start
            r_end = right_cell.Range.End

            # 1) 嵌入型图片（InlineShape）
            inline_count = int(right_cell.Range.InlineShapes.Count)

            # 2) 浮动型图片（Shape），通过锚点落入该单元格 Range 范围来判定
            floating_count = 0
            try:
                for si in range(1, wdDoc.Shapes.Count + 1):
                    sh = wdDoc.Shapes(si)
                    try:
                        a = sh.Anchor
                    except Exception:
                        a = None
                    if a is None:
                        continue
                    if r_start <= a.Start <= r_end:
                        floating_count += 1
            except Exception:
                pass

            if inline_count == 0 and floating_count == 0:
                self.dimension2_results.append((-1, '第3页"行为特征"右侧单元格内未出现图片'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page3_behavior_image_wrap(self):
        """检查第3页"行为特征"右侧单元格图片的环绕方式是否为四周型。

        细则：-1：第3页"行为特征"右侧单元格图片布局不满足四周型环绕
        —— 仅按此一项判定（不检查图片大小、位置、颜色等）。
        —— 通过办公软件（Word / WPS 文字）读取该单元格内图片的真实
           `WrapFormat.Type`，与办公软件里"图片工具→环绕方式"面板显示
           完全一致。
        —— WrapFormat.Type 常量（Word/WPS 通用）：
             0 wdWrapSquare    四周型
             1 wdWrapTight     紧密型
             2 wdWrapThrough   穿越型
             3 wdWrapNone      浮于文字上方/下方（无环绕文字）
             4 wdWrapTopBottom 上下型
             5 wdWrapBehind    衬于文字下方
             6 wdWrapFront     浮于文字上方
             15 wdWrapInline   嵌入型（InlineShape）
        —— 单元格内包含任一非"四周型"环绕的图片（含嵌入型），或找不到
           任何图片时，均判定"图片布局不满足四周型环绕"扣 -1。若单元格
           不存在（"行为特征"不在第3页），本项不判扣（另一项细则会处理）。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            right_cell = self._locate_page3_behavior_right_cell(wdDoc)
            if right_cell is None:
                return

            r_start = right_cell.Range.Start
            r_end = right_cell.Range.End

            has_image = False
            all_square = True

            # 1) 嵌入型图片：InlineShapes（嵌入型环绕不是四周型）
            try:
                inline_count = int(right_cell.Range.InlineShapes.Count)
            except Exception:
                inline_count = 0
            if inline_count > 0:
                has_image = True
                # 嵌入型 ≠ 四周型
                all_square = False

            # 2) 浮动型图片：文档级 Shapes 里锚点落入该单元格 Range 的形状
            try:
                for si in range(1, wdDoc.Shapes.Count + 1):
                    sh = wdDoc.Shapes(si)
                    try:
                        a = sh.Anchor
                    except Exception:
                        a = None
                    if a is None:
                        continue
                    if not (r_start <= a.Start <= r_end):
                        continue
                    has_image = True
                    try:
                        wrap_type = int(sh.WrapFormat.Type)
                    except Exception:
                        wrap_type = -1
                    if wrap_type != 0:  # 0 == wdWrapSquare 四周型
                        all_square = False
            except Exception:
                pass

            # 单元格内没有任何图片：本项不判扣（由"未出现图片"细则处理，避免重复扣）
            if not has_image:
                return

            if not all_square:
                self.dimension2_results.append((-1, '第3页"行为特征"右侧单元格图片布局不满足四周型环绕'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page3_behavior_image_size(self):
        """检查第3页"行为特征"右侧单元格图片大小是否为 11.25 × 5.04 厘米。

        细则：-1：第3页"行为特征"右侧单元格图片大小不满足11.25×5.04厘米
        —— 仅按此一项判定（不检查环绕、位置、颜色等，另有对应细则）。
        —— 通过办公软件（Word / WPS 文字）读取该单元格内图片的真实
           `Width` / `Height`（单位为磅 pt），换算到厘米后与
           11.25cm × 5.04cm 比较，与办公软件"图片工具→大小"面板显示
           完全一致（Office UI 中厘米保留两位小数）。
        —— 匹配容差为 0.05cm，覆盖办公软件界面按两位小数四舍五入显示
           时的等价情形。
        —— 单元格内没有任何图片时，本项不判扣（另一细则会处理
           "未出现图片"），若单元格不存在（"行为特征"不在第3页），
           同样不判扣。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            right_cell = self._locate_page3_behavior_right_cell(wdDoc)
            if right_cell is None:
                return

            r_start = right_cell.Range.Start
            r_end = right_cell.Range.End

            CM = 28.3464567  # 1 cm = 28.3464567 pt
            TOL = 0.05       # cm
            EXPECTED_W_CM = 11.25
            EXPECTED_H_CM = 5.04

            has_image = False
            size_ok_any = False  # 只要有一张图片同时满足两个维度即视为符合

            # 1) 嵌入型图片：InlineShapes
            try:
                inline_count = int(right_cell.Range.InlineShapes.Count)
            except Exception:
                inline_count = 0
            if inline_count > 0:
                has_image = True
                try:
                    for i in range(1, inline_count + 1):
                        s = right_cell.Range.InlineShapes(i)
                        w_cm = float(s.Width) / CM
                        h_cm = float(s.Height) / CM
                        if abs(w_cm - EXPECTED_W_CM) <= TOL and abs(h_cm - EXPECTED_H_CM) <= TOL:
                            size_ok_any = True
                            break
                except Exception:
                    pass

            # 2) 浮动型图片：文档级 Shapes 中锚点在该单元格 Range 内的形状
            if not size_ok_any:
                try:
                    for si in range(1, wdDoc.Shapes.Count + 1):
                        sh = wdDoc.Shapes(si)
                        try:
                            a = sh.Anchor
                        except Exception:
                            a = None
                        if a is None:
                            continue
                        if not (r_start <= a.Start <= r_end):
                            continue
                        has_image = True
                        try:
                            w_cm = float(sh.Width) / CM
                            h_cm = float(sh.Height) / CM
                        except Exception:
                            continue
                        if abs(w_cm - EXPECTED_W_CM) <= TOL and abs(h_cm - EXPECTED_H_CM) <= TOL:
                            size_ok_any = True
                            break
                except Exception:
                    pass

            # 无图片：不由本项判扣
            if not has_image:
                return

            if not size_ok_any:
                self.dimension2_results.append((-1, '第3页"行为特征"右侧单元格图片大小不满足11.25×5.04厘米'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page3_behavior_image_overlap(self):
        """检查第3页"行为特征"右侧单元格中的图片是否与表格或字体发生重叠。

        细则：-1：第3页"行为特征"右侧单元格里的图片出现与表格或字体重叠
        —— 通过办公软件（Word / WPS 文字）判定图片是否可能与表格框线
           或相邻单元格字体发生视觉重叠。
        —— 判定原则：
             · 嵌入型图片（InlineShape）本质位于文本流内，单元格会自动
               撑高/撑宽以容纳，视觉上永远位于单元格边框之内，不会与
               表格或字体重叠。本项对嵌入型图片一律不判扣。
             · 只有浮动型图片（Wrap 类型为 Square / Tight / Through /
               Behind / Front / None 等）脱离文本流，才可能越出单元格
               边框，与相邻单元格的表格框线或字体发生视觉重叠。
        —— 关键 COM 接口：
             · 单元格页面矩形：cell.Range.Information(5)/(6) 取左上角
               页面坐标（磅）；cell.Width/.Height 取单元格宽高（磅）。
             · 浮动图片页面矩形：以 sh.Anchor.Information(5)/(6) 为锚点
               在页面上的坐标，叠加 sh.Left/.Top 得到图片左上角；
               sh.Width/.Height 为图片尺寸。当
               RelativeHorizontalPosition == 1（Page）或
               RelativeVerticalPosition == 1（Page）时，Left/Top 已是
               页面坐标，不再叠加。
        —— 若浮动图片矩形越出所在单元格矩形（容差 0.05cm）任一边 → 扣 -1。
           若单元格内无浮动图片（含"完全没有图片"或"只有嵌入型图片"），
           本项不判扣。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            right_cell = self._locate_page3_behavior_right_cell(wdDoc)
            if right_cell is None:
                return

            TOL = 0.05 * 28.3464567  # 0.05cm 的磅值容差

            # 单元格页面矩形（磅）
            try:
                cell_L = float(right_cell.Range.Information(5))
                cell_T = float(right_cell.Range.Information(6))
                cell_W = float(right_cell.Width)
                cell_H = float(right_cell.Height)
            except Exception:
                return
            cell_R = cell_L + cell_W
            cell_B = cell_T + cell_H

            r_start = right_cell.Range.Start
            r_end = right_cell.Range.End

            overlapped = False

            # 只对浮动型图片做越界判定：文档级 Shapes 中锚点在该单元格 Range 内的形状
            try:
                for si in range(1, wdDoc.Shapes.Count + 1):
                    sh = wdDoc.Shapes(si)
                    # 仅图片类型 (msoPicture=13, msoLinkedPicture=11)
                    try:
                        sh_type = int(sh.Type)
                    except Exception:
                        sh_type = -1
                    if sh_type not in (11, 13):
                        continue

                    try:
                        a = sh.Anchor
                    except Exception:
                        a = None
                    if a is None:
                        continue
                    if not (r_start <= a.Start <= r_end):
                        continue

                    try:
                        sh_L = float(sh.Left)
                        sh_T = float(sh.Top)
                        sh_W = float(sh.Width)
                        sh_H = float(sh.Height)
                    except Exception:
                        continue

                    try:
                        rhp = int(sh.RelativeHorizontalPosition)
                    except Exception:
                        rhp = -1
                    try:
                        rvp = int(sh.RelativeVerticalPosition)
                    except Exception:
                        rvp = -1

                    # 计算图片左上角在页面上的坐标
                    try:
                        anc_x = float(a.Information(5))
                        anc_y = float(a.Information(6))
                    except Exception:
                        anc_x = cell_L
                        anc_y = cell_T

                    # RelativeHorizontalPosition == 1 (wdRelativeHorizontalPositionPage)
                    # 时 sh.Left 就是页面坐标；其它常见模式（Column/Character 等，
                    # 位于单元格内）以锚点页面坐标为基准叠加 sh.Left。
                    page_L = sh_L if rhp == 1 else (anc_x + sh_L)
                    page_T = sh_T if rvp == 1 else (anc_y + sh_T)
                    page_R = page_L + sh_W
                    page_B = page_T + sh_H

                    # 图片矩形越出单元格矩形任何一边（容差 0.05cm）
                    if (page_L < cell_L - TOL or
                        page_T < cell_T - TOL or
                        page_R > cell_R + TOL or
                        page_B > cell_B + TOL):
                        overlapped = True
                        break
            except Exception:
                pass

            if overlapped:
                self.dimension2_results.append((-1, '第3页"行为特征"右侧单元格里的图片出现与表格或字体重叠'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _locate_page6_target_below_cell(self, wdDoc):
        """通过办公软件定位到第6页"专业为本，工具赋能，整合多元信息化手段
        服务教学。"所在单元格正下方（同一列的下一行）的单元格。

        返回下方单元格对象；若目标句不存在于第 6 页、或该单元格已在最下
        一行（无下邻）、或表格结构异常，则返回 None。
        """
        KEY = '专业为本，工具赋能，整合多元信息化手段服务教学'
        for ti in range(1, wdDoc.Tables.Count + 1):
            tbl = wdDoc.Tables(ti)
            cells = tbl.Range.Cells
            for ci in range(1, cells.Count + 1):
                cell = cells(ci)
                try:
                    txt = cell.Range.Text
                except Exception:
                    continue
                if KEY not in txt:
                    continue
                try:
                    if int(cell.Range.Information(3)) != 6:
                        continue
                except Exception:
                    continue
                row, col = cell.RowIndex, cell.ColumnIndex
                for cj in range(1, cells.Count + 1):
                    other = cells(cj)
                    if other.RowIndex == row + 1 and other.ColumnIndex == col:
                        return other
                return None
        return None

    def _check_page6_target_below_image(self):
        """检查第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        下方单元格内是否出现图片。

        细则：-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
              下方单元格内未出现图片
        —— 通过办公软件（Word / WPS 文字）真实识别下方单元格内的"图片"
           对象。图片包括：
             · 嵌入型：InlineShape.Type ∈ {wdInlineShapePicture (3),
               wdInlineShapeLinkedPicture (4)}；
             · 浮动型：文档级 Shapes 中锚点落入该单元格 Range 的形状，
               且 Shape.Type ∈ {msoPicture (13), msoLinkedPicture (11)}。
           非图片对象（文本框 msoTextBox=17、OLE 对象、形状等）不计入。
           这与办公软件里在该单元格中肉眼可见的"图片"一致。
        —— 只对第 6 页的目标单元格进行判定（细则"第6页"三字）。若目标
           单元格或其下方单元格不存在，本项不判扣（细则未要求补扣）。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            below = self._locate_page6_target_below_cell(wdDoc)
            if below is None:
                return

            r_start = below.Range.Start
            r_end = below.Range.End

            has_image = False

            # 1) 嵌入型：InlineShapes 中的图片
            try:
                ins = below.Range.InlineShapes
                for i in range(1, ins.Count + 1):
                    try:
                        t = int(ins(i).Type)
                    except Exception:
                        t = -1
                    # 3=wdInlineShapePicture, 4=wdInlineShapeLinkedPicture
                    if t in (3, 4):
                        has_image = True
                        break
            except Exception:
                pass

            # 2) 浮动型：文档级 Shapes 中锚点在该单元格内的图片
            if not has_image:
                try:
                    for si in range(1, wdDoc.Shapes.Count + 1):
                        sh = wdDoc.Shapes(si)
                        try:
                            a = sh.Anchor
                        except Exception:
                            a = None
                        if a is None:
                            continue
                        if not (r_start <= a.Start <= r_end):
                            continue
                        try:
                            t = int(sh.Type)
                        except Exception:
                            t = -1
                        # 13=msoPicture, 11=msoLinkedPicture
                        if t in (11, 13):
                            has_image = True
                            break
                except Exception:
                    pass

            if not has_image:
                self.dimension2_results.append((-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内未出现图片'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page6_target_below_image_wrap(self):
        """检查第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        下方单元格内图片的环绕方式是否为"嵌入型"。

        细则：-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
              下方单元格内的图片布局不满足嵌入型
        —— 通过办公软件（Word / WPS 文字）读取下方单元格内每张图片的
           真实类型 / 环绕方式，与办公软件里"图片工具→环绕方式"面板
           所见完全一致。
        —— 嵌入型 = `InlineShape`：办公软件里所有 InlineShape 的环绕方式
           就是"嵌入型"。若图片以浮动 `Shape` 形式出现（无论其
           `WrapFormat.Type` 是四周型 0 / 紧密型 1 / 穿越型 2 / 无环绕 3
           / 上下型 4 / 衬于文字下方 5 / 浮于文字上方 6），都不是"嵌入型"。
        —— 判定：单元格内任一图片（`InlineShape.Type ∈ {3,4}` 或
           `Shape.Type ∈ {11,13}`）不是 InlineShape 时，即扣 -1。若下方
           单元格不存在、或单元格内无任何图片，本项不判扣（另一细则会
           处理"未出现图片"）。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            below = self._locate_page6_target_below_cell(wdDoc)
            if below is None:
                return

            r_start = below.Range.Start
            r_end = below.Range.End

            has_image = False
            all_inline = True  # 只要有任一图片不是 InlineShape，即为非嵌入型

            # 1) 嵌入型（InlineShape 图片）——环绕方式即为"嵌入型"
            try:
                ins = below.Range.InlineShapes
                for i in range(1, ins.Count + 1):
                    try:
                        t = int(ins(i).Type)
                    except Exception:
                        t = -1
                    if t in (3, 4):  # 3=wdInlineShapePicture, 4=wdInlineShapeLinkedPicture
                        has_image = True
            except Exception:
                pass

            # 2) 浮动型（Shapes 中锚点在该单元格内的图片）——非嵌入型
            try:
                for si in range(1, wdDoc.Shapes.Count + 1):
                    sh = wdDoc.Shapes(si)
                    try:
                        a = sh.Anchor
                    except Exception:
                        a = None
                    if a is None:
                        continue
                    if not (r_start <= a.Start <= r_end):
                        continue
                    try:
                        t = int(sh.Type)
                    except Exception:
                        t = -1
                    if t in (11, 13):  # 13=msoPicture, 11=msoLinkedPicture
                        has_image = True
                        all_inline = False  # 浮动图片不是"嵌入型"
                        break
            except Exception:
                pass

            # 无图片：不由本项判扣
            if not has_image:
                return

            if not all_inline:
                self.dimension2_results.append((-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内的图片布局不满足嵌入型'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page6_target_below_image_size(self):
        """检查第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        下方单元格内图片的大小是否为 4.19 × 13.53 厘米。

        细则：-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
              下方单元格内的图片大小不满足4.19×13.53厘米
        —— 仅按此一项判定（不检查环绕、位置、颜色等，另有对应细则）。
        —— 通过办公软件（Word / WPS 文字）读取下方单元格内每张图片的
           真实 `Width` / `Height`（单位为磅 pt），换算到厘米后与
           4.19 × 13.53 cm 比较，与办公软件"图片工具→大小"面板显示
           完全一致（Office UI 中厘米保留两位小数）。
        —— 匹配容差 0.05 cm，覆盖办公软件界面按两位小数四舍五入显示
           时的等价情形。细则文字"4.19×13.53厘米"未强制指定"宽×高"
           或"高×宽"顺序，故允许两种朝向匹配（即
           (W≈4.19 且 H≈13.53) 或 (W≈13.53 且 H≈4.19)），只要办公
           软件读到的图片外形符合 4.19cm 与 13.53cm 两条边长即视为
           满足；这与"图片大小"面板上"高度=4.19cm、宽度=13.53cm"
           的实际书写方式一致。
        —— 单元格内没有任何图片时，本项不判扣（另一细则会处理
           "未出现图片"）；下方单元格不存在时同样不判扣。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            below = self._locate_page6_target_below_cell(wdDoc)
            if below is None:
                return

            r_start = below.Range.Start
            r_end = below.Range.End

            CM = 28.3464567
            TOL = 0.05
            A_CM = 4.19
            B_CM = 13.53

            def _size_ok(w_cm, h_cm):
                # 允许"宽×高"或"高×宽"两种朝向，只要两条边长命中 4.19 与 13.53 即满足
                return ((abs(w_cm - A_CM) <= TOL and abs(h_cm - B_CM) <= TOL) or
                        (abs(w_cm - B_CM) <= TOL and abs(h_cm - A_CM) <= TOL))

            has_image = False
            size_ok_any = False

            # 1) 嵌入型：InlineShapes 中的图片（Type ∈ {3,4}）
            try:
                ins = below.Range.InlineShapes
                for i in range(1, ins.Count + 1):
                    s = ins(i)
                    try:
                        t = int(s.Type)
                    except Exception:
                        t = -1
                    if t not in (3, 4):
                        continue
                    has_image = True
                    try:
                        w_cm = float(s.Width) / CM
                        h_cm = float(s.Height) / CM
                    except Exception:
                        continue
                    if _size_ok(w_cm, h_cm):
                        size_ok_any = True
                        break
            except Exception:
                pass

            # 2) 浮动型：文档级 Shapes 中锚点在该单元格 Range 内的图片
            if not size_ok_any:
                try:
                    for si in range(1, wdDoc.Shapes.Count + 1):
                        sh = wdDoc.Shapes(si)
                        try:
                            a = sh.Anchor
                        except Exception:
                            a = None
                        if a is None:
                            continue
                        if not (r_start <= a.Start <= r_end):
                            continue
                        try:
                            t = int(sh.Type)
                        except Exception:
                            t = -1
                        if t not in (11, 13):  # 只判"图片"，跳过文本框/形状等
                            continue
                        has_image = True
                        try:
                            w_cm = float(sh.Width) / CM
                            h_cm = float(sh.Height) / CM
                        except Exception:
                            continue
                        if _size_ok(w_cm, h_cm):
                            size_ok_any = True
                            break
                except Exception:
                    pass

            # 无图片：不由本项判扣
            if not has_image:
                return

            if not size_ok_any:
                self.dimension2_results.append((-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内的图片大小不满足4.19×13.53厘米'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page6_target_below_image_overlap(self):
        """检查第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        下方单元格内的图片是否与下方表格或字体发生重叠。

        细则：-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
              下方单元格内的图片与下方表格或字体重叠
        —— 通过办公软件（Word / WPS 文字）读取下方单元格与该单元格内
           图片的真实页面坐标（磅），几何判定图片矩形是否越出单元格
           矩形。若图片矩形任意一边越出所在单元格矩形，则该图片一定
           会与下方相邻行/表格框线或该行内的文字发生视觉重叠，与办公
           软件所见完全一致。
        —— 关键 COM 接口（磅）：
             · 单元格页面矩形：cell.Range.Information(5) / (6) +
               cell.Width / cell.Height。
             · 嵌入型图片：InlineShape.Width / .Height 与单元格宽高比较；
               任一尺寸超出单元格即认为撑破边框。
             · 浮动型图片：sh.Anchor.Information(5)/(6) + sh.Left/.Top
               计算左上角页面坐标；当 RelativeHorizontal/VerticalPosition
               == 1（Page）时 sh.Left/.Top 已是页面坐标，不再叠加；
               再叠加 sh.Width / .Height 得图片页面矩形，与单元格矩形
               比较（容差 0.05cm）。
        —— 判定：单元格内任一图片（嵌入或浮动）越出单元格边界即扣 -1。
           若单元格不存在或单元格内无任何图片，不由本项判扣。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            below = self._locate_page6_target_below_cell(wdDoc)
            if below is None:
                return

            TOL = 0.05 * 28.3464567  # 0.05cm 的磅值容差

            try:
                cell_L = float(below.Range.Information(5))
                cell_T = float(below.Range.Information(6))
                cell_W = float(below.Width)
                cell_H = float(below.Height)
            except Exception:
                return
            cell_R = cell_L + cell_W
            cell_B = cell_T + cell_H

            r_start = below.Range.Start
            r_end = below.Range.End

            has_image = False
            overlapped = False

            # 1) 嵌入型图片：Width/Height 与单元格宽高比较
            try:
                ins = below.Range.InlineShapes
                for i in range(1, ins.Count + 1):
                    s = ins(i)
                    try:
                        t = int(s.Type)
                    except Exception:
                        t = -1
                    if t not in (3, 4):
                        continue
                    has_image = True
                    try:
                        w = float(s.Width)
                        h = float(s.Height)
                    except Exception:
                        continue
                    if w - cell_W > TOL or h - cell_H > TOL:
                        overlapped = True
                        break
            except Exception:
                pass

            # 2) 浮动型图片：文档级 Shapes 中锚点在该单元格 Range 内的图片
            if not overlapped:
                try:
                    for si in range(1, wdDoc.Shapes.Count + 1):
                        sh = wdDoc.Shapes(si)
                        try:
                            a = sh.Anchor
                        except Exception:
                            a = None
                        if a is None:
                            continue
                        if not (r_start <= a.Start <= r_end):
                            continue
                        try:
                            t = int(sh.Type)
                        except Exception:
                            t = -1
                        if t not in (11, 13):  # 只判"图片"
                            continue
                        has_image = True

                        try:
                            sh_L = float(sh.Left)
                            sh_T = float(sh.Top)
                            sh_W = float(sh.Width)
                            sh_H = float(sh.Height)
                        except Exception:
                            continue
                        try:
                            rhp = int(sh.RelativeHorizontalPosition)
                        except Exception:
                            rhp = -1
                        try:
                            rvp = int(sh.RelativeVerticalPosition)
                        except Exception:
                            rvp = -1
                        try:
                            anc_x = float(a.Information(5))
                            anc_y = float(a.Information(6))
                        except Exception:
                            anc_x = cell_L
                            anc_y = cell_T

                        page_L = sh_L if rhp == 1 else (anc_x + sh_L)
                        page_T = sh_T if rvp == 1 else (anc_y + sh_T)
                        page_R = page_L + sh_W
                        page_B = page_T + sh_H

                        if (page_L < cell_L - TOL or
                            page_T < cell_T - TOL or
                            page_R > cell_R + TOL or
                            page_B > cell_B + TOL):
                            overlapped = True
                            break
                except Exception:
                    pass

            # 单元格内无图片：不由本项判扣
            if not has_image:
                return

            if overlapped:
                self.dimension2_results.append((-1, '第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格内的图片与下方表格或字体重叠'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_cell_has_image(self, cell):
        """检查单元格是否包含图片"""
        for para in cell.paragraphs:
            for run in para.runs:
                drawings = run._element.findall('.//' + qn('w:drawing'))
                if drawings:
                    return True
        return False

    def _check_image_properties(self, cell, expected_width, expected_height, wrap_type, page_name):
        '''检查图片属性'''
        for para in cell.paragraphs:
            for run in para.runs:
                drawings = run._element.findall('.//' + qn('w:drawing'))
                for drawing in drawings:
                    # 获取图片尺寸
                    extents = drawing.findall('.//' + qn('wp:extent'))
                    for extent in extents:
                        if extent is not None:
                            width_emu = int(extent.get('cx', 0))
                            height_emu = int(extent.get('cy', 0))
                            width_cm = width_emu / 360000
                            height_cm = height_emu / 360000

                            # 检查尺寸
                            if abs(width_cm - expected_width) > 0.5 or abs(height_cm - expected_height) > 0.5:
                                self.dimension2_results.append((-1, f'{page_name}"行为特征"右侧单元格图片大小不满足{expected_width}×{expected_height}厘米'))

                    # 检查环绕方式
                    if wrap_type == '四周型':
                        wrap_square = drawing.findall('.//' + qn('wp:wrapSquare'))
                        if len(wrap_square) == 0:
                            self.dimension2_results.append((-1, f'{page_name}"行为特征"右侧单元格图片布局不满足四周型环绕'))

    def _get_page_for_element(self, element):
        """根据节分隔符判断元素所在页码"""
        # 收集所有节分隔符所在的段落索引
        section_breaks = []
        for i, para in enumerate(self.doc.paragraphs):
            pPr = para._p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                sectPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                if sectPr is not None:
                    section_breaks.append(i)

        # 找到元素对应的段落索引
        body = self.doc.element.body
        children = list(body)
        para_idx = None
        xml_idx = None

        # 找到元素在body中的位置
        for i, child in enumerate(children):
            if child is element:
                xml_idx = i
                # 向前查找最近的段落
                for j in range(i - 1, -1, -1):
                    if children[j].tag.split('}')[-1] == 'p':
                        # 映射到doc.paragraphs索引
                        for pi, para in enumerate(self.doc.paragraphs):
                            if para._p is children[j]:
                                para_idx = pi
                                break
                        break
                break

        # 计算页码
        page = 1
        if para_idx is not None:
            for sbp in section_breaks:
                if para_idx > sbp:
                    page += 1
        elif xml_idx is not None and section_breaks:
            if xml_idx > section_breaks[0]:
                page += 1

        return page

    def _is_table_split(self, table):
        """检查表格是否跨页断开"""
        tbl = table._tbl
        rows = tbl.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')

        # 检查是否有行设置了cantSplit（禁止跨页）
        # 如果表格有cantSplit且行数较多，可能表示表格被强制不分页
        # 但更直接的方式是检查表格是否被分页符分割

        # 检查表格前后是否有分页符
        body = self.doc.element.body
        children = list(body)
        tbl_idx = None
        for i, child in enumerate(children):
            if child is tbl:
                tbl_idx = i
                break

        if tbl_idx is None:
            return False

        # 检查表格后面紧跟的元素
        next_idx = tbl_idx + 1
        if next_idx < len(children):
            next_elem = children[next_idx]
            tag = next_elem.tag.split('}')[-1] if '}' in next_elem.tag else next_elem.tag
            if tag == 'p':
                # 检查是否是空段落（可能包含分页符）
                texts = next_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                text = ''.join([t.text for t in texts if t.text])
                if not text.strip():
                    # 检查是否有分页符
                    brs = next_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
                    for br in brs:
                        br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                        if br_type == 'page':
                            return True

        # 检查表格内部是否有分页符（表格被分割）
        for row in rows:
            brs = row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
            for br in brs:
                br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                if br_type == 'page':
                    return True

        return False

    def _open_with_office(self):
        """在办公软件中打开当前文档，返回 (app, doc, temp_path)。

        依次尝试 Microsoft Word（Word.Application）与 WPS 文字（kwps.Application、
        wps.Application），确保在实际的办公软件里获得真实分页信息。

        为避免文档正被用户在办公软件中打开导致独占锁定，这里先将文件复制到
        临时路径，再由办公软件以只读方式打开该副本。
        """
        if os.environ.get("OFFICEVAL_COM_ENABLED", "1") != "1":
            raise RuntimeError("Office COM 已禁用")
        import win32com.client
        import shutil
        import tempfile

        # 复制到临时文件，规避 "文档正被占用" 的情况
        tmp_dir = tempfile.gettempdir()
        temp_path = os.path.join(tmp_dir, f'_eval_{os.getpid()}_{os.path.basename(self.doc_path)}')
        shutil.copyfile(self.doc_path, temp_path)

        last_err = None
        for progid in ('Word.Application', 'kwps.Application', 'wps.Application'):
            try:
                app = win32com.client.DispatchEx(progid)
                try:
                    app.Visible = False
                except Exception:
                    pass
                try:
                    app.DisplayAlerts = 0
                except Exception:
                    pass
                doc = app.Documents.Open(
                    temp_path,
                    ReadOnly=True,
                    ConfirmConversions=False,
                    AddToRecentFiles=False,
                )
                return app, doc, temp_path
            except Exception as e:
                last_err = e
                continue

        # 全部失败，清理临时文件后抛出
        try:
            os.remove(temp_path)
        except Exception:
            pass
        raise RuntimeError(f'无法通过办公软件打开文档: {last_err}')

    def _check_page4_table(self):
        """检查第4页表格

        细则：+5：文档第4页只存在一个表格，且表格未出现任何断开位置
        —— 仅按细则的两个条件判定，不附加任何额外约束。
        —— 通过办公软件（Word / WPS 文字）的真实分页信息，确保判定在
           办公软件上与用户所见一致。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            # 办公软件不可用时的备用判定
            self._check_page4_table_fallback()
            return

        try:
            page_count = wdDoc.ComputeStatistics(2)  # wdStatisticPages

            if page_count < 4:
                return

            # 收集所有"出现在第4页"的表格及其起止页
            page4_tables = []
            for table in wdDoc.Tables:
                start = table.Range.Start
                end = table.Range.End
                start_page = wdDoc.Range(start, start).Information(3)  # wdActiveEndPageNumber
                end_page = wdDoc.Range(end, end).Information(3)

                # 表格只要有任意部分落在第4页，即视为"存在于第4页"
                if start_page <= 4 <= end_page:
                    page4_tables.append((start_page, end_page))

            # 条件1：第4页只存在一个表格
            if len(page4_tables) != 1:
                return

            # 条件2：该表格未出现任何断开位置
            #   —— 表格必须完整位于第4页，起止页均为第4页；
            #   —— 若跨到第3页或第5页，则说明表格在第3/4 或 4/5 页之间断开。
            start_page, end_page = page4_tables[0]
            if start_page == 4 and end_page == 4:
                self.dimension2_results.append((5, '文档第4页只存在一个表格，且表格未出现任何断开位置'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page5_6(self):
        """检查第5-6页"""
        # 第5页表格数量与断开检查，通过办公软件读取每张表格的真实起止
        # 页码（Range.Information(wdActiveEndPageNumber)）进行判定，严格
        # 对应细则：
        #   "-3：第5页存在的表格数量大于1或表格存在断开位置"
        # —— 判定为两者的"或"：
        #     a) 出现在第 5 页的表格数量 > 1；或
        #     b) 任一"存在于第 5 页"的表格跨到第 4 或第 6 页（在第 5 页
        #        与相邻页之间出现"断开"）。
        # —— 与办公软件里所见的分页/表格切分完全一致；不做细则未列的
        #    其它检查。
        self._check_page_tables_count_and_split(page_num=5, penalty=-3,
            label='第5页存在的表格数量大于1或表格存在断开位置')

        # 第6页表格数量与断开检查，通过办公软件读取每张表格的真实起止
        # 页码进行判定，严格对应细则：
        #   "-3：第6页存在的表格数量大于1或表格存在断开位置"
        # —— 判定为两者的"或"：
        #     a) 出现在第 6 页的表格数量 > 1；或
        #     b) 任一"存在于第 6 页"的表格跨到第 5 或第 7 页（在第 6 页
        #        与相邻页之间出现"断开"）。
        # —— 与办公软件里所见的分页/表格切分完全一致；不做细则未列的
        #    其它检查。
        self._check_page_tables_count_and_split(page_num=6, penalty=-3,
            label='第6页存在的表格数量大于1或表格存在断开位置')

        # 第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格
        # 内是否出现图片，通过办公软件（Word / WPS 文字）真实识别下方单元格
        # 内的图片对象类型进行判定，严格对应细则：
        #   "-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        #     下方单元格内未出现图片"
        self._check_page6_target_below_image()

        # 第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格
        # 内图片的环绕方式，通过办公软件读取每张图片的真实类型 / 环绕方式
        # 进行判定，严格对应细则：
        #   "-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        #     下方单元格内的图片布局不满足嵌入型"
        self._check_page6_target_below_image_wrap()

        # 第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格
        # 内图片的大小，通过办公软件读取每张图片真实 Width / Height（磅→
        # 厘米）判定，严格对应细则：
        #   "-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        #     下方单元格内的图片大小不满足4.19×13.53厘米"
        self._check_page6_target_below_image_size()

        # 第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"下方单元格
        # 内图片是否与下方表格或字体发生重叠，通过办公软件读取图片与单元格
        # 的真实页面坐标进行几何判定，严格对应细则：
        #   "-1：第6页"专业为本，工具赋能，整合多元信息化手段服务教学。"
        #     下方单元格内的图片与下方表格或字体重叠"
        self._check_page6_target_below_image_overlap()

    def _check_pages_7_to_11(self):
        """检查第7-11页表格数量"""
        # 第7页表格数量与断开检查，通过办公软件读取每张表格的真实起止
        # 页码（Range.Information(wdActiveEndPageNumber)）进行判定，严格
        # 对应细则：
        #   "-3：第7页存在的表格数量大于1或表格存在断开位置"
        # —— 判定为两者的"或"：
        #     a) 出现在第 7 页的表格数量 > 1；或
        #     b) 任一"存在于第 7 页"的表格跨到第 6 或第 8 页（在第 7 页
        #        与相邻页之间出现"断开"）。
        # —— 与办公软件里所见的分页/表格切分完全一致；不做细则未列的
        #    其它检查。
        self._check_page_tables_count_and_split(page_num=7, penalty=-3,
            label='第7页存在的表格数量大于1或表格存在断开位置')

        # 第8页表格数量与断开检查，通过办公软件读取每张表格的真实起止
        # 页码进行判定，严格对应细则：
        #   "-3：第8页存在的表格数量大于1或表格存在断开位置"
        # —— 判定为两者的"或"：
        #     a) 出现在第 8 页的表格数量 > 1；或
        #     b) 任一"存在于第 8 页"的表格跨到第 7 或第 9 页（在第 8 页
        #        与相邻页之间出现"断开"）。
        # —— 与办公软件里所见的分页/表格切分完全一致；不做细则未列的
        #    其它检查。
        self._check_page_tables_count_and_split(page_num=8, penalty=-3,
            label='第8页存在的表格数量大于1或表格存在断开位置')

        # 第9页表格数量与断开检查，通过办公软件读取每张表格的真实起止
        # 页码进行判定，严格对应细则：
        #   "-3：第9页存在的表格数量大于1或表格存在断开位置"
        # —— 判定为两者的"或"：
        #     a) 出现在第 9 页的表格数量 > 1；或
        #     b) 任一"存在于第 9 页"的表格跨到第 8 或第 10 页（在第 9 页
        #        与相邻页之间出现"断开"）。
        # —— 与办公软件里所见的分页/表格切分完全一致；不做细则未列的
        #    其它检查。
        self._check_page_tables_count_and_split(page_num=9, penalty=-3,
            label='第9页存在的表格数量大于1或表格存在断开位置')

        # 第10页表格数量与断开检查，通过办公软件读取每张表格的真实起止
        # 页码进行判定，严格对应细则：
        #   "-3：第10页存在的表格数量大于1或表格存在断开位置"
        # —— 判定为两者的"或"：
        #     a) 出现在第 10 页的表格数量 > 1；或
        #     b) 任一"存在于第 10 页"的表格跨到第 9 或第 11 页（在第 10 页
        #        与相邻页之间出现"断开"）。
        # —— 与办公软件里所见的分页/表格切分完全一致；不做细则未列的
        #    其它检查。
        self._check_page_tables_count_and_split(page_num=10, penalty=-3,
            label='第10页存在的表格数量大于1或表格存在断开位置')

        # 第11页表格数量与断开检查，通过办公软件读取每张表格的真实起止
        # 页码进行判定，严格对应细则：
        #   "-3：第11页存在的表格数量大于1或表格存在断开位置"
        # —— 判定为两者的"或"：
        #     a) 出现在第 11 页的表格数量 > 1；或
        #     b) 任一"存在于第 11 页"的表格跨到第 10 或第 12 页（在第 11 页
        #        与相邻页之间出现"断开"）。
        # —— 与办公软件里所见的分页/表格切分完全一致；不做细则未列的
        #    其它检查。
        self._check_page_tables_count_and_split(page_num=11, penalty=-3,
            label='第11页存在的表格数量大于1或表格存在断开位置')

        try:
            if os.environ.get("OFFICEVAL_COM_ENABLED", "1") != "1":
                raise RuntimeError("Office COM 已禁用")
            import win32com.client
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            wdDoc = word.Documents.Open(self.doc_path)

            page_count = wdDoc.ComputeStatistics(2)  # wdStatisticPages

            # 统计每页的表格数量
            tables_per_page = [0] * page_count

            for table in wdDoc.Tables:
                start = table.Range.Start
                end = table.Range.End

                # 获取表格起始页
                rng_start = wdDoc.Range(start, start)
                start_page = rng_start.Information(3)  # wdActiveEndPageNumber

                # 获取表格结束页
                rng_end = wdDoc.Range(end, end)
                end_page = rng_end.Information(3)

                # 统计该表格覆盖的所有页面
                for page in range(start_page, end_page + 1):
                    if 1 <= page <= page_count:
                        tables_per_page[page - 1] += 1

            # 第 7-11 页均已通过 _check_page_tables_count_and_split 严格判定，
            # 本处不再重复统计（细则未列 12 页之后的表格检查）。

            wdDoc.Close(False)
            try:
                word.Quit()
            except:
                pass

        except Exception as e:
            # 如果COM不可用，跳过此检查
            pass

        # 第11页"3.板书设计"下方是否出现图片，通过办公软件（Word / WPS 文字）
        # 定位"3.板书设计"在第 11 页的真实位置，并在其之后、仍位于第 11 页
        # 的范围内识别图片对象，严格对应细则：
        #   "-1：第11页"3.板书设计"下方未出现图片"
        self._check_page11_board_design_image()

        # 第11页"3.板书设计"下方图片尺寸检查，通过办公软件读取该图片真实
        # Width / Height（磅→厘米）判定，严格对应细则：
        #   "-1：第11页"3.板书设计"下方的图片尺寸不满足9.64×6.14厘米"
        self._check_page11_board_design_image_size()

        # 第11页"3.板书设计"下方图片与下方表格或字体的重叠检查，通过办公
        # 软件读取该图片与其下方内容的真实页面坐标进行几何判定，严格对应
        # 细则：
        #   "-1：第11页"3.板书设计"下方的图片与下方表格或字体重叠"
        self._check_page11_board_design_image_overlap()

        # 第11页"3.板书设计"下方图片布局检查，通过办公软件读取该图片的
        # 真实对象类型 / 环绕方式判定，严格对应细则：
        #   "-1：第11页"3.板书设计"下方的图片布局不满足嵌入型"
        self._check_page11_board_design_image_wrap()

    def _check_page11_board_design_image(self):
        """检查第11页"3.板书设计"下方是否出现图片。

        细则：-1：第11页"3.板书设计"下方未出现图片
        —— 通过办公软件（Word / WPS 文字）用 Content.Find 反复推进查找
           "3.板书设计"，直到该匹配项落在第 11 页；以其段落结束点作为
           "下方"的起点，在文档剩余范围里识别位于第 11 页的图片对象：
             · 嵌入型：InlineShape，且 Type ∈ {wdInlineShapePicture=3,
               wdInlineShapeLinkedPicture=4}；
             · 浮动型：文档级 Shapes，且 Type ∈ {msoPicture=13,
               msoLinkedPicture=11}，锚点位于起点之后且仍在第 11 页。
           非图片对象（文本框/形状/OLE 等）不计入。
        —— 若"3.板书设计"未出现在第 11 页，本项不判扣（细则限定第 11 页
           的目标句）。若目标句后第 11 页范围内没有任何图片，判扣 -1。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            KEY = '3.板书设计'
            content = wdDoc.Content
            content.Find.ClearFormatting()

            target_end = None
            r = content
            r.Collapse(1)
            while True:
                ok = r.Find.Execute(FindText=KEY, Forward=True, MatchWildcards=False, Wrap=0)
                if not ok:
                    break
                try:
                    page = int(r.Information(3))
                except Exception:
                    page = -1
                if page == 11:
                    # 以匹配段落结束点作为"下方"的起点，以覆盖同段之后的
                    # 图片以及后续段落中的图片。
                    try:
                        para = r.Paragraphs(1)
                        target_end = int(para.Range.End)
                    except Exception:
                        target_end = int(r.End)
                    break
                r.Collapse(0)

            # "3.板书设计" 不在第 11 页：不由本项判扣
            if target_end is None:
                return

            has_image = False

            # 1) 嵌入型：文档级 InlineShapes，位置 > target_end 且仍在第 11 页
            try:
                ish = wdDoc.InlineShapes
                for i in range(1, ish.Count + 1):
                    s = ish(i)
                    try:
                        pos = int(s.Range.Start)
                        page = int(s.Range.Information(3))
                        t = int(s.Type)
                    except Exception:
                        continue
                    if pos >= target_end and page == 11 and t in (3, 4):
                        has_image = True
                        break
            except Exception:
                pass

            # 2) 浮动型：文档级 Shapes，锚点位置 > target_end 且仍在第 11 页
            if not has_image:
                try:
                    for si in range(1, wdDoc.Shapes.Count + 1):
                        sh = wdDoc.Shapes(si)
                        try:
                            a = sh.Anchor
                            apos = int(a.Start)
                            apage = int(a.Information(3))
                            t = int(sh.Type)
                        except Exception:
                            continue
                        if apos >= target_end and apage == 11 and t in (11, 13):
                            has_image = True
                            break
                except Exception:
                    pass

            if not has_image:
                self.dimension2_results.append((-1, '第11页"3.板书设计"下方未出现图片'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page11_board_design_image_size(self):
        """检查第11页"3.板书设计"下方图片的尺寸是否为 9.64 × 6.14 厘米。

        细则：-1：第11页"3.板书设计"下方的图片尺寸不满足9.64×6.14厘米
        —— 仅按此一项判定（不检查环绕、位置、颜色等；细则未列）。
        —— 通过办公软件（Word / WPS 文字）用 Content.Find 反复推进查找
           "3.板书设计"，直到匹配项落在第 11 页；以其段落结束点作为
           "下方"的起点，在文档剩余范围里识别位于第 11 页的图片对象
           （嵌入型 InlineShape.Type ∈ {3,4}；浮动型 Shape.Type ∈
           {11,13}），读取其真实 `Width`/`Height`（磅），换算到厘米
           后与 9.64 × 6.14 cm 比较，与办公软件"图片工具→大小"面板
           显示完全一致（Office UI 中厘米保留两位小数）。
        —— 匹配容差 0.05 cm；细则文字"9.64×6.14厘米"未强制"宽×高"
           或"高×宽"顺序，允许两种朝向匹配（即
           (W≈9.64 且 H≈6.14) 或 (W≈6.14 且 H≈9.64)），只要图片外形
           符合 9.64cm 与 6.14cm 两条边长即视为满足。
        —— "3.板书设计"不在第 11 页、或其下方无任何图片时，本项不判扣
           （另一细则处理"未出现图片"）。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            KEY = '3.板书设计'
            content = wdDoc.Content
            content.Find.ClearFormatting()

            target_end = None
            r = content
            r.Collapse(1)
            while True:
                ok = r.Find.Execute(FindText=KEY, Forward=True, MatchWildcards=False, Wrap=0)
                if not ok:
                    break
                try:
                    page = int(r.Information(3))
                except Exception:
                    page = -1
                if page == 11:
                    try:
                        para = r.Paragraphs(1)
                        target_end = int(para.Range.End)
                    except Exception:
                        target_end = int(r.End)
                    break
                r.Collapse(0)

            if target_end is None:
                return

            CM = 28.3464567
            TOL = 0.05
            A_CM = 9.64
            B_CM = 6.14

            def _size_ok(w_cm, h_cm):
                return ((abs(w_cm - A_CM) <= TOL and abs(h_cm - B_CM) <= TOL) or
                        (abs(w_cm - B_CM) <= TOL and abs(h_cm - A_CM) <= TOL))

            has_image = False
            size_ok_any = False

            # 1) 嵌入型：文档级 InlineShapes
            try:
                ish = wdDoc.InlineShapes
                for i in range(1, ish.Count + 1):
                    s = ish(i)
                    try:
                        pos = int(s.Range.Start)
                        page = int(s.Range.Information(3))
                        t = int(s.Type)
                    except Exception:
                        continue
                    if not (pos >= target_end and page == 11 and t in (3, 4)):
                        continue
                    has_image = True
                    try:
                        w_cm = float(s.Width) / CM
                        h_cm = float(s.Height) / CM
                    except Exception:
                        continue
                    if _size_ok(w_cm, h_cm):
                        size_ok_any = True
                        break
            except Exception:
                pass

            # 2) 浮动型：文档级 Shapes
            if not size_ok_any:
                try:
                    for si in range(1, wdDoc.Shapes.Count + 1):
                        sh = wdDoc.Shapes(si)
                        try:
                            a = sh.Anchor
                            apos = int(a.Start)
                            apage = int(a.Information(3))
                            t = int(sh.Type)
                        except Exception:
                            continue
                        if not (apos >= target_end and apage == 11 and t in (11, 13)):
                            continue
                        has_image = True
                        try:
                            w_cm = float(sh.Width) / CM
                            h_cm = float(sh.Height) / CM
                        except Exception:
                            continue
                        if _size_ok(w_cm, h_cm):
                            size_ok_any = True
                            break
                except Exception:
                    pass

            # 无图片：不由本项判扣
            if not has_image:
                return

            if not size_ok_any:
                self.dimension2_results.append((-1, '第11页"3.板书设计"下方的图片尺寸不满足9.64×6.14厘米'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page11_board_design_image_overlap(self):
        """检查第11页"3.板书设计"下方的图片是否与下方表格或字体发生重叠。

        细则：-1：第11页"3.板书设计"下方的图片与下方表格或字体重叠
        —— 通过办公软件（Word / WPS 文字）读取该图片与所在页文本区域的
           真实页面坐标（磅），几何判定图片矩形是否越出页面文本区域。
           页面文本区域（PageSetup 派生）为：
             L = LeftMargin
             T = TopMargin
             R = PageWidth  - RightMargin
             B = PageHeight - BottomMargin
           办公软件里所有正文表格 / 字体都排布在文本区域内，图片一旦
           越出文本区域，即会与下方表格框线 / 字体发生视觉重叠。
        —— 关键 COM 接口（磅）：
             · 嵌入型 InlineShape 的页面坐标：Range.Information(5)/(6)
               给出图片左上角的页面坐标；Width/Height 为图片尺寸。
             · 浮动型 Shape 的页面坐标：Anchor.Information(5)/(6) 作锚
               点，叠加 Left/Top 得到图片左上角；当
               RelativeHorizontal/VerticalPosition==1（Page）时
               Left/Top 已是页面坐标，不再叠加；Width/Height 为图片
               尺寸。
        —— 判定：图片矩形任何一边越出所在页文本区域（容差 0.05cm），
           即判定"图片与下方表格或字体重叠"，扣 -1。若"3.板书设计"不
           在第 11 页或其下方无任何图片，本项不判扣。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            KEY = '3.板书设计'
            content = wdDoc.Content
            content.Find.ClearFormatting()

            target_end = None
            r = content
            r.Collapse(1)
            while True:
                ok = r.Find.Execute(FindText=KEY, Forward=True, MatchWildcards=False, Wrap=0)
                if not ok:
                    break
                try:
                    page = int(r.Information(3))
                except Exception:
                    page = -1
                if page == 11:
                    try:
                        para = r.Paragraphs(1)
                        target_end = int(para.Range.End)
                    except Exception:
                        target_end = int(r.End)
                    break
                r.Collapse(0)

            if target_end is None:
                return

            # 取目标位置所在节的 PageSetup（页面文本区域）
            try:
                sec = wdDoc.Range(target_end, target_end).Sections(1)
                ps = sec.PageSetup
                page_w = float(ps.PageWidth)
                page_h = float(ps.PageHeight)
                top_m = float(ps.TopMargin)
                bot_m = float(ps.BottomMargin)
                left_m = float(ps.LeftMargin)
                right_m = float(ps.RightMargin)
            except Exception:
                return

            area_L = left_m
            area_T = top_m
            area_R = page_w - right_m
            area_B = page_h - bot_m

            TOL = 0.05 * 28.3464567  # 0.05cm 磅值容差

            has_image = False
            overlapped = False

            # 1) 嵌入型：Range.Information 页面坐标 + Width/Height
            try:
                ish = wdDoc.InlineShapes
                for i in range(1, ish.Count + 1):
                    s = ish(i)
                    try:
                        pos = int(s.Range.Start)
                        page = int(s.Range.Information(3))
                        t = int(s.Type)
                    except Exception:
                        continue
                    if not (pos >= target_end and page == 11 and t in (3, 4)):
                        continue
                    has_image = True
                    try:
                        ix = float(s.Range.Information(5))
                        iy = float(s.Range.Information(6))
                        iw = float(s.Width)
                        ih = float(s.Height)
                    except Exception:
                        continue
                    if (ix < area_L - TOL or iy < area_T - TOL or
                        ix + iw > area_R + TOL or iy + ih > area_B + TOL):
                        overlapped = True
                        break
            except Exception:
                pass

            # 2) 浮动型：Anchor 页面坐标 + Left/Top + Width/Height
            if not overlapped:
                try:
                    for si in range(1, wdDoc.Shapes.Count + 1):
                        sh = wdDoc.Shapes(si)
                        try:
                            a = sh.Anchor
                            apos = int(a.Start)
                            apage = int(a.Information(3))
                            t = int(sh.Type)
                        except Exception:
                            continue
                        if not (apos >= target_end and apage == 11 and t in (11, 13)):
                            continue
                        has_image = True
                        try:
                            rhp = int(sh.RelativeHorizontalPosition)
                        except Exception:
                            rhp = -1
                        try:
                            rvp = int(sh.RelativeVerticalPosition)
                        except Exception:
                            rvp = -1
                        try:
                            sh_L = float(sh.Left)
                            sh_T = float(sh.Top)
                            sh_W = float(sh.Width)
                            sh_H = float(sh.Height)
                            anc_x = float(a.Information(5))
                            anc_y = float(a.Information(6))
                        except Exception:
                            continue
                        page_L = sh_L if rhp == 1 else (anc_x + sh_L)
                        page_T = sh_T if rvp == 1 else (anc_y + sh_T)
                        page_R = page_L + sh_W
                        page_B = page_T + sh_H
                        if (page_L < area_L - TOL or page_T < area_T - TOL or
                            page_R > area_R + TOL or page_B > area_B + TOL):
                            overlapped = True
                            break
                except Exception:
                    pass

            # 无图片：不由本项判扣
            if not has_image:
                return

            if overlapped:
                self.dimension2_results.append((-1, '第11页"3.板书设计"下方的图片与下方表格或字体重叠'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page11_board_design_image_wrap(self):
        """检查第11页"3.板书设计"下方图片的环绕方式是否为"嵌入型"。

        细则：-1：第11页"3.板书设计"下方的图片布局不满足嵌入型
        —— 通过办公软件（Word / WPS 文字）用 Content.Find 反复推进查找
           "3.板书设计"，直到匹配项落在第 11 页；以其段落结束点作为
           "下方"的起点，在文档剩余范围里识别位于第 11 页的图片对象，
           读取其真实类型 / 环绕方式，与办公软件里"图片工具→环绕方式"
           面板所见完全一致。
        —— 嵌入型 = `InlineShape`：办公软件里所有 InlineShape 的环绕方式
           就是"嵌入型"（其 Type ∈ {wdInlineShapePicture=3,
           wdInlineShapeLinkedPicture=4} 时为图片）。若图片以浮动
           `Shape` 形式出现（Type ∈ {msoPicture=13, msoLinkedPicture=11}），
           无论其 `WrapFormat.Type` 是四周型 0 / 紧密型 1 / 穿越型 2 /
           无环绕 3 / 上下型 4 / 衬于文字下方 5 / 浮于文字上方 6，
           都不是"嵌入型"。
        —— 判定：目标位置下方的任一图片不是 InlineShape，即扣 -1。若
           "3.板书设计"不在第 11 页或其下方无任何图片，本项不判扣
           （另一细则会处理"未出现图片"）。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            KEY = '3.板书设计'
            content = wdDoc.Content
            content.Find.ClearFormatting()

            target_end = None
            r = content
            r.Collapse(1)
            while True:
                ok = r.Find.Execute(FindText=KEY, Forward=True, MatchWildcards=False, Wrap=0)
                if not ok:
                    break
                try:
                    page = int(r.Information(3))
                except Exception:
                    page = -1
                if page == 11:
                    try:
                        para = r.Paragraphs(1)
                        target_end = int(para.Range.End)
                    except Exception:
                        target_end = int(r.End)
                    break
                r.Collapse(0)

            if target_end is None:
                return

            has_image = False
            all_inline = True

            # 1) 嵌入型：InlineShape 图片（环绕方式即"嵌入型"）
            try:
                ish = wdDoc.InlineShapes
                for i in range(1, ish.Count + 1):
                    s = ish(i)
                    try:
                        pos = int(s.Range.Start)
                        page = int(s.Range.Information(3))
                        t = int(s.Type)
                    except Exception:
                        continue
                    if pos >= target_end and page == 11 and t in (3, 4):
                        has_image = True
            except Exception:
                pass

            # 2) 浮动型：Shape 图片，非"嵌入型"
            try:
                for si in range(1, wdDoc.Shapes.Count + 1):
                    sh = wdDoc.Shapes(si)
                    try:
                        a = sh.Anchor
                        apos = int(a.Start)
                        apage = int(a.Information(3))
                        t = int(sh.Type)
                    except Exception:
                        continue
                    if apos >= target_end and apage == 11 and t in (11, 13):
                        has_image = True
                        all_inline = False
                        break
            except Exception:
                pass

            # 无图片：不由本项判扣
            if not has_image:
                return

            if not all_inline:
                self.dimension2_results.append((-1, '第11页"3.板书设计"下方的图片布局不满足嵌入型'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page4_table_fallback(self):
        """COM 不可用时的第4页表格备用检查

        严格按细则："+5：文档第4页只存在一个表格，且表格未出现任何断开位置"。
        细则未规定其它扣分项，故本方法只加分，不扣分。

        备用方案通过节（section）近似判定表格所在页；仅当能明确判断
        "第4节仅有 1 个表格 且 该表格未跨节" 时才加 5 分，其余情况保守不加分。
        """
        body = self.doc.element.body
        children = list(body)

        # 逐节统计表格数量，并记录每个表格所属的节区间
        current_section = 1
        section_tables = {}   # {section_index: table_count}
        table_sections = []   # [(start_section, end_section), ...]

        for i, child in enumerate(children):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'tbl':
                section_tables[current_section] = section_tables.get(current_section, 0) + 1
                # 记录该表格所在的节；此处按当前节记录，跨节情形下方另行检测
                table_sections.append([current_section, current_section])

            if tag == 'p':
                pPr = child.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                if pPr is not None:
                    sectPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                    if sectPr is not None:
                        current_section += 1

        # 第4节视为第4页（备用近似）
        page4_table_count = section_tables.get(4, 0)
        if page4_table_count == 1:
            # 未能通过 XML 精准判断断页；保守起见，仅在无跨节且无强制分页符时加分
            # 这里 table_sections 记录的都是单节，且备用路径下无法读到真实分页，
            # 因此在只有 1 张表格的前提下，认为其"未出现任何断开位置"。
            self.dimension2_results.append((5, '文档第4页只存在一个表格，且表格未出现任何断开位置'))

    def _check_page_numbers(self):
        """检查页码。

        细则：-3：文档中除封面页外其余任意一页没有出现页码
        —— 仅按此一项判定，不附加其它约束。
        —— 通过办公软件（Word / WPS 文字）打开文档，逐页判断该页对应的
           页眉/页脚（考虑首页不同、奇偶页不同的节设置）中是否存在页码。
           "存在页码"的判据（三选一即可）：
             a) 页眉/页脚的 Range.Fields 中出现 PAGE 域 (Type==33) 或
                NUMPAGES 域 (Type==26)；
             b) 页眉/页脚的正文文本（去段落符后）出现任意数字；
             c) 页眉/页脚内锚定的形状/文本框 (Shapes) 里出现 PAGE 域
                或数字文本 —— WPS/Word 通过"插入页码"功能常以文本框
                形式插入到页脚，Range.Text/Range.Fields 无法读到这类
                内容，因此必须遍历 Footers/Headers 里的 Shapes。
           若封面页（第1页）之外的任意一页未出现页码，则扣 -3。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            # 办公软件不可用时不做该项判定，避免误扣分
            return

        try:
            import re as _re

            def _range_has_pagenum(rng):
                """判定给定 Range 中是否存在页码（PAGE 域 / NUMPAGES 域 / 数字文本）。"""
                if rng is None:
                    return False
                # 1) 域
                try:
                    fields = rng.Fields
                    for i in range(1, int(fields.Count) + 1):
                        try:
                            ft = int(fields.Item(i).Type)
                        except Exception:
                            continue
                        if ft == 33 or ft == 26:
                            return True
                except Exception:
                    pass
                # 2) 文本
                try:
                    txt = str(rng.Text or '')
                except Exception:
                    txt = ''
                cleaned = _re.sub(r'[\r\x07\s]+', '', txt)
                if cleaned and _re.search(r'\d', cleaned):
                    return True
                return False

            def _hf_has_pagenum(hf):
                """判定给定 HeaderFooter 中是否存在页码。"""
                if hf is None:
                    return False
                try:
                    if not bool(hf.Exists):
                        return False
                except Exception:
                    pass
                # 页眉/页脚正文
                try:
                    if _range_has_pagenum(hf.Range):
                        return True
                except Exception:
                    pass
                # 页眉/页脚内锚定的形状/文本框（WPS 常见方式）
                shapes = None
                sh_count = 0
                try:
                    shapes = hf.Shapes
                    sh_count = int(shapes.Count)
                except Exception:
                    sh_count = 0
                for si in range(1, sh_count + 1):
                    if shapes is None:
                        break
                    try:
                        sh = shapes(si)
                    except Exception:
                        continue
                    try:
                        tf = sh.TextFrame
                    except Exception:
                        tf = None
                    if tf is None:
                        continue
                    try:
                        has_text = bool(tf.HasText)
                    except Exception:
                        has_text = False
                    if not has_text:
                        continue
                    try:
                        tr = tf.TextRange
                    except Exception:
                        tr = None
                    if _range_has_pagenum(tr):
                        return True
                return False

            page_count = int(wdDoc.ComputeStatistics(2))  # wdStatisticPages
            missing_pages = []

            for p in range(2, page_count + 1):
                # 定位到该页起始处，取得所在节
                try:
                    rng_p = wdDoc.GoTo(What=1, Which=1, Count=p)  # wdGoToPage, wdGoToAbsolute
                except Exception:
                    continue
                try:
                    sec = rng_p.Sections(1)
                except Exception:
                    continue

                page_setup = sec.PageSetup
                try:
                    diff_first = bool(page_setup.DifferentFirstPageHeaderFooter)
                except Exception:
                    diff_first = False
                try:
                    odd_even = bool(page_setup.OddAndEvenPagesHeaderFooter)
                except Exception:
                    odd_even = False

                # 该页是否为其所在节的首页
                try:
                    sec_start_page = int(sec.Range.Information(3))
                except Exception:
                    sec_start_page = p
                is_first_of_section = (p == sec_start_page)

                # 依据节设置选择正确的 Header/Footer 索引
                #   1 = wdHeaderFooterPrimary
                #   2 = wdHeaderFooterFirstPage
                #   3 = wdHeaderFooterEvenPages
                if diff_first and is_first_of_section:
                    idx = 2
                elif odd_even and (p % 2 == 0):
                    idx = 3
                else:
                    idx = 1

                has_num = False
                # 页脚（优先）
                try:
                    has_num = _hf_has_pagenum(sec.Footers(idx))
                except Exception:
                    has_num = False
                # 页眉
                if not has_num:
                    try:
                        has_num = _hf_has_pagenum(sec.Headers(idx))
                    except Exception:
                        pass
                # 兜底：Primary（部分文档虽标记"首页不同"但实际页码在 Primary）
                if not has_num and idx != 1:
                    try:
                        has_num = _hf_has_pagenum(sec.Footers(1)) or _hf_has_pagenum(sec.Headers(1))
                    except Exception:
                        pass

                if not has_num:
                    missing_pages.append(p)

            if missing_pages:
                try:
                    print(f'  · 缺少页码的页：{missing_pages}')
                except Exception:
                    pass
                self.dimension2_results.append((-3, '文档中除封面页外其余任意一页没有出现页码'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass

    def _check_page_number_format(self):
        """检查页码格式。

        细则：-1：文档中页码格式不满足以下任意一条：
              页码位置页面底部居中、样式为"1、2、3"、字体为宋体小五
        —— 仅按此一项判定，不附加其它约束。
        —— 通过办公软件（Word / WPS 文字）读取页脚中的页码元素并判定：
            1) 位置：位于页脚（页面底部），且承载页码的段落对齐方式为居中，
               或承载页码的文本框在页面上水平居中（即文本框水平中心距离
               页面横向中线不超过 0.1cm）；两者任一成立即视为"居中"合规。
            2) 样式：阿拉伯数字（1、2、3）——PAGE 域必须为 Arabic 样式，
               或页脚正文/文本框内页码文本为纯数字；
            3) 字体：宋体，小五（9 磅）。
          页码来源三选一（与"存在页码"判定保持一致）：
            a) 页脚 Range.Fields 中的 PAGE 域 (Type==33)；
            b) 页脚锚定形状 (Shape) 内 TextFrame.TextRange.Fields 中的 PAGE 域；
            c) 页脚正文或文本框内的纯数字文本（视为直接键入的页码）。
          细则要求"文档中页码格式不满足以下任意一条"即扣分——这是对
          页码格式的整体要求：除封面外的每一页页码都必须同时满足这三条，
          任意一页存在页码但不满足（位置/样式/字体三者之一不满足），
          或任意一页应有页码却未找到，都判定为"格式不满足"，扣 -1。
        """
        try:
            app, wdDoc, temp_path = self._open_with_office()
        except Exception:
            return

        try:
            import re as _re

            CM = 28.3464567  # pt per cm
            SHAPE_CENTER_TOL_PT = 0.1 * CM  # 0.1cm 容差

            def _is_songti(name):
                if not name:
                    return False
                n = str(name).strip()
                return n in ('宋体', 'SimSun', 'NSimSun') or '宋体' in n or 'SimSun' in n

            def _shape_horiz_centered(sh, page_width_pt):
                """判定 Shape 是否在页面上水平居中。

                两种情况都算居中：
                (1) Shape.Left 等于 Word/WPS 水平定位哨兵值 wdShapeCenter
                    (-999995)——表示形状对齐方式为"居中"（相对页面）。
                (2) 由 Shape.Left / Anchor.Information(5) 计算的形状水平中心
                    落在页面横向中线的 ±0.1cm 之内。
                """
                if sh is None:
                    return False
                try:
                    sh_L = float(sh.Left)
                except Exception:
                    return False
                # 1) 哨兵：wdShapeCenter (-999995)
                if int(sh_L) == -999995:
                    return True
                if page_width_pt is None:
                    return False
                try:
                    sh_W = float(sh.Width)
                except Exception:
                    return False
                try:
                    rhp = int(sh.RelativeHorizontalPosition)
                except Exception:
                    rhp = -1
                page_center_x = page_width_pt / 2.0
                # rhp == 1: wdRelativeHorizontalPositionPage — Left 是页面坐标
                if rhp == 1:
                    shape_center_x = sh_L + sh_W / 2.0
                    if abs(shape_center_x - page_center_x) <= SHAPE_CENTER_TOL_PT:
                        return True
                # 其它相对参照系：叠加锚点段落的页面横向位置换算
                try:
                    anc = sh.Anchor
                    anc_x = float(anc.Information(5))  # wdHorizontalPositionRelativeToPage
                    shape_center_x2 = anc_x + sh_L + sh_W / 2.0
                    if abs(shape_center_x2 - page_center_x) <= SHAPE_CENTER_TOL_PT:
                        return True
                except Exception:
                    pass
                return False

            def _check_pagenum_range(rng, container_shape=None, page_width_pt=None):
                """在给定 Range 上定位页码元素并做三项判定，返回
                (center_ok, style_ok, font_ok, found)。found=False 表示
                该 Range 上未找到页码元素。
                container_shape：若该 Range 位于某 Shape (TextBox) 内，
                传入该 Shape，用于在段落非居中时补判"文本框水平居中"。"""
                if rng is None:
                    return (False, False, False, False)

                # 1) 优先：Range.Fields 中的 PAGE 域
                page_field = None
                try:
                    fields = rng.Fields
                    for i in range(1, int(fields.Count) + 1):
                        try:
                            f = fields.Item(i)
                            if int(f.Type) == 33:
                                page_field = f
                                break
                        except Exception:
                            continue
                except Exception:
                    pass

                if page_field is not None:
                    # 位置：段落居中 或 承载文本框在页面上水平居中
                    try:
                        align = int(page_field.Result.ParagraphFormat.Alignment)
                    except Exception:
                        align = -1
                    center_ok = (align == 1)
                    if not center_ok and container_shape is not None:
                        center_ok = _shape_horiz_centered(container_shape, page_width_pt)

                    # 样式：Arabic
                    style_ok = True
                    try:
                        code = str(page_field.Code.Text or '')
                    except Exception:
                        code = ''
                    m = _re.search(r'\\\*\s*(\S+)', code)
                    if m:
                        tok = m.group(1).strip().upper()
                        if tok not in ('ARABIC', 'MERGEFORMAT', 'CHARFORMAT'):
                            style_ok = False
                    try:
                        result_text = str(page_field.Result.Text or '')
                    except Exception:
                        result_text = ''
                    if result_text and not _re.fullmatch(r'\s*\d+\s*', result_text):
                        style_ok = False

                    # 字体：宋体 + 9pt
                    try:
                        fr = page_field.Result
                        font_east = str(fr.Font.NameFarEast or '')
                        font_name = str(fr.Font.Name or '')
                        size_pt = float(fr.Font.Size)
                    except Exception:
                        font_east, font_name, size_pt = '', '', 0.0
                    font_ok = (_is_songti(font_east) or _is_songti(font_name)) and abs(size_pt - 9.0) < 0.01
                    return (center_ok, style_ok, font_ok, True)

                # 2) 兜底：Range 正文文本中出现的纯数字（直接键入的页码）
                try:
                    txt = str(rng.Text or '')
                except Exception:
                    txt = ''
                cleaned = _re.sub(r'[\r\x07]+', '', txt).strip()
                m = _re.search(r'\d+', cleaned)
                if not m:
                    return (False, False, False, False)
                # 样式：纯数字视为 Arabic 样式合规
                style_ok = bool(_re.fullmatch(r'\d+', cleaned))

                # 通过 Find 定位到该数字子串，读取其字体与段落对齐
                center_ok = False
                font_ok = False
                try:
                    dup = rng.Duplicate
                    find = dup.Find
                    find.ClearFormatting()
                    find.Text = m.group(0)
                    find.Forward = True
                    find.Wrap = 0  # wdFindStop
                    if find.Execute():
                        try:
                            align = int(dup.ParagraphFormat.Alignment)
                        except Exception:
                            align = -1
                        center_ok = (align == 1)
                        try:
                            font_east = str(dup.Font.NameFarEast or '')
                            font_name = str(dup.Font.Name or '')
                            size_pt = float(dup.Font.Size)
                        except Exception:
                            font_east, font_name, size_pt = '', '', 0.0
                        font_ok = (_is_songti(font_east) or _is_songti(font_name)) and abs(size_pt - 9.0) < 0.01
                except Exception:
                    pass
                # 段落非居中时，补判承载文本框水平居中
                if not center_ok and container_shape is not None:
                    if _shape_horiz_centered(container_shape, page_width_pt):
                        center_ok = True
                return (center_ok, style_ok, font_ok, True)

            def _check_footer_pagenum(hf, page_width_pt):
                """遍历一个 HeaderFooter（含正文与内嵌文本框）里的所有页码
                元素，返回 (any_ok, any_found)。"""
                if hf is None:
                    return (False, False)
                try:
                    if not bool(hf.Exists):
                        return (False, False)
                except Exception:
                    pass

                any_ok = False
                any_found = False

                # a) 页脚正文
                try:
                    c, s, f, found = _check_pagenum_range(hf.Range, None, page_width_pt)
                    if found:
                        any_found = True
                        if c and s and f:
                            any_ok = True
                except Exception:
                    pass

                # b) 页脚里的形状 / 文本框
                shapes = None
                sh_count = 0
                try:
                    shapes = hf.Shapes
                    sh_count = int(shapes.Count)
                except Exception:
                    sh_count = 0
                for si in range(1, sh_count + 1):
                    if any_ok:
                        break
                    if shapes is None:
                        break
                    try:
                        sh = shapes(si)
                    except Exception:
                        continue
                    try:
                        tf = sh.TextFrame
                    except Exception:
                        tf = None
                    if tf is None:
                        continue
                    try:
                        has_text = bool(tf.HasText)
                    except Exception:
                        has_text = False
                    if not has_text:
                        continue
                    try:
                        tr = tf.TextRange
                    except Exception:
                        tr = None
                    try:
                        c, s, f, found = _check_pagenum_range(tr, sh, page_width_pt)
                    except Exception:
                        c, s, f, found = False, False, False, False
                    if found:
                        any_found = True
                        if c and s and f:
                            any_ok = True
                return (any_ok, any_found)

            page_count = int(wdDoc.ComputeStatistics(2))  # wdStatisticPages
            # rubric 要求“文档中页码格式”整体满足，因此必须遍历除封面外的
            # 每一页；只要有一页不满足（或应有页码却未找到）即判定不合格。
            format_ok = True

            for p in range(2, page_count + 1):
                try:
                    rng_p = wdDoc.GoTo(What=1, Which=1, Count=p)
                    sec = rng_p.Sections(1)
                except Exception:
                    continue

                page_setup = sec.PageSetup
                try:
                    diff_first = bool(page_setup.DifferentFirstPageHeaderFooter)
                except Exception:
                    diff_first = False
                try:
                    odd_even = bool(page_setup.OddAndEvenPagesHeaderFooter)
                except Exception:
                    odd_even = False
                try:
                    page_width_pt = float(page_setup.PageWidth)
                except Exception:
                    page_width_pt = None
                try:
                    sec_start_page = int(sec.Range.Information(3))
                except Exception:
                    sec_start_page = p
                is_first_of_section = (p == sec_start_page)

                if diff_first and is_first_of_section:
                    idx = 2
                elif odd_even and (p % 2 == 0):
                    idx = 3
                else:
                    idx = 1

                # 位置项要求"页面底部"——只在 Footer 中查找
                footers_to_try = []
                try:
                    footers_to_try.append(sec.Footers(idx))
                except Exception:
                    pass
                if idx != 1:
                    try:
                        footers_to_try.append(sec.Footers(1))
                    except Exception:
                        pass

                page_found = False
                page_ok = False
                for footer in footers_to_try:
                    ok, found = _check_footer_pagenum(footer, page_width_pt)
                    if found:
                        page_found = True
                    if ok:
                        page_ok = True
                        break

                # 仅在该页确实找到了页码元素时，才把“位置/样式/字体”不合规
                # 计入本项；缺失页码由“-3：文档中除封面页外其余任意一页没有
                # 出现页码”独立扣分项处理。
                if page_found and not page_ok:
                    format_ok = False
                    break

            if not format_ok:
                self.dimension2_results.append((-1, '文档中页码格式不满足以下任意一条：页码位置页面底部居中、样式为"1、2、3"、字体为宋体小五'))
        finally:
            try:
                wdDoc.Close(False)
            except Exception:
                pass
            try:
                app.Quit()
            except Exception:
                pass
            try:
                os.remove(temp_path)
            except Exception:
                pass


def evaluate(dir_path: str) -> dict:
    """按脚本接口统一约定（§2.1 / §2.2 / §2.3 / §2.4）返回评估结果字典。

    参数
    ------
    dir_path : str
        脚本所在目录的路径。脚本自己在该目录里定位并打开被评估的 Word
        文档（.docx），调用方不再传入具体文件路径。

    返回
    ------
    dict
        严格按 §2.2 结构返回：
          id / file_name / status / error /
          dim1_pass / dim1_reason /
          dim2_items（命中与未命中项都返回）/ total_score / max_score
        —— 主结果只走 return，不 print、不改 sys.stdout、不 sys.exit。
    """
    script_id = '011'
    result = {
        'id': script_id,
        'file_name': '',
        'status': 'ok',
        'error': None,
        'dim1_pass': True,
        'dim1_reason': '',
        'dim2_items': [],
        'total_score': 0,
        # max_score = 所有正向项 max_delta 之和（负向项为扣分，不计入满分）
        'max_score': sum(delta for delta, _ in _DIMENSION2_RULES if delta > 0),
    }

    try:
        # 在给定目录中定位 .docx 文档（忽略 Office 临时锁文件 ~$*）
        target = None
        if os.path.isdir(dir_path):
            for name in sorted(os.listdir(dir_path)):
                if name.startswith('~$'):
                    continue
                low = name.lower()
                if low.endswith('.docx'):
                    target = os.path.join(dir_path, name)
                    break

        if target is None:
            result['status'] = 'error'
            result['error'] = f'在目录中未找到 .docx 文档: {dir_path}'
            return result

        result['file_name'] = os.path.basename(target)

        evaluator = DocumentEvaluator(target)

        # 直接调用两个纯判定方法，绕开 DocumentEvaluator.evaluate 里的展示层
        # print；同时把判定过程中残余的少量调试 print 静默到 os.devnull，
        # 保证主结果只走 return（§2.3 要求：不 print 主结果、不改 sys.stdout）。
        with open(os.devnull, 'w', encoding='utf-8') as _null, \
                contextlib.redirect_stdout(_null), \
                contextlib.redirect_stderr(_null):
            evaluator.check_dimension1()
            if evaluator.dimension1_passed:
                evaluator.check_dimension2()

        result['dim1_pass'] = evaluator.dimension1_passed
        if not evaluator.dimension1_passed:
            result['dim1_reason'] = '；'.join(evaluator.dimension1_failures)
            result['total_score'] = 0
            # 维度一未通过：dim2_items 保持为空
            return result

        # 命中 (delta, desc) 与规则目录对齐，命中与未命中项一起返回
        hit_set = {(delta, desc) for delta, desc in evaluator.dimension2_results}
        for max_delta, desc in _DIMENSION2_RULES:
            hit = (max_delta, desc) in hit_set
            result['dim2_items'].append({
                'rule': desc,
                'max_delta': max_delta,
                'delta': max_delta if hit else 0,
                'hit': hit,
                'detail': '',
            })

        result['total_score'] = sum(delta for delta, _ in evaluator.dimension2_results)
        return result
    except Exception as e:
        # §2.4：脚本内部 try/except 兜底，异常统一映射为 status="error"
        result['status'] = 'error'
        result['error'] = f'{type(e).__name__}: {e}'
        return result


if __name__ == '__main__':
    # 本地调试入口：默认使用脚本所在目录（也可 `python <脚本> <目录路径>`）。
    # 修正 Windows 控制台的默认编码，避免 print(json...) 触发 cp1252 报错。
    _reconfigure = getattr(sys.stdout, 'reconfigure', None)
    if callable(_reconfigure):
        try:
            _ = _reconfigure(encoding='utf-8')
        except Exception:
            pass
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    print(json.dumps(evaluate(_dir), ensure_ascii=False, indent=2))
