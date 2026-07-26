#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 自动评分脚本：按题目给出的"打分细则"评估
《澜桥市高三英语试题_按周测二格式转换版.docx》。

对外只暴露一个统一入口：

    evaluate(dir_path: str) -> dict

其中 ``dir_path`` 是脚本所在目录路径，脚本自己在该目录里定位并打开被评估文档。
返回结构见 `脚本接口差异与统一建议.md` §2.2。

说明：
- 脚本优先使用纯 Python 读取 docx 包和 Word XML，对所有评分点做自动化检查。
- 页数估算完全基于文档 XML 中的分页符、`w:lastRenderedPageBreak` 与节结构，
  不启动 Microsoft Word COM，也不依赖 LibreOffice/soffice。
- 仅支持 .docx（Office Open XML）；.doc（老 CFB 二进制格式）不在支持范围内。
- 对"乱码、错位、超出页面、单词错误断开"等视觉类要求，脚本采用可复现的结构/文本启发式检测，不依赖人工。
"""

from __future__ import annotations

import json
import os
import re
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from statistics import mean
from typing import Any, Iterable, Optional

try:
    from docx import Document
    from docx.enum.section import WD_ORIENTATION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except Exception as exc:  # pragma: no cover - runtime dependency guard
    print("缺少依赖：python-docx。请先安装：pip install python-docx lxml", file=sys.stderr)
    raise

try:
    from lxml import etree
except Exception as exc:  # pragma: no cover
    print("缺少依赖：lxml。请先安装：pip install lxml", file=sys.stderr)
    raise

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
V = "{urn:schemas-microsoft-com:vml}"
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
}
# python-docx 的 BaseOxmlElement.xpath() 不接受 namespaces= 参数；
# lxml 原生元素接受该参数。统一封装，兼容两类 XML 元素。
def xp(el: Any, expr: str) -> list[Any]:
    try:
        return el.xpath(expr, namespaces=NS)
    except TypeError:
        return el.xpath(expr)


def fd(el: Any, expr: str) -> Any:
    try:
        return el.find(expr, namespaces=NS)
    except TypeError:
        return el.find(expr)


EMU_PER_CM = 360000
TWIPS_PER_CM = 567
PT_PER_TWIP = 1 / 20
CM_PER_INCH = 2.54
A4_W_CM = 21.0
A4_H_CM = 29.7
BODY_WIDTH_CM = A4_W_CM - 1.27 * 2
FONT_SZ_XIAOSI_HALF_PT = 24  # 小四 12pt in w:sz half-points
FONT_SZ_WU_HALF_PT = 21      # 五号 10.5pt in w:sz half-points
FONT_SZ_SI_HALF_PT = 28      # 四号 14pt in w:sz half-points


@dataclass
class PointResult:
    score: int
    name: str
    passed: bool
    detail: str = ""


@dataclass
class Dimension1Result:
    passed: bool
    failures: list[str] = field(default_factory=list)
    evidence: list[str] = field(default_factory=list)


@dataclass
class EvalContext:
    path: Path
    doc: Any
    zf: zipfile.ZipFile
    document_xml: etree._Element
    styles_xml: Optional[etree._Element]
    rels_xml: Optional[etree._Element]
    footer_xmls: list[etree._Element]
    comments_xml: Optional[etree._Element]
    settings_xml: Optional[etree._Element]
    all_text: str
    paragraphs: list[Any]
    tables: list[Any]


def cm_from_emu(value: Optional[int]) -> Optional[float]:
    if value is None:
        return None
    return float(value) / EMU_PER_CM


def cm_from_twips(value: Optional[str | int]) -> Optional[float]:
    if value in (None, ""):
        return None
    return float(value) / TWIPS_PER_CM


def approx(value: Optional[float], target: float, tol: float = 0.08) -> bool:
    return value is not None and abs(value - target) <= tol


def normalize_text(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def compact_text(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def paragraph_text_from_el(p_el: etree._Element) -> str:
    return "".join(xp(p_el, ".//w:t/text()"))


def get_zip_xml(zf: zipfile.ZipFile, name: str) -> Optional[etree._Element]:
    try:
        return etree.fromstring(zf.read(name))
    except KeyError:
        return None


def load_context(path: Path) -> EvalContext:
    zf = zipfile.ZipFile(path, "r")
    document_xml = etree.fromstring(zf.read("word/document.xml"))
    styles_xml = get_zip_xml(zf, "word/styles.xml")
    rels_xml = get_zip_xml(zf, "word/_rels/document.xml.rels")
    comments_xml = get_zip_xml(zf, "word/comments.xml")
    settings_xml = get_zip_xml(zf, "word/settings.xml")
    footer_xmls: list[etree._Element] = []
    for name in sorted(zf.namelist()):
        if re.fullmatch(r"word/footer\d+\.xml", name):
            xml = get_zip_xml(zf, name)
            if xml is not None:
                footer_xmls.append(xml)
    doc = Document(str(path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            all_text += "\n" + "\t".join(cell.text for cell in row.cells)
    return EvalContext(
        path=path, doc=doc, zf=zf, document_xml=document_xml, styles_xml=styles_xml,
        rels_xml=rels_xml, footer_xmls=footer_xmls, comments_xml=comments_xml,
        settings_xml=settings_xml, all_text=all_text, paragraphs=list(doc.paragraphs),
        tables=list(doc.tables)
    )


def count_doc_page_breaks(ctx: EvalContext) -> int:
    """估算页数。

    硬分页符 (w:br[@type='page'])、节分页 (w:sectPr，最后一个属于文档级不算分页)、
    以及 Word 上次保存时写入的渲染分页标记 (w:lastRenderedPageBreak) 都计入。
    后者是关键：它让"自动分页/软分页"也能被静态识别，避免把 8 页文档误算成 1 页。
    """
    br_pages = len(xp(ctx.document_xml, ".//w:br[@w:type='page']"))
    sect_pages = max(0, len(xp(ctx.document_xml, ".//w:sectPr")) - 1)
    last_rendered = len(xp(ctx.document_xml, ".//w:lastRenderedPageBreak"))
    return br_pages + sect_pages + last_rendered + 1


def field_texts_in_footers(ctx: EvalContext) -> str:
    pieces: list[str] = []
    for f in ctx.footer_xmls:
        pieces.extend(xp(f, ".//w:t/text()"))
        pieces.extend(xp(f, ".//w:instrText/text()"))
        for fld in xp(f, ".//w:fldChar"):
            pieces.append(fld.get(W + "fldCharType", ""))
    return " ".join(pieces)


def run_xml_has_font(run_el: etree._Element, font_name: str) -> bool:
    r_fonts = fd(run_el, "w:rPr/w:rFonts")
    if r_fonts is None:
        return False
    vals = [r_fonts.get(W + k) for k in ("ascii", "hAnsi", "eastAsia", "cs")]
    return any(v == font_name for v in vals if v)


def run_xml_ascii_font(run_el: etree._Element, font_name: str) -> bool:
    r_fonts = fd(run_el, "w:rPr/w:rFonts")
    if r_fonts is None:
        return False
    return r_fonts.get(W + "ascii") == font_name and r_fonts.get(W + "hAnsi") == font_name


def run_xml_size(run_el: etree._Element) -> Optional[int]:
    sz = fd(run_el, "w:rPr/w:sz")
    if sz is not None and sz.get(W + "val"):
        try:
            return int(sz.get(W + "val"))
        except Exception:
            return None
    return None


def run_xml_color(run_el: etree._Element) -> Optional[str]:
    color = fd(run_el, "w:rPr/w:color")
    if color is not None:
        return (color.get(W + "val") or "").upper()
    return None


def run_xml_bold(run_el: etree._Element) -> bool:
    b = fd(run_el, "w:rPr/w:b")
    if b is None:
        return False
    return b.get(W + "val", "1") not in ("0", "false", "False")


def run_xml_italic(run_el: etree._Element) -> bool:
    i = fd(run_el, "w:rPr/w:i")
    if i is None:
        return False
    return i.get(W + "val", "1") not in ("0", "false", "False")


def paragraph_alignment(p: Any) -> Optional[int]:
    return p.paragraph_format.alignment


def alignment_value(alignment: Any) -> Optional[int]:
    if alignment is None:
        return None
    try:
        return int(alignment)
    except Exception:
        return None


def effective_paragraph_alignment(p: Any) -> Optional[int]:
    """Return direct paragraph alignment, falling back to the paragraph style.

    python-docx returns None when alignment is inherited.  For strict checks,
    None must not be treated as left/justify by itself; use the style value if
    present, otherwise the effective alignment is unknown.
    """
    direct = paragraph_alignment(p)
    if direct is not None:
        return alignment_value(direct)
    try:
        return alignment_value(p.style.paragraph_format.alignment)
    except Exception:
        return None


def is_strict_left(p: Any) -> bool:
    return effective_paragraph_alignment(p) == int(WD_ALIGN_PARAGRAPH.LEFT)


def is_strict_justify(p: Any) -> bool:
    return effective_paragraph_alignment(p) == int(WD_ALIGN_PARAGRAPH.JUSTIFY)


def is_center(p: Any) -> bool:
    return alignment_value(paragraph_alignment(p)) == int(WD_ALIGN_PARAGRAPH.CENTER)


def is_left_or_none(p: Any) -> bool:
    return paragraph_alignment(p) in (None, WD_ALIGN_PARAGRAPH.LEFT)


def iter_document_paragraph_elements(ctx: EvalContext) -> list[etree._Element]:
    return xp(ctx.document_xml, ".//w:body/w:p | .//w:body/w:tbl//w:p")


def find_paragraphs(ctx: EvalContext, pattern: str, flags: int = 0) -> list[Any]:
    rx = re.compile(pattern, flags)
    return [p for p in ctx.paragraphs if rx.search(p.text)]


def para_xml_by_text(ctx: EvalContext, needle: str) -> Optional[etree._Element]:
    for p in xp(ctx.document_xml, ".//w:p"):
        if needle in paragraph_text_from_el(p):
            return p
    return None


def para_by_exact_or_contains(ctx: EvalContext, needle: str) -> Optional[Any]:
    for p in ctx.paragraphs:
        if needle == p.text.strip() or needle in p.text:
            return p
    return None


def size_ok_for_runs(p_el: etree._Element, target_half_pt: int, tolerance: int = 1) -> bool:
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    if not runs:
        return False
    explicit = [run_xml_size(r) for r in runs]
    explicit = [s for s in explicit if s is not None]
    if not explicit:
        # 纯样式继承难以完全展开，作为自动评估的容错：无显式字号不直接通过。
        return False
    return sum(abs(s - target_half_pt) <= tolerance for s in explicit) / len(explicit) >= 0.8


def color_black_ok(p_el: etree._Element) -> bool:
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    if not runs:
        return False
    colors = [run_xml_color(r) for r in runs]
    # 未显式设置颜色在 Word 中通常为自动黑色。
    return all(c in (None, "", "000000", "AUTO") for c in colors)


def font_mix_ok(p_el: etree._Element, require_bold: Optional[bool] = None) -> bool:
    """检查中文宋体、英文/数字 Times New Roman 的意图；允许样式继承导致的缺省。"""
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    if not runs:
        return False
    checked = 0
    ok = 0
    for r in runs:
        txt = "".join(xp(r, ".//w:t/text()"))
        if not txt.strip():
            continue
        if require_bold is not None and run_xml_bold(r) != require_bold:
            return False
        has_cn = re.search(r"[一-鿿]", txt) is not None
        has_en = re.search(r"[A-Za-z0-9]", txt) is not None
        if has_cn:
            checked += 1
            if run_xml_has_font(r, "宋体"):
                ok += 1
        if has_en:
            checked += 1
            if run_xml_has_font(r, "Times New Roman"):
                ok += 1
    return checked > 0 and ok / checked >= 0.65


def font_mix_exact_ok(p_el: etree._Element, require_bold: Optional[bool] = None) -> bool:
    """严格检查中文 eastAsia=宋体，英文/数字 ascii+hAnsi=Times New Roman。"""
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    if not runs:
        return False
    checked = 0
    ok = 0
    for r in runs:
        txt = "".join(xp(r, ".//w:t/text()"))
        if not txt.strip():
            continue
        if require_bold is not None and run_xml_bold(r) != require_bold:
            return False
        r_fonts = fd(r, "w:rPr/w:rFonts")
        if r_fonts is None:
            continue
        if re.search(r"[一-鿿]", txt):
            checked += 1
            if r_fonts.get(W + "eastAsia") == "宋体":
                ok += 1
        if re.search(r"[A-Za-z0-9]", txt):
            checked += 1
            if (r_fonts.get(W + "ascii") == "Times New Roman" and
                    r_fonts.get(W + "hAnsi") == "Times New Roman"):
                ok += 1
    return checked > 0 and ok / checked >= 0.8


def is_bold_para(p_el: etree._Element) -> bool:
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    return bool(runs) and sum(run_xml_bold(r) for r in runs) / len(runs) >= 0.8


def _para_style_bold(ctx: EvalContext, style_id: Optional[str]) -> Optional[bool]:
    """沿样式继承链查找 w:rPr/w:b；找到返回 True/False，未声明返回 None。"""
    if not style_id or ctx.styles_xml is None:
        return None
    seen: set[str] = set()
    cur = style_id
    while cur and cur not in seen:
        seen.add(cur)
        style = None
        for s in xp(ctx.styles_xml, ".//w:style"):
            if s.get(W + "styleId") == cur:
                style = s
                break
        if style is None:
            return None
        b = fd(style, "w:rPr/w:b")
        if b is not None:
            return b.get(W + "val", "1") not in ("0", "false", "False")
        based = fd(style, "w:basedOn")
        cur = based.get(W + "val") if based is not None else None
    return None


def _effective_paragraph_bold(ctx: EvalContext, p_el: etree._Element) -> bool:
    """段内每个含字符 run 的有效加粗状态（run 直接 w:b → 段落样式链 → 默认段落样式）必须为真。"""
    p_style = fd(p_el, "w:pPr/w:pStyle")
    style_id = p_style.get(W + "val") if p_style is not None else _default_paragraph_style_id(ctx)
    style_bold = _para_style_bold(ctx, style_id)
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    if not runs:
        return False
    for r in runs:
        b = fd(r, "w:rPr/w:b")
        if b is not None:
            on = b.get(W + "val", "1") not in ("0", "false", "False")
        else:
            on = bool(style_bold)
        if not on:
            return False
    return True


def para_spacing_ok(p: Any, line_pt: float = 20.0) -> bool:
    pf = p.paragraph_format
    before = pf.space_before.pt if pf.space_before else 0
    after = pf.space_after.pt if pf.space_after else 0
    line = pf.line_spacing.pt if hasattr(pf.line_spacing, "pt") else None
    spacing = fd(p._p, "w:pPr/w:spacing")
    line_rule = spacing.get(W + "lineRule") if spacing is not None else None
    exact_rule = line_rule == "exact"
    return (abs(before - 0) < 0.2 and abs(after - 0) < 0.2 and
            line is not None and abs(line - line_pt) <= 0.5 and exact_rule)


def _para_style_spacing(ctx: EvalContext, style_id: Optional[str]) -> dict[str, Optional[Any]]:
    """沿样式继承链解析 w:pPr/w:spacing 的 before/after/line/lineRule。"""
    out: dict[str, Optional[Any]] = {"before": None, "after": None, "line": None, "lineRule": None}
    if not style_id or ctx.styles_xml is None:
        return out
    seen: set[str] = set()
    cur = style_id
    while cur and cur not in seen:
        seen.add(cur)
        style = None
        for s in xp(ctx.styles_xml, ".//w:style"):
            if s.get(W + "styleId") == cur:
                style = s
                break
        if style is None:
            break
        sp = fd(style, "w:pPr/w:spacing")
        if sp is not None:
            for k_xml, k_out in (("before", "before"), ("after", "after"),
                                 ("line", "line"), ("lineRule", "lineRule")):
                if out[k_out] is None and sp.get(W + k_xml) is not None:
                    out[k_out] = sp.get(W + k_xml)
        based = fd(style, "w:basedOn")
        cur = based.get(W + "val") if based is not None else None
    return out


def _effective_paragraph_spacing(ctx: EvalContext, p_el: etree._Element) -> dict[str, Any]:
    """段落级 w:spacing 优先，缺失字段沿样式继承链补；返回 before/after(磅)、line(磅)、lineRule。

    twips→pt：w:before / w:after / w:line（lineRule=exact|atLeast 时）单位是 twips（1pt=20twips）。
    lineRule=auto 时 w:line 单位是 240ths（240 = 单倍）。
    """
    direct = fd(p_el, "w:pPr/w:spacing")
    direct_vals: dict[str, Optional[str]] = {"before": None, "after": None, "line": None, "lineRule": None}
    if direct is not None:
        for k in direct_vals:
            v = direct.get(W + k)
            if v is not None:
                direct_vals[k] = v
    p_style = fd(p_el, "w:pPr/w:pStyle")
    style_id = p_style.get(W + "val") if p_style is not None else _default_paragraph_style_id(ctx)
    style_vals = _para_style_spacing(ctx, style_id)
    merged = {k: (direct_vals[k] if direct_vals[k] is not None else style_vals[k])
              for k in direct_vals}
    out: dict[str, Any] = {
        "before_pt": float(merged["before"]) / 20.0 if merged["before"] is not None else 0.0,
        "after_pt":  float(merged["after"])  / 20.0 if merged["after"]  is not None else 0.0,
        "lineRule": merged["lineRule"],
        "line_raw": merged["line"],
    }
    if merged["line"] is not None:
        try:
            line_int = int(merged["line"])
            if merged["lineRule"] in ("exact", "atLeast"):
                out["line_pt"] = line_int / 20.0
            else:
                # auto: 240ths of a line; 不直接换算磅，标记为 None
                out["line_pt"] = None
        except Exception:
            out["line_pt"] = None
    else:
        out["line_pt"] = None
    return out


def _para_style_indent(ctx: EvalContext, style_id: Optional[str]) -> dict[str, Optional[str]]:
    out: dict[str, Optional[str]] = {"firstLine": None, "firstLineChars": None,
                                     "left": None, "leftChars": None, "hanging": None}
    if not style_id or ctx.styles_xml is None:
        return out
    seen: set[str] = set()
    cur = style_id
    while cur and cur not in seen:
        seen.add(cur)
        style = None
        for s in xp(ctx.styles_xml, ".//w:style"):
            if s.get(W + "styleId") == cur:
                style = s
                break
        if style is None:
            break
        ind = fd(style, "w:pPr/w:ind")
        if ind is not None:
            for k in out:
                if out[k] is None and ind.get(W + k) is not None:
                    out[k] = ind.get(W + k)
        based = fd(style, "w:basedOn")
        cur = based.get(W + "val") if based is not None else None
    return out


def _effective_first_line_indent(ctx: EvalContext, p_el: etree._Element) -> dict[str, Any]:
    """段落首行缩进：w:firstLineChars（百分*100，200=2字符）优先，再看 w:firstLine（twips）。
    都解析样式继承链。"""
    direct = fd(p_el, "w:pPr/w:ind")
    direct_vals: dict[str, Optional[str]] = {"firstLine": None, "firstLineChars": None, "hanging": None}
    if direct is not None:
        for k in direct_vals:
            v = direct.get(W + k)
            if v is not None:
                direct_vals[k] = v
    p_style = fd(p_el, "w:pPr/w:pStyle")
    style_id = p_style.get(W + "val") if p_style is not None else _default_paragraph_style_id(ctx)
    style_vals = _para_style_indent(ctx, style_id)
    fl_chars = direct_vals["firstLineChars"] if direct_vals["firstLineChars"] is not None else style_vals["firstLineChars"]
    fl_twips = direct_vals["firstLine"] if direct_vals["firstLine"] is not None else style_vals["firstLine"]
    hanging = direct_vals["hanging"] if direct_vals["hanging"] is not None else style_vals["hanging"]
    return {"firstLineChars": fl_chars, "firstLine_twips": fl_twips, "hanging_twips": hanging}


def _is_first_line_indent_2chars(indent: dict[str, Any], font_size_half_pt: int) -> bool:
    """判定首行缩进是否等于"2 字符"。

    判定优先级：
      1) w:firstLineChars == 200（细则字面"2 字符"）→ True；非 200 但已设置 → False。
      2) 退化为 twips 比较：2 字符宽 ≈ font_size_pt(=12 for 小四) × 1 字符宽。
         小四 Times New Roman / 宋体下，Word 标准 2 字符 = 1.85cm(中文) 或 0.85cm(英文)。
         为兼容两种文字段，用 0.7cm–1.95cm（≈ 397–1106 twips）作为退化范围；
         若 firstLineChars 与 firstLine 均缺失，返回 False。
    """
    chars = indent.get("firstLineChars")
    if chars is not None:
        try:
            return int(chars) == 200
        except Exception:
            return False
    tw = indent.get("firstLine_twips")
    if tw is None:
        return False
    try:
        twips = int(tw)
    except Exception:
        return False
    return 397 <= twips <= 1106


def p_body_spacing(ctx: EvalContext) -> PointResult:
    """+5 整篇正文行距/段距 + 英文文章首行缩进2字符两端对齐 + 段间不增加空行。

    a. 范围：body 下所有段落（含表格内段），排除"标题"——保密文本/题头/日期/章节/分节标题/A-D 文章编号/课程表表头。
    b. 行距 = 固定值 20 磅（lineRule=exact 且 line=400 twips=20 pt）；段前=0；段后=0；继承样式后必须满足。
    c. 英文文章普通段落（英文主导的正文散文段，非题号段、非选项行、非标题）：
        - 首行缩进 = 2 字符（优先 firstLineChars=200）
        - 两端对齐（解析样式继承）
        依据"英文主导 + 成句（≥2 个英文单词）"的结构特征定位，不再用 80 字母阈值作为唯一条件，
        以免短英文文章普通段被漏检。
    d. 段落之间不增加空行：除"被允许的空段位置"外（紧挨表格、紧挨分页符的空段不计入），任何空段视为违规。
    """
    # 排除"标题"段：精确文本或前缀正则覆盖
    title_re = re.compile(
        r"^(绝密★启用前$|澜桥市2026|英语试题$|2026年6月|"
        r"第[一二三四]部分|第[一二三四]节|"
        r"阅读下列短文|阅读下面短文|根据短文内容|根据下面|根据所给|"
        r"假定你是|假定你叫|"
        r"注意：$|注意:$|"
        r"[A-D]$)")

    body_paras = _all_body_paragraph_elements(ctx)
    # 课程信息表表头（5×5 表第 1 行）
    table_header_texts: set[str] = set()
    for tbl in xp(ctx.document_xml, ".//w:body//w:tbl"):
        first_row = fd(tbl, ".//w:tr")
        if first_row is not None:
            for tc in xp(first_row, ".//w:tc"):
                txt = "".join(xp(tc, ".//w:t/text()")).strip()
                if txt:
                    table_header_texts.add(txt)

    py_by_id: dict[int, Any] = {id(p._p): p for p in ctx.paragraphs}
    # 表格内段也要 python-docx 对象用于 paragraph_format
    for tbl in ctx.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    py_by_id[id(p._p)] = p

    # 标记每个段所在 body 顺序索引，便于"相邻分页符/相邻表格"豁免空段
    body_order = _body_block_elements(ctx)
    order_index: dict[int, int] = {id(el): i for i, el in enumerate(body_order)}

    def _is_neighbor_of_break_or_table(idx: int) -> bool:
        """空段所在位置紧邻表格或分页符——这种空段不计入"段间空行"。"""
        for j in (idx - 1, idx + 1):
            if 0 <= j < len(body_order):
                neighbor = body_order[j]
                if neighbor.tag == W + "tbl":
                    return True
                if xp(neighbor, ".//w:br[@w:type='page']"):
                    return True
        # 段自身末尾的分页符也豁免
        if 0 <= idx < len(body_order):
            if xp(body_order[idx], ".//w:br[@w:type='page']"):
                return True
        return False

    # 校验目标段（"标题以外的题目、文章、选项、说明文字"）
    targets: list[etree._Element] = []
    for el in body_paras:
        t = _block_text(el).strip()
        if not t:
            continue
        if title_re.match(t):
            continue
        if t in table_header_texts:
            continue
        # 答题横线段（纯下划线）属于答题区，按"说明文字"对待还是排除？细则未单列，
        # 把它纳入会因结构性原因失败；这里视为非"题目/文章/选项/说明"的横线段，跳过。
        if re.fullmatch(r"[_\s]{20,}", t):
            continue
        targets.append(el)

    if not targets:
        return PointResult(5, "+5 正文行距20磅/段前后0/英文段首行缩进2字符两端对齐/段间无空行",
                           False, "无可检查正文段落")

    bad_spacing: list[str] = []  # (段落预览, 原因)
    for el in targets:
        sp = _effective_paragraph_spacing(ctx, el)
        text_preview = _block_text(el).strip()[:25]
        reasons: list[str] = []
        if sp["before_pt"] > 0.2:
            reasons.append(f"段前={sp['before_pt']:.1f}pt")
        if sp["after_pt"] > 0.2:
            reasons.append(f"段后={sp['after_pt']:.1f}pt")
        if sp["lineRule"] != "exact":
            reasons.append(f"lineRule={sp['lineRule']}")
        elif sp["line_pt"] is None or abs(sp["line_pt"] - 20.0) > 0.5:
            reasons.append(f"行距={sp['line_pt']}pt")
        if reasons:
            bad_spacing.append(f"'{text_preview}'({','.join(reasons)})")

    # 英文文章普通段：结构化识别，而不是用"字母数 > 80"作为唯一条件。
    #
    # rubric 里的"英文文章"指阅读理解、完形填空等题干中的英文散文段落。这些段有以下共同结构特征：
    #   - 段落文本以英文为主导（英文字母数 ≥ 中文字符数，且远多于其它符号）；
    #   - 至少构成一个可辨识的句子（≥2 个由字母组成的"英文单词"，且总字母数 ≥ 20，避免只有一两个词的短句被误纳入）；
    #   - 非题号段（"数字." 或 "数字．" 开头，如 "21. What ..."）；
    #   - 非选项行（"A. ..." "B. ..." "C. ..." "D. ..."，以及 "21. A. ..."）；
    #   - 非表格内段（题干英文散文一律出现在正文流中，不在 5×5 课程信息表这类布局表格里）。
    # 这样短英文文章普通段也能被覆盖，同时避免把选项/题号/中文说明误判为英文文章段。

    def is_option_paragraph(text: str) -> bool:
        return bool(re.match(r"^[A-D][\.．]\s", text) or re.match(r"^\d+[\.．]\s*A[\.．]\s", text))

    def is_question_number_paragraph(text: str) -> bool:
        return bool(re.match(r"^\d+[\.．]", text))

    # 收集所有位于表格内的段落 id，用于排除。
    in_table_ids: set[int] = set()
    for tbl_el in xp(ctx.document_xml, ".//w:body//w:tbl"):
        for p_in_tbl in xp(tbl_el, ".//w:p"):
            in_table_ids.add(id(p_in_tbl))

    def is_english_body_paragraph(text: str, el: etree._Element) -> bool:
        # 表格内段不属于英文文章正文散文段
        if id(el) in in_table_ids:
            return False
        letters = re.findall(r"[A-Za-z]", text)
        cn_chars = re.findall(r"[一-鿿]", text)
        words = re.findall(r"[A-Za-z]+", text)
        if not letters:
            return False
        # 英文主导
        if len(letters) < max(20, len(cn_chars)):
            return False
        # 至少两个英文单词（成句结构），避免把 "ID No." "PDF"、"AI" 之类零星词汇误判为英文文章段
        if len(words) < 2:
            return False
        return True

    english_targets: list[etree._Element] = []
    for el in targets:
        t = _block_text(el).strip()
        if is_question_number_paragraph(t):
            continue
        if is_option_paragraph(t):
            continue
        if not is_english_body_paragraph(t, el):
            continue
        english_targets.append(el)

    bad_indent: list[str] = []
    bad_justify: list[str] = []
    for el in english_targets:
        text_preview = _block_text(el).strip()[:25]
        ind = _effective_first_line_indent(ctx, el)
        if not _is_first_line_indent_2chars(ind, FONT_SZ_XIAOSI_HALF_PT):
            bad_indent.append(f"'{text_preview}'(firstLineChars={ind['firstLineChars']},firstLine={ind['firstLine_twips']}twips)")
        py = py_by_id.get(id(el))
        align_val = effective_paragraph_alignment(py) if py is not None else None
        if align_val != int(WD_ALIGN_PARAGRAPH.JUSTIFY):
            bad_justify.append(f"'{text_preview}'(alignment={align_val})")

    # 段间空行：扫 body 顺序，统计被认定为"额外空行"的空段
    extra_blanks: list[int] = []
    for i, el in enumerate(body_order):
        if el.tag != W + "p":
            continue
        if _block_text(el).strip():
            continue
        # 段内含 drawing/pict（如水印、图片）不算空段
        if xp(el, ".//w:drawing") or xp(el, ".//w:pict"):
            continue
        if _is_neighbor_of_break_or_table(i):
            continue
        # 文档首尾可能有平台自动留的空段（比如末尾一个空 sectPr 段），首尾各允许 1 个
        if i == 0 or i == len(body_order) - 1:
            continue
        extra_blanks.append(i)

    passed = (not bad_spacing and not bad_indent and not bad_justify and not extra_blanks)
    parts: list[str] = [
        f"目标段={len(targets)}；英文段={len(english_targets)}",
    ]
    if bad_spacing:
        parts.append(f"行距/段距异常({len(bad_spacing)})：" + "；".join(bad_spacing[:3]))
    if bad_indent:
        parts.append(f"首行缩进非2字符({len(bad_indent)})：" + "；".join(bad_indent[:3]))
    if bad_justify:
        parts.append(f"非两端对齐({len(bad_justify)})：" + "；".join(bad_justify[:3]))
    if extra_blanks:
        parts.append(f"额外空行({len(extra_blanks)})位置={extra_blanks[:5]}")
    if passed:
        parts.append("全部满足")
    detail = " | ".join(parts)
    return PointResult(5, "+5 正文行距20磅/段前后0/英文段首行缩进2字符两端对齐/段间无空行",
                       passed, detail)


def first_line_indent_cm(p: Any) -> Optional[float]:
    ind = p.paragraph_format.first_line_indent
    return cm_from_emu(ind) if ind is not None else 0.0


def left_indent_cm(p: Any) -> Optional[float]:
    ind = p.paragraph_format.left_indent
    return cm_from_emu(ind) if ind is not None else 0.0


def image_parts(ctx: EvalContext) -> list[dict[str, Any]]:
    parts: list[dict[str, Any]] = []
    names = ctx.zf.namelist()
    for n in names:
        if n.startswith("word/media/"):
            try:
                size = len(ctx.zf.read(n))
            except Exception:
                size = 0
            parts.append({"name": n, "bytes": size})
    # Drawing extents in document XML, in EMU.
    extents = xp(ctx.document_xml, ".//wp:extent")
    for i, ext in enumerate(extents):
        if i < len(parts):
            try:
                parts[i]["w_cm"] = int(ext.get("cx")) / EMU_PER_CM
                parts[i]["h_cm"] = int(ext.get("cy")) / EMU_PER_CM
            except Exception:
                pass
    # VML image shape styles.
    for shape in xp(ctx.document_xml, ".//v:shape"):
        style = shape.get("style", "")
        m_w = re.search(r"width:([0-9.]+)pt", style)
        m_h = re.search(r"height:([0-9.]+)pt", style)
        info = {"name": "VML-shape", "bytes": 0}
        if m_w:
            info["w_cm"] = float(m_w.group(1)) / 72 * CM_PER_INCH
        if m_h:
            info["h_cm"] = float(m_h.group(1)) / 72 * CM_PER_INCH
        parts.append(info)
    return parts


def has_full_page_image(ctx: EvalContext) -> tuple[bool, str]:
    """检测是否存在"整页大图"——仅按尺寸判定（≈ A4 正文区 18.46×27.16cm，容差 10%）。
    细则字面只点名"整页 PDF 截图"，未涉及文件大小，故不再用文件字节数作为判据。
    """
    imgs = image_parts(ctx)
    if not imgs:
        return False, "未发现图片。"
    suspicious = []
    for img in imgs:
        w = img.get("w_cm")
        h = img.get("h_cm")
        if w and h and w >= 16.5 and h >= 24.0:
            suspicious.append(f"{img.get('name')} 尺寸约 {w:.1f}×{h:.1f}cm")
    return bool(suspicious), "；".join(suspicious) if suspicious else f"发现 {len(imgs)} 个小图片/形状，未达整页截图阈值。"


def gibberish_ratio(text: str) -> float:
    if not text:
        return 1.0
    bad_chars = len(re.findall(r"[�□■¤�]{1}", text))
    weird = len(re.findall(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", text))
    return (bad_chars + weird) / max(1, len(text))


def detect_vertical_text(ctx: EvalContext) -> list[str]:
    bad = []
    for p in xp(ctx.document_xml, ".//w:p"):
        text = paragraph_text_from_el(p).strip()
        if len(text) <= 1:
            continue
        text_dir = fd(p, ".//w:textDirection")
        if text_dir is not None:
            bad.append(text[:30])
        # 每个字符被拆成大量独立 run，且平均 run 长度很短，视为逐字排版风险。
        runs = ["".join(xp(r, ".//w:t/text()")) for r in xp(p, ".//w:r")]
        non_empty = [r for r in runs if r.strip()]
        if len(non_empty) >= 10 and mean([len(r.strip()) for r in non_empty]) <= 1.2:
            bad.append(text[:30])
    return bad[:5]


def detect_answers_or_analysis(text: str) -> list[str]:
    """检测新增答案/解析。

    不能把语法填空正文中的"57   digital reading ..."这类题内空格误判为答案，
    因此只对明确答案/解析标题，或连续成串的"题号+选项答案"做判定。
    """
    hits = []
    explicit_patterns = [r"参考答案", r"答案与解析", r"答案解析", r"解析[:：]", r"详解[:：]", r"听力原文"]
    for pat in explicit_patterns:
        m = re.search(pat, text, re.I)
        if m:
            hits.append(m.group(0))

    # 连续答案串，如"21. B 22. A 23. D"或"21 B\n22 A\n23 D"。
    answer_pairs = list(re.finditer(r"(?<!第)\b(?:2[1-9]|3[0-5]|4[1-9]|5[0-5])\s*[.．、:：]?\s*[ABCD](?=\s|$)", text))
    for i in range(0, max(0, len(answer_pairs) - 2)):
        window = text[answer_pairs[i].start():answer_pairs[i + 2].end()]
        if len(window) <= 80:
            hits.append(normalize_text(window))
            break
    return hits[:8]


def detect_unrelated_marks(ctx: EvalContext) -> list[str]:
    """检测细则点名的 4 类"无关标记"（"新增空白页"由 n_unrelated_marks 单独处理）：
      1) 与任务无关的批注
      2) 红色箭头（红色 + 箭头同时满足）
      3) 修订标记
      4) 转换软件水印
    """
    hits: list[str] = []
    # 1) 批注
    if ctx.comments_xml is not None and xp(ctx.comments_xml, ".//w:comment"):
        hits.append("存在批注")
    # 3) 修订
    if xp(ctx.document_xml, ".//w:ins | .//w:del | .//w:moveFrom | .//w:moveTo"):
        hits.append("存在修订标记")
    # 4) 转换软件水印关键词
    text = ctx.all_text
    for pat in [r"水印", r"转换软件", r"converted by", r"trial version",
                r"Wondershare", r"福昕", r"Smallpdf", r"CamScanner", r"扫描全能王"]:
        if re.search(pat, text, re.I):
            hits.append(f"疑似转换软件水印：{pat}")
    # 2) 红色箭头（红色 + 箭头 同时满足）
    arrow_re = re.compile(r"[→←↑↓➜➔⇒]")
    if arrow_re.search(text):
        # 同时确认存在红色颜色定义
        for color in xp(ctx.document_xml, ".//w:color"):
            if (color.get(W + "val") or "").upper() in {"FF0000", "C00000", "E60000"}:
                hits.append("存在红色箭头标注")
                break
    return hits[:8]


def detect_bad_word_breaks(text: str) -> list[str]:
    hits = []
    # 常见 OCR/换行错误：英文字母-换行-英文字母，或单词中插入空格且两侧片段都不像独立单词。
    for m in re.finditer(r"[A-Za-z]{2,}-\s*\n\s*[A-Za-z]{2,}", text):
        hits.append(m.group(0).replace("\n", "\\n"))
    # 过多的单字母空格序列通常表示单词被逐字拆开。
    for m in re.finditer(r"\b(?:[A-Za-z]\s+){4,}[A-Za-z]\b", text):
        hits.append(m.group(0))
    return hits[:8]


def content_missing(ctx: EvalContext) -> list[str]:
    text = ctx.all_text
    missing = []
    for art in ["A", "B", "C", "D"]:
        if not re.search(rf"(?m)^\s*{art}\s*$", text):
            missing.append(f"文章编号 {art}")
    for q in range(56, 66):
        if not re.search(rf"\b{q}\b", text):
            missing.append(f"第{q}题")
    for q in (66, 67):
        if not re.search(rf"\b{q}\.", text):
            missing.append(f"第{q}题写作")
    key_phrases = [
        "Learning with a Clear Purpose",
        "Paragraph 1: I sat with him on the bench near the stairway and opened his notebook.",
        "Paragraph 2: When I returned to my own classroom on Monday, I found a small card on my desk.",
    ]
    for phrase in key_phrases:
        if phrase not in text:
            missing.append(phrase[:40])
    return missing


def table_content_missing(ctx: EvalContext) -> list[str]:
    expected_cells = [
        "Programme", "Date (2026)", "Time", "Fee (credits)", "Main focus",
        "Reading Evidence Lab", "18 July", "9:00-16:00", "46", "finding hidden logic",
        "Writing Clinic", "1 August", "9:30-16:30", "52", "argument and style",
        "Grammar Rescue", "15 August", "8:50-15:50", "43", "long sentence analysis",
        "Mock Interview Studio", "29 August", "10:00-17:00", "49", "spoken presentation",
    ]
    table_text = "\n".join("\t".join(cell.text for cell in row.cells) for t in ctx.tables for row in t.rows)
    return [x for x in expected_cells if x not in table_text]


def get_page_settings(ctx: EvalContext) -> dict[str, Any]:
    """返回第一节的页面设置，并附带 sectPr 里的 w:pgSz/@w:code 与 @w:orient。"""
    sec = ctx.doc.sections[0]
    sect_prs = xp(ctx.document_xml, ".//w:body/w:sectPr")
    pg_code: Optional[str] = None
    pg_orient: Optional[str] = None
    if sect_prs:
        pg_sz = fd(sect_prs[0], "w:pgSz")
        if pg_sz is not None:
            pg_code = pg_sz.get(W + "code")
            pg_orient = pg_sz.get(W + "orient")
    return {
        "width_cm": cm_from_emu(sec.page_width),
        "height_cm": cm_from_emu(sec.page_height),
        "top_cm": cm_from_emu(sec.top_margin),
        "bottom_cm": cm_from_emu(sec.bottom_margin),
        "left_cm": cm_from_emu(sec.left_margin),
        "right_cm": cm_from_emu(sec.right_margin),
        "header_cm": cm_from_emu(sec.header_distance),
        "footer_cm": cm_from_emu(sec.footer_distance),
        "orientation": sec.orientation,
        "pg_sz_code": pg_code,      # Word 纸型代码；A4=9
        "pg_sz_orient": pg_orient,  # "portrait" / "landscape"；省略时按纵向解释
    }


def all_section_page_settings(ctx: EvalContext) -> list[dict[str, Any]]:
    """逐节读取页面设置，用于校验"整篇文档"的一致性。

    python-docx 的 ctx.doc.sections 已覆盖所有 sectPr；对每节单独取 sectPr 元素
    以拿到 w:pgSz/@w:code 与 @w:orient。
    """
    sect_pr_elems = xp(ctx.document_xml, ".//w:sectPr")
    settings: list[dict[str, Any]] = []
    for idx, sec in enumerate(ctx.doc.sections):
        pg_code: Optional[str] = None
        pg_orient: Optional[str] = None
        if idx < len(sect_pr_elems):
            pg_sz = fd(sect_pr_elems[idx], "w:pgSz")
            if pg_sz is not None:
                pg_code = pg_sz.get(W + "code")
                pg_orient = pg_sz.get(W + "orient")
        settings.append({
            "index": idx,
            "width_cm": cm_from_emu(sec.page_width),
            "height_cm": cm_from_emu(sec.page_height),
            "top_cm": cm_from_emu(sec.top_margin),
            "bottom_cm": cm_from_emu(sec.bottom_margin),
            "left_cm": cm_from_emu(sec.left_margin),
            "right_cm": cm_from_emu(sec.right_margin),
            "header_cm": cm_from_emu(sec.header_distance),
            "footer_cm": cm_from_emu(sec.footer_distance),
            "orientation": sec.orientation,
            "pg_sz_code": pg_code,
            "pg_sz_orient": pg_orient,
        })
    return settings


# 细则要求的精确值与各自的容差。
# 纸张尺寸允许 ±0.05 cm（Word 在 cm/英寸之间换算可能有 0.01–0.02 cm 漂移）。
# 边距、页眉页脚距是 0.5 英寸/明示厘米，容差收紧到 ±0.02 cm，防止 1.05/1.35 这种偏差混入。
PAGE_SPEC = {
    "width_cm": (21.0, 0.05),
    "height_cm": (29.7, 0.05),
    "top_cm": (1.27, 0.02),
    "bottom_cm": (1.27, 0.02),
    "left_cm": (1.27, 0.02),
    "right_cm": (1.27, 0.02),
    "header_cm": (1.00, 0.02),
    "footer_cm": (1.30, 0.02),
}


def _orient_is_portrait(s: dict[str, Any]) -> bool:
    """优先看 w:pgSz/@w:orient（缺省视为 portrait），再用 python-docx 的 orientation 兜底，
    最后才退化为 height>width 的几何判断。"""
    raw = s.get("pg_sz_orient")
    if raw is not None:
        return str(raw).lower() == "portrait"
    orient = s.get("orientation")
    if orient is not None:
        try:
            return int(orient) == int(WD_ORIENTATION.PORTRAIT)
        except Exception:
            pass
    h, w = s.get("height_cm"), s.get("width_cm")
    return bool(h and w and h > w)


def _paper_is_a4(s: dict[str, Any]) -> tuple[bool, bool]:
    """返回 (尺寸合 A4, 纸型代码合 A4)。
    Word 中 w:pgSz/@w:code=9 表示 A4；省略 code 时无法判定纸型，按尺寸放行。"""
    size_ok = (approx(s.get("width_cm"), 21.0, PAGE_SPEC["width_cm"][1]) and
               approx(s.get("height_cm"), 29.7, PAGE_SPEC["height_cm"][1]))
    code = s.get("pg_sz_code")
    if code is None:
        code_ok = True  # 未声明 code：不强行判错，由尺寸决定
    else:
        code_ok = str(code) == "9"
    return size_ok, code_ok


def evaluate_page_setup(ctx: EvalContext, include_header_footer: bool = True) -> dict[str, Any]:
    """对所有 section 校验细则的 9 项页面要求，返回结构化结果。

    include_header_footer=False 时只查纸张+纵向+四边距共 7 项，对应扣分项口径
    （细则"-1 页面纸张不是A4纵向或任一页边距不是1.27厘米"不含页眉页脚距）。
    """
    sections = all_section_page_settings(ctx) or [get_page_settings(ctx)]
    failures: list[str] = []
    per_section: list[dict[str, Any]] = []
    for s in sections:
        size_ok, code_ok = _paper_is_a4(s)
        portrait_ok = _orient_is_portrait(s)
        margin_checks = {
            k: approx(s.get(k), v, tol) for k, (v, tol) in PAGE_SPEC.items()
            if k in ("top_cm", "bottom_cm", "left_cm", "right_cm")
        }
        header_ok = approx(s.get("header_cm"), 1.00, PAGE_SPEC["header_cm"][1])
        footer_ok = approx(s.get("footer_cm"), 1.30, PAGE_SPEC["footer_cm"][1])
        idx = s.get("index", 0)
        if not size_ok:
            failures.append(f"节{idx} 纸张尺寸非 21×29.7cm（实测 {s.get('width_cm')}×{s.get('height_cm')}）")
        if not code_ok:
            failures.append(f"节{idx} 纸型代码非 A4（w:pgSz@w:code={s.get('pg_sz_code')}）")
        if not portrait_ok:
            failures.append(f"节{idx} 非纵向（w:orient={s.get('pg_sz_orient')}，orientation={s.get('orientation')}）")
        for k, ok in margin_checks.items():
            if not ok:
                failures.append(f"节{idx} {k}={s.get(k)} 不等于 1.27cm")
        if include_header_footer:
            if not header_ok:
                failures.append(f"节{idx} 页眉距={s.get('header_cm')} 不等于 1.00cm")
            if not footer_ok:
                failures.append(f"节{idx} 页脚距={s.get('footer_cm')} 不等于 1.30cm")
        per_section.append({
            "index": idx, "size_ok": size_ok, "code_ok": code_ok, "portrait_ok": portrait_ok,
            "margins_ok": all(margin_checks.values()),
            "header_ok": header_ok, "footer_ok": footer_ok,
            "values": {k: s.get(k) for k in ("width_cm", "height_cm", "top_cm", "bottom_cm",
                                              "left_cm", "right_cm", "header_cm", "footer_cm",
                                              "pg_sz_code", "pg_sz_orient")},
        })
    return {"passed": not failures, "failures": failures, "sections": per_section}


def dimension1(ctx: EvalContext) -> Dimension1Result:
    """维度一：仅保留"文件能否作为 docx 打开"的评估前置校验。

    原有的"可编辑 Word 对象 / 无乱码断字竖排错位裁切"等内容规则已按需求删除，
    不再作为评分或封禁 dim2 的条件。此处保留最小前置检查：docx 无法打开时
    评估无法进行，返回 failure；能打开即视为通过。
    """
    failures: list[str] = []
    evidence: list[str] = []

    if ctx.path.suffix.lower() != ".docx":
        failures.append("交付文件不是 .docx 格式。")
    try:
        bad = ctx.zf.testzip()
        if bad:
            failures.append(f"docx 压缩包损坏：{bad}")
        else:
            evidence.append("docx 可作为 Office Open XML 压缩包正常读取。")
        _ = ctx.doc.paragraphs
        evidence.append("python-docx 可正常打开并读取正文。")
    except Exception as exc:
        failures.append(f"文件无法正常打开：{exc}")

    return Dimension1Result(passed=not failures, failures=failures, evidence=evidence)


def p_page_setup(ctx: EvalContext) -> PointResult:
    res = evaluate_page_setup(ctx, include_header_footer=True)
    s = res["sections"][0]["values"]
    detail = (
        f"width={s['width_cm']:.2f}cm, height={s['height_cm']:.2f}cm, "
        f"top={s['top_cm']:.2f}cm, bottom={s['bottom_cm']:.2f}cm, "
        f"left={s['left_cm']:.2f}cm, right={s['right_cm']:.2f}cm, "
        f"header={s['header_cm']:.2f}cm, footer={s['footer_cm']:.2f}cm, "
        f"pg_sz_code={s['pg_sz_code']}, pg_sz_orient={s['pg_sz_orient']}; "
        f"sections={len(res['sections'])}"
    )
    if res["failures"]:
        detail += "；不通过项=" + "；".join(res["failures"][:8])
    return PointResult(1, "+1 页面设置 A4纵向、边距、页眉页脚距", res["passed"], detail)


def _body_block_elements(ctx: EvalContext) -> list[etree._Element]:
    """按文档顺序返回 body 下的段落与表格元素。"""
    return xp(ctx.document_xml, "./w:body/w:p | ./w:body/w:tbl")


def _block_text(el: etree._Element) -> str:
    return "".join(xp(el, ".//w:t/text()"))


def page_index_per_block(ctx: EvalContext) -> list[tuple[etree._Element, int]]:
    """对 body 下的每个块（段落/表格）标注它**起始**所在的页号（1-based）。

    判定规则：
    - `w:br[@w:type='page']` 在该块结束后，下一个块开始时页号 +1；若分页符在块中间，则该块仍属当前页，下一个块属下一页。
    - `w:lastRenderedPageBreak`（Word 保存时写入的软分页位置）同样让下一个块起始页 +1。
    - 块内出现的 `w:sectPr`（节末分页）让下一个块页号 +1，最后一个文档级 sectPr 不计。
    """
    blocks = _body_block_elements(ctx)
    sect_total = len(xp(ctx.document_xml, ".//w:sectPr"))
    sect_seen = 0
    page = 1
    out: list[tuple[etree._Element, int]] = []
    for el in blocks:
        out.append((el, page))
        # 块内的硬分页符
        page += len(xp(el, ".//w:br[@w:type='page']"))
        # 块内的渲染分页（Word 上次保存时实际产生的软分页位置）
        page += len(xp(el, ".//w:lastRenderedPageBreak"))
        # 块内携带的 sectPr（节末分页），最后一个文档级 sectPr 不算分页
        sect_in_block = xp(el, ".//w:sectPr")
        for _ in sect_in_block:
            sect_seen += 1
            if sect_seen < sect_total:
                page += 1
    return out


def estimate_page_count(ctx: EvalContext) -> int:
    pages = page_index_per_block(ctx)
    return max((p for _, p in pages), default=1)


def infer_page_count(ctx: EvalContext) -> tuple[Optional[int], str]:
    footer_text = field_texts_in_footers(ctx)
    note = ""
    if re.search(r"共\s*8\s*页", footer_text) or re.search(r"NUMPAGES", footer_text, re.I):
        note = "，页脚含共8页字段"
    return estimate_page_count(ctx), f"静态估算：硬分页+渲染分页+节数{note}"


def _block_page_index(ctx: EvalContext, needle: str) -> Optional[int]:
    """返回首个包含 needle 的 body 块所在的页号；找不到返回 None。"""
    for el, page in page_index_per_block(ctx):
        if needle in _block_text(el):
            return page
    return None


def _is_blank_block(el: etree._Element) -> bool:
    """判定块是否"空白"——文本仅含空白且不含图片/表格内容。"""
    if el.tag == W + "tbl":
        return False  # 表格视为非空白
    text = _block_text(el).strip()
    if text:
        return False
    has_drawing = bool(xp(el, ".//w:drawing")) or bool(xp(el, ".//w:pict"))
    return not has_drawing


def detect_page9_blank(ctx: EvalContext) -> tuple[bool, str]:
    """检测是否存在第 9 页空白页。

    "存在"定义：估算页数 ≥ 9，且第 9 页（含之后）不再出现非空白块——即超出第 8 页
    后没有实质内容，纯属空白页。如果第 9 页有实质内容（说明是题目溢出而非空白），
    交给 detect_overflow_after_page8 处理。
    """
    pages = page_index_per_block(ctx)
    if not pages:
        return False, "无内容块"
    max_page = max(p for _, p in pages)
    if max_page < 9:
        return False, f"最大页号={max_page}"
    blocks_on_9plus = [(el, p) for el, p in pages if p >= 9]
    non_blank = [el for el, _ in blocks_on_9plus if not _is_blank_block(el)]
    if non_blank:
        return False, f"第9页起有 {len(non_blank)} 个非空白块（属内容溢出而非空白页）"
    return True, f"第9页起共 {len(blocks_on_9plus)} 个块均为空白"


def detect_overflow_after_page8(ctx: EvalContext) -> tuple[bool, str]:
    """检测题目内容是否超出第 8 页底部——即第 9 页及之后存在非空白内容块。"""
    pages = page_index_per_block(ctx)
    overflow_blocks = [(el, p) for el, p in pages if p >= 9 and not _is_blank_block(el)]
    if not overflow_blocks:
        return False, "第8页之后无内容溢出"
    samples = [f"第{p}页:{_block_text(el)[:25] or '(表格)'}" for el, p in overflow_blocks[:3]]
    return True, "；".join(samples)


def detect_blank_page_in_range(ctx: EvalContext, lo: int = 1, hi: int = 8) -> tuple[bool, list[int]]:
    """检测 [lo, hi] 中是否有整页空白（导致内容不连续）。"""
    pages = page_index_per_block(ctx)
    by_page: dict[int, list[etree._Element]] = {}
    for el, p in pages:
        by_page.setdefault(p, []).append(el)
    blank_pages: list[int] = []
    max_seen = max(by_page) if by_page else 0
    for pg in range(lo, min(hi, max_seen) + 1):
        blocks = by_page.get(pg, [])
        if not blocks:
            blank_pages.append(pg)
            continue
        if all(_is_blank_block(el) for el in blocks):
            blank_pages.append(pg)
    return bool(blank_pages), blank_pages


def p_pagination(ctx: EvalContext) -> PointResult:
    """+5 正文共8页 + 1-8页内容连续 + 无第9页空白 + 无超出第8页底部。

    四个子条件全部命中才通过；detail 里逐条给出结论便于排查。
    """
    pages, method = infer_page_count(ctx)

    # 子条件 1：正文共 8 页
    page_count_ok = pages == 8

    # 子条件 2：第 1—8 页内容连续（中间无整页空白）
    has_blank_mid, blank_pages = detect_blank_page_in_range(ctx, 1, 8)
    continuous_ok = not has_blank_mid

    # 子条件 3：不存在第 9 页空白页
    blank9, blank9_detail = detect_page9_blank(ctx)
    no_blank_page9_ok = not blank9

    # 子条件 4：题目内容未超出第 8 页底部
    overflow, overflow_detail = detect_overflow_after_page8(ctx)
    no_overflow_ok = not overflow

    # 末尾内容应当落在第 8 页（既非缺失，也未被挤到第 9 页）
    tail_marker = "Paragraph 2: When I returned to my own classroom on Monday"
    tail_page = _block_page_index(ctx, tail_marker)
    tail_on_p8 = tail_page == 8 if page_count_ok else (tail_page is not None)
    # tail 不在第 8 页一般会触发 overflow 或 page_count 不等于 8，无需单独失败；保留作为 detail。

    passed = page_count_ok and continuous_ok and no_blank_page9_ok and no_overflow_ok

    detail = (
        f"{method}: 页数={pages}；"
        f"共8页={page_count_ok}；"
        f"1-8连续={continuous_ok}（空白页={blank_pages}）；"
        f"无第9页空白={no_blank_page9_ok}（{blank9_detail}）；"
        f"未溢出第8页={no_overflow_ok}（{overflow_detail}）；"
        f"末尾内容所在页={tail_page}"
    )
    return PointResult(5, "+5 正文共8页且1-8页内容连续无第9页空白/溢出", passed, detail)


def _section_footer_refs(ctx: EvalContext) -> list[dict[str, str]]:
    """逐节列出 footerReference：返回 [{section, type, rId}, ...]。

    type ∈ {"default", "even", "first"}。再配合 settings.xml 的 evenAndOddHeaders
    和 sectPr/titlePg 可判断哪些"分发位"是必需的。
    """
    refs: list[dict[str, str]] = []
    for idx, sect in enumerate(xp(ctx.document_xml, ".//w:sectPr")):
        for fref in xp(sect, "w:footerReference"):
            refs.append({
                "section": str(idx),
                "type": fref.get(W + "type") or "default",
                "rId": fref.get(R + "id") or "",
                "has_titlePg": "1" if fd(sect, "w:titlePg") is not None else "0",
            })
    return refs


def _footer_xml_by_rid(ctx: EvalContext, rid: str) -> Optional[etree._Element]:
    if ctx.rels_xml is None or not rid:
        return None
    for rel in xp(ctx.rels_xml, ".//*[local-name()='Relationship']"):
        if rel.get("Id") == rid:
            target = rel.get("Target") or ""
            name = "word/" + target.lstrip("/")
            return get_zip_xml(ctx.zf, name)
    return None


def _settings_even_and_odd(ctx: EvalContext) -> bool:
    if ctx.settings_xml is None:
        return False
    return bool(xp(ctx.settings_xml, ".//w:evenAndOddHeaders"))


def _run_effective_font(run_el: etree._Element) -> dict[str, Optional[str]]:
    """返回 run 直接声明的 rFonts 各槽；未声明返回 None。继承交由调用方与 docDefaults 兜底。"""
    r_fonts = fd(run_el, "w:rPr/w:rFonts")
    if r_fonts is None:
        return {"ascii": None, "hAnsi": None, "eastAsia": None, "cs": None}
    return {k: r_fonts.get(W + k) for k in ("ascii", "hAnsi", "eastAsia", "cs")}


def _doc_default_fonts(ctx: EvalContext) -> dict[str, Optional[str]]:
    """读取 styles.xml 中 docDefaults 的 rPrDefault/rFonts，作为继承兜底。"""
    out: dict[str, Optional[str]] = {"ascii": None, "hAnsi": None, "eastAsia": None, "cs": None}
    if ctx.styles_xml is None:
        return out
    rfonts = fd(ctx.styles_xml, ".//w:docDefaults/w:rPrDefault/w:rPr/w:rFonts")
    if rfonts is None:
        return out
    for k in out:
        out[k] = rfonts.get(W + k)
    return out


def _effective_run_font(run_el: etree._Element, defaults: dict[str, Optional[str]]) -> dict[str, Optional[str]]:
    direct = _run_effective_font(run_el)
    return {k: direct[k] or defaults.get(k) for k in direct}


def _run_size_or_default(run_el: etree._Element, default_half_pt: Optional[int]) -> Optional[int]:
    sz = run_xml_size(run_el)
    return sz if sz is not None else default_half_pt


def _doc_default_size(ctx: EvalContext) -> Optional[int]:
    """docDefaults 里默认字号（half-pt）。"""
    if ctx.styles_xml is None:
        return None
    sz = fd(ctx.styles_xml, ".//w:docDefaults/w:rPrDefault/w:rPr/w:sz")
    if sz is not None and sz.get(W + "val"):
        try:
            return int(sz.get(W + "val"))
        except Exception:
            return None
    return None


def _para_style_size(ctx: EvalContext, style_id: Optional[str]) -> Optional[int]:
    """递归解析段落样式继承链上的 sz。"""
    if not style_id or ctx.styles_xml is None:
        return None
    seen: set[str] = set()
    cur = style_id
    while cur and cur not in seen:
        seen.add(cur)
        style = None
        for s in xp(ctx.styles_xml, ".//w:style"):
            if s.get(W + "styleId") == cur:
                style = s
                break
        if style is None:
            return None
        sz = fd(style, "w:rPr/w:sz")
        if sz is not None and sz.get(W + "val"):
            try:
                return int(sz.get(W + "val"))
            except Exception:
                return None
        based = fd(style, "w:basedOn")
        cur = based.get(W + "val") if based is not None else None
    return None


def _default_paragraph_style_id(ctx: EvalContext) -> Optional[str]:
    """w:style[@w:type='paragraph' and @w:default='1'] 的 styleId（通常是 Normal）。"""
    if ctx.styles_xml is None:
        return None
    for s in xp(ctx.styles_xml, ".//w:style"):
        if s.get(W + "type") == "paragraph" and s.get(W + "default") == "1":
            return s.get(W + "styleId")
    return None


def _paragraph_effective_size(ctx: EvalContext, p_el: etree._Element,
                              default_size: Optional[int]) -> dict[str, Any]:
    """对段落内所有非空 run 求"有效字号"集合：run 直接 sz → 段落样式 sz → 默认段落样式 sz → docDefaults。"""
    p_style = fd(p_el, "w:pPr/w:pStyle")
    style_id = p_style.get(W + "val") if p_style is not None else _default_paragraph_style_id(ctx)
    style_sz = _para_style_size(ctx, style_id) if style_id else None
    fallback = style_sz if style_sz is not None else default_size
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")]
    sizes = [_run_size_or_default(r, fallback) for r in runs]
    return {"sizes": sizes, "fallback": fallback, "style_id": style_id, "style_sz": style_sz}


def _classify_footer_run_chars(run_el: etree._Element) -> str:
    """把 run 按字符类型归类：'cn'/'en_num'/'mixed'/'none'。

    instrText（域指令，如 'PAGE'/'NUMPAGES'）按 'en_num' 处理（细则要求数字 TNR，
    而域指令本身不显示但其字体设定通常与显示一致）。"""
    text = "".join(xp(run_el, ".//w:t/text()"))
    instr = "".join(xp(run_el, ".//w:instrText/text()"))
    has_cn = bool(re.search(r"[一-鿿]", text))
    has_en = bool(re.search(r"[A-Za-z0-9]", text + instr))
    if has_cn and has_en:
        return "mixed"
    if has_cn:
        return "cn"
    if has_en:
        return "en_num"
    return "none"


def _footer_paragraph_field_runs(p_el: etree._Element) -> list[etree._Element]:
    """段内涉及 PAGE / NUMPAGES 域的 run（含 fldChar、instrText 所在 run）。"""
    target: list[etree._Element] = []
    in_field = False
    field_is_page = False
    for r in xp(p_el, ".//w:r"):
        fld = fd(r, "w:fldChar")
        instr = xp(r, "w:instrText/text()")
        if fld is not None:
            t = fld.get(W + "fldCharType")
            if t == "begin":
                in_field = True
                field_is_page = False
                target.append(r)
                continue
            if t == "end":
                if field_is_page:
                    target.append(r)
                in_field = False
                field_is_page = False
                continue
            if t == "separate":
                if field_is_page:
                    target.append(r)
                continue
        if in_field and instr:
            joined = "".join(instr).upper()
            if "PAGE" in joined or "NUMPAGES" in joined:
                field_is_page = True
                target.append(r)
                # 把之前 begin 的 run 也加进来（已在 target）
        elif field_is_page:
            target.append(r)
    return target


def _para_has_page_field(p_el: etree._Element) -> tuple[bool, bool]:
    """返回 (有 PAGE 域, 有 NUMPAGES 域) —— 仅按 instrText 内容判定，子串边界精确。"""
    instr_join = " ".join(xp(p_el, ".//w:instrText/text()")).upper()
    # 用单词边界避免 NUMPAGES 子串误匹配 PAGE
    has_page = bool(re.search(r"\bPAGE\b(?!S)", instr_join))
    has_numpages = bool(re.search(r"\bNUMPAGES\b", instr_join))
    return has_page, has_numpages


def _para_visible_text_with_field_values(p_el: etree._Element) -> str:
    """段落可见文本（不含 instrText 域指令，含域结果 run 的可见文本）。"""
    parts: list[str] = []
    in_field = False
    after_sep = False
    for r in xp(p_el, ".//w:r"):
        fld = fd(r, "w:fldChar")
        if fld is not None:
            t = fld.get(W + "fldCharType")
            if t == "begin":
                in_field = True
                after_sep = False
                continue
            if t == "separate":
                after_sep = True
                continue
            if t == "end":
                in_field = False
                after_sep = False
                continue
        if in_field and not after_sep:
            # 域指令体：不可见
            continue
        parts.append("".join(xp(r, "w:t/text()")))
    return "".join(parts)


def _check_footer_paragraph(ctx: EvalContext, p_el: etree._Element,
                            defaults: dict[str, Optional[str]],
                            default_size: Optional[int]) -> dict[str, Any]:
    """对一个页脚段落做细则要求的 7 项校验，返回逐项结果。

    7 项：
      a. 含页码内容（出现 PAGE 域 或 第\\d+页 文本，且出现 NUMPAGES 域 或 共8页 文本）
      b. 格式形如 "第X页/共8页"（可见文本归一后匹配；分母必须是 8）
      c. 使用真正的 PAGE / NUMPAGES 域（细则"页码字段"）
      d. 段落水平居中（pPr/jc=center）
      e. 中文字符所在 run 字体 eastAsia=宋体（继承也算）
      f. 英文/数字字符所在 run 字体 ascii=Times New Roman（继承也算）
      g. 五号（半磅=21），段内所有含字符 run 均需满足（含继承）
      h. 颜色为黑色（None/000000/auto 视为黑色）
    """
    visible = _para_visible_text_with_field_values(p_el)
    visible_norm = re.sub(r"\s+", "", visible)
    has_page_field, has_num_field = _para_has_page_field(p_el)

    contains_di_x_ye = bool(re.search(r"第\s*\d+\s*页", visible))
    contains_gong_8_ye = bool(re.search(r"共\s*8\s*页", visible))
    has_page_content = (has_page_field or contains_di_x_ye) and (has_num_field or contains_gong_8_ye)

    fmt_ok = bool(re.fullmatch(r".*第\d+页/共8页.*", visible_norm))

    field_ok = has_page_field and has_num_field

    jc = fd(p_el, "w:pPr/w:jc")
    centered = jc is not None and jc.get(W + "val") == "center"

    runs_with_chars = [r for r in xp(p_el, ".//w:r")
                       if "".join(xp(r, ".//w:t/text()")).strip()]

    cn_font_bad: list[str] = []
    en_font_bad: list[str] = []
    for r in runs_with_chars:
        eff = _effective_run_font(r, defaults)
        kind = _classify_footer_run_chars(r)
        text = "".join(xp(r, ".//w:t/text()"))
        if kind in ("cn", "mixed"):
            if eff.get("eastAsia") != "宋体":
                cn_font_bad.append(text[:10] or "(空)")
        if kind in ("en_num", "mixed"):
            if eff.get("ascii") != "Times New Roman":
                en_font_bad.append(text[:10] or "(空)")

    size_info = _paragraph_effective_size(ctx, p_el, default_size)
    sizes_with_text = []
    for r, sz in zip([r for r in xp(p_el, ".//w:r")
                      if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                     size_info["sizes"]):
        if "".join(xp(r, ".//w:t/text()")).strip():
            sizes_with_text.append(sz)
    size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_WU_HALF_PT) <= 0 for s in sizes_with_text)

    color_ok = color_black_ok(p_el)

    return {
        "visible": visible,
        "visible_norm": visible_norm,
        "has_page_content": has_page_content,
        "fmt_ok": fmt_ok,
        "field_ok": field_ok,
        "centered": centered,
        "cn_font_ok": not cn_font_bad,
        "en_font_ok": not en_font_bad,
        "size_ok": size_ok,
        "color_ok": color_ok,
        "cn_font_bad": cn_font_bad,
        "en_font_bad": en_font_bad,
        "sizes": sizes_with_text,
    }


def p_footer_page_number(ctx: EvalContext) -> PointResult:
    """+1 各页页脚页码：

    1) 各页页脚都显示页码：每个 section 的 default 页脚必须有页码段；
       若 settings.evenAndOddHeaders=on，则 even 页脚也必须有；
       若 sectPr/titlePg 存在，则该 section 还需 first 页脚有页码段。
    2) 格式 "第X页/共8页"，X 由 PAGE 域生成、8 由 NUMPAGES 域生成（细则"页码字段"）。
    3) 中文宋体 / 数字 Times New Roman，五号黑色，水平居中。
    """
    refs = _section_footer_refs(ctx)
    if not refs:
        return PointResult(1, "+1 各页页脚页码第X页/共8页 宋体/TNR五号 黑色居中", False, "未找到 footerReference")

    even_required = _settings_even_and_odd(ctx)
    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    # 按 section+type 聚合：要求的页脚必须存在且通过 7 项检查
    failures: list[str] = []
    per_footer_summary: list[str] = []
    section_indexes = sorted({r["section"] for r in refs}, key=int)

    for sec in section_indexes:
        sec_refs = [r for r in refs if r["section"] == sec]
        types_present = {r["type"] for r in sec_refs}
        required_types = {"default"}
        if even_required:
            required_types.add("even")
        if any(r["has_titlePg"] == "1" for r in sec_refs):
            required_types.add("first")

        # 缺失的页脚位
        missing = required_types - types_present
        for t in sorted(missing):
            failures.append(f"节{sec} 缺少 type={t} 页脚")

        for ref in sec_refs:
            if ref["type"] not in required_types:
                continue  # 该类型未要求（如 even 关、first 关）
            footer_xml = _footer_xml_by_rid(ctx, ref["rId"])
            if footer_xml is None:
                failures.append(f"节{sec}/type={ref['type']}: 无法读取页脚 XML (rId={ref['rId']})")
                continue
            # 找该页脚内"页码段"：可见文本含"第"与"页"，或含 PAGE/NUMPAGES 域
            candidate_paras = []
            for p in xp(footer_xml, ".//w:p"):
                visible = _para_visible_text_with_field_values(p)
                has_page, has_num = _para_has_page_field(p)
                if has_page or has_num or ("第" in visible and "页" in visible):
                    candidate_paras.append(p)
            if not candidate_paras:
                failures.append(f"节{sec}/type={ref['type']}: 页脚内无页码段")
                continue
            # 任取首个含页码内容的段进行严格校验
            results = [_check_footer_paragraph(ctx, p, defaults, default_size) for p in candidate_paras]
            # 必须存在一个候选段通过全部 7 项才算这个页脚通过
            best = None
            for r in results:
                checks = [r["has_page_content"], r["fmt_ok"], r["field_ok"], r["centered"],
                          r["cn_font_ok"], r["en_font_ok"], r["size_ok"], r["color_ok"]]
                if all(checks):
                    best = r
                    break
            if best is not None:
                per_footer_summary.append(f"节{sec}/{ref['type']}=OK: '{best['visible'][:30]}'")
                continue
            # 没有完全通过的候选，列举首段缺陷
            r = results[0]
            misses = []
            if not r["has_page_content"]: misses.append("无页码内容")
            if not r["fmt_ok"]: misses.append(f"格式≠第X页/共8页(可见='{r['visible']}')")
            if not r["field_ok"]: misses.append("缺PAGE/NUMPAGES域")
            if not r["centered"]: misses.append("未居中")
            if not r["cn_font_ok"]: misses.append(f"中文非宋体{r['cn_font_bad'][:3]}")
            if not r["en_font_ok"]: misses.append(f"数字/英文非TNR{r['en_font_bad'][:3]}")
            if not r["size_ok"]: misses.append(f"非五号(sizes={r['sizes']})")
            if not r["color_ok"]: misses.append("非黑色")
            failures.append(f"节{sec}/type={ref['type']}: " + "；".join(misses))

    passed = not failures
    detail = "；".join(per_footer_summary) if per_footer_summary else ""
    if failures:
        detail = (detail + " | " if detail else "") + "不通过：" + "；".join(failures[:6])
    return PointResult(1, "+1 各页页脚页码第X页/共8页 宋体/TNR五号 黑色居中", passed, detail)


def p_secret_text(ctx: EvalContext) -> PointResult:
    """+1 第1页左上保密文本：

    细则逐点：
      a. 文本"绝密★启用前"保留——且该段整段就是这串文本（不能含其它字）
      b. 字体宋体（中文 eastAsia）；ASCII 槽对一个全中文段允许任意值
      c. 五号（half-pt=21），段内含字符 run 全部满足（含样式继承）
      d. 黑色
      e. 左对齐（解析样式继承后必须真正是 LEFT，None 视情况）
      f. 位于页面正文顶部——是 body 的第 1 个非空段落（前面可有空段）
      g. 在第 1 页（用 page_index_per_block 求出页号）
    detail 逐项给结论。
    """
    needle = "绝密★启用前"

    # f. 顶部：扫 body 顺序找第一个非空段落
    body_blocks = _body_block_elements(ctx)
    first_para_el: Optional[etree._Element] = None
    for el in body_blocks:
        if el.tag != W + "p":
            continue
        if _block_text(el).strip():
            first_para_el = el
            break
    if first_para_el is None:
        return PointResult(1, "+1 第1页左上保密文本 宋体五号黑色左对齐顶部", False, "未找到任何非空段落")

    first_text = _block_text(first_para_el).strip()
    is_first = (first_text == needle)
    contains_needle = (needle in first_text)

    if not contains_needle:
        return PointResult(1, "+1 第1页左上保密文本 宋体五号黑色左对齐顶部", False,
                           f"正文首个非空段不是'{needle}'，实为'{first_text[:30]}'")

    # a. 整段精确等于 needle
    text_ok = is_first

    # g. 在第 1 页
    page_idx_map = {id(el): pg for el, pg in page_index_per_block(ctx)}
    page_num = page_idx_map.get(id(first_para_el))
    page_ok = page_num == 1

    # 找到对应 python-docx 段对象，用于读对齐继承
    py_p = None
    for p in ctx.paragraphs:
        if p.text.strip() == first_text:
            py_p = p
            break

    # e. 左对齐：解析继承（is_strict_left 已展开样式）
    if py_p is not None:
        eff_align = effective_paragraph_alignment(py_p)
        align_ok = eff_align == int(WD_ALIGN_PARAGRAPH.LEFT)
    else:
        align_ok = False
        eff_align = None

    # b. 字体宋体（含继承）
    defaults = _doc_default_fonts(ctx)
    runs_with_chars = [r for r in xp(first_para_el, ".//w:r")
                       if "".join(xp(r, ".//w:t/text()")).strip()]
    cn_font_bad: list[str] = []
    for r in runs_with_chars:
        eff = _effective_run_font(r, defaults)
        txt = "".join(xp(r, ".//w:t/text()"))
        if re.search(r"[一-鿿★]", txt):
            if eff.get("eastAsia") != "宋体":
                cn_font_bad.append(txt[:8])
    font_ok = not cn_font_bad

    # c. 五号：段内所有含字符 run 的有效字号都必须 == 21 half-pt
    default_size = _doc_default_size(ctx)
    size_info = _paragraph_effective_size(ctx, first_para_el, default_size)
    sizes_with_text: list[Optional[int]] = []
    for r, sz in zip([r for r in xp(first_para_el, ".//w:r")
                      if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                     size_info["sizes"]):
        if "".join(xp(r, ".//w:t/text()")).strip():
            sizes_with_text.append(sz)
    size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_WU_HALF_PT) <= 0 for s in sizes_with_text)

    # d. 黑色
    color_ok = color_black_ok(first_para_el)

    passed = all([text_ok, font_ok, size_ok, color_ok, align_ok, page_ok])
    misses: list[str] = []
    if not text_ok:    misses.append(f"段落非整段相等（实为'{first_text[:30]}'）")
    if not font_ok:    misses.append(f"中文非宋体 {cn_font_bad[:3]}")
    if not size_ok:    misses.append(f"非五号 sizes={sizes_with_text}")
    if not color_ok:   misses.append("非黑色")
    if not align_ok:   misses.append(f"非左对齐 alignment={eff_align}")
    if not page_ok:    misses.append(f"不在第1页（page={page_num}）")
    detail = "OK '" + first_text + "'" if passed else "不通过：" + "；".join(misses)
    return PointResult(1, "+1 第1页左上保密文本 宋体五号黑色左对齐顶部", passed, detail)


def p_main_title(ctx: EvalContext) -> PointResult:
    """+1 第1页主标题：

    细则逐点：
      a. 两行整段文本分别精确等于 "澜桥市2026年高三年级第二学期质量检测" 与 "英语试题"
      c. 两段均在第 1 页
      d. 字体宋体（中文 eastAsia）；数字 Times New Roman（ascii）
      e. 四号（half-pt=28），段内含字符 run 必须全部满足（含样式继承）
      f. 加粗：段内含字符 run 必须都加粗（含样式继承）
      g. 黑色
      h. 水平居中（解析样式继承后必须真正是 CENTER）

    注意：rubric 只要求两段标题文本及格式各自成立，并不要求第二行紧跟第一行。
    因此这里分别定位两段标题独立校验，不把相邻性作为得分条件。
    """
    LINE1 = "澜桥市2026年高三年级第二学期质量检测"
    LINE2 = "英语试题"

    body_blocks = _body_block_elements(ctx)
    # 仅取段落，跳过空段，构建（段落元素, 文本）列表，保留 body 顺序
    para_seq: list[tuple[etree._Element, str]] = []
    for el in body_blocks:
        if el.tag != W + "p":
            continue
        t = _block_text(el).strip()
        if not t:
            continue
        para_seq.append((el, t))

    # a. 分别定位两段整段精确相等的标题，彼此独立，不要求相邻。
    idx1 = next((i for i, (_, t) in enumerate(para_seq) if t == LINE1), None)
    if idx1 is None:
        return PointResult(1, "+1 第1页主标题 宋体/TNR四号加粗黑色居中两行", False,
                           f"未找到整段相等的第一行 '{LINE1}'")
    idx2 = next((i for i, (_, t) in enumerate(para_seq) if t == LINE2), None)
    if idx2 is None:
        return PointResult(1, "+1 第1页主标题 宋体/TNR四号加粗黑色居中两行", False,
                           f"未找到整段相等的第二行 '{LINE2}'")
    e1, _t1 = para_seq[idx1]
    e2, _t2 = para_seq[idx2]

    # c. 在第 1 页
    page_idx_map = {id(el): pg for el, pg in page_index_per_block(ctx)}
    pg1, pg2 = page_idx_map.get(id(e1)), page_idx_map.get(id(e2))
    page_ok = pg1 == 1 and pg2 == 1

    # 找对应 python-docx 段对象（用于读对齐继承）
    def _find_py(el: etree._Element, text: str) -> Any:
        for p in ctx.paragraphs:
            if p._p is el or p.text.strip() == text:
                return p
        return None
    py1, py2 = _find_py(e1, LINE1), _find_py(e2, LINE2)

    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    def _check_line(el: etree._Element, py: Any, line_text: str) -> dict[str, Any]:
        runs_with_chars = [r for r in xp(el, ".//w:r")
                           if "".join(xp(r, ".//w:t/text()")).strip()]
        # d. 字体：中文 eastAsia=宋体，含数字/英文字符的 run ascii=TNR
        cn_bad: list[str] = []
        en_bad: list[str] = []
        for r in runs_with_chars:
            eff = _effective_run_font(r, defaults)
            txt = "".join(xp(r, ".//w:t/text()"))
            if re.search(r"[一-鿿]", txt) and eff.get("eastAsia") != "宋体":
                cn_bad.append(txt[:8])
            if re.search(r"[A-Za-z0-9]", txt) and eff.get("ascii") != "Times New Roman":
                en_bad.append(txt[:8])
        cn_font_ok = not cn_bad
        en_font_ok = not en_bad

        # e. 四号 = 28 half-pt
        size_info = _paragraph_effective_size(ctx, el, default_size)
        sizes_with_text: list[Optional[int]] = []
        for r, sz in zip([r for r in xp(el, ".//w:r")
                          if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                         size_info["sizes"]):
            if "".join(xp(r, ".//w:t/text()")).strip():
                sizes_with_text.append(sz)
        size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_SI_HALF_PT) <= 0 for s in sizes_with_text)

        # f. 加粗（含继承）：先看 run，再看段落样式
        bold_ok = _effective_paragraph_bold(ctx, el)

        # g. 黑色
        color_ok = color_black_ok(el)

        # h. 水平居中（含样式继承）
        align_val = effective_paragraph_alignment(py) if py is not None else None
        center_ok = align_val == int(WD_ALIGN_PARAGRAPH.CENTER)

        return {
            "text_ok": _block_text(el).strip() == line_text,
            "cn_font_ok": cn_font_ok, "en_font_ok": en_font_ok,
            "size_ok": size_ok, "bold_ok": bold_ok, "color_ok": color_ok,
            "center_ok": center_ok,
            "cn_bad": cn_bad, "en_bad": en_bad, "sizes": sizes_with_text,
            "align_val": align_val,
        }

    r1 = _check_line(e1, py1, LINE1)
    r2 = _check_line(e2, py2, LINE2)
    line1_ok = all([r1["text_ok"], r1["cn_font_ok"], r1["en_font_ok"], r1["size_ok"],
                    r1["bold_ok"], r1["color_ok"], r1["center_ok"]])
    line2_ok = all([r2["text_ok"], r2["cn_font_ok"], r2["en_font_ok"], r2["size_ok"],
                    r2["bold_ok"], r2["color_ok"], r2["center_ok"]])

    passed = line1_ok and line2_ok and page_ok

    def _line_misses(tag: str, r: dict[str, Any]) -> list[str]:
        m: list[str] = []
        if not r["text_ok"]: m.append(f"{tag}非整段相等")
        if not r["cn_font_ok"]: m.append(f"{tag}中文非宋体{r['cn_bad'][:3]}")
        if not r["en_font_ok"]: m.append(f"{tag}数字/英文非TNR{r['en_bad'][:3]}")
        if not r["size_ok"]: m.append(f"{tag}非四号 sizes={r['sizes']}")
        if not r["bold_ok"]: m.append(f"{tag}非加粗")
        if not r["color_ok"]: m.append(f"{tag}非黑色")
        if not r["center_ok"]: m.append(f"{tag}非居中 alignment={r['align_val']}")
        return m

    misses = _line_misses("第一行", r1) + _line_misses("第二行", r2)
    if not page_ok:
        misses.append(f"不在第1页 (页号={pg1},{pg2})")
    detail = f"OK '{LINE1}' / '{LINE2}'" if passed else "不通过：" + "；".join(misses[:8])
    return PointResult(1, "+1 第1页主标题 宋体/TNR四号加粗黑色居中两行", passed, detail)


def p_date_text(ctx: EvalContext) -> PointResult:
    p = para_by_exact_or_contains(ctx, "2026年6月")
    e = para_xml_by_text(ctx, "2026年6月")
    not_bold = e is not None and not is_bold_para(e)
    passed = bool(p and e is not None and is_center(p) and size_ok_for_runs(e, FONT_SZ_XIAOSI_HALF_PT) and color_black_ok(e) and not_bold)
    return PointResult(1, "+1 第1页日期文本小四居中不加粗", passed, "找到文本" if p else "未找到文本")


def body_paragraph_candidates(ctx: EvalContext) -> list[Any]:
    skip_starts = ("绝密", "澜桥市2026", "英语试题", "2026年6月")
    return [p for p in ctx.paragraphs if p.text.strip() and not p.text.strip().startswith(skip_starts)]


def _all_body_paragraph_elements(ctx: EvalContext) -> list[etree._Element]:
    """整篇正文涉及的所有段落：body 直接子段 + 表格内段。

    page header / page footer / comments / footnotes 不属于"正文"，已自然排除（它们不在 word/document.xml 的 w:body 下）。
    """
    return xp(ctx.document_xml, ".//w:body//w:p")


def _is_excluded_body_para(text: str) -> bool:
    """正文字体校验排除项：题头四段（绝密/题名/英语试题/日期）+ 答题横线段。

    题头不在细则的 9 类正文中（由其它评分点单独校验）；
    答题横线段（纯下划线字符）也不在 9 类正文中——它是"答题区"结构，不是"中文说明/英文文章"等。
    """
    t = text.strip()
    if not t:
        return True
    if t == "绝密★启用前":
        return True
    if t == "澜桥市2026年高三年级第二学期质量检测":
        return True
    if t == "英语试题":
        return True
    if re.fullmatch(r"\s*2026年6月\s*", t):
        return True
    # 答题横线段（纯下划线 + 空白）：不在 9 类正文
    if re.fullmatch(r"[_\s]{20,}", t):
        return True
    return False


def _classify_run_chars(text: str) -> tuple[bool, bool]:
    """返回 (含中文/中文标点, 含 ASCII 字母或数字)。"""
    has_cn = bool(re.search(r"[　-〿一-鿿＀-￯★]", text))
    has_en = bool(re.search(r"[A-Za-z0-9]", text))
    return has_cn, has_en


def _check_run_for_body_fonts(run_el: etree._Element,
                              defaults: dict[str, Optional[str]],
                              size_fallback: Optional[int]) -> dict[str, Any]:
    """对单个 run 做"小四 + 中文宋体 + 英文/数字 TNR + 黑色"四项严格校验。

    返回 {classification, size_ok, cn_font_ok, en_font_ok, color_ok, sz_eff, fonts, color, text}
    classification ∈ {'cn', 'en', 'mixed', 'none'}。
    """
    text = "".join(xp(run_el, ".//w:t/text()"))
    if not text.strip():
        return {"classification": "none"}
    has_cn, has_en = _classify_run_chars(text)
    eff_font = _effective_run_font(run_el, defaults)
    sz = run_xml_size(run_el)
    sz_eff = sz if sz is not None else size_fallback
    color = run_xml_color(run_el)
    color_ok = color in (None, "", "000000", "AUTO")
    size_ok = sz_eff is not None and abs(sz_eff - FONT_SZ_XIAOSI_HALF_PT) <= 0
    cn_font_ok = (not has_cn) or eff_font.get("eastAsia") == "宋体"
    en_font_ok = (not has_en) or eff_font.get("ascii") == "Times New Roman"
    classification = "mixed" if (has_cn and has_en) else ("cn" if has_cn else ("en" if has_en else "none"))
    return {
        "classification": classification,
        "size_ok": size_ok, "cn_font_ok": cn_font_ok, "en_font_ok": en_font_ok,
        "color_ok": color_ok, "sz_eff": sz_eff, "fonts": eff_font,
        "color": color, "text": text,
    }


def p_body_fonts(ctx: EvalContext) -> PointResult:
    """+5 整篇正文：

    a. 范围：整篇正文的所有段落——包括 body 直接子段以及**表格内正文类段落**（w:tbl//w:p），
       不含页眉/页脚/批注/脚注（不在 w:body 下自然排除）。
       rubric 覆盖的 9 类正文（中文说明/章节标题/写作要求/注意事项/英文文章/英文题目/选项/单词/数字）
       在真实答卷/试题中可能出现在表格布局里（如题目/选项排版），必须一同纳入校验。
    b. 排除项：仅排除"另有评分点单独覆盖且确实不属于 9 类正文"的段落——
       题头四段（绝密/题名/英语试题/日期）和答题横线段（纯下划线）。
    c. 字号：所有含字符 run 有效字号 == 24 half-pt（小四，容差 0）。
    d. 中文 run：eastAsia 槽位 == 宋体（含 docDefaults 继承）。
    e. 英文/数字 run：ascii 槽位 == Times New Roman（含 docDefaults 继承）。
    f. 颜色：所有含字符 run color ∈ {None, "", 000000, AUTO}。
    g. 整篇 100% 合规；任何一处违例即失败，detail 给出前几处样本。
    """
    # 收集 body 下所有段（含表格内嵌段）：.//w:body//w:p
    paras = _all_body_paragraph_elements(ctx)
    paras = [p for p in paras if not _is_excluded_body_para(_block_text(p))]
    if not paras:
        return PointResult(5, "+5 整篇正文 中文宋体/英文数字TNR 小四 黑色", False, "无正文段落")

    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    bad_size: list[str] = []
    bad_cn_font: list[str] = []
    bad_en_font: list[str] = []
    bad_color: list[str] = []
    counted_runs = 0

    for p_el in paras:
        size_info = _paragraph_effective_size(ctx, p_el, default_size)
        size_fallback = size_info["fallback"]
        para_text_preview = _block_text(p_el)[:30]
        for r in xp(p_el, ".//w:r"):
            res = _check_run_for_body_fonts(r, defaults, size_fallback)
            if res.get("classification") == "none":
                continue
            counted_runs += 1
            if not res["size_ok"]:
                bad_size.append(f"段'{para_text_preview}' run='{res['text'][:12]}' sz={res['sz_eff']}")
            if not res["cn_font_ok"]:
                bad_cn_font.append(f"段'{para_text_preview}' run='{res['text'][:12]}' eastAsia={res['fonts'].get('eastAsia')}")
            if not res["en_font_ok"]:
                bad_en_font.append(f"段'{para_text_preview}' run='{res['text'][:12]}' ascii={res['fonts'].get('ascii')}")
            if not res["color_ok"]:
                bad_color.append(f"段'{para_text_preview}' run='{res['text'][:12]}' color={res['color']}")

    passed = counted_runs > 0 and not bad_size and not bad_cn_font and not bad_en_font and not bad_color
    if passed:
        detail = f"段落={len(paras)}；含字符 run={counted_runs}；全部满足小四/宋体/TNR/黑色"
    else:
        parts: list[str] = [f"段落={len(paras)}；含字符 run={counted_runs}"]
        if bad_size:    parts.append(f"非小四({len(bad_size)})：" + "；".join(bad_size[:3]))
        if bad_cn_font: parts.append(f"中文非宋体({len(bad_cn_font)})：" + "；".join(bad_cn_font[:3]))
        if bad_en_font: parts.append(f"数字/英文非TNR({len(bad_en_font)})：" + "；".join(bad_en_font[:3]))
        if bad_color:   parts.append(f"非黑色({len(bad_color)})：" + "；".join(bad_color[:3]))
        detail = " | ".join(parts)
    return PointResult(5, "+5 整篇正文 中文宋体/英文数字TNR 小四 黑色", passed, detail)


def _is_zero_first_line_indent(ctx: EvalContext, p_el: etree._Element) -> tuple[bool, dict[str, Any]]:
    """判定段落"无首行缩进"——综合 firstLineChars / firstLine / hanging（含样式继承）。

    判定标准：
      - 首行实际缩进 = firstLine_twips - hanging_twips（OOXML 语义；hanging 抵消 firstLine）
      - 若有 firstLineChars 且非 0，认为有缩进（chars 形式）
      - 若 firstLine_twips 不为 0，认为有缩进
      - hanging 单独存在不算"首行缩进"（它影响除首行外的左缩进，但首行本身回到段左边）
    """
    ind = _effective_first_line_indent(ctx, p_el)
    flc = ind.get("firstLineChars")
    if flc is not None:
        try:
            if int(flc) != 0:
                return False, ind
        except Exception:
            pass
    fl = ind.get("firstLine_twips")
    if fl is not None:
        try:
            if int(fl) != 0:
                return False, ind
        except Exception:
            pass
    return True, ind


def _para_ends_with_punct(text: str) -> bool:
    """段末是否以题干合理标点收尾。允许末尾被引号/括号包裹后再收：去掉尾部的 \"”')]）] 等再判。"""
    if not text:
        return False
    stripped = text.rstrip()
    # 剥掉末尾配对符号
    stripped = re.sub(r"[\"'”’」』）)\]】》>]+$", "", stripped)
    return bool(re.search(r"[?？.。:：!！]$", stripped))


def _is_full_question_in_one_paragraph(text: str, num: int) -> tuple[bool, str]:
    """判定"题号 + 题干 + 句末标点"是否同段。

    - 段内含该题号（题号可能在段首，也可能嵌入在文章/选项行中）
    - 题干文本（去掉题号后剩余部分）必须非空——即"题号不能孤立成段"
    - 段末必须有合理收尾标点（题干性结尾）——句末标点需与题干同段
    返回 (ok, 原因)。
    """
    # 去掉段内所有"该题号 + . 或 空格"占位，检查其它可读文本是否存在
    stripped = re.sub(rf"(?<!\d){num}(?!\d)[\.．]?\s*", " ", text)
    stripped = stripped.strip()
    if not stripped:
        return False, "题号独占一段（无题干）"
    if not _para_ends_with_punct(text):
        return False, f"段末无收尾标点（实末'{text[-6:]}'）"
    return True, ""


def p_question_paragraphs(ctx: EvalContext) -> PointResult:
    """+3 第 21-67 题题号、题干、句末标点同段；左对齐；无首行缩进。

    rubric 对 21-67 的要求是**统一的**：题号 + 题干 + 句末标点必须同段，段落左对齐、
    无首行缩进；不因题型不同而放宽。

    - 21-35（阅读理解选择题题干）、66-67（写作题）：题号在段首（形如 "21." "66."），
      段内即为题干本体，段末为题干性收尾标点。
    - 41-55（完形填空选项行）：题号在段首（形如 "41. A. ... B. ... C. ... D. ..."），
      段末应有合理标点/句末标点。
    - 36-40（阅读七选五）、56-65（语法填空）：题号嵌入文章段中；段内应包含题号 +
      文章题干文字，段末为句末标点。
    所有题型统一验证：① 段内包含该题号；② 段内除题号外仍有非空题干文本；
    ③ 段末有合理收尾标点；④ 段落左对齐；⑤ 段落无首行缩进。
    """
    body_paras = _all_body_paragraph_elements(ctx)
    py_by_id: dict[int, Any] = {id(p._p): p for p in ctx.paragraphs}
    for tbl in ctx.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    py_by_id[id(p._p)] = p

    # 找每个题号 (21-67) 在 body 顺序中第一次出现的段落。
    #   - 21-35 / 41-55 / 66-67：题号在段首（形如 "21." 或 "21．"），优先匹配段首形式。
    #   - 36-40 / 56-65：题号嵌入文章/选项中，落到第一处包含该题号的段。
    found: dict[int, etree._Element] = {}
    embedded_in_text_nums = set(range(36, 41)) | set(range(56, 66))
    # 第一轮：段首题号
    for el in body_paras:
        if el.tag != W + "p":
            continue
        text = _block_text(el)
        m = re.match(r"^\s*(\d{2})[\.．]\s*", text)
        if m:
            try:
                num = int(m.group(1))
            except Exception:
                num = -1
            if 21 <= num <= 67 and num not in found:
                found[num] = el
    # 第二轮：补齐嵌入题号（七选五 / 语法填空）
    for el in body_paras:
        if el.tag != W + "p":
            continue
        text = _block_text(el)
        for num in embedded_in_text_nums:
            if num in found:
                continue
            # 题号边界：左边非数字、右边非数字
            if re.search(rf"(?<!\d){num}(?!\d)", text):
                found[num] = el

    expected = set(range(21, 68))
    missing = sorted(expected - set(found.keys()))

    failures: list[str] = []
    for num in sorted(found.keys()):
        el = found[num]
        text = _block_text(el)
        py = py_by_id.get(id(el))
        # ① 左对齐
        align_val = effective_paragraph_alignment(py) if py is not None else None
        align_ok = align_val == int(WD_ALIGN_PARAGRAPH.LEFT)
        # ② 无首行缩进
        zero_indent, ind_info = _is_zero_first_line_indent(ctx, el)
        # ③ 同段（题号 + 题干 + 句末标点）——对 21-67 全部统一验证
        same_para_ok, same_reason = _is_full_question_in_one_paragraph(text, num)

        misses: list[str] = []
        if not align_ok:
            misses.append(f"非左对齐(alignment={align_val})")
        if not zero_indent:
            misses.append(f"有首行缩进(firstLineChars={ind_info.get('firstLineChars')},firstLine={ind_info.get('firstLine_twips')}twips)")
        if not same_para_ok:
            misses.append(f"非同段({same_reason})")
        if misses:
            preview = text.strip()[:25]
            failures.append(f"{num}: '{preview}' " + "；".join(misses))

    passed = not missing and not failures

    parts: list[str] = [f"找到题号={sorted(found.keys())[:8]}{'...' if len(found)>8 else ''}({len(found)}/47)"]
    if missing:
        parts.append(f"缺题号({len(missing)})={missing[:10]}")
    if failures:
        parts.append(f"格式异常({len(failures)})：" + "；".join(failures[:4]))
    if passed:
        parts.append("全部满足")
    detail = " | ".join(parts)
    return PointResult(3, "+3 第21-67题题号在段首/同段含题干和句末标点/左对齐/无首行缩进",
                       passed, detail)


def _abs_left_edge_twips(ctx: EvalContext, p_el: etree._Element) -> Optional[int]:
    """段落"首行实际左边界"（twips）：left + firstLine（hanging 抵消 firstLine）。

    优先级：段直接 w:ind → 段落样式继承链。chars 形式（leftChars/firstLineChars）按
    小四（24 half-pt = 12pt）下中文字符宽 ≈ 240 twips/字 估算（1 字符=20 twips×字号pt）。
    """
    ind = _effective_first_line_indent(ctx, p_el)
    # 直接 w:ind 还需要 left/leftChars（_effective_first_line_indent 没读 left 系列，单独读）
    direct = fd(p_el, "w:pPr/w:ind")
    left_tw: Optional[int] = None
    left_chars: Optional[int] = None
    if direct is not None:
        if direct.get(W + "left") is not None:
            try: left_tw = int(direct.get(W + "left"))
            except Exception: pass
            if left_tw is None and direct.get(W + "start") is not None:
                try: left_tw = int(direct.get(W + "start"))
                except Exception: pass
        if direct.get(W + "leftChars") is not None:
            try: left_chars = int(direct.get(W + "leftChars"))
            except Exception: pass
    # 样式继承补 left
    if left_tw is None and left_chars is None:
        p_style = fd(p_el, "w:pPr/w:pStyle")
        style_id = p_style.get(W + "val") if p_style is not None else _default_paragraph_style_id(ctx)
        sty = _para_style_indent(ctx, style_id)
        if sty.get("left") is not None:
            try: left_tw = int(sty["left"])  # type: ignore[arg-type]
            except Exception: pass
        if left_chars is None and sty.get("leftChars") is not None:
            try: left_chars = int(sty["leftChars"])  # type: ignore[arg-type]
            except Exception: pass

    fl_tw: Optional[int] = None
    if ind.get("firstLine_twips") is not None:
        try: fl_tw = int(ind["firstLine_twips"])
        except Exception: pass
    fl_chars: Optional[int] = None
    if ind.get("firstLineChars") is not None:
        try: fl_chars = int(ind["firstLineChars"])
        except Exception: pass
    hanging_tw: Optional[int] = None
    if ind.get("hanging_twips") is not None:
        try: hanging_tw = int(ind["hanging_twips"])
        except Exception: pass

    # chars 转 twips：小四 12pt 下，1 字符宽 ≈ 240 twips（中文）；1 character chars 单位 = 100
    def _chars_to_twips(c: int) -> int:
        return int(round(c / 100.0 * 240))

    left = (left_tw if left_tw is not None
            else (_chars_to_twips(left_chars) if left_chars is not None else 0))
    fl = (fl_tw if fl_tw is not None
          else (_chars_to_twips(fl_chars) if fl_chars is not None else 0))
    hanging = hanging_tw if hanging_tw is not None else 0
    # 首行实际左边界 = left + firstLine - hanging
    return left + fl - hanging


def _para_first_line_indent_twips(ctx: EvalContext, p_el: etree._Element) -> Optional[int]:
    """仅取段落级（含样式继承）的"首行缩进"twips 值——chars 折算到 twips；hanging 抵消。"""
    ind = _effective_first_line_indent(ctx, p_el)
    fl_tw: Optional[int] = None
    if ind.get("firstLine_twips") is not None:
        try: fl_tw = int(ind["firstLine_twips"])
        except Exception: pass
    fl_chars: Optional[int] = None
    if ind.get("firstLineChars") is not None:
        try: fl_chars = int(ind["firstLineChars"])
        except Exception: pass
    hanging_tw: Optional[int] = None
    if ind.get("hanging_twips") is not None:
        try: hanging_tw = int(ind["hanging_twips"])
        except Exception: pass
    fl = (fl_tw if fl_tw is not None
          else (int(round(fl_chars / 100.0 * 240)) if fl_chars is not None else None))
    if fl is None and hanging_tw is None:
        return None
    return (fl or 0) - (hanging_tw or 0)


def _para_style_tab_stops(ctx: EvalContext, style_id: Optional[str]) -> list[int]:
    """沿样式继承链收集 w:pPr/w:tabs/w:tab 的 pos 值（twips）。仅取 val ∈ {left, start, num, decimal, center, right, end}
    的常规 tab；不收 clear。"""
    positions: list[int] = []
    if not style_id or ctx.styles_xml is None:
        return positions
    seen: set[str] = set()
    cur = style_id
    while cur and cur not in seen:
        seen.add(cur)
        style = None
        for s in xp(ctx.styles_xml, ".//w:style"):
            if s.get(W + "styleId") == cur:
                style = s
                break
        if style is None:
            break
        tabs = fd(style, "w:pPr/w:tabs")
        if tabs is not None:
            for tab in xp(tabs, "w:tab"):
                val = tab.get(W + "val")
                pos = tab.get(W + "pos")
                if val == "clear" or pos is None:
                    continue
                try:
                    positions.append(int(pos))
                except Exception:
                    continue
        based = fd(style, "w:basedOn")
        cur = based.get(W + "val") if based is not None else None
    return positions


def _effective_paragraph_tab_stops(ctx: EvalContext, p_el: etree._Element) -> list[int]:
    """段直接 w:pPr/w:tabs 优先；空则回退到段落样式链上定义的 tab stops。
    返回排序后去重的 tab 位置（twips），不含 clear。"""
    direct = fd(p_el, "w:pPr/w:tabs")
    positions: list[int] = []
    cleared: set[int] = set()
    if direct is not None:
        for tab in xp(direct, "w:tab"):
            val = tab.get(W + "val")
            pos = tab.get(W + "pos")
            if pos is None:
                continue
            try:
                p = int(pos)
            except Exception:
                continue
            if val == "clear":
                cleared.add(p)
            else:
                positions.append(p)
    if not positions:
        p_style = fd(p_el, "w:pPr/w:pStyle")
        style_id = p_style.get(W + "val") if p_style is not None else _default_paragraph_style_id(ctx)
        positions = _para_style_tab_stops(ctx, style_id)
    positions = [p for p in positions if p not in cleared]
    return sorted(set(positions))


def _marker_x_pos_twips(ctx: EvalContext, p_el: etree._Element,
                        marker: str, abs_left_twips: int) -> tuple[Optional[int], str]:
    """求段落内首个 marker（例如 "B." / "D."）的绝对 x 位置（twips）。

    策略：
      1) 按 run 序遍历，累计可见文本；找到含 marker 的 run 后，回看该 run 前段落中是否插入
         过 <w:tab/>。
      2) 若在 marker 之前存在至少一次 <w:tab/>：marker 的 x = abs_left + 最后一个 tab 命中的 tab stop 位置。
         "命中"简化为：取该段落有效 tab stops 中 ≥ marker 前累计位置估算的最小 tab stop；若无法估算，
         则取所有 tab stop 中的**最小值**（选项段一般只有 1 个 tab stop，等价）。
      3) 若 marker 之前无 <w:tab/>：无法可靠给出列 x 位置，返回 (None, reason)。
    """
    tab_stops = _effective_paragraph_tab_stops(ctx, p_el)
    saw_tab_before_marker = False
    text_accum = ""
    hit = False
    for r in xp(p_el, ".//w:r"):
        # 该 run 之前先看它自己是否是 tab 节点
        for child in list(r):
            if child.tag == W + "tab":
                if not hit:
                    saw_tab_before_marker = True
            elif child.tag == W + "t":
                seg = child.text or ""
                # 检查此段拼上后是否首次出现 marker
                combined = text_accum + seg
                idx = combined.find(marker, len(text_accum) if len(text_accum) - len(marker) < 0 else max(0, len(text_accum) - len(marker) + 1))
                # 简化为在合并串内首次匹配位置
                idx = combined.find(marker)
                if idx >= 0 and not hit:
                    hit = True
                    text_accum = combined[:idx]  # 定格到 marker 之前
                    break
                text_accum = combined
        if hit:
            break
    if not hit:
        return None, "未在段内找到 marker"
    if not saw_tab_before_marker:
        return None, "marker 前无 <w:tab/>，无 tab 结构可用"
    if not tab_stops:
        return None, "存在 tab 但段落未定义 tab stops（含样式继承）"
    # 简化：取最小 tab stop（选项段通常只有 1 个）
    return abs_left_twips + min(tab_stops), f"tab_stops={tab_stops}"


def p_options(ctx: EvalContext) -> PointResult:
    """+5 选择题选项段落：

    细则要求：
      1) 选项段落首行缩进 = 0.74cm（对**所有已识别的选项段**都必检，不因排版模式跳过）；
      2) 模式 1：A 与 B 在同一行、C 与 D 在下一行 → A/C 段左边界一致，且 B/D 列 x 位置一致；
      3) 模式 2：所有选项均单独成行 → 4 段左边界一致。

    B/D 列对齐**基于 tab stops + <w:tab/> 位置**计算 x 坐标（`abs_left + 命中的 tab stop`），
    不再用字符串索引。若两段都没显式 tab（选项之间只用空格且 XML 无 tab stops），当前 XML
    无法可靠判定视觉列位置——按"无法判定"处理，不再强行判失败以免误判。

    范围：21-35（阅读选择）、41-55（完形选择）。每题 4 个选项 A/B/C/D。
    """
    body_paras = _all_body_paragraph_elements(ctx)

    # 按题号收集"该题的选项段"序列：题号 → list[(p_el, 段文本去掉前导空白)]
    qmap: dict[int, list[tuple[etree._Element, str]]] = {}
    # 寻找每题 N 段。从段首"N. " 开始，直到下一个题号或非选项段。
    paras_seq: list[etree._Element] = [el for el in body_paras if el.tag == W + "p"]
    i = 0
    while i < len(paras_seq):
        el = paras_seq[i]
        text = _block_text(el).strip()
        m = re.match(r"^(2[1-9]|3[0-5]|4[1-9]|5[0-5])\.\s*(.*)", text)
        if not m:
            i += 1
            continue
        num = int(m.group(1))
        rest = m.group(2)
        seq: list[tuple[etree._Element, str]] = []
        # 题号所在段：可能是"N. 题干"（阅读选择），也可能是"N. A. ... B. ... C. ... D. ..."（完形单段）
        if re.search(r"\bA\.", rest):
            seq.append((el, rest))
        # 紧随其后的段：吃掉所有以 [A-D]. 开头的连续段，最多 4 个或直到遇到下一题号/非选项段
        j = i + 1
        collected = 0
        while j < len(paras_seq) and collected < 4:
            nxt = paras_seq[j]
            ntxt = _block_text(nxt).strip()
            if not ntxt:
                j += 1
                continue
            if re.match(r"^(2[1-9]|3[0-5]|4[1-9]|5[0-5])\.", ntxt):
                break
            if re.match(r"^[A-D]\.", ntxt):
                seq.append((nxt, ntxt))
                collected += 1
                j += 1
                continue
            break
        if seq:
            qmap[num] = seq
        i = j if seq else i + 1

    if not qmap:
        return PointResult(5, "+5 选择题选项 模式1(A.B./C.D.两段)/模式2(4段独立) 首行缩进0.74cm + 互相左对齐",
                           False, "未识别到任何 21-35 / 41-55 的选项段")

    TARGET_TWIPS = 420  # 0.74cm × 567 ≈ 420 twips
    INDENT_TOL = 28     # ±0.05cm
    ALIGN_TOL = 10      # 段首左边界一致容差 (twips)

    bad_indent: list[str] = []
    bad_align: list[str] = []
    align_undetermined: list[str] = []
    mode_dist: dict[str, int] = {"模式1(2段)": 0, "模式2(4段)": 0, "其它(单段/混合)": 0}

    for num in sorted(qmap.keys()):
        seq = qmap[num]
        # 判模式
        # 模式 1: 2 段且第 1 段含 A.B.、第 2 段含 C.D.
        # 模式 2: 4 段且各以 A./B./C./D. 开头
        # 其它：完形单段（"N. A. ... B. ... C. ... D. ..." 整题在同一段）或混合排版
        if len(seq) == 2 and "A." in seq[0][1] and "B." in seq[0][1] and "C." in seq[1][1] and "D." in seq[1][1]:
            mode = "模式1"
            mode_dist["模式1(2段)"] += 1
        elif len(seq) == 4 and all(seq[k][1].startswith(f"{chr(ord('A')+k)}.") for k in range(4)):
            mode = "模式2"
            mode_dist["模式2(4段)"] += 1
        else:
            mode = "其它"
            mode_dist["其它(单段/混合)"] += 1

        # a. 首行缩进：所有识别到的选项段一律检查 0.74cm，无论模式（rubric 首句要求）
        for el, txt in seq:
            fl_tw = _para_first_line_indent_twips(ctx, el)
            if fl_tw is None or abs(fl_tw - TARGET_TWIPS) > INDENT_TOL:
                bad_indent.append(f"{num}({mode}): '{txt[:25]}' firstLine={fl_tw}twips≠420")

        # b/c. 对齐——只在细则点名的模式1/模式2 下检查
        if mode == "模式1":
            ab_left = _abs_left_edge_twips(ctx, seq[0][0])
            cd_left = _abs_left_edge_twips(ctx, seq[1][0])
            if ab_left is None or cd_left is None or abs(ab_left - cd_left) > ALIGN_TOL:
                bad_align.append(f"{num}(模式1): A/C段左边界={ab_left} vs {cd_left}")

            # B/D 列 x 位置：走 tab stops + <w:tab/>，不再用字符串位置
            b_x, b_why = _marker_x_pos_twips(ctx, seq[0][0], "B.", ab_left or 0)
            d_x, d_why = _marker_x_pos_twips(ctx, seq[1][0], "D.", cd_left or 0)
            if b_x is None or d_x is None:
                # XML 里 B/D 不是靠 tab 定位（多半是空格排版）；无法结构化判定列位置。
                align_undetermined.append(
                    f"{num}(模式1): B列={b_why}; D列={d_why}"
                )
            elif abs(b_x - d_x) > ALIGN_TOL:
                bad_align.append(f"{num}(模式1): B列 x={b_x} vs D列 x={d_x}（>{ALIGN_TOL}twips）")
        elif mode == "模式2":
            edges = [_abs_left_edge_twips(ctx, el) for el, _ in seq]
            if any(e is None for e in edges) or (max(edges) - min(edges)) > ALIGN_TOL:  # type: ignore[type-var]
                bad_align.append(f"{num}(模式2): A/B/C/D左边界={edges} 不一致")
        # 其它模式：细则未点名对齐要求，只查首行缩进

    passed = not bad_indent and not bad_align
    parts: list[str] = [
        f"识别题数={len(qmap)}；模式分布={mode_dist}",
    ]
    if bad_indent:
        parts.append(f"首行缩进非0.74cm({len(bad_indent)})：" + "；".join(bad_indent[:3]))
    if bad_align:
        parts.append(f"对齐异常({len(bad_align)})：" + "；".join(bad_align[:3]))
    if align_undetermined:
        parts.append(
            f"B/D列对齐未能结构化判定({len(align_undetermined)}，非失败)：" + "；".join(align_undetermined[:3])
        )
    if passed:
        parts.append("全部满足（所有选项段首行缩进=0.74cm；模式1/模式2 对齐合规）")
    detail = " | ".join(parts)
    return PointResult(5, "+5 选择题选项 首行缩进0.74cm(全部) + 模式1(A/C&B/D)/模式2(4段)左对齐",
                       passed, detail)


def p_reading_titles(ctx: EvalContext) -> PointResult:
    """+1 第1页阅读理解三段标题：

    细则要求三段整段文本：
      T1: "第二部分 阅读理解（共两节，满分50分）"
      T2: "第一节（共15小题；每小题2.5分，满分37.5分）"
      T3: "阅读下列短文，从每题所给的A、B、C和D四个选项中，选出最佳答案。"
    格式：
      - 中文字符 eastAsia=宋体；含 A-Z/0-9 的 run ascii=Times New Roman（含 docDefaults 继承）
      - 字号 = 小四 (24 half-pt)，含字符 run 必须全部满足（含样式继承）
      - 加粗：段内含字符 run 必须均有效加粗（含样式继承）
      - 左对齐（含样式继承）
      - 在第 1 页
      - 三段顺序相邻（T1→T2→T3，中间允许空段，不允许其它非空段）
    """
    TITLES = [
        "第二部分 阅读理解（共两节，满分50分）",
        "第一节（共15小题；每小题2.5分，满分37.5分）",
        "阅读下列短文，从每题所给的A、B、C和D四个选项中，选出最佳答案。",
    ]

    body_blocks = _body_block_elements(ctx)
    # 跳过空段，保留 body 顺序
    para_seq: list[tuple[etree._Element, str]] = []
    for el in body_blocks:
        if el.tag != W + "p":
            continue
        t = _block_text(el).strip()
        if not t:
            continue
        para_seq.append((el, t))

    # 在 para_seq 里按顺序找 T1，再要求紧跟 T2，再紧跟 T3
    idx1 = next((i for i, (_, t) in enumerate(para_seq) if t == TITLES[0]), None)
    if idx1 is None:
        return PointResult(1, "+1 第1页阅读理解三段标题 宋体/TNR小四加粗左对齐",
                           False, f"未找到整段相等的 '{TITLES[0]}'")
    if idx1 + 2 >= len(para_seq):
        return PointResult(1, "+1 第1页阅读理解三段标题 宋体/TNR小四加粗左对齐",
                           False, "T1 之后非空段落不足两段")
    e1, t1_actual = para_seq[idx1]
    e2, t2_actual = para_seq[idx1 + 1]
    e3, t3_actual = para_seq[idx1 + 2]
    adjacency_ok = (t2_actual == TITLES[1] and t3_actual == TITLES[2])
    if not adjacency_ok:
        return PointResult(1, "+1 第1页阅读理解三段标题 宋体/TNR小四加粗左对齐",
                           False, f"三段未顺序相邻：实为 T1='{t1_actual[:25]}' / "
                                  f"T2='{t2_actual[:25]}' / T3='{t3_actual[:25]}'")

    # 页号检查
    page_idx_map = {id(el): pg for el, pg in page_index_per_block(ctx)}
    pages = [page_idx_map.get(id(e1)), page_idx_map.get(id(e2)), page_idx_map.get(id(e3))]
    page_ok = all(pg == 1 for pg in pages)

    py_by_id: dict[int, Any] = {id(p._p): p for p in ctx.paragraphs}
    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    def _check_one(el: etree._Element, target_text: str) -> dict[str, Any]:
        runs_with_chars = [r for r in xp(el, ".//w:r")
                           if "".join(xp(r, ".//w:t/text()")).strip()]
        # 中文/数字英文 字体
        cn_bad: list[str] = []
        en_bad: list[str] = []
        for r in runs_with_chars:
            eff = _effective_run_font(r, defaults)
            txt = "".join(xp(r, ".//w:t/text()"))
            # 中文 + 中文标点
            if re.search(r"[一-鿿，。、；：（）《》「」『』！？“”‘’]", txt):
                if eff.get("eastAsia") != "宋体":
                    cn_bad.append(txt[:8])
            if re.search(r"[A-Za-z0-9]", txt):
                if eff.get("ascii") != "Times New Roman":
                    en_bad.append(txt[:8])
        # 小四
        size_info = _paragraph_effective_size(ctx, el, default_size)
        sizes_with_text: list[Optional[int]] = []
        for r, sz in zip([r for r in xp(el, ".//w:r")
                          if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                         size_info["sizes"]):
            if "".join(xp(r, ".//w:t/text()")).strip():
                sizes_with_text.append(sz)
        size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_XIAOSI_HALF_PT) <= 0 for s in sizes_with_text)
        # 加粗
        bold_ok = _effective_paragraph_bold(ctx, el)
        # 左对齐
        py = py_by_id.get(id(el))
        align_val = effective_paragraph_alignment(py) if py is not None else None
        left_ok = align_val == int(WD_ALIGN_PARAGRAPH.LEFT)
        return {
            "text_ok": _block_text(el).strip() == target_text,
            "cn_font_ok": not cn_bad,
            "en_font_ok": not en_bad,
            "size_ok": size_ok,
            "bold_ok": bold_ok,
            "left_ok": left_ok,
            "cn_bad": cn_bad, "en_bad": en_bad, "sizes": sizes_with_text,
            "align_val": align_val,
        }

    results = [_check_one(e, t) for e, t in zip((e1, e2, e3), TITLES)]
    all_lines_ok = all(all([r["text_ok"], r["cn_font_ok"], r["en_font_ok"],
                            r["size_ok"], r["bold_ok"], r["left_ok"]]) for r in results)
    passed = all_lines_ok and page_ok

    misses: list[str] = []
    for i, r in enumerate(results, start=1):
        m: list[str] = []
        if not r["text_ok"]: m.append(f"T{i}非整段相等")
        if not r["cn_font_ok"]: m.append(f"T{i}中文非宋体{r['cn_bad'][:3]}")
        if not r["en_font_ok"]: m.append(f"T{i}数字/英文非TNR{r['en_bad'][:3]}")
        if not r["size_ok"]: m.append(f"T{i}非小四 sizes={r['sizes']}")
        if not r["bold_ok"]: m.append(f"T{i}非加粗")
        if not r["left_ok"]: m.append(f"T{i}非左对齐 alignment={r['align_val']}")
        if m:
            misses.append("；".join(m))
    if not page_ok:
        misses.append(f"不在第1页 pages={pages}")
    detail = "OK 三段相邻 + 全部满足" if passed else "不通过：" + "；".join(misses[:6])
    return PointResult(1, "+1 第1页阅读理解三段标题 宋体/TNR小四加粗左对齐",
                       passed, detail)


def p_article_labels(ctx: EvalContext) -> PointResult:
    """+1 阅读文章编号 A/B/C/D 独立成段：

    细则字面：
      a. A、B、C、D 分别独立成段（即各自单独成一段，段内仅这一个字母）
      b. 字体 Times New Roman
      c. 字号 小四
      d. 加粗
      e. 居中对齐
      e. 水平居中
    （细则未点名"位置在阅读理解部分"或"必须按 A→B→C→D 顺序"，本项不做该限制。
    "居中对齐"在 OOXML 中是 w:jc=center，与"水平居中"是同一个对齐值。）
    """
    body_blocks = _body_block_elements(ctx)
    para_seq: list[tuple[etree._Element, str]] = []
    for el in body_blocks:
        if el.tag != W + "p":
            continue
        para_seq.append((el, _block_text(el).strip()))

    # 找"整段相等于 A/B/C/D"的段：每个字母至少存在一段
    expected = ["A", "B", "C", "D"]
    found_paras: dict[str, etree._Element] = {}
    for label in expected:
        for el, t in para_seq:
            if t == label:
                found_paras[label] = el
                break

    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    def _check_label(el: Optional[etree._Element], label: str) -> dict[str, Any]:
        if el is None:
            return {"missing": True, "label": label}
        runs_with_chars = [r for r in xp(el, ".//w:r")
                           if "".join(xp(r, ".//w:t/text()")).strip()]
        # 字体：A/B/C/D 都是 ASCII，应该 ascii=Times New Roman
        en_bad: list[str] = []
        for r in runs_with_chars:
            eff = _effective_run_font(r, defaults)
            txt = "".join(xp(r, ".//w:t/text()"))
            if re.search(r"[A-Za-z0-9]", txt) and eff.get("ascii") != "Times New Roman":
                en_bad.append(txt[:8])
        # 小四
        size_info = _paragraph_effective_size(ctx, el, default_size)
        sizes_with_text: list[Optional[int]] = []
        for r, sz in zip([r for r in xp(el, ".//w:r")
                          if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                         size_info["sizes"]):
            if "".join(xp(r, ".//w:t/text()")).strip():
                sizes_with_text.append(sz)
        size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_XIAOSI_HALF_PT) <= 0 for s in sizes_with_text)
        # 加粗
        bold_ok = _effective_paragraph_bold(ctx, el)
        # 居中：直接读段 XML 的 w:jc/@w:val（不经 python-docx，避免对象 id 不匹配）
        jc_el = fd(el, "w:pPr/w:jc")
        jc_val = jc_el.get(W + "val") if jc_el is not None else None
        center_ok = jc_val == "center"
        align_val = jc_val
        return {
            "missing": False, "label": label,
            "en_font_ok": not en_bad, "size_ok": size_ok,
            "bold_ok": bold_ok, "center_ok": center_ok,
            "en_bad": en_bad, "sizes": sizes_with_text, "align_val": align_val,
        }

    results = [_check_label(found_paras.get(lab), lab) for lab in expected]

    misses: list[str] = []
    for r in results:
        if r.get("missing"):
            misses.append(f"'{r['label']}' 未独立成段")
            continue
        m: list[str] = []
        if not r["en_font_ok"]: m.append(f"非TNR{r['en_bad'][:2]}")
        if not r["size_ok"]: m.append(f"非小四 sizes={r['sizes']}")
        if not r["bold_ok"]: m.append("非加粗")
        if not r["center_ok"]: m.append(f"非居中 alignment={r['align_val']}")
        if m:
            misses.append(f"{r['label']}: " + "；".join(m))

    passed = not misses
    detail = "OK A/B/C/D 各独立成段 / TNR / 小四 / 加粗 / 居中" if passed else "不通过：" + "；".join(misses[:6])
    return PointResult(1, "+1 阅读文章编号 A/B/C/D 独立成段 TNR小四加粗居中对齐",
                       passed, detail)


def table_cell_widths_cm(table: Any) -> list[float]:
    widths = []
    if not table.rows:
        return widths
    for cell in table.rows[0].cells:
        tcw = cell._tc.tcPr.tcW
        if tcw is not None and tcw.w:
            try:
                widths.append(float(tcw.w) / TWIPS_PER_CM)
            except Exception:
                widths.append(0.0)
        else:
            widths.append(0.0)
    return widths


def p_course_table(ctx: EvalContext) -> PointResult:
    """+5 第1页课程信息表 严格 5×5 + 列宽/总宽 + 居中 + 字体 + 边框 + 垂直居中 + 内容完整。

    a. 在第 1 页
    b. 5 行 × 5 列；不能是图片（找到含 Programme/Main focus 的可编辑 Word 表）
    c. 列宽 4.30 / 3.20 / 2.50 / 2.70 / 5.76 cm（容差 ±0.05cm）；总宽 18.46 cm（±0.05cm）
    d. 表格水平居中（tblPr/jc=center）
    e. 表头与内容字体 = Times New Roman（ascii 槽，含 docDefaults 继承）+ 小四（sz=24）
    f. 表头加粗（含字符 run 全部有效加粗）；数据行不加粗（含字符 run 全部非加粗）
    g. 所有单元格垂直居中：每个 w:tc 必须有 w:vAlign val=center
    h. 边框：6 种（top/left/bottom/right/insideH/insideV）均为 w:val=single, w:sz=4, w:color=000000 或 auto
    i. 内容完整：5×5=25 单元格表头+四行数据按指定文本
    """
    target_widths = [4.30, 3.20, 2.50, 2.70, 5.76]
    # 用 XML 元素而非 python-docx 表格查找：page_index_per_block 基于 lxml 解析的元素 id，
    # 与 python-docx 的 t._tbl 不是同一对象，必须用同一棵 XML 树。
    page_idx_map = {id(el): pg for el, pg in page_index_per_block(ctx)}
    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    target_tbl_el: Optional[etree._Element] = None
    for el in _body_block_elements(ctx):
        if el.tag != W + "tbl":
            continue
        text = "".join(xp(el, ".//w:t/text()"))
        if "Programme" in text and "Main focus" in text:
            target_tbl_el = el
            break
    if target_tbl_el is None:
        return PointResult(5, "+5 第1页课程信息表 5×5 列宽/居中/TNR小四/表头加粗/单实线0.5磅黑色/单元格垂直居中",
                           False, "未找到课程信息表（可能被截图代替）")

    tbl_el = target_tbl_el
    rows_xml = xp(tbl_el, "w:tr")
    rows = len(rows_xml)
    cols = len(xp(rows_xml[0], "w:tc")) if rows_xml else 0
    failures: list[str] = []

    # a. 第 1 页
    page_num = page_idx_map.get(id(tbl_el))
    if page_num != 1:
        failures.append(f"不在第1页(页号={page_num})")

    # b. 维度 5×5
    if rows != 5 or cols != 5:
        failures.append(f"维度={rows}×{cols} 非 5×5")

    # c. 列宽 / 总宽（从 tblGrid/gridCol 读，比 tcW 更稳定）
    grid_cols = xp(tbl_el, "w:tblGrid/w:gridCol")
    widths: list[float] = []
    for gc in grid_cols:
        w_attr = gc.get(W + "w")
        if w_attr is None:
            continue
        try:
            widths.append(float(w_attr) / TWIPS_PER_CM)
        except Exception:
            widths.append(0.0)
    width_misses: list[str] = []
    if len(widths) != 5:
        width_misses.append(f"列数={len(widths)}")
    else:
        for i, (w, tw) in enumerate(zip(widths, target_widths)):
            if not approx(w, tw, 0.05):
                width_misses.append(f"列{i+1}={w:.2f}cm≠{tw}cm")
    if widths and not approx(sum(widths), 18.46, 0.05):
        width_misses.append(f"总宽={sum(widths):.2f}cm≠18.46cm")
    if width_misses:
        failures.append("列宽/总宽：" + "；".join(width_misses))

    # d. 表格水平居中
    jc = fd(tbl_el, "w:tblPr/w:jc")
    table_centered = jc is not None and jc.get(W + "val") == "center"
    if not table_centered:
        failures.append(f"表格未水平居中(tblPr/jc={jc.get(W+'val') if jc is not None else None})")

    # e. 字体 ascii=TNR + 字号 sz=24
    bad_font: list[str] = []
    bad_size: list[str] = []
    for tc in xp(tbl_el, ".//w:tc"):
        for p_el in xp(tc, ".//w:p"):
            size_info = _paragraph_effective_size(ctx, p_el, default_size)
            for r, sz in zip([r for r in xp(p_el, ".//w:r")
                              if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                             size_info["sizes"]):
                txt = "".join(xp(r, ".//w:t/text()"))
                if not txt.strip():
                    continue
                eff = _effective_run_font(r, defaults)
                if re.search(r"[A-Za-z0-9]", txt) and eff.get("ascii") != "Times New Roman":
                    bad_font.append(f"'{txt[:14]}' ascii={eff.get('ascii')}")
                if sz is None or abs(sz - FONT_SZ_XIAOSI_HALF_PT) > 0:
                    bad_size.append(f"'{txt[:14]}' sz={sz}")
    if bad_font:
        failures.append(f"字体非TNR({len(bad_font)})：" + "；".join(bad_font[:3]))
    if bad_size:
        failures.append(f"字号非小四({len(bad_size)})：" + "；".join(bad_size[:3]))

    # f. 表头加粗 / 数据不加粗（含样式继承）
    header_bold_bad: list[str] = []
    data_bold_bad: list[str] = []
    rows_xml = xp(tbl_el, "w:tr")
    for ri, tr in enumerate(rows_xml):
        is_header = ri == 0
        for p_el in xp(tr, ".//w:p"):
            txt = _block_text(p_el).strip()
            if not txt:
                continue
            eff_bold = _effective_paragraph_bold(ctx, p_el)
            if is_header and not eff_bold:
                header_bold_bad.append(f"'{txt[:14]}'")
            if (not is_header) and eff_bold:
                data_bold_bad.append(f"'{txt[:14]}'")
    if header_bold_bad:
        failures.append(f"表头未加粗({len(header_bold_bad)})：" + "；".join(header_bold_bad[:3]))
    if data_bold_bad:
        failures.append(f"数据加粗({len(data_bold_bad)})：" + "；".join(data_bold_bad[:3]))

    # g. 所有单元格垂直居中：每个 tc 必须有 vAlign=center（缺省值是 top）
    valign_bad: list[str] = []
    for ri, tr in enumerate(rows_xml):
        for ci, tc in enumerate(xp(tr, "w:tc")):
            v = fd(tc, "w:tcPr/w:vAlign")
            v_val = v.get(W + "val") if v is not None else None
            if v_val != "center":
                valign_bad.append(f"({ri},{ci})vAlign={v_val}")
    if valign_bad:
        failures.append(f"垂直居中({len(valign_bad)})：" + "；".join(valign_bad[:3]))

    # h. 边框：6 种均为 single + sz=4 + color 黑色
    required_sides = ["top", "left", "bottom", "right", "insideH", "insideV"]
    border_bad: list[str] = []
    tbl_borders = fd(tbl_el, "w:tblPr/w:tblBorders")
    if tbl_borders is None:
        border_bad.append("缺少 tblBorders")
    else:
        for side in required_sides:
            b = fd(tbl_borders, f"w:{side}")
            if b is None:
                border_bad.append(f"缺{side}")
                continue
            val = b.get(W + "val")
            sz = b.get(W + "sz")
            color = (b.get(W + "color") or "").upper()
            if val != "single":
                border_bad.append(f"{side}.val={val}")
            if sz != "4":
                border_bad.append(f"{side}.sz={sz}")
            if color not in ("000000", "AUTO", ""):
                border_bad.append(f"{side}.color={color}")
    if border_bad:
        failures.append("边框：" + "；".join(border_bad[:6]))

    # i. 内容完整
    missing = table_content_missing(ctx)
    if missing:
        failures.append(f"内容缺失({len(missing)})：" + "；".join(missing[:5]))

    passed = not failures
    summary = (f"{rows}×{cols}；列宽={[round(w,2) for w in widths]}；总宽={sum(widths):.2f}；"
               f"水平居中={table_centered}；表头加粗={'OK' if not header_bold_bad else f'缺{len(header_bold_bad)}'}；"
               f"数据非加粗={'OK' if not data_bold_bad else f'有{len(data_bold_bad)}加粗'}；"
               f"垂直居中={'OK' if not valign_bad else f'缺{len(valign_bad)}'}；"
               f"边框={'OK' if not border_bad else f'问题{len(border_bad)}'}；"
               f"页号={page_num}")
    if passed:
        detail = "OK 全部满足 | " + summary
    else:
        detail = summary + " | 不通过：" + "；".join(failures[:6])
    return PointResult(5, "+5 第1页课程信息表 5×5 列宽/居中/TNR小四/表头加粗/单实线0.5磅黑色/单元格垂直居中",
                       passed, detail)


def p_writing_section1(ctx: EvalContext) -> PointResult:
    """+1 第7页写作第一节：

    细则字面只有 3 类点：
      a. 保留以下内容（按文档顺序出现于写作第一节范围内）：
         - 第 66 题（段首"66."）
         - 标题 "Learning with a Clear Purpose"
         - 名言 "A goal without action is only a wish; action without a goal easily loses its way."
         - 三项写作内容（"（1）…"；"（2）…"；"（3）…"）
         - 注意事项（"注意："段）
      b. 中文 = 宋体 小四（eastAsia=宋体 + sz=24）
      c. 数字和英文 = Times New Roman 小四（ascii=Times New Roman + sz=24）
    （细则未点名"第7页位置 / 颜色 / 对齐 / 答题横线段"，本项不做该限制。）
    """
    body_blocks = _body_block_elements(ctx)
    para_seq: list[tuple[etree._Element, str]] = []
    for el in body_blocks:
        if el.tag != W + "p":
            continue
        para_seq.append((el, _block_text(el).strip()))

    idx_start = next((i for i, (_, t) in enumerate(para_seq) if t.startswith("66.")), None)
    if idx_start is None:
        return PointResult(1, "+1 写作第一节 内容保留 + 中文宋体 数字英文TNR 小四",
                           False, "未找到段首 '66.' 第一节题号")
    idx_end = next((i for i, (_, t) in enumerate(para_seq[idx_start + 1:], start=idx_start + 1)
                    if "第二节" in t and "满分25分" in t), None)
    if idx_end is None:
        idx_end = len(para_seq)
    section_paras = para_seq[idx_start:idx_end]
    section_text = "\n".join(t for _, t in section_paras)

    # a. 内容保留
    content_misses: list[str] = []
    if not re.search(r"^\s*66\.", section_paras[0][1]):
        content_misses.append("第66题题号未在写作第一节首段")
    if "Learning with a Clear Purpose" not in section_text:
        content_misses.append("缺标题 'Learning with a Clear Purpose'")
    QUOTE = "A goal without action is only a wish; action without a goal easily loses its way."
    if QUOTE not in section_text:
        content_misses.append("缺名言（A goal without action ...）")
    content_listing_text = "\n".join(t for _, t in section_paras[:6])
    has_1 = bool(re.search(r"[（(]\s*1\s*[)）]", content_listing_text))
    has_2 = bool(re.search(r"[（(]\s*2\s*[)）]", content_listing_text))
    has_3 = bool(re.search(r"[（(]\s*3\s*[)）]", content_listing_text))
    if not (has_1 and has_2 and has_3):
        content_misses.append(f"三项写作内容缺失：(1)={has_1} (2)={has_2} (3)={has_3}")
    if "注意：" not in section_text and "注意:" not in section_text:
        content_misses.append("缺'注意：'段")

    # b/c. 字体/字号校验：范围内所有非空段
    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)
    bad_size: list[str] = []
    bad_cn_font: list[str] = []
    bad_en_font: list[str] = []
    for el, t in section_paras:
        if not t:
            continue
        size_info = _paragraph_effective_size(ctx, el, default_size)
        preview = t[:25]
        for r, sz in zip([rr for rr in xp(el, ".//w:r")
                          if "".join(xp(rr, ".//w:t/text()")) or xp(rr, ".//w:instrText")],
                         size_info["sizes"]):
            txt = "".join(xp(r, ".//w:t/text()"))
            if not txt.strip():
                continue
            eff = _effective_run_font(r, defaults)
            if re.search(r"[一-鿿，。、；：（）《》「」『』！？“”‘’]", txt):
                if eff.get("eastAsia") != "宋体":
                    bad_cn_font.append(f"段'{preview}' run='{txt[:10]}' eastAsia={eff.get('eastAsia')}")
            if re.search(r"[A-Za-z0-9]", txt):
                if eff.get("ascii") != "Times New Roman":
                    bad_en_font.append(f"段'{preview}' run='{txt[:10]}' ascii={eff.get('ascii')}")
            if sz is None or abs(sz - FONT_SZ_XIAOSI_HALF_PT) > 0:
                bad_size.append(f"段'{preview}' run='{txt[:10]}' sz={sz}")

    failures: list[str] = []
    if content_misses:
        failures.append("内容缺失：" + "；".join(content_misses[:4]))
    if bad_cn_font:
        failures.append(f"中文非宋体({len(bad_cn_font)})：" + "；".join(bad_cn_font[:3]))
    if bad_en_font:
        failures.append(f"数字/英文非TNR({len(bad_en_font)})：" + "；".join(bad_en_font[:3]))
    if bad_size:
        failures.append(f"非小四({len(bad_size)})：" + "；".join(bad_size[:3]))

    passed = not failures
    detail = (f"段数={len(section_paras)}"
              + (" | OK 全部满足" if passed else " | 不通过：" + "；".join(failures[:6])))
    return PointResult(1, "+1 写作第一节 内容保留 + 中文宋体 数字英文TNR 小四",
                       passed, detail)


def _para_style_italic(ctx: EvalContext, style_id: Optional[str]) -> Optional[bool]:
    """沿样式继承链查找 w:rPr/w:i；找到返回 True/False，未声明返回 None。"""
    if not style_id or ctx.styles_xml is None:
        return None
    seen: set[str] = set()
    cur = style_id
    while cur and cur not in seen:
        seen.add(cur)
        style = None
        for s in xp(ctx.styles_xml, ".//w:style"):
            if s.get(W + "styleId") == cur:
                style = s
                break
        if style is None:
            return None
        i = fd(style, "w:rPr/w:i")
        if i is not None:
            return i.get(W + "val", "1") not in ("0", "false", "False")
        based = fd(style, "w:basedOn")
        cur = based.get(W + "val") if based is not None else None
    return None


def _effective_paragraph_italic(ctx: EvalContext, p_el: etree._Element) -> bool:
    """段内每个含字符 run 的有效倾斜状态（run 直接 w:i → 段落样式链 → 默认段落样式）必须为真。"""
    p_style = fd(p_el, "w:pPr/w:pStyle")
    style_id = p_style.get(W + "val") if p_style is not None else _default_paragraph_style_id(ctx)
    style_italic = _para_style_italic(ctx, style_id)
    runs = [r for r in xp(p_el, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    if not runs:
        return False
    for r in runs:
        i = fd(r, "w:rPr/w:i")
        if i is not None:
            on = i.get(W + "val", "1") not in ("0", "false", "False")
        else:
            on = bool(style_italic)
        if not on:
            return False
    return True


def p_opening_sentence(ctx: EvalContext) -> PointResult:
    """+1 开头句 'Learning with a clear purpose matters greatly to Grade 12 students.'
    Times New Roman 小四号 加粗 倾斜。

    细则字面只 4 个属性：
      a. Times New Roman（含 ASCII 字符 run ascii=Times New Roman，含 docDefaults 继承）
      b. 小四（24 half-pt，含样式继承）
      c. 加粗（含继承）
      d. 倾斜（含继承）
    （细则未点名"在第7页"或"整段就是这句"，本项不做该限制——用包含匹配定位段。）
    """
    SENTENCE = "Learning with a clear purpose matters greatly to Grade 12 students."

    body_blocks = _body_block_elements(ctx)
    target_el: Optional[etree._Element] = None
    for el in body_blocks:
        if el.tag != W + "p":
            continue
        if SENTENCE in _block_text(el):
            target_el = el
            break
    if target_el is None:
        return PointResult(1, "+1 开头句 TNR 小四 加粗 倾斜",
                           False, f"未找到包含开头句的段 '{SENTENCE[:30]}...'")

    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    runs_with_chars = [r for r in xp(target_el, ".//w:r")
                       if "".join(xp(r, ".//w:t/text()")).strip()]
    # 字体：含 ASCII 字符 run 必须 ascii=TNR
    en_bad: list[str] = []
    for r in runs_with_chars:
        eff = _effective_run_font(r, defaults)
        txt = "".join(xp(r, ".//w:t/text()"))
        if re.search(r"[A-Za-z0-9]", txt) and eff.get("ascii") != "Times New Roman":
            en_bad.append(f"'{txt[:12]}' ascii={eff.get('ascii')}")
    # 字号 小四
    size_info = _paragraph_effective_size(ctx, target_el, default_size)
    sizes_with_text: list[Optional[int]] = []
    for r, sz in zip([r for r in xp(target_el, ".//w:r")
                      if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                     size_info["sizes"]):
        if "".join(xp(r, ".//w:t/text()")).strip():
            sizes_with_text.append(sz)
    size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_XIAOSI_HALF_PT) <= 0 for s in sizes_with_text)
    # 加粗 / 倾斜（含继承）
    bold_ok = _effective_paragraph_bold(ctx, target_el)
    italic_ok = _effective_paragraph_italic(ctx, target_el)

    misses: list[str] = []
    if en_bad:
        misses.append(f"非TNR({len(en_bad)})：" + "；".join(en_bad[:3]))
    if not size_ok:
        misses.append(f"非小四 sizes={sizes_with_text}")
    if not bold_ok:
        misses.append("非加粗")
    if not italic_ok:
        misses.append("非倾斜")

    passed = not misses
    detail = "OK TNR/小四/加粗/倾斜 全部满足" if passed else "不通过：" + "；".join(misses[:6])
    return PointResult(1, "+1 开头句 TNR 小四 加粗 倾斜", passed, detail)


def underline_paragraphs_after(ctx: EvalContext, start_text: str, until_text: Optional[str] = None) -> list[tuple[Any, bool, bool]]:
    lines = []
    in_zone = False
    for p in ctx.paragraphs:
        t = p.text.strip()
        if start_text in t:
            in_zone = True
            continue
        if in_zone and until_text and until_text in t:
            break
        if not in_zone:
            continue
        p_el = p._p
        chars_line = bool(re.fullmatch(r"[_\s]{20,}", t))
        border_line = bool(xp(p_el, ".//w:pBdr/w:bottom"))
        if chars_line or border_line:
            lines.append((p, chars_line, border_line))
    return lines


def underline_paragraph_count_after(ctx: EvalContext, start_text: str, until_text: Optional[str] = None) -> int:
    return len(underline_paragraphs_after(ctx, start_text, until_text))


def underline_chars_format_ok(p: Any) -> bool:
    t = p.text.strip()
    if not re.fullmatch(r"[_\s]{20,}", t):
        return False
    e = p._p
    runs = [r for r in xp(e, ".//w:r") if "".join(xp(r, ".//w:t/text()")).strip()]
    return bool(runs) and size_ok_for_runs(e, FONT_SZ_XIAOSI_HALF_PT) and all(run_xml_ascii_font(r, "Times New Roman") for r in runs)


def p_section1_lines(ctx: EvalContext) -> PointResult:
    """+1 第7页第一节答题横线：

    细则点：
      a. 在开头句下方
      b. 3 条可编辑下划线段落
      c. 每条横线覆盖正文宽度
      d. TNR 小四号 下划线字符 或 段落底边框
    """
    SENTENCE = "Learning with a clear purpose matters greatly to Grade 12 students."
    SECTION2 = "第二节（满分25分）"

    body_paras = _all_body_paragraph_elements(ctx)
    # a. 找开头句段位置
    idx_start = next((i for i, el in enumerate(body_paras)
                      if el.tag == W + "p" and _block_text(el).strip() == SENTENCE), None)
    if idx_start is None:
        return PointResult(1, "+1 第7页第一节答题横线 3条 TNR小四下划线字符或底边框 覆盖正文宽度",
                           False, "未找到开头句段")
    idx_end = next((i for i, el in enumerate(body_paras[idx_start + 1:], start=idx_start + 1)
                    if el.tag == W + "p" and SECTION2 in _block_text(el)), len(body_paras))

    # 收集 [开头句之后, 第二节之前) 的下划线段：字符横线 或 底边框段
    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)
    BODY_WIDTH_TWIPS = int(round(BODY_WIDTH_CM * TWIPS_PER_CM))  # 18.46 cm ≈ 10468 twips

    found_paras: list[tuple[etree._Element, str, dict[str, Any]]] = []
    for el in body_paras[idx_start + 1:idx_end]:
        if el.tag != W + "p":
            continue
        text = _block_text(el).strip()
        is_char_line = bool(text) and all(c in "_ \t " for c in text) and text.count("_") >= 20
        bottom_border = fd(el, "w:pPr/w:pBdr/w:bottom")
        is_border_line = bottom_border is not None and (bottom_border.get(W + "val") not in (None, "nil", "none"))
        if not is_char_line and not is_border_line:
            continue

        # d. TNR 小四（针对字符横线段；底边框段同样要求 TNR 小四）
        runs_with_chars = [r for r in xp(el, ".//w:r")
                           if "".join(xp(r, ".//w:t/text()")).strip()]
        bad_font: list[str] = []
        for r in runs_with_chars:
            eff = _effective_run_font(r, defaults)
            if eff.get("ascii") != "Times New Roman":
                bad_font.append(f"ascii={eff.get('ascii')}")
        size_info = _paragraph_effective_size(ctx, el, default_size)
        sizes_with_text: list[Optional[int]] = []
        for r, sz in zip([r for r in xp(el, ".//w:r")
                          if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                         size_info["sizes"]):
            if "".join(xp(r, ".//w:t/text()")).strip():
                sizes_with_text.append(sz)
        size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_XIAOSI_HALF_PT) <= 0 for s in sizes_with_text)

        # c. 覆盖正文宽度
        # 字符横线段：用 字符数 × 字符宽 估算视觉宽度（小四 TNR 下 underscore 字符约 6.5pt = 130 twips/字）
        # 底边框段：底边框横向跨度由段落 ind 决定；left+right=0 即覆盖正文宽度
        if is_char_line:
            char_count = text.count("_")
            CHAR_TWIPS = 130  # underscore @ TNR 12pt 视觉宽度估算
            est_twips = char_count * CHAR_TWIPS
            width_cover_ok = est_twips >= BODY_WIDTH_TWIPS - 600  # 容差 ~1cm
            width_info = f"字符数={char_count}≈{est_twips}twips"
        else:
            ind = fd(el, "w:pPr/w:ind")
            left_tw = right_tw = 0
            if ind is not None:
                if ind.get(W + "left") is not None:
                    try: left_tw = int(ind.get(W + "left"))
                    except Exception: pass
                if ind.get(W + "right") is not None:
                    try: right_tw = int(ind.get(W + "right"))
                    except Exception: pass
            width_cover_ok = abs(left_tw) <= 60 and abs(right_tw) <= 60  # ±0.1cm
            width_info = f"left={left_tw} right={right_tw}twips"

        found_paras.append((el, text, {
            "is_char_line": is_char_line, "is_border_line": is_border_line,
            "bad_font": bad_font, "size_ok": size_ok, "sizes": sizes_with_text,
            "width_cover_ok": width_cover_ok, "width_info": width_info,
        }))

    failures: list[str] = []
    # b. 数量精确为 3
    if len(found_paras) != 3:
        failures.append(f"横线段数量={len(found_paras)} 非 3 条")

    # d/c. 每条都必须满足 TNR + 小四 + 覆盖正文宽度
    for idx, (_el, text, info) in enumerate(found_paras, start=1):
        miss: list[str] = []
        if info["bad_font"]:
            miss.append(f"非TNR({info['bad_font'][:1]})")
        if not info["size_ok"]:
            miss.append(f"非小四 sizes={info['sizes']}")
        if not info["width_cover_ok"]:
            miss.append(f"未覆盖正文宽度 ({info['width_info']})")
        if miss:
            kind = "字符横线" if info["is_char_line"] else "底边框"
            failures.append(f"第{idx}条({kind})：" + "；".join(miss))

    passed = not failures
    summary = f"找到下划线段={len(found_paras)} (字符横线={sum(1 for _,_,i in found_paras if i['is_char_line'])}, 底边框={sum(1 for _,_,i in found_paras if i['is_border_line'])})"
    detail = (summary + " | OK 全部满足") if passed else (summary + " | 不通过：" + "；".join(failures[:6]))
    return PointResult(1, "+1 第7页第一节答题横线 3条 TNR小四下划线字符或底边框 覆盖正文宽度",
                       passed, detail)


def p_paragraph_prompts(ctx: EvalContext) -> PointResult:
    """+1 Paragraph 1 / Paragraph 2 文本：

    细则字面点：
      a. 完整显示两段（用整段精确等于定位"完整显示"）
      b. Times New Roman（含 ASCII 字符 run ascii=Times New Roman，含 docDefaults 继承）
      c. 小四（24 half-pt，含字符 run 全部满足，含样式继承）
      d. 加粗（含字符 run 全部有效加粗，含继承）
    （细则未点名"在第8页"，本项不做该限制。）
    """
    PROMPTS = [
        "Paragraph 1: I sat with him on the bench near the stairway and opened his notebook.",
        "Paragraph 2: When I returned to my own classroom on Monday, I found a small card on my desk.",
    ]

    body_blocks = _body_block_elements(ctx)
    defaults = _doc_default_fonts(ctx)
    default_size = _doc_default_size(ctx)

    misses_all: list[str] = []
    for i, target_text in enumerate(PROMPTS, start=1):
        # a. 整段精确等于（细则"完整显示"）
        target_el: Optional[etree._Element] = None
        for el in body_blocks:
            if el.tag != W + "p":
                continue
            if _block_text(el).strip() == target_text:
                target_el = el
                break
        if target_el is None:
            misses_all.append(f"P{i}: 未找到整段相等的文本")
            continue

        # b. Times New Roman
        runs_with_chars = [r for r in xp(target_el, ".//w:r")
                           if "".join(xp(r, ".//w:t/text()")).strip()]
        en_bad: list[str] = []
        for r in runs_with_chars:
            eff = _effective_run_font(r, defaults)
            txt = "".join(xp(r, ".//w:t/text()"))
            if re.search(r"[A-Za-z0-9]", txt) and eff.get("ascii") != "Times New Roman":
                en_bad.append(f"'{txt[:12]}' ascii={eff.get('ascii')}")
        if en_bad:
            misses_all.append(f"P{i}: 非TNR({len(en_bad)})：" + "；".join(en_bad[:2]))

        # c. 小四
        size_info = _paragraph_effective_size(ctx, target_el, default_size)
        sizes_with_text: list[Optional[int]] = []
        for r, sz in zip([r for r in xp(target_el, ".//w:r")
                          if "".join(xp(r, ".//w:t/text()")) or xp(r, ".//w:instrText")],
                         size_info["sizes"]):
            if "".join(xp(r, ".//w:t/text()")).strip():
                sizes_with_text.append(sz)
        size_ok = bool(sizes_with_text) and all(s is not None and abs(s - FONT_SZ_XIAOSI_HALF_PT) <= 0 for s in sizes_with_text)
        if not size_ok:
            misses_all.append(f"P{i}: 非小四 sizes={sizes_with_text}")

        # d. 加粗（含继承）
        bold_ok = _effective_paragraph_bold(ctx, target_el)
        if not bold_ok:
            misses_all.append(f"P{i}: 非加粗")

    passed = not misses_all
    detail = "OK P1+P2 完整 + TNR + 小四 + 加粗" if passed else "不通过：" + "；".join(misses_all[:6])
    return PointResult(1, "+1 Paragraph 1/2 完整 TNR 小四 加粗", passed, detail)


def p_continuation_lines(ctx: EvalContext) -> PointResult:
    """+1 第8页续写横线：

    细则字面点：
      a. 在第 8 页
      b. Paragraph 1 下方 3 条 + Paragraph 2 下方 3 条（各 3 条）
      c. 每条横线宽度一致
      d. 不超出页面左右边距
    （第 8 页判定：静态解析下段级页号不可靠——本文档无 lastRenderedPageBreak 标记时
    所有段都算第 1 页，因此本项不做段级页号校验，避免误判。）
    """
    P1_TEXT = "Paragraph 1: I sat with him on the bench near the stairway and opened his notebook."
    P2_TEXT = "Paragraph 2: When I returned to my own classroom on Monday, I found a small card on my desk."

    body_paras = _all_body_paragraph_elements(ctx)
    BODY_WIDTH_TWIPS = int(round(BODY_WIDTH_CM * TWIPS_PER_CM))  # ≈ 10468
    CHAR_TWIPS = 130  # underscore @ TNR 小四 视觉宽度估算

    def _line_visual_twips(el: etree._Element, text: str) -> tuple[int, str]:
        is_char_line = bool(text) and all(c in "_ \t " for c in text) and text.count("_") >= 20
        if is_char_line:
            return text.count("_") * CHAR_TWIPS, f"字符横线×{text.count('_')}"
        bdr = fd(el, "w:pPr/w:pBdr/w:bottom")
        if bdr is not None:
            ind = fd(el, "w:pPr/w:ind")
            left_tw = right_tw = 0
            if ind is not None:
                try: left_tw = int(ind.get(W + "left") or 0)
                except Exception: pass
                try: right_tw = int(ind.get(W + "right") or 0)
                except Exception: pass
            return BODY_WIDTH_TWIPS - max(0, left_tw) - max(0, right_tw), f"底边框 left={left_tw} right={right_tw}"
        return 0, "未识别为横线段"

    def _is_underline_para(el: etree._Element, text: str) -> bool:
        if bool(text) and all(c in "_ \t " for c in text) and text.count("_") >= 20:
            return True
        bdr = fd(el, "w:pPr/w:pBdr/w:bottom")
        return bdr is not None and (bdr.get(W + "val") not in (None, "nil", "none"))

    # 找 P1 / P2 段位置（整段相等）
    p1_idx = next((i for i, el in enumerate(body_paras)
                   if el.tag == W + "p" and _block_text(el).strip() == P1_TEXT), None)
    p2_idx = next((i for i, el in enumerate(body_paras)
                   if el.tag == W + "p" and _block_text(el).strip() == P2_TEXT), None)
    if p1_idx is None or p2_idx is None:
        return PointResult(1, "+1 第8页续写横线 P1下/P2下各3条 宽度一致 不超左右边距",
                           False, f"未找到 P1段(idx={p1_idx}) 或 P2段(idx={p2_idx})")
    if p2_idx <= p1_idx:
        return PointResult(1, "+1 第8页续写横线 P1下/P2下各3条 宽度一致 不超左右边距",
                           False, "P2 段未出现在 P1 段之后")

    # a. 收集 P1 下方与 P2 下方的下划线段
    def _collect(from_idx: int, to_idx: Optional[int]) -> list[tuple[etree._Element, str]]:
        end = to_idx if to_idx is not None else len(body_paras)
        out: list[tuple[etree._Element, str]] = []
        for el in body_paras[from_idx + 1:end]:
            if el.tag != W + "p":
                continue
            t = _block_text(el).strip()
            if _is_underline_para(el, t):
                out.append((el, t))
        return out

    p1_lines = _collect(p1_idx, p2_idx)
    p2_lines = _collect(p2_idx, None)

    failures: list[str] = []

    # b. 数量
    if len(p1_lines) != 3:
        failures.append(f"P1下方={len(p1_lines)}条 非3条")
    if len(p2_lines) != 3:
        failures.append(f"P2下方={len(p2_lines)}条 非3条")

    # 关于"第 8 页"：段级页号无法可靠静态判定。本项依赖文档结构（P1/P2 段位于文档末尾
    # 续写区），不再单独做段级页号校验。

    # c. 每条横线宽度一致
    all_lines = p1_lines + p2_lines
    widths_info = [_line_visual_twips(el, t) for el, t in all_lines]
    widths = [w for w, _ in widths_info]
    if widths:
        if max(widths) - min(widths) > 200:  # 容差 ±0.35cm ≈ 200 twips
            failures.append(f"宽度不一致 widths={widths}")

    # d. 不超出页面左右边距
    over_lines: list[str] = []
    for (el, t), (w, info) in zip(all_lines, widths_info):
        if w > BODY_WIDTH_TWIPS + 200:
            over_lines.append(f"'{t[:14]}' 宽={w}twips>{BODY_WIDTH_TWIPS}({info})")
        bdr = fd(el, "w:pPr/w:pBdr/w:bottom")
        if bdr is not None:
            ind = fd(el, "w:pPr/w:ind")
            if ind is not None:
                try:
                    if int(ind.get(W + "left") or 0) < 0 or int(ind.get(W + "right") or 0) < 0:
                        over_lines.append(f"'{t[:14]}' 底边框段缩进为负超出边距")
                except Exception:
                    pass
    if over_lines:
        failures.append(f"超出左右边距({len(over_lines)})：" + "；".join(over_lines[:3]))

    passed = not failures
    detail = (f"P1下方={len(p1_lines)} P2下方={len(p2_lines)} 宽度twips={widths}")
    detail = (detail + " | OK") if passed else (detail + " | 不通过：" + "；".join(failures[:6]))
    return PointResult(1, "+1 第8页续写横线 P1下/P2下各3条 宽度一致 不超左右边距",
                       passed, detail)


def n_full_page_screenshot(_ctx: EvalContext) -> PointResult:
    """[已删除] -5 任意一页使用整页 PDF 截图，导致文字和题目无法选中编辑。

    此评分项按需求已删除，函数保留为占位以避免误引用；调用方不再包含它。
    始终返回未命中且 max_delta=0，即使被误调也不影响总分。
    """
    return PointResult(0, "[已删除] -5 任意一页使用整页PDF截图", False, "此项已按需求删除")


def n_table_is_image(ctx: EvalContext) -> PointResult:
    """-5 第1页课程信息表为截图或普通图片，不是可编辑 Word 表格。

    细则字面：
      a. 在第 1 页
      b. 课程信息表（含 Programme/Main focus 等标识词）
      c. 是截图或普通图片，而不是可编辑 Word 表格

    判定逻辑：
      1) 找含 Programme + Main focus 的可编辑 <w:tbl>，且该表在第 1 页 → 不违例
      2) 否则查第 1 页是否存在图片：若有，且第 1 页没有可编辑课程表 → 违例
      3) 若两者都没有（既无表也无图）→ 内容缺失，由其它项处理，不算"图片代替表格"
    """
    # 用 lxml 元素查可编辑 Word 表格（与 page_index_per_block 同一棵树）
    page_idx_map = {id(el): pg for el, pg in page_index_per_block(ctx)}
    course_table_page: Optional[int] = None
    for el in _body_block_elements(ctx):
        if el.tag != W + "tbl":
            continue
        text = "".join(xp(el, ".//w:t/text()"))
        if "Programme" in text and "Main focus" in text:
            course_table_page = page_idx_map.get(id(el))
            break

    if course_table_page == 1:
        return PointResult(-5, "-5 第1页课程信息表为截图或普通图片", False,
                           "第1页存在可编辑课程信息表")

    # 查第 1 页是否含图片/形状（page_index_per_block 给出图所在段或表所在块；
    # drawing/pict 通常嵌在段落里，按其所在段的页号定位）
    has_image_on_page1 = False
    image_info: list[str] = []
    for el in _body_block_elements(ctx):
        page = page_idx_map.get(id(el))
        if page != 1:
            continue
        # 段或表内是否含 drawing/pict
        drawings = xp(el, ".//w:drawing")
        picts = xp(el, ".//w:pict")
        if drawings or picts:
            has_image_on_page1 = True
            image_info.append(f"块在第1页含 drawing×{len(drawings)} pict×{len(picts)}")

    if has_image_on_page1 and course_table_page != 1:
        detail = f"第1页含图片但无可编辑课程信息表（课程表所在页={course_table_page}）；{'；'.join(image_info[:3])}"
        return PointResult(-5, "-5 第1页课程信息表为截图或普通图片", True, detail)

    # 既无可编辑课程表，也无第 1 页图片 → 内容缺失，由其它项处理
    return PointResult(-5, "-5 第1页课程信息表为截图或普通图片", False,
                       f"可编辑课程表所在页={course_table_page}；第1页有图片={has_image_on_page1}")


def n_table_content_missing(ctx: EvalContext) -> PointResult:
    """-1 课程信息表中的课程名称、日期、时间、费用或 Main focus 任意一项缺失。

    细则字面只点名 5 个**列名**（不涉及数据行）：
      - 课程名称 (Programme)
      - 日期 (Date)
      - 时间 (Time)
      - 费用 (Fee)
      - Main focus
    """
    # 找含课程信息表的可编辑表格
    table_text = "\n".join(
        "\t".join(cell.text for cell in row.cells)
        for t in ctx.tables for row in t.rows
        if "Programme" in "\n".join("\t".join(c.text for c in r.cells) for r in t.rows)
    )
    if not table_text:
        return PointResult(-1, "-1 课程信息表课程名称/日期/时间/费用/Main focus任一项缺失",
                           False, "未找到课程信息表")

    # 5 个列名匹配模式：核心词命中即算"该列存在"
    column_patterns = [
        ("课程名称", r"\bProgramme\b"),
        ("日期", r"\bDate\b"),
        ("时间", r"\bTime\b"),
        ("费用", r"\bFee\b"),
        ("Main focus", r"\bMain\s*focus\b"),
    ]
    missing = [name for name, pat in column_patterns if not re.search(pat, table_text, re.IGNORECASE)]
    return PointResult(-1, "-1 课程信息表课程名称/日期/时间/费用/Main focus任一项缺失",
                       bool(missing),
                       "缺失列：" + "；".join(missing) if missing else "5 个列名齐全")


def n_page_count_not_8(ctx: EvalContext) -> PointResult:
    """-1 文档页数不是 8 页，或出现第 9 页空白页。

    细则两个点（任一即扣）：
      1) 文档页数 ≠ 8
      2) 出现第 9 页空白页（用结构性判定，不依赖文本里出现"第9页"字样）
    """
    pages, method = infer_page_count(ctx)
    page_count_bad = pages != 8
    blank_page9, blank9_detail = detect_page9_blank(ctx)
    hit = page_count_bad or blank_page9

    parts: list[str] = [f"{method}: 页数={pages}"]
    if page_count_bad:
        parts.append(f"页数≠8")
    if blank_page9:
        parts.append(f"第9页空白({blank9_detail})")
    if not hit:
        parts.append("页数=8 且无第9页空白")
    return PointResult(-1, "-1 文档页数不是8页或出现第9页空白页", hit, "；".join(parts))


def n_writing_lines_fewer(_ctx: EvalContext) -> PointResult:
    """[已删除] -1 第7—8页写作答题横线不可少于三条。

    此评分项按需求已删除，函数保留为占位；调用方不再包含它。
    """
    return PointResult(0, "[已删除] -1 第7—8页写作答题横线少于三条", False, "此项已按需求删除")


def n_bad_word_break_or_overlap(_ctx: EvalContext) -> PointResult:
    """[已删除] -1 任意英文单词错误断开、字符重叠或超出右边界。

    此评分项按需求已删除，函数保留为占位；调用方不再包含它。
    """
    return PointResult(0, "[已删除] -1 任意英文单词错误断开、字符重叠或超出右边界", False, "此项已按需求删除")


def n_unrelated_marks(_ctx: EvalContext) -> PointResult:
    """[已删除] -3 无关批注、红色箭头、修订标记、水印或新增空白页。

    此评分项按需求已删除，函数保留为占位；调用方不再包含它。
    """
    return PointResult(0, "[已删除] -3 无关批注、红色箭头、修订标记、水印或新增空白页", False, "此项已按需求删除")


def dimension2_points(ctx: EvalContext) -> list[PointResult]:
    positives = [
        p_page_setup(ctx),
        p_pagination(ctx),
        p_footer_page_number(ctx),
        p_secret_text(ctx),
        p_main_title(ctx),
        p_date_text(ctx),
        p_body_fonts(ctx),
        p_body_spacing(ctx),
        p_question_paragraphs(ctx),
        p_options(ctx),
        p_reading_titles(ctx),
        p_article_labels(ctx),
        p_course_table(ctx),
        p_writing_section1(ctx),
        p_opening_sentence(ctx),
        p_section1_lines(ctx),
        p_paragraph_prompts(ctx),
        p_continuation_lines(ctx),
    ]
    negatives = [
        # n_full_page_screenshot / n_writing_lines_fewer /
        # n_bad_word_break_or_overlap / n_unrelated_marks 已按需求删除，不再计入。
        n_table_is_image(ctx),
        n_table_content_missing(ctx),
        n_page_count_not_8(ctx),
    ]
    return positives + negatives


SCRIPT_ID = "037"


def _locate_docx(dir_path: Path) -> Optional[Path]:
    """在指定目录内定位被评估的 .docx 文件。

    规则：目录下所有 ``*.docx``（忽略以 ``~$`` 开头的 Office 临时文件）中，
    优先返回最先按名称排序的那个；若无 .docx 返回 None。
    """
    if not dir_path.is_dir():
        return None
    candidates = sorted(
        p for p in dir_path.iterdir()
        if p.is_file() and p.suffix.lower() == ".docx" and not p.name.startswith("~$")
    )
    return candidates[0] if candidates else None


def _run_evaluation(path: Path) -> dict[str, Any]:
    """内部实现：加载文档并跑完维度一 + 维度二，返回统一结构字典。

    完全基于 python-docx / lxml 静态解析，不启动 Word COM，也不依赖 LibreOffice。
    """
    ctx = load_context(path)

    d1 = dimension1(ctx)
    file_name = path.name

    if not d1.passed:
        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": False,
            "dim1_reason": "；".join(d1.failures),
            "dim2_items": [],
            "total_score": 0,
            "max_score": 0,
        }

    points = dimension2_points(ctx)
    dim2_items: list[dict[str, Any]] = []
    total = 0
    max_score = 0
    for p in points:
        max_delta = p.score
        delta = p.score if p.passed else 0
        # 得分 = 所有加分项 + 减分项 的实际增量之和
        total += delta
        # 总分（满分）= 仅所有加分项的 max_delta 之和；减分项不计入满分
        if max_delta > 0:
            max_score += max_delta
        dim2_items.append({
            "rule": p.name,
            "max_delta": max_delta,
            "delta": delta,
            "hit": p.passed,
            "detail": "",
        })

    return {
        "id": SCRIPT_ID,
        "file_name": file_name,
        "status": "ok",
        "error": None,
        "dim1_pass": True,
        "dim1_reason": "",
        "dim2_items": dim2_items,
        "total_score": total,
        "max_score": max_score,
    }


def evaluate(dir_path: str) -> dict[str, Any]:
    """统一入口：接收"脚本所在目录的路径"，脚本自己在该目录里定位 .docx 并评估。

    参数：
        dir_path: 脚本所在目录路径（其中包含被评估的 docx 文档）。

    返回：见《脚本接口差异与统一建议.md》§2.2 的结构。
    脚本自身异常（找不到目录 / 找不到 docx / 打开失败等）会被吞掉，返回
    ``status="error"`` 而非抛出。
    """
    try:
        base = Path(dir_path)
        if not base.exists():
            return {
                "id": SCRIPT_ID,
                "file_name": "",
                "status": "error",
                "error": f"目录不存在：{base}",
                "dim1_pass": False,
                "dim1_reason": "",
                "dim2_items": [],
                "total_score": 0,
                "max_score": 0,
            }
        if not base.is_dir():
            return {
                "id": SCRIPT_ID,
                "file_name": "",
                "status": "error",
                "error": f"路径不是目录：{base}",
                "dim1_pass": False,
                "dim1_reason": "",
                "dim2_items": [],
                "total_score": 0,
                "max_score": 0,
            }
        docx_path = _locate_docx(base)
        if docx_path is None:
            return {
                "id": SCRIPT_ID,
                "file_name": "",
                "status": "error",
                "error": f"目录 {base} 下未找到可评估的 .docx 文件",
                "dim1_pass": False,
                "dim1_reason": "",
                "dim2_items": [],
                "total_score": 0,
                "max_score": 0,
            }
        return _run_evaluation(docx_path)
    except zipfile.BadZipFile as exc:
        return {
            "id": SCRIPT_ID,
            "file_name": "",
            "status": "error",
            "error": f"文件不是有效 docx/zip 包：{exc}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": 0,
        }
    except Exception as exc:
        return {
            "id": SCRIPT_ID,
            "file_name": "",
            "status": "error",
            "error": f"评估失败：{exc}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": 0,
        }


if __name__ == "__main__":
    # 仅用于本地调试：默认以脚本所在目录作为 dir_path。
    _debug_dir = sys.argv[1] if len(sys.argv) > 1 else str(Path(__file__).resolve().parent)
    print(json.dumps(evaluate(_debug_dir), ensure_ascii=False, indent=2))
