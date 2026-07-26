# -*- coding: utf-8 -*-
"""
Word文档自动评估系统
基于打分细则对Word文档进行格式核验
"""

import sys

SCRIPT_ID = "013"

# 遵循"评估脚本接口差异与统一建议.md" §2.3：
#   - 不允许修改 sys.stdout
#   - 主结果只走 evaluate(...) 的 return
#   - 脚本内部大量 print 都是调试信息，统一在此收集到缓冲区里（而非落到
#     任何流），保持 stdout / stderr 干净；同时用作 dim2_items[].detail
#     的原料，命中/未命中都能给出可读的说明（§2.2）。
_DEBUG_BUF: list[str] = []

def print(*args, **kwargs):  # noqa: A001  故意遮蔽内建 print，改为收集调试信息
    try:
        _DEBUG_BUF.append(' '.join(str(a) for a in args))
    except Exception:
        pass
    return

from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import parse_xml
import re
import os

class DocumentEvaluator:
    # 维度二"加分项"规则：(方法名, 规则描述, 满分)
    # 与 check_dimension2 中的调用顺序保持一致；max_delta 与方法内部
    # `self.score += N` 中的 N 保持一致，用于逐项还原 dim2_items。
    _DIM2_ADD_RULES = [
        ('check_abstract_title', '中文摘要标题为"摘 要"或"摘要"；宋体二号、加粗、居中；与下方文本空一行', 1),
        ('check_keywords', '摘要正文下"关键词"部分：四个关键词用中文逗号分隔，最后一个词后面没有出现符号；冒号与逗号为全角形式', 1),
        ('check_english_keywords', '英文摘要页Key words部分：位于英文摘要页下方，整体与上方文本空一行；"Key words"为Times New Roman小四号加粗，冒号之后不加粗；冒号和逗号为半角，逗号后一个空格，最末尾无标点', 1),
        ('check_toc_format', '目录页标题下方的文本：一级目录采用宋体三号加粗；二级目录采用宋体小四号，缩进2字符；三级目录采用宋体小四号并缩进4字符', 5),
        ('check_heading_sequence', '目录页之后的所有页面中的标题：标题排序按"第一章/（一）/1、/1."依次排序', 5),
        ('check_header', '页眉从摘要页开始至文章尾页，页眉文字为黑体、五号、居中，页眉内容与当前页面所属章节名称一致', 3),
        ('check_page_number', '页脚页码汉字使用宋体、黑体小五号，数字使用 Times New Roman 、小五号、居中', 1),
        ('check_page_number_format', '摘要页和目录页页码使用大写罗马数字，从第一章开始，页码重新排序设置，格式用例如"第1页"', 3),
        ('check_figure_caption', '图名均位于图片正下方居中，中文采用黑体五号，英文采用 Times New Roman 五号，不加粗', 5),
        ('check_table_title', '正文表格均有表题，表题位于表格正上方居中，中文采用黑体五号，英文采用 Times New Roman 五号，不加粗', 1),
        ('check_table_three_line', '正文所有表格为三线表，仅保留顶线、栏目线、底线，无竖线；顶线和底线为1.5磅，栏目线为0.75磅', 3),
        ('check_table_content', '正文所有表格表头加粗，表内文本：中文为宋体五号，英文和数字为Times New Roman五号，行列对齐', 3),
        ('check_code_format', '程序代码段为Times New Roman五号，代码段行距为1.25倍', 1),
    ]

    # 维度二"扣分项"规则：(方法名, 规则描述, 扣分幅度[负值])
    _DIM2_SUB_RULES = [
        ('check_first_page', '第1页格式要求：顶部/底部"澄云信息科技学院"、"毕业设计（论文）"、论文标题、冒号行等', -5),
        ('check_chinese_abstract_format', '中文摘要页标题下方文本为小四号宋体、首行缩进2字符、1.5倍行距、两端对齐', -1),
        ('check_english_abstract_title', '英文摘要页标题"Abstract"采用 Times New Roman 、二号、加粗、居中', -1),
        ('check_english_abstract_content', '英文摘要页正文除"Key words"之外均采用 Times New Roman、小四号、1.5倍行距、首行缩进1字符', -1),
        ('check_toc_title', '目录页"目录"二字采用黑体二号、居中对齐、段后12磅、行距1.35倍', -1),
        ('check_body_text_format', '目录页之后除标题、图名、表题以外的文本：中文宋体小四号，英文/数字 Times New Roman 小四号；首行缩进2字符，段前段后均为0行，行距1.5倍，两端对齐', -3),
        ('check_title_spacing', '目录页后的所有页：标题与标题之间、标题与正文之间均空一行', -3),
        ('check_table_cell', '目录页后的所有页表格的单元格不满足此项：其中的文本只有单独一行', -3),
        ('check_code_indent', '程序代码段：首行缩进2字符，文本之前和文本之后无缩进', -1),
    ]

    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = None
        self.score = 0
        self.checked_points = []
        self.failed_points = []
        self.passed_dimension1 = True
        # 维度一未通过时的原因（对齐"评估脚本接口差异与统一建议.md"§2.2）
        self.dim1_reason = ''
        # 维度二逐项得分记录（命中/未命中都记录）
        self.dim2_items = []
        # 维度二 detail 截取用游标：check_dimension2 在每次 check_* 前后
        # 用它标记 _DEBUG_BUF 的新增区间。
        self._dbg_cursor: int = 0

    def evaluate(self):
        '''主评估流程

        返回结构化字典（对齐"评估脚本接口差异与统一建议.md" §2.2），
        便于批量运行器汇总到 Excel。调试信息通过遮蔽的 print 打到 stderr，
        不污染 stdout 上的主结果。
        '''
        print("=" * 60)
        print("Word文档自动评估报告")
        print("=" * 60)
        print(f"文件名: {os.path.basename(self.file_path)}")
        print("-" * 60)

        # 维度1检查
        print("\n【维度1：可用性与可修改性检查】")
        if not self.check_dimension1():
            self.passed_dimension1 = False
            # 维度一未通过时以 self.failed_points 首条作为原因说明
            self.dim1_reason = self.failed_points[0] if self.failed_points else '维度1检查未通过'
            print("\n❌ 维度1检查未通过，文档判为零分！")
            self.print_result()
            return self._build_report()

        print("\n✓ 维度1检查通过\n")

        # 维度2检查
        print("-" * 60)
        print("【维度2：完成度评分】")
        self.check_dimension2()

        # 输出结果（调试用途，主结果通过 return 返回）
        self.print_result()
        return self._build_report()

    def _build_report(self):
        '''按统一约定 §2.2 构造评估结果字典'''
        # 满分固定按加分项汇总，避免维度一未通过时 max_score 为 0
        max_score = sum(m for _, _, m in self._DIM2_ADD_RULES)
        return {
            'id': '013',
            'file_name': os.path.basename(self.file_path),
            'status': 'ok',
            'error': None,
            'dim1_pass': self.passed_dimension1,
            'dim1_reason': self.dim1_reason,
            'dim2_items': list(self.dim2_items),
            'total_score': self.score if self.passed_dimension1 else 0,
            'max_score': max_score,
        }

    def check_dimension1(self):
        '''检查维度1'''
        # 1. 检查文件格式
        if not self.check_file_format():
            return False

        # 2. 检查文件可正常打开
        if not self.check_file_openable():
            return False

        # 3. 检查正文不是整页图片
        if not self.check_not_full_page_image():
            return False

        # 4. 检查空白页和乱码
        if not self.check_no_blank_pages():
            return False

        return True

    def check_file_format(self):
        '''检查文件格式是否为.docx'''
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext == '.docx':
            print("  ✓ 文件格式: 符合要求（{}）".format(ext))
            return True
        else:
            print("  ✗ 文件格式: 不符合要求，应为.docx")
            self.failed_points.append("文件格式不符合要求")
            return False

    def check_file_openable(self):
        '''检查文件是否可正常打开'''
        try:
            self.doc = Document(self.file_path)
            print("  ✓ 文件状态: 可正常打开")
            return True
        except Exception as e:
            print(f"  ✗ 文件状态: 无法打开 - {str(e)}")
            self.failed_points.append("文件无法正常打开")
            return False

    def check_not_full_page_image(self):
        '''检查正文不是整页图片，各部分可编辑'''
        # 检查文档中的图片数量
        image_count = 0
        paragraph_count = len(self.doc.paragraphs)

        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1

        # 如果图片数接近段落数，可能存在问题
        if paragraph_count < 10 and image_count > paragraph_count * 0.5:
            print("  ✗ 正文内容: 可能是整页图片")
            self.failed_points.append("正文可能为整页图片")
            return False

        # 检查主要部分是否存在文本内容
        has_content = False
        for para in self.doc.paragraphs:
            if para.text.strip():
                has_content = True
                break

        if not has_content:
            print("  ✗ 正文内容: 未找到可编辑文本")
            self.failed_points.append("封面、摘要、目录、正文、参考文献、致谢不可编辑")
            return False

        print("  ✓ 正文内容: 文本可编辑")
        return True

    def check_no_blank_pages(self):
        '''检查是否有连续空白页和乱码'''
        consecutive_blank = 0
        max_consecutive_blank = 0
        garbled_count = 0

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                consecutive_blank += 1
                max_consecutive_blank = max(max_consecutive_blank, consecutive_blank)
            else:
                consecutive_blank = 0
                # 检查乱码（简单的启发式方法）
                if self.is_garbled(text):
                    garbled_count += 1

        # 检查连续空白页（假设一页约20个段落）
        blank_pages = max_consecutive_blank // 20
        if blank_pages >= 2:
            print(f"  ✗ 空白页检查: 发现连续{blank_pages}页以上空白")
            self.failed_points.append("存在连续2页以上空白页")
            return False

        if garbled_count > len(self.doc.paragraphs) * 0.33:
            print(f"  ✗ 乱码检查: 发现超过1/3面积乱码")
            self.failed_points.append("存在超过1/3面积乱码")
            return False

        print("  ✓ 空白页检查: 无连续2页以上空白页")
        print("  ✓ 乱码检查: 未发现大面积乱码")
        return True

    def is_garbled(self, text):
        '''简单的乱码检测'''
        # 检查是否包含大量不可打印字符
        printable_count = sum(1 for c in text if c.isprintable() or c in '\n\r\t')
        if len(text) > 0 and printable_count / len(text) < 0.7:
            return True
        return False

    def check_dimension2(self):
        '''检查维度2：完成度评分

        遵循"评估脚本接口差异与统一建议.md" §2.2：命中项与未命中项都要
        记录到 self.dim2_items 中，便于后续在 Excel 中做逐项矩阵对比。
        通过监测每次 check_* 前后 self.score 的差值来判定"命中"与实际得分。

        detail 字段：默认留空（正常命中/未命中都无需啰嗦解释）。仅当
        本次 check_* 出现"特殊情况"——例如"未找到 XXX"、"跳过检查"
        ——时，才截取相应调试行作为 detail 返回，便于批量运行时排查。
        '''
        # 触发写入 detail 的特殊情况关键字
        _SPECIAL_KWS = ('未找到', '跳过')

        def _pop_detail():
            lines = _DEBUG_BUF[self._dbg_cursor:]
            self._dbg_cursor = len(_DEBUG_BUF)
            picks: list[str] = []
            for ln in lines:
                s = ln.strip()
                # 去掉"  ✓ " / "  ✗ " / "      · " / "  - " 等前缀标记
                for prefix in ('✓ ', '✗ ', '· ', '- '):
                    if s.startswith(prefix):
                        s = s[len(prefix):]
                        break
                if not s:
                    continue
                # 仅保留"特殊情况"行
                if any(k in s for k in _SPECIAL_KWS):
                    picks.append(s)
            return '；'.join(picks)

        # 加分项：max_delta > 0
        for method_name, rule_desc, max_delta in self._DIM2_ADD_RULES:
            before = self.score
            self._dbg_cursor = len(_DEBUG_BUF)
            getattr(self, method_name)()
            delta = self.score - before
            # 仍调用 _pop_detail() 以推进 _dbg_cursor、保持调试缓冲区消费节奏，
            # 但按要求对外一律返回空字符串（不影响评分与其他字段）
            _pop_detail()
            self.dim2_items.append({
                'rule': f'+{max_delta}：{rule_desc}',
                'max_delta': max_delta,
                'delta': delta,
                'hit': delta == max_delta,
                'detail': '',
            })

        # 扣分项：max_delta < 0；命中表示实际触发扣分条件。
        for method_name, rule_desc, max_delta in self._DIM2_SUB_RULES:
            before = self.score
            self._dbg_cursor = len(_DEBUG_BUF)
            getattr(self, method_name)()
            delta = self.score - before  # 未触发时为 0；触发时为负值
            # 同上：消费调试缓冲区但丢弃 detail 文本
            _pop_detail()
            self.dim2_items.append({
                'rule': f'{max_delta}：{rule_desc}',
                'max_delta': max_delta,
                'delta': delta,
                'hit': delta == max_delta,
                'detail': '',
            })

    def get_font_info(self, run):
        '''获取run的字体信息'''
        font = run.font
        font_name = font.name
        size = font.size
        bold = font.bold
        color = font.color.rgb if font.color.rgb else None
        return {
            'name': font_name,
            'size': size.pt if size else None,
            'bold': bold,
            'color': str(color) if color else None
        }

    # ============================================================
    # 有效属性解析（考虑 run -> 段落 -> 段落样式 -> 文档默认 的继承链，
    # 以保证结果与 Office/WPS 在办公软件中的实际渲染一致）
    # ============================================================
    def _iter_style_chain(self, para):
        '''按继承顺序返回段落所用样式链'''
        style = para.style
        visited = set()
        while style is not None:
            sid = getattr(style, 'style_id', None) or id(style)
            if sid in visited:
                break
            visited.add(sid)
            yield style
            try:
                style = style.base_style
            except Exception:
                break

    def _is_code_paragraph(self, para):
        '''识别程序代码段：综合样式、等宽字体、上下文缩进和代码内容特征。'''
        text = para.text.rstrip('\r\n')
        stripped = text.strip()
        if not stripped:
            return False

        style_name = (para.style.name if para.style else '') or ''
        style_name_lower = style_name.lower()
        if re.search(r'code|代码|程序|源代码|listing', style_name_lower):
            return True

        monospace_fonts = {
            'consolas', 'courier new', 'courier', 'menlo', 'monaco',
            'lucida console', 'source code pro', 'cascadia code', '等线'
        }
        mono_runs = 0
        content_runs = 0
        for run in para.runs:
            rt = run.text
            if not rt or not rt.strip():
                continue
            content_runs += 1
            fonts = [
                self._effective_font(run, para, 'ascii'),
                self._effective_font(run, para, 'hAnsi'),
                self._effective_font(run, para, 'eastAsia'),
            ]
            if any(f and f.lower() in monospace_fonts for f in fonts):
                mono_runs += 1
        if content_runs and mono_runs == content_runs:
            return True

        code_pat = re.compile(
            r'^\s*(?:def|class|import|from|return|if|elif|else|for|while|try|except|finally|with)\b'
            r'|^\s*[A-Za-z_]\w*\s*=\s*[A-Za-z_][\w\.]*\s*\('
            r'|\.(?:withColumn|groupBy|createOrReplaceTempView|agg|filter|select|orderBy|show|collect)\s*\('
            r'|^\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|WITH)\b'
            r'|^\s*(?:public|private|protected|static|void|int|String|var|let|const)\b'
            r'|[{};]|//|#\s*\w+'
        )
        if code_pat.search(stripped):
            return True

        # 代码块中的续行常只有缩进、点号链式调用、括号/逗号结尾等弱特征；
        # 只有在具有明显代码排版上下文时才采纳，避免误判普通英文正文。
        leading_spaces = len(text) - len(text.lstrip(' '))
        has_weak_code_shape = bool(re.search(r'^\s{2,}\S|^\s*\.\w+\s*\(|[,({\[]\s*$', text))
        if leading_spaces >= 4 and has_weak_code_shape:
            return True

        return False

    def _doc_defaults_rPr(self):
        '''获取 docDefaults 中的 rPr 元素'''
        try:
            styles_elem = self.doc.styles.element
            docDefaults = styles_elem.find(qn('w:docDefaults'))
            if docDefaults is not None:
                rPrDefault = docDefaults.find(qn('w:rPrDefault'))
                if rPrDefault is not None:
                    return rPrDefault.find(qn('w:rPr'))
        except Exception:
            pass
        return None

    def _effective_font(self, run, para, script='eastAsia'):
        '''获取 run 在指定脚本(eastAsia/ascii/hAnsi)下的有效字体名'''
        # 1) run 直接属性
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                v = rFonts.get(qn('w:' + script))
                if v:
                    return v
        # 2) 段落 pPr 里的 rPr（段落标记的 run 属性）
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            p_rPr = pPr.find(qn('w:rPr'))
            if p_rPr is not None:
                rFonts = p_rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    v = rFonts.get(qn('w:' + script))
                    if v:
                        return v
        # 3) 段落样式链
        for style in self._iter_style_chain(para):
            s_rPr = style.element.find(qn('w:rPr'))
            if s_rPr is not None:
                rFonts = s_rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    v = rFonts.get(qn('w:' + script))
                    if v:
                        return v
        # 4) 文档默认
        dd_rPr = self._doc_defaults_rPr()
        if dd_rPr is not None:
            rFonts = dd_rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                v = rFonts.get(qn('w:' + script))
                if v:
                    return v
        return None

    def _effective_size(self, run, para):
        '''获取 run 的有效字号(pt)'''
        def read_sz(rPr):
            if rPr is None:
                return None
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                v = sz.get(qn('w:val'))
                if v:
                    try:
                        return float(v) / 2.0
                    except ValueError:
                        return None
            return None

        v = read_sz(run._element.find(qn('w:rPr')))
        if v is not None:
            return v
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            v = read_sz(pPr.find(qn('w:rPr')))
            if v is not None:
                return v
        for style in self._iter_style_chain(para):
            v = read_sz(style.element.find(qn('w:rPr')))
            if v is not None:
                return v
        return read_sz(self._doc_defaults_rPr())

    def _effective_bold(self, run, para):
        '''获取 run 的有效加粗状态'''
        def read_bold(rPr):
            if rPr is None:
                return None
            b = rPr.find(qn('w:b'))
            if b is None:
                return None
            v = b.get(qn('w:val'))
            if v is None:
                return True
            return v in ('1', 'true', 'True', 'on')

        v = read_bold(run._element.find(qn('w:rPr')))
        if v is not None:
            return v
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            v = read_bold(pPr.find(qn('w:rPr')))
            if v is not None:
                return v
        for style in self._iter_style_chain(para):
            v = read_bold(style.element.find(qn('w:rPr')))
            if v is not None:
                return v
        v = read_bold(self._doc_defaults_rPr())
        return v if v is not None else False

    def _effective_alignment(self, para):
        '''获取段落的有效对齐方式（返回 OOXML 的 val 字符串）'''
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                return jc.get(qn('w:val'))
        for style in self._iter_style_chain(para):
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                jc = s_pPr.find(qn('w:jc'))
                if jc is not None:
                    return jc.get(qn('w:val'))
        return None

    def check_abstract_title(self):
        '''检查中文摘要标题
        细则：+1：中文摘要标题为"摘 要"或"摘要"；宋体二号、加粗、居中；与下方文本空一行
        逐点核对，不额外附加细则未要求的约束。
        '''
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            # 细则点1：标题文本必须为"摘 要"或"摘要"
            # （Word 中"摘 要"常因排版渲染成多个空格，等价视为"摘 要"）
            if not re.match(r'^摘\s*要$', text):
                continue

            # 取第一个非空 run 作为字体/字号/加粗的检查对象
            target_run = None
            for r in para.runs:
                if r.text.strip():
                    target_run = r
                    break
            if target_run is None:
                continue

            # 细则点2：宋体（中文标题，检查 East Asian 字体，Office 依此渲染中文）
            ea_font = self._effective_font(target_run, para, 'eastAsia')
            is_songti = ea_font is not None and ('宋体' in ea_font or ea_font == 'SimSun')

            # 细则点3：二号（Office 中二号 = 精确 22pt）
            size_pt = self._effective_size(target_run, para)
            is_erhao = size_pt is not None and abs(size_pt - 22.0) < 0.01

            # 细则点4：加粗
            is_bold = self._effective_bold(target_run, para)

            # 细则点5：居中
            align_val = self._effective_alignment(para)
            is_center = align_val == 'center'

            # 细则点6：与下方文本空一行（紧邻的下一个段落为空段落）
            has_blank_line = False
            if i + 1 < len(self.doc.paragraphs):
                next_para = self.doc.paragraphs[i + 1]
                if not next_para.text.strip():
                    has_blank_line = True

            if is_songti and is_erhao and is_bold and is_center and has_blank_line:
                self.score += 1
                msg = '+1：中文摘要标题为"摘 要"或"摘要"；宋体二号、加粗、居中；与下方文本空一行'
                self.checked_points.append(msg)
                print('  ✓ ' + msg)
            else:
                reasons = []
                if not is_songti:
                    reasons.append(f'字体非宋体(当前:{ea_font})')
                if not is_erhao:
                    reasons.append(f'字号非二号22pt(当前:{size_pt}pt)')
                if not is_bold:
                    reasons.append('未加粗')
                if not is_center:
                    reasons.append(f'未居中(当前:{align_val})')
                if not has_blank_line:
                    reasons.append('下方未空一行')
                print('  ✗ 中文摘要标题不符合要求：' + '; '.join(reasons))
            return

        print('  ✗ 未找到中文摘要标题（应为"摘 要"或"摘要"）')

    def check_keywords(self):
        '''检查中文关键词格式
        细则：+1：摘要正文下"关键词"部分：四个关键词用中文逗号分隔，
        最后一个词后面没有出现符号；冒号与逗号为全角形式

        逐点核对，不额外附加细则未要求的约束。所有判断均基于办公软件
        （Word/WPS）中用户实际可见的字符，因此在办公软件上一致有效。
        '''
        import unicodedata

        # 定位"关键词"段落：以"关键词"开头，紧跟冒号（全角或半角均先接受，
        # 后续再判定是否为全角）。文档中"关键词"部分位于摘要正文之下，
        # 只取第一个匹配段落。
        kw_para = None
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^关键词\s*[：:]', text):
                kw_para = para
                break

        if kw_para is None:
            print('  ✗ 未找到"关键词"段落')
            return

        text = kw_para.text.strip()
        m = re.match(r'^关键词(\s*)([：:])(.*)$', text)
        if not m:
            print('  ✗ "关键词"段落格式异常')
            return

        colon = m.group(2)
        # 冒号之后的原始内容（保留原样，末尾空白单独处理）
        kw_content = m.group(3)

        # 细则点1：冒号为全角形式（"："）
        is_full_colon = (colon == '：')

        # 细则点2：逗号为全角形式——冒号之后不允许出现半角","作为分隔
        is_full_comma = (',' not in kw_content)

        # 细则点3：四个关键词——以全角逗号"，"分隔恰好得到 4 个非空词
        parts = kw_content.split('，')
        keywords = [p.strip() for p in parts]
        has_four = (len(keywords) == 4) and all(bool(k) for k in keywords)

        # 细则点4：最后一个词后面没有出现符号
        # 整段冒号后内容去除右侧空白后，末字符不能属于 Unicode Punctuation 类
        trimmed = kw_content.rstrip()
        last_char = trimmed[-1] if trimmed else ''
        no_end_symbol = bool(last_char) and not unicodedata.category(last_char).startswith('P')

        if is_full_colon and is_full_comma and has_four and no_end_symbol:
            self.score += 1
            msg = '+1：摘要正文下"关键词"部分：四个关键词用中文逗号分隔，最后一个词后面没有出现符号；冒号与逗号为全角形式'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            reasons = []
            if not is_full_colon:
                reasons.append(f'冒号非全角(当前:"{colon}")')
            if not is_full_comma:
                reasons.append('存在半角逗号","')
            if not has_four:
                reasons.append(f'关键词数量不为4(当前:{len(keywords)}个:{keywords})')
            if not no_end_symbol:
                reasons.append(f'末尾出现符号"{last_char}"')
            print('  ✗ 关键词格式不符合要求：' + '; '.join(reasons))

    def check_english_keywords(self):
        '''检查英文关键词格式
        细则：+1：英文摘要页Key words部分：
          (1) 位于英文摘要页下方；
          (2) 整体与上方文本空一行的位置；
          (3) "Key words"为 Times New Roman；
          (4) "Key words"为小四号；
          (5) "Key words"加粗；
          (6) 冒号之后的文本不加粗；
          (7) 冒号为半角符号；
          (8) 逗号为半角符号；
          (9) 逗号后有一个空格；
          (10) 最末尾无标点。

        逐点核对，不额外附加细则未要求的约束；均基于办公软件
        (Word/WPS) 用户实际渲染结果判定。
        '''
        import unicodedata

        # 先定位 "Abstract" 标题所在段，以便判断"英文摘要页下方"和
        # "与上方文本空一行"这两点。
        abstract_idx = -1
        for i, p in enumerate(self.doc.paragraphs):
            if p.text.strip() == 'Abstract':
                abstract_idx = i
                break

        # 找到"Key words"段落（英文摘要页内、位于 Abstract 之后）
        kw_idx = -1
        kw_para = None
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if re.match(r'^Key\s*words\b', text, flags=re.IGNORECASE):
                if abstract_idx < 0 or i > abstract_idx:
                    kw_idx = i
                    kw_para = para
                    break

        if kw_para is None:
            print('  ✗ 未找到"Key words"段落')
            return

        text = kw_para.text  # 保留原始空白，用于精确判断
        stripped = text.strip()

        # 细则点(1)：位于英文摘要页下方
        # 采用可在办公软件中稳定判定的规则：Key words 段位于 Abstract 段之后，
        # 且两者之间没有出现下一节的标题（如"目 录"/"第1章"等）。
        is_below_abstract = False
        if abstract_idx >= 0 and kw_idx > abstract_idx:
            is_below_abstract = True
            for j in range(abstract_idx + 1, kw_idx):
                mid = self.doc.paragraphs[j].text.strip()
                if re.match(r'^目\s*录$', mid) or re.match(r'^第[一二三四五六七八九十\d]+章', mid):
                    is_below_abstract = False
                    break

        # 细则点(2)：整体与上方文本空一行——上一段为空段落
        has_blank_above = False
        if kw_idx > 0:
            prev = self.doc.paragraphs[kw_idx - 1]
            if not prev.text.strip():
                has_blank_above = True

        # 定位"Key words"文字所在的 run，用于点(3)(4)(5)；
        # 以及冒号之后的第一个非空 run，用于点(6)。
        key_words_run = None
        after_colon_runs = []
        seen_colon = False
        for run in kw_para.runs:
            if not run.text:
                continue
            if not seen_colon:
                # "Key words" 字样可能与冒号同处一个 run
                if re.search(r'Key\s*words', run.text, flags=re.IGNORECASE) and key_words_run is None:
                    key_words_run = run
                if ':' in run.text or '：' in run.text:
                    seen_colon = True
                    # 如果冒号后本 run 还有其它字符，也纳入"冒号之后文本"
                    tail = run.text.split(':', 1)[-1] if ':' in run.text else run.text.split('：', 1)[-1]
                    if tail.strip():
                        after_colon_runs.append((run, tail))
            else:
                if run.text.strip():
                    after_colon_runs.append((run, run.text))

        # 细则点(3)：Key words 为 Times New Roman —— 取 ascii 字体，
        # 沿继承链解析，与 Office 渲染西文字体的取值一致。
        is_tnr = False
        if key_words_run is not None:
            ascii_font = self._effective_font(key_words_run, kw_para, 'ascii')
            is_tnr = ascii_font == 'Times New Roman'

        # 细则点(4)：小四号 —— Office 中"小四"精确为 12pt
        is_size_xs4 = False
        if key_words_run is not None:
            size_pt = self._effective_size(key_words_run, kw_para)
            is_size_xs4 = size_pt is not None and abs(size_pt - 12.0) < 0.01

        # 细则点(5)：Key words 加粗
        is_bold = False
        if key_words_run is not None:
            is_bold = bool(self._effective_bold(key_words_run, kw_para))

        # 细则点(6)：冒号之后文本不加粗
        after_colon_not_bold = True
        if not after_colon_runs:
            after_colon_not_bold = False  # 冒号后没有文本，视为不满足
        else:
            for r, _ in after_colon_runs:
                if self._effective_bold(r, kw_para):
                    after_colon_not_bold = False
                    break

        # 细则点(7)：冒号为半角
        # 以"Key words"之后第一个冒号字符判定
        m = re.search(r'Key\s*words(\s*)([：:])', stripped, flags=re.IGNORECASE)
        colon_char = m.group(2) if m else ''
        is_half_colon = (colon_char == ':')

        # 取冒号之后的内容，用于点(8)(9)(10)
        content_after_colon = ''
        if m:
            content_after_colon = stripped[m.end():]

        # 细则点(8)：逗号为半角 —— 冒号之后不允许出现全角"，"
        is_half_comma = ('，' not in content_after_colon)

        # 细则点(9)：每个半角逗号后紧跟一个空格
        # 遍历冒号之后所有半角逗号，要求其后为单个 ASCII 空格（且非连续空格）
        space_after_comma = True
        commas = [i for i, ch in enumerate(content_after_colon) if ch == ',']
        if not commas:
            space_after_comma = False
        else:
            for pos in commas:
                # 逗号不应位于末尾
                if pos + 1 >= len(content_after_colon):
                    space_after_comma = False
                    break
                if content_after_colon[pos + 1] != ' ':
                    space_after_comma = False
                    break
                # 恰好一个空格：pos+2 处不能再是空白
                if pos + 2 < len(content_after_colon) and content_after_colon[pos + 2] == ' ':
                    space_after_comma = False
                    break

        # 细则点(10)：最末尾无标点 —— 末字符不属于 Unicode Punctuation 大类
        trimmed_tail = content_after_colon.rstrip()
        last_char = trimmed_tail[-1] if trimmed_tail else ''
        no_end_punct = bool(last_char) and not unicodedata.category(last_char).startswith('P')

        all_ok = (is_below_abstract and has_blank_above and is_tnr and is_size_xs4
                  and is_bold and after_colon_not_bold and is_half_colon
                  and is_half_comma and space_after_comma and no_end_punct)

        if all_ok:
            self.score += 1
            msg = '+1：英文摘要页Key words部分：位于英文摘要页下方，整体与上方文本空一行的位置；"Key words"为Times New Roman小四号加粗，冒号之后的文本不加粗；冒号和逗号为半角符号，逗号后有一个空格，最末尾无标点。'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            reasons = []
            if not is_below_abstract:
                reasons.append('不在英文摘要页下方')
            if not has_blank_above:
                reasons.append('上方未空一行')
            if not is_tnr:
                reasons.append('"Key words"字体非Times New Roman')
            if not is_size_xs4:
                reasons.append('"Key words"字号非小四(12pt)')
            if not is_bold:
                reasons.append('"Key words"未加粗')
            if not after_colon_not_bold:
                reasons.append('冒号之后文本被加粗')
            if not is_half_colon:
                reasons.append(f'冒号非半角(当前:"{colon_char}")')
            if not is_half_comma:
                reasons.append('存在全角逗号"，"')
            if not space_after_comma:
                reasons.append('逗号后未跟单个空格')
            if not no_end_punct:
                reasons.append(f'末尾出现标点"{last_char}"')
            print('  ✗ 英文关键词格式不符合要求：' + '; '.join(reasons))

    def check_chinese_abstract_content(self):
        '''检查中文摘要页内容字体
        细则：+1：中文摘要页内容除"摘要"、"关键词"、"研究方向"，
              其余字体格式为宋体四号

        逐点核对：
          (1) 检查范围：中文摘要页内容（"摘 要"标题至该页结束）；
          (2) 排除项："摘要"标题段、"关键词"段、"研究方向"段；
          (3) 字体：宋体（中文取 w:eastAsia，英文/数字取 w:ascii）；
          (4) 字号：四号（Office 中"四号"精确为 14pt）。

        全部基于办公软件 (Word/WPS) 的样式继承链取值，确保与实际渲染一致。
        '''
        # 细则点(1)：定位"摘 要"标题段与中文摘要页结束位置
        start_idx = -1
        for i, p in enumerate(self.doc.paragraphs):
            if re.match(r'^摘\s*要$', p.text.strip()):
                start_idx = i
                break
        if start_idx < 0:
            print('  ✗ 未找到中文摘要页')
            return

        # 结束位置：遇到英文摘要页起点（Abstract 段或紧邻其上的英文标题段）、
        # 或"目 录"、或"第X章"即停止
        end_idx = len(self.doc.paragraphs)
        for j in range(start_idx + 1, len(self.doc.paragraphs)):
            t = self.doc.paragraphs[j].text.strip()
            if t == 'Abstract':
                end_idx = j
                break
            if re.match(r'^目\s*录$', t):
                end_idx = j
                break
            if re.match(r'^第[一二三四五六七八九十\d]+章', t):
                end_idx = j
                break
            # 英文论文题目段（全英文长句、字号为二号）也视作英文摘要页起点
            if t and re.match(r'^[A-Za-z0-9 ,.\-:;\'()]+$', t) and len(t) > 20:
                end_idx = j
                break

        # 细则点(2)：排除段判定
        def is_excluded(text):
            if re.match(r'^摘\s*要$', text):
                return True
            if re.match(r'^关键词\s*[：:]', text):
                return True
            if re.match(r'^研究方向\s*[：:]', text):
                return True
            return False

        errors = []
        checked = 0
        for k in range(start_idx, end_idx):
            para = self.doc.paragraphs[k]
            text = para.text.strip()
            if not text:
                continue
            if is_excluded(text):
                continue

            checked += 1
            # 遍历段内每个有实体文本的 run
            for run in para.runs:
                run_text = run.text
                if not run_text or not run_text.strip():
                    continue

                # 细则点(4)：字号 = 四号 = 14pt
                size_pt = self._effective_size(run, para)
                if size_pt is None or abs(size_pt - 14.0) >= 0.01:
                    errors.append(f'"{text[:20]}" 字号非四号14pt(当前:{size_pt}pt)')
                    break  # 该段只报一次

                # 细则点(3)：字体 = 宋体
                # 依 run 内字符类型分别检查东亚/西文字体的有效值
                has_chinese = bool(re.search(r'[一-鿿]', run_text))
                has_ascii = bool(re.search(r'[A-Za-z0-9]', run_text))

                ok = True
                if has_chinese:
                    ea = self._effective_font(run, para, 'eastAsia')
                    if not (ea and ('宋体' in ea or ea == 'SimSun')):
                        errors.append(f'"{text[:20]}" 中文字体非宋体(当前:{ea})')
                        ok = False
                if ok and has_ascii:
                    ascii_font = self._effective_font(run, para, 'ascii')
                    # 西文字体在办公软件中的"宋体"通常映射为 SimSun/宋体
                    if not (ascii_font and ('宋体' in ascii_font or ascii_font == 'SimSun')):
                        errors.append(f'"{text[:20]}" 西文字体非宋体(当前:{ascii_font})')
                        ok = False
                if not ok:
                    break

        if checked == 0:
            print('  ✗ 中文摘要页未找到需检查的文本段')
            return

        if not errors:
            self.score += 1
            msg = '+1：中文摘要页内容除"摘要"、"关键词"、"研究方向"，其余字体格式为宋体四号'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 中文摘要页字体格式不符合要求：')
            for e in errors[:8]:
                print('      · ' + e)
            if len(errors) > 8:
                print(f'      · ...共{len(errors)}项')

    def check_toc_format(self):
        '''检查目录格式
        细则：+5：目录页标题下方的文本：
          (1) 一级目录采用宋体；
          (2) 一级目录采用三号；
          (3) 一级目录加粗；
          (4) 二级目录采用宋体；
          (5) 二级目录采用小四号；
          (6) 二级目录缩进2字符；
          (7) 三级目录采用宋体；
          (8) 三级目录采用小四号；
          (9) 三级目录缩进4字符。

        逐点核对，不额外附加细则未要求的约束；所有判断均基于办公软件
        (Word/WPS) 用户实际渲染结果。
        '''
        # 1) 定位"目 录"标题所在段落（允许"目录/目 录/目  录"）
        toc_idx = -1
        for i, para in enumerate(self.doc.paragraphs):
            if re.match(r'^目\s*录$', para.text.strip()):
                toc_idx = i
                break

        if toc_idx < 0:
            print('  ✗ 未找到"目录"标题段落')
            return

        # 2) 采集目录条目：从标题下一段开始，直到遇到"正文起点"
        #    正文起点判定：不含制表符的"第X章"或"1.1"等编号段落
        #    （目录条目里"第X章 ... \t页码"含 \t，正文标题不含）
        toc_items = []
        for j in range(toc_idx + 1, len(self.doc.paragraphs)):
            p = self.doc.paragraphs[j]
            raw = p.text
            stripped = raw.strip()
            if not stripped:
                continue
            # 正文起点：不含制表符且形如"第X章 ..."
            if '\t' not in raw and re.match(r'^第[一二三四五六七八九十\d]+章', stripped):
                break
            toc_items.append(p)

        if not toc_items:
            print('  ✗ 未找到目录条目')
            return

        # 3) 按编号形态划分级别（细则只区分一级/二级/三级）
        def classify(text):
            # 三级：形如 1.1.1
            if re.match(r'^\d+\.\d+\.\d+', text):
                return 3
            # 二级：形如 1.1
            if re.match(r'^\d+\.\d+', text):
                return 2
            # 其余（如"第X章 ..."/"摘 要"/"Abstract"/"参考文献"/"致谢"）视为一级
            return 1

        # 4) 逐条按级别核对
        def approx(a, b, eps=0.5):
            return a is not None and abs(a - b) < eps

        errors = []
        for p in toc_items:
            text = p.text.strip()
            level = classify(text)

            # 取首个非空 run 作为字体/字号/加粗检查对象
            target_run = None
            for r in p.runs:
                if r.text.strip():
                    target_run = r
                    break
            if target_run is None:
                errors.append(f'"{text[:20]}" 无可检查文本')
                continue

            ea_font = self._effective_font(target_run, p, 'eastAsia')
            is_songti = ea_font is not None and ('宋体' in ea_font or ea_font == 'SimSun')

            size_pt = self._effective_size(target_run, p)
            is_bold = bool(self._effective_bold(target_run, p))

            # 左缩进（w:left / w:leftChars 由 python-docx 归一到 pt）
            left_indent = p.paragraph_format.left_indent
            left_pt = left_indent.pt if left_indent is not None else 0.0

            if level == 1:
                # 三号 = 16pt；宋体；加粗
                if not is_songti:
                    errors.append(f'一级"{text[:20]}"字体非宋体(当前:{ea_font})')
                if not approx(size_pt, 16.0, 0.01):
                    errors.append(f'一级"{text[:20]}"字号非三号16pt(当前:{size_pt}pt)')
                if not is_bold:
                    errors.append(f'一级"{text[:20]}"未加粗')

            elif level == 2:
                # 小四 = 12pt；宋体；缩进2字符 = 2×12 = 24pt
                if not is_songti:
                    errors.append(f'二级"{text[:20]}"字体非宋体(当前:{ea_font})')
                if not approx(size_pt, 12.0, 0.01):
                    errors.append(f'二级"{text[:20]}"字号非小四12pt(当前:{size_pt}pt)')
                if not approx(left_pt, 24.0, 0.5):
                    errors.append(f'二级"{text[:20]}"缩进非2字符24pt(当前:{left_pt}pt)')

            else:  # level == 3
                # 小四 = 12pt；宋体；缩进4字符 = 4×12 = 48pt
                if not is_songti:
                    errors.append(f'三级"{text[:20]}"字体非宋体(当前:{ea_font})')
                if not approx(size_pt, 12.0, 0.01):
                    errors.append(f'三级"{text[:20]}"字号非小四12pt(当前:{size_pt}pt)')
                if not approx(left_pt, 48.0, 0.5):
                    errors.append(f'三级"{text[:20]}"缩进非4字符48pt(当前:{left_pt}pt)')

        if not errors:
            self.score += 5
            msg = '+5：目录页标题下方的文本：一级目录采用宋体三号、加粗；二级目录采用宋体小四号，缩进2字符；三级目录采用宋体小四号并缩进4字符'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 目录格式不符合要求：')
            for e in errors[:12]:
                print('      · ' + e)
            if len(errors) > 12:
                print(f'      · ...共{len(errors)}项')

    def check_heading_sequence(self):
        '''检查标题排序
        细则：+5：目录页之后的所有页面中的标题：
          (1) 检查范围：目录页之后的所有页面中的"标题"；
          (2) 一级标题序号类型为"第一章"（形如"第X章"）；
          (3) 二级标题序号类型为"（一）"（形如全角括号内序数）；
          (4) 三级标题序号类型为"1、"（形如"数字+顿号"）；
          (5) 四级标题序号类型为"1."（形如"数字+点号"，非"1.1"子编号）；
          (6) 标题排序按"第一章/（一）/1、/1."依次排序——不仅不跳级，
              还要求同级序号在各自的父级范围内从1开始、依次连续递增
              （即"第一章"必须从"一"开始且逐章+1；"（一）"在其所属
              一级标题下必须从"一"开始且逐条+1；三级"1、"、四级"1."
              同理在各自父级范围内从1开始连续递增）。

        逐点核对，不额外附加细则未要求的约束；"标题"以办公软件
        （Word/WPS）的判定为准——即段落样式为 Heading 1~4/标题 1~4，
        或段落 pPr 中显式指定了 w:outlineLvl。
        '''
        # 中文数字（含"十/百/千/零"组合）→ 阿拉伯数字，用于校验章节序号连续性
        def cn_num_to_int(s):
            digit_map = {'零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
                         '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}
            if s.isdigit():
                return int(s)
            if not s:
                return None
            # 仅处理十进制以内常见组合（十/一十几/几十/几十几/百及以上不强求）
            total = 0
            unit_map = {'十': 10, '百': 100, '千': 1000}
            section = 0
            num = None
            for ch in s:
                if ch in digit_map:
                    num = digit_map[ch]
                elif ch in unit_map:
                    unit = unit_map[ch]
                    section += (num if num is not None else 1) * unit
                    num = None
                else:
                    return None
            if num is not None:
                section += num
            total += section
            return total if total > 0 or s == '零' else None

        # 细则点(1)：只检查目录页之后的段落
        toc_idx = -1
        for i, p in enumerate(self.doc.paragraphs):
            if re.match(r'^目\s*录$', p.text.strip()):
                toc_idx = i
                break
        if toc_idx < 0:
            print('  ✗ 未找到"目录"段落，无法确定检查范围')
            return

        # 判定段落的大纲级别（Office 里的"标题"就是这些）
        def outline_level(para):
            sname = (para.style.name if para.style else '') or ''
            m = re.match(r'^(?:Heading|标题)\s*(\d+)$', sname)
            if m:
                n = int(m.group(1))
                if 1 <= n <= 4:
                    return n
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                ol = pPr.find(qn('w:outlineLvl'))
                if ol is not None:
                    v = ol.get(qn('w:val'))
                    if v is not None:
                        try:
                            n = int(v) + 1  # OOXML 从 0 开始计
                            if 1 <= n <= 4:
                                return n
                        except ValueError:
                            pass
            return 0

        # 采集目录之后的标题段落及其级别
        headings = []
        for i, p in enumerate(self.doc.paragraphs):
            if i <= toc_idx:
                continue
            lvl = outline_level(p)
            if 1 <= lvl <= 4:
                headings.append((lvl, p.text.strip()))

        if not headings:
            print('  ✗ 目录之后未找到任何标题段落（Heading 1~4 / 标题 1~4）')
            return

        # 细则点(2)~(5)：各级编号"类型"正则
        # 序数允许中文数字或阿拉伯数字（细则以"第一章 / （一）"举例，指的是"类型"）。
        patterns = {
            1: re.compile(r'^第[一二三四五六七八九十百千零\d]+章'),   # 第X章
            2: re.compile(r'^（[一二三四五六七八九十百千零\d]+）'),   # （X），全角括号
            3: re.compile(r'^\d+、'),                                   # 1、（数字+顿号）
            4: re.compile(r'^\d+\.(?!\d)'),                             # 1.（数字+点号，且非1.1）
        }
        type_desc = {1: '"第X章"', 2: '"（X）"', 3: '"X、"', 4: '"X."'}

        errors = []
        # 逐条核对：每个标题的编号形态必须与其大纲级别对应
        for lvl, text in headings:
            if not patterns[lvl].match(text):
                errors.append(f'{lvl}级标题"{text[:24]}"编号类型不符（应为{type_desc[lvl]}）')

        # 细则点(6)前半：排序按 1→2→3→4 依次递进，不允许跳级
        # （同级重复、回到更高级、下一级紧跟均允许；跨级下降不允许）
        last = 0
        for lvl, text in headings:
            if lvl > last + 1:
                prev_desc = type_desc.get(last, '<起始>')
                errors.append(f'标题跳级：从{last}级({prev_desc})直接跳到{lvl}级("{text[:24]}")')
            last = lvl

        # 提取各级标题的序号数值，用于连续性/起始值校验
        extractors = {
            1: re.compile(r'^第([一二三四五六七八九十百千零\d]+)章'),
            2: re.compile(r'^（([一二三四五六七八九十百千零\d]+)）'),
            3: re.compile(r'^(\d+)、'),
            4: re.compile(r'^(\d+)\.(?!\d)'),
        }

        def extract_num(lvl, text):
            m = extractors[lvl].match(text)
            if not m:
                return None
            raw = m.group(1)
            return int(raw) if raw.isdigit() else cn_num_to_int(raw)

        # 细则点(6)后半：同级序号在各自父级范围内必须从1开始、依次连续递增。
        # 用"当前各级序号路径"模拟层级计数器：进入更高父级或新的父级序号时，
        # 该级及以下所有计数器清零重新从1开始。
        expected = {1: 1, 2: 1, 3: 1, 4: 1}
        prev_lvl = 0
        for lvl, text in headings:
            if not patterns[lvl].match(text):
                # 编号类型本身已不符，序号连续性无法可靠判断，跳过避免重复报错
                prev_lvl = lvl
                continue

            num = extract_num(lvl, text)
            if num is None:
                errors.append(f'{lvl}级标题"{text[:24]}"序号无法解析')
                prev_lvl = lvl
                continue

            # 进入比上一个更高的级别（数字更大）时，该级序号计数器重新从1开始
            if lvl > prev_lvl:
                expected[lvl] = 1

            if num != expected[lvl]:
                errors.append(
                    f'{lvl}级标题"{text[:24]}"序号不连续（应为{expected[lvl]}，实际为{num}）'
                )
                # 以实际值为基准继续校验后续序号，避免同一处偏差重复报错
                expected[lvl] = num

            expected[lvl] += 1
            # 该级序号确认后，重置所有更低级别（数字更大）的计数器，
            # 使其下一次出现时从1开始
            for lower in range(lvl + 1, 5):
                expected[lower] = 1
            prev_lvl = lvl


        if not errors:
            self.score += 5
            msg = '+5：目录页之后的所有页面中的标题：标题排序按此规则依次排序"第一章/（一）/1、/1."，一级标题序号类型为第一章、二级标题序号类型为（一）、三级标题序号类型为1、四级标题序号类型为1.'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 标题排序不符合要求：')
            for e in errors[:12]:
                print('      · ' + e)
            if len(errors) > 12:
                print(f'      · ...共{len(errors)}项')

    def check_header(self):
        '''检查页眉
        细则：+3：页眉从摘要页开始至文章尾页，页眉文字为黑体、五号、居中，
              页眉内容与当前页面所属章节名称一致

        逐点核对：
          (1) 检查范围：从摘要页所在节开始，直到文档最后一节（文章尾页）；
          (2) 该范围内每一节的有效页眉（若本节 header 链接到前节，需回溯
              取最近一个未链接的祖先页眉，即办公软件实际渲染使用的内容，
              而不是直接跳过该节）；
          (3) 页眉文字为黑体；
          (4) 五号（约10.5pt）；
          (5) 居中；
          (6) 页眉内容与"当前页面所属章节名称"一致——按 OOXML 节归属关系，
              建立"节→该节内第一个一级标题名称"的映射，而不是全文任意
              一级标题名称匹配即通过。
        '''
        # OOXML 节归属规则：段落 pPr/sectPr 存在时，本段是该节的末段；
        # 之后的段落归下一节；文档最后一节由 body 末尾的 sectPr 定义。
        para_section = []
        sec_idx = 0
        for para in self.doc.paragraphs:
            para_section.append(sec_idx)
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                sec_idx += 1

        def _section_header(idx):
            '''取第 idx 节的有效页眉：若链接到前节，回溯最近未链接的祖先页眉。'''
            while idx >= 0:
                header = self.doc.sections[idx].header
                if header is not None and not header.is_linked_to_previous:
                    return header
                idx -= 1
            # 一直链接到文档开头：以第0节页眉为准（办公软件此时也用它渲染）
            return self.doc.sections[0].header

        def _chapter_name(text):
            m = re.match(r'^第[一二三四五六七八九十百千零\d]+章\s*(.+)', text)
            if m:
                return m.group(1).strip()
            return None

        # 建立"节 → 当前章节名称"映射：每节取该节内出现的第一个一级标题
        # （不含目录条目——目录条目含制表符或本身在目录节之前）。
        section_chapter = {}
        for i, para in enumerate(self.doc.paragraphs):
            sec = para_section[i]
            if sec in section_chapter:
                continue
            if '\t' in para.text:
                continue
            chap = _chapter_name(para.text.strip())
            if chap is not None:
                section_chapter[sec] = chap

        # 章节名称对某节生效，直到下一个出现章节标题的节为止（中间无标题的
        # 节沿用上一次出现的章节名称，如同一章节内的续页）。
        last_chapter = None
        for sec in range(sec_idx + 1):
            if sec in section_chapter:
                last_chapter = section_chapter[sec]
            else:
                section_chapter[sec] = last_chapter

        # 细则点(1)：检查范围从摘要页所在节开始，到文档末节（文章尾页）
        abstract_sec = None
        for i, para in enumerate(self.doc.paragraphs):
            if re.match(r'^摘\s*要$', para.text.strip()):
                abstract_sec = para_section[i]
                break
        if abstract_sec is None:
            print('  ✗ 未找到"摘要"页，无法确定页眉检查起点')
            return

        all_valid = True
        checked_count = 0
        errors = []

        for sidx in range(abstract_sec, len(self.doc.sections)):
            header = _section_header(sidx)
            if header is None or not header.paragraphs:
                continue

            current_chapter = section_chapter.get(sidx)
            for para in header.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                checked_count += 1
                runs = para.runs
                if not runs:
                    all_valid = False
                    errors.append(f'第{sidx}节页眉"{text[:20]}"无有效run')
                    continue

                font_info = self.get_font_info(runs[0])
                is_heiti = '黑体' in str(font_info['name']) or font_info['name'] == 'SimHei'
                is_size_correct = font_info['size'] and 10 <= font_info['size'] <= 11
                is_center = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
                matches_chapter = current_chapter is not None and current_chapter in text

                if not is_heiti:
                    errors.append(f'第{sidx}节页眉"{text[:20]}"非黑体')
                if not is_size_correct:
                    errors.append(f'第{sidx}节页眉"{text[:20]}"非五号')
                if not is_center:
                    errors.append(f'第{sidx}节页眉"{text[:20]}"未居中')
                if not matches_chapter:
                    errors.append(
                        f'第{sidx}节页眉"{text[:20]}"与当前章节"{current_chapter}"不一致'
                    )

                if not (is_heiti and is_size_correct and is_center and matches_chapter):
                    all_valid = False

        if checked_count > 0 and all_valid:
            self.score += 3
            msg = '+3：页眉从摘要页开始至文章尾页，页眉文字为黑体、五号、居中，页眉内容与当前页面所属章节名称一致'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print("  ✗ 页眉格式不符合要求（需黑体、五号、居中，且与章节名称一致）")
            for e in errors[:12]:
                print('      · ' + e)
            if len(errors) > 12:
                print(f'      · ...共{len(errors)}项')

    def check_page_number(self):
        '''检查页脚页码
        细则：+1：页脚页码汉字使用宋体、黑体小五号，
              数字使用 Times New Roman 、小五号、居中

        逐点核对：
          (1) 位置：页脚中的"页码"（含 PAGE 域或形如"第X页"、纯数字）；
          (2) 汉字使用宋体、黑体（两者之一即可）；
          (3) 汉字为小五号（Office 中"小五"精确为 9pt）；
          (4) 数字使用 Times New Roman；
          (5) 数字为小五号（9pt）；
          (6) 居中。

        所有属性沿 run→段落pPr→样式链→docDefaults 继承链解析，
        与办公软件 (Word/WPS) 实际渲染取值一致。
        '''
        # 判定段落是否含"页码"内容：含 w:instrText 的 PAGE 域，或文本含数字/"页"
        def is_page_number_para(para):
            for elem in para._element.iter(qn('w:instrText')):
                if elem.text and 'PAGE' in elem.text.upper():
                    return True
            txt = para.text
            return bool(re.search(r'\d', txt)) or ('页' in txt)

        errors = []
        checked_any = False

        for si, section in enumerate(self.doc.sections):
            footer = section.footer
            if footer is None:
                continue
            for para in footer.paragraphs:
                if not is_page_number_para(para):
                    continue
                checked_any = True

                # 细则点(6)：居中
                jc_val = self._effective_alignment(para)
                if jc_val != 'center':
                    errors.append(f'第{si}节页脚未居中(当前:{jc_val})')

                # 遍历 run，按字符类型分别核对
                for run in para.runs:
                    run_text = run.text
                    if not run_text:
                        continue

                    size_pt = self._effective_size(run, para)

                    # 该 run 中包含汉字的字符
                    if re.search(r'[一-鿿]', run_text):
                        # 细则点(3)：汉字小五号（9pt）
                        if size_pt is None or abs(size_pt - 9.0) >= 0.01:
                            errors.append(f'第{si}节汉字"{run_text.strip()[:6]}"字号非小五9pt(当前:{size_pt}pt)')
                        # 细则点(2)：汉字宋体或黑体
                        ea = self._effective_font(run, para, 'eastAsia')
                        is_songti = ea is not None and ('宋体' in ea or ea == 'SimSun')
                        is_heiti = ea is not None and ('黑体' in ea or ea == 'SimHei')
                        if not (is_songti or is_heiti):
                            errors.append(f'第{si}节汉字"{run_text.strip()[:6]}"字体非宋体/黑体(当前:{ea})')

                    # 该 run 中包含数字的字符
                    if re.search(r'\d', run_text):
                        # 细则点(5)：数字小五号（9pt）
                        if size_pt is None or abs(size_pt - 9.0) >= 0.01:
                            errors.append(f'第{si}节数字"{run_text.strip()[:6]}"字号非小五9pt(当前:{size_pt}pt)')
                        # 细则点(4)：数字 Times New Roman
                        ascii_font = self._effective_font(run, para, 'ascii')
                        if ascii_font != 'Times New Roman':
                            errors.append(f'第{si}节数字"{run_text.strip()[:6]}"字体非Times New Roman(当前:{ascii_font})')

        if not checked_any:
            print('  ✗ 未找到页脚页码')
            return

        if not errors:
            self.score += 1
            msg = '+1：页脚页码汉字使用宋体、黑体小五号，数字使用 Times New Roman 、小五号、居中'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            seen = []
            for e in errors:
                if e not in seen:
                    seen.append(e)
            print('  ✗ 页脚页码格式不符合要求：')
            for e in seen[:10]:
                print('      · ' + e)
            if len(seen) > 10:
                print(f'      · ...共{len(seen)}项')

    def check_page_number_format(self):
        '''检查页码格式变化
        细则：+3：摘要页和目录页页码使用大写罗马数字，
              从第一章开始，页码重新排序设置，格式用例如"第1页"

        逐点核对：
          (1) 摘要页页码使用大写罗马数字（sectPr/pgNumType@fmt = upperRoman）；
          (2) 目录页页码使用大写罗马数字（同上）；
          (3) 从"第一章"开始，页码重新排序设置
              （即"第一章"所在节的 sectPr/pgNumType@start 显式指定为 1）；
          (4) 从"第一章"开始，页码格式用例如"第1页"
              （页脚模板文本为"第 PAGE 页"，且 PAGE 域使用十进制，
               即 sectPr/pgNumType@fmt 为 decimal 或未指定[decimal 是默认值]）。

        所有判定都基于 OOXML sectPr/pgNumType 元素与页脚 PAGE 域，
        与办公软件 (Word/WPS) 实际渲染页码的机制一致。
        '''
        # 建立 body 中"每个段落文本 → 所属节序号"的映射
        # OOXML 节归属规则：段落 pPr/sectPr 存在时，本段是该节的末段；
        # 之后的段落归下一节；文档最后一节由 body 末尾的 sectPr 定义。
        para_section = []  # 与 self.doc.paragraphs 一一对应
        sec_idx = 0
        for para in self.doc.paragraphs:
            para_section.append(sec_idx)
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                sec_idx += 1

        # 定位摘要页、目录页、第一章所在的节
        abstract_sec = None
        toc_sec = None
        chapter1_sec = None
        for i, para in enumerate(self.doc.paragraphs):
            t = para.text.strip()
            if abstract_sec is None and re.match(r'^摘\s*要$', t):
                abstract_sec = para_section[i]
            if toc_sec is None and re.match(r'^目\s*录$', t):
                toc_sec = para_section[i]
            # "第一章"以段落样式为 Heading 1/标题 1 且不含制表符（排除目录条目）为准
            if chapter1_sec is None:
                sname = (para.style.name if para.style else '') or ''
                is_h1 = re.match(r'^(?:Heading|标题)\s*1$', sname) is not None
                if is_h1 and '\t' not in para.text and re.match(r'^第[一二三四五六七八九十\d]+章', t):
                    chapter1_sec = para_section[i]

        # 读取节的 pgNumType 属性
        def pg_num_type(section):
            sectPr = section._sectPr
            if sectPr is None:
                return (None, None)
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                return (None, None)
            return (pgNumType.get(qn('w:fmt')), pgNumType.get(qn('w:start')))

        # 判断某节页脚是否为"第 PAGE 页"模板
        def footer_is_chinese_page(section):
            footer = section.footer
            if footer is None:
                return False
            for para in footer.paragraphs:
                # 页脚文本（域已缓存值）形如"第1页"、"第 12 页"
                if re.match(r'^\s*第\s*\d+\s*页\s*$', para.text):
                    return True
                # 或按结构：段前含"第"、含 PAGE 域、段后含"页"
                has_qian = False
                has_page_field = False
                has_hou = False
                for elem in para._element.iter():
                    tag = elem.tag.split('}')[1] if '}' in elem.tag else ''
                    if tag == 't' and elem.text:
                        if '第' in elem.text:
                            has_qian = True
                        if '页' in elem.text:
                            has_hou = True
                    if tag == 'instrText' and elem.text and re.search(r'\bPAGE\b', elem.text):
                        has_page_field = True
                if has_qian and has_page_field and has_hou:
                    return True
            return False

        errors = []

        # 细则点(1)：摘要页使用大写罗马数字
        if abstract_sec is None:
            errors.append('未找到摘要页所在节')
        else:
            fmt, _ = pg_num_type(self.doc.sections[abstract_sec])
            if fmt != 'upperRoman':
                errors.append(f'摘要页页码格式非大写罗马数字(当前fmt={fmt})')

        # 细则点(2)：目录页使用大写罗马数字
        if toc_sec is None:
            errors.append('未找到目录页所在节')
        else:
            fmt, _ = pg_num_type(self.doc.sections[toc_sec])
            if fmt != 'upperRoman':
                errors.append(f'目录页页码格式非大写罗马数字(当前fmt={fmt})')

        # 细则点(3)：从第一章开始，页码重新排序（start=1）
        if chapter1_sec is None:
            errors.append('未找到"第一章"所在节')
        else:
            fmt, start = pg_num_type(self.doc.sections[chapter1_sec])
            if start != '1':
                errors.append(f'"第一章"所在节未重新起始页码(当前start={start})')

            # 细则点(4)：格式用例如"第1页"
            # (a) fmt 为 decimal（或未指定，Office 默认十进制）
            if fmt not in (None, 'decimal'):
                errors.append(f'"第一章"页码非十进制阿拉伯数字(当前fmt={fmt})')
            # (b) 页脚模板为"第 PAGE 页"
            if not footer_is_chinese_page(self.doc.sections[chapter1_sec]):
                errors.append('"第一章"页脚格式非"第X页"')

        if not errors:
            self.score += 3
            msg = '+3：摘要页和目录页页码使用大写罗马数字，从第一章开始，页码重新排序设置，格式用例如"第1页"'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 页码格式不符合要求：')
            for e in errors:
                print('      · ' + e)

    def check_figure_caption(self):
        '''检查图名格式
        细则：+5：图名均位于图片正下方居中，
              中文采用黑体五号，英文采用 Times New Roman 五号，不加粗

        逐点核对（以文档中每一张图片为起点检查，而非以已存在的图名段落为起点，
        避免"有图无图名"被漏判为通过）：
          (0) 每张正文图片下方必须存在图名（形如"图X.X ..."/"图X-X ..."）；
          (1) 图名位于图片"正下方"——紧邻图片所在段落的下一段即为图名段
              （允许中间跳过至多一个空段）；
          (2) 居中；
          (3) 中文字体为黑体（w:rFonts/w:eastAsia）；
          (4) 中文字号为五号（Office 中"五号"精确为 10.5pt）；
          (5) 英文字体为 Times New Roman（w:rFonts/w:ascii）；
          (6) 英文字号为五号（10.5pt）；
          (7) 不加粗（w:b 为假或缺省）。

        所有属性沿 run→段落pPr→样式链→docDefaults 继承链解析，与办公软件
        (Word/WPS) 实际渲染取值一致。
        '''
        FIG_RE = re.compile(r'^图\s*\d+[\.\-－]\d+')

        # 判断段落是否含图片（w:drawing 或 w:pict，跳过 mc:Fallback 分支避免重复计数）
        def has_image(para):
            for elem in para._element.iter():
                tag = elem.tag.split('}')[1] if '}' in elem.tag else ''
                if tag == 'Fallback':
                    continue
                if tag in ('drawing', 'pict'):
                    return True
            return False

        para_list = self.doc.paragraphs
        image_para_idxs = [i for i, para in enumerate(para_list) if has_image(para)]

        if not image_para_idxs:
            print('  - 未找到图片，跳过图名检查')
            return

        errors = []
        checked_images = 0

        for i in image_para_idxs:
            checked_images += 1
            img_label = para_list[i].text.strip()[:24] or f'第{checked_images}张图片'

            # 细则点(0)+(1)：图片正下方必须紧邻图名段（允许跳过至多一个空段）
            cap_para = None
            j = i + 1
            hops = 0
            while j < len(para_list) and hops < 2:
                text = para_list[j].text.strip()
                if FIG_RE.match(text):
                    cap_para = para_list[j]
                    break
                if text:
                    break  # 遇到非空且非图名段，停止下探
                if has_image(para_list[j]):
                    break  # 紧接着是另一张图片，说明本图缺少图名
                j += 1
                hops += 1

            if cap_para is None:
                errors.append(f'图片"{img_label}"下方缺少图名')
                continue

            label = cap_para.text.strip()[:24]

            # 细则点(2)：居中
            if self._effective_alignment(cap_para) != 'center':
                errors.append(f'"{label}" 未居中')

            # 逐 run 按字符类型核对字体/字号/加粗
            for run in cap_para.runs:
                rt = run.text
                if not rt or not rt.strip():
                    continue

                size_pt = self._effective_size(run, cap_para)
                is_bold = bool(self._effective_bold(run, cap_para))

                has_cn = bool(re.search(r'[一-鿿]', rt))
                has_en = bool(re.search(r'[A-Za-z]', rt))

                if has_cn:
                    # 细则点(4)：中文五号 10.5pt
                    if size_pt is None or abs(size_pt - 10.5) >= 0.01:
                        errors.append(f'"{label}" 中文字号非五号10.5pt(当前:{size_pt}pt)')
                    # 细则点(3)：中文黑体
                    ea = self._effective_font(run, cap_para, 'eastAsia')
                    if not (ea and ('黑体' in ea or ea == 'SimHei')):
                        errors.append(f'"{label}" 中文字体非黑体(当前:{ea})')

                if has_en:
                    # 细则点(6)：英文五号 10.5pt
                    if size_pt is None or abs(size_pt - 10.5) >= 0.01:
                        errors.append(f'"{label}" 英文字号非五号10.5pt(当前:{size_pt}pt)')
                    # 细则点(5)：英文 Times New Roman
                    asc = self._effective_font(run, cap_para, 'ascii')
                    if asc != 'Times New Roman':
                        errors.append(f'"{label}" 英文字体非Times New Roman(当前:{asc})')

                # 细则点(7)：不加粗
                if is_bold:
                    errors.append(f'"{label}" 加粗')

        # 每类问题只保留一次
        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            self.score += 5
            msg = '+5：图名均位于图片正下方居中，中文采用黑体五号，英文采用 Times New Roman 五号，不加粗'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 图名格式不符合要求：')
            for e in seen[:12]:
                print('      · ' + e)
            if len(seen) > 12:
                print(f'      · ...共{len(seen)}项')

    def check_table_title(self):
        '''检查表题格式
        细则：+1：正文表格均有表题，表题位于表格正上方居中，
              中文采用黑体五号，英文采用 Times New Roman 五号，不加粗

        逐点核对：
          (1) 正文所有表格均有表题；
          (2) 表题位于表格正上方（紧邻表格的上一段落是表题段）；
          (3) 居中；
          (4) 中文字体为黑体（w:rFonts/w:eastAsia）；
          (5) 中文字号为五号（Office 中"五号"=10.5pt）；
          (6) 英文字体为 Times New Roman（w:rFonts/w:ascii）；
          (7) 英文字号为五号（10.5pt）；
          (8) 不加粗（w:b 为假）。

        表题识别范围：形如"表X.X ..." / "表X-X ..." / "表X－X ..."的段落。
        所有属性沿 run→段落pPr→样式链→docDefaults 继承链解析，
        与办公软件 (Word/WPS) 实际渲染取值一致。
        '''
        tables = self.doc.tables
        if not tables:
            print('  - 未找到表格，跳过表题检查')
            return

        TABLE_CAP_RE = re.compile(r'^表\s*\d+[\.\-－]\d+')

        # 遍历 body 顺序，为每个 w:tbl 定位紧邻的上一段（跳过一个空段）
        body = self.doc.element.body
        children = list(body.iterchildren())
        # 段落 XML → Paragraph 对象 的映射
        para_by_elem = {p._element: p for p in self.doc.paragraphs}

        errors = []
        found_titles = 0

        for idx, child in enumerate(children):
            tag = child.tag.split('}')[1] if '}' in child.tag else ''
            if tag != 'tbl':
                continue

            # 细则点(2)：表题位于表格正上方——上溯最多两段（允许一空段）
            title_para = None
            hops = 0
            j = idx - 1
            while j >= 0 and hops < 2:
                prev = children[j]
                ptag = prev.tag.split('}')[1] if '}' in prev.tag else ''
                if ptag != 'p':
                    break
                para_obj = para_by_elem.get(prev)
                if para_obj is None:
                    break
                text = para_obj.text.strip()
                if TABLE_CAP_RE.match(text):
                    title_para = para_obj
                    break
                if text:
                    break  # 遇到非空且非表题段，停止
                j -= 1
                hops += 1

            # 细则点(1)：每个表格必须有表题
            if title_para is None:
                errors.append(f'第{found_titles + 1}张表格上方缺少表题')
                continue

            found_titles += 1
            label = title_para.text.strip()[:24]

            # 细则点(3)：居中
            if self._effective_alignment(title_para) != 'center':
                errors.append(f'"{label}" 未居中')

            # 逐 run 按字符类型核对字体/字号/加粗
            for run in title_para.runs:
                rt = run.text
                if not rt or not rt.strip():
                    continue

                size_pt = self._effective_size(run, title_para)
                is_bold = bool(self._effective_bold(run, title_para))

                has_cn = bool(re.search(r'[一-鿿]', rt))
                has_en = bool(re.search(r'[A-Za-z]', rt))

                if has_cn:
                    # 细则点(5)：中文五号 10.5pt
                    if size_pt is None or abs(size_pt - 10.5) >= 0.01:
                        errors.append(f'"{label}" 中文字号非五号10.5pt(当前:{size_pt}pt)')
                    # 细则点(4)：中文黑体
                    ea = self._effective_font(run, title_para, 'eastAsia')
                    if not (ea and ('黑体' in ea or ea == 'SimHei')):
                        errors.append(f'"{label}" 中文字体非黑体(当前:{ea})')

                if has_en:
                    # 细则点(7)：英文五号 10.5pt
                    if size_pt is None or abs(size_pt - 10.5) >= 0.01:
                        errors.append(f'"{label}" 英文字号非五号10.5pt(当前:{size_pt}pt)')
                    # 细则点(6)：英文 Times New Roman
                    asc = self._effective_font(run, title_para, 'ascii')
                    if asc != 'Times New Roman':
                        errors.append(f'"{label}" 英文字体非Times New Roman(当前:{asc})')

                # 细则点(8)：不加粗
                if is_bold:
                    errors.append(f'"{label}" 加粗')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            self.score += 1
            msg = '+1：正文表格均有表题，表题位于表格正上方居中，中文采用黑体五号，英文采用 Times New Roman 五号，不加粗'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 表题格式不符合要求：')
            for e in seen[:12]:
                print('      · ' + e)
            if len(seen) > 12:
                print(f'      · ...共{len(seen)}项')

    def check_table_three_line(self):
        '''检查三线表格式
        细则：+3：正文所有表格为三线表，仅保留顶线、栏目线、底线，无竖线；
              顶线和底线为1.5磅，栏目线为0.75磅

        逐点核对（针对 Office 实际渲染）：
          (1) 正文所有表格；
          (2) 三线表：仅有顶线、栏目线、底线三条水平线；
          (3) 无竖线：所有单元格左右边界（外沿与内部）均为 nil/none；
          (4) 顶线 = 1.5 磅（w:sz=12，1/8 磅单位）；
          (5) 栏目线 = 0.75 磅（w:sz=6）；
          (6) 底线 = 1.5 磅（w:sz=12）；
          (7) 中间其他横线（第1/2行以外的横向边界）不存在。

        边界解析：单元格 w:tcPr/w:tcBorders 优先；未设置则回退到
        w:tblPr/w:tblBorders 的对应边（外沿用 top/bottom/left/right，
        内部用 insideH/insideV），与 Word/WPS 渲染取值一致。
        '''
        if not self.doc.tables:
            print('  - 未找到表格，跳过三线表检查')
            return

        def _val_sz(elem):
            if elem is None:
                return (None, None)
            v = elem.get(qn('w:val'))
            s = elem.get(qn('w:sz'))
            return (v, float(s) / 8 if s else None)

        def _has_line(vs):
            v, _ = vs
            return v not in (None, 'nil', 'none')

        def _line_width(vs):
            v, s = vs
            if v in (None, 'nil', 'none'):
                return 0.0
            return s or 0.0

        def _tbl_borders(table):
            d = {k: None for k in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV')}
            tblPr = table._tbl.find(qn('w:tblPr'))
            if tblPr is None:
                return d
            tb = tblPr.find(qn('w:tblBorders'))
            if tb is None:
                return d
            for k in d:
                d[k] = tb.find(qn('w:' + k))
            return d

        def _cell_eff(cell, edge, i, j, n_rows, n_cols, tbl_b):
            # 单元格级 tcBorders 优先
            tcPr = cell._tc.tcPr
            if tcPr is not None:
                tcB = tcPr.find(qn('w:tcBorders'))
                if tcB is not None:
                    b = tcB.find(qn('w:' + edge))
                    if b is not None:
                        return b
            # 回退到表级
            if edge == 'top':
                return tbl_b['top'] if i == 0 else tbl_b['insideH']
            if edge == 'bottom':
                return tbl_b['bottom'] if i == n_rows - 1 else tbl_b['insideH']
            if edge == 'left':
                return tbl_b['left'] if j == 0 else tbl_b['insideV']
            if edge == 'right':
                return tbl_b['right'] if j == n_cols - 1 else tbl_b['insideV']
            return None

        errors = []
        for ti, table in enumerate(self.doc.tables):
            rows = table.rows
            n_rows = len(rows)
            if n_rows == 0:
                errors.append(f'表{ti+1}: 无行')
                continue
            n_cols = max(len(r.cells) for r in rows)
            tbl_b = _tbl_borders(table)
            tbl_errs = []

            # 细则点(3)/(7)：竖线与顶/底线粗细，逐格四边核对
            for i, row in enumerate(rows):
                for j, cell in enumerate(row.cells):
                    # 竖线：外沿 + 内部竖向，全部不允许
                    for edge in ('left', 'right'):
                        if _has_line(_val_sz(_cell_eff(cell, edge, i, j, n_rows, n_cols, tbl_b))):
                            tbl_errs.append('出现竖线')

                    # 顶线（i==0 top）：1.5磅
                    if i == 0:
                        w = _line_width(_val_sz(_cell_eff(cell, 'top', i, j, n_rows, n_cols, tbl_b)))
                        if abs(w - 1.5) >= 0.01:
                            tbl_errs.append(f'顶线粗细≠1.5磅(当前{w}磅)')

                    # 底线（i==last bottom）：1.5磅
                    if i == n_rows - 1:
                        w = _line_width(_val_sz(_cell_eff(cell, 'bottom', i, j, n_rows, n_cols, tbl_b)))
                        if abs(w - 1.5) >= 0.01:
                            tbl_errs.append(f'底线粗细≠1.5磅(当前{w}磅)')

            # 细则点(5)：栏目线——row0/row1 边界的合成宽度=0.75磅
            if n_rows >= 2:
                for j in range(n_cols):
                    try:
                        c0, c1 = rows[0].cells[j], rows[1].cells[j]
                    except IndexError:
                        continue
                    w = max(
                        _line_width(_val_sz(_cell_eff(c0, 'bottom', 0, j, n_rows, n_cols, tbl_b))),
                        _line_width(_val_sz(_cell_eff(c1, 'top', 1, j, n_rows, n_cols, tbl_b))),
                    )
                    if abs(w - 0.75) >= 0.01:
                        tbl_errs.append(f'栏目线粗细≠0.75磅(当前{w}磅)')
                        break

            # 细则点(2)/(7)：其他中间横线不存在（rows i/i+1 之间, i∈[1, n-2]）
            for i in range(1, n_rows - 1):
                extra = False
                for j in range(n_cols):
                    try:
                        ca, cb = rows[i].cells[j], rows[i + 1].cells[j]
                    except IndexError:
                        continue
                    w = max(
                        _line_width(_val_sz(_cell_eff(ca, 'bottom', i, j, n_rows, n_cols, tbl_b))),
                        _line_width(_val_sz(_cell_eff(cb, 'top', i + 1, j, n_rows, n_cols, tbl_b))),
                    )
                    if w > 0:
                        extra = True
                        break
                if extra:
                    tbl_errs.append(f'第{i+1}/{i+2}行之间出现多余横线')
                    break

            seen = []
            for e in tbl_errs:
                if e not in seen:
                    seen.append(e)
            for e in seen:
                errors.append(f'表{ti+1}: {e}')

        if not errors:
            self.score += 3
            msg = '+3：正文所有表格为三线表，仅保留顶线、栏目线、底线，无竖线；顶线和底线为1.5磅，栏目线为0.75磅'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 表格不符合三线表格式：')
            for e in errors[:12]:
                print('      · ' + e)
            if len(errors) > 12:
                print(f'      · ...共{len(errors)}项')

    def check_table_content(self):
        '''检查表格内容格式
        细则：+3：正文所有表格表头加粗，表内文本：中文为宋体五号，
              英文和数字为Times New Roman五号，行列对齐

        逐点核对（走 run→pPr→样式链→docDefaults 的继承链，与 Office 渲染一致）：
          (1) 正文所有表格；
          (2) 表头（首行）非空 run 均加粗；
          (3) 表内所有单元格文本中的中文字符：字体为宋体（w:rFonts/w:eastAsia）；
          (4) 表内所有单元格文本中的中文字符：字号为五号 10.5pt；
          (5) 表内所有单元格文本中的英文/数字：字体为 Times New Roman（w:rFonts/w:ascii）；
          (6) 表内所有单元格文本中的英文/数字：字号为五号 10.5pt；
          (7) 行列对齐：同一列内非空单元格的段落水平对齐一致；
                       同一行内非空单元格的垂直对齐一致（w:vAlign 继承 tcPr→trPr）。
        '''
        if not self.doc.tables:
            print('  - 未找到表格，跳过表内文本检查')
            return

        def _cell_vAlign(cell):
            tcPr = cell._tc.tcPr
            if tcPr is not None:
                va = tcPr.find(qn('w:vAlign'))
                if va is not None:
                    return va.get(qn('w:val'))
            return None

        errors = []
        for ti, table in enumerate(self.doc.tables):
            rows = table.rows
            if not rows:
                continue
            n_cols = max(len(r.cells) for r in rows)

            # 细则点(2)：表头首行加粗
            for cell in rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        if not self._effective_bold(run, para):
                            errors.append(f'表{ti+1} 表头未加粗: "{run.text.strip()[:20]}"')

            # 细则点(3)-(6)：全表 run 按字符类型核对字体/字号
            for i, row in enumerate(rows):
                for j, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            rt = run.text
                            if not rt or not rt.strip():
                                continue
                            size_pt = self._effective_size(run, para)
                            has_cn = bool(re.search(r'[\u4e00-\u9fff]', rt))
                            has_en = bool(re.search(r'[A-Za-z0-9]', rt))

                            if has_cn:
                                ea = self._effective_font(run, para, 'eastAsia')
                                if not (ea and ('宋体' in ea or ea == 'SimSun')):
                                    errors.append(f'表{ti+1} 中文非宋体(当前:{ea}): "{rt.strip()[:16]}"')
                                if size_pt is None or abs(size_pt - 10.5) >= 0.01:
                                    errors.append(f'表{ti+1} 中文非五号10.5pt(当前:{size_pt}pt): "{rt.strip()[:16]}"')

                            if has_en:
                                asc = self._effective_font(run, para, 'ascii')
                                if asc != 'Times New Roman':
                                    errors.append(f'表{ti+1} 英文/数字非Times New Roman(当前:{asc}): "{rt.strip()[:16]}"')
                                if size_pt is None or abs(size_pt - 10.5) >= 0.01:
                                    errors.append(f'表{ti+1} 英文/数字非五号10.5pt(当前:{size_pt}pt): "{rt.strip()[:16]}"')

            # 细则点(7-列)：同列水平对齐一致
            for j in range(n_cols):
                aligns = []
                for i, row in enumerate(rows):
                    if j >= len(row.cells):
                        continue
                    cell = row.cells[j]
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        aligns.append(self._effective_alignment(para) or 'left')
                if aligns and len(set(aligns)) > 1:
                    errors.append(f'表{ti+1} 第{j+1}列水平对齐不一致: {sorted(set(aligns))}')

            # 细则点(7-行)：同行垂直对齐一致
            for i, row in enumerate(rows):
                valigns = []
                for cell in row.cells:
                    if not any(p.text.strip() for p in cell.paragraphs):
                        continue
                    valigns.append(_cell_vAlign(cell) or 'top')
                if valigns and len(set(valigns)) > 1:
                    errors.append(f'表{ti+1} 第{i+1}行垂直对齐不一致: {sorted(set(valigns))}')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            self.score += 3
            msg = '+3：正文所有表格表头加粗，表内文本：中文为宋体五号，英文和数字为Times New Roman五号，行列对齐'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 表格内容格式不符合要求：')
            for e in seen[:12]:
                print('      · ' + e)
            if len(seen) > 12:
                print(f'      · ...共{len(seen)}项')

    def check_code_format(self):
        '''检查代码段格式
        细则：+1：程序代码段为Times New Roman五号，代码段行距为1.25倍

        逐点核对（走 run→pPr→样式链→docDefaults 的继承链，与 Office 渲染一致）：
          (1) 定位所有程序代码段——综合样式名、等宽字体、缩进/排版上下文及
              代码内容特征识别（见 _is_code_paragraph），而非仅靠有限关键字
              正则，避免漏检导致该项被误判为通过；
          (2) 代码段字体为 Times New Roman（w:rFonts/w:ascii）；
          (3) 代码段字号为五号 10.5pt（w:sz，单位为半磅）；
          (4) 代码段行距为 1.25 倍
              —— OOXML: w:pPr/w:spacing/@w:lineRule="auto" 且 @w:line="300"
                 （auto 时 val/240 为倍数：300/240 = 1.25），
                 与 Word/WPS 的"多倍行距 1.25"一致。

        若全文未识别到任何代码段：该 rubric 项针对"程序代码段"的格式要求
        在没有代码段时不适用，按语义视为满足（不因缺少适用对象而扣分），
        但会明确提示未检测到代码段，避免与"检查过但合规"混淆。
        '''
        code_paras = [p for p in self.doc.paragraphs if self._is_code_paragraph(p)]

        if not code_paras:
            self.score += 1
            msg = '+1：程序代码段为Times New Roman五号，代码段行距为1.25倍（未检测到代码段，视为不适用）'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
            return

        errors = []
        for para in code_paras:
            snippet = para.text.strip()[:32]

            # 细则点(4)：段落行距 1.25 倍（w:lineRule=auto, w:line=300）
            pPr = para._element.find(qn('w:pPr'))
            spacing = pPr.find(qn('w:spacing')) if pPr is not None else None
            line_rule = spacing.get(qn('w:lineRule')) if spacing is not None else None
            line_val = spacing.get(qn('w:line')) if spacing is not None else None
            ok_spacing = (line_rule == 'auto' and line_val == '300')
            if not ok_spacing:
                errors.append(
                    f'代码段行距非1.25倍(lineRule={line_rule}, line={line_val}): "{snippet}"'
                )

            # 细则点(2)(3)：字体 Times New Roman、字号 10.5pt——按 run 逐一核对
            runs_iter = para.runs or []
            for run in runs_iter:
                rt = run.text
                if not rt or not rt.strip():
                    continue
                asc = self._effective_font(run, para, 'ascii')
                if asc != 'Times New Roman':
                    errors.append(f'代码段字体非Times New Roman(当前:{asc}): "{snippet}"')
                size_pt = self._effective_size(run, para)
                if size_pt is None or abs(size_pt - 10.5) >= 0.01:
                    errors.append(f'代码段字号非五号10.5pt(当前:{size_pt}pt): "{snippet}"')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            self.score += 1
            msg = '+1：程序代码段为Times New Roman五号，代码段行距为1.25倍'
            self.checked_points.append(msg)
            print('  ✓ ' + msg)
        else:
            print('  ✗ 代码段格式不符合要求：')
            for e in seen[:12]:
                print('      · ' + e)
            if len(seen) > 12:

                print(f'      · ...共{len(seen)}项')

    def check_first_page(self):
        '''检查首页格式（扣分项）
        细则：-5：第1页不满足以下任意一项：
              (A) 页面顶部"澄云信息科技学院"：小一号，宋体，加粗；
              (B) "毕业设计（论文）"为宋体，一号；
              (C) "基于Spark的园区能耗分析与可视化系统"：黑体、三号、加粗、蓝色；
              (D) 出现冒号的行：冒号及冒号前的文本为三号、宋体、加粗；
                  冒号后所有文本下方均有横线；
                  所有横线全部互相平行且长度一致；
              (E) 页面底部"澄云信息科技学院"：宋体、三号、加粗。

        字号（Office 标准）：小一=24pt，一号=26pt，三号=16pt。
        所有属性沿 run→pPr→样式链→docDefaults 继承链读取，颜色与下划线
        直接读 w:color/w:u OOXML 元素——与 Word/WPS 渲染取值一致。
        '''
        errors = []

        # 首页范围：使用真实分页/分节信号确定第1页范围，而不是用"摘要/第一章"
        # 等正文内容启发式截断，避免第一页跨多段或题名文本重合时漏检/误检。
        first_page_end = len(self.doc.paragraphs)
        for i, para in enumerate(self.doc.paragraphs):
            if i > 0 and para._element.xpath('.//w:lastRenderedPageBreak'):
                first_page_end = i
                break

            page_break_before_next = False
            for br in para._element.xpath('.//w:br'):
                if br.get(qn('w:type')) == 'page':
                    page_break_before_next = True
                    break
            if page_break_before_next:
                first_page_end = i + 1
                break

            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                first_page_end = i + 1
                break

        first_page_paras = self.doc.paragraphs[:first_page_end]

        def _run_color(run):
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                c = rPr.find(qn('w:color'))
                if c is not None:
                    return c.get(qn('w:val'))
            return None

        def _run_underline(run, para):
            # 沿 run→pPr→style chain→docDefaults 找 w:u
            for holder in [run._element.find(qn('w:rPr'))]:
                if holder is not None:
                    u = holder.find(qn('w:u'))
                    if u is not None:
                        return u.get(qn('w:val'))
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                rPr_p = pPr.find(qn('w:rPr'))
                if rPr_p is not None:
                    u = rPr_p.find(qn('w:u'))
                    if u is not None:
                        return u.get(qn('w:val'))
            for style in self._iter_style_chain(para):
                el = getattr(style, 'element', None)
                if el is None:
                    continue
                rPr_s = el.find(qn('w:rPr'))
                if rPr_s is not None:
                    u = rPr_s.find(qn('w:u'))
                    if u is not None:
                        return u.get(qn('w:val'))
            dd = self._doc_defaults_rPr()
            if dd is not None:
                u = dd.find(qn('w:u'))
                if u is not None:
                    return u.get(qn('w:val'))
            return None

        def _is_blue(hex_val):
            # 蓝色：B 通道显著大于 R/G。既覆盖纯蓝 0000FF，也覆盖 Office 默认蓝 0070C0。
            if not hex_val or hex_val == 'auto':
                return False
            hv = hex_val.lstrip('#')
            if len(hv) != 6:
                return False
            try:
                r = int(hv[0:2], 16); g = int(hv[2:4], 16); b = int(hv[4:6], 16)
            except ValueError:
                return False
            return b >= 128 and b > r + 50 and b > g + 50

        def _est_width_pt(text, size_pt):
            # 用于比较冒号后横线长度：CJK 按字号宽估计，ASCII 半宽估计
            w = 0.0
            sz = size_pt or 14.0
            for ch in text:
                if '一' <= ch <= '鿿' or '　' <= ch <= '〿' or '＀' <= ch <= '￯':
                    w += sz
                elif ch == '\t':
                    w += sz * 2
                else:
                    w += sz / 2.0
            return w

        def _line_end_pos_pt(cp):
            '''横线右端点位置（pt）：若冒号后文本以制表符结尾且段落设有
            制表位（w:tabs/w:tab@w:pos，单位 twip），横线长度由制表位置
            决定（Office 中"填线到制表位"的常见实现），此时精确到 pt；
            否则回退到按字符宽度估算（_est_width_pt 累加值），作为近似。'''
            pPr = cp._element.find(qn('w:pPr'))
            tabs = pPr.find(qn('w:tabs')) if pPr is not None else None
            if tabs is not None and cp.text.rstrip().endswith('\t') or (tabs is not None and '\t' in cp.text):
                tab_positions = []
                for tab in tabs.findall(qn('w:tab')):
                    pos = tab.get(qn('w:pos'))
                    if pos is not None:
                        try:
                            tab_positions.append(int(pos) / 20.0)  # twip -> pt
                        except ValueError:
                            pass
                if tab_positions:
                    return max(tab_positions)
            return None

        # 收集三个特征段落 & 冒号行 & 底部"澄云信息科技学院"
        top_college_para = None
        thesis_para = None
        title_para = None
        colon_paras = []
        bottom_college_para = None

        college_seen = 0
        for p in first_page_paras:
            t = p.text.strip()
            if not t:
                continue
            if t == '澄云信息科技学院':
                college_seen += 1
                if college_seen == 1:
                    top_college_para = p
                else:
                    bottom_college_para = p  # 后续出现均视为页面底部（最终以最后一次为准）
                continue
            if thesis_para is None and re.match(r'^毕业设计\s*[（(]\s*论文\s*[)）]\s*$', t):
                thesis_para = p
                continue
            if title_para is None and '基于Spark的园区能耗分析与可视化系统' in t:
                title_para = p
                continue
            if '：' in t:
                # 冒号前是短标签（如"院  校"/"专  业"/"指导教师"），排除非标签行
                head = t.split('：', 1)[0]
                if len(head.replace(' ', '')) <= 8:
                    colon_paras.append(p)

        # (A) 页面顶部"澄云信息科技学院"：小一号(24pt)、宋体、加粗
        if top_college_para is None:
            errors.append('未找到页面顶部"澄云信息科技学院"')
        else:
            for run in top_college_para.runs:
                if not run.text.strip():
                    continue
                ea = self._effective_font(run, top_college_para, 'eastAsia')
                sz = self._effective_size(run, top_college_para)
                bd = self._effective_bold(run, top_college_para)
                if not (ea and ('宋体' in ea or ea == 'SimSun')):
                    errors.append(f'页面顶部"澄云信息科技学院"非宋体(当前:{ea})')
                if sz is None or abs(sz - 24.0) >= 0.01:
                    errors.append(f'页面顶部"澄云信息科技学院"非小一号24pt(当前:{sz}pt)')
                if not bd:
                    errors.append('页面顶部"澄云信息科技学院"未加粗')

        # (B) "毕业设计（论文）"：宋体、一号(26pt)
        if thesis_para is None:
            errors.append('未找到"毕业设计（论文）"')
        else:
            for run in thesis_para.runs:
                if not run.text.strip():
                    continue
                ea = self._effective_font(run, thesis_para, 'eastAsia')
                sz = self._effective_size(run, thesis_para)
                if not (ea and ('宋体' in ea or ea == 'SimSun')):
                    errors.append(f'"毕业设计（论文）"非宋体(当前:{ea})')
                if sz is None or abs(sz - 26.0) >= 0.01:
                    errors.append(f'"毕业设计（论文）"非一号26pt(当前:{sz}pt)')

        # (C) 论文标题：黑体、三号(16pt)、加粗、蓝色
        if title_para is None:
            errors.append('未找到论文标题"基于Spark的园区能耗分析与可视化系统"')
        else:
            for run in title_para.runs:
                rt = run.text
                if not rt or not rt.strip():
                    continue
                # 中文黑体（东亚字体）；英文/数字部分沿用 ascii 字体不作限制（细则未强制英文字体）
                if re.search(r'[一-鿿]', rt):
                    ea = self._effective_font(run, title_para, 'eastAsia')
                    if not (ea and ('黑体' in ea or ea == 'SimHei')):
                        errors.append(f'论文标题中文非黑体(当前:{ea})')
                sz = self._effective_size(run, title_para)
                if sz is None or abs(sz - 16.0) >= 0.01:
                    errors.append(f'论文标题非三号16pt(当前:{sz}pt)')
                if not self._effective_bold(run, title_para):
                    errors.append('论文标题未加粗')
                color = _run_color(run)
                if not _is_blue(color):
                    errors.append(f'论文标题字体颜色非蓝色(当前:{color})')

        # (D) 冒号行
        if not colon_paras:
            errors.append('未找到冒号行')
        else:
            underline_widths = []
            for cp in colon_paras:
                # 拆分：冒号前(含冒号) vs 冒号后
                full_text = cp.text
                colon_idx = full_text.find('：')
                if colon_idx < 0:
                    continue

                pre_target_len = colon_idx + 1  # 冒号前含冒号的字符数
                acc = 0
                pre_runs, post_runs = [], []
                for run in cp.runs:
                    rt = run.text or ''
                    if not rt:
                        continue
                    start = acc
                    end = acc + len(rt)
                    acc = end
                    if end <= pre_target_len:
                        pre_runs.append((run, rt))
                    elif start >= pre_target_len:
                        post_runs.append((run, rt))
                    else:
                        # 一个 run 跨越冒号——极少见；按文本拆分并算作两部分
                        cut = pre_target_len - start
                        pre_runs.append((run, rt[:cut]))
                        post_runs.append((run, rt[cut:]))

                pre_label = full_text[:pre_target_len]

                # 冒号及冒号前的文本：三号(16pt)、宋体、加粗
                for run, seg in pre_runs:
                    if not seg.strip():
                        continue
                    ea = self._effective_font(run, cp, 'eastAsia')
                    sz = self._effective_size(run, cp)
                    bd = self._effective_bold(run, cp)
                    if not (ea and ('宋体' in ea or ea == 'SimSun')):
                        errors.append(f'冒号行"{pre_label}"冒号前非宋体(当前:{ea})')
                    if sz is None or abs(sz - 16.0) >= 0.01:
                        errors.append(f'冒号行"{pre_label}"冒号前非三号16pt(当前:{sz}pt)')
                    if not bd:
                        errors.append(f'冒号行"{pre_label}"冒号前未加粗')

                # 冒号后所有文本下方均有横线（w:u 非 none）
                post_has_underline = True
                width = 0.0
                for run, seg in post_runs:
                    if not seg.strip():
                        # 空白也计入宽度，但不强制其下划线（Office 中空格是否有横线取决于是否包含在下划线 run 中）
                        u = _run_underline(run, cp)
                        sz = self._effective_size(run, cp) or 16.0
                        if u and u != 'none':
                            width += _est_width_pt(seg, sz)
                        continue
                    u = _run_underline(run, cp)
                    if not u or u == 'none':
                        post_has_underline = False
                        errors.append(f'冒号行"{pre_label}"冒号后文本"{seg.strip()[:16]}"无下划线')
                    sz = self._effective_size(run, cp) or 16.0
                    width += _est_width_pt(seg, sz)

                if post_has_underline:
                    # 优先使用制表位精确定位横线右端点（Office 常见做法：横线
                    # 通过"下划线+制表符填充到制表位"实现），取不到时回退到按
                    # 字符宽度估算的近似值。
                    tab_end = _line_end_pos_pt(cp)
                    underline_widths.append((pre_label, tab_end if tab_end is not None else width))

            # 横线互相平行——所有冒号后文本均为水平方向的下划线（w:u），
            # 不存在斜线/竖线等非水平线型，故只要上一步逐 run 检查全部通过
            # 即视为互相平行。
            # 横线长度一致——比较所有 underline_widths（制表位精确值优先，
            # 否则为估算宽度）。
            if len(underline_widths) >= 2:
                ws = [w for _, w in underline_widths]
                if max(ws) - min(ws) > 1.0:  # 允许 1pt 舍入误差
                    detail = ', '.join(f'{lbl.strip()}≈{w:.1f}pt' for lbl, w in underline_widths)
                    errors.append(f'冒号后横线长度不一致({detail})')


        # (E) 页面底部"澄云信息科技学院"：宋体、三号(16pt)、加粗
        if bottom_college_para is None:
            errors.append('未找到页面底部"澄云信息科技学院"')
        else:
            for run in bottom_college_para.runs:
                if not run.text.strip():
                    continue
                ea = self._effective_font(run, bottom_college_para, 'eastAsia')
                sz = self._effective_size(run, bottom_college_para)
                bd = self._effective_bold(run, bottom_college_para)
                if not (ea and ('宋体' in ea or ea == 'SimSun')):
                    errors.append(f'页面底部"澄云信息科技学院"非宋体(当前:{ea})')
                if sz is None or abs(sz - 16.0) >= 0.01:
                    errors.append(f'页面底部"澄云信息科技学院"非三号16pt(当前:{sz}pt)')
                if not bd:
                    errors.append('页面底部"澄云信息科技学院"未加粗')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 首页格式符合要求')
        else:
            self.score -= 5
            msg = ('-5：第1页不满足以下任意一项：页面顶部"澄云信息科技学院"：小一号，宋体，加粗；'
                   '"毕业设计（论文）"为宋体，一号；"基于Spark的园区能耗分析与可视化系统"字号为黑体，三号，加粗，'
                   '字体颜色为蓝色；出现冒号的行：冒号及冒号前的文本为三号、宋体、加粗，冒号后所有文本下方均有横线，'
                   '所有横线全部互相平行且长度一致；页面底部"澄云信息科技学院" 宋体、三号、加粗')
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:20]:
                print('      · ' + e)
            if len(seen) > 20:
                print(f'      · ...共{len(seen)}项')

    def check_chinese_abstract_format(self):
        '''检查中文摘要页格式（扣分项）
        细则：-1：中文摘要页不满足以下任意一项：
              标题下方文本为小四号宋体、首行缩进2字符、1.5倍行距，两端对齐

        逐点核对（沿 run→pPr→样式链→docDefaults 继承链，与 Office 渲染一致）：
          (1) 目标范围：中文摘要标题下方的正文段落（到"关键词"前为止）；
          (2) 字体：中文为宋体（w:rFonts/w:eastAsia）；
          (3) 字号：小四号 12pt（w:sz，单位为半磅）；
          (4) 首行缩进 = 2 个中文字符
              —— Office 存储为 w:pPr/w:ind/@w:firstLineChars="200"，
                 或 @w:firstLine（twips）等于 2 × 字号(pt) × 20；
                 段落 pPr 未设置时，从段落样式链继承；
          (5) 行距 = 1.5 倍
              —— OOXML: w:pPr/w:spacing/@w:lineRule="auto" 且 @w:line="360"
                 （360/240 = 1.5）；段落 pPr 未设置时，从段落样式链继承；
          (6) 两端对齐
              —— w:pPr/w:jc/@w:val = "both"（Word/WPS UI 显示为"两端对齐"）；
                 段落 pPr 未设置时，从段落样式链继承。
        '''
        # 沿 段落 pPr → 段落样式链 查找指定 pPr 子元素（继承查找，与 Office 一致）
        def _eff_pPr_child(para, child_local_name):
            child_qn = qn('w:' + child_local_name)
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                c = pPr.find(child_qn)
                if c is not None:
                    return c
            for style in self._iter_style_chain(para):
                el = getattr(style, 'element', None)
                if el is None:
                    continue
                pPr_s = el.find(qn('w:pPr'))
                if pPr_s is not None:
                    c = pPr_s.find(child_qn)
                    if c is not None:
                        return c
            return None

        # 定位中文摘要标题
        title_idx = -1
        for i, para in enumerate(self.doc.paragraphs):
            if re.match(r'^摘\s*要$', para.text.strip()):
                title_idx = i
                break

        if title_idx < 0:
            print('  - 未找到中文摘要标题，跳过检查')
            return

        # 收集"摘要"标题下方到"关键词"之前的所有非空段落
        body_paras = []
        for j in range(title_idx + 1, len(self.doc.paragraphs)):
            p = self.doc.paragraphs[j]
            t = p.text.strip()
            if not t:
                continue
            if re.match(r'^关\s*键\s*词', t):
                break
            body_paras.append(p)

        if not body_paras:
            print('  - 未找到中文摘要正文段落，跳过检查')
            return

        errors = []
        for para in body_paras:
            snippet = para.text.strip()[:20]

            # 细则点(2)(3)：中文宋体、小四号 12pt——按 run 内 CJK 字符核对
            for run in para.runs:
                rt = run.text
                if not rt or not rt.strip():
                    continue
                if re.search(r'[一-鿿]', rt):
                    ea = self._effective_font(run, para, 'eastAsia')
                    if not (ea and ('宋体' in ea or ea == 'SimSun')):
                        errors.append(f'"{snippet}" 中文非宋体(当前:{ea})')
                    sz = self._effective_size(run, para)
                    if sz is None or abs(sz - 12.0) >= 0.01:
                        errors.append(f'"{snippet}" 字号非小四号12pt(当前:{sz}pt)')

            # 段落中文字号（用于把 firstLine (twips) 换算成"字符数"）
            para_cjk_size = None
            for run in para.runs:
                if re.search(r'[一-鿿]', run.text or ''):
                    para_cjk_size = self._effective_size(run, para)
                    if para_cjk_size:
                        break

            ind = _eff_pPr_child(para, 'ind')
            spacing = _eff_pPr_child(para, 'spacing')
            jc = _eff_pPr_child(para, 'jc')

            # 细则点(4)：首行缩进 2 字符
            #   优先读 w:firstLineChars（单位：字符×100），值应为 "200"；
            #   否则读 w:firstLine（twips），要求 twips == 2 * font_pt * 20。
            ok_indent = False
            cur_ind_desc = None
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                fl = ind.get(qn('w:firstLine'))
                cur_ind_desc = f'firstLineChars={flc}, firstLine={fl}'
                if flc is not None:
                    try:
                        ok_indent = int(flc) == 200
                    except ValueError:
                        pass
                if not ok_indent and fl is not None and para_cjk_size:
                    try:
                        target = 2.0 * para_cjk_size * 20.0
                        ok_indent = abs(int(fl) - target) < 1.0
                    except ValueError:
                        pass
            if not ok_indent:
                errors.append(f'"{snippet}" 首行缩进非2字符({cur_ind_desc})')

            # 细则点(5)：1.5 倍行距 (lineRule=auto, line=360)
            line_rule = spacing.get(qn('w:lineRule')) if spacing is not None else None
            line_val = spacing.get(qn('w:line')) if spacing is not None else None
            if not (line_rule == 'auto' and line_val == '360'):
                errors.append(f'"{snippet}" 行距非1.5倍(lineRule={line_rule}, line={line_val})')

            # 细则点(6)：两端对齐 (w:jc = both)
            jc_val = jc.get(qn('w:val')) if jc is not None else None
            if jc_val not in ('both', 'justify'):
                errors.append(f'"{snippet}" 非两端对齐(当前:{jc_val})')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 中文摘要页格式符合要求')
        else:
            self.score -= 1
            msg = '-1：中文摘要页不满足以下任意一项：标题下方文本为小四号宋体、首行缩进2字符、1.5倍行距，两端对齐'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:12]:
                print('      · ' + e)
            if len(seen) > 12:
                print(f'      · ...共{len(seen)}项')

    def check_english_abstract_title(self):
        '''检查英文摘要标题（扣分项）
        细则：-1：英文摘要页不满足以下任意一项：
              标题"Abstract"采用 Times New Roman、二号、加粗、居中

        逐点核对（沿 run→pPr→样式链→docDefaults 继承链，与 Office 渲染一致）：
          (1) 标题文本为"Abstract"；
          (2) 字体：Times New Roman（w:rFonts/w:ascii）；
          (3) 字号：二号 22pt（w:sz，单位为半磅）；
          (4) 加粗（w:b 继承链）；
          (5) 居中（w:jc = center 继承链，_effective_alignment）。
        '''
        # 定位 Abstract 标题
        title_para = None
        for para in self.doc.paragraphs:
            if para.text.strip() == 'Abstract':
                title_para = para
                break

        if title_para is None:
            print('  - 未找到英文摘要标题"Abstract"，跳过检查')
            return

        errors = []

        # 细则点(5)：居中
        if self._effective_alignment(title_para) != 'center':
            errors.append(f'"Abstract"未居中(当前:{self._effective_alignment(title_para)})')

        # 细则点(2)(3)(4)：字体 Times New Roman、字号 22pt、加粗——按 run 核对
        for run in title_para.runs:
            rt = run.text
            if not rt or not rt.strip():
                continue
            asc = self._effective_font(run, title_para, 'ascii')
            if asc != 'Times New Roman':
                errors.append(f'"Abstract"字体非Times New Roman(当前:{asc})')
            sz = self._effective_size(run, title_para)
            if sz is None or abs(sz - 22.0) >= 0.01:
                errors.append(f'"Abstract"字号非二号22pt(当前:{sz}pt)')
            if not self._effective_bold(run, title_para):
                errors.append('"Abstract"未加粗')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 英文摘要标题格式符合要求')
        else:
            self.score -= 1
            msg = '-1：英文摘要页不满足以下任意一项：标题"Abstract"采用 Times New Roman 、二号、加粗、居中'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:8]:
                print('      · ' + e)

    def check_english_abstract_content(self):
        '''检查英文摘要正文（扣分项）
        细则：-1：英文摘要页标题下方文本不满足以下任意一项：
              文本除"Key words"之外均采用 Times New Roman、小四号、1.5 倍行距、首行缩进1字符

        逐点核对（沿 run→pPr→样式链→docDefaults 继承链，与 Office 渲染一致）：
          (1) 目标范围：Abstract 标题下方到 "Key words" 之前的正文段落；
          (2) 字体：所有 run 的 ASCII 字体（w:rFonts/w:ascii）为 Times New Roman；
          (3) 字号：小四号 12pt（w:sz，单位为半磅）；
          (4) 行距 = 1.5 倍
              —— OOXML: w:pPr/w:spacing/@w:lineRule="auto" 且 @w:line="360"
                 （360/240 = 1.5）；段落 pPr 未设置时从段落样式链继承；
          (5) 首行缩进 = 1 个字符
              —— Office 存储为 w:pPr/w:ind/@w:firstLineChars="100"，
                 或 @w:firstLine（twips）等于 1 × 字号(pt) × 20；
                 段落 pPr 未设置时从段落样式链继承。
        '''
        # 沿 段落 pPr → 段落样式链 查找指定 pPr 子元素（继承查找）
        def _eff_pPr_child(para, child_local_name):
            child_qn = qn('w:' + child_local_name)
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                c = pPr.find(child_qn)
                if c is not None:
                    return c
            for style in self._iter_style_chain(para):
                el = getattr(style, 'element', None)
                if el is None:
                    continue
                pPr_s = el.find(qn('w:pPr'))
                if pPr_s is not None:
                    c = pPr_s.find(child_qn)
                    if c is not None:
                        return c
            return None

        # 定位 Abstract 标题
        title_idx = -1
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip() == 'Abstract':
                title_idx = i
                break

        if title_idx < 0:
            print('  - 未找到英文摘要标题，跳过检查')
            return

        # 收集 Abstract 标题下方到 "Key words" 之前的所有非空段落
        body_paras = []
        for j in range(title_idx + 1, len(self.doc.paragraphs)):
            p = self.doc.paragraphs[j]
            t = p.text.strip()
            if not t:
                continue
            if re.match(r'^Key\s*words', t, re.IGNORECASE):
                break
            body_paras.append(p)

        if not body_paras:
            print('  - 未找到英文摘要正文，跳过检查')
            return

        errors = []
        for para in body_paras:
            snippet = para.text.strip()[:28]

            # 细则点(2)(3)：Times New Roman、小四号 12pt——按 run 核对
            for run in para.runs:
                rt = run.text
                if not rt or not rt.strip():
                    continue
                asc = self._effective_font(run, para, 'ascii')
                if asc != 'Times New Roman':
                    errors.append(f'"{snippet}" 字体非Times New Roman(当前:{asc})')
                sz = self._effective_size(run, para)
                if sz is None or abs(sz - 12.0) >= 0.01:
                    errors.append(f'"{snippet}" 字号非小四号12pt(当前:{sz}pt)')

            # 段落主字号（用于换算 firstLine twips）
            para_size = None
            for run in para.runs:
                if run.text and run.text.strip():
                    para_size = self._effective_size(run, para)
                    if para_size:
                        break

            spacing = _eff_pPr_child(para, 'spacing')
            ind = _eff_pPr_child(para, 'ind')

            # 细则点(4)：1.5 倍行距 (lineRule=auto, line=360)
            line_rule = spacing.get(qn('w:lineRule')) if spacing is not None else None
            line_val = spacing.get(qn('w:line')) if spacing is not None else None
            if not (line_rule == 'auto' and line_val == '360'):
                errors.append(f'"{snippet}" 行距非1.5倍(lineRule={line_rule}, line={line_val})')

            # 细则点(5)：首行缩进 1 字符
            #   优先 w:firstLineChars="100"；否则 w:firstLine == 1 × font_pt × 20 twips
            ok_indent = False
            cur_ind_desc = None
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                fl = ind.get(qn('w:firstLine'))
                cur_ind_desc = f'firstLineChars={flc}, firstLine={fl}'
                if flc is not None:
                    try:
                        ok_indent = int(flc) == 100
                    except ValueError:
                        pass
                if not ok_indent and fl is not None and para_size:
                    try:
                        target = 1.0 * para_size * 20.0
                        ok_indent = abs(int(fl) - target) < 1.0
                    except ValueError:
                        pass
            if not ok_indent:
                errors.append(f'"{snippet}" 首行缩进非1字符({cur_ind_desc})')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 英文摘要正文格式符合要求')
        else:
            self.score -= 1
            msg = '-1：英文摘要页标题下方文本不满足以下任意一项：文本除"Key words"之外均采用Times New Roman 、小四号、1.5 倍行距、首行缩进1字符'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:12]:
                print('      · ' + e)
            if len(seen) > 12:
                print(f'      · ...共{len(seen)}项')

    def check_toc_title(self):
        '''检查目录标题格式（扣分项）
        细则：-1：目录页不满足以下任意一项：
              "目录"二字采用黑体二号、居中对齐、段后12磅，行距1.35倍

        逐点核对（沿 run→pPr→样式链→docDefaults 继承链，与 Office 渲染一致）：
          (1) 目标：正文中形如"目录"/"目 录"/"目  录"的段落；
          (2) 中文字体为黑体（w:rFonts/w:eastAsia）；
          (3) 字号为二号 22pt（w:sz，单位为半磅）；
          (4) 居中对齐（w:jc = center 继承链）；
          (5) 段后 12 磅
              —— OOXML: w:pPr/w:spacing/@w:after（单位：twentieths of a point）
                 值应为 "240"（12pt × 20）；@w:afterLines 或 @w:afterAutospacing
                 不采用；段落 pPr 未设置时从段落样式链继承；
          (6) 行距 1.35 倍
              —— OOXML: w:pPr/w:spacing/@w:lineRule="auto" 且 @w:line="324"
                 （324/240 = 1.35）；段落 pPr 未设置时从段落样式链继承。
        '''
        # 沿 段落 pPr → 段落样式链 查找指定 pPr 子元素
        def _eff_pPr_child(para, child_local_name):
            child_qn = qn('w:' + child_local_name)
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                c = pPr.find(child_qn)
                if c is not None:
                    return c
            for style in self._iter_style_chain(para):
                el = getattr(style, 'element', None)
                if el is None:
                    continue
                pPr_s = el.find(qn('w:pPr'))
                if pPr_s is not None:
                    c = pPr_s.find(child_qn)
                    if c is not None:
                        return c
            return None

        title_para = None
        for para in self.doc.paragraphs:
            if re.match(r'^目\s*录$', para.text.strip()):
                title_para = para
                break

        if title_para is None:
            print('  - 未找到"目录"标题，跳过检查')
            return

        errors = []

        # 细则点(2)(3)：中文黑体、字号二号 22pt——按 run 内 CJK 字符核对
        for run in title_para.runs:
            rt = run.text
            if not rt or not rt.strip():
                continue
            if re.search(r'[一-鿿]', rt):
                ea = self._effective_font(run, title_para, 'eastAsia')
                if not (ea and ('黑体' in ea or ea == 'SimHei')):
                    errors.append(f'"目录"中文字体非黑体(当前:{ea})')
            sz = self._effective_size(run, title_para)
            if sz is None or abs(sz - 22.0) >= 0.01:
                errors.append(f'"目录"字号非二号22pt(当前:{sz}pt)')

        # 细则点(4)：居中
        if self._effective_alignment(title_para) != 'center':
            errors.append(f'"目录"未居中(当前:{self._effective_alignment(title_para)})')

        spacing = _eff_pPr_child(title_para, 'spacing')

        # 细则点(5)：段后 12 磅（w:after == 240 twips，且未启用 afterAutospacing）
        after_val = spacing.get(qn('w:after')) if spacing is not None else None
        after_auto = spacing.get(qn('w:afterAutospacing')) if spacing is not None else None
        ok_after = False
        if after_val is not None and (after_auto in (None, '0', 'false')):
            try:
                ok_after = abs(int(after_val) - 240) < 1
            except ValueError:
                ok_after = False
        if not ok_after:
            errors.append(f'"目录"段后非12磅(after={after_val}, afterAutospacing={after_auto})')

        # 细则点(6)：行距 1.35 倍（lineRule=auto, line=324）
        line_rule = spacing.get(qn('w:lineRule')) if spacing is not None else None
        line_val = spacing.get(qn('w:line')) if spacing is not None else None
        if not (line_rule == 'auto' and line_val == '324'):
            errors.append(f'"目录"行距非1.35倍(lineRule={line_rule}, line={line_val})')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 目录标题格式符合要求')
        else:
            self.score -= 1
            msg = '-1：目录页不满足以下任意一项："目录"二字采用黑体二号、居中对齐、段后12磅，行距1.35倍'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:8]:
                print('      · ' + e)

    def check_body_text_format(self):
        """检查正文文本格式（扣分项）
        细则：-3：目录页之后除标题、图名及表题以外的文本不满足以下任意一项：
              (a) 中文文本为宋体小四号；
              (b) 英文、数字、公式说明英文部分为 Times New Roman 小四号；
              (c) 段落首行缩进 2 个中文字符；
              (d) 段前 0 行；
              (e) 段后 0 行；
              (f) 行距为 1.5 倍；
              (g) 两端对齐。

        逐点核对（沿 run→pPr→样式链→docDefaults 继承链，与 Office 渲染一致）：
          (a) 中文字符：w:rFonts/w:eastAsia == 宋体/SimSun；w:sz == 12pt；
          (b) 英文/数字字符：w:rFonts/w:ascii == Times New Roman；w:sz == 12pt；
          (c) 首行缩进：w:ind/@w:firstLineChars == "200"，或
              @w:firstLine（twips）等于 2 × 字号(pt) × 20；
          (d) 段前 = 0：w:spacing 的 @w:before/@w:beforeLines 均为 0 或缺省，
              且 @w:beforeAutospacing 不为 1；
          (e) 段后 = 0：@w:after/@w:afterLines 均为 0 或缺省，
              且 @w:afterAutospacing 不为 1；
          (f) 行距 1.5 倍：@w:lineRule == "auto" 且 @w:line == "360"；
          (g) 两端对齐：w:jc/@w:val ∈ {both, justify}。

        目标范围：从"目录页之后"开始，即找到"目录"标题后第一个章节标题
        （第X章/第X节）起，直到文档末尾。
        排除项：标题（含各级章节标题）、图名（^图...）、表题（^表...）。
        表格内文本不在 doc.paragraphs 中，自动排除。
        """
        # 沿 段落 pPr → 段落样式链 查找指定 pPr 子元素
        def _eff_pPr_child(para, child_local_name):
            child_qn = qn('w:' + child_local_name)
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                c = pPr.find(child_qn)
                if c is not None:
                    return c
            for style in self._iter_style_chain(para):
                el = getattr(style, 'element', None)
                if el is None:
                    continue
                pPr_s = el.find(qn('w:pPr'))
                if pPr_s is not None:
                    c = pPr_s.find(child_qn)
                    if c is not None:
                        return c
            return None

        # 定位"目录页之后"起点：找到"目录"标题后的第一个正文章节标题
        toc_idx = -1
        for i, para in enumerate(self.doc.paragraphs):
            if re.match(r'^目\s*录$', para.text.strip()):
                toc_idx = i
                break
        if toc_idx < 0:
            print('  - 未找到"目录"标题，跳过检查')
            return

        body_start = -1
        for j in range(toc_idx + 1, len(self.doc.paragraphs)):
            raw = self.doc.paragraphs[j].text
            t = raw.strip()
            # 目录条目形如"第1章 绪论\t1"——含制表符或末尾为页码数字，跳过
            if '\t' in raw:
                continue
            if re.search(r'\s\d+\s*$', t):
                continue
            if re.match(r'^第\s*[一二三四五六七八九十百零〇\d]+\s*[章节]', t):
                body_start = j
                break
        if body_start < 0:
            print('  - 未找到目录后的正文章节标题，跳过检查')
            return

        # 排除仍然属于目录条目的段落（含制表符或末尾页码数字）
        def _is_toc_entry(text_raw):
            t = text_raw.strip()
            if '\t' in text_raw:
                return True
            if re.search(r'\s\d+\s*$', t):
                return True
            return False

        # 标题识别：章节标题 / 二级三级四级标题 / Heading 样式
        TITLE_PAT = re.compile(
            r'^第\s*[一二三四五六七八九十百零〇\d]+\s*[章节]'
            r'|^[（(]\s*[一二三四五六七八九十]+\s*[)）]'
            r'|^\d+(?:\.\d+)+\s'
            r'|^\d+\s*[、.]\s'
        )
        FIG_PAT = re.compile(r'^图\s*\d+')
        TBL_PAT = re.compile(r'^表\s*\d+')

        def _is_title(para):
            style_name = (para.style.name if para.style else '') or ''
            if style_name.startswith('Heading'):
                return True
            if TITLE_PAT.match(para.text.strip()):
                return True
            return False

        def _spacing_zero(spacing_el, side):
            if spacing_el is None:
                return True  # 未设置视为 0
            v = spacing_el.get(qn('w:' + side))
            v_lines = spacing_el.get(qn('w:' + side + 'Lines'))
            auto = spacing_el.get(qn('w:' + side + 'Autospacing'))
            if auto in ('1', 'true'):
                return False
            if v_lines not in (None, '0'):
                return False
            if v not in (None, '0'):
                return False
            return True

        errors = []
        for para in self.doc.paragraphs[body_start:]:
            t = para.text.strip()
            if not t:
                continue
            if _is_title(para):
                continue
            if FIG_PAT.match(t):
                continue
            if TBL_PAT.match(t):
                continue

            snippet = t[:22]

            # 段落主字号（用于换算 firstLine twips）
            para_size = None
            for run in para.runs:
                if run.text and run.text.strip():
                    para_size = self._effective_size(run, para)
                    if para_size:
                        break

            # (a)(b) 按 run 内字符类型核对
            for run in para.runs:
                rt = run.text
                if not rt or not rt.strip():
                    continue
                size_pt = self._effective_size(run, para)
                has_cn = bool(re.search(r'[一-鿿]', rt))
                has_en = bool(re.search(r'[A-Za-z0-9]', rt))

                if has_cn:
                    ea = self._effective_font(run, para, 'eastAsia')
                    if not (ea and ('宋体' in ea or ea == 'SimSun')):
                        errors.append(f'"{snippet}" 中文非宋体(当前:{ea})')
                    if size_pt is None or abs(size_pt - 12.0) >= 0.01:
                        errors.append(f'"{snippet}" 中文非小四号12pt(当前:{size_pt}pt)')

                if has_en:
                    asc = self._effective_font(run, para, 'ascii')
                    if asc != 'Times New Roman':
                        errors.append(f'"{snippet}" 英文/数字非Times New Roman(当前:{asc})')
                    if size_pt is None or abs(size_pt - 12.0) >= 0.01:
                        errors.append(f'"{snippet}" 英文/数字非小四号12pt(当前:{size_pt}pt)')

            ind = _eff_pPr_child(para, 'ind')
            spacing = _eff_pPr_child(para, 'spacing')
            jc = _eff_pPr_child(para, 'jc')

            # (c) 首行缩进 2 中文字符
            ok_indent = False
            cur_ind_desc = None
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                fl = ind.get(qn('w:firstLine'))
                cur_ind_desc = f'firstLineChars={flc}, firstLine={fl}'
                if flc is not None:
                    try:
                        ok_indent = int(flc) == 200
                    except ValueError:
                        pass
                if not ok_indent and fl is not None and para_size:
                    try:
                        target = 2.0 * para_size * 20.0
                        ok_indent = abs(int(fl) - target) < 1.0
                    except ValueError:
                        pass
            if not ok_indent:
                errors.append(f'"{snippet}" 首行缩进非2字符({cur_ind_desc})')

            # (d)(e) 段前段后为 0
            if not _spacing_zero(spacing, 'before'):
                b_val = spacing.get(qn('w:before')) if spacing is not None else None
                b_lines = spacing.get(qn('w:beforeLines')) if spacing is not None else None
                b_auto = spacing.get(qn('w:beforeAutospacing')) if spacing is not None else None
                errors.append(f'"{snippet}" 段前非0(before={b_val}, beforeLines={b_lines}, beforeAutospacing={b_auto})')
            if not _spacing_zero(spacing, 'after'):
                a_val = spacing.get(qn('w:after')) if spacing is not None else None
                a_lines = spacing.get(qn('w:afterLines')) if spacing is not None else None
                a_auto = spacing.get(qn('w:afterAutospacing')) if spacing is not None else None
                errors.append(f'"{snippet}" 段后非0(after={a_val}, afterLines={a_lines}, afterAutospacing={a_auto})')

            # (f) 行距 1.5 倍
            line_rule = spacing.get(qn('w:lineRule')) if spacing is not None else None
            line_val = spacing.get(qn('w:line')) if spacing is not None else None
            if not (line_rule == 'auto' and line_val == '360'):
                errors.append(f'"{snippet}" 行距非1.5倍(lineRule={line_rule}, line={line_val})')

            # (g) 两端对齐
            jc_val = jc.get(qn('w:val')) if jc is not None else None
            if jc_val not in ('both', 'justify'):
                errors.append(f'"{snippet}" 非两端对齐(当前:{jc_val})')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 目录页之后正文文本格式符合要求')
        else:
            self.score -= 3
            msg = '-3：目录页之后除标题、图名及表题以外的文本不满足以下任意一项：中文文本为宋体小四号，英文、数字、公式说明英文部分为Times New Roman小四号；段落首行缩进2个中文字符，段前段后均为0行，行距为1.5倍，两端对齐。'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:14]:
                print('      · ' + e)
            if len(seen) > 14:
                print(f'      · ...共{len(seen)}项')

    def check_title_spacing(self):
        '''检查标题间距（扣分项）
        细则：-3：目录页后的所有页不满足以下任意一项：
              (1) 标题与标题之间空一行；
              (2) 标题与正文之间空一行。

        判定（针对 Office 实际渲染）：
          - "空一行"不等价于"恰好一个空白段落"——Word/WPS 中通过段前
            （before）/段后（after）间距设置等价的一行空白同样会在视觉
            上呈现为空一行，不应被误判为不合格。因此判定为以下任一即可：
              (i) 两个内容段之间存在恰好一个空白段落（无可见文本）；或
              (ii) 前一段的段后间距（w:after，按行）+ 后一段的段前间距
                   （w:before，按行）合计 ≈ 1 行（允许 0.9~1.1 行的容差，
                   兼容磅值换算与四舍五入误差）——按 twips 换算：
                   afterLines/beforeLines 单位为"1/100 行"，直接换算；
                   after/before（twips）按该段行距对应的"一行磅值"换算
                   为行数（缺省行距按单倍 12pt 估算）。
          - 检查范围：从目录之后的第一个真正章节标题起，到文档末尾。
          - 目录条目（含制表符或末尾页码数字）自动跳过。
          - "标题"识别：Heading 样式段，或匹配一至四级编号格式的段落：
              一级 第X章 / 第X节
              二级 （X）/(X)
              三级 X、
              四级 X.
              以及 X.X / X.X.X 编号（Heading 2/3 常用）。
          - 两个"内容段"若一个是标题，另一个是标题或正文，其间必须满足
            上述 (i) 或 (ii) 之一。
          - 找不到"目录"标题，或找不到目录后的正文章节起点时：无法验证
            本条细则是否满足，按扣分项语义处理——记为不满足，扣 3 分
            （而非跳过不扣分），并在 detail 中说明原因。
        '''
        # 定位目录
        toc_idx = -1
        for i, para in enumerate(self.doc.paragraphs):
            if re.match(r'^目\s*录$', para.text.strip()):
                toc_idx = i
                break
        if toc_idx < 0:
            self.score -= 3
            msg = '-3：目录页后的所有页不满足以下任意一项：标题与标题之间，标题与正文之间均空一行'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            print('      · 未找到"目录"标题，无法验证本条细则，按未满足处理')
            return

        # 定位目录之后第一个真正章节标题
        body_start = -1
        for j in range(toc_idx + 1, len(self.doc.paragraphs)):
            raw = self.doc.paragraphs[j].text
            t = raw.strip()
            if '\t' in raw:
                continue
            if re.search(r'\s\d+\s*$', t):
                continue
            if re.match(r'^第\s*[一二三四五六七八九十百零〇\d]+\s*[章节]', t):
                body_start = j
                break
        if body_start < 0:
            self.score -= 3
            msg = '-3：目录页后的所有页不满足以下任意一项：标题与标题之间，标题与正文之间均空一行'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            print('      · 未找到目录后的正文章节标题，无法验证本条细则，按未满足处理')
            return

        # 段前/段后间距换算为"行数"，用于判定 (ii)
        def _spacing_lines(para, side):
            '''返回段落 side（'before'/'after'）方向的间距，按"行数"估算。
            优先取 *Lines（单位=1/100行）；否则取磅值按该段行距换算。
            未设置时返回 0（视为该侧无间距）。
            '''
            pPr = para._element.find(qn('w:pPr'))
            spacing = pPr.find(qn('w:spacing')) if pPr is not None else None
            if spacing is None:
                return 0.0
            lines_val = spacing.get(qn('w:' + side + 'Lines'))
            if lines_val is not None:
                try:
                    return int(lines_val) / 100.0
                except ValueError:
                    pass
            pt_val = spacing.get(qn('w:' + side))
            if pt_val is None:
                return 0.0
            try:
                pt = int(pt_val) / 20.0  # twips -> pt
            except ValueError:
                return 0.0
            if pt <= 0:
                return 0.0
            # 用该段主字号对应的单倍行高估算"一行"磅值；取不到字号时按 12pt 估算
            line_rule = spacing.get(qn('w:lineRule'))
            line_val = spacing.get(qn('w:line'))
            if line_rule == 'auto' and line_val:
                try:
                    # w:line 240 = 单倍行距；单倍行距磅值约等于主字号 * 1.2
                    base_pt = None
                    for run in para.runs:
                        if run.text and run.text.strip():
                            base_pt = self._effective_size(run, para)
                            if base_pt:
                                break
                    base_pt = base_pt or 12.0
                    single_line_pt = base_pt * 1.2 * (int(line_val) / 240.0)
                except (ValueError, TypeError):
                    single_line_pt = 14.4
            else:
                single_line_pt = 14.4  # 12pt 字号单倍行距的近似磅值
            if single_line_pt <= 0:
                return 0.0
            return pt / single_line_pt

        def _has_gap_by_spacing(prev_para, cur_para):
            gap = _spacing_lines(prev_para, 'after') + _spacing_lines(cur_para, 'before')
            return 0.9 <= gap <= 1.1

        # 标题识别
        TITLE_PAT = re.compile(
            r'^第\s*[一二三四五六七八九十百零〇\d]+\s*[章节]'   # 一级
            r'|^[（(]\s*[一二三四五六七八九十]+\s*[)）]'         # 二级
            r'|^\d+(?:\.\d+)+\s'                                # X.X / X.X.X 编号
            r'|^\d+\s*[、.]\s'                                  # 三级 X、 / 四级 X.
        )

        def _is_title(para):
            style_name = (para.style.name if para.style else '') or ''
            if style_name.startswith('Heading'):
                return True
            if TITLE_PAT.match(para.text.strip()):
                return True
            return False

        # 收集目录后所有段落的三元组：(idx_in_paragraphs, is_empty, is_title, text)
        paras = self.doc.paragraphs
        items = []
        for i in range(body_start, len(paras)):
            raw = paras[i].text
            t = raw.strip()
            is_empty = (t == '')
            # 目录条目在正文段不应出现；一旦仍出现（如附录里的"参考文献\t18"），
            # 视为异常段跳过（不作为标题也不作为正文邻接）。
            if not is_empty and ('\t' in raw or re.search(r'\s\d+\s*$', t)):
                # 附录/参考文献目录条目——不参与判定
                continue
            items.append((i, is_empty, (not is_empty and _is_title(paras[i])), t))

        # 遍历相邻的两个非空段：如果两者中有任一是标题，其间必须满足
        # (i) 恰好 1 个空段，或 (ii) 段前/段后间距合计约 1 行
        errors = []
        prev = None  # 上一个非空段的记录
        prev_gap = 0  # 上一个非空段与当前非空段之间的空段数
        prev_para_el = None  # 上一个非空段对应的 Paragraph 对象
        for rec in items:
            i, is_empty, is_title, t = rec
            if is_empty:
                if prev is not None:
                    prev_gap += 1
                continue
            # 遇到当前非空段
            cur_para_el = paras[i]
            if prev is not None:
                p_i, _, p_title, p_text = prev
                # 只要相邻两非空段中至少一个是标题，就必须满足 (i) 或 (ii)
                if p_title or is_title:
                    gap_ok = (prev_gap == 1)
                    if not gap_ok and prev_gap == 0 and prev_para_el is not None:
                        gap_ok = _has_gap_by_spacing(prev_para_el, cur_para_el)
                    if not gap_ok:
                        label_prev = p_text[:22]
                        label_cur = t[:22]
                        kind = ('标题-标题' if (p_title and is_title)
                                else ('标题-正文' if p_title else '正文-标题'))
                        errors.append(f'{kind}未空一行(gap={prev_gap}): "{label_prev}" → "{label_cur}"')
            prev = rec
            prev_gap = 0
            prev_para_el = cur_para_el

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 标题间距符合要求')
        else:
            self.score -= 3
            msg = '-3：目录页后的所有页不满足以下任意一项：标题与标题之间，标题与正文之间均空一行'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:14]:
                print('      · ' + e)
            if len(seen) > 14:
                print(f'      · ...共{len(seen)}项')

    def check_table_cell(self):
        '''检查表格单元格文本是否只有单独一行（扣分项）
        细则：-3：目录页后的所有页表格的单元格不满足此项：其中的文本只有单独一行。

        判定（针对 Office 实际渲染）：
          - 范围：仅检查"目录"之后出现在正文中的表格；目录前的封面、
            摘要、目录等不参与检查。
          - "只有单独一行"在 OOXML 中的结构含义：
              (1) 单元格内的非空段落（`w:p`）数量 ≤ 1；出现两个及以上
                  非空段落 → Word/WPS 会视为多行；
              (2) 单元格内所有 run 中不得含有硬换行 `w:br`（换行符），
                  否则 Office 里会显示为多行。
          - 空单元格不计入判定。
          - 严格二元：任一单元格违反 (1) 或 (2)，即整项失败扣 3 分。
        '''
        # 定位目录段
        toc_para_el = None
        for para in self.doc.paragraphs:
            if re.match(r'^目\s*录$', para.text.strip()):
                toc_para_el = para._element
                break
        if toc_para_el is None:
            print('  - 未找到"目录"标题，跳过检查')
            return

        # 收集目录之后 body 中出现的所有表格（保留文档顺序）
        body = self.doc.element.body
        after_toc = False
        tables_after_toc = []
        table_qn = qn('w:tbl')
        for child in body.iterchildren():
            if child is toc_para_el:
                after_toc = True
                continue
            if not after_toc:
                continue
            if child.tag == table_qn:
                tables_after_toc.append(child)

        if not tables_after_toc:
            print('  - 目录之后未发现表格，跳过检查')
            return

        # 逐单元格审查
        violations = []
        checked_cells = 0
        tc_qn = qn('w:tc')
        p_qn = qn('w:p')
        br_qn = qn('w:br')

        for t_idx, tbl_el in enumerate(tables_after_toc, start=1):
            # findall 所有 tc，包括嵌套（保守：仍统一检查）
            for tc in tbl_el.iter(tc_qn):
                # 单元格内所有直接子 w:p
                ps = [p for p in tc.iterchildren() if p.tag == p_qn]
                # 非空段落数量
                non_empty_ps = 0
                has_br = False
                cell_text_sample = ''
                for p in ps:
                    # 收集文本
                    txt_parts = []
                    for t in p.iter(qn('w:t')):
                        if t.text:
                            txt_parts.append(t.text)
                    text = ''.join(txt_parts).strip()
                    if text:
                        non_empty_ps += 1
                        if not cell_text_sample:
                            cell_text_sample = text
                    # 硬换行 w:br（不含 type=page 分页符）
                    for br in p.iter(br_qn):
                        br_type = br.get(qn('w:type'))
                        if br_type in (None, 'textWrapping'):
                            has_br = True

                if non_empty_ps == 0:
                    continue  # 空单元格不计入
                checked_cells += 1

                if non_empty_ps > 1:
                    violations.append(
                        f'表{t_idx} 单元格含 {non_empty_ps} 个非空段落: "{cell_text_sample[:18]}"'
                    )
                elif has_br:
                    violations.append(
                        f'表{t_idx} 单元格含硬换行(w:br): "{cell_text_sample[:18]}"'
                    )

        if not violations:
            print(f'  ✓ 表格单元格格式符合要求（共检查 {checked_cells} 个非空单元格）')
        else:
            self.score -= 3
            msg = '-3：目录页后的所有页表格的单元格不满足此项：其中的文本只有单独一行'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for v in violations[:14]:
                print('      · ' + v)
            if len(violations) > 14:
                print(f'      · ...共{len(violations)}项')

    def check_code_indent(self):
        '''检查代码段缩进（扣分项）
        细则：-1：程序代码段不满足以下任意一项：
              (1) 首行缩进2字符；
              (2) 文本之前无缩进（左侧无缩进）；
              (3) 文本之后无缩进（右侧无缩进）。

        判定（针对 Office 实际渲染）：
          - 代码段识别与 check_code_format 完全一致（同一强特征正则）。
          - OOXML: 段落的缩进属性位于 w:pPr/w:ind，各属性支持
              @firstLineChars / @firstLine        —— 首行缩进
              @leftChars     / @left  / @startChars / @start  —— 左缩进
              @rightChars    / @right / @endChars  / @end     —— 右缩进
            其中 *Chars 值为字符×100（200 = 2 字符）；非 Chars 属性单位为 twip。
          - 首行缩进 2 字符：
              firstLineChars == 200， 或
              firstLine (twip) == 2 × 有效字号(pt) × 20
              （字号沿 run→pPr→样式链→docDefaults 继承；代码段字号五号=10.5pt 时
               2 字符 = 2×10.5×20 = 420 twip）
          - 左/右缩进无：所有 leftChars/left/startChars/start（右侧同理）均缺失或为 0。
          - w:ind 若在段落 pPr 中缺失，需沿样式链继承（与 Office 一致）。
          - 严格二元：任一代码段任一点违反即扣 1 分。
        '''
        # 代码段识别（与 check_code_format 保持一致）
        code_pat = re.compile(
            r'^\s*(?:def|class|import|from|return|if|elif|else|for|while|try|except|finally|with)\b'
            r'|^\s*[A-Za-z_]\w*\s*=\s*[A-Za-z_][\w\.]*\s*\('
            r'|\.(?:withColumn|groupBy|createOrReplaceTempView|agg|filter|select|orderBy|show|write|read|join|drop|na|fillna|dropna|distinct|toDF|when|otherwise)\s*\('
            r'|^\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|WITH)\s'
        )

        code_paras = []
        for para in self.doc.paragraphs:
            t = para.text
            if not t.strip():
                continue
            if code_pat.search(t):
                code_paras.append(para)

        if not code_paras:
            print('  - 未找到代码段，跳过检查')
            return

        def _eff_ind(para):
            '''沿段落 pPr → 样式链 查找 w:ind，返回首个命中的元素。'''
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    return ind
            for style in self._iter_style_chain(para):
                el = getattr(style, 'element', None)
                if el is None:
                    continue
                pPr_s = el.find(qn('w:pPr'))
                if pPr_s is not None:
                    ind = pPr_s.find(qn('w:ind'))
                    if ind is not None:
                        return ind
            return None

        def _para_size_pt(para):
            '''取段落首个非空 run 的有效字号（pt），用于将 firstLine (twip) 折算成字符数。'''
            for run in (para.runs or []):
                if run.text and run.text.strip():
                    sp = self._effective_size(run, para)
                    if sp is not None:
                        return sp
            return 10.5  # 代码段默认五号

        errors = []
        for para in code_paras:
            snippet = para.text.strip()[:32]
            ind = _eff_ind(para)

            # ---------- 点(1)：首行缩进 2 字符 ----------
            firstLineChars = ind.get(qn('w:firstLineChars')) if ind is not None else None
            firstLine = ind.get(qn('w:firstLine')) if ind is not None else None
            ok_first = False
            if firstLineChars is not None:
                try:
                    ok_first = (int(firstLineChars) == 200)
                except ValueError:
                    ok_first = False
            if not ok_first and firstLine is not None:
                try:
                    fl_twip = int(firstLine)
                    size_pt = _para_size_pt(para)
                    expected = int(round(2 * size_pt * 20))
                    ok_first = (abs(fl_twip - expected) <= 1)
                except ValueError:
                    ok_first = False
            if not ok_first:
                errors.append(
                    f'代码段首行未缩进2字符(firstLineChars={firstLineChars}, firstLine={firstLine}): "{snippet}"'
                )

            # ---------- 点(2)：文本之前无缩进（左侧） ----------
            def _zero_or_absent(el, attr_names):
                if el is None:
                    return True
                for a in attr_names:
                    v = el.get(qn('w:' + a))
                    if v is not None:
                        try:
                            if int(v) != 0:
                                return False
                        except ValueError:
                            return False
                return True

            ok_left = _zero_or_absent(ind, ['left', 'leftChars', 'start', 'startChars'])
            if not ok_left:
                vals = {a: (ind.get(qn('w:' + a)) if ind is not None else None)
                        for a in ('left', 'leftChars', 'start', 'startChars')}
                errors.append(f'代码段文本之前有缩进{vals}: "{snippet}"')

            # ---------- 点(3)：文本之后无缩进（右侧） ----------
            ok_right = _zero_or_absent(ind, ['right', 'rightChars', 'end', 'endChars'])
            if not ok_right:
                vals = {a: (ind.get(qn('w:' + a)) if ind is not None else None)
                        for a in ('right', 'rightChars', 'end', 'endChars')}
                errors.append(f'代码段文本之后有缩进{vals}: "{snippet}"')

        seen = []
        for e in errors:
            if e not in seen:
                seen.append(e)

        if not seen:
            print('  ✓ 代码段缩进符合要求')
        else:
            self.score -= 1
            msg = '-1：程序代码段不满足以下任意一项：首行缩进2字符，文本之前和文本之后无缩进'
            self.failed_points.append(msg)
            print('  ✗ ' + msg)
            for e in seen[:12]:
                print('      · ' + e)
            if len(seen) > 12:
                print(f'      · ...共{len(seen)}项')

    def print_result(self):
        '''打印评估结果'''
        print("\n" + "=" * 60)
        print("【评估结果汇总】")
        print("=" * 60)

        print(f"\n最终得分: {self.score} 分")

        if self.checked_points:
            print("\n得分/扣分明细:")
            for point in self.checked_points:
                print(f"  {point}")

        if self.failed_points:
            print("\n扣分项:")
            for item in self.failed_points:
                print(f"  {item}")

        print("\n" + "=" * 60)


def evaluate(dir_path: str) -> dict:
    '''统一入口函数（对齐"评估脚本接口差异与统一建议.md" §2.1）。

    参数：
        dir_path: **脚本所在目录的路径**（例如 ".../officeval_013"）。
                  本函数自行在该目录中定位待评估的 .docx 文档。

    返回：
        结构化 dict（字段见 §2.2）：包含 id、file_name、status、error、
        dim1_pass、dim1_reason、dim2_items、total_score、max_score。
        脚本自身崩溃 / 文档不存在等错误场景返回 status="error"。
    '''
    def _err(msg, file_name=''):
        return {
            'id': '013',
            'file_name': file_name,
            'status': 'error',
            'error': msg,
            'dim1_pass': False,
            'dim1_reason': msg,
            'dim2_items': [],
            'total_score': 0,
            'max_score': sum(m for _, _, m in DocumentEvaluator._DIM2_ADD_RULES),
        }

    try:
        # 允许直接把文档路径传进来，兼容误用
        if os.path.isfile(dir_path):
            target_path = dir_path
        else:
            if not os.path.isdir(dir_path):
                return _err(f'目录不存在或不可访问: {dir_path}')
            # 在给定目录中扫描 .docx，忽略 Office 临时文件（~$ 前缀）
            candidates = []
            for name in os.listdir(dir_path):
                if name.startswith('~$'):
                    continue
                ext = os.path.splitext(name)[1].lower()
                if ext == '.docx':
                    candidates.append(os.path.join(dir_path, name))
            if not candidates:
                return _err(f'目录中未找到 .docx 文档: {dir_path}')
            # 若存在多个，取文件名最长的（通常是完整论文而非片段）
            candidates.sort(key=lambda p: len(os.path.basename(p)), reverse=True)
            target_path = candidates[0]

        evaluator = DocumentEvaluator(target_path)
        return evaluator.evaluate()
    except Exception as e:  # 兜底：任何异常都转成 status=error，避免影响批量运行器
        file_name = ''
        try:
            file_name = os.path.basename(target_path)  # type: ignore[name-defined]
        except Exception:
            pass
        return _err(f'{type(e).__name__}: {e}', file_name=file_name)


if __name__ == "__main__":
    # 本地调试入口：接收"脚本所在目录路径"，缺省则用脚本自身所在目录。
    # 主结果通过 stdout 直接输出 UTF-8 字节，避免修改全局 sys.stdout。
    import json
    arg = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    result = evaluate(arg)
    payload = (json.dumps(result, ensure_ascii=False, indent=2) + '\n').encode('utf-8')
    sys.stdout.buffer.write(payload)
