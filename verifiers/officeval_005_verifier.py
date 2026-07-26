#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
综合设计指导书_格式修改版.docx 自动评估脚本

评估逻辑：
  维度一：可用与可修改性 — 任意一条不满足直接判 0 分，不再检查维度二。
  维度二：完成度 — 各项得分/扣分累计。
"""

import os
import sys
import json
import zipfile
import re
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 本脚本编号（用于统一返回结构中的 "id" 字段）
SCRIPT_ID = "005"


# ============================================================
# 工具函数
# ============================================================

def emu_to_cm(emu_val):
    """EMU 转厘米 (1 cm = 360000 EMU)"""
    if emu_val is None:
        return None
    try:
        return float(emu_val) / 360000.0
    except Exception:
        return None


def twip_to_cm(twip_val):
    """twip 转厘米 (1 cm = 567 twips)"""
    if twip_val is None:
        return None
    try:
        return float(twip_val) / 567.0
    except Exception:
        return None


def twip_to_pt(twip_val):
    """twip 转磅 (1 pt = 20 twips)"""
    if twip_val is None:
        return None
    try:
        return float(twip_val) / 20.0
    except Exception:
        return None


def half_pt_to_pt(hp):
    """半点转磅"""
    if hp is None:
        return None
    try:
        return float(hp) / 2.0
    except Exception:
        return None


def get_run_font(run):
    """获取 run 的字体信息（字体名、字号、加粗）"""
    rPr = run._element.find(qn("w:rPr"))
    font_name = None
    size_pt = None
    bold = False

    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            # ascii / eastAsia / hAnsi 任何一个非空
            for attr in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
                v = rFonts.get(qn(attr))
                if v:
                    font_name = v
                    break
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            size_pt = half_pt_to_pt(sz.get(qn("w:val")))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is not None and size_pt is None:
            size_pt = half_pt_to_pt(szCs.get(qn("w:val")))
        b = rPr.find(qn("w:b"))
        if b is not None and b.get(qn("w:val")) != "0":
            bold = True
    return font_name, size_pt, bold


def _get_style_element(doc, style_id):
    """从 styles part 查找指定 styleId 的 <w:style> 元素"""
    if not style_id:
        return None
    try:
        styles_el = doc.part.styles.element
    except Exception:
        return None
    for st in styles_el.findall(qn("w:style")):
        if st.get(qn("w:styleId")) == style_id:
            return st
    return None


def _resolve_spacing_from_style_chain(doc, style_id):
    """
    沿 pStyle → basedOn 链回溯查找 <w:spacing> 属性，返回 dict：
      {'after': twip_str, 'before': twip_str, 'line': twip_str,
       'lineRule': str, 'beforeLines': str, 'afterLines': str}
    某一层已提供的属性优先（子样式覆盖父样式）。全部检查过后仍缺失的键不出现。
    """
    resolved = {}
    seen = set()
    cur_id = style_id
    while cur_id and cur_id not in seen:
        seen.add(cur_id)
        st = _get_style_element(doc, cur_id)
        if st is None:
            break
        st_pPr = st.find(qn("w:pPr"))
        if st_pPr is not None:
            sp = st_pPr.find(qn("w:spacing"))
            if sp is not None:
                for attr in ("after", "before", "line", "lineRule",
                             "beforeLines", "afterLines"):
                    if attr in resolved:
                        continue
                    v = sp.get(qn("w:" + attr))
                    if v is not None:
                        resolved[attr] = v
        basedOn = st.find(qn("w:basedOn"))
        cur_id = basedOn.get(qn("w:val")) if basedOn is not None else None
    return resolved


def _resolve_spacing_from_doc_defaults(doc):
    """从 <w:docDefaults>/<w:pPrDefault>/<w:pPr>/<w:spacing> 读取属性"""
    try:
        styles_el = doc.part.styles.element
    except Exception:
        return {}
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        return {}
    pPrDefault = dd.find(qn("w:pPrDefault"))
    if pPrDefault is None:
        return {}
    default_pPr = pPrDefault.find(qn("w:pPr"))
    if default_pPr is None:
        return {}
    sp = default_pPr.find(qn("w:spacing"))
    if sp is None:
        return {}
    out = {}
    for attr in ("after", "before", "line", "lineRule",
                 "beforeLines", "afterLines"):
        v = sp.get(qn("w:" + attr))
        if v is not None:
            out[attr] = v
    return out


def get_para_format(para):
    """获取段落格式信息：行距、段后、段前、对齐、首行缩进等
    行距 / 段前 / 段后 若段落自身未设置，则沿 pStyle → basedOn 链回溯，
    最终回退到 docDefaults。这样与 Word/WPS 的实际渲染值保持一致。
    """
    info = {
        "alignment": None,
        "line_rule": None,
        "line": None,        # 倍数或磅
        "space_before": None,
        "space_after": None,
        "first_line_indent_cm": None,
        "left_indent_cm": None,
        "hanging_indent_cm": None,
        "first_line_pos_cm": None,
    }
    pf = para.paragraph_format
    # 对齐
    try:
        align = pf.alignment
        if align is not None:
            info["alignment"] = str(align).split(".")[-1]
    except Exception:
        pass
    # 行距
    try:
        line_spacing = pf.line_spacing
        line_spacing_rule = pf.line_spacing_rule
        if line_spacing is not None:
            info["line"] = float(line_spacing)
        if line_spacing_rule is not None:
            info["line_rule"] = str(line_spacing_rule).split(".")[-1]
    except Exception:
        pass
    # 段前 / 段后
    try:
        sb = pf.space_before
        sa = pf.space_after
        if sb is not None:
            info["space_before"] = float(sb.pt)
        if sa is not None:
            info["space_after"] = float(sa.pt)
    except Exception:
        pass

    # ---- 缺失字段沿样式继承链回退 ----
    # 段落自身若未设置 line/space_before/space_after，则查 pStyle → basedOn
    # → docDefaults，与 Word/WPS 显示值保持一致。
    need_line = info["line"] is None
    need_sb = info["space_before"] is None
    need_sa = info["space_after"] is None
    if need_line or need_sb or need_sa:
        try:
            doc = para.part.document
        except Exception:
            doc = None
        try:
            pPr = para._p.find(qn("w:pPr"))
            pStyle = pPr.find(qn("w:pStyle")) if pPr is not None else None
            style_id = pStyle.get(qn("w:val")) if pStyle is not None else None
            # 若段落未显式引用样式，Word 默认使用 default=paragraph 样式
            if style_id is None and doc is not None:
                try:
                    styles_el = doc.part.styles.element
                    for st in styles_el.findall(qn("w:style")):
                        if (st.get(qn("w:type")) == "paragraph"
                                and st.get(qn("w:default")) == "1"):
                            style_id = st.get(qn("w:styleId"))
                            break
                except Exception:
                    pass

            resolved = {}
            if doc is not None and style_id:
                resolved.update(_resolve_spacing_from_style_chain(doc, style_id))
            if doc is not None:
                dd_spacing = _resolve_spacing_from_doc_defaults(doc)
                # docDefaults 只填样式链未提供的属性
                for k, v in dd_spacing.items():
                    resolved.setdefault(k, v)

            # 行距：line + lineRule
            if need_line and "line" in resolved:
                try:
                    line_val_twips = int(resolved["line"])
                    line_rule = resolved.get("lineRule", "auto")
                    if line_rule == "auto":
                        info["line"] = line_val_twips / 240.0
                        if info["line_rule"] is None:
                            info["line_rule"] = "MULTIPLE"
                    elif line_rule in ("exact", "atLeast"):
                        # 磅值（用 pt）
                        info["line"] = line_val_twips / 20.0
                        if info["line_rule"] is None:
                            info["line_rule"] = ("EXACTLY" if line_rule == "exact"
                                                 else "AT_LEAST")
                except Exception:
                    pass
            # 段前 / 段后：twip → pt
            if need_sb and "before" in resolved:
                try:
                    info["space_before"] = int(resolved["before"]) / 20.0
                except Exception:
                    pass
            if need_sa and "after" in resolved:
                try:
                    info["space_after"] = int(resolved["after"]) / 20.0
                except Exception:
                    pass
        except Exception:
            pass

    # 缩进
    try:
        fli = pf.first_line_indent
        if fli is not None:
            info["first_line_indent_cm"] = emu_to_cm(fli.emu)
        li = pf.left_indent
        if li is not None:
            info["left_indent_cm"] = emu_to_cm(li.emu)
    except Exception:
        pass
    # 从底层 XML 中读取悬挂缩进（firstLine / hanging）
    try:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                if ind.get(qn("w:hanging")) is not None:
                    info["hanging_indent_cm"] = twip_to_cm(ind.get(qn("w:hanging")))
                if ind.get(qn("w:hangingChars")) is not None:
                    pass
    except Exception:
        pass
    # 首行相对页边距的绝对缩进位置（即 Word/WPS 段落对话框里显示的
    # “首行缩进/段前缩进”数值）。当只设置 w:hanging（无 w:firstLine）时，
    # python-docx 的 first_line_indent 返回的是相对 left_indent 的负偏移
    # （-hanging），并不是 WPS 对话框展示的绝对值；真正的绝对首行位置是
    # left_indent - hanging_indent。当同时不存在 hanging 时，绝对位置就
    # 等于 first_line_indent_cm 本身（此时二者语义一致）。
    try:
        li_cm = info.get("left_indent_cm")
        hi_cm = info.get("hanging_indent_cm")
        if li_cm is not None and hi_cm is not None:
            info["first_line_pos_cm"] = li_cm - hi_cm
        elif info.get("first_line_indent_cm") is not None:
            info["first_line_pos_cm"] = info["first_line_indent_cm"]
    except Exception:
        pass
    return info


def get_page_setup(doc):
    """获取页面设置：纸张大小、方向、页边距"""
    sect = doc.sections[0]
    pw = sect.page_width
    ph = sect.page_height
    orient = sect.orientation
    margins = {
        "left": sect.left_margin,
        "right": sect.right_margin,
        "top": sect.top_margin,
        "bottom": sect.bottom_margin,
    }
    return {
        "page_width_cm": emu_to_cm(pw) if pw else None,
        "page_height_cm": emu_to_cm(ph) if ph else None,
        "orientation": str(orient) if orient is not None else None,
        "margins_cm": {k: emu_to_cm(v) for k, v in margins.items()},
    }


def get_paragraph_text(para):
    return "".join(run.text for run in para.runs)


def get_all_paragraphs(doc):
    """返回所有段落（body + 表格 + 页眉页脚）"""
    return list(doc.paragraphs)


def has_continuous_blank_pages(doc, threshold=2):
    """检测连续空白页：粗略判断连续空段落数 * 段落占用行数 >= 阈值页"""
    body_paragraphs = list(doc.paragraphs)
    # 统计连续空白段落（无内容或仅空白字符）
    blank_streak = 0
    max_blank_streak = 0
    for p in body_paragraphs:
        text = p.text.strip()
        if not text:
            blank_streak += 1
            max_blank_streak = max(max_blank_streak, blank_streak)
        else:
            blank_streak = 0
    # 近似：~3 个空段 = 1 页
    estimated_blank_pages = max_blank_streak / 3.0
    return estimated_blank_pages >= threshold, max_blank_streak


def count_runs_with_text(doc):
    return sum(1 for p in doc.paragraphs if p.text.strip())


def detect_mojibake(doc):
    """检测乱码：是否含大量不可识别字符"""
    weird_pattern = re.compile(r"[�-]")
    count = 0
    for p in doc.paragraphs:
        if weird_pattern.search(p.text):
            count += 1
    return count


def get_body_paragraphs_with_page_index(doc):
    """近似为每个段落计算页码（用分节/分页符）"""
    pages = []
    cur = 1
    for p in doc.paragraphs:
        pages.append(cur)
        # 遇到分页符
        for run in p.runs:
            for br in run._element.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    cur += 1
    return pages


def get_headers_texts(doc):
    """获取所有 section 的页眉文本"""
    out = []
    for sec_idx, section in enumerate(doc.sections):
        header = section.header
        texts = []
        for p in header.paragraphs:
            t = p.text.strip()
            if t:
                texts.append(t)
        # different_first_page_header
        try:
            fp = section.first_page_header
            for p in fp.paragraphs:
                t = p.text.strip()
                if t:
                    texts.append(("first", t))
        except Exception:
            pass
        out.append((sec_idx, texts))
    return out


def get_header_borders(doc):
    """检查页眉下方的横向单线（pBdr 底部边框）"""
    results = []
    for sec_idx, section in enumerate(doc.sections):
        header = section.header
        has_line = False
        for p in header.paragraphs:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None:
                continue
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                continue
            bottom = pBdr.find(qn("w:bottom"))
            if bottom is None:
                continue
            if bottom.get(qn("w:val")) in ("single", "double"):
                has_line = True
        results.append((sec_idx, has_line))
    return results


def get_footers_texts(doc):
    """获取页脚文本（包含页码）"""
    out = []
    for sec_idx, section in enumerate(doc.sections):
        footer = section.footer
        texts = []
        for p in footer.paragraphs:
            t = p.text.strip()
            if t:
                texts.append(t)
        # 提取页码（fldChar）
        has_field = False
        for p in footer.paragraphs:
            for el in p._p.iter():
                if el.tag == qn("w:instrText") and "PAGE" in (el.text or "").upper():
                    has_field = True
        out.append((sec_idx, texts, has_field))
    return out


# ============================================================
# 维度一：可用与可修改性
# ============================================================

def check_dimension_one(doc, docx_path):
    """维度一检查，全部满足才继续"""
    issues = []
    ok = True

    # 1) 可编辑 Word 文档、格式、文件可正常打开
    # python-docx 打开没抛异常 = 满足
    if not os.path.exists(docx_path):
        issues.append("文件不存在")
        return False, issues
    try:
        with zipfile.ZipFile(docx_path) as z:
            names = z.namelist()
            if "word/document.xml" not in names:
                issues.append("文件结构不合法（缺 document.xml）")
                return False, issues
    except zipfile.BadZipFile:
        issues.append("文件不是有效的 zip/docx")
        return False, issues

    # 2) 无连续 2 页以上空白页
    too_many_blanks, blank_streak = has_continuous_blank_pages(doc)
    if too_many_blanks:
        issues.append(f"存在连续空白段落 (≈{blank_streak} 个空段) → 估 ≥2 页空白")
        ok = False

    # 3) 无超过 1/3 页面面积乱码 / 文字重叠 — 简化为：乱码 run 比例 < 30% 且无重叠
    total = 0
    weird = 0
    for p in doc.paragraphs:
        for r in p.runs:
            t = r.text
            if t.strip():
                total += len(t)
                weird += len(re.findall(r"[�-]", t))
    if total > 0 and weird / total > 0.3:
        issues.append(f"乱码字符占比 {weird/total:.1%} > 1/3")
        ok = False

    # 4) 保留主要结构（封面/说明/目录/正文/结论/参考资料/附录）
    text_all = "\n".join(p.text for p in doc.paragraphs)
    required_keywords = ["封面", "目录", "结论", "附录"]
    missing = [k for k in required_keywords if k not in text_all]
    if missing:
        issues.append(f"主要结构缺失关键字: {missing}")
        ok = False

    return ok, issues


# ============================================================
# 维度二：完成度评估
# ============================================================

def pt_approx(v, target, tol=1.5):
    if v is None:
        return False
    return abs(v - target) <= tol


def cm_approx(v, target, tol=0.2):
    if v is None:
        return False
    return abs(v - target) <= tol


def is_fixed_line_22pt(info):
    """行距为固定值 22 磅"""
    rule = info.get("line_rule")
    line = info.get("line")
    if rule is None or line is None:
        return False
    # EXACTLY / FIXED 规则：line_spacing 返回的是 EMU 值
    # 1 pt = 12700 EMU, 22 pt = 279400 EMU
    if "EXACT" in rule.upper() or "FIXED" in rule.upper():
        return abs(line - 279400) <= 12700  # ±1pt 容差
    return False


# ------------------------------------------------------------------
# 视觉判定：把 .docx 用 WPS 渲染成 PDF，再用 pdfplumber（经 pdf_backend 适配层）度量
#   - 每页可见正文行数（去掉页眉页脚区域）
#   - 每行除标点外字符数（视觉行 = PDF 中 y 坐标聚合后的行）
# 供 eval_line_spacing 的"每页最多32行 / 每行最多34字"两条使用。
# ------------------------------------------------------------------

# 中英标点集合（判定"除标点外"的字符数）
_PUNCT_SET = set("，。！？；：、“”‘’（）《》【】"
                 "…—·,.!?;:'\"()[]<>-–—…" "「」『』〈〉〖〗")


def _is_cjk(ch):
    """CJK 汉字（含扩展 A/B/兼容区）——按 Word 字符网格算 1 格"""
    cp = ord(ch)
    return (
        0x4E00 <= cp <= 0x9FFF or   # CJK Unified Ideographs
        0x3400 <= cp <= 0x4DBF or   # Ext A
        0x20000 <= cp <= 0x2A6DF or # Ext B
        0xF900 <= cp <= 0xFAFF      # Compatibility
    )


def _count_grid_chars(text):
    """
    按 Word 字符网格语义"数字"：
      - CJK 汉字     → 1 格
      - ASCII 字母/数字 → 0.5 格（半角占半格）
      - 中文标点/空白  → 0（不计入"除标点外"字数）
      - 其它符号（如 - / ~）→ 0.5 格（视为半角）
    返回 float 值。
    """
    n = 0.0
    for ch in text:
        if ch.isspace():
            continue
        if ch in _PUNCT_SET:
            continue
        if _is_cjk(ch):
            n += 1.0
        else:
            # ASCII 字母/数字，或其它半角符号
            n += 0.5
    return n


def _wrap_paragraph_lines(para, usable_pt, default_size_pt=12.0):
    """
    根据 Word 属性模拟折行，返回 [(cells, width_pt), ...]，每项代表可视行。
      - CJK 汉字宽度 ≈ font_size pt（1 em）
      - 中文全角标点宽度 ≈ font_size pt
      - ASCII/半角字符宽度 ≈ font_size * 0.5 pt
      - "cells"：CJK=1、半角非标点=0.5、标点=0、空白=0（对应"除标点外字数"）
      - 硬换行 (w:br 非 page) 起新行；page break 结束该段
      - 首行缩进：优先用 pPr/ind/@w:firstLine（twip→pt），
                  否则用 pPr/ind/@w:firstLineChars（1/100 字符 × 段落主字号）
    """
    # ---- 首行缩进 ----
    first_indent_pt = 0.0
    para_main_size = default_size_pt
    for r in para.runs:
        _, sz, _ = get_run_font(r)
        if sz is not None:
            para_main_size = sz
            break

    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            fl = ind.get(qn("w:firstLine"))
            flc = ind.get(qn("w:firstLineChars"))
            if fl is not None:
                try:
                    first_indent_pt = int(fl) / 20.0
                except Exception:
                    pass
            elif flc is not None:
                try:
                    first_indent_pt = int(flc) / 100.0 * para_main_size
                except Exception:
                    pass

    lines = [{"w": first_indent_pt, "cells": 0.0}]

    def new_line():
        lines.append({"w": 0.0, "cells": 0.0})

    for r in para.runs:
        _, sz, _ = get_run_font(r)
        if sz is None:
            sz = para_main_size
        # 处理该 run 中的 w:br（软换行 / 分页）
        for child in r._element.iter():
            if child.tag == qn("w:br"):
                # 中止：分页符结束该段视觉行序列
                if child.get(qn("w:type")) == "page":
                    return lines
                # 其它 w:br（textWrapping 等）—— 强制新行
                new_line()
        for ch in r.text or "":
            if ch in ("\n",):
                new_line()
                continue
            if ch == "\r":
                continue
            # 计算字符宽度
            if _is_cjk(ch):
                w = sz  # 1 em
                cells = 1.0
            elif ch in _PUNCT_SET:
                # 标点：占宽度按其自身宽度估算——中文全角标点算 1 em，
                # 英文半角标点算 0.5 em；但不计入"除标点外字数"
                w = sz if ord(ch) > 0x2E7F else sz * 0.5
                cells = 0.0
            elif ch.isspace():
                # 空白：宽度算半角（保守），不计入字数
                w = sz * 0.5
                cells = 0.0
            else:
                # ASCII 字母/数字/其它半角符号
                w = sz * 0.5
                cells = 0.5
            # 折行判定
            if lines[-1]["w"] + w > usable_pt + 0.01:
                new_line()
            lines[-1]["w"] += w
            lines[-1]["cells"] += cells

    return lines


def _lines_per_page_by_line_pitch(sect):
    """按 w:docGrid/@w:linePitch 计算每页可容纳的正文行数（float）"""
    sectPr = sect._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        return None
    lp = docGrid.get(qn("w:linePitch"))
    if lp is None:
        return None
    try:
        lp_twips = int(lp)
        usable_h_twips = (sect.page_height - sect.top_margin - sect.bottom_margin) / 635.0
        return usable_h_twips / lp_twips
    except Exception:
        return None


def _paragraph_section_indices(doc, para_infos):
    """返回每个段落所属 section 下标；inline sectPr 所在段落仍属上一节。"""
    indices = []
    sec_idx = 0
    max_sec_idx = max(0, len(doc.sections) - 1)
    for info in para_infos:
        indices.append(min(sec_idx, max_sec_idx))
        pPr = info["para"]._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            sec_idx += 1
    return indices


def _hard_page_break_count(para):
    """统计段落内硬分页符数量。"""
    n = 0
    for br in para._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            n += 1
    return n


def _estimated_paragraph_display_pages(doc, para_infos):
    """
    近似估算每个段落所在的实际显示页码。

    说明：python-docx 不能取得 Word/WPS 的最终渲染分页。本函数用页面可用高度、
    docGrid 行距、段落折行、段前/段后和硬分页符近似模拟分页；因此可识别
    空段落/回车撑满页面造成的自然分页，也能处理 <w:br w:type="page"/>。
    """
    sections = list(doc.sections)
    if not sections:
        return []
    sec_indices = _paragraph_section_indices(doc, para_infos)
    pages = []
    current_sec = sec_indices[0] if sec_indices else 0
    page_in_sec = 1
    used_lines = 0.0

    for idx, info in enumerate(para_infos):
        sec_idx = sec_indices[idx]
        if sec_idx != current_sec:
            current_sec = sec_idx
            page_in_sec = 1
            used_lines = 0.0

        sect = sections[min(sec_idx, len(sections) - 1)]
        lpp = _lines_per_page_by_line_pitch(sect)
        if lpp is None or lpp <= 0:
            try:
                usable_h_pt = (sect.page_height - sect.top_margin - sect.bottom_margin) / 12700.0
                lpp = usable_h_pt / 22.0
            except Exception:
                lpp = 32.0

        try:
            usable_w_pt = (sect.page_width - sect.left_margin - sect.right_margin) / 12700.0
        except Exception:
            usable_w_pt = 450.0

        para = info["para"]
        visual_lines = _wrap_paragraph_lines(para, usable_w_pt, default_size_pt=12.0)
        line_count = max(1.0, float(len(visual_lines)))

        try:
            usable_h_pt = (sect.page_height - sect.top_margin - sect.bottom_margin) / 12700.0
            line_pitch_pt = usable_h_pt / lpp if lpp else 22.0
            pf = info.get("pf", {})
            extra_pt = (pf.get("space_before") or 0.0) + (pf.get("space_after") or 0.0)
            if line_pitch_pt > 0:
                line_count += extra_pt / line_pitch_pt
        except Exception:
            pass

        if used_lines > 0 and used_lines + line_count > lpp + 0.01:
            page_in_sec += 1
            used_lines = 0.0

        pages.append(page_in_sec)
        used_lines += line_count

        while used_lines > lpp + 0.01:
            page_in_sec += 1
            used_lines -= lpp

        hard_breaks = _hard_page_break_count(para)
        if hard_breaks:
            page_in_sec += hard_breaks
            used_lines = 0.0

    display_pages = []
    for idx, page in enumerate(pages):
        sec_idx = sec_indices[idx]
        sect = sections[min(sec_idx, len(sections) - 1)]
        start = 1
        pg_num_type = sect._sectPr.find(qn("w:pgNumType"))
        if pg_num_type is not None and pg_num_type.get(qn("w:start")) is not None:
            try:
                start = int(pg_num_type.get(qn("w:start")))
            except Exception:
                start = 1
        display_pages.append(start + page - 1)
    return display_pages



def _find_wps_converter():
    """搜索本机 WPS 的 kwpsconvert.exe 路径，找不到返回 None"""
    import glob
    candidates = [
        r"C:\Users\{}\AppData\Local\Kingsoft\WPS Office".format(os.environ.get("USERNAME", "")),
        r"C:\Program Files\Kingsoft\WPS Office",
        r"C:\Program Files (x86)\Kingsoft\WPS Office",
    ]
    for base in candidates:
        if not base or not os.path.isdir(base):
            continue
        # WPS 下有版本号子目录 / office6 / kwpsconvert.exe
        hits = glob.glob(os.path.join(base, "*", "office6", "kwpsconvert.exe"))
        if hits:
            # 取版本号最大的一个（字典序基本等价）
            hits.sort()
            return hits[-1]
    return None


def measure_visual_layout(docx_path, doc):
    """
    视觉度量返回：
      {
        "ok": True/False,
        "error": "...",             # ok=False 时给出原因
        "max_lines_per_page": int,  # 全文页面正文最大行数
        "max_chars_per_line": int,  # 全文正文行"除标点外"最大字符数
        "per_page": [(lines, chars, sample_line_text), ...],
      }
    页眉/页脚过滤：使用 docx 中 section[0] 的上下页边距，
    只统计 y ∈ [top_margin_pt, page_height_pt - bottom_margin_pt] 内的行。
    """
    result = {
        "ok": False, "error": "",
        "max_lines_per_page": 0, "max_chars_per_line": 0,
        "per_page": [],
    }
    try:
        try:
            import pdf_backend
        except ImportError:
            from verifiers import pdf_backend
    except Exception as e:
        result["error"] = f"缺少 PDF 解析库: {e}"
        return result

    converter = _find_wps_converter()
    if not converter:
        result["error"] = "未找到 WPS 转换器 kwpsconvert.exe"
        return result

    import subprocess, tempfile, uuid
    # kwpsconvert 见到目标已存在会自动改名 (xxx(1).pdf)，因此
    # 必须给一个"尚未创建"的临时路径，不能用 NamedTemporaryFile 预建空文件。
    tmp_dir = tempfile.gettempdir()
    tmp_path = os.path.join(tmp_dir, f"_docx_render_{uuid.uuid4().hex}.pdf")
    try:
        proc = subprocess.run(
            [converter, "word2pdf", docx_path, "-o", tmp_path],
            capture_output=True, timeout=120,
        )
        if proc.returncode != 0 or not os.path.exists(tmp_path) \
                or os.path.getsize(tmp_path) == 0:
            err = (proc.stderr or b"").decode("utf-8", "ignore")
            out = (proc.stdout or b"").decode("utf-8", "ignore")
            result["error"] = f"WPS 转换 PDF 失败 rc={proc.returncode}: {err[:120] or out[:120]}"
            return result

        # 用 docx 的页边距做正文区域裁剪 (EMU → pt, 1pt=12700EMU)
        sect = doc.sections[0]
        top_pt = (sect.top_margin or 0) / 12700.0
        bottom_pt = (sect.bottom_margin or 0) / 12700.0

        pdoc = pdf_backend.open_pdf(tmp_path)
        try:
            for pno in range(pdoc.page_count):
                _, ph = pdoc.page_size(pno)  # pt
                body_top = top_pt * 0.75  # 略放宽：避免正文首行贴边被误剔
                body_bottom = ph - bottom_pt * 0.75

                # 视觉行（适配层已按 y 坐标聚合并做 3pt 同行合并）
                merged = []
                for ln in pdoc.extract_text_lines(pno, y_tol=3.0):
                    y0 = ln.bbox.y0
                    # 只保留正文区域内的行（剔除页眉页脚/页码）
                    if y0 < body_top or y0 > body_bottom:
                        continue
                    merged.append((y0, ln.text))

                lines_n = len(merged)
                max_chars_here = 0
                sample = ""
                for _, t in merged:
                    n = sum(1 for ch in t if (not ch.isspace()) and (ch not in _PUNCT_SET))
                    if n > max_chars_here:
                        max_chars_here = n
                        sample = t
                result["per_page"].append((lines_n, max_chars_here, sample))
                if lines_n > result["max_lines_per_page"]:
                    result["max_lines_per_page"] = lines_n
                if max_chars_here > result["max_chars_per_line"]:
                    result["max_chars_per_line"] = max_chars_here
        finally:
            pdoc.close()

        result["ok"] = True
    except subprocess.TimeoutExpired:
        result["error"] = "WPS 转换 PDF 超时"
    except Exception as e:
        result["error"] = f"视觉度量异常: {e}"
    finally:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
    return result


class Evaluator:
    def __init__(self, doc):
        self.doc = doc
        self.hits = []   # 命中的得分/扣分点（label, delta）
        self.paragraphs = list(doc.paragraphs)
        # 段落文本（清洗后）
        self.para_texts = [p.text.strip() for p in self.paragraphs]
        # 段落完整信息
        self.para_infos = []
        for p in self.paragraphs:
            txt = get_paragraph_text(p).strip()
            pf = get_para_format(p)
            # 段落主体字体/字号：取第一个 run，若无 run 给个空
            font_name, font_size, bold = None, None, False
            for r in p.runs:
                font_name, font_size, bold = get_run_font(r)
                if font_name or font_size:
                    break
            self.para_infos.append({
                "text": txt,
                "pf": pf,
                "font": font_name,
                "size": font_size,
                "bold": bold,
                "para": p,
            })

    # ---------- 辅助：根据文本找段落索引 ----------
    def find_para_idx(self, text_key, start=0, end=None):
        if end is None:
            end = len(self.para_texts)
        for i in range(start, end):
            if text_key in self.para_texts[i]:
                return i
        return -1

    # ---------- 评估项 ----------
    def check_all(self):
        self.eval_pagesize()
        self.eval_line_spacing()
        self.eval_header_text_after_cover()
        self.eval_header_font()
        self.eval_header_line()
        self.eval_footer_pagenum()
        self.eval_structure_order()
        self.eval_chapter_title()
        self.eval_conclusion_like_chapter()
        self.eval_section_title()
        self.eval_toc_title()
        self.eval_toc_items_font()
        self.eval_cover_format()
        self.eval_cover_format_pf()
        self.eval_cover_names()
        self.eval_cover_names_pf()
        self.eval_cover_college()
        self.eval_cover_college_pf()
        self.eval_cover_date()
        self.eval_cover_date_pf()
        self.eval_toc2_format()
        self.eval_toc2_format_pf()
        self.eval_toc2_content_format()
        self.eval_toc2_content_format_pf()
        self.eval_page3_title()
        self.eval_page3_title_pf()
        self.eval_page3_table_header_font()
        self.eval_page3_table_other_font()
        self.eval_page4_gc_title_font()
        self.eval_page4_gc_title_pf()
        self.eval_page4_sanfangxiang_font()
        self.eval_page4_sanfangxiang_pf()
        self.eval_page4_sanfangxiang_content_font()
        self.eval_page5_chengguo_title_font()
        self.eval_page5_chengguo_title_pf()
        self.eval_page5_yi_er_content()
        self.eval_page7_chengguo_title_font()
        self.eval_page7_chengguo_title_pf()
        self.eval_page7_yier_subtitle()
        self.eval_page7_yier_subtitle_pf()
        self.eval_page7_yier_content_font()
        self.eval_page7_yier_content_pf()
        self.eval_page7_san_table_font()
        self.eval_page8_college_font()
        self.eval_page8_college_pf()
        self.eval_page8_jiaocheng_font()
        self.eval_page8_jiaocheng_pf()
        self.eval_page8_benke_font()
        self.eval_page8_benke_pf()
        self.eval_page8_xinxi_font()
        self.eval_page8_xinxi_pf()

    # --------------------------------------------------------
    # +3 细则：
    #    文档纸张大小为A4，页面方向为纵向，
    #    页面宽度约21厘米、高度约29.7厘米；
    #    文档页边距满足：左边距约3厘米，右边距约2.5厘米，
    #                    上边距约3厘米，下边距约2.5厘米
    # --------------------------------------------------------
    def eval_pagesize(self):
        sect = self.doc.sections[0]
        # 基于 Word 文件属性读取页面设置
        pw_cm = emu_to_cm(sect.page_width)
        ph_cm = emu_to_cm(sect.page_height)
        orient = sect.orientation
        left_cm = emu_to_cm(sect.left_margin)
        right_cm = emu_to_cm(sect.right_margin)
        top_cm = emu_to_cm(sect.top_margin)
        bottom_cm = emu_to_cm(sect.bottom_margin)

        # 逐条核对细则中的每一个点（“约”按 ±0.2cm 容差）
        # 1) 页面宽度约21厘米
        width_ok = cm_approx(pw_cm, 21.0, tol=0.2)
        # 2) 页面高度约29.7厘米
        height_ok = cm_approx(ph_cm, 29.7, tol=0.2)
        # 3) 纸张大小为A4（由宽约21cm、高约29.7cm共同满足）
        is_a4 = width_ok and height_ok
        # 4) 页面方向为纵向
        is_portrait = orient is not None and "PORTRAIT" in str(orient).upper()
        # 5) 左边距约3厘米
        left_ok = cm_approx(left_cm, 3.0, tol=0.2)
        # 6) 右边距约2.5厘米
        right_ok = cm_approx(right_cm, 2.5, tol=0.2)
        # 7) 上边距约3厘米
        top_ok = cm_approx(top_cm, 3.0, tol=0.2)
        # 8) 下边距约2.5厘米
        bottom_ok = cm_approx(bottom_cm, 2.5, tol=0.2)

        if is_a4 and is_portrait and left_ok and right_ok and top_ok and bottom_ok:
            self.hits.append(("A4纵向 页宽约21cm/页高约29.7cm 左3/右2.5/上3/下2.5cm", +3))
        else:
            self.hits.append(("A4纵向 页宽约21cm/页高约29.7cm 左3/右2.5/上3/下2.5cm", 0))

    # --------------------------------------------------------
    # +3 细则：
    #    全文段落行距为固定值22磅；
    #    每页最多32行，每行除标点外最多34字。
    #    只走 Word 文件属性判定，不做 PDF 视觉度量、不依赖人工确认：
    #      - 点1：所有非空段落 w:spacing @w:lineRule=exact 且 @w:line=440 twips (22pt)
    #      - 点2：只要 w:docGrid/@w:linePitch 存在，就用
    #             (可用高度 twips / linePitch) ≤ 32.5 判定每页最多 32 行
    #      - 点3：按 Word 属性逐段模拟折行——用每个 run 的 w:sz 与 w:rFonts
    #             计算 CJK / 半角字符宽度，考虑段首缩进（w:ind @w:firstLine 或
    #             @w:firstLineChars），当宽度累计超过正文可用宽度时换行；
    #             每一可视行的"除标点外字数"（CJK=1、半角非标点=0.5、
    #             中英标点=0、空白=0）≤ 34 即通过。
    # --------------------------------------------------------
    def eval_line_spacing(self):
        LABEL = "行距固定22磅+每页最多32行+每行最多34字"

        # ----------- 点1：全文段落行距为固定值22磅 -----------
        total = 0
        hit = 0
        for info in self.para_infos:
            if not info["text"]:
                continue
            total += 1
            if is_fixed_line_22pt(info["pf"]):
                hit += 1
        all_22pt = (total > 0 and hit == total)

        # ----------- 点2：每页最多 32 行（w:docGrid/@w:linePitch） -----------
        sect = self.doc.sections[0]
        lpp = _lines_per_page_by_line_pitch(sect)
        lines_le_32 = (lpp is not None and lpp <= 32.5)

        # ----------- 点3：按 Word 属性逐段模拟折行，判定每行 ≤ 34 字 -----------
        usable_pt = (sect.page_width - sect.left_margin - sect.right_margin) / 12700.0
        chars_le_34 = True
        for info in self.para_infos:
            if not info["text"]:
                continue
            # 段落主字号（拿第一个非空 run 的字号；缺省 12pt=小四）
            default_size = info["size"] or 12.0
            lines = _wrap_paragraph_lines(info["para"], usable_pt, default_size_pt=default_size)
            for l in lines:
                # "除标点外字数"允许 34；容差 0.25（避免 0.5 累加边界抖动）
                if l["cells"] > 34.25:
                    chars_le_34 = False
                    break
            if not chars_le_34:
                break

        if all_22pt and lines_le_32 and chars_le_34:
            self.hits.append((LABEL, +3))
        else:
            self.hits.append((LABEL, 0))

    # --------------------------------------------------------
    # +3 细则：
    #    从"本科设计报告"那一页开始出现页眉，页面页眉均为
    #    "云岭数字技术学院本科课程设计报告"
    #    一切基于 Word 文件属性（各 section 的 default/first/even 页眉部件）。
    #    判定口径：
    #      - 锚点：文档中"说明"章节里的报告题目上下文中的
    #        "本科设计报告"段落（宋体12pt 那一处；与封底黑体18pt的
    #        另一处"本科设计报告"区分——封底那一处不作为锚点）；
    #        取该段落所属的 section 下标作为起始节。
    #      - 起始节之前的所有节（锚点所在页之前的页面）不检查页眉；
    #      - 起始节及以后的所有适用页眉部件（起始节的 default/first/even，
    #        以及之后各节的 default/first/even）均视为"该页及之后的页面"，
    #        逐一要求：页眉存在（非空）且文本完全等于目标；
    #      - 若锚点未找到，视为不满足；
    #      - 任一适用页眉部件为空或文本不完全等于目标 → 不给分。
    #    全部满足 → +3；否则 0。
    # --------------------------------------------------------
    def eval_header_text_after_cover(self):
        expected = "云岭数字技术学院本科课程设计报告"

        def header_texts(header):
            out = []
            try:
                for p in header.paragraphs:
                    t = p.text.strip()
                    if t:
                        out.append(t)
            except Exception:
                pass
            return out

        # 定位锚点："说明"章节里报告题目上下文中的"本科设计报告"
        # （与封底黑体18pt 那一处区分：封底那一处按第八页规则定位为
        #  "云岭数字技术学院"之后出现的第一个"本科设计报告"，这里要排除掉）。
        college_idx = -1
        for i, info in enumerate(self.para_infos):
            if info["text"] == "云岭数字技术学院":
                college_idx = i
                break

        back_cover_idx = -1
        if college_idx >= 0:
            for i in range(college_idx + 1, len(self.para_infos)):
                if self.para_infos[i]["text"] == "本科设计报告":
                    back_cover_idx = i
                    break

        anchor_idx = -1
        for i, info in enumerate(self.para_infos):
            if info["text"] != "本科设计报告":
                continue
            if i == back_cover_idx:
                # 封底那一处不作为"开始出现页眉"的锚点
                continue
            anchor_idx = i
            break

        if anchor_idx < 0:
            self.hits.append(("本科设计报告页起页眉均为指定文本", 0))
            return

        sec_indices = _paragraph_section_indices(self.doc, self.para_infos)
        start_sec = sec_indices[anchor_idx]

        # 是否启用奇偶页不同
        try:
            settings_el = self.doc.settings.element
            even_odd_enabled = settings_el.find(qn("w:evenAndOddHeaders")) is not None
        except Exception:
            even_odd_enabled = False

        ok_all = True
        sections = list(self.doc.sections)
        for sec_idx, section in enumerate(sections):
            if sec_idx < start_sec:
                # 锚点所在页之前的页面，不纳入检查
                continue

            # default 页眉：起始节及以后始终检查
            texts = header_texts(section.header)
            if not texts or texts != [expected]:
                ok_all = False
                break

            # first-page 页眉：仅在启用 titlePg 时生效
            if section.different_first_page_header_footer:
                texts = header_texts(section.first_page_header)
                if not texts or texts != [expected]:
                    ok_all = False
                    break

            # even-page 页眉：仅在启用奇偶页不同时生效
            if even_odd_enabled:
                texts = header_texts(section.even_page_header)
                if not texts or texts != [expected]:
                    ok_all = False
                    break

        self.hits.append(("本科设计报告页起页眉均为指定文本", +3 if ok_all else 0))

    # --------------------------------------------------------
    # +1 细则：
    #    页面页眉文本为宋体、五号、居中显示
    #    一切基于 Word 文件属性（节的页眉：default / first / even
    #      + 段落 w:jc + run w:rFonts / w:sz）：
    #      - 判定范围：所有实际承载页眉文本（非空）的页眉部件都要检查；
    #        没有文本的页眉部件（该页面没有页眉）不参与判定，直接跳过。
    #      - 三个点：字体=宋体、字号=五号(10.5pt)、段落对齐=居中。
    # --------------------------------------------------------
    def eval_header_font(self):
        FIVE_HAO_PT = 10.5  # Word 五号 = 10.5pt

        def check_header(header):
            """若页眉无文本则视为"该页无页眉"，返回 None（不参与判定）；
            有文本时要求所有非空段落满足：宋体 + 五号 + 居中，返回 True/False。"""
            try:
                paras = list(header.paragraphs)
            except Exception:
                return None  # 取不到页眉，视为该页无页眉，不参与判定
            if not any(p.text.strip() for p in paras):
                return None
            for p in paras:
                if not p.text.strip():
                    continue
                # 点3：居中（段落属性 w:jc = center）
                if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    return False
                # 点1 + 点2：每个非空 run 都要宋体 + 五号
                for r in p.runs:
                    if not r.text.strip():
                        continue
                    fn, sz, _ = get_run_font(r)
                    if not (fn and "宋体" in fn):
                        return False
                    if not (sz is not None and pt_approx(sz, FIVE_HAO_PT, tol=0.5)):
                        return False
            return True

        # 是否启用奇偶页不同（w:settings/w:evenAndOddHeaders）
        try:
            settings_el = self.doc.settings.element
            even_odd_enabled = settings_el.find(qn("w:evenAndOddHeaders")) is not None
        except Exception:
            even_odd_enabled = False

        ok_all = True
        checked_any = False
        sections = list(self.doc.sections)
        for section in sections:
            # default 页眉：有文本才检查
            r = check_header(section.header)
            if r is False:
                ok_all = False
                break
            if r is True:
                checked_any = True

            # first-page 页眉：仅在启用 titlePg 时生效；有文本才检查
            if section.different_first_page_header_footer:
                r = check_header(section.first_page_header)
                if r is False:
                    ok_all = False
                    break
                if r is True:
                    checked_any = True

            # even-page 页眉：仅在启用奇偶页不同时生效；有文本才检查
            if even_odd_enabled:
                r = check_header(section.even_page_header)
                if r is False:
                    ok_all = False
                    break
                if r is True:
                    checked_any = True

        # 若文档所有页眉都为空，视为不满足（没有可判定的页眉文本）
        self.hits.append(("页眉文本 宋体/五号/居中", +1 if (ok_all and checked_any) else 0))

    # --------------------------------------------------------
    # +3 页眉下横向单线 (0.5pt/单线)
    #    判定口径：只检查实际承载页眉文本（非空）的页眉部件下方是否
    #    有符合要求的横向单线；没有文本的页眉部件（该页无页眉）跳过，
    #    不参与判定。全部通过 → +3；任一有文本的页眉部件缺横线或
    #    横线不合规 → 0。
    # --------------------------------------------------------
    def eval_header_line(self):
        def header_has_text(header):
            try:
                paras = list(header.paragraphs)
            except Exception:
                return False
            return any(p.text.strip() for p in paras)

        def header_has_line(header):
            try:
                paras = list(header.paragraphs)
            except Exception:
                return False
            for p in paras:
                pPr = p._p.find(qn("w:pPr"))
                if pPr is None:
                    continue
                pBdr = pPr.find(qn("w:pBdr"))
                if pBdr is None:
                    continue
                bottom = pBdr.find(qn("w:bottom"))
                if bottom is None:
                    continue
                if bottom.get(qn("w:val")) in ("single",):
                    # 检查粗细：w:sz 单位为 1/8 pt，0.5pt = 4
                    sz = bottom.get(qn("w:sz"))
                    color = bottom.get(qn("w:color"))
                    # 检查颜色：黑色或深色（黑色=000000，深色允许其他深色值）
                    color_ok = color is None or color.lower() in ("000000", "auto") or color.lower().startswith("00")
                    # 检查粗细：约0.5pt，允许范围 2-8 (0.25pt-1pt)
                    if sz and 2 <= int(sz) <= 8 and color_ok:
                        return True
            return False

        # 是否启用奇偶页不同（w:settings/w:evenAndOddHeaders）
        try:
            settings_el = self.doc.settings.element
            even_odd_enabled = settings_el.find(qn("w:evenAndOddHeaders")) is not None
        except Exception:
            even_odd_enabled = False

        ok_all = True
        checked_any = False
        sections = list(self.doc.sections)
        for section in sections:
            # default 页眉：有文本才检查
            if header_has_text(section.header):
                checked_any = True
                if not header_has_line(section.header):
                    ok_all = False
                    break

            # first-page 页眉：仅在启用 titlePg 时生效；有文本才检查
            if section.different_first_page_header_footer:
                if header_has_text(section.first_page_header):
                    checked_any = True
                    if not header_has_line(section.first_page_header):
                        ok_all = False
                        break

            # even-page 页眉：仅在启用奇偶页不同时生效；有文本才检查
            if even_odd_enabled:
                if header_has_text(section.even_page_header):
                    checked_any = True
                    if not header_has_line(section.even_page_header):
                        ok_all = False
                        break

        self.hits.append(("页眉下横向单线(≈0.5pt)", +3 if (ok_all and checked_any) else 0))

    # --------------------------------------------------------
    # +3 细则：
    #    "第一章 应急物资调度需求概述"页下方的页脚位置出现页码"1"，
    #    且后续页面页码按2、3、4……连续递增；页码均位于页脚区域的居中位置。
    #    判定口径：
    #      - 定位页脚内实际承载 PAGE 字段的段落（而不是任意非空文本段落，
    #        因为 PAGE 字段渲染的页码不写入 w:t，p.text 常为空），
    #        对这些段落检查 w:jc = center（居中）。
    #      - 页码起始值：读取第一章标题所在 section 的
    #        <w:pgNumType w:start="N"/>（若未显式设置则默认从1续算），
    #        与 _estimated_paragraph_display_pages 的估算结果两者一致
    #        且都为 1，才认定"第一章页脚显示为1"。
    #      - 后续页码连续递增：复用估算结果的逐页 +1 检查。
    # --------------------------------------------------------
    def eval_footer_pagenum(self):
        # ---- 定位页脚中实际承载 PAGE 字段的段落，检查其居中对齐 ----
        # PAGE 字段可能是 w:fldSimple（@w:instr 含 "PAGE"）或
        # w:fldChar begin/separate/end 配合 w:instrText（文本内容含 "PAGE"）。
        # 两种形式下，字段渲染出的页码数字都不会写入该段落的 w:t，
        # 因此不能用 p.text.strip() 判断"是否需要检查居中"。
        align_ok = True
        found_page_field_para = False
        for section in self.doc.sections:
            for p in section.footer.paragraphs:
                has_page_field = False
                for el in p._p.iter():
                    if el.tag == qn("w:fldSimple"):
                        instr = el.get(qn("w:instr")) or ""
                        if "PAGE" in instr.upper():
                            has_page_field = True
                    elif el.tag == qn("w:instrText"):
                        if "PAGE" in (el.text or "").upper():
                            has_page_field = True
                if not has_page_field:
                    continue
                found_page_field_para = True
                if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    align_ok = False

        # 检查页码字段
        has_page_field = found_page_field_para

        # 检查页码从1开始，且从"第一章 应急物资调度需求概述"所在页开始，
        # 后续页码逐页递增。
        # 说明：不能只看该章所在 section 的 <w:pgNumType w:start="1"/>——
        # 这只代表"该节的页码计数器从1起"，如果该节内、计数器重置点与
        # 第一章标题之间还夹着其它自然分页（哪怕没有硬分页符，仅靠空段落/
        # 回车撑满页面），第一章实际显示的页码就不是1。因此改为用
        # _estimated_paragraph_display_pages 估算每个段落的实际显示页码
        # （同时考虑硬分页符 <w:br w:type="page"/> 和按可用行数模拟的自然
        # 分页，且该估算已把 w:pgNumType/@w:start 计入起始偏移），定位到
        # 第一章标题段落后检查其页码是否为1，并检查该页之后的页码是否
        # 单调不减、且逐步递增（无跳号、无回退）。
        starts_at_1_at_chapter = False
        # 找到"第一章 应急物资调度需求概述"章标题段落（排除目录条目：
        # 行末为 <Tab/空白>数字 的形态，例如 "第一章  应急物资调度需求概述\t3"）
        toc_entry_re = re.compile(r"(?:\t|\s)\s*\d+\s*$")
        chapter_para_idx = -1
        for i, info in enumerate(self.para_infos):
            t = info["text"]
            if "第一章" not in t or "应急物资调度需求概述" not in t:
                continue
            if toc_entry_re.search(t):
                continue  # 目录条目跳过
            chapter_para_idx = i
            break
        # 若正文里未找到独立章标题（例如仅有目录条目），退回到任意匹配段落
        if chapter_para_idx < 0:
            for i, info in enumerate(self.para_infos):
                if "第一章" in info["text"] and "应急物资调度需求概述" in info["text"]:
                    chapter_para_idx = i
                    break

        if chapter_para_idx >= 0:
            sec_indices = _paragraph_section_indices(self.doc, self.para_infos)
            display_pages = _estimated_paragraph_display_pages(self.doc, self.para_infos)
            chapter_page = display_pages[chapter_para_idx]

            # 与该段落所在 section 的 <w:pgNumType w:start="N"/> 交叉核对：
            # 若该节显式设置了 start，则第一章所在页必须恰好等于 start；
            # 未显式设置则视为承接上一节延续计数，仅要求估算值为 1。
            sections = list(self.doc.sections)
            chapter_sec_idx = sec_indices[chapter_para_idx]
            chapter_sect = sections[min(chapter_sec_idx, len(sections) - 1)]
            pg_num_type = chapter_sect._sectPr.find(qn("w:pgNumType"))
            explicit_start = None
            if pg_num_type is not None and pg_num_type.get(qn("w:start")) is not None:
                try:
                    explicit_start = int(pg_num_type.get(qn("w:start")))
                except Exception:
                    explicit_start = None

            if explicit_start is not None:
                page1_ok = (chapter_page == 1 == explicit_start)
            else:
                page1_ok = (chapter_page == 1)

            if page1_ok:
                # 逐页递增检查：从第一章开始，往后每次页码变化都应恰好 +1
                # （不允许跳号或回退；相同页内多个段落页码相同属正常）。
                monotonic_ok = True
                prev_page = chapter_page
                for i in range(chapter_para_idx + 1, len(display_pages)):
                    cur_page = display_pages[i]
                    if cur_page < prev_page or cur_page > prev_page + 1:
                        monotonic_ok = False
                        break
                    prev_page = cur_page
                starts_at_1_at_chapter = monotonic_ok
        if align_ok and has_page_field and starts_at_1_at_chapter:
            self.hits.append(("页码连续居中", +3))
        else:
            self.hits.append(("页码连续居中", 0))


    # --------------------------------------------------------
    # +3 细则：
    #    文档包含封面、目录、正文、参考文献四类结构，
    #    且先后顺序为 封面 → 目录 → 正文 → 参考文献
    #    一切基于 Word 文件属性（段落文本 w:t + 段落顺序）：
    #      - 封面：不再用"第一个非空段落"近似判断（会把任何有文字的
    #              文档误判为"有封面"）。改为要求封面标志性内容——
    #              标题"工程实践与算法应用综合设计指导书"、六位作者姓名
    #              （与 eval_cover_format / eval_cover_names 等定位一致）
    #              ——至少命中一项，取其中最早出现的段落索引作为封面位置；
    #              一项都未命中则视为无封面结构。
    #      - 目录：段落文本（去空白后）为"目录"的段落。
    #      - 正文：目录之后、以"第X章"开头且**非目录条目**的章标题段落
    #              （目录条目形如"第一章  xxx\t3"，行末带制表符+页码/仅数字，
    #                据此过滤）。
    #      - 参考文献：段落文本（去空白后）严格等于"参考文献"
    #                  （rubric 明确要求"参考文献"，"参考资料"不算）的
    #                  独立标题段落（排除目录条目）。
    #    四项都必须存在，且严格按 cover < toc < body < ref 顺序出现。
    # --------------------------------------------------------
    def eval_structure_order(self):
        # ---- 封面：命中封面标志性内容的最早段落 ----
        # 与 eval_cover_format / eval_cover_names / eval_cover_college 的
        # 定位标准保持一致，避免"任意非空段落"就被当作封面。
        cover_title = "工程实践与算法应用综合设计指导书"
        author_names = ["林沐川", "陆安琪", "许庭舟", "邵思远", "周砚初", "顾清禾"]
        idx_cover = -1
        for i, t in enumerate(self.para_texts):
            if not t:
                continue
            if cover_title in t or all(name in t for name in author_names):
                idx_cover = i
                break

        # ---- 目录：文本（去所有空白）恰为"目录" ----
        idx_toc = -1
        for i, t in enumerate(self.para_texts):
            s = t.replace(" ", "").replace("　", "").replace("\t", "")
            if s == "目录":
                idx_toc = i
                break

        # ---- 正文：目录之后、以"第X章"开头且非目录条目 ----
        # 目录条目特征：末尾是"...<Tab>数字"或"...<Tab><空格>数字"，
        # 或整行仅由"第X章 标题 数字"构成（页码结尾）。
        chapter_head_re = re.compile(r"^第\s*[一二三四五六七八九十百零0-9]+\s*章")
        toc_entry_re = re.compile(r"\t\s*\d+\s*$")  # 行尾"<Tab>数字"
        idx_body = -1
        search_from = idx_toc + 1 if idx_toc >= 0 else 0
        for i in range(search_from, len(self.para_texts)):
            t = self.para_texts[i]
            s = t.replace("　", " ").strip()
            if not s:
                continue
            if not chapter_head_re.match(s):
                continue
            # 排除目录条目（末尾 \t 数字），或整行以纯数字结尾
            if toc_entry_re.search(t) or re.search(r"\s\d+\s*$", t):
                continue
            idx_body = i
            break

        # ---- 参考文献：正文之后、独立"参考文献"章标题段落 ----
        # rubric 只认"参考文献"，"参考资料"不满足该细则。
        # 也排除目录条目（"参考文献\t8"这种）。
        idx_ref = -1
        search_from = idx_body + 1 if idx_body >= 0 else 0
        for i in range(search_from, len(self.para_texts)):
            t = self.para_texts[i]
            # 排除目录条目
            if toc_entry_re.search(t) or re.search(r"\s\d+\s*$", t):
                continue
            s = t.replace(" ", "").replace("　", "").replace("\t", "")
            if s == "参考文献":
                idx_ref = i
                break

        # ---- 判定 ----
        has_all_four = (idx_cover >= 0 and idx_toc >= 0
                        and idx_body >= 0 and idx_ref >= 0)
        order_ok = has_all_four and (idx_cover < idx_toc < idx_body < idx_ref)

        if has_all_four and order_ok:
            self.hits.append(("结构 封面→目录→正文→参考文献 齐全且有序", +3))
        else:
            self.hits.append(("结构 封面→目录→正文→参考文献 齐全且有序", 0))

    # --------------------------------------------------------
    # +3 细则：
    #    章标题段落文本匹配"第…章"开头，
    #    字体为黑体，字号为小二号
    #    一切基于 Word 文件属性（段落 w:t + run 的 w:rFonts / w:sz）：
    #      - 点1：段落文本以"第…章"开头（中文数字或阿拉伯数字均可）
    #      - 点2：字体 = 黑体
    #      - 点3：字号 = 小二号 (18pt)
    #    检查范围为章标题段落内所有非空 run（而非仅第一个 run）——
    #    避免标题中间/末尾 run 格式错误（如手动改动局部文字）被漏检。
    #    所有满足点1的章标题段落、且段落内每个非空 run 都必须同时
    #    满足点2、点3；且至少存在一个章标题段落。
    # --------------------------------------------------------
    def eval_chapter_title(self):
        pattern = re.compile(r"^第\s*[一二三四五六七八九十百零0-9]+\s*章")
        # 目录条目形如 "第一章  xxx\t3" 或 "第一章  xxx 3"
        toc_entry_re = re.compile(r"(?:\t|\s)\s*\d+\s*$")

        chapter_paras = []
        for info in self.para_infos:
            t = info["text"]
            if not t or not pattern.match(t):
                continue
            # 排除目录条目：行末为 <Tab/空白>数字
            if toc_entry_re.search(t):
                continue
            chapter_paras.append(info)

        if not chapter_paras:
            self.hits.append(("章标题 第…章/黑体/小二号", 0))
            return

        all_ok = True
        for info in chapter_paras:
            for r in info["para"].runs:
                if not r.text.strip():
                    continue
                fn, sz, _ = get_run_font(r)
                # 点2：黑体
                font_ok = fn and "黑体" in fn
                # 点3：小二号 = 18pt
                size_ok = sz is not None and pt_approx(sz, 18, tol=0.5)
                if not (font_ok and size_ok):
                    all_ok = False
                    break
            if not all_ok:
                break

        self.hits.append(("章标题 第…章/黑体/小二号", +3 if all_ok else 0))

    # --------------------------------------------------------
    # +1 细则：
    #    "结论""参考文献""附录"标题按章标题处理，
    #    字体为黑体，字号为小二号
    #    一切基于 Word 文件属性（段落 w:t + run 的 w:rFonts / w:sz）：
    #      - rubric 明确是"结论""参考文献""附录"三者（不是"参考资料"）；
    #      - 三项标题段落必须逐一存在，且每个标题段落内所有非空 run
    #        都满足黑体 + 小二号(18pt)；
    #      - 三项全部满足才给分，不是命中数 >= 2 即给分。
    # --------------------------------------------------------
    def eval_conclusion_like_chapter(self):
        keys = ["结论", "参考文献", "附录"]

        def para_font_ok(info):
            for r in info["para"].runs:
                if not r.text.strip():
                    continue
                fn, sz, _ = get_run_font(r)
                if not (fn and "黑体" in fn):
                    return False
                if not (sz is not None and pt_approx(sz, 18, tol=1.5)):
                    return False
            return True

        all_ok = True
        for k in keys:
            found_ok = False
            for info in self.para_infos:
                if info["text"] == k or info["text"].startswith(k):
                    if para_font_ok(info):
                        found_ok = True
                        break
            if not found_ok:
                all_ok = False
                break

        self.hits.append(("结论/参考文献/附录 章节化", +1 if all_ok else 0))

    # --------------------------------------------------------
    # +3 细则：
    #    一级节标题段落编号匹配"数字.数字"，例如"1.1""2.1"，
    #    字体为黑体，字号为小三号
    #    一切基于 Word 文件属性（段落 w:t + run 的 w:rFonts / w:sz）：
    #      - 点1：段落文本以"数字.数字"编号开头（严格两级，且非"数字.数字.数字"）
    #      - 点2：字体 = 黑体
    #      - 点3：字号 = 小三号 (15pt)
    #    检查范围为一级节标题段落内所有非空 run（而非仅第一个 run）——
    #    避免标题中混合 run 里后续文字格式错误被漏检。
    #    所有匹配点1的一级节标题段落、且段落内每个非空 run 都必须同时
    #    满足点2、点3；且至少存在一个一级节标题段落。
    # --------------------------------------------------------
    def eval_section_title(self):
        # "数字.数字" 后不能再接 ".数字"（那属于二级/三级节标题），
        # 且编号后必须有空白 + 非数字文本（排除封面日期 "2026.04" 这类纯数字段落）
        pattern = re.compile(r"^\d+\.\d+(?!\.\d)\s+\S")
        # 目录条目形如 "1.1  xxx\t3"，行末为 <Tab/空白>数字
        toc_entry_re = re.compile(r"(?:\t|\s)\s*\d+\s*$")

        section_paras = []
        for info in self.para_infos:
            t = info["text"]
            if not t or not pattern.match(t):
                continue
            # 排除目录条目
            if toc_entry_re.search(t):
                continue
            section_paras.append(info)

        if not section_paras:
            self.hits.append(("一级节标题 数字.数字/黑体/小三号", 0))
            return

        all_ok = True
        for info in section_paras:
            for r in info["para"].runs:
                if not r.text.strip():
                    continue
                fn, sz, _ = get_run_font(r)
                # 点2：黑体
                font_ok = fn and "黑体" in fn
                # 点3：小三号 = 15pt
                size_ok = sz is not None and pt_approx(sz, 15, tol=0.5)
                if not (font_ok and size_ok):
                    all_ok = False
                    break
            if not all_ok:
                break

        self.hits.append(("一级节标题 数字.数字/黑体/小三号", +3 if all_ok else 0))

    # --------------------------------------------------------
    # +1 目录标题"目 录" 黑体小二号 (18pt) 居中
    # --------------------------------------------------------
    def eval_toc_title(self):
        for info in self.para_infos:
            t = info["text"]
            if t.replace(" ", "").replace("\u3000", "") == "目录" or "目  录" in t:
                if info["font"] and "黑体" in info["font"] and info["size"] and pt_approx(info["size"], 18, tol=1.5):
                    if info["pf"].get("alignment") and "CENTER" in info["pf"]["alignment"].upper():
                        self.hits.append(("目录标题 黑体小二号居中", +1))
                        return
        self.hits.append(("目录标题 黑体小二号居中", 0))

    # --------------------------------------------------------
    # +1 细则：
    #    目录条目字体为宋体、小四号
    #    一切基于 Word 文件属性（段落 w:t 定位目录区 + run 的 w:rFonts / w:sz）：
    #      - 目录区：从"目录"标题段落之后，到首个"第X章"章标题段落之前
    #      - 目录条目：目录区中含数字（页码）的非空段落
    #      - 点1：字体 = 宋体
    #      - 点2：字号 = 小四号 (12pt)
    #    所有目录条目段落都必须同时满足点1、点2；
    #    且至少存在一个目录条目。
    # --------------------------------------------------------
    def eval_toc_items_font(self):
        # 目录条目通用属性：段落文本以 "<Tab>数字" 结尾（Word 目录 field 生成的
        # 条目就是这个形状——标题 + 制表符 + 页码）。以此定位目录条目，跨节/无节
        # 都能识别，不依赖"目录"和第一个"第X章"之间的段落范围。
        toc_entry_re = re.compile(r"\t\s*\d+\s*$")

        # 仍要求这些条目位于"目录"标题之后（避免误伤评分标准正文里带
        # "<Tab>数字" 的行）
        toc_start = -1
        for i, t in enumerate(self.para_texts):
            s = t.replace(" ", "").replace("　", "").replace("\t", "")
            if s == "目录":
                toc_start = i
                break
        if toc_start < 0:
            self.hits.append(("目录条目字体 宋体/小四号", 0))
            return

        toc_items = []
        # 目录顶部可能存在"前置条目"（如 "1. 成果文本编制要求	5"、
        # "2. 工程实践与算法应用综合设计评分标准	3"）——它们以"阿拉伯数字
        # + 点 + 空白 + 非数字"起头，且非章节 (第X章) 或节次 (1.1/2.3) 条目。
        # 评分细则只面向"章节目录条目"，因此这类前置条目应予以排除。
        prefix_top_level_re = re.compile(r"^\s*\d+\.\s+\D")
        for info in self.para_infos[toc_start + 1:]:
            t = info["text"]
            if not t:
                continue
            if toc_entry_re.search(t):
                if prefix_top_level_re.match(t):
                    continue  # 顶层"数字. xxx"前置条目跳过
                toc_items.append(info)
            # 一旦遇到非目录条目且是正文章标题（黑体+18pt 的"第X章"），
            # 说明目录已经结束
            elif re.match(r"^第\s*[一二三四五六七八九十百零0-9]+\s*章", t):
                # 该段非目录条目形态 → 目录结束
                if not toc_entry_re.search(t):
                    break

        if not toc_items:
            self.hits.append(("目录条目字体 宋体/小四号", 0))
            return

        all_ok = True
        for info in toc_items:
            # 点1：宋体
            font_ok = info["font"] and "宋体" in info["font"]
            # 点2：小四号 = 12pt
            size_ok = info["size"] is not None and pt_approx(info["size"], 12, tol=0.5)
            if not (font_ok and size_ok):
                all_ok = False
                break

        self.hits.append(("目录条目字体 宋体/小四号", +1 if all_ok else 0))

    # --------------------------------------------------------
    # 扣分项辅助：判断段落是否满足字体/段落格式
    # --------------------------------------------------------
    def font_check(self, info, expected_font, expected_size=None, expected_bold=None, size_tol=1.5):
        if not info["text"]:
            return True  # 空段落不算违反
        if expected_font and expected_font not in (info["font"] or ""):
            return False
        if expected_size is not None and not pt_approx(info["size"], expected_size, tol=size_tol):
            return False
        if expected_bold is True and not info["bold"]:
            return False
        if expected_bold is False and info["bold"]:
            return False
        return True

    def pf_check(self, info, alignment=None, space_after=None, line=None, line_rule=None,
                 first_line_indent_cm=None, left_indent_cm=None, hanging_indent_cm=None,
                 space_before_cm=None, first_line_pos_cm=None):
        pf = info["pf"]
        if alignment is not None:
            actual_alignment = pf.get("alignment")
            # OOXML 未显式设置 w:jc 时，Word/WPS 的渲染默认值是左对齐。
            if actual_alignment is None and alignment.upper() == "LEFT":
                pass
            elif not actual_alignment or alignment.upper() not in actual_alignment.upper():
                return False
        if space_after is not None:
            sa = pf.get("space_after")
            if sa is None or not pt_approx(sa, space_after, tol=2.0):
                return False
        if line is not None:
            ln = pf.get("line")
            if ln is None:
                return False
            # 行距允许 1.08/1.15 这样的倍数
            if not (abs(ln - line) <= 0.05):
                return False
        if first_line_indent_cm is not None:
            fli = pf.get("first_line_indent_cm")
            if fli is None or not cm_approx(fli, first_line_indent_cm, tol=0.2):
                return False
        if left_indent_cm is not None:
            li = pf.get("left_indent_cm")
            # 段前缩进即 left_indent
            if li is None or not cm_approx(li, left_indent_cm, tol=0.2):
                return False
        if hanging_indent_cm is not None:
            hi = pf.get("hanging_indent_cm")
            if hi is None or not cm_approx(hi, hanging_indent_cm, tol=0.2):
                return False
        if first_line_pos_cm is not None:
            # 首行相对页边距的绝对位置（= left_indent - hanging_indent，
            # 或仅设置 first_line_indent 时二者相同）——对应 WPS/Word
            # 段落对话框里直接显示的"首行缩进/段前缩进"数值，而不是
            # python-docx 在只有 w:hanging 时返回的相对负偏移。
            flp = pf.get("first_line_pos_cm")
            if flp is None or not cm_approx(flp, first_line_pos_cm, tol=0.2):
                return False
        return True

    # --------------------------------------------------------
    # 封面相关扣分
    # --------------------------------------------------------
    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"工程实践与算法应用综合设计指导书"
    #    不满足字体格式为 Noto Sans CJK SC 小一 加粗
    #    一切基于 Word 文件属性（段落 w:t 定位 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：段落文本包含"工程实践与算法应用综合设计指导书"，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 小一号 (24pt)
    #      - 点3：加粗 (w:b)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_format(self):
        # 判定段落是否在文档第一页（首个分页符之前）
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if "工程实践与算法应用综合设计指导书" in info["text"]:
                # 点1：Noto Sans CJK SC
                # 点2：小一 = 24pt
                # 点3：加粗
                if self.font_check(info, "Noto Sans CJK SC",
                                   expected_size=24, expected_bold=True, size_tol=0.5):
                    self.hits.append((
                        "封面标题字体 Noto Sans CJK SC 小一 加粗", 0))
                else:
                    self.hits.append((
                        "封面标题字体 Noto Sans CJK SC 小一 加粗", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面标题字体 Noto Sans CJK SC 小一 加粗", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"工程实践与算法应用综合设计指导书"
    #    段落格式不满足 段后26磅、1.08倍行距、居中对齐
    #    一切基于 Word 文件属性（段落 w:pPr 的 w:spacing / w:jc）：
    #      - 定位：段落文本包含"工程实践与算法应用综合设计指导书"，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：段后 = 26 磅  (w:spacing @w:after，单位 twips，26pt=520)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:lineRule="auto" 且 @w:line/240 ≈ 1.08)
    #      - 点3：居中对齐 (w:jc = center)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_format_pf(self):
        # 复用"第一页"判定
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if "工程实践与算法应用综合设计指导书" in info["text"]:
                # 点1：段后 26 磅
                # 点2：1.08 倍行距（Word 内部 line=1.08*240≈259，lineRule=auto）
                # 点3：居中
                if self.pf_check(info, alignment="CENTER",
                                 space_after=26, line=1.08):
                    self.hits.append((
                        "封面标题段落 段后26磅/1.08倍行距/居中", 0))
                else:
                    self.hits.append((
                        "封面标题段落 段后26磅/1.08倍行距/居中", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面标题段落 段后26磅/1.08倍行距/居中", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"林沐川 陆安琪 许庭舟 邵思远 周砚初 顾清禾"
    #    字体格式不满足 Noto Sans CJK SC 四号
    #    一切基于 Word 文件属性（段落 w:t 定位 + run 的 w:rFonts / w:sz）：
    #      - 定位：段落文本同时包含六位作者姓名，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 四号 (14pt)
    #    两点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_names(self):
        # 六位作者姓名（Word 文件属性——段落文本 w:t）
        author_names = ["林沐川", "陆安琪", "许庭舟", "邵思远", "周砚初", "顾清禾"]

        # 复用"第一页"判定：首个分页符之前
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if all(name in info["text"] for name in author_names):
                # 点1：Noto Sans CJK SC
                # 点2：四号 = 14pt
                if self.font_check(info, "Noto Sans CJK SC",
                                   expected_size=14, size_tol=0.5):
                    self.hits.append((
                        "封面作者字体 Noto Sans CJK SC 四号", 0))
                else:
                    self.hits.append((
                        "封面作者字体 Noto Sans CJK SC 四号", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面作者字体 Noto Sans CJK SC 四号", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"林沐川 陆安琪 许庭舟 邵思远 周砚初 顾清禾"
    #    段落格式不满足 段后14磅、1.08倍行距、居中对齐
    #    一切基于 Word 文件属性（段落 w:pPr 的 w:spacing / w:jc）：
    #      - 定位：段落文本同时包含六位作者姓名，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：段后 = 14 磅  (w:spacing @w:after，14pt=280 twips)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:lineRule="auto" 且 @w:line/240 ≈ 1.08)
    #      - 点3：居中对齐 (w:jc = center)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_names_pf(self):
        author_names = ["林沐川", "陆安琪", "许庭舟", "邵思远", "周砚初", "顾清禾"]

        # 复用"第一页"判定：首个分页符之前
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if all(name in info["text"] for name in author_names):
                # 点1：段后 14 磅
                # 点2：1.08 倍行距
                # 点3：居中
                if self.pf_check(info, alignment="CENTER",
                                 space_after=14, line=1.08):
                    self.hits.append((
                        "封面作者段落 段后14磅/1.08倍行距/居中", 0))
                else:
                    self.hits.append((
                        "封面作者段落 段后14磅/1.08倍行距/居中", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面作者段落 段后14磅/1.08倍行距/居中", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"云岭数字技术学院"字体格式不满足
    #    Noto Sans CJK SC 三号 加粗
    #    一切基于 Word 文件属性（段落 w:t 定位 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：段落文本为"云岭数字技术学院"，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 三号 (16pt)
    #      - 点3：加粗 (w:b)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_college(self):
        # 复用"第一页"判定：首个分页符之前
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if info["text"].strip() == "云岭数字技术学院":
                # 点1：Noto Sans CJK SC
                # 点2：三号 = 16pt
                # 点3：加粗
                if self.font_check(info, "Noto Sans CJK SC",
                                   expected_size=16, expected_bold=True, size_tol=0.5):
                    self.hits.append((
                        "封面学院字体 Noto Sans CJK SC 三号 加粗", 0))
                else:
                    self.hits.append((
                        "封面学院字体 Noto Sans CJK SC 三号 加粗", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面学院字体 Noto Sans CJK SC 三号 加粗", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"云岭数字技术学院"段落格式不满足
    #    段后28磅、1.08倍行距、居中对齐
    #    一切基于 Word 文件属性（段落 w:pPr 的 w:spacing / w:jc）：
    #      - 定位：段落文本为"云岭数字技术学院"，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：段后 = 28 磅  (w:spacing @w:after，28pt=560 twips)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:lineRule="auto" 且 @w:line/240 ≈ 1.08)
    #      - 点3：居中对齐 (w:jc = center)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_college_pf(self):
        # 复用"第一页"判定：首个分页符之前
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if info["text"].strip() == "云岭数字技术学院":
                # 点1：段后 28 磅
                # 点2：1.08 倍行距
                # 点3：居中
                if self.pf_check(info, alignment="CENTER",
                                 space_after=28, line=1.08):
                    self.hits.append((
                        "封面学院段落 段后28磅/1.08倍行距/居中", 0))
                else:
                    self.hits.append((
                        "封面学院段落 段后28磅/1.08倍行距/居中", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面学院段落 段后28磅/1.08倍行距/居中", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"2026.04"字体格式不满足 Noto Sans CJK SC 13
    #    一切基于 Word 文件属性（段落 w:t 定位 + run 的 w:rFonts / w:sz）：
    #      - 定位：段落文本为"2026.04"，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2:  字号 = 13 (13pt)
    #    两点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_date(self):
        # 复用"第一页"判定：首个分页符之前
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if info["text"].strip() == "2026.04":
                # 点1：Noto Sans CJK SC
                # 点2：13pt
                if self.font_check(info, "Noto Sans CJK SC",
                                   expected_size=13, size_tol=0.5):
                    self.hits.append((
                        "封面日期字体 Noto Sans CJK SC 13", 0))
                else:
                    self.hits.append((
                        "封面日期字体 Noto Sans CJK SC 13", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面日期字体 Noto Sans CJK SC 13", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第一页"2026.04"段落格式不满足
    #    段后16磅、1.08倍行距、居中对齐
    #    一切基于 Word 文件属性（段落 w:pPr 的 w:spacing / w:jc）：
    #      - 定位：段落文本为"2026.04"，
    #              且位于文档第一页（首个分页符之前）。
    #      - 点1：段后 = 16 磅  (w:spacing @w:after，16pt=320 twips)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:lineRule="auto" 且 @w:line/240 ≈ 1.08)
    #      - 点3：居中对齐 (w:jc = center)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_cover_date_pf(self):
        # 复用"第一页"判定：首个分页符之前
        first_page_last_idx = len(self.paragraphs) - 1
        for i, p in enumerate(self.paragraphs):
            has_page_break = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break
            if has_page_break:
                first_page_last_idx = i
                break

        for idx, info in enumerate(self.para_infos):
            if idx > first_page_last_idx:
                break
            if info["text"].strip() == "2026.04":
                # 点1：段后 16 磅
                # 点2：1.08 倍行距
                # 点3：居中
                if self.pf_check(info, alignment="CENTER",
                                 space_after=16, line=1.08):
                    self.hits.append((
                        "封面日期段落 段后16磅/1.08倍行距/居中", 0))
                else:
                    self.hits.append((
                        "封面日期段落 段后16磅/1.08倍行距/居中", -1))
                return
        # 未找到该段：不作扣分（0）
        self.hits.append(("封面日期段落 段后16磅/1.08倍行距/居中", 0))

    # --------------------------------------------------------
    # 第二页：目录两字 / 目录内容
    # --------------------------------------------------------
    # --------------------------------------------------------
    # -1 细则：
    #    文档第二页"目录"两字不满足：
    #      - 中间空一个字符
    #      - 字体格式为 Noto Sans CJK SC 二号 加粗
    #    一切基于 Word 文件属性（分页符定位第二页 + 段落 w:t + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第二页（第 1 个分页符之后、第 2 个分页符之前）
    #              且去除空白后文本为"目录"的段落。
    #      - 点1：中间空一个字符（w:t 原文含且仅含 1 个空白字符：空格或全角空格）
    #      - 点2：字体 = Noto Sans CJK SC
    #      - 点3：字号 = 二号 (22pt)
    #      - 点4：加粗 (w:b)
    #    四点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_toc2_format(self):
        # 用 Word 分页符定位第二页范围：[first_break_idx+1, second_break_idx]
        break_indices = []
        for i, p in enumerate(self.paragraphs):
            found = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        found = True
                        break
                if found:
                    break
            if found:
                break_indices.append(i)
            if len(break_indices) >= 2:
                break

        if len(break_indices) < 1:
            # 无分页符，无法定位第二页
            self.hits.append(("目录两字 中间空一字符/Noto Sans CJK SC 二号 加粗", 0))
            return
        second_page_start = break_indices[0] + 1
        second_page_end = break_indices[1] if len(break_indices) >= 2 else len(self.paragraphs) - 1

        for idx in range(second_page_start, second_page_end + 1):
            if idx >= len(self.para_infos):
                break
            info = self.para_infos[idx]
            raw = info["text"]
            stripped = raw.replace(" ", "").replace("　", "")
            if stripped != "目录":
                continue

            # 点1：中间空一个"汉字宽度"（视觉宽度 ≈ 1 em 即可，不限于 1 个字符）
            #      估算规则：U+3000 全角空格 = 1.0 em；U+0020 半角空格 = 0.5 em；
            #      制表符 \t 视为 1.0 em；其他空白按半角 0.5 em 计。
            #      视觉宽度落在 [0.75, 1.5] em 范围内视为合格。
            m = re.fullmatch(r"\s*目(\s+)录\s*", raw)
            if m is None:
                space_ok = False
            else:
                gap = m.group(1)
                width_em = 0.0
                for ch in gap:
                    if ch == "　":
                        width_em += 1.0
                    elif ch == "\t":
                        width_em += 1.0
                    elif ch == " ":
                        width_em += 0.5
                    else:
                        width_em += 0.5
                space_ok = 0.75 <= width_em <= 1.5

            # 点2/3/4：Noto Sans CJK SC / 二号(22pt) / 加粗
            font_ok = self.font_check(info, "Noto Sans CJK SC",
                                      expected_size=22, expected_bold=True, size_tol=0.5)

            if space_ok and font_ok:
                self.hits.append((
                    "目录两字 中间空一字符/Noto Sans CJK SC 二号 加粗", 0))
            else:
                self.hits.append((
                    "目录两字 中间空一字符/Noto Sans CJK SC 二号 加粗", -1))
            return

        # 第二页未找到"目录"段：不作扣分
        self.hits.append(("目录两字 中间空一字符/Noto Sans CJK SC 二号 加粗", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第二页"目录"两字段落格式不满足
    #    段后20磅、1.08倍行距、居中对齐
    #    一切基于 Word 文件属性（分页符定位第二页 + 段落 w:pPr 的 w:spacing / w:jc）：
    #      - 定位：位于第二页（第 1 个分页符之后、第 2 个分页符之前）
    #              且去除空白后文本为"目录"的段落。
    #      - 点1：段后 = 20 磅  (w:spacing @w:after，20pt=400 twips)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:lineRule="auto" 且 @w:line/240 ≈ 1.08)
    #      - 点3：居中对齐 (w:jc = center)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_toc2_format_pf(self):
        # 用 Word 分页符定位第二页范围
        break_indices = []
        for i, p in enumerate(self.paragraphs):
            found = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        found = True
                        break
                if found:
                    break
            if found:
                break_indices.append(i)
            if len(break_indices) >= 2:
                break

        if len(break_indices) < 1:
            self.hits.append(("目录两字 段落 段后20磅/1.08倍行距/居中", 0))
            return
        second_page_start = break_indices[0] + 1
        second_page_end = break_indices[1] if len(break_indices) >= 2 else len(self.paragraphs) - 1

        for idx in range(second_page_start, second_page_end + 1):
            if idx >= len(self.para_infos):
                break
            info = self.para_infos[idx]
            stripped = info["text"].replace(" ", "").replace("　", "")
            if stripped != "目录":
                continue

            # 点1：段后 20 磅
            # 点2：1.08 倍行距
            # 点3：居中
            if self.pf_check(info, alignment="CENTER",
                             space_after=20, line=1.08):
                self.hits.append((
                    "目录两字 段落 段后20磅/1.08倍行距/居中", 0))
            else:
                self.hits.append((
                    "目录两字 段落 段后20磅/1.08倍行距/居中", -1))
            return

        # 第二页未找到"目录"段：不作扣分
        self.hits.append(("目录两字 段落 段后20磅/1.08倍行距/居中", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第二页目录内容字体格式不满足 Noto Sans CJK SC 11
    #    一切基于 Word 文件属性（分页符定位第二页 + run 的 w:rFonts / w:sz）：
    #      - 定位：位于第二页（第 1 个分页符之后、第 2 个分页符之前），
    #              且为目录内容条目（非"目录"标题的非空段落）。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 11 (11pt)
    #    第二页内所有目录内容段落都必须同时满足点1、点2；
    #    任一不满足 → -1，全部满足 → 0。
    # --------------------------------------------------------
    def eval_toc2_content_format(self):
        # 用 Word 分页符定位第二页范围
        break_indices = []
        for i, p in enumerate(self.paragraphs):
            found = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        found = True
                        break
                if found:
                    break
            if found:
                break_indices.append(i)
            if len(break_indices) >= 2:
                break

        if len(break_indices) < 1:
            self.hits.append(("目录内容字体 Noto Sans CJK SC 11", 0))
            return
        second_page_start = break_indices[0] + 1
        second_page_end = break_indices[1] if len(break_indices) >= 2 else len(self.paragraphs) - 1

        # 收集第二页目录内容条目：排除"目录"标题段落
        toc_items = []
        for idx in range(second_page_start, second_page_end + 1):
            if idx >= len(self.para_infos):
                break
            info = self.para_infos[idx]
            if not info["text"]:
                continue
            stripped = info["text"].replace(" ", "").replace("　", "")
            if stripped == "目录":
                continue
            toc_items.append(info)

        if not toc_items:
            self.hits.append(("目录内容字体 Noto Sans CJK SC 11", 0))
            return

        all_ok = True
        for info in toc_items:
            # 点1：Noto Sans CJK SC
            # 点2：11pt
            if not self.font_check(info, "Noto Sans CJK SC",
                                   expected_size=11, size_tol=0.5):
                all_ok = False
                break

        self.hits.append((
            "目录内容字体 Noto Sans CJK SC 11", 0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第二页目录内容段落格式不满足
    #    段后2磅、1.15倍行距、左对齐
    #    一切基于 Word 文件属性（分页符定位第二页 + 段落 w:pPr 的 w:spacing / w:jc）：
    #      - 定位：位于第二页（第 1 个分页符之后、第 2 个分页符之前），
    #              且为目录内容条目（非"目录"标题的非空段落）。
    #      - 点1：段后 = 2 磅   (w:spacing @w:after，2pt=40 twips)
    #      - 点2：行距 = 1.15 倍 (w:spacing @w:lineRule="auto" 且 @w:line/240 ≈ 1.15)
    #      - 点3：左对齐         (w:jc = left 或未设置默认左对齐)
    #    第二页内所有目录内容段落都必须同时满足点1、点2、点3；
    #    任一不满足 → -1，全部满足 → 0。
    # --------------------------------------------------------
    def eval_toc2_content_format_pf(self):
        # 用 Word 分页符定位第二页范围
        break_indices = []
        for i, p in enumerate(self.paragraphs):
            found = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        found = True
                        break
                if found:
                    break
            if found:
                break_indices.append(i)
            if len(break_indices) >= 2:
                break

        if len(break_indices) < 1:
            self.hits.append(("目录内容段落 段后2磅/1.15倍行距/左对齐", 0))
            return
        second_page_start = break_indices[0] + 1
        second_page_end = break_indices[1] if len(break_indices) >= 2 else len(self.paragraphs) - 1

        # 收集第二页目录内容条目：排除"目录"标题段落
        toc_items = []
        for idx in range(second_page_start, second_page_end + 1):
            if idx >= len(self.para_infos):
                break
            info = self.para_infos[idx]
            if not info["text"]:
                continue
            stripped = info["text"].replace(" ", "").replace("　", "")
            if stripped == "目录":
                continue
            toc_items.append(info)

        if not toc_items:
            self.hits.append(("目录内容段落 段后2磅/1.15倍行距/左对齐", 0))
            return

        all_ok = True
        for info in toc_items:
            # 点3「左对齐」：Word 段落默认左对齐（w:jc 未设置或 = left）
            pf = info["pf"]
            align = pf.get("alignment")
            align_ok = (align is None) or ("LEFT" in align.upper())
            # 点1：段后 2 磅
            # 点2：1.15 倍行距
            if not (align_ok and
                    self.pf_check(info, space_after=2, line=1.15)):
                all_ok = False
                break

        self.hits.append((
            "目录内容段落 段后2磅/1.15倍行距/左对齐", 0 if all_ok else -1))

    # --------------------------------------------------------
    # 第三页 标题 / 表格
    # --------------------------------------------------------
    # --------------------------------------------------------
    # -1 细则：
    #    文档第三页"工程实践与算法应用综合设计评分标准"标题格式
    #    不满足 Noto Sans CJK SC 三号 加粗
    #    一切基于 Word 文件属性（分页符定位第三页 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第三页（第 2 个分页符之后、第 3 个分页符之前）
    #              且文本包含"工程实践与算法应用综合设计评分标准"的段落。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 三号 (16pt)
    #      - 点3：加粗 (w:b)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_page3_title(self):
        # 直接按段落文本属性定位——这份/其他文档的"第三页"不一定与
        # 显式分页符段落索引严格对齐，改用"文本恰为该标题、且非目录条目"
        # 的段落作为待检段落。
        toc_entry_re = re.compile(r"\t\s*\d+\s*$")
        target = "工程实践与算法应用综合设计评分标准"
        for info in self.para_infos:
            t = info["text"]
            if not t:
                continue
            # 排除目录条目（"2. 工程实践与算法应用综合设计评分标准\t3"）
            if toc_entry_re.search(t):
                continue
            # 文本主体为目标标题（去空白后完全匹配）
            s = t.replace(" ", "").replace("　", "").replace("\t", "")
            if s != target:
                continue
            # 点1：Noto Sans CJK SC   点2：三号 = 16pt   点3：加粗
            if self.font_check(info, "Noto Sans CJK SC",
                               expected_size=16, expected_bold=True, size_tol=0.5):
                self.hits.append((
                    "第三页标题字体 Noto Sans CJK SC 三号 加粗", 0))
            else:
                self.hits.append((
                    "第三页标题字体 Noto Sans CJK SC 三号 加粗", -1))
            return
        # 未找到该标题：不作扣分
        self.hits.append(("第三页标题字体 Noto Sans CJK SC 三号 加粗", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第三页"工程实践与算法应用综合设计评分标准"段落格式
    #    不满足 段后8磅、1.08倍行距、居中对齐
    #    一切基于 Word 文件属性（分页符定位第三页 + 段落 w:pPr 的 w:spacing / w:jc）：
    #      - 定位：位于第三页（第 2 个分页符之后、第 3 个分页符之前）
    #              且文本包含"工程实践与算法应用综合设计评分标准"的段落。
    #      - 点1：段后 = 8 磅  (w:spacing @w:after，8pt=160 twips)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:lineRule="auto" 且 @w:line/240 ≈ 1.08)
    #      - 点3：居中对齐 (w:jc = center)
    #    三点全部满足 → 0；任一不满足 → -1。
    # --------------------------------------------------------
    def eval_page3_title_pf(self):
        # 与 eval_page3_title 同：按段落属性定位而不依赖分页符索引
        toc_entry_re = re.compile(r"\t\s*\d+\s*$")
        target = "工程实践与算法应用综合设计评分标准"
        for info in self.para_infos:
            t = info["text"]
            if not t:
                continue
            if toc_entry_re.search(t):
                continue
            s = t.replace(" ", "").replace("　", "").replace("\t", "")
            if s != target:
                continue
            # 点1：段后 8 磅  点2：1.08 倍行距  点3：居中
            if self.pf_check(info, alignment="CENTER",
                             space_after=8, line=1.08):
                self.hits.append((
                    "第三页标题段落 段后8磅/1.08倍行距/居中", 0))
            else:
                self.hits.append((
                    "第三页标题段落 段后8磅/1.08倍行距/居中", -1))
            return
        # 未找到该标题：不作扣分
        self.hits.append(("第三页标题段落 段后8磅/1.08倍行距/居中", 0))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第三页、第四页表格第一行字体格式
    #    不满足 Noto Sans CJK SC 9.5 加粗
    #    一切基于 Word 文件属性（分页符定位第三/四页 + 表格首行 run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第三页或第四页（第 2/3 分页符之后、第 3/4 分页符之前）
    #              的所有表格的首行 (tbl.rows[0]) 中的非空文本 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 9.5 (9.5pt)
    #      - 点3：加粗 (w:b)
    #    首行所有非空 run 都必须同时满足三点；
    #    任一不满足 → -1，全部满足 → 0。
    # --------------------------------------------------------
    def eval_page3_table_header_font(self):
        # 收集前 4 个分页符位置
        break_indices = []
        for i, p in enumerate(self.paragraphs):
            found = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        found = True
                        break
                if found:
                    break
            if found:
                break_indices.append(i)
            if len(break_indices) >= 4:
                break

        if len(break_indices) < 2:
            self.hits.append((
                "第三/四页表格首行字体 Noto Sans CJK SC 9.5 加粗", 0))
            return

        # 第三页范围
        p3_start = break_indices[1] + 1
        p3_end = break_indices[2] if len(break_indices) >= 3 else len(self.paragraphs) - 1
        # 第四页范围
        p4_start = (break_indices[2] + 1) if len(break_indices) >= 3 else None
        p4_end = break_indices[3] if len(break_indices) >= 4 else (len(self.paragraphs) - 1 if p4_start is not None else None)

        def para_in_range(p_element, start, end):
            """判断 tbl 前一个 <w:p> 段落是否在 [start, end] 范围内"""
            if start is None or end is None:
                return False
            for idx in range(start, end + 1):
                if idx < len(self.paragraphs) and self.paragraphs[idx]._p is p_element:
                    return True
            return False

        # 收集第三/四页范围内的表格
        target_tables = []
        for tbl in self.doc.tables:
            # 找到 tbl 之前最近的 <w:p> 段落，判断其归属页
            prev = tbl._element.getprevious()
            while prev is not None and prev.tag != qn("w:p"):
                prev = prev.getprevious()
            if prev is None:
                continue
            if para_in_range(prev, p3_start, p3_end) or para_in_range(prev, p4_start, p4_end):
                target_tables.append(tbl)

        if not target_tables:
            self.hits.append((
                "第三/四页表格首行字体 Noto Sans CJK SC 9.5 加粗", 0))
            return

        all_ok = True
        for tbl in target_tables:
            try:
                first_row = tbl.rows[0]
            except Exception:
                continue
            for cell in first_row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not r.text.strip():
                            continue
                        fn, sz, b = get_run_font(r)
                        # 点1：Noto Sans CJK SC
                        if not (fn and "Noto Sans CJK SC" in fn):
                            all_ok = False
                            break
                        # 点2：9.5pt
                        if not (sz is not None and pt_approx(sz, 9.5, tol=0.3)):
                            all_ok = False
                            break
                        # 点3：加粗
                        if not b:
                            all_ok = False
                            break
                    if not all_ok:
                        break
                if not all_ok:
                    break
            if not all_ok:
                break

        self.hits.append((
            "第三/四页表格首行字体 Noto Sans CJK SC 9.5 加粗",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第三页表格除第一行外其他字体格式
    #    不满足 Noto Sans CJK SC 8，一、二、四列字体需额外加粗
    #    一切基于 Word 文件属性（分页符定位第三页 + 表格非首行 run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第三页（第 2 分页符之后、第 3 分页符之前）
    #              的所有表格的除首行外所有行 (tbl.rows[1:]) 中的非空文本 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 8 (8pt)
    #      - 点3：第一、二、四列（列索引 0、1、3）额外要求加粗 (w:b)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page3_table_other_font(self):
        # 收集前 3 个分页符位置（定位第三页需要第 2/3 分页符）
        break_indices = []
        for i, p in enumerate(self.paragraphs):
            found = False
            for r in p.runs:
                for br in r._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        found = True
                        break
                if found:
                    break
            if found:
                break_indices.append(i)
            if len(break_indices) >= 3:
                break

        if len(break_indices) < 2:
            self.hits.append((
                "第三页表格其他行字体 Noto Sans CJK SC 8（一二四列加粗）", 0))
            return

        p3_start = break_indices[1] + 1
        p3_end = break_indices[2] if len(break_indices) >= 3 else len(self.paragraphs) - 1

        def para_in_range(p_element, start, end):
            if start is None or end is None:
                return False
            for idx in range(start, end + 1):
                if idx < len(self.paragraphs) and self.paragraphs[idx]._p is p_element:
                    return True
            return False

        # 收集第三页范围内的表格
        target_tables = []
        for tbl in self.doc.tables:
            prev = tbl._element.getprevious()
            while prev is not None and prev.tag != qn("w:p"):
                prev = prev.getprevious()
            if prev is None:
                continue
            if para_in_range(prev, p3_start, p3_end):
                target_tables.append(tbl)

        if not target_tables:
            self.hits.append((
                "第三页表格其他行字体 Noto Sans CJK SC 8（一二四列加粗）", 0))
            return

        all_ok = True
        for tbl in target_tables:
            try:
                rows = tbl.rows
            except Exception:
                continue
            # 除第一行外
            for row in rows[1:]:
                for ci, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if not r.text.strip():
                                continue
                            fn, sz, b = get_run_font(r)
                            # 点1：Noto Sans CJK SC
                            if not (fn and "Noto Sans CJK SC" in fn):
                                all_ok = False
                                break
                            # 点2：8pt
                            if not (sz is not None and pt_approx(sz, 8, tol=0.3)):
                                all_ok = False
                                break
                            # 点3：一、二、四列（索引 0、1、3）额外加粗
                            if ci in (0, 1, 3) and not b:
                                all_ok = False
                                break
                        if not all_ok:
                            break
                    if not all_ok:
                        break
                if not all_ok:
                    break
            if not all_ok:
                break

        self.hits.append((
            "第三页表格其他行字体 Noto Sans CJK SC 8（一二四列加粗）",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # 第四页：过程性评价与备选方向 / 三综合设计备选方向
    # --------------------------------------------------------
    # --------------------------------------------------------
    # -1 细则：
    #    文档第四页"过程性评价与备选方向"
    #    字体格式不满足 Noto Sans CJK SC 三号 加粗
    #    一切基于 Word 文件属性（分页符定位第四页 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第四页（第 3 分页符之后、第 4 分页符之前）
    #              且文本包含"过程性评价与备选方向"的段落中的非空 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 三号 (16pt)
    #      - 点3：加粗 (w:b)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page4_gc_title_font(self):
        target = "过程性评价与备选方向"
        target_info = None
        for info in self.para_infos:
            s = info["text"].replace(" ", "").replace("　", "").replace("\t", "")
            if s == target:
                target_info = info
                break

        if target_info is None:
            self.hits.append((
                "第四页过程性评价标题字体 Noto Sans CJK SC 三号加粗", 0))
            return

        all_ok = True
        for r in target_info["para"].runs:
            if not r.text.strip():
                continue
            fn, sz, b = get_run_font(r)
            # 点1：Noto Sans CJK SC
            if not (fn and "Noto Sans CJK SC" in fn):
                all_ok = False
                break
            # 点2：三号 = 16pt
            if not (sz is not None and pt_approx(sz, 16, tol=0.3)):
                all_ok = False
                break
            # 点3：加粗
            if not b:
                all_ok = False
                break

        self.hits.append((
            "第四页过程性评价标题字体 Noto Sans CJK SC 三号加粗",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第四页"过程性评价与备选方向"
    #    段落格式不满足段后 8 磅、1.08 倍行距、居中对齐
    #    一切基于 Word 文件属性（分页符定位第四页 + <w:pPr> 的 w:spacing / w:jc）：
    #      - 定位：位于第四页（第 3 分页符之后、第 4 分页符之前）
    #              且文本包含"过程性评价与备选方向"的段落。
    #      - 点1：段后 = 8 磅 (w:spacing @w:after)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:line / @w:lineRule)
    #      - 点3：居中对齐 (w:jc = "center")
    #    任一不满足 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page4_gc_title_pf(self):
        target = "过程性评价与备选方向"
        target_info = None
        for info in self.para_infos:
            s = info["text"].replace(" ", "").replace("　", "").replace("\t", "")
            if s == target:
                target_info = info
                break

        if target_info is None:
            self.hits.append((
                "第四页过程性评价标题段落 段后8磅/1.08倍行距/居中", 0))
            return

        if self.pf_check(target_info, alignment="CENTER", space_after=8, line=1.08):
            self.hits.append((
                "第四页过程性评价标题段落 段后8磅/1.08倍行距/居中", 0))
        else:
            self.hits.append((
                "第四页过程性评价标题段落 段后8磅/1.08倍行距/居中", -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第四页"三、综合设计备选方向"
    #    字体格式不满足 Noto Sans CJK SC 13 加粗
    #    一切基于 Word 文件属性（分页符定位第四页 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第四页（第 3 分页符之后、第 4 分页符之前）
    #              且文本包含"三、综合设计备选方向"的段落中的非空 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 13 (13pt)
    #      - 点3：加粗 (w:b)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page4_sanfangxiang_font(self):
        target = "三、综合设计备选方向"
        target_info = None
        for info in self.para_infos:
            s = info["text"].replace(" ", "").replace("　", "").replace("\t", "")
            if s == target:
                target_info = info
                break

        if target_info is None:
            self.hits.append((
                "第四页三综合字体 Noto Sans CJK SC 13加粗", 0))
            return

        all_ok = True
        for r in target_info["para"].runs:
            if not r.text.strip():
                continue
            fn, sz, b = get_run_font(r)
            # 点1：Noto Sans CJK SC
            if not (fn and "Noto Sans CJK SC" in fn):
                all_ok = False
                break
            # 点2：13pt
            if not (sz is not None and pt_approx(sz, 13, tol=0.3)):
                all_ok = False
                break
            # 点3：加粗
            if not b:
                all_ok = False
                break

        self.hits.append((
            "第四页三综合字体 Noto Sans CJK SC 13加粗",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第四页"三、综合设计备选方向"
    #    段落格式不满足段后 4 磅、1.08 倍行距、左对齐
    #    一切基于 Word 文件属性（分页符定位第四页 + <w:pPr> 的 w:spacing / w:jc）：
    #      - 定位：位于第四页（第 3 分页符之后、第 4 分页符之前）
    #              且文本包含"三、综合设计备选方向"的段落。
    #      - 点1：段后 = 4 磅 (w:spacing @w:after)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:line / @w:lineRule)
    #      - 点3：左对齐 (w:jc = "left" 或未设置默认左对齐)
    #    任一不满足 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page4_sanfangxiang_pf(self):
        target = "三、综合设计备选方向"
        target_info = None
        for info in self.para_infos:
            s = info["text"].replace(" ", "").replace("　", "").replace("\t", "")
            if s == target:
                target_info = info
                break

        if target_info is None:
            self.hits.append((
                "第四页三综合段落 段后4磅/1.08倍行距/左对齐", 0))
            return

        if self.pf_check(target_info, alignment="LEFT", space_after=4, line=1.08):
            self.hits.append((
                "第四页三综合段落 段后4磅/1.08倍行距/左对齐", 0))
        else:
            self.hits.append((
                "第四页三综合段落 段后4磅/1.08倍行距/左对齐", -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第四页"三、综合设计备选方向"下方内容
    #    除"实施安排：第 1 周完成选题与需求边界，第 2 周完成方案设计，
    #    第 3 周完成核心实现，第 4 周完成测试、文档与答辩。"外
    #    其余字体格式不满足 Noto Sans CJK SC 10
    #    一切基于 Word 文件属性（分页符定位第四页 + run 的 w:rFonts / w:sz）：
    #      - 定位：位于第四页（第 3 分页符之后、第 4 分页符之前）
    #              且位于"三、综合设计备选方向"段落之后、"实施安排：……"段落
    #              之外的段落中的非空 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 10 (10pt)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page4_sanfangxiang_content_font(self):
        # 找到"三、综合设计备选方向"标题段落索引
        title_target = "三、综合设计备选方向"
        title_idx = None
        for i, info in enumerate(self.para_infos):
            s = info["text"].replace(" ", "").replace("　", "").replace("\t", "")
            if s == title_target:
                title_idx = i
                break

        if title_idx is None:
            self.hits.append((
                "第四页三综合下内容字体 Noto Sans CJK SC 10", 0))
            return

        # 内容区间：标题之后，直到下一个同/上级标题为止
        # 停止条件：段落文本形如 "X、..."（同级）/ "\d+\..."（父级 "1. xxx"）/
        #           "第X章..."（章级）——这些都是标题
        heading_re = re.compile(
            r"^(?:第\s*[一二三四五六七八九十百零0-9]+\s*章"
            r"|[一二三四五六七八九十]+、"
            r"|\d+\.\s*\S+)"
        )
        # 免检段落："实施安排：……" 整段排除在字体检查范围外
        exempt_text = (
            "实施安排：第1周完成选题与需求边界，第2周完成方案设计，"
            "第3周完成核心实现，第4周完成测试、文档与答辩。"
        )

        all_ok = True
        found_any = False
        for idx in range(title_idx + 1, len(self.para_infos)):
            info = self.para_infos[idx]
            t = info["text"].strip()
            if not t:
                continue
            # 遇到下一段标题就停止
            if heading_re.match(t):
                break
            # 免检段落跳过
            t_norm = t.replace(" ", "").replace("　", "")
            if t_norm == exempt_text:
                continue
            for r in info["para"].runs:
                if not r.text.strip():
                    continue
                found_any = True
                fn, sz, b = get_run_font(r)
                # 点1：Noto Sans CJK SC
                if not (fn and "Noto Sans CJK SC" in fn):
                    all_ok = False
                    break
                # 点2：10pt
                if not (sz is not None and pt_approx(sz, 10, tol=0.3)):
                    all_ok = False
                    break
            if not all_ok:
                break

        if not found_any:
            self.hits.append((
                "第四页三综合下内容字体 Noto Sans CJK SC 10", 0))
            return

        self.hits.append((
            "第四页三综合下内容字体 Noto Sans CJK SC 10",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # 第五页 1.成果文本编制要求 / 一、报告内容 / 二、打印要求
    # --------------------------------------------------------
    # --------------------------------------------------------
    # -1 细则：
    #    文档第五页"1.成果文本编制要求"
    #    字体格式不满足 Noto Sans CJK SC 三号 加粗
    #    一切基于 Word 文件属性（分页符定位第五页 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第五页（第 4 分页符之后、第 5 分页符之前）
    #              且文本包含"1.成果文本编制要求"（"1." + "成果文本编制要求"）的段落中的非空 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 三号 (16pt)
    #      - 点3：加粗 (w:b)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page5_chengguo_title_font(self):
        # "1. 成果文本编制要求" 作为章节标题的段落——排除目录条目
        toc_entry_re = re.compile(r"\t\s*\d+\s*$")
        target_info = None
        for info in self.para_infos:
            t = info["text"]
            if not t:
                continue
            if toc_entry_re.search(t):
                continue
            s = t.replace(" ", "").replace("　", "").replace("\t", "")
            # "1.成果文本编制要求" 或 "1．成果文本编制要求"
            if s == "1.成果文本编制要求" or s == "1．成果文本编制要求":
                target_info = info
                break

        if target_info is None:
            self.hits.append((
                "第五页成果编制要求字体 Noto Sans CJK SC 三号加粗", 0))
            return

        all_ok = True
        for r in target_info["para"].runs:
            if not r.text.strip():
                continue
            fn, sz, b = get_run_font(r)
            # 点1：Noto Sans CJK SC
            if not (fn and "Noto Sans CJK SC" in fn):
                all_ok = False
                break
            # 点2：三号 = 16pt
            if not (sz is not None and pt_approx(sz, 16, tol=0.3)):
                all_ok = False
                break
            # 点3：加粗
            if not b:
                all_ok = False
                break

        self.hits.append((
            "第五页成果编制要求字体 Noto Sans CJK SC 三号加粗",
            0 if all_ok else -1))

    def eval_page5_chengguo_title_pf(self):
        # 与 eval_page5_chengguo_title_font 一致：排除目录条目，
        # 只定位第五页正文里的"1. 成果文本编制要求"标题段落。
        toc_entry_re = re.compile(r"\t\s*\d+\s*$")
        target_info = None
        for info in self.para_infos:
            t = info["text"]
            if not t:
                continue
            if toc_entry_re.search(t):
                continue
            s = t.replace(" ", "").replace("　", "").replace("\t", "")
            if s == "1.成果文本编制要求" or s == "1．成果文本编制要求":
                target_info = info
                break

        if target_info is None:
            self.hits.append(("第五页成果编制要求段落 段后8磅/1.08倍行距/居中", 0))
            return

        if self.pf_check(target_info, alignment="CENTER", space_after=8, line=1.08):
            self.hits.append(("第五页成果编制要求段落 段后8磅/1.08倍行距/居中", 0))
        else:
            self.hits.append(("第五页成果编制要求段落 段后8磅/1.08倍行距/居中", -1))

    # --------------------------------------------------------
    # -3 细则：
    #    文档第五页和第六页"一、报告内容"、"二、打印要求"
    #    标题及下方内容段落格式不满足 1.15 倍行距、段后 10 磅、左对齐
    #    一切基于 Word 文件属性（分页符定位第五/六页 + <w:pPr> 的 w:spacing / w:jc）：
    #      - 定位：位于第五页（第 4 分页符之后、第 5 分页符之前）或
    #              第六页（第 5 分页符之后、第 6 分页符之前）范围。
    #              标题段落 = 文本以"一、报告内容"或"二、打印要求"开头。
    #              下方内容 = 标题段落之后、下一个"一、"/"二、"标题之前、仍在第五/六页范围内的非空段落。
    #      - 点1：行距 = 1.15 倍 (w:spacing @w:line / @w:lineRule)
    #      - 点2：段后 = 10 磅 (w:spacing @w:after)
    #      - 点3：左对齐 (w:jc = "left" 或未设置默认左对齐)
    #    任一段落违反任一点 → -3；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page5_yi_er_content(self):
        heading_re = re.compile(
            r"^(?:第\s*[一二三四五六七八九十百零0-9]+\s*章"
            r"|[一二三四五六七八九十]+、"
            r"|\d+\.\s*\S+)"
        )
        # "成果提交与附录要求"已属于第七页内容，作为额外的区间终止边界，
        # 不纳入第五/六页"一、报告内容"/"二、打印要求"下方内容的检查范围。
        extra_stop_targets = ("成果提交与附录要求",)
        title_targets = ("一、报告内容", "二、打印要求")

        title_idx_list = []
        for i, info in enumerate(self.para_infos):
            t = info["text"]
            if t and t.startswith(title_targets):
                title_idx_list.append(i)

        if not title_idx_list:
            self.hits.append((
                "第五/六页一、二 标题及内容段落 1.15倍行距/段后10磅/左对齐", 0))
            return

        all_ok = True

        # 校验标题段落格式
        for tidx in title_idx_list:
            info = self.para_infos[tidx]
            if not self.pf_check(info, alignment="LEFT", space_after=10, line=1.15):
                all_ok = False
                break

        # 校验下方内容段落格式
        if all_ok:
            for tidx in title_idx_list:
                for jdx in range(tidx + 1, len(self.para_infos)):
                    info = self.para_infos[jdx]
                    t = info["text"].strip()
                    if not t:
                        continue
                    if heading_re.match(t) or t in extra_stop_targets:
                        break
                    if not self.pf_check(info, alignment="LEFT", space_after=10, line=1.15):
                        all_ok = False
                        break
                if not all_ok:
                    break

        self.hits.append((
            "第五/六页一、二 标题及内容段落 1.15倍行距/段后10磅/左对齐",
            0 if all_ok else -3))

    # --------------------------------------------------------
    # 第七页 成果提交与附录要求 / 一、二、三
    # --------------------------------------------------------
    # --------------------------------------------------------
    # -1 细则：
    #    文档第七页"成果提交与附录要求"
    #    字体格式不满足 Noto Sans CJK SC 三号 加粗
    #    一切基于 Word 文件属性（分页符定位第七页 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第七页（第 6 分页符之后、第 7 分页符之前）
    #              且文本包含"成果提交与附录要求"的段落中的非空 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 三号 (16pt)
    #      - 点3：加粗 (w:b)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page7_chengguo_title_font(self):
        target = "成果提交与附录要求"
        target_info = None
        for info in self.para_infos:
            s = info["text"].replace(" ", "").replace("　", "").replace("\t", "")
            if s == target:
                target_info = info
                break

        if target_info is None:
            self.hits.append((
                "第七页成果提交与附录字体 Noto Sans CJK SC 三号加粗", 0))
            return

        all_ok = True
        for r in target_info["para"].runs:
            if not r.text.strip():
                continue
            fn, sz, b = get_run_font(r)
            if not (fn and "Noto Sans CJK SC" in fn):
                all_ok = False; break
            if not (sz is not None and pt_approx(sz, 16, tol=0.3)):
                all_ok = False; break
            if not b:
                all_ok = False; break

        self.hits.append((
            "第七页成果提交与附录字体 Noto Sans CJK SC 三号加粗",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第七页"成果提交与附录要求"
    #    段落格式不满足段后 8 磅、1.08 倍行距、居中对齐
    #    一切基于 Word 文件属性（分页符定位第七页 + <w:pPr> 的 w:spacing / w:jc）：
    #      - 定位：文本包含"成果提交与附录要求"的段落。
    #      - 点1：段后 = 8 磅 (w:spacing @w:after)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:line / @w:lineRule)
    #      - 点3：居中对齐 (w:jc = "center")
    #    任一不满足 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page7_chengguo_title_pf(self):
        target = "成果提交与附录要求"
        target_info = None
        for info in self.para_infos:
            s = info["text"].replace(" ", "").replace("　", "").replace("\t", "")
            if s == target:
                target_info = info
                break

        if target_info is None:
            self.hits.append((
                "第七页成果提交与附录段落 段后8磅/1.08倍行距/居中", 0))
            return

        if self.pf_check(target_info, alignment="CENTER", space_after=8, line=1.08):
            self.hits.append((
                "第七页成果提交与附录段落 段后8磅/1.08倍行距/居中", 0))
        else:
            self.hits.append((
                "第七页成果提交与附录段落 段后8磅/1.08倍行距/居中", -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第七页标题"程序与附件"、"诚信与协作"、"提交清单"
    #    字体格式不满足 Noto Sans CJK SC 小四 加粗
    #    一切基于 Word 文件属性（分页符定位第七页 + run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第七页（第 6 分页符之后、第 7 分页符之前）
    #              且文本以"一、程序与附件"/"二、诚信与协作"/"三、提交清单"开头的段落。
    #              仅检查标题主体"程序与附件"/"诚信与协作"/"提交清单"，不检查前置序号"一、"/"二、"/"三、"。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 小四 (12pt)
    #      - 点3：加粗 (w:b)
    #    标题主体中任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page7_yier_subtitle(self):
        title_prefixes = ("一、程序与附件", "二、诚信与协作", "三、提交清单")
        # 标题主体（不含前置序号"一、"/"二、"/"三、"）
        body_texts = ("程序与附件", "诚信与协作", "提交清单")
        target_infos = []
        for info in self.para_infos:
            t = info["text"]
            if t and t.startswith(title_prefixes):
                target_infos.append(info)

        if not target_infos:
            self.hits.append((
                "第七页标题字体 Noto Sans CJK SC小四加粗", 0))
            return

        all_ok = True
        for info in target_infos:
            # 按字符位置遍历 run，跳过属于前置序号"一、"/"二、"/"三、"（共2个字符）的部分，
            # 只检查标题主体文字对应的 run。
            char_pos = 0
            for r in info["para"].runs:
                run_len = len(r.text)
                run_start = char_pos
                char_pos += run_len
                if not r.text.strip():
                    continue
                # run 完全落在前置序号范围内（前2个字符）则跳过
                if run_start < 2 and run_start + run_len <= 2:
                    continue
                fn, sz, b = get_run_font(r)
                if not (fn and "Noto Sans CJK SC" in fn):
                    all_ok = False; break
                if not (sz is not None and pt_approx(sz, 12, tol=0.3)):
                    all_ok = False; break
                if not b:
                    all_ok = False; break
            if not all_ok:
                break

        self.hits.append((
            "第七页标题字体 Noto Sans CJK SC小四加粗",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第七页"一、程序与附件"、"二、诚信与协作"、"三、提交清单"
    #    段落格式不满足段后 2 磅、1.08 倍行距、左对齐
    #    一切基于 Word 文件属性（分页符定位第七页 + <w:pPr> 的 w:spacing / w:jc）：
    #      - 定位：位于第七页（第 6 分页符之后、第 7 分页符之前）
    #              且文本以"一、程序与附件"/"二、诚信与协作"/"三、提交清单"开头的段落。
    #      - 点1：段后 = 2 磅 (w:spacing @w:after)
    #      - 点2：行距 = 1.08 倍 (w:spacing @w:line / @w:lineRule)
    #      - 点3：左对齐 (w:jc = "left" 或未设置默认左对齐)
    #    任一段落违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page7_yier_subtitle_pf(self):
        title_prefixes = ("一、程序与附件", "二、诚信与协作", "三、提交清单")
        target_infos = []
        for info in self.para_infos:
            t = info["text"]
            if t and t.startswith(title_prefixes):
                target_infos.append(info)

        if not target_infos:
            self.hits.append((
                "第七页一二三小标题段落 段后2磅/1.08倍行距/左对齐", 0))
            return

        all_ok = True
        for info in target_infos:
            if not self.pf_check(info, alignment="LEFT", space_after=2, line=1.08):
                all_ok = False
                break

        self.hits.append((
            "第七页一二三小标题段落 段后2磅/1.08倍行距/左对齐",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第七页"一、程序与附件"、"二、诚信与协作"
    #    下方内容字体格式不满足 Noto Sans CJK SC 五号
    #    一切基于 Word 文件属性（分页符定位第七页 + run 的 w:rFonts / w:sz）：
    #      - 定位：位于第七页（第 6 分页符之后、第 7 分页符之前），
    #              "一、程序与附件"/"二、诚信与协作"标题之后、下一个"一、"/"二、"/"三、"标题之前、
    #              且仍在第七页范围内的所有非空段落中的非空 run。
    #      - 点1：字体 = Noto Sans CJK SC
    #      - 点2：字号 = 五号 (10.5pt)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page7_yier_content_font(self):
        # 找到"一、程序与附件" / "二、诚信与协作"标题段落索引
        title_prefixes = ("一、程序与附件", "二、诚信与协作")
        stop_prefixes = ("一、", "二、", "三、")
        heading_re = re.compile(
            r"^(?:第\s*[一二三四五六七八九十百零0-9]+\s*章"
            + r"|\d+\.\s*\S+)"
        )
        title_idx_list = []
        for i, info in enumerate(self.para_infos):
            t = info["text"]
            if t and t.startswith(title_prefixes):
                title_idx_list.append(i)

        if not title_idx_list:
            self.hits.append((
                "第七页一二内容字体 Noto Sans CJK SC 五号", 0))
            return

        all_ok = True
        found_any = False
        for tidx in title_idx_list:
            # 内容截止：下一段本级标题（"一、"/"二、"/"三、"）或上级标题
            for jdx in range(tidx + 1, len(self.para_infos)):
                info = self.para_infos[jdx]
                t = info["text"].strip()
                if not t:
                    continue
                if t.startswith(stop_prefixes) or heading_re.match(t):
                    break
                for r in self.paragraphs[jdx].runs:
                    if not r.text.strip():
                        continue
                    found_any = True
                    fn, sz, b = get_run_font(r)
                    if not (fn and "Noto Sans CJK SC" in fn):
                        all_ok = False; break
                    if not (sz is not None and pt_approx(sz, 10.5, tol=0.3)):
                        all_ok = False; break
                if not all_ok:
                    break
            if not all_ok:
                break

        if not found_any:
            self.hits.append((
                "第七页一二内容字体 Noto Sans CJK SC 五号", 0))
            return

        self.hits.append((
            "第七页一二内容字体 Noto Sans CJK SC 五号",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第七页"一、程序与附件"、"二、诚信与协作"
    #    下方内容段落格式不满足段后 1.5 磅、段前缩进 0.25 厘米、悬挂缩进 0.25 厘米、1.05 倍行距
    #    一切基于 Word 文件属性（分页符定位第七页 + <w:pPr> 的 w:spacing / w:ind）：
    #      - 定位：位于第七页（第 6 分页符之后、第 7 分页符之前），
    #              "一、程序与附件"/"二、诚信与协作"标题之后、下一个"一、"/"二、"/"三、"标题之前、
    #              且仍在第七页范围内的所有非空段落。
    #      - 点1：段后 = 1.5 磅 (w:spacing @w:after)
    #      - 点2：段前缩进（首行缩进） = 0.25 厘米 (w:ind @w:firstLine)
    #      - 点3：悬挂缩进 = 0.25 厘米 (w:ind @w:hanging)
    #      - 点4：行距 = 1.05 倍 (w:spacing @w:line / @w:lineRule)
    #    任一段落违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page7_yier_content_pf(self):
        title_prefixes = ("一、程序与附件", "二、诚信与协作")
        stop_prefixes = ("一、", "二、", "三、")
        heading_re = re.compile(
            r"^(?:第\s*[一二三四五六七八九十百零0-9]+\s*章"
            + r"|\d+\.\s*\S+)"
        )
        title_idx_list = []
        for i, info in enumerate(self.para_infos):
            t = info["text"]
            if t and t.startswith(title_prefixes):
                title_idx_list.append(i)

        if not title_idx_list:
            self.hits.append((
                "第七页一二内容段落 段后1.5磅/段前缩进0.25cm/悬挂缩进0.25cm/1.05倍行距", 0))
            return

        all_ok = True
        found_any = False
        for tidx in title_idx_list:
            for jdx in range(tidx + 1, len(self.para_infos)):
                info = self.para_infos[jdx]
                t = info["text"].strip()
                if not t:
                    continue
                if t.startswith(stop_prefixes) or heading_re.match(t):
                    break
                found_any = True
                # “段前缩进0.25cm”指 WPS/Word 段落对话框中显示的首行绝对
                # 缩进位置（= left_indent - hanging_indent），而不是
                # python-docx 的 first_line_indent 原始值（该值在只设置
                # w:hanging、未设置 w:firstLine 时返回相对左缩进的负偏移，
                # 衡量的是另一个量）。
                if not self.pf_check(
                    info,
                    space_after=1.5,
                    first_line_pos_cm=0.25,
                    hanging_indent_cm=0.25,
                    line=1.05,
                ):
                    all_ok = False
                    break
            if not all_ok:
                break

        if not found_any:
            self.hits.append((
                "第七页一二内容段落 段后1.5磅/段前缩进0.25cm/悬挂缩进0.25cm/1.05倍行距", 0))
            return

        self.hits.append((
            "第七页一二内容段落 段后1.5磅/段前缩进0.25cm/悬挂缩进0.25cm/1.05倍行距",
            0 if all_ok else -1))

    # --------------------------------------------------------
    # -1 细则：
    #    文档第七页"三、提交清单"下方表格
    #    字体格式不满足：第一行为 Noto Sans CJK SC 小五 加粗、
    #                    其余字体为 Noto Sans CJK SC 10
    #    一切基于 Word 文件属性（分页符定位第七页 + 表格 run 的 w:rFonts / w:sz / w:b）：
    #      - 定位：位于第七页（第 6 分页符之后、第 7 分页符之前），
    #              紧跟在"三、提交清单"标题段落之后（同处第七页范围内）的表格。
    #      - 第一行：
    #          点1：字体 = Noto Sans CJK SC
    #          点2：字号 = 小五 (9pt)
    #          点3：加粗 (w:b)
    #      - 其余行：
    #          点1：字体 = Noto Sans CJK SC
    #          点2：字号 = 10 (10pt)
    #    任一 run 违反任一点 → -1；全部满足 → 0。
    # --------------------------------------------------------
    def eval_page7_san_table_font(self):
        # 找到"三、提交清单"标题段落
        title_element = None
        for i, info in enumerate(self.para_infos):
            t = info["text"]
            if t and t.startswith("三、提交清单"):
                title_element = self.paragraphs[i]._p
                break

        if title_element is None:
            self.hits.append((
                "第七页三、提交清单 表格字体 首行小五加粗/其余10", 0))
            return

        # 找到位于"三、提交清单"标题之后的第一个表格
        target_tbl = None
        for tbl in self.doc.tables:
            prev = tbl._element.getprevious()
            while prev is not None and prev.tag != qn("w:p"):
                prev = prev.getprevious()
            if prev is None:
                continue
            # 表格前置段落必须位于标题之后
            if title_element is prev or self._element_before(title_element, prev):
                target_tbl = tbl
                break

        if target_tbl is None:
            self.hits.append((
                "第七页三、提交清单 表格字体 首行小五加粗/其余10", 0))
            return

        all_ok = True
        rows = target_tbl.rows

        # 第一行：Noto Sans CJK SC + 小五(9pt) + 加粗
        if len(rows) >= 1:
            for cell in rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not r.text.strip():
                            continue
                        fn, sz, b = get_run_font(r)
                        if not (fn and "Noto Sans CJK SC" in fn):
                            all_ok = False; break
                        if not (sz is not None and pt_approx(sz, 9, tol=0.3)):
                            all_ok = False; break
                        if not b:
                            all_ok = False; break
                    if not all_ok:
                        break
                if not all_ok:
                    break

        # 其余行：Noto Sans CJK SC + 10pt
        if all_ok and len(rows) > 1:
            for row in rows[1:]:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if not r.text.strip():
                                continue
                            fn, sz, b = get_run_font(r)
                            if not (fn and "Noto Sans CJK SC" in fn):
                                all_ok = False; break
                            if not (sz is not None and pt_approx(sz, 10, tol=0.3)):
                                all_ok = False; break
                        if not all_ok:
                            break
                    if not all_ok:
                        break
                if not all_ok:
                    break

        self.hits.append((
            "第七页三、提交清单 表格字体 首行小五加粗/其余10",
            0 if all_ok else -1))

    def _element_before(self, a, b):
        # 判断 element a 是否在 element b 之前（文档顺序）
        for el in b.itersiblings(preceding=True):
            if el is a:
                return True
        return False

    # --------------------------------------------------------
    # 第八页（封面附录）— 实际是文档中的"封皮"在第八页
    # 这里指文档末尾的封皮（云岭数字技术学院等）
    # --------------------------------------------------------
    def eval_page8_college_font(self):
        # -1 细则：第八页"云岭数字技术学院"字体 Noto Sans CJK SC 二号(22pt) 加粗
        # 按属性定位：文档中所有文本 == "云岭数字技术学院" 的段落，取最后一处
        # （封面首页和封底两处都有；封底是我们要检查的"第八页"位置）
        target_idx = -1
        for i, info in enumerate(self.para_infos):
            if info["text"] == "云岭数字技术学院":
                target_idx = i  # 循环到最后 → 取最后一次匹配

        if target_idx < 0:
            self.hits.append((
                "第八页云岭数字技术学院字体 Noto Sans CJK SC 二号加粗", 0))
            return

        all_ok = True
        found_any = False
        for r in self.paragraphs[target_idx].runs:
            if not r.text.strip():
                continue
            found_any = True
            fn, sz, b = get_run_font(r)
            if not (fn and "Noto Sans CJK SC" in fn):
                all_ok = False; break
            if not (sz is not None and pt_approx(sz, 22, tol=0.3)):
                all_ok = False; break
            if not b:
                all_ok = False; break

        if not found_any:
            self.hits.append((
                "第八页云岭数字技术学院字体 Noto Sans CJK SC 二号加粗", 0))
            return

        self.hits.append((
            "第八页云岭数字技术学院字体 Noto Sans CJK SC 二号加粗",
            0 if all_ok else -1))

    def eval_page8_college_pf(self):
        # -1 细则：第八页"云岭数字技术学院"段落 段后26磅/1.08倍行距/居中
        target_idx = -1
        for i, info in enumerate(self.para_infos):
            if info["text"] == "云岭数字技术学院":
                target_idx = i

        if target_idx < 0:
            self.hits.append((
                "第八页云岭数字技术学院段落 段后26磅/1.08倍行距/居中", 0))
            return

        info = self.para_infos[target_idx]
        ok = self.pf_check(info, alignment="CENTER", space_after=26, line=1.08)

        self.hits.append((
            "第八页云岭数字技术学院段落 段后26磅/1.08倍行距/居中",
            0 if ok else -1))

    def eval_page8_jiaocheng_font(self):
        # -1 细则：第八页"《工程实践与算法应用综合设计》"字体 Noto Sans CJK SC 小二(18pt) 加粗
        target_idx = -1
        for i, info in enumerate(self.para_infos):
            if "《工程实践与算法应用综合设计》" in info["text"]:
                target_idx = i
                break

        if target_idx < 0:
            self.hits.append((
                "第八页《工程实践与算法应用综合设计》字体 Noto Sans CJK SC 小二加粗", 0))
            return

        all_ok = True
        found_any = False
        for r in self.paragraphs[target_idx].runs:
            if not r.text.strip():
                continue
            found_any = True
            fn, sz, b = get_run_font(r)
            if not (fn and "Noto Sans CJK SC" in fn):
                all_ok = False; break
            if not (sz is not None and pt_approx(sz, 18, tol=0.3)):
                all_ok = False; break
            if not b:
                all_ok = False; break

        if not found_any:
            self.hits.append((
                "第八页《工程实践与算法应用综合设计》字体 Noto Sans CJK SC 小二加粗", 0))
            return

        self.hits.append((
            "第八页《工程实践与算法应用综合设计》字体 Noto Sans CJK SC 小二加粗",
            0 if all_ok else -1))

    def eval_page8_jiaocheng_pf(self):
        # -1 细则：第八页"《工程实践与算法应用综合设计》"段落 段后12磅/1.08倍行距/居中
        target_idx = -1
        for i, info in enumerate(self.para_infos):
            if "《工程实践与算法应用综合设计》" in info["text"]:
                target_idx = i
                break

        if target_idx < 0:
            self.hits.append((
                "第八页《工程实践与算法应用综合设计》段落 段后12磅/1.08倍行距/居中", 0))
            return

        info = self.para_infos[target_idx]
        ok = self.pf_check(info, alignment="CENTER", space_after=12, line=1.08)

        self.hits.append((
            "第八页《工程实践与算法应用综合设计》段落 段后12磅/1.08倍行距/居中",
            0 if ok else -1))

    def eval_page8_benke_font(self):
        # -1 细则：第八页"本科设计报告"字体 Noto Sans CJK SC 20 加粗
        # 文档里有两处"本科设计报告"：封底(黑体18pt)和"说明"章节里的
        # 报告题目上下文(宋体12pt)。以紧跟在"云岭数字技术学院"段落之后
        # 出现的那一处作为封底区块的定位锚点，而不是简单取最后一个匹配
        # （最后一个匹配实际落在"说明"上下文里，并非第八页封底）。
        college_idx = -1
        for i, info in enumerate(self.para_infos):
            if info["text"] == "云岭数字技术学院":
                college_idx = i
                break

        target_idx = -1
        if college_idx >= 0:
            for i in range(college_idx + 1, len(self.para_infos)):
                if self.para_infos[i]["text"] == "本科设计报告":
                    target_idx = i
                    break

        if target_idx < 0:
            self.hits.append((
                "第八页本科设计报告字体 Noto Sans CJK SC 20加粗", 0))
            return

        all_ok = True
        found_any = False
        for r in self.paragraphs[target_idx].runs:
            if not r.text.strip():
                continue
            found_any = True
            fn, sz, b = get_run_font(r)
            if not (fn and "Noto Sans CJK SC" in fn):
                all_ok = False; break
            if not (sz is not None and pt_approx(sz, 20, tol=0.3)):
                all_ok = False; break
            if not b:
                all_ok = False; break

        if not found_any:
            self.hits.append((
                "第八页本科设计报告字体 Noto Sans CJK SC 20加粗", 0))
            return

        self.hits.append((
            "第八页本科设计报告字体 Noto Sans CJK SC 20加粗",
            0 if all_ok else -1))

    def eval_page8_benke_pf(self):
        # -1 细则：第八页"本科设计报告"段落 段后42磅/1.08倍行距/居中
        # 同样以紧跟"云岭数字技术学院"段落之后的那一处作为定位锚点。
        college_idx = -1
        for i, info in enumerate(self.para_infos):
            if info["text"] == "云岭数字技术学院":
                college_idx = i
                break

        target_idx = -1
        if college_idx >= 0:
            for i in range(college_idx + 1, len(self.para_infos)):
                if self.para_infos[i]["text"] == "本科设计报告":
                    target_idx = i
                    break

        if target_idx < 0:
            self.hits.append((
                "第八页本科设计报告段落 段后42磅/1.08倍行距/居中", 0))
            return

        info = self.para_infos[target_idx]
        ok = self.pf_check(info, alignment="CENTER", space_after=42, line=1.08)

        self.hits.append((
            "第八页本科设计报告段落 段后42磅/1.08倍行距/居中",
            0 if ok else -1))

    def eval_page8_xinxi_font(self):
        # -1 细则：第八页 5 条信息段落
        #   字体 Noto Sans CJK SC 13 加粗、每条内容需单独成段
        expected = [
            "题    目：城市应急物资调度与路径规划系统",
            "组员姓名：程奕然  梁若澄  沈云舟",
            "组员学号：26301841  26301858  26301873",
            "班    级：信工 26-3",
            "指导教师：岳清源  沈知遥",
        ]

        all_ok = True
        found_any_target = False

        for exp in expected:
            # 点4：每条独占一段 —— 文本 == exp 的段落必须恰有 1 段
            matched_idxs = [i for i, info in enumerate(self.para_infos)
                            if info["text"] == exp]
            if len(matched_idxs) == 0:
                # 未定位到该条 → 跳过，不算 target 命中
                continue
            if len(matched_idxs) != 1:
                all_ok = False

            for idx in matched_idxs:
                found_any_target = True
                for r in self.paragraphs[idx].runs:
                    if not r.text.strip():
                        continue
                    fn, sz, b = get_run_font(r)
                    if not (fn and "Noto Sans CJK SC" in fn):
                        all_ok = False; break
                    if not (sz is not None and pt_approx(sz, 13, tol=0.3)):
                        all_ok = False; break
                    if not b:
                        all_ok = False; break

        if not found_any_target:
            self.hits.append((
                "第八页信息字体 Noto Sans CJK SC 13加粗/每条独占一段", 0))
            return

        self.hits.append((
            "第八页信息字体 Noto Sans CJK SC 13加粗/每条独占一段",
            0 if all_ok else -1))

    def eval_page8_xinxi_pf(self):
        # -1 细则：第八页 5 条信息段落 段前3.3cm/左对齐/段后5磅/1.15倍行距
        # "段前3.3cm"实际是左缩进（w:ind/left_indent），不是段前间距
        # （space_before）——文档里这5段的 space_before 均为 None，真正
        # 设置的 3.3cm 是 left_indent_cm。
        expected = [
            "题    目：城市应急物资调度与路径规划系统",
            "组员姓名：程奕然  梁若澄  沈云舟",
            "组员学号：26301841  26301858  26301873",
            "班    级：信工 26-3",
            "指导教师：岳清源  沈知遥",
        ]

        all_ok = True
        found_any_target = False

        for exp in expected:
            for info in self.para_infos:
                if info["text"] != exp:
                    continue
                found_any_target = True
                # 点1/2/3/4：段前3.3cm(左缩进) / 左对齐 / 段后5磅 / 1.15倍行距
                if not self.pf_check(
                    info,
                    alignment="LEFT",
                    space_after=5,
                    line=1.15,
                    left_indent_cm=3.3,
                ):
                    all_ok = False
                    break
            if not all_ok:
                break

        if not found_any_target:
            self.hits.append((
                "第八页信息段落 段前3.3cm/左对齐/段后5磅/1.15倍行距", 0))
            return

        self.hits.append((
            "第八页信息段落 段前3.3cm/左对齐/段后5磅/1.15倍行距",
            0 if all_ok else -1))


# ============================================================
# 主流程
# ============================================================

# 维度二各评分项的满分（max_delta）——按 §2.2 约定：
#   加分项：max_delta = +N（如 +3、+1），实际 delta ∈ {0, +N}
#   扣分项：max_delta = -N，命中扣分时 delta = max_delta，未命中时 delta = 0。
# hit = (delta == max_delta)；未命中扣分项的 delta 为 0。
RULE_MAX_DELTA = {
    # ---- 加分项 ----
    "A4纵向 页宽约21cm/页高约29.7cm 左3/右2.5/上3/下2.5cm": 3,
    "行距固定22磅+每页最多32行+每行最多34字": 3,
    "本科设计报告页起页眉均为指定文本": 3,
    "页眉文本 宋体/五号/居中": 1,
    "页眉下横向单线(≈0.5pt)": 3,
    "页码连续居中": 3,
    "结构 封面→目录→正文→参考文献 齐全且有序": 3,
    "章标题 第…章/黑体/小二号": 3,
    "结论/参考资料/附录 章节化": 1,
    "一级节标题 数字.数字/黑体/小三号": 3,
    "目录标题 黑体小二号居中": 1,
    "目录条目字体 宋体/小四号": 1,
    # ---- 扣分项 ----
    "封面标题字体 Noto Sans CJK SC 小一 加粗": -1,
    "封面标题段落 段后26磅/1.08倍行距/居中": -1,
    "封面作者字体 Noto Sans CJK SC 四号": -1,
    "封面作者段落 段后14磅/1.08倍行距/居中": -1,
    "封面学院字体 Noto Sans CJK SC 三号 加粗": -1,
    "封面学院段落 段后28磅/1.08倍行距/居中": -1,
    "封面日期字体 Noto Sans CJK SC 13": -1,
    "封面日期段落 段后16磅/1.08倍行距/居中": -1,
    "目录两字 中间空一字符/Noto Sans CJK SC 二号 加粗": -1,
    "目录两字 段落 段后20磅/1.08倍行距/居中": -1,
    "目录内容字体 Noto Sans CJK SC 11": -1,
    "目录内容段落 段后2磅/1.15倍行距/左对齐": -1,
    "第三页标题字体 Noto Sans CJK SC 三号 加粗": -1,
    "第三页标题段落 段后8磅/1.08倍行距/居中": -1,
    "第三/四页表格首行字体 Noto Sans CJK SC 9.5 加粗": -1,
    "第三页表格其他行字体 Noto Sans CJK SC 8（一二四列加粗）": -1,
    "第四页过程性评价标题字体 Noto Sans CJK SC 三号加粗": -1,
    "第四页过程性评价标题段落 段后8磅/1.08倍行距/居中": -1,
    "第四页三综合字体 Noto Sans CJK SC 13加粗": -1,
    "第四页三综合段落 段后4磅/1.08倍行距/左对齐": -1,
    "第四页三综合下内容字体 Noto Sans CJK SC 10": -1,
    "第五页成果编制要求字体 Noto Sans CJK SC 三号加粗": -1,
    "第五页成果编制要求段落 段后8磅/1.08倍行距/居中": -1,
    "第五/六页一、二 标题及内容段落 1.15倍行距/段后10磅/左对齐": -3,
    "第七页成果提交与附录字体 Noto Sans CJK SC 三号加粗": -1,
    "第七页成果提交与附录段落 段后8磅/1.08倍行距/居中": -1,
    "第七页标题字体 Noto Sans CJK SC小四加粗": -1,
    "第七页一二三小标题段落 段后2磅/1.08倍行距/左对齐": -1,
    "第七页一二内容字体 Noto Sans CJK SC 五号": -1,
    "第七页一二内容段落 段后1.5磅/段前缩进0.25cm/悬挂缩进0.25cm/1.05倍行距": -1,
    "第七页三、提交清单 表格字体 首行小五加粗/其余10": -1,
    "第八页云岭数字技术学院字体 Noto Sans CJK SC 二号加粗": -1,
    "第八页云岭数字技术学院段落 段后26磅/1.08倍行距/居中": -1,
    "第八页《工程实践与算法应用综合设计》字体 Noto Sans CJK SC 小二加粗": -1,
    "第八页《工程实践与算法应用综合设计》段落 段后12磅/1.08倍行距/居中": -1,
    "第八页本科设计报告字体 Noto Sans CJK SC 20加粗": -1,
    "第八页本科设计报告段落 段后42磅/1.08倍行距/居中": -1,
    "第八页信息字体 Noto Sans CJK SC 13加粗/每条独占一段": -1,
    "第八页信息段落 段前3.3cm/左对齐/段后5磅/1.15倍行距": -1,
}


def _find_docx_in_dir(dir_path):
    """在给定目录中定位待评估的 .docx 文件；找不到返回 None。

    选取策略：忽略 Word 临时锁文件（~$ 前缀），返回按文件名排序后
    第一个 .docx 文件。仅识别 .docx（OOXML），不再支持二进制 .doc。
    """
    if not os.path.isdir(dir_path):
        return None
    docx_candidates = []
    for name in os.listdir(dir_path):
        if name.startswith("~$"):
            continue  # Word 临时锁文件
        if name.lower().endswith(".docx"):
            docx_candidates.append(name)
    docx_candidates.sort()
    if not docx_candidates:
        return None
    return os.path.join(dir_path, docx_candidates[0])


def evaluate(dir_path: str) -> dict:
    """统一评估入口。

    参数 ``dir_path`` 是**脚本所在目录的路径**；脚本自行在该目录中定位并
    打开被评估的 .docx 文档，据此产出结构化评估结果。

    返回结构（见"脚本接口差异与统一建议 §2.2"）：
        {
            "id": "005",
            "file_name": "xxx.docx",
            "status": "ok" | "error",
            "error": None | str,
            "dim1_pass": bool,
            "dim1_reason": str,
            "dim2_items": [ {rule, max_delta, delta, hit, detail}, ... ],
            "total_score": int,
            "max_score": int,
        }
    """
    result = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": sum(v for v in RULE_MAX_DELTA.values() if v > 0),
    }

    try:
        # 1) 定位目录中的 .docx 文件
        found_path = _find_docx_in_dir(dir_path)
        if found_path is None:
            result["status"] = "error"
            result["error"] = f"目录中未找到可评估的 .docx 文件: {dir_path}"
            return result
        result["file_name"] = os.path.basename(found_path)

        # 仅支持 .docx（OOXML）；python-docx 的 Document() 只能打开 .docx。
        docx_path = found_path

        # 2) 打开文档
        try:
            doc = Document(docx_path)
        except Exception as e:
            result["status"] = "error"
            result["error"] = f"无法打开文档: {e}"
            return result

        # 3) 维度一
        ok, issues = check_dimension_one(doc, docx_path)
        result["dim1_pass"] = bool(ok)
        if not ok:
            result["dim1_reason"] = "; ".join(issues)
            # 维度一未通过：dim2_items 为空，total_score 为 0
            return result

        # 4) 维度二
        ev = Evaluator(doc)
        ev.check_all()

        total = 0
        for label, delta in ev.hits:
            max_delta = RULE_MAX_DELTA.get(label, 0)
            hit = (delta == max_delta)
            result["dim2_items"].append({
                "rule": label,
                "max_delta": max_delta,
                "delta": delta,
                "hit": hit,
                "detail": "",
            })
            total += delta
        result["total_score"] = total
        return result

    except Exception as e:
        result["status"] = "error"
        result["error"] = f"脚本执行异常: {e}"
        return result


if __name__ == "__main__":
    # 本地调试用：`python officeval_005_verifier.py <脚本所在目录>`
    # 未指定目录时，默认使用脚本自身所在目录。
    # 说明：此处仅作为脚本作者本地自测入口，不参与批量运行；主结果通过
    # evaluate() 返回值传出，本块仅打印 JSON 便于人工查看。为避免
    # Windows 默认控制台编码（cp1252）在打印中文时抛 UnicodeEncodeError，
    # 直接以 UTF-8 写入 stdout 的 buffer，而不修改全局 sys.stdout。
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(
        os.path.abspath(__file__))
    _payload = json.dumps(evaluate(_dir), ensure_ascii=False, indent=2)
    try:
        sys.stdout.buffer.write((_payload + "\n").encode("utf-8"))
    except AttributeError:
        # 某些运行环境的 stdout 无 buffer 属性，退回普通 print
        print(_payload)
