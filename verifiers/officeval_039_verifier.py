#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# pyright: reportAny=false, reportExplicitAny=false, reportUnknownMemberType=false, reportUnknownVariableType=false, reportUnknownArgumentType=false, reportUnknownParameterType=false, reportMissingTypeStubs=false, reportPrivateUsage=false, reportImplicitStringConcatenation=false, reportDeprecated=false
"""
自动评估《办公设备维护与巡检服务合同_两页排版版.docx》的排版质量。

对外接口（供批量运行器调用）：
    evaluate(dir_path: str) -> dict
        参数 `dir_path` 为“脚本所在目录的路径”，脚本自己负责在该目录里定位
        并打开被评估的 .docx 文档。返回结构见文件末尾 evaluate() 的说明，
        字段格式与《脚本接口差异与统一建议.md》§2.2 一致。

本地自测（不影响批量流程）：
    python officeval_039_verifier.py <脚本所在目录>

说明：
- 仅支持 .docx（Office Open XML）；不支持二进制 .doc。
- 需要 Windows + Microsoft Word（用于真实分页、PDF 导出、页面位置判断）。
- 需要 python-docx、lxml、pywin32：pip install python-docx lxml pywin32
- 评分逻辑严格按题目描述：维度 1 不通过直接 0 分，不继续累计维度 2；
  维度 2 中正负分点全部自动检查并累计。
"""

from __future__ import annotations

import json
import re
import sys
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except Exception as exc:  # pragma: no cover
    print("缺少依赖 python-docx，请先安装：pip install python-docx", file=sys.stderr)
    raise

try:
    from lxml import etree  # noqa: F401  # 保留以确保 python-docx 使用的 lxml 依赖已装
except Exception as exc:  # pragma: no cover
    print("缺少依赖 lxml，请先安装：pip install lxml", file=sys.stderr)
    raise

try:
    import win32com.client  # type: ignore
except Exception:  # pragma: no cover
    win32com = None  # type: ignore


PT_PER_CM = 72 / 2.54
CM_TOL = 0.08
PAGE_TOL_CM = 0.06

# Word COM 常量（避免依赖生成的 constants）
WD_STATISTIC_PAGES = 2
WD_ACTIVE_END_PAGE_NUMBER = 3
WD_NUMBER_OF_PAGES_IN_DOCUMENT = 4
WD_HORIZONTAL_POSITION_RELATIVE_TO_PAGE = 5
WD_VERTICAL_POSITION_RELATIVE_TO_PAGE = 6
WD_COLLAPSE_END = 0
WD_COLLAPSE_START = 1
WD_EXPORT_FORMAT_PDF = 17
WD_SAVE_CHANGES_FALSE = 0
WD_CHARACTER = 1


@dataclass
class Hit:
    score: int
    title: str
    detail: str = ""


@dataclass
class Miss:
    score: int
    title: str
    detail: str = ""


@dataclass
class EvalReport:
    file: str
    dimension1_passed: bool = False
    dimension1_failures: List[str] = field(default_factory=list)
    dimension1_details: List[str] = field(default_factory=list)
    hits: List[Hit] = field(default_factory=list)
    misses: List[Miss] = field(default_factory=list)
    score: int = 0
    raw: Dict[str, Any] = field(default_factory=dict)

    def add_hit(self, score: int, title: str, detail: str = "") -> None:
        self.hits.append(Hit(score, title, detail))
        self.score += score

    def add_miss(self, score: int, title: str, detail: str = "") -> None:
        self.misses.append(Miss(score, title, detail))


def cm_from_pt(pt: float) -> float:
    return pt / PT_PER_CM


def pt_from_cm(cm: float) -> float:
    return cm * PT_PER_CM


def nearly(actual: Optional[float], expected: float, tol: float = CM_TOL) -> bool:
    return actual is not None and abs(actual - expected) <= tol


def visible_text(s: str) -> str:
    """去掉 Word 表格单元格结束符等不可见控制字符。"""
    return re.sub(r"[\r\x07\x0b\x0c\t ]+", "", s or "")


def pdf_page_count(pdf_path: Path) -> Optional[int]:
    """优先用 PyPDF；没有时用 PDF 对象标记做轻量统计。"""
    try:
        from pypdf import PdfReader  # type: ignore

        return len(PdfReader(str(pdf_path)).pages)
    except Exception:
        pass
    try:
        data = pdf_path.read_bytes()
        # /Type /Page 后面用负向前瞻排除 /Pages。
        return len(re.findall(rb"/Type\s*/Page(?!s)\b", data))
    except Exception:
        return None


class WordSession:
    """用 Microsoft Word 获取真实分页、位置和 PDF 页数。"""

    def __init__(self, docx_path: Path) -> None:
        if win32com is None:  # type: ignore[name-defined]
            raise RuntimeError("未安装 pywin32，无法调用 Microsoft Word。")
        self.path = docx_path
        self.word = win32com.client.DispatchEx("Word.Application")  # type: ignore[name-defined]
        self.word.Visible = False
        self.word.DisplayAlerts = 0
        self.doc = self.word.Documents.Open(
            str(docx_path), ReadOnly=True, AddToRecentFiles=False, ConfirmConversions=False
        )
        self.doc.Repaginate()

    def close(self) -> None:
        try:
            self.doc.Close(SaveChanges=WD_SAVE_CHANGES_FALSE)
        except Exception:
            pass
        try:
            self.word.Quit()
        except Exception:
            pass

    def compute_pages(self) -> int:
        self.doc.Repaginate()
        return int(self.doc.ComputeStatistics(WD_STATISTIC_PAGES))

    def info_pages(self) -> int:
        self.doc.Repaginate()
        return int(self.doc.Range().Information(WD_NUMBER_OF_PAGES_IN_DOCUMENT))

    def export_pdf_pages(self) -> Optional[int]:
        with tempfile.TemporaryDirectory() as td:
            pdf_path = Path(td) / "contract_eval.pdf"
            try:
                self.doc.ExportAsFixedFormat(str(pdf_path), WD_EXPORT_FORMAT_PDF)
            except Exception:
                return None
            return pdf_page_count(pdf_path)

    def page_at_pos(self, pos: int) -> Optional[int]:
        try:
            end = min(pos + 1, int(self.doc.Content.End))
            rng = self.doc.Range(pos, end)
            return int(rng.Information(WD_ACTIVE_END_PAGE_NUMBER))
        except Exception:
            return None

    def page_at_range_start(self, rng: Any) -> Optional[int]:
        return self.page_at_pos(int(rng.Start))

    def y_at_pos(self, pos: int) -> Optional[float]:
        try:
            end = min(pos + 1, int(self.doc.Content.End))
            rng = self.doc.Range(pos, end)
            y = float(rng.Information(WD_VERTICAL_POSITION_RELATIVE_TO_PAGE))
            return y if y >= 0 else None
        except Exception:
            return None

    def range_pages(self, rng: Any) -> Tuple[Optional[int], Optional[int]]:
        start = int(rng.Start)
        end = max(start, int(rng.End) - 1)
        return self.page_at_pos(start), self.page_at_pos(end)

    def find_phrase(self, phrase: str) -> Optional[Dict[str, Any]]:
        rng = self.doc.Content.Duplicate
        finder = rng.Find
        finder.Text = phrase
        finder.MatchCase = False
        finder.MatchWholeWord = False
        if not finder.Execute():
            return None
        page = self.page_at_pos(int(rng.Start))
        y = self.y_at_pos(int(rng.Start))
        try:
            para_rng = rng.Paragraphs(1).Range
            bottom_pos = max(int(para_rng.Start), int(para_rng.End) - 1)
            bottom_y = self.y_at_pos(bottom_pos)
            font_size = float(para_rng.Font.Size)
            if font_size <= 0 or font_size > 100:
                font_size = 12.0
            if bottom_y is not None:
                bottom_y += max(font_size, 10.0)
        except Exception:
            bottom_y = None
        return {"start": int(rng.Start), "end": int(rng.End), "page": page, "y": y, "bottom_y": bottom_y}

    def phrase_page(self, phrase: str) -> Optional[int]:
        item = self.find_phrase(phrase)
        return None if item is None else item.get("page")

    def table_info(self, one_based_index: int) -> Dict[str, Any]:
        tbl = self.doc.Tables(one_based_index)
        start_page, end_page = self.range_pages(tbl.Range)
        start_y = self.y_at_pos(int(tbl.Range.Start))
        end_y = self.y_at_pos(max(int(tbl.Range.Start), int(tbl.Range.End) - 1))
        try:
            end_y = None if end_y is None else end_y + 12.0
        except Exception:
            pass
        widths_pt = []
        try:
            for i in range(1, int(tbl.Columns.Count) + 1):
                widths_pt.append(float(tbl.Columns(i).Width))
        except Exception:
            widths_pt = []
        left_pt = None
        try:
            left_pt = float(tbl.Range.Information(WD_HORIZONTAL_POSITION_RELATIVE_TO_PAGE))
        except Exception:
            pass
        return {
            "rows": int(tbl.Rows.Count),
            "cols": int(tbl.Columns.Count),
            "start_page": start_page,
            "end_page": end_page,
            "start_y_pt": start_y,
            "end_y_pt": end_y,
            "widths_cm_com": [cm_from_pt(x) for x in widths_pt],
            "left_pt": left_pt,
            "right_pt": None if left_pt is None or not widths_pt else left_pt + sum(widths_pt),
        }

    def paragraph_intervals_by_page(self, pages: int) -> Dict[int, List[Tuple[float, float]]]:
        intervals: Dict[int, List[Tuple[float, float]]] = {p: [] for p in range(1, pages + 1)}
        try:
            for para in self.doc.Paragraphs:
                text = visible_text(str(para.Range.Text))
                if not text:
                    continue
                pos = int(para.Range.Start)
                page = self.page_at_pos(pos)
                y = self.y_at_pos(pos)
                if page is None or y is None or page not in intervals:
                    continue
                try:
                    font_size = float(para.Range.Font.Size)
                    if font_size <= 0 or font_size > 100:
                        font_size = 12.0
                except Exception:
                    font_size = 12.0
                intervals[page].append((y, y + max(font_size, 10.0)))
        except Exception:
            pass
        try:
            for i in range(1, int(self.doc.Tables.Count) + 1):
                info = self.table_info(i)
                sp, ep = info["start_page"], info["end_page"]
                if sp == ep and sp in intervals:
                    sy, ey = info.get("start_y_pt"), info.get("end_y_pt")
                    if sy is not None and ey is not None:
                        intervals[sp].append((float(sy), max(float(ey), float(sy) + 10.0)))
        except Exception:
            pass
        return intervals


def all_doc_text(doc: Document) -> str:
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for story in (section.header, section.first_page_header, section.even_page_header, section.footer):
            for p in story.paragraphs:
                parts.append(p.text)
    return "\n".join(parts)


def table_grid_widths_cm(table: Any) -> List[Optional[float]]:
    """读取表格各列的真实宽度（cm）。

    优先取首行各单元格的 w:tcW（实际渲染列宽，会随用户调整列宽而更新），
    其次回退到 w:gridCol（仅建议网格，可能与实际不一致），
    最后回退到 python-docx 的 column.width。
    """
    # 1) 首行 tcW —— 真实列宽
    try:
        if table.rows:
            tc_widths: List[Optional[float]] = []
            for tc in table.rows[0]._tr.tc_lst:
                tcw = tc.tcPr.tcW if tc.tcPr is not None else None
                if tcw is None:
                    tc_widths.append(None)
                    continue
                w_type = tcw.get(qn("w:type"))
                raw = tcw.get(qn("w:w"))
                if raw is None or w_type not in (None, "dxa"):
                    tc_widths.append(None)
                    continue
                tc_widths.append(int(raw) / 567.0)
            if tc_widths and any(x is not None for x in tc_widths):
                return tc_widths
    except Exception:
        pass
    # 2) tblGrid 网格 —— 建议值，作为回退
    widths: List[Optional[float]] = []
    try:
        grid = table._tbl.tblGrid
        for grid_col in grid.gridCol_lst:
            raw = grid_col.get(qn("w:w"))
            widths.append(None if raw is None else int(raw) / 567.0)
        if widths:
            return widths
    except Exception:
        pass
    # 3) python-docx column.width —— 最终回退
    try:
        for col in table.columns:
            widths.append(None if col.width is None else col.width.cm)
    except Exception:
        pass
    return widths


def table_headers(table: Any) -> List[str]:
    if not table.rows:
        return []
    return [visible_text(cell.text) for cell in table.rows[0].cells]


def find_table_by_headers(doc: Document, expected_headers: Sequence[str]) -> Optional[int]:
    expected = [visible_text(x) for x in expected_headers]
    for idx, table in enumerate(doc.tables):
        headers = table_headers(table)
        if headers[: len(expected)] == expected:
            return idx
    return None


def table_contains(table: Any, *phrases: str) -> bool:
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    return all(p in text for p in phrases)


def table_shape_and_widths_ok(
    table: Any, expected_rows: int, expected_cols: int, expected_widths_cm: Sequence[float], tol: float = CM_TOL
) -> Tuple[bool, str]:
    rows = len(table.rows)
    cols = len(table.columns)
    widths = table_grid_widths_cm(table)
    reasons = []
    if rows != expected_rows or cols != expected_cols:
        reasons.append(f"实际为 {rows} 行 {cols} 列")
    if len(widths) < len(expected_widths_cm):
        reasons.append(f"未能读取到 {len(expected_widths_cm)} 个列宽，实际读取 {len(widths)} 个")
    else:
        bad = []
        for i, (actual, expected) in enumerate(zip(widths, expected_widths_cm), 1):
            if not nearly(actual, expected, tol):
                bad.append(f"第{i}列 {actual:.2f}cm≠{expected:.2f}cm" if actual is not None else f"第{i}列宽缺失")
        if bad:
            reasons.append("；".join(bad))
    return not reasons, "；".join(reasons) if reasons else "行列数和列宽均符合"


def page_setup_issue(doc: Any) -> tuple[bool, list[str]]:
    """细则：整篇文档页面设置满足任意一项即扣 1 分 ——
    纸张不是A4纵向；上下边距不是0.35cm；左右边距不是0.85cm；
    页眉距顶端不是0.20cm；页脚距底端不是0.20cm。允许微量浮点容差。
    """
    issues: list[str] = []
    tol = 0.005  # 仅吸收 EMU↔cm 换算的浮点尾差，相当于显示到 0.01cm 严格相等
    for i, sec in enumerate(doc.sections, 1):
        prefix = f"第{i}节"
        page_width_cm: float = float(sec.page_width.cm)
        page_height_cm: float = float(sec.page_height.cm)
        if sec.orientation != WD_ORIENT.PORTRAIT or page_width_cm > page_height_cm:
            issues.append(f"{prefix}纸张不是 A4 纵向")
        if not (nearly(page_width_cm, 21.0, tol) and nearly(page_height_cm, 29.7, tol)):
            issues.append(f"{prefix}纸张不是 A4（{page_width_cm:.2f}×{page_height_cm:.2f}cm）")
        checks: list[tuple[str, float, float]] = [
            ("上边距", float(sec.top_margin.cm), 0.35),
            ("下边距", float(sec.bottom_margin.cm), 0.35),
            ("左边距", float(sec.left_margin.cm), 0.85),
            ("右边距", float(sec.right_margin.cm), 0.85),
            ("页眉距顶端", float(sec.header_distance.cm), 0.20),
            ("页脚距底端", float(sec.footer_distance.cm), 0.20),
        ]
        for label, actual, expected in checks:
            if not nearly(actual, expected, tol):
                issues.append(f"{prefix}{label} {actual:.2f}cm≠{expected:.2f}cm")
    return bool(issues), issues


def run_font_name(run: Any) -> Optional[str]:
    try:
        rpr = run._element.rPr
        if rpr is not None and rpr.rFonts is not None:
            east_asia = rpr.rFonts.get(qn("w:eastAsia"))
            ascii_font = rpr.rFonts.get(qn("w:ascii"))
            h_ansi = rpr.rFonts.get(qn("w:hAnsi"))
            return east_asia or run.font.name or ascii_font or h_ansi
    except Exception:
        pass
    try:
        return run.font.name
    except Exception:
        return None


def header_issue(doc: Any, expected: str = "办公设备维护与巡检服务合同") -> tuple[bool, list[str]]:
    """细则：两页缺少任意一页页眉“办公设备维护与巡检服务合同”，
    或页眉字体不是微软雅黑、字号不是 8 磅、颜色不是灰色或黑色、未水平居中 → 扣 1 分。
    """
    issues: List[str] = []
    for si, sec in enumerate(doc.sections, 1):
        # 若没有首页/奇偶页特殊设置，regular header 同时代表两页。
        headers: list[tuple[str, Any]] = [("常规页眉", sec.header)]
        try:
            if sec.different_first_page_header_footer:
                headers.append(("首页页眉", sec.first_page_header))
        except Exception:
            pass
        try:
            doc_settings = doc._part.settings.element
            even_odd = bool(doc_settings.xpath("./w:evenAndOddHeaders"))
            if even_odd:
                headers.append(("偶数页页眉", sec.even_page_header))
        except Exception:
            pass

        for label, header in headers:
            paras = [p for p in header.paragraphs if visible_text(p.text)]
            matched = [p for p in paras if expected in p.text]
            if not matched:
                issues.append(f"第{si}节{label}缺少“{expected}”")
                continue
            para = matched[0]
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                issues.append(f"第{si}节{label}未水平居中")
            runs = [r for r in para.runs if visible_text(r.text)]
            if not runs:
                issues.append(f"第{si}节{label}没有可检查的文字 run")
                continue
            for ri, run in enumerate(runs, 1):
                font_name = run_font_name(run)
                if font_name and font_name not in {"微软雅黑", "Microsoft YaHei"}:
                    issues.append(f"第{si}节{label}第{ri}段文字字体为 {font_name}，不是微软雅黑")
                if not font_name:
                    issues.append(f"第{si}节{label}第{ri}段文字未显式设置微软雅黑")
                size = run.font.size.pt if run.font.size is not None else None
                if size is None or float(size) != 8.0:
                    issues.append(f"第{si}节{label}第{ri}段文字字号 {size}，不是 8 磅")
                rgb = run.font.color.rgb
                if rgb is not None:
                    r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
                    if not (r == g == b):
                        issues.append(f"第{si}节{label}第{ri}段文字颜色 #{rgb} 不是灰色或黑色")
    return bool(issues), issues


def right_border_exceeds_page_margins(
    doc: Any, table_infos: list[dict[str, Any]]
) -> tuple[bool, list[str]]:
    """细则：任意一个表格的右侧边框超出页面边距 → 扣 1 分。"""
    issues: List[str] = []
    sections = doc.sections
    if not sections:
        return False, issues
    sec = sections[0]
    page_width = sec.page_width
    right_margin = sec.right_margin
    page_width_pt: float = float(page_width.pt) if page_width is not None else pt_from_cm(21.0)
    right_margin_pt: float = float(right_margin.pt) if right_margin is not None else pt_from_cm(0.85)
    right_limit_pt: float = page_width_pt - right_margin_pt
    for i, info in enumerate(table_infos, 1):
        right_pt = info.get("right_pt")
        if right_pt is None:
            continue
        if float(right_pt) > right_limit_pt:
            issues.append(f"表格{i}右边框超出页面右边距")
    return bool(issues), issues


def bottom_gap_issue(word: WordSession, doc: Document, signature_table_idx: Optional[int]) -> Tuple[bool, List[str]]:
    """细则：第1页账户信息下方至页脚之间连续空白高度超过 1.5cm，
    或第2页签署表格下方至页脚之间连续空白高度超过 1.5cm → 扣 1 分。
    """
    issues: List[str] = []
    if not doc.sections:
        return False, issues
    page_height_pt = doc.sections[0].page_height.pt
    footer_top_pt = page_height_pt - doc.sections[0].footer_distance.pt

    # 第1页：账户信息下方至页脚
    acct = word.find_phrase("乙方收款账户")
    if acct and acct.get("page") == 1 and acct.get("bottom_y") is not None:
        gap_cm = cm_from_pt(footer_top_pt - float(acct["bottom_y"]))
        if gap_cm > 1.5:
            issues.append(f"第1页账户信息下方至页脚之间连续空白约 {gap_cm:.2f}cm > 1.5cm")

    # 第2页：签署表格下方至页脚
    if signature_table_idx is not None:
        info = word.table_info(signature_table_idx + 1)
        if info.get("end_page") == 2 and info.get("end_y_pt") is not None:
            gap_cm = cm_from_pt(footer_top_pt - float(info["end_y_pt"]))
            if gap_cm > 1.5:
                issues.append(f"第2页签署表格下方至页脚之间连续空白约 {gap_cm:.2f}cm > 1.5cm")
    return bool(issues), issues


def account_labels_on_page2(word: WordSession) -> Tuple[bool, List[str]]:
    # 细则要求：账户名称、开户行或账号任意一项出现在第2页 → 扣 1 分。
    # 仅当作为“标签”形式出现时才算命中：后面紧跟中/英文冒号（可有空格），
    # 避免“付款至乙方账号”这类正文中自然出现的措辞被误判。
    hits: List[str] = []
    label_re = re.compile(r"(账户名称|开户行|账[\s]*号)\s*[::]")
    try:
        full_text = str(word.doc.Content.Text)
    except Exception:
        full_text = ""
    for m in label_re.finditer(full_text):
        # Word 的 Content.Text 字符索引与 Range 的 Start 基本一致（同为 Content 范围内的字符偏移）。
        pos = m.start()
        page = word.page_at_pos(pos)
        if page == 2:
            ctx_start = max(0, pos - 20)
            ctx_end = min(len(full_text), m.end() + 20)
            ctx = visible_text(full_text[ctx_start:ctx_end])
            hits.append(f"“{m.group(1)}：”出现在第2页 [pos={pos}, 上下文: …{ctx}…]")
    return bool(hits), hits


def _evaluate_docx(docx_path: Path) -> EvalReport:
    report = EvalReport(file=str(docx_path))
    docx_path = docx_path.resolve()

    # -------------------- 维度 1：可用与可修改性 --------------------
    if docx_path.suffix.lower() != ".docx":
        report.dimension1_failures.append("交付文件不是 .docx 格式。")

    if not docx_path.exists():
        report.dimension1_failures.append("文件不存在。")
        return report

    try:
        with zipfile.ZipFile(docx_path) as zf:
            if "word/document.xml" not in zf.namelist():
                report.dimension1_failures.append("docx 包中缺少 word/document.xml。")
    except Exception as exc:
        report.dimension1_failures.append(f"文件不是可正常解析的 docx 压缩包：{exc}")
        return report

    try:
        doc = Document(str(docx_path))
        report.dimension1_details.append("python-docx 可正常读取文档结构。")
    except Exception as exc:
        report.dimension1_failures.append(f"python-docx 无法打开文档：{exc}")
        return report

    word: Optional[WordSession] = None
    try:
        word = WordSession(docx_path)
        pages_status = word.compute_pages()
        pages_info = word.info_pages()
        report.raw["pages_status_or_repaginate"] = pages_status
        report.raw["pages_word_information"] = pages_info
        if pages_status == 2 and pages_info == 2:
            report.dimension1_details.append("Microsoft Word 可正常打开，分页结果为 2 页。")
        else:
            report.dimension1_failures.append(f"Word 分页不是 2 页：ComputeStatistics={pages_status}，Information={pages_info}。")
    except Exception as exc:
        report.dimension1_failures.append(f"Microsoft Word/COM 无法正常打开文档：{exc}")

    report.dimension1_passed = not report.dimension1_failures
    if not report.dimension1_passed:
        report.score = 0
        if word is not None:
            word.close()
        return report

    assert word is not None

    # -------------------- 维度 2：完成度评分细则 --------------------
    pdf_pages = word.export_pdf_pages()
    report.raw["pages_exported_pdf"] = pdf_pages
    pages_status = report.raw.get("pages_status_or_repaginate")
    pages_info = report.raw.get("pages_word_information")

    service_scope_idx = find_table_by_headers(doc, ["服务模块", "包含内容", "交付记录", "不包含事项"])
    package_idx = find_table_by_headers(doc, ["选择", "套餐名称", "适用设备规模", "年度费用", "主要服务内容"])
    signature_idx = None
    for i, table in enumerate(doc.tables):
        if len(table.rows) == 5 and len(table.columns) == 2 and table_contains(table, "甲方（盖章）", "乙方（盖章）"):
            signature_idx = i
            break

    table_infos = [word.table_info(i + 1) for i in range(len(doc.tables))]
    report.raw["table_pages"] = table_infos

    def phrase_on_page(phrase: str, page: int) -> bool:
        return word.phrase_page(phrase) == page

    # 细则只要求：第1页包含合同标题、合同说明、第一条至第四条、合同总价段落、乙方收款账户段落；
    # 第2页包含第五条至第十条全部合同条款。这里不再对每一条款的具体小标题做约束，
    # 只检查“一、”至“十、”等条款序号是否出现在对应页。
    p1_phrases = [
        "办公设备维护与巡检服务合同",  # 合同标题
        "一、",
        "二、",
        "三、",
        "四、",
        "合同总价",
        "乙方收款账户",
    ]
    p2_phrases = ["五、", "六、", "七、", "八、", "九、", "十、"]
    p1_text_ok = all(phrase_on_page(x, 1) for x in p1_phrases)
    p2_text_ok = all(phrase_on_page(x, 2) for x in p2_phrases)
    info_table_ok = len(doc.tables) >= 1 and table_contains(doc.tables[0], "甲方", "乙方") and table_infos[0]["start_page"] == 1
    scope_table_page1 = service_scope_idx is not None and table_infos[service_scope_idx]["start_page"] == 1
    package_table_page1 = package_idx is not None and table_infos[package_idx]["start_page"] == 1
    signature_page2 = signature_idx is not None and table_infos[signature_idx]["start_page"] == 2
    all_page_counts_2 = pages_status == 2 and pages_info == 2 and pdf_pages == 2
    whole_content_ok = all(
        [all_page_counts_2, p1_text_ok, p2_text_ok, info_table_ok, scope_table_page1, package_table_page1, signature_page2]
    )
    if whole_content_ok:
        report.add_hit(
            5,
            "整篇 Word 文档页数和两页整体内容完整",
            f"Word 状态/分页={pages_status}，打印预览等价页数={pages_info}，导出 PDF={pdf_pages}；第1页和第2页关键内容均在对应页。",
        )
    else:
        missing = []
        if not all_page_counts_2:
            missing.append(f"三类页数未全部为2（Word={pages_status}/{pages_info}, PDF={pdf_pages}）")
        if not p1_text_ok:
            missing.append("第1页关键文字不全或页码不符")
        if not p2_text_ok:
            missing.append("第2页第五条至第十条不全或页码不符")
        if not info_table_ok:
            missing.append("甲乙方信息表未在第1页识别到")
        if not scope_table_page1:
            missing.append("服务范围表未在第1页识别到")
        if not package_table_page1:
            missing.append("服务套餐表未在第1页识别到")
        if not signature_page2:
            missing.append("签署信息表未在第2页识别到")
        report.add_miss(5, "整篇 Word 文档页数和两页整体内容完整", "；".join(missing))

    # +1：第1页甲乙方信息表
    # 细则要求：位于第1页；7行4列；四列宽度依次为 2.47cm、7.42cm、2.47cm、7.42cm；表格不跨页。
    # 列宽允许微量浮点容差（吸收 EMU↔cm 换算误差，相当于显示到 0.01cm 严格相等）。
    if len(doc.tables) >= 1:
        ok_shape, reason = table_shape_and_widths_ok(doc.tables[0], 7, 4, [2.47, 7.42, 2.47, 7.42], tol=0.01)
        no_cross = table_infos[0]["start_page"] == table_infos[0]["end_page"] == 1
        if ok_shape and no_cross:
            report.add_hit(1, "第1页甲乙方信息表符合 7行4列、指定列宽且不跨页", reason)
        else:
            report.add_miss(1, "第1页甲乙方信息表符合 7行4列、指定列宽且不跨页", reason + ("；表格跨页或不在第1页" if not no_cross else ""))
    else:
        report.add_miss(1, "第1页甲乙方信息表符合 7行4列、指定列宽且不跨页", "未检测到表格")

    # +1：第1页服务范围表
    # 细则要求：位于第1页；5行4列；列宽均为 4.88cm；表头依次为
    # “服务模块”“包含内容”“交付记录”“不包含事项”。列宽允许微量浮点容差。
    if service_scope_idx is not None:
        ok_shape, reason = table_shape_and_widths_ok(
            doc.tables[service_scope_idx], 5, 4, [4.88, 4.88, 4.88, 4.88], tol=0.01
        )
        headers = table_headers(doc.tables[service_scope_idx])
        expected_headers = ["服务模块", "包含内容", "交付记录", "不包含事项"]
        headers_ok: bool = headers[: len(expected_headers)] == expected_headers
        scope_start_page: int | None = table_infos[service_scope_idx]["start_page"]
        scope_end_page: int | None = table_infos[service_scope_idx]["end_page"]
        scope_page_ok: bool = scope_start_page == 1 and scope_end_page == 1
        extra: list[str] = []
        if not headers_ok:
            extra.append(f"表头不符：实际 {headers[:4]}")
        if not scope_page_ok:
            extra.append(
                f"未整体位于第1页：起始页={scope_start_page}，结束页={scope_end_page}"
            )
        if ok_shape and headers_ok and scope_page_ok:
            report.add_hit(1, "第1页服务范围表符合 5行4列、列宽和表头要求", reason)
        else:
            report.add_miss(
                1,
                "第1页服务范围表符合 5行4列、列宽和表头要求",
                "；".join([reason] + extra),
            )
    else:
        report.add_miss(1, "第1页服务范围表符合 5行4列、列宽和表头要求", "未按表头识别到服务范围表")

    # +1：第1页服务套餐表
    # 细则要求：6行5列；表头依次为“选择”“套餐名称”“适用设备规模”“年度费用”“主要服务内容”；
    # 表格整体位于第1页。列宽不作要求。
    if package_idx is not None:
        pkg_table = doc.tables[package_idx]
        pkg_rows = len(pkg_table.rows)
        pkg_cols = len(pkg_table.columns)
        shape_ok = pkg_rows == 6 and pkg_cols == 5
        package_headers = table_headers(pkg_table)
        expected_package_headers = ["选择", "套餐名称", "适用设备规模", "年度费用", "主要服务内容"]
        package_headers_ok: bool = (
            package_headers[: len(expected_package_headers)] == expected_package_headers
        )
        package_start_page: int | None = table_infos[package_idx]["start_page"]  # type: ignore[assignment]
        package_end_page: int | None = table_infos[package_idx]["end_page"]  # type: ignore[assignment]
        package_page_ok: bool = package_start_page == 1 and package_end_page == 1
        pkg_reasons: list[str] = []
        if not shape_ok:
            pkg_reasons.append(f"实际为 {pkg_rows} 行 {pkg_cols} 列（应为 6 行 5 列）")
        if not package_headers_ok:
            pkg_reasons.append(f"表头不符：实际 {package_headers[:5]}")
        if not package_page_ok:
            pkg_reasons.append(
                f"表格未整体位于第1页：起始页={package_start_page}，结束页={package_end_page}"
            )
        if shape_ok and package_headers_ok and package_page_ok:
            report.add_hit(
                1,
                "第1页服务套餐表符合 6行5列、五个指定表头且整体位于第1页",
                "6 行 5 列、表头齐全且整表位于第1页",
            )
        else:
            report.add_miss(
                1,
                "第1页服务套餐表符合 6行5列、五个指定表头且整体位于第1页",
                "；".join(pkg_reasons),
            )
    else:
        report.add_miss(
            1,
            "第1页服务套餐表符合 6行5列、五个指定表头且整体位于第1页",
            "未按表头识别到服务套餐表",
        )

    # +1：第2页签署信息表
    # 细则要求：位于第2页底部；2列5行；左右两列宽度相同（不再要求具体等于 8.10cm）。
    if signature_idx is not None:
        sig_table = doc.tables[signature_idx]
        sig_rows = len(sig_table.rows)
        sig_cols = len(sig_table.columns)
        sig_shape_ok = sig_rows == 5 and sig_cols == 2
        sig_widths = table_grid_widths_cm(sig_table)
        # 左右两列宽度相同：读取到两列且差值在容差内即通过；容差沿用其它列宽处的 0.01cm。
        widths_equal: bool = (
            len(sig_widths) >= 2
            and sig_widths[0] is not None
            and sig_widths[1] is not None
            and nearly(float(sig_widths[0]), float(sig_widths[1]), 0.01)
        )
        info = table_infos[signature_idx]
        sig_start_page: int | None = info.get("start_page")  # type: ignore[assignment]
        sig_start_y: float | None = info.get("start_y_pt")  # type: ignore[assignment]
        page_height_emu = doc.sections[0].page_height
        page_height_pt: float = page_height_emu.pt if page_height_emu is not None else pt_from_cm(29.7)
        near_bottom: bool = (
            sig_start_page == 2
            and sig_start_y is not None
            and float(sig_start_y) > page_height_pt * 0.50
        )
        sig_reasons: list[str] = []
        if not sig_shape_ok:
            sig_reasons.append(f"实际为 {sig_rows} 行 {sig_cols} 列（应为 5 行 2 列）")
        if not widths_equal:
            if len(sig_widths) < 2 or sig_widths[0] is None or sig_widths[1] is None:
                sig_reasons.append(f"未能读取到两列列宽：{sig_widths}")
            else:
                sig_reasons.append(
                    f"左右两列列宽不同：{float(sig_widths[0]):.2f}cm vs {float(sig_widths[1]):.2f}cm"
                )
        if not near_bottom:
            sig_reasons.append("未位于第2页底部区域")
        if sig_shape_ok and widths_equal and near_bottom:
            widths_note = (
                f"左右两列列宽相同：{float(sig_widths[0]):.2f}cm"
                if len(sig_widths) >= 2 and sig_widths[0] is not None
                else "左右两列列宽相同"
            )
            report.add_hit(
                1,
                "第2页底部签署信息表符合 2列5行且左右列宽相同",
                f"5 行 2 列；{widths_note}；位于第2页底部",
            )
        else:
            report.add_miss(
                1,
                "第2页底部签署信息表符合 2列5行且左右列宽相同",
                "；".join(sig_reasons),
            )
    else:
        report.add_miss(
            1,
            "第2页底部签署信息表符合 2列5行且左右列宽相同",
            "未识别到签署信息表",
        )

    # 扣分项
    acct = word.find_phrase("乙方收款账户")
    if not acct or acct.get("page") != 1:
        report.add_hit(-1, "第1页没有出现乙方收款账户段落", "未找到，或找到但不在第1页")

    labels_on_p2, label_details = account_labels_on_page2(word)
    if labels_on_p2:
        report.add_hit(-1, "账户名称、开户行或账号任意一项出现在第2页", "；".join(label_details))

    # 细则要求：第1页服务套餐表出现分页 → 扣 1 分。
    if package_idx is not None:
        pinfo = table_infos[package_idx]
        if not (pinfo["start_page"] == pinfo["end_page"] == 1):
            report.add_hit(-1, "第1页服务套餐表出现分页")

    exceeds, exceed_details = right_border_exceeds_page_margins(doc, table_infos)
    if exceeds:
        report.add_hit(-1, "任意一个表格的右侧边框超出页面边距", "；".join(exceed_details))

    setup_bad, setup_details = page_setup_issue(doc)
    if setup_bad:
        report.add_hit(-1, "页面设置存在不符合项", "；".join(setup_details))

    header_bad, header_details = header_issue(doc)
    if header_bad:
        report.add_hit(-1, "两页页眉文字、字体、字号、颜色或居中要求存在不符合项", "；".join(header_details))

    gap_bad, gap_details = bottom_gap_issue(word, doc, signature_idx)
    if gap_bad:
        report.add_hit(
            -1,
            "第1页账户信息下方或第2页签署表格下方至页脚之间连续空白高度超过1.5厘米",
            "；".join(gap_details),
        )

    word.close()
    return report


def print_report(report: EvalReport, as_json: bool = False) -> None:
    """兼容旧调用方：把内部报告按原有文本/JSON 形式打印到 stdout。

    正式批量入口是 evaluate(dir_path)，本函数仅供本地或历史脚本继续使用。
    """
    if as_json:
        payload = {
            "file": report.file,
            "dimension1_passed": report.dimension1_passed,
            "dimension1_failures": report.dimension1_failures,
            "dimension1_details": report.dimension1_details,
            "hits": [hit.__dict__ for hit in report.hits],
            "misses": [miss.__dict__ for miss in report.misses],
            "score": report.score,
            "raw": report.raw,
        }
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return

    if report.dimension1_passed:
        print("维度一：通过")
        print("维度二：评分结果")
        for hit in report.hits:
            sign = "+" if hit.score > 0 else ""
            print(f"{sign}{hit.score}：{hit.title}")
        print(f"最终得分：{report.score}")
    else:
        print("维度一：不通过")
        print("最终得分：0")


SCRIPT_ID = "039"
DEFAULT_DOCX_NAME = "办公设备维护与巡检服务合同_两页排版版.docx"

# 维度二正向评分项（命中则加分，未命中为 0）。顺序即输出顺序。
POSITIVE_RULES: List[Tuple[str, int]] = [
    ("整篇 Word 文档页数和两页整体内容完整", 5),
    ("第1页甲乙方信息表符合 7行4列、指定列宽且不跨页", 1),
    ("第1页服务范围表符合 5行4列、列宽和表头要求", 1),
    ("第1页服务套餐表符合 6行5列、五个指定表头且整体位于第1页", 1),
    ("第2页底部签署信息表符合 2列5行且左右列宽相同", 1),
]

# 维度二扣分项（命中则减分，未命中为 0）。顺序即输出顺序。
DEDUCTION_RULES: List[Tuple[str, int]] = [
    ("第1页没有出现乙方收款账户段落", -1),
    ("账户名称、开户行或账号任意一项出现在第2页", -1),
    ("第1页服务套餐表出现分页", -1),
    ("任意一个表格的右侧边框超出页面边距", -1),
    ("页面设置存在不符合项", -1),
    ("两页页眉文字、字体、字号、颜色或居中要求存在不符合项", -1),
    ("第1页账户信息下方或第2页签署表格下方至页脚之间连续空白高度超过1.5厘米", -1),
]


def _find_target_docx(dir_path: Path) -> Optional[Path]:
    """在脚本所在目录里定位被评估的 .docx 文件。

    优先匹配约定的默认文件名；否则取目录下第一个 .docx（忽略 ~$ 临时锁文件）。
    """
    if not dir_path.is_dir():
        return None
    preferred = dir_path / DEFAULT_DOCX_NAME
    if preferred.exists():
        return preferred
    candidates = sorted(
        p
        for p in dir_path.iterdir()
        if p.is_file() and p.suffix.lower() == ".docx" and not p.name.startswith("~$")
    )
    return candidates[0] if candidates else None


def _report_to_result(report: EvalReport, docx_path: Path) -> Dict[str, Any]:
    """把内部 EvalReport 转成《脚本接口差异与统一建议.md》§2.2 约定的结构。"""
    hit_titles = {hit.title: hit for hit in report.hits}
    miss_titles = {miss.title: miss for miss in report.misses}

    dim2_items: List[Dict[str, Any]] = []
    if report.dimension1_passed:
        for rule, max_delta in POSITIVE_RULES:
            if rule in hit_titles:
                _ = hit_titles[rule]
                dim2_items.append(
                    {
                        "rule": rule,
                        "max_delta": max_delta,
                        "delta": max_delta,
                        "hit": True,
                        "detail": "",
                    }
                )
            else:
                _ = miss_titles.get(rule)
                dim2_items.append(
                    {
                        "rule": rule,
                        "max_delta": max_delta,
                        "delta": 0,
                        "hit": False,
                        "detail": "",
                    }
                )
        for rule, max_delta in DEDUCTION_RULES:
            if rule in hit_titles:
                _ = hit_titles[rule]
                dim2_items.append(
                    {
                        "rule": rule,
                        "max_delta": max_delta,
                        "delta": max_delta,
                        "hit": True,
                        "detail": "",
                    }
                )
            else:
                dim2_items.append(
                    {
                        "rule": rule,
                        "max_delta": max_delta,
                        "delta": 0,
                        "hit": False,
                        "detail": "",
                    }
                )

    max_score = sum(md for _, md in POSITIVE_RULES)
    total_score = sum(item["delta"] for item in dim2_items) if report.dimension1_passed else 0

    return {
        "id": SCRIPT_ID,
        "file_name": docx_path.name,
        "status": "ok",
        "error": None,
        "dim1_pass": report.dimension1_passed,
        "dim1_reason": "" if report.dimension1_passed else "；".join(report.dimension1_failures),
        "dim2_items": dim2_items,
        "total_score": total_score,
        "max_score": max_score,
    }


def _error_result(file_name: str, message: str) -> Dict[str, Any]:
    return {
        "id": SCRIPT_ID,
        "file_name": file_name,
        "status": "error",
        "error": message,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": sum(md for _, md in POSITIVE_RULES),
    }


def evaluate(dir_path: str) -> Dict[str, Any]:
    """统一入口：接收“脚本所在目录的路径”，自行在该目录里定位 .docx 并评估。

    返回结构见《脚本接口差异与统一建议.md》§2.2。
    """
    try:
        base = Path(dir_path)
        if not base.exists():
            return _error_result("", f"目录不存在：{dir_path}")
        if not base.is_dir():
            return _error_result("", f"传入路径不是目录：{dir_path}")

        docx_path = _find_target_docx(base)
        if docx_path is None:
            return _error_result("", f"目录中未找到 .docx 文档：{dir_path}")

        report = _evaluate_docx(docx_path)
        return _report_to_result(report, docx_path)
    except Exception as exc:  # pragma: no cover - 顶层兜底
        return _error_result("", f"评估程序异常：{exc}")


if __name__ == "__main__":
    # 本地自测：默认使用脚本自身所在目录，也可通过命令行参数传入其它目录。
    target_dir = sys.argv[1] if len(sys.argv) > 1 else str(Path(__file__).resolve().parent)
    print(json.dumps(evaluate(target_dir), ensure_ascii=False, indent=2))
