# -*- coding: utf-8 -*-
"""
自动评估脚本（officeval_021）：对开题报告 .docx 文档进行打分。

统一接口约定：
- 只暴露一个入口 evaluate(dir_path: str) -> dict
- dir_path 是脚本所在目录路径；脚本自己在该目录中定位并打开被评估的 .docx
- 不 print 主结果、不改 sys.stdout、不 sys.exit、不硬编码路径
- 返回结构化 dict，字段含义见 §2.2

评估逻辑：
1. 维度一（可用与可修改性）：不满足直接判零分
2. 维度二（完成度评分）：逐项检查得分点和扣分点，累计分数
"""
import os
import sys
from docx import Document
from docx.shared import Cm, Pt, Emu
from lxml import etree

W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def emu_to_cm(emu):
    """EMU转厘米"""
    return emu / 914400 * 2.54 if emu else 0


def twip_to_cm(twip):
    """Twip转厘米"""
    return int(twip) / 567.0 if twip else 0


def get_page_count_estimate(doc):
    """估算页数：通过分页符计算"""
    body = doc.element.body
    pages = 1
    for p in body.findall(f'.//{W_NS}br'):
        if p.get(f'{W_NS}type') == 'page':
            pages += 1
    return pages


def get_eastasia_font(run_elem):
    """获取run的eastAsia字体"""
    rPr = run_elem.find(f'{W_NS}rPr')
    if rPr is not None:
        rFonts = rPr.find(f'{W_NS}rFonts')
        if rFonts is not None:
            return rFonts.get(f'{W_NS}eastAsia')
    return None


def get_font_size_half_pt(run_elem):
    """获取run的字号(半磅单位)"""
    rPr = run_elem.find(f'{W_NS}rPr')
    if rPr is not None:
        sz = rPr.find(f'{W_NS}sz')
        if sz is not None:
            return int(sz.get(f'{W_NS}val'))
    return None


def is_bold(run_elem):
    """检查run是否加粗"""
    rPr = run_elem.find(f'{W_NS}rPr')
    if rPr is not None:
        b = rPr.find(f'{W_NS}b')
        if b is not None:
            val = b.get(f'{W_NS}val')
            return val != '0' and val != 'false'
        return False
    return False


def has_underline(run_elem):
    """检查run是否有下划线"""
    rPr = run_elem.find(f'{W_NS}rPr')
    if rPr is not None:
        u = rPr.find(f'{W_NS}u')
        if u is not None:
            val = u.get(f'{W_NS}val')
            return val is not None and val != 'none'
    return False


def get_paragraph_spacing(p_elem):
    """获取段落spacing属性"""
    pPr = p_elem.find(f'{W_NS}pPr')
    if pPr is not None:
        spacing = pPr.find(f'{W_NS}spacing')
        if spacing is not None:
            return {k.split('}')[1] if '}' in k else k: v for k, v in spacing.attrib.items()}
    return {}


def get_paragraph_jc(p_elem):
    """获取段落对齐方式"""
    pPr = p_elem.find(f'{W_NS}pPr')
    if pPr is not None:
        jc = pPr.find(f'{W_NS}jc')
        if jc is not None:
            return jc.get(f'{W_NS}val')
    return None


def get_full_text(doc):
    """获取文档全文"""
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    return ' '.join(texts)


def get_body_elements_with_pages(doc):
    """获取body元素及其所在页码的映射"""
    body = doc.element.body
    elements = []
    page = 1
    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'p':
            breaks = child.findall(f'.//{W_NS}br')
            has_pb = any(b.get(f'{W_NS}type') == 'page' for b in breaks)
            if has_pb:
                page += 1
        elements.append((child, tag, page))
    return elements


# ============== 维度一检查 ==============
def check_dimension1(filepath):
    """检查维度一：可用与可修改性"""
    results = []

    # 1. 文件格式检查
    if not os.path.exists(filepath):
        results.append(("文件存在且可打开", False, "文件不存在"))
        return results, False

    ext = os.path.splitext(filepath)[1].lower()
    if ext != '.docx':
        results.append(("文件格式为.docx", False, f"格式为{ext}"))
        return results, False

    try:
        _ = Document(filepath)
        results.append(("文件可正常打开", True, ""))
    except Exception as e:
        results.append(("文件可正常打开", False, str(e)))
        return results, False

    # 说明：按用户要求，删除以下维度一检查
    #   - "页数介于5-9页"
    #   - "无连续2页以上空白页/乱码/文字重叠"
    #   - "封面文字、正文、表格、参考文献、签署意见栏均可编辑，不能整篇转成图片或PDF"

    return results, True


# ============== 维度二检查 ==============
def check_dimension2(filepath):
    """检查维度二：完成度评分"""
    doc = Document(filepath)
    body = doc.element.body
    paras = body.findall(f'{W_NS}p')
    score_items = []  # (分数, 描述, 是否满足)
    full_text = get_full_text(doc)
    elements_pages = get_body_elements_with_pages(doc)

    sec = doc.sections[0]
    # --- +3: 文档页面为A4纵向，页边距 上2.2 下2.0 左2.5 右2.3(单位:厘米) ---
    # 细则要求的每一个点：
    #   1) 纸张为 A4（21.0cm × 29.7cm）
    #   2) 方向为纵向
    #   3) 上边距 = 2.2 cm
    #   4) 下边距 = 2.0 cm
    #   5) 左边距 = 2.5 cm
    #   6) 右边距 = 2.3 cm
    # 换算：Word 内部以 twip 存储，1 cm = 1440/2.54 twip ≈ 566.929 twip
    # 容差 0.05 cm：对齐 Word「页面设置」对话框的 cm 显示精度（同一显示值的合法波动范围）
    TWIP_PER_CM = 1440.0 / 2.54
    sectPr = body.find(f'{W_NS}sectPr')
    pgSz = sectPr.find(f'{W_NS}pgSz') if sectPr is not None else None
    pgMar = sectPr.find(f'{W_NS}pgMar') if sectPr is not None else None

    # 1) A4 尺寸判断（21.0×29.7cm）
    if pgSz is not None:
        page_w_cm = int(pgSz.get(f'{W_NS}w', '0')) / TWIP_PER_CM
        page_h_cm = int(pgSz.get(f'{W_NS}h', '0')) / TWIP_PER_CM
        orient = pgSz.get(f'{W_NS}orient')  # 缺省即为 portrait
    else:
        page_w_cm = emu_to_cm(sec.page_width)
        page_h_cm = emu_to_cm(sec.page_height)
        orient = None
    is_a4 = abs(page_w_cm - 21.0) < 0.05 and abs(page_h_cm - 29.7) < 0.05
    # 2) 纵向：orient 未标注为 landscape 且 宽 < 高（Word 中改方向两个条件同步变化）
    is_portrait = (orient != 'landscape') and (page_w_cm < page_h_cm)

    # 3-6) 四个边距
    if pgMar is not None:
        top_cm = int(pgMar.get(f'{W_NS}top', '0')) / TWIP_PER_CM
        bottom_cm = int(pgMar.get(f'{W_NS}bottom', '0')) / TWIP_PER_CM
        left_cm = int(pgMar.get(f'{W_NS}left', '0')) / TWIP_PER_CM
        right_cm = int(pgMar.get(f'{W_NS}right', '0')) / TWIP_PER_CM
    else:
        top_cm = emu_to_cm(sec.top_margin)
        bottom_cm = emu_to_cm(sec.bottom_margin)
        left_cm = emu_to_cm(sec.left_margin)
        right_cm = emu_to_cm(sec.right_margin)
    top_ok = abs(top_cm - 2.2) < 0.05
    bottom_ok = abs(bottom_cm - 2.0) < 0.05
    left_ok = abs(left_cm - 2.5) < 0.05
    right_ok = abs(right_cm - 2.3) < 0.05

    passed = is_a4 and is_portrait and top_ok and bottom_ok and left_ok and right_ok
    score_items.append((3, "A4纵向+页边距上2.2/下2.0/左2.5/右2.3cm", passed,
        f"页面:{page_w_cm:.2f}x{page_h_cm:.2f}cm(A4:{is_a4},纵向:{is_portrait}) "
        f"上{top_cm:.2f}({top_ok}) 下{bottom_cm:.2f}({bottom_ok}) "
        f"左{left_cm:.2f}({left_ok}) 右{right_cm:.2f}({right_ok})"))

    # --- +3: 第1页为封面页，封面上方1/3包含"澜州应用技术学院"、"毕业设计(论文)开题报告"字样和校徽图片 ---
    # 细则要求的每一个点：
    #   1) 第 1 页为封面页（有实际内容，以分页/节结束）
    #   2) 上方 1/3 区域内出现文本 "澜州应用技术学院"
    #   3) 上方 1/3 区域内出现文本 "毕业设计(论文)开题报告"（兼容全/半角括号）
    #   4) 上方 1/3 区域内出现校徽图片（w:drawing）
    # 判定方式（对齐 Word/WPS 实际排版）：
    #   逐段累加垂直高度(twip)，加上边距后与"页面顶端起 1/3 处"比较。
    #   段落高度 = before + max(图片高度, 行高) + after
    #     - 行高：lineRule=exact/atLeast → line 直接生效；auto → 最大字号(pt) * 20 * (line/240)
    #     - 图片高度：w:drawing 中 wp:extent@cy(EMU) 换算为 twip

    EMU_PER_TWIP = 914400.0 / 1440.0  # 1 twip = 635 EMU

    def _para_height_twip(p_elem):
        pPr = p_elem.find(f'{W_NS}pPr')
        spacing = pPr.find(f'{W_NS}spacing') if pPr is not None else None
        if spacing is not None:
            line_rule = spacing.get(f'{W_NS}lineRule', 'auto')
            line_val = int(spacing.get(f'{W_NS}line', '240'))
            before = int(spacing.get(f'{W_NS}before', '0'))
            after = int(spacing.get(f'{W_NS}after', '0'))
        else:
            line_rule, line_val, before, after = 'auto', 240, 0, 0
        max_sz_hp = 0
        for r in p_elem.findall(f'{W_NS}r'):
            sz_hp = get_font_size_half_pt(r)
            if sz_hp:
                max_sz_hp = max(max_sz_hp, sz_hp)
        if max_sz_hp == 0:
            max_sz_hp = 21  # 缺省 10.5pt(五号)
        font_pt = max_sz_hp / 2.0
        if line_rule in ('exact', 'atLeast'):
            line_h = line_val
        else:  # auto
            line_h = font_pt * 20 * (line_val / 240.0)
        image_h = 0
        for node in p_elem.iter():
            if node.tag.split('}')[-1] == 'extent':
                cy = node.get('cy')
                if cy:
                    image_h = max(image_h, int(cy) / EMU_PER_TWIP)
        return before + max(image_h, line_h) + after

    def _has_page_break(p_elem):
        for br in p_elem.findall(f'.//{W_NS}br'):
            if br.get(f'{W_NS}type') == 'page':
                return True
        return False

    page_h_twip = int(pgSz.get(f'{W_NS}h', '16838')) if pgSz is not None else 16838
    top_margin_twip = int(pgMar.get(f'{W_NS}top', '1440')) if pgMar is not None else 1440
    one_third_line = page_h_twip / 3.0  # 距页面顶端 1/3 处的 twip

    school_bottom = None
    title_bottom = None
    image_bottom = None
    cursor = top_margin_twip
    page1_has_content = False
    for p in paras:
        p_h = _para_height_twip(p)
        bottom_pos = cursor + p_h
        t_elems = p.findall(f'.//{W_NS}t')
        p_text = ''.join(t.text or '' for t in t_elems)
        if p_text.strip():
            page1_has_content = True
        if '澜州应用技术学院' in p_text and school_bottom is None:
            school_bottom = bottom_pos
        norm = p_text.replace('（', '(').replace('）', ')')
        if '毕业设计' in norm and '开题报告' in norm and title_bottom is None:
            title_bottom = bottom_pos
        if p.findall(f'.//{W_NS}drawing') and image_bottom is None:
            image_bottom = bottom_pos
        cursor = bottom_pos
        if _has_page_break(p):
            break  # 第 1 页结束

    is_cover_page = page1_has_content
    school_in_top_third = school_bottom is not None and school_bottom <= one_third_line
    title_in_top_third = title_bottom is not None and title_bottom <= one_third_line
    image_in_top_third = image_bottom is not None and image_bottom <= one_third_line
    passed = is_cover_page and school_in_top_third and title_in_top_third and image_in_top_third

    def _fmt(v):
        return f"{v/566.929:.2f}cm" if v is not None else "未找到"
    score_items.append((3, "第1页封面上方1/3含校名+标题+校徽图片", passed,
        f"封面存在:{is_cover_page} 1/3线:{one_third_line/566.929:.2f}cm "
        f"校名底部:{_fmt(school_bottom)}({school_in_top_third}) "
        f"标题底部:{_fmt(title_bottom)}({title_in_top_third}) "
        f"校徽底部:{_fmt(image_bottom)}({image_in_top_third})"))

    # --- +1: 封面页中的"澜州应用技术学院"字体格式为楷体_GB2312一号 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页（第 1 页，即首个 w:br@type=page 之前）内的"澜州应用技术学院"文本
    #   2) 字体：楷体_GB2312
    #   3) 字号：一号（=26pt，OOXML 中以半磅记为 52）
    # 判定方式（对齐 Word/WPS 实际渲染）：
    #   中文字符在 Word 的字体解析中使用 w:rFonts@w:eastAsia，因此校验该属性；
    #   若 run 的 w:rPr 未直接写 rFonts，则回退取段落级 w:pPr/w:rPr（Word 的属性继承顺序：
    #   run rPr > paragraph rPr > style > default），保证与 Word/WPS 实际显示一致。
    def _run_eastasia_effective(r_elem, p_elem):
        rPr = r_elem.find(f'{W_NS}rPr')
        if rPr is not None:
            rFonts = rPr.find(f'{W_NS}rFonts')
            if rFonts is not None:
                ea = rFonts.get(f'{W_NS}eastAsia')
                if ea:
                    return ea
        pPr = p_elem.find(f'{W_NS}pPr')
        if pPr is not None:
            pRpr = pPr.find(f'{W_NS}rPr')
            if pRpr is not None:
                rFonts = pRpr.find(f'{W_NS}rFonts')
                if rFonts is not None:
                    return rFonts.get(f'{W_NS}eastAsia')
        return None

    def _run_sz_effective(r_elem, p_elem):
        sz_hp = get_font_size_half_pt(r_elem)
        if sz_hp is not None:
            return sz_hp
        pPr = p_elem.find(f'{W_NS}pPr')
        if pPr is not None:
            pRpr = pPr.find(f'{W_NS}rPr')
            if pRpr is not None:
                sz = pRpr.find(f'{W_NS}sz')
                if sz is not None:
                    return int(sz.get(f'{W_NS}val'))
        return None

    school_runs = []  # (eastAsia字体, 半磅字号)
    for p in paras:
        if _has_page_break(p):
            break  # 越过第 1 页
        # 仅锁定"标题段落"：段落文本去空白后恰为"澜州应用技术学院"，
        # 从而排除信息区字段值(如"院  校：   澜州应用技术学院")——该段有自己的字体规则。
        p_text = ''.join(t.text or '' for t in p.findall(f'.//{W_NS}t'))
        if p_text.replace(' ', '').replace('　', '').strip() != '澜州应用技术学院':
            continue
        for r in p.findall(f'{W_NS}r'):
            r_text = ''.join(t.text or '' for t in r.findall(f'{W_NS}t'))
            if '澜州应用技术学院' in r_text:
                school_runs.append((
                    _run_eastasia_effective(r, p),
                    _run_sz_effective(r, p),
                ))
    # 一号 = 26pt = 52 半磅
    all_font_ok = bool(school_runs) and all(ea == '楷体_GB2312' for ea, _ in school_runs)
    all_size_ok = bool(school_runs) and all(sz == 52 for _, sz in school_runs)
    passed = all_font_ok and all_size_ok
    score_items.append((1, "\"澜州应用技术学院\"楷体_GB2312一号", passed,
        f"命中run数={len(school_runs)} 字体全部为楷体_GB2312:{all_font_ok} "
        f"字号全部为一号(52半磅):{all_size_ok} 明细={school_runs}"))

    # --- +1: 封面页中的"澜州应用技术学院"段落格式为居中对齐、段后0.2行、1.5倍行距 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)内、段落文本恰为"澜州应用技术学院"的段
    #   2) 居中对齐
    #   3) 段后 0.2 行
    #   4) 1.5 倍行距
    # 判定方式（对齐 Word/WPS 的段落对话框）：
    #   - 居中：w:pPr/w:jc@w:val = "center"
    #   - 段后 0.2 行：Word「段落」对话框中，间距单位为"行"时写入 w:afterLines
    #                   afterLines 以 1/100 行为单位 → 0.2 行 = 20
    #   - 1.5 倍行距：Word「段落」对话框中，"1.5 倍行距"写入 lineRule=auto、line=360
    #                   （240 = 单倍行距，360 = 1.5 倍）
    school_para = None
    for p in paras:
        if _has_page_break(p):
            break
        p_text = ''.join(t.text or '' for t in p.findall(f'.//{W_NS}t'))
        if p_text.replace(' ', '').replace('　', '').strip() == '澜州应用技术学院':
            school_para = p
            break

    if school_para is None:
        passed = False
        score_items.append((1, "\"澜州应用技术学院\"居中+段后0.2行+1.5倍行距", False,
            "未定位到\"澜州应用技术学院\"标题段"))
    else:
        jc = get_paragraph_jc(school_para)
        sp = get_paragraph_spacing(school_para)
        is_center = (jc == 'center')
        # 段后 0.2 行 = afterLines=20
        after_lines_ok = sp.get('afterLines') == '20'
        # 1.5 倍行距 = lineRule=auto 且 line=360
        line_rule_ok = sp.get('lineRule', 'auto') == 'auto'
        line_val_ok = sp.get('line') == '360'
        passed = is_center and after_lines_ok and line_rule_ok and line_val_ok
        score_items.append((1, "\"澜州应用技术学院\"居中+段后0.2行+1.5倍行距", passed,
            f"jc={jc}(居中:{is_center}) afterLines={sp.get('afterLines')}(需20:{after_lines_ok}) "
            f"lineRule={sp.get('lineRule')}(需auto:{line_rule_ok}) "
            f"line={sp.get('line')}(需360:{line_val_ok})"))

    # --- +1: 封面页中的"毕业设计（论文）开题报告"字体黑体小初，位于"澜州应用技术学院"下方，居中显示 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)内、段落文本为"毕业设计（论文）开题报告"的段
    #      （兼容全角"（）"与半角"()"括号——办公软件保存/输入法差异不改变本义）
    #   2) 字体：黑体（OOXML w:rFonts@w:eastAsia = "SimHei" 或 "黑体"，Word/WPS 两种写法都会渲染为黑体）
    #   3) 字号：小初 = 36pt = 72 半磅（w:sz@w:val = "72"）
    #   4) 位置：位于"澜州应用技术学院"段落的下方（在文档流中处于其之后）
    #   5) 对齐：居中显示（w:pPr/w:jc@w:val = "center"）

    school_idx = None
    title_idx = None
    for i, p in enumerate(paras):
        if _has_page_break(p):
            break  # 只在第 1 页内查找
        p_text = ''.join(t.text or '' for t in p.findall(f'.//{W_NS}t'))
        p_clean = p_text.replace(' ', '').replace('　', '').strip()
        if school_idx is None and p_clean == '澜州应用技术学院':
            school_idx = i
        # 兼容全/半角括号
        norm = p_clean.replace('（', '(').replace('）', ')')
        if title_idx is None and norm == '毕业设计(论文)开题报告':
            title_idx = i

    if title_idx is None:
        passed = False
        score_items.append((1, "\"毕业设计（论文）开题报告\"黑体小初+位于校名下方+居中",
            False, "未在封面页定位到\"毕业设计（论文）开题报告\"段"))
    else:
        title_para = paras[title_idx]
        # 字体、字号：段内所有含目标文本的 run 都必须满足
        title_runs = []
        for r in title_para.findall(f'{W_NS}r'):
            r_text = ''.join(t.text or '' for t in r.findall(f'{W_NS}t'))
            r_norm = r_text.replace('（', '(').replace('）', ')')
            if '毕业设计' in r_norm or '开题报告' in r_norm:
                title_runs.append((
                    _run_eastasia_effective(r, title_para),
                    _run_sz_effective(r, title_para),
                ))
        font_ok = bool(title_runs) and all(
            ea in ('SimHei', '黑体') for ea, _ in title_runs)
        size_ok = bool(title_runs) and all(sz == 72 for _, sz in title_runs)
        # 位置：位于"澜州应用技术学院"之下
        below_school = (school_idx is not None) and (title_idx > school_idx)
        # 居中
        jc_val = get_paragraph_jc(title_para)
        center_ok = (jc_val == 'center')
        passed = font_ok and size_ok and below_school and center_ok
        score_items.append((1, "\"毕业设计（论文）开题报告\"黑体小初+位于校名下方+居中", passed,
            f"字体黑体:{font_ok} 字号小初(72半磅):{size_ok} "
            f"位于校名下方(校名idx={school_idx},标题idx={title_idx}):{below_school} "
            f"jc={jc_val}(居中:{center_ok}) run明细={title_runs}"))

    # --- +1: 封面页中的"毕业设计（论文）开题报告"段落格式为居中对齐、文本之后-3.01字符、1.5倍行距 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)内、段落文本为"毕业设计（论文）开题报告"的段
    #      （兼容全/半角括号，输入法/办公软件差异不改变文义）
    #   2) 居中对齐
    #   3) 文本之后 -3.01 字符（Word「段落」对话框中"缩进-文本之后"以"字符"为单位时写入 rightChars，
    #      单位为 1/100 字符 → -3.01 字符 = rightChars = -301）
    #   4) 1.5 倍行距（Word「段落」对话框选"1.5 倍行距"时写入 lineRule=auto、line=360）
    if title_idx is None:
        passed = False
        score_items.append((1, "\"毕业设计（论文）开题报告\"居中+文本之后-3.01字符+1.5倍行距",
            False, "未在封面页定位到\"毕业设计（论文）开题报告\"段"))
    else:
        title_para = paras[title_idx]
        jc_val = get_paragraph_jc(title_para)
        sp_val = get_paragraph_spacing(title_para)
        pPr_t = title_para.find(f'{W_NS}pPr')
        ind_t = pPr_t.find(f'{W_NS}ind') if pPr_t is not None else None
        right_chars = ind_t.get(f'{W_NS}rightChars') if ind_t is not None else None

        center_ok = (jc_val == 'center')
        # 文本之后 -3.01 字符 → rightChars = -301
        right_chars_ok = (right_chars == '-301')
        # 1.5 倍行距 → lineRule=auto & line=360
        line_rule_ok = (sp_val.get('lineRule', 'auto') == 'auto')
        line_val_ok = (sp_val.get('line') == '360')

        passed = center_ok and right_chars_ok and line_rule_ok and line_val_ok
        score_items.append((1, "\"毕业设计（论文）开题报告\"居中+文本之后-3.01字符+1.5倍行距", passed,
            f"jc={jc_val}(居中:{center_ok}) rightChars={right_chars}(需-301:{right_chars_ok}) "
            f"lineRule={sp_val.get('lineRule')}(需auto:{line_rule_ok}) "
            f"line={sp_val.get('line')}(需360:{line_val_ok})"))

    # --- +1: 封面页下方信息区包含字段"院校：""专业：""学号：""姓名：""指导教师："及对应内容 ---
    # 细则要求的每一个点：
    #   1) 位置：封面页(第 1 页)、且在标题"毕业设计（论文）开题报告"下方(即信息区在页面下方)
    #   2) 必须包含 5 个字段标签：院校 / 专业 / 学号 / 姓名 / 指导教师(后接冒号)
    #   3) 每个字段都必须有对应内容(冒号之后有实质文本)
    # 判定方式(对齐 Word/WPS 实际显示)：
    #   - 从"标题段"之后、第 1 页分页符之前的段落中查找字段
    #   - 字段名允许排版加空格填充(如"院    校：")→ 去空格后匹配
    #   - 冒号兼容全角"："与半角":"
    required_labels = ['院校', '专业', '学号', '姓名', '指导教师']

    # 信息区范围：标题段之后 ~ 第 1 页分页符之前
    info_start = (title_idx + 1) if title_idx is not None else 0
    info_end = len(paras)
    for i in range(info_start, len(paras)):
        if _has_page_break(paras[i]):
            info_end = i + 1  # 含当前段(分页符附着段本身可能有内容)
            break

    field_status = {}  # label -> (是否找到标签, 是否有对应内容)
    for label in required_labels:
        field_status[label] = (False, False)

    for i in range(info_start, info_end):
        p_text = ''.join(t.text or '' for t in paras[i].findall(f'.//{W_NS}t'))
        # 归一化：去除空格/全角空格，全角冒号 → 半角
        norm = p_text.replace(' ', '').replace('　', '').replace('：', ':')
        for label in required_labels:
            if field_status[label][0]:
                continue
            if label + ':' in norm:
                # 冒号之后有实质内容
                after = norm.split(label + ':', 1)[1].strip()
                has_value = len(after) > 0
                field_status[label] = (True, has_value)

    all_labels_ok = all(v[0] for v in field_status.values())
    all_values_ok = all(v[1] for v in field_status.values())
    passed = all_labels_ok and all_values_ok
    score_items.append((1, "封面下方信息区含院校/专业/学号/姓名/指导教师及对应内容", passed,
        f"标签全部命中:{all_labels_ok} 内容全部非空:{all_values_ok} "
        f"明细={ {k: {'标签':v[0],'内容':v[1]} for k,v in field_status.items()} }"))

    # --- +3: 封面页下方信息区字段"院校：""专业：""学号：""姓名：""指导教师："字体为宋体三号加粗 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中标签为"院校""专业""学号""姓名""指导教师"(含冒号)的 5 个字段
    #   2) 字体：宋体
    #   3) 字号：三号 = 16pt = 32 半磅
    #   4) 加粗
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   - 中文字体解析走 w:rFonts@w:eastAsia：宋体 = "SimSun" 或 "宋体"(Word/WPS 都识别)
    #   - 字号：w:sz@w:val == "32"
    #   - 加粗：w:b 存在且未被显式关闭(val 不为 "0"/"false")
    #   - 每个标签所在 run 单独判定，5 个字段都要满足
    required_labels_bold = ['院校', '专业', '学号', '姓名', '指导教师']

    def _label_run(p_elem, label):
        """在段落中找到显示文本(去空白后)以 label+冒号 开头的 run。"""
        for r in p_elem.findall(f'{W_NS}r'):
            r_text = ''.join(t.text or '' for t in r.findall(f'{W_NS}t'))
            r_norm = r_text.replace(' ', '').replace('　', '').replace('：', ':')
            if r_norm.startswith(label + ':'):
                return r
        return None

    label_check = {}  # label -> dict(found,font_ok,size_ok,bold_ok,detail)
    for label in required_labels_bold:
        label_check[label] = {'found': False, 'font_ok': False,
                              'size_ok': False, 'bold_ok': False,
                              'ea': None, 'sz': None, 'bold': None}
        for i in range(info_start, info_end):
            r = _label_run(paras[i], label)
            if r is None:
                continue
            # 排除"指导教师意见"这种表格里出现的同名文本(此处仅遍历信息区段落,不含表格,安全)
            ea = _run_eastasia_effective(r, paras[i])
            sz = _run_sz_effective(r, paras[i])
            bold = is_bold(r)
            label_check[label] = {
                'found': True,
                'font_ok': ea in ('SimSun', '宋体'),
                'size_ok': (sz == 32),
                'bold_ok': bool(bold),
                'ea': ea, 'sz': sz, 'bold': bold,
            }
            break

    all_found = all(v['found'] for v in label_check.values())
    all_font_ok = all_found and all(v['font_ok'] for v in label_check.values())
    all_size_ok = all_found and all(v['size_ok'] for v in label_check.values())
    all_bold_ok = all_found and all(v['bold_ok'] for v in label_check.values())
    passed = all_found and all_font_ok and all_size_ok and all_bold_ok

    detail_lines = []
    for lbl in required_labels_bold:
        v = label_check[lbl]
        detail_lines.append(f"{lbl}:ea={v['ea']}(宋体:{v['font_ok']}),"
                            f"sz={v['sz']}(三号:{v['size_ok']}),"
                            f"bold={v['bold']}(加粗:{v['bold_ok']})")
    score_items.append((3, "信息区字段标签(院校/专业/学号/姓名/指导教师)宋体三号加粗", passed,
        f"5标签全命中:{all_found} 字体全宋体:{all_font_ok} 字号全三号:{all_size_ok} "
        f"全部加粗:{all_bold_ok} | " + " | ".join(detail_lines)))

    # --- +3: 信息区内容字体为宋体三号 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中，内容为下列 4 项之一的字符片段：
    #      "澜州应用技术学院""机械装备设计与自动化""许芷涵""秦若衡"
    #   2) 中文字体：宋体（w:rFonts@w:eastAsia = "宋体" 或 "SimSun"，
    #      与 Word/WPS 中"宋体"实际渲染一致）
    #   3) 字号：三号 = 16pt = 32 半磅（w:sz@w:val == "32"）
    # 判定方式（对齐 Word/WPS 实际渲染）：
    #   - 内容可能被拆成多个 run（如"许芷涵" → "许" + "芷涵"），先在段落里
    #     把 run 文本按顺序拼接成一条 concat，定位 needle 的 [idx, end)，再
    #     回查覆盖该区间的所有 run；每个被覆盖的 run 都必须 eastAsia 命中宋体
    #     且字号为 32（半磅）——沿 run rPr → 段落 pPr rPr 继承链解析。
    #   - 4 个内容全部找到且全部合规才加分；避免旧写法用固定 paras[4:10]
    #     +"任意 run 字号≠32 才算错"这类脆弱条件误判。
    content_font_targets = ['澜州应用技术学院', '机械装备设计与自动化',
                            '许芷涵', '秦若衡']

    content_font_check = {}
    for needle in content_font_targets:
        content_font_check[needle] = {
            'found_runs': 0, 'font_ok': False, 'size_ok': False,
            'ea_list': [], 'sz_list': [],
        }
        for i in range(info_start, info_end):
            p_text_full = ''.join(t.text or '' for t in paras[i].findall(f'.//{W_NS}t'))
            if needle not in p_text_full.replace(' ', '').replace('　', ''):
                continue
            # 用与下划线检查同款做法：拼接段落 run 文本，定位 needle 覆盖的 run 集合
            runs = paras[i].findall(f'{W_NS}r')
            run_texts = []
            for r in runs:
                run_texts.append(''.join(t.text or '' for t in r.findall(f'{W_NS}t')))
            concat = ''.join(run_texts)
            idx = concat.find(needle)
            if idx < 0:
                continue
            end = idx + len(needle)
            covered = []
            cur = 0
            for ri, rt in enumerate(run_texts):
                r_start, r_end = cur, cur + len(rt)
                if r_end > idx and r_start < end:
                    covered.append(ri)
                cur = r_end
            if not covered:
                continue
            ea_list = [_run_eastasia_effective(runs[ri], paras[i]) for ri in covered]
            sz_list = [_run_sz_effective(runs[ri], paras[i]) for ri in covered]
            content_font_check[needle] = {
                'found_runs': len(covered),
                'font_ok': all(ea in ('SimSun', '宋体') for ea in ea_list),
                'size_ok': all(sz == 32 for sz in sz_list),
                'ea_list': ea_list, 'sz_list': sz_list,
            }
            break

    all_found = all(v['found_runs'] > 0 for v in content_font_check.values())
    all_font_ok = all_found and all(v['font_ok'] for v in content_font_check.values())
    all_size_ok = all_found and all(v['size_ok'] for v in content_font_check.values())
    content_font_ok = all_found and all_font_ok and all_size_ok
    detail = " | ".join(
        f"{k}:runs={v['found_runs']},ea={v['ea_list']}(宋体:{v['font_ok']}),"
        f"sz={v['sz_list']}(三号:{v['size_ok']})"
        for k, v in content_font_check.items()
    )
    score_items.append((3, "信息区4项内容(院校/专业/姓名x2)宋体三号", content_font_ok,
        f"4项全部命中:{all_found} 全宋体:{all_font_ok} 全三号:{all_size_ok} | " + detail))

    # --- +3: 封面页下方信息区字段"澜州应用技术学院""机械装备设计与自动化""2360140827"
    #          "许芷涵""秦若衡"对应内容段落格式为居中对齐、段前4磅、1.35倍行距 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中,内容为下列 5 个之一的段落：
    #      "澜州应用技术学院""机械装备设计与自动化""2360140827""许芷涵""秦若衡"
    #   2) 居中对齐：w:pPr/w:jc@w:val = "center"
    #   3) 段前 4 磅：Word「段落」对话框中"段前"以"磅"为单位时写入 w:before(单位 twip,
    #      1 磅 = 20 twip),4 磅 = 80 twip;且不写 w:beforeLines(否则单位不是"磅")
    #   4) 1.35 倍行距：Word「段落」对话框选"多倍行距 1.35"时写入 lineRule=auto、line=324
    #      (240 = 单倍行距,240 * 1.35 = 324)
    content_targets = ['澜州应用技术学院', '机械装备设计与自动化',
                       '2360140827', '许芷涵', '秦若衡']

    def _para_contains_content(p_elem, needle):
        """段落文本(去空白)包含 needle;且 needle 出现在冒号之后(排除标签本身)。"""
        p_text = ''.join(t.text or '' for t in p_elem.findall(f'.//{W_NS}t'))
        norm = p_text.replace(' ', '').replace('　', '').replace('：', ':')
        if ':' in norm:
            after = norm.split(':', 1)[1]
            return needle in after
        return needle in norm

    content_check = {}
    for needle in content_targets:
        content_check[needle] = {
            'found': False,
            'jc': None, 'before': None, 'beforeLines': None,
            'lineRule': None, 'line': None,
            'center_ok': False, 'before_ok': False, 'line_ok': False,
        }
        for i in range(info_start, info_end):
            if not _para_contains_content(paras[i], needle):
                continue
            jc_v = get_paragraph_jc(paras[i])
            sp_v = get_paragraph_spacing(paras[i])
            before_v = sp_v.get('before')
            before_lines_v = sp_v.get('beforeLines')
            line_rule_v = sp_v.get('lineRule', 'auto')
            line_v = sp_v.get('line')

            center_ok = (jc_v == 'center')
            # 段前 4 磅 = w:before = 80 twip,且不能同时用 beforeLines(否则单位变为"行")
            before_ok = (before_v == '80') and (before_lines_v in (None, '0'))
            # 1.35 倍行距 = lineRule=auto & line=324
            line_ok = (line_rule_v == 'auto') and (line_v == '324')

            content_check[needle] = {
                'found': True,
                'jc': jc_v, 'before': before_v, 'beforeLines': before_lines_v,
                'lineRule': line_rule_v, 'line': line_v,
                'center_ok': center_ok, 'before_ok': before_ok, 'line_ok': line_ok,
            }
            break

    all_found = all(v['found'] for v in content_check.values())
    all_center = all_found and all(v['center_ok'] for v in content_check.values())
    all_before = all_found and all(v['before_ok'] for v in content_check.values())
    all_line = all_found and all(v['line_ok'] for v in content_check.values())
    passed = all_found and all_center and all_before and all_line

    detail = []
    for k, v in content_check.items():
        detail.append(f"{k}:jc={v['jc']}({v['center_ok']}),"
                      f"before={v['before']}/beforeLines={v['beforeLines']}({v['before_ok']}),"
                      f"lineRule={v['lineRule']}/line={v['line']}({v['line_ok']})")
    score_items.append((3, "信息区5项内容居中+段前4磅+1.35倍行距", passed,
        f"5项全部命中:{all_found} 全居中:{all_center} 全段前4磅:{all_before} "
        f"全1.35倍行距:{all_line} | " + " | ".join(detail)))

    # --- +3: 封面页下方信息区字段"澜州应用技术学院""机械装备设计与自动化"
    #          "2360140827""许芷涵""秦若衡"内容下方带有下划线 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中,内容为下列 5 项之一的字符片段：
    #      "澜州应用技术学院""机械装备设计与自动化""2360140827""许芷涵""秦若衡"
    #   2) 每项内容的下方带有下划线(即办公软件中该文字底部有一条水平线)
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   - Word/WPS 中"下划线"以 run 级 w:rPr/w:u@w:val 表达,值不为 "none" 且存在 val 属性
    #     即为可见下划线(single/double/thick/dotted/dash/wave...)。
    #   - 一个内容可能跨多个 w:t/run(如学号"2360" + "140827"两段),因此逐 run 检查:
    #     只要"包含该内容的每个 run"都带下划线,即视为"该内容下方带有下划线"。
    #   - 中文/数字/字母都遵循 w:u,与办公软件所见即所得一致。
    underline_targets = ['澜州应用技术学院', '机械装备设计与自动化',
                         '2360140827', '许芷涵', '秦若衡']

    def _run_has_visible_underline(r_elem, p_elem):
        """判断 run 是否带可见下划线(遵循 run rPr > 段落 pPr rPr 的继承)。"""
        rPr = r_elem.find(f'{W_NS}rPr')
        if rPr is not None:
            u = rPr.find(f'{W_NS}u')
            if u is not None:
                val = u.get(f'{W_NS}val')
                # 显式设置了下划线样式(非 none)
                if val and val != 'none':
                    return True
                # 显式关闭
                if val == 'none':
                    return False
        # 回退段落级 rPr
        pPr = p_elem.find(f'{W_NS}pPr')
        if pPr is not None:
            pRpr = pPr.find(f'{W_NS}rPr')
            if pRpr is not None:
                u = pRpr.find(f'{W_NS}u')
                if u is not None:
                    val = u.get(f'{W_NS}val')
                    return bool(val) and val != 'none'
        return False

    ul_check = {}
    for needle in underline_targets:
        ul_check[needle] = {'found_runs': 0, 'all_ul': False, 'partial': []}
        for i in range(info_start, info_end):
            # 只在段落文本包含 needle 的段中检查
            p_text_full = ''.join(t.text or '' for t in paras[i].findall(f'.//{W_NS}t'))
            if needle not in p_text_full.replace(' ', '').replace('　', ''):
                continue
            # 收集"该段落里显示文本(拼接后)命中 needle 的连续 run 集合"
            # 简化并准确的做法:遍历所有 run,把 run 文本拼接成一个偏移映射,
            # 找到 needle 在整段拼接文本中的所有起止,再回查这些区间覆盖的 run。
            runs = paras[i].findall(f'{W_NS}r')
            run_texts = []
            for r in runs:
                run_texts.append(''.join(t.text or '' for t in r.findall(f'{W_NS}t')))
            concat = ''.join(run_texts)
            # 定位第一处 needle
            idx = concat.find(needle)
            if idx < 0:
                continue
            end = idx + len(needle)
            # 找出被 [idx, end) 覆盖的 run 索引
            covered = []
            cur = 0
            for ri, rt in enumerate(run_texts):
                r_start, r_end = cur, cur + len(rt)
                if r_end > idx and r_start < end:
                    covered.append(ri)
                cur = r_end
            if not covered:
                continue
            per_run_ul = [_run_has_visible_underline(runs[ri], paras[i]) for ri in covered]
            ul_check[needle] = {
                'found_runs': len(covered),
                'all_ul': all(per_run_ul),
                'partial': per_run_ul,
            }
            break

    all_found = all(v['found_runs'] > 0 for v in ul_check.values())
    all_ul_ok = all_found and all(v['all_ul'] for v in ul_check.values())
    passed = all_found and all_ul_ok
    detail = " | ".join(
        f"{k}:runs={v['found_runs']},全带下划线:{v['all_ul']},逐run={v['partial']}"
        for k, v in ul_check.items())
    score_items.append((3, "信息区5项内容下方带有下划线", passed,
        f"5项全部命中:{all_found} 全部带下划线:{all_ul_ok} | " + detail))

    # --- +5: 内容左右有空字符串且带下划线 ---
    # --- +5: 封面页下方信息区字段"澜州应用技术学院""机械设计制造及其自动化"
    #          "2360140827""许芷涵""秦若衡"内容左右两边出现个数误差不超过 2 的空字符串
    #          且空字符串的下方有下划线 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中,内容为下列 5 项之一的字符片段：
    #      "澜州应用技术学院""机械设计制造及其自动化""2360140827""许芷涵""秦若衡"
    #   2) 内容左侧存在空字符串(纯空格串,允许半角/全角空格)
    #   3) 内容右侧存在空字符串(纯空格串)
    #   4) 左右两侧空字符串的字符个数误差不超过 2 (|left_count - right_count| <= 2)
    #   5) 左侧空字符串下方有下划线(其所在 run 带 w:u,val != none)
    #   6) 右侧空字符串下方有下划线
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   - 内容可能被拆成多个 run;先把段落 run 文本拼接,定位 needle 在整段中的起止位置,
    #     再以该起止为边界向左/向右在同一段落中截取连续的空格字符,统计其字符数;
    #     同时以字符所属的 run 判断下划线属性。
    space_underline_targets = ['澜州应用技术学院', '机械设计制造及其自动化',
                               '2360140827', '许芷涵', '秦若衡']

    def _char_run_map(p_elem):
        """返回 (拼接文本, 每个字符对应的 run 元素列表)。"""
        runs = p_elem.findall(f'{W_NS}r')
        chars = []
        run_of_char = []
        for r in runs:
            r_text = ''.join(t.text or '' for t in r.findall(f'{W_NS}t'))
            for ch in r_text:
                chars.append(ch)
                run_of_char.append(r)
        return ''.join(chars), run_of_char

    space_check = {}
    for needle in space_underline_targets:
        space_check[needle] = {
            'found': False,
            'left_count': 0, 'right_count': 0,
            'diff_ok': False,
            'left_ul': False, 'right_ul': False,
        }
        for i in range(info_start, info_end):
            concat, run_map = _char_run_map(paras[i])
            idx = concat.find(needle)
            if idx < 0:
                continue
            end = idx + len(needle)
            # 向左延伸：紧邻 needle 左侧的连续空格串
            left_start = idx
            while left_start > 0 and concat[left_start - 1] in (' ', '　'):
                left_start -= 1
            left_str = concat[left_start:idx]
            # 向右延伸：紧邻 needle 右侧的连续空格串
            right_end = end
            while right_end < len(concat) and concat[right_end] in (' ', '　'):
                right_end += 1
            right_str = concat[end:right_end]

            left_count = len(left_str)
            right_count = len(right_str)
            diff_ok = abs(left_count - right_count) <= 2
            # 左空格串各字符所属 run 都必须带可见下划线;右侧同理。
            left_ul = left_count > 0 and all(
                _run_has_visible_underline(run_map[k], paras[i])
                for k in range(left_start, idx))
            right_ul = right_count > 0 and all(
                _run_has_visible_underline(run_map[k], paras[i])
                for k in range(end, right_end))

            space_check[needle] = {
                'found': True,
                'left_count': left_count, 'right_count': right_count,
                'diff_ok': diff_ok,
                'left_ul': left_ul, 'right_ul': right_ul,
            }
            break

    all_found = all(v['found'] for v in space_check.values())
    all_diff = all_found and all(v['diff_ok'] for v in space_check.values())
    all_left_ul = all_found and all(v['left_ul'] for v in space_check.values())
    all_right_ul = all_found and all(v['right_ul'] for v in space_check.values())
    passed = all_found and all_diff and all_left_ul and all_right_ul
    detail = " | ".join(
        f"{k}:左空{v['left_count']}/右空{v['right_count']}(误差≤2:{v['diff_ok']}),"
        f"左下划线:{v['left_ul']},右下划线:{v['right_ul']}"
        for k, v in space_check.items())
    score_items.append((5, "信息区5项内容左右空字符串+空串下方下划线+个数误差≤2", passed,
        f"5项全命中:{all_found} 误差全≤2:{all_diff} 左空全下划线:{all_left_ul} "
        f"右空全下划线:{all_right_ul} | " + detail))


    # --- +5: 封面页下方信息区字段"澜州应用技术学院""机械设计制造及其自动化"
    #          "2360140827""许芷涵""秦若衡"内容下方下划线右对齐 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中,内容为下列 5 项之一的段落：
    #      "澜州应用技术学院""机械设计制造及其自动化""2360140827""许芷涵""秦若衡"
    #   2) 每项内容"下方的下划线"(即该内容及其两侧带下划线的空字符串)在办公软件中
    #      呈现为"右对齐"——所有行的下划线右端在版面上位于同一横向位置。
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   在 Word/WPS 中,只有段落"对齐方式 = 右对齐"(w:pPr/w:jc@w:val = "right")时,
    #   不论内容长短,该段的文字/下划线右端都会紧贴段落的右边界(= 页面右边距 - 右缩进),
    #   多行之间在版面上完全对齐。因此严格判定条件为:
    #     a) 5 项内容所在段落 w:jc = "right"
    #     b) 5 项内容所在段落的右缩进(w:ind@w:right 与 w:ind@w:rightChars)完全一致
    #        (保证 5 行的右边界处于同一 x 坐标)
    # --- +5: 封面页下方信息区字段"澜州应用技术学院""机械设计制造及其自动化"
    #          "2360140827""许芷涵""秦若衡"内容下方下划线右对齐 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中,内容为下列 5 项之一的段落：
    #      "澜州应用技术学院""机械设计制造及其自动化""2360140827""许芷涵""秦若衡"
    #   2) 每项内容"下方的下划线"(即该内容及其两侧带下划线的空字符串)在办公软件中
    #      呈现为"右对齐"——所有行的下划线右端在版面上位于同一横向 x 坐标。
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   在 Word/WPS 中把一段文字/下划线右端固定在页面某个横向位置,不止"段落
    #   右对齐"一种做法。因此不再把"jc=right + 右缩进一致"作为唯一判据,
    #   改为分别推算每个内容所在段落的**下划线视觉右边界**(以 twip 为单位,
    #   相对于版心左侧),再比较 5 行是否落在同一 x 坐标(容差 30 twip ≈ 1.5pt):
    #     A) 段落 jc="right"：下划线尾端紧贴版心右边界 -
    #        (w:ind@w:right 或 rightChars * chars_width)
    #     B) 段落挂载"右对齐制表位" w:pPr/w:tabs/w:tab@val="right"@pos=P，
    #        且 run 流中出现 <w:tab/> 制表符：下划线右端 = P
    #        (若存在多个右对齐制表位,取位于内容之后的最靠右者)
    #     C) 其它情况(左对齐+空格填充等)：无法从 OOXML 精确推算下划线右边界,
    #        视为无法证明右对齐 —— 该项不合规
    right_align_targets = ['澜州应用技术学院', '机械设计制造及其自动化',
                           '2360140827', '许芷涵', '秦若衡']

    # 版心宽度(twip)：页面宽 - 左边距 - 右边距，用于 jc="right" 情形
    try:
        page_w_twip = int(pgSz.get(f'{W_NS}w', '0')) if pgSz is not None else 0
        left_m_twip = int(pgMar.get(f'{W_NS}left', '0')) if pgMar is not None else 0
        right_m_twip = int(pgMar.get(f'{W_NS}right', '0')) if pgMar is not None else 0
    except (TypeError, ValueError):
        page_w_twip = left_m_twip = right_m_twip = 0
    content_w_twip = max(0, page_w_twip - left_m_twip - right_m_twip)

    def _underline_right_end_twip(p_elem):
        """推算段落中下划线视觉右端的 x 坐标(twip, 相对版心左侧)。
        无法推算返回 (None, '无法证明右对齐')。"""
        pPr = p_elem.find(f'{W_NS}pPr')
        jc_v = get_paragraph_jc(p_elem)
        ind_i = pPr.find(f'{W_NS}ind') if pPr is not None else None
        # 右缩进(twip)：w:right 优先，其次 w:rightChars 按 100 折算暂无字符宽度信息，
        # 此处仅纳入 w:right(与 Word/WPS 对齐方式对话框显示单位一致)
        right_ind_twip = 0
        if ind_i is not None:
            rv = ind_i.get(f'{W_NS}right')
            if rv:
                try:
                    right_ind_twip = int(rv)
                except ValueError:
                    right_ind_twip = 0

        # A) 整段右对齐 → 下划线尾端 = 版心宽 - 右缩进
        if jc_v == 'right' and content_w_twip > 0:
            return content_w_twip - right_ind_twip, f"jc=right(right_ind={right_ind_twip})"

        # B) 右对齐制表位 + run 流中出现 <w:tab/> → 下划线尾端 = 制表位 pos
        if pPr is not None:
            tabs = pPr.find(f'{W_NS}tabs')
            right_tab_positions = []
            if tabs is not None:
                for tab in tabs.findall(f'{W_NS}tab'):
                    if tab.get(f'{W_NS}val') == 'right':
                        pv = tab.get(f'{W_NS}pos')
                        if pv:
                            try:
                                right_tab_positions.append(int(pv))
                            except ValueError:
                                pass
            has_tab_char = any(
                r.find(f'{W_NS}tab') is not None
                for r in p_elem.findall(f'{W_NS}r')
            )
            if right_tab_positions and has_tab_char:
                # 取最靠右的右对齐制表位作为下划线视觉右端
                return max(right_tab_positions), f"tab-right(pos={max(right_tab_positions)})"

        # C) 其它情况 —— 无法证明右对齐
        return None, f"jc={jc_v},right_ind={right_ind_twip},未见右对齐制表位或段落右对齐"

    ra_check = {}
    for needle in right_align_targets:
        ra_check[needle] = {'found': False, 'right_end': None, 'evidence': ''}
        for i in range(info_start, info_end):
            concat, _ = _char_run_map(paras[i])
            if needle not in concat:
                continue
            right_end, evidence = _underline_right_end_twip(paras[i])
            ra_check[needle] = {
                'found': True,
                'right_end': right_end,
                'evidence': evidence,
            }
            break

    all_found_ra = all(v['found'] for v in ra_check.values())
    all_resolved = all_found_ra and all(v['right_end'] is not None for v in ra_check.values())
    # 5 行下划线右端 x 坐标一致(容差 30 twip ≈ 1.5pt，与 Word 磅级排版精度一致)
    right_edges_uniform = False
    if all_resolved:
        edges = [v['right_end'] for v in ra_check.values()]
        right_edges_uniform = (max(edges) - min(edges)) <= 30
    passed = all_found_ra and all_resolved and right_edges_uniform
    detail = " | ".join(
        f"{k}:right_end={v['right_end']}({v['evidence']})"
        for k, v in ra_check.items()
    )
    edges_str = [v['right_end'] for v in ra_check.values() if v['found']]
    score_items.append((5, "信息区5项内容下方下划线右对齐", passed,
        f"5项全命中:{all_found_ra} 右端全部可推算:{all_resolved} "
        f"右端一致(容差30twip):{right_edges_uniform}(值:{edges_str}) | " + detail))

    # --- +5: 封面页下方信息区字段"院校：""专业：""学号：""姓名：""指导教师："字体左对齐 ---
    # 细则要求的每一个点：
    #   1) 定位对象：封面页(第 1 页)信息区中,标签为"院校""专业""学号""姓名""指导教师"
    #      (含冒号)的 5 个字段所在段落
    #   2) 字体左对齐：办公软件中该 5 个标签视觉上左对齐——即每个标签所在段落
    #      对齐方式为"左对齐",且 5 行标签起始 x 坐标一致(左缩进+首行缩进组合完全一致)
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   a) 段落 w:pPr/w:jc@w:val == "left" 或缺省(Word 默认段落对齐为左对齐,
    #      即"未设置 jc"和"jc=left"在办公软件中视觉一致)
    #   b) 段落 w:pPr/w:ind 的以下 4 个左侧位置属性组合在 5 个标签段落中完全相等,
    #      从而保证 5 行标签起始 x 坐标相同:
    #        w:left, w:leftChars, w:firstLine, w:firstLineChars
    left_align_labels = ['院校', '专业', '学号', '姓名', '指导教师']
    left_check = {}
    for label in left_align_labels:
        left_check[label] = {'found': False, 'jc': None,
                             'left': None, 'leftChars': None,
                             'firstLine': None, 'firstLineChars': None}
        for i in range(info_start, info_end):
            if _label_run(paras[i], label) is None:
                continue
            jc_v = get_paragraph_jc(paras[i])
            pPr_i = paras[i].find(f'{W_NS}pPr')
            ind_i = pPr_i.find(f'{W_NS}ind') if pPr_i is not None else None
            left_check[label] = {
                'found': True,
                'jc': jc_v,
                'left': ind_i.get(f'{W_NS}left') if ind_i is not None else None,
                'leftChars': ind_i.get(f'{W_NS}leftChars') if ind_i is not None else None,
                'firstLine': ind_i.get(f'{W_NS}firstLine') if ind_i is not None else None,
                'firstLineChars': ind_i.get(f'{W_NS}firstLineChars') if ind_i is not None else None,
            }
            break

    all_label_found = all(v['found'] for v in left_check.values())
    # 段落对齐方式：Word 默认对齐即为左对齐,jc 缺省或 jc="left" 均视为左对齐
    all_jc_left = all_label_found and all(
        v['jc'] in (None, 'left', 'start') for v in left_check.values())
    # 左侧起始位置组合必须完全一致(把 None 与 "0" 视为等价——办公软件都渲染为无缩进)
    def _norm_ind(v):
        return '0' if v in (None, '', '0') else v
    left_positions = {
        (_norm_ind(v['left']), _norm_ind(v['leftChars']),
         _norm_ind(v['firstLine']), _norm_ind(v['firstLineChars']))
        for v in left_check.values() if v['found']
    }
    positions_uniform = all_label_found and len(left_positions) == 1
    label_left_ok = all_label_found and all_jc_left and positions_uniform

    detail = " | ".join(
        f"{k}:jc={v['jc']},left={v['left']},leftChars={v['leftChars']},"
        f"firstLine={v['firstLine']},firstLineChars={v['firstLineChars']}"
        for k, v in left_check.items())
    score_items.append((5, "信息区5标签(院校/专业/学号/姓名/指导教师)左对齐", label_left_ok,
        f"5标签全命中:{all_label_found} 全部左对齐:{all_jc_left} "
        f"起始位置一致:{positions_uniform}(集合:{left_positions}) | " + detail))

    # --- +5: 第二或三页顶部出现"毕业设计（论文）开题报告"且下方有表格 ---
    # "顶部"要求该标题是该页第一个有文字的段落
    title_on_top = False
    # --- +5: 文档第二页或第三页顶部出现"毕业设计（论文）开题报告"内容,
    #          且内容下方有表格出现 ---
    # 细则要求的每一个点：
    #   1) 页码范围：第 2 页 或 第 3 页(二者满足其一即可)
    #   2) 位置：该页"顶部"——即在该页文档流中,是该页第一个含有实质文字的元素
    #      (段落或表格里的文字);在此之前不允许出现有文字的段落或有文字的表格
    #   3) 文本内容：包含"毕业设计（论文）开题报告"(兼容全/半角括号)
    #   4) 内容下方有表格：同一页内、该标题元素之后,存在至少一个 w:tbl
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   - 页码来自 get_body_elements_with_pages 按 w:br@type=page 的累积计数
    #   - "顶部"以"该页第一个有文字的元素"为界:遍历 body 子元素,首个满足文字非空的
    #     元素必须是包含标题的段落。若该页第一个有文字的元素是别的段落或有文字的表格,
    #     则视为"标题不在顶部"。
    #   - "下方有表格":同页、标题段之后的元素中出现 w:tbl 即通过。

    def _elem_text(child, tag):
        """返回元素的可见文字(拼接所有 w:t)。表格取其单元格内所有 w:t。"""
        return ''.join(t.text or '' for t in child.findall(f'.//{W_NS}t'))

    def _is_title_text(s):
        norm = s.replace('（', '(').replace('）', ')')
        return '毕业设计' in norm and '开题报告' in norm

    title_on_top = False
    title_page = None
    has_table_after = False

    for pg_check in (2, 3):
        # 收集该页所有元素(按文档流顺序)
        page_children = [(c, tg) for c, tg, pg in elements_pages if pg == pg_check]
        if not page_children:
            continue
        # 找到该页第一个"有文字的"元素
        first_text_idx = None
        for idx, (c, tg) in enumerate(page_children):
            if _elem_text(c, tg).strip():
                first_text_idx = idx
                break
        if first_text_idx is None:
            continue
        first_c, first_tg = page_children[first_text_idx]
        if first_tg != 'p':
            continue  # 顶部第一个有文字的元素是表格,不是标题段
        if not _is_title_text(_elem_text(first_c, 'p')):
            continue
        # 顶部命中标题
        title_on_top = True
        title_page = pg_check
        # 同页、标题段之后是否有 w:tbl
        for c, tg in page_children[first_text_idx + 1:]:
            if tg == 'tbl':
                has_table_after = True
                break
        if title_on_top and has_table_after:
            break  # 两个条件都满足即通过,无需再看下一个候选页

    passed = title_on_top and has_table_after
    score_items.append((5, "第2或3页顶部出现\"毕业设计（论文）开题报告\"且下方有表格", passed,
        f"顶部命中标题:{title_on_top} 命中页:{title_page} 下方有表格:{has_table_after}"))

    # --- +3: 从第三页开始所有正文内容都以表格内文字形式出现 ---
    # 细则要求的每一个点：
    #   1) 页码范围：第 3 页及之后的所有页
    #   2) "所有正文内容"：该范围内文档流中的可见文字
    #   3) 载体限定：这些文字必须出现在表格(w:tbl)内(即表格单元格中的段落),
    #      而不能作为表格外的独立段落存在
    # 判定方式(对齐 Word/WPS 实际渲染)：
    #   - 页码来自 elements_pages(按 w:br@type=page 累加)
    #   - 遍历 body 顶层子元素:
    #       * tag=='tbl' 元素本身即"表格内文字",视为合规
    #       * tag=='p' 且属于 page>=3 的段落:若该段有可见文字则违规
    #         (Word/WPS 中表格单元格内的段落不会出现在 body 顶层,而是嵌套在
    #          w:tbl/w:tr/w:tc 之下,因此顶层 w:p 都是"表格外的段落")
    #   - 允许存在空段落(纯排版占位、不显示任何文字),办公软件里视觉上不构成"正文内容"
    outside_text_page3plus = []  # (page, snippet)
    has_any_table_from_page3 = False
    for child, tag, page in elements_pages:
        if page < 3:
            continue
        if tag == 'tbl':
            has_any_table_from_page3 = True
            continue
        if tag == 'p':
            p_text = ''.join(t.text or '' for t in child.findall(f'.//{W_NS}t'))
            if p_text.strip():
                # 顶层段落且含可见文字 → 表格外的正文
                snippet = p_text.strip()
                if len(snippet) > 30:
                    snippet = snippet[:30] + '...'
                outside_text_page3plus.append((page, snippet))
    no_outside_text = (len(outside_text_page3plus) == 0)
    passed = has_any_table_from_page3 and no_outside_text
    score_items.append((3, "第三页起所有正文以表格内文字形式出现", passed,
        f"第3页起存在表格:{has_any_table_from_page3} "
        f"表格外正文段落数:{len(outside_text_page3plus)} "
        f"违规明细:{outside_text_page3plus[:5]}"))

    # --- +3: 表格外框线：上下0.5磅双实线，左右0.5磅单实线 ---
    # --- +3: 表格整体带有外框线,上下框线 0.5 磅双实线,左右框线 0.5 磅单实线 ---
    # 细则要求的每一个点：
    #   1) 范围：文档中的所有表格(每一张 w:tbl)
    #   2) 表格"整体带有外框线":上、下、左、右四条外框线全部存在(非 none/nil,即办公软件中可见)
    #   3) 上框线:样式 = 双实线(w:val="double"),粗细 = 0.5 磅
    #   4) 下框线:样式 = 双实线(w:val="double"),粗细 = 0.5 磅
    #   5) 左框线:样式 = 单实线(w:val="single"),粗细 = 0.5 磅
    #   6) 右框线:样式 = 单实线(w:val="single"),粗细 = 0.5 磅
    # 换算说明(OOXML/办公软件):
    #   w:sz 的单位是 1/8 磅 → 0.5 磅 = w:sz="4"
    #   Word/WPS「边框和底纹」对话框中,"边框宽度 0.5 磅"保存的即是 w:sz="4"
    border_check = []  # 每张表 (top_ok, bottom_ok, left_ok, right_ok, 明细)
    for ti, t in enumerate(doc.tables):
        tbl = t._tbl
        # 表格边框定义位置:w:tblPr/w:tblBorders(表格属性中的外框设置)
        tblPr = tbl.find(f'{W_NS}tblPr')
        borders = tblPr.find(f'{W_NS}tblBorders') if tblPr is not None else None

        def _check(side_elem, want_val, want_sz='4'):
            if side_elem is None:
                return False, "缺失"
            v = side_elem.get(f'{W_NS}val')
            s = side_elem.get(f'{W_NS}sz')
            # 存在性:val 不为 none/nil(否则办公软件中不显示边框)
            if v in (None, 'none', 'nil'):
                return False, f"val={v}(不可见)"
            ok = (v == want_val) and (s == want_sz)
            return ok, f"val={v}(需{want_val}),sz={s}(需{want_sz}=0.5磅)"

        if borders is None:
            border_check.append((False, False, False, False, f"表{ti}:未定义 tblBorders"))
            continue
        top_ok, top_d = _check(borders.find(f'{W_NS}top'), 'double', '4')
        bot_ok, bot_d = _check(borders.find(f'{W_NS}bottom'), 'double', '4')
        left_ok, left_d = _check(borders.find(f'{W_NS}left'), 'single', '4')
        right_ok, right_d = _check(borders.find(f'{W_NS}right'), 'single', '4')
        border_check.append((top_ok, bot_ok, left_ok, right_ok,
            f"表{ti}: 上[{top_d}] 下[{bot_d}] 左[{left_d}] 右[{right_d}]"))

    all_top = all(x[0] for x in border_check)
    all_bot = all(x[1] for x in border_check)
    all_left = all(x[2] for x in border_check)
    all_right = all(x[3] for x in border_check)
    passed = bool(border_check) and all_top and all_bot and all_left and all_right
    score_items.append((3, "表格上下0.5磅双实线+左右0.5磅单实线外框", passed,
        f"表数量:{len(border_check)} 全部上OK:{all_top} 全部下OK:{all_bot} "
        f"全部左OK:{all_left} 全部右OK:{all_right} | " +
        " || ".join(x[4] for x in border_check[:5])))

    # --- +5: 表格内标题字体宋体小四加粗 ---
    # 小四=24半磅(12pt)
    # --- +5: 表格内文字标题(8 项)字体格式为宋体小四加粗 ---
    # 细则列出的 8 个标题(严格按字面):
    #   "一．研究目的及意义" / "二．行业研究现状" / "指导教师意见" /
    #   "一、优点" / "二、缺点" /
    #   "三．本课题要研究或解决的问题和拟采用的研究手段（途径）" /
    #   "四．工作进度安排" / "五．主要参考文献"
    # 细则要求的每一个点:
    #   1) 每个标题都要在表格单元格内存在
    #   2) 字体:宋体 (w:rFonts@w:eastAsia ∈ {"SimSun","宋体"})
    #   3) 字号:小四 = 12pt = 24 半磅 (w:sz@w:val == "24")
    #   4) 加粗:w:b 存在且未被显式关闭
    # 判定方式(对齐 Word/WPS 实际渲染):
    #   - 遍历所有表格的单元格段落,归一化标点(中文点号"．" / 顿号"、"皆可,
    #     全/半角括号皆可,与办公软件输入法差异一致)后按字面查找
    #   - 命中的标题所属 run(可能跨多个 run)必须全部满足字体/字号/加粗
    #   - 中文字符走 w:eastAsia,与办公软件的字体解析路径一致
    title_titles_spec = [
        '一．研究目的及意义', '二．行业研究现状', '指导教师意见',
        '一、优点', '二、缺点',
        '三．本课题要研究或解决的问题和拟采用的研究手段（途径）',
        '四．工作进度安排', '五．主要参考文献',
    ]

    def _norm_punct(s):
        return (s.replace('．', '.').replace('。', '.').replace('、', ',')
                 .replace('（', '(').replace('）', ')')
                 .replace(' ', '').replace('　', ''))

    title_check = {}
    for spec in title_titles_spec:
        title_check[spec] = {'found': False, 'runs': [], 'font_ok': False,
                             'size_ok': False, 'bold_ok': False}
    spec_norms = {spec: _norm_punct(spec) for spec in title_titles_spec}

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_elem = p._element
                    # 用 char→run 映射跨 run 命中
                    runs = p_elem.findall(f'{W_NS}r')
                    run_texts = [''.join(t.text or '' for t in r.findall(f'{W_NS}t'))
                                 for r in runs]
                    concat = ''.join(run_texts)
                    concat_norm = _norm_punct(concat)
                    for spec, spec_norm in spec_norms.items():
                        if title_check[spec]['found']:
                            continue
                        if spec_norm not in concat_norm:
                            continue
                        # 定位标题在原始 concat 中的字符区间(用归一化字符串定位后回映到原串
                        # 需简化:因归一化只做"字符替换"(不新增/删减字符),归一化前后长度一致,
                        # 因此归一化后的下标可直接用于原串。
                        idx = concat_norm.find(spec_norm)
                        end = idx + len(spec_norm)
                        # 找覆盖 [idx,end) 的 run
                        covered = []
                        cur = 0
                        for ri, rt in enumerate(run_texts):
                            r_end = cur + len(rt)
                            if r_end > idx and cur < end:
                                covered.append(runs[ri])
                            cur = r_end
                        run_details = []
                        font_all_ok = bool(covered)
                        size_all_ok = bool(covered)
                        bold_all_ok = bool(covered)
                        for r in covered:
                            ea = _run_eastasia_effective(r, p_elem)
                            sz = _run_sz_effective(r, p_elem)
                            b = is_bold(r)
                            run_details.append((ea, sz, b))
                            if ea not in ('SimSun', '宋体'):
                                font_all_ok = False
                            if sz != 24:
                                size_all_ok = False
                            if not b:
                                bold_all_ok = False
                        title_check[spec] = {
                            'found': True,
                            'runs': run_details,
                            'font_ok': font_all_ok,
                            'size_ok': size_all_ok,
                            'bold_ok': bold_all_ok,
                        }

    all_found_t = all(v['found'] for v in title_check.values())
    all_font_t = all_found_t and all(v['font_ok'] for v in title_check.values())
    all_size_t = all_found_t and all(v['size_ok'] for v in title_check.values())
    all_bold_t = all_found_t and all(v['bold_ok'] for v in title_check.values())
    passed = all_found_t and all_font_t and all_size_t and all_bold_t
    detail = " | ".join(
        f"{spec[:16]}...:found={v['found']},字体宋体:{v['font_ok']},"
        f"小四:{v['size_ok']},加粗:{v['bold_ok']},runs={v['runs']}"
        for spec, v in title_check.items())
    score_items.append((5, "表格内8项标题宋体小四加粗", passed,
        f"8项全命中:{all_found_t} 全宋体:{all_font_t} 全小四:{all_size_t} "
        f"全加粗:{all_bold_t} | " + detail))

    # --- +5: 表格内除标题、表注外其余正文字体格式为宋体小四,
    #          英文或阿拉伯数字字体格式为 Times New Roman 小四 ---
    # 细则要求的每一个点:
    #   1) 范围:所有表格(w:tbl)的单元格内段落
    #   2) 排除对象:
    #      a) 标题(前一采分点列出的 8 个标题,以及其它同级标题写法)
    #      b) 表注:办公软件中"表注"通常为形如"表1 …""表 1-1 …"的表格说明段落,
    #         或使用 Caption 样式(w:pStyle 引用 caption/表注 相关样式)。
    #   3) 正文中的中文字符:字体 = 宋体(w:rFonts@w:eastAsia ∈ {"SimSun","宋体"}),
    #      字号 = 小四(w:sz@w:val = "24"; 12pt = 24 半磅)
    #   4) 正文中的英文字母/阿拉伯数字:字体 = Times New Roman
    #      (w:rFonts@w:ascii == "Times New Roman";兼容 w:hAnsi 同名),
    #      字号 = 小四(w:sz@w:val = "24")
    # 判定方式(对齐 Word/WPS 实际渲染):
    #   - 逐 run 判断:根据 run 文本中包含的字符种类,分别校验 eastAsia / ascii;
    #     只含空白/标点(既非中文也非英文数字)的 run 只校验字号(不涉及字体家族)。
    #   - Word 属性继承:run rPr > 段落 pPr rPr,若 run 未直接写,回退段落级默认。
    import re
    title_specs_for_exclude = [
        '一．研究目的及意义', '二．行业研究现状', '指导教师意见',
        '一、优点', '二、缺点',
        '三．本课题要研究或解决的问题和拟采用的研究手段（途径）',
        '四．工作进度安排', '五．主要参考文献',
        '研究内容与技术方案', '主要参考资料',
        '教研室审查意见', '所在二级学院审查意见',
    ]
    title_norms_for_exclude = {_norm_punct(s) for s in title_specs_for_exclude}

    def _is_caption_para(p_elem, p_text):
        # 样式引用:pStyle 名含 caption/表注/题注
        pPr = p_elem.find(f'{W_NS}pPr')
        if pPr is not None:
            pStyle = pPr.find(f'{W_NS}pStyle')
            if pStyle is not None:
                v = (pStyle.get(f'{W_NS}val') or '').lower()
                if 'caption' in v or '题注' in v or '表注' in v:
                    return True
        # 文字形态:以"表 N""表N-M""表 N.M"等开头的段落,是办公软件中常见的表注/题注写法
        stripped = p_text.strip()
        if re.match(r'^表\s*\d+([-\.．]\d+)?\s*[　 ]*.*', stripped):
            return True
        return False

    def _is_title_para(p_text):
        norm = _norm_punct(p_text.strip())
        # 完全匹配或以标题为整段(容许标题字面之外无其它文字)
        for tn in title_norms_for_exclude:
            if norm == tn:
                return True
        return False

    CHN_RE = re.compile(r'[一-鿿]')
    ENG_RE = re.compile(r'[A-Za-z0-9]')

    def _ascii_font_effective(r_elem, p_elem):
        rPr = r_elem.find(f'{W_NS}rPr')
        if rPr is not None:
            rFonts = rPr.find(f'{W_NS}rFonts')
            if rFonts is not None:
                v = rFonts.get(f'{W_NS}ascii') or rFonts.get(f'{W_NS}hAnsi')
                if v:
                    return v
        pPr = p_elem.find(f'{W_NS}pPr')
        if pPr is not None:
            pRpr = pPr.find(f'{W_NS}rPr')
            if pRpr is not None:
                rFonts = pRpr.find(f'{W_NS}rFonts')
                if rFonts is not None:
                    return rFonts.get(f'{W_NS}ascii') or rFonts.get(f'{W_NS}hAnsi')
        return None

    body_check = {'checked': 0, 'violations': []}
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_elem = p._element
                    p_text = p.text
                    stripped = p_text.strip()
                    if not stripped:
                        continue
                    if _is_title_para(stripped):
                        continue
                    if _is_caption_para(p_elem, p_text):
                        continue
                    for r in p_elem.findall(f'{W_NS}r'):
                        r_text = ''.join(t.text or '' for t in r.findall(f'{W_NS}t'))
                        if not r_text.strip():
                            continue
                        sz = _run_sz_effective(r, p_elem)
                        ea = _run_eastasia_effective(r, p_elem)
                        asc = _ascii_font_effective(r, p_elem)
                        has_chn = bool(CHN_RE.search(r_text))
                        has_eng = bool(ENG_RE.search(r_text))
                        body_check['checked'] += 1
                        problems = []
                        # 字号必须 = 24 半磅(小四),对含中文或英文/数字的 run 都要求
                        if (has_chn or has_eng) and sz != 24:
                            problems.append(f"sz={sz}(需24)")
                        # 中文字符:eastAsia 必须为宋体
                        if has_chn and ea not in ('SimSun', '宋体'):
                            problems.append(f"eastAsia={ea}(需宋体)")
                        # 英文/数字:ascii/hAnsi 必须为 Times New Roman
                        if has_eng and asc != 'Times New Roman':
                            problems.append(f"ascii={asc}(需TNR)")
                        if problems:
                            body_check['violations'].append(
                                (r_text[:20], problems))

    passed = (body_check['checked'] > 0 and not body_check['violations'])
    score_items.append((5, "表格内正文中文宋体小四+英数TNR小四(除标题/表注)", passed,
        f"检查run数:{body_check['checked']} 违规数:{len(body_check['violations'])} "
        f"前5违规:{body_check['violations'][:5]}"))

    # --- +5: 表格内文本段落行距为固定值约22磅、两端对齐、首行缩进两字符 ---
    # 细则每一个点:
    #   1) 范围:表格内的"文本段落"(有文字内容的非空段落)
    #   2) 行距 = 固定值约 22 磅
    #      对应办公软件属性: w:spacing@w:lineRule="exact" 且 w:line ≈ 440 (22pt*20=440 twip)
    #      "约"给出容差,按 ±20 twip(约 ±1pt)判定
    #   3) 段落对齐方式 = 两端对齐
    #      对应办公软件属性: w:jc@w:val="both"
    #   4) 首行缩进 = 两字符
    #      对应办公软件属性:
    #        优先字符单位: w:ind@w:firstLineChars="200"(1字符 = 100)
    #        兼容 twip 单位: w:ind@w:firstLine ≈ 2*字号pt*20 twip(小四12pt → 480 twip),容差 ±60
    #      Word/WPS 打开时按字符单位优先渲染,与字号联动;两种写法都是有效的"首行缩进 2 字符"。
    body_para = {'checked': 0, 'line_bad': [], 'jc_bad': [], 'ind_bad': []}
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_elem = p._element
                    p_text = p.text
                    if not p_text.strip():
                        continue
                    body_para['checked'] += 1
                    # 行距
                    sp = get_paragraph_spacing(p_elem)
                    line_rule = sp.get('lineRule', '')
                    try:
                        line_val = int(sp.get('line', '0') or '0')
                    except ValueError:
                        line_val = 0
                    if not (line_rule == 'exact' and abs(line_val - 440) <= 20):
                        body_para['line_bad'].append(
                            (p_text[:20], f"lineRule={line_rule},line={line_val}"))
                    # 两端对齐
                    if get_paragraph_jc(p_elem) != 'both':
                        body_para['jc_bad'].append(
                            (p_text[:20], f"jc={get_paragraph_jc(p_elem)}"))
                    # 首行缩进 2 字符
                    pPr = p_elem.find(f'{W_NS}pPr')
                    ind = pPr.find(f'{W_NS}ind') if pPr is not None else None
                    fc_raw = ind.get(f'{W_NS}firstLineChars') if ind is not None else None
                    fl_raw = ind.get(f'{W_NS}firstLine') if ind is not None else None
                    fc = int(fc_raw) if (fc_raw and fc_raw.lstrip('-').isdigit()) else None
                    fl = int(fl_raw) if (fl_raw and fl_raw.lstrip('-').isdigit()) else None
                    # 尝试用该段落有效字号(小四=12pt=24半磅)换算 twip 值
                    est_twip = None
                    for r in p_elem.findall(f'{W_NS}r'):
                        sz_hp = _run_sz_effective(r, p_elem)
                        if sz_hp:
                            est_twip = int(sz_hp / 2 * 20 * 2)  # sz_hp半磅 → pt → twip *2字符
                            break
                    if est_twip is None:
                        est_twip = 480  # 默认小四 2 字符
                    ok_char = (fc is not None and abs(fc - 200) <= 20)
                    ok_twip = (fl is not None and abs(fl - est_twip) <= 60)
                    if not (ok_char or ok_twip):
                        body_para['ind_bad'].append(
                            (p_text[:20], f"firstLineChars={fc},firstLine={fl}"))
    passed = (body_para['checked'] > 0
              and not body_para['line_bad']
              and not body_para['jc_bad']
              and not body_para['ind_bad'])
    score_items.append((5, "表格内文本段落固定22磅+两端对齐+首行缩进2字符", passed,
        f"检查段数:{body_para['checked']} "
        f"行距不合:{len(body_para['line_bad'])} "
        f"对齐不合:{len(body_para['jc_bad'])} "
        f"缩进不合:{len(body_para['ind_bad'])} "
        f"示例行距不合:{body_para['line_bad'][:2]} "
        f"示例缩进不合:{body_para['ind_bad'][:2]}"))

    # --- +1: 表格列宽为 15.00-15.90 厘米 ---
    # 细则每一个点:
    #   1) 对象: 表格的"列宽"(不是表宽,也不是行高)
    #   2) 范围: 15.00 ~ 15.90 厘米(闭区间)
    # 办公软件(Word/WPS)渲染依据:
    #   - 主属性: w:tbl / w:tblGrid / w:gridCol@w:w (twip 单位)
    #     这是 Word 表格布局算法在自动/固定布局下首选的列宽来源。
    #   - 单元格覆盖: w:tc / w:tcPr / w:tcW@w:w (仅当 @w:type="dxa"),
    #     若存在则会覆盖 gridCol,以最终渲染宽度为准。
    #   逐表逐列判断,每一列都必须落在 15.00~15.90 cm 内。
    col_width_ok = True
    col_width_report = []
    tables_all = doc.tables
    for ti, t in enumerate(tables_all):
        tbl = t._tbl
        grid = tbl.find(f'{W_NS}tblGrid')
        if grid is None:
            col_width_ok = False
            col_width_report.append((ti, 'no-tblGrid'))
            continue
        gridCols = grid.findall(f'{W_NS}gridCol')
        # 每列的默认宽度(来自 tblGrid)
        default_widths = [int(c.get(f'{W_NS}w', '0')) for c in gridCols]
        # 单元格 tcW 覆盖(取每列最常见的 tcW,若无则用默认)
        col_final = list(default_widths)
        rows = tbl.findall(f'{W_NS}tr')
        if rows:
            # 用第一行的 tc 顺序对齐列(处理合并较少的单列/规则表格已足够)
            for ri, tr in enumerate(rows):
                tcs = tr.findall(f'{W_NS}tc')
                for ci, tc in enumerate(tcs):
                    if ci >= len(col_final):
                        break
                    tcPr = tc.find(f'{W_NS}tcPr')
                    if tcPr is None:
                        continue
                    tcW = tcPr.find(f'{W_NS}tcW')
                    if tcW is None:
                        continue
                    if tcW.get(f'{W_NS}type') != 'dxa':
                        continue
                    try:
                        w_val = int(tcW.get(f'{W_NS}w', '0'))
                    except ValueError:
                        continue
                    # 第一行覆盖,后续行如与首行相等则忽略,不等则并入报告
                    if ri == 0:
                        col_final[ci] = w_val
        # 判定
        for ci, w_twip in enumerate(col_final):
            cm = w_twip / TWIP_PER_CM
            col_width_report.append((ti, ci, round(cm, 3)))
            if not (15.00 <= cm <= 15.90):
                col_width_ok = False
    score_items.append((1, "所有表格每一列列宽 15.00-15.90 cm", col_width_ok,
        f"各(表,列)宽(cm): {col_width_report}"))

    # --- +3: "指导教师意见"、"教研室审查意见"、"所在二级学院审查意见"内容都在同一表格内且表格未断页 ---
    # 细则每一个点:
    #   1) 三段字面: "指导教师意见""教研室审查意见""所在二级学院审查意见"
    #   2) 都在"同一个表格"内(必须是同一张 w:tbl,不同表格即使相邻也不满足)
    #   3) 该表格"未断页" —— 在办公软件里,"断页"包含两种情形:
    #        a) 表格内部出现强制分页(单元格段落中的 <w:br w:type="page"/>)
    #        b) Word/WPS 的自动分页 —— 当表格纵向内容超过当前页面剩余可用
    #           高度时,渲染引擎会把表格从某一行分成两页显示。
    # 判定方式(对齐 Word/WPS 实际渲染):
    #   按 body 顺序模拟渲染分页,累计每段/每行占据的 twip 高度,
    #   与"页面可用高度 = pgSz.h - pgMar.top - pgMar.bottom"比较,
    #   在需要换页时递增页号(渲染页码)。分别记录目标表格开始渲染时所在页
    #   (tbl_start_page)和结束渲染时所在页(tbl_end_page),
    #   要求 tbl_start_page == tbl_end_page 且表格内部无强制分页。
    target_labels = ['指导教师意见', '教研室审查意见', '所在二级学院审查意见']
    same_table = None
    for ti, t in enumerate(doc.tables):
        text_all = ''.join(cell.text for row in t.rows for cell in row.cells)
        if all(lbl in text_all for lbl in target_labels):
            same_table = t
            break
    all_same = same_table is not None

    # 表格内无强制分页
    no_inner_pb = True
    if same_table is not None:
        for br in same_table._tbl.findall(f'.//{W_NS}br'):
            if br.get(f'{W_NS}type') == 'page':
                no_inner_pb = False
                break

    # 渲染分页模拟:计算目标表格的起止渲染页码
    _bottom_m_twip = int(pgMar.get(f'{W_NS}bottom', '1440')) if pgMar is not None else 1440
    _usable_twip = max(1, page_h_twip - top_margin_twip - _bottom_m_twip)

    def _tr_declared_h(tr_elem):
        trPr_ = tr_elem.find(f'{W_NS}trPr')
        trH_ = trPr_.find(f'{W_NS}trHeight') if trPr_ is not None else None
        if trH_ is None:
            return None, None
        try:
            v = int(trH_.get(f'{W_NS}val', '0'))
        except ValueError:
            v = 0
        return v, trH_.get(f'{W_NS}hRule')

    def _tc_content_h(tc_elem):
        return sum(int(_para_height_twip(pp)) for pp in tc_elem.findall(f'{W_NS}p'))

    def _tr_visual_h(tr_elem):
        declared, hRule = _tr_declared_h(tr_elem)
        max_cell = 0
        for tc in tr_elem.findall(f'{W_NS}tc'):
            max_cell = max(max_cell, _tc_content_h(tc))
        if declared is None:
            return max_cell if max_cell else 240
        if hRule == 'exact':
            return declared
        return max(declared, max_cell)  # atLeast 或缺省

    cur_page_sim = 1
    y_cursor = 0
    tbl_start_page = None
    tbl_end_page = None
    tgt_tbl_elem = same_table._tbl if same_table is not None else None
    for child in body:
        tag_ = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag_ == 'p':
            # 段前分页(pPr/pageBreakBefore)与段内强制分页(<w:br w:type="page"/>)
            pPr_ = child.find(f'{W_NS}pPr')
            has_pbb = pPr_ is not None and pPr_.find(f'{W_NS}pageBreakBefore') is not None
            has_pb = any(b.get(f'{W_NS}type') == 'page'
                         for b in child.findall(f'.//{W_NS}br'))
            if (has_pbb or has_pb) and y_cursor > 0:
                cur_page_sim += 1
                y_cursor = 0
            p_h = int(_para_height_twip(child))
            if y_cursor + p_h > _usable_twip and y_cursor > 0:
                cur_page_sim += 1
                y_cursor = 0
            y_cursor += p_h
        elif tag_ == 'tbl':
            is_target = tgt_tbl_elem is not None and child is tgt_tbl_elem
            if is_target:
                tbl_start_page = cur_page_sim
            for tr in child.findall(f'{W_NS}tr'):
                row_h = int(_tr_visual_h(tr))
                if y_cursor + row_h > _usable_twip and y_cursor > 0:
                    cur_page_sim += 1
                    y_cursor = 0
                y_cursor += row_h
            if is_target:
                tbl_end_page = cur_page_sim

    if same_table is None:
        no_cross_page = False
    elif tbl_start_page is None or tbl_end_page is None:
        no_cross_page = False
    else:
        no_cross_page = (tbl_start_page == tbl_end_page)

    passed = all_same and no_inner_pb and no_cross_page
    score_items.append((3, "三段审查意见同表且表格未断页", passed,
        (f"三段同表:{all_same} 表内无强制分页:{no_inner_pb} "
         f"渲染分页(起始页={tbl_start_page},结束页={tbl_end_page},"
         f"版心高={_usable_twip}twip):未断页={no_cross_page}")))

    # 供下方 +3 行高检查复用的同表引用(兼容旧代码 t3)
    t3 = same_table if same_table is not None else doc.tables[3]

    # --- +3: "指导教师意见"所在表格单元格行高为 16.40-16.50 厘米 ---
    # 细则每一个点:
    #   1) 定位: 含"指导教师意见"文字的那一行(单元格所在的 w:tr)
    #   2) 属性: 该行的行高
    #   3) 范围: 16.40 ~ 16.50 厘米(闭区间)
    # 办公软件(Word/WPS)行高属性:
    #   w:tr / w:trPr / w:trHeight@w:val (twip 单位)
    #   @w:hRule 语义: "atLeast"=最小值, "exact"=固定值, 默认=auto;
    #   细则只说"行高为 X 厘米",不额外约束 hRule。
    row_h_cm = None
    row_h_rule = None
    found = False
    if same_table is not None:
        for row in same_table.rows:
            row_text = ''.join(cell.text for cell in row.cells)
            if '指导教师意见' in row_text:
                tr = row._tr
                trPr = tr.find(f'{W_NS}trPr')
                trHeight = trPr.find(f'{W_NS}trHeight') if trPr is not None else None
                if trHeight is not None:
                    try:
                        row_h_twip = int(trHeight.get(f'{W_NS}val', '0'))
                    except ValueError:
                        row_h_twip = 0
                    row_h_cm = row_h_twip / TWIP_PER_CM
                    row_h_rule = trHeight.get(f'{W_NS}hRule')
                found = True
                break
    passed = (found and row_h_cm is not None
              and 16.40 <= row_h_cm <= 16.50)
    score_items.append((3, "\"指导教师意见\"所在行高 16.40-16.50 cm", passed,
        f"定位到行:{found} 实际:{row_h_cm and round(row_h_cm,3)}cm hRule:{row_h_rule}"))

    # --- +3: "教研室审查意见"所在表格单元格行高为 2.70-2.90 厘米 ---
    # 细则每一个点:
    #   1) 定位: 含"教研室审查意见"文字的那一行(单元格所在的 w:tr)
    #   2) 属性: 该行的行高
    #   3) 范围: 2.70 ~ 2.90 厘米(闭区间)
    # 办公软件属性: w:tr / w:trPr / w:trHeight@w:val (twip),
    # @w:hRule 只作为报告信息,细则未额外约束。
    row2_h_cm = None
    row2_h_rule = None
    found2 = False
    if same_table is not None:
        for row in same_table.rows:
            row_text = ''.join(cell.text for cell in row.cells)
            if '教研室审查意见' in row_text:
                tr = row._tr
                trPr = tr.find(f'{W_NS}trPr')
                trHeight = trPr.find(f'{W_NS}trHeight') if trPr is not None else None
                if trHeight is not None:
                    try:
                        row_h_twip = int(trHeight.get(f'{W_NS}val', '0'))
                    except ValueError:
                        row_h_twip = 0
                    row2_h_cm = row_h_twip / TWIP_PER_CM
                    row2_h_rule = trHeight.get(f'{W_NS}hRule')
                found2 = True
                break
    passed = (found2 and row2_h_cm is not None
              and 2.70 <= row2_h_cm <= 2.90)
    score_items.append((3, "\"教研室审查意见\"所在行高 2.70-2.90 cm", passed,
        f"定位到行:{found2} 实际:{row2_h_cm and round(row2_h_cm,3)}cm hRule:{row2_h_rule}"))

    # --- +5: "五．主要参考资料"内容单独放在一个表格,且表格所在位置在文档第六页 ---
    # 细则每一个点:
    #   1) 存在包含"五．主要参考资料"内容的表格(容错"资料/文献"及全角/半角标点)
    #   2) 该内容"单独放在一个表格" —— 即:该文本仅出现在这一张 w:tbl 内,
    #      不与其它采分点所需的表格标题(如"指导教师意见""研究目的及意义"等)共处一表。
    #   3) 该表格所在位置为文档"第六页"。
    # 办公软件页码判定:
    #   Word/WPS 中"第几页"依赖分页,OOXML 层只能通过强制分页符(w:br@type="page")累加,
    #   elements_pages 已按此累加规则给出每个 body 子元素的起始页码。
    def _norm_ref_title(s):
        return _norm_punct((s or '').strip())
    target_norms = {_norm_ref_title('五．主要参考资料'), _norm_ref_title('五．主要参考文献')}

    # 找到承载该标题的表格
    ref_tbl_idx = None
    for ti, t in enumerate(doc.tables):
        text_all = ''.join(cell.text for row in t.rows for cell in row.cells)
        norm_all = _norm_ref_title(text_all)
        if any(tn in norm_all for tn in target_norms):
            ref_tbl_idx = ti
            break
    has_ref = ref_tbl_idx is not None

    # 判定"单独":只在这一张表出现,其它表格不出现该标题
    solo = False
    if has_ref:
        solo = True
        for ti, t in enumerate(doc.tables):
            if ti == ref_tbl_idx:
                continue
            text_all = ''.join(cell.text for row in t.rows for cell in row.cells)
            norm_all = _norm_ref_title(text_all)
            if any(tn in norm_all for tn in target_norms):
                solo = False
                break
        # 同时要求这一张表除该标题外不夹带其它采分点的表标题(即真正的"单独一张表")
        # 这里以其它已知同级表标题为对照,若同表出现则视为非单独
        other_titles = {
            _norm_ref_title(x) for x in [
                '指导教师意见', '教研室审查意见', '所在二级学院审查意见',
                '一．研究目的及意义', '二．行业研究现状',
                '三．本课题要研究或解决的问题和拟采用的研究手段（途径）',
                '四．工作进度安排',
            ]
        }
        this_text = _norm_ref_title(''.join(
            cell.text for row in doc.tables[ref_tbl_idx].rows for cell in row.cells))
        if any(ot in this_text for ot in other_titles):
            solo = False

    # 页码定位
    ref_page = None
    if has_ref:
        target_tbl_elem = doc.tables[ref_tbl_idx]._tbl
        for child, tag, page in elements_pages:
            if tag == 'tbl' and child is target_tbl_elem:
                ref_page = page
                break

    passed = has_ref and solo and (ref_page == 6)
    score_items.append((5, "\"五．主要参考资料\"独占一表且在第六页", passed,
        f"命中表索引:{ref_tbl_idx} 独占:{solo} 所在页:{ref_page}"))

    # --- +1: "三、研究内容与技术方案"下方没有出现图片 ---
    # 细则每一个点:
    #   1) 定位标题字面: "三、研究内容与技术方案"(容错 三. / 三． / 三、,以及可能的
    #      "本课题要研究…"扩展写法。以归一化标点后包含"研究内容与技术方案"为准。)
    #   2) "下方" —— 标题所在段落之后,直到下一个同级标题("四．工作进度安排")之前的所有段落。
    #      这一段落集合即办公软件视觉上"三"节的正文。
    #   3) "没有出现图片" —— 办公软件识别的图片对象:
    #      w:drawing (DrawingML,包括图片/形状/文本框中的图片),
    #      w:pict / v:imagedata (VML 图片),
    #      w:object (嵌入对象,包含内嵌图片时同样呈现为图片)。
    def _para_has_image(p_elem):
        if p_elem.findall(f'.//{W_NS}drawing'):
            return True
        if p_elem.findall(f'.//{W_NS}pict'):
            return True
        if p_elem.findall(f'.//{W_NS}object'):
            return True
        # v:imagedata 命名空间
        for e in p_elem.iter():
            if e.tag.endswith('}imagedata') or e.tag == 'imagedata':
                return True
        return False

    def _norm_title(s):
        return _norm_punct((s or '').strip())

    target_key = '研究内容与技术方案'
    next_section_keys = ['工作进度安排', '主要参考文献', '主要参考资料',
                         '指导教师意见']

    # 顺序遍历所有表格所有单元格所有段落,建立一个平铺段落列表
    flat_paras = []
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    flat_paras.append(p)

    title_idx = None
    for i, p in enumerate(flat_paras):
        if target_key in _norm_title(p.text):
            title_idx = i
            break

    end_idx = len(flat_paras)
    if title_idx is not None:
        for j in range(title_idx + 1, len(flat_paras)):
            nt = _norm_title(flat_paras[j].text)
            if any(k in nt for k in next_section_keys):
                end_idx = j
                break

    has_img_below = False
    img_where = None
    if title_idx is not None:
        for j in range(title_idx + 1, end_idx):
            if _para_has_image(flat_paras[j]._element):
                has_img_below = True
                img_where = j
                break

    # 若定位不到标题,严格按细则也应视为不满足(下方内容无从检查)
    passed = (title_idx is not None) and (not has_img_below)
    score_items.append((1, "\"三、研究内容与技术方案\"下方没有出现图片", passed,
        f"定位标题段:{title_idx} 节段范围:({title_idx},{end_idx}) "
        f"下方存在图片:{has_img_below} 位置:{img_where}"))

    # --- -5: 文档中出现超过页面百分之 50 的大段空白 ---
    # 细则每一个点:
    #   1) 出现"大段空白" —— 办公软件视觉上呈现的连续空白区域
    #   2) 该空白占页面比例 > 50%
    #   3) 触发一次即扣分
    # 办公软件视觉空白的三个来源(必须都算):
    #   a) 空白段落: 无文字/图片/形状的段落,其占据 line+before+after 高度即为空白
    #   b) 表格单元格内的下方空白: 行声明高度(w:trHeight@val,尤其 hRule=atLeast/exact)
    #      大于该行单元格实际内容高度时,差值即为行内下方空白
    #   c) 页面末尾空白: 一逻辑页(以 w:br@type="page" 分割)的可见内容总高度小于
    #      页面可用高度(w:pgSz@h - w:pgMar@top - w:pgMar@bottom)时,差值为页底空白
    # 判定: 对每一非封面逻辑页, 计算"空白占比 = (a+b+c) / 页可用高度",
    #       任一非封面页 > 0.5 即触发扣分。
    bottom_margin_twip = int(pgMar.get(f'{W_NS}bottom', '1440')) if pgMar is not None else 1440
    page_usable_twip = max(1, page_h_twip - top_margin_twip - bottom_margin_twip)

    def _cell_content_height(tc_elem):
        """单元格内容实际高度: 累加内部段落的 _para_height_twip(含图片)。"""
        h = 0
        for p in tc_elem.findall(f'{W_NS}p'):
            h += int(_para_height_twip(p))
        return h

    def _row_declared_height(tr_elem):
        trPr = tr_elem.find(f'{W_NS}trPr')
        trHeight = trPr.find(f'{W_NS}trHeight') if trPr is not None else None
        if trHeight is None:
            return None, None
        try:
            val = int(trHeight.get(f'{W_NS}val', '0'))
        except ValueError:
            val = 0
        return val, trHeight.get(f'{W_NS}hRule')

    def _para_is_blank(p_elem):
        txt = ''.join((t.text or '') for t in p_elem.findall(f'.//{W_NS}t'))
        if txt.strip():
            return False
        for tag_suffix in ('drawing', 'pict', 'object'):
            if p_elem.findall(f'.//{W_NS}{tag_suffix}'):
                return False
        return True

    # 按 body 顺序切页,累计 (a) 空白段落高度、(b) 行内空白、(c) 内容高度
    per_page = {1: {'blank_para': 0, 'row_gap': 0, 'content': 0}}
    cur_page = 1
    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'p':
            if _has_page_break(child):
                cur_page += 1
                per_page.setdefault(cur_page, {'blank_para': 0, 'row_gap': 0, 'content': 0})
            h = int(_para_height_twip(child))
            slot = per_page.setdefault(cur_page, {'blank_para': 0, 'row_gap': 0, 'content': 0})
            if _para_is_blank(child):
                slot['blank_para'] += h
            else:
                slot['content'] += h
        elif tag == 'tbl':
            slot = per_page.setdefault(cur_page, {'blank_para': 0, 'row_gap': 0, 'content': 0})
            for tr in child.findall(f'{W_NS}tr'):
                declared, hRule = _row_declared_height(tr)
                # 单元格最大内容高度(表格行的实际内容高度取所有列中最高的一列)
                max_cell = 0
                for tc in tr.findall(f'{W_NS}tc'):
                    max_cell = max(max_cell, _cell_content_height(tc))
                # 行的视觉高度: hRule=exact 固定;atLeast 或缺省则至少 declared,可增长
                if declared is None:
                    row_visual = max_cell if max_cell else 240
                else:
                    if hRule == 'exact':
                        row_visual = declared
                    else:  # atLeast 或未指定
                        row_visual = max(declared, max_cell)
                gap = max(0, row_visual - max_cell)  # 行内下方空白
                slot['content'] += max_cell
                slot['row_gap'] += gap

    triggered_page = None
    triggered_ratio = 0.0
    triggered_detail = None
    for pg, s in per_page.items():
        if pg == 1:
            continue  # 封面排除
        # 页底空白 (c)
        used_total = s['content'] + s['blank_para'] + s['row_gap']
        end_gap = max(0, page_usable_twip - used_total)
        blank_total = s['blank_para'] + s['row_gap'] + end_gap
        ratio = blank_total / page_usable_twip
        if ratio > 0.5:
            triggered_page = pg
            triggered_ratio = ratio
            triggered_detail = {
                'content_twip': s['content'],
                'blank_para_twip': s['blank_para'],
                'row_gap_twip': s['row_gap'],
                'end_gap_twip': end_gap,
                'usable_twip': page_usable_twip,
            }
            break

    has_large_blank = triggered_page is not None
    score_items.append((-5, "文档中出现超过页面 50% 的大段空白(扣分)", has_large_blank,
        f"触发页:{triggered_page} 空白占比:{triggered_ratio:.2%} 明细:{triggered_detail}"))

    # --- -3: 全文不可缺少核心关键词
    #         "农用颗粒播撒机构""柔性排种""防堵塞辅助结构""三维建模""运动仿真"等 ---
    # 细则每一个点:
    #   1) 范围: "全文" —— 办公软件可见的所有文本(正文段落、表格单元格、文本框内文本、页眉页脚)
    #   2) 必需关键词(细则列出): 农用颗粒播撒机构、柔性排种、防堵塞辅助结构、三维建模、运动仿真
    #      细则最后带"等"字,说明列出的这些是必须命中的示例集合;
    #      按"细则要求的每一个点都要踩到"的原则,严格判断这 5 个是否全部出现即可,
    #      不额外添加细则未列出的关键词。
    #   3) 判定: "不可缺少" —— 只要有一个未出现即触发扣分。
    # 办公软件依据:
    #   Word/WPS 中一个词可能被拆分到同段落多个 w:r/w:t(如格式变化);
    #   逐段落把所有 w:t 文本按原顺序连接后再做子串匹配,才与用户在办公软件中
    #   "查找/替换"的匹配行为一致(Word 的"查找"跨 run 但不跨段落)。
    core_keywords = ['农用颗粒播撒机构', '柔性排种', '防堵塞辅助结构',
                     '三维建模', '运动仿真']

    # 收集所有段落(含表格单元格、文本框 v:textbox/w:txbxContent、页眉页脚)
    def _collect_para_texts(doc_elem):
        texts = []
        # 主文档正文
        for p in doc_elem.iter():
            tag = p.tag.split('}')[-1] if '}' in p.tag else p.tag
            if tag != 'p':
                continue
            # 该段落所有 w:t 文本按 XML 顺序拼接(不跨段落)
            s = ''.join((t.text or '') for t in p.findall(f'.//{W_NS}t'))
            if s:
                texts.append(s)
        return texts

    para_texts = _collect_para_texts(doc.element.body)
    # 页眉页脚
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            try:
                hf_el = hf._element
            except AttributeError:
                continue
            for p in hf_el.iter():
                tag = p.tag.split('}')[-1] if '}' in p.tag else p.tag
                if tag == 'p':
                    s = ''.join((t.text or '') for t in p.findall(f'.//{W_NS}t'))
                    if s:
                        para_texts.append(s)

    missing_kw = []
    for kw in core_keywords:
        hit = any(kw in seg for seg in para_texts)
        if not hit:
            missing_kw.append(kw)

    has_missing = len(missing_kw) > 0
    score_items.append((-3, "全文缺少核心关键词(扣分)", has_missing,
        f"缺少:{missing_kw if missing_kw else '无(全部命中)'}"))

    # --- -1: 文档中残留"要改的版本""模版""模板示例""此处填写"等临时说明文字 ---
    # 细则每一个点:
    #   1) 范围: 文档中(办公软件可见的全部文本 —— 正文段、表格单元格、文本框、页眉页脚)
    #   2) 目标词(细则列出): "要改的版本""模版""模板示例""此处填写"
    #      细则末尾带"等",但按"细则没有要求的代码不加以约束",严格判断这 4 个词,
    #      不再自行扩充其它临时说明词。
    #   3) 判定: "残留" —— 只要任意一个出现即触发扣分。
    # 办公软件依据:
    #   Word/WPS 的"查找"匹配跨 run 但不跨段落,因此以段落为单位把 w:t 文本按 XML
    #   顺序拼接后再做子串匹配,与办公软件搜索行为一致。
    temp_words = ['要改的版本', '模版', '模板示例', '此处填写']

    # 复用上方 -3 中构造 para_texts 的方式;此处再次构造以保持本规则独立可读
    def _collect_all_visible_paras(d):
        segs = []
        for p in d.element.body.iter():
            tag = p.tag.split('}')[-1] if '}' in p.tag else p.tag
            if tag == 'p':
                s = ''.join((t.text or '') for t in p.findall(f'.//{W_NS}t'))
                if s:
                    segs.append(s)
        for section in d.sections:
            for hf in (section.header, section.footer,
                       section.first_page_header, section.first_page_footer,
                       section.even_page_header, section.even_page_footer):
                try:
                    hf_el = hf._element
                except AttributeError:
                    continue
                for p in hf_el.iter():
                    tag = p.tag.split('}')[-1] if '}' in p.tag else p.tag
                    if tag == 'p':
                        s = ''.join((t.text or '') for t in p.findall(f'.//{W_NS}t'))
                        if s:
                            segs.append(s)
        return segs

    all_segs = _collect_all_visible_paras(doc)

    found_temp = []
    for w in temp_words:
        if any(w in seg for seg in all_segs):
            found_temp.append(w)

    has_temp = len(found_temp) > 0
    score_items.append((-1, "残留临时说明文字(扣分)", has_temp,
        f"发现:{found_temp if found_temp else '无残留'}"))

    # --- -3: 正文缺少"一．研究目的及意义""三．研究内容与技术方案"
    #         中的任意一个一级标题 ---
    # 细则每一个点(按用户要求已删除"不可"、"五．主要参考资料"字段):
    #   1) 范围: "正文" —— 文档正文区(不含页眉页脚)。在本文档结构中,
    #      正文以表格承载,因此包含正文段落 + 表格单元格段落。
    #   2) 两个一级标题字面(细则原文):
    #        "一．研究目的及意义"
    #        "三．研究内容与技术方案"
    #      细则用的是全角"．"(dot),办公软件里也常见"、""."等写法,
    #      按 _norm_punct 归一化后比较,能对齐 Word/WPS 中的可视等价性。
    #   3) 判定: 两个中"任意一个"缺失即触发扣分。
    # 办公软件依据:
    #   一个标题可能被拆分到多个 run(格式变化);段落内 w:t 按 XML 顺序拼接后
    #   再匹配,与 Word/WPS 的查找行为一致(跨 run 不跨段落)。
    required_titles = ['一．研究目的及意义', '三．研究内容与技术方案']
    required_norms = [_norm_punct(t) for t in required_titles]

    # 收集"正文区"段落文本(主文档 body 内,不含页眉页脚)
    body_para_texts = []
    for p in doc.element.body.iter():
        tag = p.tag.split('}')[-1] if '}' in p.tag else p.tag
        if tag != 'p':
            continue
        s = ''.join((t.text or '') for t in p.findall(f'.//{W_NS}t'))
        if s:
            body_para_texts.append(s)
    body_norms = [_norm_punct(s) for s in body_para_texts]

    missing_titles = []
    for orig, norm in zip(required_titles, required_norms):
        hit = any(norm in seg for seg in body_norms)
        if not hit:
            missing_titles.append(orig)

    has_missing_title = len(missing_titles) > 0
    score_items.append((-3, "正文缺少必需一级标题(扣分)", has_missing_title,
        f"缺少:{missing_titles if missing_titles else '无(两项均在)'}"))

    return score_items


# ============== 统一入口 ==============
SCRIPT_ID = "021"


def _locate_docx(dir_path: str) -> str:
    """在给定目录中定位待评估的 .docx 文件。

    策略：
    - 仅扫描传入目录（不递归），选取扩展名为 .docx 且非 Office 临时文件（不以 ~$ 开头）的文件
    - 若目录内存在多个候选，优先取文件名中不含"副本/备份"字样的第一个
    """
    if not os.path.isdir(dir_path):
        raise FileNotFoundError(f"目录不存在: {dir_path}")
    candidates = []
    for name in os.listdir(dir_path):
        if name.startswith('~$'):
            continue
        if name.lower().endswith('.docx'):
            candidates.append(name)
    if not candidates:
        raise FileNotFoundError(f"目录中未找到 .docx 文件: {dir_path}")
    candidates.sort(key=lambda n: ('副本' in n or '备份' in n, n))
    return os.path.join(dir_path, candidates[0])


def evaluate(dir_path: str) -> dict:
    """统一评估入口。

    参数：
        dir_path: 脚本所在目录（同时也是被评估文档所在目录）。
                  脚本自己在该目录中定位并打开 .docx 文件。
    返回：
        结构化 dict，字段见文档 §2.2。
    """
    file_name = ""
    try:
        filepath = _locate_docx(dir_path)
        file_name = os.path.basename(filepath)

        # 维度一
        _d1_results, d1_passed = check_dimension1(filepath)
        if not d1_passed:
            return {
                "id": SCRIPT_ID,
                "file_name": file_name,
                "status": "ok",
                "error": None,
                "dim1_pass": False,
                "dim1_reason": "维度一未通过",
                "dim2_items": [],
                "total_score": 0,
                "max_score": 0,
            }

        # 维度二
        score_items = check_dimension2(filepath)
        dim2_items = []
        total_score = 0
        max_score = 0
        for points, desc, triggered, _detail in score_items:
            if points >= 0:
                # 加分项：命中即得 points
                delta = points if triggered else 0
                max_delta = points
                hit = bool(triggered)
            else:
                # 扣分项（points < 0）：触发才扣分
                delta = points if triggered else 0
                max_delta = points  # 扣分项使用负数 max_delta
                hit = bool(triggered)  # 命中表示触发扣分条件
            dim2_items.append({
                "rule": desc,
                "max_delta": max_delta,
                "delta": delta,
                "hit": hit,
                "detail": "",
            })
            total_score += delta
            if max_delta > 0:
                max_score += max_delta

        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": True,
            "dim1_reason": "",
            "dim2_items": dim2_items,
            "total_score": total_score,
            "max_score": max_score,
        }
    except Exception as e:
        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "error",
            "error": f"{type(e).__name__}: {e}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": 0,
        }


if __name__ == '__main__':
    import json
    # 本地自测：默认取脚本所在目录；也可通过命令行参数传入其它目录
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    print(json.dumps(evaluate(_dir), ensure_ascii=False, indent=2))
