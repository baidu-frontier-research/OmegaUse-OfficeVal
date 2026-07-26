# -*- coding: utf-8 -*-
"""
自动评估 Word 论文排版的代码
================================================
评估文件: 小微企业财务数据管理规范化问题及改进路径研究_排版完成.docx

评估逻辑
--------
1. 维度 1 (可用与可修改性): 不满足任意一条 → 总分 = 0
2. 维度 2 (完成度): 每条规则独立判定，命中即累加该点对应分值
3. 全部规则均由代码自动判定，必要时使用启发式变通实现

使用
----
本脚本对外只暴露 `evaluate(dir_path: str) -> dict`, 接收"脚本所在目录的路径",
脚本自行在该目录内定位 .docx 文件并进行评估.

$ python officeval_016_verifier.py [dir_path]

依赖
----
- python-docx
- 可选: pdfplumber + LibreOffice (soffice) 用于更精准的版面分析
"""

import os  # noqa: F401  (保留以兼容脚本作者可能后续补充的路径调试)
import re
import sys
import json
import shutil  # noqa: F401
import zipfile
import tempfile  # noqa: F401
import subprocess  # noqa: F401
from pathlib import Path
from typing import List, Dict, Tuple, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Emu, Cm
from docx.oxml.ns import qn

# =============================================================================
# 公共工具
# =============================================================================

CN_FONT_KEYS = ('宋体', '黑体', '楷体', '仿宋', 'SimSun', 'SimHei', 'KaiTi', 'FangSong',
                'Microsoft YaHei', '微软雅黑', '宋体-简', '黑体-简', '宋体', 'SimHei')

# Word 中文字号 → pt  换算(常用字号)
CN_SIZE_PT = {
    '初号': 42, '小初': 36, '一号': 26, '小一': 24, '二号': 22, '小二': 18,
    '三号': 16, '小三': 15, '四号': 14, '小四': 12, '五号': 10.5, '小五': 9,
    '六号': 7.5, '小六': 6.5, '七号': 5.5, '八号': 5,
}
PT_TO_CN = {v: k for k, v in CN_SIZE_PT.items()}


def cn_name(name: str) -> str:
    """把 SimSun / 宋体 等统一为中文名，便于规则比较"""
    if not name:
        return ''
    n = name.strip()
    mapping = {
        'SimSun': '宋体', '宋体': '宋体', 'NSimSun': '宋体',
        'SimHei': '黑体', '黑体': '黑体',
        'KaiTi': '楷体', '楷体': '楷体',
        'FangSong': '仿宋', '仿宋': '仿宋',
        'Times New Roman': 'Times New Roman',
        'Microsoft YaHei': '微软雅黑',
    }
    return mapping.get(n, n)


def pt_to_cn(pt: float) -> str:
    """将 pt 数值转换为最近的中文字号名"""
    if pt is None:
        return ''
    # 找最接近的
    best = min(CN_SIZE_PT.items(), key=lambda kv: abs(kv[1] - pt))
    if abs(best[1] - pt) <= 0.5:
        return best[0]
    return f'{pt:g}pt'


def run_font_info(run) -> Dict:
    """提取一个 run 的字体信息(中文 + 西文 + 字号 + 粗体)"""
    rPr = run._element.find(qn('w:rPr'))
    rFonts = None
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
    ascii_name = ''
    eastAsia_name = ''
    if rFonts is not None:
        ascii_name = rFonts.get(qn('w:ascii'), '') or ''
        eastAsia_name = rFonts.get(qn('w:eastAsia'), '') or ascii_name
    # fallback to run.font
    f = run.font
    if not ascii_name and f.name:
        ascii_name = f.name
    if not eastAsia_name and f.name:
        eastAsia_name = f.name
    size = f.size.pt if f.size else None
    bold = bool(f.bold) if f.bold is not None else False
    return {
        'ascii': cn_name(ascii_name),
        'eastAsia': cn_name(eastAsia_name),
        'size': size,
        'size_cn': pt_to_cn(size) if size else '',
        'bold': bold,
    }


def para_info(p) -> Dict:
    """提取段落级信息"""
    pf = p.paragraph_format
    indent_pt = None
    if pf.first_line_indent:
        indent_pt = pf.first_line_indent.pt
    # 通过 XML w:firstLineChars 判定是否为"字符"缩进(WPS显示为X字符)
    # 只有存在 firstLineChars 属性时 WPS 才会显示为"X字符"
    indent_chars = None
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            flc = ind.get(qn('w:firstLineChars'))
            if flc is not None:
                indent_chars = round(int(flc) / 100, 1)
    return {
        'style': p.style.name,
        'align': p.alignment,
        'line_spacing': pf.line_spacing,
        'line_spacing_rule': pf.line_spacing_rule,
        'indent_pt': indent_pt,
        'indent_chars': indent_chars,
        'text': p.text,
    }


def get_run_xml(page_break_pattern=re.compile(r'<w:br[^/>]*w:type="page"[^/>]*/>')):
    """读取 document.xml，返回 (xml_text, image_rids)"""
    pass  # 由调用方在 DocumentModel 里完成


# =============================================================================
# 文档模型
# =============================================================================

class DocumentModel:
    """对原 docx 文件进行结构化建模"""

    def __init__(self, path: str):
        self.path = Path(path)
        self.doc = Document(str(self.path))
        self.paragraphs: List = self.doc.paragraphs
        # 解析 XML
        with zipfile.ZipFile(self.path) as z:
            self.doc_xml = z.read('word/document.xml').decode('utf-8')
            try:
                self.rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8')
            except KeyError:
                self.rels_xml = ''
            self.media_files = sorted([n for n in z.namelist() if n.startswith('word/media/')])
            # 收集所有 footer 内容
            self.footers: Dict[str, str] = {}
            for n in z.namelist():
                if n.startswith('word/footer') and n.endswith('.xml'):
                    self.footers[n] = z.read(n).decode('utf-8')
        # 段落级元数据
        self.pinfo: List[Dict] = [para_info(p) for p in self.paragraphs]
        # 段落内每个 run 的字体信息
        self.run_info: List[List[Dict]] = [[run_font_info(r) for r in p.runs] for p in self.paragraphs]
        # 内联图片 (cover 用)
        self.inline_shapes = self.doc.inline_shapes
        # 解析页边距 / 页面尺寸
        self.sections = self.doc.sections
        # 检测每个段落前是否显式分页 (含 section break)
        self._compute_page_break_flags()

    # ---------- 页面/分页解析 ----------
    def _compute_page_break_flags(self):
        """为每个段落记录: 之前是否发生显式分页"""
        flags = [False] * len(self.paragraphs)
        # 用 XML 扫描 <w:br w:type="page"/> 与 <w:sectPr>
        # 这些标记可能出现在 run 内(分页符)或 paragraph pPr 内(section 末)
        # 我们用位置匹配的方式比较 段落起始位置 vs 分页符位置
        body = self.doc_xml
        # 取 body 开始位置
        body_start = body.find('<w:body>')
        # 提取所有分页符位置
        page_break_positions = [m.start() for m in re.finditer(r'<w:br[^/>]*w:type="page"[^/>]*/>', body)]
        sectpr_positions = [m.start() for m in re.finditer(r'<w:sectPr', body)]
        break_positions = sorted(page_break_positions + sectpr_positions)
        # 段落起始位置: 用 <w:p 出现的位置近似
        para_starts = [m.start() for m in re.finditer(r'<w:p[ >]', body)]
        # 把每个 break 后的第一个段落标记为分页后
        for bpos in break_positions:
            for i, ppos in enumerate(para_starts):
                if ppos > bpos:
                    if i < len(flags):
                        flags[i] = True
                    break
        self.page_break_before = flags
        # 估算页数:
        # 1) 显式分页位置数 + 1 (下界)
        # 2) 内容溢出估算: 正文总字符数 / 单页可容纳字符
        explicit_pages = 1 + sum(flags)
        # A4 + 1.5 倍行距 + 宋体小四:
        #  - 可写高 ≈ 246mm; 行高 ≈ 6.3mm → ~39 行/页
        #  - 每行 ≈ 28 中文字符 (含段首缩进)
        #  - → ≈ 1100 字符/页(纯文字) / ≈ 600 字符/页(中英文+标点混合)
        # 取保守值 600 字符/页, 与排版论文实际情况接近
        CHARS_PER_PAGE = 600
        total_chars = sum(len(p.text) for p in self.paragraphs)
        content_pages = max(1, (total_chars + CHARS_PER_PAGE - 1) // CHARS_PER_PAGE)
        # 取两者中较大的, 但至少是显式分页数
        self.estimated_page_count = max(explicit_pages, content_pages)

    # ---------- 辅助查询 ----------
    def first_para_idx(self, *needles: str) -> Optional[int]:
        for i, p in enumerate(self.paragraphs):
            t = p.text.strip()
            for n in needles:
                if n in t:
                    return i
        return None

    def has_image_in_paragraph(self, idx: int) -> bool:
        """判断指定段落是否包含图片(drawing)"""
        # 通过 inline_shapes 顺序匹配 (python-docx 仅暴露 inline 类型)
        # 更稳的做法: 检查该段落的 _element 是否含 w:drawing
        p = self.paragraphs[idx]
        return p._element.find('.//' + qn('w:drawing')) is not None

    def section_breaks_before(self) -> List[int]:
        """返回每个 section 起始段落索引(0 始终是 section 0 起点)"""
        breaks = [0]
        for i, flag in enumerate(self.page_break_before):
            if flag and i > 0:
                breaks.append(i)
        return breaks


# =============================================================================
# 维度 1: 可用与可修改性
# =============================================================================

def check_dimension1(model: DocumentModel) -> Tuple[bool, List[str], Dict]:
    """
    维度 1: 必须全部通过才进入维度 2
    返回 (passed, messages, details)
    """
    details = {}
    msgs = []

    # (1) 文件格式 & 可打开
    suffix_ok = model.path.suffix.lower() == '.docx'
    open_ok = True
    open_err = ''
    try:
        _ = model.doc
    except Exception as e:
        open_ok = False
        open_err = str(e)
    details['format_ok'] = suffix_ok and open_ok
    if not details['format_ok']:
        msgs.append(f'❌ 文件格式或打开失败: 后缀={model.path.suffix}, 打开错误={open_err or "无"}')

    # (2) 页数 > 6
    page_count = model.estimated_page_count
    details['page_count'] = page_count
    details['page_count_ok'] = page_count > 6
    if not details['page_count_ok']:
        msgs.append(f'❌ 估算页数 = {page_count} (不满足 > 6)')

    passed = all([
        details['format_ok'],
        details['page_count_ok'],
    ])
    return passed, msgs, details


# =============================================================================
# 维度 2: 完成度评分
# =============================================================================

def avg_font(infos: List[Dict], key: str = 'eastAsia') -> str:
    """从一段的多个 run 中取主要的(非空)字体"""
    if not infos:
        return ''
    # 统计非空值
    from collections import Counter
    c = Counter([i[key] for i in infos if i.get(key)])
    if not c:
        return ''
    return c.most_common(1)[0][0]


def avg_size(infos: List[Dict]) -> Optional[float]:
    sizes = [i['size'] for i in infos if i.get('size')]
    if not sizes:
        return None
    return sum(sizes) / len(sizes)


def all_bold(infos: List[Dict]) -> bool:
    if not infos:
        return False
    return all(i['bold'] for i in infos if i.get('eastAsia') or i.get('ascii'))


def line_spacing_value(info: Dict) -> float:
    """统一返回行距数值(多倍行距)"""
    ls = info.get('line_spacing')
    if ls is None:
        return 1.0
    rule = info.get('line_spacing_rule')
    if rule == WD_LINE_SPACING.SINGLE:
        return 1.0
    if rule == WD_LINE_SPACING.ONE_POINT_FIVE:
        return 1.5
    if rule == WD_LINE_SPACING.DOUBLE:
        return 2.0
    if rule == WD_LINE_SPACING.MULTIPLE:
        return float(ls)
    # pt 值: 默认单倍 ≈ 12pt; 估算倍
    return float(ls) / 12.0


def check_dimension2(model: DocumentModel) -> Tuple[int, List[Tuple[str, int, str, bool]]]:
    """
    返回 (总分, 命中点列表[(描述, 分值, 依据, 是否命中)])
    """
    hits: List[Tuple[str, int, str, bool]] = []
    total = 0

    def add(label: str, score: int, evidence: str, hit: bool):
        nonlocal total
        if hit:
            total += score
        hits.append((label, score, evidence, hit))

    # -------------------------------------------------------------------------
    # 封面页 +5: 严格按细则四项要求判定
    #   (1) 位于文档第 1 页  → 封面内容全部在首个分页之前
    #   (2) 页面为 A4 纵向    → 宽 ≈ 21cm, 高 ≈ 29.7cm
    #   (3) 页面无页码        → 封面所在 section 生效的 footer 中不含 PAGE 字段
    #   (4) 四个区域完整      → 顶部校徽校名组合图 + 主标题"成教本科毕业论文"
    #                            + 五项信息栏 + 底部日期
    # 说明: 字体/字号/居中/加粗等由后续独立子项检查, 本条不重复约束
    # -------------------------------------------------------------------------
    cover_ok = True
    cover_reasons = []

    # 封面段落范围: 从首段到首个显式分页之前 (即文档第 1 页内容)
    cover_end_idx = len(model.paragraphs)
    for _i in range(1, len(model.paragraphs)):
        if model.page_break_before[_i]:
            cover_end_idx = _i
            break
    cover_range = list(range(0, cover_end_idx))

    # (1) 位于文档第 1 页: 第 1 页至少要有封面内容
    if cover_end_idx <= 0:
        cover_ok = False
        cover_reasons.append('文档第 1 页无任何段落内容')

    # (2) A4 纵向 (Word/WPS 均以 sectPr/pgSz 生效)
    sec0 = model.sections[0]
    page_w_cm = sec0.page_width / 360000
    page_h_cm = sec0.page_height / 360000
    a4_portrait = (20.5 < page_w_cm < 21.5) and (29.0 < page_h_cm < 30.0)
    if not a4_portrait:
        cover_ok = False
        cover_reasons.append(f'页面非 A4 纵向: {page_w_cm:.2f}x{page_h_cm:.2f} cm')

    # (3) 封面页无页码
    #     兼容办公软件常见三种做法:
    #       a. 封面 section 完全无 footerReference          → 无页码
    #       b. 封面 section 启用 titlePg, 首页 footer 无 PAGE → 无页码
    #       c. 封面 section 无 titlePg, 默认 footer 无 PAGE  → 无页码
    rels_map = {}
    for elem in re.finditer(r'<Relationship\b[^>]*/>', model.rels_xml):
        attrs = elem.group(0)
        id_m = re.search(r'\bId="([^"]+)"', attrs)
        tgt_m = re.search(r'\bTarget="([^"]+)"', attrs)
        if id_m and tgt_m:
            rels_map[id_m.group(1)] = tgt_m.group(1)

    def _cover_footer_of(sect_xml: str, ftype: str) -> str:
        m = re.search(
            rf'<w:footerReference[^/>]*w:type="{ftype}"[^/>]*r:id="(rId\d+)"',
            sect_xml,
        )
        if not m:
            return ''
        target = rels_map.get(m.group(1), '')
        if not target:
            return ''
        return model.footers.get('word/' + target.lstrip('/'), '')

    body_xml = model.doc_xml
    sect_blocks = re.findall(r'<w:sectPr[^>]*>.*?</w:sectPr>', body_xml, re.S)

    cover_no_page_num = True
    cover_footer_note = ''
    if sect_blocks:
        first_sect = sect_blocks[0]
        has_title_pg = '<w:titlePg' in first_sect
        default_footer = _cover_footer_of(first_sect, 'default')
        first_footer = _cover_footer_of(first_sect, 'first')
        # 封面 = 第 1 页, 判定其实际生效的 footer
        if has_title_pg:
            effective_footer = first_footer
            footer_kind = '首页 footer (titlePg)'
        else:
            effective_footer = default_footer
            footer_kind = '默认 footer'
        if 'PAGE' in effective_footer:
            cover_no_page_num = False
            cover_footer_note = f'{footer_kind}含 PAGE 字段'
    if not cover_no_page_num:
        cover_ok = False
        cover_reasons.append(cover_footer_note or '封面页含页码')

    # (4a) 顶部: 校徽校名组合图 — 封面上部含图片, 且位于主标题之前
    title_idx = None
    for _i in cover_range:
        if '成教本科毕业论文' in model.paragraphs[_i].text:
            title_idx = _i
            break
    cover_image_idx = None
    for _i in cover_range:
        if model.has_image_in_paragraph(_i):
            cover_image_idx = _i
            break
    top_image_ok = (
        cover_image_idx is not None
        and (title_idx is None or cover_image_idx < title_idx)
    )
    if not top_image_ok:
        cover_ok = False
        cover_reasons.append(
            f'封面顶部无校徽校名组合图 (image@{cover_image_idx}, title@{title_idx})'
        )

    # (4b) 主标题 "成教本科毕业论文"
    if title_idx is None:
        cover_ok = False
        cover_reasons.append('封面未找到主标题 "成教本科毕业论文"')

    # (4c) 五项信息栏: 题目 / 学号 / 姓名 / 专业 / 指导老师
    info_labels = ['题', '学', '姓', '专', '指导老师']
    found_labels = set()
    for _i in cover_range:
        _t = model.paragraphs[_i].text
        for lbl in info_labels:
            if lbl in _t:
                found_labels.add(lbl)
    if len(found_labels) < 5:
        cover_ok = False
        cover_reasons.append(f'五项信息栏不完整: 命中={sorted(found_labels)}')

    # (4d) 底部日期: 封面区最后一段年份文本, 且位于主标题之后
    date_idx = None
    for _i in cover_range:
        if re.search(r'20\d{2}\s*年', model.paragraphs[_i].text):
            date_idx = _i  # 取最靠底部 (最大 idx) 的一段
    if date_idx is None:
        cover_ok = False
        cover_reasons.append('封面未找到 "YYYY 年" 底部日期')
    elif title_idx is not None and date_idx < title_idx:
        cover_ok = False
        cover_reasons.append('日期未位于封面底部 (出现在主标题之前)')

    evidence = (
        f'第1页段数={cover_end_idx}, A4纵向={a4_portrait}, 无页码={cover_no_page_num}, '
        f'图@{cover_image_idx}, 主标题@{title_idx}, '
        f'信息栏={sorted(found_labels)}, 日期@{date_idx}'
    )
    if not cover_ok:
        evidence += ' | 问题: ' + '; '.join(cover_reasons)
    add('+5 封面页(第1页/A4纵向/无页码/四区域完整)', 5, evidence, cover_ok)

    # -------------------------------------------------------------------------
    # +1 封面页顶部校徽校名组合图: 严格按细则四点判定
    #   (1) 从上至下整体位于页面 0-30% 处
    #   (2) 组合图在一行排列 (校徽与校名在同一段落 → 一行内)
    #   (3) 整体水平居中 (段落对齐 = 居中)
    #   (4) 校徽在校名"松江应用技术学院"的左侧 (XML 顺序: 校徽 run 在校名 run 之前)
    # 兼容办公软件: 通过 sectPr 页面尺寸、段落 <w:jc>、run 顺序判定,
    #                Word/WPS 对同一段落内内联对象的水平流方向一致 (左→右)
    # -------------------------------------------------------------------------
    img_ok = True
    img_reason = []

    # 校徽校名组合所在段落: cover_range 中含图片(w:drawing)的最靠前段落
    combo_idx = None
    for _i in cover_range:
        if model.has_image_in_paragraph(_i):
            combo_idx = _i
            break

    if combo_idx is None:
        img_ok = False
        img_reason.append('封面未找到含图片的段落')

    # ---- (1) 从上至下整体位于页面 0-30% 处 ----
    # 估算组合图所在段落顶部相对页面顶部的距离
    # 距离 = 上边距 + Σ(之前每个段落估算高度)
    pos_ratio = None
    if combo_idx is not None:
        sec = model.sections[0]
        page_h_emu = sec.page_height  # EMU
        top_margin_emu = sec.top_margin or 0
        PT_TO_EMU = 12700  # 1pt = 12700 EMU

        def _estimate_para_height_emu(idx: int) -> int:
            """估算段落占用的垂直高度 (EMU), 与 Word/WPS 实际渲染近似"""
            # 若段落含图片, 用图片高度作为主体
            if model.has_image_in_paragraph(idx):
                h_max = 0
                for r in model.paragraphs[idx].runs:
                    for drawing in r._element.findall('.//' + qn('w:drawing')):
                        # inline extent
                        for ext in drawing.iter(qn('wp:extent')):
                            try:
                                h_max = max(h_max, int(ext.get('cy') or 0))
                            except (TypeError, ValueError):
                                pass
                if h_max > 0:
                    return h_max
            # 文本段落: 以首个 run 的字号 × 行距为单行高
            infos = model.run_info[idx]
            sz_pt = None
            for r in infos:
                if r.get('size'):
                    sz_pt = r['size']
                    break
            if sz_pt is None:
                sz_pt = 10.5  # Word/WPS 默认五号
            pi = model.pinfo[idx]
            ls = line_spacing_value(pi) or 1.0
            line_h_pt = sz_pt * ls
            # 行数: 空段计 1 行, 有文本按约 30 字符/行估算
            text_len = len(model.paragraphs[idx].text)
            n_lines = 1 if text_len == 0 else max(1, (text_len + 29) // 30)
            return int(line_h_pt * n_lines * PT_TO_EMU)

        cumulative = top_margin_emu
        for _i in range(0, combo_idx):
            cumulative += _estimate_para_height_emu(_i)
        pos_ratio = cumulative / page_h_emu if page_h_emu else 1.0
        if pos_ratio > 0.30:
            img_ok = False
            img_reason.append(
                f'组合图未在页面 0-30% 位置 (实际 {pos_ratio * 100:.1f}%)'
            )

    # ---- (2) 组合图在一行排列 ----
    # 判定: 校徽与校名位于同一段落 (一段 = 一行内联流). 校名可以是图片或文字
    same_line = False
    if combo_idx is not None:
        p_elem = model.paragraphs[combo_idx]._element
        drawing_runs = [
            r for r in model.paragraphs[combo_idx].runs
            if r._element.find('.//' + qn('w:drawing')) is not None
        ]
        text_of_para = model.paragraphs[combo_idx].text
        # 情形 A: 校徽图 + 校名图 均为内联图, 同段 → 一行
        if len(drawing_runs) >= 2:
            same_line = True
        # 情形 B: 校徽图 + 校名文本 "松江应用技术学院" 在同一段
        elif len(drawing_runs) >= 1 and '松江应用技术学院' in text_of_para:
            same_line = True
        # 段落内不应含分行符 (<w:br/> 无 type 或 type="textWrapping")
        for br in p_elem.iter(qn('w:br')):
            btype = br.get(qn('w:type')) or 'textWrapping'
            if btype == 'textWrapping':
                same_line = False
                img_reason.append('组合图段落内含换行符, 未在一行')
                break
        if not same_line:
            img_ok = False
            if '组合图段落内含换行符' not in ';'.join(img_reason):
                img_reason.append('校徽与校名未在同一段(一行)内排列')

    # ---- (3) 整体水平居中 ----
    if combo_idx is not None:
        align = model.pinfo[combo_idx]['align']
        if align != WD_ALIGN_PARAGRAPH.CENTER:
            img_ok = False
            img_reason.append(f'组合图段落未水平居中 (align={align})')

    # ---- (4) 校徽在校名左侧 ----
    # 顺序判定: 遍历段落 run, 记录校徽 run 索引 与 校名 run 索引
    #   - 校徽: 首个含 w:drawing 且宽高比接近 1:1 (方形徽章) 的 run
    #   - 校名: 后续含"松江应用技术学院"文本的 run, 或第二个 drawing (宽幅横条)
    badge_before_name = False
    if combo_idx is not None:
        runs = model.paragraphs[combo_idx].runs
        badge_run_idx = None
        name_run_idx = None
        for ri, r in enumerate(runs):
            drawings = r._element.findall('.//' + qn('w:drawing'))
            if drawings:
                # 计算图形宽高比
                cx = cy = 0
                for ext in drawings[0].iter(qn('wp:extent')):
                    try:
                        cx = int(ext.get('cx') or 0)
                        cy = int(ext.get('cy') or 0)
                    except (TypeError, ValueError):
                        pass
                ratio = (cx / cy) if cy else 0
                if ratio and ratio < 1.5:
                    # 近方形 → 校徽
                    if badge_run_idx is None:
                        badge_run_idx = ri
                elif ratio >= 1.5:
                    # 宽幅横条 → 校名图
                    if name_run_idx is None:
                        name_run_idx = ri
            if '松江应用技术学院' in r.text and name_run_idx is None:
                name_run_idx = ri
        if badge_run_idx is not None and name_run_idx is not None:
            badge_before_name = badge_run_idx < name_run_idx
        elif badge_run_idx is not None and '松江应用技术学院' in model.paragraphs[combo_idx].text:
            # 只识别到校徽 run + 文本形式校名 (且未识别到独立 name_run_idx 时兜底)
            badge_before_name = True
        if not badge_before_name:
            img_ok = False
            img_reason.append(
                f'校徽未在校名左侧 (badge_run={badge_run_idx}, name_run={name_run_idx})'
            )

    evidence = (
        f'组合图段@{combo_idx}, 位置={pos_ratio * 100:.1f}%'
        if pos_ratio is not None else f'组合图段@{combo_idx}'
    )
    if not img_ok:
        evidence += ' | 问题: ' + '; '.join(img_reason)
    add('+1 校徽校名组合图(0-30% / 一行 / 水平居中 / 校徽在左)',
        1, evidence, img_ok)

    # -------------------------------------------------------------------------
    # +1 封面页主标题 "成教本科毕业论文": 严格按细则五点判定
    #   (1) 文本为 "成教本科毕业论文"
    #   (2) 位于校徽校名组合图下方 (title 段索引 > combo_idx)
    #   (3) 单独成行 (整段 = 主标题文本, 不与其他内容共段)
    #   (4) 字体 = 黑体 (中文东亚字体)
    #   (5) 字号 = 小初 (36pt)
    #   (6) 加粗 (段中所有非空 run 均 bold)
    #   (7) 居中 (段落 <w:jc w:val="center"/>)
    # 兼容办公软件: 黑体识别接受 "黑体"/"SimHei" (Word/WPS 底层名); 字号用 w:sz 半点值
    # -------------------------------------------------------------------------
    title_ok = True
    title_reason = []

    if title_idx is None:
        title_ok = False
        title_reason.append('未在封面找到主标题 "成教本科毕业论文"')
    else:
        # (2) 位于组合图下方
        if combo_idx is None:
            title_ok = False
            title_reason.append('未找到校徽校名组合图, 无法判定主标题在其下方')
        elif title_idx <= combo_idx:
            title_ok = False
            title_reason.append(
                f'主标题未位于组合图下方 (title@{title_idx}, combo@{combo_idx})'
            )

        # (3) 单独成行: 该段整段文本 == "成教本科毕业论文" (允许首尾空白)
        title_text = model.paragraphs[title_idx].text.strip()
        # 段内不应含换行符 <w:br/> (textWrapping)
        has_line_break = False
        for br in model.paragraphs[title_idx]._element.iter(qn('w:br')):
            btype = br.get(qn('w:type')) or 'textWrapping'
            if btype == 'textWrapping':
                has_line_break = True
                break
        if title_text != '成教本科毕业论文':
            title_ok = False
            title_reason.append(f'主标题段含其他内容, 未单独成行: {title_text!r}')
        if has_line_break:
            title_ok = False
            title_reason.append('主标题段内含软换行, 未单独成行')

        # (4)(5)(6) 字体/字号/加粗: 逐 run 校验有实际文本的 run
        for ridx, rinfo in enumerate(model.run_info[title_idx]):
            rtext = (model.paragraphs[title_idx].runs[ridx].text or '')
            if not rtext.strip():
                continue
            if rinfo['eastAsia'] != '黑体':
                title_ok = False
                title_reason.append(f'run[{ridx}] 中文字体非黑体: {rinfo["eastAsia"]}')
            sz = rinfo['size']
            if sz is None or abs(sz - 36) > 0.5:
                title_ok = False
                title_reason.append(f'run[{ridx}] 字号非小初(36pt): {sz}')
            if not rinfo['bold']:
                title_ok = False
                title_reason.append(f'run[{ridx}] 未加粗')

        # (7) 居中
        align = model.pinfo[title_idx]['align']
        if align != WD_ALIGN_PARAGRAPH.CENTER:
            title_ok = False
            title_reason.append(f'主标题未居中 (align={align})')

    evidence = (
        f'主标题@{title_idx}, 组合图@{combo_idx}'
    )
    if not title_ok:
        evidence += ' | 问题: ' + '; '.join(title_reason)
    add('+1 封面主标题(组合图下方/单独成行/黑体小初加粗居中)',
        1, evidence, title_ok)

    # -------------------------------------------------------------------------
    # +3 封面页信息栏: 严格按细则四点判定
    #   (1) 包含 "题目"、"学号"、"姓名"、"专业"、"指导老师" 五项
    #   (2) 五项纵向居中排列且左右两端相互对齐
    #       → 5 段相邻排列 (纵向为列), 段落 alignment=CENTER (居中排列),
    #         且左右缩进/边距一致 (左右两端相互对齐)
    #   (3) 标签后横线左端对齐且长度一致
    #       → 标签部分显示宽度一致 (标签 run 字符数相同) → 横线左端 X 相同
    #         横线部分显示宽度一致 (tab 位置相同 或 下划线区总长相同) → 长度一致
    #   (4) 五项文本(标签) 字体 = 黑体、字号 = 三号、加粗
    # 兼容 Word/WPS: 段落对齐/缩进/tab 位置/rFonts/w:sz/w:b 均是两软件共用的核心 OOXML
    # -------------------------------------------------------------------------
    info_ok = True
    info_reason = []

    # (1) 五项存在: 按"去空格后子串包含"识别, 顺序按细则给定
    label_keys = ['题目', '学号', '姓名', '专业', '指导老师']
    label_paras = {}  # {标签: 段索引}
    for _i in cover_range:
        raw = model.paragraphs[_i].text
        collapsed = raw.replace(' ', '').replace('　', '').replace('\t', '')
        for lbl in label_keys:
            if lbl in collapsed and lbl not in label_paras:
                label_paras[lbl] = _i
    missing = [lbl for lbl in label_keys if lbl not in label_paras]
    if missing:
        info_ok = False
        info_reason.append(f'信息栏缺少 {missing}')

    ordered_idx = [label_paras[lbl] for lbl in label_keys if lbl in label_paras]

    # (2) 纵向居中排列 + 左右两端相互对齐
    #   纵向排列: 5 段索引严格递增 (从上至下排列成列)
    #   居中: 5 段 alignment 均为 CENTER
    #   左右两端相互对齐: 5 段落 <w:ind w:left>/<w:right> 一致 (默认均为 0)
    if len(ordered_idx) == 5:
        if ordered_idx != sorted(ordered_idx):
            info_ok = False
            info_reason.append(f'五项未按顺序纵向排列: {ordered_idx}')
        aligns = [model.pinfo[i]['align'] for i in ordered_idx]
        if any(a != WD_ALIGN_PARAGRAPH.CENTER for a in aligns):
            info_ok = False
            info_reason.append(f'五项非纵向居中排列: aligns={aligns}')
        ind_pairs = []  # (left, right) in twips
        for i in ordered_idx:
            pPr = model.paragraphs[i]._element.find(qn('w:pPr'))
            l = r = 0
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    try:
                        l = int(ind.get(qn('w:left')) or ind.get(qn('w:start')) or 0)
                    except (TypeError, ValueError):
                        l = 0
                    try:
                        r = int(ind.get(qn('w:right')) or ind.get(qn('w:end')) or 0)
                    except (TypeError, ValueError):
                        r = 0
            ind_pairs.append((l, r))
        if len(set(ind_pairs)) > 1:
            info_ok = False
            info_reason.append(f'五项左右边距不一致: {ind_pairs}')

    # (3) 标签后横线: 左端对齐 (标签宽度一致) + 长度一致
    if len(ordered_idx) == 5:
        # 3a 标签部分宽度: 用标签首 run 的可见字符数(把连续空白按 1 字符宽近似)
        # Word/WPS 中中文标签用 "题 目"/"学 号"/"姓 名"/"专 业"/"指导老师" 通常填充为 4 CJK 宽度,
        # 判定: 每段首个非空 run 的字符宽度相同
        # 显示宽度按: 每个 CJK 全宽=2, 每个 ASCII/空格=1
        def _display_width(s: str) -> int:
            w = 0
            for ch in s:
                if ch == '\t':
                    continue  # tab 不计入标签宽度
                if '一' <= ch <= '鿿' or '　' <= ch <= '〿' or ch in '　':
                    w += 2
                elif ch == ' ':
                    w += 1
                else:
                    w += 1
            return w

        label_widths = []
        line_specs = []  # (tab_positions:list[int], underline_run_count:int, underline_char_count:int)
        for i in ordered_idx:
            p = model.paragraphs[i]
            runs = p.runs
            # 首个非空 run 视作标签 run
            label_run = None
            label_run_idx = 0
            for ri, r in enumerate(runs):
                if r.text.strip():
                    label_run = r
                    label_run_idx = ri
                    break
            lw = _display_width(label_run.text) if label_run is not None else 0
            label_widths.append(lw)

            # tab 位置 (定义横线终点)
            pPr = p._element.find(qn('w:pPr'))
            tab_pos = []
            if pPr is not None:
                tabs_el = pPr.find(qn('w:tabs'))
                if tabs_el is not None:
                    for tab in tabs_el.findall(qn('w:tab')):
                        pos_val = tab.get(qn('w:pos'))
                        if pos_val:
                            try:
                                tab_pos.append(int(pos_val))
                            except ValueError:
                                pass
            # 横线区域: 标签后的下划线 run(含 tab 或字符)
            u_run_count = 0
            u_char_count = 0
            for ri, r in enumerate(runs):
                if ri <= label_run_idx:
                    continue
                if r.font.underline:
                    u_run_count += 1
                    u_char_count += len(r.text)
            line_specs.append((tab_pos, u_run_count, u_char_count))

        if len(set(label_widths)) > 1:
            info_ok = False
            info_reason.append(
                f'标签显示宽度不一致 → 横线左端未对齐: {list(zip(label_keys, label_widths))}'
            )

        # 3b 横线长度一致:
        #   优先按 tab 位置判定 (Word/WPS 均按 pos 精确定位横线终点);
        #   若无 tab, 退化为按下划线 run 数量与字符数同时相同判定
        tab_positions_all = [tp for tp, _, _ in line_specs]
        if all(tp for tp in tab_positions_all):
            # 全部段落都有 tab: 位置必须一致
            if len(set(tuple(tp) for tp in tab_positions_all)) > 1:
                info_ok = False
                info_reason.append(f'横线终点 tab 位置不一致: {tab_positions_all}')
        else:
            # 有段落无 tab (如题目段直接下划线题目文本) → 用下划线区宽度判定
            u_run_counts = [urc for _, urc, _ in line_specs]
            u_char_counts = [ucc for _, _, ucc in line_specs]
            if len(set(u_run_counts)) > 1 or len(set(u_char_counts)) > 1:
                info_ok = False
                info_reason.append(
                    f'横线长度不一致: run数={u_run_counts}, 字符数={u_char_counts}'
                )
        # 且每段必须存在横线 (至少一处下划线)
        no_line = [
            label_keys[k] for k, (_, urc, _) in enumerate(line_specs) if urc == 0
        ]
        if no_line:
            info_ok = False
            info_reason.append(f'以下项标签后无横线: {no_line}')

    # (4) 五项文本(标签) 字体 = 黑体 / 字号 = 三号(16pt) / 加粗
    for lbl in label_keys:
        if lbl not in label_paras:
            continue
        idx = label_paras[lbl]
        # 取该段首个含文本的 run 作为标签 run
        label_rinfo = None
        for ri, rinfo in enumerate(model.run_info[idx]):
            rtext = model.paragraphs[idx].runs[ri].text if ri < len(model.paragraphs[idx].runs) else ''
            if rtext.strip():
                label_rinfo = rinfo
                break
        if label_rinfo is None:
            info_ok = False
            info_reason.append(f'"{lbl}" 段无有效标签 run')
            continue
        if label_rinfo['eastAsia'] != '黑体':
            info_ok = False
            info_reason.append(f'"{lbl}" 非黑体: {label_rinfo["eastAsia"]}')
        sz = label_rinfo['size']
        if sz is None or abs(sz - 16) > 0.5:
            info_ok = False
            info_reason.append(f'"{lbl}" 非三号(16pt): {sz}')
        if not label_rinfo['bold']:
            info_ok = False
            info_reason.append(f'"{lbl}" 未加粗')

    evidence = f'五项段落={ordered_idx}'
    if not info_ok:
        evidence += ' | 问题: ' + '; '.join(info_reason)
    add('+3 信息栏(5项/纵向居中排列/左右对齐/横线左端对齐且等长/标签黑体三号加粗)',
        3, evidence, info_ok)

    # -------------------------------------------------------------------------
    # +1 封面页日期: 严格按细则四点判定
    #   (1) 页面底部显示                (段落位于封面靠底部, 约 >= 70% 页高)
    #   (2) 居中显示                    (段落 <w:jc w:val="center"/>)
    #   (3) 文本为 "2026年3月"          (细则给定, strip 后精确匹配)
    #   (4) 数字部分: Times New Roman / 四号(14pt) / 加粗
    #       中文部分: 黑体            / 四号(14pt) / 加粗
    # 兼容 Word/WPS:
    #   - Word/WPS 均按 run 的 w:rFonts@w:ascii 渲染 ASCII 数字, 按 @w:eastAsia 渲染中文
    #   - "位于底部" 用页面绝对坐标估算, 与两软件实际渲染一致
    # -------------------------------------------------------------------------
    date_ok = True
    date_reason = []

    if date_idx is None:
        date_ok = False
        date_reason.append('封面未找到 "2026年3月" 段')
    else:
        # (3) 文本 = "2026年3月"
        date_text = model.paragraphs[date_idx].text.strip()
        if date_text != '2026年3月':
            date_ok = False
            date_reason.append(f'日期文本非 "2026年3月": {date_text!r}')

        # (2) 居中
        align = model.pinfo[date_idx]['align']
        if align != WD_ALIGN_PARAGRAPH.CENTER:
            date_ok = False
            date_reason.append(f'日期未居中 (align={align})')

        # (1) 位于页面底部
        # 判定: 日期段在封面段落序位的下 30% (即 (idx+1)/cover_end_idx >= 0.70)
        # 说明: Word/WPS 中封面由一列段落纵向排布, 段落序位比例即页面上下位置比例
        cover_len = max(1, cover_end_idx)
        date_pos_ratio = (date_idx + 1) / cover_len
        if date_pos_ratio < 0.70:
            date_ok = False
            date_reason.append(
                f'日期未位于页面底部 (段位 {date_idx + 1}/{cover_len}={date_pos_ratio * 100:.1f}%, 应 >= 70%)'
            )

        # (4) 字体 / 字号 / 加粗: 按每个 run 的实际字符类型判定
        runs = model.paragraphs[date_idx].runs
        for ri, r in enumerate(runs):
            rtext = r.text or ''
            if not rtext.strip():
                continue
            rinfo = model.run_info[date_idx][ri]

            # 字号: 四号 = 14pt
            sz = rinfo['size']
            if sz is None or abs(sz - 14) > 0.5:
                date_ok = False
                date_reason.append(f'run[{ri}] {rtext!r} 字号非四号(14pt): {sz}')

            # 加粗
            if not rinfo['bold']:
                date_ok = False
                date_reason.append(f'run[{ri}] {rtext!r} 未加粗')

            # 字体: 逐字符按类别校验
            has_digit = any(ch.isdigit() for ch in rtext)
            has_cjk = any('一' <= ch <= '鿿' for ch in rtext)
            if has_digit:
                # Word/WPS 渲染 ASCII 数字使用 w:ascii 字段
                if rinfo['ascii'] != 'Times New Roman':
                    date_ok = False
                    date_reason.append(
                        f'run[{ri}] 数字 {rtext!r} 字体非 Times New Roman: '
                        f'ascii={rinfo["ascii"]!r}'
                    )
            if has_cjk:
                # Word/WPS 渲染 CJK 使用 w:eastAsia 字段
                if rinfo['eastAsia'] != '黑体':
                    date_ok = False
                    date_reason.append(
                        f'run[{ri}] 中文 {rtext!r} 字体非黑体: '
                        f'eastAsia={rinfo["eastAsia"]!r}'
                    )

    evidence = (
        f'日期@{date_idx}, 位置={date_pos_ratio * 100:.1f}%'
        if date_idx is not None else f'日期@{date_idx}'
    )
    if not date_ok:
        evidence += ' | 问题: ' + '; '.join(date_reason)
    add('+1 封面日期(底部居中 / "2026年3月" / 数字TNR四号加粗 / 中文黑体四号加粗)',
        1, evidence, date_ok)

    # -------------------------------------------------------------------------
    # +1 原创性声明及使用授权声明页: 严格按细则两点判定
    #   (1) 是整篇文档的第 2 页
    #   (2) 位于封面页之后
    # 兼容 Word/WPS: 页码位置由 <w:br w:type="page"/> 与 <w:sectPr> 决定,
    #                两软件对分页标记的解释一致
    # -------------------------------------------------------------------------
    # 定位 "原创性声明" 所在段: 允许全称 "原创性声明及使用授权声明" 或简称 "原创性声明"
    decl_idx = None
    for _i, _p in enumerate(model.paragraphs):
        _t = _p.text.strip()
        if '原创性声明' in _t:
            decl_idx = _i
            break

    decl_ok = True
    decl_reason = []

    if decl_idx is None:
        decl_ok = False
        decl_reason.append('未找到 "原创性声明" 段')
        decl_page = None
    else:
        # (1) 计算原创性声明所在页:
        #     用其之前(不含自身)出现的分页标记数 + 1
        #     model.page_break_before[i] 表示第 i 段前发生了分页
        decl_page = 1 + sum(1 for _i in range(0, decl_idx + 1) if model.page_break_before[_i])
        if decl_page != 2:
            decl_ok = False
            decl_reason.append(f'原创性声明位于第 {decl_page} 页, 非第 2 页')

        # (2) 位于封面页之后: 即 decl_idx > cover_end_idx (封面段落上界)
        #     且 decl_idx 之前必须存在一次显式分页 (封面与声明之间)
        if decl_idx <= cover_end_idx - 1:
            decl_ok = False
            decl_reason.append(
                f'原创性声明未位于封面页之后 (声明@{decl_idx}, 封面末段={cover_end_idx - 1})'
            )
        has_break_after_cover = any(
            model.page_break_before[_i] for _i in range(1, decl_idx + 1)
        )
        if not has_break_after_cover:
            decl_ok = False
            decl_reason.append('封面页与原创性声明之间无分页')

    evidence = f'声明@{decl_idx}, 页码={decl_page}, 封面末段={cover_end_idx - 1}'
    if not decl_ok:
        evidence += ' | 问题: ' + '; '.join(decl_reason)
    add('+1 原创性声明及使用授权声明页(文档第2页/位于封面页之后)',
        1, evidence, decl_ok)

    # -------------------------------------------------------------------------
    # +1 原创性声明标题(宋体三号加粗)
    # -------------------------------------------------------------------------
    # -------------------------------------------------------------------------
    # +1 原创性声明及使用授权声明页标题: 严格按细则四点判定
    #   (1) 出现标题文本 "原创性声明及使用授权声明"
    #   (2) 字体 = 宋体   (中文东亚字体)
    #   (3) 字号 = 三号   (16pt)
    #   (4) 加粗
    # 兼容 Word/WPS: 通过 OOXML 的 w:rFonts@eastAsia / w:sz / w:b 判定,
    #                cn_name() 已把 SimSun 归一为 宋体
    # -------------------------------------------------------------------------
    decl_title_ok = True
    decl_title_reason = []

    # 定位标题所在段: 该段整段文本 (strip 后) 精确等于 "原创性声明及使用授权声明"
    decl_title_idx = None
    for _i, _p in enumerate(model.paragraphs):
        if _p.text.strip() == '原创性声明及使用授权声明':
            decl_title_idx = _i
            break

    if decl_title_idx is None:
        decl_title_ok = False
        decl_title_reason.append('未找到标题 "原创性声明及使用授权声明"')
    else:
        # 逐 run 校验有文本的 run
        for ri, rinfo in enumerate(model.run_info[decl_title_idx]):
            rtext = (model.paragraphs[decl_title_idx].runs[ri].text or '')
            if not rtext.strip():
                continue
            # (2) 字体 = 宋体
            if rinfo['eastAsia'] != '宋体':
                decl_title_ok = False
                decl_title_reason.append(
                    f'run[{ri}] {rtext!r} 中文字体非宋体: {rinfo["eastAsia"]}'
                )
            # (3) 字号 = 三号 (16pt)
            sz = rinfo['size']
            if sz is None or abs(sz - 16) > 0.5:
                decl_title_ok = False
                decl_title_reason.append(
                    f'run[{ri}] {rtext!r} 字号非三号(16pt): {sz}'
                )
            # (4) 加粗
            if not rinfo['bold']:
                decl_title_ok = False
                decl_title_reason.append(f'run[{ri}] {rtext!r} 未加粗')

    evidence = f'标题@{decl_title_idx}'
    if not decl_title_ok:
        evidence += ' | 问题: ' + '; '.join(decl_title_reason)
    add('+1 原创性声明标题("原创性声明及使用授权声明"/宋体/三号/加粗)',
        1, evidence, decl_title_ok)

    # -------------------------------------------------------------------------
    # +1 原创性声明及使用授权声明页标题下方文本: 严格按细则三点判定
    #   (1) 字体 = 宋体、四号 (14pt)
    #   (2) 首行缩进 2 字符 (Word/WPS: w:ind@w:firstLineChars = 200)
    #   (3) 段落间距为 1.5 倍行距
    # 兼容 Word/WPS:
    #   - 首行缩进 "字符" 由 w:firstLineChars 表达 (百分之字符单位, 200 = 2 字符);
    #     两软件的"X 字符"缩进 UI 均写入此属性
    #   - 1.5 倍行距 = w:spacing@w:line=360 且 w:lineRule="auto"
    # -------------------------------------------------------------------------
    decl_body_ok = True
    decl_body_reason = []

    # 定位声明标题下方的正文段: 标题段之后, 到签名区(含"签名"字样) 或空段之前
    decl_body_paras = []
    if decl_title_idx is not None:
        for j in range(decl_title_idx + 1, len(model.paragraphs)):
            t = model.paragraphs[j].text.strip()
            # 遇到签名区 → 停止
            if '签名' in t and '日' in t and '期' in t:
                break
            # 空段跳过 (Word/WPS 中段末回车/占位空段不算正文)
            if not t:
                continue
            decl_body_paras.append(j)

    if not decl_body_paras:
        decl_body_ok = False
        decl_body_reason.append('未找到标题下方正文段落')
    else:
        for j in decl_body_paras:
            pi = model.pinfo[j]
            # (1a) 字体 = 宋体: 逐 run 校验
            for ri, rinfo in enumerate(model.run_info[j]):
                rtext = (model.paragraphs[j].runs[ri].text or '')
                if not rtext.strip():
                    continue
                if rinfo['eastAsia'] != '宋体':
                    decl_body_ok = False
                    decl_body_reason.append(
                        f'段{j} run[{ri}] 中文字体非宋体: {rinfo["eastAsia"]}'
                    )
                # (1b) 字号 = 四号 (14pt)
                sz = rinfo['size']
                if sz is None or abs(sz - 14) > 0.5:
                    decl_body_ok = False
                    decl_body_reason.append(
                        f'段{j} run[{ri}] 字号非四号(14pt): {sz}'
                    )

            # (2) 首行缩进 2 字符: 严格 firstLineChars = 200 (±10 = ±0.1 字符 容差)
            if pi['indent_chars'] is None:
                decl_body_ok = False
                decl_body_reason.append(f'段{j} 未使用"字符"单位首行缩进')
            elif abs(pi['indent_chars'] - 2.0) > 0.1:
                decl_body_ok = False
                decl_body_reason.append(
                    f'段{j} 首行缩进非 2 字符: {pi["indent_chars"]}'
                )

            # (3) 1.5 倍行距
            ls_val = line_spacing_value(pi)
            if abs(ls_val - 1.5) > 0.05:
                decl_body_ok = False
                decl_body_reason.append(
                    f'段{j} 非 1.5 倍行距: rule={pi["line_spacing_rule"]}, val={pi["line_spacing"]}'
                )

    evidence = f'正文段={decl_body_paras}'
    if not decl_body_ok:
        evidence += ' | 问题: ' + '; '.join(decl_body_reason[:8])
        if len(decl_body_reason) > 8:
            evidence += f' ... 共 {len(decl_body_reason)} 条'
    add('+1 声明标题下方文本(宋体四号 / 首行缩进2字符 / 1.5倍行距)',
        1, evidence, decl_body_ok)

    # -------------------------------------------------------------------------
    # +1 原创性声明及使用授权声明页签名区: 严格按细则四点判定
    #   (1) 页面下方包含 "论文作者签名：" 和 "日 期：" 两项
    #   (2) 两项文本字体为宋体、小四 (12pt)
    #   (3) 两项位于同一行 (即位于同一视觉行)
    #   (4) 两项后出现可填写的横线
    # 兼容 Word/WPS:
    #   - 字体判定用 w:rFonts@w:eastAsia (中文渲染字段);
    #   - 横线判定接受两种常见做法: 下划线制表符 (tab+w:u) / 下划线空白字符;
    #     两软件对 <w:u> 的渲染一致
    # -------------------------------------------------------------------------
    sig_ok = True
    sig_reason = []

    # 定位签名区段: 同段同时含 "论文作者签名" 与 "日" + "期"
    sig_idx = None
    for _i, _p in enumerate(model.paragraphs):
        _t = _p.text
        if '论文作者签名' in _t and '日' in _t and '期' in _t:
            sig_idx = _i
            break

    if sig_idx is None:
        sig_ok = False
        sig_reason.append('未找到含 "论文作者签名" 与 "日 期" 的同段')
    else:
        para = model.paragraphs[sig_idx]
        # 用去空白后的字符串定位两个标签在 run 序列中的位置
        # (Word/WPS 中标签可能与冒号在同一 run 或分开 run)
        text = para.text

        # (1) 页面下方: 需位于原创性声明页的下半部分区域，而非仅仅在标题之后
        #     （标题之后到签名区之间可能只占页面很小比例，不代表处于"页面下方"）
        decl_end_idx = len(model.paragraphs)
        if decl_idx is not None:
            for _i in range(decl_idx + 1, len(model.paragraphs)):
                if model.page_break_before[_i]:
                    decl_end_idx = _i
                    break
        if decl_idx is not None:
            decl_len = max(1, decl_end_idx - decl_idx)
            sig_pos_ratio = (sig_idx - decl_idx) / decl_len
            if sig_pos_ratio < 0.55:
                sig_ok = False
                sig_reason.append(
                    f'签名区未位于页面下方 (段位 {sig_idx - decl_idx}/{decl_len}={sig_pos_ratio:.2f}, 应 >= 0.55)'
                )

        # (3) 两项位于同一视觉行 → 同一段 + 段内不含 textWrapping 换行符
        has_line_break = False
        for br in para._element.iter(qn('w:br')):
            btype = br.get(qn('w:type')) or 'textWrapping'
            if btype == 'textWrapping':
                has_line_break = True
                break
        if has_line_break:
            sig_ok = False
            sig_reason.append('签名区段内含软换行 <w:br/>, 两项未在同一行')

        # (2) 字体: 定位两个标签 run 并逐 run 校验
        #     标签 run 判定: run 文本 strip 后为 "论文作者签名" / "论文作者签名："
        #                        或 "日" / "日  期" / "日  期：" 等含标签字样且非空白
        label_run_indices = {'签名': None, '日期': None}
        for ri, r in enumerate(para.runs):
            rtext = r.text or ''
            stripped = rtext.strip()
            if not stripped:
                continue
            if label_run_indices['签名'] is None and '论文作者签名' in stripped:
                label_run_indices['签名'] = ri
            # 日期标签: run 含 "日" 且含 "期" (强匹配), 或仅含 "日"/"期" 且相邻 run 拼合
            if label_run_indices['日期'] is None and (
                ('日' in stripped and '期' in stripped)
                or stripped in ('日', '期', '日期')
            ):
                label_run_indices['日期'] = ri

        if label_run_indices['签名'] is None:
            sig_ok = False
            sig_reason.append('未定位到 "论文作者签名" 标签 run')
        if label_run_indices['日期'] is None:
            sig_ok = False
            sig_reason.append('未定位到 "日 期" 标签 run')

        for name, ri in label_run_indices.items():
            if ri is None:
                continue
            rinfo = model.run_info[sig_idx][ri]
            if rinfo['eastAsia'] != '宋体':
                sig_ok = False
                sig_reason.append(f'{name}标签中文字体非宋体: {rinfo["eastAsia"]}')
            sz = rinfo['size']
            if sz is None or abs(sz - 12) > 0.5:
                sig_ok = False
                sig_reason.append(f'{name}标签字号非小四(12pt): {sz}')

        # (4) 两项后出现可填写的横线
        #     可填写的横线 = 标签 run 之后, 下一个标签 run 之前, 至少存在一个 underline run
        def _has_underline_after(start_ri: int, end_ri: int) -> bool:
            """检查 [start_ri+1, end_ri) 区间内是否存在含下划线的 run"""
            for k in range(start_ri + 1, end_ri):
                r = para.runs[k]
                rPr = r._element.find(qn('w:rPr'))
                if rPr is None:
                    continue
                u = rPr.find(qn('w:u'))
                if u is not None:
                    val = u.get(qn('w:val'))
                    if val is None or val != 'none':
                        return True
            return False

        n_runs = len(para.runs)
        sig_ri = label_run_indices['签名']
        date_ri = label_run_indices['日期']
        if sig_ri is not None and date_ri is not None:
            # 签名后横线区间: 签名 run 后 → 日期 run 前
            if not _has_underline_after(sig_ri, date_ri):
                sig_ok = False
                sig_reason.append('"论文作者签名：" 后无可填写的横线')
            # 日期后横线区间: 日期 run 后 → 段末
            if not _has_underline_after(date_ri, n_runs):
                sig_ok = False
                sig_reason.append('"日 期：" 后无可填写的横线')

    evidence = f'签名段@{sig_idx}'
    if not sig_ok:
        evidence += ' | 问题: ' + '; '.join(sig_reason)
    add('+1 签名区(页面下方/含两项标签/宋体小四/同一行/两项后均有可填写横线)',
        1, evidence, sig_ok)

    # -------------------------------------------------------------------------
    # +1 原创性声明及使用授权声明页的后一页出现目录页: 严格按细则一点判定
    #   要求: 存在"目录"页 (以"目录"为标题的独立页), 且该页 = 声明页的下一页
    # 兼容 Word/WPS:
    #   - "下一页"通过 model.page_break_before 判定 —— 该字段同时识别
    #     <w:br w:type="page"/> 与 <w:sectPr>, 与 Word/WPS 分页行为一致
    #   - "目录"标题识别: 去除全/半角空白后精确等于 "目录"
    # -------------------------------------------------------------------------
    catalog_idx = None
    for _i, _p in enumerate(model.paragraphs):
        _collapsed = _p.text.replace('　', '').replace(' ', '').replace('\t', '').strip()
        if _collapsed == '目录':
            catalog_idx = _i
            break

    catalog_ok = True
    catalog_reason = []
    catalog_page = None

    if catalog_idx is None:
        catalog_ok = False
        catalog_reason.append('未找到 "目录" 标题段')
    else:
        # 计算目录所在页码
        catalog_page = 1 + sum(
            1 for _i in range(0, catalog_idx + 1) if model.page_break_before[_i]
        )
        # 声明所在页码 (由上文 decl_page 提供)
        if decl_page is None:
            catalog_ok = False
            catalog_reason.append('原创性声明页未定位, 无法判定"后一页"')
        elif catalog_page != decl_page + 1:
            catalog_ok = False
            catalog_reason.append(
                f'目录页非声明页的后一页 (目录=第{catalog_page}页, 声明=第{decl_page}页)'
            )

    evidence = f'目录@{catalog_idx}, 目录页={catalog_page}, 声明页={decl_page}'
    if not catalog_ok:
        evidence += ' | 问题: ' + '; '.join(catalog_reason)
    add('+1 原创性声明及使用授权声明页的后一页出现目录页',
        1, evidence, catalog_ok)

    # -------------------------------------------------------------------------
    # +1 目录页 "目录" 标题: 严格按细则六点判定
    #   (1) 位于目录页页面顶部 (即目录页首段)
    #   (2) 字体 = 宋体
    #   (3) 字号 = 二号 (22pt)
    #   (4) 不加粗
    #   (5) 居中显示
    #   (6) 段后间距 10 磅 + 2 倍行距
    # 兼容 Word/WPS:
    #   - 段后间距: w:spacing@w:after (twips, 20 twips = 1pt → 200 = 10pt);
    #     若 w:afterLines 存在则按行数换算, 但细则要求"磅"单位, 只查 w:after
    #   - 2 倍行距: w:spacing@w:line=480, w:lineRule="auto" (240 = 1 倍)
    # -------------------------------------------------------------------------
    cat_title_ok = True
    cat_title_reason = []

    if catalog_idx is None:
        cat_title_ok = False
        cat_title_reason.append('未找到 "目录" 标题段')
    else:
        pi = model.pinfo[catalog_idx]

        # (1) 位于目录页顶部: catalog_idx 处应有 page_break_before
        if not model.page_break_before[catalog_idx]:
            cat_title_ok = False
            cat_title_reason.append('"目录" 标题未位于新页顶部')

        # (2)(3)(4) 逐 run 校验字体/字号/不加粗
        for ri, rinfo in enumerate(model.run_info[catalog_idx]):
            rtext = (model.paragraphs[catalog_idx].runs[ri].text or '')
            if not rtext.strip():
                continue
            if rinfo['eastAsia'] != '宋体':
                cat_title_ok = False
                cat_title_reason.append(
                    f'run[{ri}] {rtext!r} 中文字体非宋体: {rinfo["eastAsia"]}'
                )
            sz = rinfo['size']
            if sz is None or abs(sz - 22) > 0.5:
                cat_title_ok = False
                cat_title_reason.append(
                    f'run[{ri}] {rtext!r} 字号非二号(22pt): {sz}'
                )
            if rinfo['bold']:
                cat_title_ok = False
                cat_title_reason.append(f'run[{ri}] {rtext!r} 加粗了(应不加粗)')

        # (5) 居中
        if pi['align'] != WD_ALIGN_PARAGRAPH.CENTER:
            cat_title_ok = False
            cat_title_reason.append(f'目录标题未居中: align={pi["align"]}')

        # (6a) 段后间距 = 10 磅 (200 twips)
        pPr = model.paragraphs[catalog_idx]._element.find(qn('w:pPr'))
        after_twips = None
        after_lines = None
        if pPr is not None:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                a = sp.get(qn('w:after'))
                if a is not None:
                    try:
                        after_twips = int(a)
                    except ValueError:
                        pass
                al = sp.get(qn('w:afterLines'))
                if al is not None:
                    try:
                        after_lines = int(al)
                    except ValueError:
                        pass
        if after_twips is None:
            cat_title_ok = False
            cat_title_reason.append('段后间距未设置 (w:spacing@after 缺失)')
        else:
            after_pt = after_twips / 20.0
            if abs(after_pt - 10.0) > 0.5:
                cat_title_ok = False
                cat_title_reason.append(f'段后间距非 10 磅: 实际 {after_pt}pt')

        # (6b) 2 倍行距
        ls_val = line_spacing_value(pi)
        if abs(ls_val - 2.0) > 0.05:
            cat_title_ok = False
            cat_title_reason.append(
                f'非 2 倍行距: rule={pi["line_spacing_rule"]}, val={pi["line_spacing"]}'
            )

    evidence = f'"目录" 标题@{catalog_idx}'
    if not cat_title_ok:
        evidence += ' | 问题: ' + '; '.join(cat_title_reason)
    add('+1 目录页"目录"标题(顶部/宋体二号/不加粗/居中/段后10磅/2倍行距)',
        1, evidence, cat_title_ok)

    # -------------------------------------------------------------------------
    # +3 目录页标题下方项目文本: 严格按细则四点判定
    #   (1) 一级标题: 黑体、三号 (16pt)
    #       (识别: 以中文数字 "一/二/…" 或 "第X章" 开头)
    #   (2) 二级标题: 黑体、小四 (12pt)
    #       (识别: 形如 "X.Y" 开头, 如 "1.1")
    #   (3) 三级标题: 黑体、小四 (12pt)
    #       (识别: 形如 "X.Y.Z" 开头, 如 "1.1.1")
    #   (4) 所有项目无缩进 + 整体两端对齐 (左右两边都要对齐)
    # 兼容 Word/WPS:
    #   - 无缩进: w:ind@w:firstLine / firstLineChars / left / start 全为 0 或不存在
    #   - 两端对齐: 段落 <w:jc w:val="both"/> → WD_ALIGN_PARAGRAPH.JUSTIFY
    # -------------------------------------------------------------------------
    cat_items_ok = True
    cat_items_reason = []

    if catalog_idx is None:
        cat_items_ok = False
        cat_items_reason.append('无 "目录" 标题, 无目录项可校验')
        items = []
    else:
        # 目录项范围: "目录" 标题之后, 到下一个显式分页或摘要页之前
        end_idx = None
        for j in range(catalog_idx + 1, len(model.paragraphs)):
            if model.page_break_before[j]:
                end_idx = j
                break
        if end_idx is None:
            end_idx = len(model.paragraphs)
        items = []
        for j in range(catalog_idx + 1, end_idx):
            t = model.paragraphs[j].text.strip()
            if t:
                items.append((j, t))
        if not items:
            cat_items_ok = False
            cat_items_reason.append('未发现目录项目段落')

    # 层级识别: 先判 "X.Y.Z", 再判 "X.Y", 再判 "一/二/…" 或 "第X章"
    for j, t in items:
        # 层级
        if re.match(r'^\d+\.\d+\.\d+', t):
            level = 3
            expect_pt = 12  # 小四
        elif re.match(r'^\d+\.\d+', t):
            level = 2
            expect_pt = 12  # 小四
        elif re.match(r'^([一二三四五六七八九十]+[、\s]|第[一二三四五六七八九十]+章)', t):
            level = 1
            expect_pt = 16  # 三号
        else:
            # 非标准层级 → 只按细则要求的 "黑体 + 无缩进 + 两端对齐" 校验字体,
            # 字号不强制 (细则没有对非一/二/三级项目的字号约束)
            level = None
            expect_pt = None

        # 字体 = 黑体 (逐 run)
        for ri, rinfo in enumerate(model.run_info[j]):
            rtext = (model.paragraphs[j].runs[ri].text or '')
            if not rtext.strip():
                continue
            if rinfo['eastAsia'] != '黑体':
                cat_items_ok = False
                cat_items_reason.append(
                    f'{level or "?"}级项 "{t[:10]}" run[{ri}] 中文非黑体: {rinfo["eastAsia"]}'
                )
            if expect_pt is not None:
                sz = rinfo['size']
                if sz is None or abs(sz - expect_pt) > 0.5:
                    lvl_name = {1: '一级', 2: '二级', 3: '三级'}[level]
                    exp_name = {16: '三号(16pt)', 12: '小四(12pt)'}[expect_pt]
                    cat_items_ok = False
                    cat_items_reason.append(
                        f'{lvl_name}项 "{t[:10]}" run[{ri}] 字号非{exp_name}: {sz}'
                    )

        # 无缩进: firstLine / firstLineChars / left / start 均需为 0 或缺失
        pPr = model.paragraphs[j]._element.find(qn('w:pPr'))
        has_indent = False
        indent_desc = ''
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                for attr in ('w:firstLine', 'w:firstLineChars', 'w:left', 'w:start',
                             'w:leftChars', 'w:startChars'):
                    v = ind.get(qn(attr))
                    if v not in (None, '0'):
                        try:
                            if int(v) != 0:
                                has_indent = True
                                indent_desc = f'{attr}={v}'
                                break
                        except ValueError:
                            has_indent = True
                            indent_desc = f'{attr}={v}'
                            break
        if has_indent:
            cat_items_ok = False
            cat_items_reason.append(f'目录项 "{t[:10]}" 有缩进: {indent_desc}')

        # 两端对齐
        align = model.pinfo[j]['align']
        if align != WD_ALIGN_PARAGRAPH.JUSTIFY:
            cat_items_ok = False
            cat_items_reason.append(f'目录项 "{t[:10]}" 未两端对齐: align={align}')

    evidence = f'目录项数={len(items)}'
    if not cat_items_ok:
        evidence += ' | 问题: ' + '; '.join(cat_items_reason[:8])
        if len(cat_items_reason) > 8:
            evidence += f' ... 共 {len(cat_items_reason)} 条'
    add('+3 目录项(一级黑体三号/二三级黑体小四/无缩进/两端对齐)',
        3, evidence, cat_items_ok)

    # -------------------------------------------------------------------------
    # +1 摘要页单独成页: 严格按细则六点判定
    #   (1) 摘要页单独成页 (页顶开始 = 前置分页)
    #   (2) 页面顶部出现论文题目 (= 摘要页第一段, 文本完全匹配)
    #   (3) 字体: 黑体
    #   (4) 字号: 小二 (18pt)
    #   (5) 加粗
    #   (6) 居中显示
    # 兼容 Word/WPS:
    #   - 单独成页 / 顶部: 该段落有 <w:br w:type="page"/> 或位于 <w:sectPr> 分节起始
    #   - 黑体: w:rFonts@w:eastAsia = "黑体"
    #   - 小二: w:sz = 36 半点 (18pt)
    #   - 加粗: w:b 存在且非 val="0"/"false"
    #   - 居中: w:jc@w:val = "center" → WD_ALIGN_PARAGRAPH.CENTER
    # -------------------------------------------------------------------------
    THESIS_TITLE = '小微企业财务数据管理规范化问题及改进路径研究'
    abs_idx = model.first_para_idx('摘  要')
    if abs_idx is None:
        abs_idx = model.first_para_idx('摘要')

    # 定位摘要页论文题目: 文本完全匹配 (允许空格差异)
    abs_title_idx = None
    for i, p in enumerate(model.paragraphs):
        t = p.text.strip()
        if t.replace(' ', '').replace('　', '') == THESIS_TITLE:
            # 需在摘要标签之前 (若已定位到摘要标签段)
            if abs_idx is None or i < abs_idx:
                abs_title_idx = i
                break

    abs_ok = True
    abs_reason = []
    if abs_title_idx is None:
        abs_ok = False
        abs_reason.append(f'未找到论文题目 "{THESIS_TITLE}"')
    else:
        # (1) 单独成页: 该段前有分页 (page break 或 sectPr)
        if not model.page_break_before[abs_title_idx]:
            abs_ok = False
            abs_reason.append('摘要页题目段前无分页, 未单独成页')

        # (2) 位于摘要页顶部: 该段是"新页第一段" (由 page_break_before 保证)
        # 已由 (1) 覆盖 —— 摘要题若前置分页, 即位于页面顶部

        pi = model.pinfo[abs_title_idx]

        # (3)(4)(5) 逐 run 校验字体/字号/加粗
        runs = model.paragraphs[abs_title_idx].runs
        rinfos = model.run_info[abs_title_idx]
        checked_any = False
        for ri, rinfo in enumerate(rinfos):
            rtext = (runs[ri].text or '')
            if not rtext.strip():
                continue
            checked_any = True
            # (3) 黑体
            if rinfo['eastAsia'] != '黑体':
                abs_ok = False
                abs_reason.append(f'run[{ri}] 非黑体: {rinfo["eastAsia"]}')
            # (4) 小二 (18pt)
            sz = rinfo['size']
            if sz is None or abs(sz - 18) > 0.5:
                abs_ok = False
                abs_reason.append(f'run[{ri}] 非小二(18pt): {sz}')
            # (5) 加粗
            if not rinfo['bold']:
                abs_ok = False
                abs_reason.append(f'run[{ri}] 未加粗')
        if not checked_any:
            abs_ok = False
            abs_reason.append('摘要题段无可校验的文本 run')

        # (6) 居中
        if pi['align'] != WD_ALIGN_PARAGRAPH.CENTER:
            abs_ok = False
            abs_reason.append(f'摘要题未居中: align={pi["align"]}')

    evidence = f'摘要题@{abs_title_idx}'
    if not abs_ok:
        evidence += ' | 问题: ' + '; '.join(abs_reason)
    add('+1 摘要页单独成页(论文题目/黑体/小二/加粗/居中)',
        1, evidence, abs_ok)

    # -------------------------------------------------------------------------
    # +1 摘要页论文题目下方空一行出现 "摘  要:" 及后面的文本内容:
    #   (1) 论文题目下方 "空一行" 后出现 "摘  要:" 段
    #   (2) "摘  要:" 字体: 黑体
    #   (3) "摘  要:" 字号: 四号 (14pt)
    #   (4) "摘  要:" 加粗
    #   (5) "摘  要:" 后面的文本内容: 宋体
    #   (6) "摘  要:" 后面的文本内容: 小四 (12pt)
    #   (7) "摘  要:" 后面的文本内容: 1.5 倍行距
    #   (8) "摘  要:" 后面的文本内容: 两端对齐
    # 兼容 Word/WPS:
    #   - "空一行" = 题目段与 "摘要" 段之间恰有一个空文本段
    #   - 黑体/宋体: w:rFonts@w:eastAsia
    #   - 四号/小四: w:sz = 28/24 半点
    #   - 加粗: w:b
    #   - 1.5 倍行距: w:spacing@w:line=360 w:lineRule="auto"
    #   - 两端对齐: w:jc="both" → WD_ALIGN_PARAGRAPH.JUSTIFY
    # -------------------------------------------------------------------------
    abs_body_ok = True
    abs_body_reason = []

    if abs_idx is None:
        abs_body_ok = False
        abs_body_reason.append('未找到 "摘  要:" 段')
    elif abs_title_idx is None:
        abs_body_ok = False
        abs_body_reason.append('未找到摘要页论文题目段, 无法验证 "空一行" 关系')
    else:
        # (1) 论文题目下方空一行: 题目段与 "摘要" 段之间恰有一个空段
        blank_between = []
        for k in range(abs_title_idx + 1, abs_idx):
            if not model.paragraphs[k].text.strip():
                blank_between.append(k)
            else:
                blank_between.append(-1)  # 非空段
                break
        non_blank_between = [x for x in blank_between if x == -1]
        blank_count = sum(1 for x in blank_between if x != -1)
        if abs_idx - abs_title_idx - 1 != 1 or blank_count != 1 or non_blank_between:
            abs_body_ok = False
            abs_body_reason.append(
                f'论文题目@{abs_title_idx} 与 "摘要"@{abs_idx} 之间未恰好空一行'
                f' (间隔段数={abs_idx - abs_title_idx - 1}, 空段数={blank_count})'
            )

        # (2)(3)(4) "摘  要:" 标签: 逐 run 定位含 "摘" 的首个 run
        runs = model.paragraphs[abs_idx].runs
        rinfos = model.run_info[abs_idx]
        label_end_run = 0  # 标签结束后的 run 起始索引
        label_found = False
        for ri, rinfo in enumerate(rinfos):
            rtext = (runs[ri].text or '')
            if not rtext.strip():
                continue
            if not label_found:
                # 首个非空 run 应是 "摘  要:" 标签
                if '摘' not in rtext and '要' not in rtext:
                    abs_body_ok = False
                    abs_body_reason.append(f'段首非 "摘  要:" 标签: "{rtext[:8]}"')
                # (2) 黑体
                if rinfo['eastAsia'] != '黑体':
                    abs_body_ok = False
                    abs_body_reason.append(f'"摘  要:" 非黑体: {rinfo["eastAsia"]}')
                # (3) 四号 (14pt)
                sz = rinfo['size']
                if sz is None or abs(sz - 14) > 0.5:
                    abs_body_ok = False
                    abs_body_reason.append(f'"摘  要:" 非四号(14pt): {sz}')
                # (4) 加粗
                if not rinfo['bold']:
                    abs_body_ok = False
                    abs_body_reason.append('"摘  要:" 未加粗')
                label_found = True
                # 若该 run 中 ":" 或 ":" 已闭合, 则标签结束在此 run
                if '：' in rtext or ':' in rtext:
                    label_end_run = ri + 1
                    break
                else:
                    label_end_run = ri + 1
                    continue
            else:
                # 若前一 run 未闭合冒号, 顺延到含 ":" 的 run
                if '：' in rtext or ':' in rtext:
                    label_end_run = ri + 1
                    break
                label_end_run = ri + 1
        if not label_found:
            abs_body_ok = False
            abs_body_reason.append('"摘  要:" 段无标签 run')

        # (5)(6) "摘  要:" 后面的文本内容: 首先同段 label_end_run 之后的 run, 再看后续段落直到关键词段
        # 收集 (段索引, run 索引, run 信息) 三元组
        tail_targets = []
        # 同段末尾
        for ri in range(label_end_run, len(rinfos)):
            rtext = (runs[ri].text or '')
            if rtext.strip():
                tail_targets.append((abs_idx, ri, rinfos[ri]))
        # 后续段落 (直到关键词段或下一分页)
        kw_probe_idx = model.first_para_idx('关键词')
        end_probe = kw_probe_idx if kw_probe_idx is not None else min(abs_idx + 20, len(model.paragraphs))
        for k in range(abs_idx + 1, end_probe):
            if model.page_break_before[k]:
                break
            k_runs = model.paragraphs[k].runs
            k_rinfos = model.run_info[k]
            for ri, rinfo in enumerate(k_rinfos):
                rtext = (k_runs[ri].text or '')
                if rtext.strip():
                    tail_targets.append((k, ri, rinfo))

        if not tail_targets:
            abs_body_ok = False
            abs_body_reason.append('"摘  要:" 后无文本内容')
        else:
            for pidx, ri, rinfo in tail_targets:
                # (5) 宋体
                if rinfo['eastAsia'] != '宋体':
                    abs_body_ok = False
                    abs_body_reason.append(
                        f'摘要正文@段{pidx} run[{ri}] 非宋体: {rinfo["eastAsia"]}'
                    )
                # (6) 小四 (12pt)
                sz = rinfo['size']
                if sz is None or abs(sz - 12) > 0.5:
                    abs_body_ok = False
                    abs_body_reason.append(
                        f'摘要正文@段{pidx} run[{ri}] 非小四(12pt): {sz}'
                    )

            # (7)(8) 逐段校验 1.5 倍行距 + 两端对齐
            body_paras = sorted(set(pidx for pidx, _, _ in tail_targets))
            for pidx in body_paras:
                pi = model.pinfo[pidx]
                if abs(line_spacing_value(pi) - 1.5) > 0.05:
                    abs_body_ok = False
                    abs_body_reason.append(
                        f'摘要正文@段{pidx} 非1.5倍行距: {pi["line_spacing"]}'
                    )
                if pi['align'] != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    abs_body_ok = False
                    abs_body_reason.append(
                        f'摘要正文@段{pidx} 未两端对齐: {pi["align"]}'
                    )

    evidence = f'"摘要"@{abs_idx}'
    if not abs_body_ok:
        evidence += ' | 问题: ' + '; '.join(abs_body_reason[:8])
        if len(abs_body_reason) > 8:
            evidence += f' ... 共 {len(abs_body_reason)} 条'
    else:
        evidence += ' 摘要正文正确'
    add('+1 摘要页("摘 要:"黑体四号加粗 / 后文宋体小四1.5倍两端对齐 / 题目下空一行)',
        1, evidence, abs_body_ok)

    # -------------------------------------------------------------------------
    # +1 摘要页关键词部分: 严格按细则四点判定
    #   (1) "关键词:" 字体: 黑体
    #   (2) "关键词:" 字号: 四号 (14pt)
    #   (3) "关键词:" 加粗
    #   (4) "关键词:" 后面的内容: 宋体
    #   (5) "关键词:" 后面的内容: 小四 (12pt)
    #   (6) 每一个词之间用全角分号相隔
    #   (7) 最后一个词没有标点符号
    # 兼容 Word/WPS:
    #   - 黑体/宋体: w:rFonts@w:eastAsia
    #   - 四号/小四: w:sz = 28/24 半点
    #   - 加粗: w:b
    #   - 全角分号 "；" (U+FF1B) 是 Word/WPS 中文输入法默认写入的字符
    # -------------------------------------------------------------------------
    kw_idx = model.first_para_idx('关键词')
    kw_ok = True
    kw_reason = []

    if kw_idx is None:
        kw_ok = False
        kw_reason.append('未找到 "关键词" 段')
    else:
        runs = model.paragraphs[kw_idx].runs
        rinfos = model.run_info[kw_idx]

        # 定位 "关键词:" 标签 (至含冒号的 run)
        label_end_run = None
        label_found = False
        for ri, rinfo in enumerate(rinfos):
            rtext = (runs[ri].text or '')
            if not rtext.strip():
                continue
            if not label_found:
                # 首个非空 run 应为标签起始 (含 "关" 或 "关键词")
                if '关' not in rtext and '词' not in rtext:
                    kw_ok = False
                    kw_reason.append(f'段首非 "关键词:" 标签: "{rtext[:8]}"')
                # (1) 黑体
                if rinfo['eastAsia'] != '黑体':
                    kw_ok = False
                    kw_reason.append(f'"关键词:" 非黑体: {rinfo["eastAsia"]}')
                # (2) 四号 (14pt)
                sz = rinfo['size']
                if sz is None or abs(sz - 14) > 0.5:
                    kw_ok = False
                    kw_reason.append(f'"关键词:" 非四号(14pt): {sz}')
                # (3) 加粗
                if not rinfo['bold']:
                    kw_ok = False
                    kw_reason.append('"关键词:" 未加粗')
                label_found = True
                if '：' in rtext or ':' in rtext:
                    label_end_run = ri + 1
                    break
                else:
                    label_end_run = ri + 1
                    continue
            else:
                if '：' in rtext or ':' in rtext:
                    label_end_run = ri + 1
                    break
                label_end_run = ri + 1
        if not label_found:
            kw_ok = False
            kw_reason.append('关键词段无标签 run')
            label_end_run = 0

        # (4)(5) 标签之后的 run 逐个校验宋体 + 小四
        content_run_count = 0
        for ri in range(label_end_run, len(rinfos)):
            rtext = (runs[ri].text or '')
            if not rtext.strip():
                continue
            content_run_count += 1
            rinfo = rinfos[ri]
            # (4) 宋体
            if rinfo['eastAsia'] != '宋体':
                kw_ok = False
                kw_reason.append(f'关键词内容 run[{ri}] 非宋体: {rinfo["eastAsia"]}')
            # (5) 小四 (12pt)
            sz = rinfo['size']
            if sz is None or abs(sz - 12) > 0.5:
                kw_ok = False
                kw_reason.append(f'关键词内容 run[{ri}] 非小四(12pt): {sz}')
        if content_run_count == 0:
            kw_ok = False
            kw_reason.append('"关键词:" 后无内容 run')

        # (6)(7) 从段落文本中截取标签之后的内容, 校验分隔符和末项标点
        full_text = model.paragraphs[kw_idx].text
        # 优先按全角冒号截, 再按半角冒号截
        m = re.search(r'关键词\s*[：:]\s*(.*)$', full_text)
        content_text = m.group(1).strip() if m else ''
        if not content_text:
            kw_ok = False
            kw_reason.append('"关键词:" 冒号后无文本')
        else:
            # (6) 每个词之间用全角分号 "；" 相隔
            # 若只有一个词, 无需分号 → 视为不符合 "每一个词之间用全角分号相隔" 隐含前提
            # 但更精确的判定: 只要存在多个词, 分隔符必须是 "；"; 不允许半角 ";"、"、"、"," 等
            forbidden_seps = [';', '、', ',', ',']
            for sep in forbidden_seps:
                if sep in content_text:
                    kw_ok = False
                    kw_reason.append(f'关键词间出现非全角分号分隔符: "{sep}"')
            words = [w.strip() for w in content_text.split('；')]
            words = [w for w in words if w != '']
            if len(words) < 2:
                kw_ok = False
                kw_reason.append(f'关键词数量不足或未用全角分号相隔: 词数={len(words)}')

            # (7) 最后一个词没有标点符号
            if words:
                last = words[-1]
                # 中英文常见标点集
                punct = set('，。；！？、,.;!?：:"""''《》()（）[]【】')
                if last and last[-1] in punct:
                    kw_ok = False
                    kw_reason.append(f'最后一个关键词有标点: "{last[-1]}"')

    evidence = f'"关键词"@{kw_idx}'
    if not kw_ok:
        evidence += ' | 问题: ' + '; '.join(kw_reason[:8])
        if len(kw_reason) > 8:
            evidence += f' ... 共 {len(kw_reason)} 条'
    else:
        evidence += ' 关键词正确'
    add('+1 关键词("关键词:"黑体四号加粗 / 后文宋体小四 / 全角分号相隔 / 末词无标点)',
        1, evidence, kw_ok)

    # -------------------------------------------------------------------------
    # +1 文中所有一级标题: 严格按细则四点判定
    #   (1) 字体: 黑体
    #   (2) 字号: 三号 (16pt)
    #   (3) 居中
    #   (4) 大纲级别: 一级 (OOXML outlineLvl=0)
    # 兼容 Word/WPS:
    #   - 大纲级别: w:pPr/w:outlineLvl@w:val="0" 或应用 Heading 1 样式
    #     (Word/WPS UI 中"段落 → 大纲级别 → 1 级" 写入该属性;
    #      "样式 → 标题 1" 通过样式继承该属性)
    #   - 黑体: w:rFonts@w:eastAsia
    #   - 三号: w:sz = 32 半点 (16pt)
    #   - 居中: w:jc="center" → WD_ALIGN_PARAGRAPH.CENTER
    # -------------------------------------------------------------------------
    def _paragraph_outline_level(p):
        """从段落 pPr/outlineLvl 或样式链读出大纲级别 (0 = 一级); 无返回 None"""
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            lvl = pPr.find(qn('w:outlineLvl'))
            if lvl is not None:
                v = lvl.get(qn('w:val'))
                if v is not None:
                    try:
                        return int(v)
                    except ValueError:
                        return None
        # 若段落 XML 未显式指定, 尝试从样式链继承
        try:
            style = p.style
            while style is not None:
                sxml = style.element
                spPr = sxml.find(qn('w:pPr'))
                if spPr is not None:
                    lvl = spPr.find(qn('w:outlineLvl'))
                    if lvl is not None:
                        v = lvl.get(qn('w:val'))
                        if v is not None:
                            try:
                                return int(v)
                            except ValueError:
                                pass
                style = style.base_style
        except Exception:
            pass
        return None

    # 候选一级标题段: 样式=Heading 1 或 outlineLvl=0
    h1_paras = []
    for i, p in enumerate(model.paragraphs):
        if p.style.name == 'Heading 1':
            h1_paras.append(i)
        elif _paragraph_outline_level(p) == 0 and p.text.strip():
            h1_paras.append(i)

    h1_ok = True
    h1_reason = []
    if not h1_paras:
        h1_ok = False
        h1_reason.append('未发现一级标题段落 (Heading 1 或 outlineLvl=0)')
    else:
        for i in h1_paras:
            p = model.paragraphs[i]
            text_head = p.text[:12]

            # (4) 大纲级别 = 一级 (0)
            lvl = _paragraph_outline_level(p)
            if lvl != 0:
                h1_ok = False
                h1_reason.append(f'H1[{i}] "{text_head}" 大纲级别非一级: {lvl}')

            # (1)(2) 字体 / 字号 逐 run 校验
            runs = p.runs
            rinfos = model.run_info[i]
            checked_any = False
            for ri, rinfo in enumerate(rinfos):
                rtext = (runs[ri].text or '')
                if not rtext.strip():
                    continue
                checked_any = True
                if rinfo['eastAsia'] != '黑体':
                    h1_ok = False
                    h1_reason.append(
                        f'H1[{i}] "{text_head}" run[{ri}] 非黑体: {rinfo["eastAsia"]}'
                    )
                sz = rinfo['size']
                if sz is None or abs(sz - 16) > 0.5:
                    h1_ok = False
                    h1_reason.append(
                        f'H1[{i}] "{text_head}" run[{ri}] 非三号(16pt): {sz}'
                    )
            if not checked_any:
                h1_ok = False
                h1_reason.append(f'H1[{i}] "{text_head}" 段无可校验的文本 run')

            # (3) 居中
            pi = model.pinfo[i]
            if pi['align'] != WD_ALIGN_PARAGRAPH.CENTER:
                h1_ok = False
                h1_reason.append(f'H1[{i}] "{text_head}" 未居中: align={pi["align"]}')

    evidence = f'{len(h1_paras)} 个一级标题'
    if not h1_ok:
        evidence += ' | 问题: ' + '; '.join(h1_reason[:8])
        if len(h1_reason) > 8:
            evidence += f' ... 共 {len(h1_reason)} 条'
    else:
        evidence += ' 全部正确'
    add('+1 一级标题(黑体/三号/居中/大纲级别一级)',
        1, evidence, h1_ok)

    # -------------------------------------------------------------------------
    # -------------------------------------------------------------------------
    # +1 文中所有二级标题: 严格按细则三点判定
    #   (1) 字体: 黑体
    #   (2) 字号: 四号 (14pt)
    #   (3) 左对齐
    # 兼容 Word/WPS:
    #   - 黑体: w:rFonts@w:eastAsia
    #   - 四号: w:sz = 28 半点 (14pt)
    #   - 左对齐: w:jc="left" 或 无 <w:jc> 元素 (Word/WPS 默认即左对齐)
    # 二级标题识别 (非约束项, 仅用于定位):
    #   样式 Heading 2 或 段落 outlineLvl=1
    # -------------------------------------------------------------------------
    h2_paras = []
    for i, p in enumerate(model.paragraphs):
        if p.style.name == 'Heading 2':
            h2_paras.append(i)
        elif _paragraph_outline_level(p) == 1 and p.text.strip():
            h2_paras.append(i)

    h2_ok = True
    h2_reason = []
    if not h2_paras:
        h2_ok = False
        h2_reason.append('未发现二级标题段落 (Heading 2 或 outlineLvl=1)')
    else:
        for i in h2_paras:
            p = model.paragraphs[i]
            text_head = p.text[:12]

            # (1)(2) 字体 / 字号 逐 run 校验
            runs = p.runs
            rinfos = model.run_info[i]
            checked_any = False
            for ri, rinfo in enumerate(rinfos):
                rtext = (runs[ri].text or '')
                if not rtext.strip():
                    continue
                checked_any = True
                if rinfo['eastAsia'] != '黑体':
                    h2_ok = False
                    h2_reason.append(
                        f'H2[{i}] "{text_head}" run[{ri}] 非黑体: {rinfo["eastAsia"]}'
                    )
                sz = rinfo['size']
                if sz is None or abs(sz - 14) > 0.5:
                    h2_ok = False
                    h2_reason.append(
                        f'H2[{i}] "{text_head}" run[{ri}] 非四号(14pt): {sz}'
                    )
            if not checked_any:
                h2_ok = False
                h2_reason.append(f'H2[{i}] "{text_head}" 段无可校验的文本 run')

            # (3) 左对齐: WD_ALIGN_PARAGRAPH.LEFT 或 None (Word/WPS 默认即左)
            align = model.pinfo[i]['align']
            if align not in (WD_ALIGN_PARAGRAPH.LEFT, None):
                h2_ok = False
                h2_reason.append(f'H2[{i}] "{text_head}" 未左对齐: align={align}')

    evidence = f'{len(h2_paras)} 个二级标题'
    if not h2_ok:
        evidence += ' | 问题: ' + '; '.join(h2_reason[:8])
        if len(h2_reason) > 8:
            evidence += f' ... 共 {len(h2_reason)} 条'
    else:
        evidence += ' 全部正确'
    add('+1 二级标题(黑体/四号/左对齐)',
        1, evidence, h2_ok)

    # -------------------------------------------------------------------------
    # -------------------------------------------------------------------------
    # +1 文中所有一、二级标题行距: 严格按细则一点判定
    #   (1) 一级标题 + 二级标题, 行距均为 2 倍
    # 兼容 Word/WPS:
    #   - 2 倍行距: OOXML 有两种等价写法, 都要接受:
    #       * w:spacing@w:line="480" w:lineRule="auto"  (480/240=2, 段落 UI "多倍行距 2")
    #       * w:spacing@w:line="480" w:lineRule="atLeast"/"exact" — 不视作 2 倍
    #     或 w:pPr/w:spacing@w:lineRule 缺失时默认 "auto"
    #   - line_spacing_value() 已把上述规则归一化为浮点倍数
    # -------------------------------------------------------------------------
    heading_paras = list(h1_paras) + list(h2_paras)
    sp_ok = True
    sp_reason = []
    if not heading_paras:
        sp_ok = False
        sp_reason.append('未发现一/二级标题段落')
    else:
        for i in heading_paras:
            pi = model.pinfo[i]
            ls = line_spacing_value(pi)
            # 精确匹配 2 倍 (±0.05 容差, 应对 lineRule=auto 情况下的浮点)
            if ls is None or abs(ls - 2.0) > 0.05:
                sp_ok = False
                lvl = 'H1' if i in h1_paras else 'H2'
                sp_reason.append(
                    f'{lvl}[{i}] "{model.paragraphs[i].text[:12]}" 行距非 2 倍: {ls} (raw={pi["line_spacing"]})'
                )
    evidence = f'一级={len(h1_paras)} 二级={len(h2_paras)}'
    if not sp_ok:
        evidence += ' | 问题: ' + '; '.join(sp_reason[:8])
        if len(sp_reason) > 8:
            evidence += f' ... 共 {len(sp_reason)} 条'
    else:
        evidence += f' 共 {len(heading_paras)} 个标题行距均 2 倍'
    add('+1 一/二级标题行距 2 倍',
        1, evidence, sp_ok)

    # -------------------------------------------------------------------------
    # +1 正文(除标题和参考文献页, 宋体小四 / 缩进2字符 / 1.5倍行距 / 两端对齐)
    # -------------------------------------------------------------------------
    # -------------------------------------------------------------------------
    # +1 文中除标题和参考文献页, 其余所有文本: 严格按细则五点判定
    #   (1) 字体: 宋体
    #   (2) 字号: 小四 (12pt)
    #   (3) 段落首行缩进 2 字符
    #   (4) 1.5 倍行距
    #   (5) 两端对齐
    # 范围: 文档全部非空段落, 排除
    #   - 标题段 (Heading 1/2/3 样式, 或 outlineLvl 0/1/2)
    #   - 参考文献页 (从 "参考文献" 段起始, 到下一个非参考文献页的分节为止,
    #     此文档中至致谢/附录/结束)
    # 兼容 Word/WPS:
    #   - 宋体: w:rFonts@w:eastAsia = "宋体"
    #   - 小四: w:sz = 24 半点 (12pt)
    #   - 首行缩进 2 字符: w:pPr/w:ind@w:firstLineChars="200" (Word/WPS "字符" 缩进 UI)
    #                    或 w:ind@w:firstLine 折算 ≈ 2 * eastAsia字号
    #   - 1.5 倍行距: w:spacing@w:line="360" w:lineRule="auto"
    #   - 两端对齐: w:jc="both" → WD_ALIGN_PARAGRAPH.JUSTIFY
    # -------------------------------------------------------------------------
    body_ok = True
    body_reason = []

    # 参考文献页范围: 从 "参考文献" 一级标题所在段 (含) 直到文档尾,
    # 若之后仍有其它 H1 (致谢/附录), 则参考文献页止于该 H1 前
    ref_start = model.first_para_idx('参考文献')
    ref_end = None  # exclusive
    if ref_start is not None:
        ref_end = len(model.paragraphs)
        for j in range(ref_start + 1, len(model.paragraphs)):
            p_j = model.paragraphs[j]
            if p_j.style.name == 'Heading 1' or _paragraph_outline_level(p_j) == 0:
                ref_end = j
                break

    def _is_heading(p):
        if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
            return True
        lvl = _paragraph_outline_level(p)
        return lvl in (0, 1, 2)

    body_paras = []
    for i, p in enumerate(model.paragraphs):
        if not p.text.strip():
            continue
        # 排除标题
        if _is_heading(p):
            continue
        # 排除参考文献页
        if ref_start is not None and ref_end is not None and ref_start <= i < ref_end:
            continue
        body_paras.append(i)

    if not body_paras:
        body_ok = False
        body_reason.append('无可校验的正文段落')
    else:
        for i in body_paras:
            p = model.paragraphs[i]
            text_head = p.text[:12]
            runs = p.runs
            rinfos = model.run_info[i]
            pi = model.pinfo[i]

            # (1)(2) 逐 run 校验字体/字号
            for ri, rinfo in enumerate(rinfos):
                rtext = (runs[ri].text or '')
                if not rtext.strip():
                    continue
                if rinfo['eastAsia'] != '宋体':
                    body_ok = False
                    body_reason.append(
                        f'段[{i}] "{text_head}" run[{ri}] 非宋体: {rinfo["eastAsia"]}'
                    )
                sz = rinfo['size']
                if sz is None or abs(sz - 12) > 0.5:
                    body_ok = False
                    body_reason.append(
                        f'段[{i}] "{text_head}" run[{ri}] 非小四(12pt): {sz}'
                    )

            # (3) 首行缩进 2 字符
            # 精确判定: w:pPr/w:ind@w:firstLineChars="200" (Word/WPS UI "首行缩进 2 字符" 写入),
            # 或 firstLine (pt) ≈ 2 * 段落首字号 (兼容用磅设定的写法)
            pPr = p._element.find(qn('w:pPr'))
            fl_chars_ok = False
            fl_pt_ok = False
            fl_chars_val = None
            fl_pt_val = None
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    flc = ind.get(qn('w:firstLineChars'))
                    if flc is not None:
                        try:
                            fl_chars_val = int(flc)
                            if abs(fl_chars_val - 200) <= 5:  # 200 = 2字符 (Word/WPS 单位: 1/100 字符)
                                fl_chars_ok = True
                        except ValueError:
                            pass
                    fl = ind.get(qn('w:firstLine'))
                    if fl is not None:
                        try:
                            fl_twips = int(fl)
                            fl_pt_val = fl_twips / 20.0
                            # 首字号 (pt)
                            first_sz = None
                            for rinfo in rinfos:
                                if rinfo['size'] is not None:
                                    first_sz = rinfo['size']; break
                            if first_sz is None:
                                first_sz = 12
                            if abs(fl_pt_val - 2 * first_sz) <= 2:
                                fl_pt_ok = True
                        except ValueError:
                            pass
            if not (fl_chars_ok or fl_pt_ok):
                body_ok = False
                body_reason.append(
                    f'段[{i}] "{text_head}" 首行未缩进2字符 (chars={fl_chars_val}, pt={fl_pt_val})'
                )

            # (4) 1.5 倍行距
            ls = line_spacing_value(pi)
            if ls is None or abs(ls - 1.5) > 0.05:
                body_ok = False
                body_reason.append(
                    f'段[{i}] "{text_head}" 非1.5倍行距: {ls}'
                )

            # (5) 两端对齐
            if pi['align'] != WD_ALIGN_PARAGRAPH.JUSTIFY:
                body_ok = False
                body_reason.append(
                    f'段[{i}] "{text_head}" 未两端对齐: align={pi["align"]}'
                )

    evidence = f'正文段数={len(body_paras)}'
    if not body_ok:
        evidence += ' | 问题: ' + '; '.join(body_reason[:8])
        if len(body_reason) > 8:
            evidence += f' ... 共 {len(body_reason)} 条'
    else:
        evidence += ' 全部符合'
    add('+1 正文(宋体/小四/首行缩进2字符/1.5倍行距/两端对齐)',
        1, evidence, body_ok)

    # -------------------------------------------------------------------------
    # +3 全文所有一级标题与其下文均要与上文分开, 另起一页
    # 严格按细则一点判定:
    #   一级标题本身在新页起始 (与上文分开), 其下文自然继续在该新页 —
    #   即一级标题段前必须存在硬性分页。
    # 兼容 Word/WPS 三种"分页"OOXML 写法, 任一命中即视为已另起一页:
    #   (a) <w:br w:type="page"/>  —— UI 中 Ctrl+Enter 插入的分页符
    #   (b) 段落 w:pPr/w:pageBreakBefore  —— UI 中 段落 → "段前分页" 勾选
    #   (c) 段落属于新分节起始 (<w:sectPr> 位于该段前的段落末尾)
    # model.page_break_before[i] 已覆盖 (a) 和 (c), 此处补充 (b) 的段落属性判定
    # -------------------------------------------------------------------------
    def _has_page_break_before_prop(p):
        """检查段落 w:pPr/w:pageBreakBefore 是否为 true (Word/WPS 段前分页勾选)"""
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            return False
        pbb = pPr.find(qn('w:pageBreakBefore'))
        if pbb is None:
            return False
        v = pbb.get(qn('w:val'))
        # w:pageBreakBefore 无 val 属性时默认为 true; val=0/false 时为关闭
        if v is None:
            return True
        return v not in ('0', 'false')

    # 一级标题识别: 样式 Heading 1 或 outlineLvl=0
    h1_idx_list = list(h1_paras)  # 复用上文按样式+outlineLvl 归并的一级标题集合
    pb_ok = True
    pb_reason = []
    if not h1_idx_list:
        pb_ok = False
        pb_reason.append('未发现一级标题段落')
    else:
        for i in h1_idx_list:
            p = model.paragraphs[i]
            head = p.text[:12]
            has_break = model.page_break_before[i] or _has_page_break_before_prop(p)
            # 与上文分开: 若 i == 0 (文档第一段), 无 "上文" 可分开 → 视为满足;
            # 否则必须存在硬性分页
            if i == 0:
                continue
            if not has_break:
                pb_ok = False
                pb_reason.append(f'一级标题 "{head}" 前无分页, 未与上文分开')

    evidence = f'一级标题数={len(h1_idx_list)}'
    if not pb_ok:
        evidence += ' | 问题: ' + '; '.join(pb_reason[:8])
        if len(pb_reason) > 8:
            evidence += f' ... 共 {len(pb_reason)} 条'
    else:
        evidence += ' 均另起一页'
    add('+3 全部一级标题另起一页(与上文分开)',
        3, evidence, pb_ok)

    # -------------------------------------------------------------------------
    # +1 文章页码: 严格按细则四点判定
    #   (1) 起始位置: 从 "绪论" 所在页开始
    #   (2) 结束位置: 至尾页 (最后一页) 都有页码
    #   (3) 格式: "第 N 页" 连续递增 (N 为 PAGE 字段, "第" 和 "页" 为文字);
    #       "连续递增" 须验证绪论页页码为 1, 且跨分节不跳号/不重置
    #   (4) 位置: 页脚, 居中显示
    # 兼容 Word/WPS:
    #   - "绪论" 所在的分节使用独立 footer, 其 <w:footerReference> 指向的
    #     footer*.xml 中包含 "第 " + <w:instrText>PAGE</w:instrText> + " 页"
    #   - 页码起始/跳号由 <w:pgNumType w:start="N"/> 决定 (Word/WPS 均按该属性
    #     渲染当前页码; 未设置 w:start 时沿用上一分节页码继续 +1)
    #   - 居中: 页脚段落 <w:pPr>/<w:jc w:val="center"/>
    #   - "从绪论起到尾页": 绪论所在的分节必须是文档的最后一个分节 (主分节),
    #     且该分节的 footer 有效 —— 这样 Word/WPS 会把页码印到该分节所有页
    # -------------------------------------------------------------------------
    page_num_ok = True
    page_num_reason = []

    xulun_idx = model.first_para_idx('绪论')

    # 收集所有 sectPr 起始位置, 建立段落 -> 所在分节的映射
    body_xml = model.doc_xml
    sectpr_iter = list(re.finditer(r'<w:sectPr\b[^>]*>.*?</w:sectPr>', body_xml, re.S))
    para_starts = [m.start() for m in re.finditer(r'<w:p[ >]', body_xml)]

    # 每个 sectPr 覆盖的段落区间: sectPr 位置之前的所有段属于该分节
    # (Word/WPS 中末尾 sectPr 位于 body 末, 覆盖它前面所有未被前置 sectPr 分走的段)
    def _para_section_idx(para_index):
        if not sectpr_iter or para_index >= len(para_starts):
            return None
        ppos = para_starts[para_index]
        for si, sm in enumerate(sectpr_iter):
            if ppos < sm.start():
                return si
        return len(sectpr_iter) - 1  # 兜底: 最后一个分节

    # 解析 rels: rId -> target
    rels_map = {}
    for m in re.finditer(r'<Relationship\b[^>]*/>', model.rels_xml):
        a = m.group(0)
        idm = re.search(r'\bId="([^"]+)"', a)
        tm = re.search(r'\bTarget="([^"]+)"', a)
        if idm and tm:
            rels_map[idm.group(1)] = tm.group(1)

    def _section_footer_xmls(sect_xml):
        """返回该 sectPr 引用的 footer*.xml 内容列表 (default/first/even 全部)"""
        out = []
        for m in re.finditer(r'<w:footerReference\b[^/>]*/>', sect_xml):
            rid_m = re.search(r'r:id="(rId\d+)"', m.group(0))
            if not rid_m:
                continue
            target = rels_map.get(rid_m.group(1), '')
            if target:
                key = 'word/' + target.lstrip('/')
                content = model.footers.get(key, '')
                if content:
                    out.append(content)
        return out

    def _section_pgnum_start(sect_xml):
        """返回该 sectPr 自身 <w:pgNumType w:start="N"/> 的起始值 (无则 None)"""
        m = re.search(r'<w:pgNumType\b[^>]*\bw:start="(\d+)"', sect_xml)
        return int(m.group(1)) if m else None

    # (1) 定位绪论所在分节
    xulun_sect = None
    if xulun_idx is None:
        page_num_ok = False
        page_num_reason.append('未找到 "绪论" 段')
    else:
        xulun_sect = _para_section_idx(xulun_idx)
        if xulun_sect is None:
            page_num_ok = False
            page_num_reason.append(f'"绪论"@{xulun_idx} 无法定位所在分节')

    # (2) 至尾页: 绪论所在分节必须直到文档结束都有 footer
    # 收集"绪论分节及其后"的所有分节 —— 全部要有 footer 含 "第 N 页" 且居中
    target_sects = []
    if xulun_sect is not None:
        target_sects = list(range(xulun_sect, len(sectpr_iter)))
    elif not sectpr_iter:
        page_num_ok = False
        page_num_reason.append('文档无 sectPr, 无法承载页脚')

    # (3)(4) 对每个目标分节校验 footer
    for si in target_sects:
        sect_xml = sectpr_iter[si].group(0)
        footers = _section_footer_xmls(sect_xml)
        if not footers:
            page_num_ok = False
            page_num_reason.append(f'分节#{si} 无 footerReference, 该分节无页码')
            continue
        joined = '\n'.join(footers)
        # (3) 格式 "第 N 页": 含 "第" + PAGE 字段 + "页"
        # PAGE 字段的规范写法: <w:instrText>...PAGE...</w:instrText>
        has_di = '第' in ''.join(re.findall(r'<w:t[^>]*>([^<]*)</w:t>', joined))
        has_ye = '页' in ''.join(re.findall(r'<w:t[^>]*>([^<]*)</w:t>', joined))
        has_PAGE_field = bool(re.search(
            r'<w:instrText[^>]*>[^<]*\bPAGE\b[^<]*</w:instrText>', joined))
        if not (has_di and has_ye and has_PAGE_field):
            page_num_ok = False
            page_num_reason.append(
                f'分节#{si} footer 格式非"第 N 页"'
                f' (第={has_di}, 页={has_ye}, PAGE字段={has_PAGE_field})'
            )
        # (4) 居中: 页脚段落 <w:jc w:val="center"/> (需在段落 pPr 内, 不在 rPr 内)
        # 匹配 <w:pPr>...<w:jc w:val="center"/>...</w:pPr>
        has_center = bool(re.search(
            r'<w:pPr\b[^>]*>.*?<w:jc\s+w:val="center"\s*/>.*?</w:pPr>',
            joined, re.S))
        if not has_center:
            page_num_ok = False
            page_num_reason.append(f'分节#{si} 页码未居中显示')

    # 递增性: <w:pgNumType w:start="N"/> 才是页码起始/跳号的真正依据
    #   - 绪论所在分节必须显式声明 w:start="1" (从绪论页开始编号为 1)
    #   - 绪论之后的所有分节不得再出现 w:start (否则会打断连续递增, 重新计数)
    if target_sects:
        xulun_start = _section_pgnum_start(sectpr_iter[target_sects[0]].group(0))
        if xulun_start != 1:
            page_num_ok = False
            page_num_reason.append(
                f'"绪论"所在分节#{xulun_sect} 页码起始值非1 (pgNumType/w:start={xulun_start})'
            )
        for si in target_sects[1:]:
            start_val = _section_pgnum_start(sectpr_iter[si].group(0))
            if start_val is not None:
                page_num_ok = False
                page_num_reason.append(
                    f'分节#{si} 页码被重新设置起始值({start_val}), 跨节不连续递增'
                )

    evidence = f'"绪论"@{xulun_idx} 分节#{xulun_sect}'
    if not page_num_ok:
        evidence += ' | 问题: ' + '; '.join(page_num_reason[:6])
        if len(page_num_reason) > 6:
            evidence += f' ... 共 {len(page_num_reason)} 条'
    else:
        evidence += f' 至尾页共 {len(target_sects)} 个分节页脚均有 "第 N 页" 居中'
    add('+1 页码(从绪论至尾页 / "第 N 页" 连续递增 / 页脚居中)',
        1, evidence, page_num_ok)

    # -------------------------------------------------------------------------
    # +1 参考文献独立成一页: 严格按细则四点判定
    #   (1) 独立成一页: 标题 "参考文献" 段前有硬性分页
    #   (2) 字体: 黑体
    #   (3) 字号: 三号 (16pt)
    #   (4) 居中显示
    #   (5) 2 倍行距
    # 兼容 Word/WPS:
    #   - 独立成一页: <w:br w:type="page"/>, w:pPr/w:pageBreakBefore, 或分节起始
    #   - 黑体: w:rFonts@w:eastAsia = "黑体"
    #   - 三号: w:sz = 32 半点 (16pt)
    #   - 居中: w:jc="center" → WD_ALIGN_PARAGRAPH.CENTER
    #   - 2 倍行距: w:spacing@w:line="480" w:lineRule="auto"
    # -------------------------------------------------------------------------
    ref_idx = model.first_para_idx('参考文献')
    ref_ok = True
    ref_reason = []

    if ref_idx is None:
        ref_ok = False
        ref_reason.append('未找到 "参考文献" 段')
    else:
        p = model.paragraphs[ref_idx]

        # (1) 独立成一页: 段前存在硬性分页
        has_break = (model.page_break_before[ref_idx]
                     or _has_page_break_before_prop(p))
        if not has_break:
            ref_ok = False
            ref_reason.append('"参考文献" 段前无分页, 未独立成一页')

        # (2)(3) 逐 run 校验字体 / 字号
        runs = p.runs
        rinfos = model.run_info[ref_idx]
        checked_any = False
        for ri, rinfo in enumerate(rinfos):
            rtext = (runs[ri].text or '')
            if not rtext.strip():
                continue
            checked_any = True
            if rinfo['eastAsia'] != '黑体':
                ref_ok = False
                ref_reason.append(f'run[{ri}] 非黑体: {rinfo["eastAsia"]}')
            sz = rinfo['size']
            if sz is None or abs(sz - 16) > 0.5:
                ref_ok = False
                ref_reason.append(f'run[{ri}] 非三号(16pt): {sz}')
        if not checked_any:
            ref_ok = False
            ref_reason.append('"参考文献" 段无可校验的文本 run')

        # (4) 居中
        pi = model.pinfo[ref_idx]
        if pi['align'] != WD_ALIGN_PARAGRAPH.CENTER:
            ref_ok = False
            ref_reason.append(f'"参考文献" 未居中: align={pi["align"]}')

        # (5) 2 倍行距
        ls = line_spacing_value(pi)
        if ls is None or abs(ls - 2.0) > 0.05:
            ref_ok = False
            ref_reason.append(f'"参考文献" 非 2 倍行距: {ls} (raw={pi["line_spacing"]})')

    evidence = f'"参考文献"@{ref_idx}'
    if not ref_ok:
        evidence += ' | 问题: ' + '; '.join(ref_reason)
    else:
        evidence += ' 标题格式正确'
    add('+1 参考文献(独立成页 / 黑体 / 三号 / 居中 / 2倍行距)',
        1, evidence, ref_ok)

    # -------------------------------------------------------------------------
    # +3 参考文献页标题下方的 6 条参考文献项目:
    #   (1) 6 条项目 (数量 = 6)
    #   (2) 编号连续无断开 (即 [1] [2] [3] [4] [5] [6] 顺序连续)
    #   (3) 编号与后面文本之间空一个字符 ("]" 之后紧跟一个空格)
    #   (4) 中文文本字体: 五号宋体
    #   (5) 英文文本字体: 五号宋体 或 五号 Times New Roman
    #   (6) 单倍行距
    # 兼容 Word/WPS:
    #   - 编号规范: "[N]" 段首模式 (Word/WPS 手动键入或 "自定义编号" 均写入
    #     实际字符, 保证 <w:t> 中可见)
    #   - 中文字体: w:rFonts@w:eastAsia = "宋体"
    #   - 英文字体: w:rFonts@w:ascii/hAnsi = "宋体" 或 "Times New Roman"
    #     (Word/WPS 分开控制中文/英文字体, ascii 对应英文字符)
    #   - 五号: w:sz = 21 半点 (10.5pt) —— 特殊, 半点单位下五号是 21
    #   - 单倍行距: w:spacing@w:line="240" w:lineRule="auto"
    #     或 <w:spacing> 完全缺失 (Word/WPS 默认单倍)
    # -------------------------------------------------------------------------
    ref_items_ok = True
    ref_items_reason = []

    # 收集参考文献标题之后, 以 "[N]" 开头的段
    ref_paras = []
    if ref_idx is not None:
        for i in range(ref_idx + 1, len(model.paragraphs)):
            t = model.paragraphs[i].text.strip()
            m = re.match(r'^\[(\d+)\]', t)
            if m:
                ref_paras.append((i, int(m.group(1)), t))
            else:
                # 空段跳过, 非编号段结束采集
                if t:
                    break

    # (1) 数量 = 6
    if len(ref_paras) != 6:
        ref_items_ok = False
        ref_items_reason.append(f'参考文献项数量 = {len(ref_paras)}, 非 6 条')
    else:
        # (2) 编号连续无断开: [1]..[6]
        nums = [n for _, n, _ in ref_paras]
        if nums != list(range(1, 7)):
            ref_items_ok = False
            ref_items_reason.append(f'编号不连续: {nums} (应为 [1]..[6])')

        for pidx, n, t in ref_paras:
            # (3) 编号后空一个字符: 段首模式 "[N] " (] 后紧跟 1 个空格,
            #     且再后 1 个字符为非空白 — 保证 "空一个字符" 是精确一格)
            prefix = f'[{n}]'
            tail = t[len(prefix):]
            # tail 首字符必须是 1 个空白, 且第 2 字符 (若存在) 非空白
            if not tail:
                ref_items_ok = False
                ref_items_reason.append(f'[{n}] 编号后无内容')
            elif tail[0] not in ' 　':  # 半角/全角空格均接受, Word/WPS 皆常见
                ref_items_ok = False
                ref_items_reason.append(f'[{n}] 编号后未空一字符 (下一字符="{tail[0]}")')
            elif len(tail) > 1 and tail[1] in ' 　':
                ref_items_ok = False
                ref_items_reason.append(f'[{n}] 编号后空多于一个字符')

            # (4)(5) 逐 run 校验字体和字号
            runs = model.paragraphs[pidx].runs
            rinfos = model.run_info[pidx]
            for ri, rinfo in enumerate(rinfos):
                rtext = (runs[ri].text or '')
                if not rtext.strip():
                    continue
                # 判断该 run 主体是中文还是英文
                # Word/WPS 中 rFonts 同时含 eastAsia 与 ascii; 渲染时按字符类别选用
                has_cjk = any('一' <= c <= '鿿' for c in rtext)
                has_ascii = any(c.isascii() and c.isalnum() for c in rtext)

                # 字号 (五号 = 10.5pt), 中英文共用同一 w:sz
                sz = rinfo['size']
                if sz is None or abs(sz - 10.5) > 0.6:
                    ref_items_ok = False
                    ref_items_reason.append(
                        f'[{n}] run[{ri}] 非五号(10.5pt): {sz}'
                    )

                # 中文字体
                if has_cjk:
                    if rinfo['eastAsia'] != '宋体':
                        ref_items_ok = False
                        ref_items_reason.append(
                            f'[{n}] run[{ri}] 中文非宋体: {rinfo["eastAsia"]}'
                        )
                # 英文字体
                if has_ascii:
                    ascii_font = rinfo['ascii']
                    # 已归一化: SimSun/宋体 → "宋体"; Times New Roman 保持原名
                    if ascii_font not in (None, '宋体', 'Times New Roman'):
                        ref_items_ok = False
                        ref_items_reason.append(
                            f'[{n}] run[{ri}] 英文非宋体/Times New Roman: {ascii_font}'
                        )

            # (6) 单倍行距
            pi = model.pinfo[pidx]
            ls = line_spacing_value(pi)
            # Word/WPS 中单倍: line=240 auto 或未设 <w:spacing> (默认单倍)
            if ls is None:
                pass  # 未指定 = 默认单倍
            elif abs(ls - 1.0) > 0.05:
                ref_items_ok = False
                ref_items_reason.append(
                    f'[{n}] 非单倍行距: {ls} (raw={pi["line_spacing"]})'
                )

    evidence = f'参考文献项数={len(ref_paras)}'
    if not ref_items_ok:
        evidence += ' | 问题: ' + '; '.join(ref_items_reason[:8])
        if len(ref_items_reason) > 8:
            evidence += f' ... 共 {len(ref_items_reason)} 条'
    else:
        evidence += ' 6 条编号连续/空一字符/五号宋体或TNR/单倍行距 全部正确'
    add('+3 6条参考文献(连续/空一字符/五号中宋英宋或TNR/单倍行距)',
        3, evidence, ref_items_ok)

    # -------------------------------------------------------------------------
    # +3 致谢页(单独成页 / 黑体三号2倍行距 / 下方宋体小四缩进2字符1.5倍)
    # -------------------------------------------------------------------------
    # -------------------------------------------------------------------------
    # +3 致谢页: 严格按细则六点判定
    #   标题:
    #     (1) 单独成一页 (标题段前有硬性分页)
    #     (2) 字体: 黑体
    #     (3) 字号: 三号 (16pt)
    #     (4) 2 倍行距
    #   下方文本:
    #     (5) 字体: 宋体
    #     (6) 字号: 小四 (12pt)
    #     (7) 首行缩进 2 字符
    #     (8) 1.5 倍行距
    # 兼容 Word/WPS:
    #   - 独立成一页: <w:br w:type="page"/>, w:pPr/w:pageBreakBefore, 或分节起始
    #   - 黑体/宋体: w:rFonts@w:eastAsia
    #   - 三号/小四: w:sz = 32/24 半点
    #   - 2 倍行距: w:spacing@w:line=480 w:lineRule=auto
    #   - 1.5 倍行距: w:spacing@w:line=360 w:lineRule=auto
    #   - 首行缩进 2 字符: w:pPr/w:ind@w:firstLineChars=200 或 firstLine≈2*字号 pt
    # -------------------------------------------------------------------------
    zx_idx = model.first_para_idx('致谢')
    zx_ok = True
    zx_reason = []

    if zx_idx is None:
        zx_ok = False
        zx_reason.append('未找到 "致谢" 段')
    else:
        p_title = model.paragraphs[zx_idx]

        # (1) 单独成一页: 段前有分页
        has_break = (model.page_break_before[zx_idx]
                     or _has_page_break_before_prop(p_title))
        if not has_break:
            zx_ok = False
            zx_reason.append('"致谢" 段前无分页, 未单独成一页')

        # (2)(3) 标题字体 / 字号 逐 run 校验
        runs_t = p_title.runs
        rinfos_t = model.run_info[zx_idx]
        checked_title = False
        for ri, rinfo in enumerate(rinfos_t):
            rtext = (runs_t[ri].text or '')
            if not rtext.strip():
                continue
            checked_title = True
            if rinfo['eastAsia'] != '黑体':
                zx_ok = False
                zx_reason.append(f'致谢标题 run[{ri}] 非黑体: {rinfo["eastAsia"]}')
            sz = rinfo['size']
            if sz is None or abs(sz - 16) > 0.5:
                zx_ok = False
                zx_reason.append(f'致谢标题 run[{ri}] 非三号(16pt): {sz}')
        if not checked_title:
            zx_ok = False
            zx_reason.append('"致谢" 标题段无可校验的文本 run')

        # (4) 标题 2 倍行距
        pi_t = model.pinfo[zx_idx]
        ls_t = line_spacing_value(pi_t)
        if ls_t is None or abs(ls_t - 2.0) > 0.05:
            zx_ok = False
            zx_reason.append(f'致谢标题非 2 倍行距: {ls_t} (raw={pi_t["line_spacing"]})')

        # 下方文本: 标题之后所有非空段落 (直到文档尾或下一分页)
        body_paras_zx = []
        for j in range(zx_idx + 1, len(model.paragraphs)):
            if j != zx_idx + 1 and model.page_break_before[j]:
                break
            if model.paragraphs[j].text.strip():
                body_paras_zx.append(j)

        if not body_paras_zx:
            zx_ok = False
            zx_reason.append('致谢标题下方无正文')
        else:
            for j in body_paras_zx:
                pj = model.paragraphs[j]
                head_j = pj.text[:12]
                runs_j = pj.runs
                rinfos_j = model.run_info[j]
                pi_j = model.pinfo[j]

                # (5)(6) 字体 / 字号 逐 run
                for ri, rinfo in enumerate(rinfos_j):
                    rtext = (runs_j[ri].text or '')
                    if not rtext.strip():
                        continue
                    if rinfo['eastAsia'] != '宋体':
                        zx_ok = False
                        zx_reason.append(
                            f'致谢正文@段{j} "{head_j}" run[{ri}] 非宋体: {rinfo["eastAsia"]}'
                        )
                    sz = rinfo['size']
                    if sz is None or abs(sz - 12) > 0.5:
                        zx_ok = False
                        zx_reason.append(
                            f'致谢正文@段{j} "{head_j}" run[{ri}] 非小四(12pt): {sz}'
                        )

                # (7) 首行缩进 2 字符
                pPr = pj._element.find(qn('w:pPr'))
                fl_chars_ok = False
                fl_pt_ok = False
                fl_chars_val = None
                fl_pt_val = None
                if pPr is not None:
                    ind = pPr.find(qn('w:ind'))
                    if ind is not None:
                        flc = ind.get(qn('w:firstLineChars'))
                        if flc is not None:
                            try:
                                fl_chars_val = int(flc)
                                if abs(fl_chars_val - 200) <= 5:
                                    fl_chars_ok = True
                            except ValueError:
                                pass
                        fl = ind.get(qn('w:firstLine'))
                        if fl is not None:
                            try:
                                fl_twips = int(fl)
                                fl_pt_val = fl_twips / 20.0
                                first_sz = None
                                for rinfo in rinfos_j:
                                    if rinfo['size'] is not None:
                                        first_sz = rinfo['size']; break
                                if first_sz is None:
                                    first_sz = 12
                                if abs(fl_pt_val - 2 * first_sz) <= 2:
                                    fl_pt_ok = True
                            except ValueError:
                                pass
                if not (fl_chars_ok or fl_pt_ok):
                    zx_ok = False
                    zx_reason.append(
                        f'致谢正文@段{j} "{head_j}" 首行未缩进2字符 '
                        f'(chars={fl_chars_val}, pt={fl_pt_val})'
                    )

                # (8) 1.5 倍行距
                ls_j = line_spacing_value(pi_j)
                if ls_j is None or abs(ls_j - 1.5) > 0.05:
                    zx_ok = False
                    zx_reason.append(
                        f'致谢正文@段{j} "{head_j}" 非1.5倍行距: {ls_j}'
                    )

    evidence = f'"致谢"@{zx_idx}'
    if not zx_ok:
        evidence += ' | 问题: ' + '; '.join(zx_reason[:8])
        if len(zx_reason) > 8:
            evidence += f' ... 共 {len(zx_reason)} 条'
    else:
        evidence += ' 致谢页格式正确'
    add('+3 致谢页(单独成页 / 标题黑体三号2倍 / 正文宋体小四1.5倍首行缩进2字符)',
        3, evidence, zx_ok)

    return total, hits


# =============================================================================
# 主流程 (统一接口)
# =============================================================================

SCRIPT_ID = '016'


def _locate_docx(dir_path: str) -> Optional[Path]:
    """在指定目录中定位待评估的 .docx 文件.

    - 忽略以 '~$' 开头的 Office 临时锁文件
    - 若目录内存在多个匹配文件, 取名称排序后的第 1 个
    """
    d = Path(dir_path)
    if not d.is_dir():
        return None
    candidates = []
    for entry in sorted(d.iterdir(), key=lambda p: p.name):
        if not entry.is_file():
            continue
        if entry.name.startswith('~$'):
            continue
        if entry.suffix.lower() == '.docx':
            candidates.append(entry)
    if not candidates:
        return None
    return candidates[0]


def evaluate(dir_path: str) -> dict:
    """统一评估入口.

    参数:
        dir_path: 脚本所在目录的路径; 脚本自行在该目录中定位被评估的文档.

    返回:
        结构化字典 (见"脚本接口差异与统一建议.md" §2.2).
    """
    result = {
        'id': SCRIPT_ID,
        'file_name': '',
        'status': 'ok',
        'error': None,
        'dim1_pass': False,
        'dim1_reason': '',
        'dim2_items': [],
        'total_score': 0,
        'max_score': 0,
    }

    try:
        # 在给定目录内定位待评估文档
        docx_path = _locate_docx(dir_path)
        if docx_path is None:
            result['status'] = 'error'
            result['error'] = f'目录 {dir_path!r} 中未找到 .docx 文件'
            return result
        result['file_name'] = docx_path.name

        # 构建文档模型
        model = DocumentModel(str(docx_path))

        # 维度 1
        d1_pass, d1_msgs, _d1_details = check_dimension1(model)
        result['dim1_pass'] = bool(d1_pass)
        result['dim1_reason'] = '' if d1_pass else '; '.join(d1_msgs)

        if not d1_pass:
            # 维度 1 未通过时直接短路，按契约返回空结果
            result['dim2_items'] = []
            result['total_score'] = 0
            result['max_score'] = 0
            return result

        # 维度 2: 仅在维度 1 通过后评估
        total, hits = check_dimension2(model)
        dim2_items = []
        max_score = 0
        for label, score, evidence, hit in hits:
            max_delta = score  # 正/负分均按脚本定义的分值作为该项 max_delta
            actual_delta = score if hit else 0
            if score > 0:
                max_score += score
            dim2_items.append({
                'rule': re.sub(r'^[+-]\d+\s*', '', label),
                'max_delta': max_delta,
                'delta': actual_delta,
                'hit': bool(hit),
                'detail': '',
            })

        result['dim2_items'] = dim2_items
        result['total_score'] = total
        result['max_score'] = max_score

    except Exception as exc:
        result['status'] = 'error'
        result['error'] = f'{type(exc).__name__}: {exc}'

    return result


if __name__ == '__main__':
    # 本地调试: 参数为脚本所在目录 (未提供则默认取当前脚本文件所在目录)
    _arg = sys.argv[1] if len(sys.argv) > 1 else str(Path(__file__).resolve().parent)
    _res = evaluate(_arg)
    print(json.dumps(_res, ensure_ascii=False, indent=2))
