#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
部门协调会记录_会议内容横线表格版.docx 自动评估代码
"""

import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import os
import re
import sys


# 本脚本对应的编号（用于 evaluate() 返回值中的 "id" 字段）
SCRIPT_ID = "003"
# 期望的被评估文档名；若目录中不存在同名文件，则回退到目录内首个 .docx
EXPECTED_DOC_NAME = "部门协调会记录_会议内容横线表格版.docx"


class DocumentEvaluator:
    """文档评估器"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = None
        self.dimension1_passed = True
        self.dimension1_details = []
        self.dimension2_score = 0
        self.dimension2_details = []

    def load_document(self):
        """加载文档"""
        try:
            self.doc = docx.Document(self.docx_path)
            self.dimension1_details.append("✓ 文档可正常打开")
            return True
        except Exception as e:
            self.dimension1_passed = False
            self.dimension1_details.append(f"✗ 文档无法正常打开: {e}")
            return False

    def evaluate_dimension1(self):
        """评估维度1：可用与可修改性"""
        # 1. 检查文件扩展名
        if not self.docx_path.endswith('.docx'):
            self.dimension1_passed = False
            self.dimension1_details.append("✗ 交付文件扩展名不是.docx")
        else:
            self.dimension1_details.append("✓ 交付文件扩展名符合要求")

        # 2. 文档可正常打开（已在load_document中检查）
        # 如果无法打开，直接返回
        if not self.doc:
            return False

        return self.dimension1_passed

    def evaluate_dimension2(self):
        """评估维度2：完成度评分细则"""
        if not self.doc:
            return

        all_text = ""
        for para in self.doc.paragraphs:
            all_text += para.text + "\n"
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "

        # +5：文档表格中没有空白行
        self._check_no_empty_rows()

        # +5：文档表格中所有文本：下方都出现横线；横线长度与表格上边框长度一致；颜色为黑色，类型为细实线
        self._check_underline_format()

        # +5：文档表格中所有横线：相邻两条横线之间的距离几乎相等，只允许有0.15cm的误差
        self._check_row_height()

        # -5：文中出现图片或者不可编辑的部分
        self._check_images_and_non_editable()

        # -5：文档第一页顶部不满足其中任意一项
        self._check_first_page_header()

        # -1："会议内容："下方第一行出现缩进
        self._check_meeting_content_indent()

        # -1：文中没有出现"侯景澜：上月两项重点现场事项整改推进如下："文本内容
        self._check_hou_jinglan_text()

        # -1："二月份作业偏差共发生若干件，具体情况如下："与"四、研究后续事项"之间的文本出现缩进
        self._check_indentation_between_sections()

    def _check_no_empty_rows(self):
        """+5：文档表格中没有空白行

        每条横线对应一个表格行（<w:tr>），要求每个表格行的文字内容不能
        为空。python-docx 的 doc.tables 只覆盖顶层表格，单元格内嵌套的
        子表格（cell.tables）不会被收进去，必须递归遍历才能覆盖表中表
        的所有行。此外，若文档中完全没有表格，该评分点无法成立（rubric
        针对"表格中"的行提出要求），不应默认通过并加分。
        """
        if self.doc is None:
            return

        invisible_chars = ' \t\r\n　\xa0​'
        empty_locations = []
        table_count = [0]

        def walk_table(tbl, prefix):
            table_count[0] += 1
            ti = table_count[0]
            label_prefix = f"{prefix}表{ti}" if prefix else f"表{ti}"
            for ri, row in enumerate(tbl.rows, 1):
                row_text = "".join(cell.text for cell in row.cells)
                if not row_text.strip(invisible_chars):
                    empty_locations.append(f"{label_prefix}第{ri}行")
                for cell in row.cells:
                    for sub in cell.tables:
                        walk_table(sub, f"{label_prefix}内嵌套")

        for table in self.doc.tables:
            walk_table(table, "")

        if table_count[0] == 0:
            self.dimension2_details.append("+5：文档表格中没有空白行 ✗（文档中不存在表格，无法验证该评分点）")
            return

        if not empty_locations:
            self.dimension2_score += 5
            self.dimension2_details.append("+5：文档表格中没有空白行 ✓")
        else:
            preview = ", ".join(empty_locations[:3])
            more = f"等{len(empty_locations)}处" if len(empty_locations) > 3 else ""
            self.dimension2_details.append(
                f"+5：文档表格中没有空白行 ✗（{preview}{more}）"
            )

    def _check_underline_format(self):
        """+5：文档表格中所有文本：下方都出现横线；横线长度与表格上边框长度一致；颜色为黑色，类型为细实线

        单元格底部边框只能保证"该单元格最后一行文字"下方有横线。
        若单元格内文字过长发生自动换行，换行后的前几行文字下方就没有横线。
        判断换行采用"字号 × 字符视觉宽度 ∑ 与 单元格净宽 比较"的估算方式，
        比按"字符数阈值"精准得多。

        "横线长度与表格上边框长度一致"：以行内各单元格 w:tcW 之和作为该
        横线实际覆盖的宽度，与表格整体宽度（各列 w:tcW 之和；缺失则用
        w:tblW）比较，要求近似相等（允许 dxa 舍入误差）——这样才能识别
        "横线只覆盖部分列宽"（例如某行被拆分成多个单元格，只有其中一个
        单元格有底部边框）的不合规情况。
        "颜色为黑色"：要求 w:color 严格等于 auto 或 000000，不再把深灰
        等非黑色（此前 <=80 的宽松阈值会误判为黑色）当作合格。
        """
        if self.doc is None:
            return

        all_ok = True
        fail_reason = ""

        # 文档默认字号（Normal 样式）作为兜底
        default_pt = 10.5
        try:
            normal_size = self.doc.styles['Normal'].font.size
            if normal_size is not None:
                default_pt = normal_size.pt
        except Exception:
            pass

        def char_visual_units(ch):
            """返回该字符相对于当前字号的宽度倍率（1 表示占一个全角字宽）。"""
            code = ord(ch)
            # 中日韩统一表意文字 / 全角标点 / 全角字符
            if (0x4E00 <= code <= 0x9FFF or 0x3400 <= code <= 0x4DBF
                    or 0x3000 <= code <= 0x303F or 0xFF00 <= code <= 0xFFEF):
                return 1.0
            if code < 128:
                # ASCII 平均按 0.55 全角宽估算（数字/字母/半角标点）
                return 0.55
            return 1.0  # 其它非 ASCII 字符按全角处理

        def paragraph_text_width_pt(para):
            """估算段落文字的视觉宽度（pt）。字号取 run 显式 → Normal 兜底。

            仅计算段落"可见内容"的宽度：先剥离段末尾随空白（Word 中尾随空白不参与换行），
            再逐 run 累加。
            """
            # 计算需要保留的字符总数（去掉尾随空白）
            full_text = "".join(run.text for run in para.runs)
            trimmed = full_text.rstrip(" \t　\xa0\r\n")
            remaining = len(trimmed)
            if remaining == 0:
                return 0.0

            width = 0.0
            for run in para.runs:
                run_pt = default_pt
                if run.font.size is not None:
                    run_pt = run.font.size.pt
                for ch in run.text:
                    if remaining <= 0:
                        break
                    remaining -= 1
                    if ch in ('\r', '\n'):
                        continue
                    width += char_visual_units(ch) * run_pt
                if remaining <= 0:
                    break
            return width

        def cell_dxa_width(cell):
            """读取单元格 w:tcW 原始宽度（dxa）。失败返回 0.0。"""
            tcPr = cell._tc.tcPr
            if tcPr is None:
                return 0.0
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                return 0.0
            try:
                w_dxa = float(tcW.get(qn('w:w'), '0'))
            except (TypeError, ValueError):
                return 0.0
            return max(0.0, w_dxa)

        def table_total_width_dxa(table):
            """表格整体宽度（dxa）：优先取首行各单元格 tcW 之和；
            取不到（单元格未显式设宽）则回退到 tblPr/w:tblW。返回 0.0
            表示无法确定（此时跳过"横线长度一致"比较，不误判失败）。
            """
            if table.rows:
                first_row_sum = sum(
                    cell_dxa_width(c) for c in table.rows[0].cells
                )
                if first_row_sum > 0:
                    return first_row_sum
            tblPr = table._tbl.find(qn('w:tblPr'))
            tblW = tblPr.find(qn('w:tblW')) if tblPr is not None else None
            if tblW is not None:
                try:
                    return max(0.0, float(tblW.get(qn('w:w'), '0')))
                except (TypeError, ValueError):
                    return 0.0
            return 0.0

        def cell_inner_width_pt(cell):
            """读取单元格净宽（pt）：tcW 减去左右单元格边距。失败返回 None。"""
            tcPr = cell._tc.tcPr
            if tcPr is None:
                return None
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                return None
            try:
                w_dxa = float(tcW.get(qn('w:w'), '0'))
            except (TypeError, ValueError):
                return None
            if w_dxa <= 0:
                return None
            # 表级默认边距
            left_dxa = right_dxa = 108.0  # Word 默认 0.19cm ≈ 108 dxa（保守值）
            try:
                tbl = cell._tc.getparent().getparent()  # tr -> tbl
                tblCellMar = tbl.find('.//' + qn('w:tblCellMar'))
                if tblCellMar is not None:
                    l = tblCellMar.find(qn('w:left'))
                    r = tblCellMar.find(qn('w:right'))
                    if l is not None and l.get(qn('w:w')) is not None:
                        left_dxa = float(l.get(qn('w:w')))
                    if r is not None and r.get(qn('w:w')) is not None:
                        right_dxa = float(r.get(qn('w:w')))
            except Exception:
                pass
            # dxa = 1/20 pt
            return max(0.0, (w_dxa - left_dxa - right_dxa) / 20.0)

        for table_index, table in enumerate(self.doc.tables, 1):
            # 表级 tblBorders/bottom：作为最后一行"底部横线"的兜底
            tblPr = table._tbl.find(qn('w:tblPr'))
            tbl_borders = tblPr.find(qn('w:tblBorders')) if tblPr is not None else None
            tbl_bottom = tbl_borders.find(qn('w:bottom')) if tbl_borders is not None else None
            last_row_index = len(table.rows)
            table_width_dxa = table_total_width_dxa(table)

            for row_index, row in enumerate(table.rows, 1):
                is_last_row = (row_index == last_row_index)

                # 该行横线实际覆盖的宽度：仅累加"有可见底部边框"的单元格
                # 宽度——只要其中一个单元格缺失/不可见边框，上面的检查会
                # 立即 all_ok=False 并 break，不会走到下面的宽度比较。
                row_line_dxa = 0.0

                for cell_index, cell in enumerate(row.cells, 1):
                    non_empty_paras = [p for p in cell.paragraphs if p.text.strip()]
                    if not non_empty_paras:
                        continue  # 跳过空单元格

                    tcPr = cell._tc.tcPr
                    tcBorders = tcPr.find(qn('w:tcBorders')) if tcPr is not None else None
                    bottom = tcBorders.find(qn('w:bottom')) if tcBorders is not None else None

                    # 最后一行：单元格自身无底部边框时，回退到表级底部边框
                    if bottom is None and is_last_row and tbl_bottom is not None:
                        bottom = tbl_bottom

                    if bottom is None:
                        all_ok = False
                        fail_reason = f"第{table_index}个表格第{row_index}行第{cell_index}列无底部边框"
                        break

                    val = bottom.get(qn('w:val'), '')
                    color = bottom.get(qn('w:color'), '').upper().lstrip('#')
                    sz = int(bottom.get(qn('w:sz'), '0'))
                    # 类型须为细实线（single），sz<=6即细线，颜色须为黑色
                    if val != 'single':
                        all_ok = False
                        fail_reason = f"第{table_index}个表格第{row_index}行底部边框类型不是细实线(val={val})"
                        break
                    if sz > 6:
                        all_ok = False
                        fail_reason = f"第{table_index}个表格第{row_index}行底部边框线宽过粗(sz={sz})"
                        break
                    # 颜色须严格为黑色：auto 或 000000。此前 RGB 各通道
                    # <=80 的宽松阈值会把深灰等非黑色误判为合格，收紧为
                    # 精确匹配。
                    if color not in ('AUTO', '000000', ''):
                        all_ok = False
                        fail_reason = f"第{table_index}个表格第{row_index}行底部边框颜色不为黑色(#{color})"
                        break

                    row_line_dxa += cell_dxa_width(cell)

                    # 若单元格内含多个非空段落：段落硬回车即视为强制换行，
                    # 除最后一个段落外，其它段落下方一定没有横线。
                    if len(non_empty_paras) > 1:
                        all_ok = False
                        fail_reason = (f"第{table_index}个表格第{row_index}行含{len(non_empty_paras)}个段落，"
                                       f"非末段下方无横线")
                        break

                    # 单段落：按视觉宽度判断是否会自动换行
                    # 不加"安全余量"——文字视觉宽度严格大于单元格净宽即判换行。
                    # （测试数据：46 字全中文 ≈ 483pt 会溢出 481.85pt 净宽而换行；
                    # 而 44 中 + 2 半 ≈ 473pt 则不会换行。）
                    inner_w = cell_inner_width_pt(cell)
                    if inner_w and inner_w > 0:
                        text_w = paragraph_text_width_pt(non_empty_paras[0])
                        if text_w > inner_w:
                            all_ok = False
                            fail_reason = (f"第{table_index}个表格第{row_index}行文字过长发生换行"
                                           f"(文字≈{text_w:.0f}pt, 单元格净宽≈{inner_w:.0f}pt)，"
                                           f"换行后的前几行文字下方无横线")
                            break
                if not all_ok:
                    break

                # "横线长度与表格上边框长度一致"：整行都有可见底部边框时，
                # 该横线覆盖宽度应与表格整体宽度一致（允许 20dxa≈1pt 舍入
                # 误差）；宽度信息缺失（table_width_dxa<=0）时无法比较，
                # 跳过该项而不误判失败。
                if table_width_dxa > 0 and row_line_dxa > 0:
                    if abs(row_line_dxa - table_width_dxa) > 20:
                        all_ok = False
                        fail_reason = (
                            f"第{table_index}个表格第{row_index}行横线宽度"
                            f"≈{row_line_dxa/20:.1f}pt，与表格上边框宽度"
                            f"≈{table_width_dxa/20:.1f}pt不一致"
                        )
                        break
            if not all_ok:
                break

        if all_ok:
            self.dimension2_score += 5
            self.dimension2_details.append("+5：文档表格中所有文本：下方都出现横线；横线长度与表格上边框长度一致；颜色为黑色，类型为细实线 ✓")
        else:
            self.dimension2_details.append(f"+5：文档表格中所有文本：下方都出现横线；横线长度与表格上边框长度一致；颜色为黑色，类型为细实线 ✗（{fail_reason}）")

    def _check_row_height(self):
        """+5：文档表格中所有横线：相邻两条横线之间的距离几乎相等，只允许有0.15cm的误差

        评分口径：把所有"相邻两条横线的距离"收集起来，最大值 - 最小值 ≤ 0.15cm。
        - 短行（无字符下划线）：距离 = 表格行的实际渲染高度
        - 长段落（含 <w:u> 字符下划线）：每一视觉行下方都有横线，
          距离 = 段落 line spacing（对每个视觉行都计入一次）
        - 递归遍历嵌套子表；含子表的表格行本身跳过、直接测子表。
        """
        if self.doc is None:
            return

        tolerance = 0.15
        gaps = []  # (label, cm)

        def paragraph_has_underline(para):
            pPr = para._p.find(qn('w:pPr'))
            if pPr is not None:
                rPr_p = pPr.find(qn('w:rPr'))
                if rPr_p is not None:
                    u = rPr_p.find(qn('w:u'))
                    if u is not None and u.get(qn('w:val')) not in (None, 'none'):
                        return True
            for r in para.runs:
                rPr_r = r._r.find(qn('w:rPr'))
                if rPr_r is None:
                    continue
                u = rPr_r.find(qn('w:u'))
                if u is not None and u.get(qn('w:val')) not in (None, 'none'):
                    return True
            return False

        def collect_underlined_gaps(row, label):
            """对带下划线段落：视觉每一行都记一次 line spacing。"""
            found_any = False
            for cell in row.cells:
                inner_w = self._cell_inner_width_pt(cell)
                for para in cell.paragraphs:
                    if not para.text.strip() or not paragraph_has_underline(para):
                        continue
                    found_any = True
                    cm = self._paragraph_line_spacing_cm(para)
                    if cm is None:
                        continue
                    # 估算视觉行数
                    lines = self._visual_line_count(para, inner_w)
                    for _ in range(lines):
                        gaps.append((label, cm))
            return found_any

        def walk(tables, prefix):
            for ti, table in enumerate(tables, 1):
                for ri, row in enumerate(table.rows, 1):
                    label = f"{prefix}表{ti}第{ri}行"
                    has_nested = any(cell.tables for cell in row.cells)
                    if not has_nested:
                        # 优先按下划线段落记录（每一视觉行一条横线间距）
                        if not collect_underlined_gaps(row, label):
                            # 无下划线：整行只贡献一条底边框——距离 = 该行实际高度
                            h_cm = self._get_effective_row_height_cm(row)
                            if h_cm is not None:
                                gaps.append((label, h_cm))
                    for cell in row.cells:
                        if cell.tables:
                            walk(cell.tables, prefix=f"{label}>")

        walk(self.doc.tables, prefix="")

        rule_text = "文档表格中所有横线：相邻两条横线之间的距离几乎相等，只允许有0.15cm的误差"
        if not gaps:
            self.dimension2_details.append(f"+5：{rule_text} ✗（未采集到任何横线间距）")
            return

        max_gap = max(cm for _, cm in gaps)
        min_gap = min(cm for _, cm in gaps)
        diff = max_gap - min_gap
        if diff <= tolerance:
            self.dimension2_score += 5
            self.dimension2_details.append(f"+5：{rule_text} ✓")
        else:
            # 挑出偏离最远的两处，便于定位
            max_label = next(l for l, cm in gaps if cm == max_gap)
            min_label = next(l for l, cm in gaps if cm == min_gap)
            self.dimension2_details.append(
                f"+5：{rule_text} ✗（"
                f"最大间距{max_gap:.2f}cm@{max_label}，"
                f"最小间距{min_gap:.2f}cm@{min_label}，"
                f"极差{diff:.2f}cm > 0.15cm）"
            )

    def _cell_inner_width_pt(self, cell):
        """单元格净宽（pt）。失败返回 None。"""
        tcPr = cell._tc.tcPr
        if tcPr is None:
            return None
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            return None
        try:
            w_dxa = float(tcW.get(qn('w:w'), '0'))
        except (TypeError, ValueError):
            return None
        if w_dxa <= 0:
            return None
        left_dxa = right_dxa = 108.0
        try:
            tbl = cell._tc.getparent().getparent()
            tblCellMar = tbl.find('.//' + qn('w:tblCellMar'))
            if tblCellMar is not None:
                l = tblCellMar.find(qn('w:left'))
                r = tblCellMar.find(qn('w:right'))
                if l is not None and l.get(qn('w:w')) is not None:
                    left_dxa = float(l.get(qn('w:w')))
                if r is not None and r.get(qn('w:w')) is not None:
                    right_dxa = float(r.get(qn('w:w')))
        except Exception:
            pass
        return max(0.0, (w_dxa - left_dxa - right_dxa) / 20.0)

    def _visual_line_count(self, para, inner_w_pt):
        """估算段落自动换行后的视觉行数（≥1）。"""
        import math
        default_pt = 10.5
        try:
            normal_size = self.doc.styles['Normal'].font.size
            if normal_size is not None:
                default_pt = normal_size.pt
        except Exception:
            pass
        para_pt = default_pt
        for r in para.runs:
            if r.font.size is not None:
                para_pt = max(para_pt, r.font.size.pt)

        def char_units(ch):
            code = ord(ch)
            if (0x4E00 <= code <= 0x9FFF or 0x3400 <= code <= 0x4DBF
                    or 0x3000 <= code <= 0x303F or 0xFF00 <= code <= 0xFFEF):
                return 1.0
            if code < 128:
                return 0.55
            return 1.0

        trimmed = "".join(run.text for run in para.runs).rstrip(" \t　\xa0\r\n")
        if not trimmed:
            return 1
        if not inner_w_pt or inner_w_pt <= 0:
            return 1
        text_w = sum(char_units(ch) * para_pt for ch in trimmed)
        return max(1, math.ceil(text_w / inner_w_pt))

    def _paragraph_line_spacing_cm(self, para):
        """段落 line spacing（cm）：exact 取值；atLeast 取 max(设置值, 字号自然行高)；
        auto/未设置按倍数换算，缺省 1.0 倍。"""
        PT_PER_CM = 28.3464567
        DXA_PER_CM = 567.0

        default_pt = 10.5
        try:
            normal_size = self.doc.styles['Normal'].font.size
            if normal_size is not None:
                default_pt = normal_size.pt
        except Exception:
            pass
        para_pt = default_pt
        for r in para.runs:
            if r.font.size is not None:
                para_pt = max(para_pt, r.font.size.pt)
        natural_cm = para_pt * 1.2 / PT_PER_CM

        pPr = para._p.find(qn('w:pPr'))
        spacing = pPr.find(qn('w:spacing')) if pPr is not None else None
        if spacing is None:
            return natural_cm
        line = spacing.get(qn('w:line'))
        line_rule = spacing.get(qn('w:lineRule'))
        if line is None:
            return natural_cm
        try:
            lv = float(line)
        except ValueError:
            return natural_cm
        if line_rule == 'exact':
            return lv / DXA_PER_CM
        if line_rule == 'atLeast':
            return max(lv / DXA_PER_CM, natural_cm)
        # auto / 未指定
        multiplier = lv / 240.0 if lv > 0 else 1.0
        return natural_cm * multiplier

    def _get_effective_row_height_cm(self, row):
        """按 Word 实际渲染语义计算行高 (cm)：
        - exact：取 trHeight 设置值
        - atLeast/auto/未设置：max(设置值, 内容需要高度)
        """
        set_cm = None
        rule = None
        try:
            trPr = row._tr.find(qn('w:trPr'))
            trHeight = trPr.find(qn('w:trHeight')) if trPr is not None else None
            if trHeight is not None:
                h_val = trHeight.get(qn('w:val'))
                rule = trHeight.get(qn('w:hRule'))
                if h_val is not None:
                    try:
                        set_cm = float(h_val) / 567.0  # 1 cm = 567 dxa
                    except (TypeError, ValueError):
                        set_cm = None
        except Exception:
            pass

        if rule == 'exact' and set_cm is not None:
            return set_cm

        content_cm = self._estimate_content_height_cm(row)
        if set_cm is None and content_cm is None:
            return None
        if set_cm is None:
            return content_cm
        if content_cm is None:
            return set_cm
        return max(set_cm, content_cm)

    def _estimate_content_height_cm(self, row):
        """估算行内容需要的高度（cm）——跨单元格取最大值。

        单元格高度 = Σ 每段（视觉行数 × 单行高度）
        - 视觉行数 = ceil(段落文字视觉宽度 / 单元格净宽)，至少 1
        - 单行高度 pt = 段落最大字号 × 1.2 × max(行距倍数, 1.0)
        """
        import math
        PT_PER_CM = 28.3464567
        DXA_PER_CM = 567.0

        default_pt = 10.5
        try:
            normal_size = self.doc.styles['Normal'].font.size
            if normal_size is not None:
                default_pt = normal_size.pt
        except Exception:
            pass

        def char_units(ch):
            code = ord(ch)
            if (0x4E00 <= code <= 0x9FFF or 0x3400 <= code <= 0x4DBF
                    or 0x3000 <= code <= 0x303F or 0xFF00 <= code <= 0xFFEF):
                return 1.0
            if code < 128:
                return 0.55
            return 1.0

        def cell_inner_width_pt(cell):
            tcPr = cell._tc.tcPr
            if tcPr is None:
                return None
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                return None
            try:
                w_dxa = float(tcW.get(qn('w:w'), '0'))
            except (TypeError, ValueError):
                return None
            if w_dxa <= 0:
                return None
            left_dxa = right_dxa = 108.0
            try:
                tbl = cell._tc.getparent().getparent()
                tblCellMar = tbl.find('.//' + qn('w:tblCellMar'))
                if tblCellMar is not None:
                    l = tblCellMar.find(qn('w:left'))
                    r = tblCellMar.find(qn('w:right'))
                    if l is not None and l.get(qn('w:w')) is not None:
                        left_dxa = float(l.get(qn('w:w')))
                    if r is not None and r.get(qn('w:w')) is not None:
                        right_dxa = float(r.get(qn('w:w')))
            except Exception:
                pass
            return max(0.0, (w_dxa - left_dxa - right_dxa) / 20.0)

        max_cell_cm = 0.0
        try:
            for cell in row.cells:
                inner_w = cell_inner_width_pt(cell)
                cell_total_pt = 0.0
                for para in cell.paragraphs:
                    para_pt = default_pt
                    for run in para.runs:
                        if run.font.size is not None:
                            para_pt = max(para_pt, run.font.size.pt)
                    # 行距倍数
                    line_mul = 1.0
                    pPr = para._p.find(qn('w:pPr'))
                    spacing = pPr.find(qn('w:spacing')) if pPr is not None else None
                    if spacing is not None:
                        line = spacing.get(qn('w:line'))
                        line_rule = spacing.get(qn('w:lineRule'))
                        if line is not None:
                            try:
                                lv = float(line)
                            except ValueError:
                                lv = None
                            if lv is not None:
                                if line_rule == 'exact':
                                    # 精确行距（dxa）：等价行距倍数 = 精确 pt / (字号 × 1.2)
                                    exact_pt = lv / 20.0
                                    denom = para_pt * 1.2
                                    line_mul = exact_pt / denom if denom > 0 else 1.0
                                elif line_rule == 'atLeast':
                                    exact_pt = lv / 20.0
                                    denom = para_pt * 1.2
                                    line_mul = max(1.0, exact_pt / denom if denom > 0 else 1.0)
                                else:  # auto 或未指定
                                    line_mul = lv / 240.0 if lv > 0 else 1.0

                    trimmed = "".join(run.text for run in para.runs).rstrip(" \t　\xa0\r\n")
                    if not trimmed:
                        lines = 1
                    else:
                        text_w = sum(char_units(ch) * para_pt for ch in trimmed)
                        if inner_w and inner_w > 0:
                            lines = max(1, math.ceil(text_w / inner_w))
                        else:
                            lines = 1
                    line_pt = para_pt * 1.2 * line_mul
                    cell_total_pt += lines * line_pt
                cm = cell_total_pt / PT_PER_CM
                if cm > max_cell_cm:
                    max_cell_cm = cm
        except Exception:
            return None

        return max_cell_cm if max_cell_cm > 0 else None

    def _check_images_and_non_editable(self):
        """-5：文中出现图片或者不可编辑的部分"""
        has_images = False

        # 检查段落中的图片
        for para in self.doc.paragraphs:
            for run in para.runs:
                # 检查是否有图片
                if hasattr(run, '_element'):
                    for pic in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        has_images = True
                        break
                if has_images:
                    break
            if has_images:
                break

        # 检查表格中的图片
        if not has_images:
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if hasattr(run, '_element'):
                                    for pic in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                                        has_images = True
                                        break
                                if has_images:
                                    break
                            if has_images:
                                break
                        if has_images:
                            break
                    if has_images:
                        break
                if has_images:
                    break

        if has_images:
            self.dimension2_score -= 5
            self.dimension2_details.append("-5：文中出现图片或者不可编辑的部分 ✗")
        else:
            self.dimension2_details.append("-5：文中出现图片或者不可编辑的部分 ✓")

    def _check_first_page_header(self):
        """-5：文档第一页顶部不满足其中任意一项：文档标题为"部门协调会"；
        主持人为许承柏；地点为综合楼二层会议室

        标准文档版式中标题独占一段（B2），随后是"时间/主持人/记录人"一行
        （B4），再是"地点/应出席数/实出席数"一行（B6），均在文档最前面
        （表格/会议内容之前）。此前的实现只对前10个段落做整体拼接后的
        子串包含匹配：既不能确认标题是否精确为"部门协调会"（可能只是
        某个较长句子里包含这四个字），也不能确认"主持人"字段本身存在
        （"许承柏"可能出现在正文其它位置而不是主持人字段）。
        改为：只在文档最前面若干段落（表格/会议内容开始之前）中，用
        "字段标签 + 分隔符 + 值"的正则做字段级精确匹配；标题段落要求
        strip 后原文本严格等于"部门协调会"。
        """
        # 字段标签内的字符之间可能插入全角/半角空格（如 "主  持  人"），
        # 用 [\s]* 把标签逐字符间的空白容忍掉。
        def label_pattern(label):
            return r'[\s]*'.join(re.escape(ch) for ch in label)

        # 顶部区域：取表格/会议内容之前的普通段落（顶部字段不可能出现在
        # 表格里），并限定在文档开头一个较小的窗口内，避免匹配到正文中
        # 出现的同名字符串。
        top_paragraphs = []
        for para in self.doc.paragraphs:
            if len(top_paragraphs) >= 15:
                break
            top_paragraphs.append(para.text)

        has_title = any(p.strip() == "部门协调会" for p in top_paragraphs)

        host_re = re.compile(label_pattern("主持人") + r'[\s:：]+许承柏')
        has_host = any(host_re.search(p) for p in top_paragraphs)

        location_re = re.compile(label_pattern("地点") + r'[\s:：]+综合楼二层会议室')
        has_location = any(location_re.search(p) for p in top_paragraphs)

        if has_title and has_host and has_location:
            self.dimension2_details.append("-5：文档第一页顶部不满足其中任意一项 ✓")
        else:
            self.dimension2_score -= 5
            missing = []
            if not has_title:
                missing.append("文档标题为'部门协调会'")
            if not has_host:
                missing.append("主持人为许承柏")
            if not has_location:
                missing.append("地点为综合楼二层会议室")
            missing_str = ", ".join(missing)
            self.dimension2_details.append(f"-5：文档第一页顶部不满足其中任意一项 ✗（缺少: {missing_str}）")

    def _check_meeting_content_indent(self):
        """-1："会议内容："下方第一行出现缩进

        - 冒号兼容全角"："(U+FF1A) 与半角":"(U+003A)，允许"会议内容"与冒号之间有空白；
        - "下方第一行"按 body 元素的版面顺序取：紧挨着的下一段；若紧跟的是表格，则
          取该表格第 1 行第 1 列的第一个非空段落。
        """
        if self.doc is None:
            return

        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        marker_re = re.compile(r'会议内容\s*[:：]')

        body = self.doc.element.body
        children = list(body)
        # Paragraph/Table 需要一个 parent 来定位 story part，正文段落的 parent 就是文档 body
        parent = self.doc.paragraphs[0]._parent if self.doc.paragraphs else self.doc.element.body

        found_marker = False
        next_line_indent = False
        marker_idx = -1

        for idx, child in enumerate(children):
            if child.tag == qn('w:p'):
                text = ''.join((t.text or '') for t in child.iter(qn('w:t')))
                if marker_re.search(text):
                    found_marker = True
                    marker_idx = idx
                    break

        def _has_indent(para):
            fli = para.paragraph_format.first_line_indent
            li = para.paragraph_format.left_indent
            return (fli is not None and fli.pt > 0) or (li is not None and li.pt > 0)

        if found_marker:
            for j in range(marker_idx + 1, len(children)):
                nxt = children[j]
                if nxt.tag == qn('w:p'):
                    next_para = Paragraph(nxt, parent)  # type: ignore[arg-type]
                    # 空段落不算"下一行"，继续向后找
                    if not next_para.text.strip():
                        continue
                    if _has_indent(next_para):
                        next_line_indent = True
                    break
                elif nxt.tag == qn('w:tbl'):
                    tbl = Table(nxt, parent)  # type: ignore[arg-type]
                    first_para = None
                    if tbl.rows and tbl.rows[0].cells:
                        for p in tbl.rows[0].cells[0].paragraphs:
                            if p.text.strip():
                                first_para = p
                                break
                        if first_para is None and tbl.rows[0].cells[0].paragraphs:
                            # 单元格全空，退而取第一个段落
                            first_para = tbl.rows[0].cells[0].paragraphs[0]
                    if first_para is not None and _has_indent(first_para):
                        next_line_indent = True
                    break

        if found_marker and next_line_indent:
            self.dimension2_score -= 1
            self.dimension2_details.append('-1："会议内容："下方第一行出现缩进 ✗')
        elif not found_marker:
            self.dimension2_details.append("-1：未找到会议内容标记，跳过检查")
        else:
            self.dimension2_details.append('-1："会议内容："下方第一行出现缩进 ✓')

    def _iter_all_text(self):
        """收集文档全部文字（含嵌套表格）。python-docx 的 doc.tables 与 cell.text
        仅覆盖顶层表格与非嵌套内容，直接拼接会漏掉表中表内的文字。"""
        if self.doc is None:
            return ""
        buf = []
        for para in self.doc.paragraphs:
            buf.append(para.text)

        def walk_table(tbl):
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        buf.append(para.text)
                    for sub in cell.tables:
                        walk_table(sub)

        for tbl in self.doc.tables:
            walk_table(tbl)
        return "\n".join(buf)

    def _check_hou_jinglan_text(self):
        """-1：侯景澜文本内容检查"""
        all_text = self._iter_all_text()

        target_text = "侯景澜：上月两项重点现场事项整改推进如下："
        if target_text in all_text:
            self.dimension2_details.append('-1：文中没有出现"侯景澜：上月两项重点现场事项整改推进如下："文本内容 ✓')
        else:
            self.dimension2_score -= 1
            self.dimension2_details.append('-1：文中没有出现"侯景澜：上月两项重点现场事项整改推进如下："文本内容 ✗')

    def _check_indentation_between_sections(self):
        """-1：区间缩进检查"""
        # 获取所有文本内容
        all_content = []
        for para in self.doc.paragraphs:
            all_content.append(('paragraph', para.text, para))
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_content.append(('cell', para.text, para))

        # 查找两个标记之间的内容
        start_text = "二月份作业偏差共发生若干件，具体情况如下："
        end_text = "四、研究后续事项"

        found_start = False
        found_end = False
        has_indentation = False

        for content_type, text, para in all_content:
            if not found_start:
                if start_text in text:
                    found_start = True
                continue

            if found_start and not found_end:
                if end_text in text or "四、研究后续事项" in text:
                    found_end = True
                    break

                # 检查此段落是否有缩进
                if text.strip():  # 只检查非空段落
                    first_line_indent = para.paragraph_format.first_line_indent
                    left_indent = para.paragraph_format.left_indent

                    if (first_line_indent is not None and first_line_indent.pt > 0) or \
                       (left_indent is not None and left_indent.pt > 0):
                        has_indentation = True

        if found_start and found_end and has_indentation:
            self.dimension2_score -= 1
            self.dimension2_details.append('-1："二月份作业偏差共发生若干件，具体情况如下："与"四、研究后续事项"之间的文本出现缩进 ✗')
        elif not found_start or not found_end:
            self.dimension2_details.append("-1：未找到指定的章节标记，跳过检查")
        else:
            self.dimension2_details.append('-1："二月份作业偏差共发生若干件，具体情况如下："与"四、研究后续事项"之间的文本出现缩进 ✓')

    def _check_xu_chengbai_text_format(self):
        """已废弃：该评分细则已删除，保留空实现以避免外部误调用。"""
        return

    def evaluate(self):
        """执行完整评估"""
        # 维度1评估
        if self.load_document():
            self.evaluate_dimension1()

        # 维度1未通过：仅输出维度1详情
        if not self.dimension1_passed:
            print("维度一：未通过")
            for detail in self.dimension1_details:
                print(detail)
            return self.dimension1_passed, self.dimension2_score, self.dimension1_details, self.dimension2_details

        # 维度1通过
        print("维度一：通过")

        # 维度2评估：仅打印命中项（+项达标 或 -项违反）
        self.evaluate_dimension2()
        print("维度二：评分结果")

        hit_pattern = re.compile(r'^([+-]\d+)：([^✓✗]+?)\s*[✓✗]')
        for detail in self.dimension2_details:
            is_plus_hit = detail.startswith('+') and '✓' in detail
            is_minus_hit = detail.startswith('-') and '✗' in detail
            if not (is_plus_hit or is_minus_hit):
                continue
            m = hit_pattern.match(detail)
            if m:
                score, rule = m.group(1), m.group(2).strip()
                print(f"{score}：{rule}")
            else:
                # 兜底：原样输出（去掉末尾说明）
                print(detail)

        print(f"总分：{self.dimension2_score}")

        return self.dimension1_passed, self.dimension2_score, self.dimension1_details, self.dimension2_details


def _build_dim2_items(details):
    """把 dimension2_details 里的可读文本解析为结构化条目。

    保留原有 `_check_*` 生成文本的逻辑不动，此处只做二次解析：
    - `+N`：✓ 计为命中（得分 +N）；✗ 计为未命中（0）
    - `-N`：✗ 计为命中（得分 -N，即触发扣分）；✓ 计为未命中（0）
    - "未找到..跳过检查" 等无标记行，视为未命中且带上"跳过检查"说明
    """
    hit_pattern = re.compile(r'^([+-])(\d+)：(.+?)\s*(✓|✗)\s*(?:（(.*?)）)?\s*$')
    skip_pattern = re.compile(r'^([+-])(\d+)：(.+)$')
    items = []
    for line in details:
        m = hit_pattern.match(line)
        if m:
            sign, num, rule, mark, extra = m.groups()
            max_delta = int(num) if sign == '+' else -int(num)
            hit = (mark == '✓') if max_delta >= 0 else (mark == '✗')
            items.append({
                "rule": rule.strip(),
                "max_delta": max_delta,
                "delta": max_delta if hit else 0,
                "hit": hit,
                "detail": "",
            })
            continue
        m2 = skip_pattern.match(line)
        if m2:
            sign, num, rest = m2.groups()
            max_delta = int(num) if sign == '+' else -int(num)
            items.append({
                "rule": rest.strip(),
                "max_delta": max_delta,
                "delta": 0,
                "hit": False,
                "detail": "",
            })
    return items


def evaluate(dir_path: str):
    """统一评估入口。

    参数:
        dir_path: 脚本所在目录的路径。脚本自行在该目录中定位并打开被评估文档。

    返回:
        结构化评估结果字典，字段规范见"脚本接口差异与统一建议.md §2.2"。
    """
    result = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }
    try:
        if not os.path.isdir(dir_path):
            raise FileNotFoundError(f"目录不存在: {dir_path}")

        # 在目录中定位被评估文档：优先精确文件名，否则回退到目录内首个 .docx
        target = os.path.join(dir_path, EXPECTED_DOC_NAME)
        if not os.path.isfile(target):
            candidates = sorted(
                f for f in os.listdir(dir_path)
                if f.lower().endswith(".docx") and not f.startswith("~$")
            )
            if not candidates:
                raise FileNotFoundError(f"目录中未找到 .docx 文件: {dir_path}")
            target = os.path.join(dir_path, candidates[0])

        result["file_name"] = os.path.basename(target)

        evaluator = DocumentEvaluator(target)
        if evaluator.load_document():
            evaluator.evaluate_dimension1()

        result["dim1_pass"] = evaluator.dimension1_passed
        if not evaluator.dimension1_passed:
            fail_msgs = [d for d in evaluator.dimension1_details if d.startswith("✗")]
            result["dim1_reason"] = "; ".join(fail_msgs) if fail_msgs else "维度一未通过"
            return result

        # 维度一通过：评估维度二
        evaluator.evaluate_dimension2()
        items = _build_dim2_items(evaluator.dimension2_details)
        result["dim2_items"] = items
        # 满分 = 所有 +N 项之和；扣分项不计入满分
        result["max_score"] = sum(it["max_delta"] for it in items if it["max_delta"] > 0)
        result["total_score"] = evaluator.dimension2_score
        return result
    except Exception as e:
        result["status"] = "error"
        result["error"] = f"{type(e).__name__}: {e}"
        return result


if __name__ == "__main__":
    # 本地自测：默认使用脚本所在目录；也可通过命令行显式传入目录
    _dir = sys.argv[1] if len(sys.argv) >= 2 else os.path.dirname(os.path.abspath(__file__))
    print(json.dumps(evaluate(_dir), ensure_ascii=False, indent=2))
