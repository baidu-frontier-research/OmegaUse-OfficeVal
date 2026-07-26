# -*- coding: utf-8 -*-
"""
自动评估脚本：对两个Word文档按打分细则进行评分
维度一：可用与可修改性（不满足则直接判零分）
维度二：完成度评分细则（逐项检查得分/扣分）
"""

import os

SCRIPT_ID = "022"
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 字号对照表 (EMU值)
FONT_SIZE = {
    '二号': 279400, '小二': 228600, '三号': 203200, '小三': 190500,
    '四号': 177800, '小四': 152400, '五号': 133350, '小五': 114300,
}
SIZE_TOLERANCE = 6400  # 约0.5pt容差


def _resolve_docs(dir_path):
    """在给定目录下定位两个被评估文档，返回 (file1, file2) 绝对路径。"""
    file1 = os.path.join(dir_path, '企业经营情况说明书_按模板格式添加封面目录.docx')
    file2 = os.path.join(dir_path, '华中区域零售项目经营资金异常及现金流风险_主题修改版.docx')
    return file1, file2


def size_match(actual, expected_name):
    """检查字体大小是否匹配"""
    if actual is None:
        return False
    expected = FONT_SIZE.get(expected_name, 0)
    return abs(actual - expected) <= SIZE_TOLERANCE


def get_east_asian_font(run):
    """获取run的中文字体名"""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:eastAsia'), '')
    return ''


def get_ascii_font(run):
    """获取run的英文字体名"""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:ascii'), '')
    return ''


def get_effective_ascii_font(run, paragraph):
    """获取有效英文字体（run直设 → 段落样式 → 样式链继承）"""
    ascii_f = get_ascii_font(run)
    if ascii_f:
        return ascii_f
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        rPr = style.element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                v = rFonts.get(qn('w:ascii'))
                if v:
                    return v
        style = style.base_style
    return ''


def get_effective_east_asian_font(run, paragraph):
    """获取有效中文字体（run直设 → 段落样式 → 样式链继承）——办公软件按此顺序解析"""
    ea = get_east_asian_font(run)
    if ea:
        return ea
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        rPr = style.element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                v = rFonts.get(qn('w:eastAsia'))
                if v:
                    return v
        style = style.base_style
    return ''


def get_effective_size(run, paragraph):
    """获取有效字号EMU（run直设 → 段落样式 → 样式链继承）"""
    if run.font.size is not None:
        return run.font.size
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        rPr = style.element.find(qn('w:rPr'))
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                if val:
                    # w:sz 单位为half-point，1pt=12700 EMU
                    return int(val) * 12700 // 2
        style = style.base_style
    return None


def get_effective_bold(run, paragraph):
    """获取有效加粗（run直设 → run 字符样式链 → 段落样式链继承）——
    与 Word/WPS 里"选中字符 → 加粗按钮亮起"的判定顺序一致。"""
    # 1) run 直设 <w:b/>
    if run.font.bold is not None:
        return bool(run.font.bold)

    def _check_style_chain(style_obj):
        """沿 style → basedOn 链找 w:b；命中返回 True/False，未命中返回 None。"""
        seen = set()
        s = style_obj
        while s is not None and s.style_id not in seen:
            seen.add(s.style_id)
            rPr = s.element.find(qn('w:rPr'))
            if rPr is not None:
                b = rPr.find(qn('w:b'))
                if b is not None:
                    val = b.get(qn('w:val'))
                    if val is None or val in ('1', 'true', 'on'):
                        return True
                    return False
            s = s.base_style
        return None

    # 2) run 的字符样式（w:rStyle）链——Word/WPS 里"Strong"这类字符样式在此层生效
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rStyle = rPr.find(qn('w:rStyle'))
        if rStyle is not None:
            style_id = rStyle.get(qn('w:val'))
            if style_id:
                char_style = None
                try:
                    for _s in paragraph.part.document.styles:
                        if getattr(_s, 'style_id', None) == style_id:
                            char_style = _s
                            break
                except Exception:
                    char_style = None
                if char_style is not None:
                    r = _check_style_chain(char_style)
                    if r is not None:
                        return r

    # 3) 段落样式链
    r = _check_style_chain(paragraph.style)
    if r is not None:
        return r
    return False


def get_first_page_paragraphs(paragraphs):
    """返回第1页范围内的段落（首个分页符/pageBreakBefore之前）——对应办公软件里第1页所见内容"""
    result = []
    for i, p in enumerate(paragraphs):
        # pageBreakBefore：本段前分页，说明本段已属新页
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None and i > 0:
            return result
        result.append(p)
        # 段落内包含 w:br type=page，之后内容属下一页
        for run in p.runs:
            for elem in run._element.iter():
                if elem.tag == qn('w:br') and elem.get(qn('w:type')) == 'page':
                    return result
    return result


def get_paragraphs_by_page(paragraphs):
    """按 pageBreakBefore / w:br type=page 将段落切分为按页分组的列表列表——
    对应办公软件里各页的实际内容顺序。"""
    pages = [[]]
    for i, p in enumerate(paragraphs):
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None and i > 0:
            pages.append([])
        pages[-1].append(p)
        # 段落内包含 w:br type=page → 之后新起一页
        broke = False
        for run in p.runs:
            for elem in run._element.iter():
                if elem.tag == qn('w:br') and elem.get(qn('w:type')) == 'page':
                    broke = True
                    break
            if broke:
                break
        if broke and i != len(paragraphs) - 1:
            pages.append([])
    return pages


def is_single_line_spacing(para):
    """检查是否单倍行距"""
    pf = para.paragraph_format
    if pf.line_spacing is None:
        return False
    if pf.line_spacing == 1.0:
        return True
    # line_spacing_rule SINGLE
    from docx.enum.text import WD_LINE_SPACING
    if pf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
        return True
    return False


def is_effective_single_line_spacing(para):
    """按办公软件行距解析规则判断单倍行距：段落直设 → 段落样式 → 样式链继承。
    单倍行距在 OOXML 里表达为：w:spacing 无 line 属性，或 line=240 且 lineRule=auto，
    或 lineRule=auto 且 line=240（Word默认240=单倍）。"""
    from docx.enum.text import WD_LINE_SPACING

    # 1) 段落直设
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        if pf.line_spacing == 1.0:
            return True
        if pf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
            return True
        return False

    # 2) 沿样式链回溯（对应办公软件里"继承样式"的行距生效逻辑）
    style = para.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                line = spacing.get(qn('w:line'))
                line_rule = spacing.get(qn('w:lineRule'))
                if line is not None:
                    # lineRule=auto 且 line=240 表示 1.0 倍（单倍）
                    if (line_rule is None or line_rule == 'auto') and line == '240':
                        return True
                    return False
        style = style.base_style

    # 3) 样式链未指定 → Word 默认为单倍
    return True


def is_effective_center_aligned(para):
    """按办公软件对齐解析规则判断居中：段落直设 → 段落样式 → 样式链继承。"""
    if para.alignment is not None:
        return para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    style = para.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                return jc.get(qn('w:val')) == 'center'
        style = style.base_style
    return False


def is_center_aligned(para):
    """检查居中对齐"""
    return para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def is_effective_justify_aligned(para):
    """按办公软件对齐解析规则判断"两端对齐"：段落直设 → 段落样式 → 样式链继承。
    Word/WPS 里"两端对齐"对应 OOXML `w:jc w:val="both"`（python-docx 映射为 WD_ALIGN_PARAGRAPH.JUSTIFY）。
    段落链上都未指定时，Word 默认对齐为"左对齐"，不算两端对齐。"""
    # 1) 段落直设
    if para.alignment is not None:
        return para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    # 2) 沿样式链回溯
    style = para.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                return jc.get(qn('w:val')) == 'both'
        style = style.base_style
    # 3) 默认非两端对齐
    return False


def is_effective_line_spacing_1_5(para):
    """按办公软件行距解析规则判断"1.5 倍行距"：段落直设 → 段落样式 → 样式链继承。
    OOXML 里 1.5 倍行距表达为 `w:spacing w:line="360" w:lineRule="auto"`
    （240 = 单倍，360 = 1.5 倍）。python-docx 中 line_spacing==1.5 亦表示 1.5 倍。"""
    from docx.enum.text import WD_LINE_SPACING

    # 1) 段落直设
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        # python-docx: 浮点数即倍数
        try:
            return abs(float(pf.line_spacing) - 1.5) < 1e-6
        except (TypeError, ValueError):
            return False
    if pf.line_spacing_rule is not None:
        # 明确设了非 multiple 规则（EXACTLY/AT_LEAST 等）→ 非 1.5 倍
        if pf.line_spacing_rule != WD_LINE_SPACING.MULTIPLE:
            return False

    # 2) 沿样式链回溯
    style = para.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                line = spacing.get(qn('w:line'))
                line_rule = spacing.get(qn('w:lineRule'))
                if line is not None:
                    if (line_rule is None or line_rule == 'auto') and line == '360':
                        return True
                    return False
        style = style.base_style

    # 3) 样式链未指定 → 非 1.5 倍
    return False


def _iter_para_style_pPr(para):
    """按办公软件属性生效顺序，产出：段落直设 pPr → 段落样式 pPr → 沿 basedOn 上溯的父样式 pPr。
    调用方在第一次读取到目标属性即返回，与 Word/WPS 属性解析行为一致。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        yield pPr
    style = para.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        sPPr = style.element.find(qn('w:pPr'))
        if sPPr is not None:
            yield sPPr
        style = style.base_style


def is_effective_left_aligned(para):
    """Word/WPS 里"左对齐"= `w:jc w:val="left"`；若段落及样式链均未设 w:jc，
    默认对齐为左对齐（对应 python-docx 的 alignment=None）。"""
    # 1) 段落直设
    if para.alignment is not None:
        return para.alignment == WD_ALIGN_PARAGRAPH.LEFT
    # 2) 沿样式链回溯
    style = para.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                return jc.get(qn('w:val')) == 'left'
        style = style.base_style
    # 3) 默认左对齐
    return True


def is_effective_right_aligned(para):
    """Word/WPS 里"右对齐"= `w:jc w:val="right"`。effective 解析：段落直设 → 样式链。"""
    # 1) 段落直设
    if para.alignment is not None:
        return para.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    # 2) 沿样式链回溯
    style = para.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                return jc.get(qn('w:val')) == 'right'
        style = style.base_style
    # 3) 默认非右对齐
    return False


def get_effective_first_line_indent_emu(para):
    """按办公软件"首行缩进"生效顺序解析：段落直设 → 样式链。返回 EMU（Length 兼容 int），
    未设置返回 None。OOXML: `w:ind w:firstLine="<twips>"`（1twip=635EMU；1cm=567twips）。"""
    pf = para.paragraph_format
    if pf.first_line_indent is not None:
        return int(pf.first_line_indent)
    # 段落直设 pPr（python-docx 已覆盖，兜底再查一次）
    for pPr in _iter_para_style_pPr(para):
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            fl = ind.get(qn('w:firstLine'))
            if fl is not None:
                try:
                    return int(fl) * 635  # twips → EMU
                except ValueError:
                    return None
            # 若设了 hanging（悬挂缩进），首行缩进按 0 计
            hg = ind.get(qn('w:hanging'))
            if hg is not None:
                return 0
    return None


def get_effective_line_spacing(para):
    """按办公软件"行距"生效顺序解析：返回 (mode, value)。
      mode: 'multiple'（倍数）、'exact'（固定值，单位=磅）、'at_least'（最小值，单位=磅）
      未设置返回 (None, None)。用于严判 1.25 倍行距。"""
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        from docx.enum.text import WD_LINE_SPACING
        rule = pf.line_spacing_rule
        if rule == WD_LINE_SPACING.EXACTLY:
            return ('exact', float(pf.line_spacing) / 12700.0)
        if rule == WD_LINE_SPACING.AT_LEAST:
            return ('at_least', float(pf.line_spacing) / 12700.0)
        # SINGLE / ONE_POINT_FIVE / DOUBLE / MULTIPLE → 倍数
        return ('multiple', float(pf.line_spacing))
    # 沿样式链
    for pPr in _iter_para_style_pPr(para):
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            line = spacing.get(qn('w:line'))
            line_rule = spacing.get(qn('w:lineRule'))
            if line is not None:
                try:
                    line_val = int(line)
                except ValueError:
                    continue
                if line_rule == 'exact':
                    return ('exact', line_val / 20.0)  # twentieths of a point
                if line_rule == 'atLeast':
                    return ('at_least', line_val / 20.0)
                # auto / 未指定 → 倍数（240=1.0）
                return ('multiple', line_val / 240.0)
    return (None, None)


def get_effective_space_before_pt(para):
    """按办公软件"段前"生效顺序解析：段落直设 → 样式链。返回磅值，未设返回 None。"""
    pf = para.paragraph_format
    if pf.space_before is not None:
        return float(pf.space_before) / 12700.0
    for pPr in _iter_para_style_pPr(para):
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            before = spacing.get(qn('w:before'))
            if before is not None:
                try:
                    return int(before) / 20.0  # twentieths of a point
                except ValueError:
                    return None
    return None


def get_effective_space_after_pt(para):
    """按办公软件"段后"生效顺序解析：段落直设 → 样式链。返回磅值，未设返回 None。"""
    pf = para.paragraph_format
    if pf.space_after is not None:
        return float(pf.space_after) / 12700.0
    for pPr in _iter_para_style_pPr(para):
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            after = spacing.get(qn('w:after'))
            if after is not None:
                try:
                    return int(after) / 20.0
                except ValueError:
                    return None
    return None


def has_page_break_before(para_idx, paragraphs):
    """检查段落前是否有分页符"""
    if para_idx == 0:
        return True
    p = paragraphs[para_idx]
    # 检查前一段是否有page break
    prev = paragraphs[para_idx - 1]
    for run in prev.runs:
        for elem in run._element:
            if elem.tag == qn('w:br'):
                if elem.get(qn('w:type')) == 'page':
                    return True
    return False


def get_separator_length_cm(paragraph, next_paragraph=None, section=None):
    """返回该段落"下方"分隔虚线在办公软件中可见的长度(厘米)；找不到则返回 None。
    支持两种实现（与 has_visible_separator_below 一致）：
      (a) 段落底部边框：边框沿段落文本区宽度延伸 = 页面宽 - 左右页边距 - 段落左右缩进
      (b) 字符模拟的分隔线：字符个数 × 每字符宽度（全角字符宽度 ≈ 字号磅值 × 0.0353 cm）
      (b') 下一段落顶部边框：同 (a)，但使用下一段落自身缩进
    """
    def _para_line_cm(p, sec):
        if sec is None:
            return None
        page_w = sec.page_width or 0
        left_m = sec.left_margin or 0
        right_m = sec.right_margin or 0
        pf = p.paragraph_format
        left_i = pf.left_indent or 0
        right_i = pf.right_indent or 0
        emu = page_w - left_m - right_m - left_i - right_i
        return emu / 360000.0  # 1cm = 360000 EMU

    # (a) 当前段落底部边框
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            bottom = pBdr.find(qn('w:bottom'))
            if bottom is not None and bottom.get(qn('w:val')) not in (None, 'nil', 'none'):
                return _para_line_cm(paragraph, section)

    if next_paragraph is None:
        return None

    # (b') 下一段落顶部边框
    npPr = next_paragraph._element.find(qn('w:pPr'))
    if npPr is not None:
        npBdr = npPr.find(qn('w:pBdr'))
        if npBdr is not None:
            top = npBdr.find(qn('w:top'))
            if top is not None and top.get(qn('w:val')) not in (None, 'nil', 'none'):
                return _para_line_cm(next_paragraph, section)

    # (b) 字符模拟：字符个数 × 每字符宽度(cm)
    sep_chars = set('━─—‒═⋯┈┅┄╌╍╴╶╸╺')
    text = next_paragraph.text.strip()
    if text and any(c in sep_chars for c in text):
        count = sum(1 for c in text if c in sep_chars)
        # 取首个分隔字符所在 run 的字号（办公软件按 run 字号渲染字宽）
        pt = None
        for r in next_paragraph.runs:
            if any(c in sep_chars for c in r.text):
                sz = get_effective_size(r, next_paragraph)
                if sz is not None:
                    pt = sz / 12700.0
                    break
        if pt is None:
            return None
        # 全角字符在 CJK 字体下宽度约等于字号(pt)，1pt = 0.0353cm
        char_w_cm = pt * 0.03527778
        return count * char_w_cm
    return None


def get_separator_width_pt(paragraph, next_paragraph=None):
    """返回该段落"下方"分隔虚线在办公软件中可见的宽度(磅)；找不到则返回 None。
    支持两种实现：
      (a) 段落底部边框 w:pBdr/w:bottom：w:sz 单位为 1/8 磅
      (b) 下一段落是分隔线：
          - 顶部段落边框：w:sz/8 磅
          - 字符模拟：由字符类型和字体字号推算笔画宽度
    """
    # (a) 当前段落底部边框
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            bottom = pBdr.find(qn('w:bottom'))
            if bottom is not None and bottom.get(qn('w:val')) not in (None, 'nil', 'none'):
                sz = bottom.get(qn('w:sz'))
                if sz is not None:
                    return int(sz) / 8.0

    if next_paragraph is None:
        return None

    # (b1) 下一段落顶部边框
    npPr = next_paragraph._element.find(qn('w:pPr'))
    if npPr is not None:
        npBdr = npPr.find(qn('w:pBdr'))
        if npBdr is not None:
            top = npBdr.find(qn('w:top'))
            if top is not None and top.get(qn('w:val')) not in (None, 'nil', 'none'):
                sz = top.get(qn('w:sz'))
                if sz is not None:
                    return int(sz) / 8.0

    # (b2) 字符模拟：按字符 Unicode 权重 × 字体字号 推算可见笔画宽度
    # 参考 Unicode Box Drawing 权重：LIGHT 约 0.06em，HEAVY 约 0.15em（在典型 CJK 字体下渲染值）
    sep_stroke_ratio = {
        '─': 0.06, '—': 0.06, '‒': 0.06,   # LIGHT / EM DASH
        '━': 0.15,                           # HEAVY
        '═': 0.06,                           # DOUBLE（每条约 0.06em）
    }
    text = next_paragraph.text.strip()
    for ch in text:
        if ch in sep_stroke_ratio:
            # 取该字符所在 run 的字号
            for r in next_paragraph.runs:
                if ch in r.text:
                    sz_emu = get_effective_size(r, next_paragraph)
                    if sz_emu is None:
                        return None
                    font_pt = sz_emu / 12700.0
                    return font_pt * sep_stroke_ratio[ch]
            break
    return None


def has_visible_separator_below(paragraph, next_paragraph=None):
    """判断某段落"下方"在办公软件里是否可见一条文档分隔虚线。
    办公软件里可见的实现方式有二：
      (a) 该段落自身设置了底部边框 w:pPr/w:pBdr/w:bottom（非 nil）；
      (b) 紧接的下一段落包含分隔线字符（━ ─ — ‒ ═ ⋯ ┈ ┅ 等），或该段落设置了顶部边框。
    只要满足其一即视为"下方有分隔虚线"。"""
    # (a) 当前段落底部边框
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            bottom = pBdr.find(qn('w:bottom'))
            if bottom is not None and bottom.get(qn('w:val')) not in (None, 'nil', 'none'):
                return True

    if next_paragraph is None:
        return False

    # (b1) 下一段落是字符模拟的分隔线
    sep_chars = set('━─—‒═⋯┈┅┄╌╍╴╶╸╺')
    text = next_paragraph.text.strip()
    if text and all((c in sep_chars or c.isspace()) for c in text) and any(c in sep_chars for c in text):
        return True

    # (b2) 下一段落顶部边框
    npPr = next_paragraph._element.find(qn('w:pPr'))
    if npPr is not None:
        npBdr = npPr.find(qn('w:pBdr'))
        if npBdr is not None:
            top = npBdr.find(qn('w:top'))
            if top is not None and top.get(qn('w:val')) not in (None, 'nil', 'none'):
                return True

    return False


def check_separator_line(para):
    """检查是否为分隔虚线（用━字符模拟的线）"""
    text = para.text.strip()
    if '━' in text or '─' in text or '—' in text:
        return True
    # 也检查是否有pBdr (paragraph border)
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            return True
    return False


def check_separator_width(para, target_pt=1.0):
    """检查分隔线宽度 - 对于文字模拟的线，通过字号近似判断"""
    # 用━字符模拟时，检查其是否存在即可（宽度由字体决定）
    for r in para.runs:
        if '━' in r.text or '─' in r.text:
            return True
    return False


def check_separator_length(para, min_cm=7.7, max_cm=8.1):
    """检查分隔线长度 - 对于━字符模拟，严格估算长度"""
    text = para.text.strip()
    dash_chars = [c for c in text if c in '━─']
    if not dash_chars:
        return False
    count = len(dash_chars)
    for r in para.runs:
        if '━' in r.text:
            # ━是全角字符，宽度约等于字号(pt转cm: 1pt=0.0353cm)
            font_size_pt = (r.font.size or 133350) / 12700
            char_width_cm = font_size_pt * 0.0353
            total_cm = count * char_width_cm
            return min_cm <= total_cm <= max_cm
    return False


def _border_is_absent(el):
    """OOXML 里"无边框"的表达：元素不存在，或 val ∈ {nil, none}，或 sz=0。"""
    if el is None:
        return True
    val = el.get(qn('w:val'))
    if val in ('nil', 'none'):
        return True
    sz = el.get(qn('w:sz'))
    if sz is not None:
        try:
            if int(sz) == 0:
                return True
        except ValueError:
            pass
    return False


def _border_sz(el):
    """取边框粗细 w:sz（1/8 磅），未设或缺失时返回 None（视作无边框）。"""
    if el is None:
        return None
    val = el.get(qn('w:val'))
    if val in ('nil', 'none'):
        return None
    sz = el.get(qn('w:sz'))
    if sz is None:
        return None
    try:
        return int(sz)
    except ValueError:
        return None


def _get_effective_border(cell, edge, table_borders):
    """返回单元格某边（top/bottom/left/right）在办公软件里的有效边框元素。
    生效顺序：tcBorders.{edge} → 表格 tblBorders 中的对应边（外沿用 top/bottom/left/right，
    内沿在此函数中不使用——由调用方按位置决定选 insideH/insideV）。返回 (element_or_None)。"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            e = tcBorders.find(qn('w:' + edge))
            if e is not None:
                return e
    if table_borders is not None:
        return table_borders.find(qn('w:' + edge))
    return None


def _get_effective_inside_border(cell, direction, table_borders):
    """内沿边框：tcBorders 里可能显式设 insideH/insideV，否则从 tblBorders 继承。"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            e = tcBorders.find(qn('w:' + direction))
            if e is not None:
                return e
    if table_borders is not None:
        return table_borders.find(qn('w:' + direction))
    return None


def check_three_line_table(table):
    """三线表严判：
      ① 表格上边框 = 1.5 磅（w:sz=12）——第一行每个单元格顶边生效为 sz=12；
      ② 表格下边框 = 1.5 磅（w:sz=12）——最后一行每个单元格底边生效为 sz=12；
      ③ 第一行下框线 = 0.5 磅（w:sz=4）——首行每个单元格底边（同时是第二行顶边）生效为 sz=4；
      ④ 其余位置无框线：单元格左/右边、非首行的顶边、非末行的底边（除首行底边外）均无边框。
    OOXML 里边框可写在 tblBorders（表格默认）或 tcBorders（单元格覆盖），
    与办公软件生效顺序一致：优先 tcBorders，其次 tblBorders（外沿用 top/bottom/left/right，
    内沿用 insideH/insideV）。"""
    rows = table.rows
    if len(rows) < 1:
        return False

    # 取表格默认 tblBorders
    tblPr = table._element.find(qn('w:tblPr'))
    tblBorders = tblPr.find(qn('w:tblBorders')) if tblPr is not None else None

    n_rows = len(rows)
    for ri, row in enumerate(rows):
        cells = row.cells
        n_cols = len(cells)
        # 处理"合并单元格"时同一 tc 会被重复暴露：用 tc 元素做去重
        seen_tcs = set()
        for ci, cell in enumerate(cells):
            if cell._tc in seen_tcs:
                continue
            seen_tcs.add(cell._tc)

            # ---- 顶边 ----
            if ri == 0:
                # ① 首行顶边 = sz 12
                top = _get_effective_border(cell, 'top', tblBorders)
                if _border_sz(top) != 12:
                    return False
            else:
                # 非首行顶边：应为无边框（等效于上一行底边——由上一行判定，这里只查 tcBorders 直设不冲突即可）
                # 直接检查生效顶边：tcBorders.top 存在且非 nil → 违反
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        top = tcBorders.find(qn('w:top'))
                        if not _border_is_absent(top):
                            return False
                # 内水平线（insideH）也要判：非首行的顶边等价于上一行底边或 insideH，
                # 由"上一行底边"判定环节处理。

            # ---- 底边 ----
            if ri == n_rows - 1:
                # ② 末行底边 = sz 12
                bottom = _get_effective_border(cell, 'bottom', tblBorders)
                if _border_sz(bottom) != 12:
                    return False
            elif ri == 0 and n_rows > 1:
                # ③ 首行底边 = sz 4（此边同时是第二行顶边——办公软件里就是"表头分隔线"）
                # 生效顺序：tcBorders.bottom → tblBorders.insideH
                tc = cell._tc
                bottom_el = None
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        bottom_el = tcBorders.find(qn('w:bottom'))
                if bottom_el is None and tblBorders is not None:
                    bottom_el = tblBorders.find(qn('w:insideH'))
                if _border_sz(bottom_el) != 4:
                    return False
            else:
                # ④ 中间行底边：应为无边框（生效顺序：tcBorders.bottom → tblBorders.insideH）
                tc = cell._tc
                bottom_el = None
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        bottom_el = tcBorders.find(qn('w:bottom'))
                if bottom_el is None and tblBorders is not None:
                    bottom_el = tblBorders.find(qn('w:insideH'))
                if not _border_is_absent(bottom_el):
                    return False

            # ---- 左边 / 右边 ----
            # ④ 所有单元格左右边均无边框
            # 左：ci==0 时外沿 = tblBorders.left；否则 = tblBorders.insideV；tcBorders.left 优先
            tc = cell._tc
            left_el = None
            right_el = None
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    left_el = tcBorders.find(qn('w:left'))
                    right_el = tcBorders.find(qn('w:right'))
            if left_el is None and tblBorders is not None:
                left_el = tblBorders.find(qn('w:left') if ci == 0 else qn('w:insideV'))
            if right_el is None and tblBorders is not None:
                right_el = tblBorders.find(qn('w:right') if ci == n_cols - 1 else qn('w:insideV'))
            if not _border_is_absent(left_el):
                return False
            if not _border_is_absent(right_el):
                return False

    return True


# =================== 维度一检查 ===================
def check_dimension_one(file1, file2):
    """可用与可修改性检查"""
    results = []

    # 1. 两个文件存在且可正常打开
    for f in [file1, file2]:
        if not os.path.exists(f):
            results.append(f"文件不存在: {os.path.basename(f)}")
            return False, results
        ext = os.path.splitext(f)[1].lower()
        if ext != '.docx':
            results.append(f"文件格式不正确: {ext}（仅支持 .docx）")
            return False, results
        try:
            Document(f)
        except Exception as e:
            results.append(f"文件无法正常打开: {os.path.basename(f)} - {e}")
            return False, results

    # 2. 可编辑性检查 - docx本身就是可编辑的，不是图片/PDF
    for f in [file1, file2]:
        doc = Document(f)
        # 检查是否有文字段落（非纯图片）
        text_paras = [p for p in doc.paragraphs if p.text.strip()]
        if len(text_paras) < 3:
            results.append(f"文档可能为纯图片，缺少可编辑文字: {os.path.basename(f)}")
            return False, results

    results.append("维度一通过：文件格式正确、可正常打开、内容可编辑")
    return True, results


# =================== 维度二检查 ===================
def check_dimension_two(file1, file2):
    """完成度评分细则"""
    doc1 = Document(file1)
    doc2 = Document(file2)
    paras1 = doc1.paragraphs
    paras2 = doc2.paragraphs

    score = 0
    details = []
    items = []  # 结构化评分项：命中(award)与未命中(fail)均记录，供 evaluate() 使用

    def award(points, desc):
        nonlocal score
        score += points
        sign = f"+{points}" if points > 0 else str(points)
        details.append(f"  {sign}：{desc}")
        items.append({
            "rule": desc,
            "max_delta": points,
            "delta": points,
            "hit": True,
            "detail": "",
        })

    def fail(points, desc):
        # 未命中项不计入 score / details（保持原有文本输出行为不变），
        # 但记录到 items，供 evaluate() 按 md §2.2 要求同时列出命中与未命中项。
        items.append({
            "rule": desc,
            "max_delta": points,
            "delta": 0,
            "hit": False,
            "detail": "",
        })

    # === 第一个文档封面检查 ===

    # +1: 第一个交付文档第1页为新增封面页，封面标题为"专项核查呈报材料"，字体为黑体二号加粗
    # 细则拆解：① 位于第1页；② 标题文本为"专项核查呈报材料"；③ 中文字体=黑体；④ 字号=二号；⑤ 加粗
    # 按办公软件实际渲染规则解析（run 直设 → 段落样式 → 样式链继承）
    first_page_paras = get_first_page_paragraphs(paras1)
    title_para = None
    for p in first_page_paras:
        if '专项核查呈报材料' in p.text.replace(' ', ''):
            title_para = p
            break
    if title_para is not None:
        # 合并该段落所有 run 检查：标题任一 run 若命中"黑体/二号/加粗"即视为格式生效
        title_runs = [r for r in title_para.runs if '专项核查呈报材料' in r.text.replace(' ', '')
                      or r.text.replace(' ', '').strip() != '']
        # 若标题被拆分到多个 run，任取覆盖标题文字的 run；否则取第一个非空 run
        target_run = None
        for r in title_runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is not None:
            ea = get_effective_east_asian_font(target_run, title_para)
            sz = get_effective_size(target_run, title_para)
            bold = get_effective_bold(target_run, title_para)
            is_heiti = (ea == '黑体')
            is_size_2 = size_match(sz, '二号')
            if is_heiti and is_size_2 and bold:
                award(1, "封面标题'专项核查呈报材料'位于第1页，黑体二号加粗")
            else:
                fail(1, f"封面标题格式不符（黑体={is_heiti}, 二号={is_size_2}, 加粗={bold}）")
        else:
            fail(1, "封面标题段落无有效run")
    else:
        fail(1, "第1页未找到封面标题'专项核查呈报材料'")

    # +1: "专项核查呈报材料"标题所在位置为封面页第二行
    # 细则拆解：① 标题为"专项核查呈报材料"；② 位于封面页（第1页）；③ 所在行为第二行
    # 办公软件里"行"= 段落（空段落也算一行），因此"第二行" = 封面页内的第2个段落（index=1）
    title_at_line2 = False
    if len(first_page_paras) >= 2:
        line2_text = first_page_paras[1].text.replace(' ', '').strip()
        if '专项核查呈报材料' in line2_text:
            title_at_line2 = True
    if title_at_line2:
        award(1, "'专项核查呈报材料'标题位于封面页第二行")
    else:
        fail(1, "'专项核查呈报材料'标题不在封面页第二行")

    # +1: 封面页"专项核查呈报材料"标题段落格式为居中对齐、单倍行距
    # 细则拆解：① 段落为封面页的"专项核查呈报材料"标题；② 居中对齐；③ 单倍行距
    # 按办公软件段落属性解析规则（段落直设 → 段落样式 → 样式链继承）
    cover_title_para = None
    for p in first_page_paras:
        if '专项核查呈报材料' in p.text.replace(' ', ''):
            cover_title_para = p
            break
    if cover_title_para is not None:
        center_ok = is_effective_center_aligned(cover_title_para)
        single_ok = is_effective_single_line_spacing(cover_title_para)
        if center_ok and single_ok:
            award(1, "封面标题段落居中对齐、单倍行距")
        else:
            fail(1, f"封面标题段落格式不符（居中={center_ok}, 单倍行距={single_ok}）")
    else:
        fail(1, "封面页未找到'专项核查呈报材料'标题段落")

    # +1: 封面副标题为"（华中区域零售项目经营资金异常及现金流风险事项）"，
    #     并位于"专项核查呈报材料"下方，字体为宋体小三
    # 细则拆解：① 副标题文本完整正确（含全角括号）；② 位于主标题下方（同一封面页内、序号更大）；
    #          ③ 中文字体=宋体；④ 字号=小三
    expected_subtitle = '（华中区域零售项目经营资金异常及现金流风险事项）'
    main_title_idx = None
    for i, p in enumerate(first_page_paras):
        if '专项核查呈报材料' in p.text.replace(' ', ''):
            main_title_idx = i
            break
    subtitle_para = None
    if main_title_idx is not None:
        for j in range(main_title_idx + 1, len(first_page_paras)):
            if expected_subtitle in first_page_paras[j].text.replace(' ', ''):
                subtitle_para = first_page_paras[j]
                break
    if subtitle_para is not None:
        target_run = None
        for r in subtitle_para.runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is not None:
            ea = get_effective_east_asian_font(target_run, subtitle_para)
            sz = get_effective_size(target_run, subtitle_para)
            is_songti = (ea == '宋体')
            is_size_xiao3 = size_match(sz, '小三')
            if is_songti and is_size_xiao3:
                award(1, "封面副标题内容正确并位于主标题下方，宋体小三")
            else:
                fail(1, f"封面副标题内容/位置正确，但字体格式不符（宋体={is_songti}, 小三={is_size_xiao3}）")
        else:
            fail(1, "封面副标题段落无有效run")
    else:
        fail(1, "封面页主标题下方未找到指定副标题")

    # +1: 封面副标题为"（华中区域零售项目经营资金异常及现金流风险事项）"段落格式为居中对齐、单倍行距
    # 细则拆解：① 段落 = 上述副标题段落；② 居中对齐；③ 单倍行距
    # 按办公软件段落属性解析规则（段落直设 → 段落样式 → 样式链继承）
    if subtitle_para is not None:
        center_ok = is_effective_center_aligned(subtitle_para)
        single_ok = is_effective_single_line_spacing(subtitle_para)
        if center_ok and single_ok:
            award(1, "封面副标题段落居中对齐、单倍行距")
        else:
            fail(1, f"封面副标题段落格式不符（居中={center_ok}, 单倍行距={single_ok}）")
    else:
        fail(1, "封面页未找到副标题段落")

    # +1: 封面副标题"（华中区域零售项目经营资金异常及现金流风险事项）"文字下方有文档分隔虚线
    # 细则拆解：① 目标段落 = 上一条已定位的副标题段落；② 其"下方"存在一条文档分隔虚线
    # 在办公软件里"下方有分隔虚线"包含两种可见形式：
    #   (a) 副标题段落自身的底部段落边框；(b) 紧邻下一段落是分隔线（字符模拟 或 顶部段落边框）
    subtitle_next = None
    if subtitle_para is not None:
        # 找到副标题在 first_page_paras 中的位置以取"下一段"
        for k, p in enumerate(first_page_paras):
            if p is subtitle_para and k + 1 < len(first_page_paras):
                subtitle_next = first_page_paras[k + 1]
                break
        if has_visible_separator_below(subtitle_para, subtitle_next):
            award(1, "副标题下方有文档分隔虚线")
        else:
            fail(1, "副标题下方无文档分隔虚线")
    else:
        fail(1, "封面页未找到副标题段落")

    # +1: 封面副标题"（华中区域零售项目经营资金异常及现金流风险事项）"文字下方文档分隔虚线的宽度为1磅
    # 细则拆解：① 目标段落 = 上条已定位的副标题段落；② 其下方存在文档分隔虚线；
    #          ③ 该分隔虚线的宽度 ≈ 1.0 磅（容差 0.5 磅）。
    # 办公软件生效性：
    #   - 段落边框实现: w:pBdr/w:bottom/@w:sz 单位为 1/8 磅, w:sz=8 即 1 磅;
    #   - 字符模拟实现: 由分隔字符 Unicode 权重 × 字体字号 推算笔画宽度 (见 get_separator_width_pt)
    if subtitle_para is not None and subtitle_next is not None:
        sep_w = get_separator_width_pt(subtitle_para, subtitle_next)
        # 容差 0.5 磅: 覆盖办公软件对 w:sz 四舍五入 / 字符模拟推算的误差
        if sep_w is not None and abs(sep_w - 1.0) <= 0.5:
            award(1, f"副标题下方分隔虚线宽度约 {sep_w:.2f} 磅（1 磅）")
        else:
            fail(1, f"副标题下方分隔虚线宽度不为约 1 磅（实测 {sep_w}）")
    else:
        fail(1, "封面页未找到副标题或其下方段落")

    # +1: 封面副标题"（华中区域零售项目经营资金异常及现金流风险事项）"文字下方文档分隔虚线的长度为7.7-8.1厘米
    # 细则拆解：① 目标段落 = 上条已定位的副标题段落；② 其下方存在分隔虚线；
    #          ③ 该分隔虚线长度在 7.7-8.1 厘米之间（闭区间）
    if subtitle_para is not None and subtitle_next is not None:
        sep_len = get_separator_length_cm(subtitle_para, subtitle_next, doc1.sections[0])
        if sep_len is not None and 7.7 <= sep_len <= 8.1:
            award(1, f"副标题下方分隔虚线长度为{sep_len:.2f}cm（7.7-8.1cm）")
        else:
            fail(1, f"副标题下方分隔虚线长度不在7.7-8.1cm（实测{sep_len}）")
    else:
        fail(1, "封面页未找到副标题或其下方段落")

    # +1: 封面"【核心核查事项】"，下方依次包含5条事项：
    #     ①320万元供应链预付款提前支付及资金占用
    #     ②118万元线上平台结算款延迟回流
    #     ③76万元运营服务费重复计费
    #     ④210万元内部资金调拨用途待核
    #     ⑤181万元开办筹备/物流代垫/促销费用待补凭证
    #     字体格式为宋体小三
    # 细则拆解：① 封面页有"【核心核查事项】"标记段；② 其下方依次出现5条事项（按细则顺序）；
    #          ③【核心核查事项】+ 5 条事项段落的中文字体=宋体；④ 字号=小三
    header_idx = None
    for i, p in enumerate(first_page_paras):
        if '核心核查事项' in p.text.replace(' ', ''):
            header_idx = i
            break

    required_items_keys = [
        ['320万元', '供应链预付款', '提前支付', '资金占用'],
        ['118万元', '线上平台', '结算款', '延迟回流'],
        ['76万元', '运营服务费', '重复计费'],
        ['210万元', '内部资金调拨', '用途待核'],
        ['181万元', '开办筹备', '物流代垫', '促销费用', '待补凭证'],
    ]

    def _match_item(text, keys):
        t = text.replace(' ', '').replace('、', '').replace('/', '').replace('，', '')
        # 允许每条至少命中 2 个关键字，且金额关键字必须命中
        hits = sum(1 for k in keys if k in t)
        return keys[0] in t and hits >= 2

    ordered_ok = False
    item_paragraphs = []
    if header_idx is not None:
        cursor = header_idx + 1
        matched_count = 0
        for keys in required_items_keys:
            found = False
            while cursor < len(first_page_paras):
                p = first_page_paras[cursor]
                if _match_item(p.text, keys):
                    item_paragraphs.append(p)
                    matched_count += 1
                    cursor += 1
                    found = True
                    break
                cursor += 1
            if not found:
                break
        ordered_ok = (matched_count == 5)

    if header_idx is not None and ordered_ok:
        header_para = first_page_paras[header_idx]
        check_paras = [header_para] + item_paragraphs
        font_ok = True
        for p in check_paras:
            target_run = None
            for r in p.runs:
                if r.text.strip():
                    target_run = r
                    break
            if target_run is None:
                font_ok = False
                break
            ea = get_effective_east_asian_font(target_run, p)
            sz = get_effective_size(target_run, p)
            if ea != '宋体' or not size_match(sz, '小三'):
                font_ok = False
                break
        if font_ok:
            award(1, "'【核心核查事项】'及下方5条事项内容完整、顺序正确，字体宋体小三")
        else:
            fail(1, "'【核心核查事项】'及5条事项内容顺序正确，但字体格式不满足宋体小三")
    else:
        fail(1, "封面未找到'【核心核查事项】'或下方5条事项顺序不完整")

    # +1: 封面"【核心核查事项】"及下方5条事项：
    #     ①320万元供应链预付款提前支付及资金占用
    #     ②118万元线上平台结算款延迟回流
    #     ③76万元运营服务费重复计费
    #     ④210万元内部资金调拨用途待核
    #     ⑤181万元开办筹备/物流代垫/促销费用待补凭证
    #     段落格式为单倍行距、居中对齐
    # 细则拆解：① 目标段落 = 上条已定位的【核心核查事项】+ 5条事项共6个段落；
    #          ② 单倍行距；③ 居中对齐
    # 按办公软件段落属性解析规则（段落直设 → 段落样式 → 样式链继承）
    if header_idx is not None and ordered_ok:
        header_para = first_page_paras[header_idx]
        check_paras = [header_para] + item_paragraphs
        para_fmt_ok = True
        for p in check_paras:
            if not is_effective_center_aligned(p) or not is_effective_single_line_spacing(p):
                para_fmt_ok = False
                break
        if para_fmt_ok:
            award(1, "'【核心核查事项】'及5条事项段落居中对齐、单倍行距")
        else:
            fail(1, "'【核心核查事项】'及5条事项段落格式不满足居中对齐、单倍行距")
    else:
        fail(1, "封面未找到'【核心核查事项】'或5条事项")

    # +1: 封面包含"待核查影响资金总额：905万元"文字字体格式为宋体小三
    # 细则拆解：① 位于封面页；② 文本包含"待核查影响资金总额：905万元"；
    #          ③ 中文字体=宋体；④ 字号=小三
    total_amount_para = None
    for p in first_page_paras:
        t = p.text.replace(' ', '')
        if '待核查影响资金总额' in t and '905万元' in t:
            total_amount_para = p
            break
    if total_amount_para is not None:
        target_run = None
        for r in total_amount_para.runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is not None:
            ea = get_effective_east_asian_font(target_run, total_amount_para)
            sz = get_effective_size(target_run, total_amount_para)
            is_songti = (ea == '宋体')
            is_size_xiao3 = size_match(sz, '小三')
            if is_songti and is_size_xiao3:
                award(1, "封面包含'待核查影响资金总额：905万元'，宋体小三")
            else:
                fail(1, f"'待核查影响资金总额：905万元'字体格式不符（宋体={is_songti}, 小三={is_size_xiao3}）")
        else:
            fail(1, "'待核查影响资金总额'段落无有效run")
    else:
        fail(1, "封面未找到'待核查影响资金总额：905万元'")

    # +1: 封面"待核查影响资金总额：905万元"文字下方有文档分隔虚线
    # 细则拆解：① 目标段落 = 上条已定位的"待核查影响资金总额：905万元"段落；
    #          ② 其下方存在一条文档分隔虚线
    # 办公软件里"下方有分隔虚线"包含两种可见形式：
    #   (a) 该段落自身底部段落边框；
    #   (b) 紧邻下一段落是分隔线（字符模拟 或 顶部段落边框）
    total_amount_next = None
    if total_amount_para is not None:
        for k, p in enumerate(first_page_paras):
            if p is total_amount_para and k + 1 < len(first_page_paras):
                total_amount_next = first_page_paras[k + 1]
                break
    if total_amount_para is not None:
        if has_visible_separator_below(total_amount_para, total_amount_next):
            award(1, "'待核查影响资金总额：905万元'下方有文档分隔虚线")
        else:
            fail(1, "'待核查影响资金总额：905万元'下方无文档分隔虚线")
    else:
        fail(1, "封面未找到'待核查影响资金总额：905万元'")

    # +1: 封面"待核查影响资金总额：905万元"文字下方文档分隔虚线的宽度为1磅
    # 细则拆解：① 目标段落 = 上条已定位的"待核查影响资金总额：905万元"段落；
    #          ② 其下方存在文档分隔虚线；③ 分隔虚线宽度 ≈ 1.0 磅（容差 0.5 磅）。
    # 办公软件生效性：
    #   - 段落边框实现: w:pBdr/w:bottom/@w:sz 单位为 1/8 磅, w:sz=8 即 1 磅;
    #   - 字符模拟实现: 由分隔字符 Unicode 权重 × 字体字号 推算笔画宽度 (见 get_separator_width_pt)
    if total_amount_para is not None and total_amount_next is not None:
        sep_w = get_separator_width_pt(total_amount_para, total_amount_next)
        # 容差 0.5 磅: 覆盖办公软件对 w:sz 四舍五入 / 字符模拟推算的误差
        if sep_w is not None and abs(sep_w - 1.0) <= 0.5:
            award(1, f"'待核查影响资金总额：905万元'下方分隔虚线宽度约 {sep_w:.2f} 磅（1 磅）")
        else:
            fail(1, f"'待核查影响资金总额：905万元'下方分隔虚线宽度不为约 1 磅（实测 {sep_w}）")
    else:
        fail(1, "封面未找到'待核查影响资金总额：905万元'或其下方段落")

    # +1: 封面"待核查影响资金总额：905万元"文字下方文档分隔虚线的长度为7.7-8.1厘米
    # 细则拆解：① 目标段落 = 上条已定位的"待核查影响资金总额：905万元"段落；
    #          ② 其下方存在分隔虚线；③ 该分隔虚线长度 ∈ [7.7, 8.1] 厘米
    if total_amount_para is not None and total_amount_next is not None:
        sep_len = get_separator_length_cm(total_amount_para, total_amount_next, doc1.sections[0])
        if sep_len is not None and 7.7 <= sep_len <= 8.1:
            award(1, f"'待核查影响资金总额：905万元'下方分隔虚线长度为{sep_len:.2f}cm（7.7-8.1cm）")
        else:
            fail(1, f"'待核查影响资金总额：905万元'下方分隔虚线长度不在7.7-8.1cm（实测{sep_len}）")
    else:
        fail(1, "封面未找到'待核查影响资金总额：905万元'或其下方段落")

    # +1: 封面包含"呈报人：林远芝""联系方式：186****2741""联系地址：华中地区羽厘商业办公区""日期：2026年5月20日"四项信息；字体格式为宋体小三
    # 细则拆解：① 位于封面页；② 四项文本完整正确；③ 中文字体=宋体；④ 字号=小三
    info_items = [
        '呈报人：林远芝',
        '联系方式：186****2741',
        '联系地址：华中地区羽厘商业办公区',
        '日期：2026年5月20日',
    ]
    info_paras = []
    for item in info_items:
        item_clean = item.replace(' ', '')
        matched = None
        for p in first_page_paras:
            if item_clean in p.text.replace(' ', ''):
                matched = p
                break
        info_paras.append(matched)
    all_found = all(x is not None for x in info_paras)
    if all_found:
        info_font_ok = True
        for p in info_paras:
            target_run = None
            for r in p.runs:
                if r.text.strip():
                    target_run = r
                    break
            if target_run is None:
                info_font_ok = False
                break
            ea = get_effective_east_asian_font(target_run, p)
            sz = get_effective_size(target_run, p)
            if ea != '宋体' or not size_match(sz, '小三'):
                info_font_ok = False
                break
        if info_font_ok:
            award(1, "封面四项信息完整，字体宋体小三")
        else:
            fail(1, "封面四项信息完整，但字体格式不满足宋体小三")
    else:
        missing = [info_items[i] for i, x in enumerate(info_paras) if x is None]
        fail(1, f"封面四项信息不完整（缺失：{missing}）")

    # +1: 封面"日期：2026年5月20日"文字下方有文档分隔虚线
    # 细则拆解：① 目标段落 = 上条已定位的"日期：2026年5月20日"段落（封面页内）；
    #          ② 其下方存在一条文档分隔虚线
    # 办公软件里"下方有分隔虚线"包含两种可见形式：
    #   (a) 该段落自身底部段落边框；
    #   (b) 紧邻下一段落是分隔线（字符模拟 或 顶部段落边框）
    date_para = info_paras[3] if len(info_paras) >= 4 else None
    date_next = None
    if date_para is not None:
        for k, p in enumerate(first_page_paras):
            if p is date_para and k + 1 < len(first_page_paras):
                date_next = first_page_paras[k + 1]
                break
    if date_para is not None:
        if has_visible_separator_below(date_para, date_next):
            award(1, "'日期：2026年5月20日'下方有文档分隔虚线")
        else:
            fail(1, "'日期：2026年5月20日'下方无文档分隔虚线")
    else:
        fail(1, "封面未找到'日期：2026年5月20日'")

    # +1: 封面"日期：2026年5月20日"文字下方文档分隔虚线的宽度为1磅
    # 细则拆解：① 目标段落 = 上条已定位的"日期：2026年5月20日"段落；
    #          ② 其下方存在文档分隔虚线；③ 分隔虚线宽度 ≈ 1.0 磅（容差 0.5 磅）。
    # 办公软件生效性：
    #   - 段落边框实现: w:pBdr/w:bottom/@w:sz 单位为 1/8 磅, w:sz=8 即 1 磅;
    #   - 字符模拟实现: 由分隔字符 Unicode 权重 × 字体字号 推算笔画宽度 (见 get_separator_width_pt)
    if date_para is not None and date_next is not None:
        sep_w = get_separator_width_pt(date_para, date_next)
        # 容差 0.5 磅: 覆盖办公软件对 w:sz 四舍五入 / 字符模拟推算的误差
        if sep_w is not None and abs(sep_w - 1.0) <= 0.5:
            award(1, f"'日期：2026年5月20日'下方分隔虚线宽度约 {sep_w:.2f} 磅（1 磅）")
        else:
            fail(1, f"'日期：2026年5月20日'下方分隔虚线宽度不为约 1 磅（实测 {sep_w}）")
    else:
        fail(1, "封面未找到'日期：2026年5月20日'或其下方段落")

    # +1: 封面"日期：2026年5月20日"文字下方文档分隔虚线的长度为7.7-8.1厘米
    # 细则拆解：① 目标段落 = 上条已定位的"日期：2026年5月20日"段落；
    #          ② 其下方存在分隔虚线；③ 该分隔虚线长度 ∈ [7.7, 8.1] 厘米
    if date_para is not None and date_next is not None:
        sep_len = get_separator_length_cm(date_para, date_next, doc1.sections[0])
        if sep_len is not None and 7.7 <= sep_len <= 8.1:
            award(1, f"'日期：2026年5月20日'下方分隔虚线长度为{sep_len:.2f}cm（7.7-8.1cm）")
        else:
            fail(1, f"'日期：2026年5月20日'下方分隔虚线长度不在7.7-8.1cm（实测{sep_len}）")
    else:
        fail(1, "封面未找到'日期：2026年5月20日'或其下方段落")

    # === 目录页检查 ===

    # +1: 第一个交付文档第2页或封面后一页为目录页，目录页标题为"目录"或"目 录"，
    #     字体为黑体小三加粗，且位于页面上方居中
    # 细则拆解：① 目录页 = 第2页 或 封面后一页；② 标题文本 ∈ {"目录","目 录"}（忽略字符间空格后匹配"目录"）；
    #          ③ 中文字体=黑体；④ 字号=小三；⑤ 加粗；⑥ 位于页面上方（该页首个非空段落）居中
    pages1 = get_paragraphs_by_page(paras1)
    toc_page_paras = pages1[1] if len(pages1) >= 2 else []
    toc_title_para = None
    toc_title_is_top = False
    top_non_empty_idx = None
    for j, p in enumerate(toc_page_paras):
        if p.text.strip():
            top_non_empty_idx = j
            break
    for j, p in enumerate(toc_page_paras):
        if p.text.replace(' ', '').replace('　', '').strip() == '目录':
            toc_title_para = p
            toc_title_is_top = (j == top_non_empty_idx)
            break

    # 兼容后续尚未修改的检查项：保留 toc_idx 变量
    toc_idx = None
    if toc_title_para is not None:
        for i, p in enumerate(paras1):
            if p is toc_title_para:
                toc_idx = i
                break

    if toc_title_para is not None:
        target_run = None
        for r in toc_title_para.runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is not None:
            ea = get_effective_east_asian_font(target_run, toc_title_para)
            sz = get_effective_size(target_run, toc_title_para)
            bold = get_effective_bold(target_run, toc_title_para)
            center = is_effective_center_aligned(toc_title_para)
            is_heiti = (ea == '黑体')
            is_size_xiao3 = size_match(sz, '小三')
            if is_heiti and is_size_xiao3 and bold and center and toc_title_is_top:
                award(1, "第2页目录页标题'目录'，黑体小三加粗，位于页面上方居中")
            else:
                fail(1, f"目录页标题格式不符（黑体={is_heiti}, 小三={is_size_xiao3}, "
                        f"加粗={bold}, 居中={center}, 页面上方={toc_title_is_top}）")
        else:
            fail(1, "目录页标题段落无有效run")
    else:
        fail(1, "第2页（封面后一页）未找到目录页标题'目录'/'目 录'")

    # +3: 目录包含六个一级标题, 字体为宋体小四
    # 细则拆解:
    #   ① 目录页里存在六个目录条目文本, 依次包含 toc_required 中六个一级标题关键词;
    #   ② 每个目录条目段落中非空 run 的中文字体 = 宋体, 字号 = 小四.
    # 办公软件生效性:
    #   - 中文字体: get_effective_east_asian_font (run 直设 → 段落 rPr → 样式链);
    #   - 字号: get_effective_size (同上继承路径); 用 size_match(sz, '小四') 判定,
    #     覆盖 152400 EMU (12pt = 小四) ± SIZE_TOLERANCE 的容差;
    #   - 只针对"目录条目里包含六个 toc_required 关键词的段落"逐段校验; 页码 run
    #     (纯数字) 交给 rubric 单独的"目录页码 Times New Roman 小四"规则, 此处不判英数;
    #   - 任一目标目录条目的任一含中文的非空 run 违反字体或字号即失败.
    toc_required = ['项目背景', '主要关联主体情况', '异常资金往来情况',
                    '业务流程变化说明', '库存与资金需求分析', '综合说明']
    toc_entries = []
    toc_entry_paras = []  # 与 toc_entries 对齐, 用于后续字体/字号校验
    if toc_idx is not None:
        for i in range(toc_idx + 1, min(toc_idx + 15, len(paras1))):
            text = paras1[i].text.strip()
            if text and any(t in text for t in toc_required):
                toc_entries.append(text)
                toc_entry_paras.append(paras1[i])
    found_toc = sum(1 for req in toc_required if any(req in e for e in toc_entries))
    if found_toc >= 6:
        toc_font_bad = []  # 收集不合格样本, 便于排错
        for p in toc_entry_paras:
            for r in p.runs:
                text = r.text or ''
                if not text.strip():
                    continue
                # 只校验含中文字符的 run 的中文字体/字号; 纯英数(页码等)不在本条约束内
                if not any('一' <= ch <= '鿿' for ch in text):
                    continue
                ea = get_effective_east_asian_font(r, p) or ''
                sz = get_effective_size(r, p)
                if ea != '宋体':
                    toc_font_bad.append(f"字体={ea!r}:{text.strip()[:8]!r}")
                    continue
                if not size_match(sz, '小四'):
                    sz_pt = (int(sz) / 12700.0) if sz else None
                    toc_font_bad.append(f"字号={sz_pt}pt:{text.strip()[:8]!r}")
        if not toc_font_bad:
            award(3, "目录包含六个一级标题, 字体宋体小四")
        else:
            fail(3, f"目录一级标题字体/字号不满足宋体小四(样本 {toc_font_bad[:5]})")
    else:
        fail(3, f"目录一级标题不完整(找到{found_toc}/6)")

    # +1: 目录条目包含对应页码数字，页码数字位于每条目录项末尾，中间用点状线连接
    # 细则拆解：① 每条目录条目包含页码数字；② 页码数字位于该条末尾；③ 中间用点状线（dot leader）连接
    # 办公软件里"点状线连接"由段落的制表位 w:tabs/w:tab/@w:leader="dot" 决定：
    # 输入 tab 后，Word/WPS 自动在标题与右侧页码之间填充点状引导线；
    # 若无 dot leader，即便含 tab 也只是空白，不满足"点状线连接"
    toc_entries = []
    if toc_title_para is not None:
        # 目录条目 = 目录页里位于"目录"标题之后的、非空段落
        started = False
        for p in toc_page_paras:
            if p is toc_title_para:
                started = True
                continue
            if not started:
                continue
            if p.text.strip():
                toc_entries.append(p)

    if toc_entries:
        all_have_page = True
        all_page_at_end = True
        all_have_dot_leader = True
        for p in toc_entries:
            text = p.text.rstrip()
            # 页码 = 末尾的一段阿拉伯数字
            m = re.search(r'(\d+)\s*$', text)
            if not m:
                all_have_page = False
                all_page_at_end = False
                continue
            # 位于末尾：数字后除空白外无其他字符（rstrip 已保证）
            # dot leader：段落制表位中至少有一个 w:leader="dot"
            has_dot = False
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                tabs = pPr.find(qn('w:tabs'))
                if tabs is not None:
                    for tab in tabs:
                        if tab.get(qn('w:leader')) == 'dot':
                            has_dot = True
                            break
            if not has_dot:
                all_have_dot_leader = False

        if all_have_page and all_page_at_end and all_have_dot_leader:
            award(1, "目录条目均含末尾页码数字，且中间以点状线连接")
        else:
            fail(1, f"目录条目/页码/点状线不满足（页码={all_have_page}, "
                    f"末尾={all_page_at_end}, 点状线={all_have_dot_leader}）")
    else:
        fail(1, "未找到目录条目")

    # +1: 目录中所有页码向右互相对齐，字体为 Times New Roman 小四
    # 细则拆解：① 目标 = 目录所有条目末尾的页码数字；
    #          ② "所有页码向右互相对齐"—— 视觉上：所有条目的页码在同一竖直位置成列（页面右侧）；
    #          ③ 英文字体 = Times New Roman；④ 字号 = 小四。
    # 视觉对齐的可见实现（Word/WPS 中各种"看起来对齐"的做法都要覆盖）：
    #   (a) tab 制表位：段落有 <w:tab/>，且段落 <w:tabs> 里最靠右的 w:pos 一致 —— 页码落到同一 X；
    #   (b) 段落右对齐 w:jc="right"：整段贴右，页码同列；
    #   (c) 敲空格/全角空格对齐：段落无 tab、无右对齐，但通过在文字与页码之间补齐 ASCII 空格 ( )
    #       或全角空格 (　) 让页码近似右侧成列 —— 需实测每条条目"页码起始 X 坐标"近似一致。
    # 视觉 X 坐标估算（基于办公软件字体排版基本模型）：
    #   x_start(pageNum) = 段落左缩进 + Σ 页码前所有 run 的可见字符宽度
    #   其中：
    #     - 中文字符（含中文标点、CJK Unified Ideographs、全角空格）宽度 ≈ 字号磅值 × 1.0 em；
    #     - ASCII 字符（半角空格、字母、数字、半角标点）宽度 ≈ 字号磅值 × 0.5 em；
    #     - 制表符 <w:tab/> 直接跳到该段 tabs 里下一个 pos（twips → 转 pt=twips/20）；
    #     - 段落左缩进：w:ind w:left / firstLine（twips）；
    #   单位统一为磅（pt）。视觉对齐容差：所有页码 X 极差 ≤ 2 磅（约 0.7 mm，办公软件肉眼可判"对齐"）。
    def _para_left_indent_twips(para):
        pf = para.paragraph_format
        left = 0
        if pf.left_indent is not None:
            left = int(pf.left_indent) / 635  # EMU → twips
        if pf.first_line_indent is not None:
            fli = int(pf.first_line_indent) / 635
            if fli > 0:
                left += fli
        return left  # twips

    def _run_char_width_pt(run, para, ch):
        sz_emu = get_effective_size(run, para)
        pt = (sz_emu / 12700.0) if sz_emu else 12.0
        # 中文/全角 = 1em；ASCII/半角 = 0.5em
        if ord(ch) > 0x2E80 or ch == '　':
            return pt
        return pt * 0.5

    def _tab_advance_twips(cur_twips, tabs_el):
        # 移动到 tabs 里第一个 pos > cur_twips 的位置；若没有则跳 720 twips（默认 tab 0.5")
        candidates = []
        if tabs_el is not None:
            for tab in tabs_el:
                pos = tab.get(qn('w:pos'))
                try:
                    if pos is not None:
                        candidates.append(int(pos))
                except ValueError:
                    pass
        candidates = [c for c in candidates if c > cur_twips]
        if candidates:
            return min(candidates)
        return cur_twips + 720

    def _pagenum_visual_x_pt(para):
        """返回目录段落里页码起始位置的视觉 X 坐标（磅）。找不到页码返回 None。"""
        pPr = para._element.find(qn('w:pPr'))
        tabs_el = pPr.find(qn('w:tabs')) if pPr is not None else None
        cur_twips = _para_left_indent_twips(para)
        cur_pt = cur_twips / 20.0
        # 找到页码所在 run
        page_run_idx = None
        for i, r in enumerate(para.runs):
            if re.search(r'\d\s*$', r.text or ''):
                page_run_idx = i
        if page_run_idx is None:
            return None
        # 遍历页码之前 + 页码 run 内部 tab（页码 run 里通常是 <w:tab/> 后 <w:t>数字</w:t>）
        for i, r in enumerate(para.runs):
            el = r._element
            for child in el:
                tag = child.tag.split('}', 1)[-1]
                if tag == 'tab':
                    new_twips = _tab_advance_twips(int(cur_pt * 20), tabs_el)
                    cur_pt = new_twips / 20.0
                    if i == page_run_idx:
                        return cur_pt  # 已到达页码起点
                elif tag == 't':
                    text = child.text or ''
                    if i == page_run_idx:
                        return cur_pt  # 页码 run 里第一个 <w:t> 前的 X 即页码起点
                    for ch in text:
                        cur_pt += _run_char_width_pt(r, para, ch)
        return cur_pt

    if toc_entries:
        all_font_ok = True
        all_size_ok = True
        visual_xs = []
        for p in toc_entries:
            page_run = None
            for r in p.runs:
                if re.search(r'\d\s*$', r.text):
                    page_run = r
            if page_run is None:
                all_font_ok = False
                all_size_ok = False
                visual_xs.append(None)
                continue
            ascii_f = get_effective_ascii_font(page_run, p)
            sz = get_effective_size(page_run, p)
            if 'Times New Roman' not in (ascii_f or ''):
                all_font_ok = False
            if not size_match(sz, '小四'):
                all_size_ok = False
            visual_xs.append(_pagenum_visual_x_pt(p))

        # 视觉对齐：所有条目 X 坐标极差 ≤ 2 磅
        xs_valid = [x for x in visual_xs if x is not None]
        mutually_aligned = (len(xs_valid) == len(visual_xs)
                            and len(xs_valid) >= 1
                            and (max(xs_valid) - min(xs_valid)) <= 2.0)

        if mutually_aligned and all_font_ok and all_size_ok:
            award(1, "目录所有页码向右互相对齐、Times New Roman 小四")
        else:
            fail(1, f"目录页码格式不符（视觉对齐={mutually_aligned}, TNR={all_font_ok}, 小四={all_size_ok}）")
    else:
        fail(1, "未找到目录条目")

    # +1: 第一个交付文档的核查材料样式，正文页顶部或页脚区域出现"第x页 共y页"格式页码
    # 细则拆解：① 位置=正文页的"顶部（页眉）"或"页脚"；② 文本格式="第x页 共y页"（x/y 为页码数字）
    # 办公软件里"页码"通常由 PAGE / NUMPAGES 域生成——按 Word/WPS 实际渲染取页眉/页脚可见文本 + 域结果，
    # 用正则 r'第\s*\d+\s*页\s*共\s*\d+\s*页' 匹配（允许中间有任意空白，与用户肉眼看到的"第 1 页 共 1 页"一致）。
    pattern = re.compile(r'第\s*\d+\s*页\s*共\s*\d+\s*页')
    page_number_found = False
    for section in doc1.sections:
        # 页眉
        header_text = ''
        for p in section.header.paragraphs:
            header_text += p.text
        if pattern.search(header_text):
            page_number_found = True
            break
        # 页脚
        footer_text = ''
        for p in section.footer.paragraphs:
            footer_text += p.text
        if pattern.search(footer_text):
            page_number_found = True
            break
    if page_number_found:
        award(1, "页眉或页脚出现'第x页 共y页'格式页码")
    else:
        fail(1, "页眉/页脚未出现'第x页 共y页'格式页码")

    # +1: 第一个交付文档正文页码字体格式：中文字体为宋体小五，阿拉伯数字为宋体小四
    # 细则拆解：① 目标 = 正文页页眉/页脚里承载"第x页 共y页"页码的段落；
    #          ② 该段落里的中文字符：中文字体=宋体，字号=小五；
    #          ③ 该段落里的阿拉伯数字：中文字体（eastAsia，Word 混排里页码数字仍走中文字体）=宋体，字号=小四
    # 按办公软件字体解析规则（run 直设 → 段落样式 → 样式链继承）
    def _classify(text):
        # 判定 run 是否为"纯阿拉伯数字"：忽略前后空白后仅由 0-9 组成
        stripped = text.strip()
        if not stripped:
            return 'empty'
        if re.fullmatch(r'\d+', stripped):
            return 'digit'
        # 含中文字符则视为中文 run
        if any('一' <= c <= '鿿' for c in stripped):
            return 'zh'
        return 'other'

    page_font_ok = True
    page_num_para_found = False
    pattern = re.compile(r'第\s*\d+\s*页\s*共\s*\d+\s*页')
    for section in doc1.sections:
        for region in (section.header, section.footer):
            for p in region.paragraphs:
                if not pattern.search(p.text):
                    continue
                page_num_para_found = True
                for r in p.runs:
                    kind = _classify(r.text)
                    if kind == 'empty' or kind == 'other':
                        continue
                    ea = get_effective_east_asian_font(r, p)
                    sz = get_effective_size(r, p)
                    if ea != '宋体':
                        page_font_ok = False
                        break
                    if kind == 'digit':
                        if not size_match(sz, '小四'):
                            page_font_ok = False
                            break
                    else:  # zh
                        if not size_match(sz, '小五'):
                            page_font_ok = False
                            break
                if not page_font_ok:
                    break
            if not page_font_ok:
                break
        if not page_font_ok:
            break

    if page_num_para_found and page_font_ok:
        award(1, "正文页码字体：中文宋体小五、阿拉伯数字宋体小四")
    else:
        if not page_num_para_found:
            fail(1, "未在页眉/页脚找到承载'第x页 共y页'的段落，无法判定页码字体")
        else:
            fail(1, "正文页码字体不满足：中文宋体小五、阿拉伯数字宋体小四")

    # === 第一个文档正文标题检查 ===

    # +3: 第一个交付文档一级标题样式应为"第一部分："，
    #     例如"第一部分：项目背景"、"第二部分：主要关联主体情况"等
    # 细则拆解：① 每个一级标题都采用"第X部分：内容"样式；
    #          ② X 为中文数字（一/二/三/四/五/六/七/八/九/十…）；
    #          ③ 冒号可为中文":"或英文":"（对应办公软件里两种常见键入方式）
    # 办公软件中"一级标题"= 应用 Word/WPS 标题1 样式的段落（Heading 1）
    h1_paras = [p for p in paras1 if p.style and p.style.name == 'Heading 1']
    heading_pattern = re.compile(r'^第[一二三四五六七八九十百]+部分[：:]\S')
    if h1_paras:
        all_ok = all(heading_pattern.match(p.text.strip()) for p in h1_paras)
        if all_ok:
            award(3, "一级标题全部采用'第X部分：内容'样式")
        else:
            fail(3, "一级标题未全部采用'第X部分：内容'样式")
    else:
        fail(3, "未找到一级标题（Heading 1）")

    # +3: 第一个交付文档包含
    #     "第一部分：项目背景"、"第二部分：主要关联主体情况"、"第三部分：异常资金往来情况"、
    #     "第四部分：业务流程变化说明"、"第五部分：库存与资金需求分析"、"第六部分：综合说明"
    #     这些一级标题且字体格式为黑体小三加粗
    # 细则拆解：① 六个一级标题（Heading 1）的文本必须依次严格匹配上述 6 条；
    #          ② 中文字体=黑体；③ 字号=小三；④ 加粗
    required_h1_titles = [
        '第一部分：项目背景',
        '第二部分：主要关联主体情况',
        '第三部分：异常资金往来情况',
        '第四部分：业务流程变化说明',
        '第五部分：库存与资金需求分析',
        '第六部分：综合说明',
    ]

    def _norm_title(s):
        # 忽略前后空白与全角/半角冒号差异，用于文本严格比对
        return s.strip().replace(':', '：')

    h1_texts = [_norm_title(p.text) for p in h1_paras]
    required_norm = [_norm_title(t) for t in required_h1_titles]
    titles_ok = (h1_texts == required_norm)

    if titles_ok:
        h1_font_ok = True
        for p in h1_paras:
            target_run = None
            for r in p.runs:
                if r.text.strip():
                    target_run = r
                    break
            if target_run is None:
                h1_font_ok = False
                break
            ea = get_effective_east_asian_font(target_run, p)
            sz = get_effective_size(target_run, p)
            bold = get_effective_bold(target_run, p)
            if ea != '黑体' or not size_match(sz, '小三') or not bold:
                h1_font_ok = False
                break
        if h1_font_ok:
            award(3, "六个一级标题按细则文本完整，黑体小三加粗")
        else:
            fail(3, "六个一级标题文本正确，但字体格式不满足黑体小三加粗")
    else:
        fail(3, "六个一级标题文本不满足细则指定的'第X部分：内容'")

    # +3: 一级标题段落格式两端对齐、1.5倍行距
    # 细则拆解：仅针对细则明确列出的 6 个一级标题（"第一部分：项目背景"…"第六部分：综合说明"），
    #          每个一级标题段落必须同时满足：① 两端对齐 ② 1.5倍行距
    # 办公软件生效性：
    #   - "两端对齐"在 Word/WPS 里对应段落 `w:jc w:val="both"`（python-docx 的 JUSTIFY），
    #     属性可来自段落直设或"标题 1"样式链 → 用 is_effective_justify_aligned 逐级回溯；
    #   - "1.5倍行距"在 Word/WPS 里对应 `w:spacing w:line="360" w:lineRule="auto"`
    #     （240=单倍，360=1.5倍） → 用 is_effective_line_spacing_1_5 逐级回溯。
    required_h1_titles_pf = [
        '第一部分：项目背景',
        '第二部分：主要关联主体情况',
        '第三部分：异常资金往来情况',
        '第四部分：业务流程变化说明',
        '第五部分：库存与资金需求分析',
        '第六部分：综合说明',
    ]
    required_norm_pf = [_norm_title(t) for t in required_h1_titles_pf]
    # 沿用之前 +3 里的 h1_paras 与 _norm_title；这里再按细则的 6 个文本严格取段落
    target_h1_paras = []
    for req in required_norm_pf:
        found = None
        for p in h1_paras:
            if _norm_title(p.text) == req:
                found = p
                break
        target_h1_paras.append(found)

    if all(p is not None for p in target_h1_paras):
        h1_para_ok = True
        for p in target_h1_paras:
            if not is_effective_justify_aligned(p):
                h1_para_ok = False
                break
            if not is_effective_line_spacing_1_5(p):
                h1_para_ok = False
                break
        if h1_para_ok:
            award(3, "一级标题段落格式两端对齐、1.5倍行距")
        else:
            fail(3, "一级标题段落格式不满足两端对齐、1.5倍行距")
    else:
        fail(3, "未找到细则指定的6个一级标题，无法核对段落格式")

    # 后续检查兼容：沿用旧变量名
    required_h1 = ['项目背景', '主要关联主体情况', '异常资金往来情况',
                   '业务流程变化说明', '库存与资金需求分析', '综合说明']

    # +3: 二级标题样式应为"一、"，如"一、提前付款情况"、"二、平台收入回流滞后情况"、"三、重复收费及费用边界不清情况"
    # 细则拆解：
    #   ① 存在二级标题（办公软件里"标题 2"样式，OOXML: p.style.name == 'Heading 2'）；
    #   ② 每个二级标题文本必须以"中文序数 + 顿号(、)"开头，如"一、"、"二、"、"三、"…；
    #   ③ 顿号后紧跟标题正文（非空），与细则示例"一、提前付款情况"结构一致。
    # 办公软件生效性：Word/WPS 的"标题 2"样式即 Heading 2；序数与顿号是段落文本的首字符，
    # 用户在编辑器里看到的就是这段字面 → 直接对 p.text 做正则匹配即可。
    h2_paras = [p for p in paras1 if p.style and p.style.name == 'Heading 2']
    h2_pattern = re.compile(r'^[一二三四五六七八九十百]+、\S')
    if h2_paras and all(h2_pattern.match(p.text.strip()) for p in h2_paras):
        award(3, "二级标题样式为'一、'（全部满足'中文序数+顿号'开头）")
    else:
        fail(3, "二级标题样式不满足'一、'格式")

    # +3: 第一个交付文档中包含"一、提前付款情况"、"二、平台收入回流滞后情况"、"三、重复收费及费用边界不清情况"
    #     这些二级标题，字体格式为黑体四号加粗
    # 细则拆解：
    #   ① 存在这 3 个字面二级标题（"标题 2"样式，Heading 2），文本按细则字面严格匹配（含中文序数与顿号）；
    #   ② 中文字体 = 黑体；
    #   ③ 字号 = 四号；
    #   ④ 加粗。
    # 办公软件生效性：
    #   - 字体/字号/加粗按"run 直设 → 段落样式 → 样式链 basedOn 回溯"解析，
    #     与 Word/WPS 里"字体对话框"实际生效顺序一致；
    #   - 通过 get_effective_east_asian_font / get_effective_size / get_effective_bold 实现。
    required_h2_titles = [
        '一、提前付款情况',
        '二、平台收入回流滞后情况',
        '三、重复收费及费用边界不清情况',
    ]

    def _norm_h2(s):
        # 去前后空白，标题字面直接比较（不做半/全角转换——细则字面即中文顿号"、"）
        return s.strip()

    h2_texts = {_norm_h2(p.text): p for p in h2_paras}
    matched_h2_paras = []
    all_present = True
    for req in required_h2_titles:
        p = h2_texts.get(_norm_h2(req))
        if p is None:
            all_present = False
            break
        matched_h2_paras.append(p)

    if all_present:
        h2_font_ok = True
        for p in matched_h2_paras:
            target_run = None
            for r in p.runs:
                if r.text.strip():
                    target_run = r
                    break
            if target_run is None:
                h2_font_ok = False
                break
            ea = get_effective_east_asian_font(target_run, p)
            sz = get_effective_size(target_run, p)
            bold = get_effective_bold(target_run, p)
            if ea != '黑体' or not size_match(sz, '四号') or not bold:
                h2_font_ok = False
                break
        if h2_font_ok:
            award(3, "包含指定三个二级标题，黑体四号加粗")
        else:
            fail(3, "三个二级标题文本齐全，但字体格式不满足黑体四号加粗")
    else:
        fail(3, "未找到细则指定的3个二级标题字面")

    # +3: 第一个交付文档中正文字体格式为宋体小四
    # 细则拆解：
    #   ① 对象为"正文"段落——办公软件里"正文"= 正文页（非封面、非目录）里非标题（非各级 Heading）、
    #      非图注（"图 N ..."）的普通段落；主标题、副标题、目录条目、页码、封面信息、图注均不属于"正文"；
    #   ② 中文字体 = 宋体；
    #   ③ 字号 = 小四（12 磅，OOXML w:sz=24）。
    # 办公软件生效性：
    #   - 字体/字号按"run 直设 → 段落 rPr → 段落样式链回溯 → 默认"解析，与 Word/WPS 实际渲染顺序一致；
    #   - 通过 get_effective_east_asian_font / get_effective_size 实现；
    #   - 遍历所有正文段落的所有非空 run，任一 run 违背即不满足（保证细则"正文字体格式为..."覆盖全部）。
    # 定位"正文"段落：
    #   - 页维度：pages1[2:] （跳过封面 pages1[0]、目录 pages1[1]）；
    #   - 段样式过滤：排除 Heading 1/2/3/…、TOC N、Title、Subtitle 等结构化样式；
    #   - 段内容过滤：跳过纯空白段落；跳过整段仅由分隔线字符 (━ ─ — ═ ⋯ ┈ 等) 组成的分隔虚线段落；
    #   - 图注过滤：排除文本以"图 N ..."开头的图注段落（细则已单独按"黑体五号"考核，不属正文）；
    #   - 页码/页眉页脚在 python-docx 的 doc.paragraphs 中不出现，无需额外排除。
    sep_chars_body = set('━─—‒═⋯┈┅┄╌╍╴╶╸╺')
    body_caption_pattern = re.compile(r'^图\s*\d+')

    def _is_body_para(p):
        text = p.text.strip()
        if not text:
            return False
        style_name = p.style.name if p.style else ''
        # 排除各级标题、目录条目样式、封面/文档标题样式
        if style_name.startswith('Heading') or style_name.startswith('TOC') \
           or style_name in ('Title', 'Subtitle', '标题', '副标题'):
            return False
        # 排除整段分隔虚线段落
        if all(ch in sep_chars_body for ch in text):
            return False
        # 排除图注段（"图 N ..."）
        if body_caption_pattern.match(text):
            return False
        return True

    body_paras = []
    for pg_idx, page_ps in enumerate(pages1):
        if pg_idx < 2:  # 跳过封面 (0)、目录 (1)
            continue
        for p in page_ps:
            if _is_body_para(p):
                body_paras.append(p)

    body_font_ok = True
    checked_runs = 0
    bad_sample = None
    for p in body_paras:
        for r in p.runs:
            if not r.text.strip():
                continue
            ea = get_effective_east_asian_font(r, p)
            sz = get_effective_size(r, p)
            if ea != '宋体':
                body_font_ok = False
                bad_sample = (r.text.strip()[:20], ea, sz)
                break
            if not size_match(sz, '小四'):
                body_font_ok = False
                bad_sample = (r.text.strip()[:20], ea, sz)
                break
            checked_runs += 1
        if not body_font_ok:
            break
    if body_paras and checked_runs > 0 and body_font_ok:
        award(3, "正文字体格式为宋体小四")
    else:
        fail(3, f"正文字体格式不满足宋体小四（首个不符样本={bad_sample}）")

    # +3: 第一个交付文档中正文段落格式为左对齐、首行缩进0.85厘米、1.25倍行距、段前0行、段后4磅
    # 细则拆解（5 个点，全部段落均须满足）：
    #   ① 左对齐（Word/WPS 里 `w:jc w:val="left"`；段落及样式链均未设则按默认左对齐）；
    #   ② 首行缩进 0.85 厘米（OOXML `w:ind w:firstLine`，单位 twips；0.85cm ≈ 482 twips）；
    #   ③ 1.25 倍行距（`w:spacing w:line="300" w:lineRule="auto"`；240=1.0，300=1.25）；
    #   ④ 段前 0 行（`w:before="0"`，或未设——效果一致：段前不留空）；
    #   ⑤ 段后 4 磅（`w:spacing w:after="80"`，20 twentieths = 4pt）。
    # 办公软件生效性：所有属性均按"段落直设 → 段落样式 → 样式链回溯"解析，
    # 与 Word/WPS "段落"对话框实际显示一致。
    body_fmt_ok = True
    fail_reason = ""
    for p in body_paras:
        # ① 左对齐
        if not is_effective_left_aligned(p):
            body_fmt_ok = False
            fail_reason = "不为左对齐"
            break
        # ② 首行缩进 0.85 厘米（允许 ±0.02 cm 内的浮点误差，对应字号磅数换算的常见小偏差）
        indent_emu = get_effective_first_line_indent_emu(p)
        if indent_emu is None:
            body_fmt_ok = False
            fail_reason = "首行缩进未设置"
            break
        # 1cm = 360000 EMU；0.85 cm = 306000 EMU
        if abs(indent_emu - 306000) > 7200:  # ±0.02 cm 容差
            body_fmt_ok = False
            fail_reason = "首行缩进不为 0.85 厘米"
            break
        # ③ 1.25 倍行距（必须为倍数模式）
        ls_mode, ls_val = get_effective_line_spacing(p)
        if ls_mode != 'multiple' or ls_val is None or abs(ls_val - 1.25) > 1e-3:
            body_fmt_ok = False
            fail_reason = "行距不为 1.25 倍"
            break
        # ④ 段前 0 行（0 磅 或 未设置——办公软件视觉一致）
        sb = get_effective_space_before_pt(p)
        if sb is not None and sb > 0.05:  # 允许极小浮点噪声
            body_fmt_ok = False
            fail_reason = "段前不为 0 行"
            break
        # ⑤ 段后 4 磅（±0.5 磅容差）
        sa = get_effective_space_after_pt(p)
        if sa is None or abs(sa - 4.0) > 0.5:
            body_fmt_ok = False
            fail_reason = "段后不为 4 磅"
            break
    if body_paras and body_fmt_ok:
        award(3, "正文段落格式为左对齐、首行缩进0.85厘米、1.25倍行距、段前0行、段后4磅")
    else:
        fail(3, f"正文段落格式不满足要求（{fail_reason or '未找到正文段落'}）")

    # +5: 表格三线表
    tables1 = doc1.tables
    all_three_line = True
    for table in tables1:
        if not check_three_line_table(table):
            all_three_line = False
            break
    if tables1 and all_three_line:
        award(5, "所有表格满足三线表格式")
    elif not tables1:
        fail(5, "文档中未找到表格")
    else:
        fail(5, "表格不满足三线表格式")

    # +3: 第一个交付文档中表格内字体格式：中文为ＭＳ 明朝小四、英文为Times New Roman 小四，表头字体需额外加粗
    # 细则拆解（4 个点，遍及表格内所有 run）：
    #   ① 中文字体 = ＭＳ 明朝（Word/WPS 里 eastAsia 字体名，支持全角"ＭＳ 明朝"与半角"MS Mincho"两种写法）；
    #   ② 英文字体 = Times New Roman（ascii 字体名）；
    #   ③ 字号 = 小四（12 磅，w:sz=24）；
    #   ④ 表头（第一行）字体额外加粗（在①②③之上再加 bold=True）。
    # 办公软件生效性：字体/字号/加粗按"run 直设 → 段落 rPr → 段落样式链 → 表格样式"解析。
    # 通过 get_effective_east_asian_font / get_effective_ascii_font / get_effective_size / get_effective_bold 实现。
    MS_MINCHO_NAMES = {'ＭＳ 明朝', 'MS 明朝', 'MS Mincho', 'ＭＳ明朝', 'MS明朝'}
    TNR_NAMES = {'Times New Roman'}

    def _check_run_font(r, p, need_bold):
        if not r.text.strip():
            return True  # 空白 run 不参与字体判定
        ea = get_effective_east_asian_font(r, p)
        asc = get_effective_ascii_font(r, p)
        sz = get_effective_size(r, p)
        bold = get_effective_bold(r, p)
        if ea not in MS_MINCHO_NAMES:
            return False
        if asc not in TNR_NAMES:
            return False
        if not size_match(sz, '小四'):
            return False
        if need_bold and not bold:
            return False
        return True

    table_font_ok = True
    if not tables1:
        table_font_ok = False
    for table in tables1:
        n_rows = len(table.rows)
        if n_rows == 0:
            table_font_ok = False
            break
        for ri, row in enumerate(table.rows):
            need_bold = (ri == 0)  # 表头行额外加粗
            seen_tcs = set()
            for cell in row.cells:
                if cell._tc in seen_tcs:
                    continue
                seen_tcs.add(cell._tc)
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not _check_run_font(r, p, need_bold):
                            table_font_ok = False
                            break
                    if not table_font_ok:
                        break
                if not table_font_ok:
                    break
            if not table_font_ok:
                break
        if not table_font_ok:
            break

    if tables1 and table_font_ok:
        award(3, "表格字体格式为ＭＳ 明朝+Times New Roman 小四，表头加粗")
    else:
        fail(3, "表格字体格式不满足要求（ＭＳ 明朝/Times New Roman/小四/表头加粗）")

    # +3: 第一个交付文档中表格第一行有淡蓝色底纹，底纹颜色编号为 D9E2F3
    # 细则拆解（3 个点）：
    #   ① 对象为"表格第一行"——遍历 doc1 所有表格的 rows[0] 的每个单元格；
    #   ② 有底纹——OOXML 里对应 `w:tcPr/w:shd`，且颜色不为 'auto'/'none'（否则办公软件里无可见底纹）；
    #   ③ 底纹颜色编号 = D9E2F3（大小写不敏感，办公软件 UI 里显示为"其他颜色 → 十六进制"D9E2F3）。
    # 办公软件生效性：Word/WPS 里第一行淡蓝色底纹通常由单元格级 `w:shd w:fill="D9E2F3"` 承载
    # （表格样式/条件格式亦可，但 python-docx 打开时 tc 上会挂到最终生效的 tcPr —— 但为对齐"要在办公软件上有效"，
    # 严格判定 tcPr.shd 或表格样式 tblStylePr[type="firstRow"] 上生效的 fill 值）。
    def _effective_first_row_fill(cell, table):
        """按办公软件生效顺序取单元格底纹填充色（大写六位十六进制），失败返回 None。
           顺序：单元格 tcPr.w:shd → 表格样式 tblStylePr type='firstRow' 的 tcPr.w:shd。"""
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                if fill and fill.lower() != 'auto':
                    return fill.upper()
                # w:color 备用（图案色）
                color = shd.get(qn('w:color'))
                if color and color.lower() != 'auto':
                    return color.upper()
        # 表格样式的 firstRow 条件格式
        style = table.style
        seen = set()
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            for tblStylePr in style.element.findall(qn('w:tblStylePr')):
                if tblStylePr.get(qn('w:type')) == 'firstRow':
                    tcPrS = tblStylePr.find(qn('w:tcPr'))
                    if tcPrS is not None:
                        shd = tcPrS.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill.lower() != 'auto':
                                return fill.upper()
            style = style.base_style
        return None

    shading_ok = True
    if not tables1:
        shading_ok = False
    for table in tables1:
        if len(table.rows) == 0:
            shading_ok = False
            break
        seen_tcs = set()
        for cell in table.rows[0].cells:
            if cell._tc in seen_tcs:
                continue
            seen_tcs.add(cell._tc)
            fill = _effective_first_row_fill(cell, table)
            if fill != 'D9E2F3':
                shading_ok = False
                break
        if not shading_ok:
            break

    if tables1 and shading_ok:
        award(3, "表格第一行底纹颜色为 D9E2F3")
    else:
        fail(3, "表格第一行底纹颜色不为 D9E2F3")

    # +1: 文档中图片有图注、图注格式为黑体五号，样式为"图1"，放置在图片下方居中
    # 细则拆解（5 个点）：
    #   ① 文档中存在图片（OOXML: run 里含 w:drawing 或 w:pict）；
    #   ② 每张图片都有对应的图注（不能只有一张图有）；
    #   ③ 图注文本样式为"图N"（如"图1"、"图 2"），允许"图"与数字间可有空白，数字部分为阿拉伯数字；
    #   ④ 图注段落位于图片段落"下方"，即紧邻在图片段落之后的下一个非空段落；
    #   ⑤ 图注段落：居中对齐 + 中文字体=黑体 + 字号=五号（首个非空 run 生效属性）。
    # 办公软件生效性：字体/字号按 effective 顺序解析（run 直设 → rPr → 样式链），
    # 与 Word/WPS 里图注段落实际渲染字体一致；对齐用 is_effective_center_aligned。
    def _paragraph_has_image(para):
        for r in para.runs:
            for elem in r._element.iter():
                if elem.tag in (qn('w:drawing'), qn('w:pict')):
                    return True
        return False

    caption_pattern = re.compile(r'^图\s*\d+')

    image_para_indices = [i for i, p in enumerate(paras1) if _paragraph_has_image(p)]

    fig_ok = None  # None=无图，True/False=判定结果
    fail_msg = ''
    if not image_para_indices:
        fig_ok = None
    else:
        fig_ok = True
        for idx in image_para_indices:
            # ④ 找图片段落"下方"的第一个非空段落作为图注候选
            caption_para = None
            j = idx + 1
            while j < len(paras1):
                if paras1[j].text.strip():
                    caption_para = paras1[j]
                    break
                j += 1
            if caption_para is None:
                fig_ok = False
                fail_msg = "存在图片下方无任何段落"
                break
            # ③ 样式为"图N"
            if not caption_pattern.match(caption_para.text.strip()):
                fig_ok = False
                fail_msg = "图注文本不满足'图N'样式"
                break
            # ⑤ 居中对齐
            if not is_effective_center_aligned(caption_para):
                fig_ok = False
                fail_msg = "图注段落未居中对齐"
                break
            # ⑤ 黑体五号（首个非空 run）
            target_run = None
            for r in caption_para.runs:
                if r.text.strip():
                    target_run = r
                    break
            if target_run is None:
                fig_ok = False
                fail_msg = "图注段落无有效文本 run"
                break
            ea = get_effective_east_asian_font(target_run, caption_para)
            sz = get_effective_size(target_run, caption_para)
            if ea != '黑体':
                fig_ok = False
                fail_msg = f"图注中文字体为'{ea}'非黑体"
                break
            if not size_match(sz, '五号'):
                fig_ok = False
                fail_msg = "图注字号不为五号"
                break

    if fig_ok is True:
        award(1, "图片图注格式正确：黑体五号、'图N'样式、图片下方居中")
    elif fig_ok is False:
        fail(1, f"图注不满足要求（{fail_msg}）")
    else:
        fail(1, "文档中未找到图片，无法评估图注")

    # === 第二个文档检查 ===

    # +1: 第二个交付文档首页主题标题改为"华中区域零售项目财务资金异常、库存占用及现金流风险专项核查材料"
    #     字体为黑体三号加粗；位于文档顶部居中
    # 细则拆解（6 个点）：
    #   ① 标题文本严格等于"华中区域零售项目财务资金异常、库存占用及现金流风险专项核查材料"；
    #   ② 中文字体 = 黑体；
    #   ③ 字号 = 三号（16 磅，w:sz=32）；
    #   ④ 加粗；
    #   ⑤ 位于"文档顶部"——即第二个交付文档第一页最上方的第一个非空段落（Word/WPS 打开后所见首个可见文本段）；
    #   ⑥ 居中对齐。
    # 办公软件生效性：字体/字号/加粗按 effective 顺序解析（run 直设 → 段落 rPr → 段落样式链）；
    # 居中用 is_effective_center_aligned；标题文本用 strip 归一化，允许标题内的空白差异，
    # 但字面必须与细则严格一致（不做子串放宽）。
    expected_title = '华中区域零售项目财务资金异常、库存占用及现金流风险专项核查材料'
    # ⑤ 定位文档顶部：第二个文档的第一页 → 第一个非空段落
    doc2_pages = get_paragraphs_by_page(paras2)
    doc2_first_page = doc2_pages[0] if doc2_pages else paras2
    title_para = None
    for p in doc2_first_page:
        if p.text.strip():
            title_para = p
            break
    if title_para is None:
        fail(1, "第二文档首页无任何可见文本段落")
    elif title_para.text.strip() != expected_title:
        fail(1, "第二文档首页标题文本与细则不符")
    else:
        # 取首个非空 run 作为字体判定基准（对应办公软件里首字符实际显示的字体）
        target_run = None
        for r in title_para.runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is None:
            fail(1, "第二文档标题段落无有效文本 run")
        else:
            ea = get_effective_east_asian_font(target_run, title_para)
            sz = get_effective_size(target_run, title_para)
            bold = get_effective_bold(target_run, title_para)
            centered = is_effective_center_aligned(title_para)
            heiti_ok = (ea == '黑体')
            size_ok = size_match(sz, '三号')
            bold_ok = bold is True
            if heiti_ok and size_ok and bold_ok and centered:
                award(1, "第二文档首页主题标题满足：文本正确、黑体三号加粗、顶部居中")
            else:
                fail(1, f"第二文档标题格式不满足（黑体={heiti_ok}, 三号={size_ok}, 加粗={bold_ok}, 居中={centered}）")

    # === 第二个文档扣分项 ===

    # -3: 第二个交付文档段落中出现空白行
    # 细则拆解：只要文档段落中"出现空白行"即扣 3 分。
    # 办公软件生效性：Word/WPS 里"空白行"= 段落里没有可见内容（无字符文本、无图片、无对象），
    # 但保留了段落回车符 —— 在编辑器里表现为一整段空白。
    # 判定：段落 text 去空白后为空，且段落内 run 不包含 w:drawing / w:pict / w:object 等可见对象。
    def _is_blank_paragraph(p):
        if p.text.strip() != '':
            return False
        for r in p.runs:
            for elem in r._element.iter():
                if elem.tag in (qn('w:drawing'), qn('w:pict'), qn('w:object')):
                    return False
        return True

    blank_count = sum(1 for p in paras2 if _is_blank_paragraph(p))
    # 契约: 命中(有空白行) → award(-3); 未命中(无空白行) → fail(-3) 以确保 dim2_items 含本项
    _rule_blank = "第二文档段落中出现空白行"
    if blank_count > 0:
        award(-3, f"{_rule_blank}（{blank_count} 处）")
    else:
        fail(-3, _rule_blank)

    # -3: 第二个交付文档一级标题字体格式不满足宋体四号加粗
    # 细则拆解（3 个点, 全部一级标题的首个非空 run 必须同时满足; 任一违背即扣 3）：
    #   ① 中文字体 = 宋体；
    #   ② 字号 = 四号（14 磅, w:sz=28）；
    #   ③ 加粗（w:b, effective bold is True）。
    # 办公软件生效性：
    #   - "一级标题" = Word/WPS 中应用"标题 1"样式的段落（p.style.name == 'Heading 1'）;
    #     不再硬编码段落索引 [3, 11, 13, ...], 避免文档结构轻微增减导致误评;
    #   - 字体/字号/加粗按"run 直设 → 段落 rPr → 段落样式链回溯"解析（effective 顺序），
    #     与 Word/WPS 打开"字体"对话框看到的最终生效值一致;
    #   - 三个维度中**任一**不满足即置 doc2_h1_font_fail = True（修正原代码只在字号不合格时才扣分的漏洞）;
    #   - 空标题段落（无非空 run）跳过, 不参与字体判定.
    doc2_h1_paras_all = [p for p in paras2 if p.style and p.style.name == 'Heading 1']
    doc2_h1_font_fail = False
    doc2_h1_bad_reason = ""
    for p in doc2_h1_paras_all:
        target_run = None
        for r in p.runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is None:
            continue
        ea = get_effective_east_asian_font(target_run, p)
        sz = get_effective_size(target_run, p)
        bold = get_effective_bold(target_run, p)
        if ea != '宋体' or not size_match(sz, '四号') or bold is not True:
            doc2_h1_font_fail = True
            sz_pt = (sz / 12700.0) if sz else None
            doc2_h1_bad_reason = f"实测 eastAsia={ea!r}, 字号={sz_pt}pt, 加粗={bold}: {p.text.strip()[:12]!r}"
            break
    # 契约: 有 Heading 1 且字体不达标 → award(-3); 其余(无 Heading 1 或字体达标) → fail(-3) 保留 item
    _rule_h1_font = "第二文档一级标题字体不满足宋体四号加粗"
    if doc2_h1_paras_all and doc2_h1_font_fail:
        award(-3, f"{_rule_h1_font}（{doc2_h1_bad_reason}）")
    else:
        fail(-3, _rule_h1_font)

    # -3: 第二个交付文档一级标题段落格式不满足两端对齐、首行缩进两字符、1.5倍行距
    # 细则拆解（3 个点，任一不满足即扣 3）：
    #   ① 两端对齐（w:jc w:val="both"，python-docx = WD_ALIGN_PARAGRAPH.JUSTIFY）；
    #   ② 首行缩进两字符（Word/WPS 里 UI 显示"以字符为单位的度量"= 2）；
    #      OOXML: `w:ind w:firstLineChars="200"`（一字符=100，两字符=200）；
    #      或等价的绝对值：2 × 段落字体磅值（全角字符宽度 ≈ 字号磅值），换算为 twips=20×pt；
    #   ③ 1.5 倍行距（w:spacing w:line="360" w:lineRule="auto"，或 python-docx line_spacing==1.5）。
    # 办公软件生效性：段落对齐/首行缩进/行距均按"段落直设 → 段落 pPr → 样式链回溯"解析；
    # 一级标题 = "标题 1"样式（Heading 1）。若段落缺失属性由样式链继承，Word/WPS 打开时按继承值渲染。
    def _get_effective_first_line_chars(para):
        """按办公软件生效顺序取 w:ind w:firstLineChars（Word/WPS UI 中"缩进量 = X 字符"的直接来源）。
        返回整数（每 100 表示 1 字符），未设置返回 None。"""
        for pPr in _iter_para_style_pPr(para):
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                flc = ind.get(qn('w:firstLineChars'))
                if flc is not None:
                    try:
                        return int(flc)
                    except ValueError:
                        return None
                # 若显式指定了绝对 firstLine 或 hanging，则以字符为单位的值不适用 → 返回 None
                if ind.get(qn('w:firstLine')) is not None or ind.get(qn('w:hanging')) is not None:
                    return None
        return None

    def _get_para_effective_size_pt(para):
        """取段落首个非空 run 的中文字号磅数；全角字符宽度 ≈ 字号磅数（用于换算"两字符"绝对缩进）。"""
        for r in para.runs:
            if r.text.strip():
                emu = get_effective_size(r, para)
                if emu is not None:
                    return emu / 12700.0
                break
        return None

    def _first_line_indent_is_two_chars(para):
        """在办公软件里判"首行缩进两字符"：
           ① 优先看 firstLineChars = 200（Word/WPS 原生"字符单位"缩进）；
           ② 否则用绝对值 firstLine（twips）与 2×字号pt×20 twips 比较（±10% 容差）。"""
        flc = _get_effective_first_line_chars(para)
        if flc is not None:
            return flc == 200
        emu = get_effective_first_line_indent_emu(para)
        if emu is None:
            return False
        # 换算：绝对缩进 (磅) = emu / 12700
        indent_pt = emu / 12700.0
        size_pt = _get_para_effective_size_pt(para)
        if size_pt is None or size_pt <= 0:
            return False
        expected_pt = 2 * size_pt  # 两个全角字符宽度
        return abs(indent_pt - expected_pt) <= 0.1 * expected_pt

    doc2_h1_paras_all = [p for p in paras2 if p.style and p.style.name == 'Heading 1']
    doc2_h1_fmt_fail = False
    for p in doc2_h1_paras_all:
        # ① 两端对齐
        if not is_effective_justify_aligned(p):
            doc2_h1_fmt_fail = True
            break
        # ② 首行缩进两字符
        if not _first_line_indent_is_two_chars(p):
            doc2_h1_fmt_fail = True
            break
        # ③ 1.5 倍行距
        if not is_effective_line_spacing_1_5(p):
            doc2_h1_fmt_fail = True
            break
    _rule_h1_fmt = "第二文档一级标题段落格式不满足两端对齐、首行缩进两字符、1.5倍行距"
    if doc2_h1_paras_all and doc2_h1_fmt_fail:
        award(-3, _rule_h1_fmt)
    else:
        fail(-3, _rule_h1_fmt)

    # -3: 第二个交付文档二级标题字体格式不满足宋体小四号加粗
    # 细则拆解（3 个点，全部二级标题的首个非空 run 必须同时满足；任一违背即扣 3）：
    #   ① 中文字体 = 宋体；
    #   ② 字号 = 小四（12 磅，w:sz=24）；
    #   ③ 加粗。
    # 办公软件生效性：
    #   - "二级标题"= Word/WPS 里应用了"标题 2"样式的段落（p.style.name == 'Heading 2'）；
    #   - 字体/字号/加粗按"run 直设 → 段落 rPr → 段落样式链回溯"解析——
    #     用户常在"标题 2"样式上定义字体/加粗，effective 解析可正确读到最终渲染字体。
    doc2_h2_paras_all = [p for p in paras2 if p.style and p.style.name == 'Heading 2']
    doc2_h2_font_fail = False
    for p in doc2_h2_paras_all:
        target_run = None
        for r in p.runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is None:
            # 空标题段落不参与字体判定（无字符可评）
            continue
        ea = get_effective_east_asian_font(target_run, p)
        sz = get_effective_size(target_run, p)
        bold = get_effective_bold(target_run, p)
        if ea != '宋体' or not size_match(sz, '小四') or bold is not True:
            doc2_h2_font_fail = True
            break
    _rule_h2_font = "第二文档二级标题字体格式不满足宋体小四号加粗"
    if doc2_h2_paras_all and doc2_h2_font_fail:
        award(-3, _rule_h2_font)
    else:
        fail(-3, _rule_h2_font)

    # 兼容后续未改造的检查项：doc2_h2_indices 仍需被引用（若存在）
    doc2_h2_indices = [i for i, p in enumerate(paras2) if p.style and p.style.name == 'Heading 2']

    # -3: 第二个交付文档二级标题段落格式不满足左对齐、首行缩进两字符、1.5倍行距
    # 细则拆解（3 个点，任一不满足即扣 3）：
    #   ① 左对齐（w:jc w:val="left"；段落及样式链未设时办公软件按默认左对齐渲染，视为满足）；
    #   ② 首行缩进两字符（Word/WPS UI 中"缩进量 = 2 字符"→ OOXML `w:ind w:firstLineChars="200"`，
    #      或等价的绝对值 firstLine ≈ 2 × 字号pt × 20 twips，±10% 容差）；
    #   ③ 1.5 倍行距（w:spacing w:line="360" w:lineRule="auto"）。
    # 办公软件生效性：
    #   - "二级标题"= 应用"标题 2"样式的段落（p.style.name == 'Heading 2'）；
    #   - 对齐/缩进/行距均按"段落直设 → 段落 pPr → 样式链回溯"解析，读取到 Word/WPS 实际渲染值。
    doc2_h2_paras_all_fmt = [p for p in paras2 if p.style and p.style.name == 'Heading 2']
    doc2_h2_fmt_fail = False
    for p in doc2_h2_paras_all_fmt:
        # ① 左对齐
        if not is_effective_left_aligned(p):
            doc2_h2_fmt_fail = True
            break
        # ② 首行缩进两字符
        if not _first_line_indent_is_two_chars(p):
            doc2_h2_fmt_fail = True
            break
        # ③ 1.5 倍行距
        if not is_effective_line_spacing_1_5(p):
            doc2_h2_fmt_fail = True
            break
    _rule_h2_fmt = "第二文档二级标题段落格式不满足左对齐、首行缩进两字符、1.5倍行距"
    if doc2_h2_paras_all_fmt and doc2_h2_fmt_fail:
        award(-3, _rule_h2_fmt)
    else:
        fail(-3, _rule_h2_fmt)

    # -3: 第二个交付文档三级标题字体格式不满足宋体小四号加粗
    # 细则拆解（3 个点，全部三级标题的首个非空 run 必须同时满足；任一违背即扣 3）：
    #   ① 中文字体 = 宋体；
    #   ② 字号 = 小四（12 磅，w:sz=24）；
    #   ③ 加粗。
    # 办公软件生效性：
    #   - "三级标题"= Word/WPS 里应用了"标题 3"样式的段落（p.style.name == 'Heading 3'）；
    #   - 字体/字号/加粗按"run 直设 → 段落 rPr → 段落样式链回溯"解析（effective 顺序），
    #     与 Word/WPS 打开"字体"对话框看到的最终生效值一致。
    doc2_h3_paras = [p for p in paras2 if p.style and p.style.name == 'Heading 3']
    doc2_h3_font_fail = False
    for p in doc2_h3_paras:
        target_run = None
        for r in p.runs:
            if r.text.strip():
                target_run = r
                break
        if target_run is None:
            continue
        ea = get_effective_east_asian_font(target_run, p)
        sz = get_effective_size(target_run, p)
        bold = get_effective_bold(target_run, p)
        if ea != '宋体' or not size_match(sz, '小四') or bold is not True:
            doc2_h3_font_fail = True
            break
    _rule_h3_font = "第二文档三级标题字体格式不满足宋体小四号加粗"
    if doc2_h3_paras and doc2_h3_font_fail:
        award(-3, _rule_h3_font)
    else:
        fail(-3, _rule_h3_font)

    # -3: 第二个交付文档三级标题段落格式不满足两端对齐、首行缩进两字符、1.5倍行距
    # 细则拆解（3 个点，任一不满足即扣 3）：
    #   ① 两端对齐（w:jc w:val="both"，python-docx = WD_ALIGN_PARAGRAPH.JUSTIFY）；
    #   ② 首行缩进两字符（Word/WPS UI 里"缩进量 = 2 字符"→ OOXML `w:ind w:firstLineChars="200"`，
    #      或等价的绝对值 firstLine ≈ 2 × 字号pt × 20 twips，±10% 容差）；
    #   ③ 1.5 倍行距（w:spacing w:line="360" w:lineRule="auto"）。
    # 办公软件生效性：
    #   - "三级标题"按 **内容视角** 识别（用户在办公软件里看到的层级，而非 OOXML style 名）：
    #     文本以阿拉伯数字 + 顿号/句点/右括号开头，如 "1、xxx" / "1．xxx" / "1) xxx"；
    #   - 若文档没有此类段落，则三级标题集合为空，此检查项前提不成立 → 不扣分；
    #   - 对齐/缩进/行距均按"段落直设 → 段落 pPr → 样式链回溯"解析，读取到 Word/WPS 实际渲染值。
    h3_content_pattern = re.compile(r'^\s*\d+\s*[、．\.\)）]')
    doc2_h3_paras_content = [p for p in paras2 if h3_content_pattern.match(p.text.strip() or '')]
    doc2_h3_fmt_fail = False
    for p in doc2_h3_paras_content:
        # ① 两端对齐
        if not is_effective_justify_aligned(p):
            doc2_h3_fmt_fail = True
            break
        # ② 首行缩进两字符
        if not _first_line_indent_is_two_chars(p):
            doc2_h3_fmt_fail = True
            break
        # ③ 1.5 倍行距
        if not is_effective_line_spacing_1_5(p):
            doc2_h3_fmt_fail = True
            break
    _rule_h3_fmt = "第二文档三级标题段落格式不满足两端对齐、首行缩进两字符、1.5倍行距"
    if doc2_h3_paras_content and doc2_h3_fmt_fail:
        award(-3, _rule_h3_fmt)
    else:
        fail(-3, _rule_h3_fmt)

    # -3: 第二个交付文档正文除数字及表格外字体格式不满足宋体小四，正文中公司名字未加粗
    #     （如："远艘零售项目部"、"华梅商业运营有限公司"、"云囍供应链有限公司"、"嘉笔数字商贸有限公司"）
    # 细则拆解（复合，两大点，任一违背即扣 3）：
    #   ① 正文（除数字及表格外）字体格式 = 宋体小四；
    #   ② 正文中出现细则列出的 4 个公司名字时，其字符必须加粗。
    # "正文"定义（只针对办公软件里视觉上真正的正文段落）：
    #   - 排除表格内容——python-docx 的 doc.paragraphs（即 paras2）**天然不含表格单元格内的段落**，
    #     刚好对应细则里"除...表格外"；
    #   - 排除文档主标题所在页（doc2 首页是黑体 18pt 的主标题，办公软件视觉上不属于正文）；
    #   - 排除 Heading 1/2/3、TOC、Title/Subtitle 等结构化样式段；
    #   - 排除**内容型标题**——doc2 里"一、基本概况"、"（一）项目概况"等一/二级标题用 style='Normal' 承载，
    #     必须按内容识别 (`^[一二三四…]+[、．.]` / `^[（(][一二三四…]+[）)]`) 才能与办公软件视觉一致；
    #   - 排除图注段（"图 N ..."，已由"图注格式"独立考核）；
    #   - 排除末尾签名段（"核查编制人："/"编制日期："，已由末尾签名相关细则独立考核）；
    #   - 排除整段仅由分隔线字符组成的段落。
    # 字体判定：只针对"非数字"字符——数字在办公软件里通常用 Times New Roman，属于细则的"除数字外"。
    # 加粗判定：以段落文本为字符流，定位公司名字面出现的位置，落到覆盖这些字符的 run 上，
    #   要求这些 run 全部 effective bold=True；这与 Word/WPS 里"选中公司名 → 加粗"的行为一致。
    company_names = ['远艘零售项目部', '华梅商业运营有限公司', '云囍供应链有限公司', '嘉笔数字商贸有限公司']

    def _run_has_non_digit_chinese(text):
        """判断 run 文本里是否含有需按'宋体小四'判定的字符：
        排除纯数字/空白/常见半全角标点后仍有内容 → 需判定。"""
        stripped = text.strip()
        stripped = stripped.replace(' ', '').replace('　', '')
        # 去掉数字与常见半/全角标点
        remains = re.sub(r'[0-9０-９.,，。；;：:%()\[\](){}【】、！!？?\-—“”"\'\'/…·]', '', stripped)
        return len(remains) > 0

    _pages2_body = get_paragraphs_by_page(paras2)
    _sep_chars_body_doc2 = set('━─—‒═⋯┈┅┄╌╍╴╶╸╺')
    _caption_pattern_doc2 = re.compile(r'^图\s*\d+')
    _h1_content_pattern_doc2 = re.compile(r'^\s*[一二三四五六七八九十百]+\s*[、．\.]')
    _h2_content_pattern_doc2 = re.compile(r'^\s*[（(][一二三四五六七八九十百]+[）)]')
    _sig_pattern_doc2 = re.compile(r'^\s*(核查编制人|编制日期)\s*[:：]')

    def _is_body_para_doc2(p):
        text = p.text.strip()
        if not text:
            return False
        style_name = p.style.name if p.style else ''
        # 排除各级标题、目录条目样式、封面/文档标题样式
        if style_name.startswith('Heading') or style_name.startswith('TOC') \
           or style_name in ('Title', 'Subtitle', '标题', '副标题'):
            return False
        # 排除整段分隔虚线段落
        if all(ch in _sep_chars_body_doc2 for ch in text):
            return False
        # 排除图注段
        if _caption_pattern_doc2.match(text):
            return False
        # 排除内容型一/二级标题（doc2 里用 style='Normal' 承载）
        if _h1_content_pattern_doc2.match(text) or _h2_content_pattern_doc2.match(text):
            return False
        # 排除末尾签名段
        if _sig_pattern_doc2.match(text):
            return False
        return True

    body_paras_doc2 = []
    for _pg_idx_b, _page_ps_b in enumerate(_pages2_body):
        if _pg_idx_b < 1:  # 跳过文档主标题所在页（doc2 首页为主标题，非正文）
            continue
        for _p_b in _page_ps_b:
            if _is_body_para_doc2(_p_b):
                body_paras_doc2.append(_p_b)

    body_font_fail_doc2 = False
    for p in body_paras_doc2:
        for r in p.runs:
            if not r.text.strip():
                continue
            if not _run_has_non_digit_chinese(r.text):
                # 纯数字/标点/空白 run → 按细则"除数字外"跳过
                continue
            ea = get_effective_east_asian_font(r, p)
            sz = get_effective_size(r, p)
            if ea != '宋体' or not size_match(sz, '小四'):
                body_font_fail_doc2 = True
                break
        if body_font_fail_doc2:
            break

    company_bold_fail = False
    for p in body_paras_doc2:
        text = p.text
        for name in company_names:
            start = 0
            while True:
                pos = text.find(name, start)
                if pos == -1:
                    break
                end = pos + len(name)
                # 遍历 runs，找出覆盖 [pos, end) 的所有 run，全部必须 effective bold
                cursor = 0
                for r in p.runs:
                    r_len = len(r.text)
                    r_start = cursor
                    r_end = cursor + r_len
                    cursor = r_end
                    # 计算与 [pos, end) 的重叠
                    if r_end <= pos or r_start >= end:
                        continue
                    if get_effective_bold(r, p) is not True:
                        company_bold_fail = True
                        break
                if company_bold_fail:
                    break
                start = pos + 1
            if company_bold_fail:
                break
        if company_bold_fail:
            break

    _rule_body_font = "第二文档正文（除数字/表格外）字体不为宋体小四或正文中公司名未全部加粗"
    if body_font_fail_doc2 or company_bold_fail:
        reasons = []
        if body_font_fail_doc2:
            reasons.append("正文（除数字/表格外）字体不为宋体小四")
        if company_bold_fail:
            reasons.append("正文中公司名未全部加粗")
        award(-3, f"第二文档{'；'.join(reasons)}")
    else:
        fail(-3, _rule_body_font)

    # -5: 第二个交付文档中出现表格不满足三线表
    # 细则拆解（三线表 4 个点，任一违背即"不满足"；只要出现任一张不满足的表即扣 5）：
    #   ① 上边框 = 1.5 磅（w:sz=12）——首行每个单元格有效顶边 = sz 12；
    #   ② 下边框 = 1.5 磅（w:sz=12）——末行每个单元格有效底边 = sz 12；
    #   ③ 表格只有上下框线、第一行存在下框线（0.5 磅，w:sz=4）——首行底边 = sz 4；
    #   ④ 其余位置无框线——单元格左/右边、非首行的顶边、非末行的底边（除首行底边）均无边框。
    # 办公软件生效性：完整判定实现于 check_three_line_table
    #   （已按 tcBorders → tblBorders 生效顺序、外沿/内沿分别处理、合并单元格去重）。
    # 前提：文档中"出现"表格 —— 若无表格，本项不成立，不扣分。
    tables2 = doc2.tables
    doc2_table_fail = False
    for table in tables2:
        if not check_three_line_table(table):
            doc2_table_fail = True
            break
    _rule_table_3line = "第二文档中出现表格不满足三线表（上下1.5磅/首行下0.5磅/其余无框线）"
    if tables2 and doc2_table_fail:
        award(-5, _rule_table_3line)
    else:
        fail(-5, _rule_table_3line)

    # -3: 第二个交付文档中表格内字体格式：中文为宋体小四、英文为 Times New Roman 小四，表头字体需额外加粗
    # 细则拆解（4 个点，任一违背即扣 3）：
    #   ① 中文字体 = 宋体（仅对 run 内实际含中文字符的情况判 eastAsia 字体）；
    #   ② 英文字体 = Times New Roman（仅对 run 内实际含 ASCII 字母/数字的情况判 ascii 字体）；
    #   ③ 字号 = 小四（12 磅，w:sz=24）——所有 run 都判；
    #   ④ 表头（第一行）在①②③之上额外加粗。
    # 办公软件生效性：
    #   - Word/WPS 渲染时 eastAsia 字段只作用于 CJK 字符、ascii 字段只作用于 ASCII 字符——
    #     run 里若无中文，eastAsia 字段视觉上不生效；若无英文/数字，ascii 字段视觉上不生效；
    #   - 因此按 run 实际字符类别分别判定，与"办公软件里看到的字体"一致；
    #   - 字体/字号/加粗按"run 直设 → 段落 rPr → 段落样式链 → 表格样式"顺序解析（effective）；
    #   - 合并单元格用 seen_tcs（按 _tc 元素）去重；
    #   - 空白 run（无字符）不参与判定。
    # 前提：文档中至少有一张表格 → 无表格不构成扣分。
    def _check_doc2_run_font(r, p, need_bold):
        text = r.text
        if not text.strip():
            return True
        has_cjk = any('⺀' <= ch <= '鿿' or '㐀' <= ch <= '䶿'
                      or '豈' <= ch <= '﫿' for ch in text)
        has_ascii_alnum = any(('A' <= ch <= 'Z') or ('a' <= ch <= 'z')
                              or ('0' <= ch <= '9') for ch in text)
        sz = get_effective_size(r, p)
        bold = get_effective_bold(r, p)
        if has_cjk:
            ea = get_effective_east_asian_font(r, p)
            if ea != '宋体':
                return False
        if has_ascii_alnum:
            asc = get_effective_ascii_font(r, p)
            if asc != 'Times New Roman':
                return False
        if not size_match(sz, '小四'):
            return False
        if need_bold and bold is not True:
            return False
        return True

    doc2_table_font_fail = False
    for table in tables2:
        n_rows = len(table.rows)
        if n_rows == 0:
            continue
        for ri, row in enumerate(table.rows):
            need_bold = (ri == 0)  # 表头行额外加粗
            seen_tcs = set()
            for cell in row.cells:
                if cell._tc in seen_tcs:
                    continue
                seen_tcs.add(cell._tc)
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not _check_doc2_run_font(r, p, need_bold):
                            doc2_table_font_fail = True
                            break
                    if doc2_table_font_fail:
                        break
                if doc2_table_font_fail:
                    break
            if doc2_table_font_fail:
                break
        if doc2_table_font_fail:
            break

    _rule_table_font = "第二文档表格字体格式不满足要求（中文宋体/英文Times New Roman/小四/表头加粗）"
    if tables2 and doc2_table_font_fail:
        award(-3, _rule_table_font)
    else:
        fail(-3, _rule_table_font)

    # -1: 第二个交付文档中末尾"核查编制人：林远芝"、"编制日期：2026年5月20日"
    #     中文字体格式不满足宋体四号加粗、阿拉伯数字字体格式不满足 Calibri 四号加粗
    # 细则拆解（针对指定两段末尾签名，任一违背即扣 1）：
    #   两段目标文本（细则字面，允许半/全角冒号差异）：
    #     A) "核查编制人：林远芝"
    #     B) "编制日期：2026年5月20日"
    #   对每一个非空 run，按其字符类型分类逐字符判定：
    #     ① 中文（含中文标点等非 ASCII 字符）→ 中文字体 = 宋体、字号 = 四号（w:sz=28）、加粗；
    #     ② 阿拉伯数字（0-9）→ 英文字体 = Calibri、字号 = 四号、加粗；
    # 办公软件生效性：
    #   - 定位末尾两段用文本匹配（不依赖脆弱的段落硬编码索引 [69, 70]），支持半/全角冒号；
    #   - 字体/字号/加粗均按 effective 顺序解析（run 直设 → 段落 rPr → 样式链回溯），
    #     与 Word/WPS 里实际渲染字体一致。
    def _norm_sig(s):
        # 归一化：strip、半角冒号→全角冒号、去掉所有空白字符（含 ASCII 空格、全角空格、Tab）——
        # Word/WPS 里"编制日期：2026年5月20日"因 autoSpaceDE/DN 常被拆成
        # "编制日期：2026 年 5 月 20 日"（中英之间自动插入空格），视觉与细则一致但文本层含空格。
        return re.sub(r'\s+', '', s.strip().replace(':', '：'))

    signature_targets = ['核查编制人：林远芝', '编制日期：2026年5月20日']
    signature_targets_norm = [_norm_sig(t) for t in signature_targets]

    signature_paras = []
    for p in paras2:
        if _norm_sig(p.text) in signature_targets_norm:
            signature_paras.append(p)

    end_font_fail = False
    if len(signature_paras) < len(signature_targets):
        # 目标签名段缺失——文本层面已不满足细则前提，字体判定无法落地
        end_font_fail = True
    else:
        for p in signature_paras:
            for r in p.runs:
                if not r.text.strip():
                    continue
                text = r.text
                has_chinese = any('一' <= ch <= '鿿' or ch in '：，。、（）【】' for ch in text)
                has_digit = any(ch.isdigit() for ch in text)
                ea = get_effective_east_asian_font(r, p)
                asc = get_effective_ascii_font(r, p)
                sz = get_effective_size(r, p)
                bold = get_effective_bold(r, p)
                if has_chinese:
                    if ea != '宋体' or not size_match(sz, '四号') or bold is not True:
                        end_font_fail = True
                        break
                if has_digit:
                    if asc != 'Calibri' or not size_match(sz, '四号') or bold is not True:
                        end_font_fail = True
                        break
            if end_font_fail:
                break

    _rule_sig_font = "第二文档末尾签名字体格式不满足要求（中文宋体四号加粗 / 数字Calibri四号加粗）"
    if end_font_fail:
        award(-1, _rule_sig_font)
    else:
        fail(-1, _rule_sig_font)

    # -1: 第二个交付文档中末尾"核查编制人：林远芝"、"编制日期：2026年5月20日"
    #     段落格式不满足右对齐、单倍行距，置于文档末尾，分别成段
    # 细则拆解（任一违背即扣 1）：
    #   ① 存在两段目标签名（字面）：
    #        A) "核查编制人：林远芝"
    #        B) "编制日期：2026年5月20日"
    #   ② 每一段段落对齐为"右对齐"（w:jc="right"），按 effective 顺序解析；
    #   ③ 每一段行距为"单倍行距"（1.0 倍），按 effective 顺序解析；
    #   ④ 置于文档末尾——两段之后不允许再出现其它非空正文段落；
    #   ⑤ 分别成段——两段各自独占一个段落（不合并成同一段）。
    # 办公软件生效性：
    #   - 对齐 / 行距均通过 effective 解析（段落 pPr → 样式链），与 Word/WPS 一致；
    #   - "置于文档末尾"用 paras2 里目标段之后的段落均为空行来判定；
    #   - 目标定位用文本匹配，兼容半/全角冒号。
    end_para_fail = False

    sig_indices = []
    for i, p in enumerate(paras2):
        if _norm_sig(p.text) in signature_targets_norm:
            sig_indices.append((i, _norm_sig(p.text)))

    # ① 两段签名都必须存在
    found_texts = {t for _, t in sig_indices}
    if not all(t in found_texts for t in signature_targets_norm):
        end_para_fail = True
    else:
        # ⑤ 分别成段：两段目标文本各自对应至少一个独立段落（不同段落索引）
        idx_a = next((i for i, t in sig_indices if t == signature_targets_norm[0]), None)
        idx_b = next((i for i, t in sig_indices if t == signature_targets_norm[1]), None)
        if idx_a is None or idx_b is None or idx_a == idx_b:
            end_para_fail = True
        else:
            # ④ 置于文档末尾：两段之后的段落必须全部为空（无正文内容）
            last_sig_idx = max(idx_a, idx_b)
            for j in range(last_sig_idx + 1, len(paras2)):
                if paras2[j].text.strip():
                    end_para_fail = True
                    break

            # ② 右对齐 + ③ 单倍行距（对两段分别校验）
            for i, _ in sig_indices:
                p = paras2[i]
                if not is_effective_right_aligned(p):
                    end_para_fail = True
                    break
                if not is_effective_single_line_spacing(p):
                    end_para_fail = True
                    break

    _rule_sig_para = "第二文档末尾签名段落格式不满足右对齐、单倍行距、置于文档末尾、分别成段"
    if end_para_fail:
        award(-1, _rule_sig_para)
    else:
        fail(-1, _rule_sig_para)

    return score, details, items


# =================== 统一入口 ===================


def evaluate(dir_path: str) -> dict:
    """统一评估入口：传入脚本所在目录路径，脚本自行在该目录内定位并打开被评估文档。"""
    result = {
        "id": SCRIPT_ID,
        "file_name": None,
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }
    try:
        if not os.path.isdir(dir_path):
            raise FileNotFoundError(f"目录不存在: {dir_path}")
        file1, file2 = _resolve_docs(dir_path)
        missing = [path for path in (file1, file2) if not os.path.isfile(path)]
        if missing:
            raise FileNotFoundError(
                "目录内缺少必需文档: " + ", ".join(os.path.basename(path) for path in missing)
            )
        result["file_name"] = f"{os.path.basename(file1)} / {os.path.basename(file2)}"

        dim1_pass, dim1_results = check_dimension_one(file1, file2)
        result["dim1_pass"] = dim1_pass
        result["dim1_reason"] = "" if dim1_pass else "；".join(dim1_results)

        if not dim1_pass:
            return result

        _score, _details, dim2_items = check_dimension_two(file1, file2)
        result["dim2_items"] = dim2_items
        # 得分 = 命中的加分项 + 命中的扣分项 之和（即 check_dimension_two 的净得分）
        result["total_score"] = sum(item["delta"] for item in dim2_items)
        # 总分 = 所有加分项（正分规则）的满分之和，不受扣分项影响
        result["max_score"] = sum(
            item["max_delta"] for item in dim2_items if item["max_delta"] > 0 and item["delta"] >= 0
        )
        return result
    except Exception as e:
        result["status"] = "error"
        result["error"] = str(e)
        return result


if __name__ == '__main__':
    import json
    target_dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    print(json.dumps(evaluate(target_dir), ensure_ascii=False, indent=2))
