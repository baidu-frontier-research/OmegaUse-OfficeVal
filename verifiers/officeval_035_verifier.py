"""
自动评估脚本：外文文献翻译_邻里阅读花园项目.docx
统一对外接口：evaluate(dir_path: str) -> dict
  参数 dir_path 为脚本所在目录，脚本自行在该目录里定位并打开被评估文档。
评分逻辑：
  维度1 不满足 → 0 分，返回 dim1_pass=False；
  维度2 逐项累计得分点，命中与未命中均记录到 dim2_items。
"""
import json
import re
import sys
from pathlib import Path


import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xml.etree.ElementTree as ET

SCRIPT_ID = "035"
DOC_FILENAME = "外文文献翻译_邻里阅读花园项目.docx"

# 由 evaluate() 在运行时赋值，供模块级辅助/校验函数以全局形式访问被评估文档。
# doc: python-docx Document 对象；_doc_path: 被评估文档的绝对路径 Path。
doc = None
_doc_path = None

# 维度二检查项静态注册表：元素为 (predicate_fn, score, desc)。
# 各 check_xxx() 定义完毕后向 _CHECKS 追加登记，evaluate() 内按登记顺序执行。
_CHECKS = []


def _dim1_check_document():
    """维度一文档级校验，返回 (True, "") 或 (False, reason)。
    1a（文件是否为可打开的 .docx）在 evaluate() 里于打开阶段直接完成。
    说明：原 1b（整篇/整页转图片）、1c（乱码）、1d（连续空白页/图表错位/正文超边界）
    对应的两条细则已按要求删除，这里不再做这些校验，只要文件是可打开的 .docx 即视为通过。"""
    return True, ""

# ─────────────────────────────────────────
# 辅助函数
# ─────────────────────────────────────────

def get_all_text():
    return "\n".join(p.text for p in doc.paragraphs)

def para_font_name(para):
    """获取段落中文字体名。
    Word 里中文字符实际使用 rFonts 的 eastAsia 字体，ascii/hAnsi 只作用于西文。
    因此判定中文字体（黑体/宋体等）时必须优先取 eastAsia，
    否则会误取到西文字体(如 Times New Roman)导致中文字体判定失败。"""
    for run in para.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'))
                if ea:
                    return ea
        # 无显式 eastAsia 时，回退到 python-docx 解析出的字体名
        if run.font.name:
            return run.font.name
    return None

def para_font_size_pt(para):
    """获取段落第一个run的字号(pt)"""
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                if val:
                    return int(val) / 2
    return None

def para_alignment(para):
    return para.alignment

def para_jc_val(para):
    """获取段落对齐方式 XML 值，如 center/left/both。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        return None
    return jc.get(qn('w:val'))

def is_center_aligned(para):
    return para.alignment == WD_ALIGN_PARAGRAPH.CENTER or para_jc_val(para) == 'center'

def is_left_aligned(para):
    jc = para_jc_val(para)
    return para.alignment == WD_ALIGN_PARAGRAPH.LEFT or jc in ('left', 'start')

def is_justify_aligned(para):
    jc = para_jc_val(para)
    return para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY or jc in ('both', 'distribute')

def is_distributed_aligned(para):
    """分散对齐：jc=distribute。python-docx 的枚举里 DISTRIBUTE 亦可。"""
    jc = para_jc_val(para)
    if jc == 'distribute':
        return True
    align = para.alignment
    return align is not None and 'DISTRIBUTE' in str(align)

def is_single_line_spacing(para):
    lineRule, line = get_line_spacing(para)
    return lineRule == 'auto' and line == '240'

def has_font_name(para, *names):
    fn = para_font_name(para)
    if not fn:
        return False
    fn_lower = fn.lower()
    return any(name in fn or name.lower() in fn_lower for name in names)

def is_size_near(para, target_pt, tolerance=0.6):
    sz = para_font_size_pt(para)
    return sz is not None and abs(sz - target_pt) <= tolerance

def is_bold(para):
    for run in para.runs:
        if run.bold:
            return True
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            b = rPr.find(qn('w:b'))
            if b is not None:
                return True
    return False

def get_line_spacing(para):
    """返回行距规则和值"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None, None
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        return None, None
    line = spacing.get(qn('w:line'))
    lineRule = spacing.get(qn('w:lineRule'))
    return lineRule, line  # lineRule='auto' line=240 => 单倍

def get_indent(para):
    """返回首行缩进(twips)"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return None
    firstLine = ind.get(qn('w:firstLine'))
    if firstLine:
        return int(firstLine)  # twips, 1cm=567, 0.85cm≈481
    return None

def find_paragraphs_containing(text, case_insensitive=False):
    result = []
    for p in doc.paragraphs:
        t = p.text
        if case_insensitive:
            if text.lower() in t.lower():
                result.append(p)
        else:
            if text in t:
                result.append(p)
    return result

def get_tables():
    return doc.tables

def para_space_before_after(para):
    """返回段前段后的“行数”。
    Word 中“段前/段后 X 行”按行设置时存于 w:beforeLines/afterLines（单位=1/100 行），
    这才是真正的“行”单位；而 w:before/after 是绝对值(twips=磅)，不能按 240 换算成行
    （240twips=12磅只在行高恰为12磅时才等于1行，对大字号并不成立）。
    因此仅当存在 beforeLines/afterLines 时才返回对应行数，否则视为未按“行”设置(None)。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None, None
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        return None, None
    before_lines = spacing.get(qn('w:beforeLines'))
    after_lines = spacing.get(qn('w:afterLines'))
    # beforeLines/afterLines 单位为 1/100 行
    b = int(before_lines)/100 if before_lines is not None else None
    a = int(after_lines)/100 if after_lines is not None else None
    return b, a

# ─────────────────────────────────────────
# 维度2
# ─────────────────────────────────────────
# 维度2 开始评分


# ── +5 文档中除"参考文献"下方内容外其余文本皆为中文 ──
def _iter_block_texts():
    """按文档正文顺序产出 (kind, text)：kind 为 'para' 或 'table'。
    这样能覆盖正文段落与表格单元格全部文本，符合细则“文档中……其余文本”的范围。"""
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map = {t._element: t for t in doc.tables}
    for child in body.iterchildren():
        if child.tag == _qn('w:p') and child in para_map:
            yield ('para', para_map[child])
        elif child.tag == _qn('w:tbl') and child in tbl_map:
            yield ('table', tbl_map[child])


def _is_top_level_boundary(p):
    """判定某段落是否构成"参考文献条目"豁免区间的终止边界：
    遇到下一个顶级标题/附录/致谢等章节时，参考文献豁免结束，恢复中文校验。"""
    style_name = ''
    try:
        style_name = (p.style.name or '') if p.style is not None else ''
    except Exception:
        style_name = ''
    if style_name.startswith('Heading 1') or style_name in ('标题 1', '标题1'):
        return True
    t = (p.text or '').strip()
    if not t:
        return False
    # 常见另起板块的顶级标题名
    if re.match(r'^(附\s*录|致\s*谢|附\s*件|索\s*引|后\s*记)\b', t):
        return True
    # 编号型一级标题："1 xxx"..."13 xxx"（数字 + 空白 + 非'.'）——与 check_h1_format 保持一致口径
    if re.match(r'^(?:1[0-3]|[1-9])[ \t　]+\S', t) and not re.match(r'^\d+\.', t):
        return True
    return False


def check_chinese_only():
    # 细则：仅“参考文献”标题下方的参考文献条目可以是英文，其余所有文本（含表格）
    # 都必须是中文。“参考文献”标题行本身仍需为中文；标题之前一律不豁免；
    # 若参考文献之后仍出现附录/致谢等其他章节，其中的文本同样按中文校验。
    body_paras = list(doc.paragraphs)
    ref_element = None
    ref_index = None
    for i, p in enumerate(body_paras):
        # 精确定位“参考文献”标题：以"参考文献"开头，避免误将正文里“……参考文献……”句子当成标题。
        if re.match(r'^\s*参考文献', p.text or ''):
            ref_element = p._element
            ref_index = i
            break

    # 计算豁免段落集合：参考文献标题下一段起，直到遇到下一个顶级标题或文档末尾为止。
    exempt_elements = set()
    if ref_element is not None and ref_index is not None:
        for p in body_paras[ref_index + 1:]:
            if _is_top_level_boundary(p):
                break
            exempt_elements.add(p._element)

    for kind, obj in _iter_block_texts():
        if kind == 'para':
            if obj._element in exempt_elements:
                continue  # 参考文献条目：豁免
            # 参考文献标题行本身与其余所有段落均按中文校验
            if _has_untranslated_english(obj.text):
                return False
        else:  # table
            # 表格不属于“参考文献条目”豁免范围，一律按中文校验。
            for row in obj.rows:
                for cell in row.cells:
                    if _has_untranslated_english(cell.text):
                        return False
    return True

def _has_untranslated_english(text):
    """判断一段文本是否含有未翻译的英文正文。
    细则明确要求“皆为中文”，且未在 rubric 中显式允许专有名词，
    因此任何长度≥2 的连续英文字母序列都视为未翻译内容——
    不再对首字母大写的人名/地名做特殊放行。
    数字、标点、空白、单字母（如 A4、N）等不构成“英文词”，仍允许出现。"""
    if not text or not text.strip():
        return False
    return bool(re.search(r'[A-Za-z]{2,}', text))

_CHECKS.append((check_chinese_only, 5, '+5 文档中除"参考文献"下方内容外其余文本皆为中文'))

# ── +1 第1页封面主标题：文本"毕业设计（论文）"字体微软雅黑初号 ──
def check_cover_main_title_font():
    # 细则要点：① 文本为"毕业设计（论文）"；② 字体微软雅黑；③ 字号初号(=42pt)。
    #          （"位于页面上方大约四分之一处"一项已按要求删除。）逐点校验，缺一不可。
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "毕业设计（论文）":
            fn = para_font_name(p)
            sz = para_font_size_pt(p)
            # ② 微软雅黑
            font_ok = bool(fn) and ("雅黑" in fn or "yahei" in fn.lower())
            # ③ 初号 = 42pt
            size_ok = sz is not None and abs(sz - 42) <= 0.6
            if font_ok and size_ok:
                return True
    return False

def _para_vertical_height_emu(para):
    """估算一个段落占用的垂直高度(EMU)，用于粗略定位其在页面中的位置。
    有文字时用字号近似行高；空段落用默认小四(12pt)行高近似。1pt=12700EMU。"""
    sz = para_font_size_pt(para)
    if sz is None:
        sz = 12  # 空行/无显式字号，按默认正文行高估算
    # 行高约为字号的 1.15 倍（Word 单倍行距近似）
    return int(sz * 1.15 * 12700)

_CHECKS.append((check_cover_main_title_font, 1, '+1 第1页封面主标题"毕业设计（论文）"微软雅黑初号'))

# ── +1 第1页封面主标题：文本"毕业设计（论文）"标题前空四行，居中对齐、单倍行距 ──
def check_cover_main_title_format():
    # 细则要点：① 文本为"毕业设计（论文）"；② 标题前空四行；③ 居中对齐；④ 单倍行距。
    # 逐点校验，缺一不可。均针对 Word 段落真实属性判定。
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        # ① 文本精确为"毕业设计（论文）"
        if p.text.strip() != "毕业设计（论文）":
            continue
        # ② 标题前空四行：其前恰好有连续 4 个空段落
        empty_before = 0
        for j in range(i - 1, -1, -1):
            if not paragraphs[j].text.strip():
                empty_before += 1
            else:
                break
        empty_ok = empty_before == 4
        # ③ 居中对齐
        align_ok = is_center_aligned(p)
        # ④ 单倍行距（lineRule=auto 且 line=240）
        spacing_ok = is_single_line_spacing(p)
        if empty_ok and align_ok and spacing_ok:
            return True
    return False

_CHECKS.append((check_cover_main_title_format, 1, '+1 第1页封面主标题"毕业设计（论文）"前空四行、居中、单倍行距'))

# ── +1 第一页"外文翻译"字体等线二号加粗，文本居中上下排列，位于页面上方约四分之一处
#      "毕业设计（论文）"文本下方，单倍行距 ──
def _para_page_number(paragraphs, target_idx):
    """依据 OOXML 静态分页信号推断某段所在页码（1-based），不使用 Word COM/渲染。
    统计从文档开头到该段之前出现的换页信号：
      · <w:br w:type="page"/>            —— 手动分页符
      · <w:lastRenderedPageBreak/>        —— 上次渲染缓存的自动换页位置（Word 保存时写入）
      · 段落级 <w:sectPr>（非 continuous）—— 分节且下一节 nextPage
    另外若目标段自身携带 <w:pageBreakBefore/>，其起始即在新页。
    这是"无渲染无COM"下最贴近真实页码的静态证据。"""
    page = 1
    for i, p in enumerate(paragraphs):
        xml = p._element.xml
        if i == target_idx:
            if '<w:pageBreakBefore' in xml:
                page += 1
            return page
        if re.search(r'<w:br[^/>]*w:type="page"', xml):
            page += 1
        page += xml.count('<w:lastRenderedPageBreak')
        # 段落级 sectPr：默认 nextPage 会换页；显式 continuous 才不换页。
        if re.search(r'<w:sectPr\b', xml) and not re.search(r'<w:type[^/>]*w:val="continuous"', xml):
            page += 1
    return page


def check_waiwenfanyi_format():
    # 细则要点：① 文本为"外文翻译"；② 字体等线；③ 字号二号(=22pt)；④ 加粗；
    #          ⑤ 居中对齐；⑥ 位于"毕业设计（论文）"文本下方，且两标题为上下排列
    #             （均居中；两段之间仅允许空段落，无其他文本插入）；
    #          ⑦ 该段与"毕业设计（论文）"均落在第1页（依据 OOXML 静态分页信号，
    #             不再使用估算段落高度的方式判断视觉位置，避免跨页误判）；
    #          ⑧ 单倍行距。逐点校验，缺一不可。
    paragraphs = doc.paragraphs

    # 先定位"毕业设计（论文）"主标题所在段落索引，用于判断"下方/上下排列/同页"。
    main_title_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "毕业设计（论文）":
            main_title_idx = i
            break
    if main_title_idx is None:
        return False

    for i, p in enumerate(paragraphs):
        if p.text.strip() != "外文翻译":
            continue
        fn = para_font_name(p)
        sz = para_font_size_pt(p)
        # ② 等线（DengXian）
        font_ok = bool(fn) and ("等线" in fn or "dengxian" in fn.lower())
        # ③ 二号 = 22pt
        size_ok = sz is not None and abs(sz - 22) <= 0.6
        # ④ 加粗
        bold_ok = is_bold(p)
        # ⑤ 居中对齐
        align_ok = is_center_aligned(p)
        # ⑧ 单倍行距
        spacing_ok = is_single_line_spacing(p)

        # ⑥ 位于主标题下方 + 上下排列：段序在主标题之后；两段均居中；
        #    二者之间只允许空段落（视觉上呈"上下两行"，不被其他正文/表格穿插）。
        below_ok = i > main_title_idx
        main_center_ok = is_center_aligned(paragraphs[main_title_idx])
        between_clean = all(not paragraphs[k].text.strip() for k in range(main_title_idx + 1, i))
        stacked_ok = below_ok and main_center_ok and between_clean

        # ⑦ 第1页：两段实际页码均为 1（基于 OOXML 静态分页信号推断）。
        main_page = _para_page_number(paragraphs, main_title_idx)
        this_page = _para_page_number(paragraphs, i)
        page1_ok = main_page == 1 and this_page == 1

        if (font_ok and size_ok and bold_ok and align_ok and
                stacked_ok and page1_ok and spacing_ok):
            return True
    return False

_CHECKS.append((check_waiwenfanyi_format, 1, '+1 第1页"外文翻译"等线二号加粗、居中、位于主标题下方并上下排列、单倍行距'))

# ── +1 第1页题目内容为"面向代际学习的社区阅读花园项目设计与评估"，宋体三号，单倍行距，两端对齐 ──
TITLE_CN = "面向代际学习的社区阅读花园项目设计与评估"
def check_cover_topic():
    # 细则要点：① 第1页题目文本为"面向代际学习的社区阅读花园项目设计与评估"；
    #          ② 宋体；③ 三号(=16pt)；④ 单倍行距；⑤ 两端对齐。逐点校验，缺一不可。
    for p in doc.paragraphs[:40]:
        # ① 题目文本精确为细则指定内容
        if p.text.strip() != TITLE_CN:
            continue
        # ② 宋体
        font_ok = has_font_name(p, "宋体", "SimSun", "Song")
        # ③ 三号 = 16pt
        size_ok = is_size_near(p, 16)
        # ④ 单倍行距
        spacing_ok = is_single_line_spacing(p)
        # ⑤ 两端对齐
        align_ok = is_justify_aligned(p)
        if font_ok and size_ok and spacing_ok and align_ok:
            return True
    return False

_CHECKS.append((check_cover_topic, 1, f'+1 第1页题目"{TITLE_CN}"宋体三号、单倍行距、两端对齐'))

# ── +3 第1页封面个人信息区表格：专业、班级、学生姓名、学号、指导教师等字段及其填写内容，
#      表格放置于"外文翻译"文本下方，每一项单独成行，居中，表格7行2列，
#      第一列列宽2.65厘米，第二列列宽7.86厘米，行高统一为2.65厘米 ──
def _tbl_index_in_body(tbl):
    """返回表格在文档正文中的块序号（含段落与表格统一计数），用于判断位置先后。"""
    from docx.oxml.ns import qn as _qn
    idx = 0
    for child in doc.element.body.iterchildren():
        if child is tbl._element:
            return idx
        if child.tag in (_qn('w:p'), _qn('w:tbl')):
            idx += 1
    return None

def _waiwen_block_index():
    """返回"外文翻译"段落在正文中的块序号。"""
    from docx.oxml.ns import qn as _qn
    idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag == _qn('w:p'):
            text = "".join(n.text or '' for n in child.iter(_qn('w:t'))).strip()
            if text == "外文翻译":
                return idx
            idx += 1
        elif child.tag == _qn('w:tbl'):
            idx += 1
    return None

def _row_height_cm(row):
    """返回行高(cm)，无显式行高返回 None。trHeight w:val 单位为 twips(1cm=567twips)。"""
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        return None
    th = trPr.find(qn('w:trHeight'))
    if th is None:
        return None
    val = th.get(qn('w:val'))
    return int(val) / 567 if val else None

def _cell_all_center(tbl):
    """表格所有非空单元格段落是否均居中。"""
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text.strip() and not is_center_aligned(p):
                    return False
    return True

def check_cover_table():
    REQUIRED_FIELDS = ["专业", "班级", "学生姓名", "学号", "指导教师"]
    waiwen_idx = _waiwen_block_index()
    for tbl in get_tables():
        rows = len(tbl.rows)
        cols = len(tbl.columns) if tbl.columns else 0
        # ① 表格 7 行 2 列
        if not (rows == 7 and cols == 2):
            continue
        # ② 表格放置于"外文翻译"文本下方
        tbl_idx = _tbl_index_in_body(tbl)
        below_ok = (waiwen_idx is not None and tbl_idx is not None and tbl_idx > waiwen_idx)
        if not below_ok:
            continue
        # ③ 含专业、班级、学生姓名、学号、指导教师等字段（第一列）
        col1_texts = [tbl.cell(r, 0).text.strip() for r in range(rows)]
        col1_joined = "".join(col1_texts)
        fields_ok = all(f in col1_joined for f in REQUIRED_FIELDS)
        # ④ 每一项单独成行：每个字段各占一行（字段数 ≤ 行数，且各字段分布在不同行）
        rows_with_field = sum(1 for t in col1_texts if any(f in t for f in REQUIRED_FIELDS))
        each_own_row_ok = rows_with_field >= len(REQUIRED_FIELDS)
        # ⑤ 居中
        center_ok = _cell_all_center(tbl)
        # ⑥ 第一列列宽 2.65cm；⑦ 第二列列宽 7.86cm（1cm=360000EMU，容差约±1.4mm）
        try:
            c1_ok = tbl.columns[0].width is not None and abs(tbl.columns[0].width - 954000) < 50000
            c2_ok = tbl.columns[1].width is not None and abs(tbl.columns[1].width - 2829600) < 50000
        except Exception:
            c1_ok = c2_ok = False
        width_ok = c1_ok and c2_ok
        # ⑧ 行高统一为 2.65cm
        heights = [_row_height_cm(r) for r in tbl.rows]
        height_ok = all(h is not None and abs(h - 2.65) < 0.05 for h in heights)

        if (fields_ok and each_own_row_ok and center_ok and
                width_ok and height_ok):
            return True
    return False

_CHECKS.append((check_cover_table, 3, "+3 封面个人信息表格7行2列、位于外文翻译下方、列宽2.65/7.86cm、行高2.65cm、居中"))

# ── +3 第1页封面个人信息区：专业、班级、学生姓名、学号、指导教师等字段为宋体三号加粗分散对齐，
#      其填写内容为宋体三号居中对齐，保持可编辑 ──
def check_cover_table_font():
    # 细则要点：
    #   ① 表格7行2列（个人信息表）；
    #   ② 字段列(第一列)：宋体、三号(=16pt)、加粗、分散对齐；
    #   ③ 内容列(第二列)：宋体、三号(=16pt)、居中对齐；
    #   ④ 保持可编辑（真实 Word 表格文本，非图片/文本框）。
    REQUIRED_FIELDS = ["专业", "班级", "学生姓名", "学号", "指导教师"]
    for tbl in get_tables():
        rows = len(tbl.rows)
        cols = len(tbl.columns) if tbl.columns else 0
        # ① 7行2列
        if not (rows == 7 and cols == 2):
            continue
        col1_texts = [tbl.cell(r, 0).text.strip() for r in range(rows)]
        if not all(f in "".join(col1_texts) for f in REQUIRED_FIELDS):
            continue

        # ② 字段列：宋体三号加粗分散对齐
        field_ok = True
        for r in range(rows):
            cell = tbl.cell(r, 0)
            for p in cell.paragraphs:
                if not p.text.strip():
                    continue
                if not (has_font_name(p, "宋体", "SimSun", "Song")
                        and is_size_near(p, 16)
                        and is_bold(p)
                        and is_distributed_aligned(p)):
                    field_ok = False
        # ③ 内容列：宋体三号居中对齐（仅校验有填写内容的单元格）
        content_ok = True
        content_seen = False
        for r in range(rows):
            cell = tbl.cell(r, 1)
            for p in cell.paragraphs:
                if not p.text.strip():
                    continue
                content_seen = True
                if not (has_font_name(p, "宋体", "SimSun", "Song")
                        and is_size_near(p, 16)
                        and is_center_aligned(p)):
                    content_ok = False
        # ④ 保持可编辑：能从 tbl.cell 读到真实文本即为可编辑的 Word 表格
        editable_ok = any(t for t in col1_texts)

        if field_ok and content_ok and content_seen and editable_ok:
            return True
    return False

_CHECKS.append((check_cover_table_font, 3, "+3 封面信息字段宋体三号加粗分散对齐、内容宋体三号居中对齐"))

# ── +1 院校年份页：宽6-8cm高2-4cm的图片和宋体三号"二〇二六 年"文本，
#      图片放置于"二〇二六 年"文本上方，位置位于封面页底部居中 ──
def _para_index_of_year():
    for i, p in enumerate(doc.paragraphs):
        if re.search(r'二〇二六\s*年|2026\s*年', p.text):
            return i, p
    return None, None

def _para_has_calligraphy_image(para):
    """返回该段落中尺寸符合(宽6-8cm、高2-4cm)的图片段落索引信息。
    命中返回 True。图片以 wp:extent 描述尺寸(EMU, 1cm=360000)。"""
    for ext in para._element.findall('.//' + qn('wp:extent')):
        cx = ext.get('cx')
        cy = ext.get('cy')
        if not cx or not cy:
            continue
        w_cm = int(cx) / 360000
        h_cm = int(cy) / 360000
        if 6 <= w_cm <= 8 and 2 <= h_cm <= 4:
            return True
    return False

def check_page2_school():
    # ① 存在"二〇二六 年"文本，且宋体三号(=16pt)、居中
    year_idx, year_para = _para_index_of_year()
    if year_para is None:
        return False
    year_ok = (has_font_name(year_para, "宋体", "SimSun", "Song")
               and is_size_near(year_para, 16)
               and is_center_aligned(year_para))
    if not year_ok:
        return False

    # ② 存在宽6-8cm、高2-4cm的图片，且图片位于年份文本上方
    img_para_idx = None
    for i, p in enumerate(doc.paragraphs):
        if _para_has_calligraphy_image(p):
            img_para_idx = i
            break
    if img_para_idx is None:
        return False
    # ③ 图片放置于"二〇二六 年"文本上方
    if year_idx is None or not (img_para_idx < year_idx):
        return False

    # ④ 位置位于封面页底部居中：图片所在段落居中，且位于封面页下半部分。
    img_para = doc.paragraphs[img_para_idx]
    img_center_ok = is_center_aligned(img_para)
    section = doc.sections[0]
    page_h = section.page_height
    top_m = section.top_margin or 0
    cursor_emu = top_m
    for i, p in enumerate(doc.paragraphs):
        if i == img_para_idx:
            break
        cursor_emu += _para_vertical_height_emu(p)
    # 底部：图片起始位置应落在页面下半部分（>1/2 页高）
    bottom_ok = (not page_h) or (cursor_emu >= page_h / 2)

    return img_center_ok and bottom_ok

_CHECKS.append((check_page2_school, 1, '+1 院校年份页图片(6-8×2-4cm)在"二〇二六 年"上方、底部居中'))

# ── +1 正文中文题目：正文第一页顶部出现中文译题（不限定具体译名），
#      字体黑体二号，水平居中，段前0.5行，段后0.5行 ──
def check_body_title():
    # 细则要点：① 正文第一页顶部出现中文译题（不限定具体译名，含中文即可）；
    #          ② 黑体；③ 二号(=22pt)；④ 水平居中；⑤ 段前0.5行；⑥ 段后0.5行。
    #          （原细则中"面向代际学习的社区阅读花园项目设计与评估"或同义译名的限定已删除。）
    # "正文第一页顶部"：正文正式内容（摘要/正文章节）出现前的首个中文译题段落。
    paras = doc.paragraphs
    # 先定位正文起点（摘要或第一个编号章节标题），译题应出现在其之前的顶部位置。
    body_start = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("【摘要】") or t.startswith("摘要") or re.match(r'^1\s+\S', t):
            body_start = i
            break

    for i, p in enumerate(paras):
        text = p.text.strip()
        if not text:
            continue
        # ① 中文译题：正文起点之前、含中文字符的标题段落（不限定具体译名）
        if not re.search(r'[一-鿿]', text):
            continue
        # 必须位于正文顶部：在正文起点之前（若能定位到正文起点）
        if body_start is not None and i >= body_start:
            continue
        # ② 黑体
        font_ok = has_font_name(p, "黑体", "HeiTi", "SimHei")
        # ③ 二号 = 22pt
        size_ok = is_size_near(p, 22)
        # ④ 水平居中
        align_ok = is_center_aligned(p)
        # ⑤ 段前0.5行；⑥ 段后0.5行（240twips=1行，0.5行=120twips）
        before, after = para_space_before_after(p)
        before_ok = before is not None and abs(before - 0.5) < 0.05
        after_ok = after is not None and abs(after - 0.5) < 0.05
        if font_ok and size_ok and align_ok and before_ok and after_ok:
            return True
    return False

_CHECKS.append((check_body_title, 1, '+1 正文中文译题黑体二号、居中、段前后各0.5行'))

# ── +1 摘要标题与摘要正文：出现"【摘要】"及完整中文摘要；
#      字体五号，中文宋体、英文和数字 Times New Roman；摘要结束后最后一段后空一行 ──
def _run_ascii_font(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))
    return run.font.name

def _run_eastasia_font(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:eastAsia'))
    return None

def check_abstract():
    # ① 出现"【摘要】"标题
    abs_para = None
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("【摘要】"):
            abs_para = p
            idx = i
            break
    if abs_para is None or idx is None:
        return False

    # 收集摘要正文（从【摘要】段到"关键词"或首个编号章节之前）
    paras = doc.paragraphs
    abstract_paras = [abs_para]
    last_abs_idx = idx  # 摘要末段的段落索引
    for j in range(idx + 1, len(paras)):
        t = paras[j].text.strip()
        if re.match(r'^关键词|^【关键词', t) or re.match(r'^\d+[\s．.]', t):
            break
        if t:
            abstract_paras.append(paras[j])
            last_abs_idx = j

    # ② 字体：中文宋体、英文和数字 Times New Roman、字号五号(=10.5pt)
    #    逐 run 校验：eastAsia=宋体、ascii=Times New Roman、sz=10.5pt。
    #    （"摘要准确覆盖八要素"一项已按要求删除。）
    font_ok = True
    for p in abstract_paras:
        for run in p.runs:
            if not run.text.strip():
                continue
            ea = _run_eastasia_font(run)
            asc = _run_ascii_font(run)
            sz = para_font_size_pt(p) if run.font.size is None else run.font.size.pt
            # 中文宋体
            if ea and not any(n in ea for n in ("宋体", "SimSun", "Song")):
                font_ok = False
            # 英文/数字 Times New Roman
            if asc and "Times New Roman" not in asc:
                font_ok = False
            # 五号 10.5pt
            if sz is None or abs(sz - 10.5) > 0.3:
                font_ok = False
    if not font_ok:
        return False

    # ③ 摘要结束后最后一段后空一行：摘要末段之后紧跟一个空段落。
    blank_after_ok = (last_abs_idx + 1 < len(paras)
                      and not paras[last_abs_idx + 1].text.strip())
    return blank_after_ok

_CHECKS.append((check_abstract, 1, "+1 摘要标题及【摘要】中文摘要（五号宋体/TNR、末段后空一行）"))

# ── +1 关键词区域：摘要下方首个关键词段，仅中文逗号分隔的 5 个关键词；
#      整段五号宋体（英文/数字 Times New Roman）；"关键词"三字加粗；段后空一行 ──
def check_keywords():
    paras = doc.paragraphs

    # 定位摘要标题所在段索引（【摘要】开头）
    abs_idx = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith("【摘要】"):
            abs_idx = i
            break
    if abs_idx is None:
        return False

    # 摘要下方首个"关键词"开头段（仅取一段，不再遍历所有含"关键词"的段）
    kw_idx = None
    kw_para = None
    for j in range(abs_idx + 1, len(paras)):
        t = paras[j].text.strip()
        if t.startswith("关键词") or t.startswith("【关键词"):
            kw_idx = j
            kw_para = paras[j]
            break
    if kw_para is None or kw_idx is None:
        return False

    text = kw_para.text.strip()

    # 分隔符：仅允许中文逗号"，"。若出现英文逗号","或顿号"、"则判失败。
    kw_body = re.sub(r'^.*关键词】?[：:]\s*', '', text)
    if re.search(r'[，,、]', kw_body):
        # 出现任一分隔符时，要求只能是中文逗号
        if re.search(r'[,、]', kw_body):
            return False
    # 关键词个数：中文逗号切分且每项非空、长度≥2
    kws = [k.strip() for k in kw_body.split('，') if k.strip()]
    if len(kws) != 5 or not all(len(k) >= 2 for k in kws):
        return False

    # 内容：期望关键词或"含义准确"——沿用五项匹配的宽松口径
    expected = ["阅读花园", "社区学习", "家庭读写", "户外图书馆", "参与式评估"]
    content_ok = sum(1 for k in expected if any(k in item for item in kws)) >= 5
    if not content_ok:
        return False

    # 逐 run 校验字体/字号（五号=10.5pt）：中文 eastAsia=宋体；英文/数字 ascii=Times New Roman。
    # 同时定位"关键词"三字所在 run 或 run 前缀，校验其加粗。
    font_ok = True
    size_ok = True
    label_bold_ok = False
    label_seen = 0  # 已匹配到的"关键词"字符数（跨 run 拼接）
    LABEL = "关键词"

    for run in kw_para.runs:
        rt = run.text or ""
        if not rt.strip():
            continue
        ea = _run_eastasia_font(run)
        asc = _run_ascii_font(run)
        # 字号：run 优先，回退到段落
        try:
            sz = run.font.size.pt if run.font.size is not None else para_font_size_pt(kw_para)
        except Exception:
            sz = para_font_size_pt(kw_para)
        if ea and not any(n in ea for n in ("宋体", "SimSun", "Song")):
            font_ok = False
        if asc and "Times New Roman" not in asc:
            font_ok = False
        if sz is None or abs(sz - 10.5) > 0.3:
            size_ok = False

        # 判断该 run 是否承载"关键词"三字（可能跨 run，逐字符消费 LABEL）。
        if label_seen < len(LABEL):
            run_covers_label = False
            k = label_seen
            for ch in rt:
                if k < len(LABEL) and ch == LABEL[k]:
                    k += 1
                    run_covers_label = True
                    if k == len(LABEL):
                        break
                else:
                    # 标签期间若混入非期望字符（如空白除外），视为标签已中断——
                    # 关键词标签必须连续出现在段首。
                    if ch.strip():
                        k = 0  # 重置
            if run_covers_label:
                # 若该 run 尚在标签范围内，其加粗属性影响 label_bold_ok。
                run_bold = bool(run.bold)
                if not run_bold:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None and rPr.find(qn('w:b')) is not None:
                        run_bold = True
                # 只有当所有承载标签的 run 都加粗，才算通过；采用"全部加粗"的合取语义。
                if label_seen == 0:
                    label_bold_ok = run_bold
                else:
                    label_bold_ok = label_bold_ok and run_bold
            label_seen = k

    if label_seen < len(LABEL) or not label_bold_ok:
        return False
    if not (font_ok and size_ok):
        return False

    # 段后空一行：紧跟一个空段落
    blank_after_ok = (kw_idx + 1 < len(paras) and not paras[kw_idx + 1].text.strip())
    return blank_after_ok

_CHECKS.append((check_keywords, 1,
                '+1 关键词区域：摘要下方5个中文逗号分隔的关键词、五号宋体、英文/数字TNR、"关键词"三字加粗、段后空一行'))

# ── +5 一级标题格式：正文第1章至第13章的一级标题均使用"三号黑体、左对齐、
#      段前1行、段后1行、单倍行距"，标题编号与标题文字之间空1个字符 ──
def check_h1_format():
    # 一级标题形如 "1 背景与目的" … "13 结论"，编号与文字间空一个字符。
    # 下面的匹配仅用于“识别”哪些段落是第1~13章的一级标题（编号1~13 + 空白 + 中文标题），
    # 不属于细则约束项；细则的 8 个点在识别出目标段落后逐一校验。
    h1_pattern = re.compile(r'^(\d{1,2})(\s+)(\S.*)$')
    h1_paras = []
    for p in doc.paragraphs:
        m = h1_pattern.match(p.text.strip())
        # 编号 1~13 且标题文字首字符为中文（排除 1.1 之类二级标题——其编号含'.'不匹配本正则）
        if m and 1 <= int(m.group(1)) <= 13 and re.match(r'[一-鿿]', m.group(3)):
            h1_paras.append((int(m.group(1)), p, m))
    if not h1_paras:
        return False

    # 【点1】范围：第1章至第13章必须全部出现且各自合规。
    nums = {n for n, _, _ in h1_paras}
    if not set(range(1, 14)).issubset(nums):
        return False

    for _, p, m in h1_paras:
        # 【点2】三号 = 16pt  【点3】黑体
        if not (has_font_name(p, "黑体", "SimHei") and is_size_near(p, 16)):
            return False
        # 【点4】左对齐
        if not is_left_aligned(p):
            return False
        # 【点5】段前1行  【点6】段后1行
        # “行”单位在 Word/WPS 中存于 w:beforeLines/afterLines（单位=1/100行）；
        # 只有真正按“行”设置才算“段前/段后 X 行”。绝对磅值(w:before/after)不算“行”。
        before, after = para_space_before_after(p)
        if not (before is not None and abs(before - 1.0) < 0.05):
            return False
        if not (after is not None and abs(after - 1.0) < 0.05):
            return False
        # 【点7】单倍行距
        if not is_single_line_spacing(p):
            return False
        # 【点8】编号与标题文字之间空1个字符（分隔恰为一个空白字符）
        if len(m.group(2)) != 1:
            return False
    return True

_CHECKS.append((check_h1_format, 5, "+5 一级标题(第1~13章)三号黑体、左对齐、段前后1行、单倍行距、编号后空1字符"))

# ── +5 二级标题格式：使用 1.1、1.2、2.1… 样式，小三号黑体，英文和数字 Times New Roman，
#      两端对齐，段前0.5行、段后0.5行，节号与节标题文字间空1格，单倍行距，勾选"对齐到文档网格" ──
def _para_snap_to_grid_on(para):
    """"对齐到文档网格"对应段落属性 w:snapToGrid。
    未设置时 Word 默认勾选(开)；显式 val=0 表示取消勾选。故只要不是 0 即视为勾选。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return True  # 默认勾选
    snap = pPr.find(qn('w:snapToGrid'))
    if snap is None:
        return True  # 默认勾选
    val = snap.get(qn('w:val'))
    return val not in ('0', 'false', 'off')

def check_h2_format():
    # 二级标题形如 "1.1 邻里背景"，节号与节标题文字间空一个字符。
    h2_pattern = re.compile(r'^(\d+\.\d+)(\s+)(\S.*)$')
    h2_paras = []
    for p in doc.paragraphs:
        m = h2_pattern.match(p.text.strip())
        if m and re.match(r'[一-鿿]', m.group(3)):
            h2_paras.append((p, m))
    if not h2_paras:
        return False

    for p, m in h2_paras:
        # ① 小三号(=15pt) 黑体（中文）
        if not (has_font_name(p, "黑体", "SimHei", "Hei") and is_size_near(p, 15)):
            return False
        # ② 英文和数字 Times New Roman（校验各 run 的 ascii/hAnsi 字体）
        for run in p.runs:
            if not run.text.strip():
                continue
            asc = _run_ascii_font(run)
            if asc and "Times New Roman" not in asc:
                return False
        # ③ 两端对齐
        if not is_justify_aligned(p):
            return False
        # ④ 段前0.5行、⑤ 段后0.5行（240twips=1行，0.5行=120twips）
        before, after = para_space_before_after(p)
        if not (before is not None and abs(before - 0.5) < 0.05):
            return False
        if not (after is not None and abs(after - 0.5) < 0.05):
            return False
        # ⑥ 节号与节标题文字间空1格
        if m.group(2) != ' ':
            return False
        # ⑦ 单倍行距
        if not is_single_line_spacing(p):
            return False
        # ⑧ 勾选"对齐到文档网格"
        if not _para_snap_to_grid_on(p):
            return False
    return True

_CHECKS.append((check_h2_format, 5, "+5 二级标题小三号黑体/TNR、两端对齐、段前后0.5行、节号后空1格、单倍行距、对齐网格"))

# ── +5 正文段落格式：中文正文统一小四号宋体，英文/数字/西文标点小四号 Times New Roman，
#      希腊字母或特殊符号 Symbol 字体；普通段落首行缩进0.85厘米或两个全角字符；
#      单倍行距；勾选"如果定义了文档网格，则对齐到网格"；中文用全角标点，
#      英文缩写/数字/SLAM等术语用半角字符；图号、表号和章节编号保持统一 ──
def _run_symbol_font(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))
    return run.font.name


# 西文（半角）标点集合：这些标点按细则应随英文/数字使用 Times New Roman。
_WESTERN_PUNCT = set(",.;:!?()[]{}\"'/\\-")
# 全角字母/数字（术语应使用半角，出现全角即不合规）。
_FULLWIDTH_ALNUM_RE = re.compile(r'[Ａ-Ｚａ-ｚ０-９]')


def _check_numbering_consistency():
    """图号、表号、章节编号一致性：
      · 图号序列应为 1,2,3… 连续且从 1 起、无重复、无跳号；
      · 表号序列同理；
      · 一级章节编号(形如 "N 标题") 应连续且从 1 起。
    任一序列不一致则返回 False。缺失某类编号（如无表）不算失败。"""
    fig_nums = []
    tbl_nums = []
    sec_nums = []
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r'^图\s*(\d+)', t)
        if m:
            fig_nums.append(int(m.group(1)))
            continue
        m = re.match(r'^表\s*(\d+)', t)
        if m:
            tbl_nums.append(int(m.group(1)))
            continue
        m = re.match(r'^(\d{1,2})[ \t　]+\S', t)
        if m and not re.match(r'^\d+\.', t):
            sec_nums.append(int(m.group(1)))

    def _seq_ok(nums):
        # 去掉重复出现的“同号”（同一图/表在正文引用与图题可能各出现一次）——
        # 取“首次出现顺序”的去重序列，要求它就是 1,2,3,…,k。
        seen = []
        for n in nums:
            if n not in seen:
                seen.append(n)
        if not seen:
            return True  # 该类编号不存在，不做约束
        return seen == list(range(1, len(seen) + 1))

    return _seq_ok(fig_nums) and _seq_ok(tbl_nums) and _seq_ok(sec_nums)


def check_body_para_format():
    # 收集正文普通段落：文字较长、非标题(编号)、非摘要/关键词/图表题、非中文译题；
    # 参考文献标题及其下方内容属于未翻译英文区，不属于"正文普通段落"，排除。
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if re.match(r'^参考文献', p.text.strip()):
            ref_idx = i
            break
    body_paras = []
    for i, p in enumerate(doc.paragraphs):
        if ref_idx is not None and i >= ref_idx:   # 参考文献及以下不算正文
            break
        t = p.text.strip()
        if len(t) < 20:
            continue
        if re.match(r'^\d+(\.\d+)*\s', t):          # 章节标题
            continue
        if t.startswith("【摘要】") or t.startswith("关键词"):
            continue
        if re.match(r'^[图表]\s*\d+', t):            # 图题/表题
            continue
        if TITLE_CN in t or ("代际学习" in t and "阅读花园" in t):  # 中文译题
            continue
        body_paras.append(p)
    if not body_paras:
        return False

    for p in body_paras:
        # ① 中文小四号(=12pt)宋体
        if not (has_font_name(p, "宋体", "SimSun", "Song") and is_size_near(p, 12)):
            return False
        for run in p.runs:
            txt = run.text
            if not txt.strip():
                continue
            # ② 英文/数字 + 西文（半角）标点 使用 Times New Roman
            has_latin_digit = bool(re.search(r'[A-Za-z0-9]', txt))
            has_western_punct = any(ch in _WESTERN_PUNCT for ch in txt)
            if has_latin_digit or has_western_punct:
                asc = _run_ascii_font(run)
                if asc and "Times New Roman" not in asc:
                    return False
            # ③ 希腊字母或特殊符号 Symbol 字体
            if re.search(r'[Α-ω]', txt):  # 希腊字母区
                sym = _run_symbol_font(run)
                if not (sym and "Symbol" in sym):
                    return False
            # ④ 术语半角：英文缩写/数字/SLAM 等不得使用全角字母或全角数字。
            if _FULLWIDTH_ALNUM_RE.search(txt):
                return False
        # ⑤ 首行缩进 0.85 厘米(=482twips) 或 两个全角字符
        indent = get_indent(p)                       # twips
        indent_chars = _get_indent_chars(p)          # 百分之一字符
        indent_ok = (indent is not None and abs(indent - 482) <= 30) or \
                    (indent_chars is not None and indent_chars == 200)
        if not indent_ok:
            return False
        # ⑥ 单倍行距
        if not is_single_line_spacing(p):
            return False
        # ⑦ 勾选"如果定义了文档网格，则对齐到网格"
        if not _para_snap_to_grid_on(p):
            return False
        # ⑧ 中文正文使用中文全角标点：中文之间若出现半角逗号/句号/分号/冒号/
        #    问号/叹号/括号等西文标点则不合规。
        if re.search(r'[一-鿿][,\.;:!?()][一-鿿]', p.text):
            return False

    # ⑨ 图号、表号、章节编号保持统一（文档级序列一致性）。
    if not _check_numbering_consistency():
        return False
    return True

def _get_indent_chars(para):
    """首行缩进的字符数(w:firstLineChars，单位百分之一字符，200=2字符)。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return None
    v = ind.get(qn('w:firstLineChars'))
    return int(v) if v else None

_CHECKS.append((check_body_para_format, 5, "+5 正文小四宋体/英数西文标点TNR/希腊Symbol、术语半角、首行缩进0.85cm、单倍行距、对齐网格、全角标点、图表章节编号统一"))

# ── +5 全文中图题和表题中文字体皆为宋体五号，图号表号数字为 Times New Roman 五号，居中对齐 ──
def check_caption_format():
    # 真正的图题/表题：以"图N"或"表N"开头，且编号后紧跟分隔符或标题文字，
    # 排除正文里"图1概述了…""图4显示了…"这类正文引用句。
    cap_pattern = re.compile(r'^([图表])\s*\d+([\s：:.．、-]|[一-鿿])')
    inline_ref = re.compile(r'^[图表]\s*\d+(概述|显示|表明|说明|给出|展示|把|将|是|中|所示)')
    caps = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if cap_pattern.match(t) and not inline_ref.match(t):
            caps.append(p)
    if not caps:
        return False

    # 细则要求"全文中图题和表题皆为"，故每一个图题/表题都必须合规。
    for p in caps:
        # ① 中文字体宋体、五号(=10.5pt)
        if not (has_font_name(p, "宋体", "SimSun", "Song") and is_size_near(p, 10.5)):
            return False
        # ② 图号表号数字为 Times New Roman、五号
        for run in p.runs:
            txt = run.text
            if not txt.strip():
                continue
            if re.search(r'[0-9A-Za-z]', txt):
                asc = _run_ascii_font(run)
                if asc and "Times New Roman" not in asc:
                    return False
            sz = para_font_size_pt(p) if run.font.size is None else run.font.size.pt
            if sz is None or abs(sz - 10.5) > 0.3:
                return False
        # ③ 居中对齐
        if not is_center_aligned(p):
            return False
    return True

_CHECKS.append((check_caption_format, 5, "+5 图题表题中文宋体五号、图号表号TNR五号、居中对齐"))

# ── 图1～图7 各+1 分（明细校验：位置、尺寸、图题内容、图文同页，见下方
#      check_figureN_detail()）──
#   历史遗留的"图题存在即可"的宽松版检查会与明细版重复登记同一分项，已删除。

def _caption_regex(fig_num):
    # 真正的图题：编号后跟分隔符或中文标题；排除"图N概述了…"等正文引用句。
    return re.compile(rf'^图\s*{fig_num}([\s：:.、．-]|[一-鿿])')

def _inline_ref_regex(fig_num):
    return re.compile(rf'^图\s*{fig_num}(概述|显示|表明|说明|给出|展示|把|将|是|中|所示)')

def _para_image_dims(para):
    """返回段落中首张图片的(宽cm,高cm)，无图片返回 None。"""
    for ext in para._element.findall('.//' + qn('wp:extent')):
        cx, cy = ext.get('cx'), ext.get('cy')
        if cx and cy:
            return int(cx) / 360000, int(cy) / 360000
    return None

# ── +1 图1 阅读花园服务流程图：插入在第1章"背景与目的"相关段落之后、1.1节之前或附近，
#      宽9-11cm高5-8cm；图片下方出现"图1 阅读花园项目的服务逻辑"或语义准确中文图题 ──
def check_figure1_detail():
    paras = doc.paragraphs
    # 定位第1章标题、1.1节标题
    ch1_idx = None
    sec11_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if ch1_idx is None and re.match(r'^1\s+背景', t):
            ch1_idx = i
        if sec11_idx is None and re.match(r'^1\.1\s', t):
            sec11_idx = i
            break
    if ch1_idx is None:
        return False

    # 定位图1图题段落（排除正文引用句）
    cap_re = _caption_regex(1)
    inline_re = _inline_ref_regex(1)
    cap_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if cap_re.match(t) and not inline_re.match(t):
            cap_idx = i
            break
    if cap_idx is None:
        return False

    # 定位图1对应的图片段落（图题上方最近的含图片段落）
    img_idx = None
    img_dims = None
    for i in range(cap_idx - 1, ch1_idx - 1, -1):
        dims = _para_image_dims(paras[i])
        if dims:
            img_idx = i
            img_dims = dims
            break
    if img_idx is None or img_dims is None:
        return False

    # ① 位置：图片在第1章相关段落之后、1.1节之前或附近
    after_ch1 = img_idx > ch1_idx
    before_sec11 = (sec11_idx is None) or (img_idx <= sec11_idx + 1)  # "或附近"给1段容差
    pos_ok = after_ch1 and before_sec11

    # ② 尺寸：宽9-11cm、高5-8cm
    w_cm, h_cm = img_dims
    size_ok = (9 <= w_cm <= 11) and (5 <= h_cm <= 8)

    # ③ 图片下方出现图1中文图题（图题在图片之后）
    cap_below_ok = cap_idx > img_idx

    # ④ 图题内容为"图1 阅读花园项目的服务逻辑"或语义准确中文图题
    cap_text = paras[cap_idx].text.strip()
    cap_content_ok = ("阅读花园" in cap_text and ("服务逻辑" in cap_text or "服务流程" in cap_text)) \
        or ("服务逻辑" in cap_text) or ("服务流程" in cap_text)

    return pos_ok and size_ok and cap_below_ok and cap_content_ok

_CHECKS.append((check_figure1_detail, 1, '+1 图1服务流程图位于第1章内、9-11×5-8cm、下方中文图题'))

# ── +1 图2 临时阅读花园示意图：插入第2章"藏书策略"附近，宽9-11cm高5-8cm；
#      图片下方出现"图2 临时阅读花园的原始矢量示意图"或语义准确中文图题 ──
def check_figure2_detail():
    paras = doc.paragraphs
    # 定位第2章标题、2.2"藏书策略"节、第3章标题（作为第2章结束边界）
    ch2_idx = sec22_idx = ch3_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if ch2_idx is None and re.match(r'^2\s+\S', t):
            ch2_idx = i
        if sec22_idx is None and re.match(r'^2\.2\s', t) and "藏书策略" in t:
            sec22_idx = i
        if ch3_idx is None and re.match(r'^3\s+\S', t):
            ch3_idx = i
    if sec22_idx is None:
        return False

    # 定位图2图题段落（排除正文引用句）
    cap_re = _caption_regex(2)
    inline_re = _inline_ref_regex(2)
    cap_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if cap_re.match(t) and not inline_re.match(t):
            cap_idx = i
            break
    if cap_idx is None:
        return False

    # 定位图2对应图片段（图题上方最近的含图片段落）
    img_idx = img_dims = None
    for i in range(cap_idx - 1, -1, -1):
        dims = _para_image_dims(paras[i])
        if dims:
            img_idx = i
            img_dims = dims
            break
    if img_idx is None or img_dims is None:
        return False

    # ① 位置：插入第2章"藏书策略"附近（在2.2节之后、第3章之前）
    upper = ch3_idx if ch3_idx is not None else len(paras)
    pos_ok = (img_idx > sec22_idx) and (img_idx < upper)

    # ② 尺寸：宽9-11cm、高5-8cm
    w_cm, h_cm = img_dims
    size_ok = (9 <= w_cm <= 11) and (5 <= h_cm <= 8)

    # ③ 图片下方出现图2中文图题
    cap_below_ok = cap_idx > img_idx

    # ④ 图题内容为"图2 临时阅读花园的原始矢量示意图"或语义准确中文图题
    cap_text = paras[cap_idx].text.strip()
    cap_content_ok = ("阅读花园" in cap_text and ("示意图" in cap_text or "矢量" in cap_text)) \
        or ("矢量示意图" in cap_text)

    return pos_ok and size_ok and cap_below_ok and cap_content_ok

_CHECKS.append((check_figure2_detail, 1, '+1 图2示意图位于第2章藏书策略附近、9-11×5-8cm、下方中文图题'))

# ── +1 图3 观察分区与移动路径图：插入第3章"观察方案"或"可靠性检查"附近，宽9-11cm高5-8cm；
#      图片下方出现"图3 工作人员在项目中使用的观察区域与参与者移动路径。"或语义准确中文图题 ──
def check_figure3_detail():
    paras = doc.paragraphs
    # 定位第3章内的"观察"节(3.1 观察规程/观察方案)与"可靠性检查"节，及第4章边界
    obs_idx = rel_idx = ch4_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if obs_idx is None and re.match(r'^3\.\d+\s', t) and "观察" in t:
            obs_idx = i
        if rel_idx is None and re.match(r'^3\.\d+\s', t) and "可靠性" in t:
            rel_idx = i
        if ch4_idx is None and re.match(r'^4\s+\S', t):
            ch4_idx = i
    anchor = obs_idx if obs_idx is not None else rel_idx
    if anchor is None:
        return False

    # 定位图3图题段落（排除正文引用句）
    cap_re = _caption_regex(3)
    inline_re = _inline_ref_regex(3)
    cap_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if cap_re.match(t) and not inline_re.match(t):
            cap_idx = i
            break
    if cap_idx is None:
        return False

    # 定位图3对应图片段（图题上方最近的含图片段落）
    img_idx = img_dims = None
    for i in range(cap_idx - 1, -1, -1):
        dims = _para_image_dims(paras[i])
        if dims:
            img_idx = i
            img_dims = dims
            break
    if img_idx is None or img_dims is None:
        return False

    # ① 位置：插入第3章观察/可靠性检查附近（在观察节之后、第4章之前）
    upper = ch4_idx if ch4_idx is not None else len(paras)
    pos_ok = (img_idx > anchor) and (img_idx < upper)

    # ② 尺寸：宽9-11cm、高5-8cm
    w_cm, h_cm = img_dims
    size_ok = (9 <= w_cm <= 11) and (5 <= h_cm <= 8)

    # ③ 图片下方出现图3中文图题
    cap_below_ok = cap_idx > img_idx

    # ④ 图题内容为"图3 工作人员在项目中使用的观察区域与参与者移动路径"或语义准确中文图题
    cap_text = paras[cap_idx].text.strip()
    cap_content_ok = ("观察区域" in cap_text and "移动路径" in cap_text) \
        or ("观察" in cap_text and "路径" in cap_text)

    return pos_ok and size_ok and cap_below_ok and cap_content_ok

_CHECKS.append((check_figure3_detail, 1, '+1 图3观察分区路径图位于第3章观察/可靠性检查附近、9-11×5-8cm、下方中文图题'))

# ── 图N 通用明细校验：位置(章节区间内)、尺寸、图题在图下、图题语义 ──
def _check_figure_detail(fig_num, lower_idx, upper_idx, caption_ok_fn,
                         w_range=(9.0, 11.0), h_range=(5.0, 8.0)):
    """lower_idx: 图片必须位于该段落索引之后（章节锚点）；
    upper_idx:  图片必须位于该段落索引之前（下一章边界，None 表示文末）；
    caption_ok_fn: 传入图题文本，返回图题语义是否准确；
    w_range/h_range: 图片宽/高允许范围(cm)。"""
    paras = doc.paragraphs
    if lower_idx is None:
        return False
    cap_re = _caption_regex(fig_num)
    inline_re = _inline_ref_regex(fig_num)
    cap_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if cap_re.match(t) and not inline_re.match(t):
            cap_idx = i
            break
    if cap_idx is None:
        return False
    img_idx = img_dims = None
    for i in range(cap_idx - 1, -1, -1):
        dims = _para_image_dims(paras[i])
        if dims:
            img_idx, img_dims = i, dims
            break
    if img_idx is None or img_dims is None:
        return False
    upper = upper_idx if upper_idx is not None else len(paras)
    pos_ok = (img_idx > lower_idx) and (img_idx < upper)
    w_cm, h_cm = img_dims
    size_ok = (w_range[0] <= w_cm <= w_range[1]) and (h_range[0] <= h_cm <= h_range[1])
    cap_below_ok = cap_idx > img_idx
    cap_content_ok = caption_ok_fn(paras[cap_idx].text.strip())
    return pos_ok and size_ok and cap_below_ok and cap_content_ok

def _chapter_idx(num):
    for i, p in enumerate(doc.paragraphs):
        if re.match(rf'^{num}\s+\S', p.text.strip()):
            return i
    return None

# ── +1 图4 十二周参与趋势图：插入第4章参与和使用模式附近，宽9-11cm高5-8cm；
#      图片下方出现"图4 十二周周期内的参与记录，综合青年访问、成人访问和回访家庭百分比。"
#      或语义准确中文图题 ──
def check_figure4_detail():
    ch4, ch5 = _chapter_idx(4), _chapter_idx(5)
    def cap_ok(t):
        return ("参与记录" in t or "参与" in t) and \
               (("访问" in t) or ("回访" in t) or ("百分比" in t))
    return _check_figure_detail(4, ch4, ch5, cap_ok)

_CHECKS.append((check_figure4_detail, 1, '+1 图4参与趋势图位于第4章附近、9-11×5-8cm、下方中文图题'))

# ── +1 图5 多指标结果图：插入第6章运行发现与第7章学习成果之间或对应分析段落附近，
#      宽14-16cm高10-12cm；图片下方出现"图5 用于指导每周调整的多指标结果。"或语义准确
#      中文图题 ──
def check_figure5_detail():
    ch6, ch7 = _chapter_idx(6), _chapter_idx(7)
    def cap_ok(t):
        return ("多指标" in t and "结果" in t) or \
               ("结果" in t and ("每周调整" in t or "指标" in t))
    return _check_figure_detail(5, ch6, ch7, cap_ok,
                                w_range=(14.0, 16.0), h_range=(10.0, 12.0))

_CHECKS.append((check_figure5_detail, 1, '+1 图5多指标结果图位于第6~7章之间、14-16×10-12cm、下方中文图题'))

# ── +1 图6 便携式资源工具包图：插入第7章家庭参与相关内容附近，宽9-11cm高5-8cm；
#      图片下方出现"图6 用于组装项目的便携式工具包，展示可移动资源与活动流程之间的关系"
#      或语义准确中文图题 ──
def check_figure6_detail():
    ch7, ch8 = _chapter_idx(7), _chapter_idx(8)
    def cap_ok(t):
        return ("便携式工具包" in t or ("工具包" in t and "便携" in t)) or \
               ("工具包" in t and ("可移动资源" in t or "资源" in t))
    return _check_figure_detail(6, ch7, ch8, cap_ok)

_CHECKS.append((check_figure6_detail, 1, '+1 图6工具包图位于第7章附近、9-11×5-8cm、下方中文图题'))

# ── +1 图7 十二周实施矩阵图：插入第9章"实施矩阵"附近，宽9-11cm高5-8cm；
#      图片下方出现"图7 供图书馆或学校改编该项目时使用的实施矩阵"或语义准确中文图题 ──
def check_figure7_detail():
    ch9, ch10 = _chapter_idx(9), _chapter_idx(10)
    def cap_ok(t):
        return "实施矩阵" in t or ("矩阵" in t and ("实施" in t or "改编" in t or "项目" in t))
    return _check_figure_detail(7, ch9, ch10, cap_ok)

_CHECKS.append((check_figure7_detail, 1, '+1 图7实施矩阵图位于第9章附近、9-11×5-8cm、下方中文图题'))

# ── 表1～表5 各+3（每张表按各自细则独立精确校验）──
# ── +3 表1 项目组成表：在第2章对应位置建立可编辑 Word 表格，包含"组成部分、主要功能、
#      证据来源"等列，并翻译移动书箱、遮阳棚、安静座垫、提示卡、录音箱和花园标签等内容 ──
def _table_index_in_body(tbl):
    idx = 0
    for child in doc.element.body.iterchildren():
        if child is tbl._element:
            return idx
        if child.tag in (qn('w:p'), qn('w:tbl')):
            idx += 1
    return None

def check_table1():
    # 第2章、第3章位置（正文块序号），用于判定"在第2章对应位置"。
    ch2_blk = ch3_blk = None
    blk = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            t = "".join(n.text or '' for n in child.iter(qn('w:t'))).strip()
            if ch2_blk is None and re.match(r'^2\s+\S', t):
                ch2_blk = blk
            if ch3_blk is None and re.match(r'^3\s+\S', t):
                ch3_blk = blk
            blk += 1
        elif child.tag == qn('w:tbl'):
            blk += 1

    # 列名同义组：组成部分/组件、主要功能、证据来源/证据
    col_groups = [
        ["组成部分", "组件", "构成", "组成"],
        ["主要功能", "功能"],
        ["证据来源", "证据", "来源"],
    ]
    # 需翻译到的六项组件内容（接受准确同义译名）
    item_groups = [
        ["移动书箱", "书箱", "移动书"],
        ["遮阳棚", "遮阴天篷", "天篷", "遮阳", "遮阴"],
        ["安静座垫", "安静坐垫", "座垫", "坐垫"],
        ["提示卡", "提示卡组"],
        ["录音箱", "录音盒", "录音"],
        ["花园标签", "标签"],
    ]

    for tbl in get_tables():
        cols = len(tbl.columns) if tbl.columns else 0
        rows = len(tbl.rows)
        if cols < 3 or rows < 2:
            continue
        # ① 表头列名（第一行）含"组成部分/主要功能/证据来源"等列
        header = "|".join(c.text.strip() for c in tbl.rows[0].cells)
        if not all(any(k in header for k in grp) for grp in col_groups):
            continue
        # ② 表体内容翻译到六项组件（每项各命中一个同义译名）
        body_text = "\n".join(
            tbl.cell(r, c).text
            for r in range(rows) for c in range(cols)
        )
        if not all(any(k in body_text for k in grp) for grp in item_groups):
            continue
        # ③ 位于第2章对应位置（在第2章之后、第3章之前）
        tbl_blk = _table_index_in_body(tbl)
        if tbl_blk is None or ch2_blk is None:
            continue
        upper = ch3_blk if ch3_blk is not None else 10 ** 9
        if not (ch2_blk < tbl_blk < upper):
            continue
        # ④ 可编辑 Word 表格：能读到真实单元格文本即为可编辑表格
        if body_text.strip():
            return True
    return False

_CHECKS.append((check_table1, 3, "+3 表1项目组成表在第2章、列含组成部分/主要功能/证据来源、六项组件译全"))

# ── +3 表2 十二周活动周期表：在第4章对应位置建立可编辑表格，包含周次、主题、青少年访问量、
#      成人访问量、回访率、员工分钟数和调整事项，共保留12周数据 ──
def check_table2():
    # 第4章、第5章位置（正文块序号）
    ch4_blk = ch5_blk = None
    blk = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            t = "".join(n.text or '' for n in child.iter(qn('w:t'))).strip()
            if ch4_blk is None and re.match(r'^4\s+\S', t):
                ch4_blk = blk
            if ch5_blk is None and re.match(r'^5\s+\S', t):
                ch5_blk = blk
            blk += 1
        elif child.tag == qn('w:tbl'):
            blk += 1

    # 七个列的同义组：周次、主题、青少年访问量、成人访问量、回访率、员工分钟数、调整事项
    col_groups = [
        ["周次", "周"],
        ["主题"],
        ["青少年访问", "青年访问", "青少年", "青年"],
        ["成人访问", "成人"],
        ["回访率", "回访"],
        ["员工分钟", "工作人员分钟", "分钟数", "分钟"],
        ["调整事项", "调整", "记录的调整"],
    ]

    for tbl in get_tables():
        cols = len(tbl.columns) if tbl.columns else 0
        rows = len(tbl.rows)
        if cols < 7 or rows < 2:
            continue
        # ① 表头含七个列
        header = "|".join(c.text.strip() for c in tbl.rows[0].cells)
        if not all(any(k in header for k in grp) for grp in col_groups):
            continue
        # ② 共保留12周数据：表体第一列出现周次 1~12。
        week_nums = set()
        for r in range(1, rows):
            first = tbl.cell(r, 0).text.strip()
            m = re.match(r'^(\d{1,2})', first)
            if m:
                week_nums.add(int(m.group(1)))
        if not set(range(1, 13)).issubset(week_nums):
            continue
        # ③ 位于第4章对应位置（第4章之后、第5章之前）
        tbl_blk = _table_index_in_body(tbl)
        if tbl_blk is None or ch4_blk is None:
            continue
        upper = ch5_blk if ch5_blk is not None else 10 ** 9
        if not (ch4_blk < tbl_blk < upper):
            continue
        # ④ 可编辑 Word 表格
        return True
    return False

_CHECKS.append((check_table2, 3, "+3 表2十二周活动周期表在第4章、含七列、保留12周数据"))

# ── +3 表3 平均工作量表：在第6章对应位置建立可编辑表格，包含场地布置、迎接、材料维护、
#      观察记录、收尾整理和合计等行，以及第1—3周、第4—8周、第9—12周等列 ──
def check_table3():
    # 第6章、第7章位置（正文块序号）
    ch6_blk = ch7_blk = None
    blk = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            t = "".join(n.text or '' for n in child.iter(qn('w:t'))).strip()
            if ch6_blk is None and re.match(r'^6\s+\S', t):
                ch6_blk = blk
            if ch7_blk is None and re.match(r'^7\s+\S', t):
                ch7_blk = blk
            blk += 1
        elif child.tag == qn('w:tbl'):
            blk += 1

    # 六个行的同义组：场地布置、迎接、材料维护、观察记录、收尾整理、合计
    row_groups = [
        ["场地布置", "布置"],
        ["迎接", "欢迎接待", "接待", "迎接接待"],
        ["材料维护", "材料照料", "材料"],
        ["观察记录", "观察"],
        ["收尾整理", "结束复位", "收尾", "结束"],
        ["合计", "总计"],
    ]
    # 三个阶段列（周区间，接受 - 或 — 连字符）
    col_groups = [
        [re.compile(r'第?\s*1\s*[-—－~至]\s*3\s*周')],
        [re.compile(r'第?\s*4\s*[-—－~至]\s*8\s*周')],
        [re.compile(r'第?\s*9\s*[-—－~至]\s*12\s*周')],
    ]

    for tbl in get_tables():
        cols = len(tbl.columns) if tbl.columns else 0
        rows = len(tbl.rows)
        if cols < 4 or rows < 2:
            continue
        # ① 列：第1-3周 / 第4-8周 / 第9-12周（校验表头行）
        header = "|".join(c.text.strip() for c in tbl.rows[0].cells)
        if not all(any(rx.search(header) for rx in grp) for grp in col_groups):
            continue
        # ② 行：六个行名（校验第一列各行）
        first_col = "\n".join(tbl.cell(r, 0).text.strip() for r in range(rows))
        if not all(any(k in first_col for k in grp) for grp in row_groups):
            continue
        # ③ 位于第6章对应位置（第6章之后、第7章之前）
        tbl_blk = _table_index_in_body(tbl)
        if tbl_blk is None or ch6_blk is None:
            continue
        upper = ch7_blk if ch7_blk is not None else 10 ** 9
        if not (ch6_blk < tbl_blk < upper):
            continue
        # ④ 可编辑 Word 表格
        return True
    return False

_CHECKS.append((check_table3, 3, "+3 表3平均工作量表在第6章、含六行、含三阶段周区间列"))

# ── +3 表4 设计权衡表：在第8章对应位置建立可编辑表格，包含权衡事项、忽视风险和运营应对等列，
#      表格文字已翻译为中文 ──
def check_table4():
    # 第8章、第9章位置（正文块序号）
    ch8_blk = ch9_blk = None
    blk = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            t = "".join(n.text or '' for n in child.iter(qn('w:t'))).strip()
            if ch8_blk is None and re.match(r'^8\s+\S', t):
                ch8_blk = blk
            if ch9_blk is None and re.match(r'^9\s+\S', t):
                ch9_blk = blk
            blk += 1
        elif child.tag == qn('w:tbl'):
            blk += 1

    # 三个列的同义组：权衡事项、忽视风险、运营应对
    col_groups = [
        ["权衡事项", "权衡"],
        ["忽视风险", "风险", "若忽视"],
        ["运营应对", "运行回应", "应对", "回应", "响应"],
    ]

    for tbl in get_tables():
        cols = len(tbl.columns) if tbl.columns else 0
        rows = len(tbl.rows)
        if cols < 3 or rows < 2:
            continue
        # ① 表头含权衡事项/忽视风险/运营应对三列
        header = "|".join(c.text.strip() for c in tbl.rows[0].cells)
        if not all(any(k in header for k in grp) for grp in col_groups):
            continue
        # ② 表格文字已翻译为中文：全表单元格不得出现未翻译英文单词。
        body_text = "\n".join(
            tbl.cell(r, c).text
            for r in range(rows) for c in range(cols)
        )
        if re.search(r'[A-Za-z]{2,}', body_text):
            continue
        # 需含中文
        if not re.search(r'[一-鿿]', body_text):
            continue
        # ③ 位于第8章对应位置（第8章之后、第9章之前）
        tbl_blk = _table_index_in_body(tbl)
        if tbl_blk is None or ch8_blk is None:
            continue
        upper = ch9_blk if ch9_blk is not None else 10 ** 9
        if not (ch8_blk < tbl_blk < upper):
            continue
        # ④ 可编辑 Word 表格
        return True
    return False

_CHECKS.append((check_table4, 3, "+3 表4设计权衡表在第8章、含权衡/忽视风险/运营应对列、全中文"))

# ── +3 表5 迁移指南表：在第12章末尾或第13章之前建立可编辑表格，包含规划领域、最低行动、
#      试点后的改进行动和复核点等列 ──
def check_table5():
    # 第12章、第13章位置（正文块序号）
    ch12_blk = ch13_blk = None
    blk = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            t = "".join(n.text or '' for n in child.iter(qn('w:t'))).strip()
            if ch12_blk is None and re.match(r'^12\s+\S', t):
                ch12_blk = blk
            if ch13_blk is None and re.match(r'^13\s+\S', t):
                ch13_blk = blk
            blk += 1
        elif child.tag == qn('w:tbl'):
            blk += 1

    # 四个列的同义组：规划领域、最低行动、试点后的改进行动、复核点
    col_groups = [
        ["规划领域", "规划"],
        ["最低行动", "最低"],
        ["试点后的改进行动", "改进行动", "改进"],
        ["复核点", "复盘点", "复核", "复盘"],
    ]

    for tbl in get_tables():
        cols = len(tbl.columns) if tbl.columns else 0
        rows = len(tbl.rows)
        if cols < 4 or rows < 2:
            continue
        # ① 表头含四列
        header = "|".join(c.text.strip() for c in tbl.rows[0].cells)
        if not all(any(k in header for k in grp) for grp in col_groups):
            continue
        # ② 位于第12章末尾或第13章之前（第12章之后、第13章之前）
        tbl_blk = _table_index_in_body(tbl)
        if tbl_blk is None or ch12_blk is None:
            continue
        upper = ch13_blk if ch13_blk is not None else 10 ** 9
        if not (ch12_blk < tbl_blk < upper):
            continue
        # ③ 可编辑 Word 表格
        return True
    return False

_CHECKS.append((check_table5, 3, "+3 表5迁移指南表在第12章末/第13章前、含规划领域/最低行动/改进行动/复核点列"))

# ── +3 图表分页控制：每幅图片与其图题位于同一页；
#      每张表格的表题、表头和首行位于同一页（不得跨页分布） ──
#   该项需要真实渲染的页码信息（尤其"表头与首行是否被自动分页拆开"），
#   纯 OOXML 只能识别显式分页符/分节符，无法覆盖 Word 自动分页；
#   在不启动 Word COM 的前提下无等价静态实现，普通模式下该得分点固定不命中。

def check_pagination():
    # 无 COM/渲染，普通模式下该得分点不命中。
    return False

_CHECKS.append((check_pagination, 3, "+3 图片与图题同页、表格表题/表头/首行同页（无渲染，默认不命中）"))

# ── +1 奇数页页眉"外文翻译" ──
# ── +1 正文页眉奇数页：正文奇数页页眉显示"外文翻译"，字体宋体，字号小五，位置位于页眉居中 ──
def _header_para_format_ok(p, expected_text):
    # ① 页眉文本为指定内容（忽略空白差异）
    text = re.sub(r'\s+', '', p.text)
    expected = re.sub(r'\s+', '', expected_text)
    if text != expected:
        return False
    # ② 字体宋体
    font_ok = has_font_name(p, "宋体", "SimSun", "Song")
    # ③ 字号小五(=9pt)
    size_ok = is_size_near(p, 9)
    # ④ 页眉居中
    align_ok = is_center_aligned(p)
    return font_ok and size_ok and align_ok

def _even_and_odd_headers_enabled():
    """检测 word/settings.xml 是否启用了"奇偶页不同"（w:evenAndOddHeaders）。
    只有该开关启用时，"奇数页页眉"与"偶数页页眉"才在渲染层面区分开来。"""
    try:
        settings_el = doc.settings.element  # type: ignore[attr-defined]
    except Exception:
        return False
    if settings_el is None:
        return False
    node = settings_el.find(qn('w:evenAndOddHeaders'))
    if node is None:
        return False
    val = node.get(qn('w:val'))
    # 缺省 val 视为启用；显式 "0"/"false" 视为未启用。
    return val not in ("0", "false")


def check_header_odd():
    # 细则：正文奇数页页眉显示"外文翻译"，宋体小五、居中。
    # 判定口径：
    #   ① 定位"正文节"——沿用本脚本 check_page_number 的约定：doc.sections[-1] 为正文节，
    #      doc.sections[0] 为封面/年份页等前置节。
    #   ② 必须实际启用"奇偶页不同"（settings.xml 的 w:evenAndOddHeaders）——
    #      否则"奇数页页眉"与"偶数页页眉"无法在渲染上分开，细则的"奇数页/偶数页"要求即无法成立。
    #   ③ 仅校验正文节的默认(奇数)页眉，不再扫描全部 sections 的 header 任一命中即通过。
    if len(doc.sections) < 1:
        return False
    if not _even_and_odd_headers_enabled():
        return False
    body_section = doc.sections[-1] if len(doc.sections) >= 2 else doc.sections[0]
    # section.header 即"奇数页/默认页眉"。若正文节设置了 linked_to_previous，
    # 实际页眉承袭自上一节，此时 python-docx 仍会解析到承袭后的 part，其段落有效。
    odd_header = body_section.header
    for p in odd_header.paragraphs:
        if _header_para_format_ok(p, "外文翻译"):
            return True
    return False

_CHECKS.append((check_header_odd, 1, "+1 正文节奇数页页眉显示'外文翻译'、宋体小五、居中（已启用奇偶页不同）"))

# ── +1 正文页眉偶数页：正文偶数页页眉显示"2022届材料科学与工程专业毕业设计（论文）"，
#      字体宋体，字号小五，位置位于页眉居中 ──
def check_header_even():
    # 细则：正文偶数页页眉显示"2022届材料科学与工程专业毕业设计（论文）"，宋体小五、居中。
    # 判定口径（与 check_header_odd 一致）：
    #   ① 定位"正文节"——doc.sections[-1] 为正文节，doc.sections[0] 为封面/年份页等前置节；
    #   ② 必须实际启用"奇偶页不同"（settings.xml 的 w:evenAndOddHeaders），否则偶数页页眉无意义；
    #   ③ 仅校验正文节的 even_page_header，不再扫描全部 sections 的 even_page_header 任一命中即通过。
    expected = "2022届材料科学与工程专业毕业设计（论文）"
    if len(doc.sections) < 1:
        return False
    if not _even_and_odd_headers_enabled():
        return False
    body_section = doc.sections[-1] if len(doc.sections) >= 2 else doc.sections[0]
    even_header = body_section.even_page_header
    for p in even_header.paragraphs:
        if _header_para_format_ok(p, expected):
            return True
    return False

_CHECKS.append((check_header_even, 1, "+1 正文节偶数页页眉显示'2022届材料科学与工程专业毕业设计（论文）'、宋体小五、居中（已启用奇偶页不同）"))

# ── +1 正文页码：正文从翻译正文摘要第一页开始使用阿拉伯数字1连续编号，页码位于页脚居中位置，
#      封面和院校年份页不显示正文页码 ──
def _xml_has_page_field(xml_text):
    return '<w:fldChar' in xml_text and re.search(r'<w:instrText[^>]*>\s*PAGE\s*</w:instrText>', xml_text)

def _footer_page_centered(footer):
    """页脚是否含 PAGE 域且该页码段落居中。"""
    xml = footer._element.xml
    if not _xml_has_page_field(xml):
        return False
    for p in footer.paragraphs:
        if _xml_has_page_field(p._element.xml) and is_center_aligned(p):
            return True
    return False

def check_page_number():
    if len(doc.sections) < 2:
        return False
    # 约定：section0 = 封面 + 院校年份页（前置节）；末节 = 正文（摘要起）。
    cover_section = doc.sections[0]
    body_section = doc.sections[-1]

    # ① 封面和院校年份页不显示正文页码：
    #    封面为该节首页(first_page_footer)，院校年份页为该节默认/偶数页脚。
    cover_footers = [
        cover_section.first_page_footer,   # 封面
        cover_section.footer,              # 院校年份页
        cover_section.even_page_footer,
    ]
    if any(_xml_has_page_field(f._element.xml) for f in cover_footers):
        return False

    # ② 正文页脚含 PAGE 域且页码居中（首页/默认/偶数页脚任一带居中页码即可）。
    body_footers = [
        body_section.footer,
        body_section.first_page_footer,
        body_section.even_page_footer,
    ]
    has_body_page_centered = any(_footer_page_centered(f) for f in body_footers)
    if not has_body_page_centered:
        return False

    # ③ 正文从第1页开始用阿拉伯数字连续编号：正文节设置 pgNumType start=1，
    #    且编号格式为阿拉伯数字(decimal，即未指定非阿拉伯 w:fmt)。
    sectPr = body_section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        return False
    if pgNumType.get(qn('w:start')) != "1":
        return False
    fmt = pgNumType.get(qn('w:fmt'))
    # 未指定 fmt 时 Word 默认阿拉伯数字；指定时必须为 decimal。
    if fmt is not None and fmt != "decimal":
        return False
    return True

_CHECKS.append((check_page_number, 1, "+1 正文自摘要首页起阿拉伯数字连续编号、页脚居中、封面与年份页不显示页码"))

# ── +1 页面设置：正文页面为 A4 纵向，页面边距上下左右皆为2.5cm、页眉距离1.5cm、页脚距离1.5cm ──
def check_page_setup():
    # 细则限定"正文页面"：仅校验正文节（doc.sections[-1]，与页眉/页码判定约定一致），
    # 不再遍历全部 sections 任一命中即通过。（1cm=360000EMU，容差±50000≈1.4mm）
    from docx.enum.section import WD_ORIENT
    if len(doc.sections) < 1:
        return False
    section = doc.sections[-1] if len(doc.sections) >= 2 else doc.sections[0]

    # ① A4 尺寸：宽21cm(=7560000)、高29.7cm(=10692000)
    pw, ph = section.page_width, section.page_height
    a4_ok = (pw is not None and abs(pw - 7560000) < 50000) and \
            (ph is not None and abs(ph - 10692000) < 50000)
    # ② 纵向
    orient_ok = section.orientation == WD_ORIENT.PORTRAIT
    # ③ 页边距上下左右皆为 2.5cm(=900000)
    margins = (section.top_margin, section.bottom_margin,
               section.left_margin, section.right_margin)
    margin_ok = all(m is not None and abs(m - 900000) < 50000 for m in margins)
    # ④ 页眉距离 1.5cm(=540000)；⑤ 页脚距离 1.5cm(=540000)
    hd, fd = section.header_distance, section.footer_distance
    header_ok = hd is not None and abs(hd - 540000) < 50000
    footer_ok = fd is not None and abs(fd - 540000) < 50000

    return a4_ok and orient_ok and margin_ok and header_ok and footer_ok

_CHECKS.append((check_page_setup, 1, "+1 正文页面A4纵向、页边距上下左右2.5cm、页眉距离1.5cm、页脚距离1.5cm"))

# ── +1 参考文献部分：正文末尾出现"参考文献（省略不翻）"或"参考文献"标题，
#      标题格式为一级标题格式，参考文献英文内容为英文 ──
def check_references():
    idx = None
    ref_para = None
    for i, p in enumerate(doc.paragraphs):
        # ① 出现"参考文献（省略不翻）"或"参考文献"标题
        t = p.text.strip()
        if t == "参考文献" or t.startswith("参考文献"):
            idx = i
            ref_para = p
            break
    if idx is None or ref_para is None:
        return False
    # ② 标题为一级标题格式：三号黑体、左对齐、段前1行、段后1行、单倍行距
    #    （与 check_h1_format 的一级标题要求一致）
    if not (has_font_name(ref_para, "黑体", "SimHei", "Hei") and is_size_near(ref_para, 16)):
        return False
    if not is_left_aligned(ref_para):
        return False
    before, after = para_space_before_after(ref_para)
    if not (before is not None and abs(before - 1.0) < 0.05):
        return False
    if not (after is not None and abs(after - 1.0) < 0.05):
        return False
    if not is_single_line_spacing(ref_para):
        return False
    # ③ 参考文献英文内容为英文：标题下方条目应含英文文本。
    after_text = "\n".join(p.text for p in doc.paragraphs[idx + 1: idx + 6])
    return bool(re.search(r'[A-Za-z]{3,}', after_text))

_CHECKS.append((check_references, 1, "+1 参考文献标题（一级标题格式）存在且下方内容为英文"))


# ─────────────────────────────────────────
# 统一对外接口
# ─────────────────────────────────────────
def _locate_docx(dir_path: str) -> Path:
    """在指定目录中定位被评估文档。
    优先返回与 DOC_FILENAME 同名的 .docx；否则回退到目录内首个非临时 .docx。"""
    d = Path(dir_path)
    target = d / DOC_FILENAME
    if target.exists():
        return target
    for f in sorted(d.glob("*.docx")):
        if f.name.startswith("~$"):
            continue
        return f
    return target  # 返回预期路径，让调用方拿到"文件不存在"错误


def evaluate(dir_path: str) -> dict:
    """脚本对外唯一入口。dir_path 为脚本所在目录，脚本自行定位并打开被评估文档。
    返回字典结构见《脚本接口差异与统一建议.md》§2.2。"""
    global doc, _doc_path

    result = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": sum(score for _, score, _ in _CHECKS),
    }

    try:
        file_path = _locate_docx(dir_path)
        result["file_name"] = file_path.name
        _doc_path = file_path.resolve()

        # 1a. 文件为 .docx 且可正常打开
        if not file_path.exists():
            result["status"] = "error"
            result["error"] = f"文件不存在: {file_path}"
            result["file_name"] = ""
            return result
        if file_path.suffix.lower() != ".docx":
            result["status"] = "error"
            result["error"] = f"文件不是 .docx 格式: {file_path.name}"
            return result
        try:
            doc = Document(str(file_path))
        except Exception as e:
            result["status"] = "error"
            result["error"] = f"文件无法打开: {e}"
            return result

        # 1b / 1c / 1d
        ok, reason = _dim1_check_document()
        if not ok:
            result["dim1_reason"] = reason
            return result

        result["dim1_pass"] = True

        # 维度二：逐项执行注册表中的检查
        total = 0
        for fn, score, desc in _CHECKS:
            try:
                is_hit = bool(fn())
            except Exception:
                is_hit = False
            rule = re.sub(r'^[+-]?\d+\s*', '', desc).strip()
            delta = score if is_hit else 0
            total += delta
            result["dim2_items"].append({
                "rule": rule,
                "max_delta": score,
                "delta": delta,
                "hit": is_hit,
                "detail": "",
            })
        result["total_score"] = total
        return result
    except Exception as e:
        result["status"] = "error"
        result["error"] = f"{type(e).__name__}: {e}"
        return result


if __name__ == "__main__":
    # 本地调试：允许命令行传入目录；默认使用脚本所在目录。
    target_dir = sys.argv[1] if len(sys.argv) > 1 else str(Path(__file__).resolve().parent)
    print(json.dumps(evaluate(target_dir), ensure_ascii=False, indent=2))
