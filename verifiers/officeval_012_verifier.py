# -*- coding: utf-8 -*-
"""
Word文档自动评估工具（officeval_012）
根据"打分细则"对Word文档进行格式检查和评分。

对外统一入口：`evaluate(dir_path: str) -> dict`
  - 参数：脚本所在目录的路径；脚本自行在该目录内定位并打开被评估的 Word 文档。
  - 返回：结构化 dict，字段见"脚本接口差异与统一建议.md" §2.2。
  - 不打印主结果、不修改 sys.stdout、不 sys.exit、不硬编码路径。
"""

import os
import re
import sys
import json

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


class WordDocumentEvaluator:
    """Word文档评估器"""

    def __init__(self, file_path):
        self.file_path = file_path
        self.document = None
        self.dimension1_passed = False
        self.dimension1_reason = ''
        self.score = 0
        self.check_results = []  # 存储所有检查结果: (item_text, passed, score_change)

    def run(self):
        """静默执行完整评估（不打印），供模块级 evaluate() 汇总结构化结果调用"""
        # 维度1检查
        dim1_result = self._check_dimension1()
        if not dim1_result['passed']:
            self.dimension1_passed = False
            self.dimension1_reason = dim1_result['reason']
            return

        self.dimension1_passed = True
        self.dimension1_reason = ''

        # 维度2检查
        self._check_dimension2()

    def _check_dimension1(self):
        """检查维度1：可用与可修改性"""
        # 检查1：文件扩展名
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext != '.docx':
            return {'passed': False, 'reason': f'文件扩展名不满足要求: {ext}'}

        # 检查2：文件能否正常打开
        # python-docx 仅支持 OOXML（.docx）；不再支持 .doc。
        try:
            self.document = Document(self.file_path)
        except Exception as e:
            return {'passed': False, 'reason': f'文件无法正常打开: {str(e)}'}

        # 检查3：是否为可编辑Word文档（非整页图片）
        if self._check_full_page_images():
            return {'passed': False, 'reason': '文件存在一整页均为图片的情况'}

        # 当前脚本不再把连续空白页、乱码/文字重叠面积作为维度一门禁。

        return {'passed': True, 'reason': ''}

    def _check_full_page_images(self):
        """检查是否存在整页图片"""
        image_count = 0
        text_count = 0

        for para in self.document.paragraphs:
            if para.text.strip():
                text_count += 1
            for run in para.runs:
                if run._element.xpath('.//a:blip'):
                    image_count += 1

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            text_count += 1
                        for run in para.runs:
                            if run._element.xpath('.//a:blip'):
                                image_count += 1

        return False

    def _check_dimension2(self):
        """检查维度2：完成度评分"""

        # 得分点检查
        self._check_toc_title_font()
        self._check_toc_title_spacing()
        self._check_toc_title_paragraph()
        self._check_toc_auto_generated()
        self._check_toc_content_font()
        self._check_toc_content_paragraph()
        self._check_toc_page_number_align()
        self._check_abstract_font()
        self._check_abstract_title_font()
        self._check_abstract_paragraph()
        self._check_academic_statement_paragraph()
        self._check_keywords_font()
        self._check_abstract_en_title_font()
        self._check_abstract_en_content_font()
        self._check_abstract_en_paragraph()
        self._check_thesis_title_font()
        self._check_thesis_title_paragraph()
        self._check_level1_title_font()
        self._check_level1_title_paragraph()
        self._check_level2_title_font()
        self._check_level2_title_paragraph()
        self._check_reference_title_font()
        self._check_reference_title_paragraph()
        self._check_reference_content_font()
        self._check_footer_margin()
        self._check_header_content()
        self._check_header_font()
        self._check_reference_mark_format()
        self._check_reference_list_format()
        self._check_body_font_paragraph()

        # 扣分点检查
        self._check_cover_title_font_penalty()
        self._check_cover_title_paragraph_penalty()
        self._check_cover_other_font_penalty()
        self._check_cover_table_font_penalty()
        self._check_cover_table_align_penalty()
        self._check_cover_address_penalty()
        self._check_statement_title_penalty()
        self._check_statement_paragraph_penalty()
        self._check_statement_font_penalty()
        self._check_abstract_align_penalty()
        self._check_keywords_paragraph_penalty()
        self._check_research_direction_penalty()
        self._check_body_font_penalty()
        self._check_page_number_format_penalty()
        self._check_page_number_style_penalty()
        self._check_header_margin_penalty()
        self._check_header_line_penalty()
        self._check_header_line_style_penalty()

    def _find_paragraph_by_text(self, text_keyword, exact=False):
        """根据文本关键词查找段落"""
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            if exact:
                if text == text_keyword:
                    return para, i
            else:
                if text_keyword in text:
                    return para, i
                # 特殊处理：查找"目录"时，也匹配"目 录"、"目  录"等带空格的情况
                if text_keyword == '目录' and '目' in text and '录' in text:
                    return para, i
        return None, -1

    def _add_result(self, item_text, passed, score_change):
        """添加检查结果"""
        self.check_results.append((item_text, passed, score_change))
        if passed:
            self.score += score_change

    # ==================== 得分点检查方法 ====================

    def _check_toc_title_font(self):
        """+1: 目录页标题字体格式为黑体二号

        细则两个点（标题所有文字须全部满足才 +1）：
          1) 字体为黑体   2) 字号为二号(22pt)
        字体/字号按办公软件继承链解析(以实际显示为准)。
        """
        item = "+1: 目录页标题字体格式为黑体二号"
        para, idx = self._find_paragraph_by_text("目录")
        if para is None:
            self._add_result(item, False, 0)
            return

        para_elem = para._element
        checked_any = False
        all_ok = True
        for run in para.runs:
            if not run.text or not run.text.strip():
                continue
            checked_any = True
            font_name, size = self._resolve_run_cn_font(run._element, para_elem)
            # 点1：黑体
            if not (font_name and '黑体' in str(font_name)):
                all_ok = False
            # 点2：二号(22pt)
            if not (size is not None and abs(size - 22) < 0.5):
                all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_toc_title_spacing(self):
        """+1: 目录标题两字中间空一个字符"""
        para, idx = self._find_paragraph_by_text("目录")
        if para is None:
            self._add_result("+1: 目录标题两字中间空一个字符", False, 0)
            return

        # 用run拼接的文本，兼容"目"、"录"被拆分到不同run的情况
        full_text = "".join(run.text for run in para.runs) or (para.text or "")
        text = full_text.strip()

        # 定位"目"与其后的"录"，取两字之间的内容
        if "目" not in text or "录" not in text:
            self._add_result("+1: 目录标题两字中间空一个字符", False, 0)
            return
        i = text.index("目")
        j = text.find("录", i + 1)
        if j == -1:
            self._add_result("+1: 目录标题两字中间空一个字符", False, 0)
            return
        gap = text[i + 1:j]

        # 细则唯一要求：两字之间空"一个字符"（即一个汉字的宽度）。
        # 办公软件中两种等效写法均渲染为一个汉字宽度：
        #   - 一个全角空格(U+3000)          → 宽度 1
        #   - 两个半角空格(U+0020)          → 0.5 × 2 = 1
        # 只允许由空格组成，且总宽度恰为一个字符。
        if gap and all(c in (" ", "　") for c in gap):
            width = sum(1.0 if c == "　" else 0.5 for c in gap)
            if abs(width - 1.0) < 0.01:
                self._add_result("+1: 目录标题两字中间空一个字符", True, 1)
                return

        self._add_result("+1: 目录标题两字中间空一个字符", False, 0)

    def _check_toc_title_paragraph(self):
        """+1: 目录页标题段落格式为居中对齐、段前0.5行、段后0.5行、两倍行距"""
        para, idx = self._find_paragraph_by_text("目录")
        if para is None:
            self._add_result("+1: 目录页标题段落格式(居中、段前段后0.5行、两倍行距)", False, 0)
            return

        pPr = para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # ---- 点1：居中对齐（段落直接 jc 优先，否则样式链）----
        jc_val = None
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                jc_val = jc.get(qn('w:val'))
        if jc_val is None and style_id:
            jc_val = self._style_jc_lookup(style_id)
        is_center = jc_val == 'center'

        # ---- 点4：两倍行距（lineRule=auto 且 line=480=2×240，取直接格式或样式链）----
        space = None
        if pPr is not None:
            space = pPr.find(qn('w:spacing'))
        if space is None and style_id:
            space = self._style_spacing_lookup(style_id)
        is_double_spacing = False
        if space is not None:
            line = space.get(qn('w:line'))
            rule = space.get(qn('w:lineRule'))
            if line and rule in (None, 'auto'):
                try:
                    is_double_spacing = abs(int(line) / 240.0 - 2.0) < 0.1
                except ValueError:
                    is_double_spacing = False

        # ---- 点2、点3：段前0.5行、段后0.5行（以办公软件显示的"行"为准）----
        has_before = self._spacing_lines_ok(pPr, style_id, 'before', 0.5)
        has_after = self._spacing_lines_ok(pPr, style_id, 'after', 0.5)

        passed = is_center and is_double_spacing and has_before and has_after
        self._add_result("+1: 目录页标题段落格式(居中、段前段后0.5行、两倍行距)", passed, 1 if passed else 0)

    def _check_toc_auto_generated(self):
        """+5: 目录为自动生成目录"""
        # 细则唯一要求：目录是"自动生成"的。
        # 在办公软件(Word/WPS)中，自动生成目录的本质是插入了一个 TOC 域(field)。
        # OOXML 中 TOC 域有两种写法，两者都应判为自动目录：
        #   1) 简单域：<w:fldSimple w:instr="TOC ..."/>
        #   2) 复杂域：<w:fldChar begin> ... <w:instrText>TOC ...</w:instrText> ... <w:fldChar end>
        # 仅凭书签/超链接/任意域来判断会误判(如页码域、普通书签)，故只认 TOC 域指令。
        has_toc = False
        doc_elem = self.document._element

        # 写法1：简单域 fldSimple，instr 含 TOC
        try:
            for fld in doc_elem.findall('.//' + qn('w:fldSimple')):
                instr = fld.get(qn('w:instr'))
                if instr and 'TOC' in instr.upper():
                    has_toc = True
                    break
        except:
            pass

        # 写法2：复杂域，指令文本 instrText 含 TOC
        if not has_toc:
            try:
                for instr_text in doc_elem.findall('.//' + qn('w:instrText')):
                    if instr_text.text and 'TOC' in instr_text.text.upper():
                        has_toc = True
                        break
            except:
                pass

        self._add_result("+5: 目录为自动生成目录", has_toc, 5 if has_toc else 0)

    def _doc_default_cn_font(self):
        """docDefaults 中的默认中文字体名与字号(pt)"""
        try:
            dd = self.document.styles.element.find(qn('w:docDefaults'))
            if dd is not None:
                rPr = dd.find('.//' + qn('w:rPr'))
                if rPr is not None:
                    font_name = None
                    size = None
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None:
                        font_name = rf.get(qn('w:eastAsia'))
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None and sz.get(qn('w:val')):
                        size = int(sz.get(qn('w:val'))) / 2
                    return font_name, size
        except:
            pass
        return None, None

    def _default_paragraph_style_id(self):
        """返回文档默认段落样式(w:default='1' 的 paragraph style)的 styleId。
        办公软件里段落若未显式挂 pStyle，会按此默认段落样式(通常是 Normal)渲染 —
        样式链回溯要从它开始，才能与办公软件显示一致。"""
        if not self.document:
            return None
        try:
            for style in self.document.styles.element.findall(qn('w:style')):
                if style.get(qn('w:type')) == 'paragraph' and style.get(qn('w:default')) == '1':
                    return style.get(qn('w:styleId'))
        except:
            pass
        return None

    def _style_rpr_lookup(self, style_id, style_type, want):
        """沿样式的 basedOn 链查找 rPr：want='font'返回eastAsia字体名，want='size'返回字号(pt)，want='bold'返回是否加粗(bool)"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, style_type)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            rPr = style.element.rPr
            if rPr is not None:
                if want == 'font':
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None and rf.get(qn('w:eastAsia')):
                        return rf.get(qn('w:eastAsia'))
                elif want == 'asciifont':
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is not None:
                        v = rf.get(qn('w:ascii')) or rf.get(qn('w:hAnsi'))
                        if v:
                            return v
                elif want == 'bold':
                    b = rPr.find(qn('w:b'))
                    if b is not None:
                        val = b.get(qn('w:val'))
                        return (val is None) or (str(val) not in ('0', 'false', 'off'))
                else:
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None and sz.get(qn('w:val')):
                        return int(sz.get(qn('w:val'))) / 2
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), style_type)
            except:
                break
        return None

    def _resolve_run_cn_font(self, run_elem, para_elem):
        """按办公软件渲染顺序解析 run 的中文字体名与字号(pt)：
        run直接格式 -> 字符样式(rStyle)链 -> 段落直接格式 -> 段落样式链 -> docDefaults"""
        font_name = None
        size = None
        rstyle_id = None
        pstyle_id = None

        rPr = run_elem.find(qn('w:rPr'))
        if rPr is not None:
            rf = rPr.find(qn('w:rFonts'))
            if rf is not None and rf.get(qn('w:eastAsia')):
                font_name = rf.get(qn('w:eastAsia'))
            sz = rPr.find(qn('w:sz'))
            if sz is not None and sz.get(qn('w:val')):
                size = int(sz.get(qn('w:val'))) / 2
            rstyle = rPr.find(qn('w:rStyle'))
            if rstyle is not None:
                rstyle_id = rstyle.get(qn('w:val'))

        if font_name is None and rstyle_id:
            font_name = self._style_rpr_lookup(rstyle_id, WD_STYLE_TYPE.CHARACTER, 'font')
        if size is None and rstyle_id:
            size = self._style_rpr_lookup(rstyle_id, WD_STYLE_TYPE.CHARACTER, 'size')

        pPr = para_elem.find(qn('w:pPr'))
        if pPr is not None:
            p_rPr = pPr.find(qn('w:rPr'))
            if p_rPr is not None:
                if font_name is None:
                    rf = p_rPr.find(qn('w:rFonts'))
                    if rf is not None and rf.get(qn('w:eastAsia')):
                        font_name = rf.get(qn('w:eastAsia'))
                if size is None:
                    sz = p_rPr.find(qn('w:sz'))
                    if sz is not None and sz.get(qn('w:val')):
                        size = int(sz.get(qn('w:val'))) / 2
            pstyle = pPr.find(qn('w:pStyle'))
            if pstyle is not None:
                pstyle_id = pstyle.get(qn('w:val'))

        if font_name is None and pstyle_id:
            font_name = self._style_rpr_lookup(pstyle_id, WD_STYLE_TYPE.PARAGRAPH, 'font')
        if size is None and pstyle_id:
            size = self._style_rpr_lookup(pstyle_id, WD_STYLE_TYPE.PARAGRAPH, 'size')

        if font_name is None or size is None:
            d_font, d_size = self._doc_default_cn_font()
            if font_name is None:
                font_name = d_font
            if size is None:
                size = d_size

        return font_name, size

    def _resolve_run_cn_bold(self, run_elem, para_elem):
        """按办公软件渲染顺序解析 run 是否加粗(bool)：
        run直接格式 -> 字符样式(rStyle)链 -> 段落直接格式 -> 段落样式链 -> docDefaults。
        w:b 无 val 或 val 非 0/false/off 视为加粗。"""
        def b_from_rPr(rPr):
            if rPr is None:
                return None
            b = rPr.find(qn('w:b'))
            if b is None:
                return None
            val = b.get(qn('w:val'))
            return (val is None) or (str(val) not in ('0', 'false', 'off'))

        # 1. run 直接格式
        rPr = run_elem.find(qn('w:rPr'))
        if rPr is not None:
            v = b_from_rPr(rPr)
            if v is not None:
                return v
            # 2. 字符样式(rStyle)链
            rstyle = rPr.find(qn('w:rStyle'))
            if rstyle is not None and rstyle.get(qn('w:val')):
                v = self._style_rpr_lookup(rstyle.get(qn('w:val')), WD_STYLE_TYPE.CHARACTER, 'bold')
                if v is not None:
                    return v

        # 3. 段落直接格式 + 4. 段落样式链
        pPr = para_elem.find(qn('w:pPr'))
        if pPr is not None:
            v = b_from_rPr(pPr.find(qn('w:rPr')))
            if v is not None:
                return v
            pstyle = pPr.find(qn('w:pStyle'))
            if pstyle is not None and pstyle.get(qn('w:val')):
                v = self._style_rpr_lookup(pstyle.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH, 'bold')
                if v is not None:
                    return v

        # 5. docDefaults
        try:
            dd = self.document.styles.element.find(qn('w:docDefaults'))
            if dd is not None:
                v = b_from_rPr(dd.find('.//' + qn('w:rPr')))
                if v is not None:
                    return v
        except:
            pass
        return False

    def _resolve_run_en_font(self, run_elem, para_elem):
        """按办公软件渲染顺序解析 run 的英文(ascii/hAnsi)字体名：
        run直接格式 -> 字符样式(rStyle)链 -> 段落直接格式 -> 段落样式链 -> docDefaults。"""
        def font_from_rPr(rPr):
            if rPr is None:
                return None
            rf = rPr.find(qn('w:rFonts'))
            if rf is not None:
                return rf.get(qn('w:ascii')) or rf.get(qn('w:hAnsi'))
            return None

        # 1. run 直接格式
        rPr = run_elem.find(qn('w:rPr'))
        if rPr is not None:
            v = font_from_rPr(rPr)
            if v:
                return v
            # 2. 字符样式(rStyle)链
            rstyle = rPr.find(qn('w:rStyle'))
            if rstyle is not None and rstyle.get(qn('w:val')):
                v = self._style_rpr_lookup(rstyle.get(qn('w:val')), WD_STYLE_TYPE.CHARACTER, 'asciifont')
                if v:
                    return v

        # 3. 段落直接格式 + 4. 段落样式链
        pPr = para_elem.find(qn('w:pPr'))
        if pPr is not None:
            v = font_from_rPr(pPr.find(qn('w:rPr')))
            if v:
                return v
            pstyle = pPr.find(qn('w:pStyle'))
            if pstyle is not None and pstyle.get(qn('w:val')):
                v = self._style_rpr_lookup(pstyle.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH, 'asciifont')
                if v:
                    return v

        # 5. docDefaults
        try:
            dd = self.document.styles.element.find(qn('w:docDefaults'))
            if dd is not None:
                v = font_from_rPr(dd.find('.//' + qn('w:rPr')))
                if v:
                    return v
        except:
            pass
        return None

    def _resolve_run_color(self, run_elem, para_elem):
        """按办公软件渲染顺序解析 run 的字体颜色(w:color 的 val, 六位十六进制大写或 'auto')：
        run直接格式 -> 字符样式(rStyle)链 -> 段落直接格式 -> 段落样式链 -> docDefaults。
        未设置任何颜色时返回 None（办公软件默认自动/黑色）。"""
        def color_from_rPr(rPr):
            if rPr is None:
                return None
            c = rPr.find(qn('w:color'))
            if c is not None and c.get(qn('w:val')):
                return c.get(qn('w:val'))
            return None

        rPr = run_elem.find(qn('w:rPr'))
        if rPr is not None:
            v = color_from_rPr(rPr)
            if v:
                return v
            rstyle = rPr.find(qn('w:rStyle'))
            if rstyle is not None and rstyle.get(qn('w:val')):
                v = self._style_color_lookup(rstyle.get(qn('w:val')), WD_STYLE_TYPE.CHARACTER)
                if v:
                    return v

        pPr = para_elem.find(qn('w:pPr'))
        if pPr is not None:
            v = color_from_rPr(pPr.find(qn('w:rPr')))
            if v:
                return v
            pstyle = pPr.find(qn('w:pStyle'))
            if pstyle is not None and pstyle.get(qn('w:val')):
                v = self._style_color_lookup(pstyle.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
                if v:
                    return v

        try:
            dd = self.document.styles.element.find(qn('w:docDefaults'))
            if dd is not None:
                v = color_from_rPr(dd.find('.//' + qn('w:rPr')))
                if v:
                    return v
        except:
            pass
        return None

    def _style_color_lookup(self, style_id, style_type):
        """沿样式 basedOn 链查找 w:color 的 val；无则返回 None"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, style_type)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            rPr = style.element.rPr
            if rPr is not None:
                c = rPr.find(qn('w:color'))
                if c is not None and c.get(qn('w:val')):
                    return c.get(qn('w:val'))
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), style_type)
            except:
                break
        return None

    def _toc_level_from_style(self, style_id):
        """若段落样式(沿 basedOn 链)对应办公软件的目录样式 toc 1/2/3/…，返回其级数 int；否则返回 None。
        目录条目段落的样式在不同办公软件里 styleId 可能只是编号(如 '14')，但样式 name 均为 'toc N'。
        因此按样式 name 识别(不区分大小写、忽略中间空白)，而非按 styleId 前缀。
        注：此函数用于识别"是否属于目录条目段落"，判分时按细则用条目文本自身确定级别。
        """
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            name_el = style.element.find(qn('w:name'))
            name = name_el.get(qn('w:val')) if name_el is not None else None
            if name:
                m = re.match(r'^toc\s*(\d+)$', name.strip(), re.IGNORECASE)
                if m:
                    return int(m.group(1))
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None

    def _toc_level_from_text(self, text):
        """按目录条目文本自身判定其"应"对应的标题级别(1/2/3)，与文档里挂的样式级别无关。
        识别规则与正文标题层级一致：
          L1: 第X章 / X、(X是汉字数字) / 摘要 / Abstract / 参考文献 / 致谢 / 结论
          L2: （X）(X是汉字数字) / X.Y(数字，非 X.Y.Z 前缀)
          L3: X.Y.Z(数字)
        text 传入目录条目的文字(末尾可能含页码，仅匹配起始不受影响)。
        """
        t = text.strip()
        # L3：X.Y.Z 数字三级
        if re.match(r'^\d+\.\d+\.\d+', t):
            return 3
        # L2：（汉字）或 X.Y
        if re.match(r'^[（(][一二三四五六七八九十]+[)）]', t) or re.match(r'^\d+\.\d+(?!\d)', t):
            return 2
        # L1：第X章 / X、汉字 / 前置或末尾常见一级页(摘要/Abstract/参考文献/致谢/结论)
        if (re.match(r'^第[一二三四五六七八九十]+章', t)
                or re.match(r'^[一二三四五六七八九十]+[、．.]', t)
                or t.startswith('摘要')
                or t.lower().startswith('abstract')
                or t.startswith('参考文献')
                or t.startswith('致谢')
                or t.startswith('致  谢')
                or t.startswith('致 谢')
                or t.startswith('结论')):
            return 1
        return None

    def _para_left_indent_chars_strict(self, pPr, style_id):
        """严格"字符"单位左缩进：只认 w:ind/@startChars 或 leftChars(1/100字符)。
        办公软件里按"字符"单位设置左缩进时才写这两个属性；按厘米/磅等绝对值(twips)设置
        的一律视为未按字符设置 → 返回 0.0，即"不是 N 字符"。
        段落直接格式优先，缺失回退样式链。"""
        ind = None
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = self._style_ind_lookup(style_id)
        if ind is None:
            return 0.0
        lc = ind.get(qn('w:startChars'))
        if lc is None:
            lc = ind.get(qn('w:leftChars'))
        if lc is not None:
            try:
                return int(lc) / 100.0
            except ValueError:
                return 0.0
        return 0.0

    def _check_toc_content_font(self):
        """+3: 目录页内容中文字体格式为宋体四号"""
        # 目录内容段落采用 toc 1/toc 2/toc 3 等目录样式（自动目录通常包裹在 sdt 中，
        # 故用 XML 遍历 body 下所有 w:p，python-docx 的 paragraphs 不含 sdt 内段落）。
        # 不同办公软件 styleId 可能只是编号(如 "14")，但样式 name 均为 "toc N"，
        # 因此用 _toc_level_from_style 沿 basedOn 链按样式名识别。
        body = self.document._element.find(qn('w:body'))
        if body is None:
            self._add_result("+3: 目录页内容中文字体格式为宋体四号", False, 0)
            return

        checked_any = False
        all_ok = True
        for p in body.findall('.//' + qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                continue
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                continue
            style_id = pStyle.get(qn('w:val'))
            if self._toc_level_from_style(style_id) is None:
                continue

            for run_elem in p.findall('.//' + qn('w:r')):
                t = run_elem.find(qn('w:t'))
                if t is None or not t.text:
                    continue
                # 仅检查中文字符（细则点针对中文字体；页码等非中文不在此项约束内）
                if not any('一' <= c <= '鿿' for c in t.text):
                    continue
                checked_any = True
                font_name, size = self._resolve_run_cn_font(run_elem, p)
                is_songti = font_name is not None and '宋体' in str(font_name)   # 点1：宋体
                is_no4 = size is not None and abs(size - 14) < 1                 # 点2：四号14pt
                if not (is_songti and is_no4):
                    all_ok = False

        passed = checked_any and all_ok
        self._add_result("+3: 目录页内容中文字体格式为宋体四号", passed, 3 if passed else 0)

    def _style_ind_lookup(self, style_id):
        """沿样式 basedOn 链查找最近的 w:ind 元素"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                ind = s_pPr.find(qn('w:ind'))
                if ind is not None:
                    return ind
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None

    def _style_linespacing_lookup(self, style_id):
        """沿样式 basedOn 链查找行距倍数(仅 lineRule=auto 时按倍数计)"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                sp = s_pPr.find(qn('w:spacing'))
                if sp is not None and sp.get(qn('w:line')):
                    if sp.get(qn('w:lineRule')) == 'auto':
                        return int(sp.get(qn('w:line'))) / 240.0
                    return None
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None

    def _para_line_spacing(self, pPr, style_id):
        """段落有效行距倍数：段落直接格式(lineRule=auto) -> 样式链"""
        if pPr is not None:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None and sp.get(qn('w:line')):
                if sp.get(qn('w:lineRule')) == 'auto':
                    return int(sp.get(qn('w:line'))) / 240.0
                return None
        return self._style_linespacing_lookup(style_id)

    def _para_left_indent_chars(self, p, pPr, style_id):
        """段落有效左缩进(以'字符'为单位)。
        办公软件两种存法：w:startChars/leftChars(单位1/100字符) 或 w:start/left(twips)。
        twips 需按该段落有效中文字号换算：1字符 = 字号pt × 20 twips。"""
        ind = None
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = self._style_ind_lookup(style_id)
        if ind is None:
            return 0.0
        # 优先字符单位
        lc = ind.get(qn('w:startChars'))
        if lc is None:
            lc = ind.get(qn('w:leftChars'))
        if lc is not None:
            return int(lc) / 100.0
        # 回退 twips，按有效字号换算
        lt = ind.get(qn('w:start'))
        if lt is None:
            lt = ind.get(qn('w:left'))
        if lt is not None:
            size_pt = self._para_cn_size(p, pPr, style_id)
            if size_pt:
                return int(lt) / (size_pt * 20.0)
        return 0.0

    def _para_cn_size(self, p, pPr, style_id):
        """段落有效中文字号(pt)：段落直接rPr -> 首个run(直接/rStyle) -> 段落样式链 -> docDefaults"""
        if pPr is not None:
            p_rPr = pPr.find(qn('w:rPr'))
            if p_rPr is not None:
                sz = p_rPr.find(qn('w:sz'))
                if sz is not None and sz.get(qn('w:val')):
                    return int(sz.get(qn('w:val'))) / 2
        for r in p.findall('.//' + qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                sz = rPr.find(qn('w:sz'))
                if sz is not None and sz.get(qn('w:val')):
                    return int(sz.get(qn('w:val'))) / 2
                rstyle = rPr.find(qn('w:rStyle'))
                if rstyle is not None:
                    s = self._style_rpr_lookup(rstyle.get(qn('w:val')), WD_STYLE_TYPE.CHARACTER, 'size')
                    if s:
                        return s
        s = self._style_rpr_lookup(style_id, WD_STYLE_TYPE.PARAGRAPH, 'size')
        if s:
            return s
        _, d = self._doc_default_cn_font()
        return d

    def _check_toc_content_paragraph(self):
        """+3: 目录页段落格式为1.5倍行距，一级标题顶格，二级标题缩进1字符，三级标题缩进2字符"""
        # 目录条目段落用 toc N 样式，通常包裹在 sdt 中，遍历 body 下所有 w:p。
        # 用样式名识别目录条目(_toc_level_from_style)；判分级别按细则改从"条目文本自身"推断
        # (_toc_level_from_text)——这样即便文档把所有条目挂在同一 toc 样式上，也按文本内容
        # 决定该条应属一级/二级/三级，检查其对应目标缩进。
        # 缩进严格按"字符"单位：只认 w:ind/@startChars|leftChars(1/100字符)；
        # 按厘米/磅等绝对值(twips)设置一律视为"不是 N 字符"。
        body = self.document._element.find(qn('w:body'))
        if body is None:
            self._add_result("+3: 目录页段落格式(1.5倍行距、标题缩进)", False, 0)
            return

        # 各级目标缩进（字符）：一级=0(顶格) 二级=1 三级=2
        target_chars = {1: 0.0, 2: 1.0, 3: 2.0}
        checked_any = False
        spacing_ok = True   # 点1：所有目录条目 1.5 倍行距
        indent_ok = True    # 点2/3/4：各级缩进

        for p in body.findall('.//' + qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                continue
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                continue
            style_id = pStyle.get(qn('w:val'))
            # 仅确认这是"目录条目段落"，不用其级别
            if self._toc_level_from_style(style_id) is None:
                continue
            text = ''.join((t.text or '') for t in p.findall('.//' + qn('w:t')))
            if not text.strip():
                continue

            # 按条目文本推断"应"对应的标题级别
            level = self._toc_level_from_text(text)
            if level not in target_chars:
                continue

            checked_any = True

            # 点1：1.5 倍行距
            ls = self._para_line_spacing(pPr, style_id)
            if not (ls is not None and abs(ls - 1.5) < 0.05):
                spacing_ok = False

            # 点2/3/4：严格按"字符"单位判缩进
            chars = self._para_left_indent_chars_strict(pPr, style_id)
            if abs(chars - target_chars[level]) > 0.15:
                indent_ok = False

        passed = checked_any and spacing_ok and indent_ok
        self._add_result("+3: 目录页段落格式(1.5倍行距、标题缩进)", passed, 3 if passed else 0)

    def _style_tabs_lookup(self, style_id):
        """沿样式 basedOn 链查找最近的 w:tabs 元素"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                tabs = s_pPr.find(qn('w:tabs'))
                if tabs is not None:
                    return tabs
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None

    def _check_toc_page_number_align(self):
        """+3: 目录页的目录条目页码向右相互对齐"""
        # 在办公软件(Word/WPS)中，目录条目的页码"向右相互对齐"是通过一个右对齐制表位
        # (w:tab val="right"，通常带 dot 引导符，位于右边距处)实现的：
        # 页码文本前有一个制表符(w:tab)，跳到该右制表位，各条目页码从而在右侧对齐到同一位置。
        # 因此逐条目录检查：其有效制表位(段落直接格式或样式链)中含 val="right" 的制表位，
        # 且段落中存在制表符将页码推到右侧。
        # 目录段落用 _toc_level_from_style 沿 basedOn 链按样式名 "toc N" 识别。
        body = self.document._element.find(qn('w:body'))
        if body is None:
            self._add_result("+3: 目录页的目录条目页码向右相互对齐", False, 0)
            return

        checked_any = False
        all_ok = True
        for p in body.findall('.//' + qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                continue
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                continue
            style_id = pStyle.get(qn('w:val'))
            if self._toc_level_from_style(style_id) is None:
                continue
            t = p.find('.//' + qn('w:t'))
            if t is None or not (t.text and t.text.strip()):
                continue
            checked_any = True

            # 有效制表位：优先段落直接格式，其次样式链
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = self._style_tabs_lookup(style_id)
            has_right_tab = False
            if tabs is not None:
                for tab in tabs.findall(qn('w:tab')):
                    if tab.get(qn('w:val')) == 'right':
                        has_right_tab = True
                        break

            # 段落中存在制表符(页码前的跳格)
            has_tab_char = p.find('.//' + qn('w:tab')) is not None

            if not (has_right_tab and has_tab_char):
                all_ok = False

        passed = checked_any and all_ok
        self._add_result("+3: 目录页的目录条目页码向右相互对齐", passed, 3 if passed else 0)

    def _check_abstract_font(self):
        """+1: 中文摘要页内容除"摘要""关键词""研究方向"，其余字体格式为宋体四号"""
        title_para, idx = self._find_paragraph_by_text("摘要")
        if title_para is None:
            self._add_result("+1: 中文摘要页内容(除标题/关键词/研究方向)字体格式为宋体四号", False, 0)
            return

        paras = self.document.paragraphs
        # 中文摘要页范围：自"摘要"标题起，至英文摘要"Abstract"标题(或文档末尾)止
        end = len(paras)
        for k in range(idx + 1, len(paras)):
            if paras[k].text.strip().lower().startswith("abstract"):
                end = k
                break

        checked_any = False
        all_ok = True
        for k in range(idx, end):
            para = paras[k]
            text = para.text.strip()
            if not text:
                continue
            # 排除项：摘要标题本身、关键词行、研究方向行（各有独立字体要求，不在本项约束内）
            if k == idx:
                continue
            if text.startswith("关键词") or text.startswith("研究方向"):
                continue

            # 其余内容：所有中文字符必须为宋体四号
            for run in para.runs:
                if not run.text:
                    continue
                if not any('一' <= c <= '鿿' for c in run.text):
                    continue
                checked_any = True
                font_name, size = self._resolve_run_cn_font(run._element, para._element)
                is_songti = font_name is not None and '宋体' in str(font_name)   # 点1：宋体
                is_no4 = size is not None and abs(size - 14) < 1                 # 点2：四号14pt
                if not (is_songti and is_no4):
                    all_ok = False

        passed = checked_any and all_ok
        self._add_result("+1: 中文摘要页内容(除标题/关键词/研究方向)字体格式为宋体四号", passed, 1 if passed else 0)

    def _check_abstract_title_font(self):
        """+1: 中文摘要页标题字体格式为黑体小三"""
        # 中文摘要页标题即"摘要"二字所在段落。document.paragraphs 不含 sdt 内的目录条目，
        # 故首个含"摘要"的段落即为正文中的标题段。
        para, idx = self._find_paragraph_by_text("摘要")
        if para is None:
            self._add_result("+1: 中文摘要页标题字体格式为黑体小三", False, 0)
            return

        checked_any = False
        all_ok = True
        for run in para.runs:
            if not run.text:
                continue
            # 仅检查中文字符（标题"摘要"二字）
            if not any('一' <= c <= '鿿' for c in run.text):
                continue
            checked_any = True
            font_name, size = self._resolve_run_cn_font(run._element, para._element)
            is_heiti = font_name is not None and '黑体' in str(font_name)   # 点1：黑体
            is_xiao3 = size is not None and abs(size - 15) < 1              # 点2：小三15pt
            if not (is_heiti and is_xiao3):
                all_ok = False

        passed = checked_any and all_ok
        self._add_result("+1: 中文摘要页标题字体格式为黑体小三", passed, 1 if passed else 0)

    def _check_abstract_paragraph(self):
        """+1: 中文摘要页段落格式满足段前0.5行、段后0.5行、两倍行距"""
        para, idx = self._find_paragraph_by_text("摘要")
        if para is None:
            self._add_result("+1: 中文摘要页段落格式(段前段后0.5行、两倍行距)", False, 0)
            return

        pPr = para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # ---- 点3：两倍行距（lineRule=auto 且 line=480=2×240，取直接格式或样式链）----
        space = None
        if pPr is not None:
            space = pPr.find(qn('w:spacing'))
        if space is None and style_id:
            space = self._style_spacing_lookup(style_id)
        is_double_spacing = False
        if space is not None:
            line = space.get(qn('w:line'))
            rule = space.get(qn('w:lineRule'))
            if line and rule in (None, 'auto'):
                try:
                    is_double_spacing = abs(int(line) / 240.0 - 2.0) < 0.1
                except ValueError:
                    is_double_spacing = False

        # ---- 点1、点2：段前0.5行、段后0.5行（以办公软件显示的"行"为准）----
        has_before = self._spacing_lines_ok(pPr, style_id, 'before', 0.5)
        has_after = self._spacing_lines_ok(pPr, style_id, 'after', 0.5)

        passed = is_double_spacing and has_before and has_after
        self._add_result("+1: 中文摘要页段落格式(段前段后0.5行、两倍行距)", passed, 1 if passed else 0)

    def _style_spacing_line_lookup(self, style_id):
        """沿样式 basedOn 链查找 (line, lineRule)；无则返回 (None, None)"""
        if not style_id or not self.document:
            return None, None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None, None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                sp = s_pPr.find(qn('w:spacing'))
                if sp is not None and sp.get(qn('w:line')):
                    return sp.get(qn('w:line')), sp.get(qn('w:lineRule'))
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None, None

    def _style_jc_lookup(self, style_id):
        """沿样式 basedOn 链查找段落对齐方式 w:jc 的 val；无则返回 None"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                jc = s_pPr.find(qn('w:jc'))
                if jc is not None and jc.get(qn('w:val')):
                    return jc.get(qn('w:val'))
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None

    def _style_spacing_lookup(self, style_id):
        """沿样式 basedOn 链查找段落 w:spacing 元素（含段前段后/行距属性）；无则返回 None"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                sp = s_pPr.find(qn('w:spacing'))
                if sp is not None:
                    return sp
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None

    def _para_spacing_lines(self, pPr, style_id, which):
        """段前/段后的"行"数(仅当办公软件按"行"为单位设置时)：which='before'/'after'。
        细则以"行"为单位，Word/WPS 按"行"设置段间距时写 w:beforeLines/afterLines(1/100行)；
        按"磅/厘米"等绝对值设置时只写 w:before/after(twips) 而无 *Lines —— 那不算"N行"。
        故只读 *Lines；无则返回 None(表示未以"行"为单位设置)。

        合并语义：按办公软件的"属性级合并"，段落直接 spacing 的 beforeLines/afterLines
        属性缺失时回退样式链；不能整块用段落 spacing 覆盖样式 spacing —— 只要段落 spacing
        没写 *Lines，样式里的 *Lines 就仍然生效(办公软件段落对话框显示的即为此合并值)。"""
        lines_attr = 'w:beforeLines' if which == 'before' else 'w:afterLines'
        direct_sp = pPr.find(qn('w:spacing')) if pPr is not None else None
        # 段落直接 spacing 上的 *Lines 优先
        if direct_sp is not None:
            lv = direct_sp.get(qn(lines_attr))
            if lv is not None:
                try:
                    return int(lv) / 100.0
                except ValueError:
                    return None
        # 段落未设 *Lines → 回退样式链
        style_sp = self._style_spacing_lookup(style_id) if style_id else None
        if style_sp is not None:
            lv = style_sp.get(qn(lines_attr))
            if lv is not None:
                try:
                    return int(lv) / 100.0
                except ValueError:
                    return None
        return None

    def _spacing_lines_ok(self, pPr, style_id, which, target):
        """段前/段后是否为 target 行(须以"行"为单位设置，且值≈target，±0.1行)。"""
        v = self._para_spacing_lines(pPr, style_id, which)
        return v is not None and abs(v - target) <= 0.1

    def _para_is_single_spacing(self, pPr, style_id):
        """段落是否为单倍行距(办公软件有效)：
        有效行距取 段落直接格式 -> 样式链；lineRule=auto 时按 line/240 计倍数，
        全部未设置时默认单倍行距。"""
        line = None
        rule = None
        if pPr is not None:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None and sp.get(qn('w:line')):
                line = sp.get(qn('w:line'))
                rule = sp.get(qn('w:lineRule'))
        if line is None:
            line, rule = self._style_spacing_line_lookup(style_id)
        if line is None:
            return True  # 未设置行距 => 办公软件默认单倍
        if rule in (None, 'auto'):
            try:
                return abs(int(line) / 240.0 - 1.0) < 0.1
            except ValueError:
                return False
        return False  # exact/atLeast 固定值行距不视为单倍

    def _para_first_line_indent_chars(self, pPr, style_id):
        """段落首行缩进(以'字符'为单位)。细则以"字符"为单位，Word/WPS 段落对话框按"字符"
        设置首行缩进时写 w:ind/@w:firstLineChars(1/100字符，2字符=200)；按"磅/厘米"等绝对值
        设置时只写 w:firstLine(twips) 而无 firstLineChars —— 那不算"N字符"。
        故只读 firstLineChars(段落直接格式优先，缺失回退样式链)；无则返回 0.0。"""
        ind = None
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = self._style_ind_lookup(style_id)
        if ind is not None:
            flc = ind.get(qn('w:firstLineChars'))
            if flc is not None:
                try:
                    return int(flc) / 100.0
                except ValueError:
                    return 0.0
        return 0.0

    def _check_academic_statement_paragraph(self):
        """+1: "学术诚信声明及作品使用授权书"页内容段落格式为首行缩进两字符，单倍行距"""
        title_para, idx = self._find_paragraph_by_text("学术诚信声明")
        if title_para is None:
            self._add_result("+1: 学术诚信声明页内容段落格式(首行缩进两字符、单倍行距)", False, 0)
            return

        paras = self.document.paragraphs
        title_style = title_para.style.name if title_para.style else None

        checked_any = False
        indent_ok = True   # 点1：首行缩进两字符
        single_ok = True   # 点2：单倍行距

        for k in range(idx, len(paras)):
            p = paras[k]
            # 限定在声明页范围内（同一声明样式的连续区块）
            if p.style and title_style and p.style.name != title_style:
                break
            text = p.text.strip()
            if not text:
                continue
            if k == idx:
                continue  # 排除页标题（标题字体/格式另有要求，不在本项约束内）
            # 排除签名、日期等短行（非正文内容段落）
            if len(text) < 20:
                continue

            checked_any = True
            pPr = p._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            chars = self._para_first_line_indent_chars(pPr, style_id)
            if abs(chars - 2.0) > 0.4:   # 首行缩进≈2字符
                indent_ok = False

            if not self._para_is_single_spacing(pPr, style_id):
                single_ok = False

        passed = checked_any and indent_ok and single_ok
        self._add_result("+1: 学术诚信声明页内容段落格式(首行缩进两字符、单倍行距)", passed, 1 if passed else 0)

    def _check_keywords_font(self):
        """+1: 中文摘要页"关键词"三字字体格式为黑体、四号、加粗

        细则三个点（全部满足才 +1）：
          1) 字体为黑体
          2) 字号为四号(14pt)
          3) 加粗
        仅约束"关键词"这三个字本身，不约束其后的关键词内容。
        字体/字号/加粗均按办公软件的继承链解析(run->字符样式->段落->段落样式->docDefaults)。
        """
        item = "+1: 关键词字体格式为黑体、四号、加粗"
        para, idx = self._find_paragraph_by_text("关键词")
        if para is None:
            self._add_result(item, False, 0)
            return

        # 定位"关键词"三字在段落中的字符区间 [start, start+3)
        full = "".join(r.text for r in para.runs)
        start = full.find("关键词")
        if start < 0:
            self._add_result(item, False, 0)
            return
        end = start + 3  # 只覆盖"关键词"三个字

        # 找出与该区间重叠的 run（即承载"关键词"三字的 run，可能被拆分为多个）
        label_runs = []
        offset = 0
        for run in para.runs:
            r_len = len(run.text)
            r_start, r_end = offset, offset + r_len
            if r_start < end and r_end > start:  # 区间重叠
                label_runs.append(run)
            offset = r_end
        if not label_runs:
            self._add_result(item, False, 0)
            return

        para_elem = para._element
        all_ok = True
        for run in label_runs:
            font_name, size = self._resolve_run_cn_font(run._element, para_elem)
            bold = self._resolve_run_cn_bold(run._element, para_elem)
            # 点1：黑体
            if not (font_name and '黑体' in str(font_name)):
                all_ok = False
            # 点2：四号(14pt)
            if not (size is not None and abs(size - 14) < 0.5):
                all_ok = False
            # 点3：加粗
            if not bold:
                all_ok = False

        self._add_result(item, all_ok, 1 if all_ok else 0)

    def _check_abstract_en_title_font(self):
        """+1: 英文摘要页标题("Abstract")字体格式为 Times New Roman、小三

        细则两个点（全部满足才 +1）：
          1) 字体为 Times New Roman（英文字体，取 ascii/hAnsi）
          2) 字号为小三(15pt)
        仅约束标题"Abstract"这一行，字体/字号均按办公软件的英文字体继承链解析
        (run直接格式 -> 字符样式 -> 段落直接格式 -> 段落样式 -> docDefaults)。
        """
        item = "+1: 英文摘要页标题字体格式为Times New Roman、小三"
        para, idx = self._find_paragraph_by_text("Abstract")
        if para is None:
            self._add_result(item, False, 0)
            return

        # 仅检查承载"Abstract"标题文字的 run
        para_elem = para._element
        checked_any = False
        all_ok = True
        for run in para.runs:
            if run.text and any(ch.isalpha() for ch in run.text):
                checked_any = True
                font_name = self._resolve_run_en_font(run._element, para_elem)
                # 字号 sz 对中英文一致，复用中文解析链取第二个返回值
                _, size = self._resolve_run_cn_font(run._element, para_elem)
                # 点1：Times New Roman
                if not (font_name and 'Times New Roman' in str(font_name)):
                    all_ok = False
                # 点2：小三(15pt)
                if not (size is not None and abs(size - 15) < 0.5):
                    all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_abstract_en_content_font(self):
        """+1: 英文摘要页内容字体格式为 Times New Roman、四号，"Key words："需加粗

        细则三个点（全部满足才 +1）：
          1) 内容英文字体为 Times New Roman
          2) 内容字号为四号(14pt)
          3) "Key words：" 标签加粗
        内容 = 英文摘要正文段落 + Key words 行（不含 Abstract 标题本身）。
        字体取 ascii/hAnsi，字号/加粗均按办公软件继承链解析。
        """
        item = "+1: 英文摘要页内容字体格式(Times New Roman四号、Key words加粗)"
        para, idx = self._find_paragraph_by_text("Abstract")
        if para is None:
            self._add_result(item, False, 0)
            return

        def is_kw_label(text):
            t = text.replace(" ", "").lower()
            return t.startswith("keywords")

        checked_content = False   # 是否检查到内容run
        content_ok = True         # 点1+点2：内容 Times New Roman 四号
        kw_found = False          # 是否找到 Key words 标签
        kw_bold_ok = True         # 点3：Key words 标签加粗

        for i in range(idx + 1, min(idx + 15, len(self.document.paragraphs))):
            p = self.document.paragraphs[i]
            para_text = p.text.strip()
            if not para_text:
                continue
            para_elem = p._element
            is_kw_para = is_kw_label(para_text)

            # 定位 Key words 标签在段落中的字符区间（"Key words"直到冒号，含中英文冒号）
            label_end = None
            if is_kw_para:
                kw_found = True
                colon = -1
                for ci, ch in enumerate(para_text):
                    if ch in ('：', ':'):
                        colon = ci
                        break
                label_end = (colon + 1) if colon >= 0 else len("Key words")

            offset = 0
            for run in p.runs:
                r_len = len(run.text)
                r_start, r_end = offset, offset + r_len
                offset = r_end
                if not run.text:
                    continue

                # 点3：Key words 标签所在 run 需加粗
                if is_kw_para and label_end is not None and r_start < label_end:
                    if any(ch.isalpha() for ch in run.text):
                        if not self._resolve_run_cn_bold(run._element, para_elem):
                            kw_bold_ok = False

                # 点1+点2：所有含字母的内容 run 需 Times New Roman 四号
                if any(ch.isalpha() for ch in run.text):
                    checked_content = True
                    font_name = self._resolve_run_en_font(run._element, para_elem)
                    _, size = self._resolve_run_cn_font(run._element, para_elem)
                    if not (font_name and 'Times New Roman' in str(font_name)):
                        content_ok = False
                    if not (size is not None and abs(size - 14) < 0.5):
                        content_ok = False

            # 英文摘要内容以 Key words 行结束，其后为中文标题/正文，不再属于本项范围
            if is_kw_para:
                break

        passed = checked_content and content_ok and kw_found and kw_bold_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_abstract_en_paragraph(self):
        """+1: 英文摘要页内容段落格式为首行缩进两字符、两端对齐、单倍行距

        细则三个点（内容正文段落须全部满足才 +1）：
          1) 首行缩进两字符
          2) 两端对齐
          3) 单倍行距
        内容 = Abstract 标题之后、Key words 行之前的英文摘要正文段落。
        缩进/行距均按办公软件继承链解析(段落直接格式 -> 样式链)，
        首行缩进兼容 firstLineChars 与 firstLine(twips 按字号换算) 两种存法。
        """
        item = "+1: 英文摘要页段落格式(首行缩进两字符、两端对齐、单倍行距)"
        para, idx = self._find_paragraph_by_text("Abstract")
        if para is None:
            self._add_result(item, False, 0)
            return

        def is_kw_label(text):
            return text.replace(" ", "").lower().startswith("keywords")

        checked_any = False
        indent_ok = True    # 点1：首行缩进两字符
        justify_ok = True   # 点2：两端对齐
        single_ok = True    # 点3：单倍行距

        for i in range(idx + 1, min(idx + 15, len(self.document.paragraphs))):
            p = self.document.paragraphs[i]
            text = p.text.strip()
            if not text:
                continue
            # Key words 行是关键词标签行，非正文内容段落，到此终止
            if is_kw_label(text):
                break

            checked_any = True
            pPr = p._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            style_id = pStyle.get(qn('w:val')) if pStyle is not None else None
            # 段落未显式挂 pStyle 时，办公软件按"默认段落样式"(通常是 Normal)渲染，
            # 样式链回溯须从默认段落样式开始，才能与显示一致。
            if style_id is None:
                style_id = self._default_paragraph_style_id()

            # 点1：首行缩进两字符
            chars = self._para_first_line_indent_chars(pPr, style_id)
            if abs(chars - 2.0) > 0.4:
                indent_ok = False

            # 点2：两端对齐（办公软件"两端对齐"存为 jc=both/justify）
            jc_val = None
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(qn('w:val'))
            if jc_val is None and style_id:
                jc_val = self._style_jc_lookup(style_id)
            if jc_val not in ('both', 'justify'):
                justify_ok = False

            # 点3：单倍行距
            if not self._para_is_single_spacing(pPr, style_id):
                single_ok = False

        passed = checked_any and indent_ok and justify_ok and single_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_thesis_title_font(self):
        """+1: 论文题目字体格式为三号、黑体、加粗

        细则三个点（题目所有文字须全部满足才 +1）：
          1) 字体为黑体
          2) 字号为三号(16pt)
          3) 加粗
        论文题目为正文首页的完整题名（含"雨水花园"与"维护机制"两个主题词的短段落，
        排除章节标题、关键词行、正文段落）。字体/字号/加粗均按办公软件继承链解析。
        """
        item = "+1: 论文题目字体格式为三号、黑体、加粗"

        # 定位论文题目段落：短段落，同时含两个主题词，且非章节标题/关键词行
        title_para = None
        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text or len(text) > 50:
                continue
            if "雨水花园" not in text or "维护机制" not in text:
                continue
            if text.startswith("关键词") or text.replace(" ", "").lower().startswith("keywords"):
                continue
            if re.match(r'^第[一二三四五六七八九十]+章', text) or re.match(r'^[一二三四五六七八九十]+[、．.]', text):
                continue
            title_para = para
            break

        if title_para is None:
            self._add_result(item, False, 0)
            return

        para_elem = title_para._element
        checked_any = False
        all_ok = True
        for run in title_para.runs:
            if not run.text or not run.text.strip():
                continue
            checked_any = True
            font_name, size = self._resolve_run_cn_font(run._element, para_elem)
            bold = self._resolve_run_cn_bold(run._element, para_elem)
            # 点1：黑体
            if not (font_name and '黑体' in str(font_name)):
                all_ok = False
            # 点2：三号(16pt)
            if not (size is not None and abs(size - 16) < 0.5):
                all_ok = False
            # 点3：加粗
            if not bold:
                all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_thesis_title_paragraph(self):
        """+1: 论文题目段落格式为居中对齐、段前段后各2行、行距2倍

        细则三个点（全部满足才 +1）：
          1) 居中对齐
          2) 段前2行、段后2行
          3) 行距2倍
        论文题目为正文首页的完整题名（含"雨水花园"与"维护机制"两个主题词的短段落，
        排除章节标题、关键词行、正文段落）。对齐/间距/行距均按办公软件继承链解析
        (段落直接格式 -> 样式链)。段前段后"行"兼容 beforeLines/afterLines(1/100行) 与
        before/after(twips，按1行=240twips换算) 两种存法。
        """
        item = "+1: 论文题目段落格式(居中、段前段后各2行、两倍行距)"

        # 定位论文题目段落（与字体检查一致的判定）
        title_para = None
        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text or len(text) > 50:
                continue
            if "雨水花园" not in text or "维护机制" not in text:
                continue
            if text.startswith("关键词") or text.replace(" ", "").lower().startswith("keywords"):
                continue
            if re.match(r'^第[一二三四五六七八九十]+章', text) or re.match(r'^[一二三四五六七八九十]+[、．.]', text):
                continue
            title_para = para
            break

        if title_para is None:
            self._add_result(item, False, 0)
            return

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # ---- 点1：居中对齐 ----
        jc_val = None
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                jc_val = jc.get(qn('w:val'))
        if jc_val is None and style_id:
            jc_val = self._style_jc_lookup(style_id)
        is_center = jc_val == 'center'

        # 取有效 spacing：段落直接格式优先，否则沿样式链查找
        space = None
        if pPr is not None:
            space = pPr.find(qn('w:spacing'))
        if space is None and style_id:
            space = self._style_spacing_lookup(style_id)

        # ---- 点2：段前2行、段后2行（以办公软件显示的"行"为准）----
        has_before = self._spacing_lines_ok(pPr, style_id, 'before', 2.0)
        has_after = self._spacing_lines_ok(pPr, style_id, 'after', 2.0)

        # ---- 点3：行距2倍（办公软件"两倍行距"存为 lineRule=auto 且 line=480=2×240）----
        is_double = False
        if space is not None:
            line = space.get(qn('w:line'))
            rule = space.get(qn('w:lineRule'))
            if line and rule in (None, 'auto'):
                try:
                    is_double = abs(int(line) / 240.0 - 2.0) < 0.1
                except ValueError:
                    is_double = False

        passed = is_center and has_before and has_after and is_double
        self._add_result(item, passed, 1 if passed else 0)

    def _check_level1_title_font(self):
        """+3: 一级标题字体格式为小三号黑体

        细则两个点（所有一级标题的所有文字须全部满足才 +3）：
          1) 字体为黑体
          2) 字号为小三(15pt)
        一级标题 = 章节标题（"第X章"/"X、"）及"致谢"等与章同级的标题。
        字体/字号按办公软件继承链解析(run->字符样式->段落->段落样式->docDefaults)。
        """
        item = "+3: 一级标题字体格式为小三号黑体"
        all_ok = True
        found_any = False

        for para in self.document.paragraphs:
            text = para.text.strip()
            if not (re.match(r'^第[一二三四五六七八九十]+章', text)
                    or re.match(r'^[一二三四五六七八九十]+[、．.]', text)
                    or ('致' in text and '谢' in text and len(text) <= 10)):
                continue

            found_any = True
            para_elem = para._element
            for run in para.runs:
                if not run.text or not run.text.strip():
                    continue
                font_name, size = self._resolve_run_cn_font(run._element, para_elem)
                # 点1：黑体
                if not (font_name and '黑体' in str(font_name)):
                    all_ok = False
                # 点2：小三(15pt)
                if not (size is not None and abs(size - 15) < 0.5):
                    all_ok = False

        passed = found_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    def _check_level1_title_paragraph(self):
        """+3: 一级标题段落格式为居中对齐、段前段后各0.5行、两倍行距

        细则三个点（所有一级标题须全部满足才 +3）：
          1) 居中对齐
          2) 段前0.5行、段后0.5行
          3) 两倍行距
        一级标题 = 章节标题（"第X章"/"X、"）及"致谢"等与章同级的标题。
        对齐/间距/行距均按办公软件继承链解析(段落直接格式 -> 样式链)；
        段前段后"0.5行"兼容 beforeLines/afterLines(1/100行，0.5行=50) 与
        before/after(twips，0.5行=150twips=7.5pt) 两种存法。
        """
        item = "+3: 一级标题段落格式(居中、段前段后0.5行、两倍行距)"
        all_ok = True
        found_any = False

        for para in self.document.paragraphs:
            text = para.text.strip()
            if not (re.match(r'^第[一二三四五六七八九十]+章', text)
                    or re.match(r'^[一二三四五六七八九十]+[、．.]', text)
                    or ('致' in text and '谢' in text and len(text) <= 10)):
                continue

            found_any = True
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            # 点1：居中对齐
            jc_val = None
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(qn('w:val'))
            if jc_val is None and style_id:
                jc_val = self._style_jc_lookup(style_id)
            if jc_val != 'center':
                all_ok = False

            # 取有效 spacing：段落直接格式优先，否则沿样式链查找
            space = None
            if pPr is not None:
                space = pPr.find(qn('w:spacing'))
            if space is None and style_id:
                space = self._style_spacing_lookup(style_id)

            # 点2：段前0.5行、段后0.5行（以办公软件显示的"行"为准）
            if not self._spacing_lines_ok(pPr, style_id, 'before', 0.5):
                all_ok = False
            if not self._spacing_lines_ok(pPr, style_id, 'after', 0.5):
                all_ok = False

            # 点3：两倍行距（办公软件"两倍行距"存为 lineRule=auto 且 line=480=2×240）
            is_double = False
            if space is not None:
                line = space.get(qn('w:line'))
                rule = space.get(qn('w:lineRule'))
                if line and rule in (None, 'auto'):
                    try:
                        is_double = abs(int(line) / 240.0 - 2.0) < 0.1
                    except ValueError:
                        is_double = False
            if not is_double:
                all_ok = False

        passed = found_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    def _check_level2_title_font(self):
        """+3: 二级标题字体格式为四号黑体、加粗

        细则三个点（所有二级标题的所有文字须全部满足才 +3）：
          1) 字体为黑体
          2) 字号为四号(14pt)
          3) 加粗
        二级标题 = "X.Y" 或 "（一）"类小节标题。
        字体/字号/加粗均按办公软件继承链解析(run->字符样式->段落->段落样式->docDefaults)。
        """
        item = "+3: 二级标题字体格式为四号黑体、加粗"
        found_any = False
        all_ok = True

        for para in self.document.paragraphs:
            text = para.text.strip()
            if not (re.match(r'^\d+\.\d+', text)
                    or re.match(r'^[（(][一二三四五六七八九十]+[)）]', text)):
                continue

            found_any = True
            para_elem = para._element
            for run in para.runs:
                if not run.text or not run.text.strip():
                    continue
                font_name, size = self._resolve_run_cn_font(run._element, para_elem)
                bold = self._resolve_run_cn_bold(run._element, para_elem)
                # 点1：黑体
                if not (font_name and '黑体' in str(font_name)):
                    all_ok = False
                # 点2：四号(14pt)
                if not (size is not None and abs(size - 14) < 0.5):
                    all_ok = False
                # 点3：加粗
                if not bold:
                    all_ok = False

        passed = found_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    def _check_level2_title_paragraph(self):
        """+3: 二级标题段落格式为顶格左对齐、段前段后各0.5行、单倍行距

        细则三个点（所有二级标题须全部满足才 +3）：
          1) 顶格左对齐（左对齐 且 无左缩进=顶格）
          2) 段前0.5行、段后0.5行
          3) 单倍行距
        二级标题 = "X.Y" 或 "（一）"类小节标题。
        对齐/缩进/间距/行距均按办公软件继承链解析(段落直接格式 -> 样式链)；
        段前段后"0.5行"兼容 beforeLines/afterLines(1/100行，0.5行=50) 与
        before/after(twips，0.5行=150twips=7.5pt) 两种存法。
        """
        item = "+3: 二级标题段落格式(顶格左对齐、段前段后0.5行、单倍行距)"
        found_any = False
        all_ok = True

        for para in self.document.paragraphs:
            text = para.text.strip()
            if not (re.match(r'^\d+\.\d+', text)
                    or re.match(r'^[（(][一二三四五六七八九十]+[)）]', text)):
                continue

            found_any = True
            pPr = para._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            # 点1：顶格左对齐 = 左对齐(left/start) 且 无左缩进(顶格)
            jc_val = None
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(qn('w:val'))
            if jc_val is None and style_id:
                jc_val = self._style_jc_lookup(style_id)
            # 办公软件左对齐存为 left/start；未设置(None)在默认样式下也按左对齐处理
            is_left = jc_val in ('left', 'start', None)
            # 顶格：有效左缩进为 0 字符
            left_chars = self._para_left_indent_chars(para._element, pPr, style_id)
            is_flush = abs(left_chars) < 0.1
            if not (is_left and is_flush):
                all_ok = False

            # 取有效 spacing：段落直接格式优先，否则沿样式链查找
            space = None
            if pPr is not None:
                space = pPr.find(qn('w:spacing'))
            if space is None and style_id:
                space = self._style_spacing_lookup(style_id)

            # 点2：段前0.5行、段后0.5行（以办公软件显示的"行"为准）
            if not self._spacing_lines_ok(pPr, style_id, 'before', 0.5):
                all_ok = False
            if not self._spacing_lines_ok(pPr, style_id, 'after', 0.5):
                all_ok = False

            # 点3：单倍行距（办公软件"单倍"存为 lineRule=auto 且 line=240=1×240，或未设行距）
            if not self._para_is_single_spacing(pPr, style_id):
                all_ok = False

        passed = found_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    def _check_reference_title_font(self):
        """+1: 参考文献页标题字体格式为小三号、黑体

        细则两个点（标题所有文字须全部满足才 +1）：
          1) 字体为黑体
          2) 字号为小三(15pt)
        仅约束"参考文献"标题这一行。字体/字号按办公软件继承链解析。
        """
        item = "+1: 参考文献页标题字体格式为小三号黑体"
        para, _ = self._find_paragraph_by_text("参考文献", exact=True)
        if para is None:
            para, _ = self._find_paragraph_by_text("参考文献")
        if para is None:
            self._add_result(item, False, 0)
            return

        para_elem = para._element
        checked_any = False
        all_ok = True
        for run in para.runs:
            if not run.text or not run.text.strip():
                continue
            checked_any = True
            font_name, size = self._resolve_run_cn_font(run._element, para_elem)
            # 点1：黑体
            if not (font_name and '黑体' in str(font_name)):
                all_ok = False
            # 点2：小三(15pt)
            if not (size is not None and abs(size - 15) < 0.5):
                all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_reference_title_paragraph(self):
        """+1: 参考文献页标题段落格式为居中对齐、段前段后各0.5行、2倍行距

        细则三个点（全部满足才 +1）：
          1) 居中对齐
          2) 段前0.5行、段后0.5行
          3) 2倍行距
        仅约束"参考文献"标题这一行。对齐/间距/行距按办公软件渲染的属性级合并解析：
        每个属性优先取段落直接格式，缺失再回退样式链；
        段前段后"0.5行"兼容 beforeLines/afterLines(1/100行，0.5行=50) 与
        before/after(twips，0.5行=150twips=7.5pt) 两种存法。
        """
        item = "+1: 参考文献页标题段落格式(居中、段前段后0.5行、两倍行距)"
        para, _ = self._find_paragraph_by_text("参考文献", exact=True)
        if para is None:
            para, _ = self._find_paragraph_by_text("参考文献")
        if para is None:
            self._add_result(item, False, 0)
            return

        pPr = para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # 段落直接 spacing 与样式链 spacing，用于属性级合并（直接格式优先，缺失回退样式）
        direct_sp = pPr.find(qn('w:spacing')) if pPr is not None else None
        style_sp = self._style_spacing_lookup(style_id) if style_id else None

        def eff_attr(attr):
            if direct_sp is not None and direct_sp.get(qn(attr)) is not None:
                return direct_sp.get(qn(attr))
            if style_sp is not None and style_sp.get(qn(attr)) is not None:
                return style_sp.get(qn(attr))
            return None

        # ---- 点1：居中对齐 ----
        jc_val = None
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                jc_val = jc.get(qn('w:val'))
        if jc_val is None and style_id:
            jc_val = self._style_jc_lookup(style_id)
        is_center = jc_val == 'center'

        # ---- 点2：段前0.5行、段后0.5行（以办公软件显示的"行"为准）----
        has_before = self._spacing_lines_ok(pPr, style_id, 'before', 0.5)
        has_after = self._spacing_lines_ok(pPr, style_id, 'after', 0.5)

        # ---- 点3：2倍行距（办公软件"两倍行距"存为 lineRule=auto 且 line=480=2×240）----
        is_double = False
        line = eff_attr('w:line')
        rule = eff_attr('w:lineRule')
        if line and rule in (None, 'auto'):
            try:
                is_double = abs(int(line) / 240.0 - 2.0) < 0.1
            except ValueError:
                is_double = False

        passed = is_center and has_before and has_after and is_double
        self._add_result(item, passed, 1 if passed else 0)

    def _check_reference_content_font(self):
        """+1: 参考文献页内容字体格式：中文用宋体、小四、黑色；英文用 Times New Roman、小四、黑色

        细则六个点（所有参考文献条目须全部满足才 +1）：
          中文：1) 宋体  2) 小四(12pt)  3) 黑色
          英文：4) Times New Roman  5) 小四(12pt)  6) 黑色
        黑色 = 未设颜色(默认自动)或 000000/black/auto。字体/字号/颜色按办公软件继承链解析。
        每个 run 按其字符类型(含中文汉字→中文规则，否则→英文规则)分别校验。
        """
        item = "+1: 参考文献页内容字体格式(中文宋体小四黑色、英文Times New Roman小四黑色)"
        para, idx = self._find_paragraph_by_text("参考文献", exact=True)
        if para is None:
            para, idx = self._find_paragraph_by_text("参考文献")
        if para is None:
            self._add_result(item, False, 0)
            return

        def is_black(color):
            # 未设颜色(None) => 办公软件默认自动/黑色；显式 auto/000000/black 也算黑
            if color is None:
                return True
            c = str(color).strip().lower()
            return c in ('auto', '000000', 'black')

        def has_cjk(s):
            return any('一' <= ch <= '鿿' for ch in s)

        checked_any = False
        all_ok = True
        for i in range(idx + 1, min(idx + 40, len(self.document.paragraphs))):
            p = self.document.paragraphs[i]
            text = p.text.strip()
            if not text:
                continue
            # 参考文献条目：以数字或 [n] 开头；遇到非条目行即视为参考文献区结束
            if not (text[0].isdigit() or text.startswith('[')):
                break

            para_elem = p._element
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                checked_any = True
                cn_font, size = self._resolve_run_cn_font(run._element, para_elem)
                en_font = self._resolve_run_en_font(run._element, para_elem)
                color = self._resolve_run_color(run._element, para_elem)

                # 点(共用)：小四(12pt)
                if not (size is not None and abs(size - 12) < 0.5):
                    all_ok = False
                # 点(共用)：黑色
                if not is_black(color):
                    all_ok = False

                if has_cjk(run.text):
                    # 中文点：宋体
                    if not (cn_font and '宋体' in str(cn_font)):
                        all_ok = False
                else:
                    # 英文点：Times New Roman
                    if not (en_font and 'Times New Roman' in str(en_font)):
                        all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_footer_margin(self):
        """+1: 页脚下边距为1.50厘米

        细则一个点：页脚下边距(页脚到页面底边的距离) = 1.50 厘米。
        办公软件将该值存于节属性 w:sectPr/w:pgMar/@w:footer(单位 twips，1cm=567twips，
        1.5cm≈850twips)，即 python-docx 的 section.footer_distance。
        文档所有节都须满足才 +1。
        """
        item = "+1: 页脚下边距为1.50厘米"
        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return

        if not sections:
            self._add_result(item, False, 0)
            return

        all_ok = True
        checked_any = False
        for section in sections:
            fd = section.footer_distance
            if fd is None:
                # 未显式设置页脚下边距 => 不能确认为1.5cm
                all_ok = False
                continue
            checked_any = True
            # 允许 ±0.05cm 误差（1.5cm 实际存为850twips=1.4993cm）
            if abs(fd.cm - 1.5) > 0.05:
                all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 1 if passed else 0)

    def _check_header_content(self):
        """+3: 页眉文本内容为"东岳技术学院高等教育自学考试毕业论文"或"东岳技术学院高等教育自学考试毕业调研报告"，论文所有页面都有页眉

        细则两个点（全部满足才 +3）：
          1) 页眉文本内容为下列任一目标文本之一：
             - "东岳技术学院高等教育自学考试毕业论文"
             - "东岳技术学院高等教育自学考试毕业调研报告"
          2) 论文所有页面都有页眉
        办公软件中页眉按节存储，每节可有独立/首页/偶数页页眉；"所有页面都有页眉"
        要求每一节的生效页眉都含目标文本(含首页页眉，若启用了首页不同)。
        """
        item = '+3: 页眉文本内容为"东岳技术学院高等教育自学考试毕业论文"或"东岳技术学院高等教育自学考试毕业调研报告"，论文所有页面都有页眉'
        targets_norm = (
            "东岳技术学院高等教育自学考试毕业论文".replace(" ", ""),
            "东岳技术学院高等教育自学考试毕业调研报告".replace(" ", ""),
        )

        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return

        if not sections:
            self._add_result(item, False, 0)
            return

        def header_text(hdr):
            if hdr is None:
                return ""
            return "".join(p.text for p in hdr.paragraphs).replace(" ", "").replace("　", "")

        all_ok = True
        for section in sections:
            # 主页眉必须存在且文本匹配任一目标
            main_ok = header_text(section.header) in targets_norm
            if not main_ok:
                all_ok = False
                break
            # 若该节启用"首页不同"，首页页眉也须匹配（否则首页无正确页眉）
            if section.different_first_page_header_footer:
                if header_text(section.first_page_header) not in targets_norm:
                    all_ok = False
                    break

        self._add_result(item, all_ok, 3 if all_ok else 0)

    def _check_header_font(self):
        """+3: 页眉字体格式为 Noto Serif CJK SC、五号、居中对齐

        细则三个点（页眉所有文字须全部满足才 +3）：
          1) 字体为 Noto Serif CJK SC
          2) 字号为五号(10.5pt)
          3) 居中对齐
        页眉按节存储；字体/字号按办公软件继承链解析，对齐按 段落直接 jc -> 样式链 解析。
        文档所有节的页眉都须满足。
        """
        item = "+3: 页眉字体格式为Noto Serif CJK SC、五号、居中对齐"
        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return

        if not sections:
            self._add_result(item, False, 0)
            return

        def has_cjk(s):
            return any('一' <= ch <= '鿿' for ch in s)

        checked_any = False
        all_ok = True

        for section in sections:
            header = section.header
            if header is None:
                continue
            for para in header.paragraphs:
                if not para.text.strip():
                    continue

                para_elem = para._element
                pPr = para_elem.find(qn('w:pPr'))
                pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
                style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

                # 点3：居中对齐（段落直接 jc 优先，否则样式链）
                jc_val = None
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None:
                        jc_val = jc.get(qn('w:val'))
                if jc_val is None and style_id:
                    jc_val = self._style_jc_lookup(style_id)
                if jc_val != 'center':
                    all_ok = False

                for run in para.runs:
                    if not run.text or not run.text.strip():
                        continue
                    checked_any = True
                    cn_font, size = self._resolve_run_cn_font(run._element, para_elem)
                    en_font = self._resolve_run_en_font(run._element, para_elem)
                    # 该 run 对应的生效字体：含中文取中文字体，否则取英文字体
                    font_name = cn_font if has_cjk(run.text) else en_font
                    # 点1：Noto Serif CJK SC
                    if not (font_name and 'Noto Serif CJK SC' in str(font_name)):
                        all_ok = False
                    # 点2：五号(10.5pt)
                    if not (size is not None and abs(size - 10.5) < 0.5):
                        all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    def _check_reference_mark_format(self):
        """+3: 参考文献标注格式：在引文或转述观点的最后一个句号前，以黑色上标的形式加以标注

        细则三个点（正文中所有引用标注须全部满足才 +3）：
          1) 黑色
          2) 上标形式
          3) 位于所在句最后一个句号前（即标注紧接在句末句号"。"之前）
        标注在办公软件中为形如 [n] 的 run；上标存于 rPr/w:vertAlign=superscript，
        颜色存于 rPr/w:color(未设或000000/black/auto 视为黑)。仅检查参考文献列表之前的正文引用。
        """
        item = "+3: 参考文献标注格式(句末句号前、黑色上标)"

        # 参考文献列表起始位置，只检查其之前的正文引用
        ref_idx = -1
        for i, para in enumerate(self.document.paragraphs):
            if para.text.strip() == '参考文献':
                ref_idx = i
                break

        def is_black(color):
            if color is None:
                return True
            c = str(color).strip().lower()
            return c in ('auto', '000000', 'black')

        checked_any = False
        all_ok = True

        for i, para in enumerate(self.document.paragraphs):
            if ref_idx >= 0 and i >= ref_idx:
                break

            full = para.text
            if not re.search(r'\[\d+\]', full):
                continue

            para_elem = para._element
            offset = 0
            for run in para.runs:
                r_text = run.text or ""
                r_end = offset + len(r_text)
                offset = r_end
                if not re.search(r'\[\d+\]', r_text):
                    continue

                checked_any = True

                # 点1：黑色
                color = self._resolve_run_color(run._element, para_elem)
                if not is_black(color):
                    all_ok = False

                # 点2：上标
                rPr = run._element.find(qn('w:rPr'))
                is_super = False
                if rPr is not None:
                    va = rPr.find(qn('w:vertAlign'))
                    if va is not None and va.get(qn('w:val')) == 'superscript':
                        is_super = True
                if not is_super:
                    all_ok = False

                # 点3：位于句末句号前（该标注 run 之后、跳过空白后的首个非空字符为"。"）
                after = full[r_end:]
                nxt = after.lstrip(' 　\t')
                if not (nxt.startswith('。') or nxt.startswith('．')):
                    all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    def _check_reference_list_format(self):
        """+3: 论文中参考文献引用格式：序号与文献间空一个汉字符；顶格左对齐、1.5倍行距

        细则三个点（所有参考文献条目须全部满足才 +3）：
          1) 序号与文献之间空一个汉字符（全角空格'　'算1字符；两个半角空格算1字符）
          2) 顶格左对齐（左对齐 且 无左缩进）
          3) 1.5倍行距
        对齐/缩进/行距按办公软件继承链解析(段落直接格式 -> 样式链)。
        """
        item = "+3: 参考文献引用格式(序号后空一汉字符、顶格左对齐、1.5倍行距)"
        para, idx = self._find_paragraph_by_text("参考文献", exact=True)
        if para is None:
            para, idx = self._find_paragraph_by_text("参考文献")
        if para is None:
            self._add_result(item, False, 0)
            return

        def sep_chars(space_str):
            # 全角空格=1字符，半角空格=0.5字符
            n = 0.0
            for ch in space_str:
                if ch == '　':
                    n += 1.0
                elif ch == ' ':
                    n += 0.5
                else:
                    break
            return n

        checked_any = False
        all_ok = True
        for i in range(idx + 1, min(idx + 40, len(self.document.paragraphs))):
            p = self.document.paragraphs[i]
            text = p.text
            stripped = text.strip()
            if not stripped:
                continue
            # 条目：以 [n] 或 n. / n、 开头；遇非条目行视为参考文献区结束
            m = re.match(r'^\s*(\[\d+\]|\d+[\.、])(\s*)', text)
            if not m:
                break

            checked_any = True
            pPr = p._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            # 点1：序号与文献间空一个汉字符
            if abs(sep_chars(m.group(2)) - 1.0) > 0.1:
                all_ok = False

            # 点2：顶格左对齐 = 左对齐(left/start) 且 无左缩进
            jc_val = None
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(qn('w:val'))
            if jc_val is None and style_id:
                jc_val = self._style_jc_lookup(style_id)
            is_left = jc_val in ('left', 'start', None)
            left_chars = self._para_left_indent_chars(p._element, pPr, style_id)
            if not (is_left and abs(left_chars) < 0.1):
                all_ok = False

            # 点3：1.5倍行距（办公软件存为 lineRule=auto 且 line=360=1.5×240）
            space = None
            if pPr is not None:
                space = pPr.find(qn('w:spacing'))
            if space is None and style_id:
                space = self._style_spacing_lookup(style_id)
            is_15 = False
            if space is not None:
                line = space.get(qn('w:line'))
                rule = space.get(qn('w:lineRule'))
                if line and rule in (None, 'auto'):
                    try:
                        is_15 = abs(int(line) / 240.0 - 1.5) < 0.1
                    except ValueError:
                        is_15 = False
            if not is_15:
                all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    def _check_body_font_paragraph(self):
        """+3: 正文段落格式为首行缩进两字符、1.5倍行距、两端对齐

        细则三个点（所有正文段落须全部满足才 +3）：
          1) 首行缩进两字符
          2) 1.5倍行距
          3) 两端对齐
        正文 = 第一个章节标题（"X、"/"第X章"）之后、"参考文献"之前的章节正文段落，
        排除各级标题行本身；不含封面/声明/摘要等前置部分。
        缩进/行距/对齐均按办公软件继承链解析(段落直接格式 -> 样式链)。
        """
        item = "+3: 正文段落格式(首行缩进两字符、1.5倍行距、两端对齐)"
        paras = self.document.paragraphs

        # 正文范围：首个一级标题之后 到 参考文献之前
        start = None
        end = len(paras)
        for i, p in enumerate(paras):
            t = p.text.strip()
            if start is None and (re.match(r'^第[一二三四五六七八九十]+章', t)
                                  or re.match(r'^[一二三四五六七八九十]+[、．.]', t)):
                start = i
            if t == '参考文献':
                end = i
                break
        if start is None:
            self._add_result(item, False, 0)
            return

        def is_title(t):
            return (re.match(r'^第[一二三四五六七八九十]+章', t)
                    or re.match(r'^[一二三四五六七八九十]+[、．.]', t)
                    or re.match(r'^\d+\.\d+', t)
                    or re.match(r'^[（(][一二三四五六七八九十]+[)）]', t))

        checked_any = False
        all_ok = True
        for i in range(start, end):
            p = paras[i]
            text = p.text.strip()
            if not text:
                continue
            if is_title(text):   # 排除各级标题行，只留正文段落
                continue

            checked_any = True
            pPr = p._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            style_id = pStyle.get(qn('w:val')) if pStyle is not None else None
            # 段落未显式挂 pStyle 时，办公软件按"默认段落样式"(通常是 Normal)渲染，
            # 样式链回溯须从默认段落样式开始，才能与显示一致。
            if style_id is None:
                style_id = self._default_paragraph_style_id()

            # 点1：首行缩进两字符
            chars = self._para_first_line_indent_chars(pPr, style_id)
            if abs(chars - 2.0) > 0.4:
                all_ok = False

            # 点2：1.5倍行距（办公软件存为 lineRule=auto 且 line=360=1.5×240）
            space = None
            if pPr is not None:
                space = pPr.find(qn('w:spacing'))
            if space is None and style_id:
                space = self._style_spacing_lookup(style_id)
            is_15 = False
            if space is not None:
                line = space.get(qn('w:line'))
                rule = space.get(qn('w:lineRule'))
                if line and rule in (None, 'auto'):
                    try:
                        is_15 = abs(int(line) / 240.0 - 1.5) < 0.1
                    except ValueError:
                        is_15 = False
            if not is_15:
                all_ok = False

            # 点3：两端对齐（办公软件"两端对齐"存为 jc=both/justify）
            jc_val = None
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(qn('w:val'))
            if jc_val is None and style_id:
                jc_val = self._style_jc_lookup(style_id)
            if jc_val not in ('both', 'justify'):
                all_ok = False

        passed = checked_any and all_ok
        self._add_result(item, passed, 3 if passed else 0)

    # ==================== 扣分点检查方法 ====================

    def _check_cover_title_font_penalty(self):
        """-1: "高等教育自学考试毕业论文封面"页标题字体不满足宋体、二号、加粗

        细则三个点（标题应同时满足；任一不满足即扣分）：
          1) 宋体  2) 二号(22pt)  3) 加粗
        封面标题 = 封面首行"…自学考试…封面"标题。字体/字号/加粗按办公软件继承链解析。
        仅当标题存在且"未全部满足"时扣 -1；标题本身找不到则不扣分。
        """
        item = "-1: 封面页标题字体不满足宋体、二号、加粗"

        title_para = None
        for para in self.document.paragraphs[:10]:
            t = para.text.strip()
            if "自学考试" in t and "封面" in t:
                title_para = para
                break
        if title_para is None:
            self._add_result(item, False, 0)
            return

        para_elem = title_para._element
        checked_any = False
        satisfied = True   # 是否满足"宋体二号加粗"
        for run in title_para.runs:
            if not run.text or not run.text.strip():
                continue
            checked_any = True
            font_name, size = self._resolve_run_cn_font(run._element, para_elem)
            bold = self._resolve_run_cn_bold(run._element, para_elem)
            # 点1：宋体
            if not (font_name and '宋体' in str(font_name)):
                satisfied = False
            # 点2：二号(22pt)
            if not (size is not None and abs(size - 22) < 0.5):
                satisfied = False
            # 点3：加粗
            if not bold:
                satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_cover_title_paragraph_penalty(self):
        """-1: "高等教育自学考试毕业论文封面"页标题段落格式不满足居中对齐、行距固定值29磅

        细则两个点（标题应同时满足；任一不满足即扣 -1）：
          1) 居中对齐
          2) 行距为固定值29磅
        办公软件"固定值行距"存为 w:spacing/@w:lineRule="exact" 且 @w:line 为磅×20 的 twips
        (29磅=580twips)。对齐/行距按 段落直接格式 -> 样式链 解析。
        仅当标题存在且"未全部满足"时扣分；标题找不到则不扣分。
        """
        item = "-1: 封面页标题段落格式不满足居中对齐、行距固定值29磅"

        title_para = None
        for para in self.document.paragraphs[:10]:
            t = para.text.strip()
            if "自学考试" in t and "封面" in t:
                title_para = para
                break
        if title_para is None:
            self._add_result(item, False, 0)
            return

        pPr = title_para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # 点1：居中对齐
        jc_val = None
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                jc_val = jc.get(qn('w:val'))
        if jc_val is None and style_id:
            jc_val = self._style_jc_lookup(style_id)
        is_center = jc_val == 'center'

        # 点2：行距固定值29磅（lineRule=exact 且 line≈580twips=29pt）
        space = None
        if pPr is not None:
            space = pPr.find(qn('w:spacing'))
        if space is None and style_id:
            space = self._style_spacing_lookup(style_id)
        is_fixed29 = False
        if space is not None:
            line = space.get(qn('w:line'))
            rule = space.get(qn('w:lineRule'))
            if rule == 'exact' and line:
                try:
                    is_fixed29 = abs(int(line) / 20.0 - 29.0) < 0.5
                except ValueError:
                    is_fixed29 = False

        # 未全部满足 => 扣分
        penalize = not (is_center and is_fixed29)
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_cover_other_font_penalty(self):
        """-3: 除"市地：…"这一行文本、"评判项目"、"评判标准"和"得分"外，封面页其余字体格式不满足宋体小四

        细则的点（封面页"其余"文字须满足宋体+小四(12pt)；任一不满足即扣 -3）：
          1) 宋体   2) 小四(12pt)
        封面页 = "学术诚信声明"之前的正文段落 + 封面首个表格（含个人信息表与评分表）。
        "其余"排除四类不计入的文字：
          - 封面大标题(另有字体要求)
          - "市地：… 20XX年 X月"这一行(该行另有段落格式扣分项覆盖，字体不在此计)
          - 表格中的"评判项目"/"评判标准"/"得分"三个表头(另有黑体要求)
        字体/字号按办公软件继承链解析。
        """
        item = "-3: 封面页其余字体格式不满足宋体小四"

        paras = self.document.paragraphs
        # 封面页范围：文首至"学术诚信声明"之前
        cover_end = len(paras)
        for i, p in enumerate(paras):
            if "学术诚信声明" in p.text:
                cover_end = i
                break

        def font_ok(run, para_elem):
            # 返回 (是否为宋体小四) —— 字号 12pt 且 生效字体为宋体
            # 生效字体按字符类型区分：含汉字的 run 取中文字体(eastAsia)，
            # 纯ASCII(数字/字母)的 run 取英文字体(ascii/hAnsi) —— 与办公软件渲染一致，
            # 否则会漏判"eastAsia=宋体 但 ascii=Calibri"这类中文标注宋体、数字实际渲染为
            # Calibri 的常见情况(如手机号/准考证号)。
            cn_font, size = self._resolve_run_cn_font(run._element, para_elem)
            en_font = self._resolve_run_en_font(run._element, para_elem)
            has_cjk = any('一' <= ch <= '鿿' for ch in run.text)
            eff_font = cn_font if has_cjk else en_font
            is_song = bool(eff_font and '宋体' in str(eff_font))
            is_xiao4 = size is not None and abs(size - 12) < 0.5
            return is_song and is_xiao4

        satisfied = True   # 其余文字是否都满足宋体小四
        checked_any = False

        # 1) 封面段落（排除大标题、"市地：…"这一行）
        for i in range(cover_end):
            p = paras[i]
            t = p.text.strip()
            if not t:
                continue
            if "自学考试" in t and "封面" in t:  # 封面大标题，另有字体要求，排除
                continue
            if t.startswith("市地"):  # "市地：… 20XX年 X月"这一行，另有段落格式扣分覆盖，字体在此不计
                continue
            para_elem = p._element
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                checked_any = True
                if not font_ok(run, para_elem):
                    satisfied = False

        # 2) 封面表格（排除"评判项目/评判标准/得分"三个表头）
        exclude_labels = ('评判项目', '评判标准', '得分')
        if self.document.tables:
            table = self.document.tables[0]
            seen = set()
            for row in table.rows:
                for cell in row.cells:
                    key = id(cell._tc)
                    if key in seen:
                        continue
                    seen.add(key)
                    cell_text = cell.text.strip()
                    if cell_text in exclude_labels:
                        continue
                    for p in cell.paragraphs:
                        para_elem = p._element
                        for run in p.runs:
                            if not run.text or not run.text.strip():
                                continue
                            checked_any = True
                            if not font_ok(run, para_elem):
                                satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -3 if penalize else 0)

    def _check_cover_table_font_penalty(self):
        """-1: 封面页"评判项目"、"评判标准"和"得分"字体不满足黑体小四号

        细则的点（这三个表头须满足黑体+小四(12pt)；任一不满足即扣 -1）：
          1) 黑体   2) 小四(12pt)
        仅约束封面表格中文本恰为"评判项目""评判标准""得分"的三个表头单元格。
        字体/字号按办公软件继承链解析。找不到这三个表头则不扣分。
        """
        item = "-1: 封面页表格特定项字体不满足黑体小四号"
        labels = ('评判项目', '评判标准', '得分')

        if not self.document.tables:
            self._add_result(item, False, 0)
            return

        checked_any = False
        satisfied = True   # 三个表头是否都满足黑体小四
        table = self.document.tables[0]
        seen = set()
        for row in table.rows:
            for cell in row.cells:
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key)
                if cell.text.strip() not in labels:
                    continue
                for para in cell.paragraphs:
                    para_elem = para._element
                    for run in para.runs:
                        if not run.text or not run.text.strip():
                            continue
                        checked_any = True
                        font_name, size = self._resolve_run_cn_font(run._element, para_elem)
                        # 点1：黑体
                        if not (font_name and '黑体' in str(font_name)):
                            satisfied = False
                        # 点2：小四(12pt)
                        if not (size is not None and abs(size - 12) < 0.5):
                            satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_cover_table_align_penalty(self):
        """-3: 封面页表格第六行第二列、第七行第二列、第八行第二列、第十二行单元格内容不满足左对齐、单倍行距

        细则两个点（这些指定单元格须同时满足；任一不满足即扣 -3）：
          1) 左对齐   2) 单倍行距
        指定单元格：第6/7/8行的第2列，以及第12行(整行)。
        对齐按 段落直接 jc -> 样式链 解析(left/start 视为左对齐)；单倍行距用继承链解析。
        表格不足12行则不扣分。
        """
        item = "-3: 封面页表格特定单元格不满足左对齐、单倍行距"

        if not self.document.tables:
            self._add_result(item, False, 0)
            return

        table = self.document.tables[0]
        rows = list(table.rows)
        if len(rows) < 12:
            self._add_result(item, False, 0)
            return

        # 收集目标单元格（去重合并单元格）：6/7/8行第2列 + 第12行整行
        target_cells = []
        seen = set()

        def add_cell(cell):
            key = id(cell._tc)
            if key not in seen:
                seen.add(key)
                target_cells.append(cell)

        for row_idx in (5, 6, 7):          # 第6/7/8行 → 索引5/6/7
            row = rows[row_idx]
            if len(row.cells) >= 2:
                add_cell(row.cells[1])     # 第2列 → 索引1
        for cell in rows[11].cells:        # 第12行整行 → 索引11
            add_cell(cell)

        checked_any = False
        satisfied = True   # 目标单元格是否都满足左对齐+单倍行距
        for cell in target_cells:
            for para in cell.paragraphs:
                if not para.text.strip():
                    continue
                checked_any = True
                pPr = para._element.find(qn('w:pPr'))
                pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
                style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

                # 点1：左对齐
                jc_val = None
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None:
                        jc_val = jc.get(qn('w:val'))
                if jc_val is None and style_id:
                    jc_val = self._style_jc_lookup(style_id)
                if jc_val not in ('left', 'start', None):
                    satisfied = False

                # 点2：单倍行距
                if not self._para_is_single_spacing(pPr, style_id):
                    satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -3 if penalize else 0)

    def _check_cover_address_penalty(self):
        """-1: 封面页地址和日期处字体段落格式不满足两端对齐、首行缩进两字符、行距29磅

        细则三个点（地址日期段落须同时满足；任一不满足即扣 -1）：
          1) 两端对齐   2) 首行缩进两字符   3) 行距固定值29磅
        地址日期 = 封面页(在"学术诚信声明"之前)含"年…月…日"的那一行。
        对齐/缩进/行距按办公软件继承链解析；行距29磅存为 lineRule=exact 且 line=580twips。
        找不到该行则不扣分。
        """
        item = "-1: 封面页地址和日期不满足两端对齐、首行缩进两字符、行距29磅"

        paras = self.document.paragraphs
        cover_end = len(paras)
        for i, p in enumerate(paras):
            if "学术诚信声明" in p.text:
                cover_end = i
                break

        addr_para = None
        for i in range(cover_end):
            t = paras[i].text.strip()
            if "年" in t and "月" in t and "日" in t:
                addr_para = paras[i]
                break
        if addr_para is None:
            self._add_result(item, False, 0)
            return

        pPr = addr_para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None
        # 段落未显式挂 pStyle 时，办公软件按"默认段落样式"(通常是 Normal)渲染，
        # 样式链回溯须从默认段落样式开始，才能与显示一致(否则 Normal 里的 jc=both 查不到)。
        if style_id is None:
            style_id = self._default_paragraph_style_id()

        # 点1：两端对齐
        jc_val = None
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                jc_val = jc.get(qn('w:val'))
        if jc_val is None and style_id:
            jc_val = self._style_jc_lookup(style_id)
        is_justify = jc_val in ('both', 'justify')

        # 点2：首行缩进两字符
        chars = self._para_first_line_indent_chars(pPr, style_id)
        is_indent2 = abs(chars - 2.0) <= 0.4

        # 点3：行距固定值29磅（lineRule=exact 且 line≈580twips=29pt）
        space = None
        if pPr is not None:
            space = pPr.find(qn('w:spacing'))
        if space is None and style_id:
            space = self._style_spacing_lookup(style_id)
        is_29 = False
        if space is not None:
            line = space.get(qn('w:line'))
            rule = space.get(qn('w:lineRule'))
            if rule == 'exact' and line:
                try:
                    is_29 = abs(int(line) / 20.0 - 29.0) < 0.5
                except ValueError:
                    is_29 = False

        # 未全部满足 => 扣分
        penalize = not (is_justify and is_indent2 and is_29)
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_statement_title_penalty(self):
        """-1: "学术诚信声明及作品使用授权书"页标题字体不满足黑体三号加粗

        细则三个点（标题应同时满足；任一不满足即扣 -1）：
          1) 黑体   2) 三号(16pt)   3) 加粗
        仅约束"学术诚信声明…"标题这一行。字体/字号/加粗按办公软件继承链解析。
        找不到标题则不扣分。
        """
        item = "-1: 学术诚信声明标题字体不满足黑体三号加粗"
        para, _ = self._find_paragraph_by_text("学术诚信声明")
        if para is None:
            self._add_result(item, False, 0)
            return

        para_elem = para._element
        checked_any = False
        satisfied = True   # 是否满足黑体+三号+加粗
        for run in para.runs:
            if not run.text or not run.text.strip():
                continue
            checked_any = True
            font_name, size = self._resolve_run_cn_font(run._element, para_elem)
            bold = self._resolve_run_cn_bold(run._element, para_elem)
            # 点1：黑体
            if not (font_name and '黑体' in str(font_name)):
                satisfied = False
            # 点2：三号(16pt)
            if not (size is not None and abs(size - 16) < 0.5):
                satisfied = False
            # 点3：加粗
            if not bold:
                satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_statement_paragraph_penalty(self):
        """-1: "学术诚信声明及作品使用授权书"标题段落格式不满足居中对齐、段前一行、段后一行、单倍行距

        细则四个点（标题段落应同时满足；任一不满足即扣 -1）：
          1) 居中对齐   2) 段前一行   3) 段后一行   4) 单倍行距
        对齐/间距/行距按办公软件继承链解析。段前段后"一行"兼容两种存法：
          beforeLines/afterLines(1/100行，1行=100) 或 before/after(twips，按标题字号换算 1行=字号pt×20)。
        """
        item = "-1: 学术诚信声明段落格式不满足居中对齐、段前段后一行、单倍行距"
        para, _ = self._find_paragraph_by_text("学术诚信声明")
        if para is None:
            self._add_result(item, False, 0)
            return

        para_elem = para._element
        pPr = para_elem.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # 点1：居中对齐
        jc_val = None
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                jc_val = jc.get(qn('w:val'))
        if jc_val is None and style_id:
            jc_val = self._style_jc_lookup(style_id)
        is_center = jc_val == 'center'

        # 取有效 spacing 与标题字号（用于 twips→行 换算）
        space = None
        if pPr is not None:
            space = pPr.find(qn('w:spacing'))
        if space is None and style_id:
            space = self._style_spacing_lookup(style_id)

        # 点2、点3：段前一行、段后一行（以办公软件显示的"行"为准）
        has_before = self._spacing_lines_ok(pPr, style_id, 'before', 1.0)
        has_after = self._spacing_lines_ok(pPr, style_id, 'after', 1.0)

        # 点4：单倍行距
        is_single = self._para_is_single_spacing(pPr, style_id)

        # 未全部满足 => 扣分
        penalize = not (is_center and has_before and has_after and is_single)
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_statement_font_penalty(self):
        """-1: "学术诚信声明及作品使用授权书"页内容字体格式不满足宋体小四号

        细则两个点（内容文字须同时满足；任一不满足即扣 -1）：
          1) 宋体   2) 小四(12pt)
        内容 = 声明页标题之后、下一页(目录)之前的正文段落（排除标题本身）。
        字体/字号按办公软件继承链解析。找不到内容则不扣分。
        """
        item = '-1: "学术诚信声明及作品使用授权书"页内容字体格式不满足宋体小四号'
        para, idx = self._find_paragraph_by_text("学术诚信声明")
        if para is None:
            self._add_result(item, False, 0)
            return

        paras = self.document.paragraphs
        # 声明页范围：标题之后 到 "目录"之前
        end = len(paras)
        for j in range(idx + 1, len(paras)):
            if paras[j].text.strip().replace(" ", "").replace("　", "") == "目录" \
               or ('目' in paras[j].text and '录' in paras[j].text and len(paras[j].text.strip()) <= 4):
                end = j
                break

        checked_any = False
        satisfied = True   # 内容是否都满足宋体小四
        for i in range(idx + 1, end):
            p = paras[i]
            if not p.text.strip():
                continue
            para_elem = p._element
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                checked_any = True
                font_name, size = self._resolve_run_cn_font(run._element, para_elem)
                # 点1：宋体
                if not (font_name and '宋体' in str(font_name)):
                    satisfied = False
                # 点2：小四(12pt)
                if not (size is not None and abs(size - 12) < 0.5):
                    satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_abstract_align_penalty(self):
        """-1: 中文摘要页内容段落格式不满足左对齐、首行缩进两字符、1.5倍行距

        细则三个点（内容段落须同时满足；任一不满足即扣 -1）：
          1) 左对齐   2) 首行缩进两字符   3) 1.5倍行距
        内容 = "摘要"标题之后、"关键词"之前的正文段落。
        对齐/缩进/行距按办公软件继承链解析；首行缩进按可见正文字号换算(以实际显示为准)。
        """
        item = "-1: 中文摘要页内容段落格式不满足左对齐、首行缩进两字符、1.5倍行距"
        para, idx = self._find_paragraph_by_text("摘要")
        if para is None:
            self._add_result(item, False, 0)
            return

        paras = self.document.paragraphs
        checked_any = False
        satisfied = True   # 是否满足 左对齐+首行缩进两字符+1.5倍行距
        for i in range(idx + 1, min(idx + 15, len(paras))):
            p = paras[i]
            text = p.text.strip()
            if "关键词" in text:
                break
            if not text:
                continue

            checked_any = True
            pPr = p._element.find(qn('w:pPr'))
            pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
            style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

            # 点1：左对齐（办公软件"左对齐"存为 left/start；未设置按默认）
            jc_val = None
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(qn('w:val'))
            if jc_val is None and style_id:
                jc_val = self._style_jc_lookup(style_id)
            if jc_val not in ('left', 'start', None):
                satisfied = False

            # 点2：首行缩进两字符（按可见字号换算）
            chars = self._para_first_line_indent_chars(pPr, style_id)
            if abs(chars - 2.0) > 0.4:
                satisfied = False

            # 点3：1.5倍行距（lineRule=auto 且 line=360=1.5×240）
            space = None
            if pPr is not None:
                space = pPr.find(qn('w:spacing'))
            if space is None and style_id:
                space = self._style_spacing_lookup(style_id)
            is_15 = False
            if space is not None:
                line = space.get(qn('w:line'))
                rule = space.get(qn('w:lineRule'))
                if line and rule in (None, 'auto'):
                    try:
                        is_15 = abs(int(line) / 240.0 - 1.5) < 0.1
                    except ValueError:
                        is_15 = False
            if not is_15:
                satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_keywords_paragraph_penalty(self):
        """-1: 中文摘要页"关键词"三字段落格式不满足左对齐、首行缩进两字符、1.5倍行距

        细则三个点（关键词所在段落须同时满足；任一不满足即扣 -1）：
          1) 左对齐   2) 首行缩进两字符   3) 1.5倍行距
        对齐/缩进/行距按办公软件继承链解析；首行缩进按可见正文字号换算(以实际显示为准)。
        找不到关键词段落则不扣分。
        """
        item = "-1: 关键词段落格式不满足左对齐、首行缩进两字符、1.5倍行距"
        para, _ = self._find_paragraph_by_text("关键词")
        if para is None:
            self._add_result(item, False, 0)
            return

        pPr = para._element.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

        # 点1：左对齐（left/start；未设置按默认）
        jc_val = None
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                jc_val = jc.get(qn('w:val'))
        if jc_val is None and style_id:
            jc_val = self._style_jc_lookup(style_id)
        is_left = jc_val in ('left', 'start', None)

        # 点2：首行缩进两字符（按可见字号换算）
        chars = self._para_first_line_indent_chars(pPr, style_id)
        is_indent2 = abs(chars - 2.0) <= 0.4

        # 点3：1.5倍行距（lineRule=auto 且 line=360=1.5×240）
        space = None
        if pPr is not None:
            space = pPr.find(qn('w:spacing'))
        if space is None and style_id:
            space = self._style_spacing_lookup(style_id)
        is_15 = False
        if space is not None:
            line = space.get(qn('w:line'))
            rule = space.get(qn('w:lineRule'))
            if line and rule in (None, 'auto'):
                try:
                    is_15 = abs(int(line) / 240.0 - 1.5) < 0.1
                except ValueError:
                    is_15 = False

        # 未全部满足 => 扣分
        penalize = not (is_left and is_indent2 and is_15)
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_research_direction_penalty(self):
        """-1: 中文摘要页"研究方向"字体不为黑体小四

        细则两个点（"研究方向"标签文字须同时满足；任一不满足即扣 -1）：
          1) 黑体   2) 小四(12pt)
        仅约束"研究方向"这四个字标签本身（不含其后内容）。字体/字号按办公软件继承链解析。
        找不到"研究方向"则不扣分。
        """
        item = "-1: 中文摘要页\"研究方向\"字体不为黑体小四"
        target = "研究方向"

        label_para = None
        for p in self.document.paragraphs:
            if target in p.text:
                label_para = p
                break
        if label_para is None:
            self._add_result(item, False, 0)
            return

        # 定位"研究方向"四字所在的字符区间 [start, start+4)
        full = "".join(r.text for r in label_para.runs)
        start = full.find(target)
        if start < 0:
            self._add_result(item, False, 0)
            return
        end = start + len(target)

        para_elem = label_para._element
        checked_any = False
        satisfied = True   # 是否满足黑体+小四
        offset = 0
        for run in label_para.runs:
            r_len = len(run.text)
            r_start, r_end = offset, offset + r_len
            offset = r_end
            if not run.text or not run.text.strip():
                continue
            # 仅检查与"研究方向"区间重叠的 run
            if not (r_start < end and r_end > start):
                continue
            checked_any = True
            font_name, size = self._resolve_run_cn_font(run._element, para_elem)
            # 点1：黑体
            if not (font_name and '黑体' in str(font_name)):
                satisfied = False
            # 点2：小四(12pt)
            if not (size is not None and abs(size - 12) < 0.5):
                satisfied = False

        # 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_body_font_penalty(self):
        """-3: 正文中文字体不满足宋体小四，英文、阿拉伯数字不满足 Times New Roman 小四

        细则的点（正文文字须满足；任一不满足即扣 -3）：
          中文：1) 宋体   2) 小四(12pt)
          英文/阿拉伯数字：3) Times New Roman   4) 小四(12pt)
        正文 = 首个一级标题之后、"参考文献"之前的章节正文段落（排除各级标题）。
        每个 run 按字符类型分别校验字体，字号对两类都要求小四；按办公软件继承链解析(以实际显示为准)。
        """
        item = "-3: 正文字体不满足(中文宋体小四、英文数字Times New Roman小四)"
        paras = self.document.paragraphs
        start = None
        end = len(paras)
        for i, p in enumerate(paras):
            t = p.text.strip()
            if start is None and (re.match(r'^第[一二三四五六七八九十]+章', t)
                                  or re.match(r'^[一二三四五六七八九十]+[、．.]', t)):
                start = i
            if t == '参考文献':
                end = i
                break

        if start is None:
            self._add_result(item, False, 0)
            return

        def is_title(t):
            return (re.match(r'^第[一二三四五六七八九十]+章', t)
                    or re.match(r'^[一二三四五六七八九十]+[、．.]', t)
                    or re.match(r'^\d+\.\d+', t)
                    or re.match(r'^[（(][一二三四五六七八九十]+[)）]', t))

        checked_any = False
        satisfied = True
        for i in range(start, end):
            p = paras[i]
            text = p.text.strip()
            if not text or is_title(text):
                continue
            para_elem = p._element
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                checked_any = True
                cn_font, size = self._resolve_run_cn_font(run._element, para_elem)
                en_font = self._resolve_run_en_font(run._element, para_elem)
                has_cjk = re.search(r'[一-鿿]', run.text)
                has_en = re.search(r'[A-Za-z0-9]', run.text)
                if not (size is not None and abs(size - 12) < 0.5):
                    satisfied = False
                if has_cjk and not (cn_font and '宋体' in str(cn_font)):
                    satisfied = False
                if has_en and not (en_font and 'Times New Roman' in str(en_font)):
                    satisfied = False

        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -3 if penalize else 0)

    def _check_page_number_format_penalty(self):
        """-3: 页码从摘要页开始，中文摘要和英文摘要页页码为大写罗马数字，正文页码为阿拉伯数字

        细则三个点（须同时满足；任一不满足即扣 -3）：
          1) 页码从摘要页开始（摘要之前的封面/声明/目录等节页脚无页码域，摘要节起有页码域）
          2) 中英文摘要所在节页码格式为大写罗马数字(upperRoman)
          3) 正文所在节页码格式为阿拉伯数字(decimal)
        办公软件中：页码格式存于节属性 w:sectPr/w:pgNumType/@w:fmt；页码是否"显示"看页脚是否含 PAGE 域。
        """
        item = "-3: 页码格式(从摘要页起、摘要大写罗马数字、正文阿拉伯数字)"
        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return
        if not sections:
            self._add_result(item, False, 0)
            return

        # 逐节收集其正文文本（按 body 子元素顺序划分节）
        body = self.document.element.body
        sec_texts = []
        buf = []
        for child in body.iterchildren():
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                t = "".join(x.text or '' for x in child.findall('.//' + qn('w:t')))
                if t.strip():
                    buf.append(t.strip())
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    sec_texts.append("\n".join(buf))
                    buf = []
            elif tag == 'sectPr':
                sec_texts.append("\n".join(buf))
                buf = []
        # 对齐长度（防御）
        while len(sec_texts) < len(sections):
            sec_texts.append("")

        def sec_fmt(section):
            pg = section._sectPr.find(qn('w:pgNumType'))
            return pg.get(qn('w:fmt')) if pg is not None else None

        def has_page_field(section):
            for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
                if ftr is None:
                    continue
                fe = ftr._element
                for it in fe.findall('.//' + qn('w:instrText')):
                    if it.text and 'PAGE' in it.text.upper():
                        return True
                for fs in fe.findall('.//' + qn('w:fldSimple')):
                    if 'PAGE' in (fs.get(qn('w:instr')) or '').upper():
                        return True
            return False

        # 分类各节：摘要节 / 正文节 / 摘要之前
        def classify(text):
            if ('摘要' in text) or ('Abstract' in text):
                return 'abstract'
            if re.search(r'(^|\n)\s*第[一二三四五六七八九十]+章', text) \
               or re.search(r'(^|\n)\s*[一二三四五六七八九十]+[、．.]', text):
                return 'body'
            return 'other'

        abstract_idx = None
        for i, txt in enumerate(sec_texts[:len(sections)]):
            if classify(txt) == 'abstract':
                abstract_idx = i
                break

        satisfied = True

        # 点1：页码从摘要页开始
        if abstract_idx is None:
            satisfied = False
        else:
            # 摘要之前的节不应显示页码
            for i in range(abstract_idx):
                if has_page_field(sections[i]):
                    satisfied = False
            # 摘要节应显示页码
            if not has_page_field(sections[abstract_idx]):
                satisfied = False

        # 点2、点3：按节内容判定页码格式
        for i, section in enumerate(sections):
            kind = classify(sec_texts[i] if i < len(sec_texts) else "")
            if kind == 'abstract':
                if sec_fmt(section) != 'upperRoman':
                    satisfied = False
            elif kind == 'body':
                if sec_fmt(section) != 'decimal':
                    satisfied = False

        # 未全部满足 => 扣分
        penalize = not satisfied
        self._add_result(item, penalize, -3 if penalize else 0)

    def _check_page_number_style_penalty(self):
        """-1: 页码格式不为 Calibri、小五、居中对齐

        细则三个点（每个显示页码的页脚段落须同时满足；任一不满足即扣 -1）：
          1) 字体 Calibri   2) 小五(9pt)   3) 居中对齐
        页码 = 页脚中含 PAGE 域的段落。字体/字号按办公软件继承链解析(以实际显示为准)，
        对齐按 段落直接 jc -> 样式链 解析。未找到任何页码则不扣分。
        """
        item = "-1: 页码格式不为Calibri小五、居中对齐"
        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return

        def para_has_page_field(p_elem):
            for it in p_elem.findall('.//' + qn('w:instrText')):
                if it.text and 'PAGE' in it.text.upper():
                    return True
            for fs in p_elem.findall('.//' + qn('w:fldSimple')):
                if 'PAGE' in (fs.get(qn('w:instr')) or '').upper():
                    return True
            return False

        checked_any = False
        satisfied = True   # 所有页码段落是否都满足 Calibri+小五+居中

        for section in sections:
            footers = [section.footer]
            if section.different_first_page_header_footer:
                footers.append(section.first_page_footer)
            for ftr in footers:
                if ftr is None:
                    continue
                for para in ftr.paragraphs:
                    para_elem = para._element
                    if not para_has_page_field(para_elem):
                        continue
                    checked_any = True

                    pPr = para_elem.find(qn('w:pPr'))
                    pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
                    style_id = pStyle.get(qn('w:val')) if pStyle is not None else None

                    # 点3：居中对齐
                    jc_val = None
                    if pPr is not None:
                        jc = pPr.find(qn('w:jc'))
                        if jc is not None:
                            jc_val = jc.get(qn('w:val'))
                    if jc_val is None and style_id:
                        jc_val = self._style_jc_lookup(style_id)
                    if jc_val != 'center':
                        satisfied = False

                    # 点1、点2：Calibri、小五(9pt)
                    for run in para.runs:
                        en_font = self._resolve_run_en_font(run._element, para_elem)
                        _, size = self._resolve_run_cn_font(run._element, para_elem)
                        if not (en_font and 'Calibri' in str(en_font)):
                            satisfied = False
                        if not (size is not None and abs(size - 9) < 0.5):
                            satisfied = False

        # 有页码 且 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -1 if penalize else 0)

    def _check_header_margin_penalty(self):
        """-1: 页眉上边距不为1.50厘米

        细则一个点：页眉上边距(页眉到页面顶边的距离) = 1.50 厘米，不满足则扣 -1。
        办公软件将该值存于节属性 w:sectPr/w:pgMar/@w:header(单位 twips，1.5cm≈851twips)，
        即 python-docx 的 section.header_distance。文档所有节都须为1.5cm；任一节不满足即扣分。
        """
        item = "-1: 页眉上边距不为1.50厘米"
        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return
        if not sections:
            self._add_result(item, False, 0)
            return

        checked_any = False
        satisfied = True
        for section in sections:
            hd = section.header_distance
            if hd is None:
                # 未显式设置 => 不能确认为1.5cm
                satisfied = False
                continue
            checked_any = True
            # 允许 ±0.05cm 误差（1.5cm 实际存为851twips≈1.501cm）
            if abs(hd.cm - 1.5) > 0.05:
                satisfied = False

        # 不满足1.5cm => 扣分
        penalize = (not checked_any) or (not satisfied)
        self._add_result(item, penalize, -1 if penalize else 0)

    def _header_bottom_border(self, section):
        """返回该节页眉中"页眉横线"对应的下边框元素 w:bottom（段落 pBdr 优先，其次段落样式链）；
        无横线返回 None。办公软件的页眉横线即页眉段落的下边框(pBdr/bottom)。"""
        hdr = section.header
        if hdr is None:
            return None
        for para in hdr.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            # 1) 段落直接 pBdr/bottom
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    bottom = pBdr.find(qn('w:bottom'))
                    if bottom is not None:
                        val = bottom.get(qn('w:val'))
                        if val and val != 'none' and val != 'nil':
                            return bottom
            # 2) 段落样式链 pBdr/bottom
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_id = pStyle.get(qn('w:val'))
            bottom = self._style_pbdr_bottom_lookup(style_id)
            if bottom is not None:
                val = bottom.get(qn('w:val'))
                if val and val != 'none' and val != 'nil':
                    return bottom
        return None

    def _style_pbdr_bottom_lookup(self, style_id):
        """沿段落样式 basedOn 链查找 pBdr/bottom 元素；无则返回 None"""
        if not style_id or not self.document:
            return None
        seen = set()
        try:
            style = self.document.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        except:
            return None
        while style is not None and style.style_id not in seen:
            seen.add(style.style_id)
            s_pPr = style.element.find(qn('w:pPr'))
            if s_pPr is not None:
                pBdr = s_pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    bottom = pBdr.find(qn('w:bottom'))
                    if bottom is not None:
                        return bottom
            based_on = style.element.find(qn('w:basedOn'))
            if based_on is None or not based_on.get(qn('w:val')):
                break
            try:
                style = self.document.styles.get_by_id(based_on.get(qn('w:val')), WD_STYLE_TYPE.PARAGRAPH)
            except:
                break
        return None

    def _check_header_line_penalty(self):
        """-3: 页眉不带有横线

        细则一个点：页眉需带有横线，不带则扣 -3。
        办公软件中"页眉横线"= 页眉段落的下边框(pBdr/bottom，val 非 none/nil)。
        任一节的页眉都没有横线即视为"不带横线"→ 扣分；所有节页眉都有横线则不扣。
        """
        item = "-3: 页眉不带有横线"
        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return
        if not sections:
            self._add_result(item, False, 0)
            return

        # 每一节页眉都须带横线；任一节无横线 => 扣分
        all_have_line = True
        for section in sections:
            if self._header_bottom_border(section) is None:
                all_have_line = False
                break

        penalize = not all_have_line
        self._add_result(item, penalize, -3 if penalize else 0)

    def _check_header_line_style_penalty(self):
        """-3: 页眉横线样式不满足双实线、宽度为3磅

        细则两个点（页眉横线须同时满足；任一不满足即扣 -3）：
          1) 样式为"两条线"（用户放宽解释：办公软件里视觉呈现为两条平行线的边框样式均算）
          2) 宽度为3磅
        办公软件中页眉横线 = 页眉段落下边框(pBdr/bottom)：
          OOXML 中呈现为两条平行线的 w:val 取值集合见 TWO_LINE_STYLES；
          宽度 w:bottom/@w:sz 单位为 1/8 磅(3磅=24)。
        无横线则本项不适用(由"页眉不带横线"另判)，不在此扣分。
        """
        item = "-3: 页眉横线样式不满足双实线、宽度为3磅"

        # 视觉上"两条线"的边框样式：等宽双线 + 各种细粗/粗细组合(不同间距)。
        # 排除 triple 和 thinThickThin*(三条线) 及所有单线样式(single/dashed/dotted/wave 等)。
        TWO_LINE_STYLES = {
            'double',
            'thinThickSmallGap', 'thickThinSmallGap',
            'thinThickMediumGap', 'thickThinMediumGap',
            'thinThickLargeGap', 'thickThinLargeGap',
        }

        try:
            sections = list(self.document.sections)
        except:
            self._add_result(item, False, 0)
            return
        if not sections:
            self._add_result(item, False, 0)
            return

        checked_any = False
        satisfied = True   # 所有带横线的页眉是否都满足 两条线+3磅
        for section in sections:
            bottom = self._header_bottom_border(section)
            if bottom is None:
                continue  # 无横线，交由 _check_header_line_penalty 处理
            checked_any = True
            # 点1：呈现为两条线的边框样式
            val = bottom.get(qn('w:val'))
            if val not in TWO_LINE_STYLES:
                satisfied = False
            # 点2：宽度3磅 (sz 单位1/8磅，3磅=24)
            sz = bottom.get(qn('w:sz'))
            try:
                is_3pt = sz is not None and abs(int(sz) / 8.0 - 3.0) < 0.2
            except ValueError:
                is_3pt = False
            if not is_3pt:
                satisfied = False

        # 有横线 且 未全部满足 => 扣分
        penalize = checked_any and not satisfied
        self._add_result(item, penalize, -3 if penalize else 0)

    def _print_final_result(self):
        """打印维度二评分结果：仅显示命中项(分数+评分细则内容)，分数在最下面。
        保留供本地调试使用；模块级 evaluate() 不再调用它。"""
        print("维度二：")

        # 命中项 = 已计入结果的项（passed=True），保持检查顺序
        # item_text 形如 "+1: xxx" / "-3: xxx"，本身已含分数前缀，直接输出
        for item_text, passed, _score in self.check_results:
            if passed:
                print(item_text)

        # 分数显示在维度二内容最下面
        print(f"评分结果：{self.score}分")


# ============================================================
# 统一对外接口
# ============================================================

# 该脚本对应的编号，供批量运行器/结果汇总使用
SCRIPT_ID = "012"


def _find_docx_in_dir(dir_path):
    """在给定目录中定位待评估的 Word 文档。
    规则：只匹配 .docx；忽略以 '~$' 开头的临时文件。
    若目录中存在多个候选，取修改时间最新的一个（与"最近改过的评估稿"直觉一致）。
    """
    if not dir_path or not os.path.isdir(dir_path):
        return None
    candidates_docx = []
    for name in os.listdir(dir_path):
        if name.startswith('~$'):
            continue
        full = os.path.join(dir_path, name)
        if not os.path.isfile(full):
            continue
        low = name.lower()
        if low.endswith('.docx'):
            candidates_docx.append(full)
    if not candidates_docx:
        return None
    candidates_docx.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return candidates_docx[0]


def _parse_item_prefix(item_text):
    """从形如 "+1: xxx" / "-3: xxx" 的检查项文本中解析出 (max_delta, rule_text)。
    max_delta：该项的分值上限（得分项为正、扣分项为负）；解析失败时为 0。"""
    m = re.match(r'^\s*([+-]?\d+)\s*[:：]\s*(.*)$', item_text)
    if not m:
        return 0, item_text
    try:
        return int(m.group(1)), m.group(2).strip()
    except ValueError:
        return 0, item_text


def evaluate(dir_path: str) -> dict:
    """统一入口：接收"脚本所在目录的路径"，脚本自身在该目录内定位并打开 Word 文档。

    返回结构（详见 §2.2）：
      {
        "id": "012", "file_name": "...", "status": "ok"|"error", "error": None|str,
        "dim1_pass": bool, "dim1_reason": str,
        "dim2_items": [ {rule, max_delta, delta, hit, detail}, ... ],
        "total_score": int, "max_score": int
      }
    """
    result = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }

    try:
        file_path = _find_docx_in_dir(dir_path)
        if not file_path:
            result["status"] = "error"
            result["error"] = f"目录中未找到可评估的 Word 文档: {dir_path}"
            return result
        result["file_name"] = os.path.basename(file_path)

        evaluator = WordDocumentEvaluator(file_path)
        evaluator.run()

        result["dim1_pass"] = bool(evaluator.dimension1_passed)
        result["dim1_reason"] = evaluator.dimension1_reason or ""

        # 维度一未通过：dim2_items 留空，总分为 0
        if not evaluator.dimension1_passed:
            result["total_score"] = 0
            result["max_score"] = 0
            return result

        # 维度二逐项：命中/未命中都返回
        items = []
        max_score = 0
        for item_text, passed, score_change in evaluator.check_results:
            max_delta, rule = _parse_item_prefix(item_text)
            items.append({
                "rule": rule,
                "max_delta": max_delta,
                "delta": score_change if passed else 0,
                "hit": bool(passed),
                "detail": "",
            })
            # 满分统计：得分项累加正分值；扣分项不计入满分
            if max_delta > 0:
                max_score += max_delta

        result["dim2_items"] = items
        result["total_score"] = int(evaluator.score)
        result["max_score"] = max_score
        return result

    except Exception as e:
        result["status"] = "error"
        result["error"] = f"{type(e).__name__}: {e}"
        return result


if __name__ == "__main__":
    # 本地调试入口：接收"脚本所在目录"，打印结构化 JSON 结果。
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    print(json.dumps(evaluate(_dir), ensure_ascii=False, indent=2))
