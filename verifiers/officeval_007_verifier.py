"""
Word文档自动评估脚本
根据评分细则评估文档格式

对外统一入口：
    evaluate(dir_path: str) -> dict

- dir_path 为脚本所在目录路径；脚本自行在该目录内定位并打开被评估的 .docx 文档。
- 返回结构化字典（含维度一通过与否、维度二逐项得分、总分），详见接口约定 §2.2。
"""
import json
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 脚本编号（与文件名 officeval_007_verifier.py 对齐）
SCRIPT_ID = "007"


def _open_word_document(file_path: str):
    """打开 .docx 文件 (python-docx 仅支持 OOXML/.docx)."""
    return Document(file_path)


class DocxEvaluator:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self.results = []  # 存储每个检查点的结果
        self.score = 0
        self.dimension1_passed = False

    def evaluate(self) -> tuple[int, list[str]]:
        """
        执行完整评估
        返回: (总分, 检查结果列表)
        """
        print(f"开始评估文件: {self.file_path}\n")

        # 维度1：可用与可修改性
        dim1_result = self.check_dimension1()

        # 打印维度1结果
        print("\n" + "="*60)
        print("维度1检查结果:")
        print("="*60)
        for result in self.results:
            status = "✓ 通过" if result['passed'] else "✗ 不通过"
            print(f"{status}: {result['description']}")
        print("="*60)

        if not dim1_result:
            print("\n维度1未通过，得分为0分")
            return 0, self.results

        self.dimension1_passed = True
        print("\n维度1通过，开始检查维度2...\n")

        # 维度2：完成度评分细则
        self.check_dimension2()

        # 打印结果
        self.print_results()

        return self.score, self.results

    def add_result(self, description: str, passed: bool, score_change: int = 0):
        """记录检查结果"""
        self.results.append({
            'description': description,
            'passed': passed,
            'score_change': score_change
        })
        if not self.dimension1_passed:
            return
        if passed:
            self.score += score_change

    def print_results(self):
        """打印评估结果"""
        print("=" * 80)
        print("评估结果:")
        print("=" * 80)
        for result in self.results:
            score_change = result['score_change']
            passed = result['passed']
            description = result['description']

            if score_change == 0:
                if passed:
                    print(f"{description}")
            elif passed:
                sign = "+" if score_change > 0 else ""
                print(f"{sign}{score_change}分：{description}")

        print("=" * 80)
        print(f"最终得分: {self.score}分")
        print("=" * 80)

    # ==================== 维度1 ====================

    def check_dimension1(self) -> bool:
        """检查维度1：可用与可修改性"""
        dimension1_passed = True

        # 1. 检查文件扩展名
        if not self.check_file_extension():
            dimension1_passed = False

        # 2. 检查文件可正常打开
        if not self.check_file_openable():
            dimension1_passed = False

        return dimension1_passed

    def check_file_extension(self) -> bool:
        """检查文件扩展名"""
        ext = os.path.splitext(self.file_path)[1].lower()
        passed = ext == '.docx'
        self.add_result("文件扩展名为.docx", passed)
        return passed

    def check_file_openable(self) -> bool:
        """检查文件可正常打开"""
        try:
            self.doc = _open_word_document(self.file_path)
            print(f"文件成功打开，段落数: {len(self.doc.paragraphs)}")
            passed = True
        except Exception as e:
            print(f"文件打开失败: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
            passed = False
        self.add_result("文件可正常打开", passed)
        return passed

    def check_toc_editable(self) -> bool:
        """检查目录内容可编辑（不是整张图片）"""
        if not self.doc:
            return False

        # 检查是否有目录段落且包含文本（去除空格后检查）
        toc_found = False
        toc_has_text = False

        for para in self.doc.paragraphs:
            text = para.text.strip()
            # 去除空格后检查
            text_no_space = text.replace(' ', '')
            if '目录' in text_no_space:
                toc_found = True
                # 检查段落是否包含文本内容（可编辑）
                if para.runs and any(run.text.strip() for run in para.runs):
                    toc_has_text = True
                    break

        # 如果找到目录，还需要检查目录区域是否全是图片
        if toc_found:
            # 检查目录附近的段落是否有文本
            for i, para in enumerate(self.doc.paragraphs):
                text_no_space = para.text.replace(' ', '')
                if '目录' in text_no_space:
                    # 检查前后几个段落
                    for j in range(max(0, i-2), min(len(self.doc.paragraphs), i+20)):
                        p = self.doc.paragraphs[j]
                        if p.runs and any(run.text.strip() for run in p.runs):
                            toc_has_text = True
                            break
                    break

        passed = toc_found and toc_has_text
        print(f"目录检查: 找到目录={toc_found}, 有可编辑文本={toc_has_text}")
        self.add_result("目录页内容可编辑，不是一张图片", passed)
        return passed

    def check_blank_pages(self) -> bool:
        """检查无连续2页及以上空白页"""
        if not self.doc:
            return False

        # 通过检查段落数量和文本内容来判断
        consecutive_blank = 0
        page_approximate_chars = 1000  # 粗略估计一页的字符数
        blank_count = 0

        # 简化检查：统计空白段落过多的情况
        empty_paras = sum(1 for p in self.doc.paragraphs if not p.text.strip())
        total_paras = len(self.doc.paragraphs)
        empty_ratio = empty_paras / total_paras if total_paras > 0 else 0

        # 如果空白段落占比过高，可能有连续空白页
        passed = empty_ratio < 0.8  # 假设80%空白段落可能有问题
        self.add_result("无连续2页及以上空白页", passed)
        return passed

    def check_garbled_text(self) -> bool:
        """检查无超过1/3面积乱码"""
        if not self.doc:
            return False

        # 检查是否有异常字符
        total_chars = 0
        garbled_chars = 0

        for para in self.doc.paragraphs:
            text = para.text
            total_chars += len(text)
            # 检查不可打印字符和异常字符
            for char in text:
                if ord(char) > 65535 or ord(char) < 32 and char not in '\n\r\t':
                    garbled_chars += 1

        garbled_ratio = garbled_chars / total_chars if total_chars > 0 else 0
        passed = garbled_ratio < 0.33
        self.add_result("无超过1/3面积乱码", passed)
        return passed

    def check_text_overlap(self) -> bool:
        """检查无超过1/3面积文字重叠"""
        # Word文档中文字重叠需要检查格式设置
        # 通过检查是否有异常的行距或位置设置来判断
        if not self.doc:
            return False

        overlap_detected = False
        for para in self.doc.paragraphs:
            # 检查是否有异常的段落格式
            if para.paragraph_format:
                line_spacing = para.paragraph_format.line_spacing
                # 如果行距设置异常（如负值或过小），可能导致文字重叠
                if line_spacing and line_spacing < 0.5:
                    overlap_detected = True
                    break

        passed = not overlap_detected
        self.add_result("无超过1/3面积文字重叠", passed)
        return passed

    # ==================== 维度2 ====================

    def check_dimension2(self):
        """检查维度2：完成度评分细则"""

        # +5：目录是自动生成目录
        self.check_auto_toc()

        # +3：目录页除"目录"两字外目录页中所有条目左对齐
        self.check_toc_left_align()

        # +1：目录页"目录"两字字体为黑体三号，居中
        self.check_toc_title_font()

        # +1：目录页"目录"两字中间空两个字符
        self.check_toc_title_spacing()

        # +1：目录页"目录"两字段落格式为两倍行距，段前1行，段后1行
        self.check_toc_title_paragraph()

        # +3：目录中包含三级标题
        self.check_toc_three_levels()

        # +3：目录中的一级标题文字为小四号、黑体
        self.check_toc_level1_font()

        # +3：目录中的一级标题段落格式为1.5倍行距、左对齐
        self.check_toc_level1_paragraph()

        # +3：目录中的二级标题文字为小四号、宋体
        self.check_toc_level2_font()

        # +3：目录中的二级标题段落格式为1.5倍行距、左对齐
        self.check_toc_level2_paragraph()

        # +3：目录中的三级标题文字为小四号、宋体
        self.check_toc_level3_font()

        # +3：目录中的三级标题段落格式为1.5倍行距、左对齐，标题后空一个字符
        self.check_toc_level3_paragraph()

        # +3：目录中的标题层级严格按照"一"、"（一）"、"1"
        self.check_toc_level_format()

        # +1：目录页的页码使用阿拉伯数字
        self.check_toc_page_number()

        # -3：目录页不在英文摘要页后面
        self.check_toc_after_english_abstract()

        # -1：封面页的图片大小不满足5.34×15.34
        self.check_cover_image_size()

        # -1：封面页的图片格式不满足嵌入型布局、两端对齐
        self.check_cover_image_layout()

        # -1：封面页"学校代码：12786"文本字体格式不满足宋体四号
        self.check_cover_school_code_font()

        # -1：封面页"学校代码：12786"文本段落格式不满足右对齐、单倍行距
        self.check_cover_school_code_paragraph()

        # -1：封面页"成人高等教育毕业论文（设计）"文本段落格式不满足居中对齐、单倍行距
        self.check_cover_graduate_paper_paragraph()

        # -1：封面页"成人高等教育毕业论文（设计）"文本字体格式不满足宋体，字号为32
        self.check_cover_graduate_paper_font()

        # -1：封面页标题字体不满足华文新魏小一
        self.check_cover_title_font()

        # -1：封面页标题段落格式不满足居中对齐、单倍行距
        self.check_cover_title_paragraph()


        # -1：封面页的表格内第二列字体段落格式不满足居中对齐、单倍行距
        self.check_cover_table_second_col_paragraph()

        # -1："启明学院毕业论文作者声明"文本所在页的标题及内容字体不满足宋体三号
        self.check_declaration_font()

        # -1："启明学院毕业论文作者声明"文本所在页的标题段落格式不满足居中对齐、1.5倍行距
        self.check_declaration_title_paragraph()

        # -1："启明学院毕业论文作者声明"文本所在页的内容段落格式不满足居中对齐、1.5倍行距
        self.check_declaration_content_paragraph()

        # -1：摘要页标题的字体格式不满足黑体三号
        self.check_abstract_title_font()

        # -1：摘要页标题的段落格式不满足段前一行、段后一行，两倍行距，居中对齐
        self.check_abstract_title_paragraph()

        # -1：摘要页的内容段落格式不满足首行缩进两字符、两端对齐、1.5倍行距
        self.check_abstract_content_paragraph()

        # -1：摘要页的关键词及冒号后内容字体格式不满足黑体四号
        self.check_abstract_keywords_font()

        # -1：摘要页的关键词及其冒号后内容段落格式不满足单倍行距、左对齐
        self.check_abstract_keywords_paragraph()

        # -1：英文摘要页的标题字体格式不满足Times New Roman三号
        self.check_english_abstract_title_font()

        # -1：英文摘要页的标题段落格式不满足段前一行、段后一行，两倍行距，居中对齐
        self.check_english_abstract_title_paragraph()

        # -1：英文摘要页的内容段落格式不满足首行缩进两字符、两端对齐、1.5倍行距
        self.check_english_abstract_content_paragraph()

        # -1：英文摘要页的关键词及其冒号后内容段落格式不满足首行缩进两字符、左对齐、1.5倍行距
        self.check_english_abstract_keywords_paragraph()

        # -1：英文摘要页的关键词及冒号后内容字体格式不满足Times New Roman四号
        self.check_english_abstract_keywords_font()

        # -3：从目录页开始出现任意一页没有页眉
        self.check_header_present()

        # -1：页眉字体格式不是宋体五号
        self.check_header_font()

        # -1：页眉段落格式不满足页眉下无横线、居中对齐、单倍行距
        self.check_header_paragraph()

        # -1：页眉上边距不满足1.50厘米
        self.check_header_top_margin()

        # -1：页脚下边距不满足1.75厘米
        self.check_footer_bottom_margin()

        # -3：文章不满足：从摘要页开始编写页码，摘要页到目录页页码格式为大写罗马数字
        self.check_page_number_roman()

        # -3：文章不满足：从正文内容页重新编写页码从1开始，页码格式为阿拉伯数字
        self.check_page_number_restart()

        # -3：文章一级标题字体格式不满足黑体三号
        self.check_level1_font()

        # -3：文章一级标题段落格式不满足居中对齐、段前一行、段后一行、两倍行距
        self.check_level1_paragraph()

        # -3：文章二级标题字体格式不满足黑体四号
        self.check_level2_font()

        # -3：文章二级标题段落格式不满足左对齐、段前0.5行、段后0.5行、1.5倍行距
        self.check_level2_paragraph()

        # -3：文章三级标题字体格式不满足黑体小四号
        self.check_level3_font()

        # -3：文章三级标题段落格式不满足左对齐、段前0.5行、段后0.5行、1.5倍行距
        self.check_level3_paragraph()

    # ==================== 得分点检查 ====================

    def check_auto_toc(self):
        """+5：目录是自动生成目录

        细则依据：办公软件（Word/WPS）判定"自动目录"的两个硬指标必须
        同时满足——目录外层由 TOC 域驱动，且目录条目由 PAGEREF 域承载
        页码。若只有 TOC 外层字段而条目已被替换为静态文本（无 PAGEREF
        字段），按 F9 无法更新、Ctrl+点击不可跳转——办公软件视觉上看
        像目录，但功能上不是自动目录。因此加分需同时满足：
          1) 文档任一段落存在 <w:instrText> 且内容含 "TOC" 指令
             （目录外层字段容器）
          2) 文档任一段落存在 <w:instrText> 且内容含 "PAGEREF" 指令
             （至少一个条目由 PAGEREF 字段生成页码）
        两条件同时成立才判为自动目录并加 +5 分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

        has_toc_field = False
        has_pageref_field = False

        for para in self.doc.paragraphs:
            for instr in para._element.findall(f'.//{W_NS}instrText'):
                text = instr.text or ''
                text_upper = text.upper()
                if 'TOC' in text_upper:
                    has_toc_field = True
                if 'PAGEREF' in text_upper:
                    has_pageref_field = True
            if has_toc_field and has_pageref_field:
                break

        passed = has_toc_field and has_pageref_field
        self.add_result("目录是自动生成目录", passed, 5)

    def check_toc_left_align(self):
        """+3：目录页除"目录"两字外目录页中所有条目左对齐

        细则依据：目录页内除"目录"标题以外的所有条目，其段落对齐方式必须为
        左对齐。Word 文件属性上，段落对齐由 w:pPr/w:jc 决定：值为 "left" 或
        "start" 视为左对齐；未设置 w:jc（None）时，Word 默认即为左对齐。
        因此本检查只判定"每个目录条目段落的对齐是否为左对齐"这一项。

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目段落。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_left = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                found_any = True

                # 直接读取 w:pPr/w:jc 判断对齐（Word 文件属性）
                jc_val = None
                pPr = p_el.find(f'{W_NS}pPr')
                if pPr is not None:
                    jc = pPr.find(f'{W_NS}jc')
                    if jc is not None:
                        jc_val = jc.get(f'{W_NS}val')

                # 左对齐：显式 left/start，或未设置（None，Word 默认左对齐）
                is_left = jc_val in (None, 'left', 'start')
                if not is_left:
                    all_left = False

        passed = found_any and all_left
        self.add_result("目录页除'目录'两字外所有条目左对齐", passed, 3)

    def check_toc_title_font(self):
        """+1：目录页'目录'两字字体为黑体三号，居中

        细则依据：目录页中"目录"两字需同时满足三点——
          1) 字体为黑体（Word 文件属性 w:rFonts 的 eastAsia/ascii/hAnsi 任一为"黑体"/SimHei）
          2) 字号为三号（Word 文件属性 w:sz 的半磅值 = 32，即 16 磅）
          3) 段落对齐为居中（Word 文件属性 w:pPr/w:jc = "center"）
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        passed = False

        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '')
            if text_ns != '目录':
                continue

            # 1) 字体：黑体
            font_ok = False
            # 2) 字号：三号 = 16 磅
            size_ok = False
            for run in para.runs:
                if not run.text.strip():
                    continue
                rpr = run._element.rPr
                font_name = ''
                if rpr is not None and rpr.rFonts is not None:
                    for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                        v = rpr.rFonts.get(f'{W_NS}{attr}')
                        if v:
                            font_name = v
                            break
                if not font_name:
                    font_name = run.font.name or ''
                if '黑体' in str(font_name) or 'SimHei' in str(font_name):
                    font_ok = True
                if run.font.size is not None and abs(run.font.size.pt - 16) < 0.5:
                    size_ok = True

            # 3) 对齐：居中（读 w:pPr/w:jc）
            align_ok = False
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None:
                jc = pPr.find(f'{W_NS}jc')
                if jc is not None and jc.get(f'{W_NS}val') == 'center':
                    align_ok = True

            passed = font_ok and size_ok and align_ok
            break

        self.add_result("目录页'目录'两字字体为黑体三号，居中", passed, 1)

    def check_toc_title_spacing(self):
        """+1：目录页'目录'两字中间空两个字符

        细则依据：目录页的"目录"标题，"目"与"录"两字之间需空两个字符。
        Word 文件属性上体现为段落文本内容中，"目"与"录"之间恰好存在两个
        空白字符（半角空格 U+0020 或全角空格 U+3000，两者混合亦可）。
        """
        if not self.doc:
            return

        passed = False
        for para in self.doc.paragraphs:
            text = para.text.strip()
            # 定位目录标题段落：去掉所有空白后正好是"目录"
            if text.replace(' ', '').replace('　', '').replace('\t', '') != '目录':
                continue

            # "目" 和 "录" 之间的字符
            m = re.match(r'^目([\s　]*)录$', text)
            if m and len(m.group(1)) == 2:
                passed = True
            break

        self.add_result("目录页'目录'两字中间空两个字符", passed, 1)

    def check_toc_title_paragraph(self):
        """+1：目录页'目录'两字段落格式为两倍行距，段前1行，段后1行

        细则依据：三项需同时满足，全部读 Word 文件属性 w:pPr/w:spacing：
          1) 两倍行距：w:line="480" 且 w:lineRule="auto"（240=单倍，480=两倍）
          2) 段前1行：w:beforeLines="100"（单位为 1/100 行，100=1行）
          3) 段后1行：w:afterLines="100"
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        passed = False

        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if text_ns != '目录':
                continue

            pPr = para._element.find(f'{W_NS}pPr')
            spacing = pPr.find(f'{W_NS}spacing') if pPr is not None else None

            line_ok = False
            before_ok = False
            after_ok = False
            if spacing is not None:
                # 1) 两倍行距
                line = spacing.get(f'{W_NS}line')
                line_rule = spacing.get(f'{W_NS}lineRule')
                if line == '480' and line_rule == 'auto':
                    line_ok = True

                # 2) 段前 1 行
                before_lines = spacing.get(f'{W_NS}beforeLines')
                if before_lines == '100':
                    before_ok = True

                # 3) 段后 1 行
                after_lines = spacing.get(f'{W_NS}afterLines')
                if after_lines == '100':
                    after_ok = True

            passed = line_ok and before_ok and after_ok
            break

        self.add_result("目录页'目录'两字段落格式为两倍行距，段前1行，段后1行", passed, 1)

    def check_toc_three_levels(self):
        """+3：目录中包含三级标题

        细则依据（宽松路线：办公软件视觉判定）：目录中的三级条目通常以
        阿拉伯数字加分隔符起始（如 "1. xxx"、"1、xxx"、"1．xxx"、
        "1 xxx"），而一级为"一/二/三…、xxx"、二级为"（一）/（二）…
        xxx"。

        本检查直接遍历底层 XML 的全部 w:p 段落。WPS/Word 生成的目录可能
        包在 w:sdt/w:sdtContent 内容控件内，python-docx 的 doc.paragraphs
        不会枚举这些段落；同时该类目录也可能没有传统 TOC 字段的
        fldChar end 结束标记。因此以"目录"标题后的、含 w:tab 的目录条目
        段落作为判断范围，避免把正文里的三级标题误当作目录条目。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        passed = False
        found_toc_entry = False

        if toc_title_idx is not None:
            level3_pattern = re.compile(r'^\d+\s*[\.．、\s]')
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if level3_pattern.match(head):
                    passed = True
                    break

        self.add_result("目录中包含三级标题", passed, 3)

    def check_toc_level1_font(self):
        """+3：目录中的一级标题文字为小四号、黑体

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目；其中以
        中文数字 + 分隔符起始的条目视为一级目录标题。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')
        L1_PAT = re.compile(r'^[一二三四五六七八九十百]+\s*[、\.．]')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        def _run_text(r_el):
            return ''.join(t.text or '' for t in r_el.findall(f'.//{W_NS}t'))

        def _run_font(r_el):
            rpr = r_el.find(f'{W_NS}rPr')
            if rpr is None:
                return ''
            rfonts = rpr.find(f'{W_NS}rFonts')
            if rfonts is None:
                return ''
            for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                value = rfonts.get(f'{W_NS}{attr}')
                if value:
                    return value
            return ''

        def _run_size_pt(r_el):
            rpr = r_el.find(f'{W_NS}rPr')
            if rpr is None:
                return None
            sz = rpr.find(f'{W_NS}sz')
            if sz is None:
                return None
            value = sz.get(f'{W_NS}val')
            try:
                return int(value) / 2
            except (TypeError, ValueError):
                return None

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_ok = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if not L1_PAT.match(head):
                    continue

                found_any = True
                for r_el in p_el.findall(f'{W_NS}r'):
                    rt = _run_text(r_el).strip()
                    if not rt or rt.isdigit():
                        continue
                    font_ok = '黑体' in _run_font(r_el) or 'SimHei' in _run_font(r_el)
                    size_pt = _run_size_pt(r_el)
                    size_ok = size_pt is not None and abs(size_pt - 12) < 0.5
                    if not (font_ok and size_ok):
                        all_ok = False
                        break
                if not all_ok:
                    break

        passed = found_any and all_ok
        self.add_result("目录中的一级标题文字为小四号、黑体", passed, 3)
    def check_toc_level1_paragraph(self):
        """+3：目录中的一级标题段落格式为1.5倍行距、左对齐

        细则依据（宽松路线：与 check_toc_level1_font 保持一致）：本文档
        目录条目未绑定内建 TOC 1 样式，无法以 pStyle 定位一级条目。改按
        办公软件视觉习惯：目录条目起始为"中文数字（一/二/三…/十/百）
        + '、' 或 '.' / '．'"的段落即视为一级条目。对这些条目段落需
        同时满足：
          1) 1.5 倍行距：w:pPr/w:spacing 的 w:line="360" 且 w:lineRule="auto"
             （240=单倍，360=1.5倍，480=两倍）
          2) 左对齐：w:pPr/w:jc 的 val 为 "left"/"start"，或未设置（Word 默认左对齐）

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')
        L1_PAT = re.compile(r'^[一二三四五六七八九十百]+\s*[、\.．]')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_ok = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if not L1_PAT.match(head):
                    continue

                found_any = True

                pPr = p_el.find(f'{W_NS}pPr')

                # 1) 1.5 倍行距
                line_ok = False
                spacing = pPr.find(f'{W_NS}spacing') if pPr is not None else None
                if spacing is not None:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line == '360' and line_rule == 'auto':
                        line_ok = True

                # 2) 左对齐
                jc = pPr.find(f'{W_NS}jc') if pPr is not None else None
                jc_val = jc.get(f'{W_NS}val') if jc is not None else None
                align_ok = jc_val in (None, 'left', 'start')

                if not (line_ok and align_ok):
                    all_ok = False
                    break

        passed = found_any and all_ok
        self.add_result("目录中的一级标题段落格式为1.5倍行距、左对齐", passed, 3)

    def check_toc_level2_font(self):
        """+3：目录中的二级标题文字为小四号、宋体

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目；其中以
        "（中文数字）"或"(中文数字)"起始的条目视为二级目录标题。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')
        L2_PAT = re.compile(r'^[（(][一二三四五六七八九十百]+[）)]')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        def _run_text(r_el):
            return ''.join(t.text or '' for t in r_el.findall(f'.//{W_NS}t'))

        def _run_font(r_el):
            rpr = r_el.find(f'{W_NS}rPr')
            if rpr is None:
                return ''
            rfonts = rpr.find(f'{W_NS}rFonts')
            if rfonts is None:
                return ''
            for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                value = rfonts.get(f'{W_NS}{attr}')
                if value:
                    return value
            return ''

        def _run_size_pt(r_el):
            rpr = r_el.find(f'{W_NS}rPr')
            if rpr is None:
                return None
            sz = rpr.find(f'{W_NS}sz')
            if sz is None:
                return None
            value = sz.get(f'{W_NS}val')
            try:
                return int(value) / 2
            except (TypeError, ValueError):
                return None

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_ok = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if not L2_PAT.match(head):
                    continue

                found_any = True
                for r_el in p_el.findall(f'{W_NS}r'):
                    rt = _run_text(r_el).strip()
                    if not rt or rt.isdigit():
                        continue
                    font_name = _run_font(r_el)
                    font_ok = '宋体' in font_name or 'SimSun' in font_name
                    size_pt = _run_size_pt(r_el)
                    size_ok = size_pt is not None and abs(size_pt - 12) < 0.5
                    if not (font_ok and size_ok):
                        all_ok = False
                        break
                if not all_ok:
                    break

        passed = found_any and all_ok
        self.add_result("目录中的二级标题文字为小四号、宋体", passed, 3)
    def check_toc_level2_paragraph(self):
        """+3：目录中的二级标题段落格式为1.5倍行距、左对齐

        细则依据（宽松路线：与 check_toc_level1_paragraph 保持一致）：本
        文档目录条目未绑定内建 TOC 2 样式，无法以 pStyle 定位二级条目。
        改按办公软件视觉习惯：目录条目起始为"（中文数字）"（如
        "（一）/（二）…"）的段落即视为二级条目。对这些条目段落需
        同时满足：
          1) 1.5 倍行距：w:pPr/w:spacing 的 w:line="360" 且 w:lineRule="auto"
             （240=单倍，360=1.5倍，480=两倍）
          2) 左对齐：w:pPr/w:jc 的 val 为 "left"/"start"，或未设置（Word 默认左对齐）

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')
        L2_PAT = re.compile(r'^[（(][一二三四五六七八九十百]+[）)]')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_ok = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if not L2_PAT.match(head):
                    continue

                found_any = True

                pPr = p_el.find(f'{W_NS}pPr')

                # 1) 1.5 倍行距
                line_ok = False
                spacing = pPr.find(f'{W_NS}spacing') if pPr is not None else None
                if spacing is not None:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line == '360' and line_rule == 'auto':
                        line_ok = True

                # 2) 左对齐
                jc = pPr.find(f'{W_NS}jc') if pPr is not None else None
                jc_val = jc.get(f'{W_NS}val') if jc is not None else None
                align_ok = jc_val in (None, 'left', 'start')

                if not (line_ok and align_ok):
                    all_ok = False
                    break

        passed = found_any and all_ok
        self.add_result("目录中的二级标题段落格式为1.5倍行距、左对齐", passed, 3)

    def check_toc_level3_font(self):
        """+3：目录中的三级标题文字为小四号、宋体

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目；其中以
        "阿拉伯数字 + 分隔符（. / 、 / ． / 空格）"起始的条目视为
        三级目录标题。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')
        L3_PAT = re.compile(r'^\d+\s*[\.．、\s]')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        def _run_text(r_el):
            return ''.join(t.text or '' for t in r_el.findall(f'.//{W_NS}t'))

        def _run_font(r_el):
            rpr = r_el.find(f'{W_NS}rPr')
            if rpr is None:
                return ''
            rfonts = rpr.find(f'{W_NS}rFonts')
            if rfonts is None:
                return ''
            for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                value = rfonts.get(f'{W_NS}{attr}')
                if value:
                    return value
            return ''

        def _run_size_pt(r_el):
            rpr = r_el.find(f'{W_NS}rPr')
            if rpr is None:
                return None
            sz = rpr.find(f'{W_NS}sz')
            if sz is None:
                return None
            value = sz.get(f'{W_NS}val')
            try:
                return int(value) / 2
            except (TypeError, ValueError):
                return None

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_ok = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if not L3_PAT.match(head):
                    continue

                found_any = True
                for r_el in p_el.findall(f'{W_NS}r'):
                    rt = _run_text(r_el).strip()
                    if not rt or rt.isdigit():
                        continue
                    font_name = _run_font(r_el)
                    font_ok = '宋体' in font_name or 'SimSun' in font_name
                    size_pt = _run_size_pt(r_el)
                    size_ok = size_pt is not None and abs(size_pt - 12) < 0.5
                    if not (font_ok and size_ok):
                        all_ok = False
                        break
                if not all_ok:
                    break

        passed = found_any and all_ok
        self.add_result("目录中的三级标题文字为小四号、宋体", passed, 3)
    def check_toc_level3_paragraph(self):
        """+3：目录中的三级标题段落格式为1.5倍行距、左对齐，标题后空一个字符

        细则依据（宽松路线：与 check_toc_level1/2_paragraph 保持一致）：
        本文档目录条目未绑定内建 TOC 3 样式，无法以 pStyle 定位三级
        条目。改按办公软件视觉习惯：目录条目起始为"阿拉伯数字 + 分隔符
        （. / 、 / ． / ' '（空格））"的段落即视为三级条目。对这些
        条目段落需同时满足：
          1) 1.5 倍行距：w:pPr/w:spacing 的 w:line="360" 且 w:lineRule="auto"
             （240=单倍，360=1.5倍，480=两倍）
          2) 左对齐：w:pPr/w:jc 的 val 为 "left"/"start"，或未设置（Word 默认左对齐）
          3) 标题后空一个字符：Word 自动目录中标题正文与页码之间由制表符
             分隔；本项要求标题正文与该制表符之间恰好存在 1 个空白字符
             （半角空格 U+0020 或全角空格 U+3000）。

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目；第 3 项
        通过按文档顺序拼接段内各 run 的 <w:t> 与 <w:tab/> 序列，取首个
        <w:tab/> 之前累计的标题文本来判断结尾是否恰好 1 个空白字符
        （因为 _paragraph_text 拼接后的字符串本身不保留 <w:tab/> 的
        位置信息）。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')
        L3_PAT = re.compile(r'^\d+\s*[\.．、\s]')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        def _title_text_before_tab(p_el):
            """按文档顺序拼接首个 <w:tab/> 之前的所有 <w:t> 文本。"""
            title = ''
            for r_el in p_el.findall(f'{W_NS}r'):
                for sub in r_el:
                    tag = sub.tag.replace(W_NS, '')
                    if tag == 'tab':
                        return title
                    if tag == 't':
                        title += sub.text or ''
            return title

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_ok = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if not L3_PAT.match(head):
                    continue

                found_any = True

                pPr = p_el.find(f'{W_NS}pPr')

                # 1) 1.5 倍行距
                line_ok = False
                spacing = pPr.find(f'{W_NS}spacing') if pPr is not None else None
                if spacing is not None:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line == '360' and line_rule == 'auto':
                        line_ok = True

                # 2) 左对齐
                jc = pPr.find(f'{W_NS}jc') if pPr is not None else None
                jc_val = jc.get(f'{W_NS}val') if jc is not None else None
                align_ok = jc_val in (None, 'left', 'start')

                # 3) 标题后空一个字符：首个 w:tab 之前的标题文本需以恰好
                #    1 个空白字符（半角/全角空格）结尾
                title_text = _title_text_before_tab(p_el)
                space_after_ok = bool(re.search(r'[^\s　][ 　]$', title_text))

                if not (line_ok and align_ok and space_after_ok):
                    all_ok = False
                    break

        passed = found_any and all_ok
        self.add_result("目录中的三级标题段落格式为1.5倍行距、左对齐，标题后空一个字符", passed, 3)

    def check_toc_level_format(self):
        """+3：目录中的标题层级严格按照"一"、"（一）"、"1" """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')

        L1_PAT = re.compile(r'^[一二三四五六七八九十百]+\s*[、\.．]')
        L2_PAT = re.compile(r'^[（(][一二三四五六七八九十百]+[）)]')
        L3_PAT = re.compile(r'^\d+\s*[\.．、\s]')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        passed = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                head = re.sub(r'\s*\d+\s*$', '', text).lstrip().lstrip('　')
                if not head:
                    continue

                found_any = True
                if L1_PAT.match(head) or L2_PAT.match(head) or L3_PAT.match(head):
                    continue

                passed = False
                break

        self.add_result("目录中的标题层级严格按照'一'、'（一）'、'1'", found_any and passed, 3)

    def check_toc_page_number(self):
        """+1：目录条目中的页码使用阿拉伯数字

        目录可能位于 w:sdt/w:sdtContent 内容控件内，不能依赖
        python-docx 的 doc.paragraphs 或传统 TOC 字段结束标记。直接从
        底层 XML 中扫描"目录"标题后的、含 w:tab 的目录条目，并检查
        条目文本末尾的页码是否为阿拉伯数字。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = self.doc.element.body.findall(f'.//{W_NS}p')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        toc_title_idx = None
        for i, p_el in enumerate(paragraphs):
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_title_idx = i
                break

        found_any = False
        all_ok = True
        found_toc_entry = False

        if toc_title_idx is not None:
            for p_el in paragraphs[toc_title_idx + 1:]:
                text = _paragraph_text(p_el).strip()
                has_tab = bool(p_el.findall(f'.//{W_NS}tab'))

                if not text:
                    continue

                if not has_tab:
                    if found_toc_entry:
                        break
                    continue

                found_toc_entry = True
                found_any = True

                page_num_match = re.search(r'(\d+)\s*$', text)
                if not page_num_match:
                    all_ok = False
                    break

                page_num = page_num_match.group(1)
                if not re.fullmatch(r'\d+', page_num):
                    all_ok = False
                    break

        self.add_result("目录条目中的页码使用阿拉伯数字", found_any and all_ok, 1)

    def check_toc_after_english_abstract(self):
        """-3：目录页不在英文摘要页后面

        细则依据（扣分项）：当"目录页"在文档流中出现在"英文摘要页"之前
        （即目录不在英文摘要后面）时扣 3 分。基于 Word 文件属性判定：
          - 目录页位置：以文档中首个 TOC 指令域（w:instrText 含 "TOC"）
            所在段落索引为准；这是 Word 自动目录在文件属性上的唯一标识。
          - 英文摘要页位置：以段落文本为 "Abstract"（忽略大小写与空白）
            的首个段落索引为准。
        当 TOC 索引 < Abstract 索引 时，判定"目录页不在英文摘要页后面"成立，
        触发扣分。若文档不含英文摘要或不含自动目录，则不触发扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        toc_idx = -1
        abstract_idx = -1

        for i, para in enumerate(self.doc.paragraphs):
            # 目录位置：首个 TOC 指令域
            if toc_idx == -1:
                for instr in para._element.findall(f'.//{W_NS}instrText'):
                    if instr.text and 'TOC' in instr.text:
                        toc_idx = i
                        break

            # 英文摘要位置：首个文本为 Abstract 的段落
            if abstract_idx == -1:
                text_ns = para.text.strip().replace(' ', '')
                if text_ns.lower() == 'abstract':
                    abstract_idx = i

            if toc_idx != -1 and abstract_idx != -1:
                break

        # 扣分成立：目录与英文摘要都存在，且目录在英文摘要之前
        passed = toc_idx != -1 and abstract_idx != -1 and toc_idx < abstract_idx
        self.add_result("目录页不在英文摘要页后面", passed, -3)

    def check_cover_image_size(self):
        """-1：封面页的图片大小不满足5.34×15.34

        细则依据（扣分项）：封面页图片的尺寸需满足 5.34cm × 15.34cm。
        Word 文件属性上，内嵌/浮动图片的尺寸存储于 wp:extent 元素的
        cx / cy 属性中，单位为 EMU（1cm = 360000 EMU）。
          - 5.34 cm = 1,922,400 EMU
          - 15.34 cm = 5,522,400 EMU
        本项判定：封面页范围内（从文档开始至首个"摘要/Abstract/目录"标题段落
        之前）的所有图片，其 (cx, cy) 换算为 cm 后是否构成 {5.34, 15.34} 集合
        （不区分宽高顺序）。任一图片尺寸不满足即触发扣分。
        """
        if not self.doc:
            return

        WP_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
        EMU_PER_CM = 360000.0

        # 圈定封面页范围：文档开始 → 首个 摘要/Abstract/目录 标题段落之前
        cover_end = len(self.doc.paragraphs)
        for i, para in enumerate(self.doc.paragraphs):
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if text_ns in ('摘要', '目录') or text_ns.lower() == 'abstract':
                cover_end = i
                break

        # 收集封面页范围内所有图片的尺寸
        found_any = False
        any_mismatch = False
        for para in self.doc.paragraphs[:cover_end]:
            for extent in para._element.findall(f'.//{WP_NS}extent'):
                cx = extent.get('cx')
                cy = extent.get('cy')
                if cx is None or cy is None:
                    continue
                found_any = True
                w_cm = int(cx) / EMU_PER_CM
                h_cm = int(cy) / EMU_PER_CM
                # 允许 5.34 与 15.34 顺序互换；容差 0.05cm 覆盖 Word 显示的舍入
                dims = sorted([w_cm, h_cm])
                if not (abs(dims[0] - 5.34) < 0.05 and abs(dims[1] - 15.34) < 0.05):
                    any_mismatch = True
                    break
            if any_mismatch:
                break

        # 扣分成立：存在图片且至少一张尺寸不满足；或封面页范围内未发现任何图片
        passed = (found_any and any_mismatch) or (not found_any)
        self.add_result("封面页的图片大小不满足5.34×15.34", passed, -1)

    def check_cover_image_layout(self):
        """-1：封面页的图片格式不满足嵌入型布局、两端对齐"""
        if not self.doc:
            return

        found_image = False
        layout_ok = True
        for para in self.doc.paragraphs:
            for run in para.runs:
                has_inline = bool(run._element.xpath('.//wp:inline'))
                has_anchor = bool(run._element.xpath('.//wp:anchor'))
                if has_inline or has_anchor:
                    found_image = True
                    inline_ok = has_inline and not has_anchor
                    align_ok = para.paragraph_format.alignment in (WD_ALIGN_PARAGRAPH.JUSTIFY, None)
                    if not inline_ok or not align_ok:
                        layout_ok = False

        if not found_image:
            layout_ok = False

        # 扣分项：只有封面图片格式不满足要求时才扣分。
        should_deduct = not layout_ok
        self.add_result("封面页的图片格式不满足嵌入型布局、两端对齐", should_deduct, -1)

    def check_cover_school_code_font(self):
        """-1：封面页"学校代码：12786"文本字体格式不满足宋体四号

        细则依据：定位段落文本包含"学校代码"与"12786"的封面段落，对该
        段落中每个非空 run，需同时满足：
          1) 字体为宋体（w:rPr/w:rFonts 的 eastAsia/ascii/hAnsi/cs 任一为"宋体"/SimSun）
          2) 字号为四号（w:rPr/w:sz 的半磅值 = 28，即 14 磅）
        只要任一 run 不满足，即"不满足宋体四号"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            if '学校代码' not in para.text or '12786' not in para.text:
                continue
            found = True

            for run in para.runs:
                if not run.text.strip():
                    continue

                # 1) 字体：宋体
                rpr = run._element.rPr
                font_name = ''
                if rpr is not None and rpr.rFonts is not None:
                    for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                        v = rpr.rFonts.get(f'{W_NS}{attr}')
                        if v:
                            font_name = v
                            break
                if not font_name:
                    font_name = run.font.name or ''
                font_ok = '宋体' in str(font_name) or 'SimSun' in str(font_name)

                # 2) 字号：四号 = 14 磅
                size_ok = run.font.size is not None and abs(run.font.size.pt - 14) < 0.5

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：目标段落存在且不满足宋体四号时触发
        passed = found and violated
        self.add_result("封面页'学校代码：12786'文本字体格式不满足宋体四号", passed, -1)

    def check_cover_school_code_paragraph(self):
        """-1：封面页"学校代码：12786"文本段落格式不满足右对齐、单倍行距

        细则依据：定位段落文本同时包含"学校代码"与"12786"的封面段落，
        其段落格式需同时满足：
          1) 右对齐：Word 文件属性 w:pPr/w:jc 的 val = "right"
          2) 单倍行距：Word 文件属性 w:pPr/w:spacing 的 w:line="240" 且
             w:lineRule="auto"；或未设置 w:spacing/w:line（Word 默认即
             为单倍行距）
        只要任一项不满足，即"不满足右对齐、单倍行距"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            if '学校代码' not in para.text or '12786' not in para.text:
                continue
            found = True

            pPr = para._element.find(f'{W_NS}pPr')

            # 1) 右对齐：w:pPr/w:jc val="right"
            align_ok = False
            if pPr is not None:
                jc = pPr.find(f'{W_NS}jc')
                if jc is not None and jc.get(f'{W_NS}val') == 'right':
                    align_ok = True

            # 2) 单倍行距：w:line="240" 且 w:lineRule="auto"，
            #    或未设置 w:spacing/w:line（Word 默认单倍行距）
            line_ok = False
            if pPr is None:
                line_ok = True
            else:
                spacing = pPr.find(f'{W_NS}spacing')
                if spacing is None:
                    line_ok = True
                else:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line is None:
                        line_ok = True
                    elif line == '240' and line_rule == 'auto':
                        line_ok = True

            if not (align_ok and line_ok):
                violated = True
                break

        # 扣分项：目标段落存在且不满足右对齐+单倍行距时触发
        passed = found and violated
        self.add_result("封面页'学校代码：12786'文本段落格式不满足右对齐、单倍行距", passed, -1)

    def check_cover_graduate_paper_paragraph(self):
        """-1：封面页"成人高等教育毕业论文（设计）"文本段落格式不满足居中对齐、单倍行距

        细则依据：定位段落文本同时包含"成人高等教育毕业论文"与"设计"的
        封面段落，其段落格式需同时满足：
          1) 居中对齐：Word 文件属性 w:pPr/w:jc 的 val = "center"
          2) 单倍行距：Word 文件属性 w:pPr/w:spacing 的 w:line="240" 且
             w:lineRule="auto"；或未设置 w:spacing/w:line（Word 默认即
             为单倍行距）
        只要任一项不满足，即"不满足居中对齐、单倍行距"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            if '成人高等教育毕业论文' not in para.text or '设计' not in para.text:
                continue
            found = True

            pPr = para._element.find(f'{W_NS}pPr')

            # 1) 居中对齐：w:pPr/w:jc val="center"
            align_ok = False
            if pPr is not None:
                jc = pPr.find(f'{W_NS}jc')
                if jc is not None and jc.get(f'{W_NS}val') == 'center':
                    align_ok = True

            # 2) 单倍行距：w:line="240" 且 w:lineRule="auto"，
            #    或未设置 w:spacing/w:line（Word 默认单倍行距）
            line_ok = False
            if pPr is None:
                line_ok = True
            else:
                spacing = pPr.find(f'{W_NS}spacing')
                if spacing is None:
                    line_ok = True
                else:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line is None:
                        line_ok = True
                    elif line == '240' and line_rule == 'auto':
                        line_ok = True

            if not (align_ok and line_ok):
                violated = True
                break

        # 扣分项：目标段落存在且不满足居中对齐+单倍行距时触发
        passed = found and violated
        self.add_result("封面页'成人高等教育毕业论文（设计）'文本段落格式不满足居中对齐、单倍行距", passed, -1)

    def check_cover_graduate_paper_font(self):
        """-1：封面页"成人高等教育毕业论文（设计）"文本字体格式不满足宋体，字号为32

        细则依据：定位段落文本同时包含"成人高等教育毕业论文"与"设计"的
        封面段落，对该段落中每个非空 run，需同时满足：
          1) 字体为宋体（Word 文件属性 w:rPr/w:rFonts 的 eastAsia/ascii/
             hAnsi/cs 任一为"宋体"/SimSun）
          2) 字号为 32（Word 文件属性 w:rPr/w:sz 的半磅值 = 64，即 32 磅）
        只要任一 run 不满足，即"不满足宋体，字号为32"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            if '成人高等教育毕业论文' not in para.text or '设计' not in para.text:
                continue
            found = True

            for run in para.runs:
                if not run.text.strip():
                    continue

                # 1) 字体：宋体
                rpr = run._element.rPr
                font_name = ''
                if rpr is not None and rpr.rFonts is not None:
                    for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                        v = rpr.rFonts.get(f'{W_NS}{attr}')
                        if v:
                            font_name = v
                            break
                if not font_name:
                    font_name = run.font.name or ''
                font_ok = '宋体' in str(font_name) or 'SimSun' in str(font_name)

                # 2) 字号：32 磅
                size_ok = run.font.size is not None and abs(run.font.size.pt - 32) < 0.5

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：目标段落存在且不满足宋体+字号32时触发
        passed = found and violated
        self.add_result("封面页'成人高等教育毕业论文（设计）'文本字体格式不满足宋体，字号为32", passed, -1)

    def check_cover_title_font(self):
        """-1：封面页标题“幼儿园自然探究活动融入一日生活的课程设计与应用”字体不满足华文新魏小一

        细则依据：封面页标题（论文题目）需同时满足：
          1) 字体为华文新魏（Word 文件属性 w:rPr/w:rFonts 的 eastAsia/ascii/
             hAnsi/cs 任一为"华文新魏"/STXinwei）
          2) 字号为小一（Word 文件属性 w:rPr/w:sz 的半磅值 = 48，即 24 磅）

        本文档标题分为两段："幼儿园自然探究活动融入一日生活的"和
        "课程设计与应用"。按标题文本精确定位这两段，只检查标题 run，
        避免把后续作者声明等非标题段落误当作封面标题候选。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        title_text = '幼儿园自然探究活动融入一日生活的课程设计与应用'

        title_paragraphs = []
        paragraphs = list(self.doc.paragraphs)
        for i in range(len(paragraphs)):
            first = paragraphs[i].text.strip().replace(' ', '').replace('　', '')
            if not first:
                continue
            combined = first
            matched = [paragraphs[i]]
            for j in range(i + 1, min(i + 4, len(paragraphs))):
                nxt = paragraphs[j].text.strip().replace(' ', '').replace('　', '')
                if not nxt:
                    continue
                combined += nxt
                matched.append(paragraphs[j])
                if len(combined) >= len(title_text):
                    break
            if combined == title_text:
                title_paragraphs = matched
                break

        violated = not bool(title_paragraphs)

        if title_paragraphs:
            for para in title_paragraphs:
                for run in para.runs:
                    if not run.text.strip():
                        continue

                    # 1) 字体：华文新魏
                    rpr = run._element.rPr
                    font_name = ''
                    if rpr is not None and rpr.rFonts is not None:
                        for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                            v = rpr.rFonts.get(f'{W_NS}{attr}')
                            if v:
                                font_name = v
                                break
                    if not font_name:
                        font_name = run.font.name or ''
                    font_ok = '华文新魏' in str(font_name) or 'STXinwei' in str(font_name)

                    # 2) 字号：小一 = 24 磅
                    size_ok = run.font.size is not None and abs(run.font.size.pt - 24) < 0.5

                    if not (font_ok and size_ok):
                        violated = True
                        break
                if violated:
                    break

        self.add_result("封面页标题“幼儿园自然探究活动融入一日生活的课程设计与应用”字体不满足华文新魏小一", violated, -1)

    def check_cover_title_paragraph(self):
        """-1：封面页标题"幼儿园自然探究活动融入一日生活的课程设计与应用"段落格式不满足居中对齐、单倍行距

        细则依据：封面页标题（论文题目）所在段落，需同时满足：
          1) 居中对齐：Word 文件属性 w:pPr/w:jc 的 val = "center"
          2) 单倍行距：Word 文件属性 w:pPr/w:spacing 的 w:line="240" 且
             w:lineRule="auto"；或未设置 w:spacing/w:line（Word 默认即
             为单倍行距）

        本文档标题分为两段："幼儿园自然探究活动融入一日生活的"和
        "课程设计与应用"。按标题文本精确定位这两段，只检查标题段落，
        避免把后续作者声明等非标题段落误当作封面标题候选。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        title_text = '幼儿园自然探究活动融入一日生活的课程设计与应用'

        title_paragraphs = []
        paragraphs = list(self.doc.paragraphs)
        for i in range(len(paragraphs)):
            first = paragraphs[i].text.strip().replace(' ', '').replace('　', '')
            if not first:
                continue
            combined = first
            matched = [paragraphs[i]]
            for j in range(i + 1, min(i + 4, len(paragraphs))):
                nxt = paragraphs[j].text.strip().replace(' ', '').replace('　', '')
                if not nxt:
                    continue
                combined += nxt
                matched.append(paragraphs[j])
                if len(combined) >= len(title_text):
                    break
            if combined == title_text:
                title_paragraphs = matched
                break

        violated = not bool(title_paragraphs)

        for para in title_paragraphs:
            pPr = para._element.find(f'{W_NS}pPr')

            # 1) 居中对齐：w:pPr/w:jc val="center"
            align_ok = False
            if pPr is not None:
                jc = pPr.find(f'{W_NS}jc')
                if jc is not None and jc.get(f'{W_NS}val') == 'center':
                    align_ok = True

            # 2) 单倍行距：w:line="240" 且 w:lineRule="auto"，
            #    或未设置 w:spacing/w:line（Word 默认单倍行距）
            line_ok = False
            if pPr is None:
                line_ok = True
            else:
                spacing = pPr.find(f'{W_NS}spacing')
                if spacing is None:
                    line_ok = True
                else:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line is None:
                        line_ok = True
                    elif line == '240' and line_rule == 'auto':
                        line_ok = True

            if not (align_ok and line_ok):
                violated = True
                break

        self.add_result("封面页标题\"幼儿园自然探究活动融入一日生活的课程设计与应用\"段落格式不满足居中对齐、单倍行距", violated, -1)

    def check_cover_table_second_col_paragraph(self):
        """-1：封面页的表格内第二列字体段落格式不满足居中对齐、单倍行距

        细则依据：封面页表格中每一行第二列单元格内的每个段落，需同时满足：
          1) 居中对齐：Word 文件属性 w:pPr/w:jc 的 val = "center"
          2) 单倍行距：Word 文件属性 w:pPr/w:spacing 的 w:line="240" 且
             w:lineRule="auto"；或未设置 w:spacing/w:line（Word 默认即
             为单倍行距）
        任一段落不满足上述任一项，即"不满足居中对齐、单倍行距"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for table in self.doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                second_cell = row.cells[1]
                for para in second_cell.paragraphs:
                    if not para.text.strip():
                        continue
                    found = True

                    pPr = para._element.find(f'{W_NS}pPr')

                    # 1) 居中对齐：w:pPr/w:jc val="center"
                    align_ok = False
                    if pPr is not None:
                        jc = pPr.find(f'{W_NS}jc')
                        if jc is not None and jc.get(f'{W_NS}val') == 'center':
                            align_ok = True

                    # 2) 单倍行距：w:line="240" 且 w:lineRule="auto"，
                    #    或未设置 w:spacing/w:line（Word 默认单倍行距）
                    line_ok = False
                    if pPr is None:
                        line_ok = True
                    else:
                        spacing = pPr.find(f'{W_NS}spacing')
                        if spacing is None:
                            line_ok = True
                        else:
                            line = spacing.get(f'{W_NS}line')
                            line_rule = spacing.get(f'{W_NS}lineRule')
                            if line is None:
                                line_ok = True
                            elif line == '240' and line_rule == 'auto':
                                line_ok = True

                    if not (align_ok and line_ok):
                        violated = True
                        break
                if violated:
                    break
            if violated:
                break

        # 扣分项：第二列段落存在且不满足居中对齐+单倍行距时触发
        passed = found and violated
        self.add_result("封面页的表格内第二列字体段落格式不满足居中对齐、单倍行距", passed, -1)

    def check_declaration_font(self):
        """-1："启明学院毕业论文作者声明"文本所在页的标题及内容字体不满足宋体三号

        细则依据：定位"启明学院毕业论文作者声明"标题段落，其所在页的标题
        与其后内容段落（直至"签名"/"日期"段落之前）中每个非空 run，需同时
        满足：
          1) 字体为宋体（Word 文件属性 w:rPr/w:rFonts 的 eastAsia/ascii/
             hAnsi/cs 任一为"宋体"/SimSun）
          2) 字号为三号（Word 文件属性 w:rPr/w:sz 的半磅值 = 32，即 16 磅）
        任一 run 不满足上述任一项，即"不满足宋体三号"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        in_declaration = False
        found = False
        violated = False

        for para in self.doc.paragraphs:
            text = para.text
            if '启明学院毕业论文作者声明' in text:
                in_declaration = True
            if not in_declaration:
                continue

            # 遇到"签名"或"日期"段落，声明标题与内容区结束
            if in_declaration and ('签名' in text or '日期' in text) \
                    and '启明学院毕业论文作者声明' not in text:
                break

            for run in para.runs:
                if not run.text.strip():
                    continue
                found = True

                # 1) 字体：宋体
                rpr = run._element.rPr
                font_name = ''
                if rpr is not None and rpr.rFonts is not None:
                    for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                        v = rpr.rFonts.get(f'{W_NS}{attr}')
                        if v:
                            font_name = v
                            break
                if not font_name:
                    font_name = run.font.name or ''
                font_ok = '宋体' in str(font_name) or 'SimSun' in str(font_name)

                # 2) 字号：三号 = 16 磅
                size_ok = run.font.size is not None and abs(run.font.size.pt - 16) < 0.5

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：声明标题/内容存在且不满足宋体三号时触发
        passed = found and violated
        self.add_result("'启明学院毕业论文作者声明'文本所在页的标题及内容字体不满足宋体三号", passed, -1)

    def check_declaration_title_paragraph(self):
        """-1："启明学院毕业论文作者声明"文本所在页的标题段落格式不满足居中对齐、1.5倍行距

        细则依据：定位段落文本包含"启明学院毕业论文作者声明"的标题段落，
        其段落格式需同时满足：
          1) 居中对齐：Word 文件属性 w:pPr/w:jc 的 val = "center"
          2) 1.5 倍行距：Word 文件属性 w:pPr/w:spacing 的 w:line="360" 且
             w:lineRule="auto"（240=单倍，360=1.5倍，480=两倍）
        只要任一项不满足，即"不满足居中对齐、1.5倍行距"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            if '启明学院毕业论文作者声明' not in para.text:
                continue
            found = True

            pPr = para._element.find(f'{W_NS}pPr')

            # 1) 居中对齐：w:pPr/w:jc val="center"
            align_ok = False
            if pPr is not None:
                jc = pPr.find(f'{W_NS}jc')
                if jc is not None and jc.get(f'{W_NS}val') == 'center':
                    align_ok = True

            # 2) 1.5 倍行距：w:pPr/w:spacing w:line="360" 且 w:lineRule="auto"
            line_ok = False
            if pPr is not None:
                spacing = pPr.find(f'{W_NS}spacing')
                if spacing is not None:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line == '360' and line_rule == 'auto':
                        line_ok = True

            if not (align_ok and line_ok):
                violated = True
                break

        # 扣分项：声明标题存在且不满足居中对齐+1.5倍行距时触发
        passed = found and violated
        self.add_result("'启明学院毕业论文作者声明'文本所在页的标题段落格式不满足居中对齐、1.5倍行距", passed, -1)

    def check_declaration_content_paragraph(self):
        """-1："启明学院毕业论文作者声明"文本所在页的内容段落格式不满足左对齐、1.5倍行距

        细则依据：定位"启明学院毕业论文作者声明"标题段落，其后至"签名"/
        "日期"段落之前的所有非空文本段落视为声明内容。这些段落需同时满足：
          1) 左对齐：Word 文件属性 w:pPr/w:jc 的 val 为 "left"/"start"，
             或未设置 w:jc（Word 默认即为左对齐）
          2) 1.5 倍行距：Word 文件属性 w:pPr/w:spacing 的 w:line="360" 且
             w:lineRule="auto"（240=单倍，360=1.5倍，480=两倍）
        只要任一段落不满足上述任一项，即"不满足左对齐、1.5倍行距"，
        触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        in_declaration = False
        found = False
        violated = False

        for para in self.doc.paragraphs:
            text = para.text
            if '启明学院毕业论文作者声明' in text:
                in_declaration = True
                continue
            if not in_declaration:
                continue

            # 遇到"签名"或"日期"段落，声明内容区结束
            if '签名' in text or '日期' in text:
                break

            if not text.strip():
                continue
            found = True

            pPr = para._element.find(f'{W_NS}pPr')

            # 1) 左对齐：w:pPr/w:jc val 为 left/start，或未设置
            align_ok = False
            if pPr is None:
                align_ok = True
            else:
                jc = pPr.find(f'{W_NS}jc')
                if jc is None:
                    align_ok = True
                else:
                    jc_val = jc.get(f'{W_NS}val')
                    if jc_val in (None, 'left', 'start'):
                        align_ok = True

            # 2) 1.5 倍行距：w:pPr/w:spacing w:line="360" 且 w:lineRule="auto"
            line_ok = False
            if pPr is not None:
                spacing = pPr.find(f'{W_NS}spacing')
                if spacing is not None:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line == '360' and line_rule == 'auto':
                        line_ok = True

            if not (align_ok and line_ok):
                violated = True
                break

        # 扣分项：声明内容存在且不满足左对齐+1.5倍行距时触发
        passed = found and violated
        self.add_result("'启明学院毕业论文作者声明'文本所在页的内容段落格式不满足左对齐、1.5倍行距", passed, -1)

    def check_abstract_title_font(self):
        """-1：摘要页标题的字体格式不满足黑体三号

        细则依据：定位段落文本为"摘要"（忽略空白后完全匹配）的标题段落，
        对该段落中每个非空 run，需同时满足：
          1) 字体为黑体（Word 文件属性 w:rPr/w:rFonts 的 eastAsia/ascii/
             hAnsi/cs 任一为"黑体"/SimHei）
          2) 字号为三号（Word 文件属性 w:rPr/w:sz 的半磅值 = 32，即 16 磅）
        只要任一 run 不满足上述任一项，即"不满足黑体三号"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if text_ns != '摘要':
                continue

            for run in para.runs:
                if not run.text.strip():
                    continue
                found = True

                # 1) 字体：黑体
                rpr = run._element.rPr
                font_name = ''
                if rpr is not None and rpr.rFonts is not None:
                    for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                        v = rpr.rFonts.get(f'{W_NS}{attr}')
                        if v:
                            font_name = v
                            break
                if not font_name:
                    font_name = run.font.name or ''
                font_ok = '黑体' in str(font_name) or 'SimHei' in str(font_name)

                # 2) 字号：三号 = 16 磅
                size_ok = run.font.size is not None and abs(run.font.size.pt - 16) < 0.5

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated or found:
                break

        # 扣分项：摘要标题存在且不满足黑体三号时触发
        passed = found and violated
        self.add_result("摘要页标题的字体格式不满足黑体三号", passed, -1)

    def check_abstract_title_paragraph(self):
        """-1：摘要页标题的段落格式不满足段前一行、段后一行，两倍行距，居中对齐

        细则依据：定位段落文本为"摘要"（忽略空白后完全匹配）的标题段落，
        其段落格式需同时满足（全部读 Word 文件属性 w:pPr）：
          1) 段前 1 行：w:pPr/w:spacing 的 w:beforeLines="100"
             （单位为 1/100 行，100=1 行）
          2) 段后 1 行：w:pPr/w:spacing 的 w:afterLines="100"
          3) 两倍行距：w:pPr/w:spacing 的 w:line="480" 且 w:lineRule="auto"
             （240=单倍，360=1.5倍，480=两倍）
          4) 居中对齐：w:pPr/w:jc 的 val = "center"
        只要任一项不满足，即"不满足段前一行、段后一行，两倍行距，居中对齐"，
        触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if text_ns != '摘要':
                continue
            found = True

            pPr = para._element.find(f'{W_NS}pPr')
            spacing = pPr.find(f'{W_NS}spacing') if pPr is not None else None

            # 1) 段前 1 行
            before_ok = False
            # 2) 段后 1 行
            after_ok = False
            # 3) 两倍行距
            line_ok = False
            if spacing is not None:
                if spacing.get(f'{W_NS}beforeLines') == '100':
                    before_ok = True
                if spacing.get(f'{W_NS}afterLines') == '100':
                    after_ok = True
                line = spacing.get(f'{W_NS}line')
                line_rule = spacing.get(f'{W_NS}lineRule')
                if line == '480' and line_rule == 'auto':
                    line_ok = True

            # 4) 居中对齐
            align_ok = False
            if pPr is not None:
                jc = pPr.find(f'{W_NS}jc')
                if jc is not None and jc.get(f'{W_NS}val') == 'center':
                    align_ok = True

            if not (before_ok and after_ok and line_ok and align_ok):
                violated = True
            break

        # 扣分项：摘要标题存在且不满足任一子项时触发
        passed = found and violated
        self.add_result("摘要页标题的段落格式不满足段前一行、段后一行，两倍行距，居中对齐", passed, -1)

    def check_abstract_content_paragraph(self):
        """-1：摘要页的内容段落格式不满足首行缩进两字符、两端对齐、1.5倍行距

        细则依据：定位"摘要"标题段落之后至"关键词"/"目录"段落之前的所有
        非空文本段落视为摘要内容。这些段落需同时满足（全部读 Word 文件
        属性 w:pPr）：
          1) 首行缩进两字符：w:pPr/w:ind 的 w:firstLineChars="200"
             （单位为 1/100 字符，200=2 字符）
          2) 两端对齐：w:pPr/w:jc 的 val = "both"（等价别名 "justify" 一并接受）；
             或未设置 w:jc（Word/WPS 正文默认即为两端对齐）
          3) 1.5 倍行距：w:pPr/w:spacing 的 w:line="360" 且 w:lineRule="auto"
             （240=单倍，360=1.5倍，480=两倍）
        只要任一段落不满足上述任一项，即"不满足首行缩进两字符、两端对齐、
        1.5倍行距"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        in_abstract = False
        found = False
        violated = False

        for para in self.doc.paragraphs:
            text = para.text
            text_ns = text.strip().replace(' ', '').replace('　', '')

            if text_ns == '摘要':
                in_abstract = True
                continue
            if not in_abstract:
                continue

            # 遇到"关键词"或"目录"段落，摘要内容区结束
            if '关键词' in text or text_ns == '目录':
                break

            if not text.strip():
                continue
            found = True

            pPr = para._element.find(f'{W_NS}pPr')

            # 1) 首行缩进两字符：w:pPr/w:ind w:firstLineChars="200"
            indent_ok = False
            if pPr is not None:
                ind = pPr.find(f'{W_NS}ind')
                if ind is not None and ind.get(f'{W_NS}firstLineChars') == '200':
                    indent_ok = True

            # 2) 两端对齐：w:pPr/w:jc val="both"（或等价别名 "justify"）；
            #    未设置 w:jc 时，Word/WPS 正文默认按两端对齐渲染，视为满足
            align_ok = False
            if pPr is None:
                align_ok = True
            else:
                jc = pPr.find(f'{W_NS}jc')
                if jc is None:
                    align_ok = True
                elif jc.get(f'{W_NS}val') in ('both', 'justify'):
                    align_ok = True

            # 3) 1.5 倍行距：w:pPr/w:spacing w:line="360" 且 w:lineRule="auto"
            line_ok = False
            if pPr is not None:
                spacing = pPr.find(f'{W_NS}spacing')
                if spacing is not None:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line == '360' and line_rule == 'auto':
                        line_ok = True

            if not (indent_ok and align_ok and line_ok):
                violated = True
                break

        # 扣分项：摘要内容存在且不满足任一子项时触发
        passed = found and violated
        self.add_result("摘要页的内容段落格式不满足首行缩进两字符、两端对齐、1.5倍行距", passed, -1)

    def check_abstract_keywords_font(self):
        """-1：摘要页的关键词及冒号后内容字体格式不满足黑体四号

        细则依据：定位段落文本包含"关键词"的摘要页段落，对该段落中每个
        非空 run，需同时满足：
          1) 字体为黑体（Word 文件属性 w:rPr/w:rFonts 的 eastAsia/ascii/
             hAnsi/cs 任一为"黑体"/SimHei）
          2) 字号为四号（Word 文件属性 w:rPr/w:sz 的半磅值 = 28，即 14 磅）
        只要任一 run 不满足上述任一项，即"不满足黑体四号"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        found = False
        violated = False

        for para in self.doc.paragraphs:
            if '关键词' not in para.text:
                continue

            for run in para.runs:
                if not run.text.strip():
                    continue
                found = True

                # 1) 字体：黑体
                rpr = run._element.rPr
                font_name = ''
                if rpr is not None and rpr.rFonts is not None:
                    for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                        v = rpr.rFonts.get(f'{W_NS}{attr}')
                        if v:
                            font_name = v
                            break
                if not font_name:
                    font_name = run.font.name or ''
                font_ok = '黑体' in str(font_name) or 'SimHei' in str(font_name)

                # 2) 字号：四号 = 14 磅
                size_ok = run.font.size is not None and abs(run.font.size.pt - 14) < 0.5

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：关键词段落存在且不满足黑体四号时触发
        passed = found and violated
        self.add_result("摘要页的关键词及冒号后内容字体格式不满足黑体四号", passed, -1)

    def check_abstract_keywords_paragraph(self):
        """-1：摘要页的关键词及其冒号后内容段落格式不满足单倍行距、左对齐

        细则依据（只判定"单倍行距"和"左对齐"两点，不加任何额外约束）：
        定位段落文本包含"关键词"的摘要页段落，判定其在办公软件（Word/WPS）
        中实际生效的段落属性——即沿"段落直接属性 → w:pStyle 样式链 →
        w:docDefaults"逐级回退取到的最终值。
          1) 单倍行距：effective w:spacing 满足 w:line="240" 且
             w:lineRule="auto"；或整条链上均未设置 w:line
             （Word 默认即为单倍行距）
          2) 左对齐：effective w:jc 的 val 为 "left"/"start"；或整条链上
             均未设置 w:jc（Word 默认即为左对齐）
        只要任一项不满足，即"不满足单倍行距、左对齐"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective(para, child_tag):
            """沿"段落直接 pPr → pStyle 链 → docDefaults/pPrDefault/pPr"
            回退查找指定 pPr 子元素（如 w:jc、w:spacing），返回首个存在者。
            """
            # 1) 段落直接属性
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None:
                el = pPr.find(f'{W_NS}{child_tag}')
                if el is not None:
                    return el
            # 2) 段落所引用样式 → basedOn 链
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(f'{W_NS}{child_tag}')
                    if el is not None:
                        return el
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            # 3) docDefaults/pPrDefault/pPr
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(f'{W_NS}{child_tag}')
                        if el is not None:
                            return el
            return None

        found = False
        violated = False

        for para in self.doc.paragraphs:
            if '关键词' not in para.text:
                continue
            found = True

            # 1) 单倍行距：effective w:spacing 的 w:line="240" 且
            #    w:lineRule="auto"；或全链未设置 w:line（Word 默认单倍）
            spacing = _find_effective(para, 'spacing')
            if spacing is None:
                line_ok = True
            else:
                line = spacing.get(f'{W_NS}line')
                line_rule = spacing.get(f'{W_NS}lineRule')
                if line is None:
                    line_ok = True
                elif line == '240' and line_rule == 'auto':
                    line_ok = True
                else:
                    line_ok = False

            # 2) 左对齐：effective w:jc val 为 left/start；或全链未设置
            jc = _find_effective(para, 'jc')
            if jc is None:
                align_ok = True
            else:
                jc_val = jc.get(f'{W_NS}val')
                align_ok = jc_val in (None, 'left', 'start')

            if not (line_ok and align_ok):
                violated = True
                break

        # 扣分项：关键词段落存在且不满足单倍行距+左对齐时触发
        passed = found and violated
        self.add_result("摘要页的关键词及其冒号后内容段落格式不满足单倍行距、左对齐", passed, -1)

    def check_english_abstract_title_font(self):
        """-1：英文摘要页的标题字体格式不满足Times New Roman三号

        细则依据（只判定"Times New Roman"和"三号"两点，不加任何额外约束）：
        定位段落文本为"Abstract"（忽略空白与大小写后完全匹配）的英文
        摘要标题段落，对该段落中每个非空 run，判定其在办公软件（Word/WPS）
        中实际生效的字体属性——沿"run 直接 rPr → 段落 pPr/rPr →
        w:pStyle 样式链（含 basedOn）→ w:docDefaults/rPrDefault/rPr"
        逐级回退取值：
          1) 字体为 Times New Roman：effective w:rFonts 的 w:ascii 值
             为 "Times New Roman"（英文标题使用 ascii 字体槽渲染）
          2) 字号为三号：effective w:sz 的半磅值 = 32（即 16 磅）
        只要任一 run 不满足上述任一项，即"不满足 Times New Roman 三号"，
        触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _effective_rpr_value(run, para, getter):
            """沿"run 直接 rPr → 段落 pPr/rPr → 段落 pStyle 链的 rPr →
            docDefaults/rPrDefault/rPr"回退，返回 getter(rPr) 的首个非空值。
            """
            # 1) run 直接 rPr
            r_rpr = run._element.find(f'{W_NS}rPr')
            if r_rpr is not None:
                v = getter(r_rpr)
                if v is not None:
                    return v
            # 2) 段落 pPr/rPr（段落级 run 属性）
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None:
                p_rpr = pPr.find(f'{W_NS}rPr')
                if p_rpr is not None:
                    v = getter(p_rpr)
                    if v is not None:
                        return v
            # 3) pStyle → basedOn 链上的 rPr
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_rpr = style.find(f'{W_NS}rPr')
                if s_rpr is not None:
                    v = getter(s_rpr)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            # 4) docDefaults/rPrDefault/rPr
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(f'{W_NS}rPrDefault')
                if rPrDefault is not None:
                    d_rpr = rPrDefault.find(f'{W_NS}rPr')
                    if d_rpr is not None:
                        v = getter(d_rpr)
                        if v is not None:
                            return v
            return None

        def _get_ascii_font(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            if rFonts is None:
                return None
            return rFonts.get(f'{W_NS}ascii')

        def _get_sz(rpr):
            sz = rpr.find(f'{W_NS}sz')
            if sz is None:
                return None
            return sz.get(f'{W_NS}val')

        found = False
        violated = False

        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if text_ns.lower() != 'abstract':
                continue
            found = True

            for run in para.runs:
                if not run.text.strip():
                    continue

                # 1) 字体：Times New Roman（ascii 字体槽）
                ascii_font = _effective_rpr_value(run, para, _get_ascii_font) or ''
                font_ok = ascii_font == 'Times New Roman'

                # 2) 字号：三号 = 半磅 32（16 磅）
                sz_val = _effective_rpr_value(run, para, _get_sz)
                size_ok = sz_val == '32'

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：英文摘要标题存在且不满足 Times New Roman 三号时触发
        passed = found and violated
        self.add_result("英文摘要页的标题字体格式不满足Times New Roman三号", passed, -1)

    def check_english_abstract_title_paragraph(self):
        """-1：英文摘要页的标题段落格式不满足段前一行、段后一行，两倍行距，居中对齐

        细则依据（只判定"段前 1 行、段后 1 行、两倍行距、居中对齐"四点，
        不加任何额外约束）：定位段落文本为"Abstract"（忽略空白与大小写后
        完全匹配）的英文摘要标题段落，判定其在办公软件（Word/WPS）中实际
        生效的段落属性——沿"段落直接 pPr → w:pStyle 样式链（含 basedOn）→
        w:docDefaults/pPrDefault/pPr"逐级回退取值：
          1) 段前 1 行：effective w:spacing 的 w:beforeLines="100"
             （单位为 1/100 行）
          2) 段后 1 行：effective w:spacing 的 w:afterLines="100"
          3) 两倍行距：effective w:spacing 的 w:line="480" 且
             w:lineRule="auto"（240=单倍，360=1.5倍，480=两倍）
          4) 居中对齐：effective w:jc 的 val = "center"
        只要任一项不满足，即"不满足段前一行、段后一行，两倍行距，居中对齐"，
        触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective(para, child_tag):
            """沿"段落直接 pPr → pStyle 链 → docDefaults/pPrDefault/pPr"
            回退查找指定 pPr 子元素，返回首个存在者。
            """
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None:
                el = pPr.find(f'{W_NS}{child_tag}')
                if el is not None:
                    return el
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(f'{W_NS}{child_tag}')
                    if el is not None:
                        return el
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(f'{W_NS}{child_tag}')
                        if el is not None:
                            return el
            return None

        found = False
        violated = False

        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if text_ns.lower() != 'abstract':
                continue
            found = True

            spacing = _find_effective(para, 'spacing')
            jc = _find_effective(para, 'jc')

            # 1) 段前 1 行
            before_ok = spacing is not None and spacing.get(f'{W_NS}beforeLines') == '100'
            # 2) 段后 1 行
            after_ok = spacing is not None and spacing.get(f'{W_NS}afterLines') == '100'
            # 3) 两倍行距
            line_ok = (
                spacing is not None
                and spacing.get(f'{W_NS}line') == '480'
                and spacing.get(f'{W_NS}lineRule') == 'auto'
            )
            # 4) 居中对齐
            align_ok = jc is not None and jc.get(f'{W_NS}val') == 'center'

            if not (before_ok and after_ok and line_ok and align_ok):
                violated = True
            break

        # 扣分项：英文摘要标题存在且不满足任一子项时触发
        passed = found and violated
        self.add_result("英文摘要页的标题段落格式不满足段前一行、段后一行，两倍行距，居中对齐", passed, -1)

    def check_english_abstract_content_paragraph(self):
        """-1：英文摘要页的内容段落格式不满足首行缩进两字符、两端对齐、1.5倍行距

        细则依据（只判定"首行缩进两字符、两端对齐、1.5 倍行距"三点，
        不加任何额外约束）：定位英文摘要"Abstract"标题段落之后至
        "Keywords"/"Key words"/"目录"段落之前的所有非空文本段落视为英文
        摘要内容。判定其在办公软件（Word/WPS）中实际生效的段落属性——
        沿"段落直接 pPr → w:pStyle 样式链（含 basedOn）→
        w:docDefaults/pPrDefault/pPr"逐级回退取值：
          1) 首行缩进两字符：effective w:ind 的 w:firstLineChars="200"
             （单位为 1/100 字符，200=2 字符）
          2) 两端对齐：effective w:jc 的 val = "both"（等价别名 "justify"
             一并接受）
          3) 1.5 倍行距：effective w:spacing 的 w:line="360" 且
             w:lineRule="auto"（240=单倍，360=1.5倍，480=两倍）
        只要任一段落不满足上述任一项，即"不满足首行缩进两字符、两端对齐、
        1.5倍行距"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective(para, child_tag):
            """沿"段落直接 pPr → pStyle 链 → docDefaults/pPrDefault/pPr"
            回退查找指定 pPr 子元素，返回首个存在者。
            """
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None:
                el = pPr.find(f'{W_NS}{child_tag}')
                if el is not None:
                    return el
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(f'{W_NS}{child_tag}')
                    if el is not None:
                        return el
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(f'{W_NS}{child_tag}')
                        if el is not None:
                            return el
            return None

        in_abstract = False
        found = False
        violated = False

        for para in self.doc.paragraphs:
            text = para.text
            text_ns = text.strip().replace(' ', '').replace('　', '')

            if text_ns.lower() == 'abstract':
                in_abstract = True
                continue
            if not in_abstract:
                continue

            # 遇到 Keywords / Key words / 目录 段落，英文摘要内容区结束
            if 'Keywords' in text or 'Key words' in text or text_ns == '目录':
                break

            if not text.strip():
                continue
            found = True

            # 1) 首行缩进两字符：effective w:ind/@w:firstLineChars == "200"
            ind = _find_effective(para, 'ind')
            indent_ok = ind is not None and ind.get(f'{W_NS}firstLineChars') == '200'

            # 2) 两端对齐：effective w:jc/@w:val ∈ {"both", "justify"}
            jc = _find_effective(para, 'jc')
            align_ok = jc is not None and jc.get(f'{W_NS}val') in ('both', 'justify')

            # 3) 1.5 倍行距：effective w:spacing/@w:line=="360" 且
            #    @w:lineRule=="auto"
            spacing = _find_effective(para, 'spacing')
            line_ok = (
                spacing is not None
                and spacing.get(f'{W_NS}line') == '360'
                and spacing.get(f'{W_NS}lineRule') == 'auto'
            )

            if not (indent_ok and align_ok and line_ok):
                violated = True
                break

        # 扣分项：英文摘要内容存在且不满足任一子项时触发
        passed = found and violated
        self.add_result("英文摘要页的内容段落格式不满足首行缩进两字符、两端对齐、1.5倍行距", passed, -1)

    def check_english_abstract_keywords_paragraph(self):
        """-1：英文摘要页的关键词及其冒号后内容段落格式不满足首行缩进两字符、左对齐、1.5倍行距

        细则依据（只判定"首行缩进两字符、左对齐、1.5 倍行距"三点，不加
        任何额外约束）：定位段落文本包含"Keywords"或"Key words"的英文
        摘要段落，判定其在办公软件（Word/WPS）中实际生效的段落属性——
        沿"段落直接 pPr → w:pStyle 样式链（含 basedOn）→
        w:docDefaults/pPrDefault/pPr"逐级回退取值：
          1) 首行缩进两字符：effective w:ind 的 w:firstLineChars="200"
             （单位为 1/100 字符，200=2 字符）
          2) 左对齐：effective w:jc 的 val 为 "left"/"start"；或整条链上
             均未设置 w:jc（Word 默认即为左对齐）
          3) 1.5 倍行距：effective w:spacing 的 w:line="360" 且
             w:lineRule="auto"（240=单倍，360=1.5倍，480=两倍）
        只要任一项不满足，即"不满足首行缩进两字符、左对齐、1.5倍行距"，
        触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective(para, child_tag):
            """沿"段落直接 pPr → pStyle 链 → docDefaults/pPrDefault/pPr"
            回退查找指定 pPr 子元素，返回首个存在者。
            """
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None:
                el = pPr.find(f'{W_NS}{child_tag}')
                if el is not None:
                    return el
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(f'{W_NS}{child_tag}')
                    if el is not None:
                        return el
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(f'{W_NS}{child_tag}')
                        if el is not None:
                            return el
            return None

        found = False
        violated = False

        for para in self.doc.paragraphs:
            if 'Keywords' not in para.text and 'Key words' not in para.text:
                continue
            found = True

            # 1) 首行缩进两字符：effective w:ind/@w:firstLineChars == "200"
            ind = _find_effective(para, 'ind')
            indent_ok = ind is not None and ind.get(f'{W_NS}firstLineChars') == '200'

            # 2) 左对齐：effective w:jc/@w:val ∈ {"left", "start"}；
            #    或整条链未设置 w:jc（Word 默认左对齐）
            jc = _find_effective(para, 'jc')
            if jc is None:
                align_ok = True
            else:
                jc_val = jc.get(f'{W_NS}val')
                align_ok = jc_val in (None, 'left', 'start')

            # 3) 1.5 倍行距：effective w:spacing/@w:line=="360" 且
            #    @w:lineRule=="auto"
            spacing = _find_effective(para, 'spacing')
            line_ok = (
                spacing is not None
                and spacing.get(f'{W_NS}line') == '360'
                and spacing.get(f'{W_NS}lineRule') == 'auto'
            )

            if not (indent_ok and align_ok and line_ok):
                violated = True
                break

        # 扣分项：英文摘要 Keywords 段落存在且不满足任一子项时触发
        passed = found and violated
        self.add_result("英文摘要页的关键词及其冒号后内容段落格式不满足首行缩进两字符、左对齐、1.5倍行距", passed, -1)

    def check_english_abstract_keywords_font(self):
        """-1：英文摘要页的关键词及冒号后内容字体格式不满足Times New Roman四号

        细则依据（只判定"Times New Roman"和"四号"两点，不加任何额外
        约束）：定位段落文本包含"Keywords"或"Key words"的英文摘要段落，
        对该段落中每个非空 run，判定其在办公软件（Word/WPS）中实际生效
        的字体属性——沿"run 直接 rPr → 段落 pPr/rPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/rPrDefault/rPr"逐级回退取值：
          1) 字体为 Times New Roman：effective w:rFonts 的 w:ascii 值
             为 "Times New Roman"（英文使用 ascii 字体槽渲染）
          2) 字号为四号：effective w:sz 的半磅值 = 28（即 14 磅）
        只要任一 run 不满足上述任一项，即"不满足 Times New Roman 四号"，
        触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _effective_rpr_value(run, para, getter):
            """沿"run 直接 rPr → 段落 pPr/rPr → pStyle 链 rPr →
            docDefaults/rPrDefault/rPr"回退，返回 getter(rPr) 的首个非空值。
            """
            r_rpr = run._element.find(f'{W_NS}rPr')
            if r_rpr is not None:
                v = getter(r_rpr)
                if v is not None:
                    return v
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None:
                p_rpr = pPr.find(f'{W_NS}rPr')
                if p_rpr is not None:
                    v = getter(p_rpr)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_rpr = style.find(f'{W_NS}rPr')
                if s_rpr is not None:
                    v = getter(s_rpr)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(f'{W_NS}rPrDefault')
                if rPrDefault is not None:
                    d_rpr = rPrDefault.find(f'{W_NS}rPr')
                    if d_rpr is not None:
                        v = getter(d_rpr)
                        if v is not None:
                            return v
            return None

        def _get_ascii_font(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            if rFonts is None:
                return None
            return rFonts.get(f'{W_NS}ascii')

        def _get_sz(rpr):
            sz = rpr.find(f'{W_NS}sz')
            if sz is None:
                return None
            return sz.get(f'{W_NS}val')

        found = False
        violated = False

        for para in self.doc.paragraphs:
            if 'Keywords' not in para.text and 'Key words' not in para.text:
                continue

            for run in para.runs:
                if not run.text.strip():
                    continue
                found = True

                # 1) 字体：Times New Roman（ascii 字体槽）
                ascii_font = _effective_rpr_value(run, para, _get_ascii_font) or ''
                font_ok = ascii_font == 'Times New Roman'

                # 2) 字号：四号 = 半磅 28（14 磅）
                sz_val = _effective_rpr_value(run, para, _get_sz)
                size_ok = sz_val == '28'

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：英文摘要 Keywords 段落存在且不满足 Times New Roman 四号时触发
        passed = found and violated
        self.add_result("英文摘要页的关键词及冒号后内容字体格式不满足Times New Roman四号", passed, -1)

    def check_header_present(self):
        """-3：目录页后出现任意一页没有页眉

        细则依据（只判定"目录页后是否存在没有页眉的页面"这一点，不加
        任何额外约束）：

        注意：不能使用 python-docx 的 doc.sections 高层接口——该文档中
        doc.sections 的条目数量与底层 XML 中实际的 <w:sectPr> 元素数量
        不一致（python-docx 在特定结构下会去重/合并分节标记，导致索引
        整体错位），会造成"目录后紧邻的、实际声明了内容页眉的分节"被
        跳过误判。因此直接遍历 self.doc.element.body 中所有 <w:sectPr>
        （包含 w:p/w:pPr 内嵌的分节标记及 body 末尾的最后一个），按其
        文档顺序索引定位目录段落所在的 sectPr，再逐一检查其后每个
        sectPr 在办公软件（Word/WPS）中实际会显示的页眉：
          1) 默认页眉（default）：分节必然会使用；从当前 sectPr 起向前
             回溯，取最近一个显式声明 headerReference type="default" 的
             sectPr，解析其引用的 header 部件内容。
          2) 首页页眉（first）：仅当该 sectPr 存在 w:titlePg 时才生效，
             同样向前回溯取最近一个显式声明 type="first" 的 sectPr。
          3) 偶数页页眉（even）：仅当 word/settings.xml 声明了
             w:evenAndOddHeaders 时才生效，同样向前回溯取最近一个显式
             声明 type="even" 的 sectPr。
        页眉"内容非空"的判定：header 部件的 XML 树中存在任一非空
        w:t 文本节点，或存在 w:drawing/w:pict/w:tbl 元素之一。
        任一"该分节会使用到"的页眉类型内容为空，即"存在页面没有页眉"，
        触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
        body = self.doc.element.body

        # 1) 按文档顺序取出所有 sectPr（分节边界标记）
        all_sectprs = body.findall(f'.//{W_NS}sectPr')

        # 2) 定位"目录"段落所属的 sectPr 序号（0 起）
        paragraphs = body.findall(f'.//{W_NS}p')

        def _paragraph_text(p_el):
            return ''.join(t.text or '' for t in p_el.findall(f'.//{W_NS}t'))

        toc_sectpr_idx = None
        sectpr_seen = 0
        for p_el in paragraphs:
            text_ns = _paragraph_text(p_el).strip().replace(' ', '').replace('　', '')
            pPr = p_el.find(f'{W_NS}pPr')
            has_sectpr = pPr is not None and pPr.find(f'{W_NS}sectPr') is not None
            if text_ns == '目录' and toc_sectpr_idx is None:
                toc_sectpr_idx = sectpr_seen
            if has_sectpr:
                sectpr_seen += 1

        # 3) settings 中是否声明 evenAndOddHeaders（偶数页页眉是否启用）
        settings_el = self.doc.settings.element
        even_and_odd = settings_el.find(f'{W_NS}evenAndOddHeaders') is not None

        # 4) 解析 r:id 对应的部件
        def _resolve_part(rid):
            try:
                return self.doc.part.rels[rid].target_part
            except KeyError:
                return None

        # 5) 从 sectpr_idx 向前回溯，取最近一个显式声明该 kind 的 headerReference
        def _effective_header_part(sectpr_idx, kind):
            idx = sectpr_idx
            while idx >= 0:
                sectPr = all_sectprs[idx]
                ref = None
                for r in sectPr.findall(f'{W_NS}headerReference'):
                    if r.get(f'{W_NS}type') == kind:
                        ref = r
                        break
                if ref is not None:
                    rid = ref.get(f'{R_NS}id')
                    return _resolve_part(rid)
                idx -= 1
            return None

        # 6) 判定 header 部件是否含内容
        def _has_header_content(part):
            if part is None:
                return False
            el = part.element
            for t in el.iter(f'{W_NS}t'):
                if t.text and t.text.strip():
                    return True
            for tag in ('drawing', 'pict', 'tbl'):
                if el.find(f'.//{W_NS}{tag}') is not None:
                    return True
            return False

        # 7) 逐个"目录后"sectPr 检查
        violated = False
        start_idx = (toc_sectpr_idx + 1) if toc_sectpr_idx is not None else 0
        for i in range(start_idx, len(all_sectprs)):
            sectPr = all_sectprs[i]
            title_pg = sectPr.find(f'{W_NS}titlePg') is not None

            # 默认页眉：始终检查
            if not _has_header_content(_effective_header_part(i, 'default')):
                violated = True
                break
            # 首页页眉：仅 titlePg 生效时
            if title_pg:
                if not _has_header_content(_effective_header_part(i, 'first')):
                    violated = True
                    break
            # 偶数页页眉：仅 settings 声明 evenAndOddHeaders 时
            if even_and_odd:
                if not _has_header_content(_effective_header_part(i, 'even')):
                    violated = True
                    break

        # 扣分项：目录页后存在没有页眉的页面
        self.add_result("目录页后出现任意一页没有页眉", violated, -3)

    def check_header_font(self):
        """-1：页眉字体格式不是宋体五号

        细则依据（只判定"宋体"和"五号"两点，不加任何额外约束）：对文档
        中所有分节实际会渲染到页面上的页眉——default 页眉始终生效；
        first 页眉仅当分节 sectPr 存在 w:titlePg 时生效；even 页眉仅当
        word/settings.xml 存在 w:evenAndOddHeaders 时生效——沿分节
        is_linked_to_previous 链回溯到显式声明该 headerReference 的
        header 部件，遍历其中所有非空 run，判定其在办公软件（Word/WPS）
        中实际生效的字体属性——沿"run 直接 rPr → 段落 pPr/rPr → w:pStyle
        样式链（含 basedOn）→ w:docDefaults/rPrDefault/rPr"逐级回退取值：
          1) 字体为宋体：run.text 中的中文字符（CJK 统一表意区）按
             effective w:rFonts/@w:eastAsia 判定；非 CJK 字符按
             effective w:rFonts/@w:ascii 判定；被判定的槽位值需为
             "宋体" 或 "SimSun"
          2) 字号为五号：effective w:sz 的半磅值 = 21（即 10.5 磅）
        只要任一非空 run 不满足上述任一项，即"不是宋体五号"，触发 -1
        扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _effective_rpr_value(run_el, para_el, getter):
            """沿"run 直接 rPr → 段落 pPr/rPr → pStyle 链 rPr →
            docDefaults/rPrDefault/rPr"回退。
            """
            r_rpr = run_el.find(f'{W_NS}rPr')
            if r_rpr is not None:
                v = getter(r_rpr)
                if v is not None:
                    return v
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                p_rpr = pPr.find(f'{W_NS}rPr')
                if p_rpr is not None:
                    v = getter(p_rpr)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_rpr = style.find(f'{W_NS}rPr')
                if s_rpr is not None:
                    v = getter(s_rpr)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(f'{W_NS}rPrDefault')
                if rPrDefault is not None:
                    d_rpr = rPrDefault.find(f'{W_NS}rPr')
                    if d_rpr is not None:
                        v = getter(d_rpr)
                        if v is not None:
                            return v
            return None

        def _get_east_asia(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}eastAsia') if rFonts is not None else None

        def _get_ascii(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}ascii') if rFonts is not None else None

        def _get_sz(rpr):
            sz = rpr.find(f'{W_NS}sz')
            return sz.get(f'{W_NS}val') if sz is not None else None

        def _is_song(name):
            return name in ('宋体', 'SimSun')

        def _is_cjk(ch):
            cp = ord(ch)
            return (
                0x4E00 <= cp <= 0x9FFF   # CJK 统一表意
                or 0x3400 <= cp <= 0x4DBF  # 扩展 A
                or 0xF900 <= cp <= 0xFAFF  # 兼容
                or 0x3000 <= cp <= 0x303F  # CJK 符号和标点
                or 0xFF00 <= cp <= 0xFFEF  # 全角
            )

        # 定位所有实际生效的 header 部件根元素
        settings_el = self.doc.settings.element
        even_and_odd = settings_el.find(f'{W_NS}evenAndOddHeaders') is not None
        sections = self.doc.sections

        def _effective_header_element(idx, kind):
            while idx >= 0:
                sec = sections[idx]
                if kind == 'default':
                    h = sec.header
                elif kind == 'first':
                    h = sec.first_page_header
                else:
                    h = sec.even_page_header
                if not h.is_linked_to_previous:
                    try:
                        return h.part.element
                    except AttributeError:
                        return None
                idx -= 1
            return None

        header_roots = []
        seen = set()
        for i, section in enumerate(sections):
            sectPr = section._sectPr
            title_pg = (
                sectPr is not None
                and sectPr.find(f'{W_NS}titlePg') is not None
            )
            kinds = ['default']
            if title_pg:
                kinds.append('first')
            if even_and_odd:
                kinds.append('even')
            for k in kinds:
                el = _effective_header_element(i, k)
                if el is not None and id(el) not in seen:
                    seen.add(id(el))
                    header_roots.append(el)

        found = False
        violated = False

        for root in header_roots:
            for para_el in root.iter(f'{W_NS}p'):
                for run_el in para_el.findall(f'{W_NS}r'):
                    text = ''.join(
                        (t.text or '') for t in run_el.findall(f'{W_NS}t')
                    )
                    if not text.strip():
                        continue
                    found = True

                    # 1) 字体：按 run 中字符类别选择字体槽位
                    has_cjk = any(_is_cjk(c) for c in text)
                    has_ascii = any(not _is_cjk(c) and not c.isspace() for c in text)
                    font_ok = True
                    if has_cjk:
                        ea = _effective_rpr_value(run_el, para_el, _get_east_asia) or ''
                        if not _is_song(ea):
                            font_ok = False
                    if font_ok and has_ascii:
                        asc = _effective_rpr_value(run_el, para_el, _get_ascii) or ''
                        if not _is_song(asc):
                            font_ok = False

                    # 2) 字号：五号 = 半磅 21（10.5 磅）
                    sz_val = _effective_rpr_value(run_el, para_el, _get_sz)
                    size_ok = sz_val == '21'

                    if not (font_ok and size_ok):
                        violated = True
                        break
                if violated:
                    break
            if violated:
                break

        # 扣分项：存在生效页眉且任一非空 run 不满足宋体五号
        passed = found and violated
        self.add_result("页眉字体格式不是宋体五号", passed, -1)

    def check_header_paragraph(self):
        """-1：页眉段落格式不满足页眉下无横线、居中对齐、单倍行距

        细则依据（只判定"页眉下无横线、居中对齐、单倍行距"三点，不加
        任何额外约束）：对文档中所有实际会渲染到页面上的页眉——default
        页眉始终生效；first 页眉仅当分节 sectPr 存在 w:titlePg 时生效；
        even 页眉仅当 word/settings.xml 存在 w:evenAndOddHeaders 时生效
        ——沿分节 is_linked_to_previous 链回溯到显式声明该 headerReference
        的 header 部件，遍历其中所有段落，判定其在办公软件（Word/WPS）
        中实际生效的段落属性——沿"段落直接 pPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/pPrDefault/pPr"逐级回退取值：
          1) 页眉下无横线：effective w:pBdr/w:bottom 缺失，或其
             w:val ∈ {"nil", "none"}，或其 w:sz == "0"
          2) 居中对齐：effective w:jc 的 val = "center"
          3) 单倍行距：effective w:spacing 的 w:line="240" 且
             w:lineRule="auto"；或整条链上均未设置 w:line
             （Word 默认即为单倍行距）
        只要任一段落不满足上述任一项，即"不满足页眉下无横线、居中对齐、
        单倍行距"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective(para_el, child_tag):
            """沿"段落直接 pPr → pStyle 链 → docDefaults/pPrDefault/pPr"
            回退查找指定 pPr 子元素，返回首个存在者。
            """
            pPr = para_el.find(f'{W_NS}pPr')
            if pPr is not None:
                el = pPr.find(f'{W_NS}{child_tag}')
                if el is not None:
                    return el
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(f'{W_NS}{child_tag}')
                    if el is not None:
                        return el
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(f'{W_NS}{child_tag}')
                        if el is not None:
                            return el
            return None

        # 收集所有实际生效的 header 部件根元素
        settings_el = self.doc.settings.element
        even_and_odd = settings_el.find(f'{W_NS}evenAndOddHeaders') is not None
        sections = self.doc.sections

        def _effective_header_element(idx, kind):
            while idx >= 0:
                sec = sections[idx]
                if kind == 'default':
                    h = sec.header
                elif kind == 'first':
                    h = sec.first_page_header
                else:
                    h = sec.even_page_header
                if not h.is_linked_to_previous:
                    try:
                        return h.part.element
                    except AttributeError:
                        return None
                idx -= 1
            return None

        def _header_has_visible_content(hdr_el):
            """header 部件是否在办公软件中会呈现可见内容——存在任一非空
            w:t 文本，或存在 w:drawing / w:pict / w:tbl 元素。
            """
            if hdr_el is None:
                return False
            for t in hdr_el.iter(f'{W_NS}t'):
                if t.text and t.text.strip():
                    return True
            for tag in ('drawing', 'pict', 'tbl'):
                if hdr_el.find(f'.//{W_NS}{tag}') is not None:
                    return True
            return False

        header_roots = []
        seen = set()
        for i, section in enumerate(sections):
            sectPr = section._sectPr
            title_pg = (
                sectPr is not None
                and sectPr.find(f'{W_NS}titlePg') is not None
            )
            kinds = ['default']
            if title_pg:
                kinds.append('first')
            if even_and_odd:
                kinds.append('even')
            for k in kinds:
                el = _effective_header_element(i, k)
                if el is not None and id(el) not in seen:
                    seen.add(id(el))
                    # 只判定"页面上实际显示的页眉段落"，跳过空 header
                    if _header_has_visible_content(el):
                        header_roots.append(el)

        found = False
        violated = False

        for root in header_roots:
            for para_el in root.iter(f'{W_NS}p'):
                found = True

                # 1) 页眉下无横线：effective w:pBdr/w:bottom 视为无线的情况
                pBdr = _find_effective(para_el, 'pBdr')
                if pBdr is None:
                    no_line_ok = True
                else:
                    bottom = pBdr.find(f'{W_NS}bottom')
                    if bottom is None:
                        no_line_ok = True
                    else:
                        val = bottom.get(f'{W_NS}val')
                        sz = bottom.get(f'{W_NS}sz')
                        no_line_ok = val in ('nil', 'none') or sz == '0'

                # 2) 居中对齐：effective w:jc/@w:val == "center"
                jc = _find_effective(para_el, 'jc')
                align_ok = jc is not None and jc.get(f'{W_NS}val') == 'center'

                # 3) 单倍行距：effective w:spacing 满足单倍行距
                spacing = _find_effective(para_el, 'spacing')
                if spacing is None:
                    line_ok = True
                else:
                    line = spacing.get(f'{W_NS}line')
                    line_rule = spacing.get(f'{W_NS}lineRule')
                    if line is None:
                        line_ok = True
                    elif line == '240' and line_rule == 'auto':
                        line_ok = True
                    else:
                        line_ok = False

                if not (no_line_ok and align_ok and line_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：存在生效页眉段落且任一段落不满足任一子项时触发
        passed = found and violated
        self.add_result("页眉段落格式不满足页眉下无横线、居中对齐、单倍行距", passed, -1)

    def check_header_top_margin(self):
        """-1：页眉上边距不满足1.50厘米

        细则依据（只判定"页眉上边距 = 1.50 厘米"这一点，不加任何额外
        约束）：对文档每个分节，读取其在办公软件（Word/WPS）"页面设置 →
        页边距 → 页眉"处所显示的值——对应 OOXML 的 w:sectPr/w:pgMar 的
        w:header 属性（单位 twip，1440 twip = 1 英寸 = 2.54 厘米）。
        将该 twip 值按办公软件默认的两位小数厘米显示进行舍入，需等于
        1.50 厘米。若任一分节的 pgMar/@w:header 缺失，或其厘米显示值
        不等于 1.50，即"不满足 1.50 厘米"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        violated = False

        for section in self.doc.sections:
            sectPr = section._sectPr
            pgMar = sectPr.find(f'{W_NS}pgMar') if sectPr is not None else None
            twips_attr = pgMar.get(f'{W_NS}header') if pgMar is not None else None
            if twips_attr is None:
                violated = True
                break
            try:
                twips = int(twips_attr)
            except (TypeError, ValueError):
                violated = True
                break
            # 与办公软件"厘米（两位小数）"显示保持一致的舍入判定
            cm_display = round(twips / 1440.0 * 2.54, 2)
            if cm_display != 1.50:
                violated = True
                break

        self.add_result("页眉上边距不满足1.50厘米", violated, -1)

    def check_footer_bottom_margin(self):
        """-1：页脚下边距不满足1.75厘米

        细则依据（只判定"页脚下边距 = 1.75 厘米"这一点，不加任何额外
        约束）：对文档每个分节，读取其在办公软件（Word/WPS）"页面设置 →
        页边距 → 页脚"处所显示的值——对应 OOXML 的 w:sectPr/w:pgMar 的
        w:footer 属性（单位 twip，1440 twip = 1 英寸 = 2.54 厘米）。
        将该 twip 值按办公软件默认的两位小数厘米显示进行舍入，需等于
        1.75 厘米。若任一分节的 pgMar/@w:footer 缺失，或其厘米显示值
        不等于 1.75，即"不满足 1.75 厘米"，触发 -1 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        violated = False

        for section in self.doc.sections:
            sectPr = section._sectPr
            pgMar = sectPr.find(f'{W_NS}pgMar') if sectPr is not None else None
            twips_attr = pgMar.get(f'{W_NS}footer') if pgMar is not None else None
            if twips_attr is None:
                violated = True
                break
            try:
                twips = int(twips_attr)
            except (TypeError, ValueError):
                violated = True
                break
            # 与办公软件"厘米（两位小数）"显示保持一致的舍入判定
            cm_display = round(twips / 1440.0 * 2.54, 2)
            if cm_display != 1.75:
                violated = True
                break

        self.add_result("页脚下边距不满足1.75厘米", violated, -1)

    def check_page_number_roman(self):
        """-3：文章不满足：从摘要页开始编写页码，摘要页到目录页页码格式为大写罗马数字

        细则依据（只判定"从摘要页开始编写页码"与"摘要页到目录页页码格式
        为大写罗马数字"两点，不加任何额外约束）：

        1) 从摘要页开始编写页码——摘要页所在分节及其后的每一分节，其在
           办公软件（Word/WPS）中实际渲染的页脚必须含"PAGE"域（对应
           OOXML 的 w:instrText 中包含 "PAGE" 关键字）；而摘要页所在
           分节之前的每一分节，其实际渲染页脚不得含 PAGE 域。
           "实际渲染的页脚"沿分节 is_linked_to_previous 链回溯到显式
           声明该 footerReference 的 footer 部件——default 始终生效；
           first 仅当分节 sectPr 存在 w:titlePg 时生效；even 仅当
           word/settings.xml 存在 w:evenAndOddHeaders 时生效。

        2) 摘要页到目录页页码格式为大写罗马数字——摘要页所在分节到
           目录页所在分节（含端点）的每一分节，其
           w:sectPr/w:pgNumType/@w:fmt == "upperRoman"。

        以上任一点不成立，即触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

        # 1) 定位摘要页和目录页所在分节序号（0 起）
        abstract_sec_idx = None
        toc_sec_idx = None
        section_idx = 0
        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if abstract_sec_idx is None and text_ns == '摘要':
                abstract_sec_idx = section_idx
            if toc_sec_idx is None and text_ns == '目录':
                toc_sec_idx = section_idx
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None and pPr.find(f'{W_NS}sectPr') is not None:
                section_idx += 1
            if abstract_sec_idx is not None and toc_sec_idx is not None:
                break

        # 无法定位摘要/目录时不触发扣分（保持保守）
        if abstract_sec_idx is None or toc_sec_idx is None:
            self.add_result("从摘要页开始编写页码，摘要页到目录页页码格式不为大写罗马数字", False, -3)
            return

        settings_el = self.doc.settings.element
        even_and_odd = settings_el.find(f'{W_NS}evenAndOddHeaders') is not None
        sections = self.doc.sections

        def _effective_footer_element(idx, kind):
            while idx >= 0:
                sec = sections[idx]
                if kind == 'default':
                    f = sec.footer
                elif kind == 'first':
                    f = sec.first_page_footer
                else:
                    f = sec.even_page_footer
                if not f.is_linked_to_previous:
                    try:
                        return f.part.element
                    except AttributeError:
                        return None
                idx -= 1
            return None

        def _footer_kinds(section):
            sectPr = section._sectPr
            title_pg = (
                sectPr is not None
                and sectPr.find(f'{W_NS}titlePg') is not None
            )
            kinds = ['default']
            if title_pg:
                kinds.append('first')
            if even_and_odd:
                kinds.append('even')
            return kinds

        def _footer_has_page_field(ftr_el):
            if ftr_el is None:
                return False
            for it in ftr_el.iter(f'{W_NS}instrText'):
                if it.text and 'PAGE' in it.text.upper():
                    return True
            return False

        violated = False

        # 2) 摘要页之前：所有会渲染的页脚都不得含 PAGE 域
        for i in range(0, abstract_sec_idx):
            for k in _footer_kinds(sections[i]):
                if _footer_has_page_field(_effective_footer_element(i, k)):
                    violated = True
                    break
            if violated:
                break

        # 3) 摘要页及之后：所有会渲染的页脚都必须含 PAGE 域
        if not violated:
            for i in range(abstract_sec_idx, len(sections)):
                for k in _footer_kinds(sections[i]):
                    if not _footer_has_page_field(_effective_footer_element(i, k)):
                        violated = True
                        break
                if violated:
                    break

        # 4) 摘要页到目录页（含端点）：pgNumType/@fmt == "upperRoman"
        if not violated:
            lo, hi = sorted((abstract_sec_idx, toc_sec_idx))
            for i in range(lo, hi + 1):
                sectPr = sections[i]._sectPr
                pgNum = sectPr.find(f'{W_NS}pgNumType') if sectPr is not None else None
                fmt = pgNum.get(f'{W_NS}fmt') if pgNum is not None else None
                if fmt != 'upperRoman':
                    violated = True
                    break

        self.add_result("从摘要页开始编写页码，摘要页到目录页页码格式不为大写罗马数字", violated, -3)

    def check_page_number_restart(self):
        """-3：文章不满足：从正文内容页重新编写页码从1开始，页码格式为阿拉伯数字

        细则依据（只判定"从正文内容页重新编写页码从 1 开始"与"页码格式
        为阿拉伯数字"两点，不加任何额外约束）：

        "正文内容页"定位为紧接目录页之后的第一个分节——即目录段落所在
        分节的下一个分节；办公软件（Word/WPS）以此分节为界重启页码。
        对该分节判定其 w:sectPr/w:pgNumType：
          1) 重新编写页码从 1 开始：@w:start == "1"（w:pgNumType 显式
             存在 @w:start=1 才表示"重新起编"；缺失时页码沿用上一节，
             办公软件不视为重起）
          2) 页码格式为阿拉伯数字：@w:fmt == "decimal"；或 pgNumType
             不存在 @w:fmt（阿拉伯数字为 OOXML 默认，办公软件默认按
             阿拉伯数字显示）
        任一点不成立，即触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

        # 1) 定位目录段落所在分节序号，正文内容页 = 目录分节 + 1
        toc_sec_idx = None
        section_idx = 0
        for para in self.doc.paragraphs:
            text_ns = para.text.strip().replace(' ', '').replace('　', '')
            if text_ns == '目录':
                toc_sec_idx = section_idx
                break
            pPr = para._element.find(f'{W_NS}pPr')
            if pPr is not None and pPr.find(f'{W_NS}sectPr') is not None:
                section_idx += 1

        sections = self.doc.sections
        body_sec_idx = (toc_sec_idx + 1) if toc_sec_idx is not None else None

        # 找不到目录分节或正文分节不存在时不触发扣分（保持保守）
        if body_sec_idx is None or body_sec_idx >= len(sections):
            self.add_result("正文内容页未重新编写页码从1开始，或页码格式不是阿拉伯数字", False, -3)
            return

        sectPr = sections[body_sec_idx]._sectPr
        pgNum = sectPr.find(f'{W_NS}pgNumType') if sectPr is not None else None
        start = pgNum.get(f'{W_NS}start') if pgNum is not None else None
        fmt = pgNum.get(f'{W_NS}fmt') if pgNum is not None else None

        # 1) 重新编写页码从 1 开始：需显式 @w:start == "1"
        restart_ok = start == '1'
        # 2) 阿拉伯数字：@w:fmt == "decimal" 或未设置（默认即阿拉伯数字）
        fmt_ok = fmt in (None, 'decimal')

        violated = not (restart_ok and fmt_ok)
        self.add_result("正文内容页未重新编写页码从1开始，或页码格式不是阿拉伯数字", violated, -3)

    def check_level1_font(self):
        """-3：文章一级标题字体格式不满足黑体三号

        细则依据（只判定"黑体"和"三号"两点，不加任何额外约束）：
        "文章一级标题" 指正文区（自动目录字段 fldChar end 之后）中所有
        办公软件（Word/WPS）识别为大纲级别 1 的段落——即段落 effective
        w:outlineLvl 的 val = "0"（沿"段落 pPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/pPrDefault/pPr"逐级回退取值）。
        对每个此类段落遍历所有非空 run，判定其在办公软件中实际生效的
        字体属性——沿"run 直接 rPr → 段落 pPr/rPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/rPrDefault/rPr"逐级回退取值：
          1) 字体为黑体：run 文本中的中文字符（CJK）按 effective
             w:rFonts/@w:eastAsia 判定；非中文字符按 effective
             w:rFonts/@w:ascii 判定；被判定槽位值需为 "黑体" 或 "SimHei"
          2) 字号为三号：effective w:sz 的半磅值 = 32（即 16 磅）
        只要存在一级标题且任一 run 不满足以上任一项，触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective_ppr_child(para_el, child_tag):
            """段落属性沿 pPr → pStyle 链 → docDefaults/pPrDefault 回退。"""
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                v = pPr.find(child_tag)
                if v is not None:
                    return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    v = s_pPr.find(child_tag)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        v = d_pPr.find(child_tag)
                        if v is not None:
                            return v
            return None

        def _effective_rpr_value(run_el, para_el, getter):
            """run 属性沿 run rPr → 段落 pPr/rPr → pStyle 链 rPr →
            docDefaults/rPrDefault/rPr 回退。
            """
            r_rpr = run_el.find(f'{W_NS}rPr')
            if r_rpr is not None:
                v = getter(r_rpr)
                if v is not None:
                    return v
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                p_rpr = pPr.find(f'{W_NS}rPr')
                if p_rpr is not None:
                    v = getter(p_rpr)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_rpr = style.find(f'{W_NS}rPr')
                if s_rpr is not None:
                    v = getter(s_rpr)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(f'{W_NS}rPrDefault')
                if rPrDefault is not None:
                    d_rpr = rPrDefault.find(f'{W_NS}rPr')
                    if d_rpr is not None:
                        v = getter(d_rpr)
                        if v is not None:
                            return v
            return None

        def _get_east_asia(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}eastAsia') if rFonts is not None else None

        def _get_ascii(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}ascii') if rFonts is not None else None

        def _get_sz(rpr):
            sz = rpr.find(f'{W_NS}sz')
            return sz.get(f'{W_NS}val') if sz is not None else None

        def _is_hei(name):
            return name in ('黑体', 'SimHei')

        def _is_cjk(ch):
            cp = ord(ch)
            return (
                0x4E00 <= cp <= 0x9FFF
                or 0x3400 <= cp <= 0x4DBF
                or 0xF900 <= cp <= 0xFAFF
                or 0x3000 <= cp <= 0x303F
                or 0xFF00 <= cp <= 0xFFEF
            )

        # 定位正文起点：自动目录字段 fldChar end 之后
        in_body = False
        found_level1 = False
        violated = False

        for para in self.doc.paragraphs:
            para_el = para._element
            if not in_body:
                if '<w:fldChar w:fldCharType="end"' in para_el.xml:
                    in_body = True
                continue

            # 判定 effective outlineLvl == "0"
            ol_el = _find_effective_ppr_child(para_el, f'{W_NS}outlineLvl')
            if ol_el is None:
                continue
            if ol_el.get(f'{W_NS}val') != '0':
                continue

            found_level1 = True

            for run_el in para_el.findall(f'{W_NS}r'):
                text = ''.join(
                    (t.text or '') for t in run_el.findall(f'{W_NS}t')
                )
                if not text.strip():
                    continue

                # 1) 字体：按 run 中字符类别选择字体槽位
                has_cjk = any(_is_cjk(c) for c in text)
                has_non_cjk = any(
                    (not _is_cjk(c)) and (not c.isspace()) for c in text
                )
                font_ok = True
                if has_cjk:
                    ea = _effective_rpr_value(run_el, para_el, _get_east_asia) or ''
                    if not _is_hei(ea):
                        font_ok = False
                if font_ok and has_non_cjk:
                    asc = _effective_rpr_value(run_el, para_el, _get_ascii) or ''
                    if not _is_hei(asc):
                        font_ok = False

                # 2) 字号：三号 = 半磅 32（16 磅）
                sz_val = _effective_rpr_value(run_el, para_el, _get_sz)
                size_ok = sz_val == '32'

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：存在一级标题且任一 run 不满足黑体三号
        passed = found_level1 and violated
        self.add_result("文章一级标题字体格式不满足黑体三号", passed, -3)

    def check_level1_paragraph(self):
        """-3：文章一级标题段落格式不满足居中对齐、段前一行、段后一行、两倍行距

        细则依据（只判定"居中对齐、段前一行、段后一行、两倍行距"四点，
        不加任何额外约束）："文章一级标题" 指正文区（自动目录字段
        fldChar end 之后）中所有办公软件（Word/WPS）识别为大纲级别 1
        的段落——即段落 effective w:outlineLvl 的 val = "0"（沿"段落
        pPr → w:pStyle 样式链（含 basedOn）→ w:docDefaults/pPrDefault/
        pPr"逐级回退取值）。对每个此类段落判定其在办公软件中实际生效
        的段落属性——对每个属性（w:jc 的 val，w:spacing 的 beforeLines
        / afterLines / line / lineRule）单独沿"段落 pPr → w:pStyle 样式
        链（含 basedOn）→ w:docDefaults/pPrDefault/pPr"逐级回退取值：
          1) 居中对齐：effective w:jc/@w:val == "center"
          2) 段前一行：effective w:spacing/@w:beforeLines == "100"
             （办公软件"段前 1 行"存储为 beforeLines=100，单位 1/100 行）
          3) 段后一行：effective w:spacing/@w:afterLines == "100"
          4) 两倍行距：effective w:spacing/@w:line == "480" 且
             effective w:spacing/@w:lineRule 属于 {"auto", 缺失}
             （OOXML lineRule 默认为 auto；两倍行距 = 480/240）
        只要存在一级标题且任一点不满足，触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective_ppr_child(para_el, child_tag):
            """段落属性沿 pPr → pStyle 链 → docDefaults/pPrDefault 回退，
            返回第一个出现的子元素（整元素）。
            """
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                v = pPr.find(child_tag)
                if v is not None:
                    return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    v = s_pPr.find(child_tag)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        v = d_pPr.find(child_tag)
                        if v is not None:
                            return v
            return None

        def _find_effective_ppr_attr(para_el, child_tag, attr_qname):
            """段落属性按 attribute 粒度回退（OOXML 中 w:spacing 各属性
            单独继承）：沿"段落 pPr → w:pStyle 样式链（含 basedOn）→
            docDefaults/pPrDefault/pPr"逐级回退，返回首次出现该属性
            的值。
            """
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                el = pPr.find(child_tag)
                if el is not None:
                    v = el.get(attr_qname)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(child_tag)
                    if el is not None:
                        v = el.get(attr_qname)
                        if v is not None:
                            return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(child_tag)
                        if el is not None:
                            v = el.get(attr_qname)
                            if v is not None:
                                return v
            return None

        # 定位正文起点：自动目录字段 fldChar end 之后
        in_body = False
        found_level1 = False
        violated = False

        for para in self.doc.paragraphs:
            para_el = para._element
            if not in_body:
                if '<w:fldChar w:fldCharType="end"' in para_el.xml:
                    in_body = True
                continue

            # 判定 effective outlineLvl == "0"
            ol_el = _find_effective_ppr_child(para_el, f'{W_NS}outlineLvl')
            if ol_el is None:
                continue
            if ol_el.get(f'{W_NS}val') != '0':
                continue

            found_level1 = True

            # 1) 居中对齐
            jc_val = _find_effective_ppr_attr(para_el, f'{W_NS}jc', f'{W_NS}val')
            align_ok = jc_val == 'center'

            # 2) 段前一行：beforeLines == "100"
            before_lines = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}beforeLines'
            )
            before_ok = before_lines == '100'

            # 3) 段后一行：afterLines == "100"
            after_lines = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}afterLines'
            )
            after_ok = after_lines == '100'

            # 4) 两倍行距：line == "480" 且 lineRule ∈ {"auto", 缺失}
            line_val = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}line'
            )
            line_rule = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}lineRule'
            )
            spacing_ok = line_val == '480' and line_rule in (None, 'auto')

            if not (align_ok and before_ok and after_ok and spacing_ok):
                violated = True
                break

        # 扣分项：存在一级标题且任一点不满足
        passed = found_level1 and violated
        self.add_result(
            "文章一级标题段落格式不满足居中对齐、段前一行、段后一行、两倍行距",
            passed, -3,
        )

    def check_level2_font(self):
        """-3：文章二级标题字体格式不满足黑体四号

        细则依据（只判定"黑体"和"四号"两点，不加任何额外约束）：
        "文章二级标题" 指正文区（自动目录字段 fldChar end 之后）中所有
        办公软件（Word/WPS）识别为大纲级别 2 的段落——即段落 effective
        w:outlineLvl 的 val = "1"（沿"段落 pPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/pPrDefault/pPr"逐级回退取值）。
        对每个此类段落遍历所有非空 run，判定其在办公软件中实际生效的
        字体属性——沿"run 直接 rPr → 段落 pPr/rPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/rPrDefault/rPr"逐级回退取值：
          1) 字体为黑体：run 文本中的中文字符（CJK）按 effective
             w:rFonts/@w:eastAsia 判定；非中文字符按 effective
             w:rFonts/@w:ascii 判定；被判定槽位值需为 "黑体" 或 "SimHei"
          2) 字号为四号：effective w:sz 的半磅值 = 28（即 14 磅）
        只要存在二级标题且任一 run 不满足以上任一项，触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective_ppr_child(para_el, child_tag):
            """段落属性沿 pPr → pStyle 链 → docDefaults/pPrDefault 回退。"""
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                v = pPr.find(child_tag)
                if v is not None:
                    return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    v = s_pPr.find(child_tag)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        v = d_pPr.find(child_tag)
                        if v is not None:
                            return v
            return None

        def _effective_rpr_value(run_el, para_el, getter):
            """run 属性沿 run rPr → 段落 pPr/rPr → pStyle 链 rPr →
            docDefaults/rPrDefault/rPr 回退。
            """
            r_rpr = run_el.find(f'{W_NS}rPr')
            if r_rpr is not None:
                v = getter(r_rpr)
                if v is not None:
                    return v
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                p_rpr = pPr.find(f'{W_NS}rPr')
                if p_rpr is not None:
                    v = getter(p_rpr)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_rpr = style.find(f'{W_NS}rPr')
                if s_rpr is not None:
                    v = getter(s_rpr)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(f'{W_NS}rPrDefault')
                if rPrDefault is not None:
                    d_rpr = rPrDefault.find(f'{W_NS}rPr')
                    if d_rpr is not None:
                        v = getter(d_rpr)
                        if v is not None:
                            return v
            return None

        def _get_east_asia(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}eastAsia') if rFonts is not None else None

        def _get_ascii(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}ascii') if rFonts is not None else None

        def _get_sz(rpr):
            sz = rpr.find(f'{W_NS}sz')
            return sz.get(f'{W_NS}val') if sz is not None else None

        def _is_hei(name):
            return name in ('黑体', 'SimHei')

        def _is_cjk(ch):
            cp = ord(ch)
            return (
                0x4E00 <= cp <= 0x9FFF
                or 0x3400 <= cp <= 0x4DBF
                or 0xF900 <= cp <= 0xFAFF
                or 0x3000 <= cp <= 0x303F
                or 0xFF00 <= cp <= 0xFFEF
            )

        # 定位正文起点：自动目录字段 fldChar end 之后
        in_body = False
        found_level2 = False
        violated = False

        for para in self.doc.paragraphs:
            para_el = para._element
            if not in_body:
                if '<w:fldChar w:fldCharType="end"' in para_el.xml:
                    in_body = True
                continue

            # 判定 effective outlineLvl == "1"
            ol_el = _find_effective_ppr_child(para_el, f'{W_NS}outlineLvl')
            if ol_el is None:
                continue
            if ol_el.get(f'{W_NS}val') != '1':
                continue

            found_level2 = True

            for run_el in para_el.findall(f'{W_NS}r'):
                text = ''.join(
                    (t.text or '') for t in run_el.findall(f'{W_NS}t')
                )
                if not text.strip():
                    continue

                # 1) 字体：按 run 中字符类别选择字体槽位
                has_cjk = any(_is_cjk(c) for c in text)
                has_non_cjk = any(
                    (not _is_cjk(c)) and (not c.isspace()) for c in text
                )
                font_ok = True
                if has_cjk:
                    ea = _effective_rpr_value(run_el, para_el, _get_east_asia) or ''
                    if not _is_hei(ea):
                        font_ok = False
                if font_ok and has_non_cjk:
                    asc = _effective_rpr_value(run_el, para_el, _get_ascii) or ''
                    if not _is_hei(asc):
                        font_ok = False

                # 2) 字号：四号 = 半磅 28（14 磅）
                sz_val = _effective_rpr_value(run_el, para_el, _get_sz)
                size_ok = sz_val == '28'

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：存在二级标题且任一 run 不满足黑体四号
        passed = found_level2 and violated
        self.add_result("文章二级标题字体格式不满足黑体四号", passed, -3)

    def check_level2_paragraph(self):
        """-3：文章二级标题段落格式不满足左对齐、段前0.5行、段后0.5行、1.5倍行距

        细则依据（只判定"左对齐、段前0.5行、段后0.5行、1.5倍行距"四点，
        不加任何额外约束）："文章二级标题" 指正文区（自动目录字段
        fldChar end 之后）中所有办公软件（Word/WPS）识别为大纲级别 2
        的段落——即段落 effective w:outlineLvl 的 val = "1"（沿"段落
        pPr → w:pStyle 样式链（含 basedOn）→ w:docDefaults/pPrDefault/
        pPr"逐级回退取值）。对每个此类段落判定其在办公软件中实际生效
        的段落属性——每个属性（w:jc 的 val，w:spacing 的 beforeLines /
        afterLines / line / lineRule）单独沿"段落 pPr → w:pStyle 样式
        链（含 basedOn）→ w:docDefaults/pPrDefault/pPr"逐级回退取值：
          1) 左对齐：effective w:jc/@w:val == "left"
          2) 段前0.5行：effective w:spacing/@w:beforeLines == "50"
             （办公软件"段前 0.5 行"存储为 beforeLines=50，单位 1/100 行）
          3) 段后0.5行：effective w:spacing/@w:afterLines == "50"
          4) 1.5倍行距：effective w:spacing/@w:line == "360" 且
             effective w:spacing/@w:lineRule 属于 {"auto", 缺失}
             （OOXML lineRule 默认为 auto；1.5 倍 = 360/240）
        只要存在二级标题且任一点不满足，触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective_ppr_child(para_el, child_tag):
            """段落属性沿 pPr → pStyle 链 → docDefaults/pPrDefault 回退。"""
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                v = pPr.find(child_tag)
                if v is not None:
                    return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    v = s_pPr.find(child_tag)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        v = d_pPr.find(child_tag)
                        if v is not None:
                            return v
            return None

        def _find_effective_ppr_attr(para_el, child_tag, attr_qname):
            """段落属性按 attribute 粒度回退（OOXML 中 w:spacing 各属性
            单独继承）。
            """
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                el = pPr.find(child_tag)
                if el is not None:
                    v = el.get(attr_qname)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(child_tag)
                    if el is not None:
                        v = el.get(attr_qname)
                        if v is not None:
                            return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(child_tag)
                        if el is not None:
                            v = el.get(attr_qname)
                            if v is not None:
                                return v
            return None

        # 定位正文起点：自动目录字段 fldChar end 之后
        in_body = False
        found_level2 = False
        violated = False

        for para in self.doc.paragraphs:
            para_el = para._element
            if not in_body:
                if '<w:fldChar w:fldCharType="end"' in para_el.xml:
                    in_body = True
                continue

            # 判定 effective outlineLvl == "1"
            ol_el = _find_effective_ppr_child(para_el, f'{W_NS}outlineLvl')
            if ol_el is None:
                continue
            if ol_el.get(f'{W_NS}val') != '1':
                continue

            found_level2 = True

            # 1) 左对齐
            jc_val = _find_effective_ppr_attr(para_el, f'{W_NS}jc', f'{W_NS}val')
            align_ok = jc_val == 'left'

            # 2) 段前 0.5 行：beforeLines == "50"
            before_lines = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}beforeLines'
            )
            before_ok = before_lines == '50'

            # 3) 段后 0.5 行：afterLines == "50"
            after_lines = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}afterLines'
            )
            after_ok = after_lines == '50'

            # 4) 1.5 倍行距：line == "360" 且 lineRule ∈ {"auto", 缺失}
            line_val = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}line'
            )
            line_rule = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}lineRule'
            )
            spacing_ok = line_val == '360' and line_rule in (None, 'auto')

            if not (align_ok and before_ok and after_ok and spacing_ok):
                violated = True
                break

        # 扣分项：存在二级标题且任一点不满足
        passed = found_level2 and violated
        self.add_result(
            "文章二级标题段落格式不满足左对齐、段前0.5行、段后0.5行、1.5倍行距",
            passed, -3,
        )

    def check_level3_font(self):
        """-3：文章三级标题字体格式不满足黑体小四号

        细则依据（只判定"黑体"和"小四号"两点，不加任何额外约束）：
        "文章三级标题" 指正文区（自动目录字段 fldChar end 之后）中所有
        办公软件（Word/WPS）识别为大纲级别 3 的段落——即段落 effective
        w:outlineLvl 的 val = "2"（沿"段落 pPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/pPrDefault/pPr"逐级回退取值）。
        对每个此类段落遍历所有非空 run，判定其在办公软件中实际生效的
        字体属性——沿"run 直接 rPr → 段落 pPr/rPr → w:pStyle 样式链
        （含 basedOn）→ w:docDefaults/rPrDefault/rPr"逐级回退取值：
          1) 字体为黑体：run 文本中的中文字符（CJK）按 effective
             w:rFonts/@w:eastAsia 判定；非中文字符按 effective
             w:rFonts/@w:ascii 判定；被判定槽位值需为 "黑体" 或 "SimHei"
          2) 字号为小四号：effective w:sz 的半磅值 = 24（即 12 磅）
        只要存在三级标题且任一 run 不满足以上任一项，触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective_ppr_child(para_el, child_tag):
            """段落属性沿 pPr → pStyle 链 → docDefaults/pPrDefault 回退。"""
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                v = pPr.find(child_tag)
                if v is not None:
                    return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    v = s_pPr.find(child_tag)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        v = d_pPr.find(child_tag)
                        if v is not None:
                            return v
            return None

        def _effective_rpr_value(run_el, para_el, getter):
            """run 属性沿 run rPr → 段落 pPr/rPr → pStyle 链 rPr →
            docDefaults/rPrDefault/rPr 回退。
            """
            r_rpr = run_el.find(f'{W_NS}rPr')
            if r_rpr is not None:
                v = getter(r_rpr)
                if v is not None:
                    return v
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                p_rpr = pPr.find(f'{W_NS}rPr')
                if p_rpr is not None:
                    v = getter(p_rpr)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_rpr = style.find(f'{W_NS}rPr')
                if s_rpr is not None:
                    v = getter(s_rpr)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(f'{W_NS}rPrDefault')
                if rPrDefault is not None:
                    d_rpr = rPrDefault.find(f'{W_NS}rPr')
                    if d_rpr is not None:
                        v = getter(d_rpr)
                        if v is not None:
                            return v
            return None

        def _get_east_asia(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}eastAsia') if rFonts is not None else None

        def _get_ascii(rpr):
            rFonts = rpr.find(f'{W_NS}rFonts')
            return rFonts.get(f'{W_NS}ascii') if rFonts is not None else None

        def _get_sz(rpr):
            sz = rpr.find(f'{W_NS}sz')
            return sz.get(f'{W_NS}val') if sz is not None else None

        def _is_hei(name):
            return name in ('黑体', 'SimHei')

        def _is_cjk(ch):
            cp = ord(ch)
            return (
                0x4E00 <= cp <= 0x9FFF
                or 0x3400 <= cp <= 0x4DBF
                or 0xF900 <= cp <= 0xFAFF
                or 0x3000 <= cp <= 0x303F
                or 0xFF00 <= cp <= 0xFFEF
            )

        # 定位正文起点：自动目录字段 fldChar end 之后
        in_body = False
        found_level3 = False
        violated = False

        for para in self.doc.paragraphs:
            para_el = para._element
            if not in_body:
                if '<w:fldChar w:fldCharType="end"' in para_el.xml:
                    in_body = True
                continue

            # 判定 effective outlineLvl == "2"
            ol_el = _find_effective_ppr_child(para_el, f'{W_NS}outlineLvl')
            if ol_el is None:
                continue
            if ol_el.get(f'{W_NS}val') != '2':
                continue

            found_level3 = True

            for run_el in para_el.findall(f'{W_NS}r'):
                text = ''.join(
                    (t.text or '') for t in run_el.findall(f'{W_NS}t')
                )
                if not text.strip():
                    continue

                # 1) 字体：按 run 中字符类别选择字体槽位
                has_cjk = any(_is_cjk(c) for c in text)
                has_non_cjk = any(
                    (not _is_cjk(c)) and (not c.isspace()) for c in text
                )
                font_ok = True
                if has_cjk:
                    ea = _effective_rpr_value(run_el, para_el, _get_east_asia) or ''
                    if not _is_hei(ea):
                        font_ok = False
                if font_ok and has_non_cjk:
                    asc = _effective_rpr_value(run_el, para_el, _get_ascii) or ''
                    if not _is_hei(asc):
                        font_ok = False

                # 2) 字号：小四号 = 半磅 24（12 磅）
                sz_val = _effective_rpr_value(run_el, para_el, _get_sz)
                size_ok = sz_val == '24'

                if not (font_ok and size_ok):
                    violated = True
                    break
            if violated:
                break

        # 扣分项：存在三级标题且任一 run 不满足黑体小四号
        passed = found_level3 and violated
        self.add_result("文章三级标题字体格式不满足黑体小四号", passed, -3)

    def check_level3_paragraph(self):
        """-3：文章三级标题段落格式不满足左对齐、段前0.5行、段后0.5行、1.5倍行距

        细则依据（只判定"左对齐、段前0.5行、段后0.5行、1.5倍行距"四点，
        不加任何额外约束）："文章三级标题" 指正文区（自动目录字段
        fldChar end 之后）中所有办公软件（Word/WPS）识别为大纲级别 3
        的段落——即段落 effective w:outlineLvl 的 val = "2"（沿"段落
        pPr → w:pStyle 样式链（含 basedOn）→ w:docDefaults/pPrDefault/
        pPr"逐级回退取值）。对每个此类段落判定其在办公软件中实际生效
        的段落属性——每个属性（w:jc 的 val，w:spacing 的 beforeLines /
        afterLines / line / lineRule）单独沿"段落 pPr → w:pStyle 样式
        链（含 basedOn）→ w:docDefaults/pPrDefault/pPr"逐级回退取值：
          1) 左对齐：effective w:jc/@w:val == "left"
          2) 段前0.5行：effective w:spacing/@w:beforeLines == "50"
             （办公软件"段前 0.5 行"存储为 beforeLines=50，单位 1/100 行）
          3) 段后0.5行：effective w:spacing/@w:afterLines == "50"
          4) 1.5倍行距：effective w:spacing/@w:line == "360" 且
             effective w:spacing/@w:lineRule 属于 {"auto", 缺失}
             （OOXML lineRule 默认为 auto；1.5 倍 = 360/240）
        只要存在三级标题且任一点不满足，触发 -3 扣分。
        """
        if not self.doc:
            return

        W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_el = self.doc.styles.element

        def _style_by_id(style_id):
            if not style_id:
                return None
            for s in styles_el.findall(f'{W_NS}style'):
                if s.get(f'{W_NS}styleId') == style_id:
                    return s
            return None

        def _find_effective_ppr_child(para_el, child_tag):
            """段落属性沿 pPr → pStyle 链 → docDefaults/pPrDefault 回退。"""
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                v = pPr.find(child_tag)
                if v is not None:
                    return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    v = s_pPr.find(child_tag)
                    if v is not None:
                        return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        v = d_pPr.find(child_tag)
                        if v is not None:
                            return v
            return None

        def _find_effective_ppr_attr(para_el, child_tag, attr_qname):
            """段落属性按 attribute 粒度回退（OOXML 中 w:spacing 各属性
            单独继承）。
            """
            pPr = para_el.find(f'{W_NS}pPr') if para_el is not None else None
            if pPr is not None:
                el = pPr.find(child_tag)
                if el is not None:
                    v = el.get(attr_qname)
                    if v is not None:
                        return v
            style_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{W_NS}pStyle')
                if pStyle is not None:
                    style_id = pStyle.get(f'{W_NS}val')
            visited = set()
            while style_id and style_id not in visited:
                visited.add(style_id)
                style = _style_by_id(style_id)
                if style is None:
                    break
                s_pPr = style.find(f'{W_NS}pPr')
                if s_pPr is not None:
                    el = s_pPr.find(child_tag)
                    if el is not None:
                        v = el.get(attr_qname)
                        if v is not None:
                            return v
                based_on = style.find(f'{W_NS}basedOn')
                style_id = based_on.get(f'{W_NS}val') if based_on is not None else None
            doc_defaults = styles_el.find(f'{W_NS}docDefaults')
            if doc_defaults is not None:
                pPrDefault = doc_defaults.find(f'{W_NS}pPrDefault')
                if pPrDefault is not None:
                    d_pPr = pPrDefault.find(f'{W_NS}pPr')
                    if d_pPr is not None:
                        el = d_pPr.find(child_tag)
                        if el is not None:
                            v = el.get(attr_qname)
                            if v is not None:
                                return v
            return None

        # 定位正文起点：自动目录字段 fldChar end 之后
        in_body = False
        found_level3 = False
        violated = False

        for para in self.doc.paragraphs:
            para_el = para._element
            if not in_body:
                if '<w:fldChar w:fldCharType="end"' in para_el.xml:
                    in_body = True
                continue

            # 判定 effective outlineLvl == "2"
            ol_el = _find_effective_ppr_child(para_el, f'{W_NS}outlineLvl')
            if ol_el is None:
                continue
            if ol_el.get(f'{W_NS}val') != '2':
                continue

            found_level3 = True

            # 1) 左对齐
            jc_val = _find_effective_ppr_attr(para_el, f'{W_NS}jc', f'{W_NS}val')
            align_ok = jc_val == 'left'

            # 2) 段前 0.5 行：beforeLines == "50"
            before_lines = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}beforeLines'
            )
            before_ok = before_lines == '50'

            # 3) 段后 0.5 行：afterLines == "50"
            after_lines = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}afterLines'
            )
            after_ok = after_lines == '50'

            # 4) 1.5 倍行距：line == "360" 且 lineRule ∈ {"auto", 缺失}
            line_val = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}line'
            )
            line_rule = _find_effective_ppr_attr(
                para_el, f'{W_NS}spacing', f'{W_NS}lineRule'
            )
            spacing_ok = line_val == '360' and line_rule in (None, 'auto')

            if not (align_ok and before_ok and after_ok and spacing_ok):
                violated = True
                break

        # 扣分项：存在三级标题且任一点不满足
        passed = found_level3 and violated
        self.add_result(
            "文章三级标题段落格式不满足左对齐、段前0.5行、段后0.5行、1.5倍行距",
            passed, -3,
        )


def _locate_docx(dir_path: str) -> str:
    """在 dir_path 目录内定位被评估的 Word 文档 (.docx)。

    约定：一个评估目录内应恰有一个待评估文档；若存在多个，选择目录中
    首个（按 os.listdir 顺序）的 .docx 文件。跳过 Office 临时文件
    （以 "~$" 开头的锁文件）。
    """
    if not os.path.isdir(dir_path):
        raise FileNotFoundError(f"目录不存在: {dir_path}")
    candidates = []
    for name in os.listdir(dir_path):
        if name.startswith('~$'):
            continue
        if name.lower().endswith('.docx'):
            candidates.append(name)
    if not candidates:
        raise FileNotFoundError(f"目录 {dir_path} 中未找到 .docx 文档")
    return os.path.join(dir_path, candidates[0])


def evaluate(dir_path: str) -> dict:
    """统一评估入口。

    参数：
        dir_path: 脚本所在目录路径；脚本自行在该目录内定位并打开被评估
                  的 .docx 文档。

    返回：结构化字典，字段定义详见接口约定 §2.2。
    """
    result = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }

    try:
        file_path = _locate_docx(dir_path)
        result["file_name"] = os.path.basename(file_path)

        evaluator = DocxEvaluator(file_path)
        # 内部检查过程中的调试 print 全部丢弃，避免污染 stdout/stderr
        # （符合接口约定 §2.3：主结果只走 return，不 print 到 stdout）
        import contextlib
        with open(os.devnull, 'w', encoding='utf-8') as _devnull, \
                contextlib.redirect_stdout(_devnull), \
                contextlib.redirect_stderr(_devnull):
            score, raw_results = evaluator.evaluate()

        # 划分维度一 / 维度二：score_change == 0 视为维度一子项
        dim1_items = [r for r in raw_results if r['score_change'] == 0]
        dim2_raw = [r for r in raw_results if r['score_change'] != 0]

        dim1_failed = [r for r in dim1_items if not r['passed']]
        result["dim1_pass"] = evaluator.dimension1_passed and not dim1_failed
        if dim1_failed:
            result["dim1_reason"] = "；".join(r['description'] for r in dim1_failed)

        # 维度二逐项：命中与未命中均输出
        for r in dim2_raw:
            delta = r['score_change'] if r['passed'] else 0
            result["dim2_items"].append({
                "rule": r['description'],
                "max_delta": r['score_change'],
                "delta": delta,
                "hit": bool(r['passed']),
                "detail": "",
            })

        # 满分：仅累加加分项（正 max_delta）；扣分项仅在命中时才拉低总分
        result["max_score"] = sum(
            item["max_delta"] for item in result["dim2_items"]
            if item["max_delta"] > 0
        )
        result["total_score"] = score
    except Exception as exc:
        result["status"] = "error"
        result["error"] = f"{type(exc).__name__}: {exc}"

    return result


def main():
    # 本地调试入口：命令行参数指定脚本所在目录；缺省时使用本脚本所在目录
    if len(sys.argv) >= 2:
        dir_path = sys.argv[1]
    else:
        dir_path = os.path.dirname(os.path.abspath(__file__))
    payload = json.dumps(evaluate(dir_path), ensure_ascii=False, indent=2)
    # 直接向 stdout.buffer 写 UTF-8 字节，避免 Windows cp1252 控制台
    # 对中文字符抛 UnicodeEncodeError；同时避免改动 sys.stdout 全局对象
    try:
        sys.stdout.buffer.write((payload + "\n").encode("utf-8"))
        sys.stdout.buffer.flush()
    except AttributeError:
        # 非二进制 stdout（如某些 IDE 捕获）时回退到普通 print
        print(payload)


if __name__ == "__main__":
    main()
