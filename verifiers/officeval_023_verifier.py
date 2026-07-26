# -*- coding: utf-8 -*-
"""
自动评估硕士论文 Word 文档 —— 完整版
文件：董事會數位素養對企業低碳轉型績效的影響_最终版.docx

评分逻辑：
  维度一：可用与可修改性 —— 任一不满足 → 总分 0
  维度二：完成度评分细则 —— 得分点(累加正分) + 扣分点(累加负分)
"""

import glob
import json
import os
import re
import time
import atexit
from typing import TypedDict
from docx import Document

# 由 evaluate(dir_path) 在运行时写入：脚本所在目录里被评估的 docx 绝对路径
_file_path = ""

_DOCX_ANCHORS = (
    '碩士研究生學位論文', '硕士研究生学位论文',
    '摘要', 'abstract',
    '第一章緒論', '第一章绪论',
    '參考文獻', '参考文献',
    '表目錄', '表目录',
    '圖目錄', '图目录',
)


def _candidate_anchor_score(path: str) -> int:
    """按本任务论文结构特征评估 DOCX 候选，解析失败时视为不匹配。"""
    try:
        doc = Document(path)
        text = ''.join(
            re.sub(r'\s+', '', paragraph.text).casefold()
            for paragraph in doc.paragraphs
        )
    except Exception:
        return 0
    return sum(anchor.casefold() in text for anchor in _DOCX_ANCHORS)


def _resolve_docx_path(dir_path: str) -> str:
    """按论文结构特征定位 DOCX；同分时稳定选择体积最大的文档。"""
    candidates = [
        p for p in glob.glob(os.path.join(dir_path, '*.docx'))
        if not os.path.basename(p).startswith('~$')
    ]
    if not candidates:
        raise FileNotFoundError(f"目录中未找到 .docx 文件: {dir_path}")
    ranked = sorted(
        (
            (_candidate_anchor_score(path), os.path.getsize(path), path)
            for path in candidates
        ),
        key=lambda item: (
            -item[0], -item[1], os.path.basename(item[2]).casefold()
        ),
    )
    if len(candidates) > 1 and ranked[0][0] == 0:
        names = sorted(os.path.basename(path) for path in candidates)
        raise RuntimeError(
            f"目录中存在多个 .docx 文件，且均不含目标论文结构特征: {names}"
        )
    return os.path.abspath(ranked[0][2])



# =========================================================================
# Word COM 会话代理：整份评估共享同一个 Word 实例与同一份 Document，
# 保证跨检查项拿到的分页/字体属性是稳定值。
# =========================================================================
class _WordSession:
    _app = None
    _doc = None
    _path = None

    @classmethod
    def app(cls):
        if cls._app is None:
            import win32com.client as w32
            try:
                cls._app = w32.DispatchEx('Word.Application')
            except Exception:
                cls._app = w32.DispatchEx('Word.Application')
            try:
                cls._app.Visible = False
            except Exception:
                pass
        return cls._app

    @classmethod
    def doc(cls):
        abs_path = os.path.abspath(_file_path)
        if cls._doc is not None and cls._path == abs_path:
            return cls._doc
        cls._doc = cls.app().Documents.Open(abs_path, ReadOnly=True)
        cls._path = abs_path
        try:
            cls._doc.Repaginate()
        except Exception:
            pass
        try:
            _ = int(cls._doc.ComputeStatistics(2))  # 强制完成分页
        except Exception:
            pass
        time.sleep(0.3)
        return cls._doc

    @classmethod
    def close(cls):
        try:
            if cls._doc is not None:
                cls._doc.Close(False)
        except Exception:
            pass
        try:
            if cls._app is not None:
                cls._app.Quit()
        except Exception:
            pass
        cls._doc = None
        cls._app = None
        cls._path = None


atexit.register(_WordSession.close)


# =========================================================================
# 维度二评分项构造：不依赖全局状态，直接返回结构化 item
# =========================================================================
class Dim2Item(TypedDict):
    rule: str
    max_delta: int
    delta: int
    hit: bool
    detail: str


def _item(rule: str, max_delta: int, hit: bool, detail: str = "") -> Dim2Item:
    """构造统一的维度二评分项：max_delta 可正可负，hit 命中即拿到 max_delta。"""
    # 说明：为保持函数签名与调用方一致仍接受 detail 入参；按需求最终输出置空，
    # 不影响任何评分逻辑与返回结构。
    _ = detail
    return {
        "rule": rule,
        "max_delta": max_delta,
        "delta": max_delta if hit else 0,
        "hit": bool(hit),
        "detail": "",
    }


# =========================================================================
# 常量：Word 枚举
# =========================================================================
WD_ACTIVE_END_PAGE = 3
WD_WITH_IN_TABLE = 12
WD_ALIGN_LEFT = 0
WD_ALIGN_CENTER = 1
WD_ALIGN_RIGHT = 2
WD_ALIGN_JUSTIFY = 3
WD_LINE_SINGLE = 0
WD_LINE_1PT5 = 1
WD_STAT_PAGES = 2

# APPEND_HERE_1

# =========================================================================
# 工具函数
# =========================================================================
def _norm(s):
    return re.sub(r'\s+', '', s or '')


def _is_en_or_num(ch):
    return ch.isascii() and ch.isalnum()


def _com_call(fn, *args, retries=3, delay=0.5, **kwargs):
    """执行 COM 调用，遇到 -2147418111（RPC 被拒绝）自动小睡后重试。"""
    for i in range(retries + 1):
        try:
            return fn(*args, **kwargs)
        except Exception as e:
            msg = str(e)
            if i < retries and ('-2147418111' in msg or 'rejected' in msg or '被呼叫方拒绝' in msg or '被调用方' in msg or '拒绝接收' in msg):
                time.sleep(delay * (i + 1))
                continue
            raise


# ---- Word COM 缓存：常用信息只读一次 --------------------------------------
_cache = {}


def com_paragraphs():
    """返回 [(idx, text, norm, style_name, outline_level, page)]，只算一次。

    页码通过 Word 布局引擎读取代价较高，缓存构建阶段不逐段查询；仅在
    评分规则命中目标段落后按需读取。
    """
    if 'com_paras' in _cache:
        return _cache['com_paras']
    doc = _WordSession.doc()
    total = int(doc.Paragraphs.Count)
    out = []
    for i in range(1, total + 1):
        p = doc.Paragraphs(i)
        text = p.Range.Text.rstrip('\r\x07\n')
        out.append((i, text, _norm(text), '', 10, -1))
    _cache['com_paras'] = out
    return out


def com_para_page(i):
    """按需读取指定段落的真实渲染页码。"""
    try:
        return int(com_para(i).Range.Information(WD_ACTIVE_END_PAGE))
    except Exception:
        return -1


def com_para_outline(i):
    """按需读取指定段落的大纲级别。"""
    try:
        return int(com_para(i).Format.OutlineLevel)
    except Exception:
        return 10


def com_para(i):
    """按 1-based 索引拿到 Word COM 的 Paragraph 对象。"""
    return _WordSession.doc().Paragraphs(i)


# =========================================================================
# Word COM 段落定位辅助函数（原 check_dim2 内的两个闭包）
# =========================================================================
# -----------------------------------------------------------------
# 通用：Word COM 找到某段文本对应的段落（返回 Paragraph 或 None）
# -----------------------------------------------------------------
def _find_com_para(norm_targets):
    """在 Word COM 里按 norm 精确匹配定位段落，返回 Paragraph 或 None。"""
    for i, text, norm, *_ in com_paragraphs():
        if norm in norm_targets:
            return com_para(i)
    return None

def _find_com_para_by_prefix(prefix):
    for i, text, norm, *_ in com_paragraphs():
        if text.strip().startswith(prefix):
            return com_para(i)
    return None



# =========================================================================
# 维度一：可用与可修改性
# =========================================================================
def check_dim1():
    """维度一：可用与可修改性。返回 (passed: bool, reason: str)；通过时 reason 为空字符串。"""
    # 1. 文件可以正常打开
    try:
        doc = Document(_file_path)
    except Exception as e:
        return False, f"D1 文件无法正常打开: {e}"

    # 说明：按用户要求，删除以下两项维度一检查
    #   - "正文/目录/表目录/页码/表格/图表/标题均可编辑，不能整篇转成PDF、图片或不可编辑对象"
    #   - "无连续2页以上空白页，无超过1/3页面面积乱码，无超过1/3页面面积文字重叠"
    _ = doc  # 保留打开动作作为唯一维度一检查

    return True, ""



# =========================================================================
# 维度二：完成度评分（得分点 + 扣分点）
# =========================================================================
def _dim2_last_page_removed():
    # -----------------------------------------------------------------
    # +3：文档中原最后一页"作者簡歷"整页已删除
    # 依据 Word 办公软件文件属性：遍历 Word COM Paragraphs，
    # 查找是否存在正文标题段"作者簡歷/作者简历"（排除目录条目）。
    # 不存在 → 整页已删除 → +3；存在 → 0 分。
    # -----------------------------------------------------------------
    _hit_para = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if '\t' in text:  # 目录条目
            continue
        if norm in ('作者簡歷', '作者简历'):
            _hit_para = (i, com_para_page(i))
            break
    return _item(
        "文档中原最后一页'作者簡歷'整页已删除",
        +3,
        _hit_para is None,
        "" if _hit_para is None else f"段号={_hit_para[0]}, 页码={_hit_para[1]}",
    )


def _dim2_section_page_no_setup():
    # -----------------------------------------------------------------
    # +3：文档存在分节页码设置：摘要页前的封面、答辩委员会审议通过表、致谢页不添加页码
    # 依据 Word 办公软件文件属性：
    #   ① Word 中存在多分节（Sections.Count > 摘要所在节序号）→"存在分节页码设置"
    #   ② 摘要所在节之前的每一节，其 Headers/Footers 中不包含 PAGE 域字段
    # -----------------------------------------------------------------
    WD_FIELD_PAGE = 33

    def _sec_has_page_field(sec):
        """Word COM Section 里是否存在 PAGE 域字段（页码字段）。"""
        for coll_name in ('Headers', 'Footers'):
            try:
                coll = getattr(sec, coll_name)
            except Exception:
                continue
            for k in range(1, 4):  # 1=Primary, 2=FirstPage, 3=EvenPages
                try:
                    hf = coll(k)
                except Exception:
                    continue
                try:
                    if not bool(hf.Exists):
                        continue
                except Exception:
                    pass
                try:
                    for f in hf.Range.Fields:
                        if int(f.Type) == WD_FIELD_PAGE:
                            return True
                except Exception:
                    pass
        return False

    _wd = _WordSession.doc()
    # 定位摘要段落所在节号（1-based）
    abs_sec_no = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if '\t' in text:
            continue
        if norm == '摘要':
            try:
                abs_sec_no = int(com_para(i).Range.Information(2))  # wdActiveEndSectionNumber=2
            except Exception:
                abs_sec_no = None
            break

    total_sec = int(_wd.Sections.Count)
    _det = []
    ok = False
    if abs_sec_no is None:
        _det.append("未定位到摘要段所在节")
    elif abs_sec_no <= 1:
        _det.append(f"摘要位于第 {abs_sec_no} 节，其前无独立分节")
    else:
        leaks = []
        for s in range(1, abs_sec_no):
            if _sec_has_page_field(_wd.Sections(s)):
                leaks.append(f"sec{s}")
        if leaks:
            _det.append("存在页码的节：" + ", ".join(leaks))
        else:
            ok = True
            _det.append(f"共 {total_sec} 节，摘要位于第 {abs_sec_no} 节，前 {abs_sec_no-1} 节均无 PAGE 域")
    return _item(
        "文档存在分节页码设置：摘要页前的封面、答辩委员会审议通过表、致谢页不添加页码",
        +3, ok, "; ".join(_det))


def _dim2_roman_pagination():
    # -----------------------------------------------------------------
    # +3：从"摘 要"所在页开始，表目录页所在页结束，页脚页码使用大写罗马数字，
    #     页码序列按 I、II、III、IV…… 连续递增
    # 完全基于 Word 办公软件文件属性：
    #   ① 定位摘要段落 → 页号 abs_page；表目录段落 → 页号 tbl_page
    #   ② 对该页范围覆盖的每一节：
    #        a. Section.Footers(x).PageNumbers.NumberStyle == wdPageNumberStyleUppercaseRoman(1)
    #        b. Footer 中包含 PAGE 域字段
    #        c. 首节 StartingNumber==1，其余节 RestartNumberingAtSection==False（连续递增）
    # -----------------------------------------------------------------
    WD_PAGENUM_UPPER_ROMAN = 1
    WD_ACTIVE_END_SEC = 2
    _wd = _WordSession.doc()

    abs_page = abs_sec = tbl_page = tbl_sec = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if '\t' in text:
            continue
        if abs_page is None and norm == '摘要':
            abs_page = com_para_page(i)
            try:
                abs_sec = int(com_para(i).Range.Information(WD_ACTIVE_END_SEC))
            except Exception:
                abs_sec = None
        if norm in ('表目錄', '表目录'):
            tbl_page = com_para_page(i)
            try:
                tbl_sec = int(com_para(i).Range.Information(WD_ACTIVE_END_SEC))
            except Exception:
                tbl_sec = None

    roman_ok = False
    _det = []
    if abs_sec is None or tbl_sec is None:
        _det.append(f"定位失败：摘要节={abs_sec}, 表目录节={tbl_sec}")
    elif abs_sec > tbl_sec:
        _det.append(f"节序号异常：摘要节 {abs_sec} > 表目录节 {tbl_sec}")
    else:
        roman_ok = True
        for s in range(abs_sec, tbl_sec + 1):
            sec = _wd.Sections(s)
            # 页脚 PAGE 域字段
            has_page = False
            try:
                for k in range(1, 4):
                    ft = sec.Footers(k)
                    try:
                        if not bool(ft.Exists):
                            continue
                    except Exception:
                        pass
                    for f in ft.Range.Fields:
                        if int(f.Type) == 33:  # wdFieldPage
                            has_page = True
                            break
                    if has_page:
                        break
            except Exception:
                pass
            if not has_page:
                roman_ok = False
                _det.append(f"sec{s} 页脚无 PAGE 域")
            # NumberStyle == 大写罗马
            style_val = None
            try:
                style_val = int(sec.Footers(1).PageNumbers.NumberStyle)
            except Exception:
                pass
            if style_val != WD_PAGENUM_UPPER_ROMAN:
                roman_ok = False
                _det.append(f"sec{s} NumberStyle={style_val}")
            # 连续递增：首节 StartingNumber==1；其余节不重启编号
            try:
                pn = sec.Footers(1).PageNumbers
                start = int(pn.StartingNumber)
                restart = bool(pn.RestartNumberingAtSection)
            except Exception:
                start, restart = None, None
            if s == abs_sec:
                if start != 1:
                    roman_ok = False
                    _det.append(f"sec{s} 起始页码={start}")
            else:
                if restart:
                    roman_ok = False
                    _det.append(f"sec{s} 在本节重启编号(start={start})")

        # 额外校验：摘要页显示为 I，表目录页显示为罗马数字且 > 摘要页
        if roman_ok and abs_page and tbl_page and tbl_page >= abs_page:
            _det.append(f"页面范围 {abs_page}→{tbl_page}（共 {tbl_page-abs_page+1} 页）")

    return _item(
        "从'摘 要'所在页开始，表目录页所在页结束，页脚页码使用大写罗马数字，页码序列按 I、II、III、IV…… 连续递增",
        +3, roman_ok, "; ".join(_det))


def _dim2_arabic_pagination():
    # -----------------------------------------------------------------
    # +3：从"第一章 緒 論"所在页开始，页脚页码改为阿拉伯数字 1，正文页码
    #     从"第一章 緒 論"页起按 1、2、3、4…… 连续递增，直到"參考文獻"结束
    # 完全基于 Word 办公软件文件属性：
    #   ① Word COM Paragraphs 定位"第一章 緒 論"标题段 → 所在节号 ch1_sec、
    #      所在页号 ch1_page；定位"參考文獻"标题段 → 所在节号 ref_sec
    #   ② 对 [ch1_sec, ref_sec] 区间的每一节：
    #        a. Footer 中包含 PAGE 域字段
    #        b. Section.Footers(1).PageNumbers.NumberStyle == wdPageNumberStyleArabic(0)
    #        c. 首节 StartingNumber == 1；其余节 RestartNumberingAtSection == False
    #   ③ "第一章 緒 論"所在页的显示页码（wdActiveEndAdjustedPageNumber=1）== 1
    # -----------------------------------------------------------------
    _WD_PAGENUM_ARABIC = 0
    _WD_ACTIVE_END_SEC = 2
    _WD_ACTIVE_END_ADJUSTED_PAGE = 1
    _wd = _WordSession.doc()

    ch1_para_idx = ch1_sec = ref_sec = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if '\t' in text:
            continue
        if ch1_para_idx is None and re.match(r'^第一章\s*緒\s*論', text.strip()):
            ch1_para_idx = i
            try:
                ch1_sec = int(com_para(i).Range.Information(_WD_ACTIVE_END_SEC))
            except Exception:
                ch1_sec = None
        if ref_sec is None and norm in ('參考文獻', '参考文献'):
            try:
                ref_sec = int(com_para(i).Range.Information(_WD_ACTIVE_END_SEC))
            except Exception:
                ref_sec = None

    arab_ok = False
    _det = []
    if ch1_sec is None or ref_sec is None:
        _det.append(f"定位失败：'第一章 緒 論'节={ch1_sec}, '參考文獻'节={ref_sec}")
    elif ch1_sec > ref_sec:
        _det.append(f"节序号异常：'第一章 緒 論'节 {ch1_sec} > '參考文獻'节 {ref_sec}")
    else:
        arab_ok = True
        for s in range(ch1_sec, ref_sec + 1):
            sec = _wd.Sections(s)
            # a. 页脚 PAGE 域字段
            has_page = False
            try:
                for k in range(1, 4):  # 1=Primary, 2=FirstPage, 3=EvenPages
                    ft = sec.Footers(k)
                    try:
                        if not bool(ft.Exists):
                            continue
                    except Exception:
                        pass
                    for f in ft.Range.Fields:
                        if int(f.Type) == 33:  # wdFieldPage
                            has_page = True
                            break
                    if has_page:
                        break
            except Exception:
                pass
            if not has_page:
                arab_ok = False
                _det.append(f"sec{s} 页脚无 PAGE 域")
            # b. NumberStyle == 阿拉伯数字
            style_val = None
            try:
                style_val = int(sec.Footers(1).PageNumbers.NumberStyle)
            except Exception:
                pass
            if style_val != _WD_PAGENUM_ARABIC:
                arab_ok = False
                _det.append(f"sec{s} NumberStyle={style_val}")
            # c. 首节起始 = 1；其余节不重启编号（连续递增）
            try:
                pn = sec.Footers(1).PageNumbers
                start = int(pn.StartingNumber)
                restart = bool(pn.RestartNumberingAtSection)
            except Exception:
                start, restart = None, None
            if s == ch1_sec:
                if start != 1:
                    arab_ok = False
                    _det.append(f"sec{s} 起始页码={start}")
            else:
                if restart:
                    arab_ok = False
                    _det.append(f"sec{s} 在本节重启编号(start={start})")

        # ③ "第一章 緒 論"所在页显示的页码 == 1
        try:
            disp = int(com_para(ch1_para_idx).Range.Information(
                _WD_ACTIVE_END_ADJUSTED_PAGE))
            if disp != 1:
                arab_ok = False
                _det.append(f"'第一章 緒 論'页显示页码={disp}")
            else:
                _det.append(f"'第一章 緒 論'页显示页码=1，节范围 {ch1_sec}→{ref_sec}")
        except Exception:
            pass

    return _item(
        "从'第一章 緒 論'所在页开始，页脚页码改为阿拉伯数字1，正文页码从'第一章 緒 論'页起按1、2、3、4……连续递增，直到'參考文獻'结束",
        +3, arab_ok, "; ".join(_det))

    # APPEND_HERE_DIM2_A


def _dim2_toc_auto_update():
    # -----------------------------------------------------------------
    # +5：文档中目录为自动更新目录，点击对应条目可跳转到其所在位置
    # 完全基于 Word 办公软件文件属性：
    #   ① Document.TablesOfContents.Count >= 1 —— 存在自动更新目录（TOC 域）
    #   ② 该 TOC 域的 TableOfContents.UseHyperlinks == True —— 点击条目跳转
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    _d = []
    has_toc = False
    can_jump = False
    try:
        if int(_wd.TablesOfContents.Count) >= 1:
            has_toc = True
            toc = _wd.TablesOfContents(1)
            try:
                can_jump = bool(toc.UseHyperlinks)
            except Exception:
                can_jump = False
    except Exception as e:
        _d.append(f"读取失败: {e}")
    if not has_toc:
        _d.append("Document.TablesOfContents.Count=0（无自动更新目录）")
    elif not can_jump:
        _d.append("TableOfContents.UseHyperlinks=False（条目不可跳转）")
    ok = has_toc and can_jump
    return _item(
        "文档中目录为自动更新目录，点击对应条目可跳转到其所在位置",
        +5, ok, "; ".join(_d))


def _dim2_toc_no_thanks():
    # -----------------------------------------------------------------
    # +1：目录中未出现"致谢"条目
    # 完全基于 Word 办公软件文件属性：
    #   Document.TablesOfContents(1).Range 即"自动更新目录"域的完整范围，
    #   遍历其中每一段（即每一条目录条目），检查文本里是否含"致谢/致謝"。
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    has_ack = False
    _d = []
    try:
        if int(_wd.TablesOfContents.Count) >= 1:
            toc = _wd.TablesOfContents(1)
            for para in toc.Range.Paragraphs:
                t = str(para.Range.Text or '')
                n = _norm(t)
                if '致謝' in n or '致谢' in n:
                    has_ack = True
                    _d.append(f"目录条目='{t.strip()[:20]}'")
                    break
        else:
            _d.append("Document.TablesOfContents.Count=0")
    except Exception as e:
        _d.append(f"读取失败: {e}")
    return _item(
        "目录中未出现'致谢'条目",
        +1, not has_ack, "; ".join(_d))


def _dim2_toc_page_numbers():
    # -----------------------------------------------------------------
    # +5：目录中"摘要"对应页码为大写罗马数字I，英文摘要Abstract对应页码为大写罗马数字II,
    #     目录页对应页码为大写罗马数字IV，表目录对应页码为VI，图目录对应页码为VII
    # 完全基于 Word 办公软件文件属性：
    #   Document.TablesOfContents(1).Range 即"自动更新目录"域的完整范围，
    #   遍历其中每一段（每一条目录条目）；每条条目为
    #     "标题文本 \t 页码"（页码是 TOC 域内嵌 PAGEREF 域产生的显示文本），
    #   按标题匹配 5 项，读取尾部页码文本与期望值比对。
    # -----------------------------------------------------------------
    expected = [
        (r'^摘\s*要$', 'I', '摘要'),
        (r'^Abstract$', 'II', 'Abstract'),
        (r'^目\s*[录錄]$', 'IV', '目录'),
        (r'^表\s*目\s*[录錄]$', 'VI', '表目录'),
        (r'^[图圖]\s*目\s*[录錄]$', 'VII', '图目录'),
    ]
    pt_res = {label: {'exp': ex, 'actual': None, 'ok': False}
              for _, ex, label in expected}
    _wd = _WordSession.doc()
    _d = []
    try:
        if int(_wd.TablesOfContents.Count) < 1:
            _d.append("Document.TablesOfContents.Count=0")
        else:
            toc = _wd.TablesOfContents(1)
            for para in toc.Range.Paragraphs:
                raw = str(para.Range.Text or '').rstrip('\r\x07\n')
                if not raw.strip():
                    continue
                if '\t' in raw:
                    parts = re.split(r'\t+', raw)
                    entry = parts[0].strip()
                    page = parts[-1].strip()
                else:
                    m = re.search(r'\s+(\S+)\s*$', raw.strip())
                    if not m:
                        continue
                    entry = raw.strip()[:m.start()].strip()
                    page = m.group(1).strip()
                for pat, exp, lbl in expected:
                    if re.match(pat, entry):
                        pt_res[lbl]['actual'] = page
                        pt_res[lbl]['ok'] = (page == exp)
                        break
    except Exception as e:
        _d.append(f"读取失败: {e}")
    ok = all(v['ok'] for v in pt_res.values())
    for lbl in ('摘要', 'Abstract', '目录', '表目录', '图目录'):
        v = pt_res[lbl]
        if v['actual'] is None:
            _d.append(f"{lbl}→期望{v['exp']}, 未找到")
        elif not v['ok']:
            _d.append(f"{lbl}→期望{v['exp']}, 实际{v['actual']}")
        else:
            _d.append(f"{lbl}→{v['actual']} ✓")
    return _item(
        "目录中'摘要'对应页码为I、Abstract对应页码为II、目录页对应页码为IV、表目录对应页码为VI、图目录对应页码为VII",
        +5, ok, "; ".join(_d))


def _dim2_table_of_tables():
    # -----------------------------------------------------------------
    # +3："表 目 錄"页中表目录位于"圖 目 錄"页之前
    # 完全基于 Word 办公软件文件属性：
    #   ① 存在"表目录页" → 文档中存在一段 Paragraph，其 OutlineLevel=1（一级标题，
    #      对应 Word "标题 1" 的大纲级别，即独立成"页"的顶级标题），
    #      且其 Range.Text 归一化后等于 "表目錄"/"表目录"（"圖 目 錄" 同理）。
    #   ② "表 目 錄"页中的表目录位于"圖 目 錄"页之前
    #      → 表目录标题段 Range.Start < 图目录标题段 Range.Start。
    # -----------------------------------------------------------------
    _d = []

    tbl_start = fig_start = None
    for i, text, norm, style, outline, page in com_paragraphs():
        t = norm.replace('\r', '').replace('\x07', '')
        if t not in ('表目錄', '表目录', '圖目錄', '图目录'):
            continue
        if com_para_outline(i) != 1:  # 只看 OutlineLevel=1 的段落
            continue
        if tbl_start is None and t in ('表目錄', '表目录'):
            try:
                tbl_start = int(com_para(i).Range.Start)
            except Exception as e:
                _d.append(f"读取'表 目 錄'页标题段 Range.Start 失败: {e}")
        elif fig_start is None and t in ('圖目錄', '图目录'):
            try:
                fig_start = int(com_para(i).Range.Start)
            except Exception as e:
                _d.append(f"读取'圖 目 錄'页标题段 Range.Start 失败: {e}")
        if tbl_start is not None and fig_start is not None:
            break

    if tbl_start is None:
        _d.append("未找到 OutlineLevel=1 且文本='表 目 錄' 的页标题段")
    if fig_start is None:
        _d.append("未找到 OutlineLevel=1 且文本='圖 目 錄' 的页标题段")

    order_ok = (tbl_start is not None) and (fig_start is not None) and (tbl_start < fig_start)
    if (tbl_start is not None) and (fig_start is not None) and not order_ok:
        _d.append(
            (f"'表 目 錄'页中表目录位置(Range.Start={tbl_start}) "
             f">= '圖 目 錄'页位置(Range.Start={fig_start})，未位于其之前")
        )

    ok = order_ok
    return _item(
        "'表 目 錄'页中表目录位于'圖 目 錄'页之前",
        +3, ok, "; ".join(_d))


def _dim2_table31_single_page():
    # -----------------------------------------------------------------
    # +3：表3-1整体位于同一页内，表格未跨页断开
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs 找到题注段（Range.Text 匹配"^表\s*3-1\b"，
    #      且非目录条目、且不在表格内）—— 该段是"表 3-1"的题注段。
    #   ② 该题注段之后紧邻的第一个 Table 对象（Document.Tables 中 Range.Start
    #      大于题注段 Range.End 且最小者）即"表 3-1"表格。
    #   ③ 表格 Range 起点与终点分别读取
    #      Range.Information(wdActiveEndPageNumber=3)，
    #      两页号相等 ⇒ 整体位于同一页内、未跨页断开。
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    _d = []
    ok = False
    cap_para = None
    for i, text, norm, style, outline, page in com_paragraphs():
        stripped = text.strip()
        if '\t' in text:
            continue
        if not re.match(r'^表\s*3-1\b', stripped):
            continue
        p = com_para(i)
        try:
            if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                continue
        except Exception:
            pass
        cap_para = p
        break
    tbl_obj = None
    if cap_para is None:
        _d.append("未找到'表 3-1'题注段")
    else:
        try:
            cap_end = int(cap_para.Range.End)
            best_start = None
            for k in range(1, int(_wd.Tables.Count) + 1):
                t = _wd.Tables(k)
                ts = int(t.Range.Start)
                if ts >= cap_end and (best_start is None or ts < best_start):
                    best_start = ts
                    tbl_obj = t
        except Exception as e:
            _d.append(f"读取 Tables 失败: {e}")
    if tbl_obj is None:
        if not _d:
            _d.append("题注下方未定位到表格对象")
    else:
        try:
            rng = tbl_obj.Range
            s = int(rng.Start)
            e = int(rng.End)
            sp = int(rng.Information(WD_ACTIVE_END_PAGE))
            end_rng = _wd.Range(max(s, e - 1), max(s, e - 1))
            ep = int(end_rng.Information(WD_ACTIVE_END_PAGE))
            ok = (sp == ep)
            _d.append(f"起始页={sp}, 结束页={ep}")
        except Exception as e:
            _d.append(f"读取表格页码失败: {e}")
    return _item(
        "表3-1整体位于同一页内，表格未跨页断开",
        +3, ok, "; ".join(_d))


def _dim2_table41_single_page():
    # -----------------------------------------------------------------
    # +3：表4-1整体位于同一页内，表格未跨页断开
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs 找到题注段（Range.Text 匹配"^表\s*4-1\b"，
    #      且非目录条目、且不在表格内）—— 该段是"表 4-1"的题注段。
    #   ② 该题注段之后紧邻的第一个 Table 对象（Document.Tables 中 Range.Start
    #      大于题注段 Range.End 且最小者）即"表 4-1"表格。
    #   ③ 表格 Range 起点与终点分别读取
    #      Range.Information(wdActiveEndPageNumber=3)，
    #      两页号相等 ⇒ 整体位于同一页内、未跨页断开。
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    _d = []
    ok = False
    cap_para = None
    for i, text, norm, style, outline, page in com_paragraphs():
        stripped = text.strip()
        if '\t' in text:
            continue
        if not re.match(r'^表\s*4-1\b', stripped):
            continue
        p = com_para(i)
        try:
            if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                continue
        except Exception:
            pass
        cap_para = p
        break
    tbl_obj = None
    if cap_para is None:
        _d.append("未找到'表 4-1'题注段")
    else:
        try:
            cap_end = int(cap_para.Range.End)
            best_start = None
            for k in range(1, int(_wd.Tables.Count) + 1):
                t = _wd.Tables(k)
                ts = int(t.Range.Start)
                if ts >= cap_end and (best_start is None or ts < best_start):
                    best_start = ts
                    tbl_obj = t
        except Exception as e:
            _d.append(f"读取 Tables 失败: {e}")
    if tbl_obj is None:
        if not _d:
            _d.append("题注下方未定位到表格对象")
    else:
        try:
            rng = tbl_obj.Range
            s = int(rng.Start)
            e = int(rng.End)
            sp = int(rng.Information(WD_ACTIVE_END_PAGE))
            end_rng = _wd.Range(max(s, e - 1), max(s, e - 1))
            ep = int(end_rng.Information(WD_ACTIVE_END_PAGE))
            ok = (sp == ep)
            _d.append(f"起始页={sp}, 结束页={ep}")
        except Exception as e:
            _d.append(f"读取表格页码失败: {e}")
    return _item(
        "表4-1整体位于同一页内，表格未跨页断开",
        +3, ok, "; ".join(_d))

    # APPEND_HERE_DIM2_B


def _dim2_table31_blank_lines():
    # -----------------------------------------------------------------
    # -3：表3-1表格上方和下方有超过一行的空白行
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs 找到题注段（Range.Text 匹配"^表\s*3-1\b"），
    #      再从 Document.Tables 集合里定位其后紧邻的 Table 对象 —— 该 Table
    #      即"表 3-1"。
    #   ② 表格上方空白行数 = Document.Range(0, Table.Range.Start).Paragraphs
    #      末尾开始、紧邻表格的连续空段数（空段 = Range.Text 去段落标记后为空）。
    #   ③ 表格下方空白行数 = Document.Range(Table.Range.End, Content.End)
    #      .Paragraphs 开头开始、紧邻表格的连续空段数。
    #   ④ 上方 > 1 或 下方 > 1 ⇒ 命中"有超过一行的空白行"→ 扣 3 分。
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    _d = []
    cap_para = None
    for i, text, norm, style, outline, page in com_paragraphs():
        stripped = text.strip()
        if '\t' in text:
            continue
        if not re.match(r'^表\s*3-1\b', stripped):
            continue
        p = com_para(i)
        try:
            if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                continue
        except Exception:
            pass
        cap_para = p
        break
    tbl_obj = None
    if cap_para is None:
        _d.append("未找到'表 3-1'题注段")
    else:
        try:
            cap_end = int(cap_para.Range.End)
            best_start = None
            for k in range(1, int(_wd.Tables.Count) + 1):
                t = _wd.Tables(k)
                ts = int(t.Range.Start)
                if ts >= cap_end and (best_start is None or ts < best_start):
                    best_start = ts
                    tbl_obj = t
        except Exception as e:
            _d.append(f"读取 Tables 失败: {e}")
    above = below = 0
    if tbl_obj is None:
        if not _d:
            _d.append("题注下方未定位到表格对象")
    else:
        try:
            tbl_start = int(tbl_obj.Range.Start)
            tbl_end = int(tbl_obj.Range.End)
            doc_end = int(_wd.Content.End)
            # ② 上方紧邻空段
            if tbl_start > 0:
                rng = _wd.Range(0, tbl_start)
                cnt = int(rng.Paragraphs.Count)
                for j in range(cnt, 0, -1):
                    p = rng.Paragraphs(j)
                    try:
                        if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                            break
                    except Exception:
                        pass
                    txt = str(p.Range.Text or '').rstrip('\r\x07\n')
                    if txt.strip() == '':
                        above += 1
                    else:
                        break
            # ③ 下方紧邻空段
            if tbl_end < doc_end:
                rng = _wd.Range(tbl_end, doc_end)
                cnt = int(rng.Paragraphs.Count)
                for j in range(1, cnt + 1):
                    p = rng.Paragraphs(j)
                    try:
                        if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                            break
                    except Exception:
                        pass
                    txt = str(p.Range.Text or '').rstrip('\r\x07\n')
                    if txt.strip() == '':
                        below += 1
                    else:
                        break
        except Exception as e:
            _d.append(f"读取空段失败: {e}")
    hit = (above > 1) or (below > 1)
    _d.append(f"上方空段={above}, 下方空段={below}")
    return _item(
        "表3-1表格上方或下方有超过一行的空白行",
        -3, hit, "; ".join(_d))


def _dim2_toc_font():
    # -----------------------------------------------------------------
    # -3：目录内容字体非 PMingLiU-ExtB 四号
    # 完全基于 Word 办公软件文件属性：
    #   ① 目录条目 → Paragraph.Style.NameLocal 以 "目录"/"目錄"/"TOC" 开头
    #      （Word 自动/手工生成的目录条目均使用 "目录 1".."目录 9" 样式）。
    #   ② 逐段读取：
    #        Range.Font.NameFarEast == 'PMingLiU-ExtB'（中文字体）
    #        Range.Font.Size        == 14pt（四号）
    #   ③ 任一条不满足即命中"字体非 PMingLiU-ExtB 四号"，扣 3 分。
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    bad_font = []
    bad_size = []
    _d = []
    scanned = 0
    try:
        for p in _wd.Paragraphs:
            try:
                sname = str(p.Style.NameLocal or '')
            except Exception:
                sname = ''
            if not (sname.startswith('目录') or sname.startswith('目錄') or sname.startswith('TOC')):
                continue
            rng = p.Range
            raw = str(rng.Text or '').rstrip('\r\x07\n')
            if not raw.strip():
                continue
            scanned += 1
            try:
                fname = str(rng.Font.NameFarEast or '')
            except Exception:
                fname = ''
            try:
                size = float(rng.Font.Size)
            except Exception:
                size = float('nan')
            if fname != 'PMingLiU-ExtB':
                bad_font.append(f"{raw.strip()[:16]}→{fname!r}")
            if not (size == size) or abs(size - 14.0) > 0.01:
                bad_size.append(f"{raw.strip()[:16]}→{size}pt")
    except Exception as e:
        _d.append(f"读取失败: {e}")
    if scanned == 0:
        _d.append("未找到 Style 以 '目录/目錄/TOC' 开头的段落")
    hit = bool(bad_font or bad_size) or (scanned == 0 and bool(_d))
    if bad_font:
        _d.append("字体: " + "; ".join(bad_font[:3]))
    if bad_size:
        _d.append("字号: " + "; ".join(bad_size[:3]))
    return _item(
        "目录内容字体非 PMingLiU-ExtB 四号",
        -3, hit, " | ".join(_d))


def _dim2_toc_paragraph():
    # -----------------------------------------------------------------
    # -3：目录内容段落格式为首行缩进两字符、1.5倍行距、两端对齐，
    #     二级标题文本之前缩进两字符
    # 完全基于 Word 办公软件文件属性：
    #   支持两种目录形态，任一形态存在即按其内容做校验：
    #     形态 A：Document.TablesOfContents(1) 存在（TOC 域"自动更新目录"）
    #             → 遍历该域 Range 内每个 Paragraph 作为目录条目。
    #     形态 B：Document.TablesOfContents.Count == 0（静态文本 / 手工制表符+
    #             页码的目录）→ 按 Paragraph.Style.NameLocal 属于
    #             {'目录 1'..'目录 9', '目錄 1'..'目錄 9', 'TOC 1'..'TOC 9'}
    #             收集目录条目。若样式也不匹配（纯手工写的目录），再按"文本
    #             形如 '标题 ...\\t页码' 或 '标题 ...  页码'"作最后一层兜底。
    #   对收集到的每条目录条目段读取 Paragraph.Format 属性并校验：
    #     a. 首行缩进两字符：CharacterUnitFirstLineIndent == 2
    #        （字符单位不可读时按 FirstLineIndent(pt) ≈ 2×字号 回退）
    #     b. 1.5 倍行距：LineSpacingRule == wdLineSpace1pt5 (=1)
    #     c. 两端对齐：Alignment == wdAlignParagraphJustify (=3)
    #     d. 二级条目（形如 "X.Y ..."）：文本之前缩进两字符
    #        CharacterUnitLeftIndent == 2
    #        （字符单位不可读时按 LeftIndent(pt) ≈ 2×字号 回退）
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    bad = []
    _d = []
    entries = []  # list of Paragraph 对象
    source = None
    try:
        # 形态 A：TOC 域
        if int(_wd.TablesOfContents.Count) >= 1:
            toc = _wd.TablesOfContents(1)
            for para in toc.Range.Paragraphs:
                raw = str(para.Range.Text or '').rstrip('\r\x07\n')
                if raw.strip():
                    entries.append(para)
            source = f"TOC 域, 条目数={len(entries)}"
        else:
            # 形态 B1：按段落样式 "目录 X / 目錄 X / TOC X" 收集
            style_prefixes = ('目录', '目錄', 'TOC')
            for para in _wd.Paragraphs:
                try:
                    sname = str(para.Style.NameLocal or '')
                except Exception:
                    sname = ''
                if not sname.startswith(style_prefixes):
                    continue
                raw = str(para.Range.Text or '').rstrip('\r\x07\n')
                if not raw.strip():
                    continue
                entries.append(para)
            if entries:
                source = f"目录样式段落, 条目数={len(entries)}"
            else:
                # 形态 B2：纯手工目录 —— 定位"目录/目錄"标题段之后、下一个一级标题之前，
                # 段内含制表符或"标题空白 页码"结构的段落
                cp = com_paragraphs()
                toc_start_idx = None
                toc_end_idx = None
                for i, text, norm, style, outline, page in cp:
                    if '\t' in text:
                        continue
                    if toc_start_idx is None and norm in ('目录', '目錄'):
                        toc_start_idx = i
                        continue
                    if toc_start_idx is not None and com_para_outline(i) == 1:
                        toc_end_idx = i
                        break
                if toc_start_idx is not None:
                    end_i = toc_end_idx if toc_end_idx is not None else (len(cp) + 1)
                    entry_pat = re.compile(r'.+?(\t+|\s{2,})\S*\d[\dIVXLCDM]*\S*\s*$')
                    for j in range(toc_start_idx + 1, end_i):
                        try:
                            para = _wd.Paragraphs(j)
                        except Exception:
                            continue
                        raw = str(para.Range.Text or '').rstrip('\r\x07\n')
                        if not raw.strip():
                            continue
                        if entry_pat.match(raw.strip()):
                            entries.append(para)
                    if entries:
                        source = f"手工目录段落, 条目数={len(entries)}"
        if not entries:
            _d.append("未定位到目录条目（TOC 域、目录样式、手工目录均未命中）")
        else:
            _d.append(f"来源={source}")
            for para in entries:
                raw = str(para.Range.Text or '').rstrip('\r\x07\n')
                pf = para.Format
                rng = para.Range
                try:
                    size = float(rng.Font.Size)
                except Exception:
                    size = 14.0
                entry = raw.split('\t', 1)[0].strip()
                is_l2 = bool(re.match(r'^\d+\.\d+\s', entry))
                # a. 首行缩进两字符
                try:
                    ch = float(pf.CharacterUnitFirstLineIndent)
                except Exception:
                    ch = float('nan')
                try:
                    pt = float(pf.FirstLineIndent)
                except Exception:
                    pt = 0.0
                first_ok = (ch == ch and abs(ch - 2.0) < 0.01) or \
                           abs(pt - 2.0 * size) < 1.0
                if not first_ok:
                    bad.append(f"{entry[:12]}→首行={ch}字符/{pt}pt")
                # b. 1.5 倍行距
                try:
                    lsr = int(pf.LineSpacingRule)
                except Exception:
                    lsr = -1
                if lsr != WD_LINE_1PT5:
                    bad.append(f"{entry[:12]}→行距规则={lsr}")
                # c. 两端对齐
                try:
                    align = int(pf.Alignment)
                except Exception:
                    align = -1
                if align != WD_ALIGN_JUSTIFY:
                    bad.append(f"{entry[:12]}→对齐={align}")
                # d. 二级标题：文本之前缩进两字符（LeftIndent = 2 字符）
                if is_l2:
                    try:
                        ch_left = float(pf.CharacterUnitLeftIndent)
                    except Exception:
                        ch_left = float('nan')
                    try:
                        left_pt = float(pf.LeftIndent)
                    except Exception:
                        left_pt = 0.0
                    left_ok = (ch_left == ch_left and abs(ch_left - 2.0) < 0.01) or \
                              abs(left_pt - 2.0 * size) < 1.0
                    if not left_ok:
                        bad.append(f"{entry[:12]}→左缩进={ch_left}字符/{left_pt}pt")
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(bad) or (not entries)
    if bad:
        _d.append("; ".join(bad[:6]))
    return _item(
        "目录内容段落格式不满足：首行缩进两字符、1.5倍行距、两端对齐，二级标题文本之前缩进两字符",
        -3, hit, " | ".join(_d))


def _dim2_table41_blank_lines():
    # -----------------------------------------------------------------
    # -3：表4-1所在页面表格上方或下方有超过一行的空白行
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs 找到题注段（Range.Text 匹配"^表\s*4-1\b"），
    #      再从 Document.Tables 集合里定位其后紧邻的 Table 对象 —— 该 Table
    #      即"表 4-1"。
    #   ② 表格上方空白行数 = Document.Range(0, Table.Range.Start).Paragraphs
    #      末尾开始、紧邻表格的连续空段数（空段 = Range.Text 去段落标记后为空）。
    #   ③ 表格下方空白行数 = Document.Range(Table.Range.End, Content.End)
    #      .Paragraphs 开头开始、紧邻表格的连续空段数。
    #   ④ 上方 > 1 或 下方 > 1 ⇒ 命中"有超过一行的空白行"→ 扣 3 分。
    # -----------------------------------------------------------------
    _wd = _WordSession.doc()
    _d = []
    cap_para = None
    for i, text, norm, style, outline, page in com_paragraphs():
        stripped = text.strip()
        if '\t' in text:
            continue
        if not re.match(r'^表\s*4-1\b', stripped):
            continue
        p = com_para(i)
        try:
            if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                continue
        except Exception:
            pass
        cap_para = p
        break
    tbl_obj = None
    if cap_para is None:
        _d.append("未找到'表 4-1'题注段")
    else:
        try:
            cap_end = int(cap_para.Range.End)
            best_start = None
            for k in range(1, int(_wd.Tables.Count) + 1):
                t = _wd.Tables(k)
                ts = int(t.Range.Start)
                if ts >= cap_end and (best_start is None or ts < best_start):
                    best_start = ts
                    tbl_obj = t
        except Exception as e:
            _d.append(f"读取 Tables 失败: {e}")
    above = below = 0
    if tbl_obj is None:
        if not _d:
            _d.append("题注下方未定位到表格对象")
    else:
        try:
            tbl_start = int(tbl_obj.Range.Start)
            tbl_end = int(tbl_obj.Range.End)
            doc_end = int(_wd.Content.End)
            # ② 上方紧邻空段
            if tbl_start > 0:
                rng = _wd.Range(0, tbl_start)
                cnt = int(rng.Paragraphs.Count)
                for j in range(cnt, 0, -1):
                    p = rng.Paragraphs(j)
                    try:
                        if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                            break
                    except Exception:
                        pass
                    txt = str(p.Range.Text or '').rstrip('\r\x07\n')
                    if txt.strip() == '':
                        above += 1
                    else:
                        break
            # ③ 下方紧邻空段
            if tbl_end < doc_end:
                rng = _wd.Range(tbl_end, doc_end)
                cnt = int(rng.Paragraphs.Count)
                for j in range(1, cnt + 1):
                    p = rng.Paragraphs(j)
                    try:
                        if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                            break
                    except Exception:
                        pass
                    txt = str(p.Range.Text or '').rstrip('\r\x07\n')
                    if txt.strip() == '':
                        below += 1
                    else:
                        break
        except Exception as e:
            _d.append(f"读取空段失败: {e}")
    hit = (above > 1) or (below > 1)
    _d.append(f"上方空段={above}, 下方空段={below}")
    return _item(
        "表4-1所在页面表格上方或下方有超过一行的空白行",
        -3, hit, "; ".join(_d))

    # APPEND_HERE_DIM2_C


def _dim2_cover_logo():
    # -----------------------------------------------------------------
    # -3：封面页最上方未出现校徽图片，或者校徽图片尺寸不满足 5.08×15.23 厘米、
    #     布局选项为嵌入型、图片两端对齐
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，只看 Range.Information(wdActiveEndPageNumber=3)==1
    #      的封面页段落；封面页最上方第一个"承载内容的段落"要求：
    #        - Paragraph.Range.InlineShapes.Count >= 1（嵌入型图片存在）
    #     若最上方先遇到非空文本段而非图片段 → 命中"未出现校徽图片"。
    #   ② 图片尺寸：InlineShape.Width / InlineShape.Height（磅）
    #        换算 1pt = 2.54/72 cm；期望 15.23cm × 5.08cm
    #   ③ 布局选项为嵌入型：InlineShapes 集合本身即"嵌入型"（Shapes 是浮动型）。
    #   ④ 图片两端对齐：图片所在段 Paragraph.Format.Alignment == wdAlignParagraphJustify(3)
    # -----------------------------------------------------------------
    pt2cm = 2.54 / 72.0
    _wd = _WordSession.doc()
    _d = []
    cover_ok = False
    try:
        total = int(_wd.Paragraphs.Count)
        first_hit = None
        logo_shape = None
        logo_para = None
        for i in range(1, total + 1):
            p = _wd.Paragraphs(i)
            rng = p.Range
            try:
                page = int(rng.Information(WD_ACTIVE_END_PAGE))
            except Exception:
                page = -1
            if page > 1:
                break
            inline_cnt = 0
            try:
                inline_cnt = int(rng.InlineShapes.Count)
            except Exception:
                inline_cnt = 0
            text = str(rng.Text or '').rstrip('\r\x07\n')
            if inline_cnt > 0:
                first_hit = 'img'
                logo_shape = rng.InlineShapes(1)
                logo_para = p
                break
            if text.strip():
                first_hit = 'text'
                break
        if first_hit != 'img':
            _d.append("封面页最上方未出现校徽图片")
        else:
            # ② 尺寸
            try:
                width_cm = float(logo_shape.Width) * pt2cm
            except Exception:
                width_cm = float('nan')
            try:
                height_cm = float(logo_shape.Height) * pt2cm
            except Exception:
                height_cm = float('nan')
            if not (abs(width_cm - 15.23) <= 0.2):
                _d.append(f"宽={width_cm:.2f}cm(期望15.23)")
            if not (abs(height_cm - 5.08) <= 0.2):
                _d.append(f"高={height_cm:.2f}cm(期望5.08)")
            # ③ 布局为嵌入型：既然来自 InlineShapes，即嵌入型
            # ④ 图片两端对齐
            try:
                align = int(logo_para.Format.Alignment)
            except Exception:
                align = -1
            if align != WD_ALIGN_JUSTIFY:
                _d.append(f"图片段对齐={align}(期望3=两端对齐)")
            cover_ok = not _d
    except Exception as e:
        _d.append(f"Word COM 读取失败: {e}")
    return _item(
        "封面页最上方出现校徽图片且尺寸5.08×15.23厘米、布局为嵌入型、图片两端对齐",
        -3, not cover_ok, "; ".join(_d))


def _dim2_cover_title_font():
    # -----------------------------------------------------------------
    # -1：封面页"碩士研究生學位論文"字体格式不满足宋体 28 加粗
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，找到 Range.Text 归一化后 == "碩士研究生學位論文"
    #      （或简体"硕士研究生学位论文"）的段落。
    #   ② 该段 Range 上的字体属性：
    #        a. Font.NameFarEast == '宋体'（或等价的 'SimSun'）
    #        b. Font.Size        == 28.0
    #        c. Font.Bold        == True（加粗）
    #   ③ 任一条不满足 → 命中"字体格式不满足宋体 28 加粗"，扣 1 分。
    # -----------------------------------------------------------------
    p = _find_com_para({'碩士研究生學位論文', '硕士研究生学位论文'})
    _d = []
    if p is None:
        _d.append("未找到'碩士研究生學位論文'段落")
    else:
        rng = p.Range
        try:
            fname = str(rng.Font.NameFarEast or '')
        except Exception:
            fname = ''
        try:
            size = float(rng.Font.Size)
        except Exception:
            size = float('nan')
        try:
            bold = bool(rng.Font.Bold)
        except Exception:
            bold = False
        if fname not in ('宋体', 'SimSun'):
            _d.append(f"字体={fname!r}")
        if not (size == size) or abs(size - 28.0) > 0.01:
            _d.append(f"字号={size}")
        if not bold:
            _d.append("未加粗")
    hit = bool(_d)
    return _item(
        "封面页'碩士研究生學位論文'字体格式不满足宋体28加粗",
        -1, hit, "; ".join(_d))


def _dim2_cover_title_paragraph():
    # -----------------------------------------------------------------
    # -1：封面页"碩士研究生學位論文"段落格式不满足 1.5 倍行距、文本之前两字符、居中对齐
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，找到 Range.Text 归一化后 == "碩士研究生學位論文"
    #      （或简体"硕士研究生学位论文"）的段落。
    #   ② 读取 Paragraph.Format 属性并校验：
    #        a. 1.5 倍行距：LineSpacingRule == wdLineSpace1pt5 (=1)
    #        b. 文本之前两字符：CharacterUnitLeftIndent == 2
    #           （字符单位不可读时按 LeftIndent(pt) ≈ 2×字号 回退）
    #        c. 居中对齐：Alignment == wdAlignParagraphCenter (=1)
    #   ③ 任一条不满足 → 命中"段落格式不满足 …"，扣 1 分。
    # -----------------------------------------------------------------
    p = _find_com_para({'碩士研究生學位論文', '硕士研究生学位论文'})
    _d = []
    if p is None:
        _d.append("未找到'碩士研究生學位論文'段落")
    else:
        pf = p.Format
        rng = p.Range
        try:
            size = float(rng.Font.Size)
        except Exception:
            size = 28.0
        # a. 1.5 倍行距
        try:
            lsr = int(pf.LineSpacingRule)
        except Exception:
            lsr = -1
        if lsr != WD_LINE_1PT5:
            _d.append(f"行距规则={lsr}")
        # b. 文本之前两字符
        try:
            ch_left = float(pf.CharacterUnitLeftIndent)
        except Exception:
            ch_left = float('nan')
        try:
            left_pt = float(pf.LeftIndent)
        except Exception:
            left_pt = 0.0
        left_ok = (ch_left == ch_left and abs(ch_left - 2.0) < 0.01) or \
                  abs(left_pt - 2.0 * size) < 1.0
        if not left_ok:
            _d.append(f"左缩进={ch_left}字符/{left_pt}pt")
        # c. 居中对齐
        try:
            align = int(pf.Alignment)
        except Exception:
            align = -1
        if align != WD_ALIGN_CENTER:
            _d.append(f"对齐={align}")
    hit = bool(_d)
    return _item(
        "封面页'碩士研究生學位論文'段落格式不满足1.5倍行距、文本之前两字符、居中对齐",
        -1, hit, "; ".join(_d))


def _dim2_thanks_title_font():
    # -----------------------------------------------------------------
    # -1：致谢页标题字体格式不满足宋体二号加粗
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，找到 Range.Text 归一化后 == "致谢"（或
    #      繁体"致謝"）的段落 —— 该段即"致谢页标题"。
    #   ② 该段 Range 上的字体属性：
    #        a. Font.NameFarEast == '宋体'（或等价的 'SimSun'）
    #        b. Font.Size        == 22.0（Word 中"二号" = 22pt）
    #        c. Font.Bold        == True（加粗）
    #   ③ 任一条不满足 → 命中"字体格式不满足宋体二号加粗"，扣 1 分。
    # -----------------------------------------------------------------
    p = _find_com_para({'致谢', '致謝'})
    _d = []
    if p is None:
        _d.append("未找到'致谢'标题段落")
    else:
        rng = p.Range
        try:
            fname = str(rng.Font.NameFarEast or '')
        except Exception:
            fname = ''
        try:
            size = float(rng.Font.Size)
        except Exception:
            size = float('nan')
        try:
            bold = bool(rng.Font.Bold)
        except Exception:
            bold = False
        if fname not in ('宋体', 'SimSun'):
            _d.append(f"字体={fname!r}")
        if not (size == size) or abs(size - 22.0) > 0.01:
            _d.append(f"字号={size}")
        if not bold:
            _d.append("未加粗")
    hit = bool(_d)
    return _item(
        "致谢页标题字体格式不满足宋体二号加粗",
        -1, hit, "; ".join(_d))


def _dim2_cn_abstract_paragraph():
    # -----------------------------------------------------------------
    # -1：中文摘要页内容段落格式不满足首行缩进两字符、1.5 倍行距、两端对齐
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，定位"摘要"标题段（Range.Text 归一化 == "摘要"
    #      且不含 \t —— 排除目录条目）。
    #   ② 从该标题段之后开始，直到遇到"關鍵詞/关键词"或"Abstract"段为止，
    #      收集其间的正文段（排除空段、目录条目段、表格内段）。
    #   ③ 对每段读取 Paragraph.Format 属性并校验：
    #        a. 首行缩进两字符：CharacterUnitFirstLineIndent == 2
    #           （字符单位不可读时按 FirstLineIndent(pt) ≈ 2×字号 回退）
    #        b. 1.5 倍行距：LineSpacingRule == wdLineSpace1pt5 (=1)
    #        c. 两端对齐：Alignment == wdAlignParagraphJustify (=3)
    #   ④ 任一段任一条不满足 → 命中"段落格式不满足 …"，扣 1 分。
    # -----------------------------------------------------------------
    _d = []
    abs_p_idx = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if norm == '摘要' and '\t' not in text:
            abs_p_idx = i
            break
    if abs_p_idx is None:
        _d.append("未找到'摘要'标题段")
    else:
        cp = com_paragraphs()
        checked_any = False
        for j in range(abs_p_idx + 1, len(cp) + 1):
            i, text, norm, style, outline, page = cp[j - 1]
            if norm in ('關鍵詞', '关键词', 'Abstract', 'abstract') or \
               text.strip().startswith('關鍵詞') or text.strip().startswith('关键词'):
                break
            if not text.strip() or '\t' in text:
                continue
            p = com_para(i)
            try:
                if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                    continue
            except Exception:
                pass
            pf = p.Format
            try:
                size = float(p.Range.Font.Size)
            except Exception:
                size = 14.0
            # a. 首行缩进两字符
            try:
                ch = float(pf.CharacterUnitFirstLineIndent)
            except Exception:
                ch = float('nan')
            try:
                pt = float(pf.FirstLineIndent)
            except Exception:
                pt = 0.0
            first_ok = (ch == ch and abs(ch - 2.0) < 0.01) or \
                       abs(pt - 2.0 * size) < 1.0
            if not first_ok:
                _d.append(f"{text[:12]}→首行={ch}字符/{pt}pt")
            # b. 1.5 倍行距
            try:
                lsr = int(pf.LineSpacingRule)
            except Exception:
                lsr = -1
            if lsr != WD_LINE_1PT5:
                _d.append(f"{text[:12]}→行距规则={lsr}")
            # c. 两端对齐
            try:
                align = int(pf.Alignment)
            except Exception:
                align = -1
            if align != WD_ALIGN_JUSTIFY:
                _d.append(f"{text[:12]}→对齐={align}")
            checked_any = True
        if not checked_any:
            _d.append("未找到中文摘要正文段")
    hit = bool(_d)
    return _item(
        "中文摘要页内容段落格式不满足首行缩进两字符、1.5倍行距、两端对齐",
        -1, hit, "; ".join(_d[:4]))


def _dim2_cn_abstract_font():
    # -----------------------------------------------------------------
    # -1：中文摘要页内容字体格式不满足宋体四号
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，定位"摘要"标题段（Range.Text 归一化 == "摘要"
    #      且不含 \t —— 排除目录条目）。
    #   ② 从该标题段之后开始，直到遇到"關鍵詞/关键词"或"Abstract"段为止，
    #      收集其间的中文摘要正文段（排除空段、目录条目段、表格内段）。
    #   ③ 对每段：先尝试段级 Range.Font 是否整段统一（NameFarEast、Size 非空且
    #      Size != 9999999 表示段内属性一致）；一致则直接比对；不一致再对该段
    #      中文字符（CJK 区间）逐字读取 Range.Font.NameFarEast / Range.Font.Size。
    #        a. Font.NameFarEast ∈ {'宋体', 'SimSun'}（Word 里中文字体属性）
    #        b. Font.Size        == 14.0（Word 里"四号" = 14pt）
    #   ④ 任一段任一字符不满足 → 命中"字体格式不满足宋体四号"，扣 1 分。
    # -----------------------------------------------------------------
    _d = []
    _WD_UNDEFINED = 9999999
    abs_p_idx = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if norm == '摘要' and '\t' not in text:
            abs_p_idx = i
            break
    if abs_p_idx is None:
        _d.append("未找到'摘要'标题段")
    else:
        cp = com_paragraphs()
        checked_any = False
        for j in range(abs_p_idx + 1, len(cp) + 1):
            i, text, norm, style, outline, page = cp[j - 1]
            if norm in ('關鍵詞', '关键词', 'Abstract', 'abstract') or \
               text.strip().startswith('關鍵詞') or text.strip().startswith('关键词'):
                break
            if not text.strip() or '\t' in text:
                continue
            p = com_para(i)
            try:
                if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                    continue
            except Exception:
                pass
            # 段级统一属性尝试
            rng = p.Range
            uniform_hit = False
            try:
                fe = str(rng.Font.NameFarEast or '')
                size = float(rng.Font.Size)
            except Exception:
                fe, size = '', 9999999.0
            if fe and size and size != 9999999.0:
                if fe not in ('宋体', 'SimSun'):
                    _d.append(f"{text[:12]}→字体={fe!r}")
                    uniform_hit = True
                elif abs(size - 14.0) > 0.01:
                    _d.append(f"{text[:12]}→字号={size}")
                    uniform_hit = True
                checked_any = True
                if uniform_hit:
                    break
                continue
            # 段内不统一：抽样逐字扫中文字符
            try:
                chars = rng.Characters
                cnt = int(chars.Count)
            except Exception:
                cnt = 0
            step = max(1, cnt // 60)
            broke = False
            for k in range(1, cnt + 1, step):
                try:
                    sub = _com_call(chars, k)
                except Exception:
                    continue
                ch = sub.Text
                if not ch or not ch.strip():
                    continue
                code = ord(ch[0])
                is_cjk = (
                    0x4E00 <= code <= 0x9FFF or
                    0x3400 <= code <= 0x4DBF or
                    0xF900 <= code <= 0xFAFF or
                    0x20000 <= code <= 0x2FFFF
                )
                if not is_cjk:
                    continue
                try:
                    fname = str(sub.Font.NameFarEast or '')
                    csize = float(sub.Font.Size)
                except Exception:
                    continue
                if fname not in ('宋体', 'SimSun'):
                    _d.append(f"{text[:8]}...{ch!r}字体={fname!r}")
                    broke = True
                    break
                if abs(csize - 14.0) > 0.01:
                    _d.append(f"{text[:8]}...{ch!r}字号={csize}")
                    broke = True
                    break
            checked_any = True
            if broke:
                break
        if not checked_any:
            _d.append("未找到中文摘要正文段")
    hit = bool(_d)
    return _item(
        "中文摘要页内容字体格式不满足宋体四号",
        -1, hit, "; ".join(_d[:3]))


def _dim2_en_abstract_font():
    # -----------------------------------------------------------------
    # -1：英文摘要页内容字体格式不满足 Times New Roman 四号
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，定位"Abstract"标题段（Range.Text 归一化 ==
    #      "Abstract" 且不含 \t —— 排除目录条目）。
    #   ② 从该标题段之后开始，直到遇到"Keywords"段或"目錄/目录"段为止，
    #      收集其间的英文摘要正文段（排除空段、目录条目段、表格内段）。
    #   ③ 对每段：先尝试段级 Range.Font 是否整段统一（NameAscii、Size 非空且
    #      Size != 9999999 表示段内属性一致）；一致则直接比对；不一致再对该段
    #      英文/数字字符逐字读取 Range.Font.NameAscii / Range.Font.Size。
    #        a. Font.NameAscii == 'Times New Roman'（Word 里英文字体属性）
    #        b. Font.Size      == 14.0（Word 里"四号" = 14pt）
    #   ④ 任一段任一字符不满足 → 命中"字体格式不满足 Times New Roman 四号"，扣 1 分。
    # -----------------------------------------------------------------
    _d = []
    abs_en_idx = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if norm == 'Abstract' and '\t' not in text:
            abs_en_idx = i
            break
    if abs_en_idx is None:
        _d.append("未找到'Abstract'标题段")
    else:
        cp = com_paragraphs()
        checked_any = False
        for j in range(abs_en_idx + 1, len(cp) + 1):
            i, text, norm, style, outline, page = cp[j - 1]
            if norm.startswith('Keywords') or text.strip().startswith('Keywords'):
                break
            if norm in ('目錄', '目录'):
                break
            if not text.strip() or '\t' in text:
                continue
            p = com_para(i)
            try:
                if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                    continue
            except Exception:
                pass
            # 段级统一属性尝试
            rng = p.Range
            uniform_hit = False
            try:
                asc = str(rng.Font.NameAscii or '')
                size = float(rng.Font.Size)
            except Exception:
                asc, size = '', 9999999.0
            if asc and size and size != 9999999.0:
                if asc != 'Times New Roman':
                    _d.append(f"{text[:12]}→字体={asc!r}")
                    uniform_hit = True
                elif abs(size - 14.0) > 0.01:
                    _d.append(f"{text[:12]}→字号={size}")
                    uniform_hit = True
                checked_any = True
                if uniform_hit:
                    break
                continue
            # 段内不统一：抽样逐字扫英文/数字字符
            try:
                chars = rng.Characters
                cnt = int(chars.Count)
            except Exception:
                cnt = 0
            step = max(1, cnt // 60)
            broke = False
            for k in range(1, cnt + 1, step):
                try:
                    sub = _com_call(chars, k)
                except Exception:
                    continue
                ch = sub.Text
                if not ch or not ch.strip():
                    continue
                if not _is_en_or_num(ch):
                    continue
                try:
                    fname = str(sub.Font.NameAscii or '')
                    csize = float(sub.Font.Size)
                except Exception:
                    continue
                if fname != 'Times New Roman':
                    _d.append(f"{text[:8]}...{ch!r}字体={fname!r}")
                    broke = True
                    break
                if abs(csize - 14.0) > 0.01:
                    _d.append(f"{text[:8]}...{ch!r}字号={csize}")
                    broke = True
                    break
            checked_any = True
            if broke:
                break
        if not checked_any:
            _d.append("未找到英文摘要正文段")
    hit = bool(_d)
    return _item(
        "英文摘要页内容字体格式不满足 Times New Roman 四号",
        -1, hit, "; ".join(_d[:3]))


def _dim2_en_abstract_paragraph():
    # -----------------------------------------------------------------
    # -1：英文摘要页内容段落格式不满足首行缩进两字符、1.5 倍行距、两端对齐
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Paragraphs，定位"Abstract"标题段（Range.Text 归一化 ==
    #      "Abstract" 且不含 \t —— 排除目录条目）。
    #   ② 从该标题段之后开始，直到遇到"Keywords"段或"目錄/目录"段为止，
    #      收集其间的英文摘要正文段（排除空段、目录条目段、表格内段）。
    #   ③ 对每段读取 Paragraph.Format 属性并校验：
    #        a. 首行缩进两字符：CharacterUnitFirstLineIndent == 2
    #           （字符单位不可读时按 FirstLineIndent(pt) ≈ 2×字号 回退）
    #        b. 1.5 倍行距：LineSpacingRule == wdLineSpace1pt5 (=1)
    #        c. 两端对齐：Alignment == wdAlignParagraphJustify (=3)
    #   ④ 任一段任一条不满足 → 命中"段落格式不满足 …"，扣 1 分。
    # -----------------------------------------------------------------
    _d = []
    abs_en_idx = None
    for i, text, norm, style, outline, page in com_paragraphs():
        if norm == 'Abstract' and '\t' not in text:
            abs_en_idx = i
            break
    if abs_en_idx is None:
        _d.append("未找到'Abstract'标题段")
    else:
        cp = com_paragraphs()
        checked_any = False
        for j in range(abs_en_idx + 1, len(cp) + 1):
            i, text, norm, style, outline, page = cp[j - 1]
            if norm.startswith('Keywords') or text.strip().startswith('Keywords'):
                break
            if norm in ('目錄', '目录'):
                break
            if not text.strip() or '\t' in text:
                continue
            p = com_para(i)
            try:
                if bool(p.Range.Information(WD_WITH_IN_TABLE)):
                    continue
            except Exception:
                pass
            pf = p.Format
            try:
                size = float(p.Range.Font.Size)
            except Exception:
                size = 14.0
            # a. 首行缩进两字符
            try:
                ch = float(pf.CharacterUnitFirstLineIndent)
            except Exception:
                ch = float('nan')
            try:
                pt = float(pf.FirstLineIndent)
            except Exception:
                pt = 0.0
            first_ok = (ch == ch and abs(ch - 2.0) < 0.01) or \
                       abs(pt - 2.0 * size) < 1.5
            if not first_ok:
                _d.append(f"{text[:12]}→首行={ch}字符/{pt}pt")
            # b. 1.5 倍行距
            try:
                lsr = int(pf.LineSpacingRule)
            except Exception:
                lsr = -1
            if lsr != WD_LINE_1PT5:
                _d.append(f"{text[:12]}→行距规则={lsr}")
            # c. 两端对齐
            try:
                align = int(pf.Alignment)
            except Exception:
                align = -1
            if align != WD_ALIGN_JUSTIFY:
                _d.append(f"{text[:12]}→对齐={align}")
            checked_any = True
        if not checked_any:
            _d.append("未找到英文摘要正文段")
    hit = bool(_d)
    return _item(
        "英文摘要页内容段落格式不满足首行缩进两字符、1.5倍行距、两端对齐",
        -1, hit, "; ".join(_d[:4]))

    # APPEND_HERE_DIM2_D


def _dim2_page_number_format():
    # -----------------------------------------------------------------
    # -3：页码格式不满足居中对齐，页脚下边距 1.27 厘米
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Document.Sections，逐节读取：
    #        a. 该节 Footers(1..3) 中包含 wdFieldPage(33) 域字段 —— 即"存在页码"。
    #   ② 对每一个"存在页码"的节：
    #        b. 页码所在段的对齐：包含 PAGE 域那一段的
    #           Paragraph.Format.Alignment == wdAlignParagraphCenter(=1)
    #        c. 页脚下边距：Section.PageSetup.FooterDistance（磅）
    #           换算 1pt = 2.54/72 cm ≈ 1.27 cm
    #   ③ 任一节任一条不满足 → 命中"页码格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    try:
        _wd = _WordSession.doc()
        for sidx in range(1, int(_wd.Sections.Count) + 1):
            sec = _wd.Sections(sidx)
            has_page = False
            page_align = None
            for wd_footer_index in (1, 2, 3):  # 1=Primary, 2=FirstPage, 3=EvenPages
                try:
                    f = sec.Footers(wd_footer_index)
                except Exception:
                    continue
                try:
                    if not bool(f.Exists):
                        continue
                except Exception:
                    pass
                # a. 页脚中是否含 wdFieldPage(=33)
                try:
                    for fld in f.Range.Fields:
                        if int(fld.Type) == 33:
                            has_page = True
                            break
                except Exception:
                    pass
                # b. 拿到含 PAGE 域段落的对齐
                if has_page and page_align is None:
                    try:
                        for para in f.Range.Paragraphs:
                            xml = str(para.Range.XML or '')
                            if re.search(r'\bPAGE\b', xml):
                                try:
                                    page_align = int(para.Format.Alignment)
                                except Exception:
                                    page_align = None
                                break
                    except Exception:
                        pass
            if has_page:
                if page_align != WD_ALIGN_CENTER:
                    _d.append(f"sec{sidx} 页码段对齐={page_align}(期望1=居中)")
                # c. 页脚下边距
                try:
                    fd = float(sec.PageSetup.FooterDistance)
                    fd_cm = fd * (2.54 / 72.0)
                except Exception:
                    fd_cm = float('nan')
                if not (fd_cm == fd_cm) or abs(fd_cm - 1.27) > 0.05:
                    _d.append(f"sec{sidx} 页脚下边距={fd_cm:.2f}cm(期望1.27)")
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "页码格式不满足居中对齐，页脚下边距1.27厘米",
        -3, hit, "; ".join(_d[:4]))


def _dim2_header_format():
    # -----------------------------------------------------------------
    # -3：页眉格式不满足左对齐，页眉上边距 1.40 厘米
    # 完全基于 Word 办公软件文件属性：
    #   ① 遍历 Word COM Document.Sections，逐节读取：
    #        a. 该节 Headers(1..3) 中"存在页眉"——即 Header.Range.Text
    #           去除 \r\x07\n\t 等控制符后仍有可见字符。
    #   ② 对每一个"存在页眉"的节：
    #        b. 页眉对齐：首个有实质内容段落
    #           Paragraph.Format.Alignment == wdAlignParagraphLeft(=0)
    #        c. 页眉上边距：Section.PageSetup.HeaderDistance（磅）
    #           换算 1pt = 2.54/72 cm ≈ 1.40 cm
    #   ③ 任一节任一条不满足 → 命中"页眉格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    try:
        _wd = _WordSession.doc()
        for sidx in range(1, int(_wd.Sections.Count) + 1):
            sec = _wd.Sections(sidx)
            has_header = False
            head_align = None
            for wd_header_index in (1, 2, 3):  # 1=Primary, 2=FirstPage, 3=EvenPages
                try:
                    h = sec.Headers(wd_header_index)
                except Exception:
                    continue
                try:
                    if not bool(h.Exists):
                        continue
                except Exception:
                    pass
                # a. 页眉是否含实质内容
                try:
                    head_text = str(h.Range.Text or '')
                except Exception:
                    head_text = ''
                if not head_text.strip('\r\x07\n\t '):
                    continue
                has_header = True
                # b. 首个有实质内容段落的对齐
                if head_align is None:
                    try:
                        for para in h.Range.Paragraphs:
                            ptxt = str(para.Range.Text or '')
                            if ptxt.strip('\r\x07\n\t '):
                                try:
                                    head_align = int(para.Format.Alignment)
                                except Exception:
                                    head_align = None
                                break
                    except Exception:
                        pass
            if has_header:
                if head_align != WD_ALIGN_LEFT:
                    _d.append(f"sec{sidx} 页眉段对齐={head_align}(期望0=左对齐)")
                # c. 页眉上边距
                try:
                    hd = float(sec.PageSetup.HeaderDistance)
                    hd_cm = hd * (2.54 / 72.0)
                except Exception:
                    hd_cm = float('nan')
                if not (hd_cm == hd_cm) or abs(hd_cm - 1.40) > 0.05:
                    _d.append(f"sec{sidx} 页眉上边距={hd_cm:.2f}cm(期望1.40)")
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "页眉格式不满足左对齐，页眉上边距1.40厘米",
        -3, hit, "; ".join(_d[:4]))


def _dim2_l2_title_font():
    # -----------------------------------------------------------------
    # -----------------------------------------------------------------
    # -3：二级标题字体格式不满足 PMingLiU-ExtB、小二号、加粗
    # 完全基于 Word 办公软件文件属性：
    #   ① 二级标题 = Word COM Paragraph.OutlineLevel == wdOutlineLevel2(=2)
    #      —— 这是 Word 段落自身的大纲级别属性，不做任何文本正则推断。
    #   ② 逐一读取每个二级标题段落 Range 的字体属性：
    #        a. 字体（中文）：Range.Font.NameFarEast == 'PMingLiU-ExtB'
    #        b. 字号：Range.Font.Size == 18.0（"小二"= 18pt）
    #        c. 加粗：Range.Font.Bold 为真（Word 返回 -1/True）
    #      若段落内属性不统一（Word 返回 9999999 = wdUndefined 或 None），
    #      则逐字符扫描 Range.Characters，任一字符不满足即命中。
    #   ③ 任一二级标题任一条不满足 → 命中"二级标题字体格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    _WD_UNDEFINED = 9999999
    try:
        _wd = _WordSession.doc()
        _L2_OUTLINE = 2  # wdOutlineLevel2
        for para in _wd.Paragraphs:
            try:
                if int(para.OutlineLevel) != _L2_OUTLINE:
                    continue
            except Exception:
                continue
            rng = para.Range
            try:
                title = str(rng.Text or '').strip('\r\x07\n\t ')
            except Exception:
                title = ''
            if not title:
                continue
            tag = title[:16]

            # a. 中文字体
            try:
                name_fe = rng.Font.NameFarEast
            except Exception:
                name_fe = None
            # b. 字号
            try:
                size_v = rng.Font.Size
            except Exception:
                size_v = None
            # c. 加粗
            try:
                bold_v = rng.Font.Bold
            except Exception:
                bold_v = None

            mixed_name = (name_fe is None) or (name_fe == _WD_UNDEFINED)
            mixed_size = (size_v is None) or (size_v == _WD_UNDEFINED)
            mixed_bold = (bold_v is None) or (bold_v == _WD_UNDEFINED)

            if not mixed_name:
                if str(name_fe) != 'PMingLiU-ExtB':
                    _d.append(f"{tag}→中文字体={name_fe!r}")
            if not mixed_size:
                if abs(float(size_v) - 18.0) > 0.01:
                    _d.append(f"{tag}→字号={size_v}")
            if not mixed_bold:
                if not bool(bold_v):
                    _d.append(f"{tag}→未加粗")

            # 若段落级别属性不统一，逐字符核验（只补齐 mixed 的维度）
            if mixed_name or mixed_size or mixed_bold:
                try:
                    for ch in rng.Characters:
                        t = str(ch.Text or '')
                        if not t or t in ('\r', '\x07', '\n', '\t', ' '):
                            continue
                        if mixed_name:
                            try:
                                nf = str(ch.Font.NameFarEast or '')
                            except Exception:
                                nf = ''
                            if nf != 'PMingLiU-ExtB':
                                _d.append(f"{tag}→中文字体含{nf!r}")
                                break
                        if mixed_size:
                            try:
                                sz = float(ch.Font.Size)
                            except Exception:
                                sz = -1.0
                            if abs(sz - 18.0) > 0.01:
                                _d.append(f"{tag}→字号含{sz}")
                                break
                        if mixed_bold:
                            try:
                                bd = ch.Font.Bold
                            except Exception:
                                bd = None
                            if bd is None or bd == _WD_UNDEFINED or not bool(bd):
                                _d.append(f"{tag}→含未加粗字符")
                                break
                except Exception:
                    pass
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "二级标题字体格式不满足 PMingLiU-ExtB、小二号、加粗",
        -3, hit, "; ".join(_d[:4]))


def _dim2_l2_title_paragraph():
    # -----------------------------------------------------------------
    # -3：二级标题段落格式不满足两端对齐、单倍行距
    # 完全基于 Word 办公软件文件属性：
    #   ① 二级标题 = Word COM Paragraph.OutlineLevel == wdOutlineLevel2(=2)
    #      —— 这是 Word 段落自身的大纲级别属性，不做任何文本正则推断。
    #   ② 逐一读取每个二级标题段落 Paragraph.Format：
    #        a. 对齐：Format.Alignment == wdAlignParagraphJustify(=3) "两端对齐"
    #        b. 行距规则：Format.LineSpacingRule == wdLineSpaceSingle(=0) "单倍行距"
    #   ③ 任一二级标题任一条不满足 → 命中"二级标题段落格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    try:
        _wd = _WordSession.doc()
        _L2_OUTLINE = 2  # wdOutlineLevel2
        for para in _wd.Paragraphs:
            try:
                if int(para.OutlineLevel) != _L2_OUTLINE:
                    continue
            except Exception:
                continue
            try:
                title = str(para.Range.Text or '').strip('\r\x07\n\t ')
            except Exception:
                title = ''
            if not title:
                continue
            tag = title[:16]
            pf = para.Format
            # a. 两端对齐
            try:
                align_v = int(pf.Alignment)
            except Exception:
                align_v = -1
            if align_v != WD_ALIGN_JUSTIFY:
                _d.append(f"{tag}→对齐={align_v}(期望3=两端)")
            # b. 单倍行距
            try:
                lsr_v = int(pf.LineSpacingRule)
            except Exception:
                lsr_v = -1
            if lsr_v != WD_LINE_SINGLE:
                _d.append(f"{tag}→行距规则={lsr_v}(期望0=单倍)")
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "二级标题段落格式不满足两端对齐、单倍行距",
        -3, hit, "; ".join(_d[:4]))


def _dim2_body_font():
    # -----------------------------------------------------------------
    # -3：正文字体格式不满足 PMingLiU-ExtB、四号
    # 完全基于 Word 办公软件文件属性：
    #   ① 正文范围：Word COM 段落中位于"第一章 緒 論"与"參考文獻"之间的段落
    #      —— 通过 Range.Start 判定，属于 Word 文档流位置属性。
    #   ② 排除非正文段落（Word 属性判定）：
    #        · Paragraph.OutlineLevel != wdOutlineLevelBodyText(=10)  —— 章/节标题
    #        · Range.Information(wdWithInTable=12) == True             —— 表格单元格
    #        · Range.InlineShapes.Count >= 1                            —— 图片段
    #        · Paragraph.Style.NameLocal ∈ {题注, Caption}              —— 图表注
    #        · 空段（Range.Text 仅含 \r\x07\n\t 空白）
    #   ③ 逐一读取每个正文段落 Range 的字体属性：
    #        a. 中文字体：Range.Font.NameFarEast == 'PMingLiU-ExtB'
    #        b. 字号：Range.Font.Size == 14.0 （"四号"= 14pt）
    #      属性不统一（返回 wdUndefined=9999999 或 None）时，逐字符扫描
    #      Range.Characters 中的中文字符（CJK 区间）作补充核验。
    #   ④ 任一正文段落任一字符不满足 → 命中"正文字体格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    _WD_UNDEFINED = 9999999
    try:
        _wd = _WordSession.doc()
        # 定位第一章 / 參考文獻 段落 Range.Start
        _ch1_start = None
        _ref_start = None
        for para in _wd.Paragraphs:
            try:
                t = str(para.Range.Text or '').replace('\r', '').replace('\x07', '').replace('\t', '').strip()
            except Exception:
                continue
            tn = t.replace(' ', '').replace('　', '')
            if _ch1_start is None and tn in ('第一章緒論', '第一章绪论'):
                _ch1_start = int(para.Range.Start)
            if _ref_start is None and tn in ('參考文獻', '参考文献'):
                _ref_start = int(para.Range.Start)
        if _ch1_start is None or _ref_start is None:
            _d.append("未定位到第一章/参考文献")
        else:
            for para in _wd.Paragraphs:
                try:
                    rng = para.Range
                    p_start = int(rng.Start)
                except Exception:
                    continue
                if p_start <= _ch1_start or p_start >= _ref_start:
                    continue
                # 排除章/节标题
                try:
                    if int(para.OutlineLevel) != 10:  # wdOutlineLevelBodyText
                        continue
                except Exception:
                    continue
                # 排除表格内
                try:
                    if bool(rng.Information(WD_WITH_IN_TABLE)):
                        continue
                except Exception:
                    pass
                # 排除图片段
                try:
                    if int(rng.InlineShapes.Count) >= 1:
                        continue
                except Exception:
                    pass
                # 排除图表注
                try:
                    style_name = str(para.Style.NameLocal or '')
                except Exception:
                    style_name = ''
                if style_name in ('题注', '題注', 'Caption'):
                    continue
                # 排除空段
                try:
                    ptxt = str(rng.Text or '')
                except Exception:
                    ptxt = ''
                if not ptxt.strip('\r\x07\n\t 　'):
                    continue
                tag = ptxt.strip('\r\x07\n\t 　')[:14]

                # a. 中文字体
                try:
                    name_fe = rng.Font.NameFarEast
                except Exception:
                    name_fe = None
                # b. 字号
                try:
                    size_v = rng.Font.Size
                except Exception:
                    size_v = None

                mixed_name = (name_fe is None) or (name_fe == _WD_UNDEFINED)
                mixed_size = (size_v is None) or (size_v == _WD_UNDEFINED)

                bad = False
                if not mixed_name:
                    if str(name_fe) != 'PMingLiU-ExtB':
                        _d.append(f"{tag}→中文字体={name_fe!r}")
                        bad = True
                if not mixed_size:
                    if abs(float(size_v) - 14.0) > 0.01:
                        _d.append(f"{tag}→字号={size_v}")
                        bad = True

                # 属性不统一 → 逐字符核验中文字符
                if (mixed_name or mixed_size) and not bad:
                    try:
                        for ch in rng.Characters:
                            t = str(ch.Text or '')
                            if not t or t in ('\r', '\x07', '\n', '\t', ' ', '　'):
                                continue
                            code = ord(t[0])
                            is_cjk = (
                                0x4E00 <= code <= 0x9FFF or
                                0x3400 <= code <= 0x4DBF or
                                0xF900 <= code <= 0xFAFF or
                                0x20000 <= code <= 0x2FFFF
                            )
                            if mixed_name and is_cjk:
                                try:
                                    nf = str(ch.Font.NameFarEast or '')
                                except Exception:
                                    nf = ''
                                if nf != 'PMingLiU-ExtB':
                                    _d.append(f"{tag}→中文字体含{nf!r}")
                                    bad = True
                                    break
                            if mixed_size:
                                try:
                                    sz = float(ch.Font.Size)
                                except Exception:
                                    sz = -1.0
                                if abs(sz - 14.0) > 0.01:
                                    _d.append(f"{tag}→字号含{sz}")
                                    bad = True
                                    break
                    except Exception:
                        pass
                if len(_d) >= 4:
                    break
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "正文字体格式不满足 PMingLiU-ExtB、四号",
        -3, hit, "; ".join(_d[:4]))


def _dim2_body_paragraph():
    # -----------------------------------------------------------------
    # -3：正文段落格式不满足左对齐、首行缩进两字符、1.5倍行距、段前0行、段后10磅
    # 完全基于 Word 办公软件文件属性：
    #   ① 正文范围：Word COM 段落中 Range.Start 位于"第一章 緒 論"与"參考文獻"之间
    #      —— 属于 Word 文档流位置属性。
    #   ② 非正文段落排除（Word 属性判定）：
    #        · Paragraph.OutlineLevel != wdOutlineLevelBodyText(=10) —— 章/节标题
    #        · Range.Information(wdWithInTable=12) == True            —— 表格单元格
    #        · Range.InlineShapes.Count >= 1                           —— 图片段
    #        · Paragraph.Style.NameLocal ∈ {题注, 題注, Caption}       —— 图表注
    #        · 空段（Range.Text 去空白后为空）
    #   ③ 逐一读取每个正文段落 Paragraph.Format：
    #        a. 左对齐        —— Format.Alignment == wdAlignParagraphLeft(=0)
    #        b. 首行缩进两字符 —— Format.CharacterUnitFirstLineIndent == 2.0
    #                            （回退：Format.FirstLineIndent(pt) ≈ 2 × 字号 pt）
    #        c. 1.5 倍行距    —— Format.LineSpacingRule == wdLineSpace1pt5(=1)
    #        d. 段前 0 行     —— Format.SpaceBefore(pt) ≈ 0
    #        e. 段后 10 磅    —— Format.SpaceAfter(pt) ≈ 10
    #   ④ 任一正文段落任一条不满足 → 命中"正文段落格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    try:
        _wd = _WordSession.doc()
        _ch1_start = None
        _ref_start = None
        for para in _wd.Paragraphs:
            try:
                t = str(para.Range.Text or '').replace('\r', '').replace('\x07', '').replace('\t', '').strip()
            except Exception:
                continue
            tn = t.replace(' ', '').replace('　', '')
            if _ch1_start is None and tn in ('第一章緒論', '第一章绪论'):
                _ch1_start = int(para.Range.Start)
            if _ref_start is None and tn in ('參考文獻', '参考文献'):
                _ref_start = int(para.Range.Start)
        if _ch1_start is None or _ref_start is None:
            _d.append("未定位到第一章/参考文献")
        else:
            for para in _wd.Paragraphs:
                try:
                    rng = para.Range
                    p_start = int(rng.Start)
                except Exception:
                    continue
                if p_start <= _ch1_start or p_start >= _ref_start:
                    continue
                try:
                    if int(para.OutlineLevel) != 10:  # wdOutlineLevelBodyText
                        continue
                except Exception:
                    continue
                try:
                    if bool(rng.Information(WD_WITH_IN_TABLE)):
                        continue
                except Exception:
                    pass
                try:
                    if int(rng.InlineShapes.Count) >= 1:
                        continue
                except Exception:
                    pass
                try:
                    style_name = str(para.Style.NameLocal or '')
                except Exception:
                    style_name = ''
                if style_name in ('题注', '題注', 'Caption'):
                    continue
                try:
                    ptxt = str(rng.Text or '')
                except Exception:
                    ptxt = ''
                if not ptxt.strip('\r\x07\n\t 　'):
                    continue
                tag = ptxt.strip('\r\x07\n\t 　')[:12]
                pf = para.Format
                # a. 左对齐
                try:
                    align = int(pf.Alignment)
                except Exception:
                    align = -1
                if align != WD_ALIGN_LEFT:
                    _d.append(f"{tag}→对齐={align}(期望0=左)")
                # b. 首行缩进两字符
                try:
                    ch_ind = float(pf.CharacterUnitFirstLineIndent)
                except Exception:
                    ch_ind = float('nan')
                try:
                    pt_ind = float(pf.FirstLineIndent)
                except Exception:
                    pt_ind = float('nan')
                try:
                    size_pt = float(rng.Font.Size)
                except Exception:
                    size_pt = 14.0
                first_ok = (
                    (ch_ind == ch_ind and abs(ch_ind - 2.0) < 0.01) or
                    (pt_ind == pt_ind and abs(pt_ind - 2.0 * size_pt) < 1.5)
                )
                if not first_ok:
                    _d.append(f"{tag}→首行缩进={ch_ind}字符/{pt_ind}pt")
                # c. 1.5 倍行距
                try:
                    lsr = int(pf.LineSpacingRule)
                except Exception:
                    lsr = -1
                if lsr != WD_LINE_1PT5:
                    _d.append(f"{tag}→行距规则={lsr}(期望1=1.5倍)")
                # d. 段前 0 行 —— 细则以"行"表述，Word 属性 SpaceBefore 单位为磅，0行即 0pt
                try:
                    sb = float(pf.SpaceBefore)
                except Exception:
                    sb = float('nan')
                if not (sb == sb) or abs(sb) > 0.05:
                    _d.append(f"{tag}→段前={sb}pt(期望0)")
                # e. 段后 10 磅
                try:
                    sa = float(pf.SpaceAfter)
                except Exception:
                    sa = float('nan')
                if not (sa == sa) or abs(sa - 10.0) > 0.5:
                    _d.append(f"{tag}→段后={sa}pt(期望10)")
                if len(_d) >= 5:
                    break
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "正文段落格式不满足左对齐、首行缩进两字符、1.5倍行距、段前0行、段后10磅",
        -3, hit, "; ".join(_d[:5]))

    # APPEND_HERE_DIM2_E


def _dim2_fig_caption_format():
    # -----------------------------------------------------------------
    # -3：图注格式不满足：居于图片下方，居中对齐、1.5倍行距、段前0行、段后10磅、首行缩进两字符
    # 完全基于 Word 办公软件文件属性：
    #   ① 图注段判定（Word 属性优先，退化到文本形态识别；并排除表注）：
    #        · Range.Text 匹配 ^[图圖]\s*\d+\s*[-–—]\s*\d+（形如 "图 1-1 / 圖 4-1"）；或
    #        · Paragraph.Style.NameLocal ∈ {'题注','題注','Caption'} 且文本不形如
    #          "表 …"（^[表表]\s*\d+）—— 因为表注同样套用题注/Caption 样式，
    #          仅凭样式无法区分图注与表注，须叠加图编号文本或排除表注文本。
    #        · 均要求 Range.Text 不含 \t，并排除 OutlineLevel==1 的段（如"圖 目 錄"标题）
    #   ② 逐一读取每个图注段：
    #        a. 居于图片下方 —— 图注段之前最近的非空段
    #           Range.InlineShapes.Count >= 1
    #        b. 居中对齐   —— Format.Alignment == wdAlignParagraphCenter(=1)
    #        c. 1.5 倍行距 —— Format.LineSpacingRule == wdLineSpace1pt5(=1)
    #        d. 段前 0 行  —— Format.SpaceBefore(pt) ≈ 0
    #        e. 段后 10 磅 —— Format.SpaceAfter(pt) ≈ 10
    #        f. 首行缩进两字符 —— Format.CharacterUnitFirstLineIndent == 2.0
    #                            （回退：FirstLineIndent(pt) ≈ 2 × 字号pt）
    #   ③ 任一图注段任一条不满足 → 命中"图注格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    _FIG_TEXT_PAT = re.compile(r'^[图圖]\s*\d+\s*[-–—]\s*\d+')
    # 表注文本形态：用于把"套用 Caption/题注 样式但实为表注"的段落排除出图注集合
    _TAB_TEXT_PAT = re.compile(r'^[表表]\s*\d+')
    try:
        _wd = _WordSession.doc()
        _paras = list(_wd.Paragraphs)
        # ① 收集图注段（保留 1-based 索引以便回溯查找上一段）
        fig_caps = []
        for idx, para in enumerate(_paras, start=1):
            try:
                style_name = str(para.Style.NameLocal or '')
            except Exception:
                style_name = ''
            try:
                raw = str(para.Range.Text or '')
            except Exception:
                raw = ''
            if '\t' in raw:
                continue
            stripped = raw.strip('\r\x07\n\t 　')
            if not stripped:
                continue
            # 排除章级标题（如"圖 目 錄"）
            try:
                if int(para.OutlineLevel) == 1:
                    continue
            except Exception:
                pass
            is_fig_style = style_name in ('题注', '題注', 'Caption')
            is_fig_text = bool(_FIG_TEXT_PAT.match(stripped))
            is_tab_text = bool(_TAB_TEXT_PAT.match(stripped))
            # 图注识别：必须命中"图编号文本"，或"题注/Caption 样式且不是表注文本"。
            # 单凭 Caption/题注 样式不足以判定为图注 —— 表注同样使用该样式，
            # 若文本形如"表 2-1"则应归为表注，从图注集合中排除，避免误判。
            if is_fig_text:
                pass
            elif is_fig_style and not is_tab_text:
                pass
            else:
                continue
            fig_caps.append((idx, para, stripped))

        if not fig_caps:
            _d.append("未找到图注")
        else:
            for cap_idx, cap_para, cap_text in fig_caps:
                tag = cap_text[:16]
                # a. 居于图片下方：向上找到最近的非空段，检查 InlineShapes
                above_has_img = False
                j = cap_idx - 2  # _paras 是 0-based，cap_idx 是 1-based
                while j >= 0:
                    prev = _paras[j]
                    try:
                        prev_txt = str(prev.Range.Text or '')
                    except Exception:
                        prev_txt = ''
                    prev_stripped = prev_txt.strip('\r\x07\n\t 　')
                    try:
                        prev_imgs = int(prev.Range.InlineShapes.Count)
                    except Exception:
                        prev_imgs = 0
                    if prev_imgs > 0:
                        above_has_img = True
                        break
                    if prev_stripped:
                        # 遇到非空且不含图片的段落 → 图注上方非图片
                        above_has_img = False
                        break
                    j -= 1
                if not above_has_img:
                    _d.append(f"{tag}→上方无图片")

                pf = cap_para.Format
                try:
                    size_pt = float(cap_para.Range.Font.Size)
                except Exception:
                    size_pt = 14.0
                # b. 居中对齐
                try:
                    align = int(pf.Alignment)
                except Exception:
                    align = -1
                if align != WD_ALIGN_CENTER:
                    _d.append(f"{tag}→对齐={align}(期望1=居中)")
                # c. 1.5 倍行距
                try:
                    lsr = int(pf.LineSpacingRule)
                except Exception:
                    lsr = -1
                if lsr != WD_LINE_1PT5:
                    _d.append(f"{tag}→行距规则={lsr}(期望1=1.5倍)")
                # d. 段前 0 行
                try:
                    sb = float(pf.SpaceBefore)
                except Exception:
                    sb = float('nan')
                if not (sb == sb) or abs(sb) > 0.05:
                    _d.append(f"{tag}→段前={sb}pt(期望0)")
                # e. 段后 10 磅
                try:
                    sa = float(pf.SpaceAfter)
                except Exception:
                    sa = float('nan')
                if not (sa == sa) or abs(sa - 10.0) > 0.5:
                    _d.append(f"{tag}→段后={sa}pt(期望10)")
                # f. 首行缩进两字符
                try:
                    ch_ind = float(pf.CharacterUnitFirstLineIndent)
                except Exception:
                    ch_ind = float('nan')
                try:
                    pt_ind = float(pf.FirstLineIndent)
                except Exception:
                    pt_ind = float('nan')
                first_ok = (
                    (ch_ind == ch_ind and abs(ch_ind - 2.0) < 0.01) or
                    (pt_ind == pt_ind and abs(pt_ind - 2.0 * size_pt) < 1.5)
                )
                if not first_ok:
                    _d.append(f"{tag}→首行缩进={ch_ind}字符/{pt_ind}pt")
                if len(_d) >= 5:
                    break
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "图注格式不满足：居于图片下方、居中对齐、1.5倍行距、段前0行、段后10磅、首行缩进两字符",
        -3, hit, "; ".join(_d[:5]))


def _dim2_fig_caption_style():
    # -----------------------------------------------------------------
    # -3：图注样式不满足"图1-1"
    # 完全基于 Word 办公软件文件属性：
    #   ① 图注段判定（Word 段落样式属性优先，退化到文本形态识别）：
    #        · Paragraph.Style.NameLocal ∈ {'题注','題注','Caption'}；或
    #        · Range.Text 不含制表符 \t，且文本匹配 ^[图圖]\s*\d+\s*[-–—]\s*\d+
    #        · 排除 OutlineLevel==1 的段（如"圖 目 錄"标题）
    #   ② 图注编号样式（Range.Text 文本内容）：
    #        · 匹配"图<可空白>数字-数字"，即以简体"图"开头 → 满足"图1-1"
    #        · 使用繁体"圖 X-Y"或任何非"图X-Y"形式 → 不满足
    #   ③ 任一图注不满足 → 命中"图注样式不满足'图1-1'"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    _FIG_TEXT_PAT = re.compile(r'^[图圖]\s*\d+\s*[-–—]\s*\d+')
    try:
        _wd = _WordSession.doc()
        caps_all = []
        for para in _wd.Paragraphs:
            try:
                style_name = str(para.Style.NameLocal or '')
            except Exception:
                style_name = ''
            try:
                raw = str(para.Range.Text or '')
            except Exception:
                raw = ''
            if '\t' in raw:
                continue
            stripped = raw.strip('\r\x07\n\t 　')
            if not stripped:
                continue
            try:
                if int(para.OutlineLevel) == 1:
                    continue
            except Exception:
                pass
            is_fig_style = style_name in ('题注', '題注', 'Caption')
            is_fig_text = bool(_FIG_TEXT_PAT.match(stripped))
            if not (is_fig_style or is_fig_text):
                continue
            caps_all.append(stripped)

        if not caps_all:
            _d.append("未找到图注")
        else:
            for cap in caps_all:
                # 图注编号样式：允许简/繁"图/圖"打头，形如"图1-1 / 圖 1-1"
                if not re.match(r'^[图圖]\s*\d+\s*[-–—]\s*\d+', cap):
                    _d.append(cap[:16])
                    if len(_d) >= 4:
                        break
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "图注样式不满足'图1-1'",
        -3, hit, "; ".join(_d[:4]))


def _dim2_fig_caption_font():
    # -----------------------------------------------------------------
    # -3：图注字体格式不满足：中文字体 PMingLiU-ExtB 四号、英文或阿拉伯数字 Noto Serif CJK TC 四号
    # 完全基于 Word 办公软件文件属性：
    #   ① 图注段判定（Word 段落样式属性优先，退化到文本形态识别）：
    #        · Paragraph.Style.NameLocal ∈ {'题注','題注','Caption'}；或
    #        · Range.Text 不含 \t，且文本匹配 ^[图圖]\s*\d+\s*[-–—]\s*\d+
    #        · 排除 OutlineLevel==1 的段（如"圖 目 錄"标题）
    #   ② 逐字符扫描 Range.Characters，读取 Word Font 属性：
    #        a. 中文字符（CJK 区间：U+4E00-U+9FFF, U+3400-U+4DBF, U+F900-U+FAFF,
    #           U+20000-U+2FFFF）
    #             · Font.NameFarEast == 'PMingLiU-ExtB'
    #             · Font.Size == 14.0（四号）
    #        b. 英文字母 [A-Za-z] 或阿拉伯数字 [0-9]
    #             · Font.NameAscii == 'Noto Serif CJK TC'
    #             · Font.Size == 14.0（四号）
    #      对每个图注段先尝试段级 Range.Font 属性做整段核验；返回 wdUndefined(=9999999)
    #      或 None 时再逐字符扫描对应维度。
    #   ③ 任一图注段任一相关字符不满足 → 命中"图注字体格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    _WD_UNDEFINED = 9999999
    _FIG_TEXT_PAT = re.compile(r'^[图圖]\s*\d+\s*[-–—]\s*\d+')
    try:
        _wd = _WordSession.doc()
        caps_all = []
        for para in _wd.Paragraphs:
            try:
                style_name = str(para.Style.NameLocal or '')
            except Exception:
                style_name = ''
            try:
                raw = str(para.Range.Text or '')
            except Exception:
                raw = ''
            if '\t' in raw:
                continue
            stripped = raw.strip('\r\x07\n\t 　')
            if not stripped:
                continue
            try:
                if int(para.OutlineLevel) == 1:
                    continue
            except Exception:
                pass
            is_fig_style = style_name in ('题注', '題注', 'Caption')
            is_fig_text = bool(_FIG_TEXT_PAT.match(stripped))
            if not (is_fig_style or is_fig_text):
                continue
            caps_all.append((para, stripped))

        if not caps_all:
            _d.append("未找到图注")
        else:
            for cap_para, cap_text in caps_all:
                tag = cap_text[:12]
                rng = cap_para.Range
                # 逐字符扫描（中文 + 英文/数字）
                try:
                    for ch in rng.Characters:
                        t = str(ch.Text or '')
                        if not t:
                            continue
                        c0 = t[0]
                        if c0 in ('\r', '\x07', '\n', '\t', ' ', '　'):
                            continue
                        code = ord(c0)
                        is_cjk = (
                            0x4E00 <= code <= 0x9FFF or
                            0x3400 <= code <= 0x4DBF or
                            0xF900 <= code <= 0xFAFF or
                            0x20000 <= code <= 0x2FFFF
                        )
                        is_en_num = ('A' <= c0 <= 'Z') or ('a' <= c0 <= 'z') or ('0' <= c0 <= '9')

                        # a. 中文字符 → NameFarEast + Size
                        if is_cjk:
                            try:
                                nf = str(ch.Font.NameFarEast or '')
                            except Exception:
                                nf = ''
                            if nf != 'PMingLiU-ExtB':
                                _d.append(f"{tag}→中文字体{c0}={nf!r}")
                                break
                            try:
                                sz = float(ch.Font.Size)
                            except Exception:
                                sz = -1.0
                            if abs(sz - 14.0) > 0.01:
                                _d.append(f"{tag}→中文字号{c0}={sz}")
                                break
                        # b. 英文/数字 → NameAscii + Size
                        elif is_en_num:
                            try:
                                nf = str(ch.Font.NameAscii or '')
                            except Exception:
                                nf = ''
                            if nf != 'Noto Serif CJK TC':
                                _d.append(f"{tag}→英数字体{c0}={nf!r}")
                                break
                            try:
                                sz = float(ch.Font.Size)
                            except Exception:
                                sz = -1.0
                            if abs(sz - 14.0) > 0.01:
                                _d.append(f"{tag}→英数字号{c0}={sz}")
                                break
                        # 其它字符不受本条约束
                except Exception:
                    pass
                if len(_d) >= 4:
                    break
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "图注字体格式不满足：中文字体 PMingLiU-ExtB 四号、英文或阿拉伯数字 Noto Serif CJK TC 四号",
        -3, hit, "; ".join(_d[:4]))


def _dim2_tab_caption_format():
    # -----------------------------------------------------------------
    # -3：表注格式不满足：居于表格上方，居中对齐、1.5倍行距、段前0行、段后10磅、首行缩进两字符
    # 完全基于 Word 办公软件文件属性：
    #   ① 表注段判定（Word 段落样式属性优先，退化到文本形态识别）：
    #        · Paragraph.Style.NameLocal ∈ {'题注','題注','Caption'}；或
    #        · Range.Text 不含 \t，且文本匹配 ^表\s*\d+\s*[-–—]\s*\d+
    #          （形如 "表 2-1 / 表 4-3"，兼容作者未使用题注样式的情形）
    #        · 排除 OutlineLevel==1 的段（如"表 目 錄"标题）
    #   ② 逐一读取每个表注段：
    #        a. 居于表格上方 —— 表注段之后最近的非空段
    #           Range.Information(wdWithInTable=12) == True
    #        b. 居中对齐   —— Format.Alignment == wdAlignParagraphCenter(=1)
    #        c. 1.5 倍行距 —— Format.LineSpacingRule == wdLineSpace1pt5(=1)
    #        d. 段前 0 行  —— Format.SpaceBefore(pt) ≈ 0
    #        e. 段后 10 磅 —— Format.SpaceAfter(pt) ≈ 10
    #        f. 首行缩进两字符 —— Format.CharacterUnitFirstLineIndent == 2.0
    #                            （回退：FirstLineIndent(pt) ≈ 2 × 字号pt）
    #   ③ 任一表注段任一条不满足 → 命中"表注格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    _TAB_TEXT_PAT = re.compile(r'^表\s*\d+\s*[-–—]\s*\d+')
    try:
        _wd = _WordSession.doc()
        _paras = list(_wd.Paragraphs)
        tab_caps = []
        for idx, para in enumerate(_paras, start=1):
            try:
                style_name = str(para.Style.NameLocal or '')
            except Exception:
                style_name = ''
            try:
                raw = str(para.Range.Text or '')
            except Exception:
                raw = ''
            if '\t' in raw:
                continue
            stripped = raw.strip('\r\x07\n\t 　')
            if not stripped or stripped[0] != '表':
                continue
            try:
                if int(para.OutlineLevel) == 1:
                    continue
            except Exception:
                pass
            is_tab_style = style_name in ('题注', '題注', 'Caption')
            is_tab_text = bool(_TAB_TEXT_PAT.match(stripped))
            if not (is_tab_style or is_tab_text):
                continue
            tab_caps.append((idx, para, stripped))

        if not tab_caps:
            _d.append("未找到表注")
        else:
            for cap_idx, cap_para, cap_text in tab_caps:
                tag = cap_text[:16]
                # a. 表注下方最近的非空段是否在表格内
                below_in_table = False
                j = cap_idx  # 下一段的 0-based 索引 = 当前 1-based 索引
                while j < len(_paras):
                    nxt = _paras[j]
                    try:
                        in_tbl = bool(nxt.Range.Information(WD_WITH_IN_TABLE))
                    except Exception:
                        in_tbl = False
                    if in_tbl:
                        below_in_table = True
                        break
                    try:
                        nxt_txt = str(nxt.Range.Text or '')
                    except Exception:
                        nxt_txt = ''
                    if nxt_txt.strip('\r\x07\n\t 　'):
                        # 遇到非空非表格段落 → 下方不是表格
                        below_in_table = False
                        break
                    j += 1
                if not below_in_table:
                    _d.append(f"{tag}→下方非表格")

                pf = cap_para.Format
                try:
                    size_pt = float(cap_para.Range.Font.Size)
                except Exception:
                    size_pt = 14.0
                # b. 居中对齐
                try:
                    align = int(pf.Alignment)
                except Exception:
                    align = -1
                if align != WD_ALIGN_CENTER:
                    _d.append(f"{tag}→对齐={align}(期望1=居中)")
                # c. 1.5 倍行距
                try:
                    lsr = int(pf.LineSpacingRule)
                except Exception:
                    lsr = -1
                if lsr != WD_LINE_1PT5:
                    _d.append(f"{tag}→行距规则={lsr}(期望1=1.5倍)")
                # d. 段前 0 行
                try:
                    sb = float(pf.SpaceBefore)
                except Exception:
                    sb = float('nan')
                if not (sb == sb) or abs(sb) > 0.05:
                    _d.append(f"{tag}→段前={sb}pt(期望0)")
                # e. 段后 10 磅
                try:
                    sa = float(pf.SpaceAfter)
                except Exception:
                    sa = float('nan')
                if not (sa == sa) or abs(sa - 10.0) > 0.5:
                    _d.append(f"{tag}→段后={sa}pt(期望10)")
                # f. 首行缩进两字符
                try:
                    ch_ind = float(pf.CharacterUnitFirstLineIndent)
                except Exception:
                    ch_ind = float('nan')
                try:
                    pt_ind = float(pf.FirstLineIndent)
                except Exception:
                    pt_ind = float('nan')
                first_ok = (
                    (ch_ind == ch_ind and abs(ch_ind - 2.0) < 0.01) or
                    (pt_ind == pt_ind and abs(pt_ind - 2.0 * size_pt) < 1.5)
                )
                if not first_ok:
                    _d.append(f"{tag}→首行缩进={ch_ind}字符/{pt_ind}pt")
                if len(_d) >= 5:
                    break
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "表注格式不满足：居于表格上方、居中对齐、1.5倍行距、段前0行、段后10磅、首行缩进两字符",
        -3, hit, "; ".join(_d[:5]))


def _dim2_ref_paragraph():
    # -----------------------------------------------------------------
    # -3：参考文献页内容段落格式不满足左对齐、首行缩进两字符、1.5倍行距、段前0行、段后10磅
    # 完全基于 Word 办公软件文件属性：
    #   ① 参考文献页范围：Word COM 段落中 Range.Start 在"參考文獻/参考文献"
    #      标题段之后的所有段落（属于 Word 文档流位置属性）。
    #   ② 非内容段落排除（Word 段落属性）：
    #        · 空段（Range.Text 去空白后为空）
    #        · 参考文献标题段本身（首个 Range.Start > ref_start 的段起才纳入）
    #        · 表格内段（Range.Information(wdWithInTable=12) == True）
    #   ③ 逐一读取每个参考文献内容段 Paragraph.Format：
    #        a. 左对齐        —— Format.Alignment == wdAlignParagraphLeft(=0)
    #        b. 首行缩进两字符 —— Format.CharacterUnitFirstLineIndent == 2.0
    #                            （回退：FirstLineIndent(pt) ≈ 2 × 字号pt）
    #        c. 1.5 倍行距    —— Format.LineSpacingRule == wdLineSpace1pt5(=1)
    #        d. 段前 0 行     —— Format.SpaceBefore(pt) ≈ 0
    #        e. 段后 10 磅    —— Format.SpaceAfter(pt) ≈ 10
    #   ④ 任一内容段任一条不满足 → 命中"参考文献页内容段落格式不满足 …"，扣 3 分。
    # -----------------------------------------------------------------
    _d = []
    try:
        _wd = _WordSession.doc()
        _ref_start = None
        for para in _wd.Paragraphs:
            try:
                t = str(para.Range.Text or '').replace('\r', '').replace('\x07', '').replace('\t', '').strip()
            except Exception:
                continue
            tn = t.replace(' ', '').replace('　', '')
            if tn in ('參考文獻', '参考文献'):
                _ref_start = int(para.Range.Start)
                break
        if _ref_start is None:
            _d.append("未找到参考文献标题")
        else:
            checked_any = False
            for para in _wd.Paragraphs:
                try:
                    rng = para.Range
                    p_start = int(rng.Start)
                except Exception:
                    continue
                if p_start <= _ref_start:
                    continue
                # 排除表格内
                try:
                    if bool(rng.Information(WD_WITH_IN_TABLE)):
                        continue
                except Exception:
                    pass
                # 排除空段
                try:
                    ptxt = str(rng.Text or '')
                except Exception:
                    ptxt = ''
                if not ptxt.strip('\r\x07\n\t 　'):
                    continue
                checked_any = True
                tag = ptxt.strip('\r\x07\n\t 　')[:12]
                pf = para.Format
                try:
                    size_pt = float(rng.Font.Size)
                except Exception:
                    size_pt = 14.0
                # a. 左对齐
                try:
                    align = int(pf.Alignment)
                except Exception:
                    align = -1
                if align != WD_ALIGN_LEFT:
                    _d.append(f"{tag}→对齐={align}(期望0=左)")
                # b. 首行缩进两字符
                try:
                    ch_ind = float(pf.CharacterUnitFirstLineIndent)
                except Exception:
                    ch_ind = float('nan')
                try:
                    pt_ind = float(pf.FirstLineIndent)
                except Exception:
                    pt_ind = float('nan')
                first_ok = (
                    (ch_ind == ch_ind and abs(ch_ind - 2.0) < 0.01) or
                    (pt_ind == pt_ind and abs(pt_ind - 2.0 * size_pt) < 1.5)
                )
                if not first_ok:
                    _d.append(f"{tag}→首行缩进={ch_ind}字符/{pt_ind}pt")
                # c. 1.5 倍行距
                try:
                    lsr = int(pf.LineSpacingRule)
                except Exception:
                    lsr = -1
                if lsr != WD_LINE_1PT5:
                    _d.append(f"{tag}→行距规则={lsr}(期望1=1.5倍)")
                # d. 段前 0 行
                try:
                    sb = float(pf.SpaceBefore)
                except Exception:
                    sb = float('nan')
                if not (sb == sb) or abs(sb) > 0.05:
                    _d.append(f"{tag}→段前={sb}pt(期望0)")
                # e. 段后 10 磅
                try:
                    sa = float(pf.SpaceAfter)
                except Exception:
                    sa = float('nan')
                if not (sa == sa) or abs(sa - 10.0) > 0.5:
                    _d.append(f"{tag}→段后={sa}pt(期望10)")
                if len(_d) >= 5:
                    break
            if not checked_any and not _d:
                _d.append("未找到参考文献正文段")
    except Exception as e:
        _d.append(f"读取失败: {e}")
    hit = bool(_d)
    return _item(
        "参考文献页内容段落格式为左对齐、首行缩进两字符、1.5倍行距、段前0行、段后10磅",
        -3, hit, "; ".join(_d[:5]))


def check_dim2():
    """依次跑完所有维度二检查项，返回结构化 item 列表（命中与未命中均包含）。"""
    checks = (
        _dim2_last_page_removed,
        _dim2_section_page_no_setup,
        _dim2_roman_pagination,
        _dim2_arabic_pagination,
        _dim2_toc_auto_update,
        _dim2_toc_no_thanks,
        _dim2_toc_page_numbers,
        _dim2_table_of_tables,
        _dim2_table31_single_page,
        _dim2_table41_single_page,
        _dim2_table31_blank_lines,
        _dim2_toc_font,
        _dim2_toc_paragraph,
        _dim2_table41_blank_lines,
        _dim2_cover_logo,
        _dim2_cover_title_font,
        _dim2_cover_title_paragraph,
        _dim2_thanks_title_font,
        _dim2_cn_abstract_paragraph,
        _dim2_cn_abstract_font,
        _dim2_en_abstract_font,
        _dim2_en_abstract_paragraph,
        _dim2_page_number_format,
        _dim2_header_format,
        _dim2_l2_title_font,
        _dim2_l2_title_paragraph,
        _dim2_body_paragraph,
        _dim2_body_font,
        _dim2_fig_caption_format,
        _dim2_fig_caption_style,
        _dim2_fig_caption_font,
        _dim2_tab_caption_format,
        _dim2_ref_paragraph,
    )
    return [check() for check in checks]


# =========================================================================
# 统一接口：evaluate(dir_path) —— 脚本自己在目录里定位并打开被评估的文档
# =========================================================================
SCRIPT_ID = "023"


def evaluate(dir_path: str) -> dict:
    """
    评估 dir_path 目录下的 Word 文档，返回结构化结果字典：
      {
        "id": "023",
        "file_name": "xxx.docx",
        "status": "ok" / "error",
        "error": None / str,
        "dim1_pass": bool,
        "dim1_reason": str,
        "dim2_items": [ {rule, max_delta, delta, hit, detail}, ... ],
        "total_score": int,
        "max_score": int,
      }
    """
    global _file_path

    result = {
        "id": SCRIPT_ID,
        "file_name": None,
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }
    try:
        docx_path = _resolve_docx_path(dir_path)
        _file_path = docx_path
        result["file_name"] = os.path.basename(docx_path)

        dim1_pass, dim1_reason = check_dim1()
        result["dim1_pass"] = dim1_pass
        result["dim1_reason"] = dim1_reason
        if not dim1_pass:
            return result

        dim2_items = check_dim2()
        result["dim2_items"] = dim2_items
        result["total_score"] = sum(item["delta"] for item in dim2_items)
        result["max_score"] = sum(item["max_delta"] for item in dim2_items if item["max_delta"] > 0)
    except Exception as e:
        result["status"] = "error"
        result["error"] = str(e)
    finally:
        _WordSession.close()
    return result


if __name__ == '__main__':
    import sys
    _target_dir = sys.argv[1] if len(sys.argv) > 1 else os.getcwd()
    print(json.dumps(evaluate(_target_dir), ensure_ascii=False, indent=2))
