# -*- coding: utf-8 -*-
"""
公共书房运营规划书自动评估脚本
根据打分细则对Word文档进行自动评估
"""

import os
import sys
import json
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


SCRIPT_ID = "018"

# 维度二评分项定义（key, rule 文本, max_delta, matched_points 中的完整文本）
# 每项 max_delta 为该项对总分的贡献值：加分项为正，扣分项为负；命中(hit=True)时
# delta=max_delta，否则 delta=0。max_score 仅累加正的 max_delta，表示"满分"。
DIM2_RULES = [
    ('title_line_spacing',
     '第1页报告题目"临海公共书房网络化运营与更新规划书"为单倍行距', 1,
     '+1:第1页报告题目"临海公共书房网络化运营与更新规划书"为单倍行距'),
    ('abstract_title_format',
     '文档中摘要页标题字体格式为黑体、四号，且在题目页之后', 1,
     '+1:文档中摘要页标题字体格式为黑体、四号，且在题目页之后'),
    ('level1_title_style',
     '一级标题序号样式为"一、"，例如"一、项目定位与服务原则"', 3,
     '+3:一级标题序号样式为"一、"，例如"一、项目定位与服务原则"'),
    ('level2_title_style',
     '二级标题序号样式为"（一）"，例如"（一）公告阅读空间的建设背景"', 3,
     '+3:二级标题序号样式为"（一）"，例如"（一）公告阅读空间的建设背景"'),
    ('level1_title_font',
     '一级标题字体格式为黑体、小三号', 3,
     '+3:一级标题字体格式为黑体、小三号'),
    ('level2_title_font',
     '二级标题字体格式为黑体、四号', 3,
     '+3:二级标题字体格式为黑体、四号'),
    ('level3_title_style',
     '三级标题序号样式为"1."，例如"1.城市文化消费与邻里阅读需求"', 3,
     '+3:三级标题序号样式为"1."，例如"1.城市文化消费与邻里阅读需求"'),
    ('level3_title_font',
     '三级标题字体格式为宋体、小四号', 3,
     '+3:三级标题字体格式为宋体、小四号'),
    ('level4_title_style',
     '四级标题序号样式为"（1）"，例如"（1）公共书房网络的服务对象与覆盖逻辑"', 3,
     '+3:四级标题序号样式为"（1）"，例如"（1）公共书房网络的服务对象与覆盖逻辑"'),
    ('level4_title_font',
     '四级标题字体格式为宋体、小四号', 3,
     '+3:四级标题字体格式为宋体、小四号'),
    ('body_font_format',
     '除标题、表格、图注、表注之外的所有字体格式统一设置为宋体、小四号', 5,
     '+5:除标题、表格、图注、表注之外的所有字体格式统一设置为宋体、小四号'),
    ('numbering_format',
     '图注、表注、附录、参考文献、公式编号均使用阿拉伯数字连续编号，如"图 1""表 1""公式（1）"', 3,
     '+3:图注、表注、附录、参考文献、公式编号均使用阿拉伯数字连续编号，如"图 1""表 1""公式（1）"'),
    ('figure_captions',
     '全文所有图片下方均出现对应图注且图注居中显示，图注按出现顺序连续编号，不出现跳号或顺序混乱', 5,
     '+5:全文所有图片下方均出现对应图注且图注居中显示，图注按出现顺序连续编号，不出现跳号或顺序混乱'),
    ('figure_caption_format',
     '图注格式为宋体、小四号，所有图注的图序与图题之间空两格，例如"图 1  公共书房功能分区关系图"', 3,
     '+3:图注格式为宋体、小四号，所有图注的图序与图题之间空两格，例如"图 1  公共书房功能分区关系图"'),
    ('table_captions',
     '所有表格上方均出现表注并居中显示，表注按出现顺序连续编号，不出现跳号或顺序混乱', 5,
     '+5:所有表格上方均出现表注并居中显示，表注按出现顺序连续编号，不出现跳号或顺序混乱'),
    ('table_caption_format',
     '表注格式为宋体、小四号；所有表注的表序与表题之间空两格，例如"表 1  公共书房运营指标表"', 3,
     '+3:表注格式为宋体、小四号；所有表注的表序与表题之间空两格，例如"表 1  公共书房运营指标表"'),
    ('table_font',
     '所有表格内文字统一设置为宋体、小四号、单倍行距', 5,
     '+5:所有表格内文字统一设置为宋体、小四号、单倍行距'),
    ('paragraph_format_deduction',
     '摘要页后的所有页，除了各级标题、表注、图注和表格中的文字之外，其余的所有文本内容段落格式不满足：1.25倍行距，段前间距为0.5行，首行缩进两字符', -5,
     '-5:摘要页后的所有页，除了各级标题、表注、图注和表格中的文字之外，其余的所有文本内容段落格式不满足：1.25倍行距，段前间距为0.5行，首行缩进两字符'),
    ('title_font_deduction',
     '第1页报告题目"临海公共书房网络化运营与更新规划书"字体格式不满足宋体、三号、加粗、居中', -1,
     '-1:第1页报告题目"临海公共书房网络化运营与更新规划书"字体格式不满足宋体、三号、加粗、居中'),
    ('large_blank_area',
     '除前两页和最后一页外文档中任意一页出现超过百分之四十的大面积空白', -5,
     '-5:除前两页和最后一页外文档中任意一页出现超过百分之四十的大面积空白'),
    ('table_page_break',
     '文档中表格出现断页、不连续', -3,
     '-3:文档中表格出现断页、不连续'),
    ('page_margin',
     '文档页面边距不满足上、下2.54厘米，左、右3.18厘米', -3,
     '-3:文档页面边距不满足上、下2.54厘米，左、右3.18厘米'),
]


class DocumentEvaluator:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = None
        self.dimension1_passed = True
        self.dimension1_issues = []
        # 命中的维度二评分项 key 集合（对应 DIM2_RULES 中的 key）
        self.matched_keys = set()
        self.font_size_map = {
            '初号': 42, '小初': 36, '一号': 26, '小一': 24,
            '二号': 22, '小二': 18, '三号': 16, '小三': 15,
            '四号': 14, '小四': 12, '五号': 10.5, '小五': 9,
            '六号': 7.5, '小六': 6.5, '七号': 5.5, '八号': 5
        }

    def is_font_size(self, run, target_size_name):
        if run.font.size is None:
            return False
        target_pt = self.font_size_map.get(target_size_name, 0)
        return abs(run.font.size.pt - target_pt) < 0.5

    def is_font_name(self, run, target_font):
        rpr = run._element.rPr
        if rpr is not None:
            east_asia = rpr.find(qn('w:eastAsia'))
            if east_asia is not None:
                cn_font = east_asia.get(qn('w:val'))
                if cn_font:
                    return target_font in cn_font
        font_name = run.font.name
        if font_name:
            return target_font in font_name
        return False

    def _para_has_content(self, para):
        if para._element.xpath('.//w:drawing') or para._element.xpath('.//w:pict'):
            return True
        return False

    def evaluate_dimension1(self):
        ext = os.path.splitext(self.file_path)[1].lower()
        # 仅识别 .docx（python-docx 只能解析 OOXML 的 .docx；旧版二进制 .doc
        # 需要 COM/LibreOffice 才能读取，按要求不支持）
        if ext != '.docx':
            self.dimension1_passed = False
            self.dimension1_issues.append(f"文件格式不符合要求：{ext}")
            return False
        try:
            self.document = Document(self.file_path)
        except Exception as e:
            self.dimension1_passed = False
            self.dimension1_issues.append(f"文件无法正常打开：{str(e)}")
            return False
        return True

    def evaluate_dimension2(self):
        self._eval_title_line_spacing()
        self._eval_abstract_title_format()
        self._eval_level1_title_style()
        self._eval_level2_title_style()
        self._eval_level1_title_font()
        self._eval_level2_title_font()
        self._eval_level3_title_style()
        self._eval_level3_title_font()
        self._eval_level4_title_style()
        self._eval_level4_title_font()
        self._eval_body_font_format()
        self._eval_numbering_format()
        self._eval_figure_captions()
        self._eval_figure_caption_format()
        self._eval_table_captions()
        self._eval_table_caption_format()
        self._eval_table_font()
        self._eval_paragraph_format_deduction()
        self._eval_title_font_deduction()
        self._eval_large_blank_area()
        self._eval_table_page_break()
        self._eval_page_margin()

    def _mark(self, key):
        """标记维度二某评分项为命中；总分与文本描述由 evaluate() 组装。"""
        self.matched_keys.add(key)

    def _eval_title_line_spacing(self):
        # 细则："第1页报告题目"临海公共书房网络化运营与更新规划书"为单倍行距"
        # 判据（针对办公软件的真实渲染）：
        #   1) 细则限定题目必须位于**第1页**——先按分页信号（显式分页符 /
        #      lastRenderedPageBreak / w:pageBreakBefore / w:sectPr / 累积
        #      内容高度超过页面可用高度）估算题目段落所在页码，非第1页则
        #      直接判定不满足，不参与后续行距判断
        #   2) 行距须为**显式**单倍行距：w:lineRule="auto" 且 w:line=240，
        #      或段落/样式继承链上出现该设置；若段落与其继承的样式链均未
        #      显式设置行距，则视为"无法确认"，不得默认按单倍行距判定通过
        title_text = "临海公共书房网络化运营与更新规划书"
        paragraphs = self.document.paragraphs

        try:
            section = self.document.sections[0]
            page_h = getattr(section, 'page_height', None)
            top_margin = getattr(section, 'top_margin', None)
            bottom_margin = getattr(section, 'bottom_margin', None)
            if page_h is None or top_margin is None or bottom_margin is None:
                raise ValueError('section metrics unavailable')
            page_h_cm = page_h.cm - top_margin.cm - bottom_margin.cm
        except Exception:
            page_h_cm = 24.62

        def _para_mark_pt(p_elem):
            ppr = p_elem.find(qn('w:pPr'))
            mark_pt = 10.5
            if ppr is not None:
                rpr = ppr.find(qn('w:rPr'))
                if rpr is not None:
                    sz_el = rpr.find(qn('w:sz'))
                    if sz_el is not None:
                        try:
                            mark_pt = int(sz_el.get(qn('w:val'))) / 2.0
                        except (TypeError, ValueError):
                            pass
            run_max_pt = 0.0
            for r in p_elem.findall(qn('w:r')):
                rp = r.find(qn('w:rPr'))
                if rp is not None:
                    s = rp.find(qn('w:sz'))
                    if s is not None:
                        try:
                            run_max_pt = max(run_max_pt, int(s.get(qn('w:val'))) / 2.0)
                        except (TypeError, ValueError):
                            pass
            return max(mark_pt, run_max_pt)

        def _para_height_cm(p_elem):
            base_pt = _para_mark_pt(p_elem)
            line_mult = 1.0
            ppr = p_elem.find(qn('w:pPr'))
            if ppr is not None:
                sp = ppr.find(qn('w:spacing'))
                if sp is not None:
                    line_val = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
                    if line_val:
                        try:
                            lv = int(line_val)
                            if line_rule in (None, 'auto'):
                                line_mult = lv / 240.0
                            elif line_rule in ('exact', 'atLeast'):
                                base_pt = max(base_pt, lv / 20.0)
                                line_mult = 1.0
                        except ValueError:
                            pass
            text_val = ''.join(n.text or '' for n in p_elem.iter() if n.text)
            num_lines = max(1, len(text_val) / 28.0) if text_val.strip() else 1
            return (base_pt * 1.3 * line_mult * num_lines) / 28.3465

        # 按 body 顺序模拟分页，定位题目段落所在的物理页码
        title_page = None
        cur_page = 1
        cur_h = 0.0
        title_para = None
        for p in paragraphs:
            if title_text in p.text:
                title_para = p
            p_elem = p._element
            ppr = p_elem.find(qn('w:pPr'))
            paged = False
            if ppr is not None:
                pb = ppr.find(qn('w:pageBreakBefore'))
                if pb is not None and pb.get(qn('w:val')) not in ('0', 'false'):
                    paged = True
            if not paged:
                brs = p_elem.findall('.//' + qn('w:br'))
                if any(b.get(qn('w:type')) == 'page' for b in brs):
                    paged = True
            if not paged and p_elem.findall('.//' + qn('w:lastRenderedPageBreak')):
                paged = True
            if paged:
                cur_page += 1
                cur_h = 0.0

            h = _para_height_cm(p_elem)
            if cur_h + h > page_h_cm and cur_h > 0:
                cur_page += 1
                cur_h = h
            else:
                cur_h += h

            if ppr is not None and ppr.find(qn('w:sectPr')) is not None:
                cur_page += 1
                cur_h = 0.0

            if title_text in p.text and title_page is None:
                title_page = cur_page if not paged else cur_page

        if title_para is None or title_page != 1:
            # 未找到题目，或题目不在第1页——按细则要求不得判定为满足
            return

        pf = title_para.paragraph_format
        ls = pf.line_spacing
        ls_rule = pf.line_spacing_rule
        is_single = False
        confirmed = False
        # 办公软件中"单倍行距"在OOXML中表现为 w:lineRule="auto" 且 w:line=240,
        # python-docx 读取时会将 line_spacing_rule 解析为 WD_LINE_SPACING.SINGLE,
        # 同时 line_spacing 为 1.0
        if ls_rule == WD_LINE_SPACING.SINGLE:
            is_single = True
            confirmed = True
        elif ls is not None and hasattr(ls, '__float__'):
            confirmed = True
            try:
                if abs(float(ls) - 1.0) < 0.05:
                    is_single = True
            except (TypeError, ValueError):
                pass
        elif ls is None and ls_rule is None:
            # 段落未显式设置——沿样式继承链查找，样式链上若仍未显式设置
            # 行距，则视为无法确认，不得默认判定为单倍行距
            try:
                style = title_para.style
                while style is not None:
                    spf = style.paragraph_format
                    if spf.line_spacing is not None or spf.line_spacing_rule is not None:
                        if spf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
                            is_single = True
                            confirmed = True
                        elif spf.line_spacing is not None and hasattr(spf.line_spacing, '__float__'):
                            confirmed = True
                            try:
                                if abs(float(spf.line_spacing) - 1.0) < 0.05:
                                    is_single = True
                            except (TypeError, ValueError):
                                pass
                        break
                    style = style.base_style
            except Exception:
                pass

        if confirmed and is_single:
            self._mark('title_line_spacing')

    def _eval_abstract_title_format(self):
        # 细则："文档中摘要页标题字体格式为黑体、四号，且在题目页之后"
        # 判据（针对办公软件的真实渲染）:
        #   1) 存在题目"临海公共书房网络化运营与更新规划书"（题目页锚点）
        #   2) 存在摘要页标题（段落文本严格为"摘要"）
        #   3) 该摘要标题的字体为"黑体"、字号"四号"
        #   4) 摘要标题在题目所在页之后。办公软件里的分页判据是:
        #      显式分页符 / w:pageBreakBefore / 中间节分隔（sectPr）
        #      / 累积高度超过一页可用高度 / 大量空段落做手动分页
        title_text = "临海公共书房网络化运营与更新规划书"
        wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

        # 页面可用高度(cm)
        try:
            section = self.document.sections[0]
            content_height_cm = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm
        except Exception:
            content_height_cm = 24.62

        title_idx = None
        abstract_idx = None
        paragraphs = self.document.paragraphs
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if title_idx is None and title_text in text:
                title_idx = i
                continue
            if title_idx is not None and abstract_idx is None:
                normalized = text.replace(" ", "").replace("　", "")
                if normalized == "摘要":
                    abstract_idx = i
                    break

        if title_idx is None or abstract_idx is None:
            return

        # 校验摘要标题字体：黑体、四号
        abstract_para = paragraphs[abstract_idx]
        font_ok = False
        for run in abstract_para.runs:
            if run.text.strip():
                if self.is_font_name(run, "黑体") and self.is_font_size(run, "四号"):
                    font_ok = True
                break
        if not font_ok:
            return

        after_title_page = False

        # 1) 摘要段落自身设置了段前分页
        if abstract_para.paragraph_format.page_break_before:
            after_title_page = True

        # 2) 题目至摘要之间存在显式分页符 / 中间段落设了段前分页 / 节分隔
        if not after_title_page:
            for k in range(title_idx, abstract_idx + 1):
                p_elem = paragraphs[k]._element
                brs = p_elem.findall('.//' + qn('w:br'))
                if any(b.get(qn('w:type')) == 'page' for b in brs):
                    after_title_page = True
                    break
                if k > title_idx and paragraphs[k].paragraph_format.page_break_before:
                    after_title_page = True
                    break
                ppr = p_elem.find(qn('w:pPr'))
                if ppr is not None and ppr.find(qn('w:sectPr')) is not None and k < abstract_idx:
                    after_title_page = True
                    break

        # 3) 累积高度估算：结合段落标记字号、行距、段前/段后间距
        if not after_title_page:
            def _para_height_cm(p):
                p_elem = p._element
                extents = p_elem.findall('.//{%s}extent' % wp_ns)
                if extents:
                    h = 0.0
                    for ext in extents:
                        cy = ext.get('cy')
                        if cy:
                            h += int(cy) / 360000.0
                    if h > 0:
                        return h
                ppr = p_elem.find(qn('w:pPr'))
                para_mark_pt = 10.5
                if ppr is not None:
                    rpr = ppr.find(qn('w:rPr'))
                    if rpr is not None:
                        sz_el = rpr.find(qn('w:sz'))
                        if sz_el is not None:
                            try:
                                para_mark_pt = int(sz_el.get(qn('w:val'))) / 2.0
                            except (TypeError, ValueError):
                                pass
                run_max_pt = 0.0
                for r in p_elem.findall(qn('w:r')):
                    rp = r.find(qn('w:rPr'))
                    if rp is not None:
                        s = rp.find(qn('w:sz'))
                        if s is not None:
                            try:
                                run_max_pt = max(run_max_pt, int(s.get(qn('w:val'))) / 2.0)
                            except (TypeError, ValueError):
                                pass
                base_pt = max(para_mark_pt, run_max_pt)
                line_mult = 1.0
                sp_before_pt = 0.0
                sp_after_pt = 0.0
                if ppr is not None:
                    sp = ppr.find(qn('w:spacing'))
                    if sp is not None:
                        line_val = sp.get(qn('w:line'))
                        line_rule = sp.get(qn('w:lineRule'))
                        if line_val:
                            try:
                                lv = int(line_val)
                                if line_rule in (None, 'auto'):
                                    line_mult = lv / 240.0
                                elif line_rule in ('exact', 'atLeast'):
                                    base_pt = max(base_pt, lv / 20.0)
                                    line_mult = 1.0
                            except ValueError:
                                pass
                        bf = sp.get(qn('w:before'))
                        af = sp.get(qn('w:after'))
                        try:
                            if bf: sp_before_pt = int(bf) / 20.0
                        except ValueError:
                            pass
                        try:
                            if af: sp_after_pt = int(af) / 20.0
                        except ValueError:
                            pass
                        bl = sp.get(qn('w:beforeLines'))
                        al = sp.get(qn('w:afterLines'))
                        try:
                            if bl:
                                sp_before_pt = max(sp_before_pt, int(bl) / 100.0 * base_pt * line_mult)
                        except ValueError:
                            pass
                        try:
                            if al:
                                sp_after_pt = max(sp_after_pt, int(al) / 100.0 * base_pt * line_mult)
                        except ValueError:
                            pass
                text_val = ''.join(n.text or '' for n in p_elem.iter() if n.text)
                if text_val.strip():
                    num_lines = max(1, len(text_val) / 28.0)
                else:
                    num_lines = 1
                line_h_pt = base_pt * 1.3 * line_mult
                total_pt = line_h_pt * num_lines + sp_before_pt + sp_after_pt
                return total_pt / 28.3465

            acc_cm = 0.0
            for k in range(title_idx, abstract_idx):
                acc_cm += _para_height_cm(paragraphs[k])
            if acc_cm >= content_height_cm:
                after_title_page = True

        # 4) 兜底：题目与摘要之间存在连续 ≥10 个空段落——办公场景下典型的手动分页占位
        if not after_title_page:
            consecutive_empty = 0
            max_consecutive = 0
            for k in range(title_idx + 1, abstract_idx):
                if not paragraphs[k].text.strip():
                    consecutive_empty += 1
                    max_consecutive = max(max_consecutive, consecutive_empty)
                else:
                    consecutive_empty = 0
            if max_consecutive >= 10:
                after_title_page = True

        if after_title_page:
            self._mark('abstract_title_format')

    def _eval_level1_title_style(self):
        # 细则：一级标题序号样式为"一、"，例如"一、项目定位与服务原则"
        # 判据（针对办公软件）:
        #   1) 定位**全部**一级标题——不用"以中文数字+、开头"作为候选筛选
        #      条件（否则相当于用结论筛选样本、必然自证合规），改按办公软件
        #      里"一级标题"的排版特征定位：
        #        · 独占一段的短文本（≤30 字符）
        #        · 段内所有非空 run 字号为"小三号"(15pt)——与配套细则
        #          "一级标题字体格式为黑体、小三号"保持一致；即便字体
        #          误设，字号通常仍是一级标题特有的小三号
        #        · 不在表格单元格内
        #        · 段末无句号、段内无句号/分号/感叹号/问号
        #   2) 序号样式须**全部**为中文数字 + "、"（"一、""二、"…"十、"）
        #      才算命中；任一一级标题序号形式不符即不加分
        #   3) 不再设"至少存在若干个一级标题"这类 rubric 之外的数量下限——
        #      只要候选集合非空，就以"全部合规"为唯一判据
        num_pattern = re.compile(r'^([一二三四五六七八九十]+)、\s*')
        titles: list[str] = []
        for p in self.document.paragraphs:
            text = p.text.strip()
            if not text or len(text) > 30:
                continue
            # 排除表格单元格中的段落
            parent = p._element.getparent()
            if parent is not None and 'tc' in parent.tag:
                continue
            # 排除内部含句号/分号/问号/感叹号的段落（真正的一级标题不带这些）
            core = text.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            if text.endswith(('。', '；', ';')):
                continue
            # 排除仅由数字/符号构成的行（如页码、公式编号占位）
            if not any('一' <= ch <= '鿿' for ch in text):
                continue
            # 一级标题字号特征：小三号 (15pt)。段内至少存在一个非空 run，
            # 且该段所有非空 run 字号均为小三号
            has_run = False
            all_xiaosan = True
            for run in p.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if not self.is_font_size(run, "小三"):
                    all_xiaosan = False
                    break
            if not has_run or not all_xiaosan:
                continue
            titles.append(text)
        if not titles:
            return
        # 序号样式：所有识别到的一级标题都必须以中文数字+"、"开头
        if all(num_pattern.match(t) for t in titles):
            self._mark('level1_title_style')



    def _eval_level2_title_style(self):
        # 细则：二级标题序号样式为"（一）"，例如"（一）公告阅读空间的建设背景"
        # 判据（针对办公软件）:
        #   1) 定位**全部**二级标题——不以"（中文数字）"起始作为候选筛选
        #      （否则等于用结论筛样本，必然自证合规）；改按办公软件里
        #      "二级标题"的排版特征定位：
        #        · 独占一段的短文本（≤40 字符）
        #        · 段内所有非空 run 字号为"四号"(14pt)——与配套细则
        #          "二级标题字体格式为黑体、四号"保持一致；即便字体
        #          误设，字号通常仍是二级标题特有的四号
        #        · 不在表格单元格内
        #        · 段末无句号、段内无句号/分号/感叹号/问号
        #        · 排除已识别为一级标题（小三号 15pt）的段落
        #   2) 序号样式须**全部**为全角括号 + 中文数字 + 全角括号
        #      （"（一）""（二）"…"（十）"，U+FF08 / U+FF09）才算命中；
        #      任一二级标题序号形式不符即不加分
        #   3) 不再设"至少 N 个标题"这类 rubric 之外的数量下限——只要
        #      候选集合非空，就以"全部合规"为唯一判据
        num_pattern = re.compile(r'^（[一二三四五六七八九十]+）\s*')
        titles: list[str] = []
        for p in self.document.paragraphs:
            text = p.text.strip()
            if not text or len(text) > 40:
                continue
            # 排除表格单元格中的段落
            parent = p._element.getparent()
            if parent is not None and 'tc' in parent.tag:
                continue
            # 排除内部含句号/分号/问号/感叹号的段落
            core = text.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            if text.endswith(('。', '；', ';')):
                continue
            # 必须包含中文字符
            if not any('一' <= ch <= '鿿' for ch in text):
                continue
            # 二级标题字号特征：四号 (14pt)；同时排除小三号（一级标题）
            has_run = False
            all_sihao = True
            for run in p.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if self.is_font_size(run, "小三"):
                    all_sihao = False
                    break
                if not self.is_font_size(run, "四号"):
                    all_sihao = False
                    break
            if not has_run or not all_sihao:
                continue
            titles.append(text)
        if not titles:
            return
        # 序号样式：所有识别到的二级标题都必须以"（中文数字）"开头
        if all(num_pattern.match(t) for t in titles):
            self._mark('level2_title_style')

    def _eval_level1_title_font(self):
        # 细则：一级标题字体格式为黑体、小三号
        # 判据（针对办公软件的真实识别）：
        #   1) 定位一级标题段落——序号"一、"…"十、"、独立短标题（与
        #      _eval_level1_title_style 使用同一套筛选，避免把正文列举误当标题）
        #   2) 段内每个非空 run 的中文字体（w:eastAsia）为"黑体"、字号为"小三"(15pt)
        #      办公软件中的字体设定通常在 w:rFonts 的 w:eastAsia 上，
        #      因此以 is_font_name 命中即可
        #   3) 必须存在至少一个一级标题，且**所有**识别到的一级标题
        #      均满足黑体+小三号才算命中（rubric 未允许 80% 阈值）
        pattern = re.compile(r'^([一二三四五六七八九十]+)、\s*(.*)$')
        total = 0
        correct = 0
        for para in self.document.paragraphs:
            text = para.text.strip()
            m = pattern.match(text)
            if not m:
                continue
            body = m.group(2).strip()
            if len(text) > 30 or not body:
                continue
            core = body.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            total += 1
            # 该段所有非空 run 都需满足 黑体 + 小三
            runs_ok = True
            has_run = False
            for run in para.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if not (self.is_font_name(run, "黑体") and self.is_font_size(run, "小三")):
                    runs_ok = False
                    break
            if has_run and runs_ok:
                correct += 1
        if total > 0 and correct == total:
            self._mark('level1_title_font')

    def _eval_level2_title_font(self):
        # 细则：二级标题字体格式为黑体、四号
        # 判据（针对办公软件的真实识别）：
        #   1) 定位二级标题段落——序号"（一）"…"（十）"（全角圆括号）、
        #      独立短标题（与 _eval_level2_title_style 使用同一套筛选，
        #      避免把正文列举误当标题）
        #   2) 段内每个非空 run 的中文字体（w:rFonts/@w:eastAsia）为"黑体"、
        #      字号为"四号"(14pt)。办公软件中常把 "（一）xxx" 拆成多个 run，
        #      必须全部合规
        #   3) 必须存在至少一个二级标题，且**所有**识别到的二级标题
        #      段内所有非空 run 均满足黑体+四号才算命中（rubric 未允许 80% 阈值）
        pattern = re.compile(r'^（([一二三四五六七八九十]+)）\s*(.*)$')
        total = 0
        correct = 0
        for para in self.document.paragraphs:
            text = para.text.strip()
            m = pattern.match(text)
            if not m:
                continue
            body = m.group(2).strip()
            if len(text) > 40 or not body:
                continue
            core = body.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            total += 1
            runs_ok = True
            has_run = False
            for run in para.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if not (self.is_font_name(run, "黑体") and self.is_font_size(run, "四号")):
                    runs_ok = False
                    break
            if has_run and runs_ok:
                correct += 1
        if total > 0 and correct == total:
            self._mark('level2_title_font')

    def _eval_level3_title_style(self):
        # 细则：三级标题序号样式为"1."，例如"1.城市文化消费与邻里阅读需求"
        # 判据（针对办公软件）:
        #   1) 定位**全部**三级标题——不以"数字+英文句点"起始作为候选筛选
        #      （否则等于用结论筛样本，误标为"1、""1)""（1）"等错误样式的
        #      三级标题永远不会被视为候选，必然自证合规）。改用广义"数字起始
        #      +分隔符"匹配作为候选：`^[（(]?\d+[.、）)．,]\s*`
        #      注：与一/二级不同，三级标题字号为小四(12pt)，与正文常见字号
        #      相同，因此不能像一/二级那样用字号特征扩大候选，只能靠段落
        #      形状（独占一段、短、以数字/带括号数字开头 + 分隔符）
        #   2) 序号样式须**全部**为阿拉伯数字 + 英文半角句点（"1.""2."…）
        #      才算命中；任一三级标题使用"1、""1)""(1)""1．"等变体即不加分
        #   3) 不再设"至少 N 个标题"这类 rubric 之外的数量下限——只要
        #      候选集合非空，就以"全部合规"为唯一判据
        cand_pattern = re.compile(r'^[（(]?\d+[.、）)．,]\s*\S')
        num_pattern = re.compile(r'^\d+\.\s*')
        titles: list[str] = []
        for p in self.document.paragraphs:
            text = p.text.strip()
            if not text or len(text) > 40:
                continue
            # 排除表格单元格中的段落
            parent = p._element.getparent()
            if parent is not None and 'tc' in parent.tag:
                continue
            if not cand_pattern.match(text):
                continue
            # 排除内部含句号/分号/问号/感叹号的段落（真正的三级标题不带这些）
            core = text.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            if text.endswith(('。', '；', ';')):
                continue
            titles.append(text)
        if not titles:
            return
        # 序号样式：所有识别到的三级标题都必须以"数字+英文半角句点"开头
        if all(num_pattern.match(t) for t in titles):
            self._mark('level3_title_style')

    def _eval_level3_title_font(self):
        # 细则：三级标题字体格式为宋体、小四号
        # 判据（针对办公软件的真实识别）：
        #   1) 定位三级标题段落——序号"1.""2."…、独立短标题（与
        #      _eval_level3_title_style 使用同一套筛选，避免把正文列举误当标题）
        #   2) 段内每个非空 run 的中文字体（w:rFonts/@w:eastAsia）为"宋体"、
        #      字号为"小四"(12pt)。办公软件中一段标题可能被拆成多个 run，
        #      必须全部合规
        #   3) 必须存在至少一个三级标题，且**所有**识别到的三级标题
        #      段内所有非空 run 均满足宋体+小四号才算命中（rubric 未允许 80% 阈值）
        pattern = re.compile(r'^(\d+)\.\s*(.*)$')
        total = 0
        correct = 0
        for para in self.document.paragraphs:
            text = para.text.strip()
            m = pattern.match(text)
            if not m:
                continue
            body = m.group(2).strip()
            if not body or len(text) > 40:
                continue
            core = body.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            total += 1
            runs_ok = True
            has_run = False
            for run in para.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if not (self.is_font_name(run, "宋体") and self.is_font_size(run, "小四")):
                    runs_ok = False
                    break
            if has_run and runs_ok:
                correct += 1
        if total > 0 and correct == total:
            self._mark('level3_title_font')

    def _eval_level4_title_style(self):
        # 细则：四级标题序号样式为"（1）"，例如"（1）公共书房网络的服务对象与覆盖逻辑"
        # 判据（针对办公软件）:
        #   1) 定位**全部**四级标题——不以"（全角括号）阿拉伯数字（全角括号）"
        #      这一结论作为候选筛选（否则误用半角括号"(1)"、或"（1）"改成
        #      "1)"".1"等错误样式的四级标题永远不会被视为候选，必然自证合规）。
        #      改用广义"带括号阿拉伯数字 + 分隔符"匹配作为候选：
        #      `^[（(]?\d+[）).、．]\s*\S`
        #      注：四级标题字号为小四(12pt)，与正文常见字号相同，无法像
        #      一/二级那样用字号特征扩大候选，只能靠段落形状（独占一段、短、
        #      以（带括号）数字开头 + 分隔符、非表格、无内部句号）
        #   2) 序号样式须**全部**为全角括号 + 阿拉伯数字 + 全角括号
        #      （"（1）""（2）"…，U+FF08 / U+FF09）才算命中；任一四级标题
        #      使用半角"(1)"或"1)"".1"等变体即不加分
        #   3) 不再设"至少 N 个标题"这类 rubric 之外的数量下限——只要
        #      候选集合非空，就以"全部合规"为唯一判据
        cand_pattern = re.compile(r'^[（(]?\d+[）).、．]\s*\S')
        num_pattern = re.compile(r'^（\d+）\s*')
        titles: list[str] = []
        for p in self.document.paragraphs:
            text = p.text.strip()
            if not text or len(text) > 40:
                continue
            # 排除表格单元格中的段落
            parent = p._element.getparent()
            if parent is not None and 'tc' in parent.tag:
                continue
            if not cand_pattern.match(text):
                continue
            core = text.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            if text.endswith(('。', '；', ';')):
                continue
            titles.append(text)
        if not titles:
            return
        # 序号样式：所有识别到的四级标题都必须以"（阿拉伯数字）"（全角括号）开头
        if all(num_pattern.match(t) for t in titles):
            self._mark('level4_title_style')

    def _eval_level4_title_font(self):
        # 细则：四级标题字体格式为宋体、小四号
        # 判据（针对办公软件的真实识别）：
        #   1) 定位四级标题段落——序号"（1）"…（全角圆括号 + 阿拉伯数字）、
        #      独立短标题（与 _eval_level4_title_style 使用同一套筛选，
        #      避免把正文列举误当标题）
        #   2) 段内每个非空 run 的中文字体（w:rFonts/@w:eastAsia）为"宋体"、
        #      字号为"小四"(12pt)。办公软件中一段标题可能被拆成多个 run，
        #      必须全部合规
        #   3) 必须存在至少一个四级标题，且**所有**识别到的四级标题
        #      段内所有非空 run 均满足宋体+小四号才算命中（rubric 未允许 80% 阈值）
        pattern = re.compile(r'^（(\d+)）\s*(.*)$')
        total = 0
        correct = 0
        for para in self.document.paragraphs:
            text = para.text.strip()
            m = pattern.match(text)
            if not m:
                continue
            body = m.group(2).strip()
            if not body or len(text) > 40:
                continue
            core = body.rstrip('。！？.')
            if any(ch in core for ch in '。；！？;'):
                continue
            total += 1
            runs_ok = True
            has_run = False
            for run in para.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if not (self.is_font_name(run, "宋体") and self.is_font_size(run, "小四")):
                    runs_ok = False
                    break
            if has_run and runs_ok:
                correct += 1
        if total > 0 and correct == total:
            self._mark('level4_title_font')

    def _eval_body_font_format(self):
        # 细则：除标题、表格、图注、表注之外的所有字体格式统一设置为宋体、小四号
        # 判据（针对办公软件的真实识别）：
        #   1) 排除 一/二/三/四 级标题——沿用各级标题的严格识别规则
        #      （序号 + 独立短标题 + 无内部句号/分号/感叹号/问号）
        #   2) 排除表格：段落父元素为 w:tc（表格单元格）时跳过
        #   3) 排除图注（段首匹配"图 数字"）与表注（段首匹配"表 数字"）
        #   4) 对其余每一段的**每一个非空 run**：解析其**实际生效**的
        #      中文字体和字号（不只看 run 显式 rPr），解析顺序遵循
        #      WordprocessingML 继承链：
        #        a. run 自身 rPr (w:rFonts/@w:eastAsia, w:sz/@w:val)
        #        b. run 引用的字符样式 (w:rStyle) 及其 basedOn 链
        #        c. 段落 rPr (w:pPr/w:rPr)
        #        d. 段落样式 (w:pStyle) 及其 basedOn 链
        #        e. 文档默认 (w:docDefaults/w:rPrDefault/w:rPr)
        #   5) rubric 用"统一设置"，无 80% 阈值——必须**所有**目标正文
        #      非空 run 解析结果都为宋体+小四号才加分

        level1 = re.compile(r'^([一二三四五六七八九十]+)、\s*(.*)$')
        level2 = re.compile(r'^（([一二三四五六七八九十]+)）\s*(.*)$')
        level3 = re.compile(r'^(\d+)\.\s*(.*)$')
        level4 = re.compile(r'^（(\d+)）\s*(.*)$')
        fig_cap = re.compile(r'^图\s*\d+')
        tbl_cap = re.compile(r'^表\s*\d+')

        def _is_title(text):
            for pat, max_len in ((level1, 30), (level2, 40), (level3, 40), (level4, 40)):
                m = pat.match(text)
                if not m:
                    continue
                body = m.group(2).strip()
                if not body or len(text) > max_len:
                    continue
                core = body.rstrip('。！？.')
                if any(ch in core for ch in '。；！？;'):
                    continue
                return True
            return False

        # ---- 字体/字号继承解析 ----
        def _ea_from_rpr(rpr):
            if rpr is None:
                return None
            r_fonts = rpr.find(qn('w:rFonts'))
            if r_fonts is None:
                return None
            val = r_fonts.get(qn('w:eastAsia'))
            return val or None

        def _sz_from_rpr(rpr):
            if rpr is None:
                return None
            sz = rpr.find(qn('w:sz'))
            if sz is None:
                return None
            val = sz.get(qn('w:val'))
            if not val:
                return None
            try:
                return float(val) / 2.0  # half-points → pt
            except (TypeError, ValueError):
                return None

        # 文档默认的 rPr（只取一次）
        default_rpr = None
        try:
            styles_elm = self.document.styles.element
            doc_defaults = styles_elm.find(qn('w:docDefaults'))
            if doc_defaults is not None:
                rpr_default = doc_defaults.find(qn('w:rPrDefault'))
                if rpr_default is not None:
                    default_rpr = rpr_default.find(qn('w:rPr'))
        except Exception:
            default_rpr = None

        def _resolve(run, para, extractor):
            # a. run 自身 rPr
            val = extractor(run._element.rPr)
            if val is not None:
                return val
            # b. run 字符样式链
            try:
                style = run.style
            except Exception:
                style = None
            while style is not None:
                try:
                    style_rpr = style.element.find(qn('w:rPr'))
                except Exception:
                    style_rpr = None
                val = extractor(style_rpr)
                if val is not None:
                    return val
                try:
                    style = style.base_style
                except Exception:
                    style = None
            # c. 段落 rPr（w:pPr/w:rPr）
            ppr = para._element.find(qn('w:pPr'))
            if ppr is not None:
                val = extractor(ppr.find(qn('w:rPr')))
                if val is not None:
                    return val
            # d. 段落样式链
            try:
                p_style = para.style
            except Exception:
                p_style = None
            while p_style is not None:
                try:
                    p_style_rpr = p_style.element.find(qn('w:rPr'))
                except Exception:
                    p_style_rpr = None
                val = extractor(p_style_rpr)
                if val is not None:
                    return val
                try:
                    p_style = p_style.base_style
                except Exception:
                    p_style = None
            # e. 文档默认
            val = extractor(default_rpr)
            return val

        total_runs = 0
        correct_runs = 0
        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 排除表格中的段落
            parent = para._element.getparent()
            if parent is not None and 'tc' in parent.tag:
                continue
            # 排除标题
            if _is_title(text):
                continue
            # 排除图注、表注
            if fig_cap.match(text) or tbl_cap.match(text):
                continue
            for run in para.runs:
                if not run.text.strip():
                    continue
                total_runs += 1
                ea = _resolve(run, para, _ea_from_rpr)
                sz = _resolve(run, para, _sz_from_rpr)
                if ea and "宋体" in ea and sz is not None and abs(sz - 12.0) < 0.5:
                    correct_runs += 1
        # rubric 要求"统一设置"：全部目标正文 run 均合规
        if total_runs > 0 and correct_runs == total_runs:
            self._mark('body_font_format')

    def _eval_numbering_format(self):
        # 细则：图注、表注、附录、参考文献、公式编号均使用阿拉伯数字连续编号，
        # 如"图 1""表 1""公式（1）"
        # 判据（针对办公软件的真实识别）：
        #   1) 图注：形如"图 N"（"图" + 可选空白 + 阿拉伯数字），且编号连续
        #   2) 表注：形如"表 N"（"表" + 可选空白 + 阿拉伯数字），且编号连续
        #   3) 公式：正文/独立段中出现"公式（N）"（全角括号 + 阿拉伯数字），且编号连续
        #   4) 附录：若存在"附录 N"（阿拉伯数字），编号连续
        #   5) 参考文献：若存在参考文献条目 [N]，编号连续
        #   6) 上述五类**只对文档中实际出现的类别**进行检查（未出现则不视为缺失，
        #      细则文本用"均"意为"凡出现的都应"）
        #   7) 至少存在图注或表注两类之一（细则示例的典型场景）
        paragraphs = self.document.paragraphs

        fig_pat = re.compile(r'^图\s*(\d+)')
        tbl_pat = re.compile(r'^表\s*(\d+)')
        formula_pat = re.compile(r'公式\s*（\s*(\d+)\s*）')  # 全角括号
        formula_bad_pat = re.compile(r'公式\s*\(\s*\d+\s*\)')  # 半角括号视为不合规
        appendix_pat = re.compile(r'^附录\s*(\d+)')
        appendix_bad_pat = re.compile(r'^附录\s*([一二三四五六七八九十]+)')  # 中文数字不合规
        ref_section_pat = re.compile(r'^参考文献$')
        ref_item_pat = re.compile(r'^\[\s*(\d+)\s*\]')

        fig_nums = []
        tbl_nums = []
        formula_nums = []
        appendix_nums = []
        ref_nums = []
        has_bad_formula = False
        has_bad_appendix = False
        has_ref_section = False

        for i, p in enumerate(paragraphs):
            t = p.text.strip()
            m = fig_pat.match(t)
            if m:
                fig_nums.append(int(m.group(1)))
            m = tbl_pat.match(t)
            if m:
                tbl_nums.append(int(m.group(1)))
            # 公式编号可能出现在正文段中，全段搜索
            for mm in formula_pat.finditer(t):
                formula_nums.append(int(mm.group(1)))
            if formula_bad_pat.search(t):
                has_bad_formula = True
            m = appendix_pat.match(t)
            if m:
                appendix_nums.append(int(m.group(1)))
            if appendix_bad_pat.match(t):
                has_bad_appendix = True
            if ref_section_pat.match(t):
                has_ref_section = True
                continue
            if has_ref_section:
                mm = ref_item_pat.match(t)
                if mm:
                    ref_nums.append(int(mm.group(1)))

        def _is_continuous(nums):
            # 空列表视为该类未出现，不参与判定
            if not nums:
                return True
            return nums == list(range(1, len(nums) + 1))

        # 至少存在图注 / 表注两类之一
        if not fig_nums and not tbl_nums:
            return
        # 各类若出现，必须使用阿拉伯数字并连续
        if not _is_continuous(fig_nums):
            return
        if not _is_continuous(tbl_nums):
            return
        if not _is_continuous(formula_nums):
            return
        if has_bad_formula:
            return
        if not _is_continuous(appendix_nums):
            return
        if has_bad_appendix:
            return
        if has_ref_section and ref_nums and not _is_continuous(ref_nums):
            return

        self._mark('numbering_format')

    def _eval_figure_captions(self):
        # 细则：全文所有图片下方均出现对应图注且图注居中显示，图注按出现顺序连续编号，
        # 不出现跳号或顺序混乱
        # 判据（针对办公软件的真实识别）：
        #   1) 所有图片段落（含 w:drawing 或 w:pict）下方紧跟一个"图 N"图注段落，
        #      办公软件里可视为"图片正下方"——中间只允许空段落
        #   2) 图注段落的对齐方式为居中 (WD_ALIGN_PARAGRAPH.CENTER)
        #   3) 图注编号按出现顺序连续（1,2,3,...）不跳号、不重复、不倒序
        #   4) 图注标题（"图 N" 之后的图题文本）不重复——两条图注共用同一图题
        #      属于"顺序混乱"（同一图被配了两条注，或不同图错用了同一图题）
        wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        fig_pattern = re.compile(r'^图\s*(\d+)\s*(.*)$')
        paragraphs = self.document.paragraphs

        # 1) 找出所有含图片的段落索引
        img_indices = []
        for i, p in enumerate(paragraphs):
            if (p._element.findall('.//' + qn('w:drawing')) or
                    p._element.findall('.//' + qn('w:pict')) or
                    p._element.findall('.//{%s}anchor' % wp_ns) or
                    p._element.findall('.//{%s}inline' % wp_ns)):
                img_indices.append(i)
        if not img_indices:
            return

        # 2) 找出所有图注段落及编号、图题
        caption_entries = []  # (para_idx, number, title, is_center)
        for i, p in enumerate(paragraphs):
            m = fig_pattern.match(p.text.strip())
            if m:
                title = m.group(2).strip()
                is_center = (p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER)
                # 若段落自身未设置，追溯样式
                if p.paragraph_format.alignment is None:
                    try:
                        st = p.style
                        while st is not None:
                            if st.paragraph_format.alignment is not None:
                                is_center = (st.paragraph_format.alignment ==
                                             WD_ALIGN_PARAGRAPH.CENTER)
                                break
                            st = st.base_style
                    except Exception:
                        pass
                caption_entries.append((i, int(m.group(1)), title, is_center))

        # 3) 每张图片下方均有对应图注：图片段落之后（跳过空段）第一个非空段应为图注段
        for ip in img_indices:
            has_caption_below = False
            for k in range(ip + 1, min(ip + 6, len(paragraphs))):
                text = paragraphs[k].text.strip()
                if not text:
                    continue
                if fig_pattern.match(text):
                    has_caption_below = True
                break
            if not has_caption_below:
                return

        # 4) 全部图注居中
        if any(not c[3] for c in caption_entries):
            return

        # 5) 按出现顺序连续编号
        numbers = [c[1] for c in caption_entries]
        if not numbers:
            return
        expected = list(range(1, len(numbers) + 1))
        if numbers != expected:
            return

        # 6) 图题不得重复（重复图题意味着两条图注共用一张图或误用同题 → 顺序混乱）
        titles = [c[2] for c in caption_entries if c[2]]
        if len(set(titles)) != len(titles):
            return

        # 7) 图注数量与图片数量一致（避免出现图片有但图注缺项的隐性错配）
        if len(caption_entries) < len(img_indices):
            return

        self._mark('figure_captions')

    def _eval_figure_caption_format(self):
        # 细则：图注格式为宋体、小四号，所有图注的图序与图题之间空两格，
        # 例如"图 1  公共书房功能分区关系图"
        # 判据（针对办公软件的真实识别）：
        #   1) 识别所有图注段落——段首匹配"图 N"（"图" + 空白 + 阿拉伯数字）
        #   2) 图序与图题之间**恰好两个空格**（半角或全角均可，共计 2 个空白字符）——
        #      细则用"所有"强调统一，任一图注不符即不加分
        #   3) 每个图注段落内所有非空 run 的中文字体 (w:rFonts/@w:eastAsia) 为"宋体"、
        #      字号 (w:sz) 为"小四"(12pt)——办公软件里图注视觉一致要求全 run 合规
        #   4) 至少存在一个图注
        #   5) 图题非空——rubric 示例"图 1  公共书房功能分区关系图"中图序后
        #      必须有图题；仅"图 1  "（空两格后无文字）不满足"图题"的存在要求

        # 段首形态：图 + 空白 + 数字 + 分隔（可能是空格） + 图题
        caption_pattern = re.compile(r'^图\s*(\d+)([ 　]*)(.*)$')

        captions = []
        for para in self.document.paragraphs:
            text = para.text.strip()
            m = caption_pattern.match(text)
            if m:
                captions.append((para, m.group(1), m.group(2), m.group(3)))
        if not captions:
            return

        # 2) 图序与图题之间空两格：分隔部分必须**恰好两个空白字符**
        #    且图题（分隔之后的文本）非空
        for _, _, sep, title in captions:
            # 细则示例是半角空格 * 2，办公软件里全角空格也常被视作"两格"，
            # 因此允许两个字符长度且都是空白
            if len(sep) != 2:
                return
            # 图题必须非空——"图 1  "后无文字不满足"图题"存在要求
            if not title.strip():
                return

        # 3) 字体：宋体 + 小四，所有非空 run 均需合规
        for para, _, _, _ in captions:
            has_run = False
            for run in para.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if not (self.is_font_name(run, "宋体") and self.is_font_size(run, "小四")):
                    return
            if not has_run:
                return

        self._mark('figure_caption_format')

    def _eval_table_captions(self):
        # 细则：所有表格上方均出现表注并居中显示，表注按出现顺序连续编号，
        # 不出现跳号或顺序混乱
        # 判据（针对办公软件的真实识别）：
        #   1) 全文档所有 <w:tbl> 表格，其**上方**紧邻位置（跳过空段落）
        #      是一段"表 N"（"表" + 空白 + 阿拉伯数字）表注段
        #   2) 该表注段落对齐方式为居中 (jc=center) —— 段落自身或样式继承
        #   3) 表注编号按出现顺序连续（1,2,3,...）不跳号、不重号、不倒序
        wp_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        tbl_cap_pattern = re.compile(r'^表\s*(\d+)')

        body = self.document.element.body
        children = list(body)

        # 1) 收集表格位置 + 每个表格上方最近的非空段落
        table_positions = [idx for idx, c in enumerate(children) if c.tag == qn('w:tbl')]
        if not table_positions:
            return

        for tp in table_positions:
            found_caption = False
            # 向上回溯：跳过空段落，遇到的第一个非空段落必须是表注段
            for k in range(tp - 1, -1, -1):
                c = children[k]
                if c.tag == qn('w:p'):
                    text = ''.join(n.text or '' for n in c.iter() if n.text).strip()
                    if not text:
                        continue
                    if tbl_cap_pattern.match(text):
                        # 校验居中
                        ppr = c.find(qn('w:pPr'))
                        jc_val = None
                        if ppr is not None:
                            jc_el = ppr.find(qn('w:jc'))
                            if jc_el is not None:
                                jc_val = jc_el.get(qn('w:val'))
                        if jc_val == 'center':
                            found_caption = True
                        else:
                            # 段落自身未设置对齐，回落到样式继承
                            # 通过 python-docx 侧的样式追溯 alignment
                            try:
                                para_obj = None
                                for pobj in self.document.paragraphs:
                                    if pobj._element is c:
                                        para_obj = pobj
                                        break
                                if para_obj is not None:
                                    st = para_obj.style
                                    while st is not None:
                                        if st.paragraph_format.alignment is not None:
                                            if st.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                                found_caption = True
                                            break
                                        st = st.base_style
                            except Exception:
                                pass
                    break
                elif c.tag == qn('w:tbl'):
                    # 上一个仍是表格，表示当前表格上方没有表注段
                    break
            if not found_caption:
                return

        # 2) 表注编号按出现顺序连续
        caption_nums = []
        for child in children:
            if child.tag == qn('w:p'):
                text = ''.join(n.text or '' for n in child.iter() if n.text).strip()
                m = tbl_cap_pattern.match(text)
                if m:
                    caption_nums.append(int(m.group(1)))
        if not caption_nums:
            return
        expected = list(range(1, len(caption_nums) + 1))
        if caption_nums != expected:
            return

        # 3) 表注数量应 ≥ 表格数量（避免有表格漏配表注但被前面遍历放过的边界情况）
        if len(caption_nums) < len(table_positions):
            return

        self._mark('table_captions')

    def _eval_table_caption_format(self):
        # 细则：表注格式为宋体、小四号；所有表注的表序与表题之间空两格，
        # 例如"表 1  公共书房运营指标表"
        # 判据（针对办公软件的真实识别）：
        #   1) 识别所有表注段落——段首匹配"表 N"（"表" + 空白 + 阿拉伯数字）
        #   2) 表序与表题之间**恰好两个空白字符**（半角或全角空格均计）——
        #      细则用"所有"强调统一，任一表注不符即不加分
        #   3) 每个表注段落内所有非空 run 的中文字体 (w:rFonts/@w:eastAsia)
        #      为"宋体"、字号 (w:sz) 为"小四"(12pt)——办公软件里视觉一致
        #      要求全 run 合规
        #   4) 至少存在一个表注
        #   5) 表题非空——rubric 示例"表 1  公共书房运营指标表"中表序后
        #      必须有表题；仅"表 1  "（空两格后无文字）不满足"表题"的存在要求

        # 段首形态：表 + 空白 + 数字 + 分隔（空白）+ 表题
        caption_pattern = re.compile(r'^表\s*(\d+)([ 　]*)(.*)$')

        captions = []
        for para in self.document.paragraphs:
            text = para.text.strip()
            m = caption_pattern.match(text)
            if m:
                captions.append((para, m.group(1), m.group(2), m.group(3)))
        if not captions:
            return

        # 2) 表序与表题之间空两格：分隔部分必须**恰好两个空白字符**
        #    且表题（分隔之后的文本）非空
        for _, _, sep, title in captions:
            if len(sep) != 2:
                return
            # 表题必须非空——"表 1  "后无文字不满足"表题"存在要求
            if not title.strip():
                return

        # 3) 字体：宋体 + 小四，所有非空 run 均需合规
        for para, _, _, _ in captions:
            has_run = False
            for run in para.runs:
                if not run.text.strip():
                    continue
                has_run = True
                if not (self.is_font_name(run, "宋体") and self.is_font_size(run, "小四")):
                    return
            if not has_run:
                return

        self._mark('table_caption_format')

    def _eval_table_font(self):
        # 细则：所有表格内文字统一设置为宋体、小四号、单倍行距
        # 判据（针对办公软件的真实识别）：
        #   1) 遍历文档中所有 <w:tbl> 表格 → 单元格 → 段落 → run
        #   2) 每个含文字的 run：中文字体 (w:rFonts/@w:eastAsia) 为"宋体"、
        #      字号 (w:sz) 为"小四"(12pt)
        #   3) 每个含文字的段落：行距为**单倍行距**——
        #      办公软件里表现为 line_spacing_rule=SINGLE 或 line_spacing≈1.0；
        #      段落未显式设置时沿样式继承，样式未设时按 Word 默认单倍行距
        #   4) 细则用"所有""统一"——任一 run/段落不符即不加分
        #   5) 至少存在一个表格且有文字

        def _is_single_line_spacing(para):
            pf = para.paragraph_format
            if pf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
                return True
            ls = pf.line_spacing
            if ls is not None and hasattr(ls, '__float__'):
                try:
                    if abs(float(ls) - 1.0) < 0.05:
                        return True
                except (TypeError, ValueError):
                    pass
            if ls is None and pf.line_spacing_rule is None:
                # 追溯样式
                try:
                    st = para.style
                    while st is not None:
                        spf = st.paragraph_format
                        if spf.line_spacing is not None or spf.line_spacing_rule is not None:
                            if spf.line_spacing_rule == WD_LINE_SPACING.SINGLE:
                                return True
                            if spf.line_spacing is not None and hasattr(spf.line_spacing, '__float__'):
                                try:
                                    if abs(float(spf.line_spacing) - 1.0) < 0.05:
                                        return True
                                except (TypeError, ValueError):
                                    pass
                            return False
                        st = st.base_style
                except Exception:
                    pass
                # 段落与样式均未设置——办公软件默认单倍行距
                return True
            return False

        tables = self.document.tables
        if not tables:
            return

        any_text = False
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_has_text = False
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                            para_has_text = True
                            any_text = True
                            # 字体：宋体 + 小四
                            if not (self.is_font_name(run, "宋体") and
                                    self.is_font_size(run, "小四")):
                                return
                        # 行距校验：仅对含文字段落检查
                        if para_has_text and not _is_single_line_spacing(para):
                            return

        if not any_text:
            return

        self._mark('table_font')

    def _eval_paragraph_format_deduction(self):
        # 细则：摘要页后的所有页，除了各级标题、表注、图注和表格中的文字之外，
        # 其余的所有文本内容段落格式不满足：1.25倍行距，段前间距为0.5行，
        # 首行缩进两字符 → 扣 -5
        # 判据（针对办公软件的真实识别）：
        #   1) 起点：从"摘要页之后"开始——按分页信号（显式分页符 /
        #      lastRenderedPageBreak / w:pageBreakBefore / w:sectPr / 累积
        #      内容高度超过页面可用高度）估算每个段落的物理页码，取"摘要"
        #      段所在页**之后**所有页的段落。不再简化为"摘要段后第一个
        #      一级标题之后"（那样会漏检摘要页后、首个一级标题之前的正文段）
        #   2) 排除：一/二/三/四 级标题（严格识别）、表注、图注、表格中的文字
        #   3) 每个被检查段落必须满足三条：
        #      a) 行距 = 1.25 倍（w:line=300 且 w:lineRule=auto，或 line_spacing≈1.25，
        #         段落未设时沿样式继承）
        #      b) 段前间距 = 0.5 行（w:beforeLines=50，即 0.5 行）
        #      c) 首行缩进 = 两字符（w:firstLineChars=200，或 firstLine 的字符数≈2）
        #   4) 三条中任一不满足即视为该段落"格式不满足"
        #   5) rubric 无 30% 容忍门槛——只要存在**任一**目标正文段落格式
        #      不满足，即判定"段落格式不满足"，扣 -5
        paragraphs = self.document.paragraphs
        title_text = "临海公共书房网络化运营与更新规划书"

        # 页面可用高度（与 _eval_title_line_spacing 同口径）
        try:
            section = self.document.sections[0]
            page_h = getattr(section, 'page_height', None)
            top_margin = getattr(section, 'top_margin', None)
            bottom_margin = getattr(section, 'bottom_margin', None)
            if page_h is None or top_margin is None or bottom_margin is None:
                raise ValueError('section metrics unavailable')
            page_h_cm = page_h.cm - top_margin.cm - bottom_margin.cm
        except Exception:
            page_h_cm = 24.62

        def _para_mark_pt(p_elem):
            ppr = p_elem.find(qn('w:pPr'))
            mark_pt = 10.5
            if ppr is not None:
                rpr = ppr.find(qn('w:rPr'))
                if rpr is not None:
                    sz_el = rpr.find(qn('w:sz'))
                    if sz_el is not None:
                        try:
                            mark_pt = int(sz_el.get(qn('w:val'))) / 2.0
                        except (TypeError, ValueError):
                            pass
            run_max_pt = 0.0
            for r in p_elem.findall(qn('w:r')):
                rp = r.find(qn('w:rPr'))
                if rp is not None:
                    s = rp.find(qn('w:sz'))
                    if s is not None:
                        try:
                            run_max_pt = max(run_max_pt, int(s.get(qn('w:val'))) / 2.0)
                        except (TypeError, ValueError):
                            pass
            return max(mark_pt, run_max_pt)

        def _para_height_cm(p_elem):
            base_pt = _para_mark_pt(p_elem)
            line_mult = 1.0
            ppr = p_elem.find(qn('w:pPr'))
            if ppr is not None:
                sp = ppr.find(qn('w:spacing'))
                if sp is not None:
                    line_val = sp.get(qn('w:line'))
                    line_rule = sp.get(qn('w:lineRule'))
                    if line_val:
                        try:
                            lv = int(line_val)
                            if line_rule in (None, 'auto'):
                                line_mult = lv / 240.0
                            elif line_rule in ('exact', 'atLeast'):
                                base_pt = max(base_pt, lv / 20.0)
                                line_mult = 1.0
                        except ValueError:
                            pass
            text_val = ''.join(n.text or '' for n in p_elem.iter() if n.text)
            num_lines = max(1, len(text_val) / 28.0) if text_val.strip() else 1
            return (base_pt * 1.3 * line_mult * num_lines) / 28.3465

        # 模拟分页：得到每个段落所在物理页码
        page_of = []
        cur_page = 1
        cur_h = 0.0
        for p in paragraphs:
            p_elem = p._element
            ppr = p_elem.find(qn('w:pPr'))
            paged = False
            if ppr is not None:
                pb = ppr.find(qn('w:pageBreakBefore'))
                if pb is not None and pb.get(qn('w:val')) not in ('0', 'false'):
                    paged = True
            if not paged:
                brs = p_elem.findall('.//' + qn('w:br'))
                if any(b.get(qn('w:type')) == 'page' for b in brs):
                    paged = True
            if not paged and p_elem.findall('.//' + qn('w:lastRenderedPageBreak')):
                paged = True
            if paged:
                cur_page += 1
                cur_h = 0.0
            h = _para_height_cm(p_elem)
            if cur_h + h > page_h_cm and cur_h > 0:
                cur_page += 1
                cur_h = h
            else:
                cur_h += h
            page_of.append(cur_page)
            if ppr is not None and ppr.find(qn('w:sectPr')) is not None:
                cur_page += 1
                cur_h = 0.0

        # 定位"摘要"段所在物理页
        abstract_page = None
        for i, p in enumerate(paragraphs):
            if title_text in p.text:
                continue
            if p.text.strip().replace(" ", "").replace("　", "") == "摘要":
                abstract_page = page_of[i]
                break
        if abstract_page is None:
            return

        # 2) 排除标题的识别（沿用之前严格规则）
        level1 = re.compile(r'^([一二三四五六七八九十]+)、\s*(.*)$')
        level2 = re.compile(r'^（([一二三四五六七八九十]+)）\s*(.*)$')
        level3 = re.compile(r'^(\d+)\.\s*(.*)$')
        level4 = re.compile(r'^（(\d+)）\s*(.*)$')
        fig_cap = re.compile(r'^图\s*\d+')
        tbl_cap = re.compile(r'^表\s*\d+')

        def _is_title(text):
            for pat, max_len in ((level1, 30), (level2, 40), (level3, 40), (level4, 40)):
                m = pat.match(text)
                if not m:
                    continue
                body = m.group(2).strip()
                if not body or len(text) > max_len:
                    continue
                core = body.rstrip('。！？.')
                if any(ch in core for ch in '。；！？;'):
                    continue
                return True
            return False

        # 3) 逐段判据
        def _get_pPr_spacing_and_ind(para):
            ppr = para._element.find(qn('w:pPr'))
            sp_attrs = {}
            ind_attrs = {}
            if ppr is not None:
                sp = ppr.find(qn('w:spacing'))
                if sp is not None:
                    for k, v in sp.attrib.items():
                        sp_attrs[k.split('}')[-1]] = v
                ind = ppr.find(qn('w:ind'))
                if ind is not None:
                    for k, v in ind.attrib.items():
                        ind_attrs[k.split('}')[-1]] = v
            return sp_attrs, ind_attrs

        def _get_eastasia_font_pt(para):
            # 段落中首个非空 run 的中文字号（单位 pt），用于把 firstLine 换算成字符数
            for r in para.runs:
                if not r.text.strip():
                    continue
                if r.font.size is not None:
                    return r.font.size.pt
                break
            # 段落标记字号
            ppr = para._element.find(qn('w:pPr'))
            if ppr is not None:
                rpr = ppr.find(qn('w:rPr'))
                if rpr is not None:
                    sz_el = rpr.find(qn('w:sz'))
                    if sz_el is not None:
                        try:
                            return int(sz_el.get(qn('w:val'))) / 2.0
                        except (TypeError, ValueError):
                            pass
            return 10.5  # Word 中文默认字号五号

        def _line_spacing_ok(para, sp_attrs):
            # (a) 段落直接设置
            line_val = sp_attrs.get('line')
            line_rule = sp_attrs.get('lineRule')
            if line_val:
                try:
                    lv = int(line_val)
                    if line_rule in (None, 'auto'):
                        return abs(lv / 240.0 - 1.25) < 0.05
                except ValueError:
                    pass
                return False
            # (b) 未设置：沿样式继承
            try:
                st = para.style
                while st is not None:
                    spf = st.paragraph_format
                    if spf.line_spacing_rule is not None or spf.line_spacing is not None:
                        if spf.line_spacing is not None and hasattr(spf.line_spacing, '__float__'):
                            try:
                                return abs(float(spf.line_spacing) - 1.25) < 0.05
                            except (TypeError, ValueError):
                                return False
                        return False
                    st = st.base_style
            except Exception:
                pass
            return False

        def _space_before_ok(para, sp_attrs):
            # 优先 beforeLines（单位 1/100 行），50 = 0.5 行
            bl = sp_attrs.get('beforeLines')
            if bl is not None:
                try:
                    return abs(int(bl) / 100.0 - 0.5) < 0.05
                except ValueError:
                    return False
            # 沿样式继承
            try:
                st = para.style
                while st is not None:
                    ppr = st.element.find(qn('w:pPr')) if hasattr(st.element, 'find') else None
                    if ppr is not None:
                        sp = ppr.find(qn('w:spacing'))
                        if sp is not None:
                            bl2 = sp.get(qn('w:beforeLines'))
                            if bl2 is not None:
                                try:
                                    return abs(int(bl2) / 100.0 - 0.5) < 0.05
                                except ValueError:
                                    return False
                    st = st.base_style
            except Exception:
                pass
            return False

        def _first_line_indent_ok(para, ind_attrs):
            # 优先 firstLineChars（单位 1/100 字符），200 = 2 字符
            flc = ind_attrs.get('firstLineChars')
            if flc is not None:
                try:
                    if abs(int(flc) / 100.0 - 2.0) < 0.05:
                        return True
                except ValueError:
                    pass
            # 其次 firstLine (twips)，按段落中文字号换算字符数
            fl = ind_attrs.get('firstLine')
            if fl is not None:
                try:
                    twips = int(fl)
                    font_pt = _get_eastasia_font_pt(para)
                    chars = twips / 20.0 / font_pt
                    # Word 中"2字符"在 12pt(小四) 与 10.5pt(五号) 下 twips 分别为 480 与 420，
                    # 均视为"2 字符"渲染
                    if abs(chars - 2.0) < 0.3:
                        return True
                    # 也允许按五号(10.5pt)基准换算：420 twips = 2 字符
                    if abs(twips / 20.0 / 10.5 - 2.0) < 0.15:
                        return True
                    return False
                except ValueError:
                    return False
            # 沿样式继承
            try:
                st = para.style
                while st is not None:
                    ppr = st.element.find(qn('w:pPr')) if hasattr(st.element, 'find') else None
                    if ppr is not None:
                        ind = ppr.find(qn('w:ind'))
                        if ind is not None:
                            flc2 = ind.get(qn('w:firstLineChars'))
                            if flc2 is not None:
                                try:
                                    if abs(int(flc2) / 100.0 - 2.0) < 0.05:
                                        return True
                                except ValueError:
                                    pass
                    st = st.base_style
            except Exception:
                pass
            return False

        total = 0
        wrong = 0
        for i in range(len(paragraphs)):
            # 只检查"摘要页之后"的页：段落物理页码 > 摘要页
            if page_of[i] <= abstract_page:
                continue
            para = paragraphs[i]
            text = para.text.strip()
            if not text:
                continue
            # 排除表格内文字
            parent = para._element.getparent()
            if parent is not None and 'tc' in parent.tag:
                continue
            # 排除标题
            if _is_title(text):
                continue
            # 排除图注/表注
            if fig_cap.match(text) or tbl_cap.match(text):
                continue

            total += 1
            sp_attrs, ind_attrs = _get_pPr_spacing_and_ind(para)
            if not (_line_spacing_ok(para, sp_attrs) and
                    _space_before_ok(para, sp_attrs) and
                    _first_line_indent_ok(para, ind_attrs)):
                wrong += 1

        # rubric 无 30% 容忍门槛：任一目标正文段落格式不满足即扣 -5
        if total > 0 and wrong > 0:
            self._mark('paragraph_format_deduction')

    def _eval_title_font_deduction(self):
        # 细则：第1页报告题目"临海公共书房网络化运营与更新规划书"
        # 字体格式不满足 宋体、三号、加粗、居中 → 扣 -1
        # 针对办公软件的取值：
        #   居中：段落 w:jc/@w:val="center"（或样式继承）
        #   宋体：run 的 w:rFonts/@w:eastAsia 包含 "宋体"（或段落 rPr/样式继承）
        #   三号：run 的 w:sz/@w:val = 32（半磅），即 16pt（±0.5 容差）
        #   加粗：run 的 w:b 存在且不为 "0"/"false"（或段落 rPr/样式继承）
        # 四项中任一不满足即扣分
        title_text = "临海公共书房网络化运营与更新规划书"
        target_para = None
        for para in self.document.paragraphs:
            if title_text in para.text:
                target_para = para
                break
        if target_para is None:
            return

        def _resolve_alignment(para):
            # 段落直接设置
            ppr = para._element.find(qn('w:pPr'))
            if ppr is not None:
                jc = ppr.find(qn('w:jc'))
                if jc is not None:
                    return jc.get(qn('w:val'))
            # 样式继承
            try:
                st = para.style
                while st is not None:
                    spr = st.element.find(qn('w:pPr')) if hasattr(st.element, 'find') else None
                    if spr is not None:
                        jc = spr.find(qn('w:jc'))
                        if jc is not None:
                            return jc.get(qn('w:val'))
                    st = st.base_style
            except Exception:
                pass
            return None

        def _run_has_songti(run, para):
            rpr = run._element.rPr
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    ea = rfonts.get(qn('w:eastAsia'))
                    if ea:
                        return '宋体' in ea or 'SimSun' in ea
                    ascii_f = rfonts.get(qn('w:ascii'))
                    if ascii_f and ('宋体' in ascii_f or 'SimSun' in ascii_f):
                        return True
            # 段落 rPr（段落标记）
            ppr = para._element.find(qn('w:pPr'))
            if ppr is not None:
                prpr = ppr.find(qn('w:rPr'))
                if prpr is not None:
                    rfonts = prpr.find(qn('w:rFonts'))
                    if rfonts is not None:
                        ea = rfonts.get(qn('w:eastAsia'))
                        if ea and ('宋体' in ea or 'SimSun' in ea):
                            return True
            return False

        def _run_size_pt(run, para):
            rpr = run._element.rPr
            if rpr is not None:
                sz = rpr.find(qn('w:sz'))
                if sz is not None:
                    try:
                        return int(sz.get(qn('w:val'))) / 2.0
                    except (TypeError, ValueError):
                        pass
            ppr = para._element.find(qn('w:pPr'))
            if ppr is not None:
                prpr = ppr.find(qn('w:rPr'))
                if prpr is not None:
                    sz = prpr.find(qn('w:sz'))
                    if sz is not None:
                        try:
                            return int(sz.get(qn('w:val'))) / 2.0
                        except (TypeError, ValueError):
                            pass
            return None

        def _run_is_bold(run, para):
            rpr = run._element.rPr
            if rpr is not None:
                b = rpr.find(qn('w:b'))
                if b is not None:
                    val = b.get(qn('w:val'))
                    return val not in ('0', 'false')
            ppr = para._element.find(qn('w:pPr'))
            if ppr is not None:
                prpr = ppr.find(qn('w:rPr'))
                if prpr is not None:
                    b = prpr.find(qn('w:b'))
                    if b is not None:
                        val = b.get(qn('w:val'))
                        return val not in ('0', 'false')
            return False

        # 逐项检查
        centered = _resolve_alignment(target_para) == 'center'

        songti_ok = True
        size_ok = True
        bold_ok = True
        any_run = False
        for run in target_para.runs:
            if not run.text.strip():
                continue
            any_run = True
            if not _run_has_songti(run, target_para):
                songti_ok = False
            pt = _run_size_pt(run, target_para)
            if pt is None or abs(pt - 16) >= 0.5:
                size_ok = False
            if not _run_is_bold(run, target_para):
                bold_ok = False
        if not any_run:
            songti_ok = size_ok = bold_ok = False

        if not (centered and songti_ok and size_ok and bold_ok):
            self._mark('title_font_deduction')

    def _eval_figure_table_mismatch(self):
        # 该项已按用户要求删除
        pass

    def _eval_text_image_overlap(self):
        # 该项已按用户要求删除
        pass

    def _eval_large_blank_area(self):
        # 细则：除前两页和最后一页外，文档中任意一页出现超过 40% 的大面积空白 → 扣 -5
        # 针对办公软件（Word/WPS）的分页与空白判定：
        #   1) 分页信号（按优先级）：
        #      a) 段落 w:pPr/w:pageBreakBefore（val 非 "0"/"false"）
        #      b) run 内的 <w:br w:type="page"/>
        #      c) <w:lastRenderedPageBreak/>（Word 保存时写回的实际分页位置）
        #      d) 累积内容高度 > 页面可用高度（正文区）
        #   2) 页面可用高度 = page_height − top_margin − bottom_margin
        #   3) 每页"空白" = (可用高度 − 内容高度) / 可用高度；> 40% 即为大面积空白
        #   4) 排除：前两页 和 最后一页 不计
        wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        section = self.document.sections[0]
        page_h = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm
        page_w = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
        if page_h <= 0 or page_w <= 0:
            return

        def _para_font_pt(p_elem):
            # 段落内首个 run 的 w:sz，或段落 rPr/w:sz
            for r in p_elem.findall('.//' + qn('w:r')):
                rpr = r.find(qn('w:rPr'))
                if rpr is not None:
                    sz = rpr.find(qn('w:sz'))
                    if sz is not None:
                        try:
                            return int(sz.get(qn('w:val'))) / 2.0
                        except (TypeError, ValueError):
                            pass
                break
            ppr = p_elem.find(qn('w:pPr'))
            if ppr is not None:
                prpr = ppr.find(qn('w:rPr'))
                if prpr is not None:
                    sz = prpr.find(qn('w:sz'))
                    if sz is not None:
                        try:
                            return int(sz.get(qn('w:val'))) / 2.0
                        except (TypeError, ValueError):
                            pass
            return 12.0  # 默认小四

        def _para_line_mult(p_elem):
            ppr = p_elem.find(qn('w:pPr'))
            if ppr is not None:
                sp = ppr.find(qn('w:spacing'))
                if sp is not None:
                    line = sp.get(qn('w:line'))
                    rule = sp.get(qn('w:lineRule'))
                    if line:
                        try:
                            lv = int(line)
                            if rule in (None, 'auto'):
                                return lv / 240.0
                            # exact/atLeast：直接给出磅值，按 240 换算成"单位磅高"
                            return (lv / 20.0) / (12.0)
                        except ValueError:
                            pass
            return 1.0

        def _para_space_before_cm(p_elem):
            ppr = p_elem.find(qn('w:pPr'))
            if ppr is not None:
                sp = ppr.find(qn('w:spacing'))
                if sp is not None:
                    bl = sp.get(qn('w:beforeLines'))
                    if bl is not None:
                        try:
                            lines = int(bl) / 100.0
                            return lines * (_para_font_pt(p_elem) * 1.25) / 28.3465
                        except ValueError:
                            pass
                    bf = sp.get(qn('w:before'))
                    if bf is not None:
                        try:
                            return (int(bf) / 20.0) / 28.3465
                        except ValueError:
                            pass
            return 0.0

        def _para_height_cm(p_elem):
            # 图片高度优先取 wp:extent/@cy（EMU → cm）
            extents = p_elem.findall('.//{%s}extent' % wp_ns)
            img_h = 0.0
            for ext in extents:
                cy = ext.get('cy')
                if cy:
                    try:
                        img_h = max(img_h, int(cy) / 360000.0)
                    except ValueError:
                        pass
            text = ''.join(n.text or '' for n in p_elem.iter() if n.text).strip()
            font_pt = _para_font_pt(p_elem)
            line_cm = (font_pt * _para_line_mult(p_elem)) / 28.3465
            if text:
                # 每行字符数按可用宽度 / 中文字宽（=字号 pt）估算
                chars_per_line = max(1, int(page_w * 28.3465 / font_pt))
                num_lines = max(1, (len(text) + chars_per_line - 1) // chars_per_line)
                text_h = num_lines * line_cm
            else:
                text_h = 0.0 if img_h > 0 else line_cm  # 空段占一行
            return _para_space_before_cm(p_elem) + max(text_h, img_h)

        def _tbl_height_cm(tbl_elem):
            total = 0.0
            for row in tbl_elem.findall(qn('w:tr')):
                row_h_specified = None
                trPr = row.find(qn('w:trPr'))
                if trPr is not None:
                    trH = trPr.find(qn('w:trHeight'))
                    if trH is not None:
                        v = trH.get(qn('w:val'))
                        if v:
                            try:
                                row_h_specified = int(v) / 20.0 / 28.3465
                            except ValueError:
                                pass
                if row_h_specified and row_h_specified > 0:
                    total += row_h_specified
                    continue
                # 未显式设定：取该行各单元格中最高文本估算
                max_cell_h = 0.7  # 默认最小一行 ≈ 0.7cm
                for cell in row.findall(qn('w:tc')):
                    cell_h = 0.0
                    for p_elem in cell.findall(qn('w:p')):
                        cell_h += _para_height_cm(p_elem)
                    if cell_h > max_cell_h:
                        max_cell_h = cell_h
                total += max_cell_h
            return total

        # 1) 遍历 body 子元素，按分页信号 + 累积高度 切分逻辑页
        body = self.document.element.body
        pages = []
        cur = 0.0
        for child in body:
            if child.tag == qn('w:sectPr'):
                continue
            if child.tag == qn('w:p'):
                ppr = child.find(qn('w:pPr'))
                # a) pageBreakBefore
                if ppr is not None:
                    pb = ppr.find(qn('w:pageBreakBefore'))
                    if pb is not None and pb.get(qn('w:val')) not in ('0', 'false'):
                        pages.append(cur)
                        cur = 0.0
                # b) 显式 <w:br w:type="page"/>
                brs = child.findall('.//' + qn('w:br'))
                if any(b.get(qn('w:type')) == 'page' for b in brs):
                    pages.append(cur)
                    cur = 0.0
                # c) lastRenderedPageBreak
                if child.findall('.//' + qn('w:lastRenderedPageBreak')):
                    pages.append(cur)
                    cur = 0.0
                h = _para_height_cm(child)
            elif child.tag == qn('w:tbl'):
                h = _tbl_height_cm(child)
            else:
                h = 0.0
            # d) 累积高度溢出
            if cur + h > page_h and cur > 0:
                pages.append(cur)
                cur = h
            else:
                cur += h
        pages.append(cur)

        # 2) 排除前两页 + 最后一页；判定 > 40% 空白
        has_large_blank = False
        total_pages = len(pages)
        for idx, h in enumerate(pages):
            if idx < 2 or idx >= total_pages - 1:
                continue
            blank_ratio = (page_h - h) / page_h
            if blank_ratio > 0.4:
                has_large_blank = True
                break

        if has_large_blank:
            self._mark('large_blank_area')

    def _eval_table_page_break(self):
        # 细则：文档中表格出现断页、不连续 → 扣 -3
        # 针对办公软件（Word/WPS）的判据：
        #   一个表格若在办公软件中被切开跨越两页及以上，即为"断页、不连续"
        #   识别方式（任一命中即算）：
        #     A. 表格内含 <w:lastRenderedPageBreak/> —— Word 保存时写回的实际分页位置
        #     B. 表格内的 run 含 <w:br w:type="page"/> —— 显式在表内插入的分页
        #     C. 按 body 子元素顺序累积高度：表格起始时页内剩余高度 < 表格总高度
        #        —— 表格自身高度使用 w:trHeight 或按单元格文本估算，
        #           页面可用高度 = page_height − top_margin − bottom_margin
        wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        section = self.document.sections[0]
        page_h = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm
        page_w = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
        if page_h <= 0 or page_w <= 0:
            return
        broken = False

        # A/B: 直接扫描每个表格是否含 lastRenderedPageBreak 或 显式分页 <w:br type=page>
        for tbl in self.document.tables:
            tel = tbl._element
            if tel.findall('.//' + qn('w:lastRenderedPageBreak')):
                broken = True
                break
            for br in tel.findall('.//' + qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    broken = True
                    break
            if broken:
                break

        # C: 按 body 顺序模拟分页（沿用 _eval_large_blank_area 中的度量），
        #    对每个 w:tbl，若其起始位置的"页内剩余高度" < 表格总高度 → 断页
        if not broken:
            def _pt_of(elem):
                # 取段落首个 run 或段落 rPr 的 w:sz
                for r in elem.findall('.//' + qn('w:r')):
                    rpr = r.find(qn('w:rPr'))
                    if rpr is not None:
                        sz = rpr.find(qn('w:sz'))
                        if sz is not None:
                            try:
                                return int(sz.get(qn('w:val'))) / 2.0
                            except (TypeError, ValueError):
                                pass
                    break
                ppr = elem.find(qn('w:pPr'))
                if ppr is not None:
                    prpr = ppr.find(qn('w:rPr'))
                    if prpr is not None:
                        sz = prpr.find(qn('w:sz'))
                        if sz is not None:
                            try:
                                return int(sz.get(qn('w:val'))) / 2.0
                            except (TypeError, ValueError):
                                pass
                return 12.0

            def _line_mult(elem):
                ppr = elem.find(qn('w:pPr'))
                if ppr is not None:
                    sp = ppr.find(qn('w:spacing'))
                    if sp is not None:
                        line = sp.get(qn('w:line'))
                        rule = sp.get(qn('w:lineRule'))
                        if line:
                            try:
                                lv = int(line)
                                if rule in (None, 'auto'):
                                    return lv / 240.0
                                return (lv / 20.0) / 12.0
                            except ValueError:
                                pass
                return 1.0

            def _space_before_cm(elem):
                ppr = elem.find(qn('w:pPr'))
                if ppr is not None:
                    sp = ppr.find(qn('w:spacing'))
                    if sp is not None:
                        bl = sp.get(qn('w:beforeLines'))
                        if bl is not None:
                            try:
                                return (int(bl) / 100.0) * (_pt_of(elem) * 1.25) / 28.3465
                            except ValueError:
                                pass
                        bf = sp.get(qn('w:before'))
                        if bf is not None:
                            try:
                                return (int(bf) / 20.0) / 28.3465
                            except ValueError:
                                pass
                return 0.0

            def _para_h(elem):
                extents = elem.findall('.//{%s}extent' % wp_ns)
                img_h = 0.0
                for ext in extents:
                    cy = ext.get('cy')
                    if cy:
                        try:
                            img_h = max(img_h, int(cy) / 360000.0)
                        except ValueError:
                            pass
                text = ''.join(n.text or '' for n in elem.iter() if n.text).strip()
                font_pt = _pt_of(elem)
                line_cm = (font_pt * _line_mult(elem)) / 28.3465
                if text:
                    chars_per_line = max(1, int(page_w * 28.3465 / font_pt))
                    n_lines = max(1, (len(text) + chars_per_line - 1) // chars_per_line)
                    txt_h = n_lines * line_cm
                else:
                    txt_h = 0.0 if img_h > 0 else line_cm
                return _space_before_cm(elem) + max(txt_h, img_h)

            def _tbl_h(tel):
                total = 0.0
                for row in tel.findall(qn('w:tr')):
                    row_h = None
                    trPr = row.find(qn('w:trPr'))
                    if trPr is not None:
                        trH = trPr.find(qn('w:trHeight'))
                        if trH is not None:
                            v = trH.get(qn('w:val'))
                            if v:
                                try:
                                    row_h = int(v) / 20.0 / 28.3465
                                except ValueError:
                                    pass
                    if row_h and row_h > 0:
                        total += row_h
                        continue
                    max_cell = 0.7
                    for cell in row.findall(qn('w:tc')):
                        ch = 0.0
                        for p in cell.findall(qn('w:p')):
                            ch += _para_h(p)
                        if ch > max_cell:
                            max_cell = ch
                    total += max_cell
                return total

            cur = 0.0
            for child in self.document.element.body:
                if child.tag == qn('w:sectPr'):
                    continue
                if child.tag == qn('w:p'):
                    ppr = child.find(qn('w:pPr'))
                    if ppr is not None:
                        pb = ppr.find(qn('w:pageBreakBefore'))
                        if pb is not None and pb.get(qn('w:val')) not in ('0', 'false'):
                            cur = 0.0
                    if any(b.get(qn('w:type')) == 'page' for b in child.findall('.//' + qn('w:br'))):
                        cur = 0.0
                    if child.findall('.//' + qn('w:lastRenderedPageBreak')):
                        cur = 0.0
                    h = _para_h(child)
                    if cur + h > page_h and cur > 0:
                        cur = h
                    else:
                        cur += h
                elif child.tag == qn('w:tbl'):
                    th = _tbl_h(child)
                    remaining = page_h - cur
                    if th > remaining:
                        broken = True
                        break
                    cur += th

        if broken:
            self._mark('table_page_break')

    def _eval_page_margin(self):
        # 细则：文档页面边距不满足 上、下 2.54 厘米，左、右 3.18 厘米 → 扣 -3
        # 针对办公软件（Word/WPS）的取值：
        #   页面边距存储于 w:sectPr/w:pgMar，属性单位为 twips(1/1440 英寸)：
        #     w:top / w:bottom / w:left / w:right
        #   换算：cm = twips / 1440 * 2.54
        #   任一节的四项之一不满足即视为不满足（一票否决）
        #   容差：0.05 cm（对应 Word/WPS 对话框"厘米"的显示分辨率）
        target_top = target_bottom = 2.54
        target_left = target_right = 3.18
        tol = 0.05

        for section in self.document.sections:
            top = section.top_margin
            bottom = section.bottom_margin
            left = section.left_margin
            right = section.right_margin
            if top is None or bottom is None or left is None or right is None:
                self._mark('page_margin')
                return
            if (abs(top.cm - target_top) > tol or
                abs(bottom.cm - target_bottom) > tol or
                abs(left.cm - target_left) > tol or
                abs(right.cm - target_right) > tol):
                self._mark('page_margin')
                return

    def _eval_text_color(self):
        # 该项已按用户要求删除
        pass

    def build_result(self, file_name):
        """将命中集合与维度一状态整理为统一的返回字典（见接口约定 §2.2）。"""
        dim2_items = []
        for key, rule, max_delta, _text in DIM2_RULES:
            hit = key in self.matched_keys
            dim2_items.append({
                "rule": rule,
                "max_delta": max_delta,
                "delta": max_delta if hit else 0,
                "hit": hit,
                "detail": "",
            })
        total_score = sum(item["delta"] for item in dim2_items)
        max_score = sum(item["max_delta"] for item in dim2_items if item["max_delta"] > 0)
        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": self.dimension1_passed,
            "dim1_reason": "" if self.dimension1_passed else "；".join(self.dimension1_issues),
            "dim2_items": dim2_items if self.dimension1_passed else [],
            "total_score": total_score if self.dimension1_passed else 0,
            "max_score": max_score,
        }


def _find_target_document(dir_path):
    """在脚本所在目录里定位被评估的 .docx 文档。

    约定：目录内应存在唯一的 .docx（忽略以 '~$' 开头的 Word 临时文件）。
    若存在多个，取修改时间最新的一个。
    """
    if not os.path.isdir(dir_path):
        return None
    candidates = []
    for name in os.listdir(dir_path):
        if name.startswith('~$'):
            continue
        if not name.lower().endswith('.docx'):
            continue
        full = os.path.join(dir_path, name)
        if os.path.isfile(full):
            candidates.append(full)
    if not candidates:
        return None
    candidates.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return candidates[0]


def evaluate(dir_path: str) -> dict:
    """脚本对外唯一入口：接收脚本所在目录路径，返回结构化评估结果。

    - dir_path：官方约定为"脚本所在目录的路径"，脚本自行在其中定位 .docx 文档。
    - 返回：符合 §2.2 的字典。评估器崩溃/文件缺失 → status="error"。
    """
    max_score = sum(md for _, _, md, _ in DIM2_RULES if md > 0)
    try:
        file_path = _find_target_document(dir_path)
        if file_path is None:
            return {
                "id": SCRIPT_ID,
                "file_name": "",
                "status": "error",
                "error": f"未在目录中找到 .docx 文档: {dir_path}",
                "dim1_pass": False,
                "dim1_reason": "",
                "dim2_items": [],
                "total_score": 0,
                "max_score": max_score,
            }
        file_name = os.path.basename(file_path)
        evaluator = DocumentEvaluator(file_path)
        if not evaluator.evaluate_dimension1():
            return evaluator.build_result(file_name)
        evaluator.evaluate_dimension2()
        return evaluator.build_result(file_name)
    except Exception as e:
        return {
            "id": SCRIPT_ID,
            "file_name": "",
            "status": "error",
            "error": f"{type(e).__name__}: {e}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": max_score,
        }


if __name__ == "__main__":
    # 本地自测：默认使用脚本所在目录，可选覆盖为 sys.argv[1]
    target_dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    _payload = json.dumps(evaluate(target_dir), ensure_ascii=False, indent=2)
    # 不修改 sys.stdout；对 Windows 控制台的非 UTF-8 默认编码做兜底
    try:
        sys.stdout.buffer.write((_payload + "\n").encode("utf-8"))
    except AttributeError:
        print(_payload)
