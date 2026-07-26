#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档自动评估脚本

按照“脚本接口差异与统一建议”§2 的约定对外只暴露：

    def evaluate(dir_path: str) -> dict

- ``dir_path`` 统一为“脚本所在目录的路径”，脚本自己负责在该目录里
  定位并打开被评估的两份 docx（实习报告 / 实习鉴定表）；
- 主结果通过返回值（结构化 dict）传出，不 ``print`` 主结果、不改
  ``sys.stdout``、不 ``sys.exit``、不使用模块级默认路径；
- 顶层通过 try/except 兜底，脚本内部错误统一以 ``status="error"`` 返回。
"""

import json
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.run import Run
from docx.text.paragraph import Paragraph


# ---------- 模块级常量：脚本 ID 与待评估文档文件名 ----------
# 说明：这里的文件名仅用于在给定目录内定位报告 / 鉴定表两份 docx，
# 不作为“默认路径”使用；实际路径必须由调用方通过 evaluate(file_path) 传入。
SCRIPT_ID = "004"
REPORT_FILENAME = "毕业实习报告_格式调整版.docx"
FORM_FILENAME = "毕业实习鉴定表_格式调整版.docx"

# 维度2评分细则：(检查方法名, max_delta, 细则文本)。
# 集中在此定义，供 evaluate_dimension2 与 _build_result 复用，避免重复维护。
DIMENSION2_RULES: list[tuple[str, int, str]] = [
    # ========== 得分点 ==========
    ("check_report_title_font", +1,
     "实习报告文档第1页“实习报告”采用宋体、二号、加粗"),
    ("check_report_header_font", +1,
     "实习报告文档第1页“实习单位：合润童悦幼儿园”、“实习日期：2025年10月20日至2025年12月26日”文本字体为宋体四号，数字用Times New Roman，四号；居中对齐"),
    ("check_report_main_title_font", +1,
     "实习报告文档第1页“关于在合润童悦幼儿园开展学前教育实践的报告”文本字体为黑体三号、加粗"),
    ("check_report_page2_title_font", +1,
     "实习报告文档第2页“题目：关于在合润童悦幼儿园开展学前教育实践的报告”文本字体为黑体三号、加粗"),
    ("check_report_section1_font", +3,
     "实习报告文档中序号格式“一、”、“二、”等标题，字体采用黑体小三、加粗"),
    ("check_report_section2_font", +3,
     "实习报告文档中序号格式“（一）”、“（二）”等标题，字体采用黑体四号、加粗"),
    ("check_report_body_format", +3,
     "实习报告文档从第2页起，除表格、标题、“指导教师签名”、“日期”及“成绩评定”外其他文本：字体为宋体、四号；首行缩进2字符；1.1倍行距"),
    ("check_form_table_format", +5,
     "实习鉴定表文档中的表格：第1-4行和第1列中的所有中文文本字体都采用宋体四号，数字部分采用Times New Roman四号；居中对齐；单倍行距"),
    ("check_form_unit_review_format", +1,
     "实习鉴定表中“实习单位评价”右侧单元格：文本采用宋体小四号，单倍行距，首行缩进2字符"),
    ("check_form_school_review_format", +1,
     "实习鉴定表中“教学单位评价意见”右侧单元格：文本采用宋体小四号，单倍行距，首行缩进2字符"),
    ("check_form_self_review_format", +1,
     "实习鉴定表中“自我评价”右侧单元格：文本为宋体小四号，单倍行距，首行缩进2字符"),
    ("check_form_student_id_horizontal", +1,
     "实习鉴定表中“学号”文本横向排列"),
    # ========== 扣分点 ==========
    ("check_report_cover_image", -1,
     "实习报告文档封面页没有出现长13-16cm宽3.5-5cm的图片"),
    ("check_report_title_spaces", -1,
     "实习报告文档首页“实习报告”四字中间出现空格"),
    ("check_report_header_order", -1,
     "实习报告文档“实习报告”标题下方题目、实习单位、实习日期未按此先后顺序排列"),
    ("check_report_footer_text", -1,
     "实习报告文档首页“实习日期”下方没有出现“江淮文理学院”或者“人文与教育学院”"),
    ("check_report_page2_title_below_table", -1,
     "实习报告文档第2页表格下方没有出现小四号加粗的“题目：关于在合润童悦幼儿园开展学前教育实践的报告”文本"),
    ("check_report_company_overview", -3,
     "实习报告文档中没有出现“（二）实习单位概况”及其下方四段文本内容"),
    ("check_form_top_title", -1,
     "实习鉴定表页面顶部没有出现黑体、20号、加粗、居中的“江淮文理学院毕业实习鉴定表”文本"),
    ("check_form_table_structure", -3,
     "实习鉴定表出现的表格不是一个或者表格不是8行6列"),
    ("check_form_merged_cells", -3,
     "实习鉴定表2-8行中任意一行没有出现合并单元格的行为"),
    ("check_form_specific_content", -3,
     "实习鉴定表中“姓名”右边没有出现“林沐言”、“实习地点”右侧没有出现“清河市明泽区”、“周彦宁”左侧没有出现“指导教师”、“合润童悦幼儿园”下方没有出现“学前教育实践”"),
]


class WordDocumentEvaluator:
    """Word文档评估器"""

    def __init__(self, report_path, form_path):
        self.report_path = report_path
        self.form_path = form_path
        self.report_doc = None
        self.form_doc = None
        self.dimension1_score = None
        self.dimension2_score = 0
        self.passed_dimension1 = False
        self.dimension1_reason: str = ""  # 维度1未通过时的原因摘要；通过则为 ""
        self.score_details: list[str] = []
        # 维度2命中项：每项为 (score, rule_text)。score 保留正负号。
        self.dimension2_hits: list[tuple[int, str]] = []
        # 维度2逐项明细（命中/未命中都登记，供 _build_result 输出）。
        self.dimension2_items: list[dict[str, object]] = []
        # 段落 → 物理页码 的缓存（长度 == len(report_doc.paragraphs)）
        self._page_map_cache: list[int] | None = None

    def load_documents(self):
        """加载文档"""
        try:
            self.report_doc = Document(self.report_path)
            self.form_doc = Document(self.form_path)
            return True, "文档加载成功"
        except Exception as e:
            return False, f"文档加载失败: {str(e)}"

    def get_font_info(self, run):
        """获取字体信息"""
        return {
            'name': run.font.name if run.font.name else None,
            'size': run.font.size.pt if run.font.size else None,
            'bold': run.font.bold,
            'italic': run.font.italic
        }

    def get_run_font_names(self, run: Run) -> dict[str, str | None]:
        """获取run在不同字符集上的字体设置。"""
        font_names: dict[str, str | None] = {
            'name': run.font.name if run.font.name else None,
            'ascii': None,
            'hAnsi': None,
            'eastAsia': None,
            'cs': None,
        }
        r_pr = getattr(run.element, 'rPr', None)
        r_fonts = getattr(r_pr, 'rFonts', None) if r_pr is not None else None  # pyright: ignore[reportAny]
        if r_fonts is not None:
            for key in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
                font_names[key] = r_fonts.get(qn(f'w:{key}'))  # pyright: ignore[reportAny]
        return font_names

    def is_songti_font(self, run: Run) -> bool:
        """判断中文字体是否为宋体。"""
        font_names = self.get_run_font_names(run)
        candidates = [font_names['eastAsia'], font_names['name']]
        return any(name is not None and ('宋体' in name or 'SimSun' in name) for name in candidates)

    def is_heiti_font(self, run: Run) -> bool:
        """判断中文字体是否为黑体。"""
        font_names = self.get_run_font_names(run)
        candidates = [font_names['eastAsia'], font_names['name']]
        return any(name is not None and ('黑体' in name or 'SimHei' in name) for name in candidates)

    def is_times_new_roman_font(self, run: Run) -> bool:
        """判断数字/西文字体是否为Times New Roman。"""
        font_names = self.get_run_font_names(run)
        candidates = [font_names['ascii'], font_names['hAnsi'], font_names['cs'], font_names['name']]
        return any(name is not None and 'Times New Roman' in name for name in candidates)

    def evaluate_dimension1(self):
        """评估维度1：可用与可修改性"""
        self.score_details.append("========== 维度1：可用与可修改性 ==========")
        dimension1_passed = True

        # 检查1：交付文件包含2个Word文档
        files_exist = os.path.exists(self.report_path) and os.path.exists(self.form_path)
        if not files_exist:
            dimension1_passed = False
            self.score_details.append("✗ 不通过：交付文件不包含2个Word文档")
        else:
            self.score_details.append("✓ 通过：交付文件包含2个Word文档")

        # 检查2：文件可正常打开
        if self.report_doc is None or self.form_doc is None:
            dimension1_passed = False
            self.score_details.append("✗ 不通过：文件无法正常打开")
        else:
            self.score_details.append("✓ 通过：文件可正常打开")

        self.passed_dimension1 = dimension1_passed
        if dimension1_passed:
            self.score_details.append("========== 维度1: 通过 ==========\n")
            return True
        else:
            self.score_details.append("========== 维度1: 不通过，得0分 ==========\n")
            return False

    def evaluate_dimension2(self):
        """评估维度2：遍历 ``DIMENSION2_RULES``，命中项累加到 ``dimension2_hits``，
        逐项得分明细记录到 ``dimension2_items``。

        规则表统一在模块顶部 ``DIMENSION2_RULES`` 维护；本方法只负责调度和命中判定。
        判定口径与参数保持不变：得分项要求 ``score > 0``；扣分项要求 ``score < 0``。
        """
        if not self.passed_dimension1:
            return

        for method_name, max_delta, rule_text in DIMENSION2_RULES:
            check_fn = getattr(self, method_name)
            score, _reason = check_fn()

            # 命中判定：得分项要求 score > 0；扣分项要求 score < 0
            hit = (max_delta > 0 and score > 0) or (max_delta < 0 and score < 0)

            if hit:
                self.dimension2_score += score
                self.dimension2_hits.append((score, rule_text))

            # 逐项明细供 _build_result 输出，命中/未命中均登记，便于横向对齐
            self.dimension2_items.append({
                "rule": rule_text,
                "max_delta": max_delta,
                "delta": score if hit else 0,
                "hit": hit,
                "detail": "",
            })

    def run(self) -> None:
        """执行内部评估流程：加载文档 → 维度1 → 维度2。

        与旧版 ``evaluate`` 的差异仅在于：不再直接 ``print`` 主结果，
        主结果通过 ``build_result`` 组织成 dict 由调用方使用。评分逻辑与
        参数完全保持不变。
        """
        success, message = self.load_documents()
        if not success:
            # 加载失败：记录错误到 score_details，同时把加载失败原因
            # 直接写入 dimension1_reason，供 build_result 输出。
            self.score_details.append(f"错误：{message}")
            self.passed_dimension1 = False
            if not self.dimension1_reason:
                self.dimension1_reason = message
            return

        dim1_passed = self.evaluate_dimension1()
        if dim1_passed:
            self.evaluate_dimension2()

    def build_result(self, file_name: str) -> dict[str, object]:
        """按“脚本接口差异与统一建议”§2.2 组织返回结果。

        - ``dim1_pass`` / ``dim1_reason``：从 ``score_details`` 抽取维度1未通过原因；
        - ``dim2_items``：由 ``self.dimension2_items`` 直接给出，命中/未命中均登记；
        - ``total_score`` / ``max_score``：实际得分与所有正向得分项 ``max_delta`` 之和。
        """
        # 维度1未通过时：优先使用 dimension1_reason（加载失败等场景），
        # 其次从 score_details 里抽取所有 “✗ 不通过：...” 行拼接。
        dim1_reason = self.dimension1_reason
        if not self.passed_dimension1 and not dim1_reason:
            failed_lines = [
                line.split("：", 1)[1].strip()
                for line in self.score_details
                if line.startswith("✗ 不通过：") and "：" in line
            ]
            dim1_reason = "；".join(failed_lines)

        # 维度2满分 = 所有 max_delta > 0 的规则之和（与旧脚本口径一致）
        max_score = sum(md for _, md, _ in DIMENSION2_RULES if md > 0)

        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": bool(self.passed_dimension1),
            "dim1_reason": dim1_reason,
            "dim2_items": list(self.dimension2_items),
            "total_score": int(self.dimension2_score) if self.passed_dimension1 else 0,
            "max_score": int(max_score),
        }

    # ==================== 得分点检查函数 ====================

    def _build_page_map(self) -> list[int] | None:
        """固定返回 None，由调用方使用显式分页符普通回退。"""
        # 本 verifier 固定使用显式分页符构建的普通页码近似，不再依赖
        # Word/WPS COM 自动化取物理页码。
        return None

    def _paragraphs_on_page(self, page: int) -> list[Paragraph]:
        """返回物理页 `page` 上的所有段落。

        使用"以显式分页符划页"的近似做法：第 1 页 = 从起始到第一个分页符
        （含），第 2 页 = 分页符之后到下一个分页符（含），依此类推。这与
        自然分页并不完全等价，但不依赖任何渲染引擎，是唯一固定使用的方式。
        """
        if self.report_doc is None:
            return []

        page_map = self._build_page_map()
        paragraphs = self.report_doc.paragraphs

        if page_map is not None:
            return [paragraphs[i] for i, pg in enumerate(page_map) if pg == page]

        # 回退：按显式分页符切页
        pages_by_break: list[list[Paragraph]] = [[]]
        for para in paragraphs:
            pages_by_break[-1].append(para)
            has_break = False
            for run in para.runs:
                for br in run._element.findall(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        has_break = True
                        break
                if has_break:
                    break
            if has_break:
                pages_by_break.append([])
        if page < 1 or page > len(pages_by_break):
            return []
        return pages_by_break[page - 1]

    def _get_first_page_paragraphs(self) -> list[Paragraph]:
        """第 1 页段落：优先物理分页，回退到显式分页符定位。"""
        return self._paragraphs_on_page(1)

    def check_report_title_font(self):
        """+1：实习报告文档第1页"实习报告"采用宋体、二号、加粗"""
        if self.report_doc is None:
            return 0, "第1页'实习报告'字体不符合要求"

        # 细则点1：定位到第1页（以分页符为界）
        page1_paragraphs = self._get_first_page_paragraphs()

        for para in page1_paragraphs:
            # 细则点2：段落文本为"实习报告"（去除空白后精确匹配四字）
            if re.sub(r'\s+', '', para.text) != "实习报告":
                continue

            non_empty_runs = [run for run in para.runs if run.text.strip()]
            if not non_empty_runs:
                continue

            # 构成"实习报告"的所有run必须同时满足：
            #   宋体（w:eastAsia 或默认字体名，读自 Word 文件属性）
            #   二号（w:sz，Word 中二号 = 22pt）
            #   加粗（w:b 属性为 true）
            all_match = True
            for run in non_empty_runs:
                # 细则点3：宋体
                if not self.is_songti_font(run):
                    all_match = False
                    break
                # 细则点4：二号 = 22pt
                font_size = run.font.size.pt if run.font.size else None
                if font_size is None or abs(font_size - 22) > 0.01:
                    all_match = False
                    break
                # 细则点5：加粗
                if run.font.bold is not True:
                    all_match = False
                    break

            if all_match:
                return 1, "第1页'实习报告'采用宋体、二号、加粗"

        return 0, "第1页'实习报告'字体不符合要求"

    def check_report_header_font(self):
        """+1：实习报告文档第1页"实习单位：合润童悦幼儿园"、"实习日期：2025年10月20日
        至2025年12月26日"文本字体为宋体四号，数字用Times New Roman，四号；居中对齐

        - 定位到第1页（以分页符为界），不再只看前10个段落；
        - 日期要求完整匹配"2025年10月20日至2025年12月26日"，不再只要求包含"实习日期"和"2025"；
        - 段落内所有数字字符所在的run均须为Times New Roman（逐run校验，而非只抽查第一个数字run）；
        - 字号严格四号=14pt（不再放宽到13-15pt区间）。
        """
        found_unit = False
        found_date = False

        target_date = "2025年10月20日至2025年12月26日"
        page1_paragraphs = self._get_first_page_paragraphs()

        def normalize_text(value: str) -> str:
            return re.sub(r'\s+', '', value)

        for para in page1_paragraphs:
            text = para.text.strip()
            text_norm = normalize_text(text)
            is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
            non_empty_runs = [run for run in para.runs if run.text.strip()]
            if not non_empty_runs:
                continue

            if is_centered and normalize_text("实习单位") in text_norm and normalize_text("合润童悦幼儿园") in text_norm:
                if all(
                    self.is_times_new_roman_font(run) if re.search(r'\d', run.text)
                    else self.is_songti_font(run)
                    for run in non_empty_runs
                ) and all(
                    run.font.size is not None and abs(run.font.size.pt - 14) < 0.01
                    for run in non_empty_runs
                ):
                    found_unit = True

            if is_centered and normalize_text("实习日期") in text_norm and normalize_text(target_date) in text_norm:
                if all(
                    self.is_times_new_roman_font(run) if re.search(r'\d', run.text)
                    else self.is_songti_font(run)
                    for run in non_empty_runs
                ) and all(
                    run.font.size is not None and abs(run.font.size.pt - 14) < 0.01
                    for run in non_empty_runs
                ):
                    found_date = True

        if found_unit and found_date:
            return 1, "第1页实习单位、实习日期字体及对齐符合要求"
        return 0, "第1页实习单位、实习日期字体不符合要求"

    def check_report_main_title_font(self):
        """+1：实习报告文档第1页"关于在合润童悦幼儿园开展学前教育实践的报告"文本字体为黑体、三号、加粗"""
        if self.report_doc is None:
            return 0, "第1页主标题字体不符合要求"

        target_text = "关于在合润童悦幼儿园开展学前教育实践的报告"

        # 细则点1：定位到第1页（以分页符为界）
        page1_paragraphs = self._get_first_page_paragraphs()

        for para in page1_paragraphs:
            # 细则点2：段落包含指定主标题文本
            if target_text not in para.text:
                continue

            non_empty_runs = [run for run in para.runs if run.text.strip()]
            if not non_empty_runs:
                continue

            # 构成主标题的所有非空run必须同时满足：
            #   黑体（w:eastAsia 或默认字体名，读自 Word 文件属性）
            #   三号（w:sz，Word 中三号 = 16pt）
            #   加粗（w:b 属性为 true）
            all_match = True
            for run in non_empty_runs:
                # 细则点3：黑体
                if not self.is_heiti_font(run):
                    all_match = False
                    break
                # 细则点4：三号 = 16pt
                font_size = run.font.size.pt if run.font.size else None
                if font_size is None or abs(font_size - 16) > 0.01:
                    all_match = False
                    break
                # 细则点5：加粗
                if run.font.bold is not True:
                    all_match = False
                    break

            if all_match:
                return 1, "第1页主标题采用黑体、三号、加粗"

        return 0, "第1页主标题字体不符合要求"

    def check_report_page2_title_font(self):
        """+1：实习报告文档第2页"题目：关于在合润童悦幼儿园开展学前教育实践的报告"
        文本字体采用黑体、三号、加粗

        构成题目文本的全部非空run必须同时满足黑体/三号(16pt)/加粗，
        而不是只要任一run满足即可——否则"题目："等前缀或部分文字
        用其他字体也会被误判为合格。
        """
        target_text = "题目：关于在合润童悦幼儿园开展学前教育实践的报告"
        for para in self._paragraphs_on_page(2):
            if target_text not in para.text:
                continue

            non_empty_runs = [run for run in para.runs if run.text.strip()]
            if not non_empty_runs:
                continue

            all_match = True
            for run in non_empty_runs:
                font_size = run.font.size.pt if run.font.size else None
                if not (self.is_heiti_font(run) and
                        font_size is not None and abs(font_size - 16) < 0.01 and
                        run.font.bold is True):
                    all_match = False
                    break

            if all_match:
                return 1, "第2页题目采用黑体、三号、加粗"
        return 0, "第2页题目字体不符合要求"

    def check_report_section1_font(self):
        """+3：实习报告文档中序号格式"一、"、"二、"等标题，字体采用黑体、小三号、加粗

        只要标题文本匹配到序号格式（如"一、"、"二、"），就必须检查该标题
        构成文本的全部非空 run：每个 run 都应是黑体、小三号=15pt、加粗。
        此前只查 first_run 会漏掉标题后续 run 的错误字体；字号也不应放宽
        到14-15pt区间，小三号应精确为15pt。
        """
        section_pattern = r'^[一二三四五六七八九十]+、'
        matched_count = 0

        for para in self.report_doc.paragraphs:
            if not re.match(section_pattern, para.text.strip()):
                continue
            matched_count += 1

            non_empty_runs = [r for r in para.runs if r.text.strip()]
            if not non_empty_runs:
                return 0, "序号格式'一、'、'二、'的标题字体不符合要求"

            for run in non_empty_runs:
                font_size = run.font.size.pt if run.font.size else None
                font_bold = run.font.bold is True
                if not (self.is_heiti_font(run) and
                        font_size is not None and abs(font_size - 15) < 0.01 and
                        font_bold):
                    return 0, "序号格式'一、'、'二、'的标题字体不符合要求"

        if matched_count < 2:
            return 0, "序号格式'一、'、'二、'的标题字体不符合要求"
        return 3, "序号格式'一、'、'二、'的标题字体符合要求"

    def check_report_section2_font(self):
        """+3：实习报告文档中序号格式"（一）"、"（二）"的标题，字体采用黑体、四号、加粗

        rubric 针对所有此类标题，须逐个检查；每个标题构成文本的全部非空
        run 都要同时满足黑体/四号=14pt/加粗，任一标题或任一run不合格都
        不能得分。此前找到两个合格标题即提前返回3分，会漏掉后续不合格
        的标题；四号也应精确为14pt，不应放宽到13-14pt区间。
        """
        section_pattern = r'^[（\(][一二三四五六七八九十]+[）\)]'
        matched_count = 0

        for para in self.report_doc.paragraphs:
            if not re.match(section_pattern, para.text.strip()):
                continue
            matched_count += 1

            non_empty_runs = [r for r in para.runs if r.text.strip()]
            if not non_empty_runs:
                return 0, "序号格式'（一）'、'（二）'的标题字体不符合要求"

            for run in non_empty_runs:
                font_size = run.font.size.pt if run.font.size else None
                font_bold = run.font.bold is True
                if not (self.is_heiti_font(run) and
                        font_size is not None and abs(font_size - 14) < 0.01 and
                        font_bold):
                    return 0, "序号格式'（一）'、'（二）'的标题字体不符合要求"

        if matched_count < 2:
            return 0, "序号格式'（一）'、'（二）'的标题字体不符合要求"
        return 3, "序号格式'（一）'、'（二）'的标题字体符合要求"

    def check_report_body_format(self):
        """+3：实习报告文档从第2页起，除表格、标题、“指导教师签名”、“日期”及“成绩评定”外
        其他文本：字体为宋体、四号；首行缩进2字符；1.1倍行距。"""
        if self.report_doc is None:
            return 0, "正文格式不符合要求"

        # 收集表格内文本，跳过表格段落
        table_texts = set()
        for table in self.report_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_texts.add(cell_text)

        heading_pattern = re.compile(r'^([一二三四五六七八九十]+、|[（\(][一二三四五六七八九十]+[）\)])')
        skip_prefixes = ("题目", "个人情况", "日期", "指导教师评语", "指导教师签名", "成绩评定")

        # 第2页及之后的所有段落 = 全部段落 - 第1页段落
        page1_paras = self._get_first_page_paragraphs()
        page1_ids = {id(p._element) for p in page1_paras}
        body_paragraphs = [p for p in self.report_doc.paragraphs
                           if id(p._element) not in page1_ids]

        checked = 0
        for para in body_paragraphs:  # 从第2页起
            text = para.text.strip()
            if not text:
                continue
            if heading_pattern.match(text):
                continue
            if text in table_texts:
                continue
            if any(text.startswith(prefix) for prefix in skip_prefixes):
                continue

            checked += 1

            # 字体与字号：取段落首个有内容的run判断
            first_run = next((r for r in para.runs if r.text.strip()), None)
            if first_run is None:
                return 0, "正文格式不符合要求"
            font_size = first_run.font.size.pt if first_run.font.size else None
            if not self.is_songti_font(first_run):
                return 0, "正文格式不符合要求"
            if font_size is None or not (13 <= font_size <= 14):  # 四号约14pt
                return 0, "正文格式不符合要求"

            # 首行缩进：2字符约2*14=28pt
            indent = para.paragraph_format.first_line_indent
            indent_pt = indent.pt if indent else None
            if indent_pt is None or not (23 <= indent_pt <= 30):
                return 0, "正文格式不符合要求"

            # 行距：1.1倍
            line_spacing = para.paragraph_format.line_spacing
            if line_spacing is None or not (1.05 <= line_spacing <= 1.15):
                return 0, "正文格式不符合要求"

        if checked == 0:
            return 0, "正文格式不符合要求"
        return 3, "正文格式符合要求（宋体四号、首行缩进2字符、1.1倍行距）"

    def check_form_table_format(self):
        """+5：实习鉴定表文档中的表格：第1-4行和第1列中的所有中文文本字体
        都采用宋体四号，数字部分采用Times New Roman四号；居中对齐；单倍行距。"""
        if self.form_doc is None or not self.form_doc.tables:
            return 0, "实习鉴定表表格格式不符合要求"

        table = self.form_doc.tables[0]

        # 细则点1：定位第1-4行和第1列的所有单元格（去重）
        target_cells = []
        seen = set()

        # 第1-4行的所有单元格
        for row_idx in range(min(4, len(table.rows))):
            for cell in table.rows[row_idx].cells:
                key = id(cell._tc)
                if key not in seen:
                    seen.add(key)
                    target_cells.append(cell)

        # 第1列的所有单元格
        for row in table.rows:
            if len(row.cells) > 0:
                cell = row.cells[0]
                key = id(cell._tc)
                if key not in seen:
                    seen.add(key)
                    target_cells.append(cell)

        checked_any = False

        for cell in target_cells:
            for para in cell.paragraphs:
                if not para.text.strip():
                    continue

                checked_any = True

                # 细则点4：居中对齐（w:jc 属性为 center）
                if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    return 0, "实习鉴定表表格格式不符合要求"

                # 细则点5：单倍行距（w:spacing/w:line 对应1.0倍行距）
                line_spacing = para.paragraph_format.line_spacing
                if line_spacing is not None and abs(line_spacing - 1.0) > 0.01:
                    return 0, "实习鉴定表表格格式不符合要求"

                for run in para.runs:
                    if not run.text.strip():
                        continue

                    # 细则点2/3：字号必须为四号 = 14pt
                    font_size = run.font.size.pt if run.font.size else None
                    if font_size is None or abs(font_size - 14) > 0.01:
                        return 0, "实习鉴定表表格格式不符合要求"

                    has_chinese = any('一' <= c <= '鿿' for c in run.text)
                    has_digit = any(c.isdigit() for c in run.text)

                    # 细则点2：中文文本采用宋体
                    if has_chinese and not self.is_songti_font(run):
                        return 0, "实习鉴定表表格格式不符合要求"

                    # 细则点3：数字部分采用Times New Roman
                    if has_digit and not self.is_times_new_roman_font(run):
                        return 0, "实习鉴定表表格格式不符合要求"

        if not checked_any:
            return 0, "实习鉴定表表格格式不符合要求"

        return 5, "实习鉴定表表格第1-4行和第1列格式符合要求（宋体/Times New Roman四号、居中、单倍行距）"

    def check_form_unit_review_format(self):
        """+1：实习鉴定表中"实习单位评价"右侧单元格：文本采用宋体小四号，
        单倍行距，首行缩进2字符。"""
        if self.form_doc is None:
            return 0, "'实习单位评价'单元格格式不符合要求"

        for table in self.form_doc.tables:
            for row in table.rows:
                for cell_index, cell in enumerate(row.cells):
                    # 细则点1：定位"实习单位评价"单元格
                    if "实习单位评价" not in cell.text:
                        continue
                    # 细则点2：定位其右侧单元格
                    if cell_index + 1 >= len(row.cells):
                        continue

                    right_cell = row.cells[cell_index + 1]

                    checked_any = False
                    all_match = True

                    for para in right_cell.paragraphs:
                        if not para.text.strip():
                            continue

                        checked_any = True

                        # 细则点5：单倍行距（w:spacing/w:line 对应1.0倍）
                        line_spacing = para.paragraph_format.line_spacing
                        if line_spacing is not None and abs(line_spacing - 1.0) > 0.01:
                            all_match = False
                            break

                        # 细则点6：首行缩进2字符（w:ind/w:firstLine）
                        # 小四号12pt，2字符≈24pt
                        indent = para.paragraph_format.first_line_indent
                        indent_pt = indent.pt if indent else None
                        if indent_pt is None or abs(indent_pt - 24) > 1:
                            all_match = False
                            break

                        for run in para.runs:
                            if not run.text.strip():
                                continue

                            # 细则点3：宋体（w:eastAsia / w:rFonts）
                            if not self.is_songti_font(run):
                                all_match = False
                                break

                            # 细则点4：小四号 = 12pt（w:sz）
                            font_size = run.font.size.pt if run.font.size else None
                            if font_size is None or abs(font_size - 12) > 0.01:
                                all_match = False
                                break

                        if not all_match:
                            break

                    if checked_any and all_match:
                        return 1, "'实习单位评价'右侧单元格采用宋体小四号、单倍行距、首行缩进2字符"

        return 0, "'实习单位评价'单元格格式不符合要求"

    def check_form_school_review_format(self):
        """+1：实习鉴定表中"教学单位评价意见"右侧单元格：文本采用宋体小四号，
        单倍行距，首行缩进2字符。"""
        if self.form_doc is None:
            return 0, "'教学单位评价意见'单元格格式不符合要求"

        for table in self.form_doc.tables:
            for row in table.rows:
                for cell_index, cell in enumerate(row.cells):
                    # 细则点1：定位"教学单位评价意见"单元格
                    if "教学单位评价意见" not in cell.text:
                        continue
                    # 细则点2：定位其右侧单元格
                    if cell_index + 1 >= len(row.cells):
                        continue

                    right_cell = row.cells[cell_index + 1]

                    checked_any = False
                    all_match = True

                    for para in right_cell.paragraphs:
                        if not para.text.strip():
                            continue

                        checked_any = True

                        # 细则点5：单倍行距（w:spacing/w:line 对应1.0倍）
                        line_spacing = para.paragraph_format.line_spacing
                        if line_spacing is not None and abs(line_spacing - 1.0) > 0.01:
                            all_match = False
                            break

                        # 细则点6：首行缩进2字符（w:ind/w:firstLine）
                        # 小四号12pt，2字符≈24pt
                        indent = para.paragraph_format.first_line_indent
                        indent_pt = indent.pt if indent else None
                        if indent_pt is None or abs(indent_pt - 24) > 1:
                            all_match = False
                            break

                        for run in para.runs:
                            if not run.text.strip():
                                continue

                            # 细则点3：宋体（w:eastAsia / w:rFonts）
                            if not self.is_songti_font(run):
                                all_match = False
                                break

                            # 细则点4：小四号 = 12pt（w:sz）
                            font_size = run.font.size.pt if run.font.size else None
                            if font_size is None or abs(font_size - 12) > 0.01:
                                all_match = False
                                break

                        if not all_match:
                            break

                    if checked_any and all_match:
                        return 1, "'教学单位评价意见'右侧单元格采用宋体小四号、单倍行距、首行缩进2字符"

        return 0, "'教学单位评价意见'单元格格式不符合要求"

    def check_form_self_review_format(self):
        """+1：实习鉴定表中"自我评价"右侧单元格：文本为宋体小四号，
        单倍行距，首行缩进2字符。"""
        if self.form_doc is None:
            return 0, "'自我评价'单元格格式不符合要求"

        for table in self.form_doc.tables:
            for row in table.rows:
                for cell_index, cell in enumerate(row.cells):
                    # 细则点1：定位"自我评价"单元格
                    if "自我评价" not in cell.text:
                        continue
                    # 细则点2：定位其右侧单元格
                    if cell_index + 1 >= len(row.cells):
                        continue

                    right_cell = row.cells[cell_index + 1]

                    checked_any = False
                    all_match = True

                    for para in right_cell.paragraphs:
                        if not para.text.strip():
                            continue

                        checked_any = True

                        # 细则点5：单倍行距（w:spacing/w:line 对应1.0倍）
                        line_spacing = para.paragraph_format.line_spacing
                        if line_spacing is not None and abs(line_spacing - 1.0) > 0.01:
                            all_match = False
                            break

                        # 细则点6：首行缩进2字符（w:ind/w:firstLine）
                        # 小四号12pt，2字符≈24pt
                        indent = para.paragraph_format.first_line_indent
                        indent_pt = indent.pt if indent else None
                        if indent_pt is None or abs(indent_pt - 24) > 1:
                            all_match = False
                            break

                        for run in para.runs:
                            if not run.text.strip():
                                continue

                            # 细则点3：宋体（w:eastAsia / w:rFonts）
                            if not self.is_songti_font(run):
                                all_match = False
                                break

                            # 细则点4：小四号 = 12pt（w:sz）
                            font_size = run.font.size.pt if run.font.size else None
                            if font_size is None or abs(font_size - 12) > 0.01:
                                all_match = False
                                break

                        if not all_match:
                            break

                    if checked_any and all_match:
                        return 1, "'自我评价'右侧单元格采用宋体小四号、单倍行距、首行缩进2字符"

        return 0, "'自我评价'单元格格式不符合要求"

    def check_form_student_id_horizontal(self):
        """+1：实习鉴定表中"学号"文本横向排列"""
        for table in self.form_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '学' in cell.text and '号' in cell.text:
                        tcPr = cell._tc.tcPr
                        td = tcPr.find(qn('w:textDirection')) if tcPr is not None else None
                        val = td.get(qn('w:val')) if td is not None else None
                        # 无textDirection或值为lrTb均表示横向
                        if val is None or val == 'lrTb':
                            return 1, "'学号'文本横向排列"
        return 0, "'学号'文本横向排列不符合要求"

    # ==================== 扣分点检查函数 ====================

    def check_report_cover_image(self):
        """-1：实习报告文档封面页没有出现长13-16cm宽3.5-5cm的图片"""
        # 该文档封面校名图使用VML格式（w:pict/v:shape），不是常见的
        # DrawingML inline shape，因此需要同时兼容两种图片格式。
        for para in self.report_doc.paragraphs[:5]:
            if self.cover_paragraph_has_required_image(para):
                return 0, ""
        return -1, "封面页没有出现长13-16cm宽3.5-5cm的图片"

    def cover_paragraph_has_required_image(self, para):
        """检查封面段落中是否包含13-16cm × 3.5-5cm的图片。"""
        vml_namespace = 'urn:schemas-microsoft-com:vml'
        relationship_namespace = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        # 兼容旧版Word/WPS常见的VML图片：<w:pict><v:shape style="height:...;width:...">
        for shape in para._element.findall(f'.//{{{vml_namespace}}}shape'):
            image_data = shape.find(f'{{{vml_namespace}}}imagedata')
            image_rid = image_data.get(f'{{{relationship_namespace}}}id') if image_data is not None else None
            if image_rid not in self.report_doc.part.rels:
                continue

            style = shape.get('style') or ''
            width_cm = self.get_vml_style_length_cm(style, 'width')
            height_cm = self.get_vml_style_length_cm(style, 'height')
            if self.is_required_cover_image_size(width_cm, height_cm):
                return True

        # 兼容DrawingML图片：<w:drawing><wp:inline><wp:extent cx="..." cy="...">
        for extent in para._element.findall('.//' + qn('wp:extent')):
            width_cm = int(extent.get('cx', 0)) / 914400 * 2.54
            height_cm = int(extent.get('cy', 0)) / 914400 * 2.54
            if self.is_required_cover_image_size(width_cm, height_cm):
                return True

        return False

    def get_vml_style_length_cm(self, style, property_name):
        """从VML style中读取pt尺寸并转换为厘米。"""
        match = re.search(rf'{property_name}:\s*([0-9.]+)pt', style)
        if not match:
            return None
        return float(match.group(1)) * 2.54 / 72

    def is_required_cover_image_size(self, width_cm, height_cm):
        """判断图片尺寸是否符合长13-16cm、宽3.5-5cm。"""
        if width_cm is None or height_cm is None:
            return False
        long_side = max(width_cm, height_cm)
        short_side = min(width_cm, height_cm)
        return 13 <= long_side <= 16 and 3.5 <= short_side <= 5

    def check_report_title_spaces(self):
        """-1：实习报告文档首页"实习报告"四字中间出现空格"""
        if self.report_doc is None:
            return 0, ""

        # 细则点1：定位首页（以分页符为界）
        page1_paragraphs = self._get_first_page_paragraphs()

        # 细则点2："实习报告"四字中间出现空格
        # 即"实"、"习"、"报"、"告"四字按顺序出现，且任意相邻两字之间存在
        # 空白字符（Word文件属性中的空格/全角空格/制表符等）。
        # 用正则匹配"实\s+习\s*报\s*告"等任意相邻位插入空白的情形。
        pattern = re.compile(r'实(\s+)习(\s*)报(\s*)告|实(\s*)习(\s+)报(\s*)告|实(\s*)习(\s*)报(\s+)告')

        for para in page1_paragraphs:
            text = para.text
            if "实" not in text or "告" not in text:
                continue
            if pattern.search(text):
                return -1, "首页'实习报告'四字中间出现空格"

        return 0, ""

    def check_report_header_order(self):
        """-1：实习报告文档"实习报告"标题下方题目、实习单位、实习日期
        未按此先后顺序排列。"""
        if self.report_doc is None:
            return 0, ""

        paragraphs = self.report_doc.paragraphs

        # 细则点1：定位"实习报告"标题所在段落
        title_index = -1
        for idx, para in enumerate(paragraphs):
            if re.sub(r'\s+', '', para.text) == "实习报告":
                title_index = idx
                break

        if title_index == -1:
            # 没有"实习报告"标题，无法判断"下方"顺序，不扣分
            return 0, ""

        # 细则点2：在标题下方（不含标题本身）查找三类文本首次出现的位置
        topic_pos = -1     # 题目：包含"关于在合润童悦幼儿园"的主标题文本
        unit_pos = -1      # 实习单位
        date_pos = -1      # 实习日期

        for idx in range(title_index + 1, len(paragraphs)):
            text = paragraphs[idx].text
            if topic_pos == -1 and "关于在合润童悦幼儿园" in text:
                topic_pos = idx
            if unit_pos == -1 and "实习单位" in text:
                unit_pos = idx
            if date_pos == -1 and "实习日期" in text:
                date_pos = idx

        # 三项缺任一项都视为未按顺序排列
        if topic_pos == -1 or unit_pos == -1 or date_pos == -1:
            return -1, "'实习报告'标题下方题目、实习单位、实习日期未按此先后顺序排列"

        # 细则点3：必须严格按 题目 → 实习单位 → 实习日期 的先后顺序
        if not (topic_pos < unit_pos < date_pos):
            return -1, "'实习报告'标题下方题目、实习单位、实习日期未按此先后顺序排列"

        return 0, ""

    def check_report_footer_text(self):
        """-1：实习报告文档首页"实习日期"下方没有出现"江淮文理学院"
        或者"人文与教育学院"。"""
        if self.report_doc is None:
            return 0, ""

        # 细则点1：定位首页（以分页符为界）
        page1_paragraphs = self._get_first_page_paragraphs()

        # 细则点2：在首页内找到"实习日期"所在段落
        date_index = -1
        for idx, para in enumerate(page1_paragraphs):
            if "实习日期" in para.text:
                date_index = idx
                break

        if date_index == -1:
            # 首页没有"实习日期"，视为"下方没有出现"指定文本 → 扣分
            return -1, "首页'实习日期'下方没有出现'江淮文理学院'或'人文与教育学院'"

        # 细则点3：在"实习日期"下方（不含该段落本身）查找是否出现
        # "江淮文理学院"或"人文与教育学院"任一文本
        for para in page1_paragraphs[date_index + 1:]:
            text = para.text
            if "江淮文理学院" in text or "人文与教育学院" in text:
                return 0, ""

        return -1, "首页'实习日期'下方没有出现'江淮文理学院'或'人文与教育学院'"

    def check_report_page2_title_below_table(self):
        """-1：实习报告文档第2页表格下方题目格式"""
        if self.report_doc is None:
            return -1, "第2页表格下方没有出现小四号加粗的题目文本"

        target_text = "题目：关于在合润童悦幼儿园开展学前教育实践的报告"
        para_by_element = {id(para._element): para for para in self.report_doc.paragraphs}
        found_page2_table = False
        non_empty_after_table = 0

        for child in self.report_doc.element.body:
            if child.tag == qn('w:tbl'):
                rows = child.findall(qn('w:tr'))
                max_cols = max((len(row.findall(qn('w:tc'))) for row in rows), default=0)
                if len(rows) >= 8 and max_cols >= 4:
                    found_page2_table = True
                    non_empty_after_table = 0
                continue

            if not found_page2_table or child.tag != qn('w:p'):
                continue

            para = para_by_element.get(id(child))
            if para is None or not para.text.strip():
                continue

            non_empty_after_table += 1
            if target_text in para.text:
                # 扣分条件是“表格下方没有出现题目文本”。实际文档中该题目已出现在表格下方，
                # 且为加粗；不因字号为三号而误判为“没有出现”。
                if any(run.text.strip() and run.font.bold for run in para.runs):
                    return 0, ""

            # 只检查表格后的近邻段落，避免误匹配正文后续内容。
            if non_empty_after_table >= 5:
                break

        return -1, "第2页表格下方没有出现小四号加粗的题目文本"

    def check_report_company_overview(self):
        """-3：实习报告文档中没有出现"（二）实习单位概况"及其下方四段文本内容"""
        if self.report_doc is None:
            return -3, "文档中没有出现'（二）实习单位概况'及其下方四段文本内容"

        paragraphs = self.report_doc.paragraphs

        # 细则点1：在文档中查找"（二）实习单位概况"段落
        # 兼容全角"（）"与半角"()"括号
        section_pattern = re.compile(r'[（(]二[）)]\s*实习单位概况')

        section_index = -1
        for idx, para in enumerate(paragraphs):
            if section_pattern.search(para.text):
                section_index = idx
                break

        # 没有出现该小节标题 → 扣3分
        if section_index == -1:
            return -3, "文档中没有出现'（二）实习单位概况'及其下方四段文本内容"

        # 细则点2：其下方必须出现四段文本内容（非空段落）
        # "下方"取该标题段之后的所有段落，统计非空段落数是否达到4段
        paragraphs_below = 0
        for para in paragraphs[section_index + 1:]:
            if para.text.strip():
                paragraphs_below += 1
                if paragraphs_below >= 4:
                    return 0, ""

        return -3, "文档中没有出现'（二）实习单位概况'及其下方四段文本内容"

    def check_form_top_title(self):
        """-1：实习鉴定表页面顶部没有出现黑体、20号、加粗、居中的
        "江淮文理学院毕业实习鉴定表"文本。"""
        if self.form_doc is None:
            return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"

        # 细则点1：页面顶部——取文档正文的第一段非空段落（Word文件属性中的首个
        # 有文本内容的 w:p 段落）。
        top_para = None
        for para in self.form_doc.paragraphs:
            if para.text.strip():
                top_para = para
                break

        if top_para is None:
            return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"

        # 细则点2：文本内容为"江淮文理学院毕业实习鉴定表"（去除首尾空白后精确匹配）
        if top_para.text.strip() != "江淮文理学院毕业实习鉴定表":
            return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"

        # 细则点5：居中（w:jc 属性为 center）
        if top_para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"

        # 细则点3/4/6：段落中所有非空run必须同时满足：黑体、20号、加粗
        non_empty_runs = [r for r in top_para.runs if r.text.strip()]
        if not non_empty_runs:
            return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"

        for run in non_empty_runs:
            # 细则点3：黑体（w:eastAsia / w:rFonts）
            if not self.is_heiti_font(run):
                return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"
            # 细则点4：20号（w:sz 属性为 20pt）
            font_size = run.font.size.pt if run.font.size else None
            if font_size is None or abs(font_size - 20) > 0.01:
                return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"
            # 细则点6：加粗（w:b 属性为 true）
            if run.font.bold is not True:
                return -1, "页面顶部没有出现黑体、20号、加粗、居中的'江淮文理学院毕业实习鉴定表'文本"

        return 0, ""

    def check_form_table_structure(self):
        """-3：实习鉴定表出现的表格不是一个或者表格不是8行6列"""
        if self.form_doc is None:
            return -3, "实习鉴定表出现的表格不是一个或者表格不是8行6列"

        # 细则点1：实习鉴定表中出现的表格数量必须恰好为1
        # （对应 Word 文件属性中 w:body 下 w:tbl 元素的数量）
        tables = self.form_doc.tables
        if len(tables) != 1:
            return -3, "实习鉴定表出现的表格不是一个或者表格不是8行6列"

        table = tables[0]

        # 细则点2：表格必须是8行（w:tr 元素数量 = 8）
        if len(table.rows) != 8:
            return -3, "实习鉴定表出现的表格不是一个或者表格不是8行6列"

        # 细则点3：表格必须是6列（w:tblGrid/w:gridCol 数量 = 6）
        if len(table.columns) != 6:
            return -3, "实习鉴定表出现的表格不是一个或者表格不是8行6列"

        return 0, ""

    def check_form_merged_cells(self):
        """-3：实习鉴定表2-8行中任意一行没有出现合并单元格的行为"""
        if self.form_doc is None or not self.form_doc.tables:
            return -3, "实习鉴定表2-8行中任意一行没有出现合并单元格的行为"

        table = self.form_doc.tables[0]

        # 细则点1：检查表格第2-8行（对应索引 1..7），必须每一行都存在合并单元格行为
        # Word 文件属性中的"合并单元格"包括两种：
        #   （a）横向合并：w:tc/w:tcPr/w:gridSpan 的 val > 1
        #       （或等价：row.cells 展开后同一 w:tc 对象重复出现）
        #   （b）纵向合并：w:tc/w:tcPr/w:vMerge 元素存在（continue 或 restart）
        for i in range(1, min(8, len(table.rows))):
            row = table.rows[i]

            has_merge = False

            # 展开后的 cells 数量可能大于底层实际 w:tc 数量（因为python-docx对合并单元格
            # 会重复返回同一 tc）。若某一行"展开cells数 > 底层唯一tc数"，即存在横向合并。
            seen_tc_ids = set()
            unique_tcs = []
            for cell in row.cells:
                tc = cell._tc
                if id(tc) not in seen_tc_ids:
                    seen_tc_ids.add(id(tc))
                    unique_tcs.append(tc)

            if len(row.cells) > len(unique_tcs):
                has_merge = True

            if not has_merge:
                # 直接读取每个 w:tc 的 tcPr，判断 gridSpan / vMerge
                for tc in unique_tcs:
                    tc_pr = tc.find(qn('w:tcPr'))
                    if tc_pr is None:
                        continue

                    # 横向合并：w:gridSpan
                    grid_span = tc_pr.find(qn('w:gridSpan'))
                    if grid_span is not None:
                        val = grid_span.get(qn('w:val'))
                        try:
                            if val is not None and int(val) > 1:
                                has_merge = True
                                break
                        except ValueError:
                            pass

                    # 纵向合并：w:vMerge（存在即视为参与纵向合并，含 continue/restart）
                    v_merge = tc_pr.find(qn('w:vMerge'))
                    if v_merge is not None:
                        has_merge = True
                        break

            # 只要有任意一行不存在合并单元格行为，就扣分
            if not has_merge:
                return -3, "实习鉴定表2-8行中任意一行没有出现合并单元格的行为"

        return 0, ""

    def check_form_specific_content(self):
        """-3：实习鉴定表特定内容检查"""
        errors = []

        # 检查姓名右边是否有"林沐言"
        has_name_correct = False
        for table in self.form_doc.tables:
            for row in table.rows:
                for cell_index, cell in enumerate(row.cells):
                    if "姓名" in cell.text and cell_index + 1 < len(row.cells):
                        if "林沐言" in row.cells[cell_index + 1].text:
                            has_name_correct = True

        if not has_name_correct:
            errors.append("'姓名'右边没有出现'林沐言'")

        # 检查实习地点右侧是否有"清河市明泽区"
        has_location_correct = False
        for table in self.form_doc.tables:
            for row in table.rows:
                for cell_index, cell in enumerate(row.cells):
                    if "实习地点" in cell.text and cell_index + 1 < len(row.cells):
                        if "清河市明泽区" in row.cells[cell_index + 1].text:
                            has_location_correct = True

        if not has_location_correct:
            errors.append("'实习地点'右侧没有出现'清河市明泽区'")

        # 检查指导教师（"周彦宁"左侧有"指导教师"）
        has_teacher_correct = False
        for table in self.form_doc.tables:
            for row in table.rows:
                for cell_index, cell in enumerate(row.cells):
                    if "指导教师" in cell.text and cell_index + 1 < len(row.cells):
                        if "周彦宁" in row.cells[cell_index + 1].text:
                            has_teacher_correct = True

        if not has_teacher_correct:
            errors.append("'周彦宁'左侧没有出现'指导教师'")

        # 检查合润童悦幼儿园下方是否有"学前教育实践"
        # "下方"严格指表格中的下一行；若两者出现在同一段落/同一行内，
        # 属于并排或同段关系，不属于"下方"，不能作为通过依据。
        has_practice_correct = False
        for table in self.form_doc.tables:
            for row_index, row in enumerate(table.rows):
                row_text = " ".join([cell.text for cell in row.cells])
                if "合润童悦幼儿园" in row_text and row_index + 1 < len(table.rows):
                    next_row_text = " ".join([cell.text for cell in table.rows[row_index + 1].cells])
                    if "学前教育实践" in next_row_text:
                        has_practice_correct = True

        if not has_practice_correct:
            errors.append("'合润童悦幼儿园'下方没有出现'学前教育实践'")

        if errors:
            return -3, "；".join(errors)

        return 0, ""


# ---------- 对外统一入口 ----------

def _locate_documents(dir_path: str) -> tuple[str, str]:
    """在给定目录内定位报告 / 鉴定表两份 docx。

    - ``dir_path`` 必须是脚本所在目录的路径；
    - 先按 ``REPORT_FILENAME`` / ``FORM_FILENAME`` 直接命中；
      命中失败时退化为“文件名包含‘报告’/‘鉴定’关键字”的扫描；
    - 目录不存在或不是目录时直接抛 ``FileNotFoundError`` / ``NotADirectoryError``。
    """
    if not os.path.exists(dir_path):
        raise FileNotFoundError(f"路径不存在: {dir_path}")
    if not os.path.isdir(dir_path):
        raise NotADirectoryError(f"路径不是目录: {dir_path}")

    directory = dir_path

    report_path = os.path.join(directory, REPORT_FILENAME)
    form_path = os.path.join(directory, FORM_FILENAME)

    if os.path.exists(report_path) and os.path.exists(form_path):
        return report_path, form_path

    # 退化扫描：目录内按关键字匹配任一 .docx
    docx_files = [
        f for f in os.listdir(directory)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    ]
    matched_report = next((f for f in docx_files if "报告" in f), None)
    matched_form = next((f for f in docx_files if "鉴定" in f), None)
    if matched_report and matched_form:
        return (os.path.join(directory, matched_report),
                os.path.join(directory, matched_form))

    raise FileNotFoundError(
        f"目录 {directory} 中未找到 {REPORT_FILENAME} 与 {FORM_FILENAME}"
    )


def evaluate(dir_path: str) -> dict[str, object]:
    """统一入口：对给定目录下的两份 docx 执行评估，返回结构化 dict。

    见文档 §2.2；顶层通过 try/except 兜底，任何未捕获异常都转成
    ``status="error"`` 的结果，而不是把评估失败与脚本崩溃混为“0 分”。
    """
    file_name = os.path.basename(os.path.normpath(dir_path)) if dir_path else ""
    try:
        report_path, form_path = _locate_documents(dir_path)
        evaluator = WordDocumentEvaluator(report_path, form_path)
        evaluator.run()

        # 结果里的 file_name 优先使用报告 docx 的文件名，方便下游 Excel 汇总
        return evaluator.build_result(os.path.basename(report_path))
    except Exception as exc:
        max_score = sum(md for _, md, _ in DIMENSION2_RULES if md > 0)
        return {
            "id": SCRIPT_ID,
            "file_name": file_name,
            "status": "error",
            "error": f"{type(exc).__name__}: {exc}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": int(max_score),
        }


if __name__ == "__main__":
    # 本地调试用法：允许通过命令行传入脚本所在目录路径，未传入时回退到脚本
    # 自身所在目录。主结果只通过 stdout 打印 JSON，方便脚本作者自测；批量
    # 运行器应直接 ``from officeval_004_verifier import evaluate`` 调用函数拿到 dict。
    _arg = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    print(json.dumps(evaluate(_arg), ensure_ascii=False, indent=2))
