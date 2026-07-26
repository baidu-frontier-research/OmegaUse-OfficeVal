"""
自动评估脚本：对《青岭山居旅居营造计划_格式修改版.docx》按给定打分细则进行自动评分。

评估逻辑：
- 维度1（可用与可修改性）：若不满足，直接判为 0 分，不再检查维度2。
- 维度2（完成度）：逐条匹配得分点/扣分点，累加分数。

所有评分点都通过解析 document.xml / 页眉页脚 / 表格 / 图片 等结构自动判断，
不依赖人工。
"""

from __future__ import annotations

import json

SCRIPT_ID = "025"
import math
import os
import re
import sys
import traceback
import zipfile
from typing import List, Dict, Tuple, Optional

import docx
from docx.oxml.ns import qn
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W_NS, "r": R_NS}


# --------------------------------------------------------------------------------------
# 工具函数
# --------------------------------------------------------------------------------------

def read_doc_xml(src: str) -> Tuple[bytes, dict, dict, list, bytes, bytes]:
    """从 docx 文件读取 document.xml、rels、footers、styles、theme。"""
    with zipfile.ZipFile(src) as z:
        names = z.namelist()
        with z.open("word/document.xml") as f:
            doc_xml = f.read()
        rels = {}
        if "word/_rels/document.xml.rels" in names:
            with z.open("word/_rels/document.xml.rels") as f:
                rels_root = etree.fromstring(f.read())
                for r in rels_root:
                    rels[r.get("Id")] = (r.get("Type", "").split("/")[-1], r.get("Target"))
        footers = {}
        for n in names:
            if n.startswith("word/footer") and n.endswith(".xml"):
                with z.open(n) as f:
                    footers[n] = f.read()
        styles_xml = b""
        if "word/styles.xml" in names:
            with z.open("word/styles.xml") as f:
                styles_xml = f.read()
        theme_xml = b""
        if "word/theme/theme1.xml" in names:
            with z.open("word/theme/theme1.xml") as f:
                theme_xml = f.read()
    return doc_xml, rels, footers, names, styles_xml, theme_xml


def para_text(p_el) -> str:
    """从 w:p 元素中抽取所有 w:t 文本。"""
    texts = p_el.findall(".//w:t", NS)
    return "".join((t.text or "") for t in texts)


def get_para_run_fonts(p_el, theme_fonts=None, default_east_font=None) -> List[dict]:
    """获取段落中所有 run 的字体信息（rFonts/sz/b/i 等）。

    - bold：仅当 w:b 存在且 val 不为 0/false/off/none 时为 True（办公软件语义：
      <w:b/> 或 <w:b w:val="1"/> 表示加粗，<w:b w:val="0"/> 表示显式取消加粗）。
    - has_text：该 run 是否含有可见文本（w:t），用于区分标题真实文字与空 run。
    - east_font：中文字体，按办公软件的解析优先级取值：
        1) run 直接写 w:eastAsia；
        2) run 用主题引用 w:eastAsiaTheme（用 theme_fonts 解析为真实字体名）；
        3) 都没有时，回退到文档默认中文字体 default_east_font
           （docDefaults/Normal 样式继承而来，办公软件即按此渲染）。
    """
    theme_fonts = theme_fonts or {}
    out = []
    for r in p_el.findall(".//w:r", NS):
        rpr = r.find("w:rPr", NS)
        info = {"bold": False, "italic": False, "size_half_pt": None,
                "ascii_font": None, "east_font": None, "has_text": False}
        if rpr is not None:
            b = rpr.find("w:b", NS)
            if b is not None and b.get(qn("w:val")) not in ("0", "false", "off", "none"):
                info["bold"] = True
            if rpr.find("w:i", NS) is not None:
                info["italic"] = True
            sz = rpr.find("w:sz", NS)
            if sz is not None and sz.get(qn("w:val")) is not None:
                info["size_half_pt"] = int(sz.get(qn("w:val")))
            rf = rpr.find("w:rFonts", NS)
            if rf is not None:
                info["ascii_font"] = rf.get(qn("w:ascii"))
                east = rf.get(qn("w:eastAsia"))
                if not east:
                    # 主题字体引用：eastAsiaTheme -> 真实中文字体
                    theme_ref = rf.get(qn("w:eastAsiaTheme"))
                    if theme_ref:
                        east = theme_fonts.get(theme_ref)
                info["east_font"] = east
        # run 上没有显式/主题字体时，回退到文档默认中文字体（继承）
        if not info["east_font"]:
            info["east_font"] = default_east_font
        info["has_text"] = any((t.text or "") for t in r.findall("w:t", NS))
        out.append(info)
    return out


def get_para_spacing(p_el) -> dict:
    """获取段落级 spacing（行距 line / 段前 before / 段后 after / beforeLines / afterLines）。"""
    info = {"line": None, "line_rule": None, "before": None, "after": None,
            "before_lines": None, "after_lines": None}
    ppr = p_el.find("w:pPr", NS)
    if ppr is None:
        return info
    sp = ppr.find("w:spacing", NS)
    if sp is not None:
        for k in ("line", "before", "after"):
            v = sp.get(qn(f"w:{k}"))
            if v is not None:
                info[k] = int(v)
        info["line_rule"] = sp.get(qn("w:lineRule"))
        bl = sp.get(qn("w:beforeLines"))
        if bl is not None:
            info["before_lines"] = int(bl)
        al = sp.get(qn("w:afterLines"))
        if al is not None:
            info["after_lines"] = int(al)
    return info


def get_para_indent(p_el) -> dict:
    """获取段落缩进（首行 leftChars/firstLineChars）。"""
    info = {"first_line_chars": None, "first_line": None, "left": None}
    ppr = p_el.find("w:pPr", NS)
    if ppr is None:
        return info
    ind = ppr.find("w:ind", NS)
    if ind is not None:
        flc = ind.get(qn("w:firstLineChars"))
        if flc is not None:
            try:
                info["first_line_chars"] = float(flc)
            except ValueError:
                pass
        fl = ind.get(qn("w:firstLine"))
        if fl is not None:
            try:
                info["first_line"] = int(fl)
            except ValueError:
                pass
        lf = ind.get(qn("w:left"))
        if lf is not None:
            try:
                info["left"] = int(lf)
            except ValueError:
                pass
    return info


def get_para_alignment(p_el) -> Optional[str]:
    ppr = p_el.find("w:pPr", NS)
    if ppr is None:
        return None
    jc = ppr.find("w:jc", NS)
    return jc.get(qn("w:val")) if jc is not None else None


def _lcs(a: str, b: str) -> str:
    """最长公共子序列（用于标题下方文本相似度判断）。"""
    m, n = len(a), len(b)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    for i in range(m):
        for j in range(n):
            if a[i] == b[j]:
                dp[i + 1][j + 1] = dp[i][j] + 1
            else:
                dp[i + 1][j + 1] = max(dp[i + 1][j], dp[i][j + 1])
    i, j = m, n
    res = []
    while i > 0 and j > 0:
        if a[i - 1] == b[j - 1]:
            res.append(a[i - 1])
            i -= 1
            j -= 1
        elif dp[i - 1][j] >= dp[i][j - 1]:
            i -= 1
        else:
            j -= 1
    return "".join(reversed(res))


# --------------------------------------------------------------------------------------
# 主类
# --------------------------------------------------------------------------------------

class Evaluator:
    def __init__(self, path: str):
        self.path = path
        self.doc_xml_bytes, self.rels, self.footers, self.zip_names, self.styles_xml_bytes, self.theme_xml_bytes = read_doc_xml(path)
        self.root = etree.fromstring(self.doc_xml_bytes)
        self.styles_root = etree.fromstring(self.styles_xml_bytes) if self.styles_xml_bytes else None
        self.theme_fonts = self._parse_theme_fonts(self.theme_xml_bytes)
        self.default_east_font = self._parse_default_east_font()
        self.body = self.root.find("w:body", NS)
        self.body_children = list(self.body)
        self.paras: List = []
        self.tables: List = []
        for ch in self.body_children:
            tag = ch.tag.split("}")[-1]
            if tag == "p":
                self.paras.append(ch)
            elif tag == "tbl":
                self.tables.append(ch)
        self.hit_results: List[str] = []
        self.dimension1_pass: bool = True
        self.dimension1_reasons: List[str] = []
        self.score: int = 0

    @staticmethod
    def _parse_theme_fonts(theme_xml_bytes: bytes) -> Dict[str, str]:
        """解析 theme1.xml，得到主题字体引用名 -> 真实中文字体名 的映射。

        Word/WPS 里 rFonts 常以主题引用指定字体（如 eastAsiaTheme="minorEastAsia"），
        真实字体在 theme1.xml 的 fontScheme 中定义：majorFont/minorFont 下
        <a:ea> 或按 <a:font script="Hans"> 给出中文字体。这里建立：
          majorEastAsia / minorEastAsia -> 对应中文(Hans)字体名。
        """
        mapping: Dict[str, str] = {}
        if not theme_xml_bytes:
            return mapping
        A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
        try:
            root = etree.fromstring(theme_xml_bytes)
        except Exception:
            return mapping
        for kind, key in (("majorFont", "majorEastAsia"), ("minorFont", "minorEastAsia")):
            fs = root.find(".//" + A + kind)
            if fs is None:
                continue
            font = None
            ea = fs.find(A + "ea")
            if ea is not None and ea.get("typeface"):
                font = ea.get("typeface")
            if not font:
                for f in fs.findall(A + "font"):
                    if f.get("script") == "Hans" and f.get("typeface"):
                        font = f.get("typeface")
                        break
            if font:
                mapping[key] = font
        return mapping

    def _parse_default_east_font(self) -> Optional[str]:
        """解析文档默认中文字体（供 run 无显式字体时回退）。

        取值优先级（贴合办公软件继承链）：
          1) styles.xml 的 docDefaults/rPrDefault/rFonts 的 w:eastAsia；
             若是主题引用 w:eastAsiaTheme，则用 theme_fonts 解析。
          2) 默认段落样式(Normal, w:default=1)的 rPr/rFonts 同样处理。
        """
        if self.styles_root is None:
            return None

        def font_from_rfonts(rf) -> Optional[str]:
            if rf is None:
                return None
            east = rf.get(qn("w:eastAsia"))
            if east:
                return east
            ref = rf.get(qn("w:eastAsiaTheme"))
            if ref:
                return self.theme_fonts.get(ref)
            return None

        # 1) docDefaults
        dd = self.styles_root.find("w:docDefaults", NS)
        if dd is not None:
            f = font_from_rfonts(dd.find(".//w:rPrDefault//w:rFonts", NS))
            if f:
                return f
        # 2) 默认段落样式(Normal)
        for s in self.styles_root.findall("w:style", NS):
            if s.get(qn("w:type")) == "paragraph" and s.get(qn("w:default")) == "1":
                f = font_from_rfonts(s.find(".//w:rPr//w:rFonts", NS))
                if f:
                    return f
        return None

    # ---------- 维度 1 ----------
    def check_dimension1(self) -> bool:
        """可用与可修改性。"""
        full_text = "".join(para_text(p) for p in self.paras)
        has_cover = bool(re.search(r"栖云青岭|乡村旅居|传统村落|营造计划", full_text))
        has_toc = ("TOC" in self.doc_xml_bytes.decode("utf-8", errors="ignore")
                   or "目录" in full_text)
        has_main_chapter = bool(re.search(
            r"一[、,．. ]市场分析|二[、,．. ]项目简介|三[、,．. ]产品与服务|"
            r"四[、,．. ]商业模式|五[、,．. ]财务分析|六[、,．. ]风险识别|"
            r"七[、,．. ]未来前景|八[、,．. ]团队介绍", full_text))
        has_images = len([n for n in self.zip_names if n.startswith("word/media/")]) > 0
        has_tables = len(self.tables) > 0
        has_page_number = any(b"PAGE" in f for f in self.footers.values())

        miss = []
        if not has_cover:
            miss.append("未保留封面")
        if not has_toc:
            miss.append("未保留目录")
        if not has_main_chapter:
            miss.append("未保留正文章节")
        if not has_images:
            miss.append("未保留图片")
        if not has_tables:
            miss.append("未保留表格")
        if not has_page_number:
            miss.append("未保留页码")


        # 说明：按用户要求，删除维度一中"无连续空白页 / 无大面积乱码 /
        # 无大面积文字重叠 / 无表格或图片严重遮挡正文"这部分检查，仅保留
        # "保留封面/目录/正文章节/图片/表格/页码等原有主要组成部分"。

        if miss:
            self.dimension1_pass = False
            self.dimension1_reasons = miss
        return self.dimension1_pass

    # ---------- 辅助：定位标题/章节 ----------
    def classify_paragraphs(self) -> Dict[int, str]:
        labels: Dict[int, str] = {}
        for i, p in enumerate(self.paras):
            t = para_text(p).strip()
            if not t:
                continue
            if re.match(r"^[一二三四五六七八九十]+[、,．.]", t):
                labels[i] = "L1"
            elif re.match(r"^（[一二三四五六七八九十]+）", t):
                labels[i] = "L2"
            # 三级标题：仅整数编号 "1." "2." "3." 形式（点后紧跟非数字，
            # 从而排除 "1.1"、"2.1" 这类更下一级(四级)标题）。
            elif re.match(r"^\d+[、,．.](?!\d)", t) and len(t) < 30:
                labels[i] = "L3"
        return labels

    def get_all_paragraph_text(self) -> str:
        return "\n".join(para_text(p) for p in self.paras)

    # ---------- 维度 2：检查项 ----------
    def _body_start_para_index(self) -> int:
        """返回"正文"起始段落在 self.paras 中的索引。

        细则口径：正文从**目录页之后**开始（封面、信息页、目录都不算正文）。
        做法：在 body 顺序中定位目录(TOC 域，通常在 <w:sdt> 里)所在位置，
        取其之后第一个段落作为正文起点；找不到目录时，退回到第二个分节符
        (封面、信息页之后)之后，仍找不到则从 0 开始。
        """
        bc = self.body_children

        def is_toc_node(node) -> bool:
            for it in node.findall(".//w:instrText", NS):
                if it.text and "TOC" in it.text:
                    return True
            for fs in node.findall(".//w:fldSimple", NS):
                if "TOC" in (fs.get(qn("w:instr")) or ""):
                    return True
            return False

        # 目录之后第一个 body 子元素的位置
        toc_body_idx = None
        for i, ch in enumerate(bc):
            if is_toc_node(ch):
                toc_body_idx = i
                break

        if toc_body_idx is not None:
            # 统计 body 中该位置之前的段落数 = 正文首段在 self.paras 的索引
            return sum(1 for ch in bc[:toc_body_idx + 1]
                       if ch.tag.split("}")[-1] == "p")

        # 无目录：退回到第二个分节符（封面、信息页）之后
        sect_para_positions = []
        for k, p in enumerate(self.paras):
            ppr = p.find("w:pPr", NS)
            if ppr is not None and ppr.find("w:sectPr", NS) is not None:
                sect_para_positions.append(k)
        if len(sect_para_positions) >= 2:
            return sect_para_positions[1] + 1
        return 0

    def check_body_line_spacing_1x(self) -> bool:
        """+3：文档中正文文本间距为1倍

        细则：文档正文文本行距为 1 倍（单倍行距）。**正文从目录页之后开始**
        （封面、信息页、目录不计入正文）。
        判定：正文区域内所有正文段落（非空、非一/二/三级标题）行距必须均为 1 倍；
        只要有任意一个正文段落不是 1 倍行距，即不得分。

        在办公软件（Word/WPS）中，1 倍行距的表示方式：
          - w:spacing line="240" lineRule="auto"（240 = 1 行，auto 表示倍数行距）
          - 或未设置任何行距（无 w:spacing/无 line），此时渲染为单倍行距。
        1.5 倍为 line="360"，2 倍为 line="480"，均判为不合格。
        """
        labels = self.classify_paragraphs()
        start = self._body_start_para_index()  # 正文起点（目录页之后）

        def is_single_spacing(sp) -> bool:
            line = sp.get("line")
            rule = sp.get("line_rule")
            # 完全未设置行距：办公软件按单倍渲染
            if line is None:
                return True
            # 倍数行距：line=240 即 1 倍
            if rule in (None, "auto"):
                return line == 240
            # atLeast / exact（固定值/最小值）不属于"1 倍"倍数行距
            return False

        count = 0
        for i, p in enumerate(self.paras):
            if i < start:            # 跳过封面、信息页、目录（非正文）
                continue
            t = para_text(p).strip()
            if not t:            # 跳过空段落
                continue
            if i in labels:      # 跳过一/二/三级标题
                continue
            # 该段属于正文文本，逐一检查其行距
            count += 1
            if not is_single_spacing(get_para_spacing(p)):
                return False
        # 若没有可判定的正文段落，则无法确认满足细则，不得分
        if count == 0:
            return False
        return True

    def check_table_line_spacing_1_25(self) -> bool:
        """+5：所有表格中的文本行距均为1.25倍

        细则：文档中所有表格里的（文本）段落行距都必须是 1.25 倍。
        判定：遍历每一个表格内的每一个段落，只要有任意一个段落的行距
        不是 1.25 倍，即不得分；全部为 1.25 倍才得分。

        在办公软件（Word/WPS）中，1.25 倍行距的表示方式：
          - w:spacing line="300" lineRule="auto"（300 = 1.25 行，auto 表示倍数行距）
        其它情况均判为不合格：
          - 未设置行距（无 w:spacing 或无 line）→ 办公软件按单倍(1.0)渲染，不是 1.25 倍
          - lineRule="exact"/"atLeast"（固定值/最小值磅数）→ 不是倍数行距
          - line 为其它值（如 240=1倍、360=1.5倍、480=2倍）
        """

        def is_1_25_spacing(sp) -> bool:
            line = sp.get("line")
            rule = sp.get("line_rule")
            # 必须是倍数行距(auto)，且 line=300（即 1.25 倍）
            if rule in (None, "auto"):
                return line == 300
            return False

        any_para = False
        for tbl in self.tables:
            for p in tbl.findall(".//w:p", NS):
                any_para = True
                if not is_1_25_spacing(get_para_spacing(p)):
                    return False
        # 无表格或表格内无段落时，无法确认满足细则，不得分
        if not any_para:
            return False
        return True

    def check_l1_heading_style(self) -> bool:
        """+3：一级标题为宋体、小三号、单倍行距；段前0.5行、段后0.5行

        细则拆解（每个点都要满足）：
          1) 字体为宋体（中文 eastAsia = 宋体/SimSun）
          2) 字号为小三号 = 15pt = 30 half-pt（w:sz val="30"）
          3) 行距为单倍（1 倍）
          4) 段前 0.5 行
          5) 段后 0.5 行
        判定：所有一级标题都必须同时满足以上 5 点；任意一个标题不满足即不得分。

        在办公软件（Word/WPS）中的有效表示：
          - 单倍行距：w:spacing line="240" lineRule="auto"，或未设置行距（默认单倍）
          - 段前/段后 0.5 行：以"行"为单位，存储为 w:beforeLines="50" / w:afterLines="50"
            （行单位以百分之一行计，50 即 0.5 行）。注意 w:before/w:after 是以
            twips(磅) 为单位的固定间距，不能等价于"0.5 行"，故必须用 *Lines 判定。
        """
        labels = self.classify_paragraphs()
        l1_indices = [i for i, lb in labels.items() if lb == "L1"]
        if not l1_indices:
            return False
        for i in l1_indices:
            p = self.paras[i]
            fonts = get_para_run_fonts(p, self.theme_fonts, self.default_east_font)
            sp = get_para_spacing(p)

            # 1) 宋体：所有设置了中文字体的 run 必须是宋体
            east_fonts = [f.get("east_font") for f in fonts
                          if f.get("east_font") is not None]
            if not east_fonts or any(ef not in ("宋体", "SimSun") for ef in east_fonts):
                return False

            # 2) 小三号(30 half-pt)：所有设置了字号的 run 必须是 30
            sizes = [f.get("size_half_pt") for f in fonts
                     if f.get("size_half_pt") is not None]
            if not sizes or any(s != 30 for s in sizes):
                return False

            # 3) 单倍行距：line=240/auto，或未设置行距
            line = sp.get("line")
            rule = sp.get("line_rule")
            if line is not None:
                if not (line == 240 and rule in (None, "auto")):
                    return False

            # 4) 段前 0.5 行；5) 段后 0.5 行（必须以"行"为单位存储）
            if sp.get("before_lines") != 50:
                return False
            if sp.get("after_lines") != 50:
                return False
        return True

    def check_l2_heading_style(self) -> bool:
        """+3：二级标题为宋体、四号、1.5倍行距

        细则拆解（每个点都要满足）：
          1) 字体为宋体（中文 eastAsia = 宋体/SimSun）
          2) 字号为四号 = 14pt = 28 half-pt（w:sz val="28"）
          3) 行距为 1.5 倍
        判定：所有二级标题都必须同时满足以上 3 点；任意一个不满足即不得分。

        在办公软件（Word/WPS）中的有效表示：
          - 1.5 倍行距：w:spacing line="360" lineRule="auto"（360 = 1.5 行，auto 表示倍数行距）
        """
        labels = self.classify_paragraphs()
        l2_indices = [i for i, lb in labels.items() if lb == "L2"]
        if not l2_indices:
            return False
        for i in l2_indices:
            p = self.paras[i]
            fonts = get_para_run_fonts(p, self.theme_fonts, self.default_east_font)
            sp = get_para_spacing(p)

            # 1) 宋体：所有设置了中文字体的 run 必须是宋体
            east_fonts = [f.get("east_font") for f in fonts
                          if f.get("east_font") is not None]
            if not east_fonts or any(ef not in ("宋体", "SimSun") for ef in east_fonts):
                return False

            # 2) 四号(28 half-pt)：所有设置了字号的 run 必须是 28
            sizes = [f.get("size_half_pt") for f in fonts
                     if f.get("size_half_pt") is not None]
            if not sizes or any(s != 28 for s in sizes):
                return False

            # 3) 1.5 倍行距：line=360 且 lineRule=auto
            if not (sp.get("line") == 360 and sp.get("line_rule") in (None, "auto")):
                return False
        return True

    def check_l3_heading_style(self) -> bool:
        """+3：三级标题为宋体、小四、单倍行距、加粗

        细则拆解（每个点都要满足）：
          1) 字体为宋体（中文 eastAsia = 宋体/SimSun）
          2) 字号为小四 = 12pt = 24 half-pt（w:sz val="24"）
          3) 行距为单倍（1 倍）
          4) 加粗
        判定：所有三级标题都必须同时满足以上 4 点；任意一个不满足即不得分。

        在办公软件（Word/WPS）中的有效表示：
          - 单倍行距：w:spacing line="240" lineRule="auto"，或未设置行距（默认单倍）
          - 加粗：run 上有 w:b（w:b val 未设或为 "1"/"true"/"on" 表示开启；
            val 为 "0"/"false"/"none" 表示显式关闭）
        """
        labels = self.classify_paragraphs()
        l3_indices = [i for i, lb in labels.items() if lb == "L3"]
        if not l3_indices:
            return False
        for i in l3_indices:
            p = self.paras[i]
            fonts = get_para_run_fonts(p, self.theme_fonts, self.default_east_font)
            sp = get_para_spacing(p)

            # 1) 宋体：所有设置了中文字体的 run 必须是宋体
            east_fonts = [f.get("east_font") for f in fonts
                          if f.get("east_font") is not None]
            if not east_fonts or any(ef not in ("宋体", "SimSun") for ef in east_fonts):
                return False

            # 2) 小四(24 half-pt)：所有设置了字号的 run 必须是 24
            sizes = [f.get("size_half_pt") for f in fonts
                     if f.get("size_half_pt") is not None]
            if not sizes or any(s != 24 for s in sizes):
                return False

            # 3) 单倍行距：line=240/auto，或未设置行距
            line = sp.get("line")
            rule = sp.get("line_rule")
            if line is not None:
                if not (line == 240 and rule in (None, "auto")):
                    return False

            # 4) 加粗：所有含文本的 run 都必须加粗
            text_fonts = [f for f in fonts if f.get("has_text")]
            if not text_fonts or any(not f.get("bold") for f in text_fonts):
                return False
        return True

    def find_tourism_resource_table(self) -> Optional:
        """定位'（二）旅游资源分析'下方的表格（首行包含'主类'/'亚类'等）。"""
        for i, ch in enumerate(self.body_children):
            if ch.tag.split("}")[-1] != "p":
                continue
            t = para_text(ch).strip()
            if "（二）旅游资源分析" in t:
                for nxt in self.body_children[i + 1:]:
                    if nxt.tag.split("}")[-1] == "tbl":
                        rows = nxt.findall("w:tr", NS)
                        if rows:
                            first_text = "".join(t.text or "" for t in rows[0].findall(".//w:t", NS))
                            if "主类" in first_text or "亚类" in first_text or "资源名称" in first_text:
                                return nxt
                break
        return None

    def find_tourism_resource_tables(self) -> list:
        """收集"（二）旅游资源分析"表格对应的**所有物理表格**。

        该表在办公软件中被拆成相邻的多个 <w:tbl>（第1页 + 分页符 + 续表），
        中间仅隔含分页符的空段落、且表头一致。为把它们当作同一张逻辑表处理，
        这里从首个物理表起，向后合并"仅隔空段落"的、表头相同的相邻表格。
        """
        first = self.find_tourism_resource_table()
        if first is None:
            return []

        def header_sig(tbl) -> str:
            rows = tbl.findall("w:tr", NS)
            if not rows:
                return ""
            return "|".join("".join(t.text or "" for t in tc.findall(".//w:t", NS))
                            for tc in rows[0].findall("w:tc", NS))

        # body 中所有表格的位置
        positions = [(i, ch) for i, ch in enumerate(self.body_children)
                     if ch.tag.split("}")[-1] == "tbl"]
        pos_of = {id(ch): i for i, ch in positions}
        start_idx = pos_of.get(id(first))
        sig = header_sig(first)

        result = [first]
        # 找到 first 在 positions 中的序号
        order = [k for k, (i, ch) in enumerate(positions) if i == start_idx]
        if not order:
            return result
        k = order[0]
        while k + 1 < len(positions):
            cur_bidx = positions[k][0]
            nxt_bidx, nxt_tbl = positions[k + 1]
            # 两表之间只夹空段落（视觉上是续表，非独立新表）
            only_blank = True
            for ch in self.body_children[cur_bidx + 1:nxt_bidx]:
                if ch.tag.split("}")[-1] != "p" or para_text(ch).strip():
                    only_blank = False
                    break
            if only_blank and header_sig(nxt_tbl) == sig and sig != "":
                result.append(nxt_tbl)
                k += 1
            else:
                break
        return result

    def check_tourism_table_two_pages_with_header(self) -> bool:
        """+5：“（二）旅游资源分析”表格排版为 2 页，第 2 页出现和第 1 页内容、格式均相同的表头

        细则要点（严格化）：
          (i)  表格恰好排版为 **2 页**（不是 3 页以上，也不是 1 页）；
          (ii) 第 2 页顶端出现与第 1 页表头**内容 + 格式**完全一致的表头。

        判定两种合规实现路径：
          A) 物理分段续表：作者把该表拆成相邻的 <w:tbl>，中间只夹含分页符 <w:br
             w:type="page"/> 的空段落；且后续表格首行（表头）与首表首行的内容签名
             和视觉签名都一致。要求恰好 2 段（>2 段视为 3 页以上，判不通过）。
          B) 自动续表：单个 <w:tbl>，首行设置 <w:tblHeader/>（Word/WPS 才会在下一
             页自动重复该表头），且估算渲染高度 total_h 满足
                 page_h < total_h <= 2 * page_h
             即恰好跨 2 页（>2*page_h 视为 3 页以上）。

        任一路径成立即通过；否则不通过。
        """
        first = self.find_tourism_resource_table()
        if first is None:
            return False

        bc = self.body_children
        positions = [(i, ch) for i, ch in enumerate(bc)
                     if ch.tag.split("}")[-1] == "tbl"]
        if not positions:
            return False

        def _row_content_sig(row) -> str:
            return "|".join("".join((t.text or "") for t in tc.findall(".//w:t", NS))
                            for tc in row.findall("w:tc", NS))

        def _row_visual_sig(row) -> tuple:
            """行的视觉签名：单元格结构 + 每个单元格首段/首 run 的字体样式。
            用于严格判定"格式相同的表头"。"""
            cells_sig = []
            for tc in row.findall("w:tc", NS):
                tcpr = tc.find("w:tcPr", NS)
                grid = w = shd_fill = valign = None
                borders_sig = ""
                if tcpr is not None:
                    gs = tcpr.find("w:gridSpan", NS)
                    if gs is not None:
                        grid = gs.get(qn("w:val"))
                    tcw = tcpr.find("w:tcW", NS)
                    if tcw is not None:
                        w = (tcw.get(qn("w:w")), tcw.get(qn("w:type")))
                    shd = tcpr.find("w:shd", NS)
                    if shd is not None:
                        shd_fill = (shd.get(qn("w:fill")), shd.get(qn("w:val")),
                                    shd.get(qn("w:color")))
                    va = tcpr.find("w:vAlign", NS)
                    if va is not None:
                        valign = va.get(qn("w:val"))
                    tcb = tcpr.find("w:tcBorders", NS)
                    if tcb is not None:
                        parts = []
                        for side in ("top", "left", "bottom", "right",
                                     "insideH", "insideV"):
                            e = tcb.find(f"w:{side}", NS)
                            if e is not None:
                                parts.append((side, e.get(qn("w:val")),
                                              e.get(qn("w:sz")),
                                              e.get(qn("w:color"))))
                        borders_sig = tuple(parts)
                # 首段的段落对齐 + 首个含文字 run 的字体样式
                p = tc.find("w:p", NS)
                jc = font_sig = None
                if p is not None:
                    ppr = p.find("w:pPr", NS)
                    if ppr is not None:
                        jc_el = ppr.find("w:jc", NS)
                        if jc_el is not None:
                            jc = jc_el.get(qn("w:val"))
                    for r in p.findall(".//w:r", NS):
                        if not any((t.text or "") for t in r.findall("w:t", NS)):
                            continue
                        rpr = r.find("w:rPr", NS)
                        b_on = i_on = False
                        sz = ascii_f = east_f = None
                        if rpr is not None:
                            b = rpr.find("w:b", NS)
                            if b is not None and b.get(qn("w:val")) not in (
                                    "0", "false", "off", "none"):
                                b_on = True
                            if rpr.find("w:i", NS) is not None:
                                i_on = True
                            sz_el = rpr.find("w:sz", NS)
                            if sz_el is not None:
                                sz = sz_el.get(qn("w:val"))
                            rf = rpr.find("w:rFonts", NS)
                            if rf is not None:
                                ascii_f = rf.get(qn("w:ascii"))
                                east_f = rf.get(qn("w:eastAsia"))
                        font_sig = (b_on, i_on, sz, ascii_f, east_f)
                        break
                cells_sig.append((grid, w, shd_fill, valign, borders_sig,
                                  jc, font_sig))
            return tuple(cells_sig)

        def _only_pagebreak_between(a_idx: int, b_idx: int) -> bool:
            """两个相邻表格之间仅含分页符的空段落（视觉即"下一页续表"）。"""
            saw_break = False
            for ch in bc[a_idx + 1:b_idx]:
                if ch.tag.split("}")[-1] != "p":
                    return False
                if para_text(ch).strip():
                    return False
                for br in ch.findall(".//w:br", NS):
                    if br.get(qn("w:type")) == "page":
                        saw_break = True
            return saw_break

        first_rows = first.findall("w:tr", NS)
        if not first_rows:
            return False
        head_content = _row_content_sig(first_rows[0])
        head_visual = _row_visual_sig(first_rows[0])
        if not head_content.strip("|"):
            return False

        # 路径 A：物理分段续表 —— 从首表起沿正文向后合并"仅隔分页符 + 表头一致"的相邻表。
        pos_of = {id(ch): i for i, ch in positions}
        start_bidx = pos_of.get(id(first))
        order = [k for k, (i, _) in enumerate(positions) if i == start_bidx]
        if not order:
            return False
        group_tbls = [first]
        k = order[0]
        while k + 1 < len(positions):
            cur_bidx = positions[k][0]
            nxt_bidx, nxt_tbl = positions[k + 1]
            nxt_rows = nxt_tbl.findall("w:tr", NS)
            if not nxt_rows:
                break
            if (_only_pagebreak_between(cur_bidx, nxt_bidx)
                    and _row_content_sig(nxt_rows[0]) == head_content
                    and _row_visual_sig(nxt_rows[0]) == head_visual):
                group_tbls.append(nxt_tbl)
                k += 1
            else:
                break

        if len(group_tbls) >= 2:
            # 恰好 2 段：第 2 段就是"第 2 页"，且表头内容+格式与第 1 页一致
            return len(group_tbls) == 2

        # 路径 B：单表 + <w:tblHeader/> 自动续表
        tbl_pr = first.find("w:tblPr", NS)  # 保留，供后续如需诊断
        _ = tbl_pr
        first_row_pr = first_rows[0].find("w:trPr", NS)
        has_tbl_header = (first_row_pr is not None
                          and first_row_pr.find("w:tblHeader", NS) is not None)
        if not has_tbl_header:
            return False  # 无 tblHeader，Word/WPS 不会在下一页重复表头
        page_h = self._page_content_height()
        total_h = self._table_height(first)
        # 严格 2 页：total_h 超过 1 页正文高度，但不超过 2 页
        return page_h < total_h <= 2 * page_h

    def _style_tbl_borders(self, style_id):
        """沿表格样式的 basedOn 链，取该样式定义的 w:tblBorders（就近优先）。"""
        if self.styles_root is None or not style_id:
            return None
        by_id = {}
        for s in self.styles_root.findall("w:style", NS):
            sid = s.get(qn("w:styleId"))
            if sid is not None:
                by_id[sid] = s
        seen = set()
        cur = style_id
        while cur and cur not in seen:
            seen.add(cur)
            s = by_id.get(cur)
            if s is None:
                break
            b = s.find(".//w:tblPr/w:tblBorders", NS)
            if b is not None:
                return b
            based = s.find("w:basedOn", NS)
            cur = based.get(qn("w:val")) if based is not None else None
        return None

    def _page_content_height(self) -> int:
        """页面正文区高度（twips）：pgSz.h - 上下页边距。"""
        page_h = 15840
        sect = self.body.find("w:sectPr", NS)
        if sect is not None:
            pg = sect.find("w:pgSz", NS)
            mar = sect.find("w:pgMar", NS)
            if pg is not None and pg.get(qn("w:h")):
                h = int(pg.get(qn("w:h")))
                top = int(mar.get(qn("w:top"))) if (mar is not None and mar.get(qn("w:top"))) else 0
                bot = int(mar.get(qn("w:bottom"))) if (mar is not None and mar.get(qn("w:bottom"))) else 0
                page_h = max(1, h - top - bot)
        return page_h

    def _table_height(self, tbl) -> int:
        """估算表格在办公软件中的渲染高度（twips）。

        逐行取"单元格中最高的一列"作为行高：每个单元格按其段落数、每段字号与
        行距（line/lineRule）折算行高，并加上单元格上下边距。字号越小、行数虽多
        也可能只占很小高度，从而正确反映"小字号表格一页放得下"的真实排版。
        """
        def para_height(p) -> float:
            szs = []
            for r in p.findall(".//w:r", NS):
                rpr = r.find("w:rPr", NS)
                sz = rpr.find("w:sz", NS) if rpr is not None else None
                if sz is not None and sz.get(qn("w:val")):
                    szs.append(int(sz.get(qn("w:val"))))
            fs_pt = (max(szs) / 2.0) if szs else 10.5  # 半磅->磅；默认约五号
            sp = get_para_spacing(p)
            line = sp.get("line")
            rule = sp.get("line_rule")
            if line and rule in (None, "auto"):
                lh = fs_pt * 20 * (line / 240.0)       # 倍数行距，240=单倍
            elif line and rule in ("exact", "atLeast"):
                lh = max(int(line), fs_pt * 20)          # 固定/最小磅值
            else:
                lh = fs_pt * 20 * 1.15                    # 默认单倍略含行间距
            n_lines = 1 + len(p.findall(".//w:br", NS))
            return lh * n_lines

        total = 0.0
        for tr in tbl.findall("w:tr", NS):
            row_h = 0.0
            for tc in tr.findall("w:tc", NS):
                cell_h = sum(para_height(p) for p in tc.findall("w:p", NS))
                row_h = max(row_h, cell_h)
            # 行高下限：w:trHeight（exact 为固定值，atLeast 为最小值）。
            # 内容折算高度小于该下限时，实际渲染以下限为准。
            trpr = tr.find("w:trPr", NS)
            th = trpr.find("w:trHeight", NS) if trpr is not None else None
            if th is not None and th.get(qn("w:val")):
                min_h = int(th.get(qn("w:val")))
                rule = th.get(qn("w:hRule"))
                if rule == "exact":
                    row_h = min_h
                else:  # atLeast / 默认
                    row_h = max(row_h, min_h)
            total += row_h + 60  # 单元格上下边距近似
        return int(total)

    def check_only_one_table_two_pages(self) -> bool:
        """+5：仅有一个表格分 2 页排版，其他表格均为一页排版

        细则拆解（严格化）：
          1) 恰好有 1 个（逻辑）表格是"分 **2** 页排版"——不是 1 页，也不是 3 页以上；
          2) 其余所有（逻辑）表格都是"**一** 页排版"（既非跨页也非多页）。

        任一表格若排到 3 页及以上，或有 2 个以上表格跨页，均不满足。

        判定（不启用 COM，仅由 OOXML 推导每个逻辑表的"占用页数"）：
          先把相邻的物理续表（仅隔含分页符的空段落、表头一致）合并成逻辑组，
          再对每个逻辑组估算页数 pages：
            (a) 物理分段续表：每一段落在各自一页，pages = 段数（len(group)）。
            (b) 单个 <w:tbl>：按字号/行距估算渲染高度 total_h，
                pages = max(1, ceil(total_h / page_h))。
                （仅凭行数会误判——小字号表格几十行仍可容于一页。）
          最终要求：恰好 1 个逻辑组 pages == 2，其余所有逻辑组 pages == 1。
        """
        NSp = NS

        def header_sig(tbl) -> str:
            rows = tbl.findall("w:tr", NSp)
            if not rows:
                return ""
            return "|".join("".join(t.text or "" for t in tc.findall(".//w:t", NSp))
                            for tc in rows[0].findall("w:tc", NSp))

        def only_pagebreak_between(a_idx: int, b_idx: int) -> bool:
            """两个相邻表格之间只夹着空段落，且其中含分页符。"""
            saw_break = False
            for ch in self.body_children[a_idx + 1:b_idx]:
                if ch.tag.split("}")[-1] != "p":
                    return False
                if para_text(ch).strip():
                    return False
                for br in ch.findall(".//w:br", NSp):
                    if br.get(qn("w:type")) == "page":
                        saw_break = True
            return saw_break

        page_h = self._page_content_height()

        tbl_positions = [(i, ch) for i, ch in enumerate(self.body_children)
                         if ch.tag.split("}")[-1] == "tbl"]
        if not tbl_positions:
            return False

        # (b) 合并"分页符续表"为逻辑组
        groups = []
        used = set()
        for k in range(len(tbl_positions)):
            if k in used:
                continue
            group = [tbl_positions[k]]
            used.add(k)
            while True:
                last_idx = group[-1][0]
                nxt = None
                for j in range(len(tbl_positions)):
                    if j in used:
                        continue
                    if tbl_positions[j][0] > last_idx:
                        nxt = j
                        break
                if nxt is None:
                    break
                nxt_idx, nxt_tbl = tbl_positions[nxt]
                if (only_pagebreak_between(last_idx, nxt_idx)
                        and header_sig(group[-1][1]) == header_sig(nxt_tbl)
                        and header_sig(nxt_tbl) != ""):
                    group.append((nxt_idx, nxt_tbl))
                    used.add(nxt)
                else:
                    break
            groups.append(group)

        def group_pages(group) -> int:
            """逻辑组占用页数：物理续表按段数；单表按渲染高度向上取整。"""
            if len(group) >= 2:
                return len(group)  # 每一续表段各占一页
            total_h = self._table_height(group[0][1])
            if page_h <= 0:
                return 1
            return max(1, math.ceil(total_h / page_h))

        pages_list = [group_pages(g) for g in groups]
        # 任一表格 3 页及以上 → 不满足（rubric 明确是"分 2 页"）
        if any(p >= 3 for p in pages_list):
            return False
        two_page_count = sum(1 for p in pages_list if p == 2)
        one_page_count = sum(1 for p in pages_list if p == 1)
        # 恰好 1 个表格 2 页，其余全部 1 页
        return two_page_count == 1 and one_page_count == len(pages_list) - 1


    def check_tourism_table_three_borders_per_page(self) -> bool:
        """+1：“（二）旅游资源分析”表格每一页只有三条边框线（三线表：顶线 + 表头下线 + 底线）

        细则语义（针对办公软件实际渲染）：
          该表被作者物理拆成相邻的多个 <w:tbl>（每个 <w:tbl> 对应"一页"），
          每一页的表格应恰好呈现 3 条水平边框线，且没有任何竖直边框线——
          这是学术论文里的"三线表"外观：顶线、表头分割线、底线。

        判定（不启用 COM，仅由 OOXML 边框声明推导）：
          对每一个物理表分别计算"实际会被绘制的边框条数"：
            水平线总数 = 顶边框(1) + 相邻行间被绘制的横线条数 + 底边框(1)
            竖直线数量 = 左/右/insideV 或任何单元格竖向边框
          单元格级 tcBorders 覆盖表级 tblBorders；表级为空时按 tblStyle
          继承（通过 _style_tbl_borders 沿 basedOn 链就近取值）。
          每张物理表都必须满足：水平线数 == 3 且 竖直线数 == 0。
        """
        tables = self.find_tourism_resource_tables()
        if not tables:
            return False

        def _drawn(el) -> bool:
            """<w:xxx w:val="..."> 是否表示"会被绘制"——非 nil/none 即视为绘制。"""
            if el is None:
                return False
            return el.get(qn("w:val")) not in ("nil", "none")

        def _tbl_default_borders(tbl):
            """表级默认 6 方位边框（tblBorders 或沿样式 basedOn 链继承）。
            返回 dict: {side: bool_drawn}"""
            sides = ("top", "left", "bottom", "right", "insideH", "insideV")
            tblpr = tbl.find("w:tblPr", NS)
            borders = tblpr.find("w:tblBorders", NS) if tblpr is not None else None
            if borders is None and tblpr is not None:
                ts = tblpr.find("w:tblStyle", NS)
                borders = self._style_tbl_borders(
                    ts.get(qn("w:val")) if ts is not None else None)
            out = {s: False for s in sides}
            if borders is None:
                return out
            for s in sides:
                out[s] = _drawn(borders.find(f"w:{s}", NS))
            return out

        def _cell_border(tc, side: str, tbl_default: dict) -> bool:
            """单元格在指定方位是否会绘制边框——tcBorders 优先，否则回退表级默认。"""
            tcpr = tc.find("w:tcPr", NS)
            tcb = tcpr.find("w:tcBorders", NS) if tcpr is not None else None
            if tcb is not None:
                el = tcb.find(f"w:{side}", NS)
                if el is not None:
                    return _drawn(el)
            return bool(tbl_default.get(side, False))

        def _page_borders_ok(tbl) -> bool:
            rows = tbl.findall("w:tr", NS)
            if not rows:
                return False
            tbl_def = _tbl_default_borders(tbl)

            # ---- 水平线：顶线 / 行间线 / 底线 ----
            top_drawn = any(_cell_border(tc, "top", tbl_def)
                            for tc in rows[0].findall("w:tc", NS))
            bot_drawn = any(_cell_border(tc, "bottom", tbl_def)
                            for tc in rows[-1].findall("w:tc", NS))
            between = 0
            for i in range(len(rows) - 1):
                upper_cells = rows[i].findall("w:tc", NS)
                lower_cells = rows[i + 1].findall("w:tc", NS)
                line_here = (
                    any(_cell_border(tc, "bottom", tbl_def) for tc in upper_cells)
                    or any(_cell_border(tc, "top", tbl_def) for tc in lower_cells)
                    or tbl_def.get("insideH", False)
                )
                if line_here:
                    between += 1
            h_total = int(top_drawn) + between + int(bot_drawn)

            # ---- 竖直线：左/右/insideV，任何单元格竖向边框都算 ----
            v_any = False
            if tbl_def.get("left") or tbl_def.get("right") or tbl_def.get("insideV"):
                v_any = True
            if not v_any:
                for tr in rows:
                    for tc in tr.findall("w:tc", NS):
                        if (_cell_border(tc, "left", tbl_def)
                                or _cell_border(tc, "right", tbl_def)):
                            v_any = True
                            break
                    if v_any:
                        break

            return h_total == 3 and not v_any

        # 每一物理表（= 每一页）都必须满足"3 条水平线 + 0 条竖直线"
        return all(_page_borders_ok(t) for t in tables)

    def check_any_page_no_page_number(self) -> bool:
        """-1：从目录页开始任意页无页码

        细则：从"目录页"开始的任意一页若没有页码，即扣分（目录页之前的封面、
        信息页不在考察范围内）。
        判定：定位目录(TOC 域)所在的节，从该节起逐节检查其生效页脚是否含页码域(PAGE)；
        只要目录页及其之后存在任一页没有带页码的页脚，即满足扣分条件。目录页之前的节
        （封面、信息页）即使无页码也不计。

        在办公软件（WPS/Word）中的有效判定要点：
          - 页脚通过节的 <w:footerReference> 关联到 footerN.xml；type 可为
            default（默认）、first（首页）、even（偶数页）。
          - 页码是页脚中的 PAGE 域：<w:fldSimple w:instr="...PAGE..."> 或
            <w:instrText>PAGE</w:instrText>（footer 的 XML 含 "PAGE" 即视为有页码）。
          - **页脚继承**：某节缺某类型 footerReference 时沿用前一节的同类页脚。
            为保证继承正确，继承状态从文档第一个节开始累积，但**只从目录页所在节起**
            才据此判违规。
          - **首页页脚**：仅当该节设置了 <w:titlePg/> 时，first 型页脚才对首页生效。
        """
        # 目录(TOC 域)所在的 body 位置
        toc_body_idx = None
        for i, ch in enumerate(self.body_children):
            found = False
            for it in ch.findall(".//w:instrText", NS):
                if it.text and "TOC" in it.text:
                    found = True
                    break
            if not found:
                for fs in ch.findall(".//w:fldSimple", NS):
                    if "TOC" in (fs.get(qn("w:instr")) or ""):
                        found = True
                        break
            if found:
                toc_body_idx = i
                break

        # 按文档顺序取出各节的 sectPr，并记录该节结束(所含最后一个 body 子元素)的位置
        sects = []       # (sectPr, end_body_idx)
        for i, ch in enumerate(self.body_children):
            if ch.tag.split("}")[-1] == "p":
                ppr = ch.find("w:pPr", NS)
                if ppr is not None and ppr.find("w:sectPr", NS) is not None:
                    sects.append((ppr.find("w:sectPr", NS), i))
        body_sect = self.body.find("w:sectPr", NS)
        if body_sect is not None:
            sects.append((body_sect, len(self.body_children)))
        if not sects:
            return True  # 没有任何节/页脚

        # 目录所在节的序号：第一个 end_body_idx >= 目录位置的节；找不到目录则从第 0 节起
        toc_section_idx = 0
        if toc_body_idx is not None:
            for k, (_, end_idx) in enumerate(sects):
                if end_idx >= toc_body_idx:
                    toc_section_idx = k
                    break

        def footer_has_page(rid) -> bool:
            if not rid:
                return False
            tgt = self.rels.get(rid)
            if not tgt or not str(tgt[1]).startswith("footer"):
                return False
            fp = "word/" + tgt[1]
            return fp in self.footers and b"PAGE" in self.footers[fp]

        def refs_by_type(sect) -> dict:
            out = {}
            for r in sect.findall("w:footerReference", NS):
                out[r.get(qn("w:type"))] = r.get(f"{{{R_NS}}}id")
            return out

        # 从第一个节起累积继承，但只从目录页所在节起判违规
        inherited_default = None
        inherited_first = None
        for k, (sect, _end) in enumerate(sects):
            refs = refs_by_type(sect)
            cur_default = refs.get("default", inherited_default)
            cur_first = refs.get("first", inherited_first)

            if k >= toc_section_idx:  # 目录页及之后的节才考察页码
                has_title_pg = sect.find("w:titlePg", NS) is not None
                if not footer_has_page(cur_default):
                    return True
                if has_title_pg and not footer_has_page(cur_first):
                    return True

            inherited_default = cur_default
            inherited_first = cur_first

        return False

    def check_cover_no_center_text(self) -> bool:
        """-1：封面居中位置没有文本内容

        细则：封面（文档第一节/首页）的"居中位置"没有任何文本内容，即扣分。
        判定：封面区域内必须存在**确实位于页面水平居中位置**的文本承载物；仅有
        文本框而位置不在居中区域的不算数。任一以下成立即"居中位置有文本" → 不扣分：

          A) 普通段落 <w:pPr><w:jc w:val="center"/> 且该段落含可见文字（含义等于
             版心水平居中的段落）。
          B) DrawingML 浮动文本框 <w:drawing><wp:anchor>：内含 <w:txbxContent>
             且其 <wp:positionH> 满足以下之一：
                (i)  <wp:align>center</wp:align>（在 relativeFrom 参考系居中）；
                (ii) <wp:posOffset>+<wp:extent cx="…"> 计算出的框中心与该
                     参考系中心相差不超过 0.5cm。
          C) DrawingML 内嵌文本框 <wp:inline>：位置跟随其所在段落，段落
             jc=center 即视为居中。
          D) VML 形状 <w:pict><v:shape/v:rect …>：其 style 含
             mso-position-horizontal:center；或由 left+width 计算得到的框中心
             与页面/版心中心相差不超过 0.5cm。

        以上均不满足才扣分。全部通过 OOXML 属性解析，不使用 COM。
        """
        WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
        V_NS = "urn:schemas-microsoft-com:vml"
        NSX = {"w": W_NS, "wp": WP_NS, "a": A_NS, "v": V_NS}

        # ---- 页面几何（twips）----
        page_w = 12240
        margin_l = margin_r = 0
        body_sect = self.body.find("w:sectPr", NS)
        if body_sect is not None:
            pg = body_sect.find("w:pgSz", NS)
            mar = body_sect.find("w:pgMar", NS)
            if pg is not None and pg.get(qn("w:w")):
                page_w = int(pg.get(qn("w:w")))
            if mar is not None:
                if mar.get(qn("w:left")):
                    margin_l = int(mar.get(qn("w:left")))
                if mar.get(qn("w:right")):
                    margin_r = int(mar.get(qn("w:right")))

        TWIP_PER_CM = 1440 / 2.54
        TOL_TWIP = int(0.5 * TWIP_PER_CM)  # 0.5cm ≈ 283 twip
        EMU_PER_TWIP = 635  # 914400 EMU/in ÷ 1440 twip/in

        def _emu_int(s) -> int:
            try:
                return int(s)
            except Exception:
                return 0

        def _emu_to_twip(v: int) -> int:
            return int(round(v / EMU_PER_TWIP))

        # ---- 定位封面区域：第一处 sectPr 所在段落之前(含) ----
        first_sect_idx = None
        for i, ch in enumerate(self.body_children):
            if ch.tag.split("}")[-1] == "p":
                ppr = ch.find("w:pPr", NS)
                if ppr is not None and ppr.find("w:sectPr", NS) is not None:
                    first_sect_idx = i
                    break
        if first_sect_idx is None:
            first_sect_idx = min(len(self.body_children) - 1, 5)

        def _drawing_anchor_centered(drawing_el) -> bool:
            """DrawingML 浮动文本框的 <wp:positionH> 是否在参考系水平居中。"""
            anchor = drawing_el.find(".//wp:anchor", NSX)
            if anchor is None:
                return False
            posH = anchor.find("wp:positionH", NSX)
            if posH is None:
                return False
            # (i) 显式 align=center
            align_el = posH.find("wp:align", NSX)
            if align_el is not None and (align_el.text or "").strip().lower() == "center":
                return True
            # (ii) posOffset + extent 数值居中
            off_el = posH.find("wp:posOffset", NSX)
            extent = anchor.find("wp:extent", NSX)
            if off_el is None or extent is None:
                return False
            off_twip = _emu_to_twip(_emu_int((off_el.text or "0").strip()))
            cx_twip = _emu_to_twip(_emu_int(extent.get("cx") or "0"))
            rel = (posH.get("relativeFrom") or "").lower()
            if rel == "margin" or rel == "column":
                ref_left, ref_right = margin_l, page_w - margin_r
            elif rel == "leftmargin":
                ref_left, ref_right = 0, margin_l
            elif rel == "rightmargin":
                ref_left, ref_right = page_w - margin_r, page_w
            elif rel == "character":
                return False  # 相对字符位置，无法判定居中
            else:  # page / outsideMargin / insideMargin / 未声明
                ref_left, ref_right = 0, page_w
            ref_center = (ref_left + ref_right) / 2
            box_center = ref_left + off_twip + cx_twip / 2
            return abs(box_center - ref_center) <= TOL_TWIP

        def _vml_shape_centered(pict_el) -> bool:
            """VML 形状：style 中 mso-position-horizontal:center，或 left+width 居中。"""
            def _parse_len_twip(s: str):
                s = (s or "").strip().lower()
                if not s:
                    return None
                num, i = "", 0
                if i < len(s) and s[i] in "+-":
                    num += s[i]; i += 1
                while i < len(s) and (s[i].isdigit() or s[i] == "."):
                    num += s[i]; i += 1
                try:
                    val = float(num)
                except Exception:
                    return None
                unit = s[i:].strip()
                if unit in ("", "pt"):
                    return val * 20  # 1pt = 20 twip
                if unit == "in":
                    return val * 1440
                if unit == "cm":
                    return val * TWIP_PER_CM
                if unit == "mm":
                    return val * TWIP_PER_CM / 10
                if unit == "px":
                    return val * 15  # 96dpi: 1px = 0.75pt = 15 twip
                return val * 20

            for shape in pict_el.iter():
                tag = shape.tag.split("}")[-1]
                if tag not in ("shape", "rect", "roundrect", "oval", "group"):
                    continue
                if shape.find(".//w:txbxContent", NS) is None:
                    continue
                style = shape.get("style") or ""
                props = {}
                for part in style.split(";"):
                    if ":" in part:
                        k, v = part.split(":", 1)
                        props[k.strip().lower()] = v.strip().lower()
                if props.get("mso-position-horizontal") == "center":
                    return True
                left = _parse_len_twip(props.get("left", ""))
                width = _parse_len_twip(props.get("width", ""))
                if left is None or width is None:
                    continue
                origin = props.get("mso-position-horizontal-relative", "page")
                if origin == "margin":
                    ref_left, ref_right = margin_l, page_w - margin_r
                else:
                    ref_left, ref_right = 0, page_w
                ref_center = (ref_left + ref_right) / 2
                box_center = ref_left + left + width / 2
                if abs(box_center - ref_center) <= TOL_TWIP:
                    return True
            return False

        def _txbx_has_visible_text(node) -> bool:
            for tb in node.findall(".//w:txbxContent", NS):
                for t in tb.findall(".//w:t", NS):
                    if (t.text or "").strip():
                        return True
            return False

        for ch in self.body_children[:first_sect_idx + 1]:
            if ch.tag.split("}")[-1] != "p":
                continue

            # 收集文本框内的 w:t，便于 A 项排除文本框文字
            txbx_texts = set()
            for tb in ch.findall(".//w:txbxContent", NS):
                for t in tb.findall(".//w:t", NS):
                    txbx_texts.add(id(t))

            # B) DrawingML 浮动文本框：由 wp:positionH 精确判居中
            for drw in ch.findall(".//w:drawing", NS):
                if drw.find(".//wp:anchor", NSX) is None:
                    continue
                if not _txbx_has_visible_text(drw):
                    continue
                if _drawing_anchor_centered(drw):
                    return False

            # C) DrawingML 内嵌文本框：位置随段落，段落 jc=center 才算居中
            for drw in ch.findall(".//w:drawing", NS):
                if drw.find(".//wp:inline", NSX) is None:
                    continue
                if not _txbx_has_visible_text(drw):
                    continue
                if get_para_alignment(ch) == "center":
                    return False

            # D) VML 文本框
            for pict in ch.findall(".//w:pict", NS):
                if not _txbx_has_visible_text(pict):
                    continue
                if _vml_shape_centered(pict):
                    return False

            # A) 段落自身水平居中且含可见文字（排除文本框内文字）
            direct = "".join(t.text or "" for t in ch.findall(".//w:t", NS)
                             if id(t) not in txbx_texts)
            if direct.strip() and get_para_alignment(ch) == "center":
                return False

        return True

    def check_page2_missing_fields(self) -> bool:
        """-1：第2页没有出现 云岭应用学院、项目名称、项目负责人、团队成员、指导老师、申报日期、项目类型

        细则：第 2 页缺少上述 7 个字段中的任意一个，即扣分（必须 7 个全部出现才不扣）。
        判定：取"第 2 页"区域文本，逐一检查 7 个字段是否出现；只要有任一缺失即扣分。

        在办公软件（WPS/Word）中的有效判定：
          - 页面由分节/分页确定。封面为第 1 节（第一处 sectPr 之前），第 2 页信息页为
            第 1 处 sectPr 之后、第 2 处 sectPr（含）之前的区域。仅统计该区域文本，
            不把封面（第 1 页）内容算进来。
          - 文字可能位于普通段落，也可能位于文本框(txbxContent)；两者都要纳入取词，
            与办公软件中"该页可见文字"一致。
        """
        sect_prs_in_paras = []
        for i, ch in enumerate(self.body_children):
            if ch.tag.split("}")[-1] == "p":
                ppr = ch.find("w:pPr", NS)
                if ppr is not None and ppr.find("w:sectPr", NS) is not None:
                    sect_prs_in_paras.append(i)
        if len(sect_prs_in_paras) < 2:
            # 没有第 2 节，无法构成"第 2 页信息页" → 视为该页缺失，扣分
            return True
        first_sect_idx = sect_prs_in_paras[0]
        second_sect_idx = sect_prs_in_paras[1]

        # 仅取第 2 页区域（第 1 处 sectPr 之后 到 第 2 处 sectPr 含）的文本，
        # 含普通段落与文本框内文字。
        page2_text = ""
        for ch in self.body_children[first_sect_idx + 1:second_sect_idx + 1]:
            if ch.tag.split("}")[-1] == "p":
                page2_text += para_text(ch) + "\n"

        fields = ["云岭应用学院", "项目名称", "项目负责人", "团队成员", "指导老师",
                  "申报日期", "项目类型"]
        missing = [f for f in fields if f not in page2_text]
        return len(missing) > 0

    def find_table_after_paragraph(self, anchor_text: str):
        """找到包含 anchor_text 的段落之后的第一个表格。"""
        anchor_idx = None
        for i, ch in enumerate(self.body_children):
            if ch.tag.split("}")[-1] == "p" and anchor_text in para_text(ch):
                anchor_idx = i
                break
        if anchor_idx is None:
            return None
        for ch in self.body_children[anchor_idx + 1:]:
            if ch.tag.split("}")[-1] == "tbl":
                return ch
        return None

    def check_9_3_table_shape(self) -> bool:
        """-3：“9.3 服务需求与年龄的交叉”下方表格不是五行六列、第一行没有出现“手作文化”

        扣分条件（满足任意一项即扣 3 分）：
          A) 该表格不是"五行六列"（行数≠5 或 列数≠6）
          B) 该表格第一行没有出现"手作文化"
        另：若 9.3 下方根本没有表格，视为不满足"五行六列/含手作文化"，同样扣分。

        在办公软件（WPS/Word）中的有效判定：
          - 行数 = <w:tr> 的个数。
          - 列数：以表格网格 <w:tblGrid>/<w:gridCol> 的列数为准（这是办公软件里表格的
            真实列数）；同时校验第一行按 gridSpan 展开后的逻辑列数一致，避免被合并单元格
            误导。两者取能反映"六列"的值来判断。
          - "手作文化"：取第一行所有单元格文本合并后是否包含该词。
        """
        tbl = self.find_table_after_paragraph("9.3 服务需求与年龄的交叉")
        if tbl is None:
            return True  # 没有表格 → 不满足要求，扣分
        rows = tbl.findall("w:tr", NS)

        # A) 五行六列
        # 行数必须为 5
        if len(rows) != 5:
            return True
        # 列数：优先按 tblGrid 的 gridCol 数；无 tblGrid 时按第一行 gridSpan 展开
        grid = tbl.find("w:tblGrid", NS)
        if grid is not None:
            n_cols = len(grid.findall("w:gridCol", NS))
        else:
            n_cols = 0
            for tc in rows[0].findall("w:tc", NS):
                tcpr = tc.find("w:tcPr", NS)
                gs = tcpr.find("w:gridSpan", NS) if tcpr is not None else None
                n_cols += int(gs.get(qn("w:val"), "1")) if gs is not None else 1
        if n_cols != 6:
            return True

        # B) 第一行是否出现"手作文化"
        first_text = "".join(t.text or "" for t in rows[0].findall(".//w:t", NS))
        if "手作文化" not in first_text:
            return True

        return False

    def check_chapter1_violations(self) -> bool:
        """-5：满足以下任意一项即扣 5 分：
        1) 文档中没有出现“第一章市场分析”标题及下方文本
        2) 文档出现相同的两个标题，并且标题下方文本有超过 50% 的内容一致
        3) “第一章市场分析”位于“第二章项目简介”之后

        三点为 OR 关系（任一满足即扣分），均针对办公软件文档实际内容判定。
        本文档章节以"一、市场分析""二、项目简介"形式书写，等价于"第一章市场分析"
        "第二章项目简介"，判定时两种写法都接受。标题以段落文本识别（办公软件里
        章节标题即独立段落），下方文本 = 该标题段之后、下一个同级标题段之前的正文。
        """
        CH1 = ("第一章市场分析", "一、市场分析")
        CH2 = ("第二章项目简介", "二、项目简介")

        def is_l1_title(t: str) -> bool:
            return (len(t) < 40
                    and (re.match(r"^[一二三四五六七八九十]+[、,．.]", t) is not None
                         or re.match(r"^第[一二三四五六七八九十]+章", t) is not None))

        # 章节标题段落索引（按文档顺序）
        title_indices = [i for i, p in enumerate(self.paras)
                         if is_l1_title(para_text(p).strip())]

        def find_first(idx_names) -> Optional[int]:
            for i in title_indices:
                t = para_text(self.paras[i]).strip()
                if any(name in t for name in idx_names):
                    return i
            return None

        def body_below(title_i: int) -> str:
            """标题段之后、到下一个 L1 标题段之前的正文（下方文本）。"""
            chunk = []
            for j in range(title_i + 1, len(self.paras)):
                tj = para_text(self.paras[j]).strip()
                if is_l1_title(tj):
                    break
                chunk.append(tj)
            return "".join(chunk)

        ch1_i = find_first(CH1)
        ch2_i = find_first(CH2)

        # 1) 没有"第一章市场分析"标题，或其下方没有文本
        if ch1_i is None:
            return True
        if not body_below(ch1_i).strip():
            return True

        # 3) "第一章市场分析" 位于 "第二章项目简介" 之后（按标题出现顺序）
        if ch2_i is not None and ch1_i > ch2_i:
            return True

        # 2) 出现相同的两个标题，且两处标题下方文本 > 50% 一致
        title_positions: Dict[str, List[int]] = {}
        for i in title_indices:
            t = para_text(self.paras[i]).strip()
            title_positions.setdefault(t, []).append(i)
        for title, positions in title_positions.items():
            if len(positions) < 2:
                continue
            a = body_below(positions[0])
            b = body_below(positions[1])
            if not a or not b:
                continue
            longer, shorter = (a, b) if len(a) >= len(b) else (b, a)
            # 以较短一方为基准的最长公共子序列占比衡量"内容一致度"
            ratio = len(_lcs(shorter, longer)) / len(shorter)
            if ratio > 0.5:
                return True

        return False

    def _toc_entries(self) -> list:
        """提取目录(TOC)的缓存条目文本，按目录中出现顺序返回。

        办公软件（WPS/Word）里目录是一个 TOC 域，其可见内容被缓存为若干段落，
        通常包裹在内容控件 <w:sdt>/<w:sdtContent> 中，每个目录项占一个段落，
        末尾带页码，如"一、市场分析2"。这里定位含 TOC 指令(instrText 里的 'TOC'
        或 fldSimple 的 instr 含 'TOC')的 sdt，取其内部段落文本作为目录条目。
        找不到 sdt 包裹时，回退到含 TOC 域的段落所在区域。
        """
        for sdt in self.root.findall(".//w:sdt", NS):
            is_toc = False
            for it in sdt.findall(".//w:instrText", NS):
                if it.text and "TOC" in it.text:
                    is_toc = True
                    break
            if not is_toc:
                for fs in sdt.findall(".//w:fldSimple", NS):
                    if "TOC" in (fs.get(qn("w:instr")) or ""):
                        is_toc = True
                        break
            if is_toc:
                content = sdt.find("w:sdtContent", NS)
                scope = content if content is not None else sdt
                return [para_text(p) for p in scope.findall(".//w:p", NS)]
        return []

    def check_toc_chapter1_violations(self) -> bool:
        """-3：文档目录部分满足以下任意一项即扣 3 分：
        1) 目录中“第一章市场分析”出现在“第二章项目简介”之后
        2) 目录中“第一章市场分析”的页码不是 1

        判定针对办公软件中的目录(TOC 域缓存条目)。本文档章节以"一、市场分析"
        "二、项目简介"书写，与"第一章市场分析""第二章项目简介"等价，两种写法都接受。
        目录每个条目为一段，形如"一、市场分析2"（末尾数字为页码）。
        """
        entries = self._toc_entries()
        if not entries:
            return False  # 无目录域可判定时不据此扣分

        CH1 = ("第一章市场分析", "一、市场分析")
        CH2 = ("第二章项目简介", "二、项目简介")

        def find_entry(names):
            for i, e in enumerate(entries):
                s = e.strip()
                if any(n in s for n in names):
                    return i, s
            return None, None

        i1, e1 = find_entry(CH1)
        i2, e2 = find_entry(CH2)

        # 目录中没有第一章条目，无法满足"页码为1/在第二章之前"，按违规处理
        if i1 is None:
            return True

        # 1) 目录顺序：第一章在第二章之后
        if i2 is not None and i1 > i2:
            return True

        # 2) 第一章页码不是 1：取该目录条目末尾的页码数字
        nums = re.findall(r"\d+", e1)
        page_num = int(nums[-1]) if nums else None
        if page_num != 1:
            return True

        return False

    def _style_first_line_chars(self, style_id):
        """沿段落样式的 basedOn 链，取该样式定义的首行缩进字符数 firstLineChars。"""
        if self.styles_root is None:
            style_id = None
        by_id = {}
        if self.styles_root is not None:
            for s in self.styles_root.findall("w:style", NS):
                sid = s.get(qn("w:styleId"))
                if sid is not None:
                    by_id[sid] = s
        # 无显式样式时，回退到默认段落样式(w:default=1 且 type=paragraph)
        cur = style_id
        if cur is None and self.styles_root is not None:
            for s in self.styles_root.findall("w:style", NS):
                if s.get(qn("w:type")) == "paragraph" and s.get(qn("w:default")) == "1":
                    cur = s.get(qn("w:styleId"))
                    break
        seen = set()
        while cur and cur not in seen:
            seen.add(cur)
            s = by_id.get(cur)
            if s is None:
                break
            ind = s.find(".//w:pPr/w:ind", NS)
            if ind is not None:
                flc = ind.get(qn("w:firstLineChars"))
                if flc is not None:
                    try:
                        return float(flc)
                    except ValueError:
                        return None
            based = s.find("w:basedOn", NS)
            cur = based.get(qn("w:val")) if based is not None else None
        return None

    def check_10_5_10_6_indent(self) -> bool:
        """-1：“10.5 决策行为”与“10.6 核心顾虑”之间的文本首行缩进不是2字符

        细则：介于"10.5 决策行为"与"10.6 核心顾虑"两个小节标题之间的正文文本，
        其首行缩进若不是 2 字符，即扣分。只要该区间存在任一正文段落的首行缩进≠2字符，
        即满足扣分条件。

        在办公软件（WPS/Word）中的有效判定：
          - "2 字符"首行缩进的规范表示是 w:ind 的 w:firstLineChars="200"
            （字符单位以百分之一字符计，200 = 2 字符）。这是与字号无关的"按字符"缩进，
            与固定磅值 w:firstLine（twips）不同。
          - 段落可能未直接设置缩进，而是**继承段落样式**（本文档默认 Normal 样式即定义了
            firstLineChars=200）。因此当段落自身无 firstLineChars 时，要沿其样式(含默认
            段落样式)的 basedOn 链解析继承值，才能得到办公软件里实际渲染的缩进。
        """
        idx_105 = None
        idx_106 = None
        for i, p in enumerate(self.paras):
            t = para_text(p).strip()
            if "10.5 决策行为" in t:
                idx_105 = i
            elif "10.6 核心顾虑" in t:
                idx_106 = i
        # 找不到两个标题或之间没有正文，无法判定 → 不扣分
        if idx_105 is None or idx_106 is None or idx_106 <= idx_105 + 1:
            return False

        for i in range(idx_105 + 1, idx_106):
            p = self.paras[i]
            if not para_text(p).strip():
                continue  # 空段落不算正文
            ind = get_para_indent(p)
            chars = ind.get("first_line_chars")
            if chars is None:
                # 段落未直接设置：取样式继承的首行缩进字符数
                ppr = p.find("w:pPr", NS)
                ps = ppr.find("w:pStyle", NS) if ppr is not None else None
                chars = self._style_first_line_chars(ps.get(qn("w:val")) if ps is not None else None)
            # 首行缩进不是 2 字符(firstLineChars=200) → 违规
            if chars is None or abs(chars - 200) > 1:
                return True
        return False

    # ---------- 主流程 ----------
    # 维度 2 评分项定义：(方法名, 规则描述, 分值增量)
    # 正数为得分点（命中即加分），负数为扣分点（命中即扣分）
    DIM2_RULES: List[Tuple[str, str, int]] = [
        ("check_body_line_spacing_1x", "文档中正文文本间距为1倍", 3),
        ("check_table_line_spacing_1_25", "所有表格中的文本行距均为1.25倍", 5),
        ("check_l1_heading_style", "一级标题为宋体、小三号、单倍行距；段前0.5行、段后0.5行", 3),
        ("check_l2_heading_style", "二级标题为宋体、四号、1.5倍行距", 3),
        ("check_l3_heading_style", "三级标题为宋体、小四、单倍行距、加粗", 3),
        ("check_tourism_table_two_pages_with_header", "“（二）旅游资源分析”表格排版为2页", 5),
        ("check_tourism_table_three_borders_per_page", "“（二）旅游资源分析”表格每一页只有三条边框线", 1),
        ("check_only_one_table_two_pages", "仅有一个表格分2页排版，其他表格均为一页排版", 5),
        ("check_any_page_no_page_number", "任意页无页码页脚", -1),
        ("check_cover_no_center_text", "封面居中位置没有文本内容", -1),
        ("check_page2_missing_fields",
         "第2页没有出现云岭应用学院、项目名称、项目负责人、团队成员、指导老师、申报日期、项目类型", -1),
        ("check_9_3_table_shape",
         "“9.3 服务需求与年龄的交叉”下方表格不是五行六列、第一行没有出现“手作文化”", -3),
        ("check_chapter1_violations",
         "满足以下任意一项：文档中没有出现“第一章市场分析”标题及下方文本；文档出现相同的两个标题并且标题下方文本有超过50%的内容一致；“第一章市场分析”位于“第二章项目简介”之后",
         -5),
        ("check_toc_chapter1_violations",
         "文档中的目录部分：出现“第一章市场分析”在“第二章项目简介”之后或者“第一章市场分析”的页码不是1",
         -3),
        ("check_10_5_10_6_indent",
         "“10.5 决策行为”与“10.6 核心顾虑”之间的文本首行缩进不是2字符", -1),
    ]

    def run(self) -> Dict:
        file_name = os.path.basename(self.path)
        # 总分 = 所有加分项之和（只计正分值，不含扣分项）
        max_score = sum(delta for _, _, delta in self.DIM2_RULES if delta > 0)

        if not self.check_dimension1():
            self.dimension1_pass = False
            return {
                "id": "025",
                "file_name": file_name,
                "status": "ok",
                "error": None,
                "dim1_pass": False,
                "dim1_reason": "；".join(self.dimension1_reasons),
                "dim2_items": [],
                "total_score": 0,
                "max_score": max_score,
            }

        dim2_items: List[dict] = []
        total_score = 0
        for method_name, rule_desc, delta in self.DIM2_RULES:
            hit = bool(getattr(self, method_name)())
            actual_delta = delta if hit else 0
            total_score += actual_delta
            dim2_items.append({
                "rule": rule_desc,
                "max_delta": delta,
                "delta": actual_delta,
                "hit": hit,
                "detail": "",
            })

        return {
            "id": "025",
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": True,
            "dim1_reason": "",
            "dim2_items": dim2_items,
            "total_score": total_score,
            "max_score": max_score,
        }


def _locate_docx(dir_path: str) -> Optional[str]:
    """在给定目录中定位待评估的 .docx 文件。

    约定：批量 runner 传入"脚本所在目录路径"，脚本自己在该目录下找唯一的 docx。
    过滤 Office 打开时留下的临时锁文件（以 ~$ 开头）。目录内若存在多个 .docx，
    优先取文件名与目录名最接近的一个，避免选到备份文件。
    """
    if not dir_path or not os.path.isdir(dir_path):
        return None
    candidates = [n for n in os.listdir(dir_path)
                  if n.lower().endswith(".docx") and not n.startswith("~$")]
    if not candidates:
        return None
    if len(candidates) == 1:
        return os.path.join(dir_path, candidates[0])
    # 多个候选：优先取不含"备份/副本/backup/copy"关键字的
    def prefer(name: str) -> int:
        lower = name.lower()
        return 0 if any(k in lower or k in name for k in
                        ("备份", "副本", "backup", "copy")) else 1
    candidates.sort(key=lambda n: (-prefer(n), len(n)))
    return os.path.join(dir_path, candidates[0])


def evaluate(dir_path: str) -> dict:
    """统一评估入口：传入脚本所在目录路径，脚本自行在该目录中定位并评估 docx。

    返回结构见《脚本接口差异与统一建议.md》§2.2。
    发生异常时返回 status="error"，不抛出，不 print 主结果。
    """
    max_score = sum(delta for _, _, delta in Evaluator.DIM2_RULES if delta > 0)
    file_path = _locate_docx(dir_path)
    file_name = os.path.basename(file_path) if file_path else ""
    if file_path is None:
        return {
            "id": "025",
            "file_name": "",
            "status": "error",
            "error": f"未在目录中找到 .docx 文件: {dir_path!r}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": max_score,
        }
    try:
        return Evaluator(file_path).run()
    except (etree.XMLSyntaxError, zipfile.BadZipFile) as e:
        return {
            "id": "025",
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": False,
            "dim1_reason": f"文档结构损坏，无法正常解析：{type(e).__name__}: {e}",
            "dim2_items": [],
            "total_score": 0,
            "max_score": max_score,
        }
    except Exception as e:
        return {
            "id": "025",
            "file_name": file_name,
            "status": "error",
            "error": f"{type(e).__name__}: {e}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": max_score,
        }


if __name__ == "__main__":
    # 仅用于本地调试：从命令行读取目录路径并打印 JSON 结果；
    # 未传参时默认使用脚本自身所在目录。
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    # ensure_ascii=True：中文以 \uXXXX 转义，避免 Windows cp1252 控制台崩溃
    print(json.dumps(evaluate(_dir), ensure_ascii=True, indent=2))
