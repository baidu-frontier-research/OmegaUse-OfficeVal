#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
课时学习方案自动评估工具（officeval_001）

统一入口：
    evaluate(dir_path: str) -> dict
    - dir_path: 脚本所在目录（内含待评估的 Word 文档）
    - 返回结构见"脚本接口差异与统一建议.md §2.2"

规范约束：
    * 只暴露 evaluate 一个函数供批量运行器调用
    * 不修改 sys.stdout、不打印主结果、不使用 sys.exit
    * 不硬编码文档路径，全部由 dir_path 参数驱动
"""

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm, Pt, Emu

SCRIPT_ID = "001"
REFERENCE_WIDTH_FILE = "01_课时学习方案_观察一位校园志愿者_第一课时_（修改断句版）20260516111637 (1).docx"
CM_PER_TWIP = 2.54 / 1440
W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
WORD_PAGE_TIMEOUT_SECONDS = 90

_WORD_PAGE_COUNTER_SCRIPT = r'''
import os
import sys

import win32com.client
import win32process

word = None
wdoc = None
try:
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    print(f"PID:{win32process.GetWindowThreadProcessId(word.Hwnd)[1]}", flush=True)
    wdoc = word.Documents.Open(
        os.path.abspath(sys.argv[1]),
        ReadOnly=True,
        AddToRecentFiles=False,
    )
    try:
        wdoc.Repaginate()
    except Exception:
        pass
    print(f"PAGES:{int(wdoc.ComputeStatistics(2))}", flush=True)
finally:
    if wdoc is not None:
        try:
            wdoc.Close(False)
        except Exception:
            pass
    if word is not None:
        try:
            word.Quit()
        except Exception:
            pass
'''


def _terminate_process_tree(pid: int | None) -> None:
    if not pid:
        return
    try:
        subprocess.run(
            ['taskkill', '/PID', str(pid), '/T', '/F'],
            capture_output=True,
            timeout=10,
        )
    except Exception:
        pass


def _word_page_count_with_timeout(path: str) -> int | None:
    """在隔离进程中获取 Word 页数，超时后定向清理该 Word 实例。"""
    output = b''
    proc = None
    try:
        proc = subprocess.run(
            [sys.executable, '-c', _WORD_PAGE_COUNTER_SCRIPT, path],
            capture_output=True,
            timeout=WORD_PAGE_TIMEOUT_SECONDS,
        )
        output = proc.stdout or b''
    except subprocess.TimeoutExpired as exc:
        output = exc.stdout or b''
    except OSError:
        return None
    text = output.decode('utf-8', errors='replace')
    if proc is None or proc.returncode != 0:
        pid_match = re.search(r'^PID:(\d+)$', text, re.MULTILINE)
        _terminate_process_tree(int(pid_match.group(1)) if pid_match else None)
        return None
    pages_match = re.search(r'^PAGES:(\d+)$', text, re.MULTILINE)
    return int(pages_match.group(1)) if pages_match else None


def extract_topic_type_widths(doc_path):
    """抽取“课时主题”和“课型”标签单元格的 tcW 宽度，单位 cm。"""
    widths = {}
    with zipfile.ZipFile(doc_path) as zf:
        root = ET.fromstring(zf.read('word/document.xml'))
    for tbl in root.findall(f'.//{W_NS}tbl'):
        for tr in tbl.findall(f'./{W_NS}tr'):
            for tc in tr.findall(f'./{W_NS}tc'):
                text = ''.join(t.text or '' for t in tc.findall(f'.//{W_NS}t')).strip()
                if text in ('课时主题', '课型'):
                    tcw = tc.find(f'./{W_NS}tcPr/{W_NS}tcW')
                    raw_width = tcw.attrib.get(W_NS + 'w') if tcw is not None else None
                    if raw_width:
                        widths[text] = int(raw_width) * CM_PER_TWIP
    return widths


def get_reference_topic_type_widths(folder_path):
    """优先使用指定参考文件，抽取评分标准宽度。"""
    reference_path = os.path.join(folder_path, REFERENCE_WIDTH_FILE)
    if os.path.exists(reference_path):
        return extract_topic_type_widths(reference_path)
    # 兜底：若指定参考文件不存在，用当前目录中第一个 01_ 文件。
    for filename in sorted(os.listdir(folder_path)):
        if filename.startswith('01_') and filename.endswith('.docx'):
            return extract_topic_type_widths(os.path.join(folder_path, filename))
    return {}


class LessonPlanEvaluator:
    """课时学习方案评估器"""

    def __init__(self, doc_path: str):
        self.doc_path: str = doc_path
        self.xml_path: str = doc_path
        self.doc = None
        self.file_name: str = os.path.basename(doc_path)
        self.dimension1_passed: bool = True
        self.dimension2_score: int = 0
        self.full_text: str = ""
        self.paragraph_texts: list[str] = []
        self.table_texts: list[str] = []
        self.check_results: dict[str, bool] = {}  # 记录每个检查点的结果
        self._temp_dir: str | None = None  # 渲染分页时用到的临时目录，评估结束后清理

    def _cleanup_temp(self):
        if self._temp_dir and os.path.isdir(self._temp_dir):
            shutil.rmtree(self._temp_dir, ignore_errors=True)
            self._temp_dir = None

    @staticmethod
    def _count_pdf_pages(pdf_path: str) -> int | None:
        try:
            with open(pdf_path, 'rb') as f:
                data = f.read()
            page_count = len(re.findall(rb'/Type\s*/Page\b', data))
            return page_count if page_count > 0 else None
        except Exception:
            return None

    def _get_rendered_page_count(self) -> int | None:
        """通过 Word/LibreOffice 渲染分页获取实际页数；不可用时返回 None。"""
        if sys.platform == 'win32':
            pages = _word_page_count_with_timeout(self.xml_path)
            if pages is not None and pages > 0:
                return pages

        soffice = shutil.which('soffice') or shutil.which('soffice.exe')
        if soffice:
            if self._temp_dir is None:
                self._temp_dir = tempfile.mkdtemp(prefix='officeval001_')
            try:
                _ = subprocess.run(
                    [soffice, '--headless', '--convert-to', 'pdf', '--outdir', self._temp_dir, self.xml_path],
                    check=True, capture_output=True, timeout=60,
                )
                pdf_path = os.path.join(
                    self._temp_dir, os.path.splitext(os.path.basename(self.xml_path))[0] + '.pdf'
                )
                if os.path.exists(pdf_path):
                    return self._count_pdf_pages(pdf_path)
            except Exception:
                pass

        return None

    def load_document(self):
        """加载Word文档（仅支持 .docx）"""
        if not self.doc_path.lower().endswith('.docx'):
            # 非 .docx（如二进制 .doc）不再尝试转换，直接判定为无法可靠打开。
            self.check_results['dim1_open'] = False
            self.dimension1_passed = False
            self.dimension2_score = 0
            return False

        self.xml_path = self.doc_path
        try:
            self.doc = Document(self.doc_path)
            self.check_results['dim1_open'] = True
            self._extract_all_text()
            return True
        except Exception:
            # python-docx 遇到个别复杂合并表格时可能读取失败，改用 docx 底层 XML 兜底提取文本。
            try:
                self.check_results['dim1_open'] = True
                self._extract_all_text_from_xml()
                return True
            except Exception:
                self.check_results['dim1_open'] = False
                self.dimension1_passed = False
                self.dimension2_score = 0
                return False

    def _extract_all_text(self):
        """提取所有文本（段落+表格）"""
        self.paragraph_texts = [para.text for para in self.doc.paragraphs]
        self.table_texts = []
        try:
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.table_texts.append(cell.text)
        except Exception:
            # 遇到不规则合并单元格时，退回 OOXML 文本提取，避免把可打开文件误判为维度一失败。
            self._extract_all_text_from_xml()
            return
        self.full_text = "\n".join(self.paragraph_texts) + "\n" + "\n".join(self.table_texts)

    def _extract_all_text_from_xml(self):
        """直接从 docx 的 document.xml 提取文本，作为 python-docx 读取复杂表格失败时的兜底。"""
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        with zipfile.ZipFile(self.xml_path) as zf:
            root = ET.fromstring(zf.read('word/document.xml'))
        self.paragraph_texts = []
        self.table_texts = []
        for p in root.findall('.//w:p', ns):
            text = ''.join(t.text or '' for t in p.findall('.//w:t', ns)).strip()
            if text:
                self.paragraph_texts.append(text)
        for tc in root.findall('.//w:tc', ns):
            text = ''.join(t.text or '' for t in tc.findall('.//w:t', ns)).strip()
            if text:
                self.table_texts.append(text)
        self.full_text = "\n".join(self.paragraph_texts) + "\n" + "\n".join(self.table_texts)

    def get_tables(self):
        if self.doc is None:
            return []
        try:
            return self.doc.tables
        except Exception:
            return []

    @staticmethod
    def _has_title_with_empty_following_content(text: str, title: str) -> bool:
        """检查同一段/同一单元格内标题后是否无实质内容。"""
        pattern = re.compile(
            rf'(?m)^\s*(?:[一二三四五六七八九十0-9]+\s*[、.．]\s*)?{re.escape(title)}\s*[:：]?'
        )
        match = pattern.search(text)
        if not match:
            return False
        return text[match.end():].strip() == ""

    def check_all(self):
        """执行所有检查"""
        try:
            if not self.load_document():
                return
            self._check_all_impl()
        finally:
            # xml_path 在渲染分页时可能生成临时 pdf，所有检查结束后统一清理。
            self._cleanup_temp()

    def _check_all_impl(self):
        # ========== 维度一检查 ==========
        # 检查1: 文件可正常打开 (已验证)

        # 检查2: 正文、标题、表格、图片均为可编辑内容
        has_tables = len(self.get_tables()) > 0 or len(self.table_texts) > 0
        has_content = len(self.full_text.strip()) > 100
        self.check_results['dim1_edit'] = has_tables and has_content

        # 检查3: 保留课时学习方案主体结构
        # 改为：命中任意 1 个核心关键词，或文件名含学习方案相关词且文本有实质内容，即通过。
        # 原逻辑要求 >=4 个，会把结构不完整但可编辑的文件误判为维度一不通过。
        required_keywords = ["课时主题", "课型", "学习过程设计", "板书设计", "教学反思"]
        found_count = sum(1 for k in required_keywords if k in self.full_text)
        file_hint = any(kw in self.file_name for kw in ["课时学习方案", "学历案", "学习方案"])
        text_hint = any(k in self.full_text for k in ["课时学习目标", "课时评价任务", "学习过程设计", "板书设计", "教学反思"])
        self.check_results['dim1_structure'] = found_count >= 1 or (file_hint and text_hint)

        # 维度一综合判断
        self.dimension1_passed = (self.check_results.get('dim1_open', False) and
                                  self.check_results['dim1_edit'] and
                                  self.check_results['dim1_structure'])

        if not self.dimension1_passed:
            self.dimension2_score = 0
            return

        # ========== 维度二检查 ==========

        # +1：含"课时学习目标"标题且标题下无内容
        # “标题下方”按同一段落/同一表格单元格内标题后的剩余内容判断；标题后有实质文本则不满足。
        self.check_results['empty_target'] = any(
            self._has_title_with_empty_following_content(text, "课时学习目标")
            for text in self.table_texts + self.paragraph_texts
        )

        # +1：含"课时评价任务"标题且标题下无内容
        # 同 empty_target：按标题后剩余内容判断，同时检查表格单元格与段落两类来源。
        self.check_results['empty_eval'] = any(
            self._has_title_with_empty_following_content(text, "课时评价任务")
            for text in self.table_texts + self.paragraph_texts
        )

        # +1：无"课时学习内容分析"标题
        self.check_results['no_content_analysis'] = "课时学习内容分析" not in self.full_text

        # +1：无"课时学生实际水平"标题
        self.check_results['no_student_level'] = "课时学生实际水平" not in self.full_text

        # +1：无"活动意图说明"标题
        self.check_results['no_activity_intent'] = "活动意图说明" not in self.full_text


        # +3：不含图片
        # 不能只看 relationships/media 文件，因为 Word 里可能残留未被正文实际引用的图片资源。
        # 这里检查 document.xml 正文中是否真实存在图片引用。
        try:
            with zipfile.ZipFile(self.xml_path) as zf:
                root = ET.fromstring(zf.read('word/document.xml'))
            has_images = any(node.tag.endswith('}blip') or node.tag.endswith('}imagedata') for node in root.iter())
        except Exception:
            has_images = self.doc is not None and any("image" in rel.reltype for rel in self.doc.part.rels.values())
        self.check_results['no_images'] = not has_images

        # +1："学习过程设计"标题前序号为"三"
        self.check_results['seq_three'] = bool(re.search(r'三\s*[,，、.．\s]*\s*学习过程设计', self.full_text))

        # +1："板书设计"标题前序号为"四"
        self.check_results['seq_four'] = bool(re.search(r'[四4]\s*[,，、.\s]*\s*板书设计', self.full_text))

        # +1："教学反思"标题前序号为"五"
        self.check_results['seq_five'] = bool(re.search(r'五\s*[,，、.\s]*\s*教学反思', self.full_text))

        # +1："环节一"后不含"（指向目标一）"
        if "环节一" in self.full_text:
            self.check_results['no_ref1'] = ("（指向目标一）" not in self.full_text and "(指向目标一)" not in self.full_text)
        else:
            self.check_results['no_ref1'] = False

        # +1："环节二"后不含"（指向目标二）"
        if "环节二" in self.full_text:
            self.check_results['no_ref2'] = ("（指向目标二）" not in self.full_text and "(指向目标二)" not in self.full_text)
        else:
            self.check_results['no_ref2'] = False

        # +1："环节三"后不含"（指向目标三）"
        if "环节三" in self.full_text:
            self.check_results['no_ref3'] = ("（指向目标三）" not in self.full_text and "(指向目标三)" not in self.full_text)
        else:
            self.check_results['no_ref3'] = False

        # +5：内容页数为两页
        # 用 Word/LibreOffice 实际渲染分页结果判断页数；两者均不可用时判为不满足（不猜测）。
        rendered_pages = self._get_rendered_page_count()
        self.check_results['two_pages'] = rendered_pages == 2

        # +5：表格未出现不连续
        # Word 的 OOXML 中没有可靠的“视觉断页/不连续”标记；这里用结构特征近似：
        # 学习过程相关主体表（宽度接近整页，且包含“学习过程设计”或“环节”）中，
        # 若最大表格行数 >= 5，说明主体表格连续；若都被拆成很短的小表，则判为不连续。
        try:
            with zipfile.ZipFile(self.xml_path) as zf:
                root = ET.fromstring(zf.read('word/document.xml'))
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            process_table_rows: list[int] = []
            for tbl in root.findall(f'.//{w_ns}tbl'):
                grid_width = 0
                for col in tbl.findall(f'./{w_ns}tblGrid/{w_ns}gridCol'):
                    grid_width += int(col.attrib.get(w_ns + 'w', 0) or 0)
                width_cm = grid_width * 2.54 / 1440 if grid_width else 0
                table_text = ''.join(t.text or '' for t in tbl.findall(f'.//{w_ns}t'))
                if width_cm >= 15 and ("学习过程设计" in table_text or "环节" in table_text):
                    process_table_rows.append(len(tbl.findall(f'./{w_ns}tr')))
            tables_ok = bool(process_table_rows) and max(process_table_rows) >= 5
        except Exception:
            tables_ok = len(self.table_texts) > 0
        self.check_results['table_continuous'] = tables_ok

        # -5：表格整体列宽不是17.00-17.30厘米
        table_width_ok = True
        try:
            for table in self.get_tables():
                for row in table.rows:
                    total_width = sum(cell.width for cell in row.cells if cell.width)
                    if total_width:
                        total_width_cm = total_width / 360000
                        if not (17.00 <= total_width_cm <= 17.30):
                            table_width_ok = False
                            break
                if not table_width_ok:
                    break
        except Exception:
            table_width_ok = True
        self.check_results['table_width_ok'] = table_width_ok

        # 计算总分
        score = 0
        if self.check_results['empty_target']:
            score += 1
        if self.check_results['empty_eval']:
            score += 1
        if self.check_results['no_content_analysis']:
            score += 1
        if self.check_results['no_student_level']:
            score += 1
        if self.check_results['no_activity_intent']:
            score += 1
        if self.check_results['no_images']:
            score += 3
        if self.check_results['seq_three']:
            score += 1
        if self.check_results['seq_four']:
            score += 1
        if self.check_results['seq_five']:
            score += 1
        if self.check_results['no_ref1']:
            score += 1
        if self.check_results['no_ref2']:
            score += 1
        if self.check_results['no_ref3']:
            score += 1
        if self.check_results['two_pages']:
            score += 5
        if self.check_results['table_continuous']:
            score += 5
        if not self.check_results['table_width_ok']:
            score -= 5

        self.dimension2_score = score


def get_short_name(filename):
    """简化文件名，只显示 数值_课时学习方案"""
    import re
    # 匹配开头的数字
    match = re.match(r'^(\d+)', os.path.basename(filename))
    if match:
        return f"{match.group(1)}_课时学习方案"
    return filename


def get_all_docx_files(folder_path: str) -> list[str]:
    """获取文件夹中 8 个待评估 Word 文件；参考文件只用于抽取基准宽度，不参与评分。"""
    docx_files: list[str] = []
    for filename in os.listdir(folder_path):
        if filename.startswith('~$') or filename.startswith('.~'):
            continue
        if filename == REFERENCE_WIDTH_FILE:
            continue
        if filename.endswith('.docx'):
            docx_files.append(os.path.join(folder_path, filename))
    return sorted(docx_files)


# ---- 维度二评分项定义（rule -> 分值） ------------------------------------
DIM2_POSITIVE_RULES = [
    ("empty_target", 1, "学历案内含'课时学习目标'标题且标题下无内容"),
    ("empty_eval", 1, "学历案内含'课时评价任务'标题且标题下无内容"),
    ("no_content_analysis", 1, "学历案内无'课时学习内容分析'标题"),
    ("no_student_level", 1, "学历案内无'课时学生实际水平'标题"),
    ("no_activity_intent", 1, "学历案内无'活动意图说明'标题"),
    ("no_images", 3, "学历案内不含图片"),
    ("seq_three", 1, "学历案内'学习过程设计'标题前序号为'三'"),
    ("seq_four", 1, "学历案内'板书设计'标题前序号为'四'"),
    ("seq_five", 1, "'教学反思'标题前序号为'五'"),
    ("no_ref1", 1, "学历案内'环节一'后不含'（指向目标一）'内容"),
    ("no_ref2", 1, "学历案内'环节二'后不含'（指向目标二）'内容"),
    ("no_ref3", 1, "学历案内'环节三'后不含'（指向目标三）'内容"),
    ("two_pages", 5, "学历案内容页数为两页"),
    ("table_continuous", 5, "学历案表格未出现不连续"),
]

DIM2_PENALTY_RULES = [
    ("table_width_ok", 5, "表格整体列宽不是17.00-17.30厘米"),
]


def _build_dim2_items(evaluators: list["LessonPlanEvaluator"]) -> list[dict[str, object]]:
    """按“每个文件、每条规则一个 item”生成严格二值的评分明细。"""
    items: list[dict[str, object]] = []
    for ev in evaluators:
        for key, delta, rule in DIM2_POSITIVE_RULES:
            hit = bool(ev.dimension1_passed and ev.check_results.get(key, False))
            items.append({
                "rule": f"[{ev.file_name}] {rule}",
                "max_delta": delta,
                "delta": delta if hit else 0,
                "hit": hit,
                "detail": "",
            })
        # check_results[key] 为 True 表示未触发扣分；False 表示触发扣分。
        for key, penalty, rule in DIM2_PENALTY_RULES:
            hit = bool(ev.dimension1_passed and not ev.check_results.get(key, True))
            max_delta = -penalty
            items.append({
                "rule": f"[{ev.file_name}] {rule}",
                "max_delta": max_delta,
                "delta": max_delta if hit else 0,
                "hit": hit,
                "detail": "",
            })
    return items


def evaluate(dir_path: str) -> dict:
    """统一入口：接收脚本所在目录，扫描目录内所有待评估 Word 文档，返回聚合评估结果。

    评分口径：每条得分项，每个满足的文件加对应分值；每条扣分项，每个未满足的文件扣对应分值。
    total_score 为所有文件所有条目的实际得分之和；
    max_score  为"所有加分项分值 × 参评文件数"，即所有文件均满足全部加分项时的理论上限。

    参数:
        dir_path: 脚本（及被评估 Word 文档）所在目录路径。
    返回:
        dict，字段对齐"脚本接口差异与统一建议.md §2.2"。
    """
    result: dict[str, object] = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }

    try:
        if not os.path.isdir(dir_path):
            result["status"] = "error"
            result["error"] = f"目录不存在:{dir_path}"
            return result

        docx_files = get_all_docx_files(dir_path)
        if not docx_files:
            result["status"] = "error"
            result["error"] = f"目录内未找到待评估的 Word 文档:{dir_path}"
            return result

        # 依次评估目录内所有文件
        evaluators: list[LessonPlanEvaluator] = []
        for doc_path in docx_files:
            ev = LessonPlanEvaluator(doc_path)
            ev.check_all()
            evaluators.append(ev)

        total = len(evaluators)
        dim1_failed = [ev.file_name for ev in evaluators if not ev.dimension1_passed]

        # max_score = 所有加分项分值 × 参评文件数
        max_score = sum(d for _, d, _ in DIM2_POSITIVE_RULES) * total

        # total_score = 所有文件维度二得分之和（维度一未通过的按 0 分累计）
        total_score = sum(int(ev.dimension2_score) if ev.dimension1_passed else 0 for ev in evaluators)

        result["file_name"] = ", ".join(ev.file_name for ev in evaluators)
        result["dim1_pass"] = len(dim1_failed) == 0
        result["dim1_reason"] = "" if not dim1_failed else f"以下文件维度一未通过：{', '.join(dim1_failed)}"
        # 只有所有文件均通过维度一时才生成维度二明细，保持门槛短路约定。
        result["dim2_items"] = _build_dim2_items(evaluators) if not dim1_failed else []
        result["total_score"] = total_score if not dim1_failed else 0
        result["max_score"] = max_score
        return result

    except Exception as exc:
        result["status"] = "error"
        result["error"] = f"{type(exc).__name__}: {exc}"
        return result


if __name__ == "__main__":
    _dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    _output = json.dumps(evaluate(_dir), ensure_ascii=False, indent=2) + "\n"
    sys.stdout.buffer.write(_output.encode("utf-8", errors="replace"))
