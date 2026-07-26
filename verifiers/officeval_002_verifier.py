#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# pyright: reportUnknownMemberType=false, reportUnknownVariableType=false, reportUnknownArgumentType=false, reportUnknownParameterType=false, reportMissingTypeArgument=false, reportImplicitStringConcatenation=false, reportDeprecated=false, reportExplicitAny=false, reportGeneralTypeIssues=false, reportArgumentType=false, reportAttributeAccessIssue=false, reportPrivateUsage=false, reportOptionalMemberAccess=false, reportOptionalSubscript=false, reportOptionalIterable=false, reportOptionalCall=false
"""
Word文档自动评估脚本
根据打分细则对Word文档进行自动评估
"""

import os
import sys
import re
import json
import glob
import contextlib
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Twips, Emu


class WordDocumentEvaluator:
    """Word文档评估器"""

    def __init__(self, file_path, skip_page_check=False):
        self.file_path = file_path
        # 供 python-docx 实际打开的路径；只支持 .docx，因此始终等于 file_path。
        self.open_path: str = file_path
        # 兼容旧调用签名保留 skip_page_check 形参，但维度一已删除页数检查，
        # 该参数在评估流程中不再被读取
        self.skip_page_check = skip_page_check
        self.doc = None
        # 穿透 w:sdt 的扁平文档索引（在维度2开始时构建）
        self._flat_paragraphs = []
        self._flat_para_elements = []
        self._flat_tables = []
        self._flat_sections = []
        self._flat_sect_pr_elements = []
        self._para_section_index = []
        self.chinese_abstract_para_idx = -1
        self.english_abstract_para_idx = -1
        self.outline_para_idx = -1
        # 分页检查固定使用节级静态回退（不依赖任何渲染引擎）。
        self._word_pagination = None  # 始终为 None，保留字段以兼容下游读取处
        self.results = {
            'dimension1_passed': True,
            'dimension1_failures': [],
            'dimension2_scores': [],
            'dimension2_deductions': [],
            'total_score': 0,
            'hit_points': []
        }

    def evaluate(self):
        """执行评估"""
        print("=" * 60)
        print("Word文档自动评估报告")
        print("=" * 60)
        print(f"文件: {os.path.basename(self.file_path)}")
        print()

        # 维度1评估
        if not self._evaluate_dimension1():
            self.results['dimension1_passed'] = False
            self._print_results()
            return

        # 维度2评估
        self._evaluate_dimension2()

        # 计算总分
        self._calculate_total_score()

        # 打印结果
        self._print_results()

    def _evaluate_dimension1(self):
        """维度1：可用与可修改性评估"""
        print("-" * 40)
        print("维度1：可用与可修改性")
        print("-" * 40)

        # 检查1：文件格式
        if not self._check_file_format():
            self.results['dimension1_failures'].append("文件格式不符合要求（需为.doc或.docx格式）")
            return False
        print("[PASS] 文件格式检查：.docx格式")

        # 检查2：文件可正常打开
        if not self._check_file_openable():
            self.results['dimension1_failures'].append("文件无法正常打开")
            return False
        print("[PASS] 文件打开检查：文件可正常打开")

        # 注：原维度一"页数=21"检查已按交付要求删除

        print()
        return True

    def _evaluate_dimension2(self):
        """维度2：完成度评估"""
        print("-" * 40)
        print("维度2：完成度")
        print("-" * 40)

        # 先深度遍历构建段落/表格文档序索引（穿透 w:sdt/w:sdtContent）
        self._build_flat_document_index()

        # 获取关键页面位置
        self._find_key_pages()

        # 初始化普通分页回退状态，不探测 Word COM。
        self._probe_word_pagination()

        # 得分点评估
        self._check_header_lines_before_abstract()  # +3
        self._check_header_lines_after_abstract()   # +5
        self._check_header_line_position()          # +5

        # 扣分点评估
        self._check_header_text()                   # -1
        self._check_footer_text()                   # -1
        self._check_page_number()                   # -1
        self._check_page_number_format()            # -1
        self._check_page_order()                    # -1
        self._check_heading1_format()               # -1
        self._check_heading2_format()               # -1
        self._check_heading3_format()               # -1
        self._check_paragraph_format()              # -1
        self._check_table_borders()                 # -3

    def _build_flat_document_index(self):
        """深度遍历 body 的顶层元素（穿透 w:sdt/w:sdtContent），
        构建按文档流顺序的段落/表格/节索引。作用域与 python-docx 的
        doc.paragraphs / doc.tables 一致：仅收集正文顶层段落与表格，
        不下沉到表格单元格内部。

        与 python-docx 内置遍历的不同：结构化文档标签（w:sdt/w:sdtContent）
        里的段落、表格、以及节末尾的 sectPr 会被穿透收集，从而覆盖被
        "内容控件"包裹的摘要/目录等标题段和被 sdt 隐藏的分节。

        产出：
          - self._flat_paragraphs: List[Paragraph]，按文档流顺序
          - self._flat_para_elements: List[<w:p>]，与上表一一对应
          - self._flat_tables: List[Table]，按文档流顺序
          - self._flat_sections: List[Section]，按文档流顺序；每个元素
            由文档流中一个 sectPr（无论嵌在 p/pPr 内还是 body 尾部，也
            无论是否被 sdt 包裹）构造而来。这是本脚本的"节"权威口径，
            所有维度2的节遍历必须使用它，而不是 self.doc.sections。
          - self._para_section_index: List[int]，与 _flat_paragraphs
            一一对应，记录该段落所属节的索引（与 _flat_sections 对齐）
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        from docx.section import Section

        self._flat_paragraphs = []
        self._flat_para_elements = []
        self._flat_tables = []
        self._flat_sections = []
        self._flat_sect_pr_elements = []
        self._para_section_index = []

        sect_idx = [0]  # 用列表包装以便在闭包内修改
        part = self.doc.part

        def walk(container):
            for child in container.iterchildren():
                tag = child.tag.split('}')[-1]
                if tag == 'p':
                    self._flat_paragraphs.append(Paragraph(child, self.doc))
                    self._flat_para_elements.append(child)
                    self._para_section_index.append(sect_idx[0])
                    pPr = child.find(qn('w:pPr'))
                    if pPr is not None:
                        sp = pPr.find(qn('w:sectPr'))
                        if sp is not None:
                            self._flat_sect_pr_elements.append(sp)
                            self._flat_sections.append(Section(sp, part))
                            sect_idx[0] += 1
                elif tag == 'tbl':
                    self._flat_tables.append(Table(child, self.doc))
                elif tag == 'sdt':
                    sdt_content = child.find(qn('w:sdtContent'))
                    if sdt_content is not None:
                        walk(sdt_content)
                elif tag == 'sectPr':
                    # body 末尾的 sectPr 对应最后一节
                    self._flat_sect_pr_elements.append(child)
                    self._flat_sections.append(Section(child, part))
                    sect_idx[0] += 1

        walk(self.doc.element.body)

    def _check_file_format(self):
        """检查文件格式"""
        ext = os.path.splitext(self.file_path)[1].lower()
        return ext == '.docx'

    def _check_file_openable(self):
        """检查文件可正常打开"""
        try:
            self.doc = Document(self.open_path)
            return True
        except Exception as e:
            print(f"打开文件失败: {e}")
            return False

    def _get_page_count(self):
        """获取页数（智能估算方法）"""
        from docx.enum.section import WD_SECTION

        # 方法1：统计段前分页（Word中段前分页的段落会开始新页）
        # 这是最可靠的分页检测方法
        page_break_before_count = 0
        for p in self.doc.paragraphs:
            if p._element.pPr is not None:
                page_break_before = p._element.pPr.find(qn('w:pageBreakBefore'))
                if page_break_before is not None:
                    val = page_break_before.get(qn('w:val'))
                    # 如果val不存在或为"1"/"on"/"true"，则启用段前分页
                    if val in [None, "1", "on", "true"]:
                        page_break_before_count += 1

        # 方法2：统计手动分页符（w:br type="page"）
        manual_breaks = 0
        for p in self.doc.paragraphs:
            xml_str = p._element.xml
            if '<w:br' in xml_str and 'w:type="page"' in xml_str:
                manual_breaks += 1

        # 方法3：统计新页分节符（WD_SECTION.NEW_PAGE）
        new_page_sections = sum(
            1 for s in self._flat_sections
            if s.start_type == WD_SECTION.NEW_PAGE
        )

        # 计算总页数
        # 首页 + 段前分页 + 手动分页符 + (新页分节符 - 1)
        # 因为第一个节不需要额外分页
        estimated_pages = 1 + page_break_before_count + manual_breaks + (new_page_sections - 1)

        # 如果页数过多（>50），可能是检测算法有问题，使用备用方法
        if estimated_pages > 50:
            # 备用方法：基于内容长度
            total_chars = sum(len(p.text) for p in self.doc.paragraphs)
            # 每页约1500-2000字符（考虑图片、表格等）
            backup_pages = max(1, total_chars // 1800)
            return backup_pages

        return estimated_pages

    def _find_key_pages(self):
        """找到中文摘要、英文摘要、目录页在文档中的位置

        依据办公软件中"看到"的标题文字定位这三个关键页：
          - 中文摘要标题段落：含"摘要"或"摘 要"字样，不含 Abstract；
            且非目录条目（无制表符引导页码、无末尾"…数字"式页码占位）
          - 英文摘要标题段落：文本以 Abstract/ABSTRACT 出现，同样排除目录条目
          - 目录标题段落：文本为"目录"或"目 录"，排除正文中偶然出现的"目录"引用
        取每类别在文档中最早出现的匹配段落索引；同时记录其所属节索引。

        使用 _build_flat_document_index 构建的扁平段落序列，穿透 w:sdt/
        w:sdtContent 结构化文档标签容器，覆盖被"内容控件"包裹的标题段。
        """
        self.chinese_abstract_para_idx = -1
        self.english_abstract_para_idx = -1
        self.outline_para_idx = -1

        for i, p in enumerate(self._flat_paragraphs):
            text = p.text.strip()
            if not text or len(text) >= 30:
                continue
            # 目录条目常见特征：含制表符、或"…"引导、或末尾以空白+数字结尾
            if ('\t' in p.text
                    or '……' in text
                    or '…' in text
                    or re.search(r'\s+\d+\s*$', text)):
                continue

            normalized = re.sub(r'\s+', '', text)

            # 中文摘要
            if (self.chinese_abstract_para_idx == -1
                    and '摘要' in normalized
                    and 'abstract' not in text.lower()):
                self.chinese_abstract_para_idx = i
                continue
            # 英文摘要
            if (self.english_abstract_para_idx == -1
                    and re.match(r'(?i)^abstract\b', text)):
                self.english_abstract_para_idx = i
                continue
            # 目录
            if self.outline_para_idx == -1 and normalized in ('目录', '目次'):
                self.outline_para_idx = i
                continue

    def _check_page_order(self):
        """检查中文摘要、英文摘要、目录页顺序 (-1分)

        细则：中文摘要页、英文摘要页、目录页不是此先后顺序

        判定规则：办公软件中三者按文档流依次出现，先后关系必须为
          中文摘要 → 英文摘要 → 目录
        - 三者任一未能在文档中定位到，则视为不满足此先后顺序
        - 使用段落在文档中的自然顺序索引（等同于办公软件从上到下阅读顺序）
        """
        has_issue = False

        if (self.chinese_abstract_para_idx == -1
                or self.english_abstract_para_idx == -1
                or self.outline_para_idx == -1):
            has_issue = True
        elif not (self.chinese_abstract_para_idx
                  < self.english_abstract_para_idx
                  < self.outline_para_idx):
            has_issue = True

        if has_issue:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：中文摘要页、英文摘要页、目录页不是此先后顺序"
            )
            print("[-1] 中文摘要页、英文摘要页、目录页不是此先后顺序")
        else:
            print("[PASS] 中文摘要页、英文摘要页、目录页顺序符合要求")

    def _get_paragraph_section_index(self, target_para_idx):
        """根据段落索引（扁平索引，穿透 w:sdt）推断其所属节索引"""
        if target_para_idx < 0 or self.doc is None:
            return -1
        if target_para_idx >= len(self._para_section_index):
            return -1
        return self._para_section_index[target_para_idx]

    def _has_page_break_before_paragraph(self, target_para_idx):
        """判断目标段落（扁平索引）所在节内，其之前是否存在物理分页。

        用于节级回退场景下识别"摘要标题与摘要所在节的节首之间还有更早
        物理页"的情况（例如封面/前置内容与摘要共享同一节，仅用分页符
        或段前分页隔开，而不是分节符）。仅在目标段落所属节的范围内向前
        扫描，扫描到该节起始处即停止，不越节判断。

        检测依据（与 _get_page_count 的分页符探测口径一致，改为扫描
        _flat_paragraphs/_flat_para_elements 以覆盖被 w:sdt 包裹的段落）：
          - 目标段落自身或其之前的同节段落设置了 w:pageBreakBefore
          - 目标段落之前的同节段落中出现手动分页符 <w:br w:type="page"/>
        任一命中即认为摘要所在节内、摘要标题之前还存在其它物理页。
        """
        if target_para_idx < 0 or target_para_idx >= len(self._flat_para_elements):
            return False
        sec_idx = self._get_paragraph_section_index(target_para_idx)
        if sec_idx < 0:
            return False

        # 目标段落自身若设置了"段前分页"，说明它另起一页，
        # 即它前面（同节内）必然还有页面。
        target_el = self._flat_para_elements[target_para_idx]
        pPr = target_el.find(qn('w:pPr'))
        if pPr is not None:
            pbb = pPr.find(qn('w:pageBreakBefore'))
            if pbb is not None:
                val = pbb.get(qn('w:val'))
                if val in (None, '1', 'on', 'true'):
                    return True

        # 向前回溯同节内的段落，寻找手动分页符或段前分页标记
        i = target_para_idx - 1
        while i >= 0 and self._para_section_index[i] == sec_idx:
            el = self._flat_para_elements[i]
            p_pPr = el.find(qn('w:pPr'))
            if p_pPr is not None:
                pbb = p_pPr.find(qn('w:pageBreakBefore'))
                if pbb is not None:
                    val = pbb.get(qn('w:val'))
                    if val in (None, '1', 'on', 'true'):
                        return True
            if '<w:br' in el.xml and 'w:type="page"' in el.xml:
                return True
            i -= 1
        return False

    def _probe_word_pagination(self):
        """分页信息固定走节级静态回退，不依赖任何渲染引擎。"""
        self._word_pagination = None

    def _page_to_section_index(self, page_no):
        """给定物理页号，返回其所属的扁平节索引；未知返回 -1。"""
        wp = self._word_pagination
        if wp is None:
            return -1
        for i, (first_p, last_p) in enumerate(wp['section_pages']):
            if first_p <= page_no <= last_p:
                return i
        return -1

    def _get_page_header(self, sec_idx, kind, odd_even):
        """按 (节, 显示场景) 返回办公软件实际渲染的页眉对象。

        - kind ∈ {'default','even','first'}
        - 未启用奇偶不同时，'even' 一律回退为 'default'
        - 未启用首页不同时，'first' 一律回退为 'default'
        - 沿"与上一节链接"关系回溯到实际显示的对象
        """
        if sec_idx < 0 or sec_idx >= len(self._flat_sections):
            return None
        section = self._flat_sections[sec_idx]
        if kind == 'even' and not odd_even:
            kind = 'default'
        if kind == 'first':
            try:
                if not getattr(
                    section, 'different_first_page_header_footer', False
                ):
                    kind = 'default'
            except Exception:
                kind = 'default'
        try:
            if kind == 'default':
                raw = section.header
            elif kind == 'even':
                raw = section.even_page_header
            elif kind == 'first':
                raw = section.first_page_header
            else:
                raw = section.header
        except Exception:
            return None
        return self._resolve_effective_header(raw, sec_idx, kind)

    def _classify_page_kind(self, page_no, sec_idx, odd_even):
        """判定某物理页面在其所在节内实际用的是哪个页眉场景。

        返回 'first' / 'even' / 'default'（default 即"奇数页/唯一页眉"）。
        """
        if sec_idx < 0 or self._word_pagination is None:
            # 无法判定；按奇偶回退
            if odd_even and page_no % 2 == 0:
                return 'even'
            return 'default'
        section = self._flat_sections[sec_idx]
        first_p, _ = self._word_pagination['section_pages'][sec_idx]
        try:
            diff_first = bool(getattr(
                section, 'different_first_page_header_footer', False
            ))
        except Exception:
            diff_first = False
        if diff_first and page_no == first_p:
            return 'first'
        if odd_even and page_no % 2 == 0:
            return 'even'
        return 'default'

    def _enumerate_pages_after_abstract(self):
        """列出"摘要页及摘要页之后"的每个物理页面对应的 (页号, 节索引, 场景)。

        依赖 self._word_pagination；无 COM 分页信息时返回 None，调用方回退到
        节级颗粒度检查。
        """
        wp = self._word_pagination
        if wp is None:
            return None
        abs_page = wp.get('abstract_zh_page')
        total = wp.get('total_pages')
        if not abs_page or not total:
            return None
        try:
            odd_even = self.doc.settings.odd_and_even_pages_header_footer
        except Exception:
            odd_even = False
        result = []
        for page_no in range(abs_page, total + 1):
            sec_i = self._page_to_section_index(page_no)
            kind = self._classify_page_kind(page_no, sec_i, odd_even)
            result.append((page_no, sec_i, kind))
        return result

    def _check_header_lines_before_abstract(self):
        """检查摘要前所有页面页眉是否没有横线 (+3分)

        细则：摘要（中文"摘要"两字标题）所在页之前的所有物理页面，页眉处
        没有横线。摘要页本身即使有横线也不构成扣分点（那是 +5 分要求的
        横线所在页）。

        覆盖范围（优先按物理页粒度）：
          - 若 Word COM 分页信息可用（_word_pagination.abstract_zh_page）：
            枚举物理页号 1 .. abstract_zh_page - 1 的每一页，对每一页取
            其所在节及其在办公软件中实际生效的页眉场景（default/even/first），
            沿"与上一节链接"关系回溯到实际显示的页眉对象；仅这些页眉
            纳入检查。摘要页所在节即使被更早的页共享，也只检查那些更早的
            物理页对应的页眉场景，不会把摘要页本身对应的页眉纳入。
          - 若 COM 分页不可用（未安装 Word / pywin32），回退到节级颗粒度：
            检查索引严格小于"摘要标题所在节"的所有节的所有页眉；若摘要
            所在节内、摘要标题段落之前还检测到分页符/段前分页（即摘要
            与更早的物理页共享同一节），该节自身的页眉也纳入检查，避免
            "没有更早的节"被误判为"没有摘要前页面"而直接放过。
          - 对每个"摘要前"节，按办公软件实际显示逻辑收集所有会出现的页眉：
              * 默认页眉（无奇偶设置时用于所有页；启用奇偶不同则用于奇数页）
              * 偶数页页眉（启用"奇偶页不同"时用于偶数页；未定义则偶数页
                复用默认页眉）
              * 首页页眉（节启用"首页不同"时用于该节首页）
          - 对每个页眉，沿"与上一节链接"关系回溯到实际显示的对象
        判定"页眉处有横线"（任一命中即视为该页眉不通过）：
          - 页眉任意段落的 w:pBdr 中存在可见的 top / bottom / between 边框
            （可见 = val 不为空/none/nil，且 sz>0）
          - 页眉内含表格，且表格任意单元格具有可见 top / bottom / insideH
            边框（这些都会在办公软件中显示为一根横线）
        """
        if self.doc is None or len(self._flat_sections) == 0:
            print("[+0] 摘要前页面页眉横线检查未通过：文档为空")
            return

        if self.chinese_abstract_para_idx < 0:
            print("[+0] 摘要前页面页眉横线检查未通过：未定位到摘要标题")
            return

        abstract_sec_idx = self._get_paragraph_section_index(
            self.chinese_abstract_para_idx
        )
        if abstract_sec_idx < 0:
            print("[+0] 摘要前页面页眉横线检查未通过：未能定位摘要所在节")
            return

        try:
            odd_even = self.doc.settings.odd_and_even_pages_header_footer
        except Exception:
            odd_even = False

        offending = []

        # 优先：按物理页粒度枚举"摘要页之前"的每一页
        wp = self._word_pagination
        abs_page_raw = wp.get('abstract_zh_page') if wp else None
        abs_page = abs_page_raw if isinstance(abs_page_raw, int) else None

        if abs_page is not None and abs_page > 0:
            if abs_page == 1:
                # 摘要就在第1页——不存在摘要前物理页面，按通过处理
                self.results['dimension2_scores'].append(3)
                self.results['hit_points'].append("+3：摘要前所有页面页眉处没有横线")
                print("[+3] 摘要前所有页面页眉处没有横线")
                return
            for page_no in range(1, abs_page):
                sec_i = self._page_to_section_index(page_no)
                kind = self._classify_page_kind(page_no, sec_i, odd_even)
                header = self._get_page_header(sec_i, kind, odd_even)
                if header is None:
                    continue
                if self._header_has_horizontal_line(header):
                    offending.append(f"第{page_no}页({kind}, 节{sec_i})页眉")

            if not offending:
                self.results['dimension2_scores'].append(3)
                self.results['hit_points'].append("+3：摘要前所有页面页眉处没有横线")
                print("[+3] 摘要前所有页面页眉处没有横线")
            else:
                print(f"[+0] 摘要前页面页眉横线检查未通过: {offending}")
            return

        # 回退：节级颗粒度——检查严格早于摘要节的所有节；
        # 此外，摘要所在节内部也可能存在"摘要前"的物理页（用分页符/
        # 段前分页与摘要标题段落分隔，而非分节符），这些页面共享
        # 摘要所在节的页眉对象，必须一并纳入检查，不能因为"没有更早的
        # 节"就直接判通过（那样会漏检与摘要同节的封面/前置页页眉）。
        sections_to_check = list(range(abstract_sec_idx))
        if self._has_page_break_before_paragraph(self.chinese_abstract_para_idx):
            sections_to_check.append(abstract_sec_idx)

        if not sections_to_check:
            # 严格早于摘要节的节为空，且摘要所在节内摘要标题前也没有
            # 检测到分页符——确认摘要就是该节第1页，不存在摘要前页面
            self.results['dimension2_scores'].append(3)
            self.results['hit_points'].append("+3：摘要前所有页面页眉处没有横线")
            print("[+3] 摘要前所有页面页眉处没有横线")
            return

        for i in sections_to_check:
            section = self._flat_sections[i]

            headers_to_check = []

            default_header = self._resolve_effective_header(
                section.header, i, 'default'
            )
            if default_header is not None:
                if odd_even:
                    headers_to_check.append(
                        (default_header, f"节{i}奇数页页眉")
                    )
                    try:
                        even_raw = section.even_page_header
                    except Exception:
                        even_raw = None
                    even_header = self._resolve_effective_header(
                        even_raw, i, 'even'
                    )
                    if even_header is not None:
                        headers_to_check.append(
                            (even_header, f"节{i}偶数页页眉")
                        )
                    else:
                        headers_to_check.append(
                            (default_header, f"节{i}偶数页页眉(复用默认)")
                        )
                else:
                    headers_to_check.append(
                        (default_header, f"节{i}页眉")
                    )

            try:
                if getattr(section, 'different_first_page_header_footer', False):
                    first_raw = section.first_page_header
                    first_header = self._resolve_effective_header(
                        first_raw, i, 'first'
                    )
                    if first_header is not None:
                        headers_to_check.append(
                            (first_header, f"节{i}首页页眉")
                        )
            except Exception:
                pass

            for header, label in headers_to_check:
                if self._header_has_horizontal_line(header):
                    offending.append(label)

        if not offending:
            self.results['dimension2_scores'].append(3)
            self.results['hit_points'].append("+3：摘要前所有页面页眉处没有横线")
            print("[+3] 摘要前所有页面页眉处没有横线")
        else:
            print(f"[+0] 摘要前页面页眉横线检查未通过: {offending}")

    def _header_has_horizontal_line(self, header):
        """页眉中是否出现横线（段落横向边框 或 表格横向边框）"""
        if header is None:
            return False
        # 段落 w:pBdr 中的 top/bottom/between 边框
        for paragraph in header.paragraphs:
            pPr = paragraph._element.pPr
            if pPr is None:
                continue
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                continue
            for border in pBdr:
                name = border.tag.split('}')[-1]
                if name in ('top', 'bottom', 'between'):
                    if self._is_visible_border(border):
                        return True
        # 页眉内表格产生的横线
        for table in header.tables:
            if self._table_has_visible_horizontal_line(table):
                return True
        return False

    def _table_has_visible_horizontal_line(self, table):
        """表格是否会在办公软件中呈现出横线"""
        # 表级默认边框：top/bottom/insideH
        tblPr = table._element.tblPr
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                for name in ('top', 'bottom', 'insideH'):
                    b = tblBorders.find(qn(f'w:{name}'))
                    if self._is_visible_border(b):
                        return True
        # 单元格边框：top/bottom
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.tcPr
                if tcPr is None:
                    continue
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    continue
                for name in ('top', 'bottom'):
                    b = tcBorders.find(qn(f'w:{name}'))
                    if self._is_visible_border(b):
                        return True
                # 嵌套表格递归
                for nested in cell.tables:
                    if self._table_has_visible_horizontal_line(nested):
                        return True
        return False

    # ==== "两条平行横线" 相关的判定辅助 =====================================

    # Word/WPS 内置的会渲染为"两条水平平行线"的复合边框样式
    # 键：val 值；值：(top_ratio, bottom_ratio) —— 办公软件中该复合边框在
    # 屏幕上从上到下呈现的两条线的粗细比例（与实际观察一致）
    # 说明：OOXML 中 thick-thin/thin-thick 命名表示从段落文本一侧向外的
    # 延展顺序；用作页眉/段落下边框(w:bottom)时，靠近文本的方向是上方，
    # 因此 "thickThin*"（thick 先/thin 后）在办公软件里实际呈现为
    # 上薄下厚，"thinThick*" 呈现为上厚下薄。
    _DOUBLE_LINE_STYLES = {
        # val: (top_ratio, bottom_ratio) —— 办公软件视觉从上到下比例
        'double':               (1, 1),  # 上下等粗
        'thickThinSmallGap':    (1, 3),  # 上薄下厚
        'thickThinMediumGap':   (1, 3),
        'thickThinLargeGap':    (1, 3),
        'thinThickSmallGap':    (3, 1),  # 上厚下薄
        'thinThickMediumGap':   (3, 1),
        'thinThickLargeGap':    (3, 1),
    }

    def _is_black_border_color(self, color):
        """办公软件中该颜色是否呈现为黑色"""
        if color is None:
            return True  # 未设 = auto = 黑
        c = str(color).strip().lower()
        if c in ('auto', 'black', '000000', '000'):
            return True
        # 允许 #000 / #000000 写法
        if c.startswith('#'):
            c = c[1:]
        if c == '000000':
            return True
        return False

    def _lines_from_double_border(self, border):
        """给定一个 val 为复合双线样式的边框，返回 (上线粗pt, 下线粗pt, color)

        计算模型：w:sz 单位为 1/8 磅，办公软件渲染时按样式内置比例
        将 sz 分配给上/下两条线（gap 忽略于粗细计算）。
        """
        val = border.get(qn('w:val'))
        if val not in self._DOUBLE_LINE_STYLES:
            return None
        try:
            sz = int(border.get(qn('w:sz')) or '0')
        except (ValueError, TypeError):
            sz = 0
        if sz <= 0:
            return None
        color = border.get(qn('w:color'))
        top_r, bottom_r = self._DOUBLE_LINE_STYLES[val]
        total_pt = sz / 8.0
        top_pt = total_pt * top_r / (top_r + bottom_r)
        bottom_pt = total_pt * bottom_r / (top_r + bottom_r)
        return top_pt, bottom_pt, color

    def _extract_two_parallel_lines(self, paragraph):
        """从段落 pBdr 中提取"两条平行横线"的粗细与颜色

        返回 (上线pt, 下线pt, [颜色列表])；找不到时返回 None。
        兼容两种表达方式：
          1) 单个复合双线边框（top/bottom/between 之一，val ∈ 双线样式集）
          2) 段落同时设置了可见的 top + bottom 两条单线边框
        """
        pPr = paragraph._element.pPr
        if pPr is None:
            return None
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            return None

        # 方式1：单个复合双线边框
        for border in pBdr:
            name = border.tag.split('}')[-1]
            if name not in ('top', 'bottom', 'between'):
                continue
            info = self._lines_from_double_border(border)
            if info is not None:
                top_pt, bottom_pt, color = info
                return top_pt, bottom_pt, [color]

        # 方式2：top + bottom 两条单线边框都可见
        top_b = pBdr.find(qn('w:top'))
        bottom_b = pBdr.find(qn('w:bottom'))
        if (top_b is not None and bottom_b is not None
                and self._is_visible_border(top_b)
                and self._is_visible_border(bottom_b)):
            try:
                top_sz = int(top_b.get(qn('w:sz')) or '0')
                bottom_sz = int(bottom_b.get(qn('w:sz')) or '0')
            except (ValueError, TypeError):
                return None
            if top_sz > 0 and bottom_sz > 0:
                return (
                    top_sz / 8.0,
                    bottom_sz / 8.0,
                    [top_b.get(qn('w:color')), bottom_b.get(qn('w:color'))],
                )
        return None

    def _header_has_valid_double_line(self, header, label, details):
        """检查页眉是否具有满足细则的双横线"""
        if header is None:
            details.append(f"{label}不存在")
            return False
        # 头里所有段落只要有一个满足即可
        for paragraph in header.paragraphs:
            info = self._extract_two_parallel_lines(paragraph)
            if info is None:
                continue
            top_pt, bottom_pt, colors = info

            # 颜色：所有涉及的颜色都必须呈黑色
            if not all(self._is_black_border_color(c) for c in colors):
                details.append(
                    f"{label}横线颜色不是黑色(colors={colors})"
                )
                return False

            # 下横线粗细不超过1.5磅
            if bottom_pt > 1.5 + 1e-6:
                details.append(
                    f"{label}下横线{bottom_pt:.3f}pt 超过1.5磅"
                )
                return False

            # 上横线粗细不低于下横线的3倍
            if bottom_pt <= 0:
                details.append(f"{label}下横线粗细为0")
                return False
            if top_pt + 1e-6 < 3 * bottom_pt:
                details.append(
                    f"{label}上横线{top_pt:.3f}pt 不足下横线{bottom_pt:.3f}pt 的3倍"
                )
                return False
            return True
        details.append(f"{label}未检出两条平行横线")
        return False

    def _check_header_lines_after_abstract(self):
        """检查摘要页及摘要页后所有页面页眉的双横线 (+5分)

        细则：摘要页及摘要页后所有页面，页眉处存在两条相互平行的横线；
              颜色为黑色；上横线的粗细不低于下横线的3倍；下横线粗细不超过
              1.5磅

        覆盖范围：
          - 通过 _get_paragraph_section_index 定位中文摘要标题所在节
            abstract_sec；摘要页及之后 = 索引 >= abstract_sec 的所有节
            （若未找到摘要，则视为从第 1 节起）
          - 每个节按办公软件实际显示逻辑收集所有会出现的页眉：
              * 默认页眉（无奇偶设置用于所有页；启用奇偶不同用于奇数页）
              * 偶数页页眉（启用"奇偶页不同"时；未定义则偶数页复用默认）
              * 首页页眉（节启用"首页不同"时）
          - 沿"与上一节链接"关系回溯到实际显示的页眉对象
        每个页眉必须至少有一个段落其 pBdr 满足全部四项要求：
          (1) 两条相互平行的横线：
              * 单个复合双线边框：val ∈ {double, thickThinSmallGap,
                thickThinMediumGap, thickThinLargeGap, thinThickSmallGap,
                thinThickMediumGap, thinThickLargeGap}
              * 或段落同时设置了可见 top + 可见 bottom 单线边框
          (2) 颜色为黑色：所有涉及边框的 color 为 auto / 000000 / black
          (3) 上横线粗细 ≥ 3 × 下横线粗细
          (4) 下横线粗细 ≤ 1.5磅
              （线粗根据 w:sz 与办公软件实际渲染比例换算：
               复合双线样式内置的粗细比例 double=1:1、thickThin*=3:1、
               thinThick*=1:3；两条单线边框直接用各自的 sz/8 磅）
        """
        if self.doc is None or len(self._flat_sections) == 0:
            print("[+0] 摘要页及摘要页后页眉双横线检查未通过：文档为空")
            return

        abstract_sec_idx = self._get_paragraph_section_index(
            self.chinese_abstract_para_idx
        )
        if abstract_sec_idx < 0:
            abstract_sec_idx = 1

        try:
            odd_even = self.doc.settings.odd_and_even_pages_header_footer
        except Exception:
            odd_even = False

        details = []
        all_passed = True

        # 优先按物理页粒度逐页检查（Word COM 精确分页可用时）
        pages_to_check = self._enumerate_pages_after_abstract()
        if pages_to_check is not None:
            for page_no, sec_i, kind in pages_to_check:
                header = self._get_page_header(sec_i, kind, odd_even)
                label = f"第{page_no}页({kind}, 节{sec_i})页眉"
                if not self._header_has_valid_double_line(
                    header, label, details
                ):
                    all_passed = False

            if all_passed:
                self.results['dimension2_scores'].append(5)
                self.results['hit_points'].append(
                    "+5：摘要页及摘要页后所有页面页眉处存在两条相互平行的黑色横线，"
                    "上横线粗细不低于下横线的3倍，下横线粗细不超过1.5磅"
                )
                print("[+5] 摘要页及摘要页后页眉双横线符合要求")
            else:
                print(f"[+0] 摘要页及摘要页后页眉双横线检查未通过: {details}")
            return

        # 回退：节级颗粒度（要求节内所有会出现的页眉均满足）
        for i in range(abstract_sec_idx, len(self._flat_sections)):
            section = self._flat_sections[i]

            headers_to_check = []
            default_header = self._resolve_effective_header(
                section.header, i, 'default'
            )
            if default_header is not None:
                if odd_even:
                    headers_to_check.append(
                        (default_header, f"节{i}奇数页页眉")
                    )
                    try:
                        even_raw = section.even_page_header
                    except Exception:
                        even_raw = None
                    even_header = self._resolve_effective_header(
                        even_raw, i, 'even'
                    )
                    if even_header is not None:
                        headers_to_check.append(
                            (even_header, f"节{i}偶数页页眉")
                        )
                    else:
                        headers_to_check.append(
                            (default_header, f"节{i}偶数页页眉(复用默认)")
                        )
                else:
                    headers_to_check.append(
                        (default_header, f"节{i}页眉")
                    )
            try:
                if getattr(section, 'different_first_page_header_footer', False):
                    first_raw = section.first_page_header
                    first_header = self._resolve_effective_header(
                        first_raw, i, 'first'
                    )
                    if first_header is not None:
                        headers_to_check.append(
                            (first_header, f"节{i}首页页眉")
                        )
            except Exception:
                pass

            for header, label in headers_to_check:
                if not self._header_has_valid_double_line(
                    header, label, details
                ):
                    all_passed = False

        if all_passed:
            self.results['dimension2_scores'].append(5)
            self.results['hit_points'].append(
                "+5：摘要页及摘要页后所有页面页眉处存在两条相互平行的黑色横线，"
                "上横线粗细不低于下横线的3倍，下横线粗细不超过1.5磅"
            )
            print("[+5] 摘要页及摘要页后页眉双横线符合要求")
        else:
            print(f"[+0] 摘要页及摘要页后页眉双横线检查未通过: {details}")

    def _check_header_line_position(self):
        """检查所有页页眉横线位于页眉文本下方，间距1-4字符 (+5分)

        细则：所有页页眉位置上的横线位于页眉文本下方，与页眉文本间距
              在1-4字符之间

        覆盖范围：
          - 遍历所有节
          - 每节按办公软件实际显示逻辑收集所有会出现的页眉：
              * 默认页眉（无奇偶设置用于所有页；启用奇偶不同用于奇数页）
              * 偶数页页眉（启用"奇偶页不同"时；未定义则偶数页复用默认）
              * 首页页眉（节启用"首页不同"时）
          - 沿"与上一节链接"关系回溯到实际显示的页眉对象
          - 只针对"页眉位置上存在横线"的页眉判定；不存在横线的页眉不适用
            该条评分点，跳过（例如摘要前页眉本身就无横线的情况）
        "横线位于页眉文本下方" 判定：
          - 段落 pBdr 中的可见边框必须出现在 <w:bottom>（或复合双线边框
            放在 bottom 位置）
          - 若同一页眉段落还带有可见 <w:top> / <w:between> 边框，则说明存在
            不在文本下方的横线，判定失败
        "与页眉文本间距在1-4字符之间" 判定：
          - 边框的 w:space 属性单位为磅（0-31），代表办公软件中横线与
            该段落文本的物理距离
          - 不再把"字符"直接等同于该页眉段落字号；此处采用 Word 中文
            文档常用的默认字符宽度口径：1字符 = 10.5磅（五号字 em 宽）
          - 将 w:space 的物理距离换算为字符数，满足 1 ≤ 字符数 ≤ 4 通过
        """
        if self.doc is None or len(self._flat_sections) == 0:
            print("[+0] 页眉横线位置检查未通过：文档为空")
            return

        try:
            odd_even = self.doc.settings.odd_and_even_pages_header_footer
        except Exception:
            odd_even = False

        details = []
        checked_any = False
        all_passed = True

        for i, section in enumerate(self._flat_sections):
            headers_to_check = []
            default_header = self._resolve_effective_header(
                section.header, i, 'default'
            )
            if default_header is not None:
                if odd_even:
                    headers_to_check.append(
                        (default_header, f"节{i}奇数页页眉")
                    )
                    try:
                        even_raw = section.even_page_header
                    except Exception:
                        even_raw = None
                    even_header = self._resolve_effective_header(
                        even_raw, i, 'even'
                    )
                    if even_header is not None:
                        headers_to_check.append(
                            (even_header, f"节{i}偶数页页眉")
                        )
                    else:
                        headers_to_check.append(
                            (default_header, f"节{i}偶数页页眉(复用默认)")
                        )
                else:
                    headers_to_check.append(
                        (default_header, f"节{i}页眉")
                    )
            try:
                if getattr(section, 'different_first_page_header_footer', False):
                    first_raw = section.first_page_header
                    first_header = self._resolve_effective_header(
                        first_raw, i, 'first'
                    )
                    if first_header is not None:
                        headers_to_check.append(
                            (first_header, f"节{i}首页页眉")
                        )
            except Exception:
                pass

            for header, label in headers_to_check:
                ok, tested, reason = self._header_line_position_ok(header, label)
                if tested:
                    checked_any = True
                    if not ok:
                        all_passed = False
                        details.append(reason)

        # 若没有任何页眉有横线，说明整份文档不存在页眉横线，此评分点不成立
        if not checked_any:
            print("[+0] 页眉横线位置检查未通过：未在任何页眉发现横线")
            return

        if all_passed:
            self.results['dimension2_scores'].append(5)
            self.results['hit_points'].append(
                "+5：所有页页眉位置上的横线位于页眉文本下方，与页眉文本间距在1-4字符之间"
            )
            print("[+5] 所有页页眉横线位置符合要求")
        else:
            print(f"[+0] 页眉横线位置检查未通过: {details}")

    def _header_line_position_ok(self, header, label):
        """检查单个页眉的横线位置

        返回 (是否通过, 是否检测到横线, 失败原因)
        """
        if header is None:
            return True, False, ""
        tested = False
        for paragraph in header.paragraphs:
            pPr = paragraph._element.pPr
            if pPr is None:
                continue
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                continue

            # 收集本段落存在的可见横向边框（top / bottom / between）
            visible_top = self._is_visible_border(pBdr.find(qn('w:top')))
            visible_between = self._is_visible_border(pBdr.find(qn('w:between')))
            bottom_border = pBdr.find(qn('w:bottom'))
            visible_bottom = self._is_visible_border(bottom_border)

            if not (visible_top or visible_between or visible_bottom):
                continue

            tested = True

            # 存在不在文本下方的横线
            if visible_top:
                return False, True, f"{label}存在文本上方的横线(w:top)"
            if visible_between:
                return False, True, f"{label}存在段落间的横线(w:between)"

            # 到此存在下方横线
            if bottom_border is None:
                return False, True, f"{label}未找到 w:bottom 边框"
            try:
                space_pt = int(bottom_border.get(qn('w:space')) or '0')
            except (ValueError, TypeError):
                space_pt = 0
            char_width_pt = 10.5
            space_chars = space_pt / char_width_pt
            if not (1 - 1e-6 <= space_chars <= 4 + 1e-6):
                return False, True, (
                    f"{label}横线间距{space_pt}磅(约{space_chars:.2f}字符)，"
                    "不在1-4字符之间"
                )
            # 该段落通过；继续检查同页眉其他段落是否有违反项
        return True, tested, ""

    def _check_header_text(self):
        """检查页眉文本是否为"澄湖理工大学硕士学位论文"且居中 (-1分)

        细则：任意一页页眉文本内容不是"澄湖理工大学硕士学位论文"或不在居中位置

        覆盖范围：
          - 遍历所有节
          - 每节检查在办公软件中实际会显示的所有页眉：
              * 默认页眉（奇数页/无奇偶设置时非首页页面显示）
              * 首页页眉（当节启用"首页不同"时首页显示）
              * 偶数页页眉（当文档启用"奇偶页不同"时偶数页显示）
          - 若页眉标记为"与上一节链接"（is_linked_to_previous），则回溯到最近
            实际定义了该类型页眉的上级节，取其真实显示的内容
          - 页眉可能含多个段落，取所有非空段落的可见文本拼接
        判定规则：
          - 该显示内容 strip 后必须严格等于"澄湖理工大学硕士学位论文"
          - 所有承载文本的非空段落对齐方式必须解析为 CENTER（居中），
            解析过程支持段落直接设置及沿样式继承链（含 base_style）追溯
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        expected_text = "澄湖理工大学硕士学位论文"

        if self.doc is None or len(self._flat_sections) == 0:
            print("[PASS] 页眉文本符合要求")
            return

        try:
            odd_even = self.doc.settings.odd_and_even_pages_header_footer
        except Exception:
            odd_even = False

        has_issue = False
        fail_reason = ""

        for i, section in enumerate(self._flat_sections):
            headers_to_check = [(section.header, 'default', f"节{i}默认页眉")]

            # 首页页眉：仅当该节启用"首页不同"时，首页实际显示的是它
            try:
                if getattr(section, 'different_first_page_header_footer', False):
                    headers_to_check.append(
                        (section.first_page_header, 'first', f"节{i}首页页眉")
                    )
            except Exception:
                pass

            # 偶数页页眉：仅当文档启用"奇偶页不同"时，偶数页实际显示的是它
            try:
                if odd_even:
                    headers_to_check.append(
                        (section.even_page_header, 'even', f"节{i}偶数页页眉")
                    )
            except Exception:
                pass

            for header, kind, label in headers_to_check:
                effective = self._resolve_effective_header(header, i, kind)
                if effective is None:
                    continue

                non_empty_paragraphs = [
                    p for p in effective.paragraphs if p.text.strip()
                ]
                visible_text = ''.join(
                    p.text.strip() for p in non_empty_paragraphs
                )

                if visible_text != expected_text:
                    has_issue = True
                    fail_reason = (
                        f"{label}文本为\"{visible_text}\"，"
                        f"不等于\"{expected_text}\""
                    )
                    break

                for p in non_empty_paragraphs:
                    align = self._resolve_paragraph_alignment(p)
                    if align != WD_ALIGN_PARAGRAPH.CENTER:
                        has_issue = True
                        fail_reason = f"{label}文本不在居中位置"
                        break
                if has_issue:
                    break
            if has_issue:
                break

        if has_issue:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append("-1：任意一页页眉文本内容不是\"澄湖理工大学硕士学位论文\"或不在居中位置")
            print(f"[-1] {fail_reason}")
        else:
            print("[PASS] 页眉文本符合要求")

    def _resolve_effective_header(self, header, section_idx, kind):
        """沿"与上一节链接"关系回溯，返回办公软件实际显示的页眉对象

        kind: 'default' | 'first' | 'even'
        """
        if header is None:
            return None
        try:
            linked = header.is_linked_to_previous
        except Exception:
            linked = False
        if not linked or section_idx <= 0:
            return header
        prev_section = self._flat_sections[section_idx - 1]
        if kind == 'first':
            prev_header = prev_section.first_page_header
        elif kind == 'even':
            prev_header = prev_section.even_page_header
        else:
            prev_header = prev_section.header
        return self._resolve_effective_header(prev_header, section_idx - 1, kind)

    def _resolve_paragraph_alignment(self, paragraph):
        """解析段落最终生效的对齐方式：直接设置 -> 段落样式 -> 样式 base_style 链"""
        align = paragraph.paragraph_format.alignment
        if align is not None:
            return align
        try:
            style = paragraph.style
        except Exception:
            return None
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            try:
                pf = style.paragraph_format
            except Exception:
                pf = None
            if pf is not None and pf.alignment is not None:
                return pf.alignment
            try:
                style = style.base_style
            except Exception:
                break
        return None

    def _check_footer_text(self):
        """检查页脚文本和对齐方式 (-1分)

        细则：目录页后所有页面不满足以下任意一项：
          - 页脚文本没有出现"城市社区微气候韧性评价与公共空间更新研究"
          - 奇数页页脚文本不是左对齐
          - 偶数页页脚文本不是右对齐

        覆盖范围：
          - 定位到"目录"所在节，其后续节视为"目录页后"页面所在节
          - 每节按办公软件实际显示逻辑取所有会出现的页脚：
              * 默认页脚 —— 未启用"奇偶页不同"时用于所有页；启用时用于奇数页
              * 偶数页页脚 —— 启用"奇偶页不同"且未链接上一节时用于偶数页
              * 首页页脚 —— 节启用"首页不同"且未链接上一节时用于该节首页
          - 沿"与上一节链接"关系回溯到实际显示的页脚对象
          - 页脚可能含多个段落，取所有非空段落
        判定规则：
          - 每个显示中的页脚，所有非空段落的可见文本合并后必须包含
            "城市社区微气候韧性评价与公共空间更新研究"
          - 奇数页对应的每个非空段落对齐方式必须解析为 LEFT
          - 偶数页对应的每个非空段落对齐方式必须解析为 RIGHT
          - 首页页脚（无法判定其在奇/偶）：允许 LEFT 或 RIGHT，文本必须存在
          - 若启用"奇偶页不同"但仅有默认页脚（无偶数页页脚定义），
            则偶数页仍复用默认页脚，需分别通过奇/偶两套对齐判定 → 无法同时满足
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        expected_text = "城市社区微气候韧性评价与公共空间更新研究"

        if self.doc is None or len(self._flat_sections) == 0:
            print("[PASS] 页脚文本符合要求")
            return

        outline_section_idx = self._get_paragraph_section_index(self.outline_para_idx)
        if outline_section_idx < 0:
            # 未找到目录，无法判定"目录页后"，按通过处理
            print("[PASS] 页脚文本符合要求")
            return

        start_section = outline_section_idx + 1
        if start_section >= len(self._flat_sections):
            print("[PASS] 页脚文本符合要求")
            return

        try:
            odd_even = self.doc.settings.odd_and_even_pages_header_footer
        except Exception:
            odd_even = False

        has_issue = False
        fail_reason = ""

        for i in range(start_section, len(self._flat_sections)):
            section = self._flat_sections[i]

            # 组装本节实际会显示的页脚：(footer对象, 显示场景, 描述)
            # 场景：'odd'=奇数页 / 'even'=偶数页 / 'all'=所有页 / 'first'=首页
            footers_to_check = []

            default_footer = self._resolve_effective_footer(section.footer, i, 'default')
            if default_footer is not None:
                if odd_even:
                    # 启用奇偶不同：默认页脚仅用于奇数页
                    footers_to_check.append((default_footer, 'odd', f"节{i}奇数页页脚"))
                    # 偶数页页脚
                    try:
                        even_raw = section.even_page_footer
                    except Exception:
                        even_raw = None
                    even_footer = self._resolve_effective_footer(even_raw, i, 'even')
                    if even_footer is not None:
                        footers_to_check.append((even_footer, 'even', f"节{i}偶数页页脚"))
                    else:
                        # 启用奇偶不同但未定义偶数页页脚 → 偶数页复用默认页脚
                        footers_to_check.append((default_footer, 'even', f"节{i}偶数页页脚(复用默认)"))
                else:
                    footers_to_check.append((default_footer, 'all', f"节{i}默认页脚"))

            # 首页页脚
            try:
                if getattr(section, 'different_first_page_header_footer', False):
                    first_raw = section.first_page_footer
                    first_footer = self._resolve_effective_footer(first_raw, i, 'first')
                    if first_footer is not None:
                        footers_to_check.append((first_footer, 'first', f"节{i}首页页脚"))
            except Exception:
                pass

            for footer, scene, label in footers_to_check:
                non_empty_paragraphs = [
                    p for p in footer.paragraphs if p.text.strip()
                ]
                visible_text = ''.join(p.text for p in non_empty_paragraphs)

                # 1) 文本必须出现
                if expected_text not in visible_text:
                    has_issue = True
                    fail_reason = f"{label}未出现\"{expected_text}\""
                    break

                # 2) 对齐方式判定
                if scene == 'odd':
                    required = {WD_ALIGN_PARAGRAPH.LEFT}
                    require_desc = "左对齐"
                elif scene == 'even':
                    required = {WD_ALIGN_PARAGRAPH.RIGHT}
                    require_desc = "右对齐"
                elif scene == 'all':
                    # 未启用奇偶页不同：同一页脚要同时满足奇数页左对齐+偶数页右对齐 → 不可能
                    has_issue = True
                    fail_reason = (
                        f"{label}未启用\"奇偶页不同\"，"
                        f"无法同时满足奇数页左对齐、偶数页右对齐"
                    )
                    break
                else:  # 'first' —— 首页可能奇可能偶，允许左或右
                    required = {WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT}
                    require_desc = "左对齐或右对齐"

                for p in non_empty_paragraphs:
                    align = self._resolve_paragraph_alignment(p)
                    if align not in required:
                        has_issue = True
                        fail_reason = f"{label}文本不是{require_desc}"
                        break
                if has_issue:
                    break
            if has_issue:
                break

        if has_issue:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：目录页后所有页面不满足以下任意一项：页脚文本没有出现"
                "\"城市社区微气候韧性评价与公共空间更新研究\"；"
                "奇数页页脚文本不是左对齐；偶数页页脚文本不是右对齐"
            )
            print(f"[-1] {fail_reason}")
        else:
            print("[PASS] 页脚文本符合要求")

    def _resolve_effective_footer(self, footer, section_idx, kind):
        """沿"与上一节链接"关系回溯，返回办公软件实际显示的页脚对象

        kind: 'default' | 'first' | 'even'
        """
        if footer is None:
            return None
        try:
            linked = footer.is_linked_to_previous
        except Exception:
            linked = False
        if not linked or section_idx <= 0:
            return footer
        prev_section = self._flat_sections[section_idx - 1]
        if kind == 'first':
            prev_footer = prev_section.first_page_footer
        elif kind == 'even':
            prev_footer = prev_section.even_page_footer
        else:
            prev_footer = prev_section.footer
        return self._resolve_effective_footer(prev_footer, section_idx - 1, kind)

    def _check_page_number(self):
        """检查页脚居中位置是否有页码 (-1分)

        细则：任意一页页脚居中位置没有出现页码

        覆盖范围：
          - 遍历所有节
          - 每节按办公软件实际显示逻辑取所有会出现的页脚：
              * 默认页脚（未启用"奇偶页不同"时用于所有页；启用时用于奇数页）
              * 偶数页页脚（启用"奇偶页不同"且已定义时用于偶数页；
                启用但未定义则偶数页复用默认页脚）
              * 首页页脚（节启用"首页不同"时用于该节首页）
          - 沿"与上一节链接"关系回溯到办公软件实际显示的页脚对象
          - 页脚中可能同时存在普通段落 和/或 表格（Word 内置三列页脚采用表格）
        判定"居中位置出现页码"（满足其一即视为该页脚通过）：
          (A) 存在页码，且其所在段落对齐方式解析为 CENTER
          (B) 存在页码，且所在段落使用制表符将其定位到该段落的居中制表位
              （制表位来自段落直接设置或样式继承链；Word 内置"Footer/页脚"
                样式默认包含页面中心的居中制表位，视为具有居中制表位）
          (C) 存在页码，且其位于页脚表格中间那一列的单元格内
              （递归包含嵌套表格）
        "页码"识别（覆盖办公软件产生页码的全部途径）：
          - <w:fldSimple w:instr="... PAGE ..."/> （简单域）
          - 三段式域：<w:fldChar begin/> + <w:instrText>...PAGE...</w:instrText>
            + <w:fldChar end/>
          - 纯数字文本（如"5"、" 5 "）
          - "第 X 页"、"Page X"、"- X -" 等常见页码文本模板
        """
        if self.doc is None or len(self._flat_sections) == 0:
            print("[PASS] 页脚有页码")
            return

        try:
            odd_even = self.doc.settings.odd_and_even_pages_header_footer
        except Exception:
            odd_even = False

        missing_label = None

        for i, section in enumerate(self._flat_sections):
            footers_to_check = []

            default_footer = self._resolve_effective_footer(
                section.footer, i, 'default'
            )
            if default_footer is not None:
                if odd_even:
                    footers_to_check.append((default_footer, f"节{i}奇数页页脚"))
                    try:
                        even_raw = section.even_page_footer
                    except Exception:
                        even_raw = None
                    even_footer = self._resolve_effective_footer(
                        even_raw, i, 'even'
                    )
                    if even_footer is not None:
                        footers_to_check.append(
                            (even_footer, f"节{i}偶数页页脚")
                        )
                    else:
                        footers_to_check.append(
                            (default_footer, f"节{i}偶数页页脚(复用默认)")
                        )
                else:
                    footers_to_check.append((default_footer, f"节{i}页脚"))

            try:
                if getattr(section, 'different_first_page_header_footer', False):
                    first_raw = section.first_page_footer
                    first_footer = self._resolve_effective_footer(
                        first_raw, i, 'first'
                    )
                    if first_footer is not None:
                        footers_to_check.append(
                            (first_footer, f"节{i}首页页脚")
                        )
            except Exception:
                pass

            for footer, label in footers_to_check:
                if not self._footer_has_center_page_number(footer):
                    missing_label = label
                    # 诊断输出：打印该页脚的实际内容，便于定位问题
                    print(f"    [调试] {label}未通过居中页码检查，内容如下：")
                    para_count = 0
                    for pi, p in enumerate(footer.paragraphs):
                        align = self._resolve_paragraph_alignment(p)
                        has_pn = self._paragraph_has_page_number(p)
                        print(
                            f"      段落{pi}: text={p.text!r} "
                            f"align={align} 含页码={has_pn}"
                        )
                        para_count += 1
                    if para_count == 0:
                        print("      (无段落)")
                    for ti, table in enumerate(footer.tables):
                        col_count = len(table.columns)
                        print(f"      表格{ti}: 列数={col_count}")
                        for ri, row in enumerate(table.rows):
                            for ci, cell in enumerate(row.cells):
                                cell_text = '\n'.join(
                                    p.text for p in cell.paragraphs
                                )
                                has_pn = any(
                                    self._paragraph_has_page_number(p)
                                    for p in cell.paragraphs
                                )
                                print(
                                    f"        行{ri}列{ci}: text={cell_text!r} "
                                    f"含页码={has_pn}"
                                )
                    break
            if missing_label:
                break

        if missing_label:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append("-1：任意一页页脚居中位置没有出现页码")
            print(f"[-1] {missing_label}居中位置未出现页码")
        else:
            print("[PASS] 页脚有页码")

    def _footer_has_center_page_number(self, footer):
        """页脚中是否存在处于'居中位置'的页码"""
        if footer is None:
            return False
        # 直接段落
        for p in footer.paragraphs:
            if self._paragraph_page_number_at_center(p):
                return True
        # 表格：递归检查所有表格
        for table in footer.tables:
            if self._table_has_center_page_number(table):
                return True
        return False

    def _table_has_center_page_number(self, table):
        """表格中间列的单元格是否包含页码（递归含嵌套表格）"""
        col_count = len(table.columns)
        if col_count == 0:
            return False
        middle_idx = col_count // 2  # 3列→1；2列→1（右列，不视为居中，故要求>=2列时中列存在）
        # 只有奇数列时"中间列"才有真正的居中含义
        if col_count % 2 == 0:
            candidate_idxs = []
        else:
            candidate_idxs = [middle_idx]
        # 无论多少列，仍需递归嵌套表格，但对本层"中列"要求列数为奇数
        for row in table.rows:
            for c_idx, cell in enumerate(row.cells):
                if c_idx in candidate_idxs:
                    for p in cell.paragraphs:
                        if self._paragraph_has_page_number(p):
                            return True
                # 嵌套表格递归
                for nested in cell.tables:
                    if self._table_has_center_page_number(nested):
                        return True
        return False

    def _paragraph_page_number_at_center(self, paragraph):
        """段落中的页码是否处于'居中位置'"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if not self._paragraph_has_page_number(paragraph):
            return False
        # 情形A：段落整体居中对齐
        align = self._resolve_paragraph_alignment(paragraph)
        if align == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        # 情形B：页码位于段落的居中制表位处
        if self._page_number_at_center_tab(paragraph):
            return True
        # 情形C：页码位于水平居中的浮动文本框内（wps:txbx / v:textbox）
        if self._page_number_in_centered_textbox(paragraph):
            return True
        return False

    def _page_number_in_centered_textbox(self, paragraph):
        """页码位于水平居中的浮动文本框内

        Word/WPS 允许把页码放进浮动文本框，通过 <wp:anchor> 的
        <wp:positionH> 把文本框锁定到页面/页边距的水平居中。这种情形
        视觉上就是"页脚居中位置出现了页码"，需按居中通过。
        """
        element = paragraph._element
        # 遍历段落内所有 <w:drawing> 或 <w:pict>（VML 回退）
        WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # 1) DrawingML: w:drawing/wp:anchor 或 wp:inline
        for drawing in element.iter(qn('w:drawing')):
            # 找到 wp:anchor / wp:inline
            for anchor in drawing.iter():
                tag = anchor.tag
                if not (tag == f'{{{WP_NS}}}anchor' or tag == f'{{{WP_NS}}}inline'):
                    continue
                # 内容里必须含 PAGE 域
                contains_page = False
                for instr in anchor.iter(f'{{{W_NS}}}instrText'):
                    if instr.text and self._instr_is_page(instr.text):
                        contains_page = True
                        break
                if not contains_page:
                    for fld in anchor.iter(f'{{{W_NS}}}fldSimple'):
                        if self._instr_is_page(
                            fld.get(f'{{{W_NS}}}instr', '')
                        ):
                            contains_page = True
                            break
                if not contains_page:
                    continue
                # 检查水平位置是否居中
                if self._anchor_horizontally_centered(anchor):
                    return True
        # 2) VML: w:pict/v:shape，style 里 mso-position-horizontal:center
        for pict in element.iter(qn('w:pict')):
            contains_page = False
            for instr in pict.iter(f'{{{W_NS}}}instrText'):
                if instr.text and self._instr_is_page(instr.text):
                    contains_page = True
                    break
            if not contains_page:
                continue
            for shape in pict.iter():
                style = shape.get('style') if hasattr(shape, 'get') else None
                if style and 'mso-position-horizontal:center' in style.replace(' ', ''):
                    return True
        return False

    def _anchor_horizontally_centered(self, anchor):
        """判断 DrawingML 锚点是否水平居中放置"""
        WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        positionH = anchor.find(f'{{{WP_NS}}}positionH')
        if positionH is None:
            return False
        align_el = positionH.find(f'{{{WP_NS}}}align')
        if align_el is not None and (align_el.text or '').strip().lower() == 'center':
            return True
        return False

    def _paragraph_has_page_number(self, paragraph):
        """段落是否包含页码（覆盖办公软件产生页码的各种途径）"""
        element = paragraph._element
        # <w:fldSimple w:instr="... PAGE ...">
        for fld in element.iter(qn('w:fldSimple')):
            if self._instr_is_page(fld.get(qn('w:instr'), '')):
                return True
        # 三段式域 <w:instrText>...PAGE...</w:instrText>
        for instr in element.iter(qn('w:instrText')):
            if instr.text and self._instr_is_page(instr.text):
                return True
        # 纯数字文本或常见页码文本模板
        text = paragraph.text.strip()
        if not text:
            return False
        if re.fullmatch(r'\d+', text):
            return True
        # "第 X 页"、"第X页"
        if re.search(r'第\s*\d+\s*页', text):
            return True
        # "Page X"、"Page X of Y"
        if re.search(r'page\s*\d+', text, re.IGNORECASE):
            return True
        # "- X -"、"— X —"
        if re.fullmatch(r'[\-—–\s]*\d+[\-—–\s]*', text):
            return True
        return False

    def _instr_is_page(self, instr):
        """域代码是否为 PAGE 类页码域"""
        if not instr:
            return False
        tokens = re.findall(r'[A-Za-z]+', instr)
        # PAGE / NUMPAGES / SECTIONPAGES 均视为页码域
        return any(t.upper() in ('PAGE', 'NUMPAGES', 'SECTIONPAGES') for t in tokens)

    def _page_number_at_center_tab(self, paragraph):
        """页码是否位于段落的居中制表位处"""
        element = paragraph._element
        tabs_before = 0
        page_seen = False

        # 按文档顺序遍历，统计页码前的 tab 字符个数
        for descendant in element.iter():
            tag = descendant.tag
            if tag == qn('w:tab'):
                parent = descendant.getparent()
                # 排除 pPr/tabs 下 tab 定义节点，仅计入 run 中的 tab 字符
                if parent is not None and parent.tag == qn('w:r'):
                    if not page_seen:
                        tabs_before += 1
            elif tag == qn('w:fldSimple'):
                if self._instr_is_page(descendant.get(qn('w:instr'), '')):
                    page_seen = True
            elif tag == qn('w:instrText'):
                if descendant.text and self._instr_is_page(descendant.text):
                    page_seen = True

        if not page_seen or tabs_before == 0:
            return False
        return self._paragraph_has_center_tab_stop(paragraph)

    def _paragraph_has_center_tab_stop(self, paragraph):
        """段落是否具有居中制表位（含样式继承，或使用内置页脚样式的默认居中制表位）"""
        from docx.enum.text import WD_TAB_ALIGNMENT
        # 直接段落设置
        try:
            for tab in paragraph.paragraph_format.tab_stops:
                if tab.alignment == WD_TAB_ALIGNMENT.CENTER:
                    return True
        except Exception:
            pass
        # 样式继承链
        try:
            style = paragraph.style
        except Exception:
            style = None
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            try:
                pf = style.paragraph_format
                for tab in pf.tab_stops:
                    if tab.alignment == WD_TAB_ALIGNMENT.CENTER:
                        return True
            except Exception:
                pass
            try:
                style = style.base_style
            except Exception:
                break
        # 内置"页脚/Footer"样式默认在页面中心提供居中制表位
        try:
            s = paragraph.style
            style_id = getattr(s, 'style_id', '') or ''
            style_name = getattr(s, 'name', '') or ''
        except Exception:
            style_id, style_name = '', ''
        if style_id.lower() == 'footer' or 'footer' in style_name.lower() or '页脚' in style_name:
            return True
        return False

    def _check_page_number_format(self):
        """检查页码格式（罗马数字/阿拉伯数字） (-1分)

        细则：目录页及目录页之前的页面页码不是罗马数字，目录页之后的页面
              页码不是阿拉伯数字

        覆盖范围：
          - 通过 _get_paragraph_section_index 定位"目录"标题所在节 outline_sec
          - 分为两段并分别校验办公软件中实际显示的页码格式：
              A) 索引 <= outline_sec 的所有节 —— 必须为罗马数字
              B) 索引 >  outline_sec 的所有节 —— 必须为阿拉伯数字
          - 每个节的页码格式来源于 sectPr/w:pgNumType/@w:fmt
              * 缺失 pgNumType 或 fmt 属性时，办公软件默认按十进制（阿拉伯数字）
                显示，视为 "decimal"
        格式识别（按办公软件"看到"的显示归类）：
          - 罗马数字：fmt ∈ {upperRoman, lowerRoman}
          - 阿拉伯数字：fmt ∈ {decimal, decimalEnclosedCircle, decimalZero,
                              decimalFullWidth, decimalHalfWidth,
                              decimalFullWidth2, decimalEnclosedFullstop,
                              decimalEnclosedParen, decimalEnclosedCircleChinese}
            —— 上述 fmt 在办公软件里最终呈现的都是"0-9"阿拉伯数字字形
        """
        if self.doc is None or len(self._flat_sections) == 0:
            print("[PASS] 页码格式符合要求")
            return

        outline_sec = self._get_paragraph_section_index(self.outline_para_idx)
        if outline_sec < 0:
            # 未找到目录，无法判定，按通过处理
            print("[PASS] 页码格式符合要求")
            return

        roman_fmts = {'upperRoman', 'lowerRoman'}
        arabic_fmts = {
            'decimal', 'decimalEnclosedCircle', 'decimalZero',
            'decimalFullWidth', 'decimalHalfWidth', 'decimalFullWidth2',
            'decimalEnclosedFullstop', 'decimalEnclosedParen',
            'decimalEnclosedCircleChinese',
        }

        has_issue = False
        fail_reason = ""

        for i, section in enumerate(self._flat_sections):
            fmt = self._get_section_page_number_format(section)
            if i <= outline_sec:
                # 目录页及目录页之前 → 必须罗马数字
                if fmt not in roman_fmts:
                    has_issue = True
                    fail_reason = (
                        f"节{i}（目录页及之前）页码格式为\"{fmt}\"，不是罗马数字"
                    )
                    break
            else:
                # 目录页之后 → 必须阿拉伯数字
                if fmt not in arabic_fmts:
                    has_issue = True
                    fail_reason = (
                        f"节{i}（目录页之后）页码格式为\"{fmt}\"，不是阿拉伯数字"
                    )
                    break

        if has_issue:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：目录页及目录页之前的页面页码不是罗马数字，"
                "目录页之后的页面页码不是阿拉伯数字"
            )
            print(f"[-1] {fail_reason}")
        else:
            print("[PASS] 页码格式符合要求")

    def _get_section_page_number_format(self, section):
        """读取节的页码格式（w:sectPr/w:pgNumType/@w:fmt）

        - 未设置 pgNumType 或未设置 fmt 时，办公软件按"decimal"显示
        """
        sectPr = section._sectPr
        if sectPr is None:
            return 'decimal'
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            return 'decimal'
        fmt = pgNumType.get(qn('w:fmt'))
        return fmt if fmt else 'decimal'

    def _check_heading1_format(self):
        """检查目录页后所有一级标题是否为黑体、三号、加粗 (-1分)

        细则：目录页后所有一级标题不是黑体、三号、加粗

        - 起始位置：目录标题段之后（未找到目录则从文档首段开始）
        - 一级标题识别（三选一）：段落样式为 Heading 1/标题 1；
          大纲级别 outlineLvl=0（含样式继承）；或文本形态属于中文论文
          常见一级标题："第X章…"、"参考文献"、"致谢"、"结论"、
          "结束语"、"附录…"。排除含制表符 / "……" / "…" 或末尾
          "空白+数字"的目录条目。
        - 每个一级标题的所有非空 run，按办公软件实际显示解析：
          (1) 中文字体 w:rFonts/@w:eastAsia：字体名包含"黑体"或"SimHei"
          (2) 字号 w:sz：等于 16 磅（三号），容差 0.5 磅
          (3) 加粗 w:b：解析为 True
          解析顺序：run 直接设置 → 段落样式链（含 base_style）
        """
        if self.doc is None:
            print("[PASS] 一级标题格式符合要求")
            return
        start = self.outline_para_idx + 1 if self.outline_para_idx >= 0 else 0
        fail_reason = None
        for p in self._flat_paragraphs[start:]:
            if not self._is_first_level_heading(p):
                continue
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                font = self._resolve_run_east_asia_font(run, p)
                size_pt = self._resolve_run_font_size_pt(run, p)
                is_bold = self._resolve_run_bold(run, p)
                if not font or ('黑体' not in font and 'SimHei' not in font):
                    fail_reason = (
                        f"一级标题「{p.text.strip()[:20]}」"
                        f"中文字体「{font}」不是黑体"
                    )
                    break
                if size_pt is None or abs(size_pt - 16) > 0.5:
                    fail_reason = (
                        f"一级标题「{p.text.strip()[:20]}」"
                        f"字号{size_pt}磅不是三号(16磅)"
                    )
                    break
                if is_bold is not True:
                    fail_reason = f"一级标题「{p.text.strip()[:20]}」未加粗"
                    break
            if fail_reason:
                break
        if fail_reason:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：目录页后所有一级标题不是黑体、三号、加粗"
            )
            print(f"[-1] {fail_reason}")
        else:
            print("[PASS] 一级标题格式符合要求")

    def _is_first_level_heading(self, paragraph):
        """判断段落是否为一级标题（仅看标题文本本身）

        按办公软件中"眼见"的一级标题文本形态判定，与样式/大纲级别无关：
          (1) 中文数字序号章节："一、xxx"、"二、xxx" ...
          (2) "第X章 xxx"（X 为中文数字或阿拉伯数字）
          (3) 论文常见独立一级标题：结论 / 结束语 / 结  语 / 参考文献 /
              致谢 / 致  谢 / 附录 / 附  录（允许中间含空白）

        排除目录条目：含制表符、"…"/"……"、末尾"空白+数字"
        """
        text = paragraph.text.strip()
        if not text:
            return False
        if '\t' in paragraph.text or '……' in text or '…' in text:
            return False
        if re.search(r'\s+\d+\s*$', text):
            return False

        normalized = re.sub(r'\s+', '', text)

        # (1) 中文数字序号章节
        if re.match(r'^[一二三四五六七八九十百千]+[、\.．]', normalized):
            return True
        # (2) 第X章
        if re.match(r'^第[一二三四五六七八九十百千0-9]+章', normalized):
            return True
        # (3) 独立标题（裸的整段标题名，不含额外序号或子标题内容）
        if normalized in ('结论', '结束语', '结语',
                          '参考文献', '致谢', '附录'):
            return True
        return False

    def _resolve_paragraph_outline_level(self, paragraph):
        """解析段落大纲级别：直接设置 → 段落样式链"""
        pPr = paragraph._element.pPr
        if pPr is not None:
            ol = pPr.find(qn('w:outlineLvl'))
            if ol is not None:
                val = ol.get(qn('w:val'))
                try:
                    return int(val) if val is not None else None
                except (ValueError, TypeError):
                    return None
        try:
            style = paragraph.style
        except Exception:
            style = None
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            try:
                sp = style.element.find(qn('w:pPr'))
            except Exception:
                sp = None
            if sp is not None:
                ol = sp.find(qn('w:outlineLvl'))
                if ol is not None:
                    val = ol.get(qn('w:val'))
                    try:
                        return int(val) if val is not None else None
                    except (ValueError, TypeError):
                        return None
            try:
                style = style.base_style
            except Exception:
                break
        return None

    def _resolve_run_east_asia_font(self, run, paragraph):
        """解析 run 最终生效的中文字体名（run → 段落样式链）"""
        rPr = run._element.rPr
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'))
                if ea:
                    return ea
                ascii_font = rFonts.get(qn('w:ascii'))
                if ascii_font:
                    return ascii_font
        try:
            style = paragraph.style
        except Exception:
            style = None
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            try:
                sr = style.element.find(qn('w:rPr'))
            except Exception:
                sr = None
            if sr is not None:
                rFonts = sr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'))
                    if ea:
                        return ea
                    ascii_font = rFonts.get(qn('w:ascii'))
                    if ascii_font:
                        return ascii_font
            try:
                style = style.base_style
            except Exception:
                break
        return None

    def _resolve_run_font_size_pt(self, run, paragraph):
        """解析 run 最终生效的字号（磅）：run → 段落样式链"""
        try:
            if run.font.size is not None:
                return run.font.size.pt
        except Exception:
            pass
        try:
            style = paragraph.style
        except Exception:
            style = None
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            try:
                sz = style.font.size
            except Exception:
                sz = None
            if sz is not None:
                return sz.pt
            try:
                style = style.base_style
            except Exception:
                break
        return None

    def _resolve_run_bold(self, run, paragraph):
        """解析 run 最终生效的加粗属性：run → 段落样式链"""
        try:
            if run.font.bold is not None:
                return run.font.bold
        except Exception:
            pass
        try:
            style = paragraph.style
        except Exception:
            style = None
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            try:
                b = style.font.bold
            except Exception:
                b = None
            if b is not None:
                return b
            try:
                style = style.base_style
            except Exception:
                break
        return None

    def _check_heading2_format(self):
        """检查目录页后所有二级标题是否为黑体、四号、加粗 (-1分)

        细则：目录页后所有二级标题不是黑体、四号、加粗

        - 起始位置：目录标题段之后（未找到目录则从文档首段开始）
        - 二级标题识别（办公软件"导航窗格"视角）：
          (1) 段落直接引用样式为 Heading 2 / 标题 2 / heading2；或
          (2) 段落大纲级别 outlineLvl = 1（含样式继承）
          排除目录条目：含制表符、"…"/"……"、末尾"空白+数字"
        - 每个二级标题的所有非空 run 按办公软件实际显示解析：
          (1) 中文字体 w:rFonts/@w:eastAsia：字体名包含"黑体"或"SimHei"
          (2) 字号 w:sz：等于 14 磅（四号），容差 0.5 磅
          (3) 加粗 w:b：解析为 True
          解析顺序：run 直接设置 → 段落样式链（含 base_style）
        """
        if self.doc is None:
            print("[PASS] 二级标题格式符合要求")
            return
        start = self.outline_para_idx + 1 if self.outline_para_idx >= 0 else 0
        fail_reason = None
        for p in self._flat_paragraphs[start:]:
            if not self._is_second_level_heading(p):
                continue
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                font = self._resolve_run_east_asia_font(run, p)
                size_pt = self._resolve_run_font_size_pt(run, p)
                is_bold = self._resolve_run_bold(run, p)
                if not font or ('黑体' not in font and 'SimHei' not in font):
                    fail_reason = (
                        f"二级标题「{p.text.strip()[:20]}」"
                        f"中文字体「{font}」不是黑体"
                    )
                    break
                if size_pt is None or abs(size_pt - 14) > 0.5:
                    fail_reason = (
                        f"二级标题「{p.text.strip()[:20]}」"
                        f"字号{size_pt}磅不是四号(14磅)"
                    )
                    break
                if is_bold is not True:
                    fail_reason = f"二级标题「{p.text.strip()[:20]}」未加粗"
                    break
            if fail_reason:
                break
        if fail_reason:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：目录页后所有二级标题不是黑体、四号、加粗"
            )
            print(f"[-1] {fail_reason}")
        else:
            print("[PASS] 二级标题格式符合要求")

    def _is_second_level_heading(self, paragraph):
        """判断段落是否为二级标题（仅看标题文本本身）

        按办公软件中"眼见"的二级标题文本形态判定，与样式/大纲级别无关：
          - "X.Y xxx"（X、Y 为阿拉伯数字，之间恰好一个点，且不再有第三段数字）

        排除目录条目：含制表符、"…"/"……"、末尾"空白+数字"
        """
        text = paragraph.text.strip()
        if not text:
            return False
        if '\t' in paragraph.text or '……' in text or '…' in text:
            return False
        if re.search(r'\s+\d+\s*$', text):
            return False

        # X.Y 开头，且不是 X.Y.Z（避免把三级标题误判为二级）
        if re.match(r'^\d+\.\d+(?!\.\d)', text):
            return True
        return False

    def _check_heading3_format(self):
        """检查目录页后所有三级标题是否为黑体、小四号、加粗 (-1分)

        细则：目录页后所有三级标题不是黑体、小四号、加粗

        - 起始位置：目录标题段之后（未找到目录则从文档首段开始）
        - 三级标题识别（办公软件"导航窗格"视角）：
          (1) 段落直接引用样式为 Heading 3 / 标题 3 / heading3；或
          (2) 段落大纲级别 outlineLvl = 2（含样式继承）
          排除目录条目：含制表符、"…"/"……"、末尾"空白+数字"
        - 每个三级标题的所有非空 run 按办公软件实际显示解析：
          (1) 中文字体 w:rFonts/@w:eastAsia：字体名包含"黑体"或"SimHei"
          (2) 字号 w:sz：等于 12 磅（小四号），容差 0.5 磅
          (3) 加粗 w:b：解析为 True
          解析顺序：run 直接设置 → 段落样式链（含 base_style）
        """
        if self.doc is None:
            print("[PASS] 三级标题格式符合要求")
            return
        start = self.outline_para_idx + 1 if self.outline_para_idx >= 0 else 0
        fail_reason = None
        for p in self._flat_paragraphs[start:]:
            if not self._is_third_level_heading(p):
                continue
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                font = self._resolve_run_east_asia_font(run, p)
                size_pt = self._resolve_run_font_size_pt(run, p)
                is_bold = self._resolve_run_bold(run, p)
                if not font or ('黑体' not in font and 'SimHei' not in font):
                    fail_reason = (
                        f"三级标题「{p.text.strip()[:20]}」"
                        f"中文字体「{font}」不是黑体"
                    )
                    break
                if size_pt is None or abs(size_pt - 12) > 0.5:
                    fail_reason = (
                        f"三级标题「{p.text.strip()[:20]}」"
                        f"字号{size_pt}磅不是小四号(12磅)"
                    )
                    break
                if is_bold is not True:
                    fail_reason = f"三级标题「{p.text.strip()[:20]}」未加粗"
                    break
            if fail_reason:
                break
        if fail_reason:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：目录页后所有三级标题不是黑体、小四号、加粗"
            )
            print(f"[-1] {fail_reason}")
        else:
            print("[PASS] 三级标题格式符合要求")

    def _is_third_level_heading(self, paragraph):
        """判断段落是否为三级标题（仅看标题文本本身）

        按办公软件中"眼见"的三级标题文本形态判定，与样式/大纲级别无关：
          - "X.Y.Z xxx"（X、Y、Z 均为阿拉伯数字，且不再有第四段数字）

        排除目录条目：含制表符、"…"/"……"、末尾"空白+数字"
        """
        text = paragraph.text.strip()
        if not text:
            return False
        if '\t' in paragraph.text or '……' in text or '…' in text:
            return False
        if re.search(r'\s+\d+\s*$', text):
            return False

        if re.match(r'^\d+\.\d+\.\d+(?!\.\d)', text):
            return True
        return False

    def _check_paragraph_format(self):
        """检查"1.1.1 研究背景"与"1.1.2 研究意义"之间的文本 (-1分)

        细则："1.1.1 研究背景"与"1.1.2 研究意义"中间的文本不满足以下
              任意一项：宋体、小四、1.5倍行距、两端对齐

        - 定位区间：找到含"1.1.1"且含"研究背景"的段落作为起始，
          含"1.1.2"且含"研究意义"的段落作为结束，取二者之间（不含边界）
          的所有非空段落进行判定
        - 若未能定位到两个边界标题，或区间内没有可检查的段落，说明
          无法验证该 rubric 项（可能是标题缺失、措辞变化导致定位失败，
          或中间确实没有正文），此时不能默认按"通过"处理——按扣分
          处理，避免标题缺失/定位失败被误判为格式合规
        - 每个段落的所有非空 run 按办公软件实际显示解析：
          (1) 中文字体 w:rFonts/@w:eastAsia：字体名包含"宋体"或"SimSun"
          (2) 字号 w:sz：等于 12 磅（小四号），容差 0.5 磅
          解析顺序均为：run 直接设置 → 段落样式链（含 base_style）
        - 每个段落自身按办公软件实际显示解析：
          (3) 1.5 倍行距：pPr/w:spacing/w:line = 360 且 w:lineRule
              ∈ {缺省, 'auto'}（办公软件中此配置显示为"1.5 倍行距"）；
              解析顺序：段落直接设置 → 段落样式链
          (4) 两端对齐：解析后的对齐方式为 WD_ALIGN_PARAGRAPH.JUSTIFY
              解析顺序：段落直接设置 → 段落样式链
        - 区间内任一段落任一项不满足即判定该项扣分
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if self.doc is None:
            print("[PASS] 段落格式符合要求")
            return

        start_idx = -1
        end_idx = -1
        for i, p in enumerate(self._flat_paragraphs):
            text = p.text.strip()
            if start_idx == -1 and '1.1.1' in text and '研究背景' in text:
                start_idx = i
            elif start_idx != -1 and '1.1.2' in text and '研究意义' in text:
                end_idx = i
                break

        if start_idx == -1 or end_idx == -1:
            # 未能定位到两个边界标题之一，无法验证该 rubric 项——
            # 不能默认按"通过"处理，按扣分处理（标题缺失/措辞变化导致
            # 定位失败本身即是不合规的信号）
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：\"1.1.1 研究背景\"与\"1.1.2 研究意义\"中间的文本"
                "不满足以下任意一项：宋体、小四、1.5倍行距、两端对齐"
            )
            print("[-1] 未定位到\"1.1.1 研究背景\"或\"1.1.2 研究意义\"边界标题，无法验证")
            return

        checkable = [
            i for i in range(start_idx + 1, end_idx)
            if self._flat_paragraphs[i].text.strip()
        ]
        if not checkable:
            # 区间为空（无可检查段落）同样无法验证该 rubric 项，
            # 不能默认按"通过"处理
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：\"1.1.1 研究背景\"与\"1.1.2 研究意义\"中间的文本"
                "不满足以下任意一项：宋体、小四、1.5倍行距、两端对齐"
            )
            print("[-1] \"1.1.1 研究背景\"与\"1.1.2 研究意义\"之间没有可检查的正文段落，无法验证")
            return

        fail_reason = None
        for i in checkable:
            p = self._flat_paragraphs[i]

            # (1)(2) 每个非空 run 的字体和字号
            for run in p.runs:
                if not run.text or not run.text.strip():
                    continue
                font = self._resolve_run_east_asia_font(run, p)
                size_pt = self._resolve_run_font_size_pt(run, p)
                if not font or ('宋体' not in font and 'SimSun' not in font):
                    fail_reason = (
                        f"段落「{p.text.strip()[:20]}」"
                        f"中文字体「{font}」不是宋体"
                    )
                    break
                if size_pt is None or abs(size_pt - 12) > 0.5:
                    fail_reason = (
                        f"段落「{p.text.strip()[:20]}」"
                        f"字号{size_pt}磅不是小四(12磅)"
                    )
                    break
            if fail_reason:
                break

            # (3) 1.5 倍行距
            line_twips, line_rule = self._resolve_paragraph_line_spacing(p)
            if not (line_twips == 360 and line_rule in (None, 'auto')):
                fail_reason = (
                    f"段落「{p.text.strip()[:20]}」"
                    f"行距(line={line_twips}, rule={line_rule})不是1.5倍行距"
                )
                break

            # (4) 两端对齐
            align = self._resolve_paragraph_alignment(p)
            if align != WD_ALIGN_PARAGRAPH.JUSTIFY:
                fail_reason = (
                    f"段落「{p.text.strip()[:20]}」对齐方式{align}不是两端对齐"
                )
                break

        if fail_reason:
            self.results['dimension2_deductions'].append(1)
            self.results['hit_points'].append(
                "-1：\"1.1.1 研究背景\"与\"1.1.2 研究意义\"中间的文本"
                "不满足以下任意一项：宋体、小四、1.5倍行距、两端对齐"
            )
            print(f"[-1] {fail_reason}")
        else:
            print("[PASS] 段落格式符合要求")

    def _resolve_paragraph_line_spacing(self, paragraph):
        """解析段落最终生效的行距：段落直接设置 → 样式链（含 base_style）

        返回 (line_twips, line_rule)：
          - line_twips: int 或 None —— w:spacing/@w:line 的原始值（twips）
          - line_rule: 'auto' / 'atLeast' / 'exact' 或 None
        """
        pPr = paragraph._element.pPr
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                line = spacing.get(qn('w:line'))
                rule = spacing.get(qn('w:lineRule'))
                if line is not None:
                    try:
                        return int(line), rule
                    except (ValueError, TypeError):
                        return None, rule

        try:
            style = paragraph.style
        except Exception:
            style = None
        visited = set()
        while style is not None and id(style) not in visited:
            visited.add(id(style))
            try:
                sp = style.element.find(qn('w:pPr'))
            except Exception:
                sp = None
            if sp is not None:
                spacing = sp.find(qn('w:spacing'))
                if spacing is not None:
                    line = spacing.get(qn('w:line'))
                    rule = spacing.get(qn('w:lineRule'))
                    if line is not None:
                        try:
                            return int(line), rule
                        except (ValueError, TypeError):
                            return None, rule
            try:
                style = style.base_style
            except Exception:
                break
        return None, None

    def _is_visible_border(self, border):
        """判断边框XML节点是否为可见边框"""
        if border is None:
            return False
        val = border.get(qn('w:val'))
        sz = border.get(qn('w:sz'))
        return val not in [None, 'none', 'nil'] and sz != '0'

    def _is_three_line_table(self, table):
        """按单元格实际边框判断是否为三线表"""
        row_count = len(table.rows)
        col_count = len(table.columns)
        if row_count < 2 or col_count == 0:
            return True  # 封面布局类表格不作为正文三线表扣分依据

        tblBorders = None
        tblPr = table._element.tblPr
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
        table_borders = {b.tag.split('}')[-1]: b for b in tblBorders} if tblBorders is not None else {}

        has_any_border = False
        has_top = self._is_visible_border(table_borders.get('top'))
        has_header_bottom = False
        has_bottom = self._is_visible_border(table_borders.get('bottom'))
        has_extra_line = False
        has_cell_borders = False

        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                tcPr = cell._tc.tcPr
                tcBorders = tcPr.find(qn('w:tcBorders')) if tcPr is not None else None
                if tcBorders is None:
                    continue

                borders = {b.tag.split('}')[-1]: b for b in tcBorders}
                for name, border in borders.items():
                    if not self._is_visible_border(border):
                        continue
                    has_any_border = True
                    if name in ['left', 'right', 'insideV', 'start', 'end']:
                        has_extra_line = True
                    elif name == 'top':
                        if r_idx == 0:
                            has_top = True
                        elif r_idx == 1:
                            # Word中表头下线可能表现为第1个正文行的上边框
                            has_header_bottom = True
                        else:
                            has_extra_line = True
                    elif name == 'bottom':
                        if r_idx == 0:
                            has_header_bottom = True
                        elif r_idx == row_count - 1:
                            has_bottom = True
                        else:
                            has_extra_line = True
                    elif name in ['insideH']:
                        has_header_bottom = True
                    elif name not in ['tl2br', 'tr2bl']:
                        has_extra_line = True

        if not has_any_border:
            return True  # 无边框布局表不按三线表扣分
        return has_top and has_header_bottom and has_bottom and not has_extra_line

    def _check_table_borders(self):
        """检查目录页后是否存在边框线数量不等于3的表格 (-3分)

        细则：目录页后任意一个表格的边框线数量不是3

        - 定位区间：按文档流顺序取"目录"标题段之后出现的所有顶层表格
          （未找到目录时，按整个文档处理）
        - 每个表格按办公软件"看到"的边框线数量单独计数（不跨表求和）：
          * 表格有 (行数+1) 个水平边框位置，(列数+1) 个垂直边框位置
          * 某位置有可见线 = 该位置对应的表级 tblBorders/@top|bottom|
            insideH|left|right|insideV|start|end 可见，或该位置任一单元格
            对应 tcBorders 中相邻的 top/bottom/left/right/start/end 可见
          * 可见 = w:val 不为 none/nil，且 w:sz > 0
          * 每个位置至多计 1 条线（与办公软件里视觉上一条线的观感一致）
        - 边框线数量 = 水平可见位置数 + 垂直可见位置数
        - 目录后只要有一个表格的边框线数量 ≠ 3，即触发 -3 扣分
        """
        if self.doc is None:
            print("[PASS] 表格边框线符合要求")
            return

        tables_to_check = self._top_level_tables_after_outline()
        if not tables_to_check:
            print("[PASS] 表格边框线符合要求")
            return

        fail_info = None
        for idx, table in enumerate(tables_to_check):
            count = self._count_table_visible_border_lines(table)
            if count != 3:
                fail_info = f"目录后第{idx + 1}个表格的边框线数量为{count}，不是3"
                break

        if fail_info:
            self.results['dimension2_deductions'].append(3)
            self.results['hit_points'].append(
                "-3：目录页后任意一个表格的边框线数量不是3"
            )
            print(f"[-3] {fail_info}")
        else:
            print("[PASS] 表格边框线符合要求")

    def _top_level_tables_after_outline(self):
        """按文档流顺序返回"目录"段之后的顶层表格对象

        使用 _build_flat_document_index 构建的扁平索引，穿透 w:sdt/
        w:sdtContent，与段落索引保持同一口径。
        """
        target_idx = self.outline_para_idx
        # 若未找到目录，纳入全部顶层表格
        if target_idx < 0:
            return list(self._flat_tables)
        # 找出目录段之后（按文档流）出现的顶层表格
        # 需要按 body 元素顺序判断表格与目录段的先后关系
        outline_p_elem = (
            self._flat_para_elements[target_idx]
            if target_idx < len(self._flat_para_elements)
            else None
        )
        if outline_p_elem is None:
            return list(self._flat_tables)

        # 深度遍历 body（穿透 sdt）按文档流顺序枚举 <w:p> / <w:tbl>
        passed = False
        result_elements = []

        def walk(container):
            nonlocal passed
            for child in container.iterchildren():
                tag = child.tag.split('}')[-1]
                if tag == 'p':
                    if not passed and child is outline_p_elem:
                        passed = True
                elif tag == 'tbl':
                    if passed:
                        result_elements.append(child)
                elif tag == 'sdt':
                    sdt_content = child.find(qn('w:sdtContent'))
                    if sdt_content is not None:
                        walk(sdt_content)

        walk(self.doc.element.body)

        result = []
        for tbl_elem in result_elements:
            for t in self._flat_tables:
                if t._element is tbl_elem:
                    result.append(t)
                    break
        return result

    def _count_table_visible_border_lines(self, table):
        """按办公软件视角统计表格可见边框线数量（水平+垂直）

        判定优先级与 OOXML/Word 一致：
          - 单元格 tcBorders 中若定义了对应边（存在即优先，无论 val），
            以该定义为准
          - 单元格未定义时才回退到表级 tblBorders 的默认边框
          - 可见 = w:val 不为 none/nil，且 w:sz > 0（缺 sz 视为不可见）
          - 一个"位置"由两侧相邻单元格共同决定：任一相邻侧判为可见即可见
          - 每个位置至多计 1 条线（与办公软件视觉观感一致）
        """
        rows = list(table.rows)
        row_count = len(rows)
        col_count = len(table.columns)
        if row_count == 0 or col_count == 0:
            return 0

        # 表级默认边框
        tbl_borders = {}
        tblPr = table._element.tblPr
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                for b in tblBorders:
                    name = b.tag.split('}')[-1]
                    tbl_borders[name] = b

        def cell_at(r_idx, c_idx):
            if not (0 <= r_idx < row_count) or not (0 <= c_idx < col_count):
                return None
            try:
                return rows[r_idx].cells[c_idx]
            except (IndexError, ValueError):
                return None

        def cell_border_element(cell, name):
            if cell is None:
                return None
            tcPr = cell._tc.tcPr
            if tcPr is None:
                return None
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                return None
            return tcBorders.find(qn(f'w:{name}'))

        def cell_side_visible(cell, side_names, table_default_names):
            """判定单元格某一侧的边框是否可见，考虑单元格覆盖表级默认"""
            if cell is None:
                return False
            # 单元格自身定义（含 nil）优先
            for name in side_names:
                b = cell_border_element(cell, name)
                if b is not None:
                    return self._is_visible_border(b)
            # 未定义则回退表级默认
            for name in table_default_names:
                if self._is_visible_border(tbl_borders.get(name)):
                    return True
            return False

        # 水平线：row_count + 1 个位置
        h_visible = 0
        for pos in range(row_count + 1):
            visible = False
            if pos == 0:
                # 顶边：由第 0 行所有单元格的 top 决定，回退 tblBorders.top
                for c in range(col_count):
                    if cell_side_visible(cell_at(0, c), ['top'], ['top']):
                        visible = True
                        break
            elif pos == row_count:
                # 底边：由末行所有单元格的 bottom 决定，回退 tblBorders.bottom
                for c in range(col_count):
                    if cell_side_visible(
                        cell_at(row_count - 1, c), ['bottom'], ['bottom']
                    ):
                        visible = True
                        break
            else:
                # 中间水平线：由上侧单元格的 bottom 或下侧单元格的 top 决定，
                # 回退 tblBorders.insideH
                for c in range(col_count):
                    if cell_side_visible(
                        cell_at(pos - 1, c), ['bottom'], ['insideH']
                    ) or cell_side_visible(
                        cell_at(pos, c), ['top'], ['insideH']
                    ):
                        visible = True
                        break
            if visible:
                h_visible += 1

        # 垂直线：col_count + 1 个位置
        v_visible = 0
        for pos in range(col_count + 1):
            visible = False
            if pos == 0:
                for r in range(row_count):
                    if cell_side_visible(
                        cell_at(r, 0), ['left', 'start'], ['left', 'start']
                    ):
                        visible = True
                        break
            elif pos == col_count:
                for r in range(row_count):
                    if cell_side_visible(
                        cell_at(r, col_count - 1),
                        ['right', 'end'],
                        ['right', 'end'],
                    ):
                        visible = True
                        break
            else:
                for r in range(row_count):
                    if cell_side_visible(
                        cell_at(r, pos - 1), ['right', 'end'], ['insideV']
                    ) or cell_side_visible(
                        cell_at(r, pos), ['left', 'start'], ['insideV']
                    ):
                        visible = True
                        break
            if visible:
                v_visible += 1

        return h_visible + v_visible

    def _calculate_total_score(self):
        """计算总分"""
        positive = sum(self.results['dimension2_scores'])
        negative = sum(self.results['dimension2_deductions'])
        self.results['total_score'] = positive - negative

    def _print_results(self):
        """打印评估结果"""
        print()
        print("=" * 60)
        print("评估结果汇总")
        print("=" * 60)

        if not self.results['dimension1_passed']:
            print("维度一：未通过，最终得分：0分")
            print("失败原因：")
            for failure in self.results['dimension1_failures']:
                print(f"  - {failure}")
            return

        print("维度一：通过")
        print("维度二：评分结果")
        for point in self.results['hit_points']:
            print(f"  {point}")

        print()
        print(f"最终得分：{self.results['total_score']}分")
        print("=" * 60)


# ---- 统一对外接口（见 "脚本接口差异与统一建议.md" §2.2 / §3） -------------
# 说明：本文件的内部评估逻辑、类结构、参数、每条评分点的判定语义均保持不变；
# 仅在此处新增模块级 evaluate(file_path) 作为唯一对外入口，把结果规范成
# §2.2 约定的 dict；类内原有 print 视为调试信息，运行时全部改走 stderr。

SCRIPT_ID = "002"


def _locate_office_file(dir_path: str) -> str:
    """在任务目录中定位唯一的非临时 DOCX 文档。"""
    if not os.path.isdir(dir_path):
        raise FileNotFoundError(f"目录不存在: {dir_path}")
    candidates = sorted(
        path for path in glob.glob(os.path.join(dir_path, "*.docx"))
        if not os.path.basename(path).startswith("~$")
    )
    if not candidates:
        raise FileNotFoundError(f"目录内未找到 .docx 文件: {dir_path}")
    return candidates[0]


def _resolve_office_file(path: str) -> str:
    """兼容批量运行器目录入口和本地调试的文件入口。"""
    if os.path.isdir(path):
        return _locate_office_file(path)
    if os.path.isfile(path):
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx" and not os.path.basename(path).startswith("~$"):
            return path
    raise FileNotFoundError(f"未找到 .docx 文件: {path}")


# 维度2各评分点静态元数据。hit_marker 与类内 hit_points 里追加的字符串
# 完全一致，用来判断该项是否命中；命令名与判定语义在类里，请不要在此改动。
# 元素结构：(kind, rule 简述, max_delta 正整数, 类内 hit_marker 原文)
_Dim2Rule = tuple[str, str, int, str]
_DIM2_RULES: list[_Dim2Rule] = [
    ("bonus", "摘要前所有页面页眉处没有横线", 3,
     "+3：摘要前所有页面页眉处没有横线"),
    ("bonus",
     ("摘要页及摘要页后所有页面页眉处存在两条相互平行的黑色横线，"
      + "上横线粗细不低于下横线的3倍，下横线粗细不超过1.5磅"), 5,
     ("+5：摘要页及摘要页后所有页面页眉处存在两条相互平行的黑色横线，"
      + "上横线粗细不低于下横线的3倍,下横线粗细不超过1.5磅")),
    ("bonus",
     "所有页页眉位置上的横线位于页眉文本下方，与页眉文本间距在1-4字符之间", 5,
     "+5：所有页页眉位置上的横线位于页眉文本下方，与页眉文本间距在1-4字符之间"),

    ("penalty",
     "任意一页页眉文本内容不是\"澄湖理工大学硕士学位论文\"或不在居中位置", 1,
     "-1：任意一页页眉文本内容不是\"澄湖理工大学硕士学位论文\"或不在居中位置"),
    ("penalty",
     ("目录页后所有页面不满足以下任意一项：页脚文本没有出现"
      + "\"城市社区微气候韧性评价与公共空间更新研究\"；"
      + "奇数页页脚文本不是左对齐；偶数页页脚文本不是右对齐"), 1,
     ("-1：目录页后所有页面不满足以下任意一项：页脚文本没有出现"
      + "\"城市社区微气候韧性评价与公共空间更新研究\"；"
      + "奇数页页脚文本不是左对齐；偶数页页脚文本不是右对齐")),
    ("penalty", "任意一页页脚居中位置没有出现页码", 1,
     "-1：任意一页页脚居中位置没有出现页码"),
    ("penalty",
     "目录页及目录页之前的页面页码不是罗马数字，目录页之后的页面页码不是阿拉伯数字", 1,
     ("-1：目录页及目录页之前的页面页码不是罗马数字，"
      + "目录页之后的页面页码不是阿拉伯数字")),
    ("penalty", "中文摘要页、英文摘要页、目录页不是此先后顺序", 1,
     "-1：中文摘要页、英文摘要页、目录页不是此先后顺序"),
    ("penalty", "目录页后所有一级标题不是黑体、三号、加粗", 1,
     "-1：目录页后所有一级标题不是黑体、三号、加粗"),
    ("penalty", "目录页后所有二级标题不是黑体、四号、加粗", 1,
     "-1：目录页后所有二级标题不是黑体、四号、加粗"),
    ("penalty", "目录页后所有三级标题不是黑体、小四号、加粗", 1,
     "-1：目录页后所有三级标题不是黑体、小四号、加粗"),
    ("penalty",
     ("\"1.1.1 研究背景\"与\"1.1.2 研究意义\"中间的文本不满足以下任意一项："
      + "宋体、小四、1.5倍行距、两端对齐"), 1,
     ("-1：\"1.1.1 研究背景\"与\"1.1.2 研究意义\"中间的文本不满足以下"
      + "任意一项：宋体、小四、1.5倍行距、两端对齐")),
    ("penalty", "目录页后任意一个表格的边框线数量不是3", 3,
     "-3：目录页后任意一个表格的边框线数量不是3"),
]


def evaluate(dir_path: str) -> dict[str, object]:
    """脚本对外唯一入口，接收任务目录并定位其中的 DOC/DOCX。

    - 不改动内部评估逻辑与参数，仍复用 WordDocumentEvaluator。
    - 类内的 print 全部丢弃（stdout/stderr 均重定向到 devnull），主结果只走 return。
    - 任何未捕获异常一律折成 status="error"，避免混淆"评估 0 分"与"脚本崩溃"。
    """
    dim2_items: list[dict[str, object]] = []
    # 满分 = 所有加分项 max_delta 之和（扣分项只影响最终得分，不计入满分）
    max_score: int = sum(m for kind, _rule, m, _mark in _DIM2_RULES
                         if kind == "bonus")
    result: dict[str, object] = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": True,
        "dim1_reason": "",
        "dim2_items": dim2_items,
        "total_score": 0,
        "max_score": max_score,
    }

    try:
        file_path = _resolve_office_file(dir_path)
        result["file_name"] = os.path.basename(file_path)

        evaluator = WordDocumentEvaluator(file_path, skip_page_check=True)
        # 类内所有 print 视为调试输出，全部丢弃；主结果仅通过 return 传出
        with open(os.devnull, "w", encoding="utf-8") as _null, \
                contextlib.redirect_stdout(_null), \
                contextlib.redirect_stderr(_null):
            evaluator.evaluate()
    except Exception as e:
        result["status"] = "error"
        result["error"] = f"{type(e).__name__}: {e}"
        return result

    res = evaluator.results  # type: ignore[assignment]

    dim1_passed_raw = res.get("dimension1_passed", True)
    if not bool(dim1_passed_raw):
        failures_raw = res.get("dimension1_failures", []) or []
        if isinstance(failures_raw, list):
            failures: list[str] = [str(x) for x in failures_raw]
        else:
            failures = []
        result["dim1_pass"] = False
        result["dim1_reason"] = "；".join(failures)
        result["total_score"] = 0
        return result

    hit_points_raw = res.get("hit_points", []) or []
    if isinstance(hit_points_raw, list):
        hit_set: set[str] = {str(x) for x in hit_points_raw}
    else:
        hit_set = set()
    for kind, rule_desc, max_delta, hit_marker in _DIM2_RULES:
        hit = hit_marker in hit_set
        # bonus 的 max_delta 天然为正，penalty 天然为负；命中时 delta = max_delta，
        # 未命中时 delta = 0。这里保留每条规则的正负号，不强行改写。
        signed_max = max_delta if kind == "bonus" else -max_delta
        dim2_items.append({
            "rule": rule_desc,
            "max_delta": signed_max,
            "delta": signed_max if hit else 0,
            "hit": hit,
            "detail": "",
        })

    total_raw = res.get("total_score", 0)
    result["total_score"] = int(total_raw) if isinstance(total_raw, (int, bool)) else 0
    return result


def main():
    """本地调试入口：evaluate(sys.argv[1]) 后打印 JSON 主结果到 stdout。

    仅用于脚本作者自测；批量运行器直接 import 本模块的 evaluate 即可。
    """
    # Windows 默认控制台可能是 cp1252/gbk，直接 print 中文 JSON 会崩溃；
    # 统一把 stdout 切到 UTF-8。
    try:
        sys.stdout.reconfigure(encoding="utf-8")  # type: ignore[attr-defined]
    except Exception:
        pass

    if len(sys.argv) >= 2:
        target = sys.argv[1]
    else:
        # 无参数：兜底扫描脚本同目录下的 .docx 用于自测，仅走第一个
        script_dir = os.path.dirname(os.path.abspath(__file__))
        office_files = sorted(
            f for f in glob.glob(os.path.join(script_dir, "*.docx"))
            if not os.path.basename(f).startswith("~$")
        )
        if not office_files:
            print(json.dumps({
                "id": SCRIPT_ID,
                "file_name": "",
                "status": "error",
                "error": "未提供 file_path，且脚本目录下无 .docx 可供自测",
            }, ensure_ascii=False, indent=2))
            return
        target = office_files[0]

    print(json.dumps(evaluate(target), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
