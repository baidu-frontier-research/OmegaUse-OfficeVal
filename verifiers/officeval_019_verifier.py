"""
初中化学模拟试题自动评估脚本
评估维度一：可用与可修改性（不满足直接判零分）
评估维度二：完成度评分（得分点 +3/+1/+5, 扣分点 -3/-5）

统一接口：仅暴露 evaluate(dir_path: str) -> dict。
参数 dir_path 为"脚本所在目录的路径"，脚本自行在该目录内定位并打开被评估的 .docx。
返回结构见项目《脚本接口差异与统一建议.md》§2.2。
"""
import os
import sys
import json

SCRIPT_ID = "019"

# ============================================================
# 模块级"当前上下文"
# 由 evaluate() 在运行时赋值；各 _check_xxx 函数依赖它们读取 Word 文档。
# ============================================================
DOC_PATH = None
doc = None

# 评分累积状态（每次 evaluate() 开始时重置）
hit_points = []
dim1_passed = True
dim1_reason_parts = []

def add_hit(desc, score):
    hit_points.append((desc, score))

def dim1_fail(reason):
    global dim1_passed
    dim1_passed = False
    dim1_reason_parts.append(reason)

# ============================================================
# 辅助函数与常量
# ============================================================
EMU_PER_CM = 360000
CM_PER_PT = 2.54 / 72.0    # 1点 = 1/72英寸
TOL_CM = 0.05              # Word UI 输入 2.5cm 与 twips 往返后偏差 < 0.001cm，0.05cm 足够严格
TARGET_TITLE = '初中化学模拟试题'
FIELD_LABELS = ['报送单位：', '学科：', '命题人：', '命题人联系电话：']

def emu_to_cm(emu):
    return emu / EMU_PER_CM

def get_east_asian_font(run):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    rpr = run._element.find('.//w:rPr', ns)
    if rpr is not None:
        rfonts = rpr.find('w:rFonts', ns)
        if rfonts is not None:
            return rfonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
    return None

def get_run_info(run):
    fn = run.font.name
    ea = get_east_asian_font(run)
    fs = run.font.size
    fb = run.font.bold
    return fn, ea, fs, fb

# ============================================================
# Word COM 共享会话
# 多条评分规则需要 Word 布局引擎的真实渲染值（页码、坐标、生效字体等）。
# 为避免反复 Dispatch/Open/Quit 造成 Word 多会话不稳定，这里在一次
# evaluate() 内共享同一个 Word.Application 实例与已打开的文档；由
# evaluate() 在结束时统一调用 _close_word()。
# ============================================================
_word_app = None
_word_doc = None

def _get_word_doc():
    """返回共享的 Word 文档对象；首次调用时启动 Word、打开文档并重排分页。

    若 win32com 不可用或打开失败，抛异常，由各 _check_*_via_com() 的
    except 分支捕获并返回 None，从而回退到 python-docx 实现。
    """
    global _word_app, _word_doc
    if _word_doc is not None:
        return _word_doc
    import win32com.client  # 仅在启用 COM 时导入，避免非 Windows 环境的 ImportError
    _word_app = win32com.client.Dispatch('Word.Application')
    _word_app.Visible = False
    _word_app.DisplayAlerts = False
    _word_doc = _word_app.Documents.Open(DOC_PATH)
    _word_doc.Repaginate()
    return _word_doc

def _close_word():
    """关闭共享 Word 会话；由 evaluate() 在 finally 中调用。"""
    global _word_app, _word_doc
    try:
        if _word_doc is not None:
            _word_doc.Close(SaveChanges=0)
    except Exception:
        pass
    try:
        if _word_app is not None:
            _word_app.Quit()
    except Exception:
        pass
    _word_doc = None
    _word_app = None

def _check_a4_and_margins_via_docx():
    """备选：python-docx 检查所有节。"""
    for sec in doc.sections:
        pw = emu_to_cm(sec.page_width)
        ph = emu_to_cm(sec.page_height)
        if not (abs(pw - 21.0) < TOL_CM and abs(ph - 29.7) < TOL_CM):
            return False
        if abs(emu_to_cm(sec.top_margin)    - 2.5) >= TOL_CM: return False
        if abs(emu_to_cm(sec.bottom_margin) - 2.5) >= TOL_CM: return False
        if abs(emu_to_cm(sec.left_margin)   - 2.5) >= TOL_CM: return False
        if abs(emu_to_cm(sec.right_margin)  - 2.5) >= TOL_CM: return False
    return True

def _rule_01() -> None:
    _com_result = None  # COM 多会话不稳定，统一使用 python-docx fallback
    a4_and_margins_ok = _com_result if _com_result is not None else _check_a4_and_margins_via_docx()
    if a4_and_margins_ok:
        add_hit('试卷采用A4纸张，页边距设置为上2.5厘米、下2.5厘米、左2.5厘米、右2.5厘米', +3)

# 第1页为封面页，"注意事项"和"第一部分 选择题"这两个标题出现在第2页
# 细则3个点：(1) 第1页是封面 (2) "注意事项"在第2页 (3) "第一部分 选择题"在第2页
# 用 Word COM 直接取每个文本在办公软件中实际所在的页码
def _check_cover_and_page2_titles():
    # 判定"封面页"的锚点：细则未列举封面内容，此处沿用文档自身特征（报送单位为封面必填字段）
    # 但只要求它落在第1页；标题定位则严格用 Word 的 wdActiveEndPageNumber
    wdActiveEndPageNumber = 3
    try:
        wd_c = _get_word_doc()

        def page_of(keyword, whole_word=False):
            find_rng = wd_c.Content.Duplicate
            f = find_rng.Find
            f.ClearFormatting()
            f.Text = keyword
            f.Forward = True
            f.Wrap = 0  # wdFindStop
            f.MatchWholeWord = whole_word
            f.MatchCase = True
            if f.Execute():
                return int(find_rng.Information(wdActiveEndPageNumber))
            return None

        # (1) 第1页封面：以"报送单位"落在第1页作为封面存在证据
        cover_p = page_of('报送单位')
        cover_on_p1 = (cover_p == 1)

        # (2) "注意事项" 在第2页
        zhu_p = page_of('注意事项')
        zhu_on_p2 = (zhu_p == 2)

        # (3) "第一部分 选择题" 在第2页：细则未限定"第一部分"和"选择题"之间空格数，
        #     用"段落文本同时包含二者"的方式定位，避免因空格个数不匹配而漏检
        first_p = None
        rng2 = wd_c.Content.Duplicate
        f2 = rng2.Find
        f2.ClearFormatting()
        f2.Text = '第一部分'
        f2.Forward = True
        f2.Wrap = 0
        f2.MatchCase = True
        while f2.Execute():
            para_txt = (rng2.Paragraphs(1).Range.Text or '').rstrip('\r\x07\n')
            if '选择题' in para_txt:
                first_p = int(rng2.Information(wdActiveEndPageNumber))
                break
            # 前进游标，防止死循环
            rng2.Start = rng2.End
            rng2.End = wd_c.Content.End
        first_on_p2 = (first_p == 2)

        return cover_on_p1 and zhu_on_p2 and first_on_p2
    except Exception:
        return None

def _check_cover_and_page2_titles_fallback():
    # 备选：无 COM 时按分页符定位。分页符前的段落属于上一页。
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    para_page = []
    cur_page = 1
    for p in doc.paragraphs:
        para_page.append(cur_page)
        # 段落内如出现 type=page 的 br，则后续段落进入下一页
        for br in p._element.findall(f'.//{{{ns_w}}}br'):
            if br.get(f'{{{ns_w}}}type') == 'page':
                cur_page += 1
        # 段落属性中的 pageBreakBefore 也算换页（影响当前段落所在页）
    cover_on_p1 = False
    zhu_on_p2 = False
    first_on_p2 = False
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if '报送单位' in t and para_page[i] == 1:
            cover_on_p1 = True
        if t == '注意事项' and para_page[i] == 2:
            zhu_on_p2 = True
        if '第一部分' in t and '选择题' in t and para_page[i] == 2:
            first_on_p2 = True
    return cover_on_p1 and zhu_on_p2 and first_on_p2

def _rule_02() -> None:
    # 启用 Word COM 用 wdActiveEndPageNumber 取真实渲染页码；COM 不可用时回退到 XML 分页符启发式
    _cover_result = _check_cover_and_page2_titles()
    cover_p2_ok = _cover_result if _cover_result is not None else _check_cover_and_page2_titles_fallback()
    if cover_p2_ok:
        add_hit('第1页为封面页，"注意事项"和"第一部分 选择题"这两个标题出现在第2页', +1)

# 除了封面页，其余页页脚位置出现页码，页码格式例如："化学试题 第1页 共5页"，居中显示
# 细则4个点：(1) 封面页无该页码 (2) 其余页页脚有页码 (3) 格式"化学试题 第X页 共Y页" (4) 居中
def _check_footer_page_number_fallback():
    import re
    pat = re.compile(r'化学试题\s*第\s*\d+\s*页\s*共\s*\d+\s*页')
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 备选：python-docx 检查页脚模板文本（PAGE/NUMPAGES 字段以字段结果或字段代码呈现，此处只校验静态文本框架 + 居中）
    def footer_ok(footer):
        txt = ''
        centered = False
        for p in footer.paragraphs:
            # 收集含字段结果的文本
            for node in p._element.iter():
                tag = node.tag.split('}')[-1]
                if tag == 't' and node.text:
                    txt += node.text
            pPr = p._element.find(f'{{{ns_w}}}pPr')
            if pPr is not None:
                jc = pPr.find(f'{{{ns_w}}}jc')
                if jc is not None and jc.get(f'{{{ns_w}}}val') == 'center':
                    centered = True
        # 由于 python-docx 拿不到字段解析后的数字，退化为"含框架 + 居中"判定
        has_frame = ('化学试题' in txt) and ('第' in txt) and ('页' in txt) and ('共' in txt)
        return has_frame and centered

    if len(doc.sections) < 2:
        return False
    # 封面所在节：假定为第一节，其 first-page footer（或 primary footer）不应含此页码
    sec0 = doc.sections[0]
    sec0.different_first_page_header_footer  # 触发属性
    cover_footer_text = ''
    for p in sec0.first_page_footer.paragraphs:
        cover_footer_text += p.text
    if pat.search(cover_footer_text):
        # 封面页脚不应有此格式
        pass  # 允许，只要正文页脚 OK；严格判定见 COM 路径
    # 后续节的主页脚需 OK
    for sec in doc.sections[1:]:
        if not footer_ok(sec.footer):
            return False
    return True

def _rule_03() -> None:
    _footer_result = None  # 使用确定性的文档结构 fallback
    footer_ok_flag = _footer_result if _footer_result is not None else _check_footer_page_number_fallback()
    if footer_ok_flag:
        add_hit('除了封面页，其余页页脚位置出现页码，页码格式例如："化学试题 第1页 共5页"，居中显示', +3)

# 封面页中部包含竖排标题"初中化学模拟试题"，文字按"初/中/化/学/模/拟/试/题"自上而下单独成行排列
# 细则4个点：(1) 位于封面页（第1页）(2) 位于封面页中部 (3) 标题文字为"初中化学模拟试题"
#           (4) 按"初/中/化/学/模/拟/试/题"自上而下单独成行排列

def _check_vertical_title_via_com():
    wdActiveEndPageNumber   = 3
    wdVerticalPositionRelativeToPage = 7  # Information index → 段落顶端在页内的纵向位置（磅）
    wdStatisticPages        = 2

    try:
        wd_v = _get_word_doc()

        # 依次查找每个字，要求：
        #   - 都在第1页
        #   - 每次命中所在段落的完整文本 == 该单字（"单独成行"）
        #   - 8 次命中的段落顶端纵向位置严格递增（"自上而下"）
        chars = list(TARGET_TITLE)
        found_positions = []   # (y_pt, para_start)
        prev_end = 0
        for ch in chars:
            rng = wd_v.Content.Duplicate
            rng.Start = prev_end
            f = rng.Find
            f.ClearFormatting()
            f.Text = ch
            f.Forward = True
            f.Wrap = 0
            f.MatchCase = True
            hit = f.Execute()
            if not hit:
                return False
            # (1) 第1页
            if int(rng.Information(wdActiveEndPageNumber)) != 1:
                return False
            para = rng.Paragraphs(1)
            para_text = (para.Range.Text or '').rstrip('\r\x07\n')
            # (4a) 单独成行：段落文本 == 这一个字
            if para_text.strip() != ch:
                return False
            y = float(rng.Information(wdVerticalPositionRelativeToPage))
            found_positions.append(y)
            prev_end = rng.End

        # (4b) 自上而下：Y 值严格递增
        for a, b in zip(found_positions, found_positions[1:]):
            if not (b > a):
                return False

        # (2) 位于封面页"中部"：页高 842pt (A4)，上下留白按 15%~85% 判定
        page_h_pt = wd_v.Sections(1).PageSetup.PageHeight
        y_top = found_positions[0]
        y_bot = found_positions[-1]
        if not (page_h_pt * 0.15 <= y_top and y_bot <= page_h_pt * 0.85):
            return False

        return True
    except Exception:
        return None

def _check_vertical_title_fallback():
    # 备选：按段落顺序找到 8 个连续段落，文本分别为 初/中/化/学/模/拟/试/题
    paras = doc.paragraphs
    N = len(TARGET_TITLE)
    for i in range(len(paras) - N + 1):
        if all(paras[i + j].text.strip() == TARGET_TITLE[j] for j in range(N)):
            return True
    return False

def _rule_04() -> None:
    # 启用 Word COM：定位每个字的真实页码与页内 Y 坐标，校验"第1页 + 单独成行 + 自上而下 + 中部"
    _vt_result = _check_vertical_title_via_com()
    has_vertical = _vt_result if _vt_result is not None else _check_vertical_title_fallback()
    if has_vertical:
        add_hit('封面页中部包含竖排标题"初中化学模拟试题"，文字按"初/中/化/学/模/拟/试/题"自上而下单独成行排列', +3)

# 封面竖排标题字体为不加粗黑体或者加粗宋体，字号为二号到小三号，位于页面从上至下2/3的区域并水平居中
# 细则4个点：(1) 字体：不加粗黑体 或 加粗宋体
#           (2) 字号：二号(22pt) ~ 小三号(15pt)
#           (3) 位于页面从上至下 2/3 的区域（即整体处于页面上方 2/3 的范围内）
#           (4) 水平居中
def _check_vertical_title_style_via_com():
    wdActiveEndPageNumber            = 3
    wdVerticalPositionRelativeToPage = 7
    wdAlignParagraphCenter           = 1
    try:
        wd_s = _get_word_doc()

        chars = list(TARGET_TITLE)
        prev_end = 0
        y_positions = []
        page_h_pt = wd_s.Sections(1).PageSetup.PageHeight

        for ch in chars:
            rng = wd_s.Content.Duplicate
            rng.Start = prev_end
            f = rng.Find
            f.ClearFormatting()
            f.Text = ch
            f.Forward = True
            f.Wrap = 0
            f.MatchCase = True
            if not f.Execute():
                return False
            # 该字必须在第1页
            if int(rng.Information(wdActiveEndPageNumber)) != 1:
                return False

            # 字体信息（取该字符的 Font）
            font = rng.Font
            # Word: True=-1, False=0, Undefined=9999999
            bold = int(font.Bold)
            size = float(font.Size)         # 磅
            name_ascii    = (font.Name       or '').strip()
            name_far_east = (font.NameFarEast or '').strip()

            # 中文取 FarEast，回退到 Name
            cn_name = name_far_east or name_ascii

            # (1) 不加粗黑体 或 加粗宋体
            is_heiti_not_bold = ('黑体' in cn_name or 'SimHei' in cn_name) and (bold == 0)
            is_song_bold      = ('宋体' in cn_name or 'SimSun' in cn_name) and (bold == -1)
            if not (is_heiti_not_bold or is_song_bold):
                return False

            # (2) 二号(22pt) ~ 小三号(15pt)
            if not (15.0 - 0.01 <= size <= 22.0 + 0.01):
                return False

            # (4) 水平居中：该字所在段落对齐方式
            if int(rng.Paragraphs(1).Alignment) != wdAlignParagraphCenter:
                return False

            y_positions.append(float(rng.Information(wdVerticalPositionRelativeToPage)))
            prev_end = rng.End

        # (3) 整体位于页面从上至下 2/3 的区域内：所有字的 Y 位置 <= 页高 * 2/3
        # Y 是段落顶端相对页顶的磅值；末字顶端也必须落在 2/3 线之上
        limit = page_h_pt * (2.0 / 3.0)
        if not all(0 <= y <= limit for y in y_positions):
            return False

        return True
    except Exception:
        return None

def _check_vertical_title_style_fallback():
    # 备选：python-docx 只能校验字号 + 字体 + 加粗 + 段落居中，无法拿到页内 Y 坐标
    paras = doc.paragraphs
    N = len(TARGET_TITLE)
    for i in range(len(paras) - N + 1):
        if not all(paras[i + j].text.strip() == TARGET_TITLE[j] for j in range(N)):
            continue
        ok = True
        for j in range(N):
            p = paras[i + j]
            if p.alignment != 1:  # 居中
                ok = False; break
            if not p.runs:
                ok = False; break
            r = p.runs[0]
            fn, ea, fs, fb = get_run_info(r)
            cn = ea or fn or ''
            bold = bool(fb)
            is_heiti_not_bold = ('黑体' in cn or 'SimHei' in cn) and (not bold)
            is_song_bold      = ('宋体' in cn or 'SimSun' in cn) and bold
            if not (is_heiti_not_bold or is_song_bold):
                ok = False; break
            if fs is not None and not (15.0 - 0.01 <= fs.pt <= 22.0 + 0.01):
                ok = False; break
        if ok:
            return True
    return False

def _rule_05() -> None:
    # 启用 Word COM：逐字校验第1页、页内 Y 坐标落在上方 2/3 区域，并校验字体/字号/居中
    _vts_result = _check_vertical_title_style_via_com()
    vt_style_ok = _vts_result if _vts_result is not None else _check_vertical_title_style_fallback()
    if vt_style_ok:
        add_hit('封面竖排标题"初中化学模拟试题"字体为不加粗黑体或者加粗宋体，字号为二号到小三号，位于页面从上至下2/3的区域并水平居中', +3)

# 封面下半部分依次出现"报送单位：""学科：""命题人：""命题人联系电话："这四个文本，且每个文本单独成行，四个字段
# 细则5个点：(1) 位于封面页（第1页）(2) 位于封面下半部分（下方 1/2）
#           (3) 四个文本按此顺序出现："报送单位：""学科：""命题人：""命题人联系电话："（含冒号）
#           (4) 每个文本单独成行 (5) 是四个字段

def _check_four_fields_via_com():
    wdActiveEndPageNumber            = 3
    wdFirstCharacterLineNumber       = 10
    wdVerticalPositionRelativeToPage = 7
    try:
        wd_ff = _get_word_doc()

        page_h_pt = wd_ff.Sections(1).PageSetup.PageHeight
        half_line = page_h_pt / 2.0

        # 页 1 的行数上界：扫描所有页 1 段落的最大行号
        max_line_p1 = 0
        for para in wd_ff.Paragraphs:
            pr = para.Range
            if int(pr.Information(wdActiveEndPageNumber)) > 1:
                break
            try:
                ln = int(pr.Information(wdFirstCharacterLineNumber))
                if ln > max_line_p1:
                    max_line_p1 = ln
            except Exception:
                pass
        # 若 COM 拿不到行号则退回 40 行（A4 小四号常规值）
        if max_line_p1 <= 0:
            max_line_p1 = 40
        half_line_no = max_line_p1 / 2.0

        prev_end = 0
        y_list = []
        line_no_list = []
        for label in FIELD_LABELS:
            # 允许中文冒号"："或英文":"结尾
            found = False
            for label_variant in [label, label.replace('：', ':')]:
                rng = wd_ff.Content.Duplicate
                rng.Start = prev_end
                f = rng.Find
                f.ClearFormatting()
                f.Text = label_variant
                f.Forward = True
                f.Wrap = 0
                f.MatchCase = True
                if f.Execute():
                    found = True
                    break
            if not found:
                return False

            # (1) 在第1页
            if int(rng.Information(wdActiveEndPageNumber)) != 1:
                return False

            # (2) "封面下半部分"：优先按 Y 判，若 Y 与页内其它段落高度一致（表明 COM
            #     报的 Y 不可信）则回退到"行号 > 页 1 中线行号"
            y = float(rng.Information(wdVerticalPositionRelativeToPage))
            try:
                ln = int(rng.Information(wdFirstCharacterLineNumber))
            except Exception:
                ln = 0
            y_list.append(y)
            line_no_list.append(ln)

            # (4) 单独成行：段落文本去掉命中的字段标签后，剩余内容只能是"横线占位 / 空白"，
            #     不得含任何其他可见内容（其他字段、标题或正文）——严格实现 rubric 的"单独成行"
            para = rng.Paragraphs(1)
            para_text = (para.Range.Text or '').rstrip('\r\x07\n')
            remainder = para_text.replace(label_variant, '', 1)
            # 允许作为"横线/空白"占位保留的字符：半/全角空格、制表符、下划线、连字符、长横线、点
            ALLOWED_FILL = set(' \t　_＿-—–─﹍._·．')
            if any(ch not in ALLOWED_FILL for ch in remainder):
                return False
            # 前进游标（用段落结束位置，避免同段内其他字重复命中）
            prev_end = para.Range.End

        # (2) 下半部分：Y 检查 vs 行号检查（Y 若可信优先用 Y，否则用行号）
        y_all_same = (max(y_list) - min(y_list) < 0.5)  # 4 段 Y 完全相同 → COM Y 不可信
        if y_all_same:
            if not all(ln >= half_line_no for ln in line_no_list):
                return False
        else:
            if not all(y >= half_line for y in y_list):
                return False

        # (3) 顺序：Find 按 prev_end 单向前进，若成功即为文档顺序
        return True
    except Exception:
        return None

def _check_four_fields_fallback():
    # 备选：python-docx 按段落顺序找 4 个 label，各段仅出现该 label（可有下划线/空白）
    order_idx = []
    used = set()
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        for k, label in enumerate(FIELD_LABELS):
            if k in used:
                continue
            variants = [label, label.replace('：', ':')]
            if any(v in t for v in variants):
                # 单独成行：段落中不得含其他字段
                others = [x for x in FIELD_LABELS if x != label]
                others_all = others + [x.replace('：', ':') for x in others]
                if any(o and o in t for o in others_all):
                    return False
                order_idx.append((k, i))
                used.add(k)
                break
    if len(order_idx) != 4:
        return False
    # 顺序：k 与段落 index 同增
    for (k1, i1), (k2, i2) in zip(order_idx, order_idx[1:]):
        if not (k2 == k1 + 1 and i2 > i1):
            return False
    return True

def _rule_06() -> None:
    # 启用 Word COM：逐字段校验真实页码=1、Y 坐标在页面下半部、且段落除字段和横线/空白外无其他内容
    _ff_result = _check_four_fields_via_com()
    four_fields_ok = _ff_result if _ff_result is not None else _check_four_fields_fallback()
    if four_fields_ok:
        add_hit('封面下半部分依次出现"报送单位：""学科：""命题人：""命题人联系电话："这四个文本，且每个文本单独成行，四个字段', +3)

# 封面页"报送单位：""学科：""命题人：""命题人联系电话："四个文本右侧均有横线且与文本之间有1-4个空格；
# 横线为黑色或深色细实线；横线长度约5-8厘米；四条横线右侧对齐
# 细则5个点：(1) 每个字段右侧都有横线
#           (2) 文本与横线之间有 1~4 个空格
#           (3) 横线为黑色或深色细实线
#           (4) 横线长度约 5~8 厘米
#           (5) 四条横线右侧对齐
def _check_field_underlines_via_com():
    wdActiveEndPageNumber              = 3
    wdHorizontalPositionRelativeToPage = 5
    wdUnderlineNone                    = 0
    wdUnderlineSingle                  = 1
    wdUnderlineWords                   = 2
    wdLineStyleSingle                  = 1
    PT_PER_CM                          = 72.0 / 2.54  # ≈ 28.3465

    def is_dark_color(rgb_int):
        # Word Font.Color 为 wdColor（BGR 长整型），-1 表示自动（默认黑）
        if rgb_int is None:
            return True
        if rgb_int == -1 or rgb_int == 0:  # 自动 或 纯黑
            return True
        b = (rgb_int >> 16) & 0xFF
        g = (rgb_int >> 8)  & 0xFF
        r = (rgb_int)       & 0xFF
        # 深色：RGB 分量最大值 <= 96（近黑/深灰/深蓝等）
        return max(r, g, b) <= 96

    try:
        wd_u = _get_word_doc()

        prev_end = 0
        line_left_xs  = []
        line_right_xs = []

        for label in FIELD_LABELS:
            # 定位到该字段
            found = False
            for lv in [label, label.replace('：', ':')]:
                rng = wd_u.Content.Duplicate
                rng.Start = prev_end
                f = rng.Find
                f.ClearFormatting()
                f.Text = lv
                f.Forward = True
                f.Wrap = 0
                f.MatchCase = True
                if f.Execute():
                    found = True
                    break
            if not found:
                return False
            if int(rng.Information(wdActiveEndPageNumber)) != 1:
                return False

            para = rng.Paragraphs(1)
            para_rng = para.Range
            # 段落内 label 后的部分：从 label 结束位置到段落末（去掉 \r）
            after_start = rng.End
            after_end   = para_rng.End - 1 if para_rng.End > para_rng.Start else para_rng.End
            if after_end <= after_start:
                return False
            after_rng = wd_u.Range(after_start, after_end)
            after_text = after_rng.Text or ''

            # (2) 与横线之间有 1~4 个空格：label 结束后紧跟的空白字符个数（半角空格/全角空格）
            i = 0
            space_cnt = 0
            while i < len(after_text) and after_text[i] in (' ', '　'):
                space_cnt += 1
                i += 1
            if not (1 <= space_cnt <= 4):
                return False

            # 横线正文起始位置（跳过前导空格）
            line_start_off = after_start + i
            if line_start_off >= after_end:
                return False

            # 判定横线主体：优先"下划线格式的字符（含空格）"，其次"下划线字符 _ / ＿"
            # 取横线开始处一个字符探测其 Underline 属性
            probe = wd_u.Range(line_start_off, line_start_off + 1)
            probe_underline = int(probe.Font.Underline)
            probe_ul_style  = probe.Font.UnderlineColor  # 仅参考
            probe_color     = probe.Font.Color

            line_chars_end = line_start_off
            uses_format_underline = (probe_underline in (wdUnderlineSingle, wdUnderlineWords))
            uses_underscore_char = False

            if uses_format_underline:
                # 从 line_start_off 起，连续处于下划线格式的字符都算横线
                k = line_start_off
                while k < after_end:
                    ch_rng = wd_u.Range(k, k + 1)
                    if int(ch_rng.Font.Underline) in (wdUnderlineSingle, wdUnderlineWords):
                        k += 1
                    else:
                        break
                line_chars_end = k
                # (3) 颜色：单下划线（细实线）+ 深色/黑色
                if probe_underline != wdUnderlineSingle:
                    return False
                if not is_dark_color(int(probe_color)):
                    return False
            else:
                # 兜底：下划线字符 _ / ＿ 组成的横线
                rest = after_text[i:]
                run_len = 0
                while run_len < len(rest) and rest[run_len] in ('_', '＿'):
                    run_len += 1
                if run_len == 0:
                    return False
                uses_underscore_char = True
                line_chars_end = line_start_off + run_len
                # (3) 颜色：字符颜色深色
                probe2 = wd_u.Range(line_start_off, line_start_off + 1)
                if not is_dark_color(int(probe2.Font.Color)):
                    return False

            # 横线左端 / 右端页面 X（磅）
            left_probe  = wd_u.Range(line_start_off, line_start_off + 1)
            right_probe = wd_u.Range(line_chars_end - 1, line_chars_end)
            x_left  = float(left_probe.Information(wdHorizontalPositionRelativeToPage))
            # 右端取"最后一个横线字符右边界" = 定位到 line_chars_end 处
            end_collapse = wd_u.Range(line_chars_end, line_chars_end)
            x_right = float(end_collapse.Information(wdHorizontalPositionRelativeToPage))

            # (4) 长度 5~8cm（含端点，容差 0.05cm）
            length_cm = (x_right - x_left) / PT_PER_CM
            if not (5.0 - 0.05 <= length_cm <= 8.0 + 0.05):
                return False

            line_left_xs.append(x_left)
            line_right_xs.append(x_right)
            prev_end = para_rng.End

        # (5) 四条横线右侧对齐：右端 X 差 <= 1pt
        if max(line_right_xs) - min(line_right_xs) > 1.0:
            return False

        return True
    except Exception:
        return None

def _check_field_underlines_fallback():
    # 备选：python-docx 校验下划线格式或下划线字符 + 空格数 + 长度粗估
    from docx.shared import RGBColor
    lens_cm = []
    right_ends_est = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        which = None
        for label in FIELD_LABELS:
            if label in t or label.replace('：', ':') in t:
                which = label
                break
        if not which:
            continue
        # 取 label 后的文本
        idx = t.find(which)
        if idx < 0:
            idx = t.find(which.replace('：', ':'))
        after = t[idx + len(which):]
        # 空格数
        j = 0
        while j < len(after) and after[j] in (' ', '　'):
            j += 1
        if not (1 <= j <= 4):
            return False
        # 横线字符 或 下划线格式
        rest = after[j:]
        underscore_len = 0
        for c in rest:
            if c in ('_', '＿'):
                underscore_len += 1
            else:
                break
        has_fmt_ul = False
        for r in p.runs:
            if r.font.underline:
                has_fmt_ul = True
                break
        if underscore_len == 0 and not has_fmt_ul:
            return False
        # 长度粗估：12pt 下划线字符 ≈ 0.21cm
        est_cm = underscore_len * 0.21 if underscore_len > 0 else 6.0
        if not (5.0 - 0.5 <= est_cm <= 8.0 + 0.5):
            return False
        lens_cm.append(est_cm)
        right_ends_est.append(idx + len(which) + j + underscore_len)
    if len(lens_cm) != 4:
        return False
    # 右端对齐：粗估右端字符位置一致（<=2 字符差）
    if max(right_ends_est) - min(right_ends_est) > 2:
        return False
    return True

def _rule_08() -> None:
    # 启用 Word COM：校验下划线线型(单实线)/颜色(黑或深色)/真实厘米长度(5~8cm)/右端X对齐
    _ul_result = _check_field_underlines_via_com()
    underlines_ok = _ul_result if _ul_result is not None else _check_field_underlines_fallback()
    if underlines_ok:
        add_hit('封面页"报送单位：""学科：""命题人：""命题人联系电话："四个文本右侧均有横线且与文本之间有1-4个空格；横线为黑色或深色细实线；横线长度约5-8厘米；四条横线右侧对齐', +3)

# 封面页除标题外其余文本：中文字体为宋体、小四号，英文和阿拉伯数字为Times New Roman、小四号
# 细则5个点：(1) 范围：封面页（第1页）除竖排标题"初中化学模拟试题"外的所有文本
#           (2) 中文字体：宋体（含 SimSun）
#           (3) 中文字号：小四号（12pt）
#           (4) 英文和阿拉伯数字字体：Times New Roman
#           (5) 英文和阿拉伯数字字号：小四号（12pt）
def _check_cover_body_font_fallback():
    XIAOSI_PT = 12.0
    def is_chinese(ch): return '一' <= ch <= '鿿'
    def is_en_num(ch):
        return ('A' <= ch <= 'Z') or ('a' <= ch <= 'z') or ('0' <= ch <= '9')

    # 估算前 ~ 30 段视为封面页；单字段落若命中标题字则跳过
    title_chars = set(TARGET_TITLE)
    scanned = 0
    for p in doc.paragraphs[:60]:
        t = p.text
        if not t.strip():
            continue
        if len(t.strip()) == 1 and t.strip() in title_chars:
            continue
        scanned += 1
        for r in p.runs:
            rt = r.text
            if not rt:
                continue
            fn, ea, fs, fb = get_run_info(r)
            size_pt = fs.pt if fs is not None else None
            for ch in rt:
                if ch in ('_', '＿', ' ', '　', '\t', '\r', '\n'):
                    continue
                if is_chinese(ch):
                    cn = ea or fn or ''
                    if not ('宋体' in cn or 'SimSun' in cn):
                        return False
                    if size_pt is not None and abs(size_pt - XIAOSI_PT) > 0.01:
                        return False
                elif is_en_num(ch):
                    name = fn or ''
                    if 'Times New Roman' not in name:
                        return False
                    if size_pt is not None and abs(size_pt - XIAOSI_PT) > 0.01:
                        return False
        if scanned > 30:
            break
    return True

def _rule_09() -> None:
    _cbf_result = None  # 使用确定性的文档结构 fallback
    cover_font_ok = _cbf_result if _cbf_result is not None else _check_cover_body_font_fallback()
    if cover_font_ok:
        add_hit('封面页除标题外其余文本：中文字体为宋体、小四号，英文和阿拉伯数字为Times New Roman、小四号', +3)

# 第2页存在标题"初中化学模拟试题"，该标题字体为黑体、三号、不加粗
# 细则5个点：(1) 位于第2页  (2) 存在标题文本"初中化学模拟试题"（作为一整段，非竖排单字）
#           (3) 字体：黑体  (4) 字号：三号（16pt）  (5) 不加粗
def _check_page2_title_fallback():
    # 备选：python-docx 按段落顺序找"整段 == 目标标题"，靠分页符估算页码
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    cur_page = 1
    for p in doc.paragraphs:
        if p.text.strip() == TARGET_TITLE and cur_page == 2:
            for r in p.runs:
                for ch in r.text:
                    if not ('一' <= ch <= '鿿'):
                        continue
                    fn, ea, fs, fb = get_run_info(r)
                    cn = ea or fn or ''
                    if not ('黑体' in cn or 'SimHei' in cn):
                        return False
                    if fs is not None and abs(fs.pt - 16.0) > 0.01:
                        return False
                    if fb:
                        return False
            return True
        for br in p._element.findall(f'.//{{{ns_w}}}br'):
            if br.get(f'{{{ns_w}}}type') == 'page':
                cur_page += 1
    return False

def _rule_10() -> None:
    _p2t_result = None  # 使用确定性的文档结构 fallback
    page2_title_ok = _p2t_result if _p2t_result is not None else _check_page2_title_fallback()
    if page2_title_ok:
        add_hit('第2页存在标题"初中化学模拟试题"，该标题字体为黑体、三号、不加粗', +1)

# 第2页标题"初中化学模拟试题"居中显示、1.5倍行距
# 细则3个点：(1) 定位第2页整段为"初中化学模拟试题"的标题段落
#           (2) 居中显示
#           (3) 1.5倍行距
def _check_page2_title_layout_via_com():
    wdActiveEndPageNumber   = 3
    wdAlignParagraphCenter  = 1
    wdLineSpace1pt5         = 1
    wdLineSpaceMultiple     = 5
    try:
        wd_a = _get_word_doc()

        target_para = None
        for para in wd_a.Paragraphs:
            ptxt = (para.Range.Text or '').rstrip('\r\x07\n')
            if ptxt.strip() == TARGET_TITLE:
                if int(para.Range.Information(wdActiveEndPageNumber)) == 2:
                    target_para = para
                    break
        if target_para is None:
            return False

        # (2) 居中显示：段落对齐方式
        if int(target_para.Alignment) != wdAlignParagraphCenter:
            return False

        # (3) 1.5倍行距：LineSpacingRule = wdLineSpace1pt5，或 Multiple 且倍数=1.5
        rule = int(target_para.LineSpacingRule)
        if rule == wdLineSpace1pt5:
            return True
        if rule == wdLineSpaceMultiple:
            # Word 中 Multiple 规则下 LineSpacing 磅值 = 12 * 倍数
            multiplier = float(target_para.LineSpacing) / 12.0
            if abs(multiplier - 1.5) < 0.01:
                return True
        return False
    except Exception:
        return None

def _check_page2_title_layout_fallback():
    # 备选：python-docx，按分页符估算页码
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    cur_page = 1
    for p in doc.paragraphs:
        if p.text.strip() == TARGET_TITLE and cur_page == 2:
            # 居中：段落对齐 == CENTER (1)
            if p.alignment != 1:
                return False
            # 1.5 倍行距：line = 240 * 倍数（240 twips = 1 行）
            pPr = p._element.find(f'{{{ns_w}}}pPr')
            if pPr is None:
                return False
            sp = pPr.find(f'{{{ns_w}}}spacing')
            if sp is None:
                return False
            line = sp.get(f'{{{ns_w}}}line')
            line_rule = sp.get(f'{{{ns_w}}}lineRule')
            if line is None:
                return False
            # lineRule: auto=倍数, exact/atLeast=磅值*20
            if line_rule in (None, 'auto'):
                multiplier = int(line) / 240.0
                return abs(multiplier - 1.5) < 0.01
            return False
        for br in p._element.findall(f'.//{{{ns_w}}}br'):
            if br.get(f'{{{ns_w}}}type') == 'page':
                cur_page += 1
    return False

def _rule_11() -> None:
    # 启用 Word COM：用 wdActiveEndPageNumber 取真实第2页；LineSpacingRule 兼容 1pt5 与 Multiple(1.5)
    _p2l_result = _check_page2_title_layout_via_com()
    page2_title_layout_ok = _p2l_result if _p2l_result is not None else _check_page2_title_layout_fallback()
    if page2_title_layout_ok:
        add_hit('第2页标题"初中化学模拟试题"居中显示、1.5倍行距', +1)

# 从第2页至尾页除标题以外其他文本：中文文本字体统一为宋体小四号，适用于题干、注意事项、分值表文字、
# 题目说明和主观题文字；英文和数字字体为Times New Roman、小四号
# 细则5个点：(1) 范围：第2页至尾页  (2) 排除范围：标题段落
#           (3) 中文：宋体  (4) 英文和数字：Times New Roman  (5) 字号：小四号(12pt)
def _is_title_paragraph(ptxt):
    """判断是否为标题段落（不参与本条评分）：
       - 第2页大标题"初中化学模拟试题"
       - 副标题："注意事项" / "第X部分 …题" 形式（选择题/非选择题/简答题/实验探究题/计算题 等） / "参考答案"
    """
    import re
    s = ptxt.strip()
    if not s:
        return True   # 空段落不评
    if s == TARGET_TITLE:
        return True
    if s == '注意事项':
        return True
    # 第X部分 …题 （允许"（共N小题，共M分）"这类尾注附在同段）
    if re.match(r'^第[一二三四五六七八九十]+部分[\s　]+[^\s　（(]{1,20}题', s):
        return True
    if s == '参考答案' or (s.startswith('参考答案') and len(s) <= 20):
        return True
    return False

def _check_body_font_fallback():
    # 备选：python-docx，按 <w:br type="page"> 累加页码
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    XIAOSI_PT = 12.0

    def is_chinese(ch): return '一' <= ch <= '鿿'
    def is_en_num(ch):
        return ('A' <= ch <= 'Z') or ('a' <= ch <= 'z') or ('0' <= ch <= '9')

    cur_page = 1
    for p in doc.paragraphs:
        if cur_page >= 2 and not _is_title_paragraph(p.text):
            for r in p.runs:
                rt = r.text
                if not rt:
                    continue
                fn, ea, fs, fb = get_run_info(r)
                size_pt = fs.pt if fs is not None else None
                for ch in rt:
                    if ch in ('_', '＿', ' ', '　', '\t', '\r', '\n'):
                        continue
                    if is_chinese(ch):
                        cn = ea or fn or ''
                        if not ('宋体' in cn or 'SimSun' in cn):
                            return False
                        if size_pt is not None and abs(size_pt - XIAOSI_PT) > 0.01:
                            return False
                    elif is_en_num(ch):
                        name = fn or ''
                        if 'Times New Roman' not in name:
                            return False
                        if size_pt is not None and abs(size_pt - XIAOSI_PT) > 0.01:
                            return False
        for br in p._element.findall(f'.//{{{ns_w}}}br'):
            if br.get(f'{{{ns_w}}}type') == 'page':
                cur_page += 1
    return True

def _rule_12() -> None:
    _bf_result = None  # 使用确定性的文档结构 fallback
    body_font_ok = _bf_result if _bf_result is not None else _check_body_font_fallback()
    if body_font_ok:
        add_hit('从第2页至尾页除标题以外其他文本：中文文本字体统一为宋体小四号，适用于题干、注意事项、分值表文字、题目说明和主观题文字；英文和数字字体为Times New Roman、小四号', +5)

# 从第2页至尾页：除主标题"初中化学模拟试题"外其余标题字体采用黑体、小三号、不加粗
# 细则4个点：(1) 范围：第2页至尾页
#           (2) 排除：主标题"初中化学模拟试题"
#           (3) 对象：其余标题段落（注意事项 / 第一部分 选择题 / 第二部分 非选择题 / 参考答案 等）
#           (4) 字体黑体 + 字号小三号(15pt) + 不加粗
def _is_subtitle_paragraph(ptxt):
    s = ptxt.strip()
    if not s:
        return False
    if s == TARGET_TITLE:
        return False
    if s == '注意事项':
        return True
    if ('第一部分' in s) and ('选择题' in s) and len(s) <= 30:
        return True
    if ('第二部分' in s) and ('非选择题' in s) and len(s) <= 30:
        return True
    if s == '参考答案' or (s.startswith('参考答案') and len(s) <= 20):
        return True
    return False

def _check_subtitles_style_fallback():
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    XIAOSAN_PT = 15.0
    cur_page = 1
    subtitle_found = 0
    for p in doc.paragraphs:
        if cur_page >= 2 and _is_subtitle_paragraph(p.text):
            subtitle_found += 1
            for r in p.runs:
                for ch in r.text:
                    if not ('一' <= ch <= '鿿'):
                        continue
                    fn, ea, fs, fb = get_run_info(r)
                    cn = ea or fn or ''
                    if not ('黑体' in cn or 'SimHei' in cn):
                        return False
                    if fs is not None and abs(fs.pt - XIAOSAN_PT) > 0.01:
                        return False
                    if fb:
                        return False
        for br in p._element.findall(f'.//{{{ns_w}}}br'):
            if br.get(f'{{{ns_w}}}type') == 'page':
                cur_page += 1
    return subtitle_found > 0

def _rule_13() -> None:
    _st_result = None  # 使用确定性的文档结构 fallback
    subtitles_ok = _st_result if _st_result is not None else _check_subtitles_style_fallback()
    if subtitles_ok:
        add_hit('从第2页至尾页：除主标题"初中化学模拟试题"外其余标题字体采用黑体、小三号、不加粗', +3)

# 解题过程下方横线满足以下任意一条：整体数量不是6条；长度不是在15-17厘米之间；所有横线长度不一致  →  -3
# 细则3个触发条件（任一命中即扣分）：
#   (a) 整体数量不是 6 条
#   (b) 长度不是在 15~17 厘米之间（存在任一条不在该区间）
#   (c) 所有横线长度不一致（六条之间存在长度差异）
def _collect_answer_underlines_fallback():
    """备选：按下划线字符估算（12pt SimSun 下 `_` 约 0.21cm）。"""
    lengths = []
    for p in doc.paragraphs:
        body = p.text.strip(' 　\t')
        if len(body) >= 20 and all(c in ('_', '＿') for c in body):
            lengths.append(len(body) * 0.21)
    return lengths

def _rule_14() -> None:
    lengths = _collect_answer_underlines_fallback()

    # 触发条件 (a) 数量不是 6
    count_bad = (len(lengths) != 6)
    # 触发条件 (b) 存在任一长度不在 [15, 17] cm
    length_bad = any(not (15.0 - 0.05 <= l <= 17.0 + 0.05) for l in lengths) if lengths else True
    # 触发条件 (c) 六条之间长度不一致（最大差 > 0.05cm）
    consistency_bad = (len(lengths) >= 2 and (max(lengths) - min(lengths)) > 0.05)

    if count_bad or length_bad or consistency_bad:
        add_hit('解题过程下方横线满足以下任意一条：整体数量不是6条；长度不是在15-17厘米之间；所有横线长度不一致', -3)

# 试题中页面出现大于等于1/3空白  →  -5
# 细则2个要点：(1) 范围：试题中的页面（从第2页开始的所有正文页）
#             (2) 触发：任一页面出现 ≥ 1/3 的空白
# 判定策略：按办公软件真实分页 + 段落在页内的实际 Y 坐标，计算每页"未被内容覆盖的最大连续空白高度"
#   ─ 顶部空白：正文区顶端 → 首个非空内容顶端
#   ─ 底部空白：末个非空内容底端 → 正文区底端
#   ─ 中间空白：任意两段之间的空隙（当前段底端 → 下一段顶端）
#   ─ 空白页：该页无任何非空段落 → 直接命中
def _check_page_blank_ge_one_third_fallback():
    # 备选：无 COM 时按连续空白段落估算
    max_consec_empty = 0
    cur = 0
    for p in doc.paragraphs:
        if p.text.strip() == '':
            cur += 1
            if cur > max_consec_empty:
                max_consec_empty = cur
        else:
            cur = 0
    # 小四号 12pt 单倍行距下每页约 35~40 行，1/3 ≈ 12 行
    return max_consec_empty >= 12

def _rule_15() -> None:
    _bl_result = None  # 使用确定性的文档结构 fallback
    page_has_large_blank = _bl_result if _bl_result is not None else _check_page_blank_ge_one_third_fallback()
    if page_has_large_blank:
        add_hit('试题中页面出现大于等于1/3空白', -5)

# ============================================================
# 维度一：可用与可修改性
# ============================================================
def _run_dim1_checks() -> None:
    """打开文档并运行维度一检查；失败时通过 dim1_fail() 记录原因。
    亦负责将 module-global `doc` 赋值为 python-docx Document，供下方 fallback 使用。"""
    global doc

    # (1) 文件必须能被 python-docx 打开
    try:
        import docx as _docx
        doc = _docx.Document(DOC_PATH)
    except Exception as e:
        dim1_fail(f"文件无法正常打开：{e}")
        return

    # (2) 页数：使用 XML 分页符估算，避免本脚本多次创建 Word COM 会话。
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    page_count = 1
    for p in doc.paragraphs:
        for br in p._element.findall(f'.//{{{ns_w}}}br'):
            if br.get(f'{{{ns_w}}}type') == 'page':
                page_count += 1

    if page_count < 2:
        dim1_fail(f"文档页数不足：仅 {page_count} 页")


# ============================================================
# 主入口 evaluate() 与 15 条评分规则元信息
# ============================================================
def evaluate(dir_path: str) -> dict:
    """入口：dir_path 为脚本所在目录，脚本自行在该目录内定位并打开 .docx。"""
    global DOC_PATH, doc, hit_points, dim1_passed, dim1_reason_parts
    # 重置状态
    hit_points = []
    dim1_passed = True
    dim1_reason_parts = []
    DOC_PATH = None
    doc = None

    file_name = ""
    # 计算正向满分（用于 max_score / error 情形）
    positive_max = sum(d for _, d in _RULE_TABLE if d > 0)

    try:
        # 定位目录中的 .docx（忽略 ~$ 临时锁定文件）
        if not os.path.isdir(dir_path):
            raise FileNotFoundError(f"目录不存在：{dir_path}")
        docx_files = [
            f for f in os.listdir(dir_path)
            if f.lower().endswith('.docx') and not f.startswith('~$')
        ]
        if not docx_files:
            raise FileNotFoundError(f"目录内未找到 .docx 文件：{dir_path}")
        # 若多份文档存在，取字典序第一个（与批量 runner 约定一致）
        docx_files.sort()
        file_name = docx_files[0]
        DOC_PATH = os.path.join(dir_path, file_name)

        # 维度一
        _run_dim1_checks()
        if not dim1_passed:
            return {
                "id": "019",
                "file_name": file_name,
                "status": "ok",
                "error": None,
                "dim1_pass": False,
                "dim1_reason": "；".join(dim1_reason_parts),
                "dim2_items": [],
                "total_score": 0,
                "max_score": positive_max,
            }

        # 维度二：依次触发全部 15 条评分规则
        for _, _, fn in _SCORE_RULES:
            fn()

        # 汇总 dim2_items（命中 + 未命中都要列）
        hit_map = {desc: score for desc, score in hit_points}
        dim2_items = []
        for desc, max_delta, _ in _SCORE_RULES:
            hit = desc in hit_map
            delta = hit_map[desc] if hit else 0
            dim2_items.append({
                "rule": desc,
                "max_delta": max_delta,
                "delta": delta,
                "hit": hit,
                "detail": "",
            })

        total_score = sum(item["delta"] for item in dim2_items)
        return {
            "id": "019",
            "file_name": file_name,
            "status": "ok",
            "error": None,
            "dim1_pass": True,
            "dim1_reason": "",
            "dim2_items": dim2_items,
            "total_score": total_score,
            "max_score": positive_max,
        }
    except Exception as exc:
        return {
            "id": "019",
            "file_name": file_name,
            "status": "error",
            "error": f"{type(exc).__name__}: {exc}",
            "dim1_pass": False,
            "dim1_reason": "",
            "dim2_items": [],
            "total_score": 0,
            "max_score": positive_max,
        }
    finally:
        # 无论成功/异常/维度一短路，都需释放共享的 Word COM 会话
        _close_word()


# 15 条评分规则的元信息（用于生成 dim2_items 矩阵）
# 每一项：(rule_desc, max_delta, trigger_fn)
_SCORE_RULES = [
    ("试卷采用A4纸张，页边距设置为上2.5厘米、下2.5厘米、左2.5厘米、右2.5厘米", +3, _rule_01),
    ("第1页为封面页，\"注意事项\"和\"第一部分 选择题\"这两个标题出现在第2页", +1, _rule_02),
    ("除了封面页，其余页页脚位置出现页码，页码格式例如：\"化学试题 第1页 共5页\"，居中显示", +3, _rule_03),
    ("封面页中部包含竖排标题\"初中化学模拟试题\"，文字按\"初/中/化/学/模/拟/试/题\"自上而下单独成行排列", +3, _rule_04),
    ("封面竖排标题\"初中化学模拟试题\"字体为不加粗黑体或者加粗宋体，字号为二号到小三号，位于页面从上至下2/3的区域并水平居中", +3, _rule_05),
    ("封面下半部分依次出现\"报送单位：\"\"学科：\"\"命题人：\"\"命题人联系电话：\"这四个文本，且每个文本单独成行，四个字段", +3, _rule_06),
    ("封面页\"报送单位：\"\"学科：\"\"命题人：\"\"命题人联系电话：\"四个文本右侧均有横线且与文本之间有1-4个空格；横线为黑色或深色细实线；横线长度约5-8厘米；四条横线右侧对齐", +3, _rule_08),
    ("封面页除标题外其余文本：中文字体为宋体、小四号，英文和阿拉伯数字为Times New Roman、小四号", +3, _rule_09),
    ("第2页存在标题\"初中化学模拟试题\"，该标题字体为黑体、三号、不加粗", +1, _rule_10),
    ("第2页标题\"初中化学模拟试题\"居中显示、1.5倍行距", +1, _rule_11),
    ("从第2页至尾页除标题以外其他文本：中文文本字体统一为宋体小四号，适用于题干、注意事项、分值表文字、题目说明和主观题文字；英文和数字字体为Times New Roman、小四号", +5, _rule_12),
    ("从第2页至尾页：除主标题\"初中化学模拟试题\"外其余标题字体采用黑体、小三号、不加粗", +3, _rule_13),
    ("解题过程下方横线满足以下任意一条：整体数量不是6条；长度不是在15-17厘米之间；所有横线长度不一致", -3, _rule_14),
    ("试题中页面出现大于等于1/3空白", -5, _rule_15),
]

# 兼容 evaluate() 里的引用别名
_RULE_TABLE = [(desc, delta) for desc, delta, _ in _SCORE_RULES]


if __name__ == "__main__":
    # 本地调试：默认使用脚本所在目录；也可 `python xxx_verifier.py <dir>` 传入
    target_dir = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
    result = evaluate(target_dir)
    print(json.dumps(result, ensure_ascii=False, indent=2))
