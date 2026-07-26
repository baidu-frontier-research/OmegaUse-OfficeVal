#!/usr/bin/env python3
"""
自动评估脚本：评估 Word 文档是否符合评分细则。
维度1：可用与可修改性（不满足直接0分）
维度2：完成度评分细则（逐项检查得分/扣分点）
"""

import os

SCRIPT_ID = "009"
import re
import sys
import json
from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor, Twips
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def evaluate_document(filepath):
    """评估文档，返回得分和命中的规则列表。"""
    results = []  # [(rule_description, score, passed: bool)]

    # ============================
    # 维度1：可用与可修改性
    # ============================
    dim1_pass = True

    # 1.1 文件为.docx格式，文件可正常打开
    # python-docx 只能解析 OOXML（.docx）；传统二进制 .doc 不再支持。
    ext = os.path.splitext(filepath)[1].lower()
    if ext != '.docx':
        results.append(("维度1：文件为.docx格式", 0, False))
        dim1_pass = False
    else:
        try:
            doc = Document(filepath)
            results.append(("维度1：文件为.docx格式，可正常打开", 0, True))
        except Exception:
            results.append(("维度1：文件可正常打开", 0, False))
            dim1_pass = False

    if not dim1_pass:
        return results, 0

    # 1.2 文件没有出现大面积无法编辑的情况
    # 检查是否有内容保护或文档保护
    body = doc.element.body
    protection = doc.element.find(qn('w:body'))
    doc_protection = doc.settings.element.find(qn('w:documentProtection'))
    if doc_protection is not None:
        edit_val = doc_protection.get(qn('w:edit'))
        if edit_val in ('readOnly', 'comments', 'trackedChanges'):
            results.append(("维度1：文件没有大面积无法编辑", 0, False))
            dim1_pass = False
        else:
            results.append(("维度1：文件没有大面积无法编辑", 0, True))
    else:
        # 检查是否存在大量锁定的内容控件
        sdt_count = len(list(body.iter(qn('w:sdt'))))
        total_paras = len(doc.paragraphs)
        if sdt_count > total_paras * 0.5 and sdt_count > 10:
            results.append(("维度1：文件没有大面积无法编辑", 0, False))
            dim1_pass = False
        else:
            results.append(("维度1：文件没有大面积无法编辑", 0, True))

    if not dim1_pass:
        return results, 0

    # 1.3 不存在文字重叠、表格错位、连续空白页的问题
    # 检查连续空白段落（模拟连续空白页）
    consecutive_empty = 0
    max_consecutive_empty = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            consecutive_empty += 1
            max_consecutive_empty = max(max_consecutive_empty, consecutive_empty)
        else:
            consecutive_empty = 0

    # 检查是否有连续的分页符（连续空白页）
    page_break_count = 0
    consecutive_breaks = 0
    max_consecutive_breaks = 0
    for para in doc.paragraphs:
        has_break = False
        for run in para.runs:
            for br in run._r.iter(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_break = True
        if has_break and not para.text.strip():
            consecutive_breaks += 1
            max_consecutive_breaks = max(max_consecutive_breaks, consecutive_breaks)
        else:
            consecutive_breaks = 0

    # 连续超过3个空段落视为潜在连续空白页问题（宽松判定）
    if max_consecutive_breaks >= 2 or max_consecutive_empty > 15:
        results.append(("维度1：不存在文字重叠、表格错位、连续空白页", 0, False))
        dim1_pass = False
    else:
        results.append(("维度1：不存在文字重叠、表格错位、连续空白页", 0, True))

    if not dim1_pass:
        return results, 0

    # ============================
    # 维度2：完成度评分细则
    # ============================
    total_score = 0

    # +3：跨页表格在表格底部页面右下侧有"续下页"文本
    score, passed = check_xuyexia_exists(doc)
    results.append(('+3：跨页表格在表格底部页面右下侧有"续下页"文本', score, passed))
    total_score += score

    # +5：文中所有"续下页"：位于页面右下侧、表格底部附近，靠近表格右边界和下边线
    score, passed = check_xuyexia_position(doc)
    results.append(('+5：所有"续下页"位于页面右下侧、表格底部附近', score, passed))
    total_score += score

    # +1："续下页"字体颜色为黑色，无其他样式
    score, passed = check_xuyexia_style(doc)
    results.append(('+1："续下页"字体颜色为黑色，无其他样式', score, passed))
    total_score += score

    # +5：每一页的表格：都有三条平行的边线，上边线和下边线1.5磅，内边线0.75磅
    score, passed = check_table_borders(doc)
    results.append(('+5：表格有三条平行边线，上下1.5磅，内0.75磅', score, passed))
    total_score += score

    # +5：每一页最上方的表格：都要有表头且表头中文字体字号为8pt宋体加粗
    # （该项已按需求移除，不再计入评分）

    # +5："续下页"上方的表格表头与下方的表格表头：内容、中文和英文数字的字体字号格式均一致
    score, passed = check_header_consistency(doc)
    results.append(('+5："续下页"上方的表格表头与下方的表格表头：内容、中文和英文数字的字体字号格式均一致', score, passed))
    total_score += score

    # -5：交付文件出现图片
    score, passed = check_no_images(doc)
    results.append(('-5：交付文件出现图片', score, passed))
    total_score += score

    # -3：表格上方加粗的表题：其中数字部分字体为宋体
    score, passed = check_caption_number_font(doc)
    results.append(('-3：表题中数字部分字体为宋体', score, passed))
    total_score += score

    # -1：表5不是9行5列的表格
    score, passed = check_table5_dimensions(doc)
    results.append(('-1：表5不是9行5列的表格', score, passed))
    total_score += score

    # -1："3.3 分层比较与模型校准"格式检查
    score, passed = check_heading_33_format(doc)
    results.append(('-1："3.3 分层比较与模型校准"格式不符', score, passed))
    total_score += score

    # -1：表1下方空一行的位置没有出现注释
    score, passed = check_table1_note(doc)
    results.append(('-1：表1下方空一行位置没有注释', score, passed))
    total_score += score

    # -3：表2整体少于40行
    score, passed = check_table2_rows(doc)
    results.append(('-3：表2整体少于40行', score, passed))
    total_score += score

    # -3：至少有一个表格出现大于3条边框
    score, passed = check_table_border_count(doc)
    results.append(('-3：至少有一个表格出现大于3条边框', score, passed))
    total_score += score

    return results, total_score


def identify_cross_page_tables(doc):
    """识别跨页表格。

    办公软件（Word/WPS）里实现跨页表格有两种等价结构，均需识别：
      结构一（单表 + 重复表头）：一张 <w:tbl> 首行带 <w:trPr><w:tblHeader/>，
        软件跨页渲染时自动在下一页顶部重复该表头行；这种情况在 XML 里始终只有
        一个物理 <w:tbl>，不会拆成两个相邻 <w:tbl>。
      结构二（逻辑分表）：作者手动把一张跨页表拆成两个相邻的 <w:tbl>（中间无
        非空段落），且两表首行内容一致，办公软件里视觉上仍是同一张表跨页。

    返回 [(idx_a, idx_b, body_pos_a, body_pos_b), ...]：
    - idx_a/idx_b：在 doc.tables 中的索引；结构一时 idx_a == idx_b（同一张表）。
    - body_pos_a/body_pos_b：在 body 直接子节点中的位置（稳定标识）；结构一时
      body_pos_a == body_pos_b。
    """
    body = doc.element.body
    cross_page_pairs = []

    # 收集body直接子节点中的所有tbl：(body_position, table_idx)
    table_records = []
    table_idx = -1
    for body_pos, child in enumerate(body):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            table_idx += 1
            table_records.append((body_pos, table_idx))

    # 结构一：单表首行 <w:tblHeader/> → 表本身就是跨页表格
    for body_pos, idx in table_records:
        tbl_elem = body[body_pos]
        first_tr = tbl_elem.find(qn('w:tr'))
        if first_tr is None:
            continue
        trPr = first_tr.find(qn('w:trPr'))
        if trPr is not None and trPr.find(qn('w:tblHeader')) is not None:
            cross_page_pairs.append((idx, idx, body_pos, body_pos))

    # 结构二：相邻两个物理表格，中间无非空段落，且首行内容一致
    for i in range(len(table_records) - 1):
        body_pos_a, idx_a = table_records[i]
        body_pos_b, idx_b = table_records[i + 1]

        # 检查两个表格之间是否仅有空段落
        between_has_content = False
        for j in range(body_pos_a + 1, body_pos_b):
            sibling = body[j]
            stag = sibling.tag.split('}')[1] if '}' in sibling.tag else sibling.tag
            if stag == 'p':
                text = ''
                for t in sibling.iter(qn('w:t')):
                    if t.text:
                        text += t.text
                if text.strip():
                    between_has_content = True
                    break
        if between_has_content:
            continue

        # 表头一致性
        table_a = doc.tables[idx_a]
        table_b = doc.tables[idx_b]
        if len(table_a.rows) == 0 or len(table_b.rows) == 0:
            continue
        header_a = [cell.text.strip() for cell in table_a.rows[0].cells]
        header_b = [cell.text.strip() for cell in table_b.rows[0].cells]
        if header_a == header_b and any(h for h in header_a):
            cross_page_pairs.append((idx_a, idx_b, body_pos_a, body_pos_b))

    return cross_page_pairs



def get_xuyexia_occurrences(doc):
    """获取所有"续下页"出现的位置（含 anchor 信息）。

    返回 [(anchor_body_pos, anchor_table_body_pos, position_info), ...]
    - anchor_body_pos：包含该"续下页"图形的 body 直接子节点位置。
    - anchor_table_body_pos：若该 anchor 在某个表格的单元格内（body直接子tbl），
      则为该 tbl 的 body 位置；否则为 None（如游离段落）。
    - position_info：浮动文本框的 (left_pt, top_pt)，若无则 None。

    注：续下页 通常以浮动 VML 文本框形式存在，其 anchor 段落决定它实际渲染到
    哪一页（mso-position-vertical-relative=page 时落在 anchor 段落所在页面）。
    """
    body = doc.element.body
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body_children = list(body)

    occurrences = []
    seen_picts = set()

    for t in body.iter(f'{{{ns_w}}}t'):
        if not (t.text and '续下页' in t.text):
            continue

        # 找到包含此文本的 <w:pict>（去重，每个图形只算一次）
        anc = t.getparent()
        pict_elem = None
        while anc is not None:
            atag = anc.tag.split('}')[1] if '}' in anc.tag else anc.tag
            if atag == 'pict':
                pict_elem = anc
                break
            anc = anc.getparent()
        if pict_elem is None or id(pict_elem) in seen_picts:
            continue
        seen_picts.add(id(pict_elem))

        # 找到 anchor 段落（pict 所在的 w:p）
        anchor_p = pict_elem.getparent()
        while anchor_p is not None:
            atag = anchor_p.tag.split('}')[1] if '}' in anchor_p.tag else anchor_p.tag
            if atag == 'p':
                break
            anchor_p = anchor_p.getparent()
        if anchor_p is None:
            continue

        # 求 anchor 段落所对应的 body 直接子节点位置
        anchor_body_pos = None
        anc2 = anchor_p
        while anc2 is not None:
            parent = anc2.getparent()
            if parent is None:
                break
            if parent.tag == body.tag:
                for i, c in enumerate(body_children):
                    if c is anc2:
                        anchor_body_pos = i
                        break
                break
            anc2 = parent

        # 求 anchor 段落所在表格（body 直接子 tbl）的 body 位置
        anchor_table_body_pos = None
        anc3 = anchor_p.getparent()
        while anc3 is not None:
            parent = anc3.getparent()
            if parent is None:
                break
            atag = anc3.tag.split('}')[1] if '}' in anc3.tag else anc3.tag
            if atag == 'tbl' and parent.tag == body.tag:
                for i, c in enumerate(body_children):
                    if c is anc3:
                        anchor_table_body_pos = i
                        break
                break
            anc3 = parent

        # 文本框物理位置（VML shape）
        position = None
        shape_elem = pict_elem.find('{urn:schemas-microsoft-com:vml}shape')
        if shape_elem is not None:
            style = shape_elem.get('style', '') or ''
            lm = re.search(r'left:\s*([\d.]+)pt', style)
            tm = re.search(r'top:\s*([\d.]+)pt', style)
            if lm and tm:
                position = (float(lm.group(1)), float(tm.group(1)))

        occurrences.append((anchor_body_pos, anchor_table_body_pos, position))

    return occurrences


def check_xuyexia_exists(doc):
    """+3：跨页表格在表格底部页面右下侧有"续下页"文本。

    细则的每一个点：
      (1) 跨页表格：必须是 identify_cross_page_tables 判定出的真实跨页表格
          （结构一：单表首行 <w:tblHeader/>；结构二：相邻两表且首行内容一致），
          不是"任意锚定了续下页浮框的表格" —— 后者只能证明浮框位置，不能
          证明该表格真的跨页；
      (2) "续下页"文本：跨页表格中出现"续下页"文字；
      (3) 表格底部：位于跨页表格锚定行的下方（视觉上是该页表格底部）；
      (4) 页面右下侧：文本框水平中心位于页面水平中线右侧。

    办公软件（Word/WPS）中的表现（与本文档一致的主流做法）：
      - 结构一：一张 <w:tbl> + <w:trPr><w:tblHeader/> —— 软件在跨页时自动
        重复表头；每次跨页边界前的行末以 <w:drawing>（含 <wp:anchor>）
        或 <w:pict>（含 <v:shape>）锚定的浮动文本框标注"续下页"。
      - 结构二：两个连续 <w:tbl> 之间由作者手动分割 —— 上表末尾同样通过
        锚定的浮动文本框放"续下页"。
      两种结构在办公软件里最终效果一致：读者看到"续下页"贴在当前页面
      表格右下角，翻页后表格继续。

    判定步骤（不引入细则未要求的额外约束）：
      1) 用 identify_cross_page_tables 得到真实跨页表格集合（结构一收敛为
         单个宿主表；结构二取"上表"，即续下页应锚定的那一半）；
      2) 若该集合为空 → 文档中无跨页表格，不适用，返回 0；
      3) 收集所有"续下页"锚定的浮动文本框（跳过 <mc:Fallback>，办公软件
         走 <mc:Choice>），只保留锚定宿主属于步骤1集合的文本框；
      4) 逐张跨页表格判定：该表至少一个"续下页"同时满足
         (3) 位于锚定段落下方（posOffset_V > 0）与
         (4) 文本框水平中心 ≥ 页面中线。
         任一跨页表格没有满足条件的"续下页"（包括完全没有锚定其上的
         "续下页"）→ 未命中，返回 0；全部满足 → 得 3 分。

    位置解析（办公软件几何等价）：
      - DrawingML：<wp:positionH relativeFrom="page|column|margin"> +
        <wp:posOffset>（EMU，12700 EMU=1pt）+ <wp:extent cx="…"/>。
        本文档 anchor 位于单元格内（layoutInCell="1"），relativeFrom="column"
        的基准是该单元格左内边——用 compute_table_horizontal_extent 得到
        宿主表的页面 X 起点，再按 <w:tblGrid>/<w:gridCol> 累加锚定单元格
        之前列的宽度，得到单元格左内边的页面 X。
      - VML：<v:shape style="margin-left;width;mso-position-horizontal-relative">，
        margin-left 相对基准同 DrawingML 的 relativeFrom 语义。
    """
    body = doc.element.body
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    ns_mc = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    ns_v = 'urn:schemas-microsoft-com:vml'

    # (1) 先识别真实跨页表格；结构一 idx_a==idx_b 收敛为单个宿主表，
    # 结构二取"上表"（idx_a）——续下页应锚定在跨页边界前的上半表。
    cross_pairs = identify_cross_page_tables(doc)
    if not cross_pairs:
        return 0, False  # 文档中没有真实跨页表格 → 不适用

    section = doc.sections[0]
    page_w_pt = section.page_width.pt
    left_margin_pt = section.left_margin.pt
    page_mid_x = page_w_pt / 2.0
    EMU_PER_PT = 12700.0

    def _parse_pt(v):
        if v is None:
            return None
        m = re.match(r'([-\d.]+)pt', v.strip())
        return float(m.group(1)) if m else None

    def _in_fallback(elem):
        p = elem.getparent()
        while p is not None:
            if p.tag == f'{{{ns_mc}}}Fallback':
                return True
            p = p.getparent()
        return False

    # 建 body 直接子 tbl id → Table 对象 映射
    tbl_id_to_table = {}
    ti = -1
    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            ti += 1
            tbl_id_to_table[id(child)] = doc.tables[ti]

    # 真实跨页表格的宿主 body 直接子 <w:tbl> 元素集合：
    #   结构一（idx_a==idx_b）：该表自身；
    #   结构二（idx_a!=idx_b）：续下页应锚定在跨页边界前的"上表"，即 body_pos_a。
    body_children = list(body)
    cross_page_host_ids = set()
    for idx_a, idx_b, body_pos_a, body_pos_b in cross_pairs:
        cross_page_host_ids.add(id(body_children[body_pos_a]))

    def _host_body_tbl(elem):
        """返回锚定元素所在的 body 直接子 <w:tbl> 元素（若无则 None）。
        对嵌套表以最外层为准 —— 但办公软件按最内层 cell 布局，两种情形
        本处只用于识别宿主，最外层已足够。"""
        e = elem.getparent()
        host = None
        while e is not None:
            parent = e.getparent()
            tag = e.tag.split('}')[1] if '}' in e.tag else e.tag
            if tag == 'tbl' and parent is not None and parent.tag == body.tag:
                host = e
                break
            e = parent
        return host

    def _anchor_cell_index(elem):
        """返回锚定元素所在 <w:tc> 在其行内的列索引（0-based）；若不在
        单元格里返回 None。"""
        e = elem.getparent()
        tc = None
        while e is not None:
            tag = e.tag.split('}')[1] if '}' in e.tag else e.tag
            if tag == 'tc':
                tc = e
                break
            e = e.getparent()
        if tc is None:
            return None
        tr = tc.getparent()
        tcs = list(tr.findall(qn('w:tc')))
        try:
            return tcs.index(tc)
        except ValueError:
            return None

    def _cell_left_page_x(host_tbl_elem, cell_idx):
        """得到宿主表格中第 cell_idx 列左内边的页面 X（pt）。"""
        table_obj = tbl_id_to_table.get(id(host_tbl_elem))
        if table_obj is None:
            return None
        try:
            tbl_left, _tbl_right = compute_table_horizontal_extent(doc, table_obj)
        except Exception:
            return None
        tblGrid = host_tbl_elem.find(qn('w:tblGrid'))
        if tblGrid is None:
            return tbl_left
        widths_pt = []
        for gc in tblGrid.findall(qn('w:gridCol')):
            try:
                widths_pt.append(int(gc.get(qn('w:w'))) / 20.0)
            except (TypeError, ValueError):
                widths_pt.append(0.0)
        if cell_idx is None or cell_idx < 0 or cell_idx >= len(widths_pt):
            return tbl_left
        return tbl_left + sum(widths_pt[:cell_idx])

    def _drawing_center_x_and_below(drawing_elem, host_tbl_elem, cell_idx):
        anchor = drawing_elem.find(f'{{{ns_wp}}}anchor')
        if anchor is None:
            return None, False
        posH = anchor.find(f'{{{ns_wp}}}positionH')
        posV = anchor.find(f'{{{ns_wp}}}positionV')
        extent = anchor.find(f'{{{ns_wp}}}extent')
        if posH is None or posV is None or extent is None:
            return None, False

        relH = posH.get('relativeFrom') or 'column'
        offH_elem = posH.find(f'{{{ns_wp}}}posOffset')
        try:
            offH_pt = float(offH_elem.text.strip()) / EMU_PER_PT
        except (AttributeError, TypeError, ValueError):
            return None, False
        try:
            cx_pt = float(extent.get('cx')) / EMU_PER_PT
        except (TypeError, ValueError):
            return None, False

        if relH == 'page':
            base_x = 0.0
        elif relH == 'margin':
            base_x = left_margin_pt
        else:  # column / character / insideMargin / outsideMargin / 缺省
            # layoutInCell=1 时 column 基准 = 锚定单元格的左内边
            cell_left = _cell_left_page_x(host_tbl_elem, cell_idx)
            base_x = cell_left if cell_left is not None else left_margin_pt
        center_x = base_x + offH_pt + cx_pt / 2.0

        offV_elem = posV.find(f'{{{ns_wp}}}posOffset')
        try:
            offV = float(offV_elem.text.strip())
        except (AttributeError, TypeError, ValueError):
            offV = None
        is_below = (offV is not None and offV > 0)
        return center_x, is_below

    def _pict_center_x_and_below(pict_elem, host_tbl_elem, cell_idx):
        shape = pict_elem.find(f'{{{ns_v}}}shape')
        if shape is None:
            return None, False
        style = shape.get('style', '') or ''
        attrs = {}
        for kv in style.split(';'):
            if ':' in kv:
                k, v = kv.split(':', 1)
                attrs[k.strip().lower()] = v.strip()
        # 严格只认 VML 官方坐标属性：margin-left / margin-top。
        # CSS 的 left/top 在 Word/WPS 里对浮动图形不生效——即便 XML 里存在，
        # 办公软件也不会渲染到期望位置。为避免"XML 有但用户看不到"的假加分，
        # 这里不再回退到 left/top。
        ml = _parse_pt(attrs.get('margin-left'))
        mt = _parse_pt(attrs.get('margin-top'))
        w = _parse_pt(attrs.get('width')) or 0.0
        if ml is None:
            return None, False
        h_rel = (attrs.get('mso-position-horizontal-relative') or 'column').lower()
        if h_rel == 'page':
            base_x = 0.0
        elif h_rel == 'margin':
            base_x = left_margin_pt
        else:
            cell_left = _cell_left_page_x(host_tbl_elem, cell_idx)
            base_x = cell_left if cell_left is not None else left_margin_pt
        center_x = base_x + ml + w / 2.0
        is_below = (mt is not None and mt > 0)
        return center_x, is_below

    # 收集"续下页"锚定的浮动文本框 → 宿主表 → 是否已找到满足条件的标记。
    # 只对真实跨页表格集合初始化，未锚定任何"续下页"的跨页表格保持 False，
    # 从而在最终判定里被正确视为"未命中"。
    host_status = {hid: False for hid in cross_page_host_ids}

    def _register(elem, is_drawing):
        host = _host_body_tbl(elem)
        if host is None:
            return  # 不在表格里，非"跨页表格里的续下页"，忽略
        host_id = id(host)
        if host_id not in cross_page_host_ids:
            return  # 锚定的宿主表不是真实跨页表格，不计入判定
        if host_status.get(host_id):
            return  # 该宿主表已有满足条件的标记
        cell_idx = _anchor_cell_index(elem)
        if is_drawing:
            cx, is_below = _drawing_center_x_and_below(elem, host, cell_idx)
        else:
            cx, is_below = _pict_center_x_and_below(elem, host, cell_idx)
        if cx is None:
            return
        if is_below and cx >= page_mid_x:
            host_status[host_id] = True

    # DrawingML（AlternateContent/Choice 主渲染）
    for drawing in body.iter(f'{{{ns_w}}}drawing'):
        if _in_fallback(drawing):
            continue
        has_xy = False
        for t in drawing.iter(f'{{{ns_w}}}t'):
            if t.text and '续下页' in t.text:
                has_xy = True
                break
        if not has_xy:
            continue
        _register(drawing, is_drawing=True)

    # VML（旧版兼容渲染）
    for pict in body.iter(f'{{{ns_w}}}pict'):
        if _in_fallback(pict):
            continue
        has_xy = False
        for t in pict.iter(f'{{{ns_w}}}t'):
            if t.text and '续下页' in t.text:
                has_xy = True
                break
        if not has_xy:
            continue
        _register(pict, is_drawing=False)

    if not host_status:
        return 0, False
    for _hid, ok in host_status.items():
        if not ok:
            return 0, False
    return 3, True


def get_shapes_with_xuyexia(doc):
    """获取包含"续下页"的浮动文本框容器（供其他检查项复用）。

    办公软件（Word/WPS）里文本框有两种等价渲染形式，必须都覆盖，否则只用
    DrawingML 的合规文档会被漏检，只用 VML 的也一样：
      - DrawingML：<w:drawing>...<wps:txbx>/<w:txbxContent>...</w:drawing>
      - VML（旧版兼容）：<w:pict><v:shape>...<v:textbox><w:txbxContent>...

    跳过 <mc:Fallback> 内的重复内容（<mc:AlternateContent> 现代软件走
    <mc:Choice>，<mc:Fallback> 只是老版本阅读器的备用渲染，与 <mc:Choice>
    内容重复，不应重复计入）。

    返回的每个元素都是"容器"（<w:drawing> 或 <w:pict>），而不是 VML 特有的
    <v:shape>，以便调用方统一通过 .//w:txbxContent 取文本框内容。
    """
    body = doc.element.body
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_mc = 'http://schemas.openxmlformats.org/markup-compatibility/2006'

    def _in_fallback(elem):
        p = elem.getparent()
        while p is not None:
            if p.tag == f'{{{ns_mc}}}Fallback':
                return True
            p = p.getparent()
        return False

    containers = []

    for drawing in body.iter(f'{{{ns_w}}}drawing'):
        if _in_fallback(drawing):
            continue
        has_xy = any(t.text and '续下页' in t.text
                     for t in drawing.iter(f'{{{ns_w}}}t'))
        if has_xy:
            containers.append(drawing)

    for pict in body.iter(f'{{{ns_w}}}pict'):
        if _in_fallback(pict):
            continue
        has_xy = any(t.text and '续下页' in t.text
                     for t in pict.iter(f'{{{ns_w}}}t'))
        if has_xy:
            containers.append(pict)

    return containers


def compute_table_horizontal_extent(doc, table):
    """计算表格在页面上的水平范围 (table_left_pt, table_right_pt)。

    依据：
    - 表格列宽来自 <w:tblGrid>/<w:gridCol w:w="..."/>（单位 twips）。
    - 总宽 = sum(gridCol)，若 <w:tblW w:type="dxa"> 显式给出则优先用它。
    - 表格水平定位由 <w:tblPr>/<w:jc>（center/left/right/start/end）和 <w:tblInd> 决定，
      在页面正文区（[left_margin, page_width - right_margin]）内放置。
    - 浮动表（含 <w:tblpPr>）一般以 tblpX/tblpY 直接定位；这里若存在则使用 tblpX/tblpY+宽度。

    若信息不足，回退为内容区全宽。
    """
    section = doc.sections[0]
    page_w = section.page_width.pt
    left_m = section.left_margin.pt
    right_m = section.right_margin.pt
    content_left = left_m
    content_right = page_w - right_m

    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))

    # 总宽（twips → pt：除以 20）
    width_pt = None
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            w_type = tblW.get(qn('w:type'))
            w_val = tblW.get(qn('w:w'))
            if w_type == 'dxa' and w_val:
                try:
                    width_pt = int(w_val) / 20.0
                except ValueError:
                    pass

    if width_pt is None:
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            total = 0
            for gc in tblGrid.findall(qn('w:gridCol')):
                try:
                    total += int(gc.get(qn('w:w')))
                except (TypeError, ValueError):
                    pass
            if total > 0:
                width_pt = total / 20.0

    if width_pt is None:
        return (content_left, content_right)

    # 浮动表（绝对定位）
    if tblPr is not None:
        tblpPr = tblPr.find(qn('w:tblpPr'))
        if tblpPr is not None:
            tblpX = tblpPr.get(qn('w:tblpX'))
            if tblpX:
                try:
                    left_abs = int(tblpX) / 20.0
                    return (left_abs, left_abs + width_pt)
                except ValueError:
                    pass

    # 水平对齐
    jc_val = None
    if tblPr is not None:
        jc = tblPr.find(qn('w:jc'))
        if jc is not None:
            jc_val = jc.get(qn('w:val'))

    # 表格缩进
    ind_pt = 0.0
    if tblPr is not None:
        tblInd = tblPr.find(qn('w:tblInd'))
        if tblInd is not None and tblInd.get(qn('w:type')) == 'dxa':
            try:
                ind_pt = int(tblInd.get(qn('w:w'))) / 20.0
            except (TypeError, ValueError):
                pass

    if jc_val == 'center':
        # 居中在内容区
        avail = (content_right - content_left)
        offset = (avail - width_pt) / 2.0
        tbl_left = content_left + max(offset, 0)
    elif jc_val in ('right', 'end'):
        tbl_left = content_right - width_pt
    else:  # left / start / None
        tbl_left = content_left + ind_pt

    tbl_right = tbl_left + width_pt
    return (tbl_left, tbl_right)


def check_xuyexia_position(doc):
    """+5：文中所有"续下页"：位于页面右下侧、表格底部附近，靠近表格右边界和下边线。

    细则的每一个点：
      (1) 文中所有"续下页"：遍历文档全部"续下页"浮动文本框，任一不合规 → 0 分；
      (2) 位于页面右下侧：文本框水平中心位于页面水平中线右侧
          （在办公软件里，续下页的宿主行由于跨页触发本身位于该页底部，
          垂直方向只要"位于锚定段落下方"即满足"下"）；
      (3) 表格底部附近：文本框所在锚定段落 posOffset_V > 0（浮于锚定行下方），
          且锚定行不是宿主表格的首行（首行是表头，不会在其下贴"续下页"）；
      (4) 靠近表格右边界：文本框右沿与宿主 <w:tbl> 右边界的水平距离 ≤ TOL_X；
      (5) 靠近表格下边线：文本框底沿垂直越过锚定段落（posOffset_V > 0 已满足；
          在办公软件里锚定行即该页表格的下边线所在处）。

    办公软件（Word/WPS）中的表现：
      - 一张 <w:tbl> 用 <w:trPr><w:tblHeader/> 实现跨页时，跨页边界处的行末
        单元格内锚定一个浮动文本框（DrawingML 主渲染 / VML 兼容渲染），
        posOffset_H 让文本框水平贴在单元格右侧、越出宿主单元格右边界并停在
        表格右边界附近；posOffset_V 让文本框浮到锚定行下方，恰好压在
        跨页处表格下边线之外一点，读者视觉上就是"表格右下角外侧"。
      - 判定不引入细则未要求的额外约束，例外只 skip <mc:Fallback>
        （办公软件不渲染此分支）。
    """
    body = doc.element.body
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    ns_mc = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    ns_v = 'urn:schemas-microsoft-com:vml'

    section = doc.sections[0]
    page_w_pt = section.page_width.pt
    left_margin_pt = section.left_margin.pt
    page_mid_x = page_w_pt / 2.0
    EMU_PER_PT = 12700.0

    # 容差（pt）：靠近表格右边界的最大水平偏离
    TOL_X = 40.0

    def _parse_pt(v):
        if v is None:
            return None
        m = re.match(r'([-\d.]+)pt', v.strip())
        return float(m.group(1)) if m else None

    def _in_fallback(elem):
        p = elem.getparent()
        while p is not None:
            if p.tag == f'{{{ns_mc}}}Fallback':
                return True
            p = p.getparent()
        return False

    # body 直接子 tbl id → Table 对象
    tbl_id_to_table = {}
    ti = -1
    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            ti += 1
            tbl_id_to_table[id(child)] = doc.tables[ti]

    def _host_body_tbl(elem):
        e = elem.getparent()
        while e is not None:
            parent = e.getparent()
            tag = e.tag.split('}')[1] if '}' in e.tag else e.tag
            if tag == 'tbl' and parent is not None and parent.tag == body.tag:
                return e
            e = parent
        return None

    def _anchor_row_and_cell_index(elem):
        """返回 (row_index_in_host_tbl, cell_index_in_row)。"""
        e = elem.getparent()
        tc = None
        while e is not None:
            tag = e.tag.split('}')[1] if '}' in e.tag else e.tag
            if tag == 'tc':
                tc = e
                break
            e = e.getparent()
        if tc is None:
            return None, None
        tr = tc.getparent()
        host = tr.getparent()
        # host 可能是嵌套；找到 body 直接子 tbl 之下 tr 的索引
        # 但对本细则来说，只要不是首行即可
        rows = list(host.findall(qn('w:tr')))
        try:
            r_idx = rows.index(tr)
        except ValueError:
            r_idx = None
        tcs = list(tr.findall(qn('w:tc')))
        try:
            c_idx = tcs.index(tc)
        except ValueError:
            c_idx = None
        return r_idx, c_idx

    def _cell_left_page_x(host_tbl_elem, cell_idx):
        table_obj = tbl_id_to_table.get(id(host_tbl_elem))
        if table_obj is None:
            return None
        try:
            tbl_left, _ = compute_table_horizontal_extent(doc, table_obj)
        except Exception:
            return None
        tblGrid = host_tbl_elem.find(qn('w:tblGrid'))
        if tblGrid is None:
            return tbl_left
        widths_pt = []
        for gc in tblGrid.findall(qn('w:gridCol')):
            try:
                widths_pt.append(int(gc.get(qn('w:w'))) / 20.0)
            except (TypeError, ValueError):
                widths_pt.append(0.0)
        if cell_idx is None or cell_idx < 0 or cell_idx >= len(widths_pt):
            return tbl_left
        return tbl_left + sum(widths_pt[:cell_idx])

    def _drawing_metrics(drawing_elem, host_tbl_elem, cell_idx):
        anchor = drawing_elem.find(f'{{{ns_wp}}}anchor')
        if anchor is None:
            return None
        posH = anchor.find(f'{{{ns_wp}}}positionH')
        posV = anchor.find(f'{{{ns_wp}}}positionV')
        extent = anchor.find(f'{{{ns_wp}}}extent')
        if posH is None or posV is None or extent is None:
            return None
        try:
            cx_pt = float(extent.get('cx')) / EMU_PER_PT
        except (TypeError, ValueError):
            return None
        offH_elem = posH.find(f'{{{ns_wp}}}posOffset')
        try:
            offH_pt = float(offH_elem.text.strip()) / EMU_PER_PT
        except (AttributeError, TypeError, ValueError):
            return None
        relH = posH.get('relativeFrom') or 'column'
        if relH == 'page':
            base_x = 0.0
        elif relH == 'margin':
            base_x = left_margin_pt
        else:
            cell_left = _cell_left_page_x(host_tbl_elem, cell_idx)
            base_x = cell_left if cell_left is not None else left_margin_pt
        left_x = base_x + offH_pt
        right_x = left_x + cx_pt
        center_x = left_x + cx_pt / 2.0
        offV_elem = posV.find(f'{{{ns_wp}}}posOffset')
        try:
            offV_pt = float(offV_elem.text.strip()) / EMU_PER_PT
        except (AttributeError, TypeError, ValueError):
            offV_pt = None
        return {'left': left_x, 'right': right_x, 'center_x': center_x, 'offV': offV_pt}

    def _pict_metrics(pict_elem, host_tbl_elem, cell_idx):
        shape = pict_elem.find(f'{{{ns_v}}}shape')
        if shape is None:
            return None
        style = shape.get('style', '') or ''
        attrs = {}
        for kv in style.split(';'):
            if ':' in kv:
                k, v = kv.split(':', 1)
                attrs[k.strip().lower()] = v.strip()
        # 严格只认 VML 官方坐标属性：margin-left / margin-top。
        # CSS 的 left/top 在办公软件里不参与浮动图形定位，回退到它们会导致
        # "XML 里有坐标但用户看不到"的假加分/假合规。
        ml = _parse_pt(attrs.get('margin-left'))
        mt = _parse_pt(attrs.get('margin-top'))
        w = _parse_pt(attrs.get('width')) or 0.0
        if ml is None:
            return None
        h_rel = (attrs.get('mso-position-horizontal-relative') or 'column').lower()
        if h_rel == 'page':
            base_x = 0.0
        elif h_rel == 'margin':
            base_x = left_margin_pt
        else:
            cell_left = _cell_left_page_x(host_tbl_elem, cell_idx)
            base_x = cell_left if cell_left is not None else left_margin_pt
        left_x = base_x + ml
        right_x = left_x + w
        center_x = left_x + w / 2.0
        return {'left': left_x, 'right': right_x, 'center_x': center_x, 'offV': mt}

    def _check_one(elem, is_drawing):
        host = _host_body_tbl(elem)
        if host is None:
            return False  # 不在表格里 → 不符合"续下页"位置语义
        r_idx, c_idx = _anchor_row_and_cell_index(elem)
        if r_idx is None or r_idx == 0:
            return False  # 锚定在首行/表头之外 → 不符合语义
        if is_drawing:
            m = _drawing_metrics(elem, host, c_idx)
        else:
            m = _pict_metrics(elem, host, c_idx)
        if m is None:
            return False
        # (2) 位于页面右下侧（水平右侧 + 垂直下方 via posOffset_V > 0）
        if m['center_x'] < page_mid_x:
            return False
        if m['offV'] is None or m['offV'] <= 0:
            return False
        # (4) 靠近表格右边界
        table_obj = tbl_id_to_table.get(id(host))
        if table_obj is None:
            return False
        try:
            _, tbl_right = compute_table_horizontal_extent(doc, table_obj)
        except Exception:
            return False
        if abs(tbl_right - m['right']) > TOL_X:
            return False
        # (5) 靠近表格下边线：posOffset_V > 0 已满足（锚定行即该页表格下边线所在处）
        return True

    checked_any = False
    for drawing in body.iter(f'{{{ns_w}}}drawing'):
        if _in_fallback(drawing):
            continue
        has_xy = any(t.text and '续下页' in t.text for t in drawing.iter(f'{{{ns_w}}}t'))
        if not has_xy:
            continue
        checked_any = True
        if not _check_one(drawing, is_drawing=True):
            return 0, False
    for pict in body.iter(f'{{{ns_w}}}pict'):
        if _in_fallback(pict):
            continue
        has_xy = any(t.text and '续下页' in t.text for t in pict.iter(f'{{{ns_w}}}t'))
        if not has_xy:
            continue
        checked_any = True
        if not _check_one(pict, is_drawing=False):
            return 0, False

    if not checked_any:
        return 0, False
    return 5, True


def check_xuyexia_style(doc):
    """+1："续下页"字体颜色为黑色，无其他样式。

    严格按细则的每一个点判定（不引入细则未要求的额外约束）：
      (1) 字体颜色为黑色：办公软件（Word/WPS）中渲染为黑色。
          - 显式 w:color val 为 "000000" / "auto" / 空 → 黑色；
          - 未显式设置（rPr 内无 w:color，且样式链/docDefaults 也未覆盖）→ 默认为
            "auto"，办公软件白底黑字渲染 → 视为黑色；
          - 其它 hex 值（如 FF0000）或 w:themeColor 指向非黑色 → 不通过。
      (2) 无其他样式：办公软件字体对话框里可切换的显式样式均不能存在。
          与细则严格对应，检查项覆盖：
          加粗 w:b / w:bCs、斜体 w:i / w:iCs、下划线 w:u (val≠none)、
          删除线 w:strike / 双删除线 w:dstrike、突出显示 w:highlight (val≠none)、
          底纹 w:shd (val≠clear/nil 或 fill≠auto)、
          阴文 w:vanish、隐藏 w:webHidden、
          上下标 w:vertAlign (val≠baseline)、
          阴影 w:shadow、空心 w:outline、阳文 w:emboss、阴文 w:imprint、
          小型大写 w:smallCaps、全部大写 w:caps。
          解析层级：直接 rPr → 段落 rPr → 字符样式 rStyle → 段落样式 pStyle
          → docDefaults。任一层引入非默认样式即不通过。

    办公软件（Word/WPS）有效性：
      - 文本框内文字在办公软件里最终渲染由 rPr 与样式继承共同决定。
        因此不仅检查 run 直接 rPr，还沿样式链回溯，确保办公软件实际渲染
        效果与细则一致。
      - 未指定 color 时办公软件按 "auto" 渲染为黑色，遵循 OOXML 规范。
    """
    containers = get_shapes_with_xuyexia(doc)
    if not containers:
        return 0, False

    styles = doc.styles

    def _get_style_by_id(style_id):
        if not style_id:
            return None
        for s in styles:
            try:
                if s.style_id == style_id:
                    return s
            except Exception:
                continue
        return None

    def _rPr_of_style_chain(style_obj, kind):
        """沿样式链（含 basedOn）逐级返回 rPr 元素列表，kind ∈ {'char','para'}。
        对字符样式取 style.element/w:rPr；对段落样式取 style.element/w:rPr（段落样式也可含默认 run 属性）。"""
        chain = []
        cur = style_obj
        depth = 0
        while cur is not None and depth < 20:
            rPr_elem = cur.element.find(qn('w:rPr'))
            if rPr_elem is not None:
                chain.append(rPr_elem)
            try:
                cur = cur.base_style
            except Exception:
                cur = None
            depth += 1
        return chain

    def _doc_defaults_rPr():
        styles_elem = styles.element
        dd = styles_elem.find(qn('w:docDefaults'))
        if dd is None:
            return None
        rPrDefault = dd.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            return None
        return rPrDefault.find(qn('w:rPr'))

    def _collect_rPr_chain(r_elem):
        """按办公软件解析顺序收集所有生效的 rPr（越前面优先级越高，用于"是否存在某属性"检查取合集）。
        顺序：run rPr → rStyle 链 → 段落 rPr → 段落样式 pStyle 链 → docDefaults。"""
        chain = []
        run_rPr = r_elem.find(qn('w:rPr'))
        if run_rPr is not None:
            chain.append(run_rPr)
            rStyle = run_rPr.find(qn('w:rStyle'))
            if rStyle is not None:
                sid = rStyle.get(qn('w:val'))
                chain.extend(_rPr_of_style_chain(_get_style_by_id(sid), 'char'))

        # 段落
        p = r_elem.getparent()
        while p is not None:
            ptag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if ptag == 'p':
                break
            p = p.getparent()
        if p is not None:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                para_rPr = pPr.find(qn('w:rPr'))
                if para_rPr is not None:
                    chain.append(para_rPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'))
                    chain.extend(_rPr_of_style_chain(_get_style_by_id(sid), 'para'))

        dd_rPr = _doc_defaults_rPr()
        if dd_rPr is not None:
            chain.append(dd_rPr)
        return chain

    def _first_of(chain, tag_qn):
        """按办公软件覆盖顺序返回首个非 None 的子元素（最高优先级层的取值）。"""
        for rPr_elem in chain:
            elt = rPr_elem.find(tag_qn)
            if elt is not None:
                return elt
        return None

    def _any_active(chain, tag_qn):
        """是否任何一层显式激活了该 boolean 类样式（w:val 缺省或非 '0'/'false' 视为激活）。"""
        elt = _first_of(chain, tag_qn)
        if elt is None:
            return False
        val = elt.get(qn('w:val'))
        if val is None:
            return True
        return val.lower() not in ('0', 'false', 'off')

    def _color_is_black(chain):
        color_elem = _first_of(chain, qn('w:color'))
        if color_elem is None:
            return True  # 办公软件默认 auto → 黑色
        val = (color_elem.get(qn('w:val')) or '').strip().lower()
        theme = color_elem.get(qn('w:themeColor'))
        if theme:
            # 非明确黑色主题 → 不视为黑色
            return theme.lower() in ('text1', 'dark1')
        if val in ('', 'auto', '000000'):
            return True
        return False

    def _underline_active(chain):
        elt = _first_of(chain, qn('w:u'))
        if elt is None:
            return False
        val = (elt.get(qn('w:val')) or '').lower()
        return val not in ('', 'none')

    def _highlight_active(chain):
        elt = _first_of(chain, qn('w:highlight'))
        if elt is None:
            return False
        val = (elt.get(qn('w:val')) or '').lower()
        return val not in ('', 'none')

    def _shd_active(chain):
        elt = _first_of(chain, qn('w:shd'))
        if elt is None:
            return False
        val = (elt.get(qn('w:val')) or '').lower()
        fill = (elt.get(qn('w:fill')) or '').lower()
        # clear + fill=auto 视为无底纹
        if val in ('', 'nil', 'clear') and fill in ('', 'auto'):
            return False
        return True

    def _vertalign_nonbaseline(chain):
        elt = _first_of(chain, qn('w:vertAlign'))
        if elt is None:
            return False
        val = (elt.get(qn('w:val')) or '').lower()
        return val not in ('', 'baseline')

    for container in containers:
        txbx_content = container.find(qn('w:txbxContent'))
        if txbx_content is None:
            # DrawingML 的 txbxContent 深度不同（wps:txbx/w:txbxContent 或
            # a:graphicData/wps:wsp/wps:txbx/w:txbxContent），用 .// 兜底查找。
            txbx_content = container.find(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent')
        if txbx_content is None:
            return 0, False

        for r_elem in txbx_content.iter(qn('w:r')):
            t_elem = r_elem.find(qn('w:t'))
            if t_elem is None or not t_elem.text or '续下页' not in t_elem.text:
                continue

            chain = _collect_rPr_chain(r_elem)

            # (1) 字体颜色为黑色
            if not _color_is_black(chain):
                return 0, False

            # (2) 无其他样式
            if _any_active(chain, qn('w:b')):
                return 0, False
            if _any_active(chain, qn('w:bCs')):
                return 0, False
            if _any_active(chain, qn('w:i')):
                return 0, False
            if _any_active(chain, qn('w:iCs')):
                return 0, False
            if _underline_active(chain):
                return 0, False
            if _any_active(chain, qn('w:strike')):
                return 0, False
            if _any_active(chain, qn('w:dstrike')):
                return 0, False
            if _highlight_active(chain):
                return 0, False
            if _shd_active(chain):
                return 0, False
            if _any_active(chain, qn('w:vanish')):
                return 0, False
            if _any_active(chain, qn('w:webHidden')):
                return 0, False
            if _vertalign_nonbaseline(chain):
                return 0, False
            if _any_active(chain, qn('w:shadow')):
                return 0, False
            if _any_active(chain, qn('w:outline')):
                return 0, False
            if _any_active(chain, qn('w:emboss')):
                return 0, False
            if _any_active(chain, qn('w:imprint')):
                return 0, False
            if _any_active(chain, qn('w:smallCaps')):
                return 0, False
            if _any_active(chain, qn('w:caps')):
                return 0, False

    return 1, True


def check_table_borders(doc):
    """+5：每一页的表格：都有三条平行的边线，上边线和下边线1.5磅，内边线0.75磅。

    严格按细则的每一个点判定（不引入细则未要求的额外约束）：
      (1) 每一页的表格：遍历文档中每一个物理表格。跨页表格在文件中体现为
          多个物理表格（每个物理表格对应一页里的表格），任一表格不合规即失败。
      (2) 都有三条平行的边线：
          - 上边线：表格第一行的 top 必须实际渲染（val 非 nil、sz>0）；
          - 内边线：表格第一行的 bottom（即表头下方的水平分隔线）必须实际渲染；
          - 下边线：表格最后一行的 bottom 必须实际渲染。
      (3) 上边线 1.5 磅：sz = 12 半磅。
      (4) 下边线 1.5 磅：sz = 12 半磅。
      (5) 内边线 0.75 磅：sz = 6 半磅。

    办公软件（Word/WPS）有效性：
      - Word/WPS 渲染单元格边线按 tcBorders > tblBorders > tblStyle 依次继承合成；
        任何一层显式设为 val="nil" 都会隐藏该边线。
      - 只有当边线在办公软件里实际"看得见"（val 非 nil 且 sz>0）时才计入
        "三条平行边线"。
      - 上边线取表格第一行的顶边；下边线取最后一行的底边；内边线取第一行的
        底边（即办公软件中三线表模式下表头与数据的分隔线）。
      - 表格所有单元格在同一边上的边线共同渲染成一整条线，所以要求每一列
        对应的单元格都满足对应边线的存在与磅数。
    """
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _resolve_cell_border(cell, table, edge, is_first_row, is_last_row):
        """按 Word 渲染优先级解析单元格某条边线，返回 (val, sz_int_or_None)。

        优先级：cell tcBorders > table tblBorders > table style tblBorders。
        对 tblBorders 层：top 边在第一行取 top、否则取 insideH；
                       bottom 边在最后一行取 bottom、否则取 insideH。
        """
        tc = cell._tc
        # (1) 单元格级
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                b = tcBorders.find(qn(f'w:{edge}'))
                if b is not None:
                    val = b.get(qn('w:val'))
                    sz = b.get(qn('w:sz'))
                    try:
                        sz_i = int(sz) if sz is not None else None
                    except ValueError:
                        sz_i = None
                    return val, sz_i

        # (2) 表格级 tblBorders
        def _pick_from_tblBorders(tblBorders_elem):
            if tblBorders_elem is None:
                return None
            if edge == 'top':
                tag = 'top' if is_first_row else 'insideH'
            elif edge == 'bottom':
                tag = 'bottom' if is_last_row else 'insideH'
            else:
                tag = edge
            b = tblBorders_elem.find(qn(f'w:{tag}'))
            if b is None:
                return None
            v = b.get(qn('w:val'))
            s = b.get(qn('w:sz'))
            try:
                s_i = int(s) if s is not None else None
            except ValueError:
                s_i = None
            return v, s_i

        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            got = _pick_from_tblBorders(tblPr.find(qn('w:tblBorders')))
            if got is not None:
                return got

            # (3) 表格样式
            tblStyle = tblPr.find(qn('w:tblStyle'))
            if tblStyle is not None:
                sid = tblStyle.get(qn('w:val'))
                if sid:
                    for st in doc.styles:
                        try:
                            if st.style_id != sid:
                                continue
                        except Exception:
                            continue
                        st_tblPr = st.element.find(qn('w:tblPr'))
                        if st_tblPr is not None:
                            got = _pick_from_tblBorders(st_tblPr.find(qn('w:tblBorders')))
                            if got is not None:
                                return got
                        break
        return None, None

    def _line_visible(val, sz):
        """办公软件中该边线是否实际可见。"""
        if not val or val.lower() == 'nil':
            return False
        if sz is None or sz <= 0:
            return False
        return True

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            # 少于 2 行无法形成"三条平行边线"（无内边线位置）→ 不合规
            return 0, False

        n = len(rows)

        # (2)-(5)：逐列检查上边线、内边线、下边线
        first_row_cells = rows[0].cells
        last_row_cells = rows[n - 1].cells

        # 上边线：第一行每个单元格 top，必须可见 且 sz=12
        for cell in first_row_cells:
            val, sz = _resolve_cell_border(cell, table, 'top',
                                           is_first_row=True, is_last_row=(n == 1))
            if not _line_visible(val, sz):
                return 0, False
            if sz != 12:
                return 0, False

        # 内边线：第一行每个单元格 bottom（表头分隔线），必须可见 且 sz=6
        for cell in first_row_cells:
            val, sz = _resolve_cell_border(cell, table, 'bottom',
                                           is_first_row=True, is_last_row=(n == 1))
            if not _line_visible(val, sz):
                return 0, False
            if sz != 6:
                return 0, False

        # 下边线：最后一行每个单元格 bottom，必须可见 且 sz=12
        for cell in last_row_cells:
            val, sz = _resolve_cell_border(cell, table, 'bottom',
                                           is_first_row=(n == 1), is_last_row=True)
            if not _line_visible(val, sz):
                return 0, False
            if sz != 12:
                return 0, False

    return 5, True


def check_table_headers(doc):
    """+5：每一页最上方的表格：都要有表头且表头中文字体字号为8pt宋体加粗。

    严格按细则的每一个点判定（不引入细则未要求的额外约束）：
      (1) 每一页最上方的表格：遍历文档中每一个物理表格 <w:tbl>。
          - 单页表格：该页只有它自己 → 就是该页最上方的表格；
          - 跨页表格续页：续页开头即该表格，它也是该页最上方的表格。
          文档中每个物理表格都对应"某一页最上方的表格"，故逐个检查。
      (2) 有表头：表格第一行必须存在带实际文字内容的 run。
      (3) 表头中文字号为 8pt：办公软件中生效的 w:sz = 16 半磅（8 pt）。
          仅对含中文字符（U+4E00–U+9FFF）的 run 判定字号（细则限定"中文字体字号"）。
      (4) 表头中文字体为宋体：办公软件中生效的中文字体（w:rFonts/w:eastAsia）= "宋体"。
          仅对含中文字符的 run 判定字体。
      (5) 表头加粗：办公软件中生效的 w:b 为激活态（对表头所有非空 run 判定，
          "加粗"不区分中英文）。

    办公软件（Word/WPS）有效性：
      - 字体/字号/加粗按 OOXML 覆盖顺序生效：
        run rPr → 字符样式（含 basedOn 链） → 段落 rPr → 段落样式（含 basedOn 链）
        → docDefaults/rPrDefault/rPr。代码沿此链回溯，取首个显式值为准，
        无显式值时按办公软件默认渲染判断。
    """
    styles = doc.styles

    def _get_style_by_id(style_id):
        if not style_id:
            return None
        for s in styles:
            try:
                if s.style_id == style_id:
                    return s
            except Exception:
                continue
        return None

    def _rPr_of_style_chain(style_obj):
        chain = []
        cur = style_obj
        depth = 0
        while cur is not None and depth < 20:
            rPr_elem = cur.element.find(qn('w:rPr'))
            if rPr_elem is not None:
                chain.append(rPr_elem)
            try:
                cur = cur.base_style
            except Exception:
                cur = None
            depth += 1
        return chain

    def _doc_defaults_rPr():
        styles_elem = styles.element
        dd = styles_elem.find(qn('w:docDefaults'))
        if dd is None:
            return None
        rPrDefault = dd.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            return None
        return rPrDefault.find(qn('w:rPr'))

    def _collect_rPr_chain(r_elem):
        """按办公软件解析顺序返回生效的 rPr 元素列表（先出现的层级优先级更高）。
        顺序：run rPr → rStyle 链 → 段落 rPr → pStyle 链 → docDefaults。
        """
        chain = []
        run_rPr = r_elem.find(qn('w:rPr'))
        if run_rPr is not None:
            chain.append(run_rPr)
            rStyle = run_rPr.find(qn('w:rStyle'))
            if rStyle is not None:
                sid = rStyle.get(qn('w:val'))
                chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        p = r_elem.getparent()
        while p is not None:
            ptag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if ptag == 'p':
                break
            p = p.getparent()
        if p is not None:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                para_rPr = pPr.find(qn('w:rPr'))
                if para_rPr is not None:
                    chain.append(para_rPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'))
                    chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        dd_rPr = _doc_defaults_rPr()
        if dd_rPr is not None:
            chain.append(dd_rPr)
        return chain

    def _first_of(chain, tag_qn):
        for rPr_elem in chain:
            elt = rPr_elem.find(tag_qn)
            if elt is not None:
                return elt
        return None

    def _bold_active(chain):
        elt = _first_of(chain, qn('w:b'))
        if elt is None:
            return False
        val = elt.get(qn('w:val'))
        if val is None:
            return True
        return val.lower() not in ('0', 'false', 'off')

    def _eastAsia_font(chain):
        elt = _first_of(chain, qn('w:rFonts'))
        if elt is None:
            return None
        return elt.get(qn('w:eastAsia'))

    def _sz_halfpt(chain):
        elt = _first_of(chain, qn('w:sz'))
        if elt is None:
            return None
        try:
            return int(elt.get(qn('w:val')))
        except (TypeError, ValueError):
            return None

    for table in doc.tables:
        if len(table.rows) == 0:
            return 0, False

        header_row = table.rows[0]
        has_any_content = False

        for cell in header_row.cells:
            for para in cell.paragraphs:
                for r_elem in para._p.iter(qn('w:r')):
                    t_elem = r_elem.find(qn('w:t'))
                    if t_elem is None or not (t_elem.text and t_elem.text.strip()):
                        continue
                    has_any_content = True

                    chain = _collect_rPr_chain(r_elem)
                    text = t_elem.text
                    has_chinese = any('一' <= ch <= '鿿' for ch in text)

                    # (4) 中文字体为宋体（仅对含中文字符的 run 判定）
                    if has_chinese and _eastAsia_font(chain) != '宋体':
                        return 0, False

                    # (3) 中文字号 8pt = sz=16 半磅（仅对含中文字符的 run 判定）
                    if has_chinese and _sz_halfpt(chain) != 16:
                        return 0, False

                    # (5) 加粗（表头所有非空 run）
                    if not _bold_active(chain):
                        return 0, False

        # (2) 有表头 —— 第一行必须存在实际文字内容
        if not has_any_content:
            return 0, False

    return 5, True


def check_header_consistency(doc):
    """+5："续下页"上方的表格表头与下方的表格表头：内容、中文和英文数字的字体字号格式均一致。

    办公软件里"续下页"上下的表格有两种等价实现：
      结构一（行重复表头 / 单 <w:tbl>）：一张 <w:tbl> 首行含 <w:trPr>
        <w:tblHeader/>，办公软件在跨页时会把该首行自动重复渲染到下一页顶部；
        "续下页"用浮动文本框锚定在该表跨页边界所在行的单元格内。此时
        "上方的表格表头" = 该表首行（本页底部渲染的实例），
        "下方的表格表头" = 该表首行（下一页顶部自动重复渲染的实例），
        二者是同一个 <w:tr> 元素在两页各渲染一次，办公软件保证内容与格式
        逐字节一致。
      结构二（物理分表）：作者手动把一张跨页表拆成两张相邻 <w:tbl>，
        以浮动文本框放"续下页"于上表末端。此时
        "上方的表格表头" = 上表第一行，
        "下方的表格表头" = 下表第一行，
        需按第一行内容/格式逐单元格逐 run 比对。

    严格按细则的每一个点判定（不引入细则未要求的额外约束）：
      (1) 定位每个"续下页"：遍历所有"续下页"浮动文本框（DrawingML + VML，
          跳过 <mc:Fallback>），取其 anchor 段落所在的 body 直接子 <w:tbl>。
      (2) 找到"下方的表格"：
          - 若"上方"表首行有 <w:tblHeader/> → 下方即同一表首行（结构一）；
          - 否则若 body 之后紧邻另一张 <w:tbl> → 下方即该表（结构二）；
          - 否则该续下页不构成"上下表头对"，跳过。
      (3) 内容一致 & (4) 格式一致：
          - 结构一：同一 <w:tr> 元素，办公软件保证两页渲染逐格一致，判通过；
          - 结构二：两表第一行按列比较，规范化后文本相等且每单元格非空 run
            序列的生效格式（字体、字号、加粗、斜体、下划线、删除线、字体
            颜色、突出显示、上下标）逐项相同。

    办公软件（Word/WPS）有效性：
      - 格式解析按 OOXML 覆盖顺序：run rPr → rStyle 链 → 段落 rPr → pStyle
        链 → docDefaults，取每一层首个显式值，等价于 Word/WPS 最终渲染值。
      - 排除浮动对象（<w:pict>/<w:drawing>）内部的 run/段落，避免"续下页"
        浮框文本被计入宿主单元格。
    """
    body = doc.element.body
    body_children = list(body)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_mc = 'http://schemas.openxmlformats.org/markup-compatibility/2006'

    # body_pos → docx Table 对象
    body_pos_to_table = {}
    ti_ = -1
    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            ti_ += 1
            body_pos_to_table[i] = doc.tables[ti_]

    styles = doc.styles

    def _get_style_by_id(style_id):
        if not style_id:
            return None
        for s in styles:
            try:
                if s.style_id == style_id:
                    return s
            except Exception:
                continue
        return None

    def _rPr_of_style_chain(style_obj):
        chain = []
        cur = style_obj
        depth = 0
        while cur is not None and depth < 20:
            rPr_elem = cur.element.find(qn('w:rPr'))
            if rPr_elem is not None:
                chain.append(rPr_elem)
            try:
                cur = cur.base_style
            except Exception:
                cur = None
            depth += 1
        return chain

    def _doc_defaults_rPr():
        dd = styles.element.find(qn('w:docDefaults'))
        if dd is None:
            return None
        rPrDefault = dd.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            return None
        return rPrDefault.find(qn('w:rPr'))

    def _collect_rPr_chain(r_elem):
        chain = []
        run_rPr = r_elem.find(qn('w:rPr'))
        if run_rPr is not None:
            chain.append(run_rPr)
            rStyle = run_rPr.find(qn('w:rStyle'))
            if rStyle is not None:
                sid = rStyle.get(qn('w:val'))
                chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        p = r_elem.getparent()
        while p is not None:
            ptag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if ptag == 'p':
                break
            p = p.getparent()
        if p is not None:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                para_rPr = pPr.find(qn('w:rPr'))
                if para_rPr is not None:
                    chain.append(para_rPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'))
                    chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        dd_rPr = _doc_defaults_rPr()
        if dd_rPr is not None:
            chain.append(dd_rPr)
        return chain

    def _first_of(chain, tag_qn):
        for rPr_elem in chain:
            elt = rPr_elem.find(tag_qn)
            if elt is not None:
                return elt
        return None

    def _bool_state(chain, tag_qn):
        elt = _first_of(chain, tag_qn)
        if elt is None:
            return False
        v = elt.get(qn('w:val'))
        if v is None:
            return True
        return v.lower() not in ('0', 'false', 'off')

    def _resolved_run_format(r_elem):
        chain = _collect_rPr_chain(r_elem)
        fonts_elem = _first_of(chain, qn('w:rFonts'))
        ascii_f = fonts_elem.get(qn('w:ascii')) if fonts_elem is not None else None
        hAnsi_f = fonts_elem.get(qn('w:hAnsi')) if fonts_elem is not None else None
        eastAsia_f = fonts_elem.get(qn('w:eastAsia')) if fonts_elem is not None else None
        cs_f = fonts_elem.get(qn('w:cs')) if fonts_elem is not None else None

        sz_elem = _first_of(chain, qn('w:sz'))
        try:
            sz = int(sz_elem.get(qn('w:val'))) if sz_elem is not None else None
        except (TypeError, ValueError):
            sz = None

        color_elem = _first_of(chain, qn('w:color'))
        if color_elem is None:
            color = 'auto'
        else:
            v = (color_elem.get(qn('w:val')) or '').strip().lower()
            theme = color_elem.get(qn('w:themeColor'))
            color = f"theme:{theme}" if theme else (v or 'auto')

        u_elem = _first_of(chain, qn('w:u'))
        underline = (u_elem.get(qn('w:val')) or 'single').lower() if u_elem is not None else 'none'
        if underline == '':
            underline = 'single'

        hl_elem = _first_of(chain, qn('w:highlight'))
        highlight = (hl_elem.get(qn('w:val')) or 'none').lower() if hl_elem is not None else 'none'

        va_elem = _first_of(chain, qn('w:vertAlign'))
        vert_align = (va_elem.get(qn('w:val')) or 'baseline').lower() if va_elem is not None else 'baseline'

        return (
            ascii_f, hAnsi_f, eastAsia_f, cs_f,
            sz,
            _bool_state(chain, qn('w:b')),
            _bool_state(chain, qn('w:bCs')),
            _bool_state(chain, qn('w:i')),
            _bool_state(chain, qn('w:iCs')),
            _bool_state(chain, qn('w:strike')),
            _bool_state(chain, qn('w:dstrike')),
            underline,
            color,
            highlight,
            vert_align,
        )

    def _is_inside_float(elem):
        p = elem.getparent()
        while p is not None:
            tag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if tag in ('pict', 'drawing'):
                return True
            p = p.getparent()
        return False

    def _in_fallback(elem):
        p = elem.getparent()
        while p is not None:
            if p.tag == f'{{{ns_mc}}}Fallback':
                return True
            p = p.getparent()
        return False

    def _cell_text(cell):
        parts = []
        for p_elem in cell._tc.iter(qn('w:p')):
            if _is_inside_float(p_elem):
                continue
            for t_elem in p_elem.iter(qn('w:t')):
                if _is_inside_float(t_elem):
                    continue
                if t_elem.text:
                    parts.append(t_elem.text)
            parts.append('\n')
        return ''.join(parts).strip()

    def _cell_run_formats(cell):
        fmts = []
        for r_elem in cell._tc.iter(qn('w:r')):
            if _is_inside_float(r_elem):
                continue
            t_elem = r_elem.find(qn('w:t'))
            if t_elem is None or not (t_elem.text and t_elem.text.strip()):
                continue
            fmts.append(_resolved_run_format(r_elem))
        return fmts

    def _first_row_has_tblHeader(tbl_elem):
        """判断该 <w:tbl> 首行是否为重复表头（<w:trPr><w:tblHeader/>）。"""
        first_tr = tbl_elem.find(qn('w:tr'))
        if first_tr is None:
            return False
        trPr = first_tr.find(qn('w:trPr'))
        if trPr is None:
            return False
        return trPr.find(qn('w:tblHeader')) is not None

    def _host_body_tbl(elem):
        e = elem.getparent()
        while e is not None:
            parent = e.getparent()
            tag = e.tag.split('}')[1] if '}' in e.tag else e.tag
            if tag == 'tbl' and parent is not None and parent.tag == body.tag:
                return e
            e = parent
        return None

    # 收集所有"续下页"浮框（DrawingML + VML，跳过 mc:Fallback）
    seen_floats = set()
    checked_keys = set()
    any_checked = False

    def _iter_xyye_floats():
        for drawing in body.iter(f'{{{ns_w}}}drawing'):
            if _in_fallback(drawing):
                continue
            if any(t.text and '续下页' in t.text for t in drawing.iter(f'{{{ns_w}}}t')):
                yield drawing
        for pict in body.iter(f'{{{ns_w}}}pict'):
            if _in_fallback(pict):
                continue
            if any(t.text and '续下页' in t.text for t in pict.iter(f'{{{ns_w}}}t')):
                yield pict

    for float_elem in _iter_xyye_floats():
        if id(float_elem) in seen_floats:
            continue
        seen_floats.add(id(float_elem))

        upper_tbl = _host_body_tbl(float_elem)
        if upper_tbl is None:
            continue  # 不在表格里，跳过

        upper_body_pos = None
        for i, c in enumerate(body_children):
            if c is upper_tbl:
                upper_body_pos = i
                break
        if upper_body_pos is None:
            continue

        # 结构一：首行 <w:tblHeader/> → 上下表头是同一 <w:tr>，办公软件保证一致
        if _first_row_has_tblHeader(upper_tbl):
            key = ('repeat', upper_body_pos)
            if key in checked_keys:
                continue
            checked_keys.add(key)
            any_checked = True
            # 同一 <w:tr> 在两页各渲染一次，内容与格式必然一致
            continue

        # 结构二：body 之后紧邻另一张 <w:tbl>（跳过纯空段落）
        lower_body_pos = None
        for j in range(upper_body_pos + 1, len(body_children)):
            sib = body_children[j]
            stag = sib.tag.split('}')[1] if '}' in sib.tag else sib.tag
            if stag == 'p':
                text = ''.join((tx.text or '') for tx in sib.iter(qn('w:t')))
                if text.strip():
                    break
                continue
            if stag == 'tbl':
                lower_body_pos = j
                break

        if lower_body_pos is None:
            continue  # 不构成上下表头对，跳过

        key = ('pair', upper_body_pos, lower_body_pos)
        if key in checked_keys:
            continue
        checked_keys.add(key)
        any_checked = True

        ta = body_pos_to_table.get(upper_body_pos)
        tb = body_pos_to_table.get(lower_body_pos)
        if ta is None or tb is None:
            return 0, False

        cells_a = ta.rows[0].cells
        cells_b = tb.rows[0].cells
        if len(cells_a) != len(cells_b):
            return 0, False

        for ci in range(len(cells_a)):
            if _cell_text(cells_a[ci]) != _cell_text(cells_b[ci]):
                return 0, False
            if _cell_run_formats(cells_a[ci]) != _cell_run_formats(cells_b[ci]):
                return 0, False

    if not any_checked:
        return 0, False
    return 5, True
    body = doc.element.body
    body_children = list(body)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # body_pos → docx Table 对象
    body_pos_to_table = {}
    ti_ = -1
    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            ti_ += 1
            body_pos_to_table[i] = doc.tables[ti_]

    styles = doc.styles

    def _get_style_by_id(style_id):
        if not style_id:
            return None
        for s in styles:
            try:
                if s.style_id == style_id:
                    return s
            except Exception:
                continue
        return None

    def _rPr_of_style_chain(style_obj):
        chain = []
        cur = style_obj
        depth = 0
        while cur is not None and depth < 20:
            rPr_elem = cur.element.find(qn('w:rPr'))
            if rPr_elem is not None:
                chain.append(rPr_elem)
            try:
                cur = cur.base_style
            except Exception:
                cur = None
            depth += 1
        return chain

    def _doc_defaults_rPr():
        dd = styles.element.find(qn('w:docDefaults'))
        if dd is None:
            return None
        rPrDefault = dd.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            return None
        return rPrDefault.find(qn('w:rPr'))

    def _collect_rPr_chain(r_elem):
        chain = []
        run_rPr = r_elem.find(qn('w:rPr'))
        if run_rPr is not None:
            chain.append(run_rPr)
            rStyle = run_rPr.find(qn('w:rStyle'))
            if rStyle is not None:
                sid = rStyle.get(qn('w:val'))
                chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        p = r_elem.getparent()
        while p is not None:
            ptag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if ptag == 'p':
                break
            p = p.getparent()
        if p is not None:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                para_rPr = pPr.find(qn('w:rPr'))
                if para_rPr is not None:
                    chain.append(para_rPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'))
                    chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        dd_rPr = _doc_defaults_rPr()
        if dd_rPr is not None:
            chain.append(dd_rPr)
        return chain

    def _first_of(chain, tag_qn):
        for rPr_elem in chain:
            elt = rPr_elem.find(tag_qn)
            if elt is not None:
                return elt
        return None

    def _bool_state(chain, tag_qn):
        elt = _first_of(chain, tag_qn)
        if elt is None:
            return False
        v = elt.get(qn('w:val'))
        if v is None:
            return True
        return v.lower() not in ('0', 'false', 'off')

    def _resolved_run_format(r_elem):
        chain = _collect_rPr_chain(r_elem)
        fonts_elem = _first_of(chain, qn('w:rFonts'))
        ascii_f = fonts_elem.get(qn('w:ascii')) if fonts_elem is not None else None
        hAnsi_f = fonts_elem.get(qn('w:hAnsi')) if fonts_elem is not None else None
        eastAsia_f = fonts_elem.get(qn('w:eastAsia')) if fonts_elem is not None else None
        cs_f = fonts_elem.get(qn('w:cs')) if fonts_elem is not None else None

        sz_elem = _first_of(chain, qn('w:sz'))
        try:
            sz = int(sz_elem.get(qn('w:val'))) if sz_elem is not None else None
        except (TypeError, ValueError):
            sz = None

        color_elem = _first_of(chain, qn('w:color'))
        if color_elem is None:
            color = 'auto'
        else:
            v = (color_elem.get(qn('w:val')) or '').strip().lower()
            theme = color_elem.get(qn('w:themeColor'))
            color = f"theme:{theme}" if theme else (v or 'auto')

        u_elem = _first_of(chain, qn('w:u'))
        underline = (u_elem.get(qn('w:val')) or 'single').lower() if u_elem is not None else 'none'
        if underline == '':
            underline = 'single'

        hl_elem = _first_of(chain, qn('w:highlight'))
        highlight = (hl_elem.get(qn('w:val')) or 'none').lower() if hl_elem is not None else 'none'

        va_elem = _first_of(chain, qn('w:vertAlign'))
        vert_align = (va_elem.get(qn('w:val')) or 'baseline').lower() if va_elem is not None else 'baseline'

        return (
            ascii_f, hAnsi_f, eastAsia_f, cs_f,
            sz,
            _bool_state(chain, qn('w:b')),
            _bool_state(chain, qn('w:bCs')),
            _bool_state(chain, qn('w:i')),
            _bool_state(chain, qn('w:iCs')),
            _bool_state(chain, qn('w:strike')),
            _bool_state(chain, qn('w:dstrike')),
            underline,
            color,
            highlight,
            vert_align,
        )

    def _is_inside_float(elem):
        """判断 elem 是否位于浮动对象（VML pict / DrawingML）内。
        办公软件里 <w:pict>/<w:drawing> 为浮动文本框或图形，其内部段落/run
        不是宿主单元格 header 行的组成部分，比对表头时须排除。"""
        p = elem.getparent()
        while p is not None:
            tag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if tag in ('pict', 'drawing'):
                return True
            p = p.getparent()
        return False

    def _cell_text(cell):
        parts = []
        for p_elem in cell._tc.iter(qn('w:p')):
            if _is_inside_float(p_elem):
                continue
            for t_elem in p_elem.iter(qn('w:t')):
                if _is_inside_float(t_elem):
                    continue
                if t_elem.text:
                    parts.append(t_elem.text)
            parts.append('\n')
        return ''.join(parts).strip()

    def _cell_run_formats(cell):
        fmts = []
        for r_elem in cell._tc.iter(qn('w:r')):
            if _is_inside_float(r_elem):
                continue
            t_elem = r_elem.find(qn('w:t'))
            if t_elem is None or not (t_elem.text and t_elem.text.strip()):
                continue
            fmts.append(_resolved_run_format(r_elem))
        return fmts

    # 遍历所有"续下页"文本框，确定 (上方表格, 下方表格) 对
    seen_picts = set()
    checked_pairs = set()
    any_checked = False

    for t_elem in body.iter(f'{{{ns_w}}}t'):
        if not (t_elem.text and '续下页' in t_elem.text):
            continue

        anc = t_elem.getparent()
        pict_elem = None
        while anc is not None:
            atag = anc.tag.split('}')[1] if '}' in anc.tag else anc.tag
            if atag == 'pict':
                pict_elem = anc
                break
            anc = anc.getparent()
        if pict_elem is None or id(pict_elem) in seen_picts:
            continue
        seen_picts.add(id(pict_elem))

        # anchor 段落
        anchor_p = pict_elem.getparent()
        while anchor_p is not None:
            atag = anchor_p.tag.split('}')[1] if '}' in anchor_p.tag else anchor_p.tag
            if atag == 'p':
                break
            anchor_p = anchor_p.getparent()
        if anchor_p is None:
            continue

        # anchor 所在的 body 直接子 tbl 位置（"上方的表格"）
        upper_body_pos = None
        a = anchor_p.getparent()
        while a is not None:
            parent = a.getparent()
            if parent is None:
                break
            atag = a.tag.split('}')[1] if '}' in a.tag else a.tag
            if atag == 'tbl' and parent.tag == body.tag:
                for i, c in enumerate(body_children):
                    if c is a:
                        upper_body_pos = i
                        break
                break
            a = parent

        if upper_body_pos is None:
            # anchor 不在表格里 → 该"续下页"不指示"上方/下方"表格对，跳过
            continue

        # "下方的表格"：body 中 upper_body_pos 之后紧邻的下一个 tbl
        lower_body_pos = None
        for j in range(upper_body_pos + 1, len(body_children)):
            sib = body_children[j]
            stag = sib.tag.split('}')[1] if '}' in sib.tag else sib.tag
            if stag == 'p':
                # 非空段落 → 中间已插入正文段，跨页拆分中断
                text = ''.join((tx.text or '') for tx in sib.iter(qn('w:t')))
                if text.strip():
                    break
                continue
            if stag == 'tbl':
                lower_body_pos = j
                break

        if lower_body_pos is None:
            # 该"续下页"后不存在紧邻的下方表格 → 不构成"上方/下方"表格对，跳过
            continue

        pair_key = (upper_body_pos, lower_body_pos)
        if pair_key in checked_pairs:
            continue
        checked_pairs.add(pair_key)
        any_checked = True

        ta = body_pos_to_table.get(upper_body_pos)
        tb = body_pos_to_table.get(lower_body_pos)
        if ta is None or tb is None:
            return 0, False

        cells_a = ta.rows[0].cells
        cells_b = tb.rows[0].cells

        # 列数一致
        if len(cells_a) != len(cells_b):
            return 0, False

        # (3) 内容一致：逐单元格文本
        # (4) 格式一致：逐单元格 run 序列的生效格式
        for ci in range(len(cells_a)):
            if _cell_text(cells_a[ci]) != _cell_text(cells_b[ci]):
                return 0, False
            if _cell_run_formats(cells_a[ci]) != _cell_run_formats(cells_b[ci]):
                return 0, False

    if not any_checked:
        return 0, False
    return 5, True


def check_no_images(doc):
    """-5：交付文件出现图片。

    严格按细则的一个点判定（不引入细则未要求的额外约束）：
      (1) 交付文件出现图片 → 扣 5 分。

    办公软件（Word/WPS）有效性：
      "图片"以办公软件打开文档时用户能实际看到的图形对象为准。判定条件：
      (a) 该图片渲染元素不位于 <mc:Fallback> 中；<mc:AlternateContent> 里
          现代 Office 使用 <mc:Choice>，<mc:Fallback> 只服务老版本阅读器，
          在 Word/WPS 中不渲染。
      (b) 该元素持有真实的图片关系引用，并且引用的关系目标类型是 image：
          - DrawingML：<a:blip> 的 r:embed 或 r:link 指向 image relationship
          - VML：<v:imagedata> 的 r:id 或 o:relid 指向 image relationship
          - VML：<v:image> 元素本身
          仅有空 imagedata（无 r:id）不引用任何图片，办公软件不显示任何内容，
          不应扣分（如本文档中作为 mc:Fallback 备用的空 imagedata）。
      扫描范围包含主文档、所有页眉/页脚部件（办公软件会渲染这些部件中的
      图片）。
    """
    ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    ns_v = 'urn:schemas-microsoft-com:vml'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    ns_o = 'urn:schemas-microsoft-com:office:office'
    ns_mc = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    IMAGE_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'

    def _in_fallback(elem):
        p = elem.getparent()
        while p is not None:
            tag = p.tag
            if tag == f'{{{ns_mc}}}Fallback':
                return True
            p = p.getparent()
        return False

    def _rel_is_image(part, rid):
        if not rid:
            return False
        try:
            rel = part.rels.get(rid)
        except Exception:
            return False
        if rel is None:
            return False
        try:
            return rel.reltype == IMAGE_REL_TYPE
        except Exception:
            return False

    def _part_has_image(part):
        try:
            root = part.element
        except AttributeError:
            return False

        # DrawingML: <a:blip r:embed="…"/> or r:link="…"
        for blip in root.iter(f'{{{ns_a}}}blip'):
            if _in_fallback(blip):
                continue
            rid = blip.get(f'{{{ns_r}}}embed') or blip.get(f'{{{ns_r}}}link')
            if _rel_is_image(part, rid):
                return True

        # VML: <v:imagedata r:id/o:relid="…"/>
        for im in root.iter(f'{{{ns_v}}}imagedata'):
            if _in_fallback(im):
                continue
            rid = im.get(f'{{{ns_r}}}id') or im.get(f'{{{ns_o}}}relid')
            if _rel_is_image(part, rid):
                return True

        # VML: <v:image> 元素直接表示图片渲染
        for _ in root.iter(f'{{{ns_v}}}image'):
            if _in_fallback(_):
                continue
            return True

        return False

    # 主文档
    if _part_has_image(doc.part):
        return -5, True

    # 页眉/页脚
    for section in doc.sections:
        for hf in (
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ):
            try:
                if _part_has_image(hf.part):
                    return -5, True
            except Exception:
                continue

    return 0, False


def check_caption_number_font(doc):
    """-3：表格上方加粗的表题：其中数字部分字体为宋体。

    严格按细则的每一个点判定（不引入细则未要求的额外约束）：
      (1) 表格上方：从 body 中每个 <w:tbl> 向上回溯，跳过纯空段落后
          遇到的第一个段落 = 该表格上方的段落（候选表题）。
      (2) 表题：候选段落文字必须形如 "表 <数字>…"（如"表 1"、"表 12"），
          与办公软件里我们通常识别的表题一致；不匹配则不是"表题"，跳过。
      (3) 加粗：候选段落里承载"表题正文"的 run 生效状态为加粗
          （沿 run rPr → rStyle 链 → 段落 rPr → pStyle 链 → docDefaults 解析）。
          非加粗段落不属于"加粗的表题"，跳过。
      (4) 数字部分：只针对 run.text 中包含数字字符（0-9）的 run。
      (5) 字体为宋体：办公软件（Word/WPS）中数字字符按 ASCII 类别渲染，
          使用 w:rFonts/w:ascii（若无则 asciiTheme），沿样式链回溯生效值；
          该值等于"宋体"或"SimSun"（不区分大小写）即命中扣分项。
          - 若 asciiTheme 指向主题字体且主题字体最终解析为宋体/SimSun，也算。
          - Word 对 `w:hint` 的处理：hint="eastAsia" 时数字仍按 ASCII 类别用
            ascii 字体渲染（仅在 ascii/asciiTheme 均缺省时才回退 eastAsia）；
            故此处按"ascii 值 → 无则 asciiTheme 主题解析 → 无则回退 eastAsia
            (仅当 hint=eastAsia)"顺序取用，与办公软件实际渲染一致。

    任一表格上方的加粗表题里、任一含数字的 run 数字部分字体为宋体
    → 命中扣分项，返回 -3。
    """
    styles = doc.styles

    # 主题字体解析：<w:themeElements>/<a:fontScheme>/<a:majorFont|a:minorFont>/<a:latin typeface=...>
    theme_latin = {}  # {'majorAscii': 'xxx', 'minorAscii': 'xxx'}
    try:
        theme_part = doc.part.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme')
    except Exception:
        theme_part = None
    if theme_part is not None:
        try:
            theme_root = theme_part.element
            ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            for major in theme_root.iter(f'{{{ns_a}}}majorFont'):
                latin = major.find(f'{{{ns_a}}}latin')
                if latin is not None:
                    theme_latin['majorAscii'] = latin.get('typeface')
            for minor in theme_root.iter(f'{{{ns_a}}}minorFont'):
                latin = minor.find(f'{{{ns_a}}}latin')
                if latin is not None:
                    theme_latin['minorAscii'] = latin.get('typeface')
        except Exception:
            pass

    def _get_style_by_id(style_id):
        if not style_id:
            return None
        for s in styles:
            try:
                if s.style_id == style_id:
                    return s
            except Exception:
                continue
        return None

    def _rPr_of_style_chain(style_obj):
        chain = []
        cur = style_obj
        depth = 0
        while cur is not None and depth < 20:
            rPr_elem = cur.element.find(qn('w:rPr'))
            if rPr_elem is not None:
                chain.append(rPr_elem)
            try:
                cur = cur.base_style
            except Exception:
                cur = None
            depth += 1
        return chain

    def _doc_defaults_rPr():
        dd = styles.element.find(qn('w:docDefaults'))
        if dd is None:
            return None
        rPrDefault = dd.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            return None
        return rPrDefault.find(qn('w:rPr'))

    def _collect_rPr_chain(r_elem):
        chain = []
        run_rPr = r_elem.find(qn('w:rPr'))
        if run_rPr is not None:
            chain.append(run_rPr)
            rStyle = run_rPr.find(qn('w:rStyle'))
            if rStyle is not None:
                sid = rStyle.get(qn('w:val'))
                chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        p = r_elem.getparent()
        while p is not None:
            ptag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if ptag == 'p':
                break
            p = p.getparent()
        if p is not None:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                para_rPr = pPr.find(qn('w:rPr'))
                if para_rPr is not None:
                    chain.append(para_rPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'))
                    chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        dd_rPr = _doc_defaults_rPr()
        if dd_rPr is not None:
            chain.append(dd_rPr)
        return chain

    def _first_of(chain, tag_qn):
        for rPr_elem in chain:
            elt = rPr_elem.find(tag_qn)
            if elt is not None:
                return elt
        return None

    def _bold_active(chain):
        elt = _first_of(chain, qn('w:b'))
        if elt is None:
            return False
        v = elt.get(qn('w:val'))
        if v is None:
            return True
        return v.lower() not in ('0', 'false', 'off')

    def _digit_font(chain):
        """办公软件中数字（ASCII 0-9）实际使用的字体名（若可解析），否则 None。"""
        # 沿链找到第一个 rFonts；ascii > asciiTheme > (hint=eastAsia 时) eastAsia/eastAsiaTheme
        for rPr_elem in chain:
            rFonts = rPr_elem.find(qn('w:rFonts'))
            if rFonts is None:
                continue
            ascii_val = rFonts.get(qn('w:ascii'))
            if ascii_val:
                return ascii_val
            ascii_theme = rFonts.get(qn('w:asciiTheme'))
            if ascii_theme:
                if ascii_theme in ('majorAscii', 'majorHAnsi'):
                    return theme_latin.get('majorAscii')
                if ascii_theme in ('minorAscii', 'minorHAnsi'):
                    return theme_latin.get('minorAscii')
                return ascii_theme
            # ascii 缺省，检查 hint 是否强制走 eastAsia 分支
            hint = rFonts.get(qn('w:hint'))
            if hint == 'eastAsia':
                ea = rFonts.get(qn('w:eastAsia'))
                if ea:
                    return ea
        return None

    def _is_songti(name):
        if not name:
            return False
        n = name.strip().lower()
        return n in ('宋体', 'simsun', 'simsun-extb', 'nsimsun')

    # 遍历 body 中每一个 <w:tbl>，找其"上方"最近的非空段落作为候选表题
    body = doc.element.body
    body_children = list(body)
    p_by_body_pos = {}
    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'p':
            p_by_body_pos[i] = child

    # 建 body_pos → docx Paragraph 映射（用于 style 名）
    paragraph_pos_to_obj = {}
    for para in doc.paragraphs:
        p_elem = para._p
        for i, c in enumerate(body_children):
            if c is p_elem:
                paragraph_pos_to_obj[i] = para
                break

    caption_re = re.compile(r'^表\s*\d')

    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag != 'tbl':
            continue

        # 向上寻找最近的非空段落
        caption_pos = None
        for j in range(i - 1, -1, -1):
            sib = body_children[j]
            stag = sib.tag.split('}')[1] if '}' in sib.tag else sib.tag
            if stag == 'p':
                text = ''.join((t.text or '') for t in sib.iter(qn('w:t')))
                if text.strip():
                    caption_pos = j
                    break
                else:
                    continue
            else:
                # 遇到另一个表格，说明当前表格是紧跟前一个表格的（跨页续页），无独立表题
                break

        if caption_pos is None:
            continue

        cap_para = paragraph_pos_to_obj.get(caption_pos)
        if cap_para is None:
            continue

        # (2) 必须形如"表 <数字>…"
        full_text = ''.join((t.text or '') for t in cap_para._p.iter(qn('w:t'))).strip()
        if not caption_re.match(full_text):
            continue

        # (3)+(4)+(5) 遍历段落里包含数字的 run，判断"加粗 + 数字字体"
        for r_elem in cap_para._p.iter(qn('w:r')):
            t_elem = r_elem.find(qn('w:t'))
            if t_elem is None or not t_elem.text:
                continue
            if not re.search(r'\d', t_elem.text):
                continue

            chain = _collect_rPr_chain(r_elem)

            # (3) 加粗
            if not _bold_active(chain):
                continue

            # (5) 数字字体为宋体
            font_name = _digit_font(chain)
            if _is_songti(font_name):
                return -3, True

    return 0, False


def check_table5_dimensions(doc):
    """-1：表5不是9行5列的表格。

    严格按细则的每一个点判定（不引入细则未要求的额外约束）：
      (1) 表5：找到形如"表 5"或"表5"的段落文字（后面不再跟数字，以避免误配
          "表 50"/"表55"），紧邻其后的第一个物理表格 <w:tbl> 即"表5"的物理起点。
      (2) 若"表5"被作者手动拆分为跨页两半（第二个物理表格与第一半表头一致、
          中间仅空段落），在办公软件（Word/WPS）里用户视觉上仍是同一张逻辑表 →
          行数按 (前半行数 + 后半行数 - 重复表头 1 行) 合并；列数取两半的
          <w:tblGrid> 列数最大值（正常情况下二者相同）。
      (3) 行数 = 9：办公软件里"行数"= <w:tr> 元素个数。
      (4) 列数 = 5：办公软件里"列数"= <w:tblGrid>/<w:gridCol> 个数
          （即 Word 表格属性里显示的"列数"，与合并单元格无关）。
      (5) 行数≠9 或 列数≠5 → 命中扣分项，返回 -1；找不到"表5"按不适用处理（不扣分）。
    """
    body = doc.element.body

    def _col_count(table):
        """办公软件中显示的列数 = <w:tblGrid>/<w:gridCol> 个数。"""
        tblGrid = table._tbl.find(qn('w:tblGrid'))
        if tblGrid is None:
            return len(table.columns)
        gridCols = tblGrid.findall(qn('w:gridCol'))
        if not gridCols:
            return len(table.columns)
        return len(gridCols)

    # 定位"表5"：正则要求 "5" 后不再接数字，避免匹配 "表 50" / "表55"
    caption_re = re.compile(r'^表\s*5(?!\d)')
    table5_idx = None
    pending = False
    table_counter = 0

    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = ''
            for t_elem in child.iter(qn('w:t')):
                if t_elem.text:
                    text += t_elem.text
            if caption_re.match(text.strip()):
                pending = True
        elif tag == 'tbl':
            if pending and table5_idx is None:
                table5_idx = table_counter
                pending = False
            table_counter += 1

    if table5_idx is None or table5_idx >= len(doc.tables):
        return 0, False  # 找不到"表5"→ 不适用

    # 跨页拆分合并：仅当当前表格是某跨页对的前半时合并后半
    cross_pairs = identify_cross_page_tables(doc)
    merged_partner_idx = None
    for idx_a, idx_b, _, _ in cross_pairs:
        if idx_a == table5_idx:
            merged_partner_idx = idx_b
            break

    t5 = doc.tables[table5_idx]
    rows = len(t5.rows)
    cols = _col_count(t5)

    if merged_partner_idx is not None:
        partner = doc.tables[merged_partner_idx]
        rows = rows + len(partner.rows) - 1  # 去掉重复表头
        cols = max(cols, _col_count(partner))

    if rows != 9 or cols != 5:
        return -1, True
    return 0, False


def _resolved_paragraph_line_spacing(para):
    """解析段落最终生效的行距（line, lineRule）。

    解析顺序：段落级 pPr → 段落样式（含 basedOn 链）→ docDefaults。
    返回 (line_val:int|None, line_rule:str|None)；找不到时为 (None, None)。
    """
    # 段落级
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            line = spacing.get(qn('w:line'))
            if line:
                return int(line), spacing.get(qn('w:lineRule'))

    # 样式链
    cur_style = para.style
    while cur_style is not None:
        style_pPr = cur_style.element.find(qn('w:pPr'))
        if style_pPr is not None:
            spacing = style_pPr.find(qn('w:spacing'))
            if spacing is not None:
                line = spacing.get(qn('w:line'))
                if line:
                    return int(line), spacing.get(qn('w:lineRule'))
        cur_style = cur_style.base_style

    # docDefaults
    styles_elem = para.part.document.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        pPrDefault = doc_defaults.find(qn('w:pPrDefault'))
        if pPrDefault is not None:
            pPr_d = pPrDefault.find(qn('w:pPr'))
            if pPr_d is not None:
                spacing = pPr_d.find(qn('w:spacing'))
                if spacing is not None:
                    line = spacing.get(qn('w:line'))
                    if line:
                        return int(line), spacing.get(qn('w:lineRule'))

    return None, None


def _document_is_rtl(doc):
    """判断文档默认双向方向：任意 sectPr/pPrDefault 中存在 <w:bidi/> 即视为 RTL。"""
    styles_elem = doc.styles.element
    # docDefaults/pPrDefault
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        pPrDefault = doc_defaults.find(qn('w:pPrDefault'))
        if pPrDefault is not None:
            pPr_d = pPrDefault.find(qn('w:pPr'))
            if pPr_d is not None and pPr_d.find(qn('w:bidi')) is not None:
                return True
    # sectPr
    body = doc.element.body
    for sp in body.iter(qn('w:sectPr')):
        if sp.find(qn('w:bidi')) is not None:
            return True
    return False


def _resolved_paragraph_alignment(para):
    """解析段落最终生效的水平对齐，归一化为 'left'|'right'|'center'|'both'|'distribute'。

    解析顺序：段落级 pPr → 段落样式链（含 basedOn）→ docDefaults。
    若四层都未显式设置：按 OOXML 规范，由文档双向方向决定（LTR 默认 left，RTL 默认 right）。
    段落自身的 <w:bidi/> 或文档级 bidi 触发 RTL 翻转，将 'start'/'end' 映射到正确的物理方向。
    """
    def _jc_in(pPr_elem):
        if pPr_elem is None:
            return None
        jc = pPr_elem.find(qn('w:jc'))
        return jc.get(qn('w:val')) if jc is not None else None

    raw = None

    # 1) 段落级
    pPr = para._p.find(qn('w:pPr'))
    raw = _jc_in(pPr)

    # 2) 样式链
    if raw is None:
        cur_style = para.style
        while cur_style is not None:
            v = _jc_in(cur_style.element.find(qn('w:pPr')))
            if v is not None:
                raw = v
                break
            cur_style = cur_style.base_style

    # 3) docDefaults
    if raw is None:
        styles_elem = para.part.document.styles.element
        doc_defaults = styles_elem.find(qn('w:docDefaults'))
        if doc_defaults is not None:
            pPrDefault = doc_defaults.find(qn('w:pPrDefault'))
            if pPrDefault is not None:
                raw = _jc_in(pPrDefault.find(qn('w:pPr')))

    # 4) 判定文档/段落 RTL 方向
    doc = para.part.document
    is_rtl = _document_is_rtl(doc)
    if pPr is not None and pPr.find(qn('w:bidi')) is not None:
        is_rtl = True

    # 5) 物理方向归一化
    if raw is None:
        return 'right' if is_rtl else 'left'
    if raw == 'start':
        return 'right' if is_rtl else 'left'
    if raw == 'end':
        return 'left' if is_rtl else 'right'
    return raw


def check_heading_33_format(doc):
    """-1："3.3 分层比较与模型校准"：中文文本不是四号黑体，不是1.3倍行距，不是左对齐。

    严格按细则的每一个点判定（不引入细则未要求的额外约束）：
      (1) 目标段落：文本包含"3.3 分层比较与模型校准"的段落（取第一个匹配段）。
      (2) 中文文本是四号黑体：
          - "中文文本"专指段落中的中文字符（Word/WPS 按 ASCII/eastAsia
            分类渲染，中文按 eastAsia 分类）。仅检查含中文字符的 run。
          - "黑体"= 该 run 生效的 eastAsia 字体 = "黑体"。
          - "四号"= 该 run 生效字号 = sz=28 半磅 (14 pt)。
          - 任一含中文的 run 违反上述任一 → 命中"不是四号黑体"。
      (3) 1.3 倍行距：段落生效 line=312 (=240×1.3) 且 lineRule 为 auto 或缺省
          (auto = 倍数模式；exact/atLeast 是绝对/最小点数模式，不算"倍数")。
      (4) 左对齐：段落生效对齐为 left 或 start（LTR 文档中未显式设置视为默认左对齐）。
      (5) 三条中任一条不满足 → 命中扣分项 -1。

    办公软件（Word/WPS）有效性：
      - 字体/字号解析沿 OOXML 覆盖顺序：run rPr → rStyle 链 → 段落 rPr →
        pStyle 链（含 basedOn） → docDefaults；取首个显式值，与办公软件
        最终渲染的字体/字号等价。
      - 行距/对齐通过既有的 `_resolved_paragraph_line_spacing` /
        `_resolved_paragraph_alignment` 沿 pPr → 样式链 → docDefaults 解析，
        并处理 LTR/RTL 方向对 start/end 的物理映射。
      - "中文文本"用 Unicode 范围 U+4E00–U+9FFF 判定，与办公软件识别中文字符
        用于选择 eastAsia 字体的规则一致（覆盖 CJK 基本区）。
    """
    styles = doc.styles

    def _get_style_by_id(style_id):
        if not style_id:
            return None
        for s in styles:
            try:
                if s.style_id == style_id:
                    return s
            except Exception:
                continue
        return None

    def _rPr_of_style_chain(style_obj):
        chain = []
        cur = style_obj
        depth = 0
        while cur is not None and depth < 20:
            rPr_elem = cur.element.find(qn('w:rPr'))
            if rPr_elem is not None:
                chain.append(rPr_elem)
            try:
                cur = cur.base_style
            except Exception:
                cur = None
            depth += 1
        return chain

    def _doc_defaults_rPr():
        dd = styles.element.find(qn('w:docDefaults'))
        if dd is None:
            return None
        rPrDefault = dd.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            return None
        return rPrDefault.find(qn('w:rPr'))

    def _collect_rPr_chain(r_elem):
        chain = []
        run_rPr = r_elem.find(qn('w:rPr'))
        if run_rPr is not None:
            chain.append(run_rPr)
            rStyle = run_rPr.find(qn('w:rStyle'))
            if rStyle is not None:
                sid = rStyle.get(qn('w:val'))
                chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        p = r_elem.getparent()
        while p is not None:
            ptag = p.tag.split('}')[1] if '}' in p.tag else p.tag
            if ptag == 'p':
                break
            p = p.getparent()
        if p is not None:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                para_rPr = pPr.find(qn('w:rPr'))
                if para_rPr is not None:
                    chain.append(para_rPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    sid = pStyle.get(qn('w:val'))
                    chain.extend(_rPr_of_style_chain(_get_style_by_id(sid)))

        dd_rPr = _doc_defaults_rPr()
        if dd_rPr is not None:
            chain.append(dd_rPr)
        return chain

    def _first_of(chain, tag_qn):
        for rPr_elem in chain:
            elt = rPr_elem.find(tag_qn)
            if elt is not None:
                return elt
        return None

    def _eastAsia_font(chain):
        elt = _first_of(chain, qn('w:rFonts'))
        return elt.get(qn('w:eastAsia')) if elt is not None else None

    def _sz_halfpt(chain):
        elt = _first_of(chain, qn('w:sz'))
        if elt is None:
            return None
        try:
            return int(elt.get(qn('w:val')))
        except (TypeError, ValueError):
            return None

    def _has_chinese(s):
        if not s:
            return False
        for ch in s:
            if '一' <= ch <= '鿿':
                return True
        return False

    # 定位目标段落
    target_para = None
    for para in doc.paragraphs:
        if "3.3 分层比较与模型校准" in para.text:
            target_para = para
            break
    if target_para is None:
        return 0, False  # 找不到 → 不适用

    # (2) 中文文本是否为四号黑体
    is_sihao_heiti = True
    checked_any_chinese_run = False
    for r_elem in target_para._p.iter(qn('w:r')):
        t_elem = r_elem.find(qn('w:t'))
        if t_elem is None or not t_elem.text:
            continue
        if not _has_chinese(t_elem.text):
            continue
        checked_any_chinese_run = True
        chain = _collect_rPr_chain(r_elem)
        if _eastAsia_font(chain) != '黑体':
            is_sihao_heiti = False
            break
        if _sz_halfpt(chain) != 28:
            is_sihao_heiti = False
            break
    if not checked_any_chinese_run:
        # 段落里没有中文字符 → "中文文本"这一小项不适用；按不违反处理
        pass

    # (3) 1.3 倍行距
    line, line_rule = _resolved_paragraph_line_spacing(target_para)
    is_1_3_spacing = (
        line is not None
        and (line_rule in (None, 'auto'))
        and line == 312
    )

    # (4) 左对齐
    jc_val = _resolved_paragraph_alignment(target_para)
    is_left = jc_val in ('left', 'start')

    # (5) 任一条不满足 → 命中扣分项
    if (not is_sihao_heiti) or (not is_1_3_spacing) or (not is_left):
        return -1, True
    return 0, False


def check_table1_note(doc):
    """-1：表1下方空一行的位置没有出现指定注释文本，则命中扣分。

    细则原文（两个点）：
      A. 位置："表 1 下方空一行的位置" —— Word/WPS 中表 1 视觉结束之后的
         那个空段落（渲染为一个空白行）再往下的那段文字。
      B. 内容：完整出现下述注释（细则给出的完整原文）：
         "注：分类变量以 n（%）表示，连续变量根据分布特征以 Mean±SD 或
          M（Q1，Q3）表示；统计量为χ²、t 或 Z 值。"

    判定步骤（严格对齐两个点、不引入细则未要求的额外约束）：
      1) 定位表 1 在 body 中的物理结束位置：
         若表 1 为跨页拆分（identify_cross_page_tables 命中 idx_a=0），
         物理表为 tables[0]+tables[1]，结束于第二个 <w:tbl>；否则结束于
         第一个 <w:tbl>。跨页拆分对应办公软件里视觉连续的同一张表 1，
         其"下方"从最后一个物理 <w:tbl> 之后开始。
      2) 从表 1 结束位置起向后收集段落，直到遇到下一张表格或段落用尽。
      3) A 点判定："表 1 下方"的紧邻段落存在且为空段落（空一行）。
      4) B 点判定：紧随空段落的下一段（也即"空一行的位置"下方那段），
         其规范化文本等于细则的注释原文；不同则视为"没有出现"。
      5) A 或 B 任一不满足 → 扣 -1。

    办公软件（Word/WPS）有效性：
      - "空一行"依 OOXML 段落级建模判定为无可见文本的段落（<w:p> 内
        所有 <w:t> 拼接去空白后为空），与办公软件渲染出的空白行一致。
      - 文本比较统一归一化：合并 Unicode 空白（含半/全角空格、U+00A0、
        U+3000 等）为单个 ASCII 空格并去首尾空白。半/全角标点（如"："
        vs ":"、"，" vs ","、"（"vs "("）与细则原文一致，办公软件按原样
        存储和渲染，不做二次替换以免放宽细则。
    """
    body = doc.element.body

    # 1) 表 1 物理表数量（跨页 → 2；否则 1）
    table1_phys_count = 1
    for idx_a, _idx_b, _bpa, _bpb in identify_cross_page_tables(doc):
        if idx_a == 0:
            table1_phys_count = 2
            break

    # 2) 表 1 结束后逐段收集，直到下一个 <w:tbl> 或段落用尽
    seen_tbls = 0
    paras_after = []
    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            seen_tbls += 1
            if seen_tbls > table1_phys_count:
                break
            continue
        if seen_tbls < table1_phys_count:
            continue
        if tag == 'p':
            text = ''
            for t_elem in child.iter(qn('w:t')):
                if t_elem.text:
                    text += t_elem.text
            paras_after.append(text)

    if len(paras_after) < 2:
        return -1, True

    # 3) A 点：紧邻段落为空一行
    if paras_after[0].strip() != '':
        return -1, True

    # 4) B 点：空一行下方那段的规范化文本 = 细则原文的规范化形式
    import re
    def _normalize(s):
        # 合并所有 Unicode 空白为单个空格；去首尾空白
        return re.sub(r'\s+', ' ', s).strip()

    expected = (
        "注：分类变量以 n（%）表示，"
        "连续变量根据分布特征以 Mean±SD 或 M（Q1，Q3）表示；"
        "统计量为χ²、t 或 Z 值。"
    )
    if _normalize(paras_after[1]) != _normalize(expected):
        return -1, True

    return 0, False


def check_table2_rows(doc):
    """-3：表 2 整体少于 40 行 → 扣 3 分。

    定位逻辑（与表 5 检查一致）：
      1. 遍历 body 直接子节点，找到以 "表 2" 开头的段落（正则避免误匹配"表 20"等）；
         紧随其后出现的第一个 <w:tbl> 即"表 2"的物理起点 table2_idx。
      2. 若 table2_idx 在 identify_cross_page_tables 的跨页对中作为前半部分，
         则合并后半部分行数：rows = len(主表) + len(后半) - 1（去重表头）。
      3. 行数 < 40 → 命中扣分项；否则不扣分；找不到"表 2"按不适用处理（不扣分）。
    """
    body = doc.element.body
    table2_idx = None
    pending = False
    table_counter = 0

    for child in body:
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = ''
            for t_elem in child.iter(qn('w:t')):
                if t_elem.text:
                    text += t_elem.text
            if re.match(r'^表\s*2(?:\s|$)', text.strip()):
                pending = True
        elif tag == 'tbl':
            if pending and table2_idx is None:
                table2_idx = table_counter
                pending = False
            table_counter += 1

    if table2_idx is None or table2_idx >= len(doc.tables):
        return 0, False

    # 合并跨页后半部分
    cross_pairs = identify_cross_page_tables(doc)
    merged_partner_idx = None
    for idx_a, idx_b, _, _ in cross_pairs:
        if idx_a == table2_idx:
            merged_partner_idx = idx_b
            break

    t2 = doc.tables[table2_idx]
    total_rows = len(t2.rows)
    if merged_partner_idx is not None:
        total_rows += len(doc.tables[merged_partner_idx].rows) - 1  # 去掉重复表头

    if total_rows < 40:
        return -3, True
    return 0, False


def check_table_border_count(doc):
    """检查是否有表格出现大于3条水平边框线。有则扣3分。"""
    has_issue = False

    for ti, table in enumerate(doc.tables):
        # 统计每个表格的不同水平线位置
        border_positions = set()

        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    continue
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    continue

                for border in tcBorders:
                    tag = border.tag.split('}')[1]
                    val = border.get(qn('w:val'))

                    if val == 'nil' or not val:
                        continue

                    if tag == 'top':
                        border_positions.add(('h', ri, 'top'))
                    elif tag == 'bottom':
                        border_positions.add(('h', ri, 'bottom'))
                    elif tag in ('left', 'right'):
                        border_positions.add(('v', ri, tag))

        # 合并水平线：top of row N = bottom of row N-1
        h_lines = set()
        v_exists = False
        for pos_type, ri, direction in border_positions:
            if pos_type == 'h':
                if direction == 'top':
                    h_lines.add(ri)
                else:
                    h_lines.add(ri + 1)
            else:
                v_exists = True

        # 总边框线数 = 水平线 + 垂直线（如果有的话）
        total_border_lines = len(h_lines)
        if v_exists:
            total_border_lines += 1  # 简化：有任何垂直线算一条

        if total_border_lines > 3:
            has_issue = True
            break

    if has_issue:
        return -3, True  # 命中扣分项
    return 0, False


def _find_docx_in_dir(dir_path):
    """在给定目录中定位待评估的 Word 文档。

    规则：
      - 只考虑 .docx 后缀；
      - 忽略以 "~$" 开头的 Office 临时锁定文件；
      - 若存在多个候选，按文件名排序取第一个（保持行为可复现）。
    找不到时返回 None。
    """
    if not os.path.isdir(dir_path):
        return None
    candidates = []
    for name in os.listdir(dir_path):
        if name.startswith('~$'):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext == '.docx':
            candidates.append(name)
    if not candidates:
        return None
    candidates.sort()
    return os.path.join(dir_path, candidates[0])


# 维度二规则清单：与 evaluate_document 返回的 results 顺序一一对应，用于稳定
# 提取每一项的 max_delta（+N/-N）以及命中时的 delta，便于组装统一返回结构。
_DIM2_RULE_PATTERN = re.compile(r'^([+-])(\d+)：')


def evaluate(dir_path: str) -> dict:
    """统一入口：接收"脚本所在目录的路径"，脚本自己在该目录里定位并打开被评估的文档。

    返回结构见《脚本接口差异与统一建议.md》§2.2。
    """
    result = {
        "id": SCRIPT_ID,
        "file_name": "",
        "status": "ok",
        "error": None,
        "dim1_pass": False,
        "dim1_reason": "",
        "dim2_items": [],
        "total_score": 0,
        "max_score": 0,
    }

    try:
        filepath = _find_docx_in_dir(dir_path)
        if filepath is None:
            result["status"] = "error"
            result["error"] = f"目录中未找到 .docx 文件：{dir_path}"
            return result

        result["file_name"] = os.path.basename(filepath)

        results, total_score = evaluate_document(filepath)

        # 维度一：所有以 "维度1" 开头的结果项必须全部 passed
        dim1_items = [(desc, passed) for desc, _score, passed in results
                      if desc.startswith("维度1")]
        dim1_failed = [desc for desc, passed in dim1_items if not passed]
        result["dim1_pass"] = (len(dim1_failed) == 0)
        result["dim1_reason"] = "" if result["dim1_pass"] else "；".join(dim1_failed)

        # 维度二：逐项列出（命中和未命中都列出）
        dim2_items = []
        max_score = 0
        for desc, score, passed in results:
            if desc.startswith("维度1"):
                continue
            m = _DIM2_RULE_PATTERN.match(desc)
            if m:
                sign = 1 if m.group(1) == '+' else -1
                max_delta = sign * int(m.group(2))
            else:
                raise ValueError(f"维度二规则缺少正负分值前缀: {desc}")
            dim2_items.append({
                "rule": desc,
                "max_delta": max_delta,
                "delta": score,
                "hit": bool(passed),
                "detail": "",
            })
            # max_score = 所有加分项（max_delta > 0）的 max_delta 之和，
            # 不计入惩罚项（-N），与"满分"的直觉一致。
            if max_delta > 0:
                max_score += max_delta

        result["dim2_items"] = dim2_items
        result["total_score"] = total_score
        result["max_score"] = max_score
        return result

    except Exception as e:
        result["status"] = "error"
        result["error"] = f"{type(e).__name__}: {e}"
        return result


if __name__ == "__main__":
    # 本地调试入口：接收"脚本所在目录的路径"，默认取脚本自身所在目录。
    if len(sys.argv) > 1:
        _dir = sys.argv[1]
    else:
        _dir = os.path.dirname(os.path.abspath(__file__))
    _out = json.dumps(evaluate(_dir), ensure_ascii=False, indent=2)
    try:
        print(_out)
    except UnicodeEncodeError:
        # Windows 控制台默认 cp1252 时，直接按 UTF-8 写入二进制 stdout
        sys.stdout.buffer.write(_out.encode('utf-8'))
        sys.stdout.buffer.write(b'\n')
